from __future__ import annotations

import re
import json
import shutil
from pathlib import Path
from typing import Any

from docx import Document

from app.services.onlyoffice_documents import document_path
from app.services.store import now_iso, store
from app.services.workspace_artifacts import business_workspace_dir


def apply_controlled_business_rewrite(
    project_id: str,
    *,
    original_text: str,
    replacement_text: str,
    operator: str = "当前用户",
) -> dict[str, Any]:
    """Apply a single, exact, human-approved rewrite to the business bid docx.

    This intentionally avoids broad AI edits. The original text must match one
    and only one body/table paragraph, or appear literally once inside one
    paragraph. Otherwise no document mutation is performed.
    """

    original = _normalize_input(original_text)
    replacement = _normalize_input(replacement_text)
    if not original:
        raise ValueError("原文段落不能为空。")
    if not replacement:
        raise ValueError("替换文本不能为空。")

    path = document_path(project_id)
    if not path.exists():
        state = store.get_document_state(project_id)
        raise FileNotFoundError(f"商务标正文文件不存在：{state.get('fileName') or path.name}")

    document = Document(str(path))
    match = _find_unique_rewrite_match(document, original)
    backup_path = _backup_document(project_id, path)

    paragraph = match["paragraph"]
    if match["mode"] == "literal_substring":
        next_text = paragraph.text.replace(original, replacement, 1)
    else:
        next_text = replacement
    _replace_paragraph_text(paragraph, next_text)
    document.save(str(path))

    record = {
        "schemaVersion": "business-controlled-rewrite-v1",
        "appliedAt": now_iso(),
        "operator": operator,
        "originalText": original,
        "replacementText": replacement,
        "matchMode": match["mode"],
        "matchLocation": match["location"],
        "backupFile": str(backup_path),
        "outputFile": str(path),
    }
    record["historyFile"] = str(_append_rewrite_history(project_id, record))
    return record


def _normalize_input(value: str) -> str:
    text = str(value or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def _match_key(value: str) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _find_unique_rewrite_match(document: Document, original: str) -> dict[str, Any]:
    paragraphs = _iter_body_and_table_paragraphs(document)
    original_key = _match_key(original)
    exact_matches = [
        item for item in paragraphs
        if _match_key(item["paragraph"].text) == original_key
    ]
    if len(exact_matches) == 1:
        return {**exact_matches[0], "mode": "exact_paragraph"}
    if len(exact_matches) > 1:
        raise ValueError(f"原文段落在正文中匹配到 {len(exact_matches)} 处，请补充更精确的原文。")

    literal_matches = [
        item for item in paragraphs
        if original in item["paragraph"].text
    ]
    if len(literal_matches) == 1:
        return {**literal_matches[0], "mode": "literal_substring"}
    if len(literal_matches) > 1:
        raise ValueError(f"原文片段在正文中匹配到 {len(literal_matches)} 处，请补充更精确的原文。")

    raise ValueError("未在当前商务标正文中找到该原文段落，请从左侧 Word 预览中复制完整段落后重试。")


def _iter_body_and_table_paragraphs(document: Document) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            items.append({"paragraph": paragraph, "location": f"正文段落 {index}"})

    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                for paragraph_index, paragraph in enumerate(cell.paragraphs, start=1):
                    text = paragraph.text.strip()
                    if text:
                        items.append(
                            {
                                "paragraph": paragraph,
                                "location": (
                                    f"表格 {table_index} / 行 {row_index} / 列 {cell_index}"
                                    f" / 段落 {paragraph_index}"
                                ),
                            }
                        )
    return items


def _replace_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def _backup_document(project_id: str, path: Path) -> Path:
    backup_dir = business_workspace_dir(project_id) / "s4_ai_rewrite_backups"
    backup_dir.mkdir(parents=True, exist_ok=True)
    backup_path = backup_dir / f"{now_iso().replace(':', '').replace('-', '').replace('Z', '')}-{path.name}"
    shutil.copy2(path, backup_path)
    return backup_path


def _append_rewrite_history(project_id: str, record: dict[str, Any]) -> Path:
    history_path = business_workspace_dir(project_id) / "s4_ai_rewrite_backups" / "rewrite-history.jsonl"
    history_path.parent.mkdir(parents=True, exist_ok=True)
    with history_path.open("a", encoding="utf-8") as handle:
        handle.write(json.dumps(record, ensure_ascii=False) + "\n")
    return history_path
