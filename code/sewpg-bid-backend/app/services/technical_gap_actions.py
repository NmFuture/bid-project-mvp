from __future__ import annotations

import base64
import binascii
import copy
import re
import shutil
from uuid import uuid4
from pathlib import Path
from typing import Any

from docx import Document

from app.services.bid_runtime_state import now_iso
from app.services.minio_client import minio_client
from app.services.technical_gap_ai_fill import (
    TECHNICAL_TABLE_FILL_SKILL_NAME,
    TECHNICAL_WORD_FILL_SKILL_NAME,
    run_technical_ai_fill_for_gap,
)
from app.services.technical_gap_planner import build_technical_gap_plan_for_project
from app.services.technical_gap_domain import (
    summarize_technical_gap_plan,
    technical_gap_artifact_onlyoffice_payload,
)
from app.services.technical_gap_state import legacy_technical_gap_items_from_plan
from app.services.technical_material_store import technical_material_store
from app.services.workspace_artifacts import technical_workspace_dir

def _safe_filename(value: str, fallback: str) -> str:
    text = re.sub(r"[\\/:*?\"<>|]+", "-", str(value or "").strip())
    text = re.sub(r"\s+", " ", text).strip(" .")
    return text or fallback


def _project_dir(project: dict[str, Any]) -> Path:
    project_id = str(project.get("id") or "")
    project_dir = technical_workspace_dir(project_id)
    project_dir.mkdir(parents=True, exist_ok=True)
    return project_dir


def _decode_uploaded_docx(content: str) -> bytes | None:
    text = str(content or "").strip()
    if not text.startswith("data:"):
        return None
    header, separator, payload = text.partition(",")
    if not separator or ";base64" not in header.lower():
        return None
    try:
        decoded = base64.b64decode(payload, validate=True)
    except (binascii.Error, ValueError):
        return None
    if decoded.startswith(b"PK\x03\x04"):
        return decoded
    return None


def _write_technical_manual_upload_docx(path: Path, *, title: str, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    uploaded_bytes = _decode_uploaded_docx(content)
    if uploaded_bytes is not None:
        path.write_bytes(uploaded_bytes)
        return
    doc = Document()
    doc.add_heading(title or path.stem, level=1)
    text = content.strip() or "人工上传客户资料，原始文件内容请以项目素材库归档为准。"
    for paragraph in text.splitlines() or [text]:
        doc.add_paragraph(paragraph)
    doc.save(path)


def _selected_technical_material_entries(data: dict[str, Any]) -> list[dict[str, Any]]:
    raw_entries = data.get("materials")
    if not raw_entries:
        raw_entries = data.get("materialIds") or data.get("referenceMaterialIds") or []
    entries: list[dict[str, Any]] = []
    for entry in raw_entries if isinstance(raw_entries, list) else []:
        if isinstance(entry, str):
            entries.append({"id": entry})
        elif isinstance(entry, dict):
            entries.append(entry)
    return entries


def _payload_is_docx(payload: dict[str, Any]) -> bool:
    mime_type = str(payload.get("mimeType") or "")
    file_name = str(payload.get("fileName") or "")
    return "wordprocessingml" in mime_type or file_name.lower().endswith(".docx")


def _cleanup_partial_download(target_path: Path, download_error: Exception) -> Exception:
    cleanup_errors: list[Exception] = []
    partial_path = target_path.with_suffix(f"{target_path.suffix}.download")
    for path in (target_path, partial_path):
        try:
            path.unlink(missing_ok=True)
        except Exception as cleanup_error:
            cleanup_errors.append(cleanup_error)
    if cleanup_errors:
        return ExceptionGroup(
            f"下载失败且未能清理部分文件：{target_path}",
            [download_error, *cleanup_errors],
        )
    return download_error


def _remove_technical_material_batch_dirs(batch_dirs: set[Path]) -> None:
    cleanup_errors: list[Exception] = []
    for batch_dir in batch_dirs:
        try:
            shutil.rmtree(batch_dir)
        except FileNotFoundError:
            continue
        except Exception as cleanup_error:
            cleanup_errors.append(cleanup_error)
    if cleanup_errors:
        raise ExceptionGroup("已选择素材批次清理失败。", cleanup_errors)


def cleanup_prepared_technical_gap_material_files(prepared_files: list[dict[str, Any]]) -> None:
    batch_dirs = {
        Path(str(prepared.get("batchDir")))
        for prepared in prepared_files
        if str(prepared.get("batchDir") or "").strip()
    }
    _remove_technical_material_batch_dirs(batch_dirs)


async def _cleaned_technical_material_payload(material_id: str) -> dict[str, Any]:
    payload = await technical_material_store.raw_download_cleaned_content(material_id)
    if not _payload_is_docx(payload):
        raise ValueError(f"素材 {material_id} 没有可用于拼接的 Word 文件或清洗稿。")
    return payload


async def _downloadable_technical_material_payload(material_id: str) -> tuple[dict[str, Any], str]:
    # 优先原始 Word：素材清洗版可能误改标题层级，组装只认原始 docx 的真实
    # Heading/outlineLvl/TOC；原始文件缺失或不是 docx（如 .doc）时回退清洗稿。
    raw_error: Exception | None = None
    try:
        payload = await technical_material_store.raw_download_content(material_id)
        if _payload_is_docx(payload):
            return payload, "raw"
    except Exception as exc:
        raw_error = exc

    try:
        return await _cleaned_technical_material_payload(material_id), "cleaned"
    except Exception as cleaned_error:
        if raw_error is None:
            raise
        raise ExceptionGroup(
            f"素材 {material_id} 原始 Word 与清洗稿元数据均不可用。",
            [raw_error, cleaned_error],
        )


def register_technical_manual_gap_upload(
    project: dict[str, Any],
    gap_id: str,
    data: dict[str, Any],
    *,
    browser_base_url: str = "",
    onlyoffice_base_url: str = "",
) -> dict[str, Any]:
    gap_state = project.get("gap_state") or {}
    plan = gap_state.get("plan") if isinstance(gap_state.get("plan"), dict) else {}
    items = plan.get("items") if isinstance(plan.get("items"), list) else []
    item = next((entry for entry in items if str(entry.get("id") or "") == gap_id), None)
    if item is None:
        raise KeyError(gap_id)
    files = [entry for entry in data.get("files") or [] if isinstance(entry, dict)]
    if not files:
        raise ValueError("至少需要提交一个文件。")

    work_dir = _project_dir(project) / "s4_gap_workdir" / "manual_upload" / gap_id
    work_dir.mkdir(parents=True, exist_ok=True)
    created_at = now_iso()
    artifacts: list[dict[str, Any]] = []
    for index, file in enumerate(files, start=1):
        name = _safe_filename(str(file.get("name") or f"{gap_id}-{index}.docx"), f"{gap_id}-{index}.docx")
        if not name.lower().endswith(".docx"):
            name = f"{Path(name).stem}.docx"
        output_file = work_dir / name
        content = str(file.get("data") or file.get("text") or file.get("content") or "")
        _write_technical_manual_upload_docx(output_file, title=str(item.get("title") or name), content=content)
        artifact_id = f"ART-{gap_id}-UPLOAD-{index}"
        artifacts.append(
            {
                "id": artifact_id,
                "source": "manual_upload",
                "skill": "",
                "title": str(item.get("title") or output_file.stem),
                "fileName": output_file.name,
                "path": str(output_file),
                "createdAt": created_at,
                "operator": str(data.get("operator") or "当前用户"),
                "unfilledFields": [],
                "evidenceRefs": [],
                "onlyoffice": technical_gap_artifact_onlyoffice_payload(
                    project_id=str(project.get("id") or ""),
                    artifact_id=artifact_id,
                    file_name=output_file.name,
                    browser_base_url=browser_base_url,
                    onlyoffice_base_url=onlyoffice_base_url,
                ),
                "s7Ready": True,
            }
        )

    item["status"] = "resolved"
    item.setdefault("resolvedArtifacts", []).extend(artifacts)
    item["resolvedAt"] = created_at
    item["resolvedSource"] = artifacts[0]["fileName"]
    item["latestUploadAt"] = created_at
    item["latestSubmissionId"] = artifacts[0]["id"]
    # 选定即定案（产品裁决 2026-08-04）：人工上传本身就是人工决策，不再要求二次点「确认」。
    item["humanConfirmed"] = True
    item["humanConfirmedAt"] = created_at
    item["humanConfirmedBy"] = str(data.get("operator") or "当前用户")
    plan["updatedAt"] = created_at
    plan["summary"] = summarize_technical_gap_plan(plan)
    gap_state["plan"] = plan
    gap_state["items"] = legacy_technical_gap_items_from_plan(plan)
    gap_state["submittedForReview"] = False
    gap_state["reviewConfirmed"] = False
    gap_state["reviewedAt"] = ""
    return {"item": item, "artifact": artifacts[0], "artifacts": artifacts, "gapPlan": plan}


async def prepare_technical_existing_gap_material_files(
    project: dict[str, Any],
    gap_id: str,
    data: dict[str, Any],
) -> list[dict[str, Any]]:
    selected = _selected_technical_material_entries(data)
    if not selected:
        raise ValueError("至少需要选择一份素材。")

    work_root = _project_dir(project) / "s4_gap_workdir" / "selected_material" / gap_id
    work_root.mkdir(parents=True, exist_ok=True)
    batch_dir = work_root / f"batch-{uuid4().hex}"
    batch_dir.mkdir()
    prepared: list[dict[str, Any]] = []
    try:
        for index, material in enumerate(selected, start=1):
            material_id = str(material.get("id") or material.get("materialId") or "").strip()
            if not material_id:
                continue
            payload, source_kind = await _downloadable_technical_material_payload(material_id)
            file_name = _safe_filename(
                str(
                    material.get("cleanedFileName")
                    or payload.get("fileName")
                    or material.get("name")
                    or f"{material_id}.docx"
                ),
                f"{material_id}.docx",
            )
            if not file_name.lower().endswith(".docx"):
                file_name = f"{Path(file_name).stem}.docx"
            target_path = batch_dir / f"{index:02d}-{file_name}"
            try:
                minio_client.download_file(str(payload["bucket"]), str(payload["key"]), target_path)
            except Exception as raw_download_error:
                if source_kind != "raw":
                    raise _cleanup_partial_download(target_path, raw_download_error)
                raw_failure = _cleanup_partial_download(target_path, raw_download_error)
                try:
                    payload = await _cleaned_technical_material_payload(material_id)
                    minio_client.download_file(str(payload["bucket"]), str(payload["key"]), target_path)
                    source_kind = "cleaned"
                except Exception as cleaned_error:
                    cleaned_failure = _cleanup_partial_download(target_path, cleaned_error)
                    raise ExceptionGroup(
                        f"素材 {material_id} 原始 Word 与清洗稿均无法下载。",
                        [raw_failure, cleaned_failure],
                    )
            prepared.append(
                {
                    "materialId": material_id,
                    "materialName": str(material.get("name") or payload.get("fileName") or material_id),
                    "fileName": target_path.name,
                    "path": str(target_path),
                    "batchDir": str(batch_dir),
                    "folderPath": str(material.get("folderPath") or ""),
                    "materialTier": str(material.get("materialTier") or ""),
                    "sourceKind": source_kind,
                    "originalMaterial": material,
                }
            )
        if not prepared:
            raise ValueError("至少需要选择一份有效素材。")
        return prepared
    except Exception as prepare_error:
        try:
            _remove_technical_material_batch_dirs({batch_dir})
        except Exception as cleanup_error:
            raise ExceptionGroup(
                "准备已选择素材失败，且当前批次未能完整清理。",
                [prepare_error, cleanup_error],
            ) from prepare_error
        raise


def register_technical_existing_gap_material(
    project: dict[str, Any],
    gap_id: str,
    data: dict[str, Any],
    prepared_files: list[dict[str, Any]],
    *,
    browser_base_url: str = "",
    onlyoffice_base_url: str = "",
) -> dict[str, Any]:
    gap_state = project.get("gap_state") or {}
    plan = gap_state.get("plan") if isinstance(gap_state.get("plan"), dict) else {}
    items = plan.get("items") if isinstance(plan.get("items"), list) else []
    item = next((entry for entry in items if str(entry.get("id") or "") == gap_id), None)
    if item is None:
        raise KeyError(gap_id)
    if not prepared_files:
        raise ValueError("至少需要选择一份素材。")

    created_at = now_iso()
    existing_count = len(item.get("resolvedArtifacts") or [])
    artifacts: list[dict[str, Any]] = []
    for index, prepared in enumerate(prepared_files, start=1):
        path = Path(str(prepared.get("path") or ""))
        if not path.exists():
            raise ValueError(f"已选择素材文件不存在：{path}")
        artifact_id = f"ART-{gap_id}-MAT-{existing_count + index}"
        artifacts.append(
            {
                "id": artifact_id,
                "source": "material_library",
                "skill": "",
                "title": str(prepared.get("materialName") or path.stem),
                "fileName": str(prepared.get("fileName") or path.name),
                "path": str(path),
                "materialId": str(prepared.get("materialId") or ""),
                "folderPath": str(prepared.get("folderPath") or ""),
                "materialTier": str(prepared.get("materialTier") or ""),
                "sourceKind": str(prepared.get("sourceKind") or ""),
                "createdAt": created_at,
                "operator": str(data.get("operator") or "当前用户"),
                "unfilledFields": [],
                "evidenceRefs": [{"type": "material", "id": str(prepared.get("materialId") or "")}],
                "onlyoffice": technical_gap_artifact_onlyoffice_payload(
                    project_id=str(project.get("id") or ""),
                    artifact_id=artifact_id,
                    file_name=str(prepared.get("fileName") or path.name),
                    browser_base_url=browser_base_url,
                    onlyoffice_base_url=onlyoffice_base_url,
                ),
                "s7Ready": True,
            }
        )

    item["status"] = "resolved"
    item.setdefault("resolvedArtifacts", []).extend(artifacts)
    item["resolvedAt"] = created_at
    item["resolvedSource"] = artifacts[0]["fileName"]
    item.setdefault("reviewNotes", []).append(f"人工选择已有素材：{len(artifacts)} 份")
    # 选定即定案（产品裁决 2026-08-04）：人工亲手选素材本身就是人工决策，不再要求二次点「确认」。
    item["humanConfirmed"] = True
    item["humanConfirmedAt"] = created_at
    item["humanConfirmedBy"] = str(data.get("operator") or "当前用户")
    plan["updatedAt"] = created_at
    plan["summary"] = summarize_technical_gap_plan(plan)
    gap_state["plan"] = plan
    gap_state["items"] = legacy_technical_gap_items_from_plan(plan)
    gap_state["submittedForReview"] = False
    gap_state["reviewConfirmed"] = False
    gap_state["reviewedAt"] = ""
    return {"item": item, "artifact": artifacts[0], "artifacts": artifacts, "gapPlan": plan}


# ---------------------------------------------------------------------------
# 人工设置「父章节覆盖」（产品需求 2026-07-27）
#
# planner 只在拿得出证据时才自动设整章覆盖；实际评审里人看一眼就知道「这一章选了
# 这份素材，下面的小节都跟着它写」。这里把这个判断开放给人工：以本节点为覆盖源，
# 把其后代目录项统一标成 covered_by_parent，可撤销。
# ---------------------------------------------------------------------------

_CHAPTER_NUMBER_RE = re.compile(r"^第\s*([一二三四五六七八九十百千万0-9]+)\s*章$")
_CHINESE_DIGITS = {"零": 0, "一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9}
# 撤销时需要还原的字段；与下方写入的字段一一对应。
_PARENT_COVERAGE_FIELDS = (
    "status",
    "decision",
    "usage",
    "coverageRole",
    "coveredByParent",
    "matchedMaterials",
    "gapReason",
    "nextActions",
)


def _chinese_number_to_int(value: str) -> int | None:
    text = str(value or "").strip()
    if text.isdigit():
        return int(text)
    if text in _CHINESE_DIGITS:
        return _CHINESE_DIGITS[text]
    if text == "十":
        return 10
    if "十" in text:
        left, _, right = text.partition("十")
        tens = _CHINESE_DIGITS.get(left, 1) if left else 1
        ones = _CHINESE_DIGITS.get(right, 0) if right else 0
        return tens * 10 + ones
    return None


def technical_gap_number_key(number: Any) -> str:
    """目录号归一化：「第3章」→「3」，其余原样；与 planner 的 toc_number_key 同口径。"""
    text = str(number or "").strip()
    match = _CHAPTER_NUMBER_RE.fullmatch(text)
    if match:
        value = _chinese_number_to_int(match.group(1))
        return str(value) if value is not None else text
    return text


def technical_gap_descendant_items(items: list[dict[str, Any]], parent: dict[str, Any]) -> list[dict[str, Any]]:
    """本节点的全部后代目录项（按目录号前缀，第3章 → 3.1 / 3.1.2 …）。"""
    parent_key = technical_gap_number_key(parent.get("number"))
    if not parent_key:
        return []
    prefix = f"{parent_key}."
    return [
        item
        for item in items
        if item is not parent and technical_gap_number_key(item.get("number")).startswith(prefix)
    ]


def apply_technical_gap_parent_coverage(
    project: dict[str, Any],
    gap_id: str,
    data: dict[str, Any] | None = None,
) -> dict[str, Any]:
    payload = data or {}
    covered = payload.get("covered", True) is not False
    operator = str(payload.get("operator") or "当前用户")

    gap_state = project.get("gap_state") or {}
    plan = gap_state.get("plan") if isinstance(gap_state.get("plan"), dict) else {}
    items = plan.get("items") if isinstance(plan.get("items"), list) else []
    parent = next((entry for entry in items if str(entry.get("id") or "") == gap_id), None)
    if parent is None:
        raise KeyError(gap_id)

    descendants = technical_gap_descendant_items(items, parent)
    if not descendants:
        raise ValueError("本目录项没有下级目录，无法设置父章节覆盖。")

    parent_title = str(parent.get("title") or "")
    applied: list[str] = []
    skipped: list[str] = []

    if covered:
        # 覆盖源必须自己先有素材，否则子节会指向一个空的父节点。
        has_material = bool(parent.get("matchedMaterials")) or bool(parent.get("resolvedArtifacts"))
        if not has_material:
            raise ValueError("本章还没有选用素材，请先选好素材再设置父章节覆盖。")
        parent_decision = str(parent.get("decision") or "ready")
        for child in descendants:
            # 人工已经单独选过素材的子节不动，避免一键覆盖抹掉别人的选择。
            if child.get("resolvedArtifacts"):
                skipped.append(str(child.get("number") or child.get("id") or ""))
                continue
            if not isinstance(child.get("parentCoverageBackup"), dict):
                child["parentCoverageBackup"] = {
                    field: copy.deepcopy(child.get(field)) for field in _PARENT_COVERAGE_FIELDS
                }
            child["status"] = "needs_input" if parent_decision == "fill_required" else "matched"
            child["decision"] = parent_decision
            child["usage"] = "covered_by_parent"
            child["coverageRole"] = "covered_by_parent"
            child["coveredByParent"] = gap_id
            child["matchedMaterials"] = []
            child["gapReason"] = f"已由父章节“{parent_title}”整章素材覆盖。"
            child["nextActions"] = ["ai_fill_word"] if parent_decision == "fill_required" else ["s4_merge_material"]
            child["parentCoverageSource"] = "manual"
            child.setdefault("reviewNotes", []).append(f"人工设为父章节覆盖（{parent_title}）：{operator}")
            applied.append(str(child.get("number") or child.get("id") or ""))
    else:
        for child in descendants:
            if str(child.get("coveredByParent") or "") != gap_id:
                continue
            if str(child.get("parentCoverageSource") or "") != "manual":
                # planner 自动判定的覆盖不由这个按钮撤销，避免和识别结果打架。
                skipped.append(str(child.get("number") or child.get("id") or ""))
                continue
            backup = child.get("parentCoverageBackup")
            if isinstance(backup, dict):
                for field in _PARENT_COVERAGE_FIELDS:
                    child[field] = copy.deepcopy(backup.get(field))
            child.pop("parentCoverageBackup", None)
            child.pop("parentCoverageSource", None)
            child.setdefault("reviewNotes", []).append(f"撤销父章节覆盖（{parent_title}）：{operator}")
            applied.append(str(child.get("number") or child.get("id") or ""))

    parent["parentCoverageApplied"] = covered and bool(applied)
    updated_at = now_iso()
    plan["updatedAt"] = updated_at
    plan["summary"] = summarize_technical_gap_plan(plan)
    gap_state["plan"] = plan
    gap_state["items"] = legacy_technical_gap_items_from_plan(plan)
    gap_state["submittedForReview"] = False
    gap_state["reviewConfirmed"] = False
    gap_state["reviewedAt"] = ""
    return {
        "item": parent,
        "applied": applied,
        "skipped": skipped,
        "gapPlan": plan,
    }
