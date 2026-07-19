from __future__ import annotations

import hashlib
import json
import logging
import re
import shutil
import subprocess
import tempfile
import unicodedata
from copy import deepcopy
from io import BytesIO
from pathlib import Path
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile
from xml.etree import ElementTree as ET

from docx import Document
from docx.shared import Pt
from sqlalchemy import bindparam, text

from app.core.config import settings
from app.models import async_session
from app.services.material_runtime_tables import ensure_material_runtime_tables
from app.services.material_tags import normalize_material_tags
from app.services.minio_client import minio_client
from app.services.peripheral import PeripheralError


PERFORMANCE_CATEGORY_SCOPES = {"standard", "customer", "project"}
PERFORMANCE_CATEGORY_REVIEW_STATUSES = {"draft", "reviewed", "disabled"}
PERFORMANCE_CATEGORY_STATUSES = {"enabled", "disabled"}
SUMMARY_ATTACHMENT_TYPE = "summary_table"
CONTRACT_ATTACHMENT_TYPE = "contract_bundle"
ITEM_CONTRACT_ATTACHMENT_TYPE = "contract_item"
DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
ITEM_CONTRACT_FORMAT_VERSION = 15
CONTRACT_OUTPUT_EAST_ASIA_FONT = "Songti SC"
CONTRACT_OUTPUT_WESTERN_FONT = "Times New Roman"
CONTRACT_OUTPUT_SYMBOL_FONT = "Symbol"
RELATIONSHIP_NAMESPACE = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
WORD_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
MC_NAMESPACE = "http://schemas.openxmlformats.org/markup-compatibility/2006"
A_NAMESPACE = "http://schemas.openxmlformats.org/drawingml/2006/main"
WP_NAMESPACE = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
PIC_NAMESPACE = "http://schemas.openxmlformats.org/drawingml/2006/picture"
THEME_EAST_ASIA_SCRIPTS = {"Hans", "Hant", "Jpan", "Hang"}
OOXML_NAMESPACES = {
    "w": WORD_NAMESPACE,
    "mc": MC_NAMESPACE,
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp": WP_NAMESPACE,
    "a": A_NAMESPACE,
    "pic": PIC_NAMESPACE,
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
    "wp14": "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing",
    "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
}
logger = logging.getLogger(__name__)

for _ooxml_prefix, _ooxml_namespace in OOXML_NAMESPACES.items():
    ET.register_namespace(_ooxml_prefix, _ooxml_namespace)


class PerformancePackageService:
    async def list_categories(
        self,
        *,
        keyword: str = "",
        scene: str = "",
        power_rating: str = "",
        turbine_model: str = "",
        time_keyword: str = "",
        contract_year: str = "",
        delivery_year: str = "",
        operation_year: str = "",
        tag: str = "",
        status: str = "enabled",
        sort_by: str = "updatedAt",
        sort_order: str = "desc",
        page: int = 1,
        page_size: int = 20,
    ) -> dict[str, Any]:
        current_page = max(1, int(page or 1))
        current_page_size = max(1, min(100, int(page_size or 20)))
        offset = (current_page - 1) * current_page_size
        filters, params = self._category_filters(
            keyword=keyword,
            scene=scene,
            power_rating=power_rating,
            turbine_model=turbine_model,
            time_keyword=time_keyword,
            contract_year=contract_year,
            delivery_year=delivery_year,
            operation_year=operation_year,
            tag=tag,
            status=status,
        )
        where_sql = f"WHERE {' AND '.join(filters)}" if filters else ""
        order_sql = _category_order_sql(sort_by, sort_order)
        async with async_session() as session:
            await ensure_material_runtime_tables(session)
            await self._backfill_derived_item_fields(session)
            total_result = await session.execute(
                text(f"SELECT COUNT(*) FROM performance_categories c {where_sql}"),
                params,
            )
            total = int(total_result.scalar_one() or 0)
            rows = await session.execute(
                text(
                    f"""
                    SELECT
                        c.*,
                        (SELECT COUNT(*) FROM performance_items i WHERE i.category_id = c.id) AS item_count,
                        (
                            SELECT COALESCE(jsonb_agg(DISTINCT model_value ORDER BY model_value), '[]'::jsonb)
                            FROM performance_items i
                            CROSS JOIN LATERAL jsonb_array_elements_text(COALESCE(i.turbine_models, '[]'::jsonb)) AS model_value
                            WHERE i.category_id = c.id AND model_value <> ''
                        ) AS turbine_models,
                        (
                            SELECT COALESCE(jsonb_agg(DISTINCT i.contract_year ORDER BY i.contract_year), '[]'::jsonb)
                            FROM performance_items i
                            WHERE i.category_id = c.id AND i.contract_year IS NOT NULL
                        ) AS contract_years,
                        (
                            SELECT COALESCE(jsonb_agg(DISTINCT i.delivery_year ORDER BY i.delivery_year), '[]'::jsonb)
                            FROM performance_items i
                            WHERE i.category_id = c.id AND i.delivery_year IS NOT NULL
                        ) AS delivery_years,
                        (
                            SELECT COALESCE(jsonb_agg(DISTINCT i.operation_year ORDER BY i.operation_year), '[]'::jsonb)
                            FROM performance_items i
                            WHERE i.category_id = c.id AND i.operation_year IS NOT NULL
                        ) AS operation_years,
                        (SELECT COUNT(*) FROM performance_attachments a WHERE a.category_id = c.id) AS attachment_count,
                        (
                            SELECT COUNT(*)
                            FROM performance_attachments a
                            WHERE a.category_id = c.id AND a.attachment_type = :contract_attachment_type
                        ) AS contract_attachment_count,
                        (
                            SELECT COUNT(*)
                            FROM performance_item_attachments ia
                            WHERE ia.category_id = c.id AND ia.attachment_type = :item_contract_attachment_type
                        ) AS item_contract_attachment_count,
                        (
                            SELECT a.file_name
                            FROM performance_attachments a
                            WHERE a.category_id = c.id AND a.attachment_type = :summary_attachment_type
                            ORDER BY a.id DESC
                            LIMIT 1
                        ) AS summary_file_name,
                        (
                            SELECT a.file_name
                            FROM performance_attachments a
                            WHERE a.category_id = c.id AND a.attachment_type = :contract_attachment_type
                            ORDER BY a.id DESC
                            LIMIT 1
                        ) AS contract_file_name
                    FROM performance_categories c
                    {where_sql}
                    {order_sql}
                    LIMIT :limit OFFSET :offset
                    """
                ),
                {
                    **params,
                    "limit": current_page_size,
                    "offset": offset,
                    "summary_attachment_type": SUMMARY_ATTACHMENT_TYPE,
                    "contract_attachment_type": CONTRACT_ATTACHMENT_TYPE,
                    "item_contract_attachment_type": ITEM_CONTRACT_ATTACHMENT_TYPE,
                },
            )
        return {
            "items": [self._category_row_to_dict(row._mapping) for row in rows],
            "total": total,
            "page": current_page,
            "pageSize": current_page_size,
        }

    async def list_items(
        self,
        *,
        keyword: str = "",
        turbine_model: str = "",
        time_keyword: str = "",
        contract_year: str = "",
        delivery_year: str = "",
        operation_year: str = "",
        category_id: str = "",
        status: str = "enabled",
        sort_by: str = "updatedAt",
        sort_order: str = "desc",
        page: int = 1,
        page_size: int = 20,
    ) -> dict[str, Any]:
        current_page = max(1, int(page or 1))
        current_page_size = max(1, min(100, int(page_size or 20)))
        offset = (current_page - 1) * current_page_size
        filters, params = self._item_filters(
            keyword=keyword,
            turbine_model=turbine_model,
            time_keyword=time_keyword,
            contract_year=contract_year,
            delivery_year=delivery_year,
            operation_year=operation_year,
            category_id=category_id,
            status=status,
        )
        where_sql = f"WHERE {' AND '.join(filters)}" if filters else ""
        order_sql = _item_order_sql(sort_by, sort_order)
        async with async_session() as session:
            await ensure_material_runtime_tables(session)
            await self._backfill_derived_item_fields(session)
            total_result = await session.execute(
                text(
                    f"""
                    SELECT COUNT(*)
                    FROM performance_items i
                    JOIN performance_categories c ON c.id = i.category_id
                    {where_sql}
                    """
                ),
                params,
            )
            total = int(total_result.scalar_one() or 0)
            rows = await session.execute(
                text(
                    f"""
                    SELECT
                        i.*,
                        c.name AS category_name,
                        COALESCE(c.status, CASE WHEN c.review_status = 'disabled' THEN 'disabled' ELSE 'enabled' END) AS category_status
                    FROM performance_items i
                    JOIN performance_categories c ON c.id = i.category_id
                    {where_sql}
                    {order_sql}
                    LIMIT :limit OFFSET :offset
                    """
                ),
                {**params, "limit": current_page_size, "offset": offset},
            )
            item_payload: list[dict[str, Any]] = []
            numeric_ids: list[int] = []
            for row in rows:
                row_dict = self._item_row_to_dict(row._mapping)
                row_dict["categoryName"] = row._mapping.get("category_name") or ""
                row_dict["categoryStatus"] = row._mapping.get("category_status") or "enabled"
                row_dict["attachments"] = []
                item_payload.append(row_dict)
                numeric_ids.append(int(row._mapping.get("id") or 0))
            if numeric_ids:
                attachment_rows = await session.execute(
                    text(
                        """
                        SELECT *
                        FROM performance_item_attachments
                        WHERE item_id IN :item_ids
                        ORDER BY item_id ASC, created_at DESC, id DESC
                        """
                    ).bindparams(bindparam("item_ids", expanding=True)),
                    {"item_ids": numeric_ids},
                )
                attachments_by_item: dict[int, list[dict[str, Any]]] = {}
                for row in attachment_rows:
                    attachments_by_item.setdefault(int(row._mapping.get("item_id") or 0), []).append(
                        self._item_attachment_row_to_dict(row._mapping)
                    )
                for row_dict, numeric_item_id in zip(item_payload, numeric_ids):
                    row_dict["attachments"] = attachments_by_item.get(numeric_item_id, [])
        return {
            "items": item_payload,
            "total": total,
            "page": current_page,
            "pageSize": current_page_size,
        }

    def _item_filters(
        self,
        *,
        keyword: str,
        turbine_model: str,
        time_keyword: str,
        contract_year: str,
        delivery_year: str,
        operation_year: str,
        category_id: str,
        status: str,
    ) -> tuple[list[str], dict[str, Any]]:
        filters: list[str] = []
        params: dict[str, Any] = {}
        selected_status = str(status or "").strip()
        if selected_status in PERFORMANCE_CATEGORY_STATUSES:
            filters.append("COALESCE(c.status, CASE WHEN c.review_status = 'disabled' THEN 'disabled' ELSE 'enabled' END) = :status")
            params["status"] = selected_status
        kw = str(keyword or "").strip()
        if kw:
            filters.append(
                "("
                "i.project_name ILIKE :keyword OR i.customer_name ILIKE :keyword OR "
                "i.turbine_model ILIKE :keyword OR i.row_values::text ILIKE :keyword OR "
                "c.name ILIKE :keyword"
                ")"
            )
            params["keyword"] = f"%{kw}%"
        selected_model = str(turbine_model or "").strip()
        if selected_model:
            filters.append(
                "(i.turbine_model ILIKE :turbine_model OR i.turbine_models::text ILIKE :turbine_model OR i.row_values::text ILIKE :turbine_model)"
            )
            params["turbine_model"] = f"%{selected_model}%"
        selected_time = str(time_keyword or "").strip()
        if selected_time:
            filters.append(
                "(i.delivery_or_operation_time ILIKE :time_keyword OR i.time_facts::text ILIKE :time_keyword OR i.row_values::text ILIKE :time_keyword)"
            )
            params["time_keyword"] = f"%{selected_time}%"
        contract_year_value = _parse_year_filter(contract_year)
        if contract_year_value is not None:
            filters.append("i.contract_year = :contract_year")
            params["contract_year"] = contract_year_value
        delivery_year_value = _parse_year_filter(delivery_year)
        if delivery_year_value is not None:
            filters.append("i.delivery_year = :delivery_year")
            params["delivery_year"] = delivery_year_value
        operation_year_value = _parse_year_filter(operation_year)
        if operation_year_value is not None:
            filters.append("i.operation_year = :operation_year")
            params["operation_year"] = operation_year_value
        selected_category = str(category_id or "").strip()
        if selected_category:
            filters.append("i.category_id = :category_id")
            params["category_id"] = self._numeric_category_id(selected_category)
        return filters, params

    async def get_category(self, category_id: str) -> dict[str, Any]:
        numeric_id = self._numeric_category_id(category_id)
        async with async_session() as session:
            await ensure_material_runtime_tables(session)
            await self._backfill_derived_item_fields(session, category_id=numeric_id)
            await self._ensure_contract_item_attachments(session, category_id=numeric_id)
            category_result = await session.execute(
                text(
                        """
                    SELECT
                        c.*,
                        (SELECT COUNT(*) FROM performance_items i WHERE i.category_id = c.id) AS item_count,
                        (
                            SELECT COALESCE(jsonb_agg(DISTINCT model_value ORDER BY model_value), '[]'::jsonb)
                            FROM performance_items i
                            CROSS JOIN LATERAL jsonb_array_elements_text(COALESCE(i.turbine_models, '[]'::jsonb)) AS model_value
                            WHERE i.category_id = c.id AND model_value <> ''
                        ) AS turbine_models,
                        (
                            SELECT COALESCE(jsonb_agg(DISTINCT i.contract_year ORDER BY i.contract_year), '[]'::jsonb)
                            FROM performance_items i
                            WHERE i.category_id = c.id AND i.contract_year IS NOT NULL
                        ) AS contract_years,
                        (
                            SELECT COALESCE(jsonb_agg(DISTINCT i.delivery_year ORDER BY i.delivery_year), '[]'::jsonb)
                            FROM performance_items i
                            WHERE i.category_id = c.id AND i.delivery_year IS NOT NULL
                        ) AS delivery_years,
                        (
                            SELECT COALESCE(jsonb_agg(DISTINCT i.operation_year ORDER BY i.operation_year), '[]'::jsonb)
                            FROM performance_items i
                            WHERE i.category_id = c.id AND i.operation_year IS NOT NULL
                        ) AS operation_years,
                        (SELECT COUNT(*) FROM performance_attachments a WHERE a.category_id = c.id) AS attachment_count,
                        (
                            SELECT COUNT(*)
                            FROM performance_attachments a
                            WHERE a.category_id = c.id AND a.attachment_type = :contract_attachment_type
                        ) AS contract_attachment_count,
                        (
                            SELECT COUNT(*)
                            FROM performance_item_attachments ia
                            WHERE ia.category_id = c.id AND ia.attachment_type = :item_contract_attachment_type
                        ) AS item_contract_attachment_count,
                        (
                            SELECT a.file_name
                            FROM performance_attachments a
                            WHERE a.category_id = c.id AND a.attachment_type = :summary_attachment_type
                            ORDER BY a.id DESC
                            LIMIT 1
                        ) AS summary_file_name,
                        (
                            SELECT a.file_name
                            FROM performance_attachments a
                            WHERE a.category_id = c.id AND a.attachment_type = :contract_attachment_type
                            ORDER BY a.id DESC
                            LIMIT 1
                        ) AS contract_file_name
                    FROM performance_categories c
                    WHERE c.id = :id
                    """
                ),
                {
                    "id": numeric_id,
                    "summary_attachment_type": SUMMARY_ATTACHMENT_TYPE,
                    "contract_attachment_type": CONTRACT_ATTACHMENT_TYPE,
                    "item_contract_attachment_type": ITEM_CONTRACT_ATTACHMENT_TYPE,
                },
            )
            category_row = category_result.first()
            if category_row is None:
                raise PeripheralError(404, "业绩类别不存在。", "PERFORMANCE_CATEGORY_NOT_FOUND")
            item_rows = await session.execute(
                text(
                    """
                    SELECT *
                    FROM performance_items
                    WHERE category_id = :id
                    ORDER BY row_index ASC, id ASC
                    """
                ),
                {"id": numeric_id},
            )
            attachment_rows = await session.execute(
                text(
                    """
                    SELECT *
                    FROM performance_attachments
                    WHERE category_id = :id
                    ORDER BY created_at DESC, id DESC
                    """
                ),
                {"id": numeric_id},
            )
            item_attachment_rows = await session.execute(
                text(
                    """
                    SELECT *
                    FROM performance_item_attachments
                    WHERE category_id = :id
                    ORDER BY item_id ASC, created_at DESC, id DESC
                    """
                ),
                {"id": numeric_id},
            )
            item_attachments: dict[int, list[dict[str, Any]]] = {}
            for row in item_attachment_rows:
                item_attachments.setdefault(int(row._mapping.get("item_id") or 0), []).append(self._item_attachment_row_to_dict(row._mapping))
            item_payload = []
            for row in item_rows:
                row_dict = self._item_row_to_dict(row._mapping)
                row_dict["attachments"] = item_attachments.get(int(row._mapping.get("id") or 0), [])
                item_payload.append(row_dict)
        return {
            "item": self._category_row_to_dict(category_row._mapping),
            "rows": item_payload,
            "attachments": [self._attachment_row_to_dict(row._mapping) for row in attachment_rows],
        }

    async def preview_summary(self, upload: Any) -> dict[str, Any]:
        content, file_name, _mime_type = await _read_upload(upload)
        parsed = parse_performance_summary_docx(content, file_name=file_name)
        return {"message": "业绩汇总表解析完成", "preview": parsed}

    async def import_summary(
        self,
        upload: Any,
        *,
        contract_uploads: list[Any] | None = None,
        category_name: str = "",
        scene: str = "",
        power_rating: str = "",
        tags: Any = None,
        scope: str = "standard",
        review_status: str = "draft",
    ) -> dict[str, Any]:
        content, file_name, mime_type = await _read_upload(upload)
        contract_files: list[tuple[bytes, str, str]] = []
        for contract_upload in contract_uploads or []:
            contract_files.append(await _read_upload(contract_upload))
        if not contract_files:
            raise PeripheralError(
                400,
                "请同时上传合同附件：汇总表与合同需一次导入，不能拆开提交。",
                "PERFORMANCE_CONTRACT_FILES_REQUIRED",
            )
        parsed = parse_performance_summary_docx(content, file_name=file_name)
        category_name = str(category_name or parsed.get("categoryName") or Path(file_name).stem).strip()
        if not category_name:
            raise PeripheralError(400, "业绩类别名称不能为空。", "PERFORMANCE_CATEGORY_NAME_REQUIRED")
        selected_scene = str(scene or parsed.get("scene") or "").strip()
        selected_power_rating = str(power_rating or parsed.get("powerRating") or "").strip()
        selected_scope = scope if scope in PERFORMANCE_CATEGORY_SCOPES else "standard"
        selected_review_status = review_status if review_status in PERFORMANCE_CATEGORY_REVIEW_STATUSES else "draft"
        normalized_tags = normalize_material_tags(tags)

        async with async_session() as session:
            await ensure_material_runtime_tables(session)
            category_result = await session.execute(
                text(
                    """
                    INSERT INTO performance_categories (
                        name, scene, power_rating, summary, field_schema, tags, scope, review_status
                    )
                    VALUES (
                        :name, :scene, :power_rating, :summary, CAST(:field_schema AS JSONB),
                        CAST(:tags AS JSONB), :scope, :review_status
                    )
                    RETURNING *
                    """
                ),
                {
                    "name": category_name,
                    "scene": selected_scene,
                    "power_rating": selected_power_rating,
                    "summary": str(parsed.get("summary") or ""),
                    "field_schema": _json(parsed.get("fieldSchema") or []),
                    "tags": _json(normalized_tags),
                    "scope": selected_scope,
                    "review_status": selected_review_status,
                },
            )
            category_row = category_result.first()
            if category_row is None:
                raise PeripheralError(500, "业绩类别创建失败。", "PERFORMANCE_CATEGORY_CREATE_FAILED")
            numeric_id = int(category_row._mapping["id"])
            if parsed.get("rows"):
                await session.execute(
                    text(
                        """
                        INSERT INTO performance_items (
                            category_id, row_index, project_name, customer_name, turbine_model,
                            turbine_models,
                            contract_quantity, trial_operation_quantity, commissioned_capacity_mw,
                            delivery_or_operation_time, contract_year, delivery_year, operation_year,
                            time_facts, contact_info, row_values
                        )
                        VALUES (
                            :category_id, :row_index, :project_name, :customer_name, :turbine_model,
                            CAST(:turbine_models AS JSONB),
                            :contract_quantity, :trial_operation_quantity, :commissioned_capacity_mw,
                            :delivery_or_operation_time, :contract_year, :delivery_year, :operation_year,
                            CAST(:time_facts AS JSONB), :contact_info, CAST(:row_values AS JSONB)
                        )
                        """
                    ),
                    [
                        {
                            "category_id": numeric_id,
                            "row_index": int(row.get("rowIndex") or index + 1),
                            "project_name": row.get("projectName") or "",
                            "customer_name": row.get("customerName") or "",
                            "turbine_model": row.get("turbineModel") or "",
                            "turbine_models": _json(row.get("turbineModels") or []),
                            "contract_quantity": row.get("contractQuantity") or "",
                            "trial_operation_quantity": row.get("trialOperationQuantity") or "",
                            "commissioned_capacity_mw": row.get("commissionedCapacityMw") or "",
                            "delivery_or_operation_time": row.get("deliveryOrOperationTime") or "",
                            "contract_year": row.get("contractYear"),
                            "delivery_year": row.get("deliveryYear"),
                            "operation_year": row.get("operationYear"),
                            "time_facts": _json(row.get("timeFacts") or {}),
                            "contact_info": row.get("contactInfo") or "",
                            "row_values": _json(row.get("values") or {}),
                        }
                        for index, row in enumerate(parsed.get("rows") or [])
                    ],
                )

            bucket = settings.minio_buckets["materials"]
            object_key = f"performance-categories/PERCAT-{numeric_id:04d}/summary/{_safe_file_name(file_name)}"
            uploaded_objects: list[tuple[str, str]] = []
            try:
                minio_client.put_object(bucket, object_key, content, content_type=mime_type)
                uploaded_objects.append((bucket, object_key))
                await session.execute(
                    text(
                        """
                        INSERT INTO performance_attachments (
                            category_id, attachment_type, file_name, minio_key, minio_bucket,
                            mime_type, size_bytes
                        )
                        VALUES (
                            :category_id, :attachment_type, :file_name, :minio_key, :minio_bucket,
                            :mime_type, :size_bytes
                        )
                        """
                    ),
                    {
                        "category_id": numeric_id,
                        "attachment_type": SUMMARY_ATTACHMENT_TYPE,
                        "file_name": file_name,
                        "minio_key": object_key,
                        "minio_bucket": bucket,
                        "mime_type": mime_type,
                        "size_bytes": len(content),
                    },
                )
                for contract_index, (contract_content, contract_file_name, contract_mime_type) in enumerate(contract_files):
                    contract_object_key = (
                        f"performance-categories/PERCAT-{numeric_id:04d}/{CONTRACT_ATTACHMENT_TYPE}/"
                        f"{contract_index + 1:02d}-{_safe_file_name(contract_file_name)}"
                    )
                    minio_client.put_object(bucket, contract_object_key, contract_content, content_type=contract_mime_type)
                    uploaded_objects.append((bucket, contract_object_key))
                    contract_attachment_result = await session.execute(
                        text(
                            """
                            INSERT INTO performance_attachments (
                                category_id, attachment_type, file_name, minio_key, minio_bucket,
                                mime_type, size_bytes
                            )
                            VALUES (
                                :category_id, :attachment_type, :file_name, :minio_key, :minio_bucket,
                                :mime_type, :size_bytes
                            )
                            RETURNING id
                            """
                        ),
                        {
                            "category_id": numeric_id,
                            "attachment_type": CONTRACT_ATTACHMENT_TYPE,
                            "file_name": contract_file_name,
                            "minio_key": contract_object_key,
                            "minio_bucket": bucket,
                            "mime_type": contract_mime_type,
                            "size_bytes": len(contract_content),
                        },
                    )
                    contract_attachment_row = contract_attachment_result.first()
                    if contract_attachment_row is None:
                        raise PeripheralError(500, "业绩合同附件保存失败。", "PERFORMANCE_ATTACHMENT_CREATE_FAILED")
                    if contract_file_name.lower().endswith(".docx"):
                        uploaded_objects.extend(
                            await self._replace_contract_item_attachments_for_source(
                                session,
                                category_id=numeric_id,
                                source_attachment_id=int(contract_attachment_row._mapping["id"]),
                                content=contract_content,
                                source_file_name=contract_file_name,
                            )
                        )
                await session.commit()
            except Exception:
                for uploaded_bucket, uploaded_key in uploaded_objects:
                    try:
                        minio_client.remove_object(uploaded_bucket, uploaded_key)
                    except Exception as exc:  # pragma: no cover - cleanup should not mask original error
                        logger.warning("Failed to remove performance import object %s/%s: %s", uploaded_bucket, uploaded_key, exc)
                raise
        return {
            "message": "业绩包已导入",
            "item": await self.get_category(f"PERCAT-{numeric_id:04d}"),
        }

    async def upload_attachment(
        self,
        category_id: str,
        upload: Any,
        *,
        attachment_type: str = CONTRACT_ATTACHMENT_TYPE,
    ) -> dict[str, Any]:
        numeric_id = self._numeric_category_id(category_id)
        content, file_name, mime_type = await _read_upload(upload)
        normalized_type = str(attachment_type or CONTRACT_ATTACHMENT_TYPE).strip() or CONTRACT_ATTACHMENT_TYPE
        if normalized_type not in {SUMMARY_ATTACHMENT_TYPE, CONTRACT_ATTACHMENT_TYPE, "other"}:
            normalized_type = "other"
        object_key = f"performance-categories/PERCAT-{numeric_id:04d}/{normalized_type}/{_safe_file_name(file_name)}"
        async with async_session() as session:
            await ensure_material_runtime_tables(session)
            category_result = await session.execute(
                text("SELECT id FROM performance_categories WHERE id = :id"),
                {"id": numeric_id},
            )
            if category_result.first() is None:
                raise PeripheralError(404, "业绩类别不存在。", "PERFORMANCE_CATEGORY_NOT_FOUND")
            uploaded_child_objects: list[tuple[str, str]] = []
            try:
                minio_client.put_object(
                    settings.minio_buckets["materials"],
                    object_key,
                    content,
                    content_type=mime_type,
                )
                attachment_result = await session.execute(
                    text(
                        """
                        INSERT INTO performance_attachments (
                            category_id, attachment_type, file_name, minio_key, minio_bucket,
                            mime_type, size_bytes
                        )
                        VALUES (
                            :category_id, :attachment_type, :file_name, :minio_key, :minio_bucket,
                            :mime_type, :size_bytes
                        )
                        RETURNING *
                        """
                    ),
                    {
                        "category_id": numeric_id,
                        "attachment_type": normalized_type,
                        "file_name": file_name,
                        "minio_key": object_key,
                        "minio_bucket": settings.minio_buckets["materials"],
                        "mime_type": mime_type,
                        "size_bytes": len(content),
                    },
                )
                attachment_row = attachment_result.first()
                if attachment_row is None:
                    raise PeripheralError(500, "业绩附件保存失败。", "PERFORMANCE_ATTACHMENT_CREATE_FAILED")
                if normalized_type == CONTRACT_ATTACHMENT_TYPE and file_name.lower().endswith(".docx"):
                    uploaded_child_objects = await self._replace_contract_item_attachments_for_source(
                        session,
                        category_id=numeric_id,
                        source_attachment_id=int(attachment_row._mapping["id"]),
                        content=content,
                        source_file_name=file_name,
                    )
                await session.execute(
                    text("UPDATE performance_categories SET updated_at = NOW() WHERE id = :id"),
                    {"id": numeric_id},
                )
                await session.commit()
            except Exception:
                minio_client.remove_object(settings.minio_buckets["materials"], object_key)
                for bucket, key in uploaded_child_objects:
                    try:
                        minio_client.remove_object(bucket, key)
                    except Exception as exc:  # pragma: no cover - cleanup should not mask original error
                        logger.warning("Failed to remove split performance contract object %s/%s: %s", bucket, key, exc)
                raise
        return {"message": "业绩附件已上传", "item": await self.get_category(category_id)}

    async def download_attachment(self, category_id: str, attachment_id: str) -> dict[str, Any]:
        numeric_category_id = self._numeric_category_id(category_id)
        numeric_attachment_id = self._numeric_attachment_id(attachment_id)
        async with async_session() as session:
            await ensure_material_runtime_tables(session)
            result = await session.execute(
                text(
                    """
                    SELECT *
                    FROM performance_attachments
                    WHERE id = :attachment_id AND category_id = :category_id
                    """
                ),
                {"attachment_id": numeric_attachment_id, "category_id": numeric_category_id},
            )
            row = result.first()
        if row is None:
            raise PeripheralError(404, "业绩附件不存在。", "PERFORMANCE_ATTACHMENT_NOT_FOUND")
        item = self._attachment_row_to_dict(row._mapping)
        return {
            "bucket": item.get("minioBucket") or settings.minio_buckets["materials"],
            "key": item.get("minioKey") or "",
            "fileName": item.get("fileName") or f"{item['id']}.docx",
            "mimeType": item.get("mimeType") or "application/octet-stream",
            "attachmentType": item.get("attachmentType") or "",
            "sizeBytes": int(item.get("sizeBytes") or 0),
        }

    async def preview_attachment(
        self,
        category_id: str,
        attachment_id: str,
        *,
        browser_base_url: str = "",
        onlyoffice_base_url: str = "",
    ) -> dict[str, Any]:
        payload = await self.download_attachment(category_id, attachment_id)
        file_name = str(payload.get("fileName") or f"{attachment_id}.docx")
        file_path = f"/api/materials/performance/categories/{category_id}/attachments/{attachment_id}"
        return self._preview_payload(
            attachment_id=attachment_id,
            file_name=file_name,
            file_path=file_path,
            minio_key=str(payload.get("key") or ""),
            browser_base_url=browser_base_url,
            onlyoffice_base_url=onlyoffice_base_url,
        )

    async def download_item_attachment(self, category_id: str, item_id: str, attachment_id: str) -> dict[str, Any]:
        numeric_category_id = self._numeric_category_id(category_id)
        numeric_item_id = self._numeric_item_id(item_id)
        numeric_attachment_id = self._numeric_item_attachment_id(attachment_id)
        async with async_session() as session:
            await ensure_material_runtime_tables(session)
            result = await session.execute(
                text(
                    """
                    SELECT *
                    FROM performance_item_attachments
                    WHERE id = :attachment_id
                      AND item_id = :item_id
                      AND category_id = :category_id
                    """
                ),
                {
                    "attachment_id": numeric_attachment_id,
                    "item_id": numeric_item_id,
                    "category_id": numeric_category_id,
                },
            )
            row = result.first()
        if row is None:
            raise PeripheralError(404, "项目合同附件不存在。", "PERFORMANCE_ITEM_ATTACHMENT_NOT_FOUND")
        item = self._item_attachment_row_to_dict(row._mapping)
        return {
            "bucket": item.get("minioBucket") or settings.minio_buckets["materials"],
            "key": item.get("minioKey") or "",
            "fileName": item.get("fileName") or f"{item['id']}.docx",
            "mimeType": item.get("mimeType") or "application/octet-stream",
            "attachmentType": item.get("attachmentType") or "",
            "sizeBytes": int(item.get("sizeBytes") or 0),
        }

    async def preview_item_attachment(
        self,
        category_id: str,
        item_id: str,
        attachment_id: str,
        *,
        browser_base_url: str = "",
        onlyoffice_base_url: str = "",
    ) -> dict[str, Any]:
        payload = await self.download_item_attachment(category_id, item_id, attachment_id)
        file_name = str(payload.get("fileName") or f"{attachment_id}.docx")
        file_path = f"/api/materials/performance/categories/{category_id}/items/{item_id}/attachments/{attachment_id}"
        return self._preview_payload(
            attachment_id=attachment_id,
            file_name=file_name,
            file_path=file_path,
            minio_key=str(payload.get("key") or ""),
            browser_base_url=browser_base_url,
            onlyoffice_base_url=onlyoffice_base_url,
        )

    def _preview_payload(
        self,
        *,
        attachment_id: str,
        file_name: str,
        file_path: str,
        minio_key: str,
        browser_base_url: str = "",
        onlyoffice_base_url: str = "",
    ) -> dict[str, Any]:
        suffix = Path(file_name).suffix.lower().lstrip(".") or "docx"
        document_type = (
            "cell" if suffix in {"xls", "xlsx", "csv"}
            else "slide" if suffix in {"ppt", "pptx"}
            else "pdf" if suffix == "pdf"
            else "word"
        )
        if suffix not in {"doc", "docx", "xls", "xlsx", "csv", "ppt", "pptx", "pdf"}:
            raise PeripheralError(400, "该附件类型暂不支持在线预览。", "PERFORMANCE_ATTACHMENT_PREVIEW_UNSUPPORTED")
        browser_file_url = f"{browser_base_url.rstrip('/')}{file_path}" if browser_base_url else file_path
        onlyoffice_file_url = (
            f"{onlyoffice_base_url.rstrip('/')}{file_path}"
            if onlyoffice_base_url
            else browser_file_url
        )
        digest = hashlib.sha1("|".join([attachment_id, minio_key, file_name]).encode("utf-8")).hexdigest()[:16]
        return {
            "status": "ready",
            "attachmentId": attachment_id,
            "fileName": file_name,
            "fileType": suffix,
            "documentType": document_type,
            "fileUrl": browser_file_url,
            "onlyoffice": {
                "documentKey": f"performance-{attachment_id}-{digest}",
                "title": file_name,
                "fileUrl": onlyoffice_file_url,
                "browserFileUrl": browser_file_url,
                "documentServerFileUrl": onlyoffice_file_url,
                "fileType": suffix,
                "documentType": document_type,
                "user": {
                    "id": "user-1",
                    "name": "当前用户",
                },
            },
        }

    async def update_category_status(self, category_id: str, status: str) -> dict[str, Any]:
        numeric_id = self._numeric_category_id(category_id)
        selected_status = str(status or "").strip()
        if selected_status not in PERFORMANCE_CATEGORY_STATUSES:
            raise PeripheralError(400, "业绩类别状态无效。", "PERFORMANCE_CATEGORY_STATUS_INVALID")
        async with async_session() as session:
            await ensure_material_runtime_tables(session)
            result = await session.execute(
                text(
                    """
                    UPDATE performance_categories
                    SET status = CAST(:status AS VARCHAR),
                        review_status = CASE
                            WHEN CAST(:status AS VARCHAR) = 'enabled' AND review_status = 'disabled' THEN 'draft'
                            ELSE review_status
                        END,
                        updated_at = NOW()
                    WHERE id = :id
                    RETURNING *
                    """
                ),
                {"id": numeric_id, "status": selected_status},
            )
            row = result.first()
            if row is None:
                raise PeripheralError(404, "业绩类别不存在。", "PERFORMANCE_CATEGORY_NOT_FOUND")
            await session.commit()
        return {"message": "业绩类别已启用" if selected_status == "enabled" else "业绩类别已停用", "item": self._category_row_to_dict(row._mapping)}

    async def disable_category(self, category_id: str) -> dict[str, Any]:
        return await self.update_category_status(category_id, "disabled")

    async def delete_category(self, category_id: str, confirm_name: str = "") -> dict[str, Any]:
        numeric_id = self._numeric_category_id(category_id)
        selected_confirm_name = str(confirm_name or "").strip()
        async with async_session() as session:
            await ensure_material_runtime_tables(session)
            category_result = await session.execute(
                text(
                    """
                    SELECT id, name
                    FROM performance_categories
                    WHERE id = :id
                    """
                ),
                {"id": numeric_id},
            )
            category_row = category_result.first()
            if category_row is None:
                raise PeripheralError(404, "业绩类别不存在。", "PERFORMANCE_CATEGORY_NOT_FOUND")
            category_name = str(category_row._mapping.get("name") or "").strip()
            if not selected_confirm_name or selected_confirm_name != category_name:
                raise PeripheralError(400, "请输入完整业绩类别名称后再删除。", "PERFORMANCE_CATEGORY_DELETE_CONFIRM_REQUIRED")
            attachment_result = await session.execute(
                text(
                    """
                    SELECT minio_bucket, minio_key
                    FROM performance_attachments
                    WHERE category_id = :id
                    """
                ),
                {"id": numeric_id},
            )
            attachments = [dict(row._mapping) for row in attachment_result]
            item_attachment_result = await session.execute(
                text(
                    """
                    SELECT minio_bucket, minio_key
                    FROM performance_item_attachments
                    WHERE category_id = :id
                    """
                ),
                {"id": numeric_id},
            )
            attachments.extend(dict(row._mapping) for row in item_attachment_result)
            result = await session.execute(
                text(
                    """
                    DELETE FROM performance_categories
                    WHERE id = :id
                    RETURNING id
                    """
                ),
                {"id": numeric_id},
            )
            if result.first() is None:
                raise PeripheralError(404, "业绩类别不存在。", "PERFORMANCE_CATEGORY_NOT_FOUND")
            await session.commit()

        for attachment in attachments:
            bucket = attachment.get("minio_bucket") or settings.minio_buckets["materials"]
            key = attachment.get("minio_key") or ""
            if not key:
                continue
            try:
                minio_client.remove_object(bucket, key)
            except Exception as exc:  # pragma: no cover - object cleanup should not roll back DB delete
                logger.warning("Failed to remove performance attachment object %s/%s: %s", bucket, key, exc)
        return {"message": "业绩类别已删除", "id": f"PERCAT-{numeric_id:04d}"}

    def _category_filters(
        self,
        *,
        keyword: str,
        scene: str,
        power_rating: str,
        turbine_model: str,
        time_keyword: str,
        contract_year: str,
        delivery_year: str,
        operation_year: str,
        tag: str,
        status: str,
    ) -> tuple[list[str], dict[str, Any]]:
        filters: list[str] = []
        params: dict[str, Any] = {}
        selected_status = str(status or "").strip()
        if selected_status in PERFORMANCE_CATEGORY_STATUSES:
            filters.append("COALESCE(c.status, CASE WHEN c.review_status = 'disabled' THEN 'disabled' ELSE 'enabled' END) = :status")
            params["status"] = selected_status
        kw = str(keyword or "").strip()
        if kw:
            filters.append(
                "("
                "c.name ILIKE :keyword OR c.scene ILIKE :keyword OR c.power_rating ILIKE :keyword OR "
                "c.summary ILIKE :keyword OR EXISTS ("
                "SELECT 1 FROM performance_items i WHERE i.category_id = c.id AND "
                "(i.project_name ILIKE :keyword OR i.customer_name ILIKE :keyword OR "
                "i.turbine_model ILIKE :keyword OR i.row_values::text ILIKE :keyword)"
                ")"
                ")"
            )
            params["keyword"] = f"%{kw}%"
        selected_scene = str(scene or "").strip()
        if selected_scene:
            filters.append("c.scene ILIKE :scene")
            params["scene"] = f"%{selected_scene}%"
        selected_power = str(power_rating or "").strip()
        if selected_power:
            filters.append("c.power_rating ILIKE :power_rating")
            params["power_rating"] = f"%{selected_power}%"
        selected_model = str(turbine_model or "").strip()
        if selected_model:
            filters.append(
                """
                EXISTS (
                    SELECT 1
                    FROM performance_items i
                    WHERE i.category_id = c.id
                      AND (
                        i.turbine_model ILIKE :turbine_model OR
                        i.turbine_models::text ILIKE :turbine_model OR
                        i.row_values::text ILIKE :turbine_model
                      )
                )
                """
            )
            params["turbine_model"] = f"%{selected_model}%"
        selected_time = str(time_keyword or "").strip()
        if selected_time:
            filters.append(
                """
                EXISTS (
                    SELECT 1
                    FROM performance_items i
                    WHERE i.category_id = c.id
                      AND (
                        i.delivery_or_operation_time ILIKE :time_keyword OR
                        i.time_facts::text ILIKE :time_keyword OR
                        i.row_values::text ILIKE :time_keyword
                      )
                )
                """
            )
            params["time_keyword"] = f"%{selected_time}%"
        contract_year_value = _parse_year_filter(contract_year)
        if contract_year_value is not None:
            filters.append(
                """
                EXISTS (
                    SELECT 1
                    FROM performance_items i
                    WHERE i.category_id = c.id AND i.contract_year = :contract_year
                )
                """
            )
            params["contract_year"] = contract_year_value
        delivery_year_value = _parse_year_filter(delivery_year)
        if delivery_year_value is not None:
            filters.append(
                """
                EXISTS (
                    SELECT 1
                    FROM performance_items i
                    WHERE i.category_id = c.id AND i.delivery_year = :delivery_year
                )
                """
            )
            params["delivery_year"] = delivery_year_value
        operation_year_value = _parse_year_filter(operation_year)
        if operation_year_value is not None:
            filters.append(
                """
                EXISTS (
                    SELECT 1
                    FROM performance_items i
                    WHERE i.category_id = c.id AND i.operation_year = :operation_year
                )
                """
            )
            params["operation_year"] = operation_year_value
        selected_tag = str(tag or "").strip()
        if selected_tag:
            filters.append("c.tags @> CAST(:tag_json AS JSONB)")
            params["tag_json"] = _json([selected_tag])
        return filters, params

    async def _backfill_derived_item_fields(self, session: Any, *, category_id: int | None = None) -> None:
        filters = [
            "("
            "COALESCE(jsonb_array_length(COALESCE(turbine_models, '[]'::jsonb)), 0) = 0 OR "
            "COALESCE(time_facts, '{}'::jsonb) = '{}'::jsonb"
            ")"
        ]
        params: dict[str, Any] = {}
        if category_id is not None:
            filters.append("category_id = :category_id")
            params["category_id"] = category_id
        rows = await session.execute(
            text(
                f"""
                SELECT id, row_values, turbine_model, delivery_or_operation_time
                FROM performance_items
                WHERE {' AND '.join(filters)}
                LIMIT 2000
                """
            ),
            params,
        )
        updates: list[dict[str, Any]] = []
        for row in rows:
            row_dict = dict(row._mapping)
            row_values = dict(row_dict.get("row_values") or {})
            if not row_values:
                turbine_model = str(row_dict.get("turbine_model") or "").strip()
                delivery_or_operation_time = str(row_dict.get("delivery_or_operation_time") or "").strip()
                if turbine_model:
                    row_values["型号"] = turbine_model
                if delivery_or_operation_time:
                    row_values["交货期/投运时间"] = delivery_or_operation_time
            core = _core_values(row_values)
            updates.append(
                {
                    "id": row_dict["id"],
                    "turbine_models": _json(core.get("turbineModels") or []),
                    "contract_year": core.get("contractYear"),
                    "delivery_year": core.get("deliveryYear"),
                    "operation_year": core.get("operationYear"),
                    "time_facts": _json(core.get("timeFacts") or {}),
                }
            )
        if not updates:
            return
        await session.execute(
            text(
                """
                UPDATE performance_items
                SET turbine_models = CAST(:turbine_models AS JSONB),
                    contract_year = :contract_year,
                    delivery_year = :delivery_year,
                    operation_year = :operation_year,
                    time_facts = CAST(:time_facts AS JSONB),
                    updated_at = NOW()
                WHERE id = :id
                """
            ),
            updates,
        )
        await session.commit()

    async def _ensure_contract_item_attachments(self, session: Any, *, category_id: int) -> None:
        source_rows = await session.execute(
            text(
                """
                SELECT *
                FROM performance_attachments
                WHERE category_id = :category_id
                  AND attachment_type = :attachment_type
                  AND LOWER(file_name) LIKE '%.docx'
                ORDER BY id ASC
                LIMIT 20
                """
            ),
            {"category_id": category_id, "attachment_type": CONTRACT_ATTACHMENT_TYPE},
        )
        for row in source_rows:
            source = dict(row._mapping)
            existing_result = await session.execute(
                text(
                    """
                    SELECT COUNT(*)
                    FROM performance_item_attachments
                    WHERE source_attachment_id = :source_attachment_id
                      AND COALESCE(format_version, 1) >= :format_version
                    """
                ),
                {
                    "source_attachment_id": int(source.get("id") or 0),
                    "format_version": ITEM_CONTRACT_FORMAT_VERSION,
                },
            )
            if int(existing_result.scalar_one() or 0) > 0:
                continue
            try:
                content = minio_client.get_object(
                    source.get("minio_bucket") or settings.minio_buckets["materials"],
                    source.get("minio_key") or "",
                )
                await self._replace_contract_item_attachments_for_source(
                    session,
                    category_id=category_id,
                    source_attachment_id=int(source["id"]),
                    content=content,
                    source_file_name=str(source.get("file_name") or "合同附件.docx"),
                )
                await session.commit()
            except Exception as exc:  # pragma: no cover - lazy split should not break detail loading
                logger.warning("Failed to lazily split performance contract attachment %s: %s", source.get("id"), exc)
                await session.rollback()
                await ensure_material_runtime_tables(session)

    async def _replace_contract_item_attachments_for_source(
        self,
        session: Any,
        *,
        category_id: int,
        source_attachment_id: int,
        content: bytes,
        source_file_name: str,
    ) -> list[tuple[str, str]]:
        old_rows = await session.execute(
            text(
                """
                DELETE FROM performance_item_attachments
                WHERE source_attachment_id = :source_attachment_id
                RETURNING minio_bucket, minio_key
                """
            ),
            {"source_attachment_id": source_attachment_id},
        )
        for row in old_rows:
            bucket = row._mapping.get("minio_bucket") or settings.minio_buckets["materials"]
            key = row._mapping.get("minio_key") or ""
            if not key:
                continue
            try:
                minio_client.remove_object(bucket, key)
            except Exception as exc:  # pragma: no cover - cleanup should not block a replacement split
                logger.warning("Failed to remove old split performance contract object %s/%s: %s", bucket, key, exc)

        chunks = split_performance_contract_docx(content, file_name=source_file_name)
        if not chunks:
            return []
        item_rows = await session.execute(
            text(
                """
                SELECT id, row_index, project_name, customer_name, turbine_model, turbine_models, row_values
                FROM performance_items
                WHERE category_id = :category_id
                ORDER BY row_index ASC, id ASC
                """
            ),
            {"category_id": category_id},
        )
        items = [self._item_row_to_dict(row._mapping) for row in item_rows]
        matches = match_contract_chunks_to_items(chunks, items)
        uploaded_objects: list[tuple[str, str]] = []
        inserts: list[dict[str, Any]] = []
        bucket = settings.minio_buckets["materials"]
        for match in matches:
            item = match.get("item") or {}
            chunk = match.get("chunk") or {}
            item_numeric_id = self._numeric_item_id(str(item.get("id") or ""))
            if not item_numeric_id:
                continue
            chunk_content = chunk.get("content")
            if not isinstance(chunk_content, bytes):
                chunk_content = render_contract_item_docx(
                    chunk,
                    output_title=str(item.get("projectName") or chunk.get("title") or ""),
                )
            file_name = _contract_item_file_name(
                item.get("rowIndex") or 0,
                item.get("projectName") or chunk.get("title") or Path(source_file_name).stem,
                chunk.get("index") or 0,
            )
            object_key = (
                f"performance-categories/PERCAT-{category_id:04d}/"
                f"item-contracts/PERITEM-{item_numeric_id:04d}/"
                f"SRC-{source_attachment_id:04d}-{_safe_file_name(file_name)}"
            )
            minio_client.put_object(bucket, object_key, chunk_content, content_type=DOCX_MIME_TYPE)
            uploaded_objects.append((bucket, object_key))
            inserts.append(
                {
                    "category_id": category_id,
                    "item_id": item_numeric_id,
                    "source_attachment_id": source_attachment_id,
                    "attachment_type": ITEM_CONTRACT_ATTACHMENT_TYPE,
                    "file_name": file_name,
                    "minio_key": object_key,
                    "minio_bucket": bucket,
                    "mime_type": DOCX_MIME_TYPE,
                    "size_bytes": len(chunk_content),
                    "format_version": ITEM_CONTRACT_FORMAT_VERSION,
                    "match_confidence": int(match.get("confidence") or 0),
                    "match_method": str(match.get("method") or ""),
                    "source_title": str(chunk.get("title") or ""),
                    "source_block_start": chunk.get("blockStart"),
                    "source_block_end": chunk.get("blockEnd"),
                }
            )
        if inserts:
            await session.execute(
                text(
                    """
                    INSERT INTO performance_item_attachments (
                        category_id, item_id, source_attachment_id, attachment_type, file_name,
                        minio_key, minio_bucket, mime_type, size_bytes, format_version, match_confidence,
                        match_method, source_title, source_block_start, source_block_end
                    )
                    VALUES (
                        :category_id, :item_id, :source_attachment_id, :attachment_type, :file_name,
                        :minio_key, :minio_bucket, :mime_type, :size_bytes, :format_version, :match_confidence,
                        :match_method, :source_title, :source_block_start, :source_block_end
                    )
                    """
                ),
                inserts,
            )
        return uploaded_objects

    def _category_row_to_dict(self, row: Any) -> dict[str, Any]:
        row_dict = dict(row)
        numeric_id = int(row_dict.get("id") or 0)
        return {
            "id": f"PERCAT-{numeric_id:04d}",
            "name": row_dict.get("name") or "",
            "scene": row_dict.get("scene") or "",
            "powerRating": row_dict.get("power_rating") or "",
            "summary": row_dict.get("summary") or "",
            "fieldSchema": list(row_dict.get("field_schema") or []),
            "turbineModels": list(row_dict.get("turbine_models") or []),
            "contractYears": list(row_dict.get("contract_years") or []),
            "deliveryYears": list(row_dict.get("delivery_years") or []),
            "operationYears": list(row_dict.get("operation_years") or []),
            "tags": list(row_dict.get("tags") or []),
            "scope": row_dict.get("scope") or "standard",
            "status": row_dict.get("status") or ("disabled" if row_dict.get("review_status") == "disabled" else "enabled"),
            "reviewStatus": row_dict.get("review_status") or "draft",
            "itemCount": int(row_dict.get("item_count") or 0),
            "attachmentCount": int(row_dict.get("attachment_count") or 0),
            "contractAttachmentCount": int(row_dict.get("contract_attachment_count") or 0),
            "itemContractAttachmentCount": int(row_dict.get("item_contract_attachment_count") or 0),
            "summaryFileName": row_dict.get("summary_file_name") or "",
            "contractFileName": row_dict.get("contract_file_name") or "",
            "createdAt": row_dict.get("created_at").isoformat() if row_dict.get("created_at") else "",
            "updatedAt": row_dict.get("updated_at").isoformat() if row_dict.get("updated_at") else "",
        }

    def _item_row_to_dict(self, row: Any) -> dict[str, Any]:
        row_dict = dict(row)
        numeric_id = int(row_dict.get("id") or 0)
        return {
            "id": f"PERITEM-{numeric_id:04d}",
            "categoryId": f"PERCAT-{int(row_dict.get('category_id') or 0):04d}",
            "rowIndex": int(row_dict.get("row_index") or 0),
            "projectName": row_dict.get("project_name") or "",
            "customerName": row_dict.get("customer_name") or "",
            "turbineModel": row_dict.get("turbine_model") or "",
            "turbineModels": list(row_dict.get("turbine_models") or []),
            "contractQuantity": row_dict.get("contract_quantity") or "",
            "trialOperationQuantity": row_dict.get("trial_operation_quantity") or "",
            "commissionedCapacityMw": row_dict.get("commissioned_capacity_mw") or "",
            "deliveryOrOperationTime": row_dict.get("delivery_or_operation_time") or "",
            "contractYear": row_dict.get("contract_year"),
            "deliveryYear": row_dict.get("delivery_year"),
            "operationYear": row_dict.get("operation_year"),
            "timeFacts": dict(row_dict.get("time_facts") or {}),
            "contactInfo": row_dict.get("contact_info") or "",
            "values": dict(row_dict.get("row_values") or {}),
            "createdAt": row_dict.get("created_at").isoformat() if row_dict.get("created_at") else "",
            "updatedAt": row_dict.get("updated_at").isoformat() if row_dict.get("updated_at") else "",
        }

    def _attachment_row_to_dict(self, row: Any) -> dict[str, Any]:
        row_dict = dict(row)
        numeric_id = int(row_dict.get("id") or 0)
        return {
            "id": f"PERATT-{numeric_id:04d}",
            "categoryId": f"PERCAT-{int(row_dict.get('category_id') or 0):04d}",
            "attachmentType": row_dict.get("attachment_type") or "",
            "fileName": row_dict.get("file_name") or "",
            "minioKey": row_dict.get("minio_key") or "",
            "minioBucket": row_dict.get("minio_bucket") or settings.minio_buckets["materials"],
            "mimeType": row_dict.get("mime_type") or "",
            "sizeBytes": int(row_dict.get("size_bytes") or 0),
            "createdAt": row_dict.get("created_at").isoformat() if row_dict.get("created_at") else "",
        }

    def _item_attachment_row_to_dict(self, row: Any) -> dict[str, Any]:
        row_dict = dict(row)
        numeric_id = int(row_dict.get("id") or 0)
        return {
            "id": f"PERITEMATT-{numeric_id:04d}",
            "categoryId": f"PERCAT-{int(row_dict.get('category_id') or 0):04d}",
            "itemId": f"PERITEM-{int(row_dict.get('item_id') or 0):04d}",
            "sourceAttachmentId": f"PERATT-{int(row_dict.get('source_attachment_id') or 0):04d}" if row_dict.get("source_attachment_id") else "",
            "attachmentType": row_dict.get("attachment_type") or "",
            "fileName": row_dict.get("file_name") or "",
            "minioKey": row_dict.get("minio_key") or "",
            "minioBucket": row_dict.get("minio_bucket") or settings.minio_buckets["materials"],
            "mimeType": row_dict.get("mime_type") or "",
            "sizeBytes": int(row_dict.get("size_bytes") or 0),
            "matchConfidence": int(row_dict.get("match_confidence") or 0),
            "matchMethod": row_dict.get("match_method") or "",
            "sourceTitle": row_dict.get("source_title") or "",
            "sourceBlockStart": row_dict.get("source_block_start"),
            "sourceBlockEnd": row_dict.get("source_block_end"),
            "createdAt": row_dict.get("created_at").isoformat() if row_dict.get("created_at") else "",
        }

    def _numeric_category_id(self, category_id: str) -> int:
        try:
            return int(str(category_id or "").replace("PERCAT-", ""))
        except ValueError as exc:
            raise PeripheralError(400, "业绩类别 ID 无效。", "PERFORMANCE_CATEGORY_ID_INVALID") from exc

    def _numeric_item_id(self, item_id: str) -> int:
        try:
            return int(str(item_id or "").replace("PERITEM-", ""))
        except ValueError as exc:
            raise PeripheralError(400, "业绩明细 ID 无效。", "PERFORMANCE_ITEM_ID_INVALID") from exc

    def _numeric_attachment_id(self, attachment_id: str) -> int:
        try:
            return int(str(attachment_id or "").replace("PERATT-", ""))
        except ValueError as exc:
            raise PeripheralError(400, "业绩附件 ID 无效。", "PERFORMANCE_ATTACHMENT_ID_INVALID") from exc

    def _numeric_item_attachment_id(self, attachment_id: str) -> int:
        try:
            return int(str(attachment_id or "").replace("PERITEMATT-", ""))
        except ValueError as exc:
            raise PeripheralError(400, "项目合同附件 ID 无效。", "PERFORMANCE_ITEM_ATTACHMENT_ID_INVALID") from exc


def parse_performance_summary_docx(content: bytes, *, file_name: str = "performance.docx") -> dict[str, Any]:
    if not file_name.lower().endswith(".docx"):
        raise PeripheralError(400, "业绩汇总表解析仅支持 .docx 文件。", "PERFORMANCE_SUMMARY_DOCX_REQUIRED")
    try:
        doc = Document(BytesIO(content))
    except Exception as exc:
        raise PeripheralError(400, "业绩汇总表无法读取，请确认文件格式。", "PERFORMANCE_SUMMARY_DOCX_INVALID") from exc

    paragraphs = [_clean_text(paragraph.text) for paragraph in doc.paragraphs]
    category_name = _infer_category_name(paragraphs, file_name=file_name)
    summary = _infer_summary(paragraphs, category_name=category_name)
    scene = _infer_scene(category_name)
    power_rating = _infer_power_rating(category_name)

    selected_table = _select_summary_table(doc.tables)
    if selected_table is None:
        raise PeripheralError(400, "未在文件中识别到业绩汇总表。", "PERFORMANCE_SUMMARY_TABLE_NOT_FOUND")

    table_rows = _table_to_rows(selected_table)
    header_index = _find_header_row(table_rows)
    if header_index < 0:
        raise PeripheralError(400, "未识别到业绩汇总表表头。", "PERFORMANCE_SUMMARY_HEADER_NOT_FOUND")
    headers = _unique_headers(table_rows[header_index])
    field_schema = [_field_schema_item(header) for header in headers]
    detail_rows: list[dict[str, Any]] = []
    for index, values in enumerate(table_rows[header_index + 1 :], start=1):
        if not _is_data_row(values):
            continue
        row_values = {
            header: _normalize_empty(values[column_index] if column_index < len(values) else "")
            for column_index, header in enumerate(headers)
        }
        core = _core_values(row_values)
        detail_rows.append(
            {
                "rowIndex": index,
                "values": row_values,
                **core,
            }
        )

    if not detail_rows:
        raise PeripheralError(400, "业绩汇总表未识别到有效明细行。", "PERFORMANCE_SUMMARY_ROWS_EMPTY")

    return {
        "categoryName": category_name,
        "scene": scene,
        "powerRating": power_rating,
        "summary": summary,
        "fieldSchema": field_schema,
        "rows": detail_rows,
        "rowCount": len(detail_rows),
        "sourceFileName": file_name,
    }


def split_performance_contract_docx(content: bytes, *, file_name: str = "contract.docx") -> list[dict[str, Any]]:
    if not file_name.lower().endswith(".docx"):
        return []
    try:
        source_doc = Document(BytesIO(content))
    except Exception as exc:
        raise PeripheralError(400, "合同附件无法读取，请确认文件格式。", "PERFORMANCE_CONTRACT_DOCX_INVALID") from exc

    body_children = list(source_doc.element.body.iterchildren())
    title_indexes = [
        index
        for index, child in enumerate(body_children)
        if _is_contract_title_block(child)
    ]
    if not title_indexes:
        title_indexes = _contract_title_indexes_from_page_breaks(body_children)
    if not title_indexes:
        title_indexes = [0] if body_children else []

    chunks: list[dict[str, Any]] = []
    for chunk_index, start in enumerate(title_indexes):
        end = title_indexes[chunk_index + 1] if chunk_index + 1 < len(title_indexes) else len(body_children)
        selected_blocks = [
            child
            for child in body_children[start:end]
            if _block_kind(child) != "sectPr"
        ]
        if not selected_blocks:
            continue
        title = _first_meaningful_block_text(selected_blocks) or f"{Path(file_name).stem}-{chunk_index + 1}"
        chunks.append(
            {
                "index": chunk_index + 1,
                "title": _dedupe_repeated_title(title)[:300],
                "blocks": selected_blocks,
                "sourceDoc": source_doc,
                "blockStart": int(start),
                "blockEnd": int(max(start, end - 1)),
            }
        )
    for chunk in chunks:
        content_bytes = render_contract_item_docx(chunk)
        chunk["content"] = content_bytes
        chunk["sizeBytes"] = len(content_bytes)
    return chunks


def match_contract_chunks_to_items(chunks: list[dict[str, Any]], items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    if not chunks or not items:
        return []
    matches: list[dict[str, Any]] = []
    used_chunk_indexes: set[int] = set()
    for item_index, item in enumerate(items):
        best_chunk: dict[str, Any] | None = None
        best_score = 0
        for chunk_index, chunk in enumerate(chunks):
            if chunk_index in used_chunk_indexes:
                continue
            score = _contract_match_score(item, chunk)
            if score > best_score:
                best_score = score
                best_chunk = {**chunk, "_chunkIndex": chunk_index}
        if best_chunk is not None and best_score >= 58:
            used_chunk_indexes.add(int(best_chunk["_chunkIndex"]))
            matches.append({"item": item, "chunk": _without_private_keys(best_chunk), "confidence": best_score, "method": "project_name"})
            continue
        if item_index < len(chunks) and item_index not in used_chunk_indexes:
            used_chunk_indexes.add(item_index)
            matches.append({"item": item, "chunk": chunks[item_index], "confidence": 45, "method": "row_order"})
    return matches


def render_contract_item_docx(chunk: dict[str, Any], *, output_title: str = "") -> bytes:
    source_doc = chunk.get("sourceDoc")
    blocks = list(chunk.get("blocks") or [])
    if source_doc is None:
        content = chunk.get("content")
        if isinstance(content, bytes):
            return content
        return b""
    title = _clean_contract_output_title(output_title or str(chunk.get("title") or "项目合同"))
    content_blocks = _contract_content_blocks(blocks)
    return _docx_blocks_to_bytes(source_doc, content_blocks, title=title)


def _docx_blocks_to_bytes(source_doc: Any, blocks: list[Any], *, title: str = "") -> bytes:
    target_doc = Document()
    target_body = target_doc.element.body
    for child in list(target_body):
        target_body.remove(child)
    if title:
        title_paragraph = target_doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(12)
        title_format = title_paragraph.paragraph_format
        title_format.space_before = Pt(0)
        title_format.space_after = Pt(6)
        title_format.line_spacing = 1.5
        target_body.remove(title_paragraph._p)
        target_body.append(title_paragraph._p)
    for block in blocks:
        cloned = deepcopy(block)
        _copy_related_parts(source_doc.part, target_doc.part, cloned)
        _normalize_contract_block_format(cloned)
        _sanitize_contract_drawingml(cloned)
        _stabilize_contract_table_pagination(cloned)
        target_body.append(cloned)
    source_sect = source_doc.element.body.sectPr
    if source_sect is not None:
        target_body.append(deepcopy(source_sect))
    output = BytesIO()
    target_doc.save(output)
    sanitized = _sanitize_contract_docx_fonts(output.getvalue())
    return _normalize_contract_docx_with_soffice(sanitized)


def _normalize_contract_docx_with_soffice(content: bytes) -> bytes:
    executable = shutil.which("soffice") or shutil.which("libreoffice")
    if not executable:
        logger.warning("Skip performance contract DOCX normalization: soffice/libreoffice not found.")
        return content
    try:
        with tempfile.TemporaryDirectory(prefix="performance-contract-docx-") as temp_root:
            root = Path(temp_root)
            input_dir = root / "input"
            output_dir = root / "output"
            profile_dir = root / "profile"
            input_dir.mkdir()
            output_dir.mkdir()
            profile_dir.mkdir()
            source_path = input_dir / "source.docx"
            source_path.write_bytes(content)
            completed = subprocess.run(
                [
                    executable,
                    "--headless",
                    f"-env:UserInstallation=file://{profile_dir.as_posix()}",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    str(output_dir),
                    str(source_path),
                ],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                timeout=90,
            )
            converted_path = output_dir / source_path.name
            if not converted_path.exists():
                candidates = list(output_dir.glob("*.docx"))
                converted_path = candidates[0] if candidates else converted_path
            if not converted_path.exists():
                logger.warning(
                    "Skip performance contract DOCX normalization: no output file. stdout=%s stderr=%s",
                    (completed.stdout or "").strip(),
                    (completed.stderr or "").strip(),
                )
                return content
            normalized = converted_path.read_bytes()
            return _sanitize_contract_docx_fonts(normalized) if normalized else content
    except FileNotFoundError:
        logger.warning("Skip performance contract DOCX normalization: soffice/libreoffice not found.")
    except subprocess.TimeoutExpired:
        logger.warning("Skip performance contract DOCX normalization: conversion timed out.")
    except subprocess.CalledProcessError as exc:
        logger.warning(
            "Skip performance contract DOCX normalization: conversion failed. stdout=%s stderr=%s",
            (exc.stdout or "").strip(),
            (exc.stderr or "").strip(),
        )
    except Exception as exc:  # pragma: no cover - normalization must not break import
        logger.warning("Skip performance contract DOCX normalization: %s", exc)
    return content


def _contract_content_blocks(blocks: list[Any]) -> list[Any]:
    result: list[Any] = []
    skipped_title = False
    had_title = False
    for block in blocks:
        if not skipped_title and _is_contract_title_block(block):
            skipped_title = True
            had_title = True
            continue
        result.append(block)
    if not had_title:
        result = _trim_leading_layout_blocks(result)
    return _trim_trailing_layout_blocks(result)


def _trim_leading_layout_blocks(blocks: list[Any]) -> list[Any]:
    start = 0
    while start < len(blocks) and _is_layout_only_paragraph(blocks[start]):
        start += 1
    return blocks[start:]


def _trim_trailing_layout_blocks(blocks: list[Any]) -> list[Any]:
    end = len(blocks)
    while end > 0 and _is_layout_only_paragraph(blocks[end - 1]):
        end -= 1
    return blocks[:end]


def _normalize_contract_block_format(block: Any) -> None:
    paragraph_nodes = [block] if _block_kind(block) == "p" else list(block.xpath('.//*[local-name()="p"]'))
    for paragraph in paragraph_nodes:
        if _is_layout_only_paragraph(paragraph):
            if _has_ooxml_ancestor(paragraph, "tc"):
                _normalize_layout_paragraph_spacing(paragraph)
            continue
        ppr = paragraph.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
        if ppr is None:
            ppr = _insert_ooxml_child(paragraph, "pPr", 0)
        spacing = ppr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spacing")
        if spacing is None:
            spacing = _append_ooxml_child(ppr, "spacing")
        spacing.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}before", "0")
        spacing.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}after", "0")
        spacing.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}line", "360")
        spacing.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lineRule", "auto")
        rpr = ppr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
        if rpr is None:
            rpr = _append_ooxml_child(ppr, "rPr")
        _set_ooxml_rfonts(rpr)
    for run_properties in block.xpath('.//*[local-name()="rPr"]'):
        _set_ooxml_rfonts(run_properties)


def _sanitize_contract_drawingml(block: Any) -> None:
    for picture_index, inline in enumerate(block.xpath('.//*[local-name()="inline"]'), start=1):
        extent = next((node for node in inline.iter() if node.tag.rsplit("}", 1)[-1] == "extent"), None)
        extent_attrs = {attr: extent.attrib.get(attr) for attr in ("cx", "cy")} if extent is not None else {}
        for node in inline.iter():
            local_name = node.tag.rsplit("}", 1)[-1]
            if local_name in {"docPr", "cNvPr"} and "id" in node.attrib:
                node.set("id", str(picture_index))
            elif local_name == "ext" and extent_attrs.get("cx") and extent_attrs.get("cy"):
                node.set("cx", str(extent_attrs["cx"]))
                node.set("cy", str(extent_attrs["cy"]))
    for element in list(block.iter()):
        local_name = element.tag.rsplit("}", 1)[-1]
        if local_name == "effectExtent":
            for attr_name in ("l", "t", "r", "b"):
                element.set(attr_name, "0")
        elif local_name == "blip":
            for child in list(element):
                if child.tag.rsplit("}", 1)[-1] == "extLst":
                    element.remove(child)
        elif local_name == "spPr":
            for child in list(element):
                if child.tag.rsplit("}", 1)[-1] == "ln":
                    element.remove(child)


def _stabilize_contract_table_pagination(block: Any) -> None:
    if _block_kind(block) != "tbl":
        return
    rows = [child for child in block.iterchildren() if _block_kind(child) == "tr"]
    for index, row in enumerate(rows[:-1]):
        next_row = rows[index + 1]
        if not _is_contract_table_caption_row(row):
            continue
        if not _has_drawings(next_row):
            continue
        _set_row_cant_split(row)
        _set_row_cant_split(next_row)
        _set_row_keep_next(row)


def _is_contract_table_caption_row(row: Any) -> bool:
    text_value = _dedupe_repeated_title(_block_text(row))
    if not text_value or len(text_value) > 120:
        return False
    if _has_drawings(row):
        return False
    return any(keyword in text_value for keyword in ("页", "合同", "参数", "机型", "盖章", "首页"))


def _set_row_cant_split(row: Any) -> None:
    trpr = row.find(f"{{{WORD_NAMESPACE}}}trPr")
    if trpr is None:
        trpr = _insert_ooxml_child(row, "trPr", 0)
    if trpr.find(f"{{{WORD_NAMESPACE}}}cantSplit") is None:
        trpr.insert(0, trpr.makeelement(f"{{{WORD_NAMESPACE}}}cantSplit"))


def _set_row_keep_next(row: Any) -> None:
    for paragraph in row.xpath('.//*[local-name()="p"]'):
        ppr = paragraph.find(f"{{{WORD_NAMESPACE}}}pPr")
        if ppr is None:
            ppr = _insert_ooxml_child(paragraph, "pPr", 0)
        if ppr.find(f"{{{WORD_NAMESPACE}}}keepNext") is not None:
            continue
        keep_next = ppr.makeelement(f"{{{WORD_NAMESPACE}}}keepNext")
        insert_at = 0
        pstyle = ppr.find(f"{{{WORD_NAMESPACE}}}pStyle")
        if pstyle is not None:
            insert_at = list(ppr).index(pstyle) + 1
        ppr.insert(insert_at, keep_next)


def _normalize_layout_paragraph_spacing(paragraph: Any) -> None:
    ppr = paragraph.find(f"{{{WORD_NAMESPACE}}}pPr")
    if ppr is None:
        ppr = _insert_ooxml_child(paragraph, "pPr", 0)
    spacing = ppr.find(f"{{{WORD_NAMESPACE}}}spacing")
    if spacing is None:
        spacing = _append_ooxml_child(ppr, "spacing")
    spacing.set(f"{{{WORD_NAMESPACE}}}before", "0")
    spacing.set(f"{{{WORD_NAMESPACE}}}after", "0")
    spacing.set(f"{{{WORD_NAMESPACE}}}line", "360")
    spacing.set(f"{{{WORD_NAMESPACE}}}lineRule", "auto")


def _has_ooxml_ancestor(node: Any, local_name: str) -> bool:
    return any(_block_kind(ancestor) == local_name for ancestor in node.iterancestors())


def _set_ooxml_rfonts(run_properties: Any) -> None:
    rfonts = run_properties.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
    if rfonts is None:
        rfonts = _insert_ooxml_child(run_properties, "rFonts", 0)
    rfonts.set(f"{{{WORD_NAMESPACE}}}ascii", CONTRACT_OUTPUT_WESTERN_FONT)
    rfonts.set(f"{{{WORD_NAMESPACE}}}hAnsi", CONTRACT_OUTPUT_WESTERN_FONT)
    rfonts.set(f"{{{WORD_NAMESPACE}}}eastAsia", CONTRACT_OUTPUT_EAST_ASIA_FONT)
    rfonts.set(f"{{{WORD_NAMESPACE}}}cs", CONTRACT_OUTPUT_WESTERN_FONT)
    for attr_name in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme", "hint"):
        rfonts.attrib.pop(f"{{{WORD_NAMESPACE}}}{attr_name}", None)


def _sanitize_contract_docx_fonts(content: bytes) -> bytes:
    source = BytesIO(content)
    target = BytesIO()
    with ZipFile(source, "r") as src_zip, ZipFile(target, "w", compression=ZIP_DEFLATED) as dst_zip:
        for info in src_zip.infolist():
            data = src_zip.read(info.filename)
            if info.filename == "word/fontTable.xml":
                data = _contract_font_table_xml()
            elif info.filename == "word/theme/theme1.xml":
                data = _sanitize_theme_fonts_xml(data)
            elif info.filename == "word/settings.xml":
                data = _sanitize_contract_settings_xml(data)
            elif info.filename.startswith("word/") and info.filename.endswith(".xml"):
                data = _sanitize_word_fonts_xml(data)
            dst_zip.writestr(info, data)
    return target.getvalue()


def _sanitize_contract_settings_xml(data: bytes) -> bytes:
    try:
        root = ET.fromstring(data)
    except ET.ParseError:
        return data
    for local_name in ("compat", "rsids"):
        for element in list(root):
            if element.tag == f"{{{WORD_NAMESPACE}}}{local_name}":
                root.remove(element)
    return _serialize_ooxml(root)


def _sanitize_word_fonts_xml(data: bytes) -> bytes:
    try:
        root = ET.fromstring(data)
    except ET.ParseError:
        return data
    changed = False
    for element in root.iter():
        local_name = element.tag.rsplit("}", 1)[-1]
        if local_name == "rFonts":
            element.attrib[f"{{{WORD_NAMESPACE}}}ascii"] = CONTRACT_OUTPUT_WESTERN_FONT
            element.attrib[f"{{{WORD_NAMESPACE}}}hAnsi"] = CONTRACT_OUTPUT_WESTERN_FONT
            element.attrib[f"{{{WORD_NAMESPACE}}}eastAsia"] = CONTRACT_OUTPUT_EAST_ASIA_FONT
            element.attrib[f"{{{WORD_NAMESPACE}}}cs"] = CONTRACT_OUTPUT_WESTERN_FONT
            for attr_name in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme", "hint"):
                element.attrib.pop(f"{{{WORD_NAMESPACE}}}{attr_name}", None)
            changed = True
    if not changed:
        return data
    return _serialize_ooxml(root)


def _sanitize_theme_fonts_xml(data: bytes) -> bytes:
    try:
        root = ET.fromstring(data)
    except ET.ParseError:
        return data
    for tag_name in ("latin", "ea", "cs"):
        for element in root.iter(f"{{{A_NAMESPACE}}}{tag_name}"):
            element.set("typeface", CONTRACT_OUTPUT_EAST_ASIA_FONT if tag_name == "ea" else CONTRACT_OUTPUT_WESTERN_FONT)
    for element in root.iter(f"{{{A_NAMESPACE}}}font"):
        script = str(element.get("script") or "")
        element.set(
            "typeface",
            CONTRACT_OUTPUT_EAST_ASIA_FONT if script in THEME_EAST_ASIA_SCRIPTS else CONTRACT_OUTPUT_WESTERN_FONT,
        )
    return _serialize_ooxml(root)


def _serialize_ooxml(root: ET.Element) -> bytes:
    _sanitize_mc_ignorable(root)
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _sanitize_mc_ignorable(root: ET.Element) -> None:
    ignorable_attr = f"{{{MC_NAMESPACE}}}Ignorable"
    value = str(root.attrib.get(ignorable_attr) or "").strip()
    if not value:
        return
    used_namespaces = {str(element.tag).split("}", 1)[0][1:] for element in root.iter() if str(element.tag).startswith("{")}
    kept = [
        prefix
        for prefix in value.split()
        if OOXML_NAMESPACES.get(prefix) in used_namespaces
    ]
    if kept:
        root.attrib[ignorable_attr] = " ".join(kept)
    else:
        root.attrib.pop(ignorable_attr, None)


def _contract_font_table_xml() -> bytes:
    return f"""<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<w:fonts xmlns:w="{WORD_NAMESPACE}">
  <w:font w:name="{CONTRACT_OUTPUT_EAST_ASIA_FONT}">
    <w:charset w:val="86"/>
    <w:family w:val="roman"/>
    <w:pitch w:val="variable"/>
  </w:font>
  <w:font w:name="{CONTRACT_OUTPUT_WESTERN_FONT}">
    <w:charset w:val="00"/>
    <w:family w:val="roman"/>
    <w:pitch w:val="variable"/>
  </w:font>
  <w:font w:name="{CONTRACT_OUTPUT_SYMBOL_FONT}">
    <w:charset w:val="02"/>
    <w:family w:val="auto"/>
    <w:pitch w:val="variable"/>
  </w:font>
</w:fonts>
""".encode("utf-8")


def _append_ooxml_child(parent: Any, local_name: str) -> Any:
    child = parent.makeelement(f"{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}{local_name}")
    parent.append(child)
    return child


def _insert_ooxml_child(parent: Any, local_name: str, index: int) -> Any:
    child = parent.makeelement(f"{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}{local_name}")
    parent.insert(index, child)
    return child


def _copy_related_parts(source_part: Any, target_part: Any, element: Any) -> None:
    for node in element.iter():
        for attr_name, attr_value in list(node.attrib.items()):
            if not attr_name.startswith(RELATIONSHIP_NAMESPACE):
                continue
            relation = source_part.rels.get(attr_value)
            if relation is None:
                continue
            if relation.is_external:
                new_rid = target_part.relate_to(relation.target_ref, relation.reltype, is_external=True)
            else:
                new_rid = target_part.relate_to(relation.target_part, relation.reltype)
            node.attrib[attr_name] = new_rid


def _clean_contract_output_title(value: str) -> str:
    text_value = _dedupe_repeated_title(_clean_text(value))
    text_value = re.sub(r"(采购)?合同$", "", text_value).strip(" ，,。；;:-_")
    if len(text_value) > 80:
        text_value = text_value[:80].rstrip(" ，,。；;:-_")
    return text_value or "项目合同"


def _is_contract_title_block(block: Any) -> bool:
    if _block_kind(block) != "p":
        return False
    text_value = _dedupe_repeated_title(_block_text(block))
    if len(text_value) < 8:
        return False
    if _has_drawings(block):
        return False
    if _is_empty_page_break_paragraph(block):
        return False
    if re.fullmatch(r"标段[一二三四五六七八九十\d]+[:：]?", text_value):
        return False
    if not any(keyword in text_value for keyword in ("合同", "项目", "工程", "风电", "设备", "采购", "供货", "EPC")):
        return False
    if any(keyword in text_value for keyword in ("合同首页", "签字盖章页", "技术数据页", "预验收证明", "运行证明", "并网证明")):
        return False
    return True


def _contract_title_indexes_from_page_breaks(blocks: list[Any]) -> list[int]:
    indexes: list[int] = []
    next_meaningful_starts_chunk = True
    for index, block in enumerate(blocks):
        if _block_kind(block) == "sectPr":
            continue
        text_value = _dedupe_repeated_title(_block_text(block))
        meaningful = bool(text_value) or _block_kind(block) == "tbl" or _has_drawings(block)
        if next_meaningful_starts_chunk and meaningful:
            indexes.append(index)
            next_meaningful_starts_chunk = False
        if _block_has_page_break(block):
            next_meaningful_starts_chunk = True
    return indexes


def _first_meaningful_block_text(blocks: list[Any]) -> str:
    for block in blocks:
        text_value = _dedupe_repeated_title(_block_text(block))
        if text_value:
            return text_value
    return ""


def _block_kind(block: Any) -> str:
    return str(getattr(block, "tag", "")).rsplit("}", 1)[-1]


def _block_text(block: Any) -> str:
    return _clean_text("".join(block.itertext()))


def _block_has_page_break(block: Any) -> bool:
    return bool(block.xpath('.//*[local-name()="br" and @*[local-name()="type"]="page"]'))


def _has_drawings(block: Any) -> bool:
    return bool(block.xpath('.//*[local-name()="drawing" or local-name()="pict"]'))


def _is_empty_page_break_paragraph(block: Any) -> bool:
    return _block_kind(block) == "p" and not _block_text(block) and _block_has_page_break(block)


def _is_layout_only_paragraph(block: Any) -> bool:
    return _block_kind(block) == "p" and not _block_text(block) and not _has_drawings(block)


def _dedupe_repeated_title(value: str) -> str:
    text_value = _clean_text(value)
    if len(text_value) < 8:
        return text_value
    for unit_length in range(4, max(5, len(text_value) // 2 + 1)):
        if len(text_value) % unit_length:
            continue
        unit = text_value[:unit_length]
        repeats = len(text_value) // unit_length
        if repeats >= 2 and unit * repeats == text_value:
            return unit
    half = len(text_value) // 2
    if half >= 8 and text_value[:half] == text_value[half:]:
        return text_value[:half]
    for end in range(8, min(len(text_value), 180)):
        unit = text_value[:end]
        if text_value.startswith(unit + unit):
            return unit
    return text_value


def _contract_match_score(item: dict[str, Any], chunk: dict[str, Any]) -> int:
    project_name = str(item.get("projectName") or "")
    chunk_title = str(chunk.get("title") or "")
    if not project_name or not chunk_title:
        return 0
    project_norm = _match_text(project_name)
    chunk_norm = _match_text(chunk_title)
    if not project_norm or not chunk_norm:
        return 0
    if project_norm in chunk_norm or chunk_norm in project_norm:
        return 96
    project_tokens = _match_tokens(project_name)
    chunk_tokens = _match_tokens(chunk_title)
    if not project_tokens:
        return 0
    overlap = project_tokens & chunk_tokens
    score = int(len(overlap) / max(1, len(project_tokens)) * 100)
    longest = _longest_common_substring_length(project_norm, chunk_norm)
    score = max(score, int(longest / max(1, len(project_norm)) * 100))
    ordered = _ordered_character_coverage(project_norm, chunk_norm)
    score = max(score, int(ordered * 100))
    model_values = item.get("turbineModels") or []
    if isinstance(model_values, list) and any(_match_text(model) and _match_text(model) in chunk_norm for model in model_values):
        score += 8
    customer = _match_text(item.get("customerName") or "")
    if customer and customer in chunk_norm:
        score += 8
    return max(0, min(100, score))


def _match_text(value: Any) -> str:
    text_value = unicodedata.normalize("NFKC", str(value or "")).lower()
    return re.sub(r"[^0-9a-z\u4e00-\u9fff]+", "", text_value)


def _match_tokens(value: Any) -> set[str]:
    normalized = _match_text(value)
    if not normalized:
        return set()
    tokens = {match.group(0) for match in re.finditer(r"[a-z]+\d*(?:\.\d+)?|\d+(?:\.\d+)?mw|\d+(?:\.\d+)?万千瓦|\d+(?:\.\d+)?千瓦|\d+x\d+|\d+×\d+", normalized)}
    for width in (2, 3, 4):
        for start in range(0, max(0, len(normalized) - width + 1)):
            token = normalized[start : start + width]
            if len(token) == width:
                tokens.add(token)
    return tokens


def _ordered_character_coverage(left: str, right: str) -> float:
    if not left or not right:
        return 0.0
    right_index = 0
    matched = 0
    for char in left:
        found_at = right.find(char, right_index)
        if found_at < 0:
            continue
        matched += 1
        right_index = found_at + 1
    return matched / max(1, len(left))


def _longest_common_substring_length(left: str, right: str) -> int:
    if not left or not right:
        return 0
    previous = [0] * (len(right) + 1)
    best = 0
    for left_char in left:
        current = [0] * (len(right) + 1)
        for index, right_char in enumerate(right, start=1):
            if left_char == right_char:
                current[index] = previous[index - 1] + 1
                best = max(best, current[index])
        previous = current
    return best


def _without_private_keys(value: dict[str, Any]) -> dict[str, Any]:
    return {key: item for key, item in value.items() if not str(key).startswith("_")}


def _contract_item_file_name(row_index: Any, project_name: str, chunk_index: Any = 0) -> str:
    safe_project = _safe_file_name(_clean_text(project_name))[:96].strip(" .-")
    if not safe_project:
        safe_project = f"项目{int(chunk_index or 0) or 1}"
    try:
        row_number = int(row_index or 0)
    except (TypeError, ValueError):
        row_number = 0
    prefix = f"{row_number:03d}-" if row_number > 0 else ""
    return f"{prefix}{safe_project}_合同.docx"


def _select_summary_table(tables: Any) -> Any | None:
    best_table = None
    best_score = -1
    for table in tables:
        rows = _table_to_rows(table)
        if not rows:
            continue
        header_index = _find_header_row(rows)
        if header_index < 0:
            continue
        header = " ".join(rows[header_index])
        score = _header_score(rows[header_index]) * 10 + max(0, len(rows) - header_index - 1)
        if any(keyword in header for keyword in ("项目", "合同", "买方", "型号")):
            score += 10
        if score > best_score:
            best_score = score
            best_table = table
    return best_table


def _table_to_rows(table: Any) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [_clean_text(cell.text) for cell in row.cells]
        if any(cell for cell in cells):
            rows.append(cells)
    return rows


def _find_header_row(rows: list[list[str]]) -> int:
    best_index = -1
    best_score = 0
    for index, row in enumerate(rows[:8]):
        score = _header_score(row)
        non_empty = sum(1 for cell in row if cell)
        if non_empty >= 3 and score > best_score:
            best_index = index
            best_score = score
    return best_index if best_score >= 2 else -1


def _header_score(row: list[str]) -> int:
    joined = " ".join(row)
    score = 0
    for keyword in ("项目名称", "合同名称", "买方", "业主", "客户", "型号", "合同台数", "试运行", "投运容量", "联系人"):
        if keyword in joined:
            score += 1
    return score


def _infer_category_name(paragraphs: list[str], *, file_name: str) -> str:
    for paragraph in paragraphs[:6]:
        if paragraph and not paragraph.startswith("合同业绩台数"):
            return paragraph[:300]
    return Path(file_name).stem[:300]


def _infer_summary(paragraphs: list[str], *, category_name: str) -> str:
    for paragraph in paragraphs[:10]:
        if paragraph and paragraph != category_name and any(keyword in paragraph for keyword in ("合同业绩", "台数", "投运容量")):
            return paragraph
    return ""


def _infer_scene(text_value: str) -> str:
    text_value = str(text_value or "")
    if "海上" in text_value:
        return "海上"
    if "陆上" in text_value:
        return "陆上"
    return ""


def _infer_power_rating(text_value: str) -> str:
    match = re.search(r"(\d+(?:\.\d+)?)\s*MW\s*(?:及以上|以上)?", str(text_value or ""), flags=re.IGNORECASE)
    if not match:
        return ""
    suffix = "及以上" if "及以上" in match.group(0) or "以上" in match.group(0) else ""
    return f"{match.group(1)}MW{suffix}"


def _unique_headers(headers: list[str]) -> list[str]:
    result: list[str] = []
    counts: dict[str, int] = {}
    for index, header in enumerate(headers, start=1):
        value = _clean_text(header) or f"字段{index}"
        count = counts.get(value, 0) + 1
        counts[value] = count
        result.append(value if count == 1 else f"{value}_{count}")
    return result


def _field_schema_item(header: str) -> dict[str, str]:
    key = _field_key(header)
    return {"key": key, "label": header, "sourceHeader": header}


def _field_key(header: str) -> str:
    normalized = re.sub(r"[^0-9A-Za-z一-龥]+", "_", str(header or "").strip()).strip("_")
    return normalized or "field"


def _core_values(row_values: dict[str, str]) -> dict[str, str]:
    turbine_model = _first_value(row_values, ("型号和规格", "型号", "机型", "风机型号"))
    time_facts = _time_facts(row_values)
    return {
        "projectName": _first_value(row_values, ("项目名称", "合同名称", "工程名称")),
        "customerName": _first_value(row_values, ("买方名称", "业主", "客户", "采购方")),
        "turbineModel": turbine_model,
        "turbineModels": _split_turbine_models(turbine_model),
        "contractQuantity": _first_value(row_values, ("合同台数", "台数", "数量")),
        "trialOperationQuantity": _first_value(row_values, ("试运行台数", "240", "试运")),
        "commissionedCapacityMw": _first_value(row_values, ("投运容量", "容量")),
        "deliveryOrOperationTime": time_facts["deliveryOrOperationTimeRaw"],
        "contractYear": time_facts.get("contractYear"),
        "deliveryYear": time_facts.get("deliveryYear"),
        "operationYear": time_facts.get("operationYear"),
        "timeFacts": time_facts,
        "contactInfo": _first_value(row_values, ("联系人", "电话", "联系方式")),
    }


def _first_value(row_values: dict[str, str], keywords: tuple[str, ...]) -> str:
    for key, value in row_values.items():
        if any(keyword in key for keyword in keywords):
            return value
    return ""


def _split_turbine_models(value: str) -> list[str]:
    text_value = _normalize_empty(value)
    if not text_value:
        return []
    normalized = re.sub(r"[，,、/；;]+", " ", text_value)
    parts = re.split(r"\s+", normalized)
    result: list[str] = []
    seen: set[str] = set()
    for part in parts:
        candidate = part.strip("()（）[]【】")
        if not candidate:
            continue
        if not re.search(r"\d", candidate):
            continue
        if not re.search(r"[A-Za-z]|MW|mw|-", candidate):
            continue
        key = candidate.upper()
        if key in seen:
            continue
        seen.add(key)
        result.append(candidate)
    return result


def _time_facts(row_values: dict[str, str]) -> dict[str, Any]:
    contract_raw = _value_by_header(row_values, ("合同时间", "合同日期", "签约时间", "签订时间", "签订日期"))
    operation_raw = _value_by_header(row_values, ("投运时间", "投运日期", "并网时间", "运行时间"))
    delivery_raw = _value_by_header(row_values, ("交货期", "交付时间", "交货时间", "交付日期"))
    delivery_or_operation_raw = _value_by_header(row_values, ("交货期/投运时间", "交货/投运", "交付/投运"))
    combined_raw = delivery_or_operation_raw or operation_raw or delivery_raw
    contract_year = _extract_year(contract_raw)
    delivery_year = _extract_year(delivery_raw)
    operation_year = _extract_year(operation_raw)
    combined_year = _extract_year(combined_raw)
    if delivery_or_operation_raw:
        if delivery_year is None:
            delivery_year = combined_year
        if operation_year is None:
            operation_year = combined_year
    elif not operation_raw and delivery_raw and operation_year is None:
        operation_year = None
    return {
        "contractTimeRaw": contract_raw,
        "deliveryTimeRaw": delivery_raw,
        "operationTimeRaw": operation_raw,
        "deliveryOrOperationTimeRaw": combined_raw,
        "contractYear": contract_year,
        "deliveryYear": delivery_year,
        "operationYear": operation_year,
        "years": _unique_ints([contract_year, delivery_year, operation_year, combined_year]),
    }


def _value_by_header(row_values: dict[str, str], keywords: tuple[str, ...]) -> str:
    for key, value in row_values.items():
        normalized_key = re.sub(r"\s+", "", str(key or ""))
        if any(keyword in normalized_key for keyword in keywords):
            return value
    return ""


def _extract_year(value: str) -> int | None:
    text_value = _normalize_empty(value)
    if not text_value:
        return None
    match = re.search(r"(19\d{2}|20\d{2})", text_value)
    if not match:
        return None
    year = int(match.group(1))
    if 1990 <= year <= 2100:
        return year
    return None


def _unique_ints(values: list[int | None]) -> list[int]:
    result: list[int] = []
    seen: set[int] = set()
    for value in values:
        if value is None or value in seen:
            continue
        seen.add(value)
        result.append(value)
    return result


def _parse_year_filter(value: Any) -> int | None:
    year = _extract_year(str(value or ""))
    return year


def _is_data_row(values: list[str]) -> bool:
    meaningful = [value for value in values if _normalize_empty(value)]
    if len(meaningful) < 2:
        return False
    joined = " ".join(meaningful)
    return not all(keyword in joined for keyword in ("序号", "型号", "项目"))


def _normalize_empty(value: Any) -> str:
    text_value = _clean_text(str(value or ""))
    if text_value in {"/", "-", "—", "无", "N/A", "NA"}:
        return ""
    return text_value


def _clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", str(value or "").replace("\u3000", " ")).strip()


async def _read_upload(upload: Any) -> tuple[bytes, str, str]:
    file_name = _safe_file_name(str(getattr(upload, "filename", "") or "performance.docx"))
    if not file_name.lower().endswith((".doc", ".docx")):
        raise PeripheralError(400, "业绩包仅支持上传 Word 文件。", "PERFORMANCE_WORD_REQUIRED")
    content = await upload.read()
    if len(content) > settings.max_upload_file_size_bytes:
        limit_mb = settings.max_upload_file_size_bytes // 1024 // 1024
        raise PeripheralError(413, f"文件超过 {limit_mb}MB 上限。", "PERFORMANCE_FILE_TOO_LARGE")
    mime_type = str(getattr(upload, "content_type", "") or "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    return content, file_name, mime_type


def _json(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False)


def _category_order_sql(sort_by: str, sort_order: str) -> str:
    direction = "ASC" if str(sort_order or "").lower() == "asc" else "DESC"
    order_map = {
        "name": "c.name",
        "scene": "c.scene",
        "powerRating": "c.power_rating",
        "itemCount": "item_count",
        "fieldCount": "jsonb_array_length(COALESCE(c.field_schema, '[]'::jsonb))",
        "attachmentCount": "attachment_count",
        "status": "COALESCE(c.status, CASE WHEN c.review_status = 'disabled' THEN 'disabled' ELSE 'enabled' END)",
        "reviewStatus": "c.review_status",
        "updatedAt": "c.updated_at",
        "createdAt": "c.created_at",
    }
    expression = order_map.get(str(sort_by or "").strip(), "c.updated_at")
    return f"ORDER BY {expression} {direction} NULLS LAST, c.id DESC"


def _item_order_sql(sort_by: str, sort_order: str) -> str:
    direction = "ASC" if str(sort_order or "").lower() == "asc" else "DESC"
    order_map = {
        "projectName": "i.project_name",
        "customerName": "i.customer_name",
        "turbineModel": "i.turbine_model",
        "contractYear": "i.contract_year",
        "deliveryYear": "i.delivery_year",
        "operationYear": "i.operation_year",
        "categoryName": "c.name",
        "rowIndex": "i.row_index",
        "updatedAt": "i.updated_at",
        "createdAt": "i.created_at",
    }
    expression = order_map.get(str(sort_by or "").strip(), "i.updated_at")
    return f"ORDER BY {expression} {direction} NULLS LAST, i.id DESC"


def _safe_file_name(value: str) -> str:
    text_value = re.sub(r"[\\/:*?\"<>|]+", "-", str(value or "").strip())
    text_value = re.sub(r"\s+", " ", text_value).strip(" .")
    return text_value or "performance.docx"


performance_package_service = PerformancePackageService()
