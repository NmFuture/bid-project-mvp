from __future__ import annotations

"""技术标项目事实表专项抽取器。

针对通用启发式（facts_from_docx_material / facts_from_xlsx_material）无法覆盖的
专项素材提供结构化解析，产出与 technical_gap_fact_table.material_fact 同形的事实，
label 严格等于 app/data/technical_fact_field_specs.json 中的 spec label，
build 末尾的 reconcile_fact_fields_with_specs 据此零改动归位。

覆盖范围：
- 风资源评估报告 docx（承诺考核/保证双列表、代表年风速表、项目参数表）
- 塔架与基础工程量 docx/xlsx（钢塔 3 列大表状态机 + 混塔变体）
- 基础弯矩表 xlsx（表头行 + 单位行 + 工况数据行）
- 认证证书 pdf/docx（型式认证优先于设计认证；仅取文本层，OCR 为异步 worker，
  同步上下文不可用，文本层取不到时直接放弃该文件，绝不编造）
- 发电小时数承诺函（版本标识）与项目生产制造基地专题（名称/介绍）
"""

import re
from pathlib import Path
from typing import Any, Callable

from docx import Document
from openpyxl import load_workbook

from app.services.parsing import extract_pdf_text
from app.services.technical_fact_field_specs import load_specs
from app.services.technical_gap_fact_table import (
    clean_fact_text,
    facts_from_docx_material,
    material_fact,
)

# 专项置信度：结构化表头指纹命中，高于通用启发式
SPECIAL_FACT_CONFIDENCE = 0.88


def _spec_labels(reference_keyword: str, *, source_kind: str = "material") -> set[str]:
    """按 referenceFile 关键词取目标 spec label 集合（与清单原文严格一致）。"""
    labels: set[str] = set()
    for spec in load_specs():
        if str(spec.get("sourceKind") or "") != source_kind:
            continue
        if reference_keyword and reference_keyword not in str(spec.get("referenceFile") or ""):
            continue
        label = str(spec.get("label") or "").strip()
        if label:
            labels.add(label)
    return labels


def _dedupe_cells(cells: list[str]) -> list[str]:
    """合并单元格在 python-docx/openpyxl 中会重复出现，折叠相邻重复值。"""
    result: list[str] = []
    for cell in cells:
        if result and result[-1] == cell:
            continue
        result.append(cell)
    return result


def _clean_special_value(value: Any) -> str:
    """剥 ~/空格/全角空格前缀，保留负号、小数与括号注解。"""
    text = re.sub(r"\s+", "", str(value or ""))
    text = text.lstrip("~～")
    return text.strip("：:；;，,、")


def _first_number(value: Any) -> float | None:
    match = re.search(r"-?[0-9]+(?:\.[0-9]+)?", str(value or ""))
    return float(match.group(0)) if match else None


def _format_number(value: Any) -> str:
    """xlsx 数值统一成短字符串（46791.3 / -149696.92），非数值原样清洗。"""
    if isinstance(value, bool):
        return ""
    if isinstance(value, (int, float)):
        text = f"{float(value):.2f}".rstrip("0").rstrip(".")
        return text or "0"
    return _clean_special_value(value)


class _SpecialFactSink:
    """按 label 去重收集专项事实；只允许产出目标 spec label。"""

    def __init__(self, spec_labels: set[str], material: dict[str, Any]):
        self.spec_labels = spec_labels
        self.material = material
        self.facts: dict[str, dict[str, Any]] = {}

    def emit(self, label: str, value: Any, *, location: str = "", confidence: float = SPECIAL_FACT_CONFIDENCE) -> bool:
        label_text = str(label or "").strip()
        value_text = str(value or "").strip()
        if not label_text or not value_text:
            return False
        if label_text not in self.spec_labels or label_text in self.facts:
            return False
        self.facts[label_text] = material_fact(
            label_text, value_text, self.material, location=location, confidence=confidence
        )
        return True

    def __contains__(self, label: str) -> bool:
        return label in self.facts

    def result(self) -> list[dict[str, Any]]:
        return list(self.facts.values())


# ---------------------------------------------------------------------------
# a) 风资源评估报告
# ---------------------------------------------------------------------------

# 承诺值表行关键词 → 双列（承诺考核值/承诺保证值）取值后的 spec label 映射
_WIND_COMMITMENT_ROWS = [
    # (行关键词, 排除词, 考核值 labels, 保证值 labels, 是否剥 % )
    (
        "折减",
        "",
        ["折减系数（考核值，%）", "折减值（考核值）"],
        ["折减系数（保证值，%）", "折减值（保证值）"],
        True,
    ),
    (
        "全场净上网电量",
        "扇区",
        ["全场净上网电量（考核值，MWh/y）", "净发电量（第2列，MWh/y）"],
        ["全场净上网电量（保证值，MWh/y）", "净发电量（第1列，MWh/y）"],
        False,
    ),
    (
        "小时数",
        "扇区",
        [
            "等效上网小时数（考核值，h）",
            "年等效满负荷小时数（考核值，h）",
            "年等效满发小时数（考核值，h）",
            "有效小时数（第1列，h）",
        ],
        [
            "等效上网小时数（保证值，h）",
            "年等效满负荷小时数（保证值，h）",
            "年等效满发小时数（保证值，h）",
            "有效小时数（第2列，h）",
        ],
        False,
    ),
]


def _wind_commitment_table(
    rows: list[list[str]], sink: _SpecialFactSink, *, table_idx: int
) -> None:
    """承诺值双列表：定位"承诺考核值|承诺保证值"相邻列后按行关键词取值。"""
    assess_col = guarantee_col = -1
    for row in rows:
        for idx in range(len(row) - 1):
            if "承诺考核" in row[idx] and "承诺保证" in row[idx + 1]:
                assess_col, guarantee_col = idx, idx + 1
                break
        if assess_col >= 0:
            break
    if assess_col < 0:
        return
    header_row_idx = next(
        (i for i, row in enumerate(rows) if any("承诺考核" in cell for cell in row)), 0
    )
    header = rows[header_row_idx]
    # 承诺方式（第1列/第2列）：取双列表头原文
    if assess_col < len(header):
        sink.emit("承诺方式（第1列）", header[assess_col], location=f"T{table_idx}/R{header_row_idx + 1}")
    if guarantee_col < len(header):
        sink.emit("承诺方式（第2列）", header[guarantee_col], location=f"T{table_idx}/R{header_row_idx + 1}")
    for row_idx, row in enumerate(rows, start=1):
        if max(assess_col, guarantee_col) >= len(row):
            continue
        row_label = " ".join(cell for cell in row[:assess_col] if cell)
        if not row_label or "承诺" in row_label:
            continue
        assess_value = _clean_special_value(row[assess_col])
        guarantee_value = _clean_special_value(row[guarantee_col])
        if not assess_value and not guarantee_value:
            continue
        location = f"T{table_idx}/R{row_idx}"
        for keyword, exclude, assess_labels, guarantee_labels, strip_pct in _WIND_COMMITMENT_ROWS:
            if keyword not in row_label or (exclude and exclude in row_label):
                continue
            pair_assess, pair_guarantee = assess_value, guarantee_value
            if strip_pct:
                pair_assess = pair_assess.rstrip("%％")
                pair_guarantee = pair_guarantee.rstrip("%％")
            for label in assess_labels:
                sink.emit(label, pair_assess, location=location)
            for label in guarantee_labels:
                sink.emit(label, pair_guarantee, location=location)
        if re.search(r"^(?:总)?容量$", row_label.replace(" ", "")) or (
            "容量" in row_label and "单机" not in row_label and "容量系数" not in row_label
        ):
            sink.emit("投标总容量（MW）", assess_value or guarantee_value, location=location)


def _wind_rep_year_table(
    rows: list[list[str]], sink: _SpecialFactSink, *, table_idx: int
) -> None:
    """代表年风速表：表头含"代表年风速"，逐测风塔取 轮毂高度/代表年风速。"""
    header = rows[0] if rows else []
    rep_col = next((i for i, cell in enumerate(header) if "代表年风速" in cell), -1)
    if rep_col < 0:
        return
    hub_col = next((i for i, cell in enumerate(header) if "轮毂" in cell and "高度" in cell), -1)
    towers: list[str] = []
    speeds: list[str] = []
    heights: list[str] = []
    for row_idx, row in enumerate(rows[1:], start=2):
        if rep_col >= len(row):
            continue
        speed = _clean_special_value(row[rep_col])
        if _first_number(speed) is None:
            continue
        tower = _clean_special_value(row[0]) if row else ""
        height = _clean_special_value(row[hub_col]) if 0 <= hub_col < len(row) else ""
        if tower:
            towers.append(tower)
        speeds.append(speed)
        if height:
            heights.append(height)
    if not speeds:
        return
    location = f"T{table_idx}/R2"
    sink.emit("参考高度处年平均风速（m/s）", "/".join(speeds), location=location)
    if heights:
        unique_heights = list(dict.fromkeys(heights))
        sink.emit("测风塔参考高度（m）", "/".join(unique_heights), location=location)
    if towers:
        summary = "、".join(towers)
        if heights:
            summary += f"（轮毂高度{'/'.join(dict.fromkeys(heights))}m，代表年风速{'/'.join(speeds)}m/s）"
        sink.emit("指定测风塔/测风条件", summary, location=location)


def _wind_config_rows(
    rows: list[list[str]], config: dict[str, str], *, table_idx: int
) -> None:
    """项目参数/风电场配置表：扫描已知行标签（任意列），取行尾值。"""
    row_patterns = [
        (r"^轮毂(?:中心)?高度$", "hubHeight"),
        (r"^(?:适用等级|设计等级|安全等级)$", "safetyClass"),
        (r"^(?:风轮直径|叶轮直径)$", "rotorDiameter"),
        (r"^方案$", "scheme"),
        (r"^(?:机位数量|机组台数|风机台数|台数)$", "turbineCount"),
        (r"^(?:风场容量|总容量|建设容量|总装机容量|容量)$", "capacity"),
        (r"^(?:机型|机组型号|机组类型|投标机型)$", "model"),
    ]

    def valid(key: str, value: str) -> bool:
        if not value or value in {"-", "—", "/"}:
            return False
        if key in {"hubHeight", "rotorDiameter", "capacity", "turbineCount"}:
            return _first_number(value) is not None
        if key == "model":
            return bool(re.search(r"[A-Za-z]", value) and re.search(r"\d", value))
        if key == "safetyClass":
            return bool(re.search(r"[A-Za-z0-9]", value))
        return True

    for row in rows:
        cells = [cell for cell in row if cell]
        if len(cells) < 2:
            continue
        value = _clean_special_value(cells[-1])
        for cell in cells[:-1]:
            label = re.sub(r"[（(][^（）()]*[）)]", "", cell)
            for pattern, key in row_patterns:
                if re.match(pattern, label) and key not in config and valid(key, value):
                    config[key] = value


def _wind_tower_info_towers(rows: list[list[str]]) -> list[str]:
    """测风塔信息表（表头"测风塔|XX#|YY#"）→ 塔编号列表。"""
    header = rows[0] if rows else []
    if not header or header[0] not in {"测风塔", "测风塔编号"}:
        return []
    return [cell for cell in header[1:] if cell and re.search(r"#|号|塔", cell)]


def facts_from_wind_resource_docx(
    path: Path, material: dict[str, Any], project: dict[str, Any]
) -> list[dict[str, Any]] | None:
    """风资源评估报告专项抽取。非 docx 载体返回 None（回退通用启发式）。"""
    if path.suffix.lower() not in {".docx", ".doc"}:
        return None
    sink = _SpecialFactSink(_spec_labels("风资源报告"), material)
    try:
        document = Document(str(path))
    except Exception:
        return sink.result()
    config: dict[str, str] = {}
    fallback_towers: list[str] = []
    fallback_location = ""
    for table_idx, table in enumerate(document.tables, start=1):
        # 不做相邻去重：承诺考核/保证两列经常同值（600/600），去重会错位丢列
        rows = [[clean_fact_text(cell.text) for cell in row.cells] for row in table.rows]
        rows = [row for row in rows if any(row)]
        if not rows:
            continue
        _wind_commitment_table(rows, sink, table_idx=table_idx)
        _wind_rep_year_table(rows, sink, table_idx=table_idx)
        _wind_config_rows(rows, config, table_idx=table_idx)
        if not fallback_towers:
            towers = _wind_tower_info_towers(rows)
            if towers:
                fallback_towers = towers
                fallback_location = f"T{table_idx}/R1"
    # 测风塔信息表兜底（代表年风速表已给出带高度的完整口径时不覆盖）
    if fallback_towers and "指定测风塔/测风条件" not in sink:
        sink.emit(
            "指定测风塔/测风条件",
            "、".join(fallback_towers),
            location=fallback_location,
            confidence=0.85,
        )
    # 段落兜底：适用/设计等级（部分项目安全等级只在正文出现）
    if "safetyClass" not in config:
        for paragraph in document.paragraphs:
            text = clean_fact_text(paragraph.text)
            match = re.search(r"(?:适用等级|设计等级|安全等级)[^A-Za-z0-9]{0,8}(IEC\s*[A-Z0-9+]{1,4})", text)
            if match:
                config["safetyClass"] = match.group(1)
                break
    if "safetyClass" not in config:
        # 正文直接描述场址等级（"为IEC III类风场"+"湍流…属于IEC C类" → IEC IIIC）
        full_text = "\n".join(p.text for p in document.paragraphs)
        site = re.search(r"IEC\s*(I{1,3})\s*类风场", full_text)
        turb = re.search(r"湍流[^。\n]{0,40}IEC\s*([ABC])\s*类", full_text)
        if site and turb:
            config["safetyClass"] = f"IEC {site.group(1)}{turb.group(1)}"
        elif site:
            config["safetyClass"] = f"IEC {site.group(1)}类"
    if config.get("safetyClass"):
        safety = re.sub(r"^IEC\s*", "IEC ", config["safetyClass"]).strip()
        sink.emit("场址要求安全等级", safety, location="项目参数表")
    if config.get("hubHeight"):
        sink.emit("投标轮毂中心高度（m）", config["hubHeight"], location="项目参数表")
    if config.get("capacity"):
        sink.emit("投标总容量（MW）", config["capacity"], location="项目参数表")
    scheme = config.get("scheme")
    if not scheme and config.get("turbineCount") and config.get("model"):
        scheme = f"{config['turbineCount']}台{config['model']}"
    if scheme:
        sink.emit("投标方案配置/技术路线说明", scheme, location="项目参数表")
    # 投标叶尖高度 = 轮毂高度 + 叶轮直径/2（叶轮直径缺时从机型串或项目机型参数补）
    rotor = _first_number(config.get("rotorDiameter"))
    if rotor is None and config.get("model"):
        model_match = re.search(r"-(\d{3})(?:-\d+)?$", config["model"])
        rotor = float(model_match.group(1)) if model_match else None
    hub = _first_number(config.get("hubHeight"))
    if rotor is not None and hub is not None:
        tip = hub + rotor / 2
        tip_text = f"{tip:g}"
        sink.emit("投标叶尖高度（m）", tip_text, location="项目参数表", confidence=0.85)
    return sink.result()


# ---------------------------------------------------------------------------
# b) 塔架与基础工程量
# ---------------------------------------------------------------------------

# 直接行（合并单元格 col0==col1）命名差异 → spec label
_TOWER_DIRECT_LABEL_MAP = {
    "挖方（m3）": "基础挖方量（m³）",
    "填方（m3）": "基础填方量（m³）",
    "开挖量（m3）": "基础挖方量（m³）",
    "回填量（m3）": "基础填方量（m³）",
    "钢材型号Q355NE的筒节质量（kg）": "Q355NE筒节质量（kg）",
    "筒壁+法兰重量（kg）": "筒壁加法兰重量（kg）",
    "塔架总重（筒壁+法兰+内附件）（kg）": "塔架总重（筒壁+法兰+内附件，kg）",
    "基础混凝土用量（m3）": "基础混凝土用量（m³）",
    "垫层混凝土用量（m3）": "垫层混凝土用量（m³）",
    "垫层混凝土等级": "垫层混凝土型号",
    "垫层混凝土体积（m3）": "垫层混凝土用量（m³）",
    "基础混凝土等级": "基础混凝土型号",
    "基础混凝土体积（m3）": "基础混凝土用量（m³）",
    "钢筋标号": "基础钢筋型号",
    "钢筋用量（t）": "基础钢筋用量（kg）",
}

# 值需要 t → kg 换算由"原 label 以（t）结尾且映射后以（kg）结尾"判定，见下方直接行分支

_CHINESE_DIGITS = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9, "十": 10}


def _tower_section_prefix(text: str) -> str:
    """col0 段名（如"第5段(顶) 塔节Q355NE"）→ spec 段前缀"第5段（顶）"。"""
    match = re.match(r"^第([0-9一二三四五六七八九十]+)段\s*(?:[（(]\s*(顶|底)\s*[）)])?", text)
    if not match:
        return ""
    token = match.group(1)
    number = _CHINESE_DIGITS.get(token, None)
    if number is None:
        try:
            number = int(token)
        except ValueError:
            return ""
    suffix = f"（{match.group(2)}）" if match.group(2) else ""
    return f"第{number}段{suffix}"


def _tower_rows_from_path(path: Path) -> list[tuple[str, list[str]]] | None:
    """返回 (location 前缀, 行 cells) 序列；不支持的载体返回 None。"""
    suffix = path.suffix.lower()
    rows: list[tuple[str, list[str]]] = []
    if suffix in {".docx", ".doc"}:
        try:
            document = Document(str(path))
        except Exception:
            return rows
        for table_idx, table in enumerate(document.tables, start=1):
            for row_idx, row in enumerate(table.rows, start=1):
                cells = _dedupe_cells([clean_fact_text(cell.text) for cell in row.cells])
                if any(cells):
                    rows.append((f"T{table_idx}/R{row_idx}", cells))
        return rows
    if suffix in {".xlsx", ".xlsm"}:
        try:
            workbook = load_workbook(path, data_only=True, read_only=True)
        except Exception:
            return rows
        for worksheet in workbook.worksheets:
            for row_idx, row in enumerate(worksheet.iter_rows(values_only=True), start=1):
                cells = _dedupe_cells([clean_fact_text(cell) for cell in row])
                if any(cells):
                    rows.append((f"{worksheet.title}!R{row_idx}", cells))
        return rows
    return None


def facts_from_tower_quantity_docx(
    path: Path, material: dict[str, Any], project: dict[str, Any]
) -> list[dict[str, Any]] | None:
    """塔架与基础工程量专项抽取（钢塔 3 列大表 + 混塔变体）。"""
    rows = _tower_rows_from_path(path)
    if rows is None:
        return None
    sink = _SpecialFactSink(_spec_labels("塔架与基础工程量"), material)
    current_section = ""  # 钢塔：当前"第N段（顶/底）"
    mixed_mode = ""  # 混塔：concrete / steel / foundation
    for location, cells in rows:
        first = cells[0]
        # 混塔分区标题（"一、混凝土塔段" / "二、钢塔段（含法兰）" / "三、基础参数"）
        if len(cells) == 1:
            if re.search(r"钢塔段", first):
                mixed_mode = "steel"
            elif re.search(r"混凝土塔段", first):
                mixed_mode = "concrete"
            elif re.search(r"基础", first):
                mixed_mode = "foundation"
            continue
        section = _tower_section_prefix(first)
        plain_section = bool(re.fullmatch(r"第[0-9一二三四五六七八九十]+段", first))
        if section and len(cells) >= 3 and not plain_section:
            # 钢塔段明细行：col0=段名（含塔节/材质），col1=子项，col2=值
            if re.search(r"塔节|筒节|Q355|Q\d{3}", first):
                current_section = section
                item = cells[1]
                value = _clean_special_value(cells[2])
                sink.emit(f"{current_section}塔节{item}", value, location=location)
                continue
        if section and plain_section:
            # 混塔钢塔段"第一段"行：材料/重量/长度按列位置取
            current_section = section
            if mixed_mode == "steel":
                # 混塔钢塔段属于塔架顶部，段号与 spec 的"第N段"映射不可靠，
                # 重量/长度/直径不猜，只取确定无疑的钢板材规格。
                material_cells = [cell for cell in cells[1:] if re.search(r"Q\d{3}", cell)]
                if material_cells:
                    sink.emit("筒节钢板材规格型号", material_cells[0], location=location)
            continue
        if len(cells) == 2:
            label, raw_value = cells
            value = _clean_special_value(raw_value)
            mapped = _TOWER_DIRECT_LABEL_MAP.get(label, label.replace("（m3）", "（m³）"))
            if label.endswith("（t）") and mapped.endswith("（kg）"):
                # 混塔钢材/钢筋按 t 计，spec 单位为 kg，×1000 换算
                number = _first_number(value)
                if number is not None:
                    value = f"{number * 1000:g}"
            sink.emit(mapped, value, location=location)
            continue
        if len(cells) >= 3 and "塔筒内附件" in first:
            number = next((n for n in (_first_number(cell) for cell in cells[1:]) if n is not None), None)
            if number is not None:
                # 混塔单位为 t，钢塔样本直接给 kg；按数量级判断（<1000 视为 t）
                kg = number * 1000 if number < 1000 else number
                sink.emit("塔筒内附件近似重量（kg）", f"{kg:g}", location=location)
            continue
    return sink.result()


# ---------------------------------------------------------------------------
# c) 基础弯矩表
# ---------------------------------------------------------------------------

# 工况名 → spec 工况前缀（可扩充）
_MOMENT_CONDITION_MAP = {
    "正常运行载荷工况": "正常工况",
    "极端载荷工况": "极端工况",
}

# 弯矩分量表头 → spec 分量名（Safety factor 无单位，单独处理）
_MOMENT_COMPONENT_ALIASES = {
    "safetyfactor": "安全系数",
    "安全系数": "安全系数",
}


def _moment_condition(text: str) -> str:
    normalized = re.sub(r"\s+", "", str(text or ""))
    if normalized in _MOMENT_CONDITION_MAP:
        return _MOMENT_CONDITION_MAP[normalized]
    if "极端" in normalized:
        return "极端工况"
    if "正常" in normalized:
        return "正常工况"
    return ""


def facts_from_foundation_moment_xlsx(
    path: Path, material: dict[str, Any], project: dict[str, Any]
) -> list[dict[str, Any]] | None:
    """基础弯矩表专项抽取：定位含 Mx 的表头行 + 单位行，按工况行取值。"""
    if path.suffix.lower() not in {".xlsx", ".xlsm"}:
        return None
    sink = _SpecialFactSink(_spec_labels("弯矩"), material)
    try:
        workbook = load_workbook(path, data_only=True, read_only=True)
    except Exception:
        return sink.result()
    for worksheet in workbook.worksheets:
        rows = [list(row) for row in worksheet.iter_rows(values_only=True)]
        header_idx = -1
        for idx, row in enumerate(rows[:30]):
            cells = [clean_fact_text(cell) for cell in row]
            if "Mx" in cells and "Fx" in cells:
                header_idx = idx
                break
        if header_idx < 0:
            continue
        header = [clean_fact_text(cell) for cell in rows[header_idx]]
        unit_row = (
            [clean_fact_text(cell) for cell in rows[header_idx + 1]]
            if header_idx + 1 < len(rows)
            else []
        )
        # 列映射：分量名 → (列号, 单位)
        columns: list[tuple[str, str, int]] = []
        for col_idx, name in enumerate(header):
            if not name or name in {"载荷工况", "-"}:
                continue
            component = _MOMENT_COMPONENT_ALIASES.get(name.lower(), name)
            if not re.fullmatch(r"(?:M|F)(?:x|y|z|xy)|安全系数", component):
                continue
            unit = unit_row[col_idx] if col_idx < len(unit_row) else ""
            if component == "安全系数":
                unit = ""
            columns.append((component, unit, col_idx))
        for row_idx in range(header_idx + 2, len(rows)):
            row = rows[row_idx]
            if not row:
                continue
            condition_cell = next((cell for cell in row if clean_fact_text(cell)), "")
            condition = _moment_condition(condition_cell)
            if not condition:
                continue
            location = f"{worksheet.title}!R{row_idx + 1}"
            for component, unit, col_idx in columns:
                if col_idx >= len(row):
                    continue
                value = _format_number(row[col_idx])
                if not value:
                    continue
                if component == "安全系数":
                    label = f"{condition}-安全系数"
                else:
                    label = f"{condition}-{component}（{unit}）"
                sink.emit(label, value, location=location)
    return sink.result()


# ---------------------------------------------------------------------------
# d) 认证证书（型式认证 > 设计认证）
# ---------------------------------------------------------------------------

_CERT_TEXT_MIN_LENGTH = 200  # 低于该长度视为扫描件/水印页（如"供招投标使用"），放弃


def _certificate_text(path: Path) -> str:
    """取证书文本层（pdf/docx）。OCR 是异步 worker，同步上下文不可用，
    文本层为空或过短时直接返回 ""（标缺少来源，绝不编造）。"""
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            text, info = extract_pdf_text(path)
            if info.get("requiresOcr"):
                return ""
            return text or ""
        if suffix in {".docx", ".doc"}:
            document = Document(str(path))
            parts = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    parts.extend(cell.text for cell in row.cells)
            return "\n".join(parts)
    except Exception:
        return ""
    return ""


def _certificate_pdf_pages(path: Path) -> list[str]:
    """按页提取 PDF 文本层（pypdf，与 extract_pdf_text 同库，不加新依赖）。

    仅 PDF 有物理页码概念；非 PDF 或提取失败返回 []（不写 page）。"""
    if path.suffix.lower() != ".pdf":
        return []
    try:
        from pypdf import PdfReader
    except ImportError:  # pragma: no cover - 与 extract_pdf_text 同依赖
        return []
    try:
        reader = PdfReader(str(path))
        return [(page.extract_text() or "") for page in reader.pages]
    except Exception:
        return []


def _certificate_param_pages(pages: list[str], params: dict[str, str]) -> dict[str, int]:
    """定位每个命中参数首次出现的页码（1 起始）。找不到则不记。"""
    hits: dict[str, int] = {}
    if not pages or not params:
        return hits
    for index, page_text in enumerate(pages, start=1):
        page_params = parse_certificate_wind_params(page_text)
        for key, value in params.items():
            if key not in hits and page_params.get(key) == value:
                hits[key] = index
    return hits


def parse_certificate_wind_params(text: str) -> dict[str, str]:
    """从证书文本提取核心风资源参数（纯函数，便于测试）。

    返回键：vref（10min 平均极限风速 m/s）、iecClass、turbulence、extreme50。
    """
    params: dict[str, str] = {}
    if not text or len(text.strip()) < 20:
        return params
    vref = re.search(
        r"(?:参考风速|Vref|10\s*(?:min|分钟)\s*平均(?:极限|最大)?风速)\s*[:：为=]?\s*([0-9]+(?:\.[0-9]+)?)",
        text,
        flags=re.I,
    )
    if vref:
        params["vref"] = vref.group(1)
    extreme50 = re.search(r"50\s*年一遇[^0-9]{0,15}([0-9]+(?:\.[0-9]+)?)\s*m/?s", text, flags=re.I)
    if extreme50:
        params["extreme50"] = extreme50.group(1)
    iec = re.search(
        r"(?:安全等级|设计等级|认证等级|IEC\s*标准等级)[^A-Za-z0-9]{0,10}(IEC\s*[A-Z0-9+]{1,5})",
        text,
        flags=re.I,
    ) or re.search(r"IEC\s*(I{1,3}\s*[ABC]\+?|S)\b", text, flags=re.I)
    if iec:
        params["iecClass"] = re.sub(r"\s+", " ", iec.group(1)).strip()
    turbulence = re.search(
        r"湍流(?:强度)?(?:期望值)?\s*[:：为=]?\s*(0?\.\d+|[0-9]+(?:\.[0-9]+)?)\s*%?",
        text,
        flags=re.I,
    )
    if turbulence:
        params["turbulence"] = turbulence.group(1)
    return params


def _certificate_rank(material: dict[str, Any]) -> int:
    """型式认证优先于设计认证。"""
    name = str(material.get("name") or material.get("cleanedFileName") or "")
    if "型式认证" in name:
        return 0
    if "设计认证" in name:
        return 1
    return 2


def facts_from_certificate_materials(
    cert_materials: list[tuple[dict[str, Any], Path]], project: dict[str, Any]
) -> list[dict[str, Any]]:
    """同项目多本证书按"型式认证 > 设计认证"排序，同字段先中先得。"""
    sink_specs = _spec_labels("", source_kind="cert")
    ordered = sorted(cert_materials, key=lambda item: _certificate_rank(item[0]))
    emitted: dict[str, dict[str, Any]] = {}
    for material, path in ordered:
        text = _certificate_text(path)
        if len(text.strip()) < _CERT_TEXT_MIN_LENGTH:
            continue
        params = parse_certificate_wind_params(text)
        # 仅 PDF 证书补来源页码（docx 无物理页码概念，不写 page 键）
        param_pages = _certificate_param_pages(_certificate_pdf_pages(path), params)
        candidates: list[tuple[str, str, str]] = []
        if params.get("vref"):
            candidates.append(("机型认证10分钟平均极限风速（m/s）", params["vref"], "vref"))
        if params.get("iecClass"):
            candidates.append(("机型认证安全等级", params["iecClass"], "iecClass"))
        if params.get("turbulence"):
            candidates.append(("机型认证湍流强度（%）", params["turbulence"], "turbulence"))
        # 设计认证核心风资源参数：按 spec note 拼摘要串（有几个拼几个，不编造）
        pieces: list[str] = []
        piece_param_keys: list[str] = []
        if params.get("vref"):
            pieces.append(f"参考风速{params['vref']}m/s")
            piece_param_keys.append("vref")
        if params.get("turbulence"):
            pieces.append(f"湍流强度{params['turbulence']}")
            piece_param_keys.append("turbulence")
        if params.get("extreme50"):
            pieces.append(f"50年一遇极端风速{params['extreme50']}m/s")
            piece_param_keys.append("extreme50")
        if pieces:
            candidates.append(("设计认证核心风资源参数", "，".join(pieces), piece_param_keys[0]))
        for label, value, param_key in candidates:
            if label not in sink_specs or label in emitted:
                continue
            fact = material_fact(
                label, value, material, location="证书文本层", confidence=0.85
            )
            page = param_pages.get(param_key)
            if page is not None:
                fact["sourceRef"]["page"] = page
            emitted[label] = fact
    return list(emitted.values())


# ---------------------------------------------------------------------------
# e) 零星来源：发电小时数承诺函 / 生产制造基地专题
# ---------------------------------------------------------------------------


def facts_from_hours_commitment_docx(
    path: Path, material: dict[str, Any], project: dict[str, Any]
) -> list[dict[str, Any]] | None:
    """发电小时数承诺函：版本标识（承诺保证值/承诺考核值）+ 通用启发式（保证矩阵等）。"""
    if path.suffix.lower() not in {".docx", ".doc"}:
        return None
    facts = facts_from_docx_material(path, material)
    name = str(material.get("name") or material.get("cleanedFileName") or "")
    title = name
    try:
        document = Document(str(path))
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                title = f"{name} {paragraph.text.strip()}"
                break
    except Exception:
        pass
    version = ""
    if "承诺保证值" in title:
        version = "承诺保证值"
    elif "承诺考核值" in title:
        version = "承诺考核值"
    # spec label"发电小时数/电量承诺函版本"含"发电小时"，canonical 会坍缩成"保证有效小时数"，
    # 因此产出别名"电量承诺函版本"，经 SPEC_LABEL_ALIASES 归位到该 spec。
    if version:
        facts.append(
            material_fact(
                "电量承诺函版本",
                version,
                material,
                location="标题",
                confidence=0.9,
            )
        )
    return facts


def facts_from_production_base_docx(
    path: Path, material: dict[str, Any], project: dict[str, Any]
) -> list[dict[str, Any]] | None:
    """生产制造基地专题：基地名称（标题"-"后缀）与基地介绍（概况段，截 200 字）。"""
    if path.suffix.lower() not in {".docx", ".doc"}:
        return None
    sink = _SpecialFactSink(_spec_labels("生产制造基地"), material)
    try:
        document = Document(str(path))
    except Exception:
        return sink.result()
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    if not paragraphs:
        return sink.result()
    name = str(material.get("name") or material.get("cleanedFileName") or "")
    is_blade = bool(re.search(r"叶片", name)) or any("叶片" in p for p in paragraphs[:1])
    name_label = "叶片供货制造基地名称" if is_blade else "主机供货制造基地名称"
    intro_label = "叶片供货制造基地介绍" if is_blade else "主机供货制造基地介绍"
    title = paragraphs[0].lstrip("*")
    base_name = title.split("-")[-1].strip() if "-" in title else title
    if base_name:
        sink.emit(name_label, base_name, location="P1", confidence=0.85)
    intro_parts: list[str] = []
    capture = False
    for paragraph in paragraphs[1:]:
        if "概况" in paragraph and len(paragraph) <= 12:
            capture = True
            continue
        if capture:
            intro_parts.append(paragraph)
            if sum(len(p) for p in intro_parts) >= 200:
                break
    if intro_parts:
        intro = "".join(intro_parts)[:200]
        sink.emit(intro_label, intro, location="生产基地概况", confidence=0.85)
    return sink.result()


# ---------------------------------------------------------------------------
# 路由器
# ---------------------------------------------------------------------------


def special_extractor_for_material(material: dict[str, Any]) -> str | None:
    """按 material 名称/路径关键词路由到专项解析器种类，未命中返回 None。"""
    name = str(material.get("name") or material.get("cleanedFileName") or "")
    text = " ".join(
        str(material.get(key) or "")
        for key in ("name", "cleanedFileName", "folderPath", "path")
    )
    if re.search(r"型式认证|设计认证", name):
        return "certificate"
    if re.search(r"弯矩", text):
        return "foundation_moment"
    if re.search(r"发电小时数承诺函|电量承诺", text):
        return "hours_commitment"
    if re.search(r"生产制造基地|供货制造基地", text):
        return "production_base"
    if re.search(r"风资源", text):
        return "wind_resource"
    if re.search(r"工程量|塔架", text):
        return "tower_quantity"
    return None


_SPECIAL_EXTRACTORS: dict[str, Callable[[Path, dict[str, Any], dict[str, Any]], list[dict[str, Any]] | None]] = {
    "wind_resource": facts_from_wind_resource_docx,
    "tower_quantity": facts_from_tower_quantity_docx,
    "foundation_moment": facts_from_foundation_moment_xlsx,
    "hours_commitment": facts_from_hours_commitment_docx,
    "production_base": facts_from_production_base_docx,
}


def run_special_extractor(
    kind: str, path: Path, material: dict[str, Any], project: dict[str, Any]
) -> list[dict[str, Any]] | None:
    """执行专项解析器；返回 None 表示该文件不适用（调用方回退通用启发式）。"""
    extractor = _SPECIAL_EXTRACTORS.get(kind)
    if extractor is None:
        return None
    return extractor(path, material, project)
