from __future__ import annotations

import base64
import binascii
import asyncio
import json
import re
import shutil
import subprocess
import sys
import threading
from datetime import UTC, datetime
from pathlib import Path
from typing import Any, Callable
from urllib.parse import quote

from docx import Document

from app.core.config import BASE_DIR, settings
from app.services.identity import build_project_material_scope
from app.services.material_store import material_store
from app.services.minio_client import minio_client
from app.services.opencode_client import OpencodeClient
from app.services.turbine_models import project_turbine_model
from app.services.workspace_artifacts import legacy_workspace_roots, technical_workspace_dir, technical_workspace_stage_dir


GAP_PLAN_SCHEMA_VERSION = "bid-tech-gap-plan-v1"
TABLE_FILL_SCHEMA_VERSION = "bid-tech-table-fill-v1"
WORD_FILL_SCHEMA_VERSION = "bid-tech-word-placeholder-fill-v1"
GAP_PLANNER_SKILL_NAME = "bid-tech-gap-planner"
TABLE_FILL_SKILL_NAME = "bid-tech-table-filler"
WORD_FILL_SKILL_NAME = "bid-tech-word-placeholder-filler"
GAP_PLANNER_RUNNER = BASE_DIR / "opencode" / "skill" / GAP_PLANNER_SKILL_NAME / "scripts" / "run_from_manifest.py"
TABLE_FILL_RUNNER = BASE_DIR / "opencode" / "skill" / TABLE_FILL_SKILL_NAME / "scripts" / "run_from_manifest.py"
WORD_FILL_RUNNER = BASE_DIR / "opencode" / "skill" / WORD_FILL_SKILL_NAME / "scripts" / "run_from_manifest.py"


def now_iso() -> str:
    return datetime.now(UTC).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def safe_filename(value: str, fallback: str) -> str:
    text = re.sub(r"[\\/:*?\"<>|]+", "-", str(value or "").strip())
    text = re.sub(r"\s+", " ", text).strip(" .")
    return text or fallback


def build_gap_plan_for_project(project: dict[str, Any]) -> dict[str, Any]:
    """Create a real plan from the confirmed directory and parse/material refs.

    The OpenCode skill owns the contract and is available in the image. The
    backend uses the same manifest and runner locally so tests and offline
    deployments do not depend on a live model just to materialize the JSON
    contract.
    """

    project_id = str(project.get("id") or "")
    project_dir = _project_dir(project)
    work_dir = technical_workspace_stage_dir(project_id, "s4_gap_workdir")
    if work_dir.exists():
        shutil.rmtree(work_dir)
    work_dir.mkdir(parents=True, exist_ok=True)

    toc_json_path = _resolve_toc_json(project, work_dir)
    parse_result_path = work_dir / "parse_result.json"
    parse_result_path.write_text(
        json.dumps(project.get("parse_result") or {}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    output_file = work_dir / "gap_plan.json"
    manifest_path = work_dir / "s4_gap_input.json"
    wiki_dir = _resolve_wiki_dir(project, project_dir, work_dir)
    turbine_model = project_turbine_model(project)
    material_scope = build_project_material_scope(project)
    material_index = _allowed_material_index(material_scope, turbine_model)
    manifest = {
        "projectId": project_id,
        "projectName": str(project.get("name") or project_id),
        "bidType": str(project.get("bidType") or "技术标"),
        "workDir": str(work_dir),
        "tocJsonPath": str(toc_json_path),
        "wikiDir": str(wiki_dir) if wiki_dir else "",
        "parseResultPath": str(parse_result_path),
        "projectIdentity": project.get("identity") or {},
        "materialScope": material_scope,
        "materialIndex": material_index,
        "projectTurbineModel": turbine_model,
        "outputFile": str(output_file),
    }
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    result = run_gap_planner_skill(manifest_path)
    plan_path = Path(str(result.get("outputFile") or output_file))
    plan = json.loads(plan_path.read_text(encoding="utf-8"))
    _validate_gap_plan_toc_coverage(plan, toc_json_path)
    plan["projectTurbineModel"] = turbine_model
    plan["planFile"] = str(plan_path)
    plan["manifestPath"] = str(manifest_path)
    plan["phase"] = "gap_detection"
    plan["scopeBoundary"] = material_scope
    normalize_gap_plan_fill_task_skills(plan)
    plan["opencodeOutput"] = result.get("opencodeOutput") or {
        "status": "received",
        "sessionId": str(manifest_path),
        "providerId": "local-skill",
        "modelId": GAP_PLANNER_SKILL_NAME,
        "receivedAt": now_iso(),
        "parts": [{"type": "text", "text": json.dumps({"outputFile": str(plan_path)}, ensure_ascii=False)}],
    }
    return plan


def _is_material_word_fill_task(item: dict[str, Any], task: dict[str, Any]) -> bool:
    blank = task.get("blankSource") if isinstance(task.get("blankSource"), dict) else {}
    source_type = str(blank.get("sourceType") or "")
    usage = str(item.get("usage") or "")
    try:
        placeholder_count = int(blank.get("placeholderCount") or 0)
    except (TypeError, ValueError):
        placeholder_count = 0
    material_id = str(blank.get("materialId") or blank.get("id") or "")
    return (
        source_type == "material_fill_template"
        or (
            usage in {"section_fill", "chapter_fill"}
            and placeholder_count > 0
            and material_id.startswith("RAW-")
        )
    )


def normalize_gap_plan_fill_task_skills(plan: dict[str, Any]) -> int:
    """Repair stale gap plans whose material Word fill tasks still use table filler."""

    repaired = 0
    for item in _object_items(plan.get("items")):
        for task in _object_items(item.get("fillTasks")):
            current = str(task.get("skill") or "")
            expected = WORD_FILL_SKILL_NAME if _is_material_word_fill_task(item, task) else ""
            if not expected or current == expected:
                continue
            task["skill"] = expected
            repaired += 1
    if repaired:
        plan["summary"] = summarize_gap_plan(plan)
    return repaired


def _validate_gap_plan_toc_coverage(plan: dict[str, Any], toc_json_path: Path) -> None:
    try:
        toc = json.loads(toc_json_path.read_text(encoding="utf-8"))
    except Exception as exc:  # pragma: no cover - malformed workspace input
        raise RuntimeError(f"缺口识别无法读取审核目录：{toc_json_path}") from exc
    toc_items = _object_items(toc.get("items"))
    plan_items = _object_items(plan.get("items"))
    expected_count = len(toc_items)
    actual_count = len(plan_items)
    summary = plan.get("summary") if isinstance(plan.get("summary"), dict) else {}
    summary_count = int(summary.get("totalTocItems") or 0)
    if expected_count == actual_count and (summary_count in (0, expected_count)):
        return
    raise RuntimeError(
        "缺口识别结果不完整："
        f"S2 审核目录有 {expected_count} 个目录项，"
        f"Skill 输出 {actual_count} 个目录项，"
        f"summary.totalTocItems={summary_count}。请重新运行 bid-tech-gap-planner。"
    )


def _run_async(awaitable: Any) -> Any:
    try:
        asyncio.get_running_loop()
    except RuntimeError:
        return asyncio.run(awaitable)
    result: dict[str, Any] = {}
    error: dict[str, BaseException] = {}

    def runner() -> None:
        try:
            result["value"] = asyncio.run(awaitable)
        except BaseException as exc:  # pragma: no cover - re-raised in caller
            error["value"] = exc

    thread = threading.Thread(target=runner, daemon=True)
    thread.start()
    thread.join()
    if error:
        raise error["value"]
    return result.get("value")


def _allowed_material_index(material_scope: dict[str, Any], turbine_model: dict[str, Any]) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    seen: set[str] = set()
    for scope in material_scope.get("readableScopes") or []:
        if not isinstance(scope, dict):
            continue
        folder_path = str(scope.get("path") or "").strip()
        if not folder_path:
            continue
        payload = _run_async(
            material_store.raw_files(
                folder_path=folder_path,
                bid_type=str(material_scope.get("bidType") or "技术标"),
                material_tier=str(scope.get("materialTier") or ""),
                turbine_model=turbine_model,
                recursive=True,
                page=1,
                page_size=1000,
            )
        )
        for raw in payload.get("items") or []:
            if not isinstance(raw, dict):
                continue
            material_id = str(raw.get("id") or "")
            if not material_id or material_id in seen:
                continue
            seen.add(material_id)
            items.append(
                {
                    "id": material_id,
                    "name": str(raw.get("name") or ""),
                    "folderPath": str(raw.get("folderPath") or ""),
                    "materialTier": str(raw.get("materialTier") or scope.get("materialTier") or ""),
                    "hasCleanedWord": bool(raw.get("hasCleanedWord")),
                    "cleanedFileName": str(raw.get("cleanedFileName") or ""),
                    "cleanStatus": str(raw.get("cleanStatus") or ""),
                    "turbineModelLabel": str(raw.get("turbineModelLabel") or ""),
                    "updatedAt": str(raw.get("updatedAt") or ""),
                }
            )
    return items


def _object_items(value: Any) -> list[dict[str, Any]]:
    return [item for item in value if isinstance(item, dict)] if isinstance(value, list) else []


def _string_items(value: Any) -> list[str]:
    if not isinstance(value, list):
        return []
    result: list[str] = []
    seen: set[str] = set()
    for item in value:
        text = str(item or "").strip()
        if not text or text in seen:
            continue
        seen.add(text)
        result.append(text)
    return result


def _material_key(material: dict[str, Any]) -> str:
    return str(material.get("id") or material.get("materialId") or material.get("path") or material.get("docx") or "").strip()


def _material_summary(material: dict[str, Any]) -> dict[str, Any]:
    return {
        "id": str(material.get("id") or material.get("materialId") or ""),
        "name": str(material.get("name") or material.get("title") or material.get("fileName") or ""),
        "path": str(material.get("path") or material.get("docx") or ""),
        "folderPath": str(material.get("folderPath") or ""),
        "materialTier": str(material.get("materialTier") or material.get("materialScope") or ""),
        "usage": str(material.get("usage") or ""),
        "source": str(material.get("source") or ""),
        "hasCleanedWord": bool(material.get("hasCleanedWord")),
        "cleanedFileName": str(material.get("cleanedFileName") or ""),
        "turbineModelLabel": str(material.get("turbineModelLabel") or ""),
        "matchReason": str(material.get("matchReason") or ""),
        "turbineFit": str(material.get("turbineFit") or ""),
    }


def _dedupe_material_summaries(materials: list[dict[str, Any]]) -> list[dict[str, Any]]:
    result: list[dict[str, Any]] = []
    seen: set[str] = set()
    for material in materials:
        key = _material_key(material)
        if not key or key in seen:
            continue
        seen.add(key)
        result.append(_material_summary(material))
    return result


def _appendix_task_for_fill(item: dict[str, Any], task: dict[str, Any]) -> dict[str, Any]:
    blank_source = task.get("blankSource") if isinstance(task.get("blankSource"), dict) else {}
    blank_id = str(blank_source.get("id") or "").strip()
    appendix_tasks = _object_items(item.get("appendixTasks"))
    for appendix_task in appendix_tasks:
        if blank_id and str(appendix_task.get("id") or "") == blank_id:
            return dict(appendix_task)
    return dict(appendix_tasks[0]) if appendix_tasks else {}


def _selected_reference_material_ids(
    item: dict[str, Any],
    appendix_task: dict[str, Any],
    data: dict[str, Any],
) -> list[str]:
    requested = _string_items(data.get("referenceMaterialIds"))
    if requested:
        return requested

    matched = [_material_key(material) for material in _object_items(item.get("matchedMaterials"))]
    matched = [item for item in matched if item]
    if matched:
        return matched

    recommended = [
        _material_key(material)
        for material in _object_items(appendix_task.get("recommendedMaterials"))[:1]
    ]
    return [item for item in recommended if item]


def _reference_materials_for_fill(
    item: dict[str, Any],
    appendix_task: dict[str, Any],
    data: dict[str, Any],
    selected_ids: list[str],
) -> list[dict[str, Any]]:
    context = _dedupe_material_summaries(
        _object_items(data.get("referenceMaterials"))
        + _object_items(item.get("matchedMaterials"))
        + _object_items(item.get("candidateMaterials"))
        + _object_items(appendix_task.get("recommendedMaterials"))
        + [
            material
            for task in _object_items(item.get("appendixTasks"))
            for material in _object_items(task.get("recommendedMaterials"))
        ]
    )
    by_id = {
        str(material.get("id") or "").strip(): material
        for material in context
        if str(material.get("id") or "").strip()
    }
    result: list[dict[str, Any]] = []
    seen: set[str] = set()
    for material_id in selected_ids:
        if material_id in seen:
            continue
        seen.add(material_id)
        result.append(by_id.get(material_id) or {"id": material_id, "name": material_id})
    return result


def _material_index_for_fill(project: dict[str, Any], plan: dict[str, Any], item: dict[str, Any]) -> list[dict[str, Any]]:
    item_candidates = _dedupe_material_summaries(
        _object_items(plan.get("materialIndex"))
        + _object_items(item.get("candidateMaterials"))
        + _object_items(item.get("matchedMaterials"))
        + [
            material
            for task in _object_items(item.get("appendixTasks"))
            for material in _object_items(task.get("recommendedMaterials"))
        ]
    )

    material_scope = plan.get("scopeBoundary") if isinstance(plan.get("scopeBoundary"), dict) else {}
    if not material_scope:
        material_scope = build_project_material_scope(project)
    scoped_index = _allowed_material_index(material_scope, project_turbine_model(project))
    return _dedupe_material_summaries(item_candidates + scoped_index)


def _prepare_material_index_files(
    material_index: list[dict[str, Any]],
    work_dir: Path,
    *,
    cache_dir: Path | None = None,
    limit: int = 80,
) -> list[dict[str, Any]]:
    prepared: list[dict[str, Any]] = []
    cache_dir = cache_dir or (work_dir / "material_index")
    for material in material_index[:limit]:
        item = dict(material)
        material_id = str(item.get("id") or item.get("materialId") or "").strip()
        if not material_id:
            prepared.append(item)
            continue
        try:
            payload, source_kind = _run_async(_downloadable_fill_source_payload(material_id))
        except Exception:
            prepared.append(item)
            continue
        file_name = safe_filename(
            str(item.get("cleanedFileName") or payload.get("fileName") or item.get("name") or f"{material_id}.docx"),
            f"{material_id}.docx",
        )
        target_path = cache_dir / f"{material_id}-{file_name}"
        if not target_path.exists():
            minio_client.download_file(str(payload["bucket"]), str(payload["key"]), target_path)
        item.update(
            {
                "path": str(target_path),
                "sourceKind": source_kind,
                "fileName": target_path.name,
            }
        )
        prepared.append(item)
    return prepared


def _field_key(field: dict[str, Any]) -> str:
    return str(field.get("id") or field.get("key") or field.get("label") or field.get("title") or "").strip()


def _field_summary(field: dict[str, Any]) -> dict[str, Any]:
    return {
        "id": str(field.get("id") or field.get("key") or field.get("label") or ""),
        "label": str(field.get("label") or field.get("title") or field.get("keyEntity") or field.get("id") or ""),
        "value": str(field.get("value") or field.get("keyValue") or ""),
        "sourceFile": str(field.get("sourceFile") or ""),
        "evidence": str(field.get("evidence") or ""),
        "evidenceLocation": str(field.get("evidenceLocation") or ""),
    }


def _parse_fields_for_fill(appendix_task: dict[str, Any], task: dict[str, Any], data: dict[str, Any]) -> list[dict[str, Any]]:
    candidates = _object_items(appendix_task.get("availableParseFields")) + _object_items(appendix_task.get("fields"))
    requested = _string_items(data.get("parseFieldIds"))
    blank_source = task.get("blankSource") if isinstance(task.get("blankSource"), dict) else {}
    blank_id = str(blank_source.get("id") or "").strip()
    requested_field_ids = [item for item in requested if item != blank_id]
    if not requested_field_ids:
        return [_field_summary(field) for field in candidates]

    by_key: dict[str, dict[str, Any]] = {}
    for field in candidates:
        key = _field_key(field)
        if key:
            by_key[key] = field
    result: list[dict[str, Any]] = []
    for field_id in requested_field_ids:
        result.append(_field_summary(by_key.get(field_id) or {"id": field_id, "label": field_id}))
    return result


def run_gap_planner_skill(manifest_path: Path) -> dict[str, Any]:
    prompt = _build_gap_planner_prompt(manifest_path)
    try:
        return OpencodeClient().run_bid_tech_gap_planner_with_trace(prompt)
    except Exception:
        # Keep S4 usable in offline tests and local environments. The fallback
        # executes the same skill command contract that OpenCode is instructed
        # to run, so the generated gap plan schema stays identical.
        return _run_local_skill_runner(GAP_PLANNER_RUNNER, manifest_path, GAP_PLAN_SCHEMA_VERSION)


def run_table_filler_skill(
    manifest_path: Path,
    progress_callback: Callable[[str, dict[str, Any] | None], None] | None = None,
) -> dict[str, Any]:
    prompt = _build_table_filler_prompt(manifest_path)
    try:
        result = OpencodeClient().run_bid_tech_table_filler_with_trace(
            prompt,
            stream_callback=(
                (lambda details: progress_callback("table_filler_delta", details))
                if progress_callback
                else None
            ),
            early_tool_command="s4fill",
        )
        return result
    except Exception:
        # Keep the feature usable in offline test/deploy environments. The
        # fallback executes the same skill runner and records local-skill trace.
        return _run_local_skill_runner(TABLE_FILL_RUNNER, manifest_path, TABLE_FILL_SCHEMA_VERSION)


def run_word_placeholder_filler_skill(
    manifest_path: Path,
    progress_callback: Callable[[str, dict[str, Any] | None], None] | None = None,
) -> dict[str, Any]:
    prompt = _build_word_placeholder_filler_prompt(manifest_path)
    try:
        result = OpencodeClient().run_bid_tech_table_filler_with_trace(
            prompt,
            stream_callback=(
                (lambda details: progress_callback("word_filler_delta", details))
                if progress_callback
                else None
            ),
            early_tool_command="s4wordfill",
        )
        return result
    except Exception:
        return _run_local_skill_runner(WORD_FILL_RUNNER, manifest_path, WORD_FILL_SCHEMA_VERSION)


def _numeric_report_value(report: dict[str, Any], *keys: str) -> int:
    for key in keys:
        value = report.get(key)
        if isinstance(value, (int, float)):
            return max(0, int(value))
        if isinstance(value, str) and value.strip().isdigit():
            return max(0, int(value.strip()))
    return 0


def _merge_fill_sidecar_report(result: dict[str, Any], output_file: Path) -> dict[str, Any]:
    sidecar = output_file.with_suffix(".fill_report.json")
    if not sidecar.exists():
        return result
    try:
        data = json.loads(sidecar.read_text(encoding="utf-8"))
    except Exception:
        return result
    if not isinstance(data, dict):
        return result
    merged = dict(result)
    for key in ("evidenceRefs", "filledFieldDetails", "unfilledFieldDetails", "unfilledFields"):
        sidecar_value = data.get(key)
        current_value = merged.get(key)
        if isinstance(sidecar_value, list) and (
            not isinstance(current_value, list) or len(sidecar_value) > len(current_value)
        ):
            merged[key] = sidecar_value
    sidecar_report = data.get("fillReport")
    if isinstance(sidecar_report, dict):
        current_report = merged.get("fillReport") if isinstance(merged.get("fillReport"), dict) else {}
        merged["fillReport"] = {**current_report, **sidecar_report}
    return merged


def _build_fill_quality_report(result: dict[str, Any], *, output_exists: bool) -> dict[str, Any]:
    report = result.get("fillReport") if isinstance(result.get("fillReport"), dict) else {}
    unfilled_fields = result.get("unfilledFields") if isinstance(result.get("unfilledFields"), list) else []
    evidence_refs = result.get("evidenceRefs") if isinstance(result.get("evidenceRefs"), list) else []
    filled_count = _numeric_report_value(
        report,
        "filledFieldCount",
        "filledPlaceholderCount",
        "filledCellCount",
    )
    unfilled_count = max(
        _numeric_report_value(
            report,
            "unfilledFieldCount",
            "unfilledPlaceholderCount",
            "unfilledCellCount",
            "residualPlaceholderCount",
        ),
        len(unfilled_fields),
    )
    expected_count = _numeric_report_value(
        report,
        "targetFieldCount",
        "targetPlaceholderCount",
        "fieldCount",
        "placeholderCount",
        "totalFieldCount",
    )
    if expected_count <= 0:
        expected_count = filled_count + unfilled_count
    coverage_rate = filled_count / expected_count if expected_count else 0.0
    evidence_chain_rate = min(1.0, len(evidence_refs) / filled_count) if filled_count else 0.0
    semantic_check_count = _numeric_report_value(report, "semanticCheckCount")
    semantic_failed_count = _numeric_report_value(report, "semanticFailedCount")
    semantic_rate_value = report.get("semanticValidationRate")
    try:
        semantic_validation_rate = float(semantic_rate_value) if semantic_rate_value is not None else None
    except (TypeError, ValueError):
        semantic_validation_rate = None
    correctness_rate = (
        min(evidence_chain_rate, semantic_validation_rate)
        if semantic_check_count > 0 and semantic_validation_rate is not None
        else evidence_chain_rate
    )
    completeness_rate = 1.0 if output_exists and unfilled_count == 0 and coverage_rate >= 0.85 else min(coverage_rate, 1.0 if output_exists else 0.0)
    thresholds = {
        "coverageRate": 0.85,
        "correctnessRate": 0.85,
        "completenessRate": 0.85,
    }
    status = "passed" if (
        coverage_rate >= thresholds["coverageRate"]
        and correctness_rate >= thresholds["correctnessRate"]
        and completeness_rate >= thresholds["completenessRate"]
        and unfilled_count == 0
        and semantic_failed_count == 0
        and output_exists
    ) else "needs_review"
    return {
        "schemaVersion": "bid-fill-quality-report-v1",
        "status": status,
        "coverageRate": round(coverage_rate, 4),
        "correctnessRate": round(correctness_rate, 4),
        "completenessRate": round(completeness_rate, 4),
        "evidenceChainRate": round(evidence_chain_rate, 4),
        "expectedFieldCount": expected_count,
        "filledFieldCount": filled_count,
        "unfilledFieldCount": unfilled_count,
        "evidenceRefCount": len(evidence_refs),
        "residualPlaceholderCount": unfilled_count,
        "semanticCheckCount": semantic_check_count,
        "semanticFailedCount": semantic_failed_count,
        "semanticValidationRate": round(semantic_validation_rate, 4) if semantic_validation_rate is not None else None,
        "thresholds": thresholds,
    }


def run_ai_fill_for_gap(
    project: dict[str, Any],
    gap_id: str,
    data: dict[str, Any],
    *,
    browser_base_url: str = "",
    onlyoffice_base_url: str = "",
) -> dict[str, Any]:
    gap_state = project.get("gap_state") or {}
    plan = gap_state.get("plan") if isinstance(gap_state.get("plan"), dict) else {}
    project_fact_table = gap_state.get("projectFactTable") if isinstance(gap_state.get("projectFactTable"), dict) else {}
    normalize_gap_plan_fill_task_skills(plan)
    items = plan.get("items") if isinstance(plan.get("items"), list) else []
    item = next((entry for entry in items if str(entry.get("id") or "") == gap_id), None)
    if item is None:
        raise KeyError(gap_id)

    fill_tasks = item.get("fillTasks") if isinstance(item.get("fillTasks"), list) else []
    requested_task_id = str(data.get("fillTaskId") or "")
    task = next(
        (
            entry
            for entry in fill_tasks
            if not requested_task_id or str(entry.get("id") or "") == requested_task_id
        ),
        None,
    )
    if task is None:
        raise ValueError("当前缺口没有可执行的 AI 填写任务。")

    skill_name = str(task.get("skill") or TABLE_FILL_SKILL_NAME)
    appendix_task = _appendix_task_for_fill(item, task)
    blank_source = dict(task.get("blankSource") or {}) if isinstance(task.get("blankSource"), dict) else {}
    selected_reference_ids = _selected_reference_material_ids(item, appendix_task, data)
    reference_materials = _reference_materials_for_fill(item, appendix_task, data, selected_reference_ids)
    blank_source_id = str(blank_source.get("id") or blank_source.get("materialId") or "").strip()
    reference_materials = [
        material
        for material in reference_materials
        if str(material.get("id") or material.get("materialId") or "").strip() != blank_source_id
    ]
    recommended_materials = _dedupe_material_summaries(_object_items(appendix_task.get("recommendedMaterials")))
    material_index = _material_index_for_fill(project, plan, item)
    parse_fields = _parse_fields_for_fill(appendix_task, task, data)
    work_dir = _project_dir(project) / "s4_gap_workdir" / "ai_fill" / gap_id
    work_dir.mkdir(parents=True, exist_ok=True)
    material_index = _prepare_material_index_files(
        material_index,
        work_dir,
        cache_dir=_project_dir(project) / "s4_gap_workdir" / "ai_fill" / "_material_index_cache",
    )
    artifact_task_id = safe_filename(str(task.get("id") or "task"), "task")
    artifact_id = f"ART-{gap_id}-{artifact_task_id}-{datetime.now(UTC).strftime('%Y%m%d%H%M%S%f')}"
    output_stem = safe_filename(str(item.get("title") or gap_id), gap_id)
    if len(fill_tasks) > 1:
        output_suffix = safe_filename(str(task.get("id") or blank_source_id or "task"), "task")
        output_file = work_dir / f"{output_stem}_{output_suffix}_AI填写.docx"
    else:
        output_file = work_dir / f"{output_stem}_AI填写.docx"
    manifest_path = work_dir / ("word_fill_input.json" if skill_name == WORD_FILL_SKILL_NAME else "table_fill_input.json")
    if skill_name == WORD_FILL_SKILL_NAME:
        blank_source = _prepare_word_blank_source(blank_source, work_dir)
    manifest = {
        "schemaVersion": WORD_FILL_SCHEMA_VERSION if skill_name == WORD_FILL_SKILL_NAME else TABLE_FILL_SCHEMA_VERSION,
        "projectId": str(project.get("id") or ""),
        "projectName": str(project.get("name") or ""),
        "projectIdentity": project.get("identity") or {},
        "customerName": str(project.get("customerName") or ""),
        "projectTurbineModel": project_turbine_model(project),
        "gapId": gap_id,
        "fillTaskId": str(task.get("id") or ""),
        "title": str(item.get("title") or ""),
        "gapItem": {
            "id": gap_id,
            "number": str(item.get("number") or item.get("section") or ""),
            "title": str(item.get("title") or ""),
            "decision": str(item.get("decision") or ""),
            "usage": str(item.get("usage") or ""),
            "gapReason": str(item.get("gapReason") or item.get("reason") or ""),
            "materialScope": item.get("materialScope") or {},
            "turbineCheck": item.get("turbineCheck") or {},
        },
        "appendixTask": {
            "id": str(appendix_task.get("id") or ""),
            "title": str(appendix_task.get("title") or ""),
            "sourceFile": str(appendix_task.get("sourceFile") or ""),
            "docxPath": str(appendix_task.get("docxPath") or ""),
            "workspacePath": str(appendix_task.get("workspacePath") or ""),
            "rowCount": appendix_task.get("rowCount") or 0,
            "availableParseFields": parse_fields,
        },
        "blankSource": blank_source,
        "referenceMaterialIds": selected_reference_ids,
        "referenceMaterials": reference_materials,
        "projectFactTable": project_fact_table,
        "materialIndex": material_index,
        "recommendedMaterials": recommended_materials,
        "parseFieldIds": _string_items(data.get("parseFieldIds")),
        "parseFields": parse_fields,
        "constraints": str(data.get("constraints") or ""),
        "operator": str(data.get("operator") or "当前用户"),
        "outputFile": str(output_file),
    }
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    if skill_name == WORD_FILL_SKILL_NAME:
        result = run_word_placeholder_filler_skill(manifest_path)
    else:
        result = run_table_filler_skill(manifest_path)
    resolved_output = Path(str(result.get("outputFile") or output_file))
    if not resolved_output.exists():
        raise RuntimeError(f"AI 填写未生成输出文件：{resolved_output}")
    result = _merge_fill_sidecar_report(result, resolved_output)
    quality_report = _build_fill_quality_report(result, output_exists=True)

    artifact = {
        "id": artifact_id,
        "source": "ai_fill",
        "skill": skill_name,
        "gapId": gap_id,
        "fillTaskId": str(task.get("id") or ""),
        "title": str(item.get("title") or resolved_output.stem),
        "fileName": resolved_output.name,
        "path": str(resolved_output),
        "createdAt": now_iso(),
        "operator": str(data.get("operator") or "当前用户"),
        "unfilledFields": list(result.get("unfilledFields") or []),
        "evidenceRefs": list(result.get("evidenceRefs") or []),
        "fillReport": result.get("fillReport") or {},
        "qualityReport": quality_report,
        "referenceMaterials": reference_materials,
        "recommendedMaterials": recommended_materials,
        "parseFields": parse_fields,
        "manifestPath": str(manifest_path),
        "opencodeOutput": result.get("opencodeOutput") or {},
        "onlyoffice": _artifact_onlyoffice_payload(
            project_id=str(project.get("id") or ""),
            artifact_id=artifact_id,
            file_name=resolved_output.name,
            browser_base_url=browser_base_url,
            onlyoffice_base_url=onlyoffice_base_url,
        ),
        "s7Ready": True,
    }

    task["status"] = "completed"
    task["outputArtifactId"] = artifact_id
    task["completedAt"] = artifact["createdAt"]
    item["status"] = "resolved"
    item["qualityStatus"] = quality_report["status"]
    item["qualityReport"] = quality_report
    item["resolvedArtifacts"] = _replace_resolved_artifact(
        item.get("resolvedArtifacts"),
        artifact,
        fill_task_id=str(task.get("id") or ""),
        skill_name=skill_name,
        file_name=resolved_output.name,
    )
    item["resolvedAt"] = artifact["createdAt"]
    item["resolvedSource"] = artifact["fileName"]
    item["reviewNotes"] = list(item.get("reviewNotes") or [])
    if artifact["unfilledFields"]:
        item["reviewNotes"].append(f"AI 填写仍有未填字段：{len(artifact['unfilledFields'])} 项")
    if quality_report["status"] != "passed":
        item["reviewNotes"].append("AI 填写质量验收未达标，请人工复核或补充事实表后重填。")

    plan["updatedAt"] = artifact["createdAt"]
    plan["summary"] = summarize_gap_plan(plan)
    gap_state["plan"] = plan
    gap_state["items"] = _legacy_items_from_plan(plan)
    gap_state["submittedForReview"] = False
    gap_state["reviewConfirmed"] = False
    gap_state["reviewedAt"] = ""
    return {"item": item, "artifact": artifact, "gapPlan": plan}


def _compact_resolved_artifact(artifact: dict[str, Any]) -> dict[str, Any]:
    return {
        "id": artifact.get("id"),
        "source": artifact.get("source"),
        "skill": artifact.get("skill"),
        "gapId": artifact.get("gapId"),
        "fillTaskId": artifact.get("fillTaskId"),
        "title": artifact.get("title"),
        "fileName": artifact.get("fileName"),
        "path": artifact.get("path"),
        "createdAt": artifact.get("createdAt"),
        "operator": artifact.get("operator"),
        "unfilledFields": list(artifact.get("unfilledFields") or []),
        "evidenceRefs": list(artifact.get("evidenceRefs") or []),
        "fillReport": artifact.get("fillReport") or {},
        "qualityReport": artifact.get("qualityReport") or {},
        "manifestPath": artifact.get("manifestPath"),
        "onlyoffice": artifact.get("onlyoffice") or {},
        "s7Ready": artifact.get("s7Ready", True),
        "referenceMaterialCount": len(artifact.get("referenceMaterials") or []),
        "recommendedMaterialCount": len(artifact.get("recommendedMaterials") or []),
    }


def _replace_resolved_artifact(
    current: Any,
    artifact: dict[str, Any],
    *,
    fill_task_id: str,
    skill_name: str,
    file_name: str,
) -> list[dict[str, Any]]:
    def compact_key(value: Any) -> str:
        return str(value or "").strip()

    compact = _compact_resolved_artifact(artifact)
    kept: list[dict[str, Any]] = []
    for item in _object_items(current):
        if fill_task_id and compact_key(item.get("fillTaskId")) == fill_task_id:
            continue
        if not item.get("fillTaskId") and compact_key(item.get("skill")) == skill_name and compact_key(item.get("fileName")) == file_name:
            continue
        kept.append(_compact_resolved_artifact(item))
    kept.append(compact)
    return kept[-12:]

def register_manual_gap_upload(
    project: dict[str, Any],
    gap_id: str,
    data: dict[str, Any],
    *,
    browser_base_url: str = "",
    onlyoffice_base_url: str = "",
) -> dict[str, Any]:
    gap_state = project.get("gap_state") or {}
    plan = gap_state.get("plan") if isinstance(gap_state.get("plan"), dict) else {}
    items = plan.get("items") if isinstance(plan.get("items"), list) else []
    item = next((entry for entry in items if str(entry.get("id") or "") == gap_id), None)
    if item is None:
        raise KeyError(gap_id)
    files = [entry for entry in data.get("files") or [] if isinstance(entry, dict)]
    if not files:
        raise ValueError("至少需要提交一个文件。")

    work_dir = _project_dir(project) / "s4_gap_workdir" / "manual_upload" / gap_id
    work_dir.mkdir(parents=True, exist_ok=True)
    created_at = now_iso()
    artifacts: list[dict[str, Any]] = []
    for index, file in enumerate(files, start=1):
        name = safe_filename(str(file.get("name") or f"{gap_id}-{index}.docx"), f"{gap_id}-{index}.docx")
        if not name.lower().endswith(".docx"):
            name = f"{Path(name).stem}.docx"
        output_file = work_dir / name
        content = str(file.get("data") or file.get("text") or file.get("content") or "")
        _write_manual_upload_docx(output_file, title=str(item.get("title") or name), content=content)
        artifact_id = f"ART-{gap_id}-UPLOAD-{index}"
        artifacts.append(
            {
                "id": artifact_id,
                "source": "manual_upload",
                "skill": "",
                "title": str(item.get("title") or output_file.stem),
                "fileName": output_file.name,
                "path": str(output_file),
                "createdAt": created_at,
                "operator": str(data.get("operator") or "当前用户"),
                "unfilledFields": [],
                "evidenceRefs": [],
                "onlyoffice": _artifact_onlyoffice_payload(
                    project_id=str(project.get("id") or ""),
                    artifact_id=artifact_id,
                    file_name=output_file.name,
                    browser_base_url=browser_base_url,
                    onlyoffice_base_url=onlyoffice_base_url,
                ),
                "s7Ready": True,
            }
        )

    item["status"] = "resolved"
    item.setdefault("resolvedArtifacts", []).extend(artifacts)
    item["resolvedAt"] = created_at
    item["resolvedSource"] = artifacts[0]["fileName"]
    item["latestUploadAt"] = created_at
    item["latestSubmissionId"] = artifacts[0]["id"]
    plan["updatedAt"] = created_at
    plan["summary"] = summarize_gap_plan(plan)
    gap_state["plan"] = plan
    gap_state["items"] = _legacy_items_from_plan(plan)
    gap_state["submittedForReview"] = False
    gap_state["reviewConfirmed"] = False
    gap_state["reviewedAt"] = ""
    return {"item": item, "artifact": artifacts[0], "artifacts": artifacts, "gapPlan": plan}


async def prepare_existing_gap_material_files(
    project: dict[str, Any],
    gap_id: str,
    data: dict[str, Any],
) -> list[dict[str, Any]]:
    selected = _selected_material_entries(data)
    if not selected:
        raise ValueError("至少需要选择一份素材。")

    work_dir = _project_dir(project) / "s4_gap_workdir" / "selected_material" / gap_id
    work_dir.mkdir(parents=True, exist_ok=True)
    prepared: list[dict[str, Any]] = []
    for index, material in enumerate(selected, start=1):
        material_id = str(material.get("id") or material.get("materialId") or "").strip()
        if not material_id:
            continue
        payload, source_kind = await _downloadable_material_payload(material_id)
        file_name = safe_filename(
            str(material.get("cleanedFileName") or payload.get("fileName") or material.get("name") or f"{material_id}.docx"),
            f"{material_id}.docx",
        )
        if not file_name.lower().endswith(".docx"):
            file_name = f"{Path(file_name).stem}.docx"
        target_path = work_dir / f"{index:02d}-{file_name}"
        minio_client.download_file(str(payload["bucket"]), str(payload["key"]), target_path)
        prepared.append(
            {
                "materialId": material_id,
                "materialName": str(material.get("name") or payload.get("fileName") or material_id),
                "fileName": target_path.name,
                "path": str(target_path),
                "folderPath": str(material.get("folderPath") or ""),
                "materialTier": str(material.get("materialTier") or ""),
                "sourceKind": source_kind,
                "originalMaterial": material,
            }
        )
    if not prepared:
        raise ValueError("至少需要选择一份有效素材。")
    return prepared


def register_existing_gap_material(
    project: dict[str, Any],
    gap_id: str,
    data: dict[str, Any],
    prepared_files: list[dict[str, Any]],
    *,
    browser_base_url: str = "",
    onlyoffice_base_url: str = "",
) -> dict[str, Any]:
    gap_state = project.get("gap_state") or {}
    plan = gap_state.get("plan") if isinstance(gap_state.get("plan"), dict) else {}
    items = plan.get("items") if isinstance(plan.get("items"), list) else []
    item = next((entry for entry in items if str(entry.get("id") or "") == gap_id), None)
    if item is None:
        raise KeyError(gap_id)
    if not prepared_files:
        raise ValueError("至少需要选择一份素材。")

    created_at = now_iso()
    existing_count = len(item.get("resolvedArtifacts") or [])
    artifacts: list[dict[str, Any]] = []
    for index, prepared in enumerate(prepared_files, start=1):
        path = Path(str(prepared.get("path") or ""))
        if not path.exists():
            raise ValueError(f"已选择素材文件不存在：{path}")
        artifact_id = f"ART-{gap_id}-MAT-{existing_count + index}"
        artifacts.append(
            {
                "id": artifact_id,
                "source": "material_library",
                "skill": "",
                "title": str(prepared.get("materialName") or path.stem),
                "fileName": str(prepared.get("fileName") or path.name),
                "path": str(path),
                "materialId": str(prepared.get("materialId") or ""),
                "folderPath": str(prepared.get("folderPath") or ""),
                "materialTier": str(prepared.get("materialTier") or ""),
                "sourceKind": str(prepared.get("sourceKind") or ""),
                "createdAt": created_at,
                "operator": str(data.get("operator") or "当前用户"),
                "unfilledFields": [],
                "evidenceRefs": [{"type": "material", "id": str(prepared.get("materialId") or "")}],
                "onlyoffice": _artifact_onlyoffice_payload(
                    project_id=str(project.get("id") or ""),
                    artifact_id=artifact_id,
                    file_name=str(prepared.get("fileName") or path.name),
                    browser_base_url=browser_base_url,
                    onlyoffice_base_url=onlyoffice_base_url,
                ),
                "s7Ready": True,
            }
        )

    item["status"] = "resolved"
    item.setdefault("resolvedArtifacts", []).extend(artifacts)
    item["resolvedAt"] = created_at
    item["resolvedSource"] = artifacts[0]["fileName"]
    item.setdefault("reviewNotes", []).append(f"人工选择已有素材：{len(artifacts)} 份")
    plan["updatedAt"] = created_at
    plan["summary"] = summarize_gap_plan(plan)
    gap_state["plan"] = plan
    gap_state["items"] = _legacy_items_from_plan(plan)
    gap_state["submittedForReview"] = False
    gap_state["reviewConfirmed"] = False
    gap_state["reviewedAt"] = ""
    return {"item": item, "artifact": artifacts[0], "artifacts": artifacts, "gapPlan": plan}


def _write_manual_upload_docx(path: Path, *, title: str, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    uploaded_bytes = _decode_uploaded_docx(content)
    if uploaded_bytes is not None:
        path.write_bytes(uploaded_bytes)
        return
    doc = Document()
    doc.add_heading(title or path.stem, level=1)
    text = content.strip() or "人工上传客户资料，原始文件内容请以项目素材库归档为准。"
    for paragraph in text.splitlines() or [text]:
        doc.add_paragraph(paragraph)
    doc.save(path)


def _decode_uploaded_docx(content: str) -> bytes | None:
    text = str(content or "").strip()
    if not text.startswith("data:"):
        return None
    header, separator, payload = text.partition(",")
    if not separator or ";base64" not in header.lower():
        return None
    try:
        decoded = base64.b64decode(payload, validate=True)
    except (binascii.Error, ValueError):
        return None
    if decoded.startswith(b"PK\x03\x04"):
        return decoded
    return None


def _selected_material_entries(data: dict[str, Any]) -> list[dict[str, Any]]:
    raw_entries = data.get("materials")
    if not raw_entries:
        raw_entries = data.get("materialIds") or data.get("referenceMaterialIds") or []
    entries: list[dict[str, Any]] = []
    for entry in raw_entries if isinstance(raw_entries, list) else []:
        if isinstance(entry, str):
            entries.append({"id": entry})
        elif isinstance(entry, dict):
            entries.append(entry)
    return entries


async def _downloadable_material_payload(material_id: str) -> tuple[dict[str, Any], str]:
    try:
        payload = await material_store.raw_download_cleaned_content(material_id)
        return payload, "cleaned"
    except Exception:
        payload = await material_store.raw_download_content(material_id)
    mime_type = str(payload.get("mimeType") or "")
    file_name = str(payload.get("fileName") or "")
    if "wordprocessingml" not in mime_type and not file_name.lower().endswith(".docx"):
        raise ValueError(f"素材 {material_id} 没有可用于拼接的 Word 文件或清洗稿。")
    return payload, "raw"


async def _downloadable_fill_source_payload(material_id: str) -> tuple[dict[str, Any], str]:
    try:
        payload = await material_store.raw_download_cleaned_content(material_id)
        return payload, "cleaned"
    except Exception:
        payload = await material_store.raw_download_content(material_id)
    mime_type = str(payload.get("mimeType") or "")
    file_name = str(payload.get("fileName") or "").lower()
    allowed_ext = file_name.endswith((".docx", ".xlsx", ".xlsm"))
    allowed_mime = (
        "wordprocessingml" in mime_type
        or "spreadsheetml" in mime_type
        or "ms-excel" in mime_type
    )
    if not allowed_ext and not allowed_mime:
        raise ValueError(f"素材 {material_id} 不是填写 Skill 可读取的 Word/Excel。")
    return payload, "raw"


def _prepare_word_blank_source(blank_source: dict[str, Any], work_dir: Path) -> dict[str, Any]:
    source = dict(blank_source)
    for key in ("docxPath", "path", "workspacePath"):
        candidate = Path(str(source.get(key) or ""))
        if candidate.exists() and candidate.suffix.lower() == ".docx":
            source["docxPath"] = str(candidate)
            return source

    material_id = str(source.get("materialId") or source.get("id") or "").strip()
    if not material_id:
        raise ValueError("待填写 Word 缺少 materialId/docxPath，无法准备模板。")
    payload, source_kind = _run_async(_downloadable_material_payload(material_id))
    file_name = safe_filename(
        str(source.get("cleanedFileName") or payload.get("fileName") or source.get("title") or f"{material_id}.docx"),
        f"{material_id}.docx",
    )
    if not file_name.lower().endswith(".docx"):
        file_name = f"{Path(file_name).stem}.docx"
    target_path = work_dir / "blank_source" / f"{material_id}-{file_name}"
    if not target_path.exists():
        minio_client.download_file(str(payload["bucket"]), str(payload["key"]), target_path)
    source.update(
        {
            "docxPath": str(target_path),
            "path": str(target_path),
            "sourceKind": source_kind,
            "fileName": target_path.name,
        }
    )
    return source


def summarize_gap_plan(plan: dict[str, Any]) -> dict[str, Any]:
    items = [item for item in plan.get("items") or [] if isinstance(item, dict)]
    decision_counts: dict[str, int] = {}
    for item in items:
        decision = str(item.get("decision") or "")
        if decision:
            decision_counts[decision] = decision_counts.get(decision, 0) + 1
    return {
        "totalTocItems": len(items),
        "matchedCount": sum(1 for item in items if item.get("status") == "matched"),
        "missingCount": sum(1 for item in items if item.get("status") in {"missing", "needs_input"}),
        "resolvedCount": sum(1 for item in items if item.get("status") == "resolved"),
        "ignoredCount": sum(1 for item in items if item.get("status") == "ignored"),
        "structuralCount": sum(1 for item in items if item.get("status") == "structural"),
        "fillableTaskCount": sum(len(item.get("fillTasks") or []) for item in items),
        "blockingCount": sum(1 for item in items if item.get("status") in {"missing", "needs_input", "filling"}),
        "readyCount": decision_counts.get("ready", 0),
        "fillRequiredCount": decision_counts.get("fill_required", 0),
        "materialRequiredCount": decision_counts.get("material_required", 0),
        "reviewRequiredCount": decision_counts.get("review_required", 0),
        "appendixTaskCount": sum(len(item.get("appendixTasks") or []) for item in items),
    }


def check_gap_integrity(plan: dict[str, Any]) -> dict[str, Any]:
    summary = summarize_gap_plan(plan)
    blocking_items = [
        {
            "id": str(item.get("id") or ""),
            "number": str(item.get("number") or ""),
            "title": str(item.get("title") or ""),
            "status": str(item.get("status") or ""),
        }
        for item in plan.get("items") or []
        if str(item.get("status") or "") in {"missing", "needs_input", "filling"}
    ]
    return {
        "status": "passed" if not blocking_items else "blocked",
        "checkedAt": now_iso(),
        "blockingCount": len(blocking_items),
        "blockingItems": blocking_items,
        "summary": summary,
    }


def _project_dir(project: dict[str, Any]) -> Path:
    project_id = str(project.get("id") or "")
    project_dir = technical_workspace_dir(project_id)
    project_dir.mkdir(parents=True, exist_ok=True)
    return project_dir


def _resolve_toc_json(project: dict[str, Any], work_dir: Path) -> Path:
    project_id = str(project.get("id") or "")
    parse_storage = project.get("parse_storage") or {}
    candidates = []
    directory_output = ((project.get("directory_state") or {}).get("opencodeOutput") or {})
    for value in (directory_output.get("tocJsonPath"), directory_output.get("outputFile")):
        if value:
            candidates.append(Path(str(value)))
    for root in legacy_workspace_roots(project_id, parse_storage):
        s2_work_dir = root / "s2_toc_workdir"
        candidates.extend(path for path in sorted(s2_work_dir.glob("*.json")) if "evidence" not in path.name.lower())
    for candidate in candidates:
        if candidate.exists() and candidate.suffix.lower() == ".json":
            target = work_dir / settings.s2_toc_output_file_name
            shutil.copy2(candidate, target)
            return target

    outline_nodes = list((project.get("outline_state") or {}).get("nodes") or [])
    output = {
        "schema_version": "bid-toc-json-v1",
        "document_title": f"{project.get('name') or project_id}投标文件总目录",
        "project": {
            "owner": project.get("customerName") or "",
            "name": project.get("name") or project_id,
            "code": project.get("projectCode") or project_id,
        },
        "items": _outline_nodes_to_toc_items(outline_nodes),
    }
    target = work_dir / settings.s2_toc_output_file_name
    target.write_text(json.dumps(output, ensure_ascii=False, indent=2), encoding="utf-8")
    return target


def _resolve_wiki_dir(project: dict[str, Any], project_dir: Path, work_dir: Path) -> Path | None:
    project_id = str(project.get("id") or "")
    parse_storage = project.get("parse_storage") or {}
    candidates = [root / "s2_toc_workdir" / "wiki" for root in legacy_workspace_roots(project_id, parse_storage)]
    for candidate in candidates:
        if (candidate / "卡片").exists():
            target = work_dir / "wiki"
            if target.exists():
                shutil.rmtree(target)
            shutil.copytree(candidate, target, dirs_exist_ok=True)
            return target
    return None


def _outline_nodes_to_toc_items(nodes: list[dict[str, Any]], prefix: str = "", level: int = 1) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    for index, node in enumerate(nodes, start=1):
        number = f"{prefix}.{index}" if prefix else str(index)
        items.append(
            {
                "order": len(items) + 1,
                "number": number,
                "title": str(node.get("title") or "").strip(),
                "level": level,
                "annotation": str(node.get("annotation") or "保留"),
                "source": "outline_state",
                "reason": "",
                "material_refs": list(node.get("material_refs") or []),
            }
        )
        children = node.get("children") or []
        if isinstance(children, list):
            child_items = _outline_nodes_to_toc_items(children, number, level + 1)
            for child in child_items:
                child["order"] = len(items) + 1
                items.append(child)
    return items


def _run_local_skill_runner(runner: Path, manifest_path: Path, schema_version: str) -> dict[str, Any]:
    if not runner.exists():
        raise RuntimeError(f"Skill runner 不存在：{runner}")
    result = subprocess.run(
        [sys.executable, str(runner), "--manifest", str(manifest_path), "--response", "summary"],
        capture_output=True,
        text=True,
        encoding="utf-8",
    )
    if result.returncode != 0:
        detail = "\n".join(part for part in ((result.stdout or "").strip(), (result.stderr or "").strip()) if part)
        raise RuntimeError(f"Skill runner 执行失败（{result.returncode}）：{detail}")
    payload = json.loads(result.stdout or "{}")
    payload.setdefault("schema_version", schema_version)
    payload.setdefault(
        "opencodeOutput",
        {
            "status": "received",
            "sessionId": str(manifest_path),
            "providerId": "local-skill",
            "modelId": runner.parent.parent.name,
            "receivedAt": now_iso(),
            "parts": [{"type": "text", "text": result.stdout.strip()}],
        },
    )
    return payload


def _build_table_filler_prompt(manifest_path: Path) -> str:
    return f"""
Use the {TABLE_FILL_SKILL_NAME} skill.

你现在在做技术标缺口项 AI 填写。后端已经准备好 manifest，其中包含待填写空表/Word、人工指定的参考素材、解析字段和输出路径。

manifest：{manifest_path}

请直接调用一次 Bash 工具执行下面命令，Bash 工具 timeout 必须设置为 1800000 毫秒或更高。不要先检查工作目录，不要先执行 pwd/ls/cat/read/glob，不要拆成多条命令，不要改写命令或路径：

s4fill {manifest_path}

只返回命令 stdout 中的小型 JSON，不要返回解释文字，不要使用 Markdown 代码块。
""".strip()


def _build_word_placeholder_filler_prompt(manifest_path: Path) -> str:
    return f"""
Use the {WORD_FILL_SKILL_NAME} skill.

你现在在做技术标素材库待填写 Word 的 AI 填写。后端已经准备好 manifest，其中包含待填写 Word、人工指定参考素材、解析字段、项目投标机型和输出路径。

manifest：{manifest_path}

请直接调用一次 Bash 工具执行下面命令，Bash 工具 timeout 必须设置为 1800000 毫秒或更高。不要先检查工作目录，不要先执行 pwd/ls/cat/read/glob，不要拆成多条命令，不要改写命令或路径：

s4wordfill {manifest_path}

只返回命令 stdout 中的小型 JSON，不要返回解释文字，不要使用 Markdown 代码块。
""".strip()


def _build_gap_planner_prompt(manifest_path: Path) -> str:
    return f"""
Use the {GAP_PLANNER_SKILL_NAME} skill.

你现在在做 S3 技术标缺口识别。后端已经准备好 manifest，其中包含人工确认后的目录 JSON、招标解析结构化结果、S2 素材 Wiki 副本、项目/客户/通用素材边界、素材索引、项目身份信息和人工确认的投标机型信息。

manifest：{manifest_path}

请直接调用一次 Bash 工具执行下面命令，Bash 工具 timeout 必须设置为 1800000 毫秒或更高。不要先检查工作目录，不要先执行 pwd/ls/cat/read/glob，不要拆成多条命令，不要改写命令或路径。命令会把完整 gap_plan.json 写入 manifest 指定路径，并只在 stdout 打印小型摘要 JSON：

s4gap {manifest_path}

只返回命令 stdout 中的小型 JSON，不要返回解释文字，不要使用 Markdown 代码块。
返回格式必须是：
{{
  "schema_version": "{GAP_PLAN_SCHEMA_VERSION}",
  "outputFile": "/data/documents/PRJ-0001/technical-workspace/s4_gap_workdir/gap_plan.json",
  "summary": {{"totalTocItems": 0, "matchedCount": 0, "missingCount": 0, "resolvedCount": 0, "ignoredCount": 0, "structuralCount": 0, "fillableTaskCount": 0, "blockingCount": 0}},
  "itemCount": 0
}}
""".strip()


def _artifact_onlyoffice_payload(
    *,
    project_id: str,
    artifact_id: str,
    file_name: str,
    browser_base_url: str = "",
    onlyoffice_base_url: str = "",
) -> dict[str, Any]:
    file_url = f"/api/projects/{project_id}/gaps/artifacts/{artifact_id}/content/{quote(file_name)}"
    browser_url = f"{browser_base_url.rstrip('/')}{file_url}" if browser_base_url else file_url
    document_server_url = (
        f"{onlyoffice_base_url.rstrip('/')}{file_url}"
        if onlyoffice_base_url
        else browser_url
    )
    return {
        "status": "ready",
        "mode": "view",
        "fileUrl": document_server_url,
        "browserFileUrl": browser_url,
        "documentServerFileUrl": document_server_url,
        "documentKey": f"{project_id}-{artifact_id}",
        "title": file_name,
    }


def _legacy_items_from_plan(plan: dict[str, Any]) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    for index, item in enumerate(plan.get("items") or [], start=1):
        status = str(item.get("status") or "")
        if status in {"matched", "structural"}:
            continue
        items.append(
            {
                "id": str(item.get("id") or f"GAP-{index}"),
                "section": str(item.get("section") or item.get("parentTitle") or ""),
                "title": str(item.get("title") or ""),
                "desc": str(item.get("gapReason") or item.get("reason") or "请补充该目录项所需素材。"),
                "priority": str(item.get("priority") or "medium"),
                "bidType": "技术标",
                "status": "resolved" if status == "resolved" else "skipped" if status == "ignored" else "pending",
                "skipReason": str(item.get("skipReason") or ""),
                "resolvedSource": str(item.get("resolvedSource") or ""),
                "resolvedAt": str(item.get("resolvedAt") or ""),
                "latestUploadAt": str(item.get("latestUploadAt") or ""),
                "latestSubmissionId": str(item.get("latestSubmissionId") or ""),
            }
        )
    return items
