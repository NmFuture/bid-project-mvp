from __future__ import annotations

import hashlib
import os
import re
import time
from pathlib import Path

import httpx
from docx import Document

from app.core.config import settings
from app.services.minio_client import minio_client

WORD_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def document_path(project_id: str) -> Path:
    return settings.documents_dir / f"{project_id}.docx"


def review_document_path(project_id: str) -> Path:
    return settings.documents_dir / f"{project_id}-review.docx"


def document_object_key(project_id: str) -> str:
    return f"documents/{project_id}/draft.docx"


def review_document_object_key(project_id: str) -> str:
    return f"documents/{project_id}/review.docx"


def build_document_key(path: Path) -> str:
    stat = path.stat()
    raw_key = f"{path.resolve()}-{stat.st_mtime_ns}-{stat.st_size}"
    digest = hashlib.sha256(raw_key.encode("utf-8")).hexdigest()[:24]
    safe_stem = re.sub(r"[^A-Za-z0-9._=-]+", "-", path.stem).strip(".-_= ")
    safe_stem = safe_stem[:56].strip(".-_= ")
    prefix = f"{safe_stem}-" if safe_stem else ""
    return f"{prefix}{digest}-{stat.st_size}"


def build_editor_session_key(path: Path, version: int | str | None = None) -> str:
    base_key = build_document_key(path)
    if version in (None, ""):
        return base_key
    return f"{base_key}-v{version}"


def refresh_document_session(path: Path) -> None:
    current = path.stat().st_mtime_ns
    refreshed = max(time.time_ns(), current + 1)
    os.utime(path, ns=(refreshed, refreshed))


def ensure_document(project_id: str, title: str, content: str) -> Path:
    path = document_path(project_id)
    if not path.exists():
        bucket = settings.minio_buckets["documents"]
        key = document_object_key(project_id)
        try:
            if key and minio_client.object_exists(bucket, key):
                minio_client.download_file(bucket, key, path)
            else:
                write_document(path, title, content)
                minio_client.upload_file(bucket, key, path, WORD_MEDIA_TYPE)
        except Exception:
            write_document(path, title, content)
    return path


def ensure_review_document(project_id: str, title: str, content: str) -> Path:
    path = review_document_path(project_id)
    if not path.exists():
        bucket = settings.minio_buckets["documents"]
        key = review_document_object_key(project_id)
        try:
            if key and minio_client.object_exists(bucket, key):
                minio_client.download_file(bucket, key, path)
            else:
                write_document(path, title, content)
                minio_client.upload_file(bucket, key, path, WORD_MEDIA_TYPE)
        except Exception:
            write_document(path, title, content)
    return path


def write_document(path: Path, title: str, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    title_text = Path(title).stem.strip()
    if title_text:
        doc.add_heading(title_text, level=0)

    for raw_line in (content or "").replace("\r\n", "\n").split("\n"):
        line = raw_line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        doc.add_paragraph(line)

    doc.save(path)


def sync_document_to_minio(path: Path, key: str) -> str:
    try:
        return minio_client.upload_file(settings.minio_buckets["documents"], key, path, WORD_MEDIA_TYPE)
    except Exception:
        return key


async def download_document_from_onlyoffice(
    download_url: str,
    target_path: Path,
    *,
    max_bytes: int | None = None,
) -> None:
    """Download the OnlyOffice-saved document with bounded size and no redirects."""

    byte_limit = max_bytes or settings.onlyoffice_download_max_bytes
    target_path.parent.mkdir(parents=True, exist_ok=True)
    temp_path = target_path.with_suffix(f"{target_path.suffix}.download")

    try:
        async with httpx.AsyncClient(timeout=120.0, follow_redirects=False, trust_env=False) as client:
            async with client.stream("GET", download_url) as response:
                if 300 <= response.status_code < 400:
                    raise RuntimeError("OnlyOffice document download redirect is not allowed.")
                response.raise_for_status()

                written = 0
                with temp_path.open("wb") as handle:
                    async for chunk in response.aiter_bytes():
                        if not chunk:
                            continue
                        written += len(chunk)
                        if written > byte_limit:
                            raise RuntimeError("OnlyOffice document download exceeds the configured size limit.")
                        handle.write(chunk)

        temp_path.replace(target_path)
    except Exception:
        temp_path.unlink(missing_ok=True)
        raise
