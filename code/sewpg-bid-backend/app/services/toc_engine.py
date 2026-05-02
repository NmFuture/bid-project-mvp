from __future__ import annotations

import json
import re
import zipfile
from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from app.core.config import settings

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None  # type: ignore[assignment]


WORD_NAMESPACE = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

CHINESE_NUMERAL_PATTERN = "[一二三四五六七八九十百千万零〇两]+"
HEADING_PATTERNS: tuple[re.Pattern[str], ...] = (
    re.compile(rf"^第(?P<num>{CHINESE_NUMERAL_PATTERN}|\d+)章[\s　:：、.-]*(?P<title>.+)$"),
    re.compile(r"^(?P<num>\d+(?:\.\d+){0,5})[\s　:：、.-]+(?P<title>.+)$"),
    re.compile(rf"^(?P<num>{CHINESE_NUMERAL_PATTERN})[、.．]\s*(?P<title>.+)$"),
    re.compile(r"^[（(](?P<num>\d+)[）)]\s*(?P<title>.+)$"),
    re.compile(r"^(?P<num>(?:技术\s*)?附(?:件|表|录)\s*[A-Za-z0-9一二三四五六七八九十百千万零〇两.-]*)[\s　:：、.-]*(?P<title>.+)$"),
)
REQUIREMENT_PATTERNS: tuple[re.Pattern[str], ...] = (
    re.compile(r"(?:投标人|供应商|响应方|报价人|申请人)?(?:应|须|必须|需|需要|应当|不得|禁止)(?P<title>[^。；;\n]{2,80})"),
    re.compile(r"(?:提供|提交|编制|说明|承诺|响应)(?P<title>[^。；;\n]{2,80})"),
)
APPENDIX_TITLE_PATTERN = re.compile(
    r"^(?P<prefix>(?:技术\s*)?附(?:件|表|录)\s*[A-Za-z0-9一二三四五六七八九十百千万零〇两.-]*)[\s　:：、.-]*(?P<title>[^。；;\n]{2,100})"
)
TABLE_TITLE_PATTERN = re.compile(
    r"^(?P<prefix>表\s*\d+(?:\.\d+)*)[\s　:：、.-]*(?P<title>[^。；;\n]{2,100}(?:表|清单|一览表|响应表|承诺表|偏差表|明细表))$"
)
NOISE_TITLES = {
    "目录",
    "正文",
    "附件",
    "附录",
    "投标文件",
    "招标文件",
    "要求",
    "内容",
    "说明",
    "响应",
}
GENERIC_PREFIX_PATTERN = re.compile(
    r"^(?:投标人|供应商|响应方|报价人|申请人)?(?:应|须|必须|需|需要|应当|不得|禁止|提供|提交|编制|说明|承诺|响应|按要求|按照要求)"
)


@dataclass(frozen=True)
class OutlineEntry:
    number: str
    title: str
    level: int
    source_file: str
    paragraph_index: int
    raw_text: str


@dataclass(frozen=True)
class TenderCandidate:
    title: str
    raw_text: str
    source_file: str
    paragraph_index: int
    kind: str = "requirement"
    number: str = ""
    file_id: str = ""
    file_name: str = ""
    auto_add: bool = False
    context_title: str = ""


@dataclass(frozen=True)
class DocxParagraph:
    text: str
    style_id: str
    style_name: str
    outline_level: int | None
    paragraph_index: int
    tab_count: int


@dataclass(frozen=True)
class TocRules:
    max_level: int
    max_tender_candidates: int
    max_title_chars: int
    auto_append_tender_requirements: bool


def _default_rules() -> TocRules:
    return TocRules(
        max_level=settings.s2_toc_max_level,
        max_tender_candidates=settings.s2_toc_max_tender_candidates,
        max_title_chars=settings.s2_toc_max_title_chars,
        auto_append_tender_requirements=settings.s2_toc_auto_append_tender_requirements,
    )


def _rules_from_manifest(manifest: dict[str, Any]) -> TocRules:
    raw = manifest.get("rules")
    data = raw if isinstance(raw, dict) else {}
    defaults = _default_rules()
    return TocRules(
        max_level=_positive_int(data.get("maxLevel"), defaults.max_level),
        max_tender_candidates=_non_negative_int(
            data.get("maxTenderCandidates"),
            defaults.max_tender_candidates,
        ),
        max_title_chars=_positive_int(data.get("maxTitleChars"), defaults.max_title_chars),
        auto_append_tender_requirements=_bool_value(
            data.get("autoAppendTenderRequirements"),
            defaults.auto_append_tender_requirements,
        ),
    )


def _rules_to_dict(rules: TocRules) -> dict[str, Any]:
    return {
        "maxLevel": rules.max_level,
        "maxTenderCandidates": rules.max_tender_candidates,
        "maxTitleChars": rules.max_title_chars,
        "autoAppendTenderRequirements": rules.auto_append_tender_requirements,
    }


def _positive_int(value: Any, fallback: int) -> int:
    try:
        parsed = int(str(value))
    except (TypeError, ValueError):
        return fallback
    return parsed if parsed > 0 else fallback


def _non_negative_int(value: Any, fallback: int) -> int:
    try:
        parsed = int(str(value))
    except (TypeError, ValueError):
        return fallback
    return parsed if parsed >= 0 else fallback


def _bool_value(value: Any, fallback: bool) -> bool:
    if value is None:
        return fallback
    if isinstance(value, bool):
        return value
    return str(value).strip().lower() in {"1", "true", "yes", "on"}


def generate_toc_from_manifest(manifest: dict[str, Any]) -> dict[str, Any]:
    work_dir = Path(str(manifest.get("workDir") or "")).expanduser()
    if not str(work_dir):
        raise ValueError("workDir is required")
    work_dir.mkdir(parents=True, exist_ok=True)

    rules = _rules_from_manifest(manifest)
    template_file = _existing_path(manifest.get("templateFile"), "投标模板")
    tender_files = _tender_inputs(manifest)
    attach_file = _optional_existing_path(manifest.get("attachFile"))
    output_file = Path(str(manifest.get("outputFile") or work_dir / settings.s2_toc_output_file_name)).expanduser()
    evidence_file = Path(
        str(manifest.get("evidenceFile") or work_dir / settings.s2_toc_evidence_file_name)
    ).expanduser()

    template_outline = extract_docx_outline(template_file, rules=rules)
    if not template_outline:
        raise ValueError("投标模板未识别到可用目录标题，请检查模板是否包含 Word 标题或编号章节。")

    tender_candidates = extract_tender_candidates(tender_files, rules=rules)
    attach_outline = extract_docx_outline(attach_file, rules=rules) if attach_file else []
    items, decisions = build_toc_items(
        template_outline=template_outline,
        tender_candidates=tender_candidates,
        attach_outline=attach_outline,
        rules=rules,
    )
    annotation_counts = Counter(str(item.get("annotation") or "") for item in items)

    toc = {
        "schema_version": "bid-toc-json-v1",
        "document_title": "投标文件总目录",
        "project": {
            "projectId": str(manifest.get("projectId") or ""),
            "projectCode": str(manifest.get("projectCode") or ""),
            "projectName": str(manifest.get("projectName") or ""),
            "bidType": str(manifest.get("bidType") or ""),
        },
        "source_files": {
            "tender": [_source_file_payload(item) for item in tender_files],
            "template": str(template_file),
            "attach": str(attach_file) if attach_file else "",
            "output": str(output_file),
            "evidence": str(evidence_file),
        },
        "summary": {
            "total_items": len(items),
            "annotation_counts": dict(annotation_counts),
            "template_heading_count": len(template_outline),
            "tender_candidate_count": len(tender_candidates),
            "rule_config": _rules_to_dict(rules),
        },
        "items": items,
        "outputFile": str(output_file),
        "evidenceFile": str(evidence_file),
    }
    evidence = {
        "schema_version": "bid-toc-evidence-v1",
        "engine": "local-rule-engine",
        "inputs": toc["source_files"],
        "ruleConfig": _rules_to_dict(rules),
        "rules": [
            {
                "id": "template-first",
                "description": "投标模板标题作为目录主骨架。",
            },
            {
                "id": "uncovered-tender-requirements",
                "description": "招标文件中未被模板标题覆盖的通用要求追加为待审核目录项。",
            },
        ],
        "templateOutline": [_entry_evidence(item) for item in template_outline],
        "tenderCandidates": [_candidate_evidence(item) for item in tender_candidates],
        "decisions": decisions,
    }

    output_file.parent.mkdir(parents=True, exist_ok=True)
    output_file.write_text(json.dumps(toc, ensure_ascii=False, indent=2), encoding="utf-8")
    evidence_file.parent.mkdir(parents=True, exist_ok=True)
    evidence_file.write_text(json.dumps(evidence, ensure_ascii=False, indent=2), encoding="utf-8")
    toc["ruleEvidence"] = _public_rule_evidence(evidence)
    return toc


def extract_docx_outline(path: Path, *, rules: TocRules | None = None) -> list[OutlineEntry]:
    active_rules = rules or _default_rules()
    entries = _extract_docx_outline_from_xml(path, rules=active_rules)
    if entries:
        return entries
    if Document is not None:
        return _extract_docx_outline_with_python_docx(path, rules=active_rules)
    return []


def extract_tender_candidates(
    sources: list[dict[str, Any] | Path],
    *,
    rules: TocRules | None = None,
) -> list[TenderCandidate]:
    active_rules = rules or _default_rules()
    candidates: list[TenderCandidate] = []
    seen: set[str] = set()
    for source in sources:
        meta = _source_file_meta(source)
        path = meta["path"]
        lines = extract_docx_paragraph_text(path)
        appendix_group_titles: dict[str, str] = {}
        for index, line in enumerate(lines, start=1):
            for candidate in _candidates_from_line(line, rules=active_rules):
                group_key = _appendix_group_key(str(candidate.get("title") or ""))
                context_title = appendix_group_titles.get(group_key, "") if group_key else ""
                if group_key and not context_title and _is_appendix_group_heading(candidate):
                    context_title = str(candidate.get("title") or "")
                key = _candidate_identity_key(candidate, meta)
                if not key or key in seen:
                    continue
                seen.add(key)
                candidates.append(
                    TenderCandidate(
                        title=str(candidate.get("title") or ""),
                        raw_text=line,
                        source_file=str(path),
                        paragraph_index=index,
                        kind=str(candidate.get("kind") or "requirement"),
                        number=str(candidate.get("number") or ""),
                        file_id=str(meta.get("id") or ""),
                        file_name=str(meta.get("name") or path.name),
                        auto_add=_bool_value(candidate.get("autoAdd"), False),
                        context_title=context_title,
                    )
                )
                if group_key and _is_appendix_group_heading(candidate):
                    appendix_group_titles[group_key] = str(candidate.get("title") or "")
                if (
                    active_rules.max_tender_candidates > 0
                    and len(candidates) >= active_rules.max_tender_candidates
                ):
                    return candidates
    return candidates


def build_toc_items(
    *,
    template_outline: list[OutlineEntry],
    tender_candidates: list[TenderCandidate],
    attach_outline: list[OutlineEntry],
    rules: TocRules | None = None,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    active_rules = rules or _default_rules()
    items: list[dict[str, Any]] = []
    decisions: list[dict[str, Any]] = []
    template_index: list[dict[str, Any]] = []
    pending_appendix_items: list[dict[str, Any]] = []

    for order, entry in enumerate(template_outline, start=1):
        item = _toc_item_from_template(order, entry)
        items.append(item)
        template_index.append(_template_match_record(item, entry))
        decisions.append(
            {
                "title": entry.title,
                "action": "keep",
                "source": "template",
                "reason": "模板标题作为目录主骨架。",
                "sourceRef": _entry_evidence(entry),
            }
        )

    for candidate in tender_candidates:
        key = _title_key(candidate.title)
        match = _best_template_match(candidate, template_index)
        if match is not None:
            matched_item = match["item"]
            matched_item.setdefault("source_refs", []).append(_source_ref_from_candidate(candidate, role="basis"))
            decisions.append(
                {
                    "title": candidate.title,
                    "action": "covered",
                    "source": "tender",
                    "reason": "招标要求已被模板标题覆盖。",
                    "sourceRef": _candidate_evidence(candidate),
                    "matchedTitle": str(matched_item.get("title") or ""),
                }
            )
            continue
        if not _should_add_candidate_to_toc(candidate, rules=active_rules):
            decisions.append(
                {
                    "title": candidate.title,
                    "action": "candidate",
                    "source": "tender",
                    "reason": "招标要求作为审核证据保留，S2 保守模式不自动追加目录项。",
                    "sourceRef": _candidate_evidence(candidate),
                }
            )
            continue
        if candidate.kind == "appendix":
            pending_appendix_items.append(
                _toc_item_from_tender(
                    order=0,
                    number=candidate.number or _next_root_number(items + pending_appendix_items),
                    level=1,
                    candidate=candidate,
                )
            )
            decisions.append(
                {
                    "title": candidate.title,
                    "action": "add",
                    "source": "tender",
                    "reason": _add_reason_for_candidate(candidate),
                    "sourceRef": _candidate_evidence(candidate),
                }
            )
            continue
        parent_for_additions = _best_parent_for_candidate(candidate, template_outline)
        number = _next_child_number(items, parent_for_additions["number"])
        item = _toc_item_from_tender(
            order=len(items) + 1,
            number=number,
            level=min(int(parent_for_additions["level"]) + 1, active_rules.max_level),
            candidate=candidate,
        )
        items.insert(_insertion_index_for_parent(items, parent_for_additions), item)
        decisions.append(
            {
                "title": candidate.title,
                "action": "add",
                "source": "tender",
                "reason": _add_reason_for_candidate(candidate),
                "sourceRef": _candidate_evidence(candidate),
            }
        )

    covered_titles = {_title_key(entry.title) for entry in template_outline}
    for entry in attach_outline:
        if _is_covered_by_template(_title_key(entry.title), covered_titles):
            continue
        items.append(_toc_item_from_template(len(items) + 1, entry, annotation="保留", source="attachment"))

    items.extend(pending_appendix_items)

    for index, item in enumerate(items, start=1):
        item["order"] = index
    return items, decisions


def extract_docx_paragraph_text(path: Path) -> list[str]:
    paragraphs: list[str] = []

    current_parts: list[str] = []

    def flush_paragraph() -> None:
        line = _normalize_space("".join(current_parts))
        current_parts.clear()
        if line:
            paragraphs.append(line)

    try:
        with zipfile.ZipFile(path) as archive:
            with archive.open("word/document.xml") as xml_file:
                for _, element in ET.iterparse(xml_file, events=("end",)):
                    if element.tag == f"{WORD_NAMESPACE}t":
                        current_parts.append(element.text or "")
                    elif element.tag == f"{WORD_NAMESPACE}tab":
                        current_parts.append(" ")
                    elif element.tag in {f"{WORD_NAMESPACE}br", f"{WORD_NAMESPACE}cr"}:
                        current_parts.append(" ")
                    elif element.tag == f"{WORD_NAMESPACE}p":
                        flush_paragraph()
                        element.clear()
    except Exception:
        paragraphs = []
    if paragraphs or Document is None:
        return paragraphs

    try:
        document = Document(str(path))
        for paragraph in document.paragraphs:
            text = _normalize_space(paragraph.text)
            if text:
                paragraphs.append(text)
    except Exception:
        return []
    return paragraphs


def _extract_docx_outline_with_python_docx(path: Path, *, rules: TocRules) -> list[OutlineEntry]:
    document = Document(str(path))
    entries: list[OutlineEntry] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = _normalize_space(paragraph.text)
        if not text:
            continue
        style_name = str(paragraph.style.name if paragraph.style else "")
        parsed = _parse_heading_line(text, rules=rules)
        if not parsed and style_name.lower().startswith("heading"):
            style_level = _level_from_style(style_name, rules=rules)
            parsed = {"number": "", "title": _strip_heading_number(text), "level": style_level}
        if not parsed:
            continue
        title = _clean_title(parsed["title"])
        if not _usable_title(title, rules=rules):
            continue
        entries.append(
            OutlineEntry(
                number=str(parsed.get("number") or ""),
                title=title,
                level=int(parsed.get("level") or 1),
                source_file=str(path),
                paragraph_index=index,
                raw_text=text,
            )
        )
    return _normalize_outline_levels(entries, rules=rules)


def _extract_docx_outline_from_xml(path: Path, *, rules: TocRules) -> list[OutlineEntry]:
    toc_entries = _extract_docx_toc_outline_from_xml(path, rules=rules)
    if toc_entries:
        return _normalize_outline_levels(toc_entries, rules=rules)

    styled_entries = _extract_styled_docx_outline_from_xml(path, rules=rules)
    if styled_entries:
        return _normalize_outline_levels(styled_entries, rules=rules)

    entries: list[OutlineEntry] = []
    for index, line in enumerate(extract_docx_paragraph_text(path), start=1):
        parsed = _parse_heading_line(line, rules=rules)
        if not parsed:
            continue
        title = _clean_title(parsed["title"])
        if not _usable_title(title, rules=rules):
            continue
        entries.append(
            OutlineEntry(
                number=str(parsed.get("number") or ""),
                title=title,
                level=int(parsed.get("level") or 1),
                source_file=str(path),
                paragraph_index=index,
                raw_text=line,
            )
        )
    return _normalize_outline_levels(entries, rules=rules)


def _extract_docx_toc_outline_from_xml(path: Path, *, rules: TocRules) -> list[OutlineEntry]:
    entries: list[OutlineEntry] = []
    seen_toc_title = False
    for paragraph in _iter_docx_paragraphs(path):
        text = paragraph.text.strip()
        if _is_toc_title(text, paragraph.style_name):
            seen_toc_title = True
            continue

        toc_level = _toc_level_from_style(paragraph.style_id, paragraph.style_name, rules=rules)
        if toc_level is None:
            if entries and seen_toc_title:
                break
            continue
        if not _has_toc_page_reference(text, paragraph.tab_count):
            if entries and seen_toc_title:
                break
            continue
        if not seen_toc_title and not entries:
            seen_toc_title = True

        entry = _toc_entry_from_paragraph(path, paragraph, toc_level, rules=rules)
        if entry is not None:
            entries.append(entry)
    return entries


def _extract_styled_docx_outline_from_xml(path: Path, *, rules: TocRules) -> list[OutlineEntry]:
    entries: list[OutlineEntry] = []
    for paragraph in _iter_docx_paragraphs(path):
        if paragraph.outline_level is None or paragraph.outline_level > rules.max_level:
            continue
        text = _strip_toc_page_reference(paragraph.text)
        parsed = _parse_heading_line(text, rules=rules)
        number = ""
        title = text
        level = paragraph.outline_level
        if parsed:
            number = str(parsed.get("number") or "")
            title = str(parsed.get("title") or "")
        title = _clean_title(title)
        if not _usable_title(title, rules=rules):
            continue
        entries.append(
            OutlineEntry(
                number=number,
                title=title,
                level=level,
                source_file=str(path),
                paragraph_index=paragraph.paragraph_index,
                raw_text=text,
            )
        )
    return entries


def _iter_docx_paragraphs(path: Path) -> list[DocxParagraph]:
    try:
        styles = _docx_paragraph_styles(path)
        paragraphs: list[DocxParagraph] = []
        with zipfile.ZipFile(path) as archive:
            with archive.open("word/document.xml") as xml_file:
                for index, element in enumerate(_iter_docx_paragraph_elements(xml_file), start=1):
                    text, tab_count = _docx_paragraph_text_with_tabs(element)
                    if not text.strip():
                        element.clear()
                        continue
                    style_id = _paragraph_style_id(element)
                    direct_outline = _paragraph_direct_outline_level(element)
                    style = styles.get(style_id, {})
                    outline_level = direct_outline or style.get("outline_level")
                    paragraphs.append(
                        DocxParagraph(
                            text=text,
                            style_id=style_id,
                            style_name=str(style.get("name") or ""),
                            outline_level=outline_level if isinstance(outline_level, int) else None,
                            paragraph_index=index,
                            tab_count=tab_count,
                        )
                    )
                    element.clear()
        return paragraphs
    except Exception:
        return []


def _iter_docx_paragraph_elements(xml_file: Any):
    for _, element in ET.iterparse(xml_file, events=("end",)):
        if element.tag == f"{WORD_NAMESPACE}p":
            yield element


def _docx_paragraph_styles(path: Path) -> dict[str, dict[str, Any]]:
    styles: dict[str, dict[str, Any]] = {}
    try:
        with zipfile.ZipFile(path) as archive:
            with archive.open("word/styles.xml") as xml_file:
                root = ET.parse(xml_file).getroot()
    except Exception:
        return styles

    for style in root.findall(f"{WORD_NAMESPACE}style"):
        style_id = style.attrib.get(f"{WORD_NAMESPACE}styleId", "")
        if not style_id:
            continue
        name_el = style.find(f"{WORD_NAMESPACE}name")
        based_on_el = style.find(f"{WORD_NAMESPACE}basedOn")
        outline_el = style.find(f".//{WORD_NAMESPACE}outlineLvl")
        styles[style_id] = {
            "name": name_el.attrib.get(f"{WORD_NAMESPACE}val", "") if name_el is not None else "",
            "based_on": based_on_el.attrib.get(f"{WORD_NAMESPACE}val", "") if based_on_el is not None else "",
            "outline_raw": outline_el.attrib.get(f"{WORD_NAMESPACE}val", "") if outline_el is not None else "",
        }

    def resolve_outline(style_id: str, seen: set[str] | None = None) -> int | None:
        if not style_id or style_id not in styles:
            return None
        seen = set() if seen is None else seen
        if style_id in seen:
            return None
        seen.add(style_id)
        direct = _word_outline_level(styles[style_id].get("outline_raw"))
        if direct is not None:
            return direct
        return resolve_outline(str(styles[style_id].get("based_on") or ""), seen)

    for style_id, style in styles.items():
        style["outline_level"] = resolve_outline(style_id)
    return styles


def _docx_paragraph_text_with_tabs(element: ET.Element) -> tuple[str, int]:
    parts: list[str] = []
    tab_count = 0
    for node in element.iter():
        if node.tag == f"{WORD_NAMESPACE}t":
            parts.append(node.text or "")
        elif node.tag == f"{WORD_NAMESPACE}tab":
            parts.append("\t")
            tab_count += 1
        elif node.tag in {f"{WORD_NAMESPACE}br", f"{WORD_NAMESPACE}cr"}:
            parts.append(" ")
    return "".join(parts).strip(), tab_count


def _paragraph_style_id(element: ET.Element) -> str:
    style_el = element.find(f"./{WORD_NAMESPACE}pPr/{WORD_NAMESPACE}pStyle")
    return style_el.attrib.get(f"{WORD_NAMESPACE}val", "") if style_el is not None else ""


def _paragraph_direct_outline_level(element: ET.Element) -> int | None:
    outline_el = element.find(f"./{WORD_NAMESPACE}pPr/{WORD_NAMESPACE}outlineLvl")
    value = outline_el.attrib.get(f"{WORD_NAMESPACE}val", "") if outline_el is not None else ""
    return _word_outline_level(value)


def _word_outline_level(value: Any) -> int | None:
    try:
        level = int(str(value))
    except (TypeError, ValueError):
        return None
    return level + 1


def _is_toc_title(text: str, style_name: str) -> bool:
    normalized = _clean_title(_strip_toc_page_reference(text))
    if normalized in {"目录", "目 录", "目次"}:
        return True
    return normalized in {"目录", "目次"} and "toc" in style_name.lower()


def _toc_level_from_style(style_id: str, style_name: str, *, rules: TocRules) -> int | None:
    text = f"{style_id} {style_name}".lower()
    if "toc heading" in text and not re.search(r"(?:toc|目录)\s*(?:标题)?\s*[1-9]", text, re.I):
        return None
    patterns = (
        r"(?:^|[^a-z])toc\s*([1-9])(?:\b|[^0-9])",
        r"目录\s*([1-9])",
        r"目录\s*标题\s*([1-9])",
        r"toc\s*标题\s*([1-9])",
    )
    for pattern in patterns:
        match = re.search(pattern, text, re.I)
        if match:
            return max(1, min(int(match.group(1)), rules.max_level))
    return None


def _has_toc_page_reference(text: str, tab_count: int) -> bool:
    if tab_count > 0:
        parts = [part.strip() for part in text.split("\t") if part.strip()]
        if len(parts) >= 2 and re.fullmatch(r"[ivxlcdmIVXLCDM]*\d+[A-Za-z-]*", parts[-1]):
            return True
    return bool(re.search(r"(?:\.{2,}|…{2,})\s*\d+\s*$", text))


def _toc_entry_from_paragraph(
    path: Path,
    paragraph: DocxParagraph,
    toc_level: int,
    *,
    rules: TocRules,
) -> OutlineEntry | None:
    text = _strip_toc_page_reference(paragraph.text)
    if not text:
        return None
    parsed = _parse_heading_line(text, rules=rules)
    number = ""
    title = text
    if parsed:
        number = str(parsed.get("number") or "")
        title = str(parsed.get("title") or "")
    title = _clean_title(title)
    if not _usable_title(title, rules=rules):
        return None
    return OutlineEntry(
        number=number,
        title=title,
        level=toc_level,
        source_file=str(path),
        paragraph_index=paragraph.paragraph_index,
        raw_text=text,
    )


def _strip_toc_page_reference(text: str) -> str:
    value = str(text or "").strip()
    if "\t" in value:
        parts = [part.strip(" .…·•") for part in value.split("\t") if part.strip(" .…·•")]
        if len(parts) >= 2 and re.fullmatch(r"[ivxlcdmIVXLCDM]*\d+[A-Za-z-]*", parts[-1]):
            return parts[0].strip()
    value = re.sub(r"(?:\.{2,}|…{2,}|·{2,}|•{2,})\s*\d+\s*$", "", value).strip()
    if _looks_like_toc_numbered_line(value):
        value = re.sub(r"\s+\d{1,4}\s*$", "", value).strip()
    value = re.sub(r"(?<=\D)\s+\d{1,4}\s*$", "", value).strip()
    return _normalize_space(value)


def _looks_like_toc_numbered_line(value: str) -> bool:
    text = _normalize_space(value)
    return bool(
        re.match(rf"^(?:第(?:{CHINESE_NUMERAL_PATTERN}|\d+)章|\d+(?:\.\d+){{1,6}}|(?:技术\s*)?附(?:件|表|录)\s*[A-Za-z0-9一二三四五六七八九十百千万零〇两.-]+)", text)
    )


def _parse_heading_line(text: str, *, rules: TocRules | None = None) -> dict[str, Any] | None:
    active_rules = rules or _default_rules()
    line = _strip_toc_page_reference(text)
    if not line or len(line) > 120:
        return None
    for pattern in HEADING_PATTERNS:
        match = pattern.match(line)
        if not match:
            continue
        number = str(match.group("num") or "").strip()
        title = _clean_title(str(match.group("title") or ""))
        if not _usable_title(title, rules=active_rules):
            return None
        return {
            "number": _display_number(line, number),
            "title": title,
            "level": _level_from_number(line, number, rules=active_rules),
        }
    return None


def _candidate_titles_from_line(line: str, *, rules: TocRules | None = None) -> list[str]:
    return [
        str(item.get("title") or "")
        for item in _candidates_from_line(line, rules=rules)
        if str(item.get("title") or "").strip()
    ]


def _candidates_from_line(line: str, *, rules: TocRules | None = None) -> list[dict[str, str]]:
    active_rules = rules or _default_rules()
    text = _normalize_space(line)
    if not text or len(text) > 180:
        return []
    toc_text = _strip_toc_page_reference(text)
    candidates: list[dict[str, str]] = []

    parsed_heading = _parse_heading_line(text, rules=active_rules)
    if parsed_heading and _looks_like_tender_heading(text, parsed_heading):
        candidate_kind = "appendix" if _is_appendix_title(text) else "heading"
        candidate_number = str(parsed_heading.get("number") or "")
        candidate_title = str(parsed_heading.get("title") or "")
        if candidate_kind == "appendix" and candidate_number:
            candidate_title = _clean_title(f"{candidate_number} {candidate_title}")
        candidates.append(
            {
                "title": candidate_title,
                "kind": candidate_kind,
                "number": candidate_number,
                "autoAdd": "true" if _is_technical_appendix_title(text) else "false",
            }
        )

    appendix_match = APPENDIX_TITLE_PATTERN.match(toc_text)
    if appendix_match:
        appendix_title = _clean_title(
            f"{appendix_match.group('prefix')} {appendix_match.group('title')}"
        )
        if _usable_title(appendix_title, rules=active_rules):
            candidates.append(
                {
                    "title": appendix_title[: active_rules.max_title_chars].strip(),
                    "kind": "appendix",
                    "number": str(appendix_match.group("prefix") or ""),
                    "autoAdd": "true" if _is_technical_appendix_title(toc_text) else "false",
                }
            )

    table_match = TABLE_TITLE_PATTERN.match(toc_text)
    if table_match:
        table_title = _clean_title(table_match.group("title"))
        if _usable_title(table_title, rules=active_rules):
            candidates.append(
                {
                    "title": table_title[: active_rules.max_title_chars].strip(),
                    "kind": "appendix",
                    "number": str(table_match.group("prefix") or ""),
                    "autoAdd": "false",
                }
            )

    for pattern in REQUIREMENT_PATTERNS:
        for match in pattern.finditer(text):
            title = _clean_requirement_title(match.group("title"), rules=active_rules)
            if _usable_title(title, rules=active_rules):
                candidates.append(
                    {
                        "title": title,
                        "kind": "requirement",
                        "number": "",
                    }
                )
    return _dedupe_candidates(candidates)


def _looks_like_tender_heading(text: str, parsed: dict[str, Any]) -> bool:
    title = str(parsed.get("title") or "")
    number = str(parsed.get("number") or "")
    if _is_appendix_title(text):
        return True
    if not re.fullmatch(r"\d+(?:\.\d+){0,5}|第(?:\d+|[一二三四五六七八九十百千万零〇两]+)章", number):
        return False
    if re.search(r"(目录|附表|附件|附录|技术规范|供货范围|评分|评审|响应|承诺|方案|要求|偏差|表)$", title):
        return True
    return len(title) >= 4 and len(title) <= 40


def _is_appendix_title(text: str) -> bool:
    return bool(re.match(r"^(?:技术\s*)?附(?:件|表|录)", _normalize_space(text)))


def _is_technical_appendix_title(text: str) -> bool:
    return bool(re.match(r"^(?:技术\s*)?附表", _normalize_space(text)))


def _dedupe_candidates(candidates: list[dict[str, str]]) -> list[dict[str, str]]:
    result: list[dict[str, str]] = []
    seen: set[str] = set()
    for candidate in candidates:
        title = str(candidate.get("title") or "").strip()
        kind = str(candidate.get("kind") or "requirement")
        key = f"{kind}:{_title_key(title)}"
        if not title or not key or key in seen:
            continue
        seen.add(key)
        result.append(candidate)
    return result


def _candidate_identity_key(candidate: dict[str, str], meta: dict[str, Any]) -> str:
    title = str(candidate.get("title") or "")
    key = _title_key(title)
    if not key:
        return ""
    return f"{meta.get('id') or meta.get('name') or meta.get('path')}:{candidate.get('kind') or 'requirement'}:{key}"


def _appendix_group_key(value: str) -> str:
    match = re.search(r"(?:技术\s*)?附表\s*([A-Za-z一二三四五六七八九十百千万零〇两]+)", _normalize_space(value))
    return match.group(1).upper() if match else ""


def _is_appendix_group_heading(candidate: dict[str, str]) -> bool:
    title = str(candidate.get("title") or "")
    return "附表" in title and not re.match(r"^(?:技术\s*)?附表\s*[A-Za-z一二三四五六七八九十百千万零〇两]+\.\d+", title)


def _clean_requirement_title(value: str, *, rules: TocRules) -> str:
    title = _clean_title(value)
    title = GENERIC_PREFIX_PATTERN.sub("", title).strip()
    title = re.sub(r"^(?:全部|相关|相应|本项目|项目|完整的?)", "", title).strip()
    title = re.split(r"[，,：:]", title, maxsplit=1)[0].strip()
    title = re.sub(r"^[的地得]", "", title).strip()
    title = re.sub(r"(?:等|及相关|相关)?(?:材料|文件|内容|要求)$", "", title).strip() or title
    return title[: rules.max_title_chars].strip()


def _clean_title(value: str) -> str:
    title = _normalize_space(value)
    title = re.sub(r"[。；;]+$", "", title).strip()
    title = re.sub(r"^[：:、.．\-\s]+", "", title).strip()
    return title


def _strip_heading_number(value: str) -> str:
    parsed = _parse_heading_line(value)
    if parsed:
        return str(parsed.get("title") or "")
    return _clean_title(value)


def _usable_title(title: str, *, rules: TocRules | None = None) -> bool:
    active_rules = rules or _default_rules()
    text = _clean_title(title)
    if len(text) < 2 or len(text) > active_rules.max_title_chars:
        return False
    if len(text) <= 4 and re.search(r"(性|性要求|要求|专题|表)\s*\d*$", text):
        return False
    if re.fullmatch(r"[\u4e00-\u9fff]{1,3}\s*\d{1,4}", text):
        return False
    if text in NOISE_TITLES:
        return False
    if re.fullmatch(r"[\d.．、（）()\s]+", text):
        return False
    return True


def _normalize_outline_levels(entries: list[OutlineEntry], *, rules: TocRules) -> list[OutlineEntry]:
    normalized: list[OutlineEntry] = []
    previous_level = 0
    for entry in entries:
        level = max(1, min(entry.level, rules.max_level))
        if not normalized:
            level = 1
        elif level > previous_level + 1:
            level = previous_level + 1
        normalized.append(
            OutlineEntry(
                number=entry.number,
                title=entry.title,
                level=level,
                source_file=entry.source_file,
                paragraph_index=entry.paragraph_index,
                raw_text=entry.raw_text,
            )
        )
        previous_level = level
    return normalized


def _level_from_style(style_name: str, *, rules: TocRules) -> int:
    match = re.search(r"(\d+)", style_name)
    if not match:
        return 1
    return max(1, min(int(match.group(1)), rules.max_level))


def _level_from_number(line: str, number: str, *, rules: TocRules | None = None) -> int:
    active_rules = rules or _default_rules()
    if re.match(rf"^第(?:{CHINESE_NUMERAL_PATTERN}|\d+)章", line):
        return 1
    if "." in number:
        return max(1, min(number.count(".") + 1, active_rules.max_level))
    if re.match(rf"^{CHINESE_NUMERAL_PATTERN}[、.．]", line):
        return 1
    return 1


def _display_number(line: str, number: str) -> str:
    chapter = re.match(rf"^(第(?:{CHINESE_NUMERAL_PATTERN}|\d+)章)", line)
    if chapter:
        return chapter.group(1)
    return number


def _toc_item_from_template(
    order: int,
    entry: OutlineEntry,
    *,
    annotation: str = "保留",
    source: str = "template",
) -> dict[str, Any]:
    return {
        "order": order,
        "number": entry.number,
        "title": entry.title,
        "level": entry.level,
        "annotation": annotation,
        "source": source,
        "reason": "来自投标模板目录骨架。",
        "source_refs": [
            {
                "type": "template",
                "role": "skeleton",
                "path": entry.source_file,
                "paragraphIndex": entry.paragraph_index,
                "raw_text": entry.raw_text,
                "rawText": entry.raw_text,
                "basisText": entry.raw_text,
            }
        ],
        "material_refs": [],
    }


def _toc_item_from_tender(
    *,
    order: int,
    number: str,
    level: int,
    candidate: TenderCandidate,
) -> dict[str, Any]:
    return {
        "order": order,
        "number": number,
        "title": candidate.title,
        "level": level,
        "annotation": "新增-副表" if candidate.kind == "appendix" else "新增-招标要求",
        "source": "tender",
        "reason": _add_reason_for_candidate(candidate),
        "source_refs": [_source_ref_from_candidate(candidate, role="basis")],
        "material_refs": [],
    }


def _best_parent_for_additions(template_outline: list[OutlineEntry]) -> dict[str, Any]:
    level_one_entries = [entry for entry in template_outline if entry.level == 1]
    parent = level_one_entries[-1] if level_one_entries else template_outline[-1]
    return {
        "number": parent.number or str(len(level_one_entries) or 1),
        "title": parent.title,
        "level": parent.level,
    }


def _best_parent_for_candidate(candidate: TenderCandidate, template_outline: list[OutlineEntry]) -> dict[str, Any]:
    if not template_outline:
        return {"number": "", "title": "", "level": 1}

    best = template_outline[-1]
    best_score = 0.0
    candidate_context = " ".join(
        item for item in (candidate.title, candidate.context_title, candidate.raw_text) if item
    )
    candidate_key = _title_key(candidate_context)
    candidate_tokens = _title_tokens(candidate_context)
    for entry in template_outline:
        title_key = _title_key(entry.title)
        title_tokens = _title_tokens(entry.title)
        score = _match_score(candidate_key, candidate_tokens, candidate_context, title_key, title_tokens, candidate_tokens)
        if entry.level == 1:
            score *= 0.92
        if score > best_score:
            best = entry
            best_score = score
    return {
        "number": best.number or str(template_outline.index(best) + 1),
        "title": best.title,
        "level": best.level,
    }


def _insertion_index_for_parent(items: list[dict[str, Any]], parent: dict[str, Any]) -> int:
    parent_level = int(parent.get("level") or 1)
    parent_number = str(parent.get("number") or "").strip()
    parent_title = str(parent.get("title") or "").strip()
    parent_index: int | None = None

    for index, item in enumerate(items):
        if int(item.get("level") or 1) != parent_level:
            continue
        item_number = str(item.get("number") or "").strip()
        item_title = str(item.get("title") or "").strip()
        if (parent_number and item_number == parent_number) or (parent_title and item_title == parent_title):
            parent_index = index

    if parent_index is None:
        return len(items)

    insert_at = parent_index + 1
    while insert_at < len(items) and int(items[insert_at].get("level") or 1) > parent_level:
        insert_at += 1
    return insert_at


def _next_child_number(items: list[dict[str, Any]], parent_number: str) -> str:
    clean_parent = str(parent_number or "").strip()
    if not clean_parent:
        return str(len(items) + 1)
    if re.fullmatch(r"\d+(?:\.\d+)*", clean_parent):
        child_count = sum(1 for item in items if str(item.get("number") or "").startswith(clean_parent + "."))
        return f"{clean_parent}.{child_count + 1}"
    return ""


def _next_root_number(items: list[dict[str, Any]]) -> str:
    root_numbers: list[int] = []
    for item in items:
        if int(item.get("level") or 1) != 1:
            continue
        number = str(item.get("number") or "").strip()
        match = re.fullmatch(r"\d+", number)
        if match:
            root_numbers.append(int(match.group(0)))
    return str(max(root_numbers, default=0) + 1)


def _template_match_record(item: dict[str, Any], entry: OutlineEntry) -> dict[str, Any]:
    return {
        "item": item,
        "title": entry.title,
        "titleKey": _title_key(entry.title),
        "tokens": _title_tokens(entry.title),
    }


def _best_template_match(
    candidate: TenderCandidate,
    template_index: list[dict[str, Any]],
) -> dict[str, Any] | None:
    candidate_key = _title_key(candidate.title)
    candidate_tokens = _title_tokens(candidate.title)
    raw_tokens = _title_tokens(candidate.raw_text)
    best: dict[str, Any] | None = None
    best_score = 0.0

    for record in template_index:
        title_key = str(record.get("titleKey") or "")
        if not candidate_key or not title_key:
            continue
        title_tokens = set(record.get("tokens") or set())
        score = _match_score(candidate_key, candidate_tokens, candidate.raw_text, title_key, title_tokens, raw_tokens)
        if score > best_score:
            best = record
            best_score = score

    return best if best is not None and best_score >= _match_threshold(candidate) else None


def _match_score(
    candidate_key: str,
    candidate_tokens: set[str],
    raw_text: str,
    title_key: str,
    title_tokens: set[str],
    raw_tokens: set[str],
) -> float:
    if candidate_key in title_key or title_key in candidate_key:
        return 1.0
    if title_key and title_key in _title_key(raw_text):
        return 0.92
    if not candidate_tokens or not title_tokens:
        return 0.0
    overlap = candidate_tokens & title_tokens
    score = len(overlap) / max(1, min(len(candidate_tokens), len(title_tokens)))
    if title_tokens and title_tokens <= raw_tokens:
        score = max(score, 0.86)
    elif overlap and raw_tokens:
        score = max(score, len(overlap) / max(1, len(title_tokens)))
    return score


def _match_threshold(candidate: TenderCandidate) -> float:
    if candidate.kind in {"heading", "appendix"}:
        return 0.58
    return 0.66


def _title_tokens(value: str) -> set[str]:
    text = _clean_title(value)
    tokens = re.findall(r"[\u4e00-\u9fff]{2,}|[A-Za-z0-9]+", text)
    result: set[str] = set()
    for token in tokens:
        token = GENERIC_PREFIX_PATTERN.sub("", token).strip().lower()
        if not token or token in NOISE_TITLES:
            continue
        if len(token) <= 6:
            result.add(token)
            continue
        for size in (2, 3, 4):
            for index in range(0, max(0, len(token) - size + 1)):
                result.add(token[index : index + size])
    return result


def _should_add_candidate_to_toc(candidate: TenderCandidate, *, rules: TocRules) -> bool:
    if candidate.kind == "appendix" and candidate.auto_add:
        return True
    return rules.auto_append_tender_requirements


def _add_reason_for_candidate(candidate: TenderCandidate) -> str:
    if candidate.kind == "appendix":
        return "招标文件副表/附表未在投标模板目录中覆盖，追加为待审核目录项。"
    if candidate.kind == "heading":
        return "招标文件章节未在投标模板目录中覆盖，追加为待审核目录项。"
    return "招标要求未在投标模板目录中覆盖，追加为待审核目录项。"


def _source_ref_from_candidate(candidate: TenderCandidate, *, role: str) -> dict[str, Any]:
    basis_text = _basis_text(candidate.raw_text, candidate.title)
    return {
        "type": "tender",
        "role": role,
        "kind": candidate.kind,
        "fileId": candidate.file_id,
        "fileName": candidate.file_name or Path(candidate.source_file).name,
        "path": candidate.source_file,
        "paragraphIndex": candidate.paragraph_index,
        "raw_text": candidate.raw_text,
        "rawText": candidate.raw_text,
        "basisText": basis_text,
        "searchText": basis_text,
        "title": candidate.title,
        "number": candidate.number,
        "contextTitle": candidate.context_title,
    }


def _basis_text(raw_text: str, title: str) -> str:
    text = _normalize_space(raw_text)
    if text:
        return text[:80]
    return _normalize_space(title)[:80]


def _title_key(title: str) -> str:
    text = _clean_title(title).lower()
    text = GENERIC_PREFIX_PATTERN.sub("", text)
    text = re.sub(r"[\s　,，.。:：;；、（）()\[\]【】《》<>\"'“”‘’\-_/\\]+", "", text)
    text = re.sub(r"(方案|计划|安排|要求|说明|响应|材料|文件|内容)$", "", text)
    return text


def _is_covered_by_template(candidate_key: str, covered_titles: set[str]) -> bool:
    if not candidate_key:
        return True
    for title_key in covered_titles:
        if not title_key:
            continue
        if candidate_key in title_key or title_key in candidate_key:
            return True
    return False


def _entry_evidence(entry: OutlineEntry) -> dict[str, Any]:
    return {
        "number": entry.number,
        "title": entry.title,
        "level": entry.level,
        "sourceFile": entry.source_file,
        "paragraphIndex": entry.paragraph_index,
        "rawText": entry.raw_text,
    }


def _candidate_evidence(candidate: TenderCandidate) -> dict[str, Any]:
    return {
        "title": candidate.title,
        "kind": candidate.kind,
        "number": candidate.number,
        "fileId": candidate.file_id,
        "fileName": candidate.file_name,
        "sourceFile": candidate.source_file,
        "paragraphIndex": candidate.paragraph_index,
        "rawText": candidate.raw_text,
        "basisText": _basis_text(candidate.raw_text, candidate.title),
        "contextTitle": candidate.context_title,
    }


def _public_rule_evidence(evidence: dict[str, Any]) -> dict[str, Any]:
    decisions = evidence.get("decisions") if isinstance(evidence.get("decisions"), list) else []
    candidates = evidence.get("tenderCandidates") if isinstance(evidence.get("tenderCandidates"), list) else []
    template_outline = evidence.get("templateOutline") if isinstance(evidence.get("templateOutline"), list) else []
    return {
        "schemaVersion": str(evidence.get("schema_version") or ""),
        "engine": str(evidence.get("engine") or ""),
        "ruleConfig": evidence.get("ruleConfig") if isinstance(evidence.get("ruleConfig"), dict) else {},
        "templateOutlineCount": len(template_outline),
        "tenderCandidateCount": len(candidates),
        "decisions": copy_decisions(decisions),
    }


def copy_decisions(decisions: list[Any]) -> list[dict[str, Any]]:
    copied: list[dict[str, Any]] = []
    for item in decisions:
        if isinstance(item, dict):
            copied.append(dict(item))
    return copied


def _existing_path(value: Any, label: str) -> Path:
    path = Path(str(value or "")).expanduser()
    if not str(value or "").strip() or not path.exists():
        raise ValueError(f"{label}不存在，请先上传可读取的投标模板文件。")
    return path


def _optional_existing_path(value: Any) -> Path | None:
    text = str(value or "").strip()
    if not text:
        return None
    path = Path(text).expanduser()
    return path if path.exists() else None


def _tender_inputs(manifest: dict[str, Any]) -> list[dict[str, Any]]:
    sources: list[dict[str, Any]] = []
    for item in manifest.get("tenderFiles") or []:
        raw_path = item.get("path") if isinstance(item, dict) else item
        path = Path(str(raw_path or "")).expanduser()
        if path.exists() and path.suffix.lower() == ".docx":
            sources.append(
                {
                    "id": str(item.get("id") or "") if isinstance(item, dict) else "",
                    "name": str(item.get("name") or path.name) if isinstance(item, dict) else path.name,
                    "path": path,
                    "originalPath": str(item.get("originalPath") or "") if isinstance(item, dict) else "",
                }
            )
    fallback = Path(str(manifest.get("tenderFile") or "")).expanduser()
    known_paths = {Path(str(item.get("path") or "")).expanduser() for item in sources}
    if fallback.exists() and fallback.suffix.lower() == ".docx" and fallback not in known_paths:
        sources.append({"id": "", "name": fallback.name, "path": fallback, "originalPath": ""})
    if not sources:
        raise ValueError("招标文件不存在，请先完成 S1 解析并上传可读取的 docx 招标文件。")
    return sources


def _source_file_meta(source: dict[str, Any] | Path) -> dict[str, Any]:
    if isinstance(source, dict):
        path = Path(str(source.get("path") or "")).expanduser()
        return {
            "id": str(source.get("id") or ""),
            "name": str(source.get("name") or path.name),
            "path": path,
            "originalPath": str(source.get("originalPath") or ""),
        }
    path = Path(str(source)).expanduser()
    return {"id": "", "name": path.name, "path": path, "originalPath": ""}


def _source_file_payload(source: dict[str, Any] | Path) -> dict[str, str]:
    meta = _source_file_meta(source)
    return {
        "id": str(meta.get("id") or ""),
        "name": str(meta.get("name") or ""),
        "path": str(meta.get("path") or ""),
        "originalPath": str(meta.get("originalPath") or ""),
    }


def _normalize_space(value: str) -> str:
    return re.sub(r"\s+", " ", str(value or "").replace("\u3000", " ")).strip()
