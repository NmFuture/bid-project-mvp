from __future__ import annotations

import json
import copy
import asyncio
import mimetypes
import re
import shutil
import subprocess
import sys
import threading
import zipfile
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any, Callable
from xml.etree import ElementTree as ET

from docx import Document

from app.core.config import settings
from app.services.business_section_tree import write_business_section_tree
from app.services.business_template_extractor import run_business_template_extractor
from app.services.bid_type import BUSINESS_BID_TYPE, TECHNICAL_BID_TYPE
from app.services.ocr_service import IMAGE_SUFFIXES, ocr_service
from app.services.opencode_client import OpencodeClient
from app.services.parse_profiles import (
    BUSINESS_PARSE_PROFILE,
    TECHNICAL_PARSE_PROFILE,
    ParseProfile,
    TECHNICAL_PARSE_SKILL_NAME,
    resolve_parse_profile,
)
from app.services.peripheral import PeripheralError

PARSER_CORE_DIR = (
    Path(__file__).resolve().parents[2] / "opencode" / "skill" / TECHNICAL_PARSE_SKILL_NAME / "scripts"
)
if str(PARSER_CORE_DIR) not in sys.path:
    sys.path.insert(0, str(PARSER_CORE_DIR))

from parser_core import parse_documents as parse_structured_documents  # noqa: E402

WORD_NAMESPACE = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
TEXT_PREVIEW_LIMIT = 600


@dataclass(frozen=True)
class ParseCategory:
    key: str
    label: str
    keywords: tuple[str, ...]


PARSE_CATEGORIES: tuple[ParseCategory, ...] = (
    ParseCategory("scoring_criteria", "评分细则", ("评分", "得分", "分值", "证明材料", "商务评分", "技术评分")),
    ParseCategory(
        "project_basics",
        "项目基础信息",
        (
            "项目名称",
            "招标编号",
            "项目编号",
            "招标人",
            "管理单位",
            "建设单位",
            "业主",
            "标段规模",
            "招标规模",
            "建设规模",
            "交货周期",
            "交货期",
            "质保期",
            "质量保证期",
            "技术承诺",
            "起始日期",
            "开工日期",
            "截止日期",
            "截止时间",
            "结束日期",
        ),
    ),
    ParseCategory(
        "turbine_parameters",
        "风机核心参数",
        (
            "单机容量",
            "叶轮直径",
            "轮毂高度",
            "叶片最低点",
            "塔筒型式",
            "箱变型式",
            "安全等级",
            "空气密度",
            "风速",
            "湍流强度",
        ),
    ),
    ParseCategory(
        "performance_guarantees",
        "性能保证指标",
        ("功率曲线", "可利用率", "发电量", "涉网性能", "保证率", "性能保证"),
    ),
    ParseCategory(
        "environment_adaptation",
        "环境适应性要求",
        ("低温", "覆冰", "防凝露", "潮湿", "防雷暴", "防雷", "风沙", "高温", "环境适应性"),
    ),
    ParseCategory(
        "topic_plans",
        "专题方案要求",
        ("专题方案", "叶片", "变桨系统", "主轴", "齿轮箱", "总体方案", "技术先进性"),
    ),
    ParseCategory("tables_and_scope", "附表和供货范围", ("附表", "供货范围", "供货清单", "供货界面")),
    ParseCategory(
        "assessment_terms",
        "考核条款",
        ("考核", "发电量考核", "可利用率考核", "功率曲线考核", "部件考核", "认证考核"),
    ),
)

DATE_PATTERN = re.compile(
    r"(?P<year>20\d{2})\s*(?:年|[-/.])\s*(?P<month>\d{1,2})\s*(?:月|[-/.])\s*(?P<day>\d{1,2})\s*日?"
)
BID_DEADLINE_DATE_PATTERN = re.compile(
    r"(?P<year>20\d{2})\s*(?:年|[-/.])\s*(?P<month>\d{1,2})\s*(?:月|[-/.])\s*(?P<day>\d{1,2})\s*日?"
    r"(?:[\sT]*(?P<hour>\d{1,2})\s*(?:时|:|：)\s*(?P<minute>\d{1,2})\s*分?)?"
)
LABEL_VALUE_PATTERN = re.compile(r"^\s*(?P<label>[^:：]{2,40})\s*[:：]\s*(?P<value>.+?)\s*$")
LEADING_NUMBER_PATTERN = re.compile(r"^\s*(?:第?[一二三四五六七八九十百千0-9]+[章节条]?|[（(]?\d+[）)]?)\s*[、.．\s]+")

START_DATE_KEYWORDS = ("起始日期", "开始日期", "开工日期", "计划开工", "服务期自", "合同开始")
END_DATE_KEYWORDS = ("投标截止", "截止日期", "截止时间", "结束日期", "竣工日期", "完成日期", "服务期至")
BID_DEADLINE_CONTEXT = ("递交截止时间", "投标文件递交截止时间", "投标截止时间", "提交截止时间")
OPENING_TIME_CONTEXT = ("开标时间", "开标日期", "开标时间和地点", "开标地点")


@dataclass(frozen=True)
class FieldSpec:
    key: str
    label: str
    aliases: tuple[str, ...]


PROJECT_BASIC_FIELDS: tuple[FieldSpec, ...] = (
    FieldSpec("projectName", "项目名称", ("项目名称",)),
    FieldSpec("tenderNo", "招标编号", ("招标编号", "项目编号")),
    FieldSpec("tenderer", "招标人", ("招标人", "业主", "建设单位")),
    FieldSpec("managementUnit", "管理单位", ("管理单位",)),
    FieldSpec("bidSectionScale", "标段规模", ("标段规模", "招标规模", "建设规模")),
    FieldSpec("deliveryPeriod", "交货周期", ("交货周期", "交货期")),
    FieldSpec("warrantyPeriod", "质保期", ("质保期", "质量保证期")),
    FieldSpec("technicalCommitment", "技术承诺", ("技术承诺",)),
)

TURBINE_CORE_FIELDS: tuple[FieldSpec, ...] = (
    FieldSpec("singleCapacity", "单机容量", ("单机容量",)),
    FieldSpec("rotorDiameter", "叶轮直径", ("叶轮直径",)),
    FieldSpec("hubHeight", "轮毂高度", ("轮毂高度",)),
    FieldSpec("bladeTipClearance", "叶片最低点距地", ("叶片最低点距地", "叶片最低点")),
    FieldSpec("towerType", "塔筒型式", ("塔筒型式",)),
    FieldSpec("boxTransformerType", "箱变型式", ("箱变型式",)),
    FieldSpec("safetyClass", "安全等级", ("安全等级",)),
    FieldSpec("airDensity", "空气密度", ("空气密度",)),
    FieldSpec("windSpeed", "风速", ("风速",)),
    FieldSpec("turbulenceIntensity", "湍流强度", ("湍流强度",)),
)

PERFORMANCE_FIELDS: tuple[FieldSpec, ...] = (
    FieldSpec("powerCurve", "功率曲线", ("功率曲线",)),
    FieldSpec("availability", "可利用率", ("可利用率",)),
    FieldSpec("generation", "发电量", ("发电量", "上网电量")),
    FieldSpec("gridPerformance", "涉网性能", ("涉网性能", "高低电压穿越", "电压穿越")),
)

ENVIRONMENT_FIELDS: tuple[FieldSpec, ...] = (
    FieldSpec("lowTemperature", "抗低温", ("抗低温", "低温")),
    FieldSpec("icingCondensation", "抗覆冰防凝露", ("抗覆冰", "覆冰", "防凝露")),
    FieldSpec("humidity", "防潮湿", ("防潮湿", "潮湿")),
    FieldSpec("lightning", "防雷暴", ("防雷暴", "防雷")),
    FieldSpec("sandstorm", "防风沙", ("防风沙", "风沙")),
    FieldSpec("highTemperature", "抗高温", ("抗高温", "高温")),
)

BUSINESS_RESPONSE_FIELDS: tuple[FieldSpec, ...] = (
    FieldSpec("bidLetterRequired", "投标函要求", ("投标函",)),
    FieldSpec("authorizationLetterRequired", "授权委托书要求", ("授权委托书", "法定代表人授权委托书")),
    FieldSpec("integrityCommitmentRequired", "廉洁承诺要求", ("廉洁", "廉洁自律承诺", "廉洁承诺")),
    FieldSpec("sealValidityStatementRequired", "投标专用章效力说明要求", ("投标专用章效力说明",)),
    FieldSpec("bidPriceTableRequired", "投标价格表要求", ("投标价格", "投标价格表")),
    FieldSpec("openingPriceTableRequired", "开标价格表要求", ("开标价格表",)),
    FieldSpec("specificationTableRequired", "货物规格表要求", ("货物规格", "规格表")),
    FieldSpec("commercialDeviationTableRequired", "商务偏差表要求", ("商务偏差", "偏差表")),
    FieldSpec("supplyScopeTableRequired", "供货范围表要求", ("供货范围", "供货范围表")),
    FieldSpec("bidSecurityRequired", "投标保证金要求", ("投标保证金", "保证金", "保函")),
    FieldSpec("performanceBondCommitmentRequired", "履约保证承诺要求", ("履约保证函", "履约承诺", "履约保证")),
    FieldSpec("attachment9Required", "附件9要求", ("附件9", "附件九")),
)

QUALIFICATION_SUPPORT_FIELDS: tuple[FieldSpec, ...] = (
    FieldSpec("qualificationDocumentRequired", "资格证明文件要求", ("资格证明", "合格投标人", "资格审查")),
    FieldSpec("performanceDocumentRequired", "业绩证明文件要求", ("业绩证明", "合同扫描件", "中标通知书", "验收报告")),
    FieldSpec("financialDocumentRequired", "财务文件要求", ("审计报告", "财务报表", "财务状况")),
    FieldSpec("creditDocumentRequired", "资信诚信文件要求", ("资信证明", "信用中国", "纳税信用", "失信")),
    FieldSpec("certificationDocumentRequired", "证书文件要求", ("认证证书", "资质证书", "体系认证")),
    FieldSpec("customerSpecificProofRequired", "客户专项证明要求", ("战略协议", "框架协议", "评价信", "优秀供应商证明", "示范应用证明")),
)

COMMITMENT_REQUIREMENT_FIELDS: tuple[FieldSpec, ...] = (
    FieldSpec("generalCommitmentCount", "承诺线索识别数", ("承诺",)),
    FieldSpec("generatedCommitmentCount", "自动生成承诺文件数", ("承诺函", "承诺书", "不得存在下列情形")),
    FieldSpec("pendingCommitmentCount", "待确认承诺线索数", ("承诺",)),
    FieldSpec("disqualificationCommitmentRequired", "不得存在下列情形承诺要求", ("不得存在下列情形", "不得存在下列情形之一")),
    FieldSpec("otherCommitmentSectionRequired", "其他承诺章节要求", ("投标人需要说明的其他内容", "其他内容", "其他承诺")),
    FieldSpec("commitmentGenerationBasis", "承诺文件生成依据", ("承诺函", "承诺书", "不得存在下列情形")),
)

QUALIFICATION_SECTION_ANCHORS = (
    "投标人资格要求",
    "投标人资格条件",
    "资格能力要求",
    "专用资格条件",
    "通用资格条件",
    "合格投标人资格",
    "供应商资格要求",
    "资质条件、能力和信誉",
)

QUALIFICATION_STOP_ANCHORS = (
    "投标文件的组成",
    "投标报价",
    "投标保证金",
    "投标人须知",
    "资格审查资料",
    "评标办法",
    "符合性审查",
    "商务评分",
    "技术评分",
    "投标文件格式",
    "合同条款",
)

QUALIFICATION_EXCLUDE_KEYWORDS = (
    "资格审查资料",
    "复印件",
    "扫描件",
    "附件",
    "评分",
    "得分",
    "分值",
    "满分",
    "基础分",
    "加分",
    "否决",
    "废标",
    "不予受理",
    "目录",
    "页码",
    "见投标人须知前附表",
    "见评标办法前附表",
    "同招标公告",
)

QUALIFICATION_REQUIRED_CUES = (
    "投标人",
    "投标机型",
    "供应商",
    "联合体",
    "须",
    "应",
    "需",
    "具有",
    "具备",
    "不得",
    "不允许",
    "不接受",
    "没有处于",
    "未被",
)

SCOPE_PATTERN = re.compile(
    r"^(?:"
    r"标段[一二三四五六七八九十\d]+(?:[、至和及,\-]+标段?[一二三四五六七八九十\d]+)*"
    r"|第[一二三四五六七八九十\d]+标段"
    r"|所有标段"
    r"|全部标段"
    r"|本项目"
    r")(?:（[^）]*）)?[:：]?$"
)
CLAUSE_PATTERN = re.compile(r"^(?:\d+(?:\.\d+){1,4}|[（(][一二三四五六七八九十\d]+[）)]|[一二三四五六七八九十\d]+[、.．])\s*")

BUSINESS_CORE_PROJECT_FIELDS: tuple[FieldSpec, ...] = (
    FieldSpec("projectName", "项目名称", ("项目名称", "招标项目名称", "采购项目名称")),
    FieldSpec("tenderNo", "招标编号", ("招标编号", "项目编号", "招标文件编号", "采购编号", "采购项目编号")),
    FieldSpec("tenderer", "招标人", ("招标人", "采购人", "采购单位", "业主", "建设单位", "项目单位")),
    FieldSpec("tenderAgency", "招标代理机构", ("招标代理机构", "采购代理机构", "代理机构")),
    FieldSpec("bidDeadline", "递交截止时间", ("递交截止时间", "投标截止时间", "投标文件递交截止时间", "提交截止时间", "响应文件提交截止时间", "响应截止时间", "提交响应文件截止时间")),
)

COMMERCIAL_REJECTION_KEYWORDS = (
    "否决",
    "废标",
    "无效投标",
    "不予受理",
    "★",
    "实质性响应",
    "投标人不得存在",
    "不得存在下列情形",
)

SCORING_SCORE_PATTERN = re.compile(r"(?P<item>[\u4e00-\u9fa5A-Za-z0-9（）()、/]+?)\s*(?P<score>\d+(?:\.\d+)?\s*分)")
MARKDOWN_TABLE_LINE_PATTERN = re.compile(r"^\s*\|.*\|\s*$")
BIDDER_INSTRUCTION_TABLE_TITLE = "投标人须知前附表"


def parsed_project_dir(project_id: str) -> Path:
    path = settings.parsed_dir / project_id
    path.mkdir(parents=True, exist_ok=True)
    return path


def parsed_appendix_dir(project_id: str) -> Path:
    path = settings.parsed_dir / project_id / "s1_appendices"
    path.mkdir(parents=True, exist_ok=True)
    return path


def parsed_appendix_path(project_id: str) -> Path:
    return settings.parsed_dir / project_id / "s1_appendices"


def _normalize_text(raw_text: str) -> str:
    lines = [line.rstrip() for line in raw_text.replace("\r\n", "\n").split("\n")]
    compact: list[str] = []
    previous_blank = False
    for line in lines:
        blank = not line.strip()
        if blank and previous_blank:
            continue
        compact.append(line)
        previous_blank = blank
    return "\n".join(compact).strip()


def extract_docx_text(path: Path) -> str:
    pieces: list[str] = []
    with zipfile.ZipFile(path) as archive:
        with archive.open("word/document.xml") as xml_file:
            for _, element in ET.iterparse(xml_file, events=("end",)):
                if element.tag == f"{WORD_NAMESPACE}t":
                    pieces.append(element.text or "")
                elif element.tag == f"{WORD_NAMESPACE}tab":
                    pieces.append("\t")
                elif element.tag in {f"{WORD_NAMESPACE}br", f"{WORD_NAMESPACE}cr"}:
                    pieces.append("\n")
                elif element.tag == f"{WORD_NAMESPACE}p":
                    pieces.append("\n")
                    element.clear()
    return _normalize_text("".join(pieces))


def extract_pdf_text(path: Path) -> tuple[str, dict[str, Any]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("未安装 pypdf，当前无法解析 PDF。") from exc

    reader = PdfReader(str(path))
    page_count = len(reader.pages)
    empty_pages = 0
    page_texts: list[str] = []
    for page in reader.pages:
        text = (page.extract_text() or "").strip()
        if not text:
            empty_pages += 1
        page_texts.append(text)

    warnings: list[str] = []
    if page_count and empty_pages == page_count:
        warnings.append("PDF 未提取到文本，疑似扫描件，已进入 OCR 兜底识别。")
    elif empty_pages:
        warnings.append(f"PDF 有 {empty_pages} 页未提取到文本。")

    return _normalize_text("\n\n".join(page_texts)), {
        "pageCount": page_count,
        "warnings": warnings,
        "requiresOcr": bool(page_count and empty_pages == page_count),
    }


def _run_async_ocr(coro: Any) -> Any:
    try:
        asyncio.get_running_loop()
    except RuntimeError:
        return asyncio.run(coro)

    result: Any = None
    error: BaseException | None = None

    def run() -> None:
        nonlocal result, error
        try:
            result = asyncio.run(coro)
        except BaseException as exc:  # pragma: no cover - defensive bridge
            error = exc

    thread = threading.Thread(target=run, daemon=True, name="parse-visual-recognition")
    thread.start()
    thread.join()
    if error:
        raise error
    return result


def _ocr_fallback_text(project_id: str, file_record: dict[str, Any], file_path: Path) -> tuple[str, dict[str, Any]]:
    _ = project_id
    try:
        text, raw = _run_async_ocr(
            ocr_service.recognize_text_for_parse(
                file_name=str(file_record.get("name") or file_path.name),
                content=file_path.read_bytes(),
                mime_type=str(file_record.get("content_type") or mimetypes.guess_type(file_path.name)[0] or ""),
            )
        )
        return _normalize_text(text), {
            "status": "completed",
            "pageCount": raw.get("pageCount") or "-",
        }
    except PeripheralError as exc:
        return "", {
            "status": "failed",
            "code": exc.code,
            "message": exc.detail,
        }
    except Exception as exc:
        return "", {
            "status": "failed",
            "code": "OCR_PARSE_FALLBACK_FAILED",
            "message": str(exc),
        }


def _normalize_date_match(match: re.Match[str]) -> str:
    try:
        parsed = date(
            int(match.group("year")),
            int(match.group("month")),
            int(match.group("day")),
        )
    except ValueError:
        return ""
    return parsed.isoformat()


def _find_dates(text: str) -> list[str]:
    dates: list[str] = []
    for match in DATE_PATTERN.finditer(text):
        normalized = _normalize_date_match(match)
        if normalized:
            dates.append(normalized)
    return dates


def _strip_leading_number(text: str) -> str:
    return LEADING_NUMBER_PATTERN.sub("", text).strip()


def _split_label_value(line: str, fallback_label: str) -> tuple[str, str]:
    normalized = _strip_leading_number(line)
    match = LABEL_VALUE_PATTERN.match(normalized)
    if match:
        label = _strip_leading_number(match.group("label")).strip()
        value = match.group("value").strip(" ；;。")
        return label or fallback_label, value or normalized
    return fallback_label, normalized


def _first_matching_category(line: str) -> ParseCategory | None:
    if "考核" in line:
        return next(category for category in PARSE_CATEGORIES if category.key == "assessment_terms")
    for category in PARSE_CATEGORIES:
        if any(keyword in line for keyword in category.keywords):
            return category
    return None


def _record_project_date(line: str, dates: list[str], project_dates: dict[str, str]) -> None:
    if not dates:
        return
    has_start_keyword = any(keyword in line for keyword in START_DATE_KEYWORDS)
    has_end_keyword = any(keyword in line for keyword in END_DATE_KEYWORDS)
    has_range_hint = any(token in line for token in ("至", "到", "止", "~", "—", "-"))

    if has_start_keyword and not project_dates["startDate"]:
        project_dates["startDate"] = dates[0]
    if has_end_keyword and not project_dates["endDate"]:
        project_dates["endDate"] = dates[-1]
    if has_range_hint and len(dates) >= 2:
        if not project_dates["startDate"]:
            project_dates["startDate"] = dates[0]
        if not project_dates["endDate"]:
            project_dates["endDate"] = dates[-1]


def _empty_field(spec: FieldSpec) -> dict[str, Any]:
    return {
        "key": spec.key,
        "label": spec.label,
        "value": "",
        "status": "missing",
        "sourceFile": "",
        "evidence": "",
        "evidenceLocation": "",
    }


def _field_from_item(spec: FieldSpec, item: dict[str, Any]) -> dict[str, Any]:
    return {
        "key": spec.key,
        "label": spec.label,
        "value": str(item.get("value") or item.get("keyValue") or "").strip(),
        "status": "found",
        "sourceFile": str(item.get("sourceFile") or ""),
        "evidence": str(item.get("evidence") or ""),
        "evidenceLocation": str(item.get("evidenceLocation") or ""),
    }


def _build_fixed_field_group(items: list[dict[str, Any]], specs: tuple[FieldSpec, ...]) -> list[dict[str, Any]]:
    fields: list[dict[str, Any]] = []
    for spec in specs:
        matched = next(
            (
                item
                for item in items
                if any(alias in str(item.get("keyEntity") or item.get("title") or item.get("evidence") or "") for alias in spec.aliases)
            ),
            None,
        )
        fields.append(_field_from_item(spec, matched) if matched else _empty_field(spec))
    return fields


def _build_scoring_fields(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    scoring_items = [item for item in items if item.get("category") == "scoring_criteria"]
    rows: list[dict[str, Any]] = []
    for index, item in enumerate(scoring_items, start=1):
        evidence = str(item.get("evidence") or item.get("value") or "")
        matches = list(SCORING_SCORE_PATTERN.finditer(evidence))
        if matches:
            for match_index, match in enumerate(matches, start=1):
                next_match = matches[match_index] if match_index < len(matches) else None
                segment = evidence[match.start() : next_match.start() if next_match else len(evidence)]
                segment = segment.strip(" ；;。")
                proof_requirement = segment if any(keyword in segment for keyword in ("证明", "材料", "提供")) else ""
                rows.append(
                    {
                        "id": f"{item.get('id') or f'SCORE-{index}'}-{match_index}",
                        "scoringItem": match.group("item").strip(" ，,：:；;"),
                        "score": match.group("score").replace(" ", ""),
                        "scorePoint": segment or str(item.get("value") or evidence),
                        "proofRequirement": proof_requirement,
                        "sourceFile": item.get("sourceFile") or "",
                        "evidence": evidence,
                        "evidenceLocation": item.get("evidenceLocation") or "",
                    }
                )
            continue

        scoring_item = str(item.get("keyEntity") or item.get("title") or f"评分项{index}")
        proof_requirement = evidence if "证明" in evidence or "材料" in evidence else ""
        rows.append(
            {
                "id": item.get("id") or f"SCORE-{index}",
                "scoringItem": scoring_item,
                "score": "",
                "scorePoint": str(item.get("value") or evidence),
                "proofRequirement": proof_requirement,
                "sourceFile": item.get("sourceFile") or "",
                "evidence": evidence,
                "evidenceLocation": item.get("evidenceLocation") or "",
            }
        )
    return rows


def _presence_from_items(
    items: list[dict[str, Any]],
    category: str,
    *,
    keywords: tuple[str, ...] = (),
) -> dict[str, Any]:
    matched = [item for item in items if item.get("category") == category]
    if keywords:
        matched = [
            item
            for item in matched
            if any(
                keyword in str(item.get("title") or item.get("value") or item.get("evidence") or "")
                for keyword in keywords
            )
        ]
    if not matched:
        return {
            "status": "missing",
            "summary": "招标文件中暂未识别到明确要求。",
            "evidences": [],
        }
    evidences = [
        {
            "sourceFile": item.get("sourceFile") or "",
            "evidence": item.get("evidence") or item.get("value") or "",
            "evidenceLocation": item.get("evidenceLocation") or "",
        }
        for item in matched[:5]
    ]
    summary_parts = [str(item.get("value") or item.get("evidence") or "").strip() for item in matched]
    return {
        "status": "present",
        "summary": "；".join(part for part in summary_parts if part)[:500],
        "evidences": evidences,
    }


def _build_requirement_presence(items: list[dict[str, Any]]) -> dict[str, Any]:
    return {
        "topicPlans": _presence_from_items(items, "topic_plans"),
        "supplyScope": _presence_from_items(
            items,
            "tables_and_scope",
            keywords=("供货范围", "供货清单", "供货界面", "供货"),
        ),
        "assessmentTerms": _presence_from_items(items, "assessment_terms"),
    }


def _build_field_groups(items: list[dict[str, Any]]) -> dict[str, Any]:
    return {
        "scoringCriteria": _build_scoring_fields(items),
        "projectBasics": _build_fixed_field_group(items, PROJECT_BASIC_FIELDS),
        "turbineCoreParameters": _build_fixed_field_group(items, TURBINE_CORE_FIELDS),
        "performanceGuarantees": _build_fixed_field_group(items, PERFORMANCE_FIELDS),
        "environmentAdaptation": _build_fixed_field_group(items, ENVIRONMENT_FIELDS),
    }


def _copy_meta_fields(item: dict[str, Any]) -> dict[str, Any]:
    return {
        "sourceFile": str(item.get("sourceFile") or ""),
        "sourceDocumentId": str(item.get("sourceDocumentId") or ""),
        "section": str(item.get("section") or ""),
        "evidence": str(item.get("evidence") or ""),
        "evidenceLocation": str(item.get("evidenceLocation") or ""),
    }


def _qualification_source_text(*, source_file: str, section: str, clause_no: str = "") -> str:
    parts = [part.strip(" ：:") for part in (section, clause_no) if str(part or "").strip()]
    readable = " > ".join(dict.fromkeys(parts))
    if readable:
        return f"{source_file}：{readable}"
    return source_file


def _business_field_from_item(spec: FieldSpec, item: dict[str, Any], *, value_override: str | None = None) -> dict[str, Any]:
    field = {
        "key": spec.key,
        "label": spec.label,
        "value": (value_override if value_override is not None else str(item.get("value") or item.get("keyValue") or "")).strip(),
        "status": "found",
        **_copy_meta_fields(item),
        "confidence": float(item.get("confidence") or 0.86),
    }
    return field


def _empty_business_field(spec: FieldSpec, *, value: str = "") -> dict[str, Any]:
    return {
        "key": spec.key,
        "label": spec.label,
        "value": value,
        "status": "missing" if not value else "derived",
        "sourceFile": "",
        "sourceDocumentId": "",
        "section": "",
        "evidence": "",
        "evidenceLocation": "",
        "confidence": 0.0,
    }


def _find_business_item(items: list[dict[str, Any]], spec: FieldSpec) -> dict[str, Any] | None:
    for item in items:
        haystack = " ".join(
            str(item.get(key) or "")
            for key in ("title", "keyEntity", "value", "evidence", "section")
        )
        if any(alias in haystack for alias in spec.aliases):
            return item
    return None


def _clean(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "").replace("\u3000", " ")).strip()


def _is_reference_only_value(value: str) -> bool:
    normalized = re.sub(r"\s+", "", str(value or ""))
    if not normalized:
        return False
    if normalized in {
        "见招标公告",
        "见采购公告",
        "详见采购公告",
        "见投标人须知前附表",
        "详见招标公告",
        "详见技术规范书",
        "详见招标文件",
        "按招标文件要求",
    }:
        return True
    return bool(re.match(r"^(?:详见|见|参见|按|同)", normalized)) and any(
        token in normalized
        for token in (
            "招标公告",
            "采购公告",
            "投标人须知前附表",
            "供应商须知前附表",
            "招标文件",
            "采购文件",
            "技术规范书",
        )
    )


def _is_bid_deadline_relative_context(text: str) -> bool:
    compact = re.sub(r"\s+", "", str(text or ""))
    return bool(
        re.search(r"投标截止时间\d{1,3}(?:日|天)前", compact)
        or re.search(r"收到(?:澄清|修改)后\d{1,3}小时内", compact)
        or re.search(r"开标结束后\d{1,3}分钟内", compact)
    )


def _business_project_value_usable(spec: FieldSpec, value: str, evidence: str = "") -> bool:
    cleaned = _clean(value).strip(" ：:；;，,。")
    if not cleaned or _is_reference_only_value(cleaned):
        return False
    if spec.key == "bidDeadline":
        normalized = _normalize_bid_deadline(cleaned)
        combined = f"{evidence} {cleaned}"
        is_opening_time = "开标" in combined and not any(token in combined for token in ("递交截止", "投标截止", "提交截止", "响应截止"))
        return _is_normalized_bid_deadline(normalized) and not is_opening_time and not _is_bid_deadline_relative_context(combined)
    if spec.key in {"tenderer", "tenderAgency"}:
        if len(cleaned) > 100:
            return False
        if any(keyword in cleaned for keyword in ("联系人", "联系方式", "联系电话", "电话", "邮箱", "电子邮件", "地址")):
            return False
        combined = f"{evidence} {cleaned}"
        if any(keyword in combined for keyword in ("招标人代表", "采购人代表", "异议", "投诉", "质疑", "服务费", "代理服务费", "招标人不接受", "采购人不接受")):
            return False
        if re.search(r"(?:现)?委托.+(?:招标|采购|代理)", cleaned):
            return False
        if "，" in cleaned and any(token in cleaned for token in ("进行公开招标", "进行采购", "项目业主为")):
            return False
    if spec.key == "projectName" and len(cleaned) > 160:
        return False
    return True


def _block_number(location: str) -> int:
    match = re.match(r"B(\d+)", str(location or ""))
    return int(match.group(1)) if match else 0


def _business_core_field_score(item: dict[str, Any], spec: FieldSpec) -> int:
    value = str(item.get("value") or item.get("keyValue") or "").strip()
    section = str(item.get("section") or "")
    evidence = str(item.get("evidence") or "")
    location = str(item.get("evidenceLocation") or "")
    block_no = _block_number(location)
    score = 0
    if location.startswith("B"):
        score += 40
        if 0 < block_no <= 30:
            score += 25
    if section == "封面":
        score += 90
    if "投标人须知前附表" in section:
        score += 80
    if "招标公告" in section or "联系方式" in section:
        score += 50
    if _is_reference_only_value(value):
        score -= 220
    if re.search(r"\d{4}-\d{2}-\d{2}|20\d{2}年", value):
        score += 30
    if spec.key == "bidDeadline":
        haystack = " ".join(str(item.get(key) or "") for key in ("title", "keyEntity", "value", "evidence", "section"))
        if any(token in haystack for token in BID_DEADLINE_CONTEXT):
            score += 120
        if any(token in haystack for token in OPENING_TIME_CONTEXT):
            score -= 180
        if _is_normalized_bid_deadline(_normalize_bid_deadline(" ".join(part for part in (value, evidence) if part))):
            score += 120
        else:
            score -= 260
    if spec.key == "projectName":
        if "项目" in value and len(value) <= 120:
            score += 45
        if value.endswith("招标") or "采购" in value:
            score += 20
    if spec.key == "tenderNo":
        if re.search(r"[A-Z]{2,}.*\d", value):
            score += 70
        if value in {"招标编号", "项目编号", "招标文件编号"}:
            score -= 100
    if spec.key in {"tenderer", "tenderAgency"}:
        if len(value) <= 100:
            score += 35
        if "：" in value or ":" in value:
            score -= 10
    if len(value) > 180:
        score -= 50
    if not value:
        score -= 300
    if "目录" in evidence:
        score -= 60
    return score


def _normalize_bid_deadline(value: str) -> str:
    text = str(value or "").strip()
    matches = list(BID_DEADLINE_DATE_PATTERN.finditer(text))
    if not matches:
        return text
    first_date = ""
    for match in matches:
        try:
            parsed = date(
                int(match.group("year")),
                int(match.group("month")),
                int(match.group("day")),
            )
        except ValueError:
            continue
        normalized = parsed.isoformat()
        if not first_date:
            first_date = normalized
        hour = match.group("hour")
        minute = match.group("minute")
        if hour is None or minute is None:
            continue
        hour_int = int(hour)
        minute_int = int(minute)
        if hour_int > 23 or minute_int > 59:
            continue
        return f"{normalized} {hour_int:02d}:{minute_int:02d}"
    return first_date or text


def _is_normalized_bid_deadline(value: str) -> bool:
    return bool(re.fullmatch(r"20\d{2}-\d{2}-\d{2}(?: \d{2}:\d{2})?", str(value or "").strip()))


def _strip_core_party_contact_tail(value: str) -> str:
    text = _clean(value).strip(" ：:")
    tail_match = re.search(
        r"\s*(?:地\s*址|联\s*系\s*人|电\s*话|电子\s*邮\s*件|邮\s*箱|传\s*真|网\s*址)\s*[：:]",
        text,
    )
    if tail_match:
        text = text[: tail_match.start()]
    return text.strip(" ：:")


def _normalize_core_field_value(spec: FieldSpec, value: str) -> str:
    cleaned = _clean(value)
    if spec.key == "bidDeadline":
        return _normalize_bid_deadline(cleaned)
    stripped = cleaned
    for alias in sorted(spec.aliases, key=len, reverse=True):
        alias_pattern = r"\s*".join(re.escape(char) for char in alias)
        stripped_candidate = re.sub(rf"^\s*{alias_pattern}\s*[：:]\s*", "", cleaned, count=1)
        if stripped_candidate != cleaned:
            stripped = stripped_candidate.strip()
            break
    if spec.key in {"tenderer", "tenderAgency"}:
        return _strip_core_party_contact_tail(stripped)
    return stripped


def _build_business_project_basics(
    items: list[dict[str, Any]],
    project_dates: dict[str, Any] | None = None,
) -> list[dict[str, Any]]:
    project_dates = project_dates or {}
    fields: list[dict[str, Any]] = []
    for spec in BUSINESS_CORE_PROJECT_FIELDS:
        project_deadline_value = ""
        if spec.key == "bidDeadline":
            project_deadline_value = _normalize_bid_deadline(str(project_dates.get("endDate") or ""))
        candidates = [
            item
            for item in items
            if any(
                alias in " ".join(str(item.get(key) or "") for key in ("title", "keyEntity", "evidence", "section"))
                for alias in spec.aliases
            )
        ]
        candidates = [
            item
            for item in candidates
            if _business_project_value_usable(
                spec,
                str(item.get("value") or item.get("keyValue") or ""),
                str(item.get("evidence") or ""),
            )
        ]
        if spec.key == "bidDeadline":
            candidates = [
                item
                for item in candidates
                if (
                    any(
                        token in " ".join(str(item.get(key) or "") for key in ("title", "keyEntity", "evidence", "section", "value"))
                        for token in BID_DEADLINE_CONTEXT
                    )
                    and not any(
                        token in " ".join(str(item.get(key) or "") for key in ("title", "keyEntity", "evidence", "section"))
                        for token in OPENING_TIME_CONTEXT
                    )
                    and _is_normalized_bid_deadline(
                        _normalize_bid_deadline(
                            " ".join(
                                str(item.get(key) or "")
                                for key in ("value", "keyValue", "evidence")
                                if str(item.get(key) or "").strip()
                            )
                        )
                    )
                )
            ]
        matched = max(candidates, key=lambda item: _business_core_field_score(item, spec)) if candidates else None
        if matched and _business_core_field_score(matched, spec) > -50:
            if spec.key == "bidDeadline":
                normalized = _normalize_core_field_value(
                    spec,
                    " ".join(
                        str(matched.get(key) or "")
                        for key in ("value", "keyValue", "evidence")
                        if str(matched.get(key) or "").strip()
                    ),
                )
                project_dates["endDate"] = normalized
            else:
                normalized = _normalize_core_field_value(spec, str(matched.get("value") or matched.get("keyValue") or ""))
            fields.append(_business_field_from_item(spec, matched, value_override=normalized))
        elif spec.key == "bidDeadline" and project_deadline_value:
            project_dates["endDate"] = project_deadline_value
            field = _empty_business_field(spec, value=project_deadline_value)
            field["status"] = "found"
            field["confidence"] = 0.78
            fields.append(field)
        else:
            fields.append(_empty_business_field(spec))
    return fields


def _build_business_response_fields(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    fields: list[dict[str, Any]] = []
    for spec in BUSINESS_RESPONSE_FIELDS:
        matched = _find_business_item(items, spec)
        fields.append(_business_field_from_item(spec, matched) if matched else _empty_business_field(spec))
    return fields


def _build_qualification_support_fields(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    fields: list[dict[str, Any]] = []
    for spec in QUALIFICATION_SUPPORT_FIELDS:
        matched = _find_business_item(items, spec)
        fields.append(_business_field_from_item(spec, matched) if matched else _empty_business_field(spec))
    return fields


def _new_docx_candidate_item(
    items: list[dict[str, Any]],
    *,
    document: dict[str, Any],
    label: str,
    value: str,
    section: str,
    evidence: str,
    location: str,
    confidence: float = 0.86,
) -> None:
    cleaned_label = _clean(label)
    cleaned_value = _clean(value)
    if not cleaned_label or not cleaned_value:
        return
    items.append(
        {
            "id": f"DOCX-{len(items) + 1:04d}",
            "type": "商务核心字段候选",
            "category": "business_core_candidate",
            "title": cleaned_label,
            "keyEntity": cleaned_label,
            "keyValue": cleaned_value,
            "value": cleaned_value,
            "sourceFile": str(document.get("name") or document.get("id") or "招标文件"),
            "sourceDocumentId": str(document.get("id") or ""),
            "section": section,
            "evidence": evidence,
            "evidenceLocation": location,
            "confidence": confidence,
        }
    )


def _is_docx_source(path: Path) -> bool:
    return path.is_file() and path.suffix.lower() == ".docx"


def _document_text_lines(document: dict[str, Any], texts_by_id: dict[str, str]) -> list[dict[str, Any]]:
    document_id = str(document.get("id") or "")
    source_file = str(document.get("name") or document_id or "招标文件")
    text = str(texts_by_id.get(document_id) or "")
    lines: list[dict[str, Any]] = []
    for line_number, raw_line in enumerate(text.splitlines(), start=1):
        line = _clean(raw_line)
        if not line:
            continue
        lines.append(
            {
                "text": line,
                "sourceFile": source_file,
                "sourceDocumentId": document_id,
                "evidenceLocation": f"L{line_number}",
            }
        )
    return lines


def _qualification_heading_level(text: str) -> int:
    stripped = str(text or "").strip()
    match = re.match(r"^(\d+(?:\.\d+){0,4})\s+", stripped)
    if not match:
        return 99
    return match.group(1).count(".") + 1


def _is_qualification_anchor(text: str) -> bool:
    return any(anchor in text for anchor in QUALIFICATION_SECTION_ANCHORS)


def _is_qualification_stop(text: str, active_root_level: int) -> bool:
    if not text:
        return False
    if any(anchor in text for anchor in QUALIFICATION_STOP_ANCHORS):
        return True
    level = _qualification_heading_level(text)
    return level <= active_root_level and not _is_qualification_anchor(text)


def _normalize_qualification_content(text: str) -> str:
    value = _clean(text)
    value = re.sub(r"^\d+(?:\.\d+){1,4}\s*", "", value)
    value = re.sub(r"^[（(][一二三四五六七八九十\d]+[）)]\s*", "", value)
    value = re.sub(r"^[一二三四五六七八九十\d]+[、.．]\s*", "", value)
    value = value.strip(" ：:；;。")
    return value


def _looks_like_scope_heading(text: str) -> bool:
    return bool(SCOPE_PATTERN.match(str(text or "").strip()))


def _looks_like_qualification_intro_line(text: str) -> bool:
    value = _normalize_qualification_content(text)
    return bool(
        re.search(r"(?:下列|如下|以下)(?:条件|要求|规定)$", value)
        or re.search(r"(?:应|需)(?:具备|满足|符合).*(?:下列|如下|以下)(?:条件|要求|规定)$", value)
    )


def _normalize_qualification_scope(text: str) -> str:
    value = str(text or "").strip()
    value = re.split(r"[（(]", value, maxsplit=1)[0]
    return value.strip(" ：:")


def _looks_like_qualification_requirement(text: str) -> bool:
    value = _normalize_qualification_content(text)
    if len(value) < 8:
        return False
    if _looks_like_scope_heading(value):
        return False
    if _looks_like_qualification_intro_line(value):
        return False
    compact = re.sub(r"\s+", "", value)
    if re.search(r"\t\d+$|\.{3,}\d+$", value):
        return False
    if any(keyword in value for keyword in QUALIFICATION_EXCLUDE_KEYWORDS):
        return False
    if compact in {"见评标办法前附表", "见投标人须知前附表", "同招标公告"}:
        return False
    return any(cue in value for cue in QUALIFICATION_REQUIRED_CUES)


def _extract_qualification_requirements_from_documents(
    documents: list[dict[str, Any]],
    texts_by_id: dict[str, str],
) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    seen: set[tuple[str, str, str]] = set()

    for document in documents:
        source_file = str(document.get("name") or document.get("id") or "招标文件")
        document_id = str(document.get("id") or "")
        active = False
        active_root_level = 99
        section_path: list[tuple[int, str]] = []
        applicable_scope = "全部标段"

        for line in _document_text_lines(document, texts_by_id):
            text = str(line["text"])
            level = _qualification_heading_level(text)

            if _is_qualification_anchor(text):
                was_active = active
                active = True
                if level < 99 and (not was_active or active_root_level == 99):
                    active_root_level = level
                if level < 99:
                    section_path = [(old_level, title) for old_level, title in section_path if old_level < level]
                    section_path.append((level, text))
                elif not section_path:
                    section_path = [(1, text)]
                continue

            if not active:
                continue

            if _is_qualification_stop(text, active_root_level):
                active = False
                section_path = []
                applicable_scope = "全部标段"
                continue

            if level < 99:
                section_path = [(old_level, title) for old_level, title in section_path if old_level < level]
                section_path.append((level, text))
                applicable_scope = "全部标段"

            if _looks_like_scope_heading(text):
                applicable_scope = _normalize_qualification_scope(text)
                continue

            if not _looks_like_qualification_requirement(text):
                continue

            content = _normalize_qualification_content(text)
            section = " > ".join(title for _, title in section_path) or "投标人资格要求"
            clause_no_match = re.match(r"^(\d+(?:\.\d+){1,4}|[（(][一二三四五六七八九十\d]+[）)])", text)
            clause_no = clause_no_match.group(1) if clause_no_match else ""
            dedupe_key = (content, applicable_scope, section)
            if dedupe_key in seen:
                continue
            seen.add(dedupe_key)
            rows.append(
                {
                    "id": f"QUAL-{len(rows) + 1:04d}",
                    "order": len(rows) + 1,
                    "content": content,
                    "applicableScope": applicable_scope or "全部标段",
                    "sourceText": _qualification_source_text(
                        source_file=source_file,
                        section=section,
                        clause_no=clause_no,
                    ),
                    "sourceFile": source_file,
                    "sourceDocumentId": document_id,
                    "section": section,
                    "evidence": text,
                    "evidenceLocation": str(line.get("evidenceLocation") or ""),
                    "confidence": 0.9,
                }
            )

    return rows


def _extract_docx_core_candidate_items(documents: list[dict[str, Any]]) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    for document in documents:
        source_path = Path(str(document.get("sourcePath") or ""))
        if not _is_docx_source(source_path):
            continue
        blocks = _iter_docx_blocks(source_path)
        current_section = "封面"
        for block_index, block in enumerate(blocks, start=1):
            if block.get("type") == "paragraph":
                text = _clean(block.get("text"))
                if text:
                    current_section = text
                continue
            if block.get("type") != "table":
                continue
            rows = block.get("rows") or []
            section = current_section or ("封面" if block_index <= 30 else "")
            for row_index, row in enumerate(rows, start=1):
                cells = [_clean(cell) for cell in row if _clean(cell)]
                if len(cells) < 2:
                    continue
                _new_docx_candidate_item(
                    items,
                    document=document,
                    label=cells[0].strip(" ：:"),
                    value=cells[-1].strip(" ：:"),
                    section=section,
                    evidence=" | ".join(cells),
                    location=f"B{block_index}/R{row_index}",
                    confidence=0.9 if block_index <= 30 or "投标人须知前附表" in section else 0.82,
                )
    return items


def _strip_bidder_instruction_title_prefix(text: str) -> str:
    cleaned = _clean(text).strip("# ").strip(" ：:。；;")
    cleaned = re.sub(r"^\s*第?[一二三四五六七八九十百千0-9]+[章节条]?\s*[、.．\s]+", "", cleaned)
    return re.sub(r"\s+", "", cleaned)


def _is_bidder_instruction_title_anchor(text: str) -> bool:
    compact = _strip_bidder_instruction_title_prefix(text)
    if not compact.startswith(BIDDER_INSTRUCTION_TABLE_TITLE):
        return False
    suffix = compact.removeprefix(BIDDER_INSTRUCTION_TABLE_TITLE)
    return bool(re.fullmatch(r"(?:[0-9一二三四五六七八九十百千页（）()、.．:-]*)", suffix))


def _docx_table_after_bidder_instruction_anchor(
    blocks: list[dict[str, Any]],
    anchor_index: int,
    *,
    max_text_gap: int,
) -> tuple[int, list[list[str]]] | None:
    gap = 0
    index = anchor_index + 1
    while index < len(blocks):
        block = blocks[index]
        if block.get("type") == "table":
            rows = [[_clean(cell) for cell in row] for row in block.get("rows") or [] if any(_clean(cell) for cell in row)]
            return (index, rows) if gap <= max_text_gap and len(rows) > 1 else None
        text = _clean(block.get("text")) if block.get("type") == "paragraph" else ""
        if text:
            gap += 1
            if gap > max_text_gap:
                return None
        index += 1
    return None


def _parse_bidder_instruction_rows(
    rows: list[list[str]],
    *,
    document: dict[str, Any],
    section: str,
    block_index: int,
    table_title: str = BIDDER_INSTRUCTION_TABLE_TITLE,
) -> list[dict[str, Any]]:
    cleaned_rows = [[_clean(cell) for cell in row] for row in rows if any(_clean(cell) for cell in row)]
    if len(cleaned_rows) <= 1:
        return []
    header = cleaned_rows[0]
    parsed: list[dict[str, Any]] = []
    for row_index, row in enumerate(cleaned_rows[1:], start=2):
        if len(row) < 2:
            continue
        clause_no = row[0]
        clause_name = row[1]
        content = "；".join(cell for cell in row[2:] if cell).strip()
        if not clause_no and not clause_name and not content:
            continue
        parsed.append(
            {
                "id": f"BIDDER-INST-{len(parsed) + 1:04d}",
                "clauseNo": clause_no,
                "clauseName": clause_name,
                "content": content,
                "headers": header,
                "cells": row,
                "tableTitle": table_title or BIDDER_INSTRUCTION_TABLE_TITLE,
                "tableLocation": f"B{block_index}",
                "sourceFile": str(document.get("name") or ""),
                "sourceDocumentId": str(document.get("id") or ""),
                "section": section,
                "evidence": "；".join(
                    f"{header[index]}：{cell}" if index < len(header) and header[index] else cell
                    for index, cell in enumerate(row)
                    if cell
                ),
                "evidenceLocation": f"B{block_index}/R{row_index}",
                "confidence": 0.9,
            }
        )
    return parsed


def _extract_bidder_instruction_rows(documents: list[dict[str, Any]]) -> list[dict[str, Any]]:
    for document in documents:
        source_path = Path(str(document.get("sourcePath") or ""))
        if not _is_docx_source(source_path):
            continue
        blocks = _iter_docx_blocks(source_path)
        anchors = [
            (index, _clean(block.get("text")))
            for index, block in enumerate(blocks)
            if block.get("type") == "paragraph" and _is_bidder_instruction_title_anchor(str(block.get("text") or ""))
        ]
        for max_text_gap in (0, 3):
            for anchor_index, anchor_title in anchors:
                match = _docx_table_after_bidder_instruction_anchor(blocks, anchor_index, max_text_gap=max_text_gap)
                if not match:
                    continue
                table_index, rows = match
                return _parse_bidder_instruction_rows(
                    rows,
                    document=document,
                    section=anchor_title or BIDDER_INSTRUCTION_TABLE_TITLE,
                    block_index=table_index + 1,
                    table_title=anchor_title or BIDDER_INSTRUCTION_TABLE_TITLE,
                )
    return []


def _build_qualification_requirements(
    items: list[dict[str, Any]],
    *,
    documents: list[dict[str, Any]] | None = None,
    texts_by_id: dict[str, str] | None = None,
) -> list[dict[str, Any]]:
    if documents is not None and texts_by_id is not None:
        rows = _extract_qualification_requirements_from_documents(documents, texts_by_id)
        if rows:
            return rows

    keywords = ("投标人资格要求", "资格要求", "资格能力要求", "投标人资质条件", "合格投标人")
    matched: list[dict[str, Any]] = []
    seen: set[str] = set()
    for item in items:
        text = " ".join(str(item.get(key) or "") for key in ("title", "keyEntity", "value", "evidence", "section"))
        if not any(keyword in text for keyword in keywords):
            continue
        content = str(item.get("evidence") or item.get("value") or "").strip()
        if not _looks_like_qualification_requirement(content):
            continue
        content = _normalize_qualification_content(content)
        if not content or content in seen:
            continue
        seen.add(content)
        source_file = str(item.get("sourceFile") or "招标文件")
        section = str(item.get("section") or item.get("title") or "投标人资格要求")
        matched.append(
            {
                "id": f"QUAL-{len(matched) + 1:04d}",
                "order": len(matched) + 1,
                "content": content,
                "applicableScope": "全部标段",
                "sourceText": _qualification_source_text(source_file=source_file, section=section),
                **_copy_meta_fields(item),
                "confidence": float(item.get("confidence") or 0.78),
            }
        )
    return matched[:12]


def _extract_commercial_rejection_clauses(
    documents: list[dict[str, Any]],
    texts_by_id: dict[str, str],
) -> list[dict[str, Any]]:
    clauses: list[dict[str, Any]] = []
    seen: set[str] = set()
    for document in documents:
        document_id = str(document.get("id") or "")
        source_file = str(document.get("name") or document_id or "招标文件")
        current_section = ""
        for line_number, raw_line in enumerate(str(texts_by_id.get(document_id) or "").splitlines(), start=1):
            line = _clean(raw_line)
            if not line:
                continue
            if _looks_like_section_heading(line):
                current_section = line
            if not any(keyword in line for keyword in COMMERCIAL_REJECTION_KEYWORDS):
                continue
            if line in seen:
                continue
            seen.add(line)
            matched_keywords = [keyword for keyword in COMMERCIAL_REJECTION_KEYWORDS if keyword in line]
            clauses.append(
                {
                    "id": f"REJECT-{len(clauses) + 1:04d}",
                    "title": current_section or "商务废标项",
                    "content": line,
                    "matchedKeywords": matched_keywords,
                    "riskLevel": "high"
                    if any(keyword in matched_keywords for keyword in ("否决", "废标", "无效投标", "不予受理"))
                    else "medium",
                    "sourceFile": source_file,
                    "sourceDocumentId": document_id,
                    "section": current_section,
                    "evidence": line,
                    "evidenceLocation": f"L{line_number}",
                    "confidence": 0.82,
                }
            )
    return clauses


def _line_has_explicit_commitment_obligation(text: str) -> bool:
    normalized = re.sub(r"\s+", "", str(text or ""))
    if not normalized:
        return False
    if "不得存在下列情形" in normalized:
        return True
    if any(hint in normalized for hint in COMMITMENT_GENERATION_HINTS):
        return True
    has_commitment = "承诺" in normalized
    has_obligation = any(token in normalized for token in ("须", "应", "需", "必须", "无条件"))
    has_doc_action = any(token in normalized for token in ("提供", "提交", "出具", "附", "另附", "递交"))
    has_doc_name = any(token in normalized for token in (*COMMITMENT_DOC_KEYWORDS, "书面承诺", "承诺材料"))
    return has_commitment and has_obligation and (has_doc_action or has_doc_name)


def _find_commitment_items(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    matched: list[dict[str, Any]] = []
    seen: set[str] = set()
    for item in items:
        text = " ".join(str(item.get(key) or "") for key in ("title", "value", "evidence", "section"))
        category = str(item.get("category") or "")
        key_entity = str(item.get("keyEntity") or "")
        title = str(item.get("title") or "")
        evidence = str(item.get("evidence") or "")
        if "承诺" not in text and "不得存在下列情形" not in text:
            continue
        if category == "project_basics" and key_entity == "项目名称":
            continue
        if title == "项目名称" and evidence.startswith("项目名称"):
            continue
        item_id = str(item.get("id") or "")
        if item_id in seen:
            continue
        seen.add(item_id)
        matched.append(item)
    return matched


COMMITMENT_DOC_KEYWORDS = ("承诺函", "承诺书")
COMMITMENT_GENERATION_HINTS = (
    "另附承诺函",
    "另附承诺书",
    "单独提供承诺函",
    "单独提供承诺书",
    "须提供承诺函",
    "须提供承诺书",
    "应提供承诺函",
    "应提供承诺书",
    "应出具承诺函",
    "应出具承诺书",
    "需提供承诺函",
    "需提供承诺书",
    "需提供书面承诺",
    "须出具承诺函",
    "须出具承诺书",
    "须无条件承诺",
)
COMMITMENT_REQUIREMENT_CONTEXT_HINTS = (
    "提供",
    "提交",
    "出具",
    "附",
    "另附",
    "单独",
    "递交",
    "响应",
    "按要求",
    "须",
    "应",
    "需",
    "必须",
)
COMMITMENT_IGNORE_KEYWORDS = (
    "技术承诺",
    "廉洁承诺",
    "廉洁自律承诺",
    "履约承诺",
    "履约保证承诺",
    "投标函",
    "授权委托书",
    "评分标准",
    "评分办法",
    "证明材料",
    "目录",
)
COMMITMENT_NON_REQUIREMENT_TITLE_HINTS = (
    "格式",
    "模板",
    "目录",
    "附件",
    "附录",
    "详见",
    "示例",
    "参考",
)
TECHNICAL_COMMITMENT_KEYWORDS = (
    "等效满负荷小时",
    "满负荷小时",
    "保证年等效",
    "保证小时",
    "发电小时",
    "发电量",
    "上网电量",
    "电量",
    "功率曲线",
    "功率",
    "可利用率",
    "涉网性能",
    "机组",
    "叶轮",
    "轮毂",
    "塔筒",
    "箱变",
    "风机",
    "载荷",
    "噪声",
    "振动",
    "发电性能",
    "性能保证",
    "技术指标",
)
COMMITMENT_TOPIC_TITLE_KEYWORDS: tuple[tuple[str, tuple[str, ...], str], ...] = (
    ("confidentiality", ("保密",), "保密承诺书"),
    ("disqualification", ("不得存在下列情形",), "投标人不存在下列情形之一承诺函"),
    ("certificate_obtainment", ("取得本条", "取得材料", "取得证书", "取得认证", "供货前取得"), "材料取得承诺书"),
    ("delivery", ("交货", "供货周期", "交付"), "交货周期承诺书"),
    ("quality", ("质量", "质保", "售后", "服务"), "质量服务承诺书"),
    ("security", ("投标保证金", "保函", "保证金"), "投标保证金承诺书"),
    ("compliance", ("合规", "守法", "违法", "违规", "信用"), "合规承诺书"),
)
COMMITMENT_TOPIC_KEYWORDS: tuple[tuple[str, tuple[str, ...]], ...] = (
    ("confidentiality", ("保密",)),
    ("compliance", ("合规", "守法", "违法", "违规", "信用")),
    ("security", ("投标保证金", "保函", "保证金")),
    ("disqualification", ("不得存在下列情形",)),
    ("certificate_obtainment", ("取得本条", "取得材料", "取得证书", "取得认证", "供货前取得")),
    ("integrity", ("廉洁",)),
    ("performance_bond", ("履约保证", "履约承诺")),
    ("delivery_commitment", ("交货", "工期", "供货周期")),
    ("quality_commitment", ("质量", "质保", "售后", "服务承诺")),
)
COMMITMENT_SEMANTIC_REVIEW_MAX_ITEMS = 12


def _commitment_text(item: dict[str, Any]) -> str:
    return " ".join(str(item.get(key) or "") for key in ("title", "value", "evidence", "section"))


def _is_technical_commitment_item(item: dict[str, Any]) -> bool:
    normalized = re.sub(r"\s+", "", _commitment_text(item))
    if not normalized:
        return False
    return any(keyword in normalized for keyword in TECHNICAL_COMMITMENT_KEYWORDS)


def _preferred_commitment_title(item: dict[str, Any], trigger_text: str) -> str:
    raw_title = str(item.get("title") or "").strip()
    raw_value = str(item.get("value") or "").strip()
    raw_evidence = str(item.get("evidence") or "").strip()
    source_text = raw_evidence or raw_value or raw_title or trigger_text
    normalized_source = re.sub(r"\s+", "", source_text)

    for _, keywords, preferred_title in COMMITMENT_TOPIC_TITLE_KEYWORDS:
        if any(keyword in normalized_source for keyword in keywords):
            return preferred_title

    title = raw_title or trigger_text or "承诺文件"
    if title in COMMITMENT_DOC_KEYWORDS or title == "承诺":
        for keyword in COMMITMENT_GENERATION_HINTS + COMMITMENT_DOC_KEYWORDS:
            if keyword in source_text:
                prefix = _normalize_commitment_title_prefix(source_text.split(keyword, 1)[0])
                if prefix:
                    return f"{prefix}{keyword}"
                return keyword

    title = _normalize_commitment_title_prefix(title)
    if not any(keyword in title for keyword in COMMITMENT_DOC_KEYWORDS):
        title = f"{title}承诺书"
    return title


def _normalize_commitment_title_by_topic(topic_key: str, title: str, item: dict[str, Any]) -> str:
    if topic_key == "certificate_obtainment":
        return "材料取得承诺书"
    if topic_key == "disqualification":
        return "投标人不存在下列情形之一承诺函"
    return title.strip() or _preferred_commitment_title(item, str(item.get("triggerText") or "承诺"))


def _normalize_commitment_topic(text: str) -> str:
    normalized = re.sub(r"\s+", "", str(text or ""))
    for topic, keywords in COMMITMENT_TOPIC_KEYWORDS:
        if any(keyword in normalized for keyword in keywords):
            return topic
    if "承诺函" in normalized or "承诺书" in normalized:
        return normalized[:80]
    return normalized[:80] or "commitment"


def _normalize_commitment_title_prefix(text: str) -> str:
    normalized = str(text or "").strip(" ：:，,、。.；;（）()")
    for prefix in ("投标人应出具", "投标人应提供", "投标人须提供", "投标人需提供", "投标人须出具", "投标人另附", "应出具", "应提供", "须提供", "需提供", "须出具", "另附"):
        if normalized.startswith(prefix):
            normalized = normalized[len(prefix):].strip(" ：:，,、。.；;（）()")
            break
    return normalized


def _contains_commitment_requirement_context(text: str) -> bool:
    normalized = re.sub(r"\s+", "", str(text or ""))
    return any(hint in normalized for hint in COMMITMENT_REQUIREMENT_CONTEXT_HINTS)


def _looks_like_bare_commitment_title(item: dict[str, Any]) -> bool:
    evidence = str(item.get("evidence") or item.get("value") or item.get("title") or "").strip()
    normalized = re.sub(r"\s+", "", evidence)
    if not normalized:
        return False
    if not any(keyword in normalized for keyword in COMMITMENT_DOC_KEYWORDS):
        return False
    if _contains_commitment_requirement_context(normalized):
        return False
    if any(token in normalized for token in COMMITMENT_NON_REQUIREMENT_TITLE_HINTS):
        return True
    if _looks_like_section_heading(evidence):
        return True
    return len(normalized) <= 18 and normalized.endswith(COMMITMENT_DOC_KEYWORDS)


def _extract_commitment_trigger_phrase(text: str) -> str:
    normalized = str(text or "").strip()
    if "不得存在下列情形" in normalized:
        return "投标人不得存在下列情形之一"
    if "须无条件承诺" in normalized:
        return "须无条件承诺"
    for hint in COMMITMENT_GENERATION_HINTS:
        if hint in normalized:
            return hint
    for keyword in COMMITMENT_DOC_KEYWORDS:
        if keyword in normalized:
            return keyword
    return "承诺"


def _build_commitment_semantic_review_prompt(candidates: list[dict[str, Any]]) -> str:
    records = []
    for item in candidates[:COMMITMENT_SEMANTIC_REVIEW_MAX_ITEMS]:
        records.append(
            {
                "id": str(item.get("semanticReviewId") or item.get("id") or ""),
                "title": str(item.get("title") or ""),
                "section": str(item.get("section") or ""),
                "evidence": str(item.get("evidence") or item.get("value") or ""),
                "sourceFile": str(item.get("sourceFile") or ""),
                "evidenceLocation": str(item.get("evidenceLocation") or ""),
                "contextBefore": str(item.get("contextBefore") or ""),
                "contextAfter": str(item.get("contextAfter") or ""),
                "topicKey": str(item.get("topicKey") or ""),
                "triggerText": str(item.get("triggerText") or ""),
            }
        )

    return (
        "你在做商务标承诺文件语义复核。请只根据输入文本判断，该条是否要求投标人单独形成一份承诺函/承诺书。\n"
        "输出 JSON，格式为：\n"
        "{\n"
        '  "decisions": [\n'
        '    {"id":"RAW-0001","action":"generate|clue|ignore","topicKey":"confidentiality","preferredTitle":"保密承诺书","reason":"一句简短原因"}\n'
        "  ]\n"
        "}\n"
        "判断规则：\n"
        "1. 如果只是章节标题、目录项、模板名、格式名、附件名，不要生成，action=ignore 或 clue。\n"
        "2. 如果语义上明确要求投标人单独提供/提交/出具一份承诺函或承诺书，action=generate。\n"
        "3. 如果存在承诺字样，但看不出是否必须单独成文，action=clue。\n"
        "4. 如果 evidence 只是类似“保密承诺书”这类短标题，必须结合 section、contextBefore、contextAfter 判断；上下文没有明确“提供/提交/出具/另附/单独成文”要求时，不要生成。\n"
        "5. 同一主题如果只是重复标题、重复要求或同一事项的不同表述，最多保留一个 generate，其余用 ignore 或 clue。\n"
        "6. topicKey 尽量归一，例如 confidentiality、compliance、security、delivery_commitment、quality_commitment、disqualification。\n"
        "7. 发电量、功率、满负荷小时数、性能保证、机组技术参数等技术标承诺不要生成商务承诺文件，action=ignore。\n"
        "8. preferredTitle 只在 action=generate 时填写，且应是适合最终文件名的主题名称。\n\n"
        f"候选列表：\n{json.dumps(records, ensure_ascii=False, indent=2)}"
    )


def _review_commitment_candidates_semantically(candidates: list[dict[str, Any]]) -> dict[str, dict[str, Any]]:
    if not candidates:
        return {}
    try:
        result = OpencodeClient().review_business_commitments_with_trace(
            _build_commitment_semantic_review_prompt(candidates)
        )
    except RuntimeError:
        return {}
    decisions = result.get("decisions")
    if not isinstance(decisions, list):
        return {}
    reviewed: dict[str, dict[str, Any]] = {}
    for decision in decisions:
        if not isinstance(decision, dict):
            continue
        item_id = str(decision.get("id") or "").strip()
        if not item_id:
            continue
        reviewed[item_id] = decision
    return reviewed


def _commitment_review_signature(item: dict[str, Any]) -> tuple[str, str, str]:
    return (
        str(item.get("sourceFile") or ""),
        str(item.get("evidenceLocation") or ""),
        re.sub(r"\s+", "", str(item.get("evidence") or item.get("value") or item.get("title") or "")),
    )


def _commitment_decision_for_item(
    item: dict[str, Any],
    reviewed: dict[str, dict[str, Any]],
    reviewed_by_signature: dict[tuple[str, str, str], dict[str, Any]],
) -> tuple[dict[str, Any], bool]:
    for item_id in (str(item.get("semanticReviewId") or ""), str(item.get("id") or "")):
        if item_id and item_id in reviewed:
            return reviewed[item_id], True
    signature = _commitment_review_signature(item)
    if signature in reviewed_by_signature:
        return reviewed_by_signature[signature], True
    return {}, False


def _commitment_decision_action(decision: dict[str, Any], default: str = "clue") -> str:
    action = str(decision.get("action") or default).strip().lower()
    return action if action in {"generate", "clue", "ignore"} else default


def _append_semantic_commitment_candidate(
    item: dict[str, Any],
    *,
    base: dict[str, Any],
    semantic_candidates: list[dict[str, Any]],
    semantic_candidate_ids: set[str],
    semantic_candidate_signatures: set[tuple[str, str, str]],
    force_generate_on_fallback: bool = False,
) -> None:
    item_id = str(item.get("id") or "")
    signature = _commitment_review_signature(item)
    if item_id and (item_id in semantic_candidate_ids or signature in semantic_candidate_signatures):
        return
    if item_id:
        semantic_candidate_ids.add(item_id)
    semantic_candidate_signatures.add(signature)
    semantic_candidates.append(
        {
            **item,
            **base,
            "forceGenerateOnFallback": force_generate_on_fallback,
        }
    )


def _append_commitment_letter(
    letters: list[dict[str, Any]],
    generated_topics: set[str],
    item: dict[str, Any],
    *,
    topic_key: str,
    title: str,
    commitment_type: str,
    risk_flags: list[str],
) -> bool:
    if topic_key in generated_topics:
        return False
    generated_topics.add(topic_key)
    letters.append(
        {
            "id": f"CL-{len(letters) + 1:04d}",
            "artifactType": "commitment_letter",
            "title": title,
            "commitmentType": commitment_type,
            "status": "pending_review",
            **_copy_meta_fields(item),
            "topic": topic_key,
            "topicKey": topic_key,
            "triggerText": str(item.get("triggerText") or "承诺"),
            "triggerContext": str(item.get("triggerContext") or item.get("evidence") or item.get("value") or "").strip(),
            "docxPath": "",
            "workspacePath": "",
            "placementHint": "投标人需要说明的其他内容",
            "needsHumanReview": True,
            "riskFlags": risk_flags,
            "previewType": "onlyoffice",
        }
    )
    return True


def _append_commitment_clue(
    clues: list[dict[str, Any]],
    clue_topics: set[tuple[str, str]],
    item: dict[str, Any],
    *,
    topic_key: str,
    recommended_action: str,
    risk_flags: list[str],
) -> bool:
    trigger_context = str(item.get("triggerContext") or item.get("evidence") or item.get("value") or "").strip()
    clue_key = (topic_key, trigger_context)
    if clue_key in clue_topics:
        return False
    clue_topics.add(clue_key)
    clues.append(
        {
            "id": f"CC-{len(clues) + 1:04d}",
            "artifactType": "commitment_clue",
            "title": str(item.get("title") or item.get("triggerText") or "承诺线索").strip() or "承诺线索",
            "clueType": "pending_manual_review",
            "status": "needs_review",
            **_copy_meta_fields(item),
            "topic": topic_key,
            "topicKey": topic_key,
            "triggerText": str(item.get("triggerText") or "承诺"),
            "triggerContext": trigger_context,
            "recommendedAction": recommended_action,
            "riskFlags": risk_flags,
        }
    )
    return True


def _is_commitment_doc_required(item: dict[str, Any]) -> bool:
    text = _commitment_text(item)
    normalized = re.sub(r"\s+", "", text)
    if "不得存在下列情形" in normalized:
        return True
    if any(hint in normalized for hint in COMMITMENT_GENERATION_HINTS):
        return True
    if _line_has_explicit_commitment_obligation(normalized):
        return True
    if any(keyword in normalized for keyword in COMMITMENT_DOC_KEYWORDS):
        if _looks_like_bare_commitment_title(item):
            return False
        if str(item.get("category") or "") == "business_hint" and str(item.get("section") or "").strip() == str(item.get("evidence") or "").strip():
            return False
        if _contains_commitment_requirement_context(normalized):
            return True
        return False
    if any(keyword in normalized for keyword in COMMITMENT_NON_REQUIREMENT_TITLE_HINTS):
        return False
    if _looks_like_section_heading(str(item.get("evidence") or item.get("title") or "")):
        return False
    if ("投标人" in normalized or "投标方" in normalized) and any(token in normalized for token in ("须承诺", "应承诺", "承诺如下")):
        return True
    return False


def _is_commitment_item_ignored(item: dict[str, Any]) -> bool:
    text = _commitment_text(item)
    normalized = re.sub(r"\s+", "", text)
    if "承诺" not in normalized and "不得存在下列情形" not in normalized:
        return True
    if _is_technical_commitment_item(item):
        return True
    if _is_commitment_doc_required(item):
        return False
    if any(keyword in normalized for keyword in COMMITMENT_IGNORE_KEYWORDS):
        if "不得存在下列情形" not in normalized:
            return True
    if any(token in normalized for token in ("评分", "得分", "分值", "证明材料要求")):
        return True
    return False


def _build_business_commitment_analysis(
    items: list[dict[str, Any]],
    *,
    run_semantic_review: bool = True,
) -> dict[str, list[dict[str, Any]]]:
    all_items = _find_commitment_items(items)
    clues: list[dict[str, Any]] = []
    letters: list[dict[str, Any]] = []
    generated_topics: set[str] = set()
    clue_topics: set[tuple[str, str]] = set()
    semantic_candidates: list[dict[str, Any]] = []
    semantic_candidate_ids: set[str] = set()
    semantic_candidate_signatures: set[tuple[str, str, str]] = set()

    for item in all_items:
        if _is_commitment_item_ignored(item):
            continue
        text = _commitment_text(item)
        topic = _normalize_commitment_topic(text)
        trigger_text = _extract_commitment_trigger_phrase(text)
        topic_key = topic or "commitment"
        base = {
            **_copy_meta_fields(item),
            "topic": topic,
            "topicKey": topic_key,
            "triggerText": trigger_text,
            "triggerContext": str(item.get("evidence") or item.get("value") or "").strip(),
        }
        normalized = re.sub(r"\s+", "", text)
        if topic_key == "disqualification" or "不得存在下列情形" in normalized:
            _append_commitment_letter(
                letters,
                generated_topics,
                {**item, **base},
                topic_key="disqualification",
                title="投标人不存在下列情形之一承诺函",
                commitment_type="disqualification",
                risk_flags=["template_pending", "legal_wording_review_required"],
            )
            continue

        if _is_commitment_doc_required(item):
            _append_semantic_commitment_candidate(
                item,
                base=base,
                semantic_candidates=semantic_candidates,
                semantic_candidate_ids=semantic_candidate_ids,
                semantic_candidate_signatures=semantic_candidate_signatures,
                force_generate_on_fallback=True,
            )
            continue

        if any(keyword in normalized for keyword in COMMITMENT_DOC_KEYWORDS):
            _append_semantic_commitment_candidate(
                item,
                base=base,
                semantic_candidates=semantic_candidates,
                semantic_candidate_ids=semantic_candidate_ids,
                semantic_candidate_signatures=semantic_candidate_signatures,
            )
            continue

        _append_commitment_clue(
            clues,
            clue_topics,
            {**item, **base},
            topic_key=topic_key,
            recommended_action="暂不自动生成，请人工确认是否需要单独承诺函/承诺书。",
            risk_flags=["ambiguous_requirement"],
        )

    reviewed = _review_commitment_candidates_semantically(semantic_candidates) if run_semantic_review else {}
    for index, item in enumerate(semantic_candidates[:COMMITMENT_SEMANTIC_REVIEW_MAX_ITEMS], start=1):
        item["semanticReviewId"] = f"RAW-{index:04d}"
    semantic_candidate_by_id: dict[str, dict[str, Any]] = {}
    for item in semantic_candidates:
        for key in (str(item.get("id") or ""), str(item.get("semanticReviewId") or "")):
            if key:
                semantic_candidate_by_id[key] = item
    reviewed_by_signature: dict[tuple[str, str, str], dict[str, Any]] = {}
    for reviewed_id, decision in reviewed.items():
        source_item = semantic_candidate_by_id.get(reviewed_id)
        if source_item:
            reviewed_by_signature[_commitment_review_signature(source_item)] = decision
    for item in semantic_candidates:
        decision, has_ai_decision = _commitment_decision_for_item(item, reviewed, reviewed_by_signature)
        has_ai_decision = bool(run_semantic_review and has_ai_decision)
        if has_ai_decision:
            action = _commitment_decision_action(decision)
        elif bool(item.get("forceGenerateOnFallback")):
            action = "generate"
            decision = {
                "topicKey": str(item.get("topicKey") or item.get("topic") or "commitment"),
                "preferredTitle": _preferred_commitment_title(item, str(item.get("triggerText") or "承诺")),
                "reason": "AI 语义复核未返回结果，按明确承诺要求规则兜底生成。",
            }
        else:
            action = "clue"
            decision = {
                "topicKey": str(item.get("topicKey") or item.get("topic") or "commitment"),
                "reason": "AI 语义复核未返回结果，保留为人工确认线索。",
            }

        topic_key = str(decision.get("topicKey") or item.get("topicKey") or item.get("topic") or "commitment").strip() or "commitment"
        if action == "ignore":
            continue
        if action == "generate":
            title = str(decision.get("preferredTitle") or "").strip() or _preferred_commitment_title(item, str(item.get("triggerText") or "承诺"))
            title = _normalize_commitment_title_by_topic(topic_key, title, item)
            commitment_type = "disqualification" if topic_key == "disqualification" else "general_commitment"
            risk_flags = ["semantic_review_passed", "legal_wording_review_required"] if has_ai_decision else ["rule_fallback_generated", "legal_wording_review_required"]
            _append_commitment_letter(
                letters,
                generated_topics,
                item,
                topic_key=topic_key,
                title=title,
                commitment_type=commitment_type,
                risk_flags=risk_flags,
            )
            continue

        _append_commitment_clue(
            clues,
            clue_topics,
            item,
            topic_key=topic_key,
            recommended_action=str(decision.get("reason") or "暂不自动生成，请人工确认是否需要单独承诺函/承诺书。").strip(),
            risk_flags=["semantic_review_required"] if has_ai_decision else ["semantic_review_unavailable"],
        )

    return {"letters": letters, "clues": clues}


def _looks_like_section_heading(line: str) -> bool:
    text = str(line or "").strip()
    if not text:
        return False
    return bool(
        re.match(r"^(?:第[一二三四五六七八九十百千0-9]+[章节条]|[（(]?[一二三四五六七八九十0-9]+[）)])", text)
        or text.startswith("附件")
        or text.startswith("附表")
    )


def _scan_business_hint_items(documents: list[dict[str, Any]], texts_by_id: dict[str, str]) -> list[dict[str, Any]]:
    keywords = sorted(
        {
            *[alias for spec in PROJECT_BASIC_FIELDS for alias in spec.aliases if spec.key != "technicalCommitment"],
            *[alias for spec in BUSINESS_RESPONSE_FIELDS for alias in spec.aliases],
            *[alias for spec in QUALIFICATION_SUPPORT_FIELDS for alias in spec.aliases],
            *[alias for spec in COMMITMENT_REQUIREMENT_FIELDS for alias in spec.aliases],
            "投标人不得存在下列情形之一",
            "不得存在下列情形",
            "投标人需要说明的其他内容",
            "书面承诺",
            "无条件承诺",
        },
        key=len,
        reverse=True,
    )
    items: list[dict[str, Any]] = []
    seen: set[tuple[str, str]] = set()
    for document in documents:
        document_id = str(document.get("id") or "")
        source_file = str(document.get("name") or document_id or "招标文件")
        lines = [line.strip() for line in str(texts_by_id.get(document_id) or "").splitlines()]
        current_section = ""
        for line_number, line in enumerate(lines, start=1):
            if not line:
                continue
            if _looks_like_section_heading(line):
                current_section = line
            matched_keyword = next((keyword for keyword in keywords if keyword and keyword in line), "")
            is_commitment_obligation = _line_has_explicit_commitment_obligation(line)
            if not matched_keyword and not is_commitment_obligation:
                continue
            dedupe_key = (document_id, line)
            if dedupe_key in seen:
                continue
            seen.add(dedupe_key)
            label, value = _split_label_value(line, matched_keyword or "承诺要求")
            line_index = line_number - 1
            context_before = next(
                (candidate for candidate in reversed(lines[max(0, line_index - 2):line_index]) if candidate),
                "",
            )
            context_after = next(
                (candidate for candidate in lines[line_index + 1:line_index + 3] if candidate),
                "",
            )
            items.append(
                {
                    "id": f"RAW-{len(items) + 1:04d}",
                    "type": "商务提示",
                    "category": "business_hint",
                    "title": label or matched_keyword or "承诺要求",
                    "keyEntity": matched_keyword or "承诺要求",
                    "keyValue": value,
                    "value": value or line,
                    "sourceFile": source_file,
                    "sourceDocumentId": document_id,
                    "section": current_section,
                    "evidence": line,
                    "evidenceLocation": f"L{line_number}",
                    "contextBefore": context_before,
                    "contextAfter": context_after,
                    "confidence": 0.72,
                }
            )
    return items


def _build_commitment_requirement_fields(
    items: list[dict[str, Any]],
    *,
    analysis: dict[str, list[dict[str, Any]]] | None = None,
) -> list[dict[str, Any]]:
    analysis = analysis or _build_business_commitment_analysis(items)
    commitment_items = _find_commitment_items(items)
    disqualification_item = next(
        (
            item
            for item in commitment_items
            if "不得存在下列情形" in " ".join(str(item.get(key) or "") for key in ("title", "value", "evidence"))
        ),
        None,
    )
    other_content_item = next(
        (
            item
            for item in items
            if "投标人需要说明的其他内容" in " ".join(str(item.get(key) or "") for key in ("title", "value", "evidence", "section"))
        ),
        None,
    )

    fields: list[dict[str, Any]] = []
    for spec in COMMITMENT_REQUIREMENT_FIELDS:
        if spec.key == "generalCommitmentCount":
            count_text = str(len(analysis["letters"]) + len(analysis["clues"]))
            matched = commitment_items[0] if commitment_items else None
            fields.append(
                _business_field_from_item(spec, matched, value_override=count_text)
                if matched
                else _empty_business_field(spec, value="0")
            )
            continue
        if spec.key == "generatedCommitmentCount":
            matched = analysis["letters"][0] if analysis["letters"] else None
            fields.append(
                _business_field_from_item(spec, matched, value_override=str(len(analysis["letters"])))
                if matched
                else _empty_business_field(spec, value="0")
            )
            continue
        if spec.key == "pendingCommitmentCount":
            matched = analysis["clues"][0] if analysis["clues"] else None
            fields.append(
                _business_field_from_item(spec, matched, value_override=str(len(analysis["clues"])))
                if matched
                else _empty_business_field(spec, value="0")
            )
            continue
        if spec.key == "disqualificationCommitmentRequired":
            fields.append(
                _business_field_from_item(spec, disqualification_item)
                if disqualification_item
                else _empty_business_field(spec)
            )
            continue
        if spec.key == "otherCommitmentSectionRequired":
            fields.append(
                _business_field_from_item(spec, other_content_item)
                if other_content_item
                else _empty_business_field(spec)
            )
            continue
        if spec.key == "commitmentGenerationBasis":
            matched = analysis["letters"][0] if analysis["letters"] else None
            value = "；".join(str(item.get("triggerContext") or "").strip() for item in analysis["letters"][:3] if str(item.get("triggerContext") or "").strip()).strip("；")
            fields.append(
                _business_field_from_item(spec, matched, value_override=value or "未识别到明确的承诺函/承诺书生成依据")
                if matched
                else _empty_business_field(spec, value="未识别到明确的承诺函/承诺书生成依据")
            )
            continue
        fields.append(_empty_business_field(spec))
    return fields


def _business_presence_from_keywords(items: list[dict[str, Any]], *, keywords: tuple[str, ...]) -> dict[str, Any]:
    matched = [
        item
        for item in items
        if any(keyword in " ".join(str(item.get(key) or "") for key in ("title", "value", "evidence", "section")) for keyword in keywords)
    ]
    if not matched:
        return {
            "status": "missing",
            "summary": "招标文件中暂未识别到明确要求。",
            "evidences": [],
        }
    evidences = [
        {
            **_copy_meta_fields(item),
        }
        for item in matched[:8]
    ]
    summary = "；".join(str(item.get("value") or item.get("evidence") or "").strip() for item in matched if str(item.get("value") or item.get("evidence") or "").strip())
    return {
        "status": "present",
        "summary": summary[:800],
        "evidences": evidences,
    }


def _build_business_requirement_presence(items: list[dict[str, Any]]) -> dict[str, Any]:
    return {
        "qualificationDocuments": _business_presence_from_keywords(
            items,
            keywords=("资格证明", "合格投标人", "资格审查"),
        ),
        "performanceDocuments": _business_presence_from_keywords(
            items,
            keywords=("业绩", "合同", "中标通知书", "验收报告", "试运行"),
        ),
        "deviationResponse": _business_presence_from_keywords(
            items,
            keywords=("商务偏差", "偏差表", "偏离表"),
        ),
        "bidSecurity": _business_presence_from_keywords(
            items,
            keywords=("投标保证金", "保证金", "保函"),
        ),
        "otherCommitments": _business_presence_from_keywords(
            items,
            keywords=("承诺", "投标人需要说明的其他内容", "履约保证"),
        ),
        "disqualificationClauses": _business_presence_from_keywords(
            items,
            keywords=("投标人不得存在下列情形之一", "不得存在下列情形"),
        ),
    }


def _filter_business_scoring(scoring: dict[str, Any]) -> dict[str, list[dict[str, Any]]]:
    if not isinstance(scoring, dict):
        return {"business": [], "price": [], "compliance": []}
    return {
        "business": copy.deepcopy(scoring.get("business") or []),
        "price": copy.deepcopy(scoring.get("price") or []),
        "compliance": copy.deepcopy(scoring.get("compliance") or []),
    }


def _scoring_bucket_from_title(title: str) -> str:
    text = re.sub(r"\s+", " ", str(title or "").replace("\u3000", " ")).strip()
    if not text:
        return ""
    if any(keyword in text for keyword in ("投标报价评分", "报价评分", "价格评分", "开标价格表", "报价表")):
        return "price"
    if any(keyword in text for keyword in ("符合性审查", "合规", "符合性", "审查标准")):
        return "compliance"
    if any(keyword in text for keyword in ("商务评分", "商务评审", "商务打分", "评标办法")):
        return "business"
    return ""


def _collect_markdown_table(lines: list[str], start_index: int) -> tuple[list[tuple[int, str]], int]:
    rows: list[tuple[int, str]] = []
    index = start_index
    while index < len(lines):
        line = lines[index]
        if not MARKDOWN_TABLE_LINE_PATTERN.match(line.strip()):
            break
        rows.append((index + 1, line))
        index += 1
    return rows, index


def _parse_markdown_scoring_rows(
    *,
    rows: list[tuple[int, str]],
    document: dict[str, Any],
    section: str,
    start_index: int,
) -> list[dict[str, Any]]:
    parsed_rows: list[dict[str, Any]] = []
    table_rows = [(line_no, _parse_markdown_table_row(line)) for line_no, line in rows]
    filtered_rows = [(line_no, cells) for line_no, cells in table_rows if cells and not _is_markdown_separator_row(cells)]
    if len(filtered_rows) <= 1:
        return parsed_rows
    data_rows = filtered_rows[1:]
    order = start_index
    for line_no, cells in data_rows:
        if len(cells) < 4:
            continue
        scoring_item = cells[1].strip() if len(cells) > 1 else ""
        score = cells[2].strip() if len(cells) > 2 else ""
        score_point = cells[3].strip() if len(cells) > 3 else ""
        proof_requirement = cells[4].strip() if len(cells) > 4 else ""
        if not scoring_item and not score_point:
            continue
        parsed_rows.append(
            {
                "order": str(order),
                "scoringItem": scoring_item,
                "score": score,
                "scorePoint": score_point,
                "proofRequirement": proof_requirement,
                "status": "found",
                "sourceFile": str(document.get("name") or ""),
                "sourceDocumentId": str(document.get("id") or ""),
                "section": section,
                "evidence": " | ".join(cell for cell in cells if cell),
                "evidenceLocation": f"L{line_no}",
            }
        )
        order += 1
    return parsed_rows


def _extract_markdown_scoring(documents: list[dict[str, Any]], texts_by_id: dict[str, str]) -> dict[str, list[dict[str, Any]]]:
    scoring = {"business": [], "price": [], "compliance": []}
    for document in documents:
        text = str(texts_by_id.get(str(document.get("id") or "")) or "")
        if "|" not in text:
            continue
        lines = text.splitlines()
        current_section = ""
        pending_scoring_title = ""
        index = 0
        while index < len(lines):
            raw_line = lines[index]
            line = raw_line.strip()
            if not line:
                index += 1
                continue
            if line.startswith("#"):
                current_section = line.strip("# ").strip()
                pending_scoring_title = ""
                index += 1
                continue

            bucket_from_line = _scoring_bucket_from_title(line)
            if _looks_like_section_heading(line):
                current_section = line
                pending_scoring_title = line if bucket_from_line else ""
                index += 1
                continue

            if bucket_from_line:
                pending_scoring_title = line
                if index + 1 < len(lines) and MARKDOWN_TABLE_LINE_PATTERN.match(lines[index + 1].strip()):
                    rows, next_index = _collect_markdown_table(lines, index + 1)
                    scoring[bucket_from_line].extend(
                        _parse_markdown_scoring_rows(
                            rows=rows,
                            document=document,
                            section=line,
                            start_index=len(scoring[bucket_from_line]) + 1,
                        )
                    )
                    pending_scoring_title = ""
                    index = next_index
                    continue

            if MARKDOWN_TABLE_LINE_PATTERN.match(line):
                bucket = _scoring_bucket_from_title(pending_scoring_title or current_section)
                if bucket:
                    rows, next_index = _collect_markdown_table(lines, index)
                    scoring[bucket].extend(
                        _parse_markdown_scoring_rows(
                            rows=rows,
                            document=document,
                            section=pending_scoring_title or current_section,
                            start_index=len(scoring[bucket]) + 1,
                        )
                    )
                    pending_scoring_title = ""
                    index = next_index
                    continue

            index += 1
    return scoring


def _merge_business_scoring(
    base_scoring: dict[str, Any],
    markdown_scoring: dict[str, list[dict[str, Any]]],
) -> dict[str, list[dict[str, Any]]]:
    merged = _filter_business_scoring(base_scoring)
    seen = {
        bucket: {
            (
                str(row.get("sourceDocumentId") or ""),
                str(row.get("evidenceLocation") or ""),
                str(row.get("scoringItem") or ""),
                str(row.get("score") or ""),
            )
            for row in merged.get(bucket) or []
            if isinstance(row, dict)
        }
        for bucket in merged
    }
    for bucket, rows in markdown_scoring.items():
        for row in rows:
            key = (
                str(row.get("sourceDocumentId") or ""),
                str(row.get("evidenceLocation") or ""),
                str(row.get("scoringItem") or ""),
                str(row.get("score") or ""),
            )
            if key in seen[bucket]:
                continue
            seen[bucket].add(key)
            merged[bucket].append(row)
    return merged


def _build_business_coverage(
    field_groups: dict[str, Any],
    scoring: dict[str, list[dict[str, Any]]],
    presence: dict[str, Any],
) -> list[dict[str, Any]]:
    checks = [
        ("商务评分要求", len(scoring.get("business") or []) > 0),
        ("报价与价格表", any(field.get("status") == "found" for field in field_groups.get("businessResponse") or [] if field.get("key") in {"bidPriceTableRequired", "openingPriceTableRequired"})),
        ("偏差响应", bool((presence.get("deviationResponse") or {}).get("status") == "present")),
        ("资格证明", bool((presence.get("qualificationDocuments") or {}).get("status") == "present")),
        ("业绩证明", bool((presence.get("performanceDocuments") or {}).get("status") == "present")),
        ("保证金", bool((presence.get("bidSecurity") or {}).get("status") == "present")),
        ("其他承诺", bool((presence.get("otherCommitments") or {}).get("status") == "present")),
    ]
    return [
        {
            "label": label,
            "status": "covered" if covered else "missing",
        }
        for label, covered in checks
    ]


def _build_business_commitment_letters(
    project_id: str,
    items: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    _ = project_id
    return _build_business_commitment_analysis(items)["letters"]



def _build_business_commitment_clues(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return _build_business_commitment_analysis(items)["clues"]


def _normalized_business_match_text(value: str) -> str:
    return re.sub(
        r"(?:附件|附表)?[A-Za-z0-9一二三四五六七八九十]+|承诺函|承诺书|格式|模板|投标人|投标方|我方|本公司|\s+",
        "",
        str(value or ""),
    )


def _commitment_alignment_topics_for_template(appendix: dict[str, Any]) -> set[str]:
    template_type = str(appendix.get("templateType") or "").strip()
    text = re.sub(
        r"\s+",
        "",
        " ".join(
            str(appendix.get(key) or "")
            for key in ("title", "templateType", "templateTypeLabel", "evidence")
        ),
    )
    topics: set[str] = set()
    if template_type == "integrity_commitment" or "廉洁" in text:
        topics.add("integrity")
    if template_type == "performance_bond" or "履约保证" in text or "履约保函" in text:
        topics.add("performance_bond")
    if template_type == "bid_security" or "投标保证金" in text or "保证金" in text:
        topics.add("security")
    if "保密" in text:
        topics.add("confidentiality")
    if "不得存在下列情形" in text:
        topics.add("disqualification")
    if any(token in text for token in ("取得本条", "取得材料", "材料取得", "取得证书", "证书取得", "取得认证", "供货前取得")):
        topics.add("certificate_obtainment")
    if any(token in text for token in ("合规", "守法", "违法", "违规", "信用")):
        topics.add("compliance")
    if any(token in text for token in ("交货", "工期", "供货周期", "交付")):
        topics.update({"delivery", "delivery_commitment"})
    if any(token in text for token in ("质量", "质保", "售后", "服务承诺")):
        topics.update({"quality", "quality_commitment"})
    if template_type == "commitment" and not topics:
        topics.add("generic_commitment_template")
    return topics


def _commitment_template_can_cover_letter(letter: dict[str, Any], appendix: dict[str, Any]) -> bool:
    topic_key = str(letter.get("topicKey") or letter.get("topic") or "").strip()
    if not topic_key:
        return False
    template_topics = _commitment_alignment_topics_for_template(appendix)
    if topic_key in template_topics:
        return True
    if topic_key == "security" and "bid_security" in template_topics:
        return True
    if topic_key in {"delivery", "delivery_commitment"} and template_topics & {"delivery", "delivery_commitment"}:
        return True
    if topic_key in {"quality", "quality_commitment"} and template_topics & {"quality", "quality_commitment"}:
        return True
    return False


def _commitment_template_match_score(letter: dict[str, Any], appendix: dict[str, Any]) -> float:
    if not _commitment_template_can_cover_letter(letter, appendix):
        return 0.0
    letter_text = " ".join(
        str(letter.get(key) or "")
        for key in ("title", "topicKey", "triggerText", "triggerContext")
    )
    appendix_text = " ".join(
        str(appendix.get(key) or "")
        for key in ("title", "templateType", "templateTypeLabel", "evidence")
    )
    letter_normalized = _normalized_business_match_text(letter_text)
    appendix_normalized = _normalized_business_match_text(appendix_text)
    score = 0.74
    if not letter_normalized or not appendix_normalized:
        return score
    if letter_normalized in appendix_normalized or appendix_normalized in letter_normalized:
        return 0.92
    topic_key = str(letter.get("topicKey") or "")
    for key, keywords in COMMITMENT_TOPIC_KEYWORDS:
        if key == topic_key and any(keyword in appendix_text for keyword in keywords):
            score = max(score, 0.86)
    letter_tokens = {token for token in re.split(r"[、，,；;：:（）()]+", letter_normalized) if len(token) >= 2}
    appendix_tokens = {token for token in re.split(r"[、，,；;：:（）()]+", appendix_normalized) if len(token) >= 2}
    overlap = letter_tokens & appendix_tokens
    if letter_tokens and appendix_tokens and overlap:
        score = max(score, min(0.82, len(overlap) / max(len(letter_tokens), len(appendix_tokens)) + 0.35))
    return score


def _align_commitment_letters_with_existing_templates(
    letters: list[dict[str, Any]],
    appendices: list[dict[str, Any]],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    templates = [
        item
        for item in appendices
        if isinstance(item, dict)
        and item.get("artifactType") == "business_attachment_template"
        and (
            str(item.get("templateType") or "") in {"commitment", "integrity_commitment", "performance_bond"}
            or any(token in str(item.get("title") or "") for token in COMMITMENT_DOC_KEYWORDS)
        )
    ]
    if not letters or not templates:
        return letters, []

    remaining: list[dict[str, Any]] = []
    alignments: list[dict[str, Any]] = []
    for letter in letters:
        scored = [
            (_commitment_template_match_score(letter, template), template)
            for template in templates
        ]
        score, template = max(scored, key=lambda item: item[0], default=(0.0, {}))
        if score >= 0.62 and isinstance(template, dict):
            alignments.append(
                {
                    "status": "covered_by_existing_template",
                    "requirementTitle": str(letter.get("title") or ""),
                    "requirementTopicKey": str(letter.get("topicKey") or ""),
                    "requirementSource": str(letter.get("triggerContext") or letter.get("evidence") or ""),
                    "matchedTemplateId": str(template.get("id") or ""),
                    "matchedTemplateTitle": str(template.get("title") or ""),
                    "matchedTemplateType": str(template.get("templateType") or ""),
                    "confidence": round(score, 3),
                }
            )
            continue
        remaining.append(letter)
    return remaining, alignments


BUSINESS_PROJECT_FACT_FIELD_SPECS: tuple[dict[str, Any], ...] = (
    {"fieldKey": "projectName", "label": "项目名称", "required": True},
    {"fieldKey": "tenderNo", "label": "招标编号", "required": True},
    {"fieldKey": "tenderer", "label": "招标人", "required": True},
    {"fieldKey": "managementUnit", "label": "管理单位", "required": False},
    {"fieldKey": "bidSectionScale", "label": "标段规模", "required": False},
    {"fieldKey": "deliveryPeriod", "label": "交货周期", "required": False},
    {"fieldKey": "warrantyPeriod", "label": "质保期", "required": False},
)


def _build_business_project_fact_fields(
    field_groups: dict[str, Any],
    project_dates: dict[str, Any],
) -> list[dict[str, Any]]:
    project_basics = field_groups.get("projectBasics") if isinstance(field_groups.get("projectBasics"), list) else []
    fields_by_key = {
        str(field.get("key") or ""): field
        for field in project_basics
        if isinstance(field, dict)
    }
    fact_fields: list[dict[str, Any]] = []
    for spec in BUSINESS_PROJECT_FACT_FIELD_SPECS:
        field_key = str(spec["fieldKey"])
        source = fields_by_key.get(field_key) or {}
        value = str(source.get("value") or "").strip()
        fact_fields.append(
            {
                "fieldKey": field_key,
                "label": str(spec["label"]),
                "value": value,
                "category": "项目基础信息",
                "status": "found" if value else "missing",
                "required": bool(spec.get("required", False)),
                "confidence": float(source.get("confidence") or (0.86 if value else 0.0)),
                "sourceFile": str(source.get("sourceFile") or ""),
                "sourceDocumentId": str(source.get("sourceDocumentId") or ""),
                "section": str(source.get("section") or ""),
                "evidence": str(source.get("evidence") or ""),
                "evidenceLocation": str(source.get("evidenceLocation") or ""),
            }
        )

    date_specs = (
        ("bidStartDate", "投标起始日期", "startDate", False),
        ("bidDeadline", "投标截止日期", "endDate", True),
    )
    for field_key, label, date_key, required in date_specs:
        value = str(project_dates.get(date_key) or "").strip()
        fact_fields.append(
            {
                "fieldKey": field_key,
                "label": label,
                "value": value,
                "category": "投标时间信息",
                "status": "found" if value else "missing",
                "required": required,
                "confidence": 0.78 if value else 0.0,
                "sourceFile": "",
                "sourceDocumentId": "",
                "section": "",
                "evidence": "",
                "evidenceLocation": "",
            }
        )
    return fact_fields


def _transform_to_business_contract(
    project_id: str,
    payload: dict[str, Any],
    *,
    profile: ParseProfile,
    documents: list[dict[str, Any]],
    texts_by_id: dict[str, str],
    run_semantic_review: bool = True,
) -> dict[str, Any]:
    result = copy.deepcopy(payload if isinstance(payload, dict) else {})
    items = result.get("items") if isinstance(result.get("items"), list) else []
    hint_items = _scan_business_hint_items(documents, texts_by_id)
    docx_candidate_items = _extract_docx_core_candidate_items(documents)
    merged_items = [*copy.deepcopy(items), *docx_candidate_items, *hint_items]
    result["items"] = merged_items
    structured = result.get("structured") if isinstance(result.get("structured"), dict) else {}
    source_documents = copy.deepcopy(structured.get("sourceDocuments") or [])
    project_dates = copy.deepcopy(structured.get("projectDates") or {"startDate": "", "endDate": ""})
    appendices = copy.deepcopy(structured.get("appendices") or [])
    scoring = _merge_business_scoring(
        structured.get("scoringCriteria") or {},
        _extract_markdown_scoring(documents, texts_by_id),
    )
    commitment_analysis = _build_business_commitment_analysis(
        merged_items,
        run_semantic_review=run_semantic_review,
    )
    field_groups = {
        "projectBasics": _build_business_project_basics(merged_items, project_dates),
        "businessResponse": _build_business_response_fields(merged_items),
        "qualificationSupport": _build_qualification_support_fields(merged_items),
        "qualificationRequirements": _build_qualification_requirements(
            merged_items,
            documents=documents,
            texts_by_id=texts_by_id,
        ),
        "bidderInstructions": _extract_bidder_instruction_rows(documents),
        "commercialRejectionClauses": _extract_commercial_rejection_clauses(documents, texts_by_id),
        "commitmentRequirements": _build_commitment_requirement_fields(
            merged_items,
            analysis=commitment_analysis,
        ),
    }
    presence = _build_business_requirement_presence(merged_items)
    _ = project_id
    commitment_letters = copy.deepcopy(commitment_analysis["letters"])
    commitment_clues = copy.deepcopy(commitment_analysis["clues"])
    commitment_letters, commitment_template_alignments = _align_commitment_letters_with_existing_templates(
        commitment_letters,
        appendices,
    )
    project_fact_fields = _build_business_project_fact_fields(field_groups, project_dates)

    result["structured"] = {
        "schemaVersion": profile.schema_version,
        "targetSkill": profile.skill_name,
        "mode": str(structured.get("mode") or "local-structured-parser"),
        "sourceDocuments": source_documents,
        "scoringCriteria": scoring,
        "fieldGroups": field_groups,
        "requirementPresence": presence,
        "coverage": _build_business_coverage(field_groups, scoring, presence),
        "projectDates": {
            "startDate": str(project_dates.get("startDate") or ""),
            "endDate": str(project_dates.get("endDate") or ""),
        },
        "appendices": appendices,
        "commitmentLetters": commitment_letters,
        "commitmentClues": commitment_clues,
        "commitmentTemplateAlignments": commitment_template_alignments,
        "businessFormatRegions": [
            {
                "sourceFile": str(document.get("name") or ""),
                "regionCount": len(_detect_business_format_regions(_iter_docx_blocks(Path(str(document.get("sourcePath") or "")))))
                if str(document.get("sourcePath") or "").lower().endswith(".docx")
                and Path(str(document.get("sourcePath") or "")).exists()
                else 0,
            }
            for document in documents
        ],
        "projectFactFields": project_fact_fields,
        "categoryCounts": {
            "商务评分": len(scoring.get("business") or []),
            "报价评分": len(scoring.get("price") or []),
            "合规审查": len(scoring.get("compliance") or []),
            "商务附表": len(appendices),
            "承诺文件": len(commitment_letters),
            "待确认承诺线索": len(commitment_clues),
        },
        "opencodeOutput": copy.deepcopy(structured.get("opencodeOutput") or {}),
    }
    return result


def _business_project_name_from_structured(structured: dict[str, Any]) -> str:
    field_groups = structured.get("fieldGroups") if isinstance(structured, dict) else {}
    project_basics = field_groups.get("projectBasics") if isinstance(field_groups, dict) else []
    for field in project_basics if isinstance(project_basics, list) else []:
        if not isinstance(field, dict):
            continue
        if str(field.get("key") or "") != "projectName":
            continue
        value = str(field.get("value") or "").strip()
        if value:
            return value
    raw_basics = structured.get("projectBasics") if isinstance(structured, dict) else {}
    if isinstance(raw_basics, dict):
        return str(raw_basics.get("项目名称") or raw_basics.get("projectName") or "").strip()
    return ""


def _business_tenderer_name_from_structured(structured: dict[str, Any]) -> str:
    field_groups = structured.get("fieldGroups") if isinstance(structured, dict) else {}
    project_basics = field_groups.get("projectBasics") if isinstance(field_groups, dict) else []
    for field in project_basics if isinstance(project_basics, list) else []:
        if not isinstance(field, dict):
            continue
        if str(field.get("key") or "") != "tenderer":
            continue
        value = str(field.get("value") or "").strip()
        if value:
            return value
    return ""


def _business_text_lines(text: str) -> list[tuple[int, str]]:
    return [
        (line_number, line.strip())
        for line_number, line in enumerate(str(text or "").splitlines(), start=1)
        if line.strip()
    ]


def _looks_like_toc_or_directory_line(text: str) -> bool:
    normalized = re.sub(r"\s+", " ", str(text or "").strip().lstrip("#").strip())
    if not normalized:
        return False
    if any(token in normalized for token in ("……", "....", "----")):
        return True
    if re.search(r"(?:\s|\.{2,}|…{2,})\d{1,4}$", normalized):
        return True
    return False


def _looks_like_business_template_body_sentence(text: str) -> bool:
    normalized = str(text or "").strip().lstrip("#").strip()
    if not normalized:
        return False
    if any(mark in normalized for mark in ("。", "；", ";")):
        return True
    return any(
        token in normalized
        for token in (
            "我方",
            "本公司",
            "投标人承诺",
            "按招标文件",
            "按照招标文件",
            "愿意",
            "承担",
            "提交",
            "提供",
        )
    )


def _looks_like_business_template_stop_heading(text: str, profile: ParseProfile) -> bool:
    normalized = str(text or "").strip().lstrip("#").strip()
    if not normalized:
        return False
    if _is_scoring_appendix_heading(normalized):
        return True
    if profile.key == "business" and _is_business_major_section_heading(normalized):
        return True
    if _is_relevant_appendix_heading(normalized, profile):
        return True
    if profile.key == "business" and _looks_like_business_attachment_template_title(normalized, in_template_section=True):
        return True
    return bool(
        re.match(r"^(?:第[一二三四五六七八九十百千0-9]+[章节条]|[一二三四五六七八九十]+[、.．]|[（(][一二三四五六七八九十0-9]+[）)])", normalized)
        and any(
            token in normalized
            for token in (
                "投标函", "法定代表人", "授权", "廉洁", "专用章", "投标价格", "开标价格",
                "商务偏差", "货物规格", "供货范围", "保证金", "履约", "附件",
                "资格", "证明", "其他内容", "承诺函", "承诺书",
            )
        )
    )


def _appendix_has_material_content(rows: list[list[str]], content_blocks: list[dict[str, Any]]) -> bool:
    if rows:
        return True
    for block in content_blocks:
        if not isinstance(block, dict):
            continue
        if block.get("type") == "paragraph" and len(str(block.get("text") or "").strip()) >= 6:
            return True
        if block.get("type") == "table" and isinstance(block.get("rows"), list) and block.get("rows"):
            return True
    return False


def _business_template_should_materialize(metadata: dict[str, Any], rows: list[list[str]], content_blocks: list[dict[str, Any]]) -> bool:
    if not _appendix_has_material_content(rows, content_blocks):
        return False
    quality = str(metadata.get("extractionQuality") or "").strip()
    issues = metadata.get("qualityIssues") if isinstance(metadata.get("qualityIssues"), list) else []
    if quality in {"title_only", "probably_incomplete"}:
        return False
    if any("未识别到明显表格、签章栏或待填写占位" in str(issue) for issue in issues):
        return False
    return True


def _business_text_template_appendix_title(line: str, profile: ParseProfile) -> str:
    title = str(line or "").strip().lstrip("#").strip()
    if _is_relevant_appendix_heading(title, profile):
        return title.strip(" #") or "商务附件模板"
    match = re.match(
        r"^(?:[一二三四五六七八九十0-9]+[、.．]\s*|[（(][一二三四五六七八九十0-9]+[）)]\s*)?"
        r"(?P<title>(?:投标函|法定代表人(?:（单位负责人）)?身份证明|法定代表人授权书|法定代表人授权委托书|授权委托书|投标人廉洁自律承诺书|廉洁承诺书|投标专用章效力说明|投标价格表|开标价格表|商务偏差表|货物规格表|供货范围表|投标保证金|履约保证函格式承诺书|履约承诺书|否决项响应|投标人需要说明的其他内容|其他说明|附件\s*[0-9一二三四五六七八九十]+)[^：:。；;]*)",
        title,
    )
    return (match.group("title").strip() if match else title) or "商务附件模板"


def _extract_text_business_appendices(
    project_id: str,
    documents: list[dict[str, Any]],
    texts_by_id: dict[str, str],
    *,
    start_index: int = 0,
    profile: ParseProfile = BUSINESS_PARSE_PROFILE,
) -> list[dict[str, Any]]:
    if profile.key != "business":
        return []
    appendices: list[dict[str, Any]] = []
    for document in documents:
        document_id = str(document.get("id") or "")
        source_file = str(document.get("name") or document_id or "招标文件")
        source_path = Path(str(document.get("sourcePath") or ""))
        if source_path.suffix.lower() in {".md", ".docx"}:
            continue
        lines = _business_text_lines(texts_by_id.get(document_id, ""))
        in_template_section = False
        template_section_title = ""
        for index, (line_number, line) in enumerate(lines):
            if _is_business_template_section_heading(line):
                in_template_section = True
                template_section_title = line
                continue
            if in_template_section and _is_business_major_section_heading(line):
                in_template_section = False
                template_section_title = ""
            if not _business_template_title_allowed(line, in_template_section=in_template_section):
                continue
            title = _business_text_template_appendix_title(line, profile)
            content_blocks: list[dict[str, Any]] = []
            next_boundary_index = len(lines)
            for lookahead_index, (_, next_line) in enumerate(lines[index + 1:], start=index + 1):
                if _is_business_template_section_heading(next_line):
                    next_boundary_index = lookahead_index
                    break
                if _is_business_major_section_heading(next_line):
                    next_boundary_index = lookahead_index
                    break
                if _looks_like_business_template_stop_heading(next_line, profile):
                    next_boundary_index = lookahead_index
                    break
                content_blocks.append({"type": "paragraph", "text": next_line})
                if len(content_blocks) >= 80:
                    next_boundary_index = lookahead_index + 1
                    break
            metadata = _business_template_metadata(
                title=title,
                source_file=source_file,
                evidence=line,
                evidence_location=f"L{line_number}",
                template_section_title=template_section_title,
                source_start=_business_template_source_start(template_section_title, f"L{line_number}", line),
                source_end=_business_template_source_end_from_lines([text for _, text in lines], next_boundary_index),
                rows=[],
                content_blocks=content_blocks,
                extraction_mode="text_slice",
            )
            if not _business_template_should_materialize(metadata, [], content_blocks):
                continue
            appendix_id = f"APPX-{start_index + len(appendices) + 1:04d}"
            appendices.append(
                materialize_appendix_docx(
                    project_id,
                    {
                        "id": appendix_id,
                        "title": title,
                        "status": "generated",
                        "rows": [],
                        "contentBlocks": content_blocks,
                        "rowCount": 0,
                        "docxPath": "",
                        **metadata,
                    },
                    profile=profile,
                )
            )
    return appendices


def _sanitize_docx_name(value: str, fallback: str) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|]+", "-", str(value or "").strip())
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" .")
    while len(cleaned.encode("utf-8")) > 120:
        cleaned = cleaned[:-1].strip(" .")
    return cleaned or fallback


def _parse_markdown_table_row(line: str) -> list[str]:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return cells


def _is_markdown_separator_row(cells: list[str]) -> bool:
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def _write_table_to_docx(doc: Document, rows: list[list[str]]) -> None:
    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index in range(column_count):
            table.cell(row_index, col_index).text = row[col_index] if col_index < len(row) else ""


def _write_appendix_docx(
    path: Path,
    title: str,
    rows: list[list[str]],
    content_blocks: list[dict[str, Any]] | None = None,
) -> None:
    """Fallback: build a fresh docx from flattened rows/content blocks.

    Source docx slicing is preferred when available because it preserves merged
    cells and styles. Business attachments can additionally pass content blocks
    so template body text is not reduced to a title-only Word file.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(title, level=1)
    wrote_content = False
    for block in content_blocks or []:
        if not isinstance(block, dict):
            continue
        block_type = str(block.get("type") or "")
        if block_type == "paragraph":
            block_text = str(block.get("text") or "").strip()
            if block_text:
                doc.add_paragraph(block_text)
                wrote_content = True
        elif block_type == "table":
            block_rows = block.get("rows") if isinstance(block.get("rows"), list) else []
            if block_rows:
                _write_table_to_docx(doc, block_rows)
                wrote_content = True
    if rows and not wrote_content:
        _write_table_to_docx(doc, rows)
    doc.save(path)

def _build_appendix_slice_state(source_docx: Path) -> dict[str, Any] | None:
    """Read ``source_docx`` once and return cached state for repeated calls
    to :func:`_slice_appendix_from_source`.

    A real RFP can be 20+ MB with millions of XML elements and dozens of
    appendices. Cloning the entire document tree per appendix (the previous
    approach) burned multiple seconds each — the caller froze for minutes
    on production-sized files. The current approach is genuinely a "cut":
    we hold direct references to every body child of the source, and each
    slice MOVES the kept children into a fresh per-appendix body element
    rather than cloning. Children are O(1) to detach; only the trailing
    ``<w:sectPr>`` is deepcopied (it's a single small element shared across
    all outputs).

    The returned dict carries:
    - ``sourcePath``  — original file path (kept for diagnostics)
    - ``parts``       — every zip entry of the source as ``{name: (info, bytes)}``,
      so each output zip can be assembled without re-reading the source
    - ``rootTag`` / ``rootAttrib`` / ``rootNsmap`` — used to clone the
      ``<w:document>`` shell (without its body) for each output
    - ``rootChildrenBeforeBody`` — non-body siblings of body (rare, e.g.
      ``<w:background>``) cached as deepcopies so each output gets a fresh
      copy
    - ``bodyTag`` / ``bodyAttrib`` — used to build the new body element
    - ``bodyChildren`` — list of references to each body child of the source,
      indexable by ``body_index``. Children are MOVED out of this list into
      per-appendix bodies, mutating the source tree in the process.
    - ``sectPr`` — reference to the source's trailing ``<w:sectPr>`` so each
      output can deepcopy its own copy
    """

    if not source_docx.is_file():
        return None
    try:
        from lxml import etree as _etree
        from copy import deepcopy as _deepcopy

        with zipfile.ZipFile(source_docx, "r") as zf:
            parts: dict[str, tuple[Any, bytes]] = {
                info.filename: (info, zf.read(info.filename))
                for info in zf.infolist()
            }
        doc_xml = parts["word/document.xml"][1]
        doc_tree = _etree.fromstring(doc_xml)
        body_tag = f"{WORD_NAMESPACE}body"
        sect_pr_tag = f"{WORD_NAMESPACE}sectPr"
        body = doc_tree.find(body_tag)
        if body is None:
            return None

        # Cache non-body siblings of <w:document> (e.g. <w:background>) as
        # deepcopies so each output gets a fresh, independent copy.
        root_children_before_body: list[Any] = []
        for child in doc_tree.iterchildren():
            if child.tag == body_tag:
                break
            root_children_before_body.append(_deepcopy(child))

        body_children = list(body.iterchildren())
        sect_pr = body.find(sect_pr_tag)
    except Exception:
        return None

    return {
        "sourcePath": str(source_docx),
        "parts": parts,
        "rootTag": doc_tree.tag,
        "rootAttrib": dict(doc_tree.attrib),
        "rootNsmap": dict(doc_tree.nsmap),
        "rootChildrenBeforeBody": root_children_before_body,
        "bodyTag": body_tag,
        "bodyAttrib": dict(body.attrib),
        "bodyChildren": body_children,
        "sectPr": sect_pr,
    }


def _slice_appendix_from_source(
    source_docx: Path,
    target_docx: Path,
    keep_start: int,
    keep_end: int,
    source_state: dict[str, Any] | None = None,
) -> bool:
    """Produce ``target_docx`` by literally CUTTING children
    ``[keep_start, keep_end]`` (inclusive, by ``body_index``) out of the
    source docx body and dropping them into a fresh ``<w:document>`` shell.

    The kept children are MOVED, not cloned: lxml's ``new_parent.append(elem)``
    detaches ``elem`` from its old parent in O(1). For a 21 MB RFP with 50
    appendices this brings per-appendix CPU work from ~3 s (full-tree
    deepcopy) to a few ms (deepcopy of one ``<w:sectPr>``).

    All non-document parts of the docx (``word/styles.xml``,
    ``word/numbering.xml``, ``word/_rels/document.xml.rels``,
    ``word/media/*`` …) are written to the output zip from cached source
    bytes, so cell merges, fonts, embedded images, list numbering and
    hyperlinks render exactly as in the source.

    Caveat: because moves mutate the shared ``source_state``, processing
    appendices that overlap on body indices is not supported (``S0`` already
    enforces disjoint ranges via ``used_tables``). Children that have been
    moved into one output's body cannot be moved again.

    Returns ``True`` on success, ``False`` if the source is missing or the
    indices are not valid."""

    if keep_end < keep_start:
        return False
    if source_state is None:
        source_state = _build_appendix_slice_state(source_docx)
    if source_state is None:
        return False
    target_docx.parent.mkdir(parents=True, exist_ok=True)

    from copy import deepcopy as _deepcopy
    from lxml import etree as _etree

    body_children: list[Any] = source_state["bodyChildren"]
    sect_pr = source_state["sectPr"]
    sect_pr_tag = f"{WORD_NAMESPACE}sectPr"

    # Build a fresh <w:document> shell with the source's namespaces / attribs.
    new_root = _etree.Element(
        source_state["rootTag"],
        attrib=source_state["rootAttrib"],
        nsmap=source_state["rootNsmap"],
    )
    # Re-attach any non-body siblings that lived above <w:body> in the source
    # (rare, but safe).
    for sibling in source_state["rootChildrenBeforeBody"]:
        new_root.append(_deepcopy(sibling))
    new_body = _etree.SubElement(
        new_root,
        source_state["bodyTag"],
        attrib=source_state["bodyAttrib"],
    )

    # MOVE each kept child from source body into the new body. lxml semantics:
    # ``new_body.append(child)`` detaches ``child`` from its prior parent.
    n = len(body_children)
    start = max(0, keep_start)
    end = min(n - 1, keep_end)
    for idx in range(start, end + 1):
        child = body_children[idx]
        if child.tag == sect_pr_tag:
            # sect_pr is appended at the end as a deepcopy so subsequent
            # appendices can also reference it.
            continue
        # If this child has already been moved (overlapping ranges, which
        # shouldn't happen but we guard anyway), skip it.
        if child.getparent() is None or child.getparent() is not None and child.getparent().tag != source_state["bodyTag"]:
            # Already moved out of the source body — skip silently.
            if child.getparent() is None:
                continue
        new_body.append(child)

    if sect_pr is not None:
        new_body.append(_deepcopy(sect_pr))

    new_doc_xml = _etree.tostring(
        new_root,
        xml_declaration=True,
        encoding="UTF-8",
        standalone=True,
    )

    # Assemble the output zip from cached source parts. The only entry that
    # changes per appendix is ``word/document.xml``; everything else
    # (styles, numbering, rels, media) ships byte-for-byte.
    parts: dict[str, tuple[Any, bytes]] = source_state["parts"]
    try:
        with zipfile.ZipFile(target_docx, "w", zipfile.ZIP_DEFLATED) as dst:
            for filename, (info, data) in parts.items():
                if filename == "word/document.xml":
                    dst.writestr(info, new_doc_xml)
                else:
                    dst.writestr(info, data)
    except Exception:
        target_docx.unlink(missing_ok=True)
        return False
    return True


def _appendix_output_dir(project_id: str) -> Path:
    return parsed_appendix_path(project_id)


def _commitment_letter_output_dir(project_id: str) -> Path:
    path = settings.parsed_dir / project_id / "s1_commitment_letters"
    path.mkdir(parents=True, exist_ok=True)
    return path


def _workspace_appendix_output_dir(project_id: str, profile: ParseProfile) -> Path:
    return settings.documents_dir / project_id / profile.workspace_dirname / "appendices"


def _workspace_commitment_letter_output_dir(project_id: str, profile: ParseProfile) -> Path:
    return settings.documents_dir / project_id / profile.workspace_dirname / "commitment-letters"


def _appendix_asset_path(project_id: str, appendix_id: str, title: str) -> tuple[Path, str]:
    file_name = f"{appendix_id}-{_sanitize_docx_name(title, '附表')}.docx"
    return _appendix_output_dir(project_id) / file_name, f"s1_appendices/{file_name}"


def _commitment_letter_asset_path(project_id: str, letter_id: str, title: str) -> tuple[Path, str]:
    file_name = f"{letter_id}-{_sanitize_docx_name(title, '承诺函')}.docx"
    return _commitment_letter_output_dir(project_id) / file_name, f"s1_commitment_letters/{file_name}"


def _path_is_inside(path: Path, root: Path) -> bool:
    try:
        path.resolve().relative_to(root.resolve())
        return True
    except ValueError:
        return False


def _write_commitment_letter_docx(path: Path, letter: dict[str, Any], *, project_name: str = "", tenderer_name: str = "") -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    title = str(letter.get("title") or "承诺函").strip() or "承诺函"
    trigger_context = str(letter.get("triggerContext") or letter.get("evidence") or letter.get("triggerText") or "").strip()
    commitment_type = str(letter.get("commitmentType") or "").strip()

    doc = Document()
    doc.add_heading(title, level=1)
    if project_name:
        doc.add_paragraph(f"项目名称：{project_name}")
    if tenderer_name:
        doc.add_paragraph(f"招标人：{tenderer_name}")
    if trigger_context:
        doc.add_paragraph(f"触发依据：{trigger_context}")

    doc.add_paragraph("致：招标人")
    doc.add_paragraph("说明：本文件为商务标解析阶段自动生成的承诺函草稿，请结合招标文件原文、法务要求和项目实际情况复核后使用。")

    if commitment_type == "disqualification":
        doc.add_paragraph("我方在参加本项目投标过程中，郑重承诺如下：")
        doc.add_paragraph("1. 我方不存在招标文件中列示的“投标人不得存在下列情形之一”所述任一情形。")
        doc.add_paragraph("2. 如本承诺与实际情况不符，我方愿意按照招标文件和相关规定承担相应责任。")
    else:
        doc.add_paragraph("根据招标文件中的承诺要求，我方郑重承诺如下：")
        doc.add_paragraph("1. 我方将严格按照招标文件要求对相关事项进行响应并履行承诺。")
        if trigger_context:
            doc.add_paragraph(f"2. 本承诺函重点对应的招标文件原文为：{trigger_context}")
        doc.add_paragraph("3. 如本承诺与实际情况不符，我方愿意按照招标文件和相关规定承担相应责任。")

    doc.add_paragraph("")
    doc.add_paragraph("投标人（盖章）：________________")
    doc.add_paragraph("法定代表人或授权代表（签字或盖章）：________________")
    doc.add_paragraph("日期：________________")
    doc.save(path)


def materialize_appendix_docx(project_id: str, appendix: dict[str, Any], *, profile: ParseProfile = TECHNICAL_PARSE_PROFILE) -> dict[str, Any]:
    """Ensure an appendix entry has a generated Word asset, even when no template table was found.

    The slice metadata, if present under the private ``_slice`` key, is consumed
    here and stripped from the returned dict so it never enters JSON / DB. The
    expected shape is::

        {"sourcePath": str, "keepStart": int, "keepEnd": int}

    When provided and the slice succeeds, the appendix docx is produced by
    cutting the body of ``sourcePath`` directly, which preserves merges, styles,
    media, numbering and section properties verbatim. Otherwise we fall back to
    rebuilding from ``rows`` (Markdown/PDF inputs)."""

    # Peel off the private ``_slice`` key BEFORE deep-copying, since it can
    # carry an in-memory lxml tree (``sourceState``) that is large and
    # expensive to copy element-by-element. Shallow-copy the rest of the
    # dict so we don't mutate the caller's payload.
    slice_info = None
    if isinstance(appendix, dict) and "_slice" in appendix:
        appendix_without_slice = dict(appendix)
        slice_info = appendix_without_slice.pop("_slice", None)
        item = copy.deepcopy(appendix_without_slice)
    else:
        item = copy.deepcopy(appendix)
    appendix_id = str(item.get("id") or "").strip() or "APPX-0000"
    title = str(item.get("title") or item.get("evidence") or "附表").strip() or "附表"
    rows = item.get("rows") if isinstance(item.get("rows"), list) else []
    content_blocks = item.get("contentBlocks") if isinstance(item.get("contentBlocks"), list) else []
    output_dir = _appendix_output_dir(project_id)
    workspace_output_dir = _workspace_appendix_output_dir(project_id, profile)
    project_workspace_root = settings.documents_dir / project_id

    existing_path = Path(str(item.get("docxPath") or ""))
    if not existing_path.is_absolute():
        existing_path = Path()
    extractor_docx_path = existing_path if str(item.get("extractionMode") or "") == "business_template_extractor_skill" else Path()
    existing_path_allowed = (
        bool(existing_path)
        and (
            _path_is_inside(existing_path, output_dir)
            or _path_is_inside(existing_path, workspace_output_dir)
            or _path_is_inside(existing_path, project_workspace_root)
        )
    )
    if not existing_path_allowed:
        existing_path, workspace_path = _appendix_asset_path(project_id, appendix_id, title)
    else:
        workspace_path = str(item.get("workspacePath") or f"{profile.workspace_dirname}/appendices/{existing_path.name}")
    if (
        extractor_docx_path
        and extractor_docx_path.is_file()
        and existing_path != extractor_docx_path
        and not existing_path.exists()
    ):
        existing_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(str(extractor_docx_path), str(existing_path))

    should_rewrite = False
    if should_rewrite or isinstance(slice_info, dict) or not existing_path.exists():
        sliced = False
        if isinstance(slice_info, dict):
            source_path = Path(str(slice_info.get("sourcePath") or ""))
            keep_start_raw = slice_info.get("keepStart")
            keep_end_raw = slice_info.get("keepEnd")
            source_state = slice_info.get("sourceState")
            if (
                source_path.is_file()
                and isinstance(keep_start_raw, int)
                and isinstance(keep_end_raw, int)
            ):
                sliced = _slice_appendix_from_source(
                    source_path,
                    existing_path,
                    keep_start_raw,
                    keep_end_raw,
                    source_state=source_state if isinstance(source_state, dict) else None,
                )
        if sliced and profile.key == "business":
            item["extractionMode"] = "source_docx_slice"
            quality = _business_template_quality(
                title=title,
                rows=rows,
                content_blocks=content_blocks,
                extraction_mode="source_docx_slice",
                template_section_title=str(item.get("templateSectionTitle") or ""),
            )
            item.update(quality)
        if not sliced:
            if isinstance(slice_info, dict) and profile.key == "business":
                item["extractionMode"] = "source_docx_rebuild_fallback"
                existing_issues = item.get("qualityIssues") if isinstance(item.get("qualityIssues"), list) else []
                item["qualityIssues"] = [*existing_issues, "原 DOCX 切片失败，已退回重建 Word"]
                item["needsReview"] = True
            _write_appendix_docx(existing_path, title, rows, content_blocks)

    item.update(
        {
            "id": appendix_id,
            "title": title,
            "status": "generated",
            "rows": rows,
            "contentBlocks": content_blocks,
            "rowCount": len(rows),
            "docxPath": str(existing_path),
            "workspacePath": workspace_path,
        }
    )
    return item


def materialize_business_commitment_letter_docx(
    project_id: str,
    letter: dict[str, Any],
    *,
    profile: ParseProfile = BUSINESS_PARSE_PROFILE,
    project_name: str = "",
    tenderer_name: str = "",
) -> dict[str, Any]:
    item = copy.deepcopy(letter)
    letter_id = str(item.get("id") or "").strip() or "CL-0000"
    title = str(item.get("title") or item.get("triggerText") or "承诺函").strip() or "承诺函"
    output_dir = _commitment_letter_output_dir(project_id)
    workspace_output_dir = _workspace_commitment_letter_output_dir(project_id, profile)
    project_workspace_root = settings.documents_dir / project_id

    existing_path = Path(str(item.get("docxPath") or ""))
    if not existing_path.is_absolute():
        existing_path = Path()
    existing_path_allowed = (
        bool(existing_path)
        and (
            _path_is_inside(existing_path, output_dir)
            or _path_is_inside(existing_path, workspace_output_dir)
            or _path_is_inside(existing_path, project_workspace_root)
        )
    )
    if not existing_path_allowed:
        existing_path, workspace_path = _commitment_letter_asset_path(project_id, letter_id, title)
    else:
        if str(item.get("workspacePath") or "").strip():
            workspace_path = str(item.get("workspacePath") or "")
        elif _path_is_inside(existing_path, workspace_output_dir):
            workspace_path = f"{profile.workspace_dirname}/commitment-letters/{existing_path.name}"
        else:
            workspace_path = f"s1_commitment_letters/{existing_path.name}"

    if not existing_path.exists():
        _write_commitment_letter_docx(existing_path, item, project_name=project_name, tenderer_name=tenderer_name)

    item.update(
        {
            "id": letter_id,
            "title": title,
            "status": "generated",
            "docxPath": str(existing_path),
            "workspacePath": workspace_path,
            "previewType": "onlyoffice",
        }
    )
    return item


def _appendix_title_for_match(value: str) -> str:
    title = str(value or "").strip().lstrip("#").strip()
    return re.sub(r"\s+", " ", title)


def _appendix_row_count(item: dict[str, Any]) -> int:
    row_count = item.get("rowCount")
    if isinstance(row_count, int):
        return row_count
    rows = item.get("rows")
    return len(rows) if isinstance(rows, list) else 0


def _is_toc_page_number_appendix(item: dict[str, Any], candidates: list[dict[str, Any]]) -> bool:
    if str(item.get("extractionMode") or "") == "source_docx_table_fingerprint":
        return False
    fingerprint = item.get("tableFingerprint") if isinstance(item.get("tableFingerprint"), dict) else {}
    if str(fingerprint.get("type") or "").strip():
        return False
    if _appendix_row_count(item) > 0:
        return False

    title = _appendix_title_for_match(str(item.get("title") or item.get("evidence") or ""))
    if not re.search(r"\d{1,4}$", title):
        return False

    for candidate in candidates:
        if candidate is item:
            continue
        candidate_title = _appendix_title_for_match(str(candidate.get("title") or candidate.get("evidence") or ""))
        if not candidate_title or candidate_title == title or not title.startswith(candidate_title):
            continue

        suffix = title[len(candidate_title):]
        if not re.fullmatch(r"\d{1,4}", suffix):
            continue
        if len(suffix) == 1 and re.search(r"[\w.]$", candidate_title):
            continue
        return True

    return False


def _dedupe_appendix_page_number_artifacts(
    appendices: list[dict[str, Any]],
    *,
    include_attachments: bool = False,
    include_business_templates: bool = False,
) -> list[dict[str, Any]]:
    candidates = [
        item
        for item in appendices
        if isinstance(item, dict) and (
            str(item.get("extractionMode") or "") in {
                "source_docx_table_fingerprint",
                "business_template_extractor_skill",
            }
            or (
                isinstance(item.get("tableFingerprint"), dict)
                and bool(str(item["tableFingerprint"].get("type") or "").strip())
            )
            or
            _looks_like_business_attachment_template_title(
                str(item.get("title") or item.get("evidence") or ""),
                in_template_section=True,
            )
            if include_business_templates
            else _is_appendix_heading(
                str(item.get("title") or item.get("evidence") or ""),
                include_attachments=include_attachments,
            )
        )
    ]
    return [item for item in candidates if not _is_toc_page_number_appendix(item, candidates)]


def _prepare_appendix_outputs(
    project_id: str,
    appendices: list[dict[str, Any]],
    *,
    renumber: bool,
    profile: ParseProfile = TECHNICAL_PARSE_PROFILE,
) -> list[dict[str, Any]]:
    """Materialize appendix docx assets, renumbering IDs sequentially when asked.

    On renumber, an appendix may already have a docx generated under its old ID
    by an earlier pass (typically the slicing pass in ``_extract_docx_appendices``).
    Re-running ``materialize_appendix_docx`` with a cleared ``docxPath`` would
    fall back to a rows-rebuild and lose the carefully preserved formatting.
    Instead, when the existing file is on disk, move it to its new ID-aligned
    path so the materializer sees a valid ``docxPath`` and skips regeneration."""

    prepared: list[dict[str, Any]] = []
    for index, appendix in enumerate(
        _dedupe_appendix_page_number_artifacts(
            appendices,
            include_attachments=profile.key == "business",
            include_business_templates=profile.key == "business",
        ),
        start=1,
    ):
        item = copy.deepcopy(appendix)
        if profile.key == "business":
            rows = item.get("rows") if isinstance(item.get("rows"), list) else []
            content_blocks = item.get("contentBlocks") if isinstance(item.get("contentBlocks"), list) else []
            if (
                str(item.get("extractionMode") or "") != "business_template_extractor_skill"
                and not _business_template_should_materialize(item, rows, content_blocks)
            ):
                continue
        if renumber:
            new_id = f"APPX-{index:04d}"
            old_id = str(item.get("id") or "").strip()
            old_docx = Path(str(item.get("docxPath") or ""))
            old_docx_valid = bool(str(item.get("docxPath") or "")) and old_docx.is_file()
            renamed = False
            if old_docx_valid and old_id and old_id != new_id:
                title = str(item.get("title") or item.get("evidence") or "附表").strip() or "附表"
                new_docx_path, new_workspace_path = _appendix_asset_path(project_id, new_id, title)
                if old_docx.resolve() != new_docx_path.resolve():
                    new_docx_path.parent.mkdir(parents=True, exist_ok=True)
                    try:
                        shutil.move(str(old_docx), str(new_docx_path))
                    except Exception:
                        # Move can fail across filesystems; fall back to copy + unlink.
                        shutil.copy2(str(old_docx), str(new_docx_path))
                        old_docx.unlink(missing_ok=True)
                    item["docxPath"] = str(new_docx_path)
                    item["workspacePath"] = new_workspace_path
                    renamed = True
                else:
                    # Path already matches the new ID — just pin metadata.
                    item["docxPath"] = str(new_docx_path)
                    item["workspacePath"] = new_workspace_path
                    renamed = True
            elif old_docx_valid and old_id == new_id:
                # ID unchanged (re-numbering produced the same value) — keep file as-is.
                renamed = True
            item["id"] = new_id
            if not renamed:
                item["docxPath"] = ""
                item.pop("workspacePath", None)
        prepared.append(materialize_appendix_docx(project_id, item, profile=profile))
    return prepared


def _prepare_commitment_letter_outputs(
    project_id: str,
    letters: list[dict[str, Any]],
    *,
    renumber: bool,
    profile: ParseProfile = BUSINESS_PARSE_PROFILE,
    project_name: str = "",
    tenderer_name: str = "",
) -> list[dict[str, Any]]:
    prepared: list[dict[str, Any]] = []
    for index, letter in enumerate(letters, start=1):
        item = copy.deepcopy(letter)
        if renumber:
            item["id"] = f"CL-{index:04d}"
            item["docxPath"] = ""
            item.pop("workspacePath", None)
        prepared.append(
            materialize_business_commitment_letter_docx(
                project_id,
                item,
                profile=profile,
                project_name=project_name,
                tenderer_name=tenderer_name,
            )
        )
    return prepared


def materialize_parse_appendix_docx_assets(
    project_id: str,
    parse_result: dict[str, Any],
    *,
    bid_type: str,
) -> dict[str, Any]:
    payload = copy.deepcopy(parse_result)
    structured = payload.get("structured")
    if not isinstance(structured, dict):
        return payload
    appendices = structured.get("appendices")
    if not isinstance(appendices, list):
        return payload
    profile = resolve_parse_profile(bid_type)
    structured["appendices"] = _prepare_appendix_outputs(project_id, appendices, renumber=False, profile=profile)
    return payload


def materialize_parse_business_commitment_letter_docx_assets(
    project_id: str,
    parse_result: dict[str, Any],
    *,
    bid_type: str,
) -> dict[str, Any]:
    payload = copy.deepcopy(parse_result)
    structured = payload.get("structured")
    if not isinstance(structured, dict):
        return payload
    if resolve_parse_profile(bid_type).key != "business":
        return payload
    letters = structured.get("commitmentLetters")
    if not isinstance(letters, list):
        return payload
    project_name = _business_project_name_from_structured(structured)
    tenderer_name = _business_tenderer_name_from_structured(structured)
    structured["commitmentLetters"] = _prepare_commitment_letter_outputs(
        project_id,
        letters,
        renumber=False,
        profile=BUSINESS_PARSE_PROFILE,
        project_name=project_name,
        tenderer_name=tenderer_name,
    )
    return payload


def _is_appendix_heading(text: str, *, include_attachments: bool = False) -> bool:
    normalized = str(text or "").strip().lstrip("#").strip()
    if re.match(r"^(?:技术)?[附副]表(?:\s*[A-Za-z0-9一二三四五六七八九十]+)?(?:[：:、.．\s]|$)", normalized):
        return True
    if include_attachments and re.match(r"^附件\s*[A-Za-z0-9一二三四五六七八九十]+(?:[：:、.．\s]|$)", normalized):
        return True
    return False


def _is_scoring_appendix_heading(text: str) -> bool:
    return any(keyword in text for keyword in ("评分标准", "评分细则", "评标办法", "符合性审查", "投标报价评分", "度电成本评分"))


BUSINESS_ATTACHMENT_TEMPLATE_TOPICS = (
    "投标函", "法定代表人", "单位负责人", "身份证明", "授权书", "授权委托书",
    "廉洁", "投标专用章", "效力说明", "投标价格", "开标价格", "商务偏差",
    "货物规格", "规格表", "供货范围", "投标保证金", "履约保证函", "履约承诺",
    "资格证明", "合格投标人", "资格履行合同", "业绩情况", "业绩表",
    "财务状况", "制造商授权", "联合体协议", "分包", "其他内容", "其他说明",
    "否决项", "承诺书", "承诺函",
)

BUSINESS_ATTACHMENT_TEMPLATE_CONTEXT_TOKENS = (
    "投标文件格式",
    "响应文件格式",
    "商务文件格式",
    "投标文件组成",
    "第六章",
    "第6章",
)

BUSINESS_FORMAT_START_KEYWORDS = (
    "投标文件格式",
    "响应文件格式",
    "商务文件格式",
    "商务标格式",
    "商务响应文件格式",
)
BUSINESS_FORMAT_END_KEYWORDS = (
    "评标办法",
    "评审办法",
    "合同条款",
    "技术要求",
    "技术规范",
    "供货要求",
    "用户需求",
)

BUSINESS_ATTACHMENT_TEMPLATE_TYPE_RULES: tuple[tuple[str, str, tuple[str, ...]], ...] = (
    ("bid_letter", "投标函", ("投标函",)),
    ("legal_rep_id", "法定代表人身份证明", ("法定代表人（单位负责人）身份证明", "法定代表人身份证明", "单位负责人身份证明", "身份证明")),
    ("authorization", "授权书", ("法定代表人授权委托书", "法定代表人授权书", "授权委托书", "授权书")),
    ("integrity_commitment", "廉洁承诺", ("投标人廉洁自律承诺书", "廉洁自律承诺书", "廉洁承诺书", "廉洁承诺")),
    ("seal_validity", "投标专用章效力说明", ("投标专用章效力说明", "专用章效力说明", "投标专用章")),
    ("bid_price", "投标价格表", ("投标价格表", "投标价格")),
    ("opening_price", "开标价格表", ("开标价格表", "开标报价表")),
    ("commercial_deviation", "商务偏差表", ("商务偏差表", "商务偏差", "商务偏离表", "偏差表", "偏离表")),
    ("specification", "货物规格表", ("货物规格表", "货物规格", "规格表")),
    ("supply_scope", "供货范围表", ("供货范围表", "供货范围", "供货清单")),
    ("bid_security", "投标保证金", ("投标保证金", "保证金", "投标保函", "保函")),
    ("performance_bond", "履约保证", ("履约保证函格式承诺书", "履约保证函", "履约承诺书", "履约承诺", "履约保证")),
    ("qualification", "资格证明文件", ("资格证明", "合格投标人", "资格履行合同")),
    ("performance_table", "业绩情况表", ("业绩情况表", "业绩表", "业绩证明")),
    ("financial", "财务状况表", ("财务状况", "财务报表", "审计报告")),
    ("manufacturer_authorization", "制造商授权", ("制造商授权", "厂家授权")),
    ("joint_venture", "联合体协议", ("联合体协议",)),
    ("other_notes", "其他说明", ("投标人需要说明的其他内容", "其他内容", "其他说明")),
    ("commitment", "承诺函/承诺书", ("承诺函", "承诺书")),
)

BUSINESS_TEMPLATE_COMPLETENESS_TOKENS = (
    "盖章",
    "签字",
    "签章",
    "签名",
    "年月日",
    "年  月  日",
    "身份证号",
    "报价",
    "金额",
    "偏差",
    "保证金",
    "开户行",
    "账号",
)
BUSINESS_TEMPLATE_PLACEHOLDER_TOKENS = ("____", "＿＿", "（盖章）", "(盖章)", "签字", "日期", "年  月  日")
BUSINESS_TEMPLATE_SEMANTIC_REVIEW_MAX_ITEMS = 16

BUSINESS_TABLE_FINGERPRINTS: tuple[tuple[str, str, tuple[tuple[str, ...], ...]], ...] = (
    ("opening_price", "开标价格表", (("序号",), ("投标报价", "报价", "总价", "合价", "单价"), ("项目名称", "名称", "货物名称"))),
    ("bid_price", "投标价格表", (("序号",), ("投标报价", "报价", "总价", "合价", "单价"), ("备注", "说明", "项目名称"))),
    ("specification", "货物规格表", (("序号",), ("货物名称", "设备名称", "名称"), ("规格", "规格型号", "型号"))),
    ("commercial_deviation", "商务偏差表", (("序号",), ("条款", "招标文件"), ("偏差", "响应"))),
    ("supply_scope", "供货范围表", (("序号",), ("名称", "设备", "货物"), ("数量", "单位"))),
    ("performance_table", "业绩情况表", (("序号",), ("项目名称", "工程名称"), ("合同", "容量", "投运", "业主"))),
    ("legal_rep_id", "法定代表人身份证明", (("姓名",), ("身份证", "证件"), ("职务", "职称"))),
    ("authorization", "授权书", (("姓名",), ("身份证", "证件"), ("授权", "代理人", "委托"))),
    ("bid_security", "投标保证金", (("保证金", "保函"), ("金额", "账号", "开户行", "银行"))),
)


def _is_business_template_section_heading(text: str) -> bool:
    normalized = str(text or "").strip().lstrip("#").strip()
    if not normalized:
        return False
    if _looks_like_toc_or_directory_line(normalized):
        return False
    if _looks_like_business_template_body_sentence(normalized):
        return False
    if any(token in normalized for token in BUSINESS_FORMAT_START_KEYWORDS):
        return True
    return bool(re.match(r"^第[六6]章", normalized)) and any(
        token in normalized for token in ("格式", "投标文件", "响应文件", "商务文件")
    )


def _is_business_format_end_heading(text: str) -> bool:
    normalized = str(text or "").strip().lstrip("#").strip()
    if not normalized or _looks_like_toc_or_directory_line(normalized):
        return False
    return any(keyword in normalized for keyword in BUSINESS_FORMAT_END_KEYWORDS)


def _block_heading_rank(block: dict[str, Any], text: str) -> int | None:
    raw_level = block.get("headingLevel")
    if isinstance(raw_level, int):
        return raw_level
    normalized = str(text or "").strip().lstrip("#").strip()
    if re.match(r"^第[一二三四五六七八九十0-9]+章", normalized):
        return 0
    if re.match(r"^(?:[一二三四五六七八九十0-9]+[、.．]|[（(][一二三四五六七八九十0-9]+[）)])", normalized):
        return 1
    if bool(block.get("isLikelyHeading")):
        return 2
    return None


def _is_business_major_section_heading(text: str) -> bool:
    normalized = str(text or "").strip().lstrip("#").strip()
    return bool(re.match(r"^第[一二三四五六七八九十0-9]+章", normalized))


def _has_business_attachment_topic(text: str) -> bool:
    normalized = str(text or "").strip()
    return any(keyword in normalized for keyword in BUSINESS_ATTACHMENT_TEMPLATE_TOPICS)


def _business_template_type(title: str) -> tuple[str, str]:
    normalized = re.sub(r"\s+", "", str(title or ""))
    for key, label, keywords in BUSINESS_ATTACHMENT_TEMPLATE_TYPE_RULES:
        if any(keyword and keyword in normalized for keyword in keywords):
            return key, label
    return "generic_attachment", "商务附件模板"


def _business_template_type_label(template_type: str) -> str:
    normalized = str(template_type or "").strip()
    for key, label, _ in BUSINESS_ATTACHMENT_TEMPLATE_TYPE_RULES:
        if key == normalized:
            return label
    return "商务附件模板"


def _business_template_type_label_for_key(template_type: str) -> str:
    normalized = str(template_type or "").strip()
    for key, label, _ in BUSINESS_ATTACHMENT_TEMPLATE_TYPE_RULES:
        if key == normalized:
            return label
    for key, label, _ in BUSINESS_TABLE_FINGERPRINTS:
        if key == normalized:
            return label
    return "商务附件模板"


def _flatten_table_text(rows: list[list[str]]) -> str:
    return re.sub(r"\s+", "", " ".join(str(cell or "") for row in rows for cell in row))


def _classify_business_table(rows: list[list[str]], title_hint: str = "") -> tuple[str, str, float, list[str]]:
    table_text = _flatten_table_text(rows)
    title_text = re.sub(r"\s+", "", str(title_hint or ""))
    combined = f"{title_text}{table_text}"
    if not combined:
        return "", "", 0.0, []
    candidates: list[tuple[float, str, str, list[str]]] = []
    for key, label, groups in BUSINESS_TABLE_FINGERPRINTS:
        matched_groups = []
        for group in groups:
            matched = next((token for token in group if token and token in combined), "")
            if matched:
                matched_groups.append(matched)
        if len(matched_groups) >= max(2, len(groups) - 1):
            confidence = min(0.98, 0.62 + 0.12 * len(matched_groups))
            if any(token in title_text for token in (label, label.replace("表", ""), "报价", "规格", "偏差", "供货", "业绩")):
                confidence = min(0.99, confidence + 0.04)
            candidates.append((confidence, key, label, matched_groups))
    if candidates:
        confidence, key, label, matched_groups = max(candidates, key=lambda item: (item[0], len(item[3])))
        return key, label, confidence, matched_groups
    if any(token in combined for token in BUSINESS_TEMPLATE_PLACEHOLDER_TOKENS):
        return "generic_attachment", "商务附件模板", 0.58, ["placeholder"]
    return "", "", 0.0, []


def _business_template_title_allowed(text: str, *, in_template_section: bool = False) -> bool:
    normalized = str(text or "").strip().lstrip("#").strip()
    if _looks_like_toc_or_directory_line(normalized):
        return False
    if not _looks_like_business_attachment_template_title(normalized, in_template_section=in_template_section):
        return False
    if in_template_section:
        return True
    has_explicit_template = any(token in normalized for token in ("格式", "模板", "样式"))
    return has_explicit_template and _is_appendix_heading(normalized, include_attachments=True)


def _looks_like_business_attachment_template_title(text: str, *, in_template_section: bool = False) -> bool:
    normalized = str(text or "").strip().lstrip("#").strip()
    if _looks_like_business_template_body_sentence(normalized):
        return False
    if _is_scoring_appendix_heading(normalized):
        return False
    has_appendix_prefix = _is_appendix_heading(normalized, include_attachments=True)
    has_template_suffix = any(token in normalized for token in ("格式", "模板", "样式"))
    has_response_format_title = _is_business_template_section_heading(normalized)
    if has_response_format_title:
        return False
    has_numbered_business_prefix = bool(
        re.match(r"^(?:[一二三四五六七八九十0-9]+[、.．]|[（(][一二三四五六七八九十0-9]+[）)])", normalized)
    )
    has_business_topic = _has_business_attachment_topic(normalized)
    if has_appendix_prefix and (has_business_topic or has_template_suffix or in_template_section):
        return True
    if in_template_section and has_business_topic and (has_numbered_business_prefix or has_template_suffix or len(normalized) <= 24):
        return True
    return has_business_topic and has_template_suffix


def _business_template_quality(
    *,
    title: str,
    rows: list[list[str]],
    content_blocks: list[dict[str, Any]],
    extraction_mode: str,
    template_section_title: str,
) -> dict[str, Any]:
    paragraph_text = " ".join(
        str(block.get("text") or "").strip()
        for block in content_blocks
        if isinstance(block, dict) and block.get("type") == "paragraph"
    )
    table_count = sum(
        1
        for block in content_blocks
        if isinstance(block, dict) and block.get("type") == "table" and isinstance(block.get("rows"), list) and block.get("rows")
    )
    row_count = len(rows) if rows else sum(
        len(block.get("rows") or [])
        for block in content_blocks
        if isinstance(block, dict) and block.get("type") == "table" and isinstance(block.get("rows"), list)
    )
    text_for_quality = f"{title} {paragraph_text} " + " ".join(" ".join(row) for row in rows if isinstance(row, list))
    issues: list[str] = []
    if not template_section_title:
        issues.append("未定位到投标文件格式/响应文件格式章节")
    if not _appendix_has_material_content(rows, content_blocks):
        issues.append("正文内容过少")
    if extraction_mode != "source_docx_slice" and str(extraction_mode).startswith("source_docx"):
        issues.append("原 DOCX 切片未成功")
    has_placeholder = any(token in text_for_quality for token in BUSINESS_TEMPLATE_PLACEHOLDER_TOKENS)
    has_completeness_token = any(token in text_for_quality for token in BUSINESS_TEMPLATE_COMPLETENESS_TOKENS)
    if issues:
        quality = "title_only" if "正文内容过少" in issues else "needs_review"
    elif table_count > 0 or row_count > 0 or (len(paragraph_text) >= 20 and (has_placeholder or has_completeness_token)):
        quality = "complete"
    else:
        quality = "probably_incomplete"
        issues.append("未识别到明显表格、签章栏或待填写占位")
    return {
        "extractionQuality": quality,
        "qualityIssues": issues,
        "needsReview": quality != "complete" or bool(issues),
        "contentStats": {
            "paragraphTextLength": len(paragraph_text),
            "tableCount": table_count,
            "rowCount": row_count,
            "hasPlaceholder": has_placeholder,
        },
    }


def _business_template_metadata(
    *,
    title: str,
    source_file: str,
    evidence: str,
    evidence_location: str,
    template_section_title: str,
    source_start: str,
    source_end: str,
    rows: list[list[str]],
    content_blocks: list[dict[str, Any]],
    extraction_mode: str,
) -> dict[str, Any]:
    template_type, template_type_label = _business_template_type(title)
    table_type, table_type_label, table_confidence, matched_tokens = _classify_business_table(rows, title)
    if table_type and template_type == "generic_attachment":
        template_type, template_type_label = table_type, table_type_label
    quality = _business_template_quality(
        title=title,
        rows=rows,
        content_blocks=content_blocks,
        extraction_mode=extraction_mode,
        template_section_title=template_section_title,
    )
    return {
        "artifactType": "business_attachment_template",
        "sourceMode": "parsed_from_tender_attachment_template",
        "templateType": template_type,
        "templateTypeLabel": template_type_label,
        "tableFingerprint": {
            "type": table_type,
            "typeLabel": table_type_label,
            "confidence": table_confidence,
            "matchedTokens": matched_tokens,
        },
        "templateSectionTitle": template_section_title,
        "templateSectionDetected": bool(template_section_title),
        "sourceFile": source_file,
        "evidence": evidence,
        "evidenceLocation": evidence_location,
        "sourceStart": source_start,
        "sourceEnd": source_end,
        "extractionMode": extraction_mode,
        "assetReviewStatus": "pending_review",
        "assetSyncStatus": "pending",
        "previewType": "onlyoffice",
        **quality,
    }


def _build_business_template_review_prompt(appendices: list[dict[str, Any]]) -> str:
    records = []
    for item in appendices[:BUSINESS_TEMPLATE_SEMANTIC_REVIEW_MAX_ITEMS]:
        content_blocks = item.get("contentBlocks") if isinstance(item.get("contentBlocks"), list) else []
        text_preview_parts: list[str] = []
        for block in content_blocks:
            if not isinstance(block, dict):
                continue
            if block.get("type") == "paragraph":
                text_preview_parts.append(str(block.get("text") or "").strip())
            elif block.get("type") == "table":
                rows = block.get("rows") if isinstance(block.get("rows"), list) else []
                text_preview_parts.append(" | ".join(" / ".join(str(cell) for cell in row) for row in rows[:3] if isinstance(row, list)))
        records.append(
            {
                "id": str(item.get("id") or ""),
                "title": str(item.get("title") or ""),
                "templateType": str(item.get("templateType") or ""),
                "templateSectionTitle": str(item.get("templateSectionTitle") or ""),
                "sourceStart": str(item.get("sourceStart") or ""),
                "sourceEnd": str(item.get("sourceEnd") or ""),
                "extractionMode": str(item.get("extractionMode") or ""),
                "extractionQuality": str(item.get("extractionQuality") or ""),
                "qualityIssues": item.get("qualityIssues") if isinstance(item.get("qualityIssues"), list) else [],
                "textPreview": "\n".join(part for part in text_preview_parts if part)[:1200],
            }
        )
    return (
        "你在做商务标解析阶段的附件模板语义校验。请判断每个候选是否确实是投标文件格式章节中可用于后续商务投标文件的模板/表格/承诺书格式。\n"
        "只输出 JSON，格式为：\n"
        "{\n"
        '  "decisions": [\n'
        '    {"id":"APPX-0001","action":"accept|review|reject","templateType":"bid_letter","quality":"complete|probably_incomplete|title_only","reason":"一句简短原因"}\n'
        "  ]\n"
        "}\n"
        "判断规则：\n"
        "1. 第六章、投标文件格式、响应文件格式、商务文件格式中的投标函、授权书、廉洁承诺、报价表、偏差表、资格/业绩表等通常 action=accept。\n"
        "2. 只有标题没有正文或表格的候选 action=review，quality=title_only，不要直接 reject。\n"
        "3. 普通条款、目录页、评分标准、技术参数承诺、非投标文件模板正文 action=reject。\n"
        "4. templateType 可按语义修正，例如 bid_letter、authorization、opening_price、commercial_deviation、performance_table、commitment、generic_attachment。\n\n"
        f"候选列表：\n{json.dumps(records, ensure_ascii=False, indent=2)}"
    )


def _review_business_attachment_templates_semantically(appendices: list[dict[str, Any]]) -> dict[str, dict[str, Any]]:
    candidates = [item for item in appendices if isinstance(item, dict) and item.get("artifactType") == "business_attachment_template"]
    if not candidates:
        return {}
    try:
        result = OpencodeClient().review_business_attachment_templates_with_trace(
            _build_business_template_review_prompt(candidates)
        )
    except RuntimeError:
        return {}
    decisions = result.get("decisions")
    if not isinstance(decisions, list):
        return {}
    reviewed: dict[str, dict[str, Any]] = {}
    for decision in decisions:
        if not isinstance(decision, dict):
            continue
        item_id = str(decision.get("id") or "").strip()
        if item_id:
            reviewed[item_id] = decision
    return reviewed


def _apply_business_template_semantic_review(
    appendices: list[dict[str, Any]],
    *,
    run_semantic_review: bool,
) -> list[dict[str, Any]]:
    if not appendices:
        return appendices
    reviewed = _review_business_attachment_templates_semantically(appendices) if run_semantic_review else {}
    prepared: list[dict[str, Any]] = []
    for item in appendices:
        if not isinstance(item, dict) or item.get("artifactType") != "business_attachment_template":
            prepared.append(item)
            continue
        result = copy.deepcopy(item)
        decision = reviewed.get(str(result.get("id") or ""))
        if isinstance(decision, dict):
            action = str(decision.get("action") or "").strip().lower()
            if action == "reject":
                result["semanticReviewStatus"] = "rejected"
                result["needsReview"] = True
                issues = result.get("qualityIssues") if isinstance(result.get("qualityIssues"), list) else []
                result["qualityIssues"] = [*issues, str(decision.get("reason") or "AI 判断该候选不是可用商务附件模板")]
                continue
            if action in {"accept", "review"}:
                result["semanticReviewStatus"] = action
                if str(decision.get("templateType") or "").strip():
                    result["templateType"] = str(decision.get("templateType") or "").strip()
                    result["templateTypeLabel"] = _business_template_type_label(str(result.get("templateType") or ""))
                if str(decision.get("quality") or "").strip():
                    result["extractionQuality"] = str(decision.get("quality") or "").strip()
                if action == "review" or result.get("extractionQuality") != "complete":
                    result["needsReview"] = True
                    issues = result.get("qualityIssues") if isinstance(result.get("qualityIssues"), list) else []
                    reason = str(decision.get("reason") or "").strip()
                    result["qualityIssues"] = [*issues, reason] if reason and reason not in issues else issues
                else:
                    result["needsReview"] = False
        else:
            result.setdefault("semanticReviewStatus", "not_run" if not run_semantic_review else "unavailable")
        prepared.append(result)
    return prepared


def _business_template_source_end_from_lines(lines: list[str], next_index: int) -> str:
    if 0 <= next_index < len(lines):
        return f"L{next_index + 1}: {str(lines[next_index]).strip()}"
    return "EOF"


def _business_template_source_start(section_title: str, location: str, text: str) -> str:
    start = f"{location}: {str(text or '').strip()}"
    section = str(section_title or "").strip()
    return f"{section} / {start}" if section else start


def _business_template_source_end_from_blocks(blocks: list[dict[str, Any]], next_index: int) -> str:
    if 0 <= next_index < len(blocks):
        block = blocks[next_index]
        return f"B{next_index + 1}: {str(block.get('text') or '').strip()}"
    return "EOF"


def _business_template_previous_state(blocks: list[dict[str, Any]], start_index: int) -> tuple[bool, str]:
    in_template_section = False
    template_section_title = ""
    for previous in range(0, min(start_index + 1, len(blocks))):
        block = blocks[previous]
        if block.get("type") != "paragraph":
            continue
        text = str(block.get("text") or "").strip()
        if _is_business_template_section_heading(text):
            in_template_section = True
            template_section_title = text
        elif in_template_section and _is_business_major_section_heading(text):
            in_template_section = False
            template_section_title = ""
    return in_template_section, template_section_title


def _detect_business_format_regions(blocks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    regions: list[dict[str, Any]] = []
    active: dict[str, Any] | None = None
    for index, block in enumerate(blocks):
        if block.get("type") != "paragraph":
            continue
        text = str(block.get("text") or "").strip()
        if not text:
            continue
        rank = _block_heading_rank(block, text)
        is_start = _is_business_template_section_heading(text)
        if is_start:
            if active is not None:
                active["end"] = max(active["start"] + 1, index)
                regions.append(active)
            active = {
                "start": index,
                "end": len(blocks),
                "title": text,
                "rank": rank if rank is not None else 0,
            }
            continue
        if active is None:
            continue
        active_rank = int(active.get("rank") if isinstance(active.get("rank"), int) else 0)
        same_or_higher_heading = rank is not None and rank <= active_rank
        if same_or_higher_heading and (_is_business_format_end_heading(text) or _is_business_major_section_heading(text)):
            active["end"] = index
            regions.append(active)
            active = None
    if active is not None:
        regions.append(active)
    return regions


def _business_region_for_index(regions: list[dict[str, Any]], index: int) -> dict[str, Any] | None:
    for region in regions:
        start = region.get("start")
        end = region.get("end")
        if isinstance(start, int) and isinstance(end, int) and start <= index < end:
            return region
    return None


def _is_business_toc_or_directory_context(blocks: list[dict[str, Any]], index: int) -> bool:
    line = str(blocks[index].get("text") or "").strip()
    if _looks_like_toc_or_directory_line(line):
        return True

    for previous in range(index - 1, max(-1, index - 9), -1):
        previous_block = blocks[previous]
        if previous_block.get("type") != "paragraph":
            continue
        text = str(previous_block.get("text") or "").strip()
        if not text:
            continue
        if _is_business_template_section_heading(text) or _is_business_major_section_heading(text):
            return False
        if re.fullmatch(r"(?:目\s*录|目录|contents)", text, flags=re.IGNORECASE):
            return True
    return False


def _last_body_index_before(blocks: list[dict[str, Any]], end_index: int, fallback: int | None) -> int | None:
    for previous in range(min(end_index, len(blocks)) - 1, -1, -1):
        raw_body_index = blocks[previous].get("body_index")
        if isinstance(raw_body_index, int):
            return raw_body_index
    return fallback


def _is_relevant_appendix_heading(text: str, profile: ParseProfile) -> bool:
    if profile.key == "business":
        return _business_template_title_allowed(text)
    return _is_appendix_heading(text)


def _next_appendix_heading_index(blocks: list[dict[str, Any]], start_index: int, profile: ParseProfile) -> int:
    in_template_section = False
    if profile.key == "business":
        in_template_section, _ = _business_template_previous_state(blocks, start_index)
    for lookahead in range(start_index + 1, len(blocks)):
        next_block = blocks[lookahead]
        if next_block.get("type") != "paragraph":
            continue
        next_text = str(next_block.get("text") or "")
        if profile.key == "business":
            if _is_business_template_section_heading(next_text) or _is_business_major_section_heading(next_text):
                return lookahead
            if _business_template_title_allowed(next_text, in_template_section=in_template_section):
                return lookahead
        elif _is_relevant_appendix_heading(next_text, profile):
            return lookahead
    return len(blocks)


def _slice_content_blocks(blocks: list[dict[str, Any]], start_index: int, end_index: int) -> list[dict[str, Any]]:
    content: list[dict[str, Any]] = []
    for block in blocks[start_index + 1:end_index]:
        if block.get("type") == "paragraph":
            text = str(block.get("text") or "").strip()
            if text:
                content.append({"type": "paragraph", "text": text})
        elif block.get("type") == "table":
            rows = block.get("rows") if isinstance(block.get("rows"), list) else []
            if rows:
                content.append({"type": "table", "rows": rows})
    return content


def _previous_business_title_block_index(blocks: list[dict[str, Any]], table_index: int, region: dict[str, Any] | None) -> int | None:
    start = int(region.get("start")) if isinstance(region, dict) and isinstance(region.get("start"), int) else 0
    for previous in range(table_index - 1, max(start, table_index - 8) - 1, -1):
        block = blocks[previous]
        if block.get("type") != "paragraph":
            continue
        text = str(block.get("text") or "").strip()
        if not text or _looks_like_toc_or_directory_line(text):
            continue
        if _business_template_title_allowed(text, in_template_section=True):
            return previous
        if bool(block.get("isLikelyHeading")) and _has_business_attachment_topic(text):
            return previous
        if len(re.sub(r"\s+", "", text)) <= 36 and not _looks_like_business_template_body_sentence(text):
            return previous
    return None


def _appendix_title_has_trailing_page_number(text: str) -> bool:
    title = _appendix_title_for_match(text)
    return bool(re.search(r"\D\d{2,4}$", title))


def _is_docx_appendix_toc_artifact(blocks: list[dict[str, Any]], index: int) -> bool:
    line = str(blocks[index].get("text") or "").strip()
    if not _appendix_title_has_trailing_page_number(line):
        return False

    for lookahead in range(index + 1, min(len(blocks), index + 8)):
        next_block = blocks[lookahead]
        if next_block.get("type") == "table":
            return False
        if next_block.get("type") == "paragraph" and _is_appendix_heading(str(next_block.get("text") or "")):
            break

    nearby_headings = 0
    for nearby in range(max(0, index - 3), min(len(blocks), index + 4)):
        if nearby == index:
            continue
        block = blocks[nearby]
        if block.get("type") == "paragraph" and _is_appendix_heading(str(block.get("text") or "")):
            nearby_headings += 1

    return nearby_headings >= 1


def _extract_markdown_appendices(
    project_id: str,
    documents: list[dict[str, Any]],
    texts_by_id: dict[str, str],
    *,
    start_index: int = 0,
    profile: ParseProfile = TECHNICAL_PARSE_PROFILE,
) -> list[dict[str, Any]]:
    appendices: list[dict[str, Any]] = []
    for document in documents:
        document_id = str(document.get("id") or "")
        source_file = str(document.get("name") or document_id or "招标文件")
        source_path = Path(str(document.get("sourcePath") or ""))
        if source_path.suffix.lower() != ".md":
            continue
        lines = texts_by_id.get(document_id, "").splitlines()
        index = 0
        in_template_section = False
        template_section_title = ""
        while index < len(lines):
            line = lines[index].strip()
            if profile.key == "business" and _is_business_template_section_heading(line):
                in_template_section = True
                template_section_title = line.strip(" #")
                index += 1
                continue
            if profile.key == "business" and in_template_section and _is_business_major_section_heading(line):
                in_template_section = False
                template_section_title = ""
            if profile.key == "business":
                is_heading = _business_template_title_allowed(
                    line,
                    in_template_section=in_template_section,
                )
            else:
                is_heading = _is_relevant_appendix_heading(line, profile)
            if not is_heading:
                index += 1
                continue
            if profile.key == "business" and not in_template_section:
                index += 1
                continue
            if _is_scoring_appendix_heading(line):
                index += 1
                continue

            title = line.strip(" #")
            table_start = index + 1
            while table_start < len(lines) and not MARKDOWN_TABLE_LINE_PATTERN.match(lines[table_start]):
                if profile.key != "business" and lines[table_start].strip() and not _is_relevant_appendix_heading(lines[table_start], profile):
                    break
                if (
                    profile.key == "business"
                    and lines[table_start].strip()
                    and (
                        _is_business_major_section_heading(lines[table_start])
                        or _business_template_title_allowed(
                            lines[table_start],
                            in_template_section=in_template_section,
                        )
                    )
                ):
                    break
                table_start += 1
            if profile.key == "business":
                next_heading = len(lines)
                for lookahead in range(index + 1, len(lines)):
                    lookahead_line = lines[lookahead].strip()
                    if (
                        _is_business_template_section_heading(lookahead_line)
                        or _is_business_major_section_heading(lookahead_line)
                        or _business_template_title_allowed(
                            lookahead_line,
                            in_template_section=in_template_section,
                        )
                    ):
                        next_heading = lookahead
                        break
            else:
                next_heading = table_start + 1 if table_start < len(lines) else index + 1
            if table_start >= len(lines) or not MARKDOWN_TABLE_LINE_PATTERN.match(lines[table_start]):
                content_blocks = [
                    {"type": "paragraph", "text": text.strip()}
                    for text in lines[index + 1:next_heading]
                    if text.strip()
                ]
                metadata = {}
                if profile.key == "business":
                    metadata = _business_template_metadata(
                        title=title,
                        source_file=source_file,
                        evidence=line,
                        evidence_location=f"L{index + 1}",
                        template_section_title=template_section_title,
                        source_start=_business_template_source_start(template_section_title, f"L{index + 1}", line),
                        source_end=_business_template_source_end_from_lines(lines, next_heading),
                        rows=[],
                        content_blocks=content_blocks,
                        extraction_mode="markdown_slice",
                    )
                    if not _business_template_should_materialize(metadata, [], content_blocks):
                        index = max(index + 1, next_heading)
                        continue
                appendix_id = f"APPX-{start_index + len(appendices) + 1:04d}"
                appendices.append(materialize_appendix_docx(
                    project_id,
                    {
                        "id": appendix_id,
                        "title": title,
                        "status": "generated",
                        "rows": [],
                        "contentBlocks": content_blocks,
                        "rowCount": 0,
                        "docxPath": "",
                        **metadata,
                    },
                    profile=profile,
                ))
                index = max(index + 1, next_heading)
                continue

            rows: list[list[str]] = []
            table_end = table_start
            table_limit = next_heading if profile.key == "business" else len(lines)
            while table_end < table_limit and MARKDOWN_TABLE_LINE_PATTERN.match(lines[table_end]):
                cells = _parse_markdown_table_row(lines[table_end])
                if not _is_markdown_separator_row(cells):
                    rows.append(cells)
                table_end += 1
            content_blocks = [
                {"type": "paragraph", "text": text.strip()}
                for text in lines[index + 1:table_start]
                if text.strip()
            ]
            if rows:
                content_blocks.append({"type": "table", "rows": rows})
            if profile.key == "business":
                content_blocks.extend(
                    {"type": "paragraph", "text": text.strip()}
                    for text in lines[table_end:next_heading]
                    if text.strip()
                )
            metadata = (
                _business_template_metadata(
                    title=title,
                    source_file=source_file,
                    evidence=line,
                    evidence_location=f"L{index + 1}",
                    template_section_title=template_section_title,
                    source_start=_business_template_source_start(template_section_title, f"L{index + 1}", line),
                    source_end=_business_template_source_end_from_lines(lines, next_heading),
                    rows=rows,
                    content_blocks=content_blocks,
                    extraction_mode="markdown_slice",
                )
                if profile.key == "business"
                else {}
            )
            if profile.key == "business" and not _business_template_should_materialize(metadata, rows, content_blocks):
                index = max(table_end, next_heading)
                continue

            appendix_id = f"APPX-{start_index + len(appendices) + 1:04d}"
            appendices.append(materialize_appendix_docx(
                project_id,
                {
                    "id": appendix_id,
                    "title": title,
                    "status": "generated",
                    "rows": rows,
                    "contentBlocks": content_blocks,
                    "rowCount": len(rows),
                    "docxPath": "",
                    **metadata,
                },
                profile=profile,
            ))
            index = max(table_end, next_heading)
    return appendices


def _docx_paragraph_text(element: Any) -> str:
    return "".join(node.text or "" for node in element.iter(f"{WORD_NAMESPACE}t")).strip()


def _docx_xml_attr(element: Any, name: str) -> str:
    if element is None:
        return ""
    return str(element.get(f"{WORD_NAMESPACE}{name}") or element.get(name) or "").strip()


def _docx_paragraph_xml_outline_level(element: Any) -> int | None:
    p_pr = element.find(f"{WORD_NAMESPACE}pPr")
    if p_pr is None:
        return None
    outline = p_pr.find(f"{WORD_NAMESPACE}outlineLvl")
    if outline is not None:
        raw = _docx_xml_attr(outline, "val")
        if raw.isdigit():
            return int(raw)
    p_style = p_pr.find(f"{WORD_NAMESPACE}pStyle")
    raw_style = _docx_xml_attr(p_style, "val")
    match = re.search(r"(?:Heading|标题)\s*([1-9])", raw_style, flags=re.IGNORECASE)
    if match:
        return max(0, int(match.group(1)) - 1)
    return None


def _docx_paragraph_alignment_value(element: Any, paragraph: Any) -> str:
    alignment = getattr(paragraph, "alignment", None) if paragraph is not None else None
    if alignment is not None:
        value = getattr(alignment, "value", alignment)
        if str(value) in {"1", "CENTER", "WD_ALIGN_PARAGRAPH.CENTER"}:
            return "center"
    p_pr = element.find(f"{WORD_NAMESPACE}pPr")
    jc = p_pr.find(f"{WORD_NAMESPACE}jc") if p_pr is not None else None
    raw = _docx_xml_attr(jc, "val").lower()
    return raw


def _docx_paragraph_style_level(style_name: str) -> int | None:
    normalized = str(style_name or "").strip()
    match = re.search(r"(?:Heading|标题)\s*([1-9])", normalized, flags=re.IGNORECASE)
    if match:
        return max(0, int(match.group(1)) - 1)
    return None


def _docx_paragraph_metadata(element: Any, paragraph: Any, text: str) -> dict[str, Any]:
    style_name = str(getattr(getattr(paragraph, "style", None), "name", "") or "")
    outline_level = _docx_paragraph_xml_outline_level(element)
    style_level = _docx_paragraph_style_level(style_name)
    heading_level = outline_level if outline_level is not None else style_level
    alignment = _docx_paragraph_alignment_value(element, paragraph)
    runs = list(getattr(paragraph, "runs", []) or []) if paragraph is not None else []
    non_empty_runs = [run for run in runs if str(getattr(run, "text", "") or "").strip()]
    bold_runs = [run for run in non_empty_runs if bool(getattr(getattr(run, "font", None), "bold", False) or getattr(run, "bold", False))]
    font_sizes = []
    for run in non_empty_runs:
        size = getattr(getattr(run, "font", None), "size", None)
        if size is not None and getattr(size, "pt", None):
            font_sizes.append(float(size.pt))
    bold_ratio = (len(bold_runs) / len(non_empty_runs)) if non_empty_runs else 0.0
    normalized = re.sub(r"\s+", "", str(text or ""))
    is_short = 0 < len(normalized) <= 48
    is_likely_heading = bool(
        heading_level is not None
        or style_name.lower().startswith("heading")
        or "标题" in style_name
        or (alignment == "center" and is_short)
        or (bold_ratio >= 0.6 and is_short)
    )
    return {
        "styleName": style_name,
        "outlineLevel": outline_level,
        "styleLevel": style_level,
        "headingLevel": heading_level,
        "alignment": alignment,
        "isCentered": alignment == "center",
        "boldRatio": round(bold_ratio, 3),
        "fontSizeMax": max(font_sizes) if font_sizes else None,
        "isLikelyHeading": is_likely_heading,
    }


def _docx_table_rows(table: Any) -> list[list[str]]:
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _iter_docx_blocks(path: Path) -> list[dict[str, Any]]:
    """Walk the docx body and yield {paragraph,table} blocks.

    Each block records ``body_index`` — its position among ``body.iterchildren()`` —
    so callers can map a block back to the original ``<w:p>`` / ``<w:tbl>`` element
    when slicing the source docx (see ``_slice_appendix_from_source``)."""

    doc = Document(str(path))
    tables = iter(doc.tables)
    paragraphs = iter(doc.paragraphs)
    blocks: list[dict[str, Any]] = []
    for body_index, child in enumerate(doc.element.body.iterchildren()):
        if child.tag == f"{WORD_NAMESPACE}p":
            paragraph = next(paragraphs, None)
            text = _docx_paragraph_text(child)
            blocks.append(
                {
                    "type": "paragraph",
                    "text": text,
                    "body_index": body_index,
                    **_docx_paragraph_metadata(child, paragraph, text),
                }
            )
        elif child.tag == f"{WORD_NAMESPACE}tbl":
            table = next(tables, None)
            if table is not None:
                blocks.append(
                    {
                        "type": "table",
                        "rows": _docx_table_rows(table),
                        "body_index": body_index,
                    }
                )
    return blocks


def _extract_docx_appendices(
    project_id: str,
    documents: list[dict[str, Any]],
    *,
    start_index: int = 0,
    profile: ParseProfile = TECHNICAL_PARSE_PROFILE,
) -> list[dict[str, Any]]:
    appendices: list[dict[str, Any]] = []
    for document in documents:
        source_path = Path(str(document.get("sourcePath") or ""))
        if source_path.suffix.lower() != ".docx" or not source_path.exists():
            continue

        source_file = str(document.get("name") or source_path.name or "招标文件")
        blocks = _iter_docx_blocks(source_path)
        business_regions = _detect_business_format_regions(blocks) if profile.key == "business" else []
        # Parse the source docx once and hand the cached lxml tree + zip
        # path to every per-appendix slicer. Without this cache, a 21 MB RFP
        # with 50 appendices would re-parse the source ~50 times (each open
        # + parse takes seconds, freezing the request for minutes).
        source_state = _build_appendix_slice_state(source_path)
        used_tables: set[int] = set()
        in_template_section = False
        template_section_title = ""
        for index, block in enumerate(blocks):
            if block.get("type") != "paragraph":
                continue
            line = str(block.get("text") or "").strip()
            region = _business_region_for_index(business_regions, index) if profile.key == "business" else None
            if profile.key == "business" and _is_business_template_section_heading(line):
                in_template_section = True
                template_section_title = line
                continue
            if profile.key == "business" and in_template_section and _is_business_major_section_heading(line):
                in_template_section = False
                template_section_title = ""
            if profile.key == "business":
                is_heading = _business_template_title_allowed(
                    line,
                    in_template_section=in_template_section or region is not None,
                )
            else:
                is_heading = _is_relevant_appendix_heading(line, profile)
            if not is_heading:
                continue
            if _is_scoring_appendix_heading(line):
                continue
            if profile.key == "business":
                if not in_template_section and region is None:
                    continue
                if _is_business_toc_or_directory_context(blocks, index):
                    continue
            elif _is_docx_appendix_toc_artifact(blocks, index):
                continue

            title = line.strip(" #") or "附表"
            table_index = -1
            rows: list[list[str]] = []
            table_body_index: int | None = None
            next_heading_index = _next_appendix_heading_index(blocks, index, profile)
            if profile.key == "business" and isinstance(region, dict) and isinstance(region.get("end"), int):
                next_heading_index = min(next_heading_index, int(region["end"]))
            search_limit = next_heading_index if profile.key == "business" else min(len(blocks), index + 8)
            for lookahead in range(index + 1, min(search_limit, index + 12)):
                next_block = blocks[lookahead]
                if next_block.get("type") == "table" and lookahead not in used_tables:
                    table_index = lookahead
                    rows = next_block.get("rows") or []
                    raw_body_index = next_block.get("body_index")
                    table_body_index = raw_body_index if isinstance(raw_body_index, int) else None
                    break
            content_blocks = _slice_content_blocks(blocks, index, next_heading_index)

            appendix_id = f"APPX-{start_index + len(appendices) + 1:04d}"
            heading_body_index_raw = block.get("body_index")
            heading_body_index = heading_body_index_raw if isinstance(heading_body_index_raw, int) else None
            slice_end_body_index = _last_body_index_before(blocks, next_heading_index, heading_body_index)
            metadata: dict[str, Any] = {}
            if profile.key == "business":
                section_title = template_section_title or str((region or {}).get("title") or "")
                metadata = _business_template_metadata(
                    title=title,
                    source_file=source_file,
                    evidence=line,
                    evidence_location=f"B{index + 1}",
                    template_section_title=section_title,
                    source_start=_business_template_source_start(section_title, f"B{index + 1}", line),
                    source_end=_business_template_source_end_from_blocks(blocks, next_heading_index),
                    rows=rows,
                    content_blocks=content_blocks,
                    extraction_mode="source_docx_slice",
                )
                if not _business_template_should_materialize(metadata, rows, content_blocks):
                    continue

            if table_index == -1 or not rows:
                appendix_payload = {
                    "id": appendix_id,
                    "title": title,
                    "status": "generated",
                    "sourceFile": source_file,
                    "evidence": line,
                    "evidenceLocation": f"B{index + 1}",
                    "rows": [],
                    "contentBlocks": content_blocks,
                    "rowCount": 0,
                    "docxPath": "",
                    **metadata,
                }
                if (
                    profile.key == "business"
                    and heading_body_index is not None
                    and slice_end_body_index is not None
                ):
                    appendix_payload["_slice"] = {
                        "sourcePath": str(source_path),
                        "keepStart": heading_body_index,
                        "keepEnd": slice_end_body_index,
                        "sourceState": source_state,
                    }
                appendices.append(materialize_appendix_docx(project_id, appendix_payload, profile=profile))
                continue

            used_tables.add(table_index)
            appendix_payload: dict[str, Any] = {
                "id": appendix_id,
                "title": title,
                "status": "generated",
                "sourceFile": source_file,
                "evidence": line,
                "evidenceLocation": f"B{index + 1}",
                "rows": rows,
                "contentBlocks": content_blocks,
                "rowCount": len(rows),
                "docxPath": "",
                **metadata,
            }
            slice_keep_end = slice_end_body_index if profile.key == "business" else table_body_index
            if heading_body_index is not None and slice_keep_end is not None:
                appendix_payload["_slice"] = {
                    "sourcePath": str(source_path),
                    "keepStart": heading_body_index,
                    "keepEnd": slice_keep_end,
                    "sourceState": source_state,
                }
            appendices.append(materialize_appendix_docx(project_id, appendix_payload, profile=profile))
        if profile.key == "business":
            used_heading_indices = {
                int(str(item.get("evidenceLocation") or "B0").lstrip("B") or "0") - 1
                for item in appendices
                if isinstance(item, dict) and str(item.get("evidenceLocation") or "").startswith("B")
            }
            for table_index, table_block in enumerate(blocks):
                if table_index in used_tables or table_block.get("type") != "table":
                    continue
                region = _business_region_for_index(business_regions, table_index)
                if region is None:
                    continue
                rows = table_block.get("rows") if isinstance(table_block.get("rows"), list) else []
                table_type, table_label, table_confidence, matched_tokens = _classify_business_table(rows)
                if not table_type or table_confidence < 0.62:
                    continue
                title_index = _previous_business_title_block_index(blocks, table_index, region)
                heading_index = title_index if title_index is not None else table_index
                if title_index is not None and title_index in used_heading_indices:
                    continue
                title = (
                    str(blocks[title_index].get("text") or "").strip()
                    if title_index is not None
                    else table_label
                ) or table_label
                content_blocks = _slice_content_blocks(blocks, heading_index, min(table_index + 1, len(blocks)))
                if not any(block.get("type") == "table" for block in content_blocks if isinstance(block, dict)):
                    content_blocks.append({"type": "table", "rows": rows})
                section_title = str(region.get("title") or "")
                metadata = _business_template_metadata(
                    title=title,
                    source_file=source_file,
                    evidence=title,
                    evidence_location=f"B{heading_index + 1}",
                    template_section_title=section_title,
                    source_start=_business_template_source_start(section_title, f"B{heading_index + 1}", title),
                    source_end=_business_template_source_end_from_blocks(blocks, min(table_index + 1, len(blocks))),
                    rows=rows,
                    content_blocks=content_blocks,
                    extraction_mode="source_docx_table_fingerprint",
                )
                metadata["templateType"] = table_type
                metadata["templateTypeLabel"] = _business_template_type_label_for_key(table_type)
                metadata["tableFingerprint"] = {
                    "type": table_type,
                    "typeLabel": table_label,
                    "confidence": table_confidence,
                    "matchedTokens": matched_tokens,
                }
                if not _business_template_should_materialize(metadata, rows, content_blocks):
                    continue
                appendix_id = f"APPX-{start_index + len(appendices) + 1:04d}"
                heading_body_index_raw = blocks[heading_index].get("body_index") if 0 <= heading_index < len(blocks) else table_block.get("body_index")
                table_body_index_raw = table_block.get("body_index")
                heading_body_index = heading_body_index_raw if isinstance(heading_body_index_raw, int) else None
                table_body_index = table_body_index_raw if isinstance(table_body_index_raw, int) else None
                payload: dict[str, Any] = {
                    "id": appendix_id,
                    "title": title,
                    "status": "generated",
                    "sourceFile": source_file,
                    "evidence": title,
                    "evidenceLocation": f"B{heading_index + 1}",
                    "rows": rows,
                    "contentBlocks": content_blocks,
                    "rowCount": len(rows),
                    "docxPath": "",
                    **metadata,
                }
                if heading_body_index is not None and table_body_index is not None:
                    payload["_slice"] = {
                        "sourcePath": str(source_path),
                        "keepStart": heading_body_index,
                        "keepEnd": table_body_index,
                        "sourceState": source_state,
                    }
                used_tables.add(table_index)
                appendices.append(materialize_appendix_docx(project_id, payload, profile=profile))
    return appendices


def _extract_structured_requirements(documents: list[dict[str, Any]], texts_by_id: dict[str, str]) -> dict[str, Any]:
    return parse_structured_documents(
        documents,
        texts_by_id,
        mode="local-structured-parser",
    )


def _business_local_contract_result(
    project_id: str,
    structured_result: dict[str, Any],
    *,
    profile: ParseProfile,
    documents: list[dict[str, Any]],
    texts_by_id: dict[str, str],
) -> dict[str, Any]:
    if profile.key != "business":
        return structured_result
    return _transform_to_business_contract(
        project_id,
        structured_result,
        profile=profile,
        documents=documents,
        texts_by_id=texts_by_id,
        run_semantic_review=True,
    )


def _business_template_extractor_allows_preview_fallback(warning: str) -> bool:
    text = str(warning or "")
    if not text:
        return False
    if any(
        token in text
        for token in (
            "Agent 裁决未完成",
            "agent 裁决未完成",
            "缺少 Agent 裁决文件",
            "缺少 agent 裁决文件",
            "btplbound boundary-decision failed",
            "opencode incomplete/stalled",
            "futurecode 创建 session 失败",
            "getaddrinfo failed",
        )
    ):
        return False
    return any(
        token in text
        for token in (
            "未找到可用于商务模板提取 skill 的 DOCX 招标文件",
        )
    )


def _merge_business_local_artifacts(
    structured_result: dict[str, Any],
    local_business_result: dict[str, Any],
    *,
    profile: ParseProfile,
) -> dict[str, Any]:
    if profile.key != "business":
        return structured_result
    structured = structured_result.get("structured") if isinstance(structured_result.get("structured"), dict) else {}
    local_structured = local_business_result.get("structured") if isinstance(local_business_result.get("structured"), dict) else {}
    if not isinstance(structured, dict) or not isinstance(local_structured, dict):
        return structured_result

    field_groups = structured.setdefault("fieldGroups", {})
    local_field_groups = local_structured.get("fieldGroups") if isinstance(local_structured.get("fieldGroups"), dict) else {}
    if isinstance(field_groups, dict):
        for key in ("businessResponse", "qualificationSupport", "commitmentRequirements"):
            if key not in field_groups and isinstance(local_field_groups.get(key), list):
                field_groups[key] = copy.deepcopy(local_field_groups.get(key) or [])

    scoring = structured.setdefault("scoringCriteria", {})
    if isinstance(scoring, dict):
        for key in ("price", "compliance"):
            scoring.setdefault(key, [])

    for key in ("appendices", "commitmentLetters", "commitmentClues", "commitmentTemplateAlignments", "businessFormatRegions"):
        local_value = local_structured.get(key)
        if isinstance(local_value, list) and (local_value or not isinstance(structured.get(key), list)):
            if not structured.get(key):
                structured[key] = copy.deepcopy(local_value)

    local_presence = local_structured.get("requirementPresence")
    if isinstance(local_presence, dict):
        current_presence = structured.get("requirementPresence")
        if not isinstance(current_presence, dict) or not current_presence:
            structured["requirementPresence"] = copy.deepcopy(local_presence)

    category_counts = structured.setdefault("categoryCounts", {})
    local_category_counts = local_structured.get("categoryCounts") if isinstance(local_structured.get("categoryCounts"), dict) else {}
    if isinstance(category_counts, dict):
        for key, value in local_category_counts.items():
            category_counts.setdefault(key, value)

    return structured_result


def _build_tender_parse_prompt(skill_manifest_path: Path, profile: ParseProfile) -> str:
    if profile.key == "business":
        return f"""
Use the {profile.skill_name} skill.

你在做 S1 商务招标文件结构化解析。业务任务书、交付清单和语义原则以 skill 内的 `SKILL.md` 为准；本提示只约束执行链路。

manifest：{skill_manifest_path}

必须用 Bash 按顺序执行 `s1parse` 小输出链路，timeout 设置为 600000 毫秒或更高：

s1parse prepare {skill_manifest_path}
s1parse overview {skill_manifest_path} --page 1 --page-size 30
s1parse search {skill_manifest_path} "<query>" --limit 20
s1parse read {skill_manifest_path} <evidenceId> --mode summary --max-chars 2000
s1parse window {skill_manifest_path} <evidenceId> --before 4 --after 6
s1parse table {skill_manifest_path} <tableId> --rows 1-12 --max-chars 4000
s1parse submit {skill_manifest_path} projectBasics '<json>'
s1parse submit {skill_manifest_path} qualificationRequirements '<json>'
s1parse submit {skill_manifest_path} bidderInstructions '<json>'
s1parse submit {skill_manifest_path} commercialRejectionClauses '<json>'
s1parse submit {skill_manifest_path} businessScoringCriteria '<json>'
s1parse validate {skill_manifest_path}
s1parse status {skill_manifest_path}
s1parse finalize {skill_manifest_path}

禁止用 opencode 的 read 工具读取或打印解析中间产物的大 JSON；证据定位必须通过 s1parse 小输出导航命令完成。禁止调用 Task/subagent/子代理/任务委派工具。

只使用 s1parse 返回过的 evidenceId，不要编造证据。提交值必须能被对应证据文本直接支撑，项目基础信息至少为项目名称、招标人、递交截止时间提交字段级 evidenceIds。validate 失败时继续回查并重新 submit；若仍失败，必须让 workflow 暴露 missingTargets 或 validationErrors，不能把失败结果说成成功。必须执行 finalize，最终结构化 JSON 必须由 finalize 写入 manifest.structuredResultPath。

最后只返回 finalize 命令 stdout 中的小型 JSON，不要返回解释文字，不要使用 Markdown 代码块。
返回格式必须是：
{{
  "schemaVersion": "{profile.schema_version}",
  "targetSkill": "{profile.skill_name}",
  "outputFile": "manifest 中的 structuredResultPath",
  "summary": {{"itemCount": 0, "targetCounts": {{}}, "scoringCounts": {{"business": 0}}, "workflowStage": "finalized", "projectDates": {{"startDate": "", "endDate": ""}}}}
}}

完整 JSON 必须包含 structured.sourceDocuments、structured.scoringCriteria、structured.fieldGroups、structured.projectFactFields、structured.projectDates、structured.coverage、structured.workflow，且 workflow.mode 为 opencode-agentic-navigation。
""".strip()
    return f"""
Use the {profile.skill_name} skill.

你现在在做 S1 招标文件结构化解析。请调用解析 Skill 读取 manifest 中的多份招标文件文本，输出可直接给后端使用的结构化 JSON。

manifest：{skill_manifest_path}

请直接调用一次 Bash 工具执行下面命令，Bash 工具 timeout 必须设置为 600000 毫秒或更高。不要先检查工作目录，不要先执行 pwd/ls/cat/read/glob，不要拆成多条命令，不要改写命令或路径。命令会把完整结构化 JSON 写入 manifest.structuredResultPath，并只在 stdout 打印小型摘要 JSON：

s1parse {skill_manifest_path}

只返回命令 stdout 中的小型 JSON，不要返回解释文字，不要使用 Markdown 代码块。
返回格式必须是：
{{
  "schemaVersion": "{profile.schema_version}",
  "targetSkill": "{profile.skill_name}",
  "outputFile": "manifest 中的 structuredResultPath",
  "summary": {{"itemCount": 0, "categoryCounts": {{}}, "scoringCounts": {{"technical": 0, "business": 0, "price": 0, "lcoe": 0, "compliance": 0}}, "projectDates": {{"startDate": "", "endDate": ""}}}}
}}

解析目标必须覆盖：
1. 评分细则：评分项、得分点、证明材料要求。
2. 项目基础信息：项目名称、招标编号、招标人/管理单位、规模、交货周期、质保期、技术承诺。
3. 风机核心参数：单机容量、叶轮直径、轮毂高度、叶片最低点距地、塔筒型式、箱变型式、安全等级、空气密度、风速、湍流强度。
4. 性能保证指标：功率曲线、可利用率、发电量、涉网性能。
5. 环境适应性要求：低温、覆冰防凝露、潮湿、防雷暴、防风沙、高温。
6. 专题方案要求：叶片、变桨系统、主轴、齿轮箱等专题。
7. 附表、供货范围和考核条款。
8. 投标相关日期：招标文件获取/报名起始日期、投标文件递交截止日期或开标日期。不要把交货、供货、服务期、工期、竣工、安装调试等履约日期写入 projectDates。

完整 JSON 必须包含 structured.sourceDocuments、structured.scoringCriteria、structured.fieldGroups、structured.requirementPresence、structured.coverage。每条 item、评分行和字段必须保留 sourceFile、sourceDocumentId、section、evidence、evidenceLocation。
""".strip()


def _build_tender_parse_retry_prompt(
    skill_manifest_path: Path,
    profile: ParseProfile,
    first_error: RuntimeError,
) -> str:
    if profile.key != "business":
        return _build_tender_parse_prompt(skill_manifest_path, profile)
    failure = str(first_error)
    return f"""
Use the {profile.skill_name} skill.

这是同一个 S1 商务招标文件结构化解析任务的一次恢复重试。第一次 opencode 会话没有完成 finalize，失败原因如下：
{failure}

manifest：{skill_manifest_path}

不要重新发散式探索，不要读取或打印大 JSON，不要调用 Task/subagent/子代理/任务委派工具。必须使用 Bash 继续完成同一条 `s1parse` 工作流，timeout 设置为 600000 毫秒或更高。

按下面顺序执行：

s1parse status {skill_manifest_path}
s1parse submit {skill_manifest_path} projectBasics '<json>'
s1parse submit {skill_manifest_path} qualificationRequirements '<json>'
s1parse submit {skill_manifest_path} bidderInstructions '<json>'
s1parse submit {skill_manifest_path} commercialRejectionClauses '<json>'
s1parse submit {skill_manifest_path} businessScoringCriteria '<json>'
s1parse validate {skill_manifest_path}
s1parse status {skill_manifest_path}
s1parse finalize {skill_manifest_path}

如果 status 显示已有提交项，只补缺失项；如果证据不足，只用 s1parse search/read/window/table 做最小回查。validate 失败时继续补 submit 并重新 validate。必须执行 finalize，最终结构化 JSON 必须由 finalize 写入 manifest.structuredResultPath。

最后只返回 finalize 命令 stdout 中的小型 JSON，不要返回解释文字，不要使用 Markdown 代码块。返回格式必须是：
{{
  "schemaVersion": "{profile.schema_version}",
  "targetSkill": "{profile.skill_name}",
  "outputFile": "manifest 中的 structuredResultPath",
  "summary": {{"itemCount": 0, "targetCounts": {{}}, "scoringCounts": {{"business": 0}}, "workflowStage": "finalized", "projectDates": {{"startDate": "", "endDate": ""}}}}
}}
""".strip()


def _resolve_skill_structured_result(
    result: dict[str, Any],
    *,
    local_result: dict[str, Any],
    profile: ParseProfile,
) -> dict[str, Any]:
    output_file = Path(str(result.get("outputFile") or ""))
    if output_file.exists():
        loaded = json.loads(output_file.read_text(encoding="utf-8"))
        if isinstance(loaded, dict) and isinstance(loaded.get("items"), list):
            resolved = loaded
        elif profile.key == "business":
            raise RuntimeError(f"S1 商务解析 Skill 输出结构不合法：{output_file}")
        else:
            resolved = local_result
    elif isinstance(result.get("items"), list):
        resolved = {
            "items": result.get("items") or [],
            "structured": result.get("structured") if isinstance(result.get("structured"), dict) else {},
        }
    elif profile.key == "business":
        raise RuntimeError("S1 商务解析 Skill 未返回结构化 items。")
    else:
        resolved = local_result

    structured = resolved.setdefault("structured", {})
    if isinstance(structured, dict):
        local_structured = local_result.get("structured") if isinstance(local_result, dict) else {}
        if profile.key != "business" and isinstance(local_structured, dict):
            for key in ["sourceDocuments", "fieldGroups", "scoringCriteria", "requirementPresence", "coverage", "appendices"]:
                if not structured.get(key) and local_structured.get(key):
                    structured[key] = local_structured[key]
        structured["targetSkill"] = profile.skill_name
        structured["mode"] = "opencode-skill"
        if isinstance(result.get("opencodeOutput"), dict):
            structured["opencodeOutput"] = copy.deepcopy(result.get("opencodeOutput") or {})
        elif isinstance(local_structured, dict) and isinstance(local_structured.get("opencodeOutput"), dict):
            structured["opencodeOutput"] = copy.deepcopy(local_structured.get("opencodeOutput") or {})
        _apply_opencode_trace_to_workflow(structured)
        structured["schemaVersion"] = str(structured.get("schemaVersion") or profile.schema_version)
    return resolved


def _apply_opencode_trace_to_workflow(structured: dict[str, Any]) -> None:
    trace = structured.get("opencodeOutput") if isinstance(structured.get("opencodeOutput"), dict) else {}
    if not isinstance(trace, dict) or not trace:
        return
    workflow = structured.get("workflow") if isinstance(structured.get("workflow"), dict) else {}
    workflow = copy.deepcopy(workflow)
    session_id = str(trace.get("sessionId") or "").strip()
    if session_id:
        workflow["opencodeSessionId"] = session_id
    agent_status = str(trace.get("agentStatus") or trace.get("status") or "").strip()
    if agent_status:
        workflow["opencodeAgentStatus"] = agent_status
    last_tool = str(trace.get("lastTool") or "").strip()
    if last_tool:
        workflow["opencodeLastTool"] = last_tool
    last_tool_status = str(trace.get("lastToolStatus") or "").strip()
    if last_tool_status:
        workflow["opencodeLastToolStatus"] = last_tool_status
    failure_reason = str(trace.get("failureReason") or "").strip()
    if failure_reason:
        workflow["opencodeFailureReason"] = failure_reason
    elif "opencodeFailureReason" not in workflow:
        workflow["opencodeFailureReason"] = ""
    if workflow:
        structured["workflow"] = workflow


def _opencode_attempt_from_error(exc: RuntimeError, attempt: int) -> dict[str, Any]:
    trace = getattr(exc, "opencode_trace", None)
    if not isinstance(trace, dict):
        trace = {"status": "error", "failureReason": str(exc)}
    attempt_payload = {
        "attempt": attempt,
        "status": str(trace.get("status") or trace.get("agentStatus") or "error"),
        "sessionId": str(trace.get("sessionId") or ""),
        "providerId": str(trace.get("providerId") or ""),
        "modelId": str(trace.get("modelId") or ""),
        "agentStatus": str(trace.get("agentStatus") or trace.get("status") or ""),
        "failureReason": str(trace.get("failureReason") or str(exc)),
    }
    for key in ("lastTool", "lastToolStatus", "lastToolInput", "errorName", "errorStatusCode"):
        if key in trace:
            attempt_payload[key] = copy.deepcopy(trace.get(key))
    return attempt_payload


def _attach_opencode_attempts(structured_result: dict[str, Any], attempts: list[dict[str, Any]]) -> dict[str, Any]:
    if not attempts:
        return structured_result
    structured = structured_result.setdefault("structured", {})
    if not isinstance(structured, dict):
        return structured_result
    workflow = structured.get("workflow") if isinstance(structured.get("workflow"), dict) else {}
    workflow = copy.deepcopy(workflow)
    workflow["opencodeAttempts"] = copy.deepcopy(attempts)
    structured["workflow"] = workflow
    return structured_result


def _fallback_parse_skill_result(
    exc: RuntimeError,
    *,
    local_result: dict[str, Any],
    progress_callback: Callable[[str, dict[str, Any] | None], None] | None = None,
    attempts: list[dict[str, Any]] | None = None,
) -> tuple[dict[str, Any], str]:
    fallback = json.loads(json.dumps(local_result, ensure_ascii=False))
    structured = fallback.setdefault("structured", {})
    if isinstance(structured, dict):
        structured["mode"] = "local-structured-parser"
        structured["opencodeError"] = str(exc)
        trace = getattr(exc, "opencode_trace", None)
        if isinstance(trace, dict):
            structured["opencodeOutput"] = copy.deepcopy(trace)
            if progress_callback:
                progress_callback("opencode_delta", copy.deepcopy(trace))
            if not trace.get("failureReason"):
                structured["opencodeOutput"]["failureReason"] = str(exc)
            _apply_opencode_trace_to_workflow(structured)
        if attempts:
            workflow = structured.get("workflow") if isinstance(structured.get("workflow"), dict) else {}
            workflow = copy.deepcopy(workflow)
            workflow["opencodeAttempts"] = copy.deepcopy(attempts)
            structured["workflow"] = workflow
    return fallback, f"S1 解析 Skill 调用失败，已使用本地结构化解析兜底：{exc}"


def _run_parse_skill(
    skill_manifest_path: Path,
    *,
    local_result: dict[str, Any],
    profile: ParseProfile,
    progress_callback: Callable[[str, dict[str, Any] | None], None] | None = None,
) -> tuple[dict[str, Any], str]:
    if not settings.s1_parse_opencode_enabled:
        return local_result, ""
    client = OpencodeClient()
    stream_callback = (
        (lambda details: progress_callback("opencode_delta", details))
        if progress_callback
        else None
    )
    try:
        result = client.generate_tender_parse_with_trace(
            _build_tender_parse_prompt(skill_manifest_path, profile),
            stream_callback=stream_callback,
        )
        return _resolve_skill_structured_result(result, local_result=local_result, profile=profile), ""
    except RuntimeError as exc:
        attempts = [_opencode_attempt_from_error(exc, 1)]
        trace = getattr(exc, "opencode_trace", None)
        if progress_callback and isinstance(trace, dict):
            progress_callback("opencode_delta", copy.deepcopy(trace))
        try:
            retry_result = client.generate_tender_parse_with_trace(
                _build_tender_parse_retry_prompt(skill_manifest_path, profile, exc),
                stream_callback=stream_callback,
            )
            resolved = _resolve_skill_structured_result(retry_result, local_result=local_result, profile=profile)
            retry_trace = (resolved.get("structured") or {}).get("opencodeOutput")
            attempts.append(
                {
                    "attempt": 2,
                    "status": "succeeded",
                    "sessionId": str(retry_trace.get("sessionId") or "") if isinstance(retry_trace, dict) else "",
                    "providerId": str(retry_trace.get("providerId") or "") if isinstance(retry_trace, dict) else "",
                    "modelId": str(retry_trace.get("modelId") or "") if isinstance(retry_trace, dict) else "",
                }
            )
            return _attach_opencode_attempts(resolved, attempts), ""
        except RuntimeError as retry_exc:
            attempts.append(_opencode_attempt_from_error(retry_exc, 2))
            return _fallback_parse_skill_result(
                retry_exc,
                local_result=local_result,
                progress_callback=progress_callback,
                attempts=attempts,
            )


def _business_s1_runner_path() -> Path:
    return (
        Path(__file__).resolve().parents[2]
        / "opencode"
        / "skill"
        / BUSINESS_PARSE_PROFILE.skill_name
        / "scripts"
        / "run_from_manifest.py"
    )


def _workflow_from_result(structured_result: dict[str, Any]) -> dict[str, Any]:
    structured = structured_result.get("structured") if isinstance(structured_result, dict) else {}
    workflow = structured.get("workflow") if isinstance(structured, dict) else {}
    return workflow if isinstance(workflow, dict) else {}


def _business_validation_report_path(skill_manifest_path: Path, workflow: dict[str, Any]) -> Path:
    workflow_path = str(workflow.get("validationReportPath") or "").strip()
    if workflow_path:
        return Path(workflow_path)
    return skill_manifest_path.with_name("validation_report.json")


def _needs_business_s1_finalize_guard(
    *,
    profile: ParseProfile,
    structured_result: dict[str, Any],
    skill_manifest_path: Path,
) -> bool:
    if profile.key != "business":
        return False
    workflow = _workflow_from_result(structured_result)
    if str(workflow.get("mode") or "").strip() != "opencode-agentic-navigation":
        return False
    workflow_stage = str(workflow.get("stage") or "").strip()
    if workflow_stage != "finalized":
        return True
    return not _business_validation_report_path(skill_manifest_path, workflow).is_file()


def _is_business_skill_workflow_payload(
    *,
    profile: ParseProfile,
    structured_payload: dict[str, Any] | None,
    skill_manifest_path: Path,
) -> bool:
    if profile.key != "business" or not isinstance(structured_payload, dict):
        return False
    if str(structured_payload.get("targetSkill") or "").strip() == BUSINESS_PARSE_PROFILE.skill_name:
        return True
    workflow = structured_payload.get("workflow") if isinstance(structured_payload.get("workflow"), dict) else {}
    if str(workflow.get("mode") or "").strip() == "opencode-agentic-navigation":
        return True
    artifact_keys = ("navStorePath", "documentMapPath", "submissionPath", "validationReportPath")
    if any(str(workflow.get(key) or "").strip() for key in artifact_keys):
        return True
    return False


def _mark_business_skill_workflow_guard(structured_payload: dict[str, Any]) -> None:
    workflow = structured_payload.get("workflow") if isinstance(structured_payload.get("workflow"), dict) else {}
    workflow = copy.deepcopy(workflow)
    workflow["backendFinalizeGuardApplied"] = True
    structured_payload["workflow"] = workflow


def _business_finalize_error_result(
    skill_manifest_path: Path,
    structured_result: dict[str, Any],
    error: str,
) -> dict[str, Any]:
    fallback = copy.deepcopy(structured_result if isinstance(structured_result, dict) else {})
    structured = fallback.setdefault("structured", {})
    if not isinstance(structured, dict):
        fallback["structured"] = structured = {}

    existing_workflow = structured.get("workflow") if isinstance(structured.get("workflow"), dict) else {}
    workflow = copy.deepcopy(existing_workflow)
    nav_store_path = Path(str(workflow.get("navStorePath") or skill_manifest_path.with_name("s1_nav.sqlite")))
    document_map_path = Path(str(workflow.get("documentMapPath") or skill_manifest_path.with_name("document_map.json")))
    submission_path = Path(str(workflow.get("submissionPath") or skill_manifest_path.with_name("agentic_submissions.json")))
    validation_report_path = _business_validation_report_path(skill_manifest_path, workflow)
    workflow.update(
        {
            "stage": "failed",
            "mode": str(workflow.get("mode") or "opencode-agentic-navigation"),
            "aiReviewTrusted": False,
            "navStorePath": str(nav_store_path),
            "documentMapPath": str(document_map_path),
            "submissionPath": str(submission_path),
            "validationReportPath": str(validation_report_path),
            "submittedTargetCount": int(workflow.get("submittedTargetCount") or 0),
            "missingTargets": list(workflow.get("missingTargets") or []),
            "validationErrors": list(workflow.get("validationErrors") or []),
            "backendFinalizeGuardApplied": True,
            "backendFinalizeError": error,
        }
    )
    structured["workflow"] = workflow
    structured["mode"] = str(structured.get("mode") or "opencode-skill")
    return fallback


def _finalize_business_s1_result(
    skill_manifest_path: Path,
    structured_result: dict[str, Any],
    profile: ParseProfile,
) -> tuple[dict[str, Any], str]:
    runner_path = _business_s1_runner_path()
    try:
        if not runner_path.is_file():
            raise RuntimeError(f"商务 S1 finalize runner 不存在: {runner_path}")
        completed = subprocess.run(
            [sys.executable, str(runner_path), "finalize", str(skill_manifest_path)],
            check=False,
            capture_output=True,
            text=True,
            encoding="utf-8",
            timeout=600,
        )
        if completed.returncode != 0:
            detail = (completed.stderr or completed.stdout or "").strip()
            raise RuntimeError(f"s1parse finalize failed with exit {completed.returncode}: {detail}")
        manifest = json.loads(skill_manifest_path.read_text(encoding="utf-8"))
        output_path = Path(str(manifest.get("structuredResultPath") or skill_manifest_path.with_name("s1_structured_result.json")))
        if not output_path.is_file():
            raise RuntimeError(f"s1parse finalize 未写入结果文件: {output_path}")
        resolved = _resolve_skill_structured_result(
            {
                "outputFile": str(output_path),
            },
            local_result=structured_result,
            profile=profile,
        )
        structured = resolved.setdefault("structured", {})
        if isinstance(structured, dict):
            local_structured = structured_result.get("structured") if isinstance(structured_result, dict) else {}
            if isinstance(local_structured, dict) and isinstance(local_structured.get("opencodeOutput"), dict):
                structured["opencodeOutput"] = copy.deepcopy(local_structured.get("opencodeOutput") or {})
            structured["backendFinalizeOutput"] = {
                "backendFinalizeGuardApplied": True,
                "stdout": completed.stdout.strip(),
            }
            _apply_opencode_trace_to_workflow(structured)
            workflow = structured.get("workflow") if isinstance(structured.get("workflow"), dict) else {}
            workflow = copy.deepcopy(workflow)
            workflow["backendFinalizeGuardApplied"] = True
            structured["workflow"] = workflow
        return resolved, ""
    except Exception as exc:
        message = str(exc)
        return (
            _business_finalize_error_result(skill_manifest_path, structured_result, message),
            f"S1 商务 finalize 收口失败，已保留错误 workflow：{message}",
        )


def parse_tender_documents(
    project_id: str,
    tender_files: list[dict[str, Any]],
    *,
    bid_type: str,
    progress_callback: Callable[[str, dict[str, Any] | None], None] | None = None,
) -> tuple[dict[str, Any], dict[str, Any]]:
    profile = resolve_parse_profile(bid_type)
    project_dir = parsed_project_dir(project_id)
    appendix_temp_dir = project_dir / "s1_appendices"
    if appendix_temp_dir.exists():
        shutil.rmtree(appendix_temp_dir)
    documents: list[dict[str, Any]] = []
    combined_parts: list[str] = []
    texts_by_id: dict[str, str] = {}
    warnings: list[str] = []
    if progress_callback:
        progress_callback("extract_started", {"fileCount": len(tender_files)})

    for file_record in tender_files:
        file_path = Path(str(file_record["path"]))
        extension = file_path.suffix.lower()
        page_count: int | str = "-"
        file_warnings: list[str] = []
        ocr_meta: dict[str, Any] | None = None
        if progress_callback:
            progress_callback("extracting_file", {"fileName": file_record.get("name") or file_path.name})

        if extension == ".docx":
            text = extract_docx_text(file_path)
        elif extension == ".md":
            text = file_path.read_text(encoding="utf-8", errors="replace")
        elif extension == ".txt":
            text = file_path.read_text(encoding="utf-8", errors="replace")
        elif extension == ".pdf":
            text, pdf_meta = extract_pdf_text(file_path)
            page_count = pdf_meta["pageCount"]
            file_warnings.extend(pdf_meta["warnings"])
            if pdf_meta.get("requiresOcr"):
                ocr_text, ocr_meta = _ocr_fallback_text(project_id, file_record, file_path)
                if ocr_text:
                    text = ocr_text
                    page_count = ocr_meta.get("pageCount") or page_count
                    file_warnings.append("扫描型 PDF 已通过 OCR/视觉模型转为可解析文本。")
                else:
                    file_warnings.append(
                        f"OCR 兜底识别未完成：{ocr_meta.get('message') if ocr_meta else '未知错误'}"
                    )
        elif extension in IMAGE_SUFFIXES:
            ocr_text, ocr_meta = _ocr_fallback_text(project_id, file_record, file_path)
            if ocr_text:
                text = ocr_text
                page_count = ocr_meta.get("pageCount") or 1
                file_warnings.append("图片文件已通过 OCR/视觉模型转为可解析文本。")
            else:
                text = ""
                page_count = 1
                file_warnings.append(
                    f"图片文件需要 OCR 识别，但兜底识别未完成：{ocr_meta.get('message') if ocr_meta else '未知错误'}"
                )
        else:
            text = ""
            file_warnings.append(f"当前 MVP 暂不解析 {extension or '未知'} 类型文件。")

        text = _normalize_text(text)
        text_length = len(text)
        text_path = project_dir / f"{file_record['id']}.txt"
        text_path.write_text(text, encoding="utf-8")

        metadata: dict[str, Any] = {
            "id": file_record["id"],
            "name": file_record["name"],
            "sourcePath": str(file_path),
            "textPath": str(text_path),
            "textLength": text_length,
            "pageCount": page_count,
            "warnings": file_warnings,
        }
        if ocr_meta:
            metadata["ocr"] = ocr_meta
        documents.append(metadata)
        texts_by_id[str(file_record["id"])] = text
        warnings.extend(file_warnings)

        if text:
            combined_parts.append(f"# 文件：{file_record['name']}\n\n{text}")
        if progress_callback:
            progress_callback(
                "file_extracted",
                {
                    "fileName": file_record.get("name") or file_path.name,
                    "textLength": text_length,
                    "warnings": file_warnings,
                },
            )

    combined_text = _normalize_text("\n\n".join(combined_parts))
    combined_text_path = project_dir / "combined.txt"
    combined_text_path.write_text(combined_text, encoding="utf-8")

    structured_result = _extract_structured_requirements(documents, texts_by_id)
    template_extraction_payload: dict[str, Any] | None = None
    template_extraction_warning = ""
    template_extraction_path = project_dir / "business_template_extraction" / "business_template_extraction.json"
    business_section_tree_path = ""
    business_section_tree_summary: dict[str, Any] = {}

    if profile.key == "business":
        section_tree_path, section_tree_payload = write_business_section_tree(documents, project_dir)
        business_section_tree_path = str(section_tree_path)
        business_section_tree_summary = (
            section_tree_payload.get("summary")
            if isinstance(section_tree_payload, dict) and isinstance(section_tree_payload.get("summary"), dict)
            else {}
        )
        if progress_callback:
            progress_callback("business_template_extraction_started", {"documentCount": len(documents)})
        appendices, template_extraction_payload, template_extraction_warning = run_business_template_extractor(
            project_id=project_id,
            documents=documents,
            project_dir=project_dir,
            progress_callback=progress_callback,
        )
        if progress_callback:
            progress_callback(
                "business_template_extraction_finished",
                {
                    "appendixCount": len(appendices),
                    "warningCount": len(template_extraction_payload.get("warnings") or [])
                    if isinstance(template_extraction_payload, dict)
                    else 0,
                },
            )
        if not appendices and _business_template_extractor_allows_preview_fallback(template_extraction_warning):
            appendices = _extract_markdown_appendices(project_id, documents, texts_by_id, profile=profile)
            appendices.extend(_extract_docx_appendices(project_id, documents, start_index=len(appendices), profile=profile))
            appendices.extend(
                _extract_text_business_appendices(
                    project_id,
                    documents,
                    texts_by_id,
                    start_index=len(appendices),
                    profile=profile,
                )
            )
    else:
        appendices = _extract_markdown_appendices(project_id, documents, texts_by_id, profile=profile)
        appendices.extend(_extract_docx_appendices(project_id, documents, start_index=len(appendices), profile=profile))
    appendices = _prepare_appendix_outputs(project_id, appendices, renumber=True, profile=profile)
    structured_result["structured"]["appendices"] = appendices
    if profile.key == "business" and not settings.s1_parse_opencode_enabled:
        local_business_result = _business_local_contract_result(
            project_id,
            structured_result,
            profile=profile,
            documents=documents,
            texts_by_id=texts_by_id,
        )
        structured_result = local_business_result
    else:
        local_business_result = structured_result
    if progress_callback:
        progress_callback(
            "appendices_extracted",
            {
                "appendixCount": len(appendices),
                "generatedCount": sum(1 for item in appendices if item.get("status") == "generated"),
            },
        )
    structured_path = project_dir / "s1_structured_result.json"
    structured_path.write_text(json.dumps(structured_result, ensure_ascii=False, indent=2), encoding="utf-8")
    skill_manifest_path = project_dir / "s1_parse_manifest.json"
    skill_manifest = {
        "projectId": project_id,
        "bidType": profile.bid_type,
        "parseProfile": profile.key,
        "targetSkill": profile.skill_name,
        "combinedTextPath": str(combined_text_path),
        "structuredResultPath": str(structured_path),
        "businessTemplateExtractionPath": str(template_extraction_path) if profile.key == "business" and template_extraction_path.is_file() else "",
        "businessTemplateExtractionSummary": (
            template_extraction_payload.get("summary")
            if profile.key == "business" and isinstance(template_extraction_payload, dict)
            else {}
        ),
        "businessSectionTreePath": business_section_tree_path,
        "businessSectionTreeSummary": business_section_tree_summary,
        "documents": documents,
        "targets": list(profile.targets),
    }
    skill_manifest_path.write_text(json.dumps(skill_manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    if progress_callback:
        progress_callback("skill_manifest_ready", {"manifestPath": str(skill_manifest_path)})
    structured_result, skill_warning = _run_parse_skill(
        skill_manifest_path,
        local_result=structured_result,
        profile=profile,
        progress_callback=progress_callback,
    )
    if _needs_business_s1_finalize_guard(
        profile=profile,
        structured_result=structured_result,
        skill_manifest_path=skill_manifest_path,
    ):
        structured_result, finalize_warning = _finalize_business_s1_result(
            skill_manifest_path,
            structured_result,
            profile,
        )
        if finalize_warning:
            skill_warning = f"{skill_warning}；{finalize_warning}" if skill_warning else finalize_warning
    if template_extraction_warning:
        warnings.append(template_extraction_warning)
    structured_result = _merge_business_local_artifacts(
        structured_result,
        local_business_result,
        profile=profile,
    )
    resolved_structured = structured_result.setdefault("structured", {})
    if profile.key == "business" and not resolved_structured.get("appendices") and appendices:
        resolved_structured["appendices"] = appendices
    elif profile.key != "business" and not resolved_structured.get("appendices"):
        resolved_structured["appendices"] = appendices
    if isinstance(resolved_structured.get("appendices"), list):
        resolved_structured["appendices"] = _prepare_appendix_outputs(
            project_id,
            resolved_structured["appendices"],
            renumber=True,
            profile=profile,
        )
    if profile.key == "business" and isinstance(resolved_structured.get("commitmentLetters"), list):
        resolved_structured["commitmentLetters"] = _prepare_commitment_letter_outputs(
            project_id,
            resolved_structured.get("commitmentLetters") or [],
            renumber=True,
            profile=profile,
            project_name=_business_project_name_from_structured(resolved_structured),
            tenderer_name=_business_tenderer_name_from_structured(resolved_structured),
        )
    if skill_warning:
        warnings.append(skill_warning)
    structured_path.write_text(json.dumps(structured_result, ensure_ascii=False, indent=2), encoding="utf-8")

    manifest_path = project_dir / "manifest.json"
    manifest = {
        "projectId": project_id,
        "combinedTextPath": str(combined_text_path),
        "documents": documents,
        "structuredResultPath": str(structured_path),
        "skillManifestPath": str(skill_manifest_path),
    }
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    items = structured_result.get("items") if isinstance(structured_result.get("items"), list) else []
    structured = structured_result.get("structured") if isinstance(structured_result.get("structured"), dict) else {}
    project_dates = structured.get("projectDates") if isinstance(structured.get("projectDates"), dict) else {}

    summary = {
        "fileCount": len(tender_files),
        "extractedCount": len(items),
        "textLength": len(combined_text),
        "textPreview": combined_text[:TEXT_PREVIEW_LIMIT],
        "warnings": warnings,
        "targetSkill": profile.skill_name,
        "categoryCounts": structured.get("categoryCounts") or {},
        "projectDates": {
            "startDate": project_dates.get("startDate") or "",
            "endDate": project_dates.get("endDate") or "",
        },
        "appendixCount": len(structured.get("appendices") or []),
    }
    if profile.key == "business":
        summary["commitmentLetterCount"] = len(structured.get("commitmentLetters") or [])

    if progress_callback:
        progress_callback(
            "complete",
            {
                "extractedCount": len(items),
                "appendixCount": len(structured.get("appendices") or []),
            },
        )

    return summary, {
        "projectDir": str(project_dir),
        "combinedTextPath": str(combined_text_path),
        "manifestPath": str(manifest_path),
        "structuredResultPath": str(structured_path),
        "skillManifestPath": str(skill_manifest_path),
        "bidType": profile.bid_type,
        "parseProfile": profile.key,
        "documents": documents,
        "items": items,
        "structured": structured,
        "projectUpdates": {
            "startDate": project_dates.get("startDate") or "",
            "endDate": project_dates.get("endDate") or "",
            "deadline": project_dates.get("endDate") or "",
        },
    }
