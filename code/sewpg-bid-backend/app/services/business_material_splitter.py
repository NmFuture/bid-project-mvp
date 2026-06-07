from __future__ import annotations

import asyncio
import json
import re
import time
import zipfile
from copy import deepcopy
from io import BytesIO
from pathlib import PurePosixPath
from typing import Any

import httpx
from lxml import etree
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from sqlalchemy import select
from sqlalchemy.orm import selectinload

from app.core.config import settings
from app.models import async_session
from app.models.materials import RawFile
from app.services.bid_type import BUSINESS_BID_TYPE
from app.services.file_utils import safe_segment
from app.services.minio_client import minio_client
from app.services.opencode_client import OpencodeClient
from app.services.peripheral import PeripheralError
from app.services.system_settings import system_settings_service


BUSINESS_SPLIT_SCHEMA = "bid-business-material-split-plan-v1"

TEMPLATE_KEYWORDS = (
    "投标函",
    "授权书",
    "身份证明",
    "廉洁",
    "承诺函",
    "承诺书",
    "价格表",
    "开标价格",
    "规格一览表",
    "商务偏差",
    "供货范围",
    "保证金",
    "履约保证函",
    "质量保函",
    "业绩情况表",
    "资格证明",
    "声明函",
)
CERTIFICATE_KEYWORDS = (
    "认证证书",
    "型式认证",
    "设计认证",
    "符合证明",
    "产品认证",
    "大部件",
    "证书",
)
ASSET_KEYWORDS = (
    "订单",
    "合同",
    "业绩",
    "中标通知书",
    "运行证明",
    "验收报告",
    "试运行",
    "供货业绩",
)
ORDER_TABLE_HEADER_GROUPS = (
    ("项目名称", "项目", "工程名称", "风场", "电场"),
    ("订单", "订单号", "订单编号", "采购订单", "PO", "合同", "合同编号", "中标通知书", "通知书"),
    ("客户", "业主", "买方", "卖方", "招标人", "采购人", "单位名称"),
    ("机型", "型号", "规格", "产品型号", "设备名称", "供货范围"),
    ("容量", "MW", "兆瓦", "数量", "台数", "金额", "装机", "功率"),
    ("时间", "日期", "年份", "签订", "投运", "验收", "供货", "交货"),
)
ORDER_ROW_SKIP_RE = re.compile(r"^\s*(?:合计|小计|总计|备注|说明|序号|编号)\s*$")
BUSINESS_SECTION_RE = re.compile(
    r"^\s*(?:(?:附件|附表)\s*[0-9０-９一二三四五六七八九十]+|[0-9０-９]+(?:[.．、][0-9０-９]+)*|[一二三四五六七八九十]+[、.．])\s*[-—:：、.．]?\s*(.+)$"
)


async def preview_business_material_split(file_id: str, *, target_path: str = "", ai_mode: str = "auto") -> dict[str, Any]:
    source = await _load_business_source_file(file_id)
    fragments = [] if ai_mode == "force" else _build_split_plan(source, target_path=target_path)
    diagnostics: dict[str, Any] = {"strategy": "deterministic", "aiAttempted": False, "aiError": "", "aiMode": ai_mode}
    if ai_mode != "force" and (not fragments or _looks_like_single_collection_fragment(fragments)):
        semantic_fragments = _build_local_semantic_split_plan(source, target_path=target_path)
        if semantic_fragments:
            fragments = semantic_fragments
            diagnostics["strategy"] = "local_semantic"
        elif not fragments:
            fragments = []
            diagnostics["strategy"] = "deterministic_empty"
    if ai_mode != "force" and not fragments:
        fragments = _build_local_semantic_split_plan(source, target_path=target_path)
        diagnostics["strategy"] = "local_semantic" if fragments else "deterministic_empty"
    should_run_ai = ai_mode == "force" or not fragments or _should_ai_enhance_split(source, fragments)
    if should_run_ai:
        diagnostics["aiAttempted"] = True
        try:
            ai_fragments = await asyncio.to_thread(
                _build_ai_semantic_split_plan,
                source,
                target_path=target_path,
                exhaustive=ai_mode == "force",
            )
            if ai_fragments:
                fragments = _merge_split_fragments(fragments if ai_mode != "force" else [], ai_fragments)
                diagnostics["strategy"] = "ai_enhanced" if fragments else "ai_semantic_empty"
                diagnostics["aiFragmentCount"] = len(ai_fragments)
            elif not fragments:
                diagnostics["strategy"] = "ai_semantic_empty"
        except Exception as exc:  # AI is a fallback; the user can still manually split when it is unavailable.
            diagnostics["aiError"] = str(exc)[:300]

    return {
        "schemaVersion": BUSINESS_SPLIT_SCHEMA,
        "message": f"已生成 {len(fragments)} 个切分建议，请审核后确认入库。",
        "source": _source_payload(source),
        "fragments": fragments,
        "limits": {
            "supported": "支持商务标 .docx 的标题/附件/证书边界切分、订单/合同/业绩汇总表按行切分；多订单集合会自动调用 opencode AI 增强切片，也可手动强制 AI 重新识别。",
            "requiresReview": True,
        },
        "diagnostics": diagnostics,
    }


async def confirm_business_material_split(
    file_id: str,
    *,
    fragments: list[dict[str, Any]],
    default_target_path: str = "",
    on_conflict: str = "",
) -> dict[str, Any]:
    source = await _load_business_source_file(file_id)
    plan_items = _build_split_plan(source, target_path=default_target_path)
    if not plan_items:
        plan_items = _build_local_semantic_split_plan(source, target_path=default_target_path)
    by_id = {str(item["id"]): item for item in plan_items}
    document = Document(BytesIO(source["content"]))
    blocks = list(_iter_docx_blocks(document))
    selected: list[dict[str, Any]] = []
    for item in fragments:
        fragment_id = str(item.get("id") or "")
        planned = by_id.get(fragment_id)
        if not planned and isinstance(item.get("sourceLocation"), dict):
            planned = _client_fragment_from_payload(source, blocks, item, default_target_path=default_target_path)
        if not planned:
            continue
        if item.get("selected") is False:
            continue
        material_type = str(item.get("materialType") or planned.get("materialType") or "商务素材片段")
        if material_type == "业绩订单":
            raise PeripheralError(
                400,
                "业绩订单类片段请走业绩库导入，不再写入商务标通用素材目录。",
                "BUSINESS_SPLIT_PERFORMANCE_REQUIRES_LIBRARY",
            )
        selected.append(
            {
                **planned,
                "title": str(item.get("title") or planned.get("title") or "").strip() or planned["title"],
                "suggestedFileName": safe_segment(
                    str(item.get("fileName") or item.get("suggestedFileName") or planned.get("suggestedFileName") or ""),
                    planned["suggestedFileName"],
                ),
                "suggestedPath": str(item.get("targetPath") or item.get("suggestedPath") or planned.get("suggestedPath") or default_target_path or source["folderPath"]).strip().strip("/"),
                "materialType": material_type,
            }
        )

    if not selected:
        raise PeripheralError(400, "请选择至少一个切分片段。", "BUSINESS_SPLIT_NO_FRAGMENT_SELECTED")

    uploaded_items: list[dict[str, Any]] = []
    grouped: dict[str, list[dict[str, Any]]] = {}
    for fragment in selected:
        docx_bytes = _render_fragment_docx(source, document, fragment)
        target_path = fragment["suggestedPath"] or source["folderPath"]
        grouped.setdefault(target_path, []).append(
            {
                "name": _ensure_docx_name(fragment["suggestedFileName"]),
                "type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "data": docx_bytes,
                "relativePath": "",
                "extFields": _split_ext_fields(source, fragment),
            }
        )

    for target_path, files in grouped.items():
        result = await _upload_business_split_files(
            target_path=target_path,
            on_conflict=on_conflict,
            files=files,
        )
        uploaded_items.extend(result.get("items") or [])

    return {
        "schemaVersion": BUSINESS_SPLIT_SCHEMA,
        "message": f"切分入库完成，共生成 {len(uploaded_items)} 个子素材。",
        "source": _source_payload(source),
        "items": uploaded_items,
        "fragments": selected,
    }


async def _upload_business_split_files(
    *,
    target_path: str,
    on_conflict: str,
    files: list[dict[str, Any]],
) -> dict[str, Any]:
    from app.services.business_material_store import business_material_store

    return await business_material_store.raw_upload(
        target_path=target_path,
        on_conflict=on_conflict,
        files=files,
    )


async def _load_business_source_file(file_id: str) -> dict[str, Any]:
    numeric_id = _numeric_raw_id(file_id)
    async with async_session() as session:
        result = await session.execute(
            select(RawFile).where(RawFile.id == numeric_id).options(selectinload(RawFile.folder))
        )
        item = result.scalar_one_or_none()
        if item is None:
            raise PeripheralError(404, "素材文件不存在。", "RAW_FILE_NOT_FOUND")
        if item.folder is None:
            raise PeripheralError(400, "素材目录信息缺失。", "RAW_FILE_FOLDER_MISSING")
        payload = item.to_dict()
        ext = item.ext_fields or {}
        if str(ext.get("bidType") or payload.get("bidType") or "") != BUSINESS_BID_TYPE:
            raise PeripheralError(400, "素材切分入口当前仅支持商务标素材。", "BUSINESS_SPLIT_BID_TYPE_ONLY")
        if PurePosixPath(item.name).suffix.lower() != ".docx":
            raise PeripheralError(400, "第一版切分仅支持 .docx 文件；PDF/扫描件请先人工拆分或后续 OCR 增强。", "BUSINESS_SPLIT_DOCX_ONLY")
        if not item.minio_key:
            raise PeripheralError(400, "素材文件对象缺失，无法切分。", "RAW_FILE_OBJECT_MISSING")
        content = minio_client.get_object(str(item.minio_bucket or settings.minio_buckets["materials"]), str(item.minio_key))
        return {
            "id": f"RAW-{item.id:04d}",
            "name": item.name,
            "folderPath": item.folder.path,
            "materialTier": payload.get("materialTier") or "",
            "materialTierLabel": payload.get("materialTierLabel") or "",
            "customerId": payload.get("customerId") or "",
            "customerName": payload.get("customerName") or "",
            "projectId": payload.get("projectId") or "",
            "projectCode": payload.get("projectCode") or "",
            "projectName": payload.get("projectName") or "",
            "minioBucket": item.minio_bucket,
            "minioKey": item.minio_key,
            "content": content,
            "size": int(item.size_bytes or 0),
        }


def _build_split_plan(source: dict[str, Any], *, target_path: str = "") -> list[dict[str, Any]]:
    document = Document(BytesIO(source["content"]))
    blocks = list(_iter_docx_blocks(document))
    fragments: list[dict[str, Any]] = []
    boundaries: list[int] = []
    for index, block in enumerate(blocks):
        if block["kind"] != "paragraph":
            continue
        text = _clean_text(block["text"])
        if _is_fragment_boundary(text, block.get("style", "")):
            boundaries.append(index)

    for ordinal, start in enumerate(boundaries, start=1):
        end = boundaries[ordinal] if ordinal < len(boundaries) else len(blocks)
        title = _fragment_title(blocks[start]["text"], ordinal)
        material_type = _classify_fragment(title)
        suggested_path = _suggest_target_path(source, material_type, title, target_path)
        content_preview = _fragment_preview(blocks[start:end])
        risks = _risk_tips(title, content_preview, material_type)
        fragments.append(
            {
                "id": f"frag-{ordinal:03d}",
                "title": title,
                "materialType": material_type,
                "suggestedFileName": _suggest_file_name(title, ordinal),
                "suggestedPath": suggested_path,
                "sourceLocation": {
                    "blockStart": start,
                    "blockEnd": end,
                    "heading": title,
                },
                "contentPreview": content_preview,
                "confidence": _confidence(title, content_preview, risks),
                "riskTips": risks,
                "selected": material_type != "业绩订单" and (len(content_preview) >= 20 or material_type in {"证书", "承诺书模板", "商务附件模板"}),
            }
        )
    fragments.extend(_build_order_table_fragments(source, blocks, target_path=target_path, start_ordinal=len(fragments) + 1))
    return fragments


def _should_ai_enhance_split(source: dict[str, Any], fragments: list[dict[str, Any]]) -> bool:
    if not _looks_like_asset_collection_source(source):
        return False
    document = Document(BytesIO(source["content"]))
    blocks = list(_iter_docx_blocks(document))
    signals = _asset_signal_count(blocks)
    current_count = len([item for item in fragments if item.get("materialType") == "业绩订单"])
    if signals >= 5 and current_count < max(3, signals // 2):
        return True
    if signals >= 3 and len(fragments) <= 2:
        return True
    return False


def _looks_like_asset_collection_source(source: dict[str, Any]) -> bool:
    text = f"{source.get('name') or ''} {source.get('folderPath') or ''}"
    if any(keyword in text for keyword in (*ASSET_KEYWORDS, "订单", "合同", "业绩", "台账", "清单", "明细", "汇总")):
        return True
    try:
        document = Document(BytesIO(source["content"]))
        blocks = list(_iter_docx_blocks(document))
    except Exception:
        return False
    return _asset_signal_count(blocks) >= 3


def _asset_signal_count(blocks: list[dict[str, Any]]) -> int:
    signal_text = "\n".join(_clean_text(block.get("text") or "") for block in blocks[:220])
    patterns = (
        r"20[0-9]{2}\s*年",
        r"\b(?:HT|PO|SO|CG|ZB|NO|No)[-_:/]?[A-Za-z0-9-]{3,}\b",
        r"[\u4e00-\u9fff]{2,}(?:项目|风电|风场|电场)",
        r"(?:EW|SE|GWH|MY)[A-Za-z0-9.\\-]*",
        r"[0-9]+(?:\\.[0-9]+)?\\s*(?:MW|兆瓦|台)",
        r"(?:订单|合同|业绩|中标|供货|验收|投运)",
    )
    count = 0
    for pattern in patterns:
        count += len(re.findall(pattern, signal_text, flags=re.IGNORECASE))
    return count


def _merge_split_fragments(base: list[dict[str, Any]], ai_fragments: list[dict[str, Any]]) -> list[dict[str, Any]]:
    merged: list[dict[str, Any]] = []
    seen: set[str] = set()
    for item in [*ai_fragments, *base]:
        key = _fragment_dedupe_key(item)
        if key in seen:
            continue
        seen.add(key)
        next_item = dict(item)
        next_item["id"] = f"frag-{len(merged) + 1:03d}"
        next_item["suggestedFileName"] = _suggest_file_name(str(next_item.get("title") or "商务素材片段"), len(merged) + 1)
        merged.append(next_item)
    return merged


def _fragment_dedupe_key(item: dict[str, Any]) -> str:
    source_location = item.get("sourceLocation") or {}
    quote = _clean_text(source_location.get("quote") or item.get("contentPreview") or "")
    if quote:
        return quote[:120]
    return f"{source_location.get('blockStart')}:{source_location.get('blockEnd')}:{item.get('title')}"


def _looks_like_single_collection_fragment(fragments: list[dict[str, Any]]) -> bool:
    if len(fragments) != 1:
        return False
    item = fragments[0]
    title = str(item.get("title") or "")
    preview = str(item.get("contentPreview") or "")
    if _looks_like_asset_summary_title(title):
        return True
    marker_count = len(re.findall(r"(?:订单|合同|项目|风电|20[0-9]{2}|HT[-_0-9A-Za-z]+|PO[-_0-9A-Za-z]+)", preview, flags=re.IGNORECASE))
    return marker_count >= 6 and len(preview) > 80


def _build_local_semantic_split_plan(source: dict[str, Any], *, target_path: str = "") -> list[dict[str, Any]]:
    document = Document(BytesIO(source["content"]))
    blocks = list(_iter_docx_blocks(document))
    boundary_indexes = _semantic_boundary_indexes(source, blocks)
    if not boundary_indexes:
        return []
    return _fragments_from_boundaries(source, blocks, boundary_indexes, target_path=target_path, method="localSemantic", confidence_base=0.62)


def _build_ai_semantic_split_plan(source: dict[str, Any], *, target_path: str = "", exhaustive: bool = False) -> list[dict[str, Any]]:
    document = Document(BytesIO(source["content"]))
    blocks = list(_iter_docx_blocks(document))
    samples = _block_samples_for_ai(blocks, exhaustive=exhaustive)
    if not samples:
        return []
    batches = _ai_sample_batches(samples, max_batches=6 if exhaustive else 3)
    fragments: list[dict[str, Any]] = []
    ordinal = 1
    for batch_index, batch in enumerate(batches, start=1):
        prompt = _build_ai_split_prompt(source, batch, batch_index=batch_index, batch_count=len(batches))
        result = _send_business_split_ai_prompt(prompt)
        payload = _parse_ai_split_reply(str(result.get("reply") or ""))
        raw_fragments = payload.get("fragments") if isinstance(payload, dict) else None
        if not isinstance(raw_fragments, list):
            continue
        for item in raw_fragments[:40]:
            if not isinstance(item, dict):
                continue
            fragment = _client_fragment_from_payload(source, blocks, item, default_target_path=target_path, ordinal=ordinal)
            if fragment:
                fragments.append(fragment)
                ordinal += 1
    return fragments


def _send_business_split_ai_prompt(prompt: str) -> dict[str, Any]:
    config = system_settings_service.get_opencode_model_config_sync()
    base_url = str(config.get("baseUrl") or "").strip()
    model = str(config.get("model") or config.get("modelId") or "").strip()
    enabled = bool(config.get("enabled"))
    if enabled and base_url and model:
        try:
            return _send_openai_compatible_prompt(config, prompt)
        except Exception as exc:
            direct_error = str(exc)
        provider_id = str(config.get("providerId") or "").strip().lower()
        # A configured external LLM gateway should not fall back to opencode: opencode may not know
        # the provider/model and would hide the actual gateway error with a second unrelated error.
        if provider_id and provider_id not in {"opencode", "futurecode"}:
            raise RuntimeError(f"LLM 直连失败：{direct_error}") from None
        try:
            result = OpencodeClient(timeout_ms=90_000).send_text_prompt("商务素材语义切片", prompt)
            if not str(result.get("reply") or "").strip() and direct_error:
                raise RuntimeError(f"LLM 直连失败：{direct_error}；opencode 返回空响应。")
            return result
        except Exception as opencode_exc:
            raise RuntimeError(f"LLM 直连失败：{direct_error}；opencode 调用失败：{opencode_exc}") from opencode_exc
    return OpencodeClient(timeout_ms=90_000).send_text_prompt("商务素材语义切片", prompt)


def _send_openai_compatible_prompt(config: dict[str, Any], prompt: str) -> dict[str, Any]:
    base_url = str(config.get("baseUrl") or "").strip().rstrip("/")
    model = str(config.get("model") or config.get("modelId") or "").strip()
    if not base_url or not model:
        raise RuntimeError("LLM 配置缺少 Base URL 或模型。")
    url = _chat_completions_url(base_url)
    timeout_ms = min(300_000, max(90_000, int(config.get("timeoutMs") or 90_000)))
    # The global settings page may use a very small value for connection tests.
    # Material splitting needs enough room to return many JSON fragments.
    max_tokens = min(8192, max(2048, int(config.get("maxTokens") or 4096)))
    headers = {"Content-Type": "application/json"}
    api_key = str(config.get("apiKey") or "").strip()
    if api_key:
        headers["Authorization"] = f"Bearer {api_key}"
    payload: dict[str, Any] = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": "你是商务标素材库的资料切分助手。必须只输出严格 JSON，不要输出解释。",
            },
            {"role": "user", "content": prompt},
        ],
        "temperature": 0,
        "stream": False,
        "max_tokens": max_tokens,
        "response_format": {"type": "json_object"},
    }
    start = time.perf_counter()
    with httpx.Client(timeout=httpx.Timeout(max(1.0, timeout_ms / 1000), connect=10.0), trust_env=False) as client:
        response = client.post(url, headers=headers, json=payload)
        if response.status_code >= 400 and _looks_like_response_format_error(response.text):
            retry_payload = dict(payload)
            retry_payload.pop("response_format", None)
            response = client.post(url, headers=headers, json=retry_payload)
    if response.status_code >= 400:
        raise RuntimeError(f"HTTP {response.status_code}: {response.text[:500]}")
    try:
        raw = response.json()
    except ValueError as exc:
        raise RuntimeError(f"返回了非 JSON 响应：{response.text[:300]}") from exc
    reply = _extract_openai_compatible_reply(raw)
    if not reply:
        raise RuntimeError(f"返回了空内容（{_openai_response_debug(raw)}）。")
    return {
        "sessionId": "",
        "providerId": str(config.get("providerId") or "llm-gateway"),
        "modelId": model,
        "reply": reply,
        "opencodeOutput": {
            "summary": f"LLM 直连完成，用时 {int((time.perf_counter() - start) * 1000)} ms。",
            "raw": raw,
        },
    }


def _chat_completions_url(base_url: str) -> str:
    url = str(base_url or "").strip().rstrip("/")
    if not url:
        return ""
    if url.endswith("/chat/completions"):
        return url
    if url.endswith("/v1"):
        return f"{url}/chat/completions"
    return f"{url}/chat/completions"


def _looks_like_response_format_error(text: str) -> bool:
    lowered = str(text or "").lower()
    return "response_format" in lowered and any(marker in lowered for marker in ("unsupported", "not support", "unknown", "invalid", "extra"))


def _looks_like_json_object(text: str) -> bool:
    stripped = str(text or "").strip()
    return stripped.startswith("{") and stripped.endswith("}")


def _extract_openai_compatible_reply(raw: dict[str, Any]) -> str:
    choices = raw.get("choices") if isinstance(raw, dict) else None
    if isinstance(choices, list):
        for choice in choices:
            if not isinstance(choice, dict):
                continue
            message = choice.get("message") if isinstance(choice.get("message"), dict) else {}
            for value in (
                message.get("content"),
                message.get("text"),
                choice.get("text"),
                choice.get("output_text"),
                (choice.get("delta") or {}).get("content") if isinstance(choice.get("delta"), dict) else "",
            ):
                text = _content_to_text(value)
                if text:
                    return text
            reasoning = _content_to_text(message.get("reasoning_content"))
            if reasoning and _looks_like_json_object(reasoning):
                return reasoning
    for value in (raw.get("output_text"), raw.get("content"), raw.get("text")):
        text = _content_to_text(value)
        if text:
            return text
    return ""


def _content_to_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, list):
        parts: list[str] = []
        for item in value:
            if isinstance(item, str):
                parts.append(item)
            elif isinstance(item, dict):
                parts.append(str(item.get("text") or item.get("content") or ""))
        return "\n".join(part for part in (part.strip() for part in parts) if part).strip()
    return str(value).strip()


def _openai_response_debug(raw: dict[str, Any]) -> str:
    if not isinstance(raw, dict):
        return "响应不是 JSON 对象"
    choices = raw.get("choices") if isinstance(raw.get("choices"), list) else []
    choice = choices[0] if choices and isinstance(choices[0], dict) else {}
    message = choice.get("message") if isinstance(choice.get("message"), dict) else {}
    usage = raw.get("usage") if isinstance(raw.get("usage"), dict) else {}
    compact_usage = {
        key: usage.get(key)
        for key in ("prompt_tokens", "completion_tokens", "total_tokens", "completion_tokens_details")
        if key in usage
    }
    return (
        f"finish_reason={choice.get('finish_reason') or '-'}, "
        f"messageKeys={list(message.keys())}, "
        f"choiceKeys={list(choice.keys())}, "
        f"usage={compact_usage}"
    )


def _fragments_from_boundaries(
    source: dict[str, Any],
    blocks: list[dict[str, Any]],
    boundaries: list[int],
    *,
    target_path: str,
    method: str,
    confidence_base: float,
) -> list[dict[str, Any]]:
    unique_boundaries = sorted({idx for idx in boundaries if 0 <= idx < len(blocks)})
    if not unique_boundaries:
        return []
    fragments: list[dict[str, Any]] = []
    for ordinal, start in enumerate(unique_boundaries, start=1):
        next_boundary = unique_boundaries[ordinal] if ordinal < len(unique_boundaries) else len(blocks)
        end = _semantic_fragment_end(blocks, start, next_boundary)
        if end <= start:
            continue
        title = _semantic_fragment_title(blocks[start], ordinal)
        preview = _fragment_preview(blocks[start:end])
        if len(preview) < 12:
            continue
        material_type = _classify_fragment(title if title else preview)
        suggested_path = _suggest_target_path(source, material_type, title, target_path)
        risks = _risk_tips(title, preview, material_type)
        fragments.append(
            {
                "id": f"frag-{ordinal:03d}",
                "title": title,
                "materialType": material_type,
                "suggestedFileName": _suggest_file_name(title, ordinal),
                "suggestedPath": suggested_path,
                "sourceLocation": {
                    "mode": method,
                    "blockStart": start,
                    "blockEnd": end,
                    "heading": title,
                },
                "contentPreview": preview,
                "confidence": _confidence_with_base(title, preview, risks, confidence_base),
                "riskTips": [*risks, "该片段来自语义兜底切分，建议重点人工复核边界。"],
                "selected": True,
            }
        )
    return fragments


def _semantic_boundary_indexes(source: dict[str, Any], blocks: list[dict[str, Any]]) -> list[int]:
    indexes: list[int] = []
    source_text = f"{source.get('name') or ''} {source.get('folderPath') or ''}"
    asset_hint = any(keyword in source_text for keyword in ASSET_KEYWORDS)
    for index, block in enumerate(blocks):
        text = _clean_text(block.get("text") or "")
        if not text:
            continue
        if block.get("kind") == "paragraph":
            if _looks_like_semantic_material_start(text, asset_hint=asset_hint):
                indexes.append(index)
        elif block.get("kind") == "table":
            rows = block.get("rows") or []
            if _order_table_header_index(rows, source) >= 0:
                indexes.append(index)
    if len(indexes) == 1:
        return indexes
    return _dedupe_close_indexes(indexes)


def _looks_like_semantic_material_start(text: str, *, asset_hint: bool) -> bool:
    if len(text) > 180:
        return False
    if _is_fragment_boundary(text):
        return True
    patterns = (
        r"^\s*(?:订单|合同|业绩|项目|风场|电场|中标|验收|运行证明|试运行)\s*[:：#编号]*\s*[\wA-Za-z0-9一-龥-]{2,}",
        r"^\s*(?:PO|SO|HT|ZB|CG|NO\.?|No\.?)\s*[-_:：]?\s*[A-Za-z0-9-]{4,}",
        r".*(?:订单编号|合同编号|项目名称|客户名称|业主单位|买方|卖方|中标通知书编号)\s*[:：].+",
        r".*20[0-9]{2}\s*年.*(?:订单|合同|中标|供货|验收|投运).*",
    )
    if any(re.match(pattern, text, flags=re.IGNORECASE) for pattern in patterns):
        return True
    if asset_hint and 12 <= len(text) <= 120:
        keyword_hits = sum(1 for keyword in ASSET_KEYWORDS if keyword in text)
        if keyword_hits >= 1 and bool(re.search(r"20[0-9]{2}|[A-Za-z]{2,}[-_0-9]", text)):
            return True
        if _looks_like_asset_data_line(text):
            return True
    return False


def _looks_like_asset_data_line(text: str) -> bool:
    if len(text) < 12 or len(text) > 180:
        return False
    has_year_or_code = bool(re.search(r"20[0-9]{2}|[A-Za-z]{2,}[-_/]?[0-9]{3,}|[0-9]{4,}[-_/][0-9]{2,}", text))
    has_company_or_project = bool(re.search(r"[\u4e00-\u9fff]{2,}(?:项目|风电|风场|电场|公司|集团|有限|能源|华能|国能|大唐|华电|中广核)", text))
    has_capacity_or_model = bool(re.search(r"(?:EW|SE|GWH|MY|[0-9]+(?:\\.[0-9]+)?\\s*MW|[0-9]+(?:\\.[0-9]+)?兆瓦|[0-9]+台)", text, re.IGNORECASE))
    return has_year_or_code and (has_company_or_project or has_capacity_or_model)


def _semantic_fragment_end(blocks: list[dict[str, Any]], start: int, next_boundary: int) -> int:
    if next_boundary > start:
        return next_boundary
    return min(len(blocks), start + 8)


def _semantic_fragment_title(block: dict[str, Any], ordinal: int) -> str:
    text = _clean_text(block.get("text") or "")
    if block.get("kind") == "table":
        rows = block.get("rows") or []
        if len(rows) >= 2:
            return _order_row_title(rows[0], rows[1], ordinal)
        return f"表格素材片段{ordinal:02d}"
    return _fragment_title(text, ordinal)


def _dedupe_close_indexes(indexes: list[int]) -> list[int]:
    result: list[int] = []
    for index in sorted(set(indexes)):
        if result and index - result[-1] <= 1:
            continue
        result.append(index)
    return result


def _confidence_with_base(title: str, preview: str, risks: list[str], base: float) -> float:
    score = base
    if len(preview) >= 80:
        score += 0.12
    if any(keyword in f"{title}\n{preview}" for keyword in (*TEMPLATE_KEYWORDS, *CERTIFICATE_KEYWORDS, *ASSET_KEYWORDS)):
        score += 0.08
    if risks:
        score -= min(0.15, len(risks) * 0.05)
    return round(max(0.35, min(0.88, score)), 2)


def _block_samples_for_ai(
    blocks: list[dict[str, Any]],
    *,
    max_blocks: int = 260,
    max_chars: int = 36000,
    exhaustive: bool = False,
) -> list[dict[str, Any]]:
    samples: list[dict[str, Any]] = []
    total = 0
    for index, block in enumerate(blocks):
        text = _clean_text(block.get("text") or "")
        if not text:
            continue
        chunks = _ai_text_chunks(text, chunk_chars=1200 if exhaustive else 900)
        for chunk_index, sample_text in enumerate(chunks, start=1):
            total += len(sample_text)
            samples.append(
                {
                    "index": index,
                    "kind": block.get("kind") or "",
                    "part": chunk_index,
                    "parts": len(chunks),
                    "text": sample_text,
                }
            )
            if len(samples) >= max_blocks or total >= max_chars:
                return samples
        if len(samples) >= max_blocks or total >= max_chars:
            break
    return samples


def _ai_text_chunks(text: str, *, chunk_chars: int, overlap_chars: int = 120) -> list[str]:
    cleaned = _clean_text(text)
    if len(cleaned) <= chunk_chars:
        return [cleaned]
    chunks: list[str] = []
    start = 0
    while start < len(cleaned):
        end = min(len(cleaned), start + chunk_chars)
        if end < len(cleaned):
            end = _ai_chunk_boundary(cleaned, start + int(chunk_chars * 0.6), end)
        chunk = cleaned[start:end].strip(" ；;。,\n\t")
        if chunk:
            chunks.append(chunk)
        if end >= len(cleaned):
            break
        start = max(start + 1, end - overlap_chars)
    return chunks


def _ai_chunk_boundary(text: str, lower: int, upper: int) -> int:
    for separator in ("；", ";", "。", "\n", "，", ",", " "):
        pos = text.rfind(separator, lower, upper)
        if pos >= lower:
            return pos + 1
    return upper


def _ai_sample_batches(
    samples: list[dict[str, Any]],
    *,
    max_batch_chars: int = 12000,
    max_batch_samples: int = 80,
    max_batches: int = 3,
) -> list[list[dict[str, Any]]]:
    batches: list[list[dict[str, Any]]] = []
    current: list[dict[str, Any]] = []
    current_chars = 0
    for sample in samples:
        sample_chars = len(str(sample.get("text") or ""))
        if current and (current_chars + sample_chars > max_batch_chars or len(current) >= max_batch_samples):
            batches.append(current)
            if len(batches) >= max_batches:
                return batches
            current = []
            current_chars = 0
        current.append(sample)
        current_chars += sample_chars
    if current and len(batches) < max_batches:
        batches.append(current)
    return batches


def _build_ai_split_prompt(source: dict[str, Any], samples: list[dict[str, Any]], *, batch_index: int = 1, batch_count: int = 1) -> str:
    sample_lines = "\n".join(
        f"[{item['index']}#{item.get('part', 1)}/{item.get('parts', 1)}] {item['kind']}: {item['text']}"
        for item in samples
    )
    return f"""你是商务标素材库的资料切分助手。请根据 Word 文档的文本块列表，判断是否存在可独立入库的商务素材片段。

目标：
1. 优先识别订单、合同、业绩、证书、承诺书、投标附件模板等可以作为独立素材的片段。
2. 如果一个文本块里包含多条订单/合同/业绩，请逐条拆分。此时 blockStart/blockEnd 可以相同，也可以填 blockEnd=blockStart+1，但必须填写 quote，quote 必须是原文中的连续摘录。
3. 只返回原文切片边界或原文 quote，不要改写正文。
4. 一个片段必须有明确主题，不能把整份汇总文件作为一个片段。
5. 请尽量完整识别本批次内的所有独立订单/合同/业绩，不要只返回示例。
6. 文本块编号里的 #1/3 表示同一个 Word 块被拆成多个展示片段；输出 blockStart/blockEnd 时仍然使用 # 前面的原始块编号。
7. 如果不能可靠切分，返回空数组。

输出必须是严格 JSON：
{{
  "fragments": [
    {{
      "title": "片段标题",
      "materialType": "业绩订单/证书/承诺书模板/商务附件模板/商务素材片段",
      "blockStart": 0,
      "blockEnd": 3,
      "quote": "如果需要在同一文本块内细切，填写原文连续摘录",
      "reason": "为什么这样切"
    }}
  ]
}}

文件名：{source.get('name') or ''}
所在路径：{source.get('folderPath') or ''}
当前批次：{batch_index}/{batch_count}

文本块：
{sample_lines}
"""


def _parse_ai_split_reply(reply: str) -> dict[str, Any]:
    text = str(reply or "").strip()
    if not text:
        return {}
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", text, re.S)
        if not match:
            return {}
        try:
            return json.loads(match.group(0))
        except json.JSONDecodeError:
            return {}


def _client_fragment_from_payload(
    source: dict[str, Any],
    blocks: list[dict[str, Any]],
    item: dict[str, Any],
    *,
    default_target_path: str,
    ordinal: int | None = None,
) -> dict[str, Any] | None:
    try:
        raw_start = item.get("blockStart") if item.get("blockStart") is not None else (item.get("sourceLocation") or {}).get("blockStart")
        raw_end = item.get("blockEnd") if item.get("blockEnd") is not None else (item.get("sourceLocation") or {}).get("blockEnd")
        start = _coerce_ai_block_index(raw_start)
        end = _coerce_ai_block_index(raw_end)
    except (TypeError, ValueError):
        return None
    if start < 0 or start >= len(blocks):
        return None
    end = max(start + 1, min(end, len(blocks)))
    title = safe_segment(str(item.get("title") or _semantic_fragment_title(blocks[start], ordinal or 1)), f"商务素材片段{ordinal or 1:02d}")[:90]
    material_type = str(item.get("materialType") or _classify_fragment(title) or "商务素材片段")
    if material_type not in {"业绩订单", "证书", "承诺书模板", "商务附件模板", "商务素材片段"}:
        material_type = _classify_fragment(title)
    quote = _clean_text(item.get("quote") or item.get("text") or "")
    preview = quote or _fragment_preview(blocks[start:end])
    if len(preview) < 8:
        return None
    risks = _risk_tips(title, preview, material_type)
    reason = str(item.get("reason") or "").strip()
    if reason:
        risks.append(f"AI切分理由：{reason[:120]}")
    risks.append("该片段来自语义兜底切分，建议重点人工复核边界。")
    frag_ordinal = ordinal or _extract_fragment_ordinal(str(item.get("id") or "")) or 1
    return {
        "id": str(item.get("id") or f"frag-{frag_ordinal:03d}"),
        "title": title,
        "materialType": material_type,
        "suggestedFileName": _suggest_file_name(title, frag_ordinal),
        "suggestedPath": str(item.get("targetPath") or item.get("suggestedPath") or default_target_path or _suggest_target_path(source, material_type, title, "")),
        "sourceLocation": {
            "mode": str((item.get("sourceLocation") or {}).get("mode") or "aiSemantic"),
            "blockStart": start,
            "blockEnd": end,
            "heading": title,
            "quote": quote,
        },
        "contentPreview": preview,
        "confidence": _confidence_with_base(title, preview, risks, 0.58),
        "riskTips": risks,
        "selected": item.get("selected") is not False,
    }


def _coerce_ai_block_index(value: Any) -> int:
    if isinstance(value, int):
        return value
    text = str(value or "").strip()
    match = re.match(r"^([0-9]+)(?:#.*)?$", text)
    if not match:
        raise ValueError(f"invalid block index: {value}")
    return int(match.group(1))


def _extract_fragment_ordinal(fragment_id: str) -> int:
    match = re.search(r"([0-9]+)$", fragment_id)
    return int(match.group(1)) if match else 0


def _iter_docx_blocks(document: DocxDocument) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, document)
            items.append(
                {
                    "kind": "paragraph",
                    "text": paragraph.text or "",
                    "style": paragraph.style.name if paragraph.style is not None else "",
                    "element": child,
                }
            )
        elif child.tag == qn("w:tbl"):
            table = Table(child, document)
            rows = _table_rows(table)
            items.append(
                {
                    "kind": "table",
                    "text": _table_text(table),
                    "rows": rows,
                    "element": child,
                }
            )
    return items


def _is_fragment_boundary(text: str, style_name: str = "") -> bool:
    if not text or len(text) > 120:
        return False
    all_heading_keywords = (*TEMPLATE_KEYWORDS, *CERTIFICATE_KEYWORDS, *ASSET_KEYWORDS)
    if _looks_like_asset_summary_title(text):
        return False
    if str(style_name or "").lower().startswith("heading") and any(keyword in text for keyword in all_heading_keywords):
        return True
    if BUSINESS_SECTION_RE.match(text) and any(keyword in text for keyword in all_heading_keywords):
        return True
    if re.match(r"^\s*(?:附件|附表)\s*[0-9０-９一二三四五六七八九十]+", text):
        return True
    if _looks_like_asset_title(text):
        return True
    if any(keyword in text for keyword in CERTIFICATE_KEYWORDS) and len(text) <= 60 and not _looks_like_certificate_field(text):
        return True
    return any(keyword in text for keyword in TEMPLATE_KEYWORDS) and len(text) <= 60


def _looks_like_certificate_field(text: str) -> bool:
    return bool(re.match(r"^\s*(?:证书编号|证书号|编号|有效期|发证日期|颁发日期|证书状态)\s*[:：]", text))


def _looks_like_asset_title(text: str) -> bool:
    if len(text) > 80:
        return False
    if _looks_like_asset_summary_title(text):
        return False
    if re.match(r"^\s*(?:订单|合同|业绩|中标通知书)\s*[0-9０-９一二三四五六七八九十]+", text):
        return True
    if any(keyword in text for keyword in ("供货业绩", "中标业绩")):
        return True
    return False


def _looks_like_asset_summary_title(text: str) -> bool:
    return any(
        keyword in text
        for keyword in (
            "订单汇总",
            "合同汇总",
            "业绩汇总",
            "订单清单",
            "合同清单",
            "业绩清单",
            "订单明细",
            "合同明细",
            "业绩明细",
            "订单列表",
            "合同列表",
            "业绩列表",
            "订单台账",
            "合同台账",
            "业绩台账",
            "订单资料",
            "合同资料",
            "业绩资料",
            "历史订单",
            "历史合同",
            "历史业绩",
        )
    )


def _fragment_title(text: str, ordinal: int) -> str:
    cleaned = _clean_text(text).strip(" ：:；;，,、")
    match = BUSINESS_SECTION_RE.match(cleaned)
    if match and match.group(1):
        cleaned = _clean_text(match.group(1)).strip(" ：:；;，,、")
    return safe_segment(cleaned, f"商务素材片段{ordinal:02d}")[:80]


def _classify_fragment(title: str) -> str:
    if any(keyword in title for keyword in CERTIFICATE_KEYWORDS):
        return "证书"
    if any(keyword in title for keyword in ASSET_KEYWORDS) or _looks_like_asset_data_line(title):
        return "业绩订单"
    if "承诺" in title:
        return "承诺书模板"
    if any(keyword in title for keyword in ("投标函", "授权书", "身份证明", "廉洁", "价格表", "规格", "偏差", "供货范围", "业绩情况表")):
        return "商务附件模板"
    return "商务素材片段"


def _suggest_target_path(source: dict[str, Any], material_type: str, title: str, target_path: str) -> str:
    if target_path:
        return target_path.strip().strip("/")
    folder_path = str(source.get("folderPath") or "").strip().strip("/")
    tier_path = _tier_root_path(source) or f"{BUSINESS_BID_TYPE}/通用素材"
    if material_type == "证书":
        return f"{BUSINESS_BID_TYPE}/通用素材/机型认证与测试报告"
    if material_type == "业绩订单":
        return ""
    if material_type in {"承诺书模板", "商务附件模板"}:
        if "/项目素材/" in folder_path:
            return f"{tier_path}/项目模板与过程稿"
        if "/客户素材/" in folder_path:
            return f"{tier_path}/客户模板与过程稿"
        return f"{BUSINESS_BID_TYPE}/通用素材/通用表单与模板"
    return folder_path or tier_path


def _tier_root_path(source: dict[str, Any]) -> str:
    folder_path = str(source.get("folderPath") or "")
    parts = [part for part in folder_path.split("/") if part]
    if len(parts) >= 3 and parts[0] == BUSINESS_BID_TYPE and parts[1] in {"客户素材", "项目素材"}:
        return "/".join(parts[:3])
    if len(parts) >= 2 and parts[0] == BUSINESS_BID_TYPE and parts[1] == "通用素材":
        return f"{BUSINESS_BID_TYPE}/通用素材"
    return ""


def _suggest_file_name(title: str, ordinal: int) -> str:
    return _ensure_docx_name(safe_segment(title, f"商务素材片段{ordinal:02d}")[:90])


def _ensure_docx_name(value: str) -> str:
    name = safe_segment(value, "商务素材片段.docx")
    if not name.lower().endswith(".docx"):
        name = f"{PurePosixPath(name).stem}.docx"
    return name


def _fragment_preview(blocks: list[dict[str, Any]], limit: int = 420) -> str:
    chunks: list[str] = []
    for block in blocks[:10]:
        text = _clean_text(block.get("text") or "")
        if text:
            chunks.append(text)
    preview = "\n".join(chunks)
    return preview[:limit]


def _risk_tips(title: str, preview: str, material_type: str) -> list[str]:
    tips: list[str] = []
    if len(preview) < 20:
        tips.append("片段文本较少，可能只识别到标题，需人工确认是否入库。")
    if material_type == "证书" and not re.search(r"20[0-9]{2}[-./年]", preview):
        tips.append("未识别到明显有效期，证书有效期需人工复核。")
    if material_type == "业绩订单":
        tips.append("业绩订单类片段不再写入通用素材目录；请走业绩库导入或合同包上传链路。")
        if not any(keyword in preview for keyword in ("合同", "订单", "中标", "项目")):
            tips.append("未识别到明显项目/合同/订单字段，需人工确认是否为独立业绩素材。")
        if not re.search(r"20[0-9]{2}", preview):
            tips.append("未识别到明显年份，业绩时间需人工复核。")
    if "[embedded_image" in preview or "图片" in preview:
        tips.append("片段可能包含图片，第一版切分以文字/表格为主，图片需人工复核。")
    return tips


def _confidence(title: str, preview: str, risks: list[str]) -> float:
    score = 0.72
    if len(preview) >= 80:
        score += 0.12
    if any(keyword in title for keyword in (*TEMPLATE_KEYWORDS, *CERTIFICATE_KEYWORDS, *ASSET_KEYWORDS)):
        score += 0.08
    if risks:
        score -= min(0.2, len(risks) * 0.08)
    return round(max(0.35, min(0.95, score)), 2)


def _render_fragment_docx(source: dict[str, Any], document: DocxDocument, fragment: dict[str, Any]) -> bytes:
    blocks = list(_iter_docx_blocks(document))
    source_location = fragment.get("sourceLocation") or {}
    if source_location.get("mode") == "tableRow":
        return _render_table_row_fragment(source, blocks, fragment)
    if source_location.get("quote"):
        return _render_quote_fragment(fragment)
    start = int((fragment.get("sourceLocation") or {}).get("blockStart") or 0)
    end = int((fragment.get("sourceLocation") or {}).get("blockEnd") or len(blocks))
    preserved = _render_body_slice_preserving_media(source["content"], start, end, title=str(fragment.get("title") or "商务素材片段"))
    if preserved:
        return preserved
    output = Document()
    output.add_heading(str(fragment.get("title") or "商务素材片段"), level=1)
    for block in blocks[start:end]:
        if block["kind"] == "paragraph":
            text = block.get("text") or ""
            if text:
                output.add_paragraph(text)
        elif block["kind"] == "table":
            _append_table(output, block.get("text") or "")
    buffer = BytesIO()
    output.save(buffer)
    return buffer.getvalue()


def _render_quote_fragment(fragment: dict[str, Any]) -> bytes:
    output = Document()
    output.add_heading(str(fragment.get("title") or "商务素材片段"), level=1)
    output.add_paragraph(str((fragment.get("sourceLocation") or {}).get("quote") or fragment.get("contentPreview") or ""))
    buffer = BytesIO()
    output.save(buffer)
    return buffer.getvalue()


def _render_table_row_fragment(source: dict[str, Any], blocks: list[dict[str, Any]], fragment: dict[str, Any]) -> bytes:
    source_location = fragment.get("sourceLocation") or {}
    block_index = int(source_location.get("tableBlockIndex") or source_location.get("blockStart") or 0)
    header_index = int(source_location.get("headerRowIndex") or 0)
    row_index = int(source_location.get("rowIndex") or 0)
    block = blocks[block_index] if 0 <= block_index < len(blocks) else {}
    rows = block.get("rows") or []
    header = rows[header_index] if 0 <= header_index < len(rows) else []
    row = rows[row_index] if 0 <= row_index < len(rows) else []
    preserved = _render_table_row_preserving_media(
        source["content"],
        block_index=block_index,
        header_index=header_index,
        row_index=row_index,
        title=str(fragment.get("title") or "商务业绩订单片段"),
    )
    if preserved:
        return preserved

    output = Document()
    output.add_heading(str(fragment.get("title") or "商务业绩订单片段"), level=1)
    if header and row:
        column_count = max(len(header), len(row))
        table = output.add_table(rows=2, cols=column_count)
        table.style = "Table Grid"
        for col_index in range(column_count):
            table.cell(0, col_index).text = header[col_index] if col_index < len(header) else ""
            table.cell(1, col_index).text = row[col_index] if col_index < len(row) else ""
    else:
        output.add_paragraph(str(fragment.get("contentPreview") or ""))

    buffer = BytesIO()
    output.save(buffer)
    return buffer.getvalue()


WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WORD_BODY_TAG = f"{{{WORD_NS}}}body"
WORD_SECTPR_TAG = f"{{{WORD_NS}}}sectPr"
WORD_P_TAG = f"{{{WORD_NS}}}p"
WORD_TBL_TAG = f"{{{WORD_NS}}}tbl"
WORD_TR_TAG = f"{{{WORD_NS}}}tr"


def _render_body_slice_preserving_media(source_docx: bytes, start: int, end: int, *, title: str) -> bytes:
    try:
        state = _docx_zip_state(source_docx)
        body_children = state["bodyChildren"]
        if start < 0 or start >= len(body_children) or end <= start:
            return b""
        selected_children = [
            deepcopy(child)
            for child in body_children[start : min(end, len(body_children))]
            if child.tag != WORD_SECTPR_TAG
        ]
        if not selected_children:
            return b""
        return _assemble_docx_from_body_children(state, selected_children, title=title)
    except Exception:
        return b""


def _render_table_row_preserving_media(source_docx: bytes, *, block_index: int, header_index: int, row_index: int, title: str) -> bytes:
    try:
        state = _docx_zip_state(source_docx)
        body_children = state["bodyChildren"]
        if block_index < 0 or block_index >= len(body_children):
            return b""
        table = body_children[block_index]
        if table.tag != WORD_TBL_TAG:
            return b""
        rows = [child for child in table.iterchildren() if child.tag == WORD_TR_TAG]
        if header_index < 0 or row_index < 0 or header_index >= len(rows) or row_index >= len(rows):
            return b""
        new_table = deepcopy(table)
        for child in list(new_table.iterchildren()):
            if child.tag == WORD_TR_TAG:
                new_table.remove(child)
        new_table.append(deepcopy(rows[header_index]))
        if row_index != header_index:
            new_table.append(deepcopy(rows[row_index]))
        return _assemble_docx_from_body_children(state, [new_table], title=title)
    except Exception:
        return b""


def _docx_zip_state(source_docx: bytes) -> dict[str, Any]:
    with zipfile.ZipFile(BytesIO(source_docx), "r") as zf:
        parts: dict[str, tuple[zipfile.ZipInfo, bytes]] = {
            info.filename: (info, zf.read(info.filename))
            for info in zf.infolist()
        }
    root = etree.fromstring(parts["word/document.xml"][1])
    body = root.find(WORD_BODY_TAG)
    if body is None:
        raise ValueError("docx body missing")
    return {
        "parts": parts,
        "rootTag": root.tag,
        "rootAttrib": dict(root.attrib),
        "rootNsmap": dict(root.nsmap),
        "rootChildrenBeforeBody": [deepcopy(child) for child in root.iterchildren() if child.tag != WORD_BODY_TAG],
        "bodyTag": body.tag,
        "bodyAttrib": dict(body.attrib),
        "bodyChildren": list(body.iterchildren()),
        "sectPr": body.find(WORD_SECTPR_TAG),
    }


def _assemble_docx_from_body_children(state: dict[str, Any], body_children: list[Any], *, title: str) -> bytes:
    root = etree.Element(state["rootTag"], attrib=state["rootAttrib"], nsmap=state["rootNsmap"])
    for sibling in state["rootChildrenBeforeBody"]:
        root.append(deepcopy(sibling))
    body = etree.SubElement(root, state["bodyTag"], attrib=state["bodyAttrib"])
    for child in _title_paragraph_elements(title):
        body.append(child)
    for child in body_children:
        if child.tag != WORD_SECTPR_TAG:
            body.append(deepcopy(child))
    if state.get("sectPr") is not None:
        body.append(deepcopy(state["sectPr"]))
    new_doc_xml = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as dst:
        for filename, (info, data) in state["parts"].items():
            dst.writestr(info, new_doc_xml if filename == "word/document.xml" else data)
    return buffer.getvalue()


def _title_paragraph_elements(title: str) -> list[Any]:
    tmp = Document()
    tmp.add_heading(str(title or "商务素材片段"), level=1)
    body = tmp.element.body
    return [deepcopy(child) for child in body.iterchildren() if child.tag == WORD_P_TAG]


def _append_table(document: DocxDocument, table_text: str) -> None:
    rows = [
        [cell.strip() for cell in row.split("\t")]
        for row in str(table_text or "").split("\n")
        if row.strip()
    ]
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index in range(column_count):
            table.cell(row_index, col_index).text = row[col_index] if col_index < len(row) else ""


def _table_text(table: Table) -> str:
    return "\n".join("\t".join(row) for row in _table_rows(table))


def _table_rows(table: Table) -> list[list[str]]:
    rows: list[str] = []
    parsed: list[list[str]] = []
    for row in table.rows:
        parsed.append([_clean_text(cell.text) for cell in row.cells])
    return parsed


def _build_order_table_fragments(
    source: dict[str, Any],
    blocks: list[dict[str, Any]],
    *,
    target_path: str,
    start_ordinal: int,
) -> list[dict[str, Any]]:
    fragments: list[dict[str, Any]] = []
    ordinal = start_ordinal
    for block_index, block in enumerate(blocks):
        if block.get("kind") != "table":
            continue
        rows = block.get("rows") or []
        header_index = _order_table_header_index(rows, source)
        if header_index < 0:
            continue
        header = rows[header_index]
        for row_index in range(header_index + 1, len(rows)):
            row = rows[row_index]
            if not _is_valid_order_data_row(row, header):
                continue
            title = _order_row_title(header, row, ordinal)
            preview = _order_row_preview(header, row)
            material_type = "业绩订单"
            risks = _risk_tips(title, preview, material_type)
            fragments.append(
                {
                    "id": f"frag-{ordinal:03d}",
                    "title": title,
                    "materialType": material_type,
                    "suggestedFileName": _suggest_file_name(title, ordinal),
                    "suggestedPath": _suggest_target_path(source, material_type, title, target_path),
                    "sourceLocation": {
                        "mode": "tableRow",
                        "blockStart": block_index,
                        "blockEnd": block_index + 1,
                        "tableBlockIndex": block_index,
                        "headerRowIndex": header_index,
                        "rowIndex": row_index,
                        "heading": title,
                    },
                    "contentPreview": preview[:420],
                    "confidence": _confidence(title, preview, risks),
                    "riskTips": risks,
                    "selected": False,
                }
            )
            ordinal += 1
    return fragments


def _order_table_header_index(rows: list[list[str]], source: dict[str, Any]) -> int:
    if len(rows) < 2:
        return -1
    source_text = f"{source.get('name') or ''} {source.get('folderPath') or ''}"
    source_has_asset_hint = any(keyword in source_text for keyword in ASSET_KEYWORDS)
    for index, row in enumerate(rows[: min(4, len(rows) - 1)]):
        score = _order_header_score(row)
        next_valid_rows = sum(1 for candidate in rows[index + 1 :] if _is_valid_order_data_row(candidate, row))
        if score >= 2 and next_valid_rows >= 1:
            return index
        if source_has_asset_hint and score >= 1 and next_valid_rows >= 2:
            return index
    return -1


def _order_header_score(row: list[str]) -> int:
    cells = [_clean_text(cell) for cell in row if _clean_text(cell)]
    text = " ".join(cells)
    score = 0
    for group in ORDER_TABLE_HEADER_GROUPS:
        if any(keyword in text for keyword in group):
            score += 1
    if any(cell in {"序号", "编号"} for cell in cells):
        score += 1
    return score


def _is_valid_order_data_row(row: list[str], header: list[str]) -> bool:
    cells = [_clean_text(cell) for cell in row]
    non_empty = [cell for cell in cells if cell]
    if len(non_empty) < 2:
        return False
    row_text = " ".join(non_empty)
    header_text = " ".join(_clean_text(cell) for cell in header if _clean_text(cell))
    if row_text == header_text:
        return False
    if ORDER_ROW_SKIP_RE.match(non_empty[0]) and len(non_empty) <= 2:
        return False
    if any(keyword in row_text for keyword in ("合计", "小计", "总计")) and len(non_empty) <= 3:
        return False
    return True


def _order_row_title(header: list[str], row: list[str], ordinal: int) -> str:
    mapping = _row_mapping(header, row)
    project = _first_mapping_value(mapping, ("项目名称", "工程名称", "项目", "风场", "电场"))
    contract = _first_mapping_value(mapping, ("订单号", "订单", "合同编号", "合同", "中标通知书", "通知书"))
    customer = _first_mapping_value(mapping, ("客户", "业主", "买方", "招标人", "采购人", "单位名称"))
    model = _first_mapping_value(mapping, ("机型", "型号", "规格", "产品型号", "设备名称"))
    parts = _dedupe_non_empty([project, contract, customer, model])
    title = " - ".join(parts[:3]) if parts else ""
    if not title:
        title = _clean_text(" ".join(cell for cell in row if cell))[:50]
    return safe_segment(title, f"业绩订单素材{ordinal:02d}")[:90]


def _order_row_preview(header: list[str], row: list[str]) -> str:
    mapping = _row_mapping(header, row)
    if mapping:
        return "\n".join(f"{key}：{value}" for key, value in mapping if value)
    return _clean_text(" ".join(cell for cell in row if cell))


def _row_mapping(header: list[str], row: list[str]) -> list[tuple[str, str]]:
    pairs: list[tuple[str, str]] = []
    width = max(len(header), len(row))
    for index in range(width):
        key = _clean_text(header[index] if index < len(header) else f"字段{index + 1}")
        value = _clean_text(row[index] if index < len(row) else "")
        if key or value:
            pairs.append((key or f"字段{index + 1}", value))
    return pairs


def _first_mapping_value(mapping: list[tuple[str, str]], keys: tuple[str, ...]) -> str:
    for key, value in mapping:
        if value and any(keyword in key for keyword in keys):
            return value
    return ""


def _dedupe_non_empty(values: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for value in values:
        cleaned = _clean_text(value)
        if cleaned and cleaned not in seen:
            seen.add(cleaned)
            result.append(cleaned)
    return result


def _split_ext_fields(source: dict[str, Any], fragment: dict[str, Any]) -> dict[str, Any]:
    source_location = fragment.get("sourceLocation") or {}
    split_method = {
        "tableRow": "business_docx_order_table_row_v1",
        "localSemantic": "business_docx_local_semantic_v1",
        "aiSemantic": "business_docx_ai_semantic_v1",
    }.get(str(source_location.get("mode") or ""), "business_docx_boundary_v1")
    return {
        "splitParentMaterialId": source.get("id") or "",
        "splitParentFileName": source.get("name") or "",
        "splitParentFolderPath": source.get("folderPath") or "",
        "splitMethod": split_method,
        "splitFragmentId": fragment.get("id") or "",
        "splitFragmentTitle": fragment.get("title") or "",
        "splitMaterialType": fragment.get("materialType") or "",
        "splitSourceLocation": source_location,
        "splitConfidence": fragment.get("confidence"),
        "splitRiskTips": fragment.get("riskTips") or [],
        "reviewStatus": "pending_review",
    }


def _source_payload(source: dict[str, Any]) -> dict[str, Any]:
    return {key: value for key, value in source.items() if key != "content"}


def _clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", str(value or "").replace("\u3000", " ")).strip()


def _numeric_raw_id(file_id: str) -> int:
    try:
        return int(str(file_id or "").replace("RAW-", ""))
    except ValueError as exc:
        raise PeripheralError(400, "素材文件 ID 无效。", "RAW_FILE_ID_INVALID") from exc
