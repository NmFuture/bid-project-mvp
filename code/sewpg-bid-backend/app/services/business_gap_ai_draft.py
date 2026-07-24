from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document


def write_business_ai_draft_docx(
    output_path: Path,
    project: dict[str, Any],
    task: dict[str, Any],
    facts: dict[str, str],
    data: dict[str, Any],
) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    title = str(task.get("title") or "商务响应文件")
    normalized = title.replace("函格式", "函").replace("格式", "").strip() or "商务响应文件"
    doc = Document()
    doc.add_heading(normalized, level=1)

    tenderer = fact_value(facts, "招标人") or str(project.get("customerName") or project.get("owner") or "")
    bidder = fact_value(facts, "投标人") or str(data.get("bidderName") or "投标人")
    project_name = fact_value(facts, "项目名称") or str(project.get("name") or "")
    tender_no = fact_value(facts, "招标编号") or str(project.get("projectCode") or project.get("id") or "")
    date_value = fact_value(facts, "日期") or "    年    月    日"

    if tenderer:
        doc.add_paragraph(f"致：{tenderer}")
    for paragraph in business_ai_draft_body(normalized, project_name, tender_no):
        doc.add_paragraph(paragraph)

    doc.add_paragraph(f"项目名称：{project_name or '[待填写：项目名称]'}")
    doc.add_paragraph(f"招标编号：{tender_no or '[待填写：招标编号]'}")
    if "报价" in normalized or "价格" in normalized:
        doc.add_paragraph(f"投标报价：{fact_value(facts, '投标报价') or '[待填写：投标报价]'}")
        doc.add_paragraph(f"币种：{fact_value(facts, '币种') or '人民币'}")
    if "规格" in normalized or "供货" in normalized:
        doc.add_paragraph(f"投标机型：{fact_value(facts, '投标机型') or '[待填写：投标机型]'}")
        doc.add_paragraph(f"总装机容量：{fact_value(facts, '总装机容量') or '[待填写：总装机容量]'}")
    doc.add_paragraph("")
    doc.add_paragraph(f"投标人：{bidder}")
    doc.add_paragraph("法定代表人或授权代表：________________")
    doc.add_paragraph("日期：" + date_value)
    doc.add_paragraph("")
    doc.add_paragraph("注：本稿由商务标 S3 根据目录任务和项目事实表受控起草，签字、盖章、报价、金额及承诺范围须人工复核。")
    doc.save(output_path)


def business_ai_draft_body(title: str, project_name: str, tender_no: str) -> list[str]:
    project_label = project_name or "本项目"
    if "授权" in title:
        return [
            f"我单位现授权本单位工作人员作为 {project_label}（招标编号：{tender_no or '待填写'}）投标事宜的合法代理人。",
            "代理人在本项目投标、澄清、签署相关文件及处理投标事务中的行为，我单位均予以认可并承担相应法律责任。",
        ]
    if "廉洁" in title:
        return [
            f"我单位参加 {project_label} 投标活动，承诺严格遵守国家法律法规、招标文件及廉洁从业要求。",
            "我单位不以任何不正当方式影响评审结果，不向相关人员输送利益，并接受招标人及监管机构监督。",
        ]
    if "承诺" in title or "声明" in title or "说明" in title:
        return [
            f"我单位已充分理解 {project_label} 招标文件要求，并承诺按招标文件、投标文件及后续合同约定履行相关义务。",
            "如本承诺与招标文件专用条款或澄清文件存在不一致，以经人工复核后的最终投标文件为准。",
        ]
    if "履约" in title:
        return [
            f"我单位承诺在 {project_label} 中标后，按招标文件和合同约定及时提交履约保证金或履约保函。",
            "若未按要求提交，我单位愿承担招标文件及合同约定的相关责任。",
        ]
    if "投标函" in title:
        return [
            f"我单位已认真阅读 {project_label} 招标文件及其补遗、澄清文件，愿按招标文件要求提交商务投标文件。",
            "我单位承诺投标文件所载内容真实、准确、完整，并在投标有效期内保持投标文件有效。",
        ]
    return [
        f"我单位针对 {project_label} 编制本商务响应文件。",
        "本文件内容依据招标文件目录任务和项目事实表生成，需结合最终招标要求人工复核完善。",
    ]


def fact_value(facts: dict[str, str], label: str) -> str:
    key = re.sub(r"[\s　：:（）()【】\\[\\]]+", "", str(label or ""))
    if label in facts:
        return facts[label]
    if key in facts:
        return facts[key]
    for existing, value in facts.items():
        if key and (key in existing or existing in key):
            return value
    return ""
