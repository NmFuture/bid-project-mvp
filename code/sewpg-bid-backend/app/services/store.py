from __future__ import annotations

import base64
import binascii
import copy
import itertools
import json
import mimetypes
import re
import shutil
import asyncio
import threading
from contextlib import closing
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

import psycopg
from docx import Document
from openpyxl import load_workbook
from psycopg.rows import dict_row
from psycopg.types.json import Jsonb

from app.core.config import settings
from app.services.identity import build_project_identity, build_project_material_scope
from app.services.business_gap_planning import (
    _business_template_index,
    build_business_gap_material_picker_index,
    build_business_gap_plan_for_project,
    refresh_business_gap_artifact_urls,
    summarize_business_gap_plan,
)
from app.services.gap_planning import (
    TABLE_FILL_SKILL_NAME,
    WORD_FILL_SKILL_NAME,
    _artifact_onlyoffice_payload,
    _allowed_material_index,
    build_gap_plan_for_project,
    check_gap_integrity,
    normalize_gap_plan_fill_task_skills,
    _prepare_material_index_files,
    prepare_existing_gap_material_files,
    register_manual_gap_upload,
    register_existing_gap_material,
    run_ai_fill_for_gap,
    safe_filename,
    summarize_gap_plan,
)
from app.services.minio_client import minio_client
from app.services.parse_profiles import normalize_bid_type
from app.services.material_store import material_store
from app.services.workspace_artifacts import (
    business_workspace_dir,
    cleanup_parse_temp_workspace,
    promote_parse_artifacts_to_workspace,
    workspace_parse_dir,
)
from app.services.turbine_models import normalize_project_turbine_model, project_turbine_model


STAGE_SCHEME = "S0_S6"
MAX_PROJECT_STAGE = 6

STAGE_NAMES = {
    1: "模板与目录",
    2: "审核目录",
    3: "缺口处理",
    4: "生成标书",
    5: "共创",
    6: "导出",
}

STAGE_PROGRESS_NAMES = STAGE_NAMES

STAGE_PROGRESS_GROUPS = [
    {
        "id": 1,
        "name": "模板与目录",
        "routeStageId": 1,
        "isHuman": False,
    },
    {
        "id": 2,
        "name": "审核目录",
        "routeStageId": 2,
        "isHuman": True,
    },
    {
        "id": 3,
        "name": "缺口处理",
        "routeStageId": 3,
        "isHuman": True,
    },
    {
        "id": 4,
        "name": "生成标书",
        "routeStageId": 4,
        "isHuman": False,
    },
    {
        "id": 5,
        "name": "共创",
        "routeStageId": 5,
        "isHuman": True,
    },
    {
        "id": 6,
        "name": "导出",
        "routeStageId": 6,
        "isHuman": False,
    },
]

PROJECT_FACT_TABLE_SCHEMA_VERSION = "bid-project-fact-table-v1"
PROJECT_FACT_CONFIRMED_STATUSES = {"confirmed"}
FACT_TABLE_HEADER_WORDS = {
    "编号",
    "序号",
    "项目",
    "名称",
    "内容",
    "备注",
    "说明",
    "单位",
    "计量单位",
    "技术参数与规格",
    "主要项目",
    "投标机型1",
    "投标机型2",
    "保证值",
    "授权人签名",
}
COMMON_PROJECT_FACT_LABELS = {
    "项目名称",
    "招标编号",
    "招标人",
    "招标方",
    "客户名称",
    "投标方案",
    "投标机型",
    "机组类型",
    "机组台数",
    "总装机容量",
    "单机容量",
    "叶轮直径",
    "轮毂高度",
    "扫风面积",
    "比功率",
    "安全等级",
    "设计寿命",
    "空气密度",
    "湍流强度",
    "极端风速",
    "年平均风速",
    "风剪切",
    "保证发电量",
    "保证有效小时数",
    "功率曲线保证率",
    "全场可利用率",
    "单台可利用率",
    "主要部件更换率",
}
FACT_MATERIAL_SOURCE_PRIORITIES = {
    "project": 300,
    "customer": 200,
    "standard": 100,
}

LEGACY_STAGE_TO_CURRENT = {
    0: 1,
    1: 1,
    2: 1,
    3: 2,
    4: 3,
    5: 3,
    6: 3,
    7: 4,
    8: 4,
    9: 5,
    10: 6,
}

REVIEW_DECISION_LABELS = {
    "pending": "待审核",
    "participate": "参与投标",
    "abandon": "不参与",
}


def now_iso() -> str:
    return datetime.now(UTC).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def default_outline_nodes(project_name: str) -> list[dict[str, Any]]:
    return [
        {
            "id": "OL-1",
            "title": "项目概况",
            "children": [
                {"id": "OL-1-1", "title": "项目背景", "children": []},
                {"id": "OL-1-2", "title": "建设目标", "children": []},
            ],
        },
        {
            "id": "OL-2",
            "title": "技术方案",
            "children": [
                {"id": "OL-2-1", "title": f"{project_name}总体方案", "children": []},
            ],
        },
        {
            "id": "OL-3",
            "title": "实施与保障",
            "children": [],
        },
    ]


def build_directory_event(
    message: str,
    *,
    level: str = "info",
    step: str = "general",
    at: str | None = None,
) -> dict[str, Any]:
    return {
        "at": at or now_iso(),
        "level": level,
        "step": step,
        "message": message,
    }


def build_directory_opencode_output(
    *,
    status: str = "idle",
    session_id: str = "",
    provider_id: str = "",
    model_id: str = "",
    received_at: str = "",
    parts: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    return {
        "status": status,
        "sessionId": session_id,
        "providerId": provider_id,
        "modelId": model_id,
        "receivedAt": received_at,
        "parts": copy.deepcopy(parts or []),
    }


def build_parse_event(
    message: str,
    *,
    level: str = "info",
    step: str = "general",
    at: str | None = None,
) -> dict[str, Any]:
    return {
        "at": at or now_iso(),
        "level": level,
        "step": step,
        "message": message,
    }


def fill_task_label(project: dict[str, Any] | None = None) -> str:
    bid_type = str((project or {}).get("bidType") or "技术标")
    return f"调用{bid_type}正文拼装 skill"


def fill_document_label(project: dict[str, Any] | None = None) -> str:
    bid_type = str((project or {}).get("bidType") or "技术标")
    return f"{bid_type}正文"


def default_fill_tasks(project: dict[str, Any] | None = None) -> list[dict[str, Any]]:
    return [
        {"id": "task-1", "label": "准备 S2 目录、Wiki 与素材库", "status": "pending"},
        {"id": "task-2", "label": fill_task_label(project), "status": "pending"},
        {"id": "task-3", "label": "写入 Word 正文", "status": "pending"},
    ]


def default_fill_state(project: dict[str, Any] | None = None) -> dict[str, Any]:
    return {
        "status": "idle",
        "percentage": 0,
        "filledAt": "",
        "runDurationSec": 0,
        "runDuration": "",
        "summary": "尚未生成标书，请点击“生成标书”后继续。",
        "output": None,
        "sections": [],
        "opencodeOutput": build_directory_opencode_output(),
        "events": [],
        "tasks": default_fill_tasks(project),
    }


def _run_async_material_upload(**kwargs: Any) -> dict[str, Any]:
    awaitable = material_store.raw_upload(**kwargs)
    try:
        asyncio.get_running_loop()
    except RuntimeError:
        return asyncio.run(awaitable)
    result: dict[str, Any] = {}
    error: dict[str, BaseException] = {}

    def runner() -> None:
        try:
            result["value"] = asyncio.run(awaitable)
        except BaseException as exc:  # pragma: no cover
            error["value"] = exc

    thread = threading.Thread(target=runner, daemon=True)
    thread.start()
    thread.join()
    if error:
        raise error["value"]
    return result.get("value") if isinstance(result.get("value"), dict) else {}


class AppStore:
    def __init__(self, storage_backend: str | None = None) -> None:
        settings.ensure_dirs()
        self._storage_backend = (storage_backend or settings.project_store_backend or "postgres").strip().lower()
        self._projects: dict[str, dict[str, Any]] = {}
        self._ensure_db()
        self._load_projects()
        self._counter = itertools.count(self._next_project_number())

    @staticmethod
    def _parse_project_number(project_id: str) -> int:
        match = re.fullmatch(r"PRJ-(\d{4,})", project_id)
        if not match:
            return 0
        return int(match.group(1))

    @staticmethod
    def _normalize_stage_value(value: Any, *, scheme: str = "") -> int:
        try:
            raw_stage = int(value or 1)
        except (TypeError, ValueError):
            raw_stage = 1
        if str(scheme or "") == STAGE_SCHEME:
            return max(1, min(MAX_PROJECT_STAGE, raw_stage))
        return LEGACY_STAGE_TO_CURRENT.get(raw_stage, max(1, min(MAX_PROJECT_STAGE, raw_stage)))

    @staticmethod
    def _normalize_stage_request(value: Any) -> int:
        try:
            raw_stage = int(value or 1)
        except (TypeError, ValueError):
            raw_stage = 1
        if raw_stage > MAX_PROJECT_STAGE:
            return LEGACY_STAGE_TO_CURRENT.get(raw_stage, MAX_PROJECT_STAGE)
        return max(1, min(MAX_PROJECT_STAGE, raw_stage))

    def _next_project_number(self) -> int:
        max_id = max((self._parse_project_number(project_id) for project_id in self._projects), default=0)
        return max_id + 1

    @staticmethod
    def _postgres_dsn() -> str:
        return settings.database_url.replace("postgresql+asyncpg://", "postgresql://", 1).replace(
            "postgresql+psycopg://",
            "postgresql://",
            1,
        )

    @property
    def _uses_postgres(self) -> bool:
        return self._storage_backend == "postgres"

    def _connect(self) -> psycopg.Connection:
        return psycopg.connect(self._postgres_dsn(), row_factory=dict_row)

    def _ensure_db(self) -> None:
        if not self._uses_postgres:
            return
        with closing(self._connect()) as connection:
            connection.execute(
                """
                CREATE TABLE IF NOT EXISTS projects (
                    id VARCHAR(50) PRIMARY KEY,
                    payload JSONB NOT NULL,
                    updated_at TEXT NOT NULL
                )
                """
            )
            connection.commit()

    def _load_projects(self) -> None:
        if not self._uses_postgres:
            return
        self._ensure_db()
        with closing(self._connect()) as connection:
            rows = connection.execute("SELECT id, payload FROM projects").fetchall()
        self._projects = {
            str(row["id"]): self._normalize_project_identity(
                row["payload"] if isinstance(row["payload"], dict) else json.loads(str(row["payload"]))
            )
            for row in rows
        }

    def _load_project(self, project_id: str) -> dict[str, Any] | None:
        if not self._uses_postgres:
            return self._projects.get(project_id)
        self._ensure_db()
        with closing(self._connect()) as connection:
            row = connection.execute("SELECT id, payload FROM projects WHERE id = %s", (project_id,)).fetchone()
        if row is None:
            self._projects.pop(project_id, None)
            return None
        payload = row["payload"] if isinstance(row["payload"], dict) else json.loads(str(row["payload"]))
        project = self._normalize_project_identity(payload)
        self._projects[project_id] = project
        return project

    def _persist_project(self, project: dict[str, Any]) -> None:
        if not self._uses_postgres:
            return
        self._ensure_db()
        with closing(self._connect()) as connection:
            connection.execute(
                """
                INSERT INTO projects (id, payload, updated_at)
                VALUES (%s, %s, %s)
                ON CONFLICT(id) DO UPDATE SET
                    payload=excluded.payload,
                    updated_at=excluded.updated_at
                """,
                (project["id"], Jsonb(project), project["updatedAt"]),
            )
            connection.commit()

    def _delete_project_record(self, project_id: str) -> None:
        if not self._uses_postgres:
            return
        self._ensure_db()
        with closing(self._connect()) as connection:
            connection.execute("DELETE FROM projects WHERE id = %s", (project_id,))
            connection.commit()

    def reset_for_tests(self, *, clear_persistent: bool = False) -> None:
        self._projects = {}
        if clear_persistent and self._uses_postgres:
            self._ensure_db()
            with closing(self._connect()) as connection:
                connection.execute("DELETE FROM projects")
                connection.commit()
        self._counter = itertools.count(1)

    @staticmethod
    def _normalize_project_identity(project: dict[str, Any]) -> dict[str, Any]:
        project_id = str(project.get("id") or "")
        project["currentStage"] = AppStore._normalize_stage_value(
            project.get("currentStage"),
            scheme=str(project.get("stageScheme") or ""),
        )
        project["stageScheme"] = STAGE_SCHEME
        project["projectCode"] = str(project.get("projectCode") or project_id)
        project["startDate"] = str(project.get("startDate") or "")
        project["endDate"] = str(project.get("endDate") or project.get("deadline") or "")
        project["deadline"] = str(project.get("deadline") or project.get("endDate") or "")
        identity = build_project_identity(project)
        project["identity"] = identity
        project["customerId"] = identity.get("customerId") or ""
        project["customerCanonicalName"] = identity.get("customerCanonicalName") or ""
        project["materialCustomerId"] = identity.get("customerId") or ""
        project["materialCustomerName"] = identity.get("customerCanonicalName") or ""
        project["materialProjectId"] = identity.get("projectId") or ""
        project["materialProjectCode"] = identity.get("projectCode") or ""
        project["materialProjectName"] = identity.get("projectName") or ""
        project["materialProjectMode"] = identity.get("materialProjectMode") or project.get("materialProjectMode") or ""
        turbine = project_turbine_model(project)
        project["turbineModel"] = turbine
        project["selectedTurbineModel"] = copy.deepcopy(turbine)
        project["turbineModelLabel"] = str(turbine.get("model") or "")
        return project

    @staticmethod
    def _should_skip_technical_gap_stage(project: dict[str, Any]) -> bool:
        if normalize_bid_type(str(project.get("bidType") or "")) == "商务标":
            return False
        outline_state = project.get("outline_state") if isinstance(project.get("outline_state"), dict) else {}
        return str(outline_state.get("reviewStatus") or "") == "confirmed"

    @staticmethod
    def format_size(size_bytes: int) -> str:
        if size_bytes <= 0:
            return "0 MB"
        return f"{size_bytes / 1024 / 1024:.1f} MB"

    @staticmethod
    def _source_file_type(file_name: str) -> str:
        lowered = str(file_name or "").lower()
        if lowered.endswith(".pdf"):
            return "PDF"
        if lowered.endswith(".md"):
            return "MD"
        if lowered.endswith((".doc", ".docx")):
            return "DOCX"
        return "文件"

    def _require(self, project_id: str) -> dict[str, Any]:
        project = self._load_project(project_id) or self._projects.get(project_id)
        if not project:
            raise KeyError(project_id)
        self._ensure_runtime_states(project)
        return project

    def _ensure_runtime_states(self, project: dict[str, Any]) -> dict[str, Any]:
        if not isinstance(project.get("parse_result"), dict):
            project["parse_result"] = self._recover_parse_result(project)
        if not isinstance(project.get("parse_storage"), dict):
            project["parse_storage"] = self._recover_parse_storage(project)
        if not isinstance(project.get("parse_progress"), dict):
            status = str((project.get("parse_result") or {}).get("status") or "idle")
            project["parse_progress"] = {
                "status": "completed" if status == "completed" else "idle",
                "percentage": 100 if status == "completed" else 0,
                "summary": "已从项目工作区恢复解析结果。" if status == "completed" else "尚未触发招标文件解析。",
                "startedAt": "",
                "completedAt": str((project.get("parse_result") or {}).get("parsedAt") or ""),
                "events": [],
                "opencodeOutput": build_directory_opencode_output(),
            }
        if not isinstance(project.get("fill_state"), dict):
            project["fill_state"] = default_fill_state(project)
        if not isinstance(project.get("directory_state"), dict):
            project["directory_state"] = self._recover_directory_state(project)
        if not isinstance(project.get("outline_state"), dict):
            project["outline_state"] = self._recover_outline_state(project)
        if not isinstance(project.get("document_state"), dict):
            project["document_state"] = {
                "status": "ready",
                "documentId": f"DOC-{project['id']}",
                "sourceFileName": f"{str(project.get('name') or project['id'])}_正文.docx",
                "fileName": f"{str(project.get('name') or project['id'])}_正文.docx",
                "fileType": "docx",
                "version": 1,
                "lastSavedAt": "",
                "onlyoffice": {
                    "documentKey": f"{project['id']}-v1",
                    "title": f"{str(project.get('name') or project['id'])}_正文.docx",
                    "user": {"id": "user-1", "name": "当前用户"},
                },
                "fallback": {"content": ""},
            }
        project["fill_state"].setdefault("opencodeOutput", build_directory_opencode_output())
        project["fill_state"].setdefault("events", [])
        project["fill_state"].setdefault("tasks", default_fill_tasks(project))
        return project

    def _recover_parse_result(self, project: dict[str, Any]) -> dict[str, Any]:
        project_id = str(project.get("id") or "")
        bid_type = normalize_bid_type(str(project.get("bidType") or "技术标"))
        parse_dir = workspace_parse_dir(project_id, bid_type)
        candidates = [
            parse_dir / "parse-result.workspace.json",
            parse_dir / "s1_structured_result.json",
            business_workspace_dir(project_id) / "parse" / "parse-result.workspace.json",
            settings.documents_dir / project_id / "technical-workspace" / "parse" / "parse-result.workspace.json",
            settings.parsed_dir / project_id / "parse-result.workspace.json",
            settings.parsed_dir / project_id / "s1_structured_result.json",
        ]
        for path in candidates:
            payload = self._read_json_file(path)
            if not payload:
                continue
            if path.name == "s1_structured_result.json":
                payload = {
                    "status": "completed",
                    "parsedAt": now_iso(),
                    "sourceFiles": [],
                    "items": copy.deepcopy(payload.get("items") or []),
                    "structured": copy.deepcopy(payload.get("structured") or {}),
                    "summary": payload.get("summary") if isinstance(payload.get("summary"), dict) else {},
                }
            payload.setdefault("status", "completed")
            payload.setdefault("parsedAt", now_iso())
            payload.setdefault("sourceFiles", [])
            payload.setdefault("items", [])
            payload.setdefault("structured", {})
            payload.setdefault(
                "summary",
                {
                    "fileCount": len(payload.get("sourceFiles") or []),
                    "extractedCount": len(payload.get("items") or []),
                    "textLength": 0,
                    "textPreview": "",
                    "warnings": ["解析结果由项目工作区自动恢复。"],
                },
            )
            return payload
        return self._empty_parse_result(project)

    def _recover_parse_storage(self, project: dict[str, Any]) -> dict[str, Any]:
        project_id = str(project.get("id") or "")
        bid_type = normalize_bid_type(str(project.get("bidType") or "技术标"))
        parse_dir = workspace_parse_dir(project_id, bid_type)
        manifest = self._read_json_file(parse_dir / "manifest.json")
        structured_path = parse_dir / "s1_structured_result.json"
        return {
            "projectDir": str(parse_dir.parent) if parse_dir.exists() else "",
            "parseDir": str(parse_dir) if parse_dir.exists() else "",
            "combinedTextPath": str(parse_dir / "combined.txt") if (parse_dir / "combined.txt").exists() else "",
            "manifestPath": str(parse_dir / "manifest.json") if (parse_dir / "manifest.json").exists() else "",
            "structuredResultPath": str(structured_path) if structured_path.exists() else "",
            "skillManifestPath": str(parse_dir / "s1_parse_manifest.json") if (parse_dir / "s1_parse_manifest.json").exists() else "",
            "documents": list(manifest.get("documents") or []) if isinstance(manifest.get("documents"), list) else [],
        }

    @staticmethod
    def _empty_parse_result(project: dict[str, Any]) -> dict[str, Any]:
        return {
            "status": "idle",
            "parsedAt": "",
            "project": {
                "id": str(project.get("id") or ""),
                "files": copy.deepcopy(project.get("files") or []),
                "templateFiles": copy.deepcopy(project.get("templateFiles") or []),
                "startDate": str(project.get("startDate") or ""),
                "endDate": str(project.get("endDate") or project.get("deadline") or ""),
                "deadline": str(project.get("deadline") or project.get("endDate") or ""),
                "currentStage": project.get("currentStage") or 1,
            },
            "sourceFiles": [],
            "items": [],
            "structured": {},
            "summary": {
                "fileCount": 0,
                "extractedCount": 0,
                "textLength": 0,
                "textPreview": "",
                "warnings": [],
            },
        }

    @staticmethod
    def _read_json_file(path: Path) -> dict[str, Any]:
        try:
            if not path.exists() or not path.is_file():
                return {}
            payload = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            return {}
        return payload if isinstance(payload, dict) else {}

    def _recover_directory_state(self, project: dict[str, Any]) -> dict[str, Any]:
        project_id = str(project.get("id") or "")
        toc_path = self._recover_business_toc_path(project)
        generated_at = now_iso()
        if toc_path:
            return {
                "status": "completed",
                "percentage": 100,
                "summary": "已从商务标 S2 目录产物恢复目录状态。",
                "generatedAt": generated_at,
                "output": {
                    "fileName": f"{str(project.get('name') or project_id)}_目录.docx",
                    "fileType": "docx",
                    "chapterCount": 0,
                },
                "opencodeOutput": {
                    **build_directory_opencode_output(),
                    "tocJsonPath": str(toc_path),
                    "outputFile": str(toc_path),
                },
                "ruleEvidence": {},
                "events": [
                    build_directory_event(
                        "已从商务标 S2 目录产物恢复目录状态。",
                        level="info",
                        step="recovered",
                        at=generated_at,
                    )
                ],
                "tasks": [
                    {"id": "task-1", "label": "准备目录候选", "status": "done"},
                    {"id": "task-2", "label": "futurecode 语义审核", "status": "done"},
                    {"id": "task-3", "label": "保存审核目录", "status": "done"},
                ],
            }
        return {
            "status": "idle",
            "percentage": 0,
            "summary": "尚未生成目录。",
            "generatedAt": "",
            "output": None,
            "opencodeOutput": build_directory_opencode_output(),
            "ruleEvidence": {},
            "events": [],
            "tasks": [
                {"id": "task-1", "label": "准备目录候选", "status": "pending"},
                {"id": "task-2", "label": "futurecode 语义审核", "status": "pending"},
                {"id": "task-3", "label": "保存审核目录", "status": "pending"},
            ],
        }

    def _recover_outline_state(self, project: dict[str, Any]) -> dict[str, Any]:
        project_id = str(project.get("id") or "")
        toc_path = self._recover_business_toc_path(project)
        nodes: list[dict[str, Any]] = []
        generated_at = now_iso()
        if toc_path:
            try:
                toc = json.loads(toc_path.read_text(encoding="utf-8"))
                items = toc.get("items") if isinstance(toc, dict) else []
                if isinstance(items, list):
                    nodes = self._outline_nodes_from_toc_items(items)
                generated_at = str(toc.get("generatedAt") or toc.get("generated_at") or generated_at) if isinstance(toc, dict) else generated_at
            except Exception:
                nodes = []
        return {
            "outlineVersion": 1,
            "reviewStatus": "confirmed" if nodes else "draft",
            "generatedAt": generated_at if nodes else "",
            "summary": {"totalNodeCount": self._count_nodes(nodes)},
            "nodes": copy.deepcopy(nodes),
            "recoveredFrom": str(toc_path or ""),
        }

    @staticmethod
    def _recover_business_toc_path(project: dict[str, Any]) -> Path | None:
        project_id = str(project.get("id") or "")
        if not project_id:
            return None
        candidates: list[Path] = []
        directory_state = project.get("directory_state") if isinstance(project.get("directory_state"), dict) else {}
        opencode_output = directory_state.get("opencodeOutput") if isinstance(directory_state.get("opencodeOutput"), dict) else {}
        for value in (opencode_output.get("tocJsonPath"), opencode_output.get("outputFile")):
            if value:
                candidates.append(Path(str(value)).expanduser())
        candidates.extend(
            [
                settings.documents_dir / project_id / "business-workspace" / "s2_toc_workdir" / settings.s2_toc_output_file_name,
                settings.documents_dir / project_id / "business-workspace" / "s2_toc_workdir" / "toc.json",
                settings.documents_dir / project_id / "business-workspace" / "gaps" / settings.s2_toc_output_file_name,
                settings.documents_dir / project_id / "business-workspace" / "gaps" / "toc.json",
            ]
        )
        for candidate in candidates:
            if candidate.exists() and candidate.is_file() and candidate.suffix.lower() == ".json":
                return candidate
        return None

    def _summary(self, project: dict[str, Any]) -> dict[str, Any]:
        project = self._normalize_project_identity(project)
        if self._should_skip_technical_gap_stage(project) and int(project.get("currentStage") or 1) == 3:
            project["currentStage"] = 4
        identity = project.get("identity") or {}
        review_decision = str(project.get("reviewDecision") or "participate")
        if review_decision not in REVIEW_DECISION_LABELS:
            review_decision = "pending"
        stage_label = "审核终止" if review_decision == "abandon" else STAGE_NAMES[project["currentStage"]]
        return {
            "id": project["id"],
            "projectCode": project.get("projectCode") or project["id"],
            "name": project["name"],
            "customerName": project["customerName"],
            "customerId": identity.get("customerId") or "",
            "customerCanonicalName": identity.get("customerCanonicalName") or "",
            "customerAliases": copy.deepcopy(identity.get("customerAliases") or []),
            "materialCustomerId": project.get("materialCustomerId") or identity.get("customerId") or "",
            "materialCustomerName": project.get("materialCustomerName") or identity.get("customerCanonicalName") or "",
            "materialProjectId": project.get("materialProjectId") or identity.get("projectId") or "",
            "materialProjectCode": project.get("materialProjectCode") or identity.get("projectCode") or "",
            "materialProjectName": project.get("materialProjectName") or identity.get("projectName") or "",
            "materialProjectMode": project.get("materialProjectMode") or identity.get("materialProjectMode") or "",
            "turbineModel": copy.deepcopy(project_turbine_model(project)),
            "selectedTurbineModel": copy.deepcopy(project_turbine_model(project)),
            "turbineModelLabel": str(project_turbine_model(project).get("model") or ""),
            "owner": project["owner"],
            "manager": project["manager"],
            "startDate": project.get("startDate") or "",
            "endDate": project.get("endDate") or project.get("deadline") or "",
            "deadline": project.get("deadline") or project.get("endDate") or "",
            "bidType": project["bidType"],
            "currentStage": project["currentStage"],
            "stageLabel": stage_label,
            "reviewDecision": review_decision,
            "reviewDecisionLabel": REVIEW_DECISION_LABELS[review_decision],
            "reviewDecidedAt": project.get("reviewDecidedAt") or "",
            "updatedAt": project["updatedAt"],
            "identity": copy.deepcopy(identity),
        }

    def _detail(self, project: dict[str, Any]) -> dict[str, Any]:
        return {
            **self._summary(project),
            "files": copy.deepcopy(project["files"]),
            "templateFiles": copy.deepcopy(project["templateFiles"]),
            "templateFallback": copy.deepcopy(self._normalize_template_fallback(project)),
            "isKeyAccount": project["isKeyAccount"],
            "keyAccountId": project["keyAccountId"],
            "reviewComment": str(project.get("reviewComment") or ""),
        }

    @staticmethod
    def _normalize_template_fallback(project: dict[str, Any]) -> dict[str, Any]:
        raw = project.get("templateFallback")
        fallback = raw if isinstance(raw, dict) else {}
        enabled = fallback.get("enabled")
        if enabled is None:
            enabled = True
        return {
            "enabled": bool(enabled),
            "sourceId": str(fallback.get("sourceId") or "system-default"),
        }

    def list_projects(
        self,
        status: str = "",
        bid_type: str = "",
        date_range: str = "",
        page: int = 1,
        page_size: int = 12,
    ) -> dict[str, Any]:
        self._load_projects()
        items = [self._summary(project) for project in self._projects.values()]
        normalized_bid_type = str(bid_type or "").strip()
        if normalized_bid_type:
            items = [item for item in items if str(item.get("bidType") or "").strip() == normalized_bid_type]
        items.sort(key=lambda item: item["updatedAt"], reverse=True)
        start = max(0, (page - 1) * page_size)
        end = start + page_size
        return {
            "items": items[start:end],
            "total": len(items),
            "page": page,
            "pageSize": page_size,
        }

    def create_project(self, data: dict[str, Any]) -> dict[str, Any]:
        project_id = f"PRJ-{next(self._counter):04d}"
        review_decision = str(data.get("reviewDecision") or "pending").strip().lower()
        if review_decision not in REVIEW_DECISION_LABELS:
            review_decision = "pending"
        project = {
            "id": project_id,
            "projectCode": str(data.get("projectCode") or project_id),
            "name": str(data.get("name") or project_id),
            "customerName": str(data.get("customerName") or ""),
            "customerId": str(data.get("customerId") or ""),
            "customerCanonicalName": str(data.get("customerCanonicalName") or ""),
            "materialCustomerId": str(data.get("materialCustomerId") or data.get("customerId") or ""),
            "materialCustomerName": str(data.get("materialCustomerName") or data.get("customerCanonicalName") or data.get("customerName") or ""),
            "materialProjectMode": str(data.get("materialProjectMode") or ""),
            "materialProjectId": str(data.get("materialProjectId") or ""),
            "materialProjectCode": str(data.get("materialProjectCode") or ""),
            "materialProjectName": str(data.get("materialProjectName") or data.get("name") or ""),
            "owner": str(data.get("owner") or data.get("customerName") or ""),
            "manager": str(data.get("manager") or ""),
            "startDate": str(data.get("startDate") or ""),
            "endDate": str(data.get("endDate") or data.get("deadline") or ""),
            "deadline": str(data.get("deadline") or data.get("endDate") or ""),
            "bidType": str(data.get("bidType") or "技术标"),
            "turbineModel": normalize_project_turbine_model(
                data.get("turbineModel") or data.get("selectedTurbineModel") or data.get("machineModel")
            ),
            "isKeyAccount": bool(data.get("isKeyAccount")),
            "keyAccountId": str(data.get("keyAccountId") or ""),
            "reviewDecision": review_decision,
            "reviewDecidedAt": now_iso() if review_decision in {"participate", "abandon"} else "",
            "reviewComment": str(data.get("reviewComment") or ""),
            "files": [],
            "templateFiles": [],
            "fileRecords": [],
            "templateFileRecords": [],
            "templateFallback": {
                "enabled": True,
                "sourceId": "system-default",
            },
            "stageScheme": STAGE_SCHEME,
            "currentStage": 1,
            "updatedAt": now_iso(),
            "parse_result": {
                "status": "idle",
                "parsedAt": "",
                "sourceFiles": [],
                "items": [],
                "summary": {
                    "fileCount": 0,
                    "extractedCount": 0,
                    "textLength": 0,
                    "textPreview": "",
                    "warnings": [],
                },
            },
            "parse_storage": {
                "projectDir": "",
                "combinedTextPath": "",
                "manifestPath": "",
                "documents": [],
            },
            "parse_progress": {
                "status": "idle",
                "percentage": 0,
                "summary": "尚未触发招标文件解析。",
                "startedAt": "",
                "completedAt": "",
                "events": [],
                "opencodeOutput": build_directory_opencode_output(),
            },
            "directory_state": {
                "status": "idle",
                "percentage": 0,
                "summary": "尚未生成目录。",
                "generatedAt": "",
                "output": None,
                "opencodeOutput": build_directory_opencode_output(),
                "events": [],
                "tasks": [
                    {"id": "task-1", "label": "准备目录候选", "status": "pending"},
                    {"id": "task-2", "label": "futurecode 语义审核", "status": "pending"},
                    {"id": "task-3", "label": "保存审核目录", "status": "pending"},
                ],
            },
            "outline_state": {
                "outlineVersion": 1,
                "reviewStatus": "draft",
                "generatedAt": "",
                "summary": {"totalNodeCount": 0},
                "nodes": [],
            },
            "fill_state": default_fill_state({"bidType": str(data.get("bidType") or "技术标")}),
            "document_state": {
                "status": "ready",
                "documentId": f"DOC-{project_id}",
                "sourceFileName": f"{str(data.get('name') or project_id)}_正文.docx",
                "fileName": f"{str(data.get('name') or project_id)}_正文.docx",
                "fileType": "docx",
                "version": 1,
                "lastSavedAt": "",
                "onlyoffice": {
                    "documentKey": f"{project_id}-v1",
                    "title": f"{str(data.get('name') or project_id)}_正文.docx",
                    "user": {"id": "user-1", "name": "当前用户"},
                },
                "fallback": {
                    "content": "OnlyOffice 未接入前，这里先展示后端生成的占位文档内容。",
                },
            },
            "gap_state": {
                "recognitionStatus": "idle",
                "recognizedAt": "",
                "submittedForReview": False,
                "reviewConfirmed": False,
                "reviewedAt": "",
                "items": [],
                "submissions": [],
                "plan": {},
                "planFile": "",
                "integrity": {},
                "projectFactTable": {},
            },
            "business_gap_state": {
                "recognitionStatus": "idle",
                "recognizedAt": "",
                "submittedForReview": False,
                "reviewConfirmed": False,
                "reviewedAt": "",
                "plan": {},
                "planFile": "",
                "integrity": {},
            },
            "review_document_state": {
                "parseStatus": "idle",
                "parsedAt": "",
                "documentId": f"REV-DOC-{project_id}",
                "sourceFileName": "",
                "fileName": f"{str(data.get('name') or project_id)}_缺口处理确认预览.docx",
                "fileType": "docx",
                "content": "",
                "lastSavedAt": "",
                "version": 1,
                "documentKey": f"{project_id}-s6-v1",
            },
        }
        self._normalize_project_identity(project)
        self._projects[project_id] = project
        self._persist_project(project)
        return self._detail(project)

    def get_project(self, project_id: str) -> dict[str, Any]:
        return self._detail(self._require(project_id))

    def get_project_runtime_state(self, project_id: str) -> dict[str, Any]:
        """Return the full persisted project payload for backend pipeline services."""
        return copy.deepcopy(self._require(project_id))

    def update_project(self, project_id: str, data: dict[str, Any]) -> dict[str, Any]:
        project = self._require(project_id)
        for field in [
            "name",
            "projectCode",
            "customerName",
            "customerId",
            "customerCanonicalName",
            "materialCustomerId",
            "materialCustomerName",
            "materialProjectMode",
            "materialProjectId",
            "materialProjectCode",
            "materialProjectName",
            "owner",
            "manager",
            "startDate",
            "endDate",
            "deadline",
            "bidType",
        ]:
            if field in data:
                project[field] = str(data[field] or "") if field == "projectCode" else data[field]
        if any(field in data for field in ("turbineModel", "selectedTurbineModel", "machineModel")):
            project["turbineModel"] = normalize_project_turbine_model(
                data.get("turbineModel") or data.get("selectedTurbineModel") or data.get("machineModel")
            )
        if "endDate" in data and "deadline" not in data:
            project["deadline"] = str(data.get("endDate") or "")
        if "deadline" in data and "endDate" not in data:
            project["endDate"] = str(data.get("deadline") or "")
        if "reviewDecision" in data:
            decision = str(data.get("reviewDecision") or "").strip().lower()
            if decision not in REVIEW_DECISION_LABELS:
                decision = "pending"
            project["reviewDecision"] = decision
            project["reviewDecidedAt"] = now_iso() if decision in {"participate", "abandon"} else ""
            if decision == "participate":
                parse_result = project.get("parse_result") if isinstance(project.get("parse_result"), dict) else {}
                parse_storage = project.get("parse_storage") if isinstance(project.get("parse_storage"), dict) else {}
                if parse_result.get("status") == "completed":
                    promoted = promote_parse_artifacts_to_workspace(
                        project_id,
                        parse_result,
                        parse_storage,
                        bid_type=normalize_bid_type(str(project.get("bidType") or "技术标")),
                    )
                    project["parse_result"] = promoted["parseResult"]
                    project["parse_storage"] = promoted["parseStorage"]
                    project["workspaceArtifacts"] = promoted["artifacts"]
                    cleanup_parse_temp_workspace(project_id)
            elif decision == "abandon":
                cleanup_parse_temp_workspace(project_id)
        if "reviewComment" in data:
            project["reviewComment"] = str(data.get("reviewComment") or "")
        self._normalize_project_identity(project)
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        return self._detail(project)

    def delete_project(self, project_id: str) -> None:
        self._require(project_id)
        cleanup_parse_temp_workspace(project_id)
        self._projects.pop(project_id, None)
        self._delete_project_record(project_id)

    def get_stages(self, project_id: str) -> list[dict[str, Any]]:
        project = self._require(project_id)
        current_stage = self._normalize_stage_value(project.get("currentStage"), scheme=STAGE_SCHEME)
        skip_technical_gap_stage = self._should_skip_technical_gap_stage(project)
        if skip_technical_gap_stage and current_stage == 3:
            current_stage = 4
        stages: list[dict[str, Any]] = []
        for group in STAGE_PROGRESS_GROUPS:
            stage_id = int(group["id"])
            if skip_technical_gap_stage and stage_id == 3:
                status = "completed"
            elif stage_id < current_stage:
                status = "completed"
            elif stage_id == current_stage:
                status = "active"
            else:
                status = "pending"
            stages.append(
                {
                    "id": group["id"],
                    "name": group["name"],
                    "status": status,
                    "isHuman": bool(group["isHuman"]),
                    "routeStageId": group["routeStageId"],
                    "isSkipped": bool(skip_technical_gap_stage and stage_id == 3),
                }
            )
        return stages

    def update_stage(self, project_id: str, stage: int, data: dict[str, Any]) -> dict[str, Any]:
        project = self._require(project_id)
        status = str(data.get("status") or "").strip()
        normalized_stage = self._normalize_stage_request(stage)
        current_stage = self._normalize_stage_value(project.get("currentStage"), scheme=STAGE_SCHEME)
        if status == "completed":
            next_stage = normalized_stage + 1
            if normalized_stage == 2 and self._should_skip_technical_gap_stage(project):
                next_stage = 4
            project["currentStage"] = min(MAX_PROJECT_STAGE, max(current_stage, next_stage))
        elif status == "active":
            if normalized_stage == 3 and self._should_skip_technical_gap_stage(project):
                normalized_stage = 4
            project["currentStage"] = max(1, min(MAX_PROJECT_STAGE, normalized_stage))
        project["stageScheme"] = STAGE_SCHEME
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        return {
            "message": "阶段状态已更新",
            "currentStage": project["currentStage"],
            "stageLabel": STAGE_NAMES[project["currentStage"]],
        }

    def complete_parse(
        self,
        project_id: str,
        tender_files: list[dict[str, Any]],
        template_files: list[dict[str, Any]],
        summary: dict[str, Any] | None = None,
        parse_storage: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        project = self._require(project_id)
        parsed_at = now_iso()
        project_updates = parse_storage.get("projectUpdates") if isinstance(parse_storage, dict) else {}
        if isinstance(project_updates, dict):
            for field in ["startDate", "endDate", "deadline"]:
                value = str(project_updates.get(field) or "").strip()
                if value and not str(project.get(field) or "").strip():
                    project[field] = value
        project["files"] = [item["name"] for item in tender_files]
        project["fileRecords"] = copy.deepcopy(tender_files)
        project["templateFileRecords"] = copy.deepcopy(template_files)
        project["templateFiles"] = [
            {
                "id": item["id"],
                "name": item["name"],
                "sizeLabel": item["size_label"],
            }
            for item in template_files
        ]
        source_files = [
            {
                "id": item["id"].replace("TEN", "SRC"),
                "name": item["name"],
                "type": self._source_file_type(item["name"]),
                "pageCount": 12,
                "size": item["size_label"],
            }
            for item in tender_files
        ]
        items: list[dict[str, Any]] = []
        structured: dict[str, Any] = {}
        if isinstance(parse_storage, dict):
            raw_items = parse_storage.get("items")
            raw_structured = parse_storage.get("structured")
            if isinstance(raw_items, list):
                items = copy.deepcopy(raw_items)
            if isinstance(raw_structured, dict):
                structured = copy.deepcopy(raw_structured)
        source_file_lookup = {item["name"]: item for item in source_files}
        if summary and parse_storage:
            for document in parse_storage.get("documents", []):
                if source_file_lookup.get(document["name"]):
                    source_file_lookup[document["name"]]["pageCount"] = document.get("pageCount", "-")
                    source_file_lookup[document["name"]]["textLength"] = document.get("textLength", 0)
        project["parse_result"] = {
            "status": "completed",
            "parsedAt": parsed_at,
            "project": {
                "id": project["id"],
                "files": copy.deepcopy(project["files"]),
                "templateFiles": copy.deepcopy(project["templateFiles"]),
                "startDate": project.get("startDate") or "",
                "endDate": project.get("endDate") or project.get("deadline") or "",
                "deadline": project.get("deadline") or project.get("endDate") or "",
                "currentStage": project["currentStage"],
                "stageLabel": STAGE_NAMES[project["currentStage"]],
            },
            "sourceFiles": source_files,
            "items": items,
            "structured": structured,
            "summary": summary or {
                "fileCount": len(source_files),
                "extractedCount": len(items),
                "textLength": 0,
                "textPreview": "",
                "warnings": [],
            },
        }
        project["parse_storage"] = copy.deepcopy(parse_storage or project.get("parse_storage") or {})
        project["updatedAt"] = parsed_at
        self._persist_project(project)
        return copy.deepcopy(project["parse_result"])

    def get_parse_result(self, project_id: str) -> dict[str, Any]:
        return copy.deepcopy(self._require(project_id)["parse_result"])

    def update_parse_result(
        self,
        project_id: str,
        parse_result: dict[str, Any],
        *,
        parse_storage: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        project = self._require(project_id)
        project["parse_result"] = copy.deepcopy(parse_result)
        if parse_storage is not None:
            project["parse_storage"] = copy.deepcopy(parse_storage)
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        return copy.deepcopy(project["parse_result"])

    def get_parse_progress(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        progress = project.get("parse_progress")
        if not isinstance(progress, dict):
            progress = {
                "status": "idle",
                "percentage": 0,
                "summary": "尚未触发招标文件解析。",
                "startedAt": "",
                "completedAt": "",
                "events": [],
                "opencodeOutput": build_directory_opencode_output(),
            }
            project["parse_progress"] = progress
            self._persist_project(project)
        return copy.deepcopy(progress)

    def start_parse_progress(self, project_id: str, message: str = "已开始招标文件解析。") -> dict[str, Any]:
        project = self._require(project_id)
        started_at = now_iso()
        progress = {
            "status": "running",
            "percentage": 5,
            "summary": message,
            "startedAt": started_at,
            "completedAt": "",
            "events": [build_parse_event(message, step="start", at=started_at)],
            "opencodeOutput": build_directory_opencode_output(status="idle"),
        }
        project["parse_progress"] = progress
        project["updatedAt"] = started_at
        self._persist_project(project)
        return copy.deepcopy(progress)

    def update_parse_progress(
        self,
        project_id: str,
        *,
        status: str | None = None,
        percentage: int | None = None,
        summary: str | None = None,
        event_message: str = "",
        event_step: str = "general",
        event_level: str = "info",
        opencode_output: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        project = self._require(project_id)
        progress = project.get("parse_progress") if isinstance(project.get("parse_progress"), dict) else {}
        if not progress:
            progress = self.start_parse_progress(project_id)
            project = self._require(project_id)
            progress = project["parse_progress"]
        if status:
            progress["status"] = status
        if percentage is not None:
            progress["percentage"] = max(0, min(100, int(percentage)))
        if summary is not None:
            progress["summary"] = summary
        if opencode_output:
            progress["opencodeOutput"] = {
                **build_directory_opencode_output(),
                **copy.deepcopy(opencode_output),
            }
        if event_message:
            events = progress.setdefault("events", [])
            events.append(build_parse_event(event_message, level=event_level, step=event_step))
            progress["events"] = events[-80:]
        if status == "completed":
            progress["completedAt"] = now_iso()
            progress["percentage"] = 100
        project["parse_progress"] = progress
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        return copy.deepcopy(progress)

    def get_parse_storage(self, project_id: str) -> dict[str, Any]:
        return copy.deepcopy(self._require(project_id)["parse_storage"])

    def get_parse_inputs(
        self,
        project_id: str,
        *,
        include_fallback: bool = True,
    ) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
        project = self._require(project_id)
        template_file_records = copy.deepcopy(project.get("templateFileRecords") or [])
        if include_fallback and not template_file_records and self._normalize_template_fallback(project)["enabled"]:
            from app.services import template_store as template_store_module

            fallback_record = template_store_module.resolve_fallback_bid_template_file_sync(
                project_id,
                str(project.get("bidType") or "技术标"),
            )
            if fallback_record is not None:
                template_file_records = [fallback_record]
        return (
            copy.deepcopy(project.get("fileRecords") or []),
            template_file_records,
        )

    def get_template_fallback(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        from app.services.template_store import template_fallback_payload

        current = self._normalize_template_fallback(project)
        import asyncio

        return asyncio.run(
            template_fallback_payload(
                project_id=project_id,
                bid_type=str(project.get("bidType") or "技术标"),
                enabled=bool(current["enabled"]),
                source_id=str(current["sourceId"]),
                has_project_template=bool(project.get("templateFileRecords")),
            )
        )

    def template_fallback_context(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        current = self._normalize_template_fallback(project)
        return {
            "projectId": project_id,
            "bidType": str(project.get("bidType") or "技术标"),
            "enabled": bool(current["enabled"]),
            "sourceId": str(current["sourceId"]),
            "hasProjectTemplate": bool(project.get("templateFileRecords")),
        }

    def update_template_fallback(self, project_id: str, data: dict[str, Any]) -> dict[str, Any]:
        project = self._require(project_id)
        current = self._normalize_template_fallback(project)
        if "enabled" in data:
            current["enabled"] = bool(data.get("enabled"))
        if "sourceId" in data:
            current["sourceId"] = str(data.get("sourceId") or "system-default")
        project["templateFallback"] = current
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        return self.template_fallback_context(project_id)

    def update_template_files(self, project_id: str, template_files: list[dict[str, Any]]) -> dict[str, Any]:
        project = self._require(project_id)
        project["templateFileRecords"] = copy.deepcopy(template_files)
        project["templateFiles"] = [
            {
                "id": item["id"],
                "name": item["name"],
                "sizeLabel": item["size_label"],
            }
            for item in template_files
        ]
        parse_result = project.get("parse_result") or {}
        if isinstance(parse_result, dict):
            parse_project = parse_result.get("project") or {}
            parse_project["id"] = project["id"]
            parse_project["templateFiles"] = copy.deepcopy(project["templateFiles"])
            parse_result["project"] = parse_project
            project["parse_result"] = parse_result
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        return {
            "project": {
                "id": project["id"],
                "templateFiles": copy.deepcopy(project["templateFiles"]),
                "currentStage": project["currentStage"],
                "stageLabel": STAGE_NAMES[project["currentStage"]],
            },
            "templateFiles": copy.deepcopy(project["templateFiles"]),
            "updatedAt": project["updatedAt"],
        }

    def complete_directory_generation(self, project_id: str, data: dict[str, Any]) -> dict[str, Any]:
        project = self._require(project_id)
        generated_at = now_iso()
        nodes = default_outline_nodes(project["name"])
        payload = {
            "status": "completed",
            "percentage": 100,
            "summary": "目录生成完成。",
            "generatedAt": generated_at,
            "output": {
                "fileName": f"{project['name']}_目录.docx",
                "fileType": "docx",
                "chapterCount": len(nodes),
            },
            "opencodeOutput": build_directory_opencode_output(),
            "ruleEvidence": {},
            "events": [
                build_directory_event("目录生成完成。", level="success", step="done", at=generated_at),
            ],
            "tasks": [
                {"id": "task-1", "label": "准备目录候选", "status": "done"},
                {"id": "task-2", "label": "futurecode 语义审核", "status": "done"},
                {"id": "task-3", "label": "保存审核目录", "status": "done"},
            ],
        }
        project["directory_state"] = payload
        project["outline_state"] = {
            "outlineVersion": 1,
            "reviewStatus": "draft",
            "generatedAt": generated_at,
            "summary": {"totalNodeCount": self._count_nodes(nodes)},
            "nodes": nodes,
        }
        project["updatedAt"] = generated_at
        self._persist_project(project)
        return copy.deepcopy(payload)

    def start_directory_generation(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        payload = {
            "status": "running",
            "percentage": 5,
            "summary": "已开始生成目录，正在准备招标文件与投标模板候选。",
            "generatedAt": "",
            "output": None,
            "opencodeOutput": build_directory_opencode_output(),
            "ruleEvidence": {},
            "events": [
                build_directory_event(
                    "已开始生成目录任务，正在准备招标文件与投标模板候选。",
                    step="bootstrap",
                ),
            ],
            "tasks": [
                {"id": "task-1", "label": "准备目录候选", "status": "running"},
                {"id": "task-2", "label": "futurecode 语义审核", "status": "pending"},
                {"id": "task-3", "label": "保存审核目录", "status": "pending"},
            ],
        }
        project["directory_state"] = payload
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        return copy.deepcopy(payload)

    def update_directory_generation_state(
        self,
        project_id: str,
        *,
        percentage: int | None = None,
        summary: str | None = None,
        tasks: list[dict[str, Any]] | None = None,
        status: str | None = None,
        event_message: str | None = None,
        event_level: str = "info",
        event_step: str = "general",
        opencode_output: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        project = self._require(project_id)
        current = copy.deepcopy(project.get("directory_state") or self._recover_directory_state(project))
        if percentage is not None:
            current["percentage"] = max(0, min(100, int(percentage)))
        if summary is not None:
            current["summary"] = summary
        if tasks is not None:
            current["tasks"] = copy.deepcopy(tasks)
        if status is not None:
            current["status"] = status
        if opencode_output is not None:
            merged_output = build_directory_opencode_output()
            merged_output.update(copy.deepcopy(current.get("opencodeOutput") or {}))
            merged_output.update(copy.deepcopy(opencode_output))
            merged_output["parts"] = copy.deepcopy(merged_output.get("parts") or [])[-20:]
            current["opencodeOutput"] = merged_output
        if event_message:
            events = list(current.get("events") or [])
            events.append(
                build_directory_event(
                    event_message,
                    level=event_level,
                    step=event_step,
                )
            )
            current["events"] = events[-20:]
        project["directory_state"] = current
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        return copy.deepcopy(current)

    def fail_directory_generation(
        self,
        project_id: str,
        message: str,
        tasks: list[dict[str, Any]] | None = None,
    ) -> dict[str, Any]:
        project = self._require(project_id)
        current = copy.deepcopy(project.get("directory_state") or self._recover_directory_state(project))
        current["status"] = "failed"
        current["summary"] = message
        current_output = build_directory_opencode_output()
        current_output.update(copy.deepcopy(current.get("opencodeOutput") or {}))
        if (
            current_output.get("status") != "idle"
            or current_output.get("sessionId")
            or current_output.get("providerId")
            or current_output.get("modelId")
            or current_output.get("parts")
        ):
            current_output["status"] = "failed"
        current["opencodeOutput"] = current_output
        if tasks is not None:
            current["tasks"] = copy.deepcopy(tasks)
        events = list(current.get("events") or [])
        events.append(build_directory_event(message, level="error", step="failed"))
        current["events"] = events[-20:]
        project["directory_state"] = current
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        return copy.deepcopy(current)

    def save_generated_outline(
        self,
        project_id: str,
        nodes: list[dict[str, Any]],
        generated_at: str,
        summary: str,
        opencode_output: dict[str, Any] | None = None,
        rule_evidence: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        project = self._require(project_id)
        current_state = copy.deepcopy(project.get("directory_state") or {})
        current_events = list(copy.deepcopy(current_state.get("events") or []))
        current_output = build_directory_opencode_output()
        current_output.update(copy.deepcopy(current_state.get("opencodeOutput") or {}))
        if opencode_output is not None:
            current_output.update(copy.deepcopy(opencode_output))
        current_output["parts"] = copy.deepcopy(current_output.get("parts") or [])[-20:]
        current_events.append(
            build_directory_event(
                f"目录生成完成，已输出 {len(nodes)} 个一级章节。",
                level="success",
                step="done",
                at=generated_at,
            )
        )
        payload = {
            "status": "completed",
            "percentage": 100,
            "summary": summary,
            "generatedAt": generated_at,
            "output": {
                "fileName": f"{project['name']}_目录.docx",
                "fileType": "docx",
                "chapterCount": len(nodes),
            },
            "opencodeOutput": current_output,
            "ruleEvidence": copy.deepcopy(rule_evidence or current_state.get("ruleEvidence") or {}),
            "events": current_events[-20:],
            "tasks": [
                {"id": "task-1", "label": "准备目录候选", "status": "done"},
                {"id": "task-2", "label": "futurecode 语义审核", "status": "done"},
                {"id": "task-3", "label": "保存审核目录", "status": "done"},
            ],
        }
        project["directory_state"] = payload
        current_outline = project.get("outline_state") if isinstance(project.get("outline_state"), dict) else {}
        project["outline_state"] = {
            "outlineVersion": int(current_outline.get("outlineVersion") or 0) + 1,
            "reviewStatus": "draft",
            "generatedAt": generated_at,
            "summary": {"totalNodeCount": self._count_nodes(nodes)},
            "nodes": copy.deepcopy(nodes),
        }
        project["updatedAt"] = generated_at
        self._persist_project(project)
        return copy.deepcopy(payload)

    def get_directory_state(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        state = copy.deepcopy(project["directory_state"])
        if not isinstance(state.get("ruleEvidence"), dict) or not state.get("ruleEvidence"):
            evidence = self._load_directory_rule_evidence(state)
            if evidence:
                state["ruleEvidence"] = evidence
        return state

    def _load_directory_rule_evidence(self, state: dict[str, Any]) -> dict[str, Any]:
        opencode_output = state.get("opencodeOutput") if isinstance(state.get("opencodeOutput"), dict) else {}
        evidence_path = Path(str(opencode_output.get("evidencePath") or "")).expanduser()
        if not str(evidence_path).strip() or not evidence_path.exists():
            return {}
        try:
            evidence = json.loads(evidence_path.read_text(encoding="utf-8"))
        except Exception:
            return {}
        if not isinstance(evidence, dict):
            return {}
        decisions = evidence.get("decisions") if isinstance(evidence.get("decisions"), list) else []
        candidates = evidence.get("tenderCandidates") if isinstance(evidence.get("tenderCandidates"), list) else []
        template_outline = evidence.get("templateOutline") if isinstance(evidence.get("templateOutline"), list) else []
        decision_limit = 80
        return {
            "schemaVersion": str(evidence.get("schema_version") or ""),
            "engine": str(evidence.get("engine") or ""),
            "ruleConfig": copy.deepcopy(evidence.get("ruleConfig") if isinstance(evidence.get("ruleConfig"), dict) else {}),
            "templateOutlineCount": len(template_outline),
            "tenderCandidateCount": len(candidates),
            "decisionCount": len(decisions),
            "decisions": [
                copy.deepcopy(item)
                for item in decisions
                if isinstance(item, dict)
            ][:decision_limit],
        }

    def get_outline_state(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        return copy.deepcopy(project["outline_state"])

    def save_outline(self, project_id: str, nodes: list[dict[str, Any]]) -> dict[str, Any]:
        project = self._require(project_id)
        current = copy.deepcopy(project.get("outline_state") or self._recover_outline_state(project))
        current["outlineVersion"] = int(current.get("outlineVersion") or 1) + 1
        current["reviewStatus"] = "draft"
        current["summary"] = {"totalNodeCount": self._count_nodes(nodes)}
        current["nodes"] = copy.deepcopy(nodes)
        project["outline_state"] = current
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        return copy.deepcopy(current)

    def regenerate_outline(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        nodes = self._regenerated_outline_nodes(project)
        current_outline = project.get("outline_state") if isinstance(project.get("outline_state"), dict) else {}
        project["outline_state"] = {
            "outlineVersion": int(current_outline.get("outlineVersion") or 1) + 1,
            "reviewStatus": "draft",
            "generatedAt": now_iso(),
            "summary": {"totalNodeCount": self._count_nodes(nodes)},
            "nodes": nodes,
        }
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        return copy.deepcopy(project["outline_state"])

    def _regenerated_outline_nodes(self, project: dict[str, Any]) -> list[dict[str, Any]]:
        if str(project.get("bidType") or "").strip() != "商务标":
            return default_outline_nodes(project["name"])
        nodes = self._outline_nodes_from_directory_toc(project)
        return nodes or copy.deepcopy(project.get("outline_state", {}).get("nodes") or [])

    def _outline_nodes_from_directory_toc(self, project: dict[str, Any]) -> list[dict[str, Any]]:
        directory_state = project.get("directory_state") if isinstance(project.get("directory_state"), dict) else {}
        opencode_output = directory_state.get("opencodeOutput") if isinstance(directory_state.get("opencodeOutput"), dict) else {}
        toc_path = Path(str(opencode_output.get("tocJsonPath") or "")).expanduser()
        if not str(toc_path).strip() or not toc_path.exists():
            return []
        try:
            toc = json.loads(toc_path.read_text(encoding="utf-8"))
        except Exception:
            return []
        items = toc.get("items") if isinstance(toc, dict) else None
        if not isinstance(items, list):
            return []
        return self._outline_nodes_from_toc_items(items)

    @staticmethod
    def _outline_nodes_from_toc_items(items: list[Any]) -> list[dict[str, Any]]:
        roots: list[dict[str, Any]] = []
        stack: list[tuple[int, dict[str, Any]]] = []
        counters: list[int] = []
        for fallback_order, item in enumerate(items, start=1):
            if not isinstance(item, dict):
                continue
            try:
                level = max(1, int(item.get("level") or 1))
            except (TypeError, ValueError):
                level = 1
            if not stack and level > 1:
                level = 1
            elif stack and level > stack[-1][0] + 1:
                level = stack[-1][0] + 1
            while stack and stack[-1][0] >= level:
                stack.pop()
            counters = counters[:level]
            if len(counters) < level:
                counters.extend([0] * (level - len(counters)))
            counters[level - 1] += 1
            title = str(item.get("title") or item.get("name") or f"未命名章节{fallback_order}").strip()
            node = {
                "id": "OL-" + "-".join(str(part) for part in counters[:level] if part),
                "title": title,
                "children": [],
                "tocNumber": str(item.get("number") or "").strip(),
                "annotation": str(item.get("annotation") or "").strip(),
                "required_status": str(item.get("required_status") or item.get("requiredStatus") or "").strip(),
                "requiredStatus": str(item.get("requiredStatus") or item.get("required_status") or "").strip(),
                "source_text": str(item.get("source_text") or item.get("sourceText") or "").strip(),
                "sourceText": str(item.get("sourceText") or item.get("source_text") or "").strip(),
                "source": str(item.get("source") or "").strip(),
                "reason": str(item.get("reason") or "").strip(),
            }
            if isinstance(item.get("source_refs"), list):
                node["sourceRefs"] = copy.deepcopy(item["source_refs"])
            if isinstance(item.get("material_refs"), list):
                node["materialRefs"] = copy.deepcopy(item["material_refs"])
            if stack:
                stack[-1][1].setdefault("children", []).append(node)
            else:
                roots.append(node)
            stack.append((level, node))
        return roots

    def confirm_outline(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        if not isinstance(project.get("outline_state"), dict):
            project["outline_state"] = self._recover_outline_state(project)
        project["outline_state"]["reviewStatus"] = "confirmed"
        if self._should_skip_technical_gap_stage(project):
            current_stage = self._normalize_stage_value(project.get("currentStage"), scheme=STAGE_SCHEME)
            project["currentStage"] = min(MAX_PROJECT_STAGE, max(current_stage, 4))
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        return {
            "message": "目录已确认",
            "outlineVersion": project["outline_state"]["outlineVersion"],
            "reviewStatus": "confirmed",
        }

    def get_gap_detection(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        gap_state = self._ensure_gap_state(project)
        if self._repair_gap_state_fill_task_skills(gap_state):
            project["updatedAt"] = now_iso()
            self._persist_project(project)
        return self._build_gap_detection_payload(project, gap_state)

    def run_gap_detection(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        gap_state = self._ensure_gap_state(project)
        plan = build_gap_plan_for_project(project)
        items = self._legacy_gap_items_from_plan(plan)
        recognized_at = now_iso()
        plan["summary"] = summarize_gap_plan(plan)
        plan["integrity"] = {}
        gap_state.update(
            {
                "recognitionStatus": "completed",
                "recognizedAt": recognized_at,
                "submittedForReview": False,
                "reviewConfirmed": False,
                "reviewedAt": "",
                "items": items,
                "plan": plan,
                "planFile": str(plan.get("planFile") or ""),
                "integrity": {},
            }
        )
        project["review_document_state"] = self._default_review_document_state(project)
        project["updatedAt"] = recognized_at
        self._persist_project(project)
        return self._build_gap_detection_payload(project, gap_state)

    def get_gap_filling(
        self,
        project_id: str,
        *,
        browser_base_url: str = "",
        onlyoffice_base_url: str = "",
    ) -> dict[str, Any]:
        project = self._require(project_id)
        gap_state = self._ensure_gap_state(project)
        if gap_state["recognitionStatus"] != "completed":
            raise ValueError("请先触发缺口识别后再进入缺口处理。")
        if self._repair_gap_state_fill_task_skills(gap_state):
            project["updatedAt"] = now_iso()
            self._persist_project(project)
        gap_plan = copy.deepcopy(gap_state.get("plan") or {})
        self._refresh_gap_plan_artifact_urls(
            project_id,
            gap_plan,
            browser_base_url=browser_base_url,
            onlyoffice_base_url=onlyoffice_base_url,
        )
        return {
            "status": "ready",
            "recognizedAt": gap_state["recognizedAt"],
            "submittedForReview": bool(gap_state["submittedForReview"]),
            "items": copy.deepcopy(gap_state["items"]),
            "submissions": copy.deepcopy(gap_state["submissions"]),
            "gapPlan": gap_plan,
            "integrity": copy.deepcopy(gap_state.get("integrity") or {}),
            "projectFactTable": copy.deepcopy(gap_state.get("projectFactTable") or {}),
        }

    def get_gap_fact_table(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        gap_state = self._ensure_gap_state(project)
        table = gap_state.get("projectFactTable") if isinstance(gap_state.get("projectFactTable"), dict) else {}
        if table.get("schemaVersion") == PROJECT_FACT_TABLE_SCHEMA_VERSION:
            return copy.deepcopy(table)
        return self._empty_project_fact_table(project_id)

    def build_gap_fact_table(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        gap_state = self._ensure_gap_state(project)
        if gap_state["recognitionStatus"] != "completed":
            raise ValueError("请先完成缺口识别，再维护项目事实表。")
        table = self._build_project_fact_table(project, gap_state)
        gap_state["projectFactTable"] = table
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        return copy.deepcopy(table)

    def save_gap_fact_table(self, project_id: str, data: dict[str, Any]) -> dict[str, Any]:
        project = self._require(project_id)
        gap_state = self._ensure_gap_state(project)
        if gap_state["recognitionStatus"] != "completed":
            raise ValueError("请先完成缺口识别，再维护项目事实表。")
        current = gap_state.get("projectFactTable")
        if not isinstance(current, dict) or current.get("schemaVersion") != PROJECT_FACT_TABLE_SCHEMA_VERSION:
            current = self._build_project_fact_table(project, gap_state)
        incoming_fields = data.get("fields") if isinstance(data.get("fields"), list) else current.get("fields") or []
        confirm = bool(data.get("confirm") or data.get("confirmed"))
        operator = str(data.get("operator") or "当前用户")
        saved_at = now_iso()
        fields = [
            self._normalize_project_fact_field(field, index=index, confirm=confirm, operator=operator, saved_at=saved_at)
            for index, field in enumerate(incoming_fields, start=1)
            if isinstance(field, dict)
        ]
        table = {
            "schemaVersion": PROJECT_FACT_TABLE_SCHEMA_VERSION,
            "projectId": project_id,
            "status": "confirmed" if confirm else "draft",
            "builtAt": str(current.get("builtAt") or saved_at),
            "updatedAt": saved_at,
            "confirmedAt": saved_at if confirm else str(current.get("confirmedAt") or ""),
            "confirmedBy": operator if confirm else str(current.get("confirmedBy") or ""),
            "fields": fields,
            "summary": self._summarize_project_fact_fields(fields),
        }
        gap_state["projectFactTable"] = table
        project["updatedAt"] = saved_at
        self._persist_project(project)
        return copy.deepcopy(table)

    def run_gap_ai_fill(
        self,
        project_id: str,
        gap_id: str,
        data: dict[str, Any],
        *,
        browser_base_url: str = "",
        onlyoffice_base_url: str = "",
    ) -> dict[str, Any]:
        project = self._require(project_id)
        gap_state = self._ensure_gap_state(project)
        if gap_state["recognitionStatus"] != "completed":
            raise ValueError("请先完成缺口识别。")
        if self._repair_gap_state_fill_task_skills(gap_state):
            project["updatedAt"] = now_iso()
            self._persist_project(project)
        self._require_confirmed_project_fact_table(gap_state)
        result = run_ai_fill_for_gap(
            project,
            gap_id,
            data,
            browser_base_url=browser_base_url,
            onlyoffice_base_url=onlyoffice_base_url,
        )
        gap_state = self._ensure_gap_state(project)
        gap_state["integrity"] = check_gap_integrity(gap_state.get("plan") or {})
        if isinstance(gap_state.get("plan"), dict):
            gap_state["plan"]["integrity"] = gap_state["integrity"]
            gap_state["plan"]["summary"] = summarize_gap_plan(gap_state["plan"])
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        return copy.deepcopy(result)

    def run_gap_ai_fill_all(
        self,
        project_id: str,
        data: dict[str, Any],
        *,
        browser_base_url: str = "",
        onlyoffice_base_url: str = "",
    ) -> dict[str, Any]:
        project = self._require(project_id)
        gap_state = self._ensure_gap_state(project)
        if gap_state["recognitionStatus"] != "completed":
            raise ValueError("请先完成缺口识别。")
        if self._repair_gap_state_fill_task_skills(gap_state):
            project["updatedAt"] = now_iso()
            self._persist_project(project)
        self._require_confirmed_project_fact_table(gap_state)
        plan = gap_state.get("plan") if isinstance(gap_state.get("plan"), dict) else {}
        requested_gap_ids = {
            str(item or "").strip()
            for item in (data.get("gapIds") if isinstance(data.get("gapIds"), list) else [])
            if str(item or "").strip()
        }
        tasks: list[tuple[int, int, int, str, str]] = []
        for index, item in enumerate(plan.get("items") or [], start=1):
            if not isinstance(item, dict):
                continue
            gap_id = str(item.get("id") or "")
            if requested_gap_ids and gap_id not in requested_gap_ids:
                continue
            if str(item.get("decision") or "") != "fill_required":
                continue
            for task_index, task in enumerate(item.get("fillTasks") or [], start=1):
                if not isinstance(task, dict):
                    continue
                if str(task.get("status") or "pending") == "completed" and not data.get("rerun"):
                    continue
                skill = str(task.get("skill") or TABLE_FILL_SKILL_NAME)
                rank = 0 if skill == WORD_FILL_SKILL_NAME else 1
                tasks.append((rank, index, task_index, gap_id, str(task.get("id") or "")))
        tasks.sort(key=lambda item: (item[0], item[1], item[2]))
        base_data = {
            key: value
            for key, value in data.items()
            if key not in {"fillTaskId", "gapIds", "rerun"}
        }
        results: list[dict[str, Any]] = []
        errors: list[dict[str, str]] = []
        for _, _, _, gap_id, fill_task_id in tasks:
            try:
                result = run_ai_fill_for_gap(
                    project,
                    gap_id,
                    {**base_data, "fillTaskId": fill_task_id, "operator": str(data.get("operator") or "当前用户")},
                    browser_base_url=browser_base_url,
                    onlyoffice_base_url=onlyoffice_base_url,
                )
                artifact = result.get("artifact") if isinstance(result.get("artifact"), dict) else {}
                artifacts = [
                    item
                    for item in (result.get("artifacts") if isinstance(result.get("artifacts"), list) else [artifact])
                    if isinstance(item, dict) and item
                ]
                for artifact_item in artifacts:
                    results.append(
                        {
                            "gapId": gap_id,
                            "artifactId": str(artifact_item.get("id") or ""),
                            "artifactIds": [str(item.get("id") or "") for item in artifacts if str(item.get("id") or "")],
                            "skill": str(artifact_item.get("skill") or ""),
                            "fileName": str(artifact_item.get("fileName") or ""),
                            "batchTargetIndex": artifact_item.get("batchTargetIndex") or 0,
                            "batchTargetCount": artifact_item.get("batchTargetCount") or 0,
                            "qualityReport": copy.deepcopy(artifact_item.get("qualityReport") or {}),
                        }
                    )
                project["updatedAt"] = now_iso()
                self._persist_project(project)
            except Exception as exc:  # pragma: no cover - batch must report failures instead of hiding progress
                errors.append({"gapId": gap_id, "message": str(exc)})
        gap_state = self._ensure_gap_state(project)
        gap_state["integrity"] = check_gap_integrity(gap_state.get("plan") or {})
        if isinstance(gap_state.get("plan"), dict):
            gap_state["plan"]["integrity"] = gap_state["integrity"]
            gap_state["plan"]["summary"] = summarize_gap_plan(gap_state["plan"])
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        aggregate = self._aggregate_gap_fill_quality(results, errors)
        return {
            "status": "completed" if not errors else "needs_review",
            "summary": {
                "total": len(results) + len(errors),
                "passed": sum(1 for result in results if result.get("qualityReport", {}).get("status") == "passed"),
                "needsReview": sum(1 for result in results if result.get("qualityReport", {}).get("status") != "passed"),
                "failed": len(errors),
            },
            "qualityReport": aggregate,
            "results": results,
            "errors": errors,
            "gapPlan": copy.deepcopy(gap_state.get("plan") or {}),
            "projectFactTable": copy.deepcopy(gap_state.get("projectFactTable") or {}),
        }

    def upload_gap_artifact(
        self,
        project_id: str,
        gap_id: str,
        data: dict[str, Any],
        *,
        browser_base_url: str = "",
        onlyoffice_base_url: str = "",
    ) -> dict[str, Any]:
        project = self._require(project_id)
        gap_state = self._ensure_gap_state(project)
        if gap_state["recognitionStatus"] != "completed":
            raise ValueError("请先完成缺口识别。")
        result = register_manual_gap_upload(
            project,
            gap_id,
            data,
            browser_base_url=browser_base_url,
            onlyoffice_base_url=onlyoffice_base_url,
        )
        gap_state = self._ensure_gap_state(project)
        gap_state["integrity"] = check_gap_integrity(gap_state.get("plan") or {})
        if isinstance(gap_state.get("plan"), dict):
            gap_state["plan"]["integrity"] = gap_state["integrity"]
            gap_state["plan"]["summary"] = summarize_gap_plan(gap_state["plan"])
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        return copy.deepcopy(result)

    async def select_gap_material(
        self,
        project_id: str,
        gap_id: str,
        data: dict[str, Any],
        *,
        browser_base_url: str = "",
        onlyoffice_base_url: str = "",
    ) -> dict[str, Any]:
        project = self._require(project_id)
        gap_state = self._ensure_gap_state(project)
        if gap_state["recognitionStatus"] != "completed":
            raise ValueError("请先完成缺口识别。")
        prepared_files = await prepare_existing_gap_material_files(project, gap_id, data)
        result = register_existing_gap_material(
            project,
            gap_id,
            data,
            prepared_files,
            browser_base_url=browser_base_url,
            onlyoffice_base_url=onlyoffice_base_url,
        )
        gap_state = self._ensure_gap_state(project)
        gap_state["integrity"] = check_gap_integrity(gap_state.get("plan") or {})
        if isinstance(gap_state.get("plan"), dict):
            gap_state["plan"]["integrity"] = gap_state["integrity"]
            gap_state["plan"]["summary"] = summarize_gap_plan(gap_state["plan"])
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        return copy.deepcopy(result)

    def check_gap_plan_integrity(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        gap_state = self._ensure_gap_state(project)
        if gap_state["recognitionStatus"] != "completed":
            raise ValueError("请先完成缺口识别。")
        integrity = check_gap_integrity(gap_state.get("plan") or {})
        gap_state["integrity"] = integrity
        plan = gap_state.get("plan")
        if isinstance(plan, dict):
            plan["integrity"] = integrity
            plan["summary"] = summarize_gap_plan(plan)
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        return copy.deepcopy(integrity)

    def run_business_gap_detection(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        if normalize_bid_type(str(project.get("bidType") or "")) != "商务标":
            raise ValueError("商务标缺口处理仅支持商务标项目。")
        if str((project.get("outline_state") or {}).get("reviewStatus") or "") != "confirmed":
            raise ValueError("请先完成人工目录审核，再生成商务标缺口计划。")
        business_gap_state = self._ensure_business_gap_state(project)
        plan = build_business_gap_plan_for_project(project)
        recognized_at = now_iso()
        plan["summary"] = summarize_business_gap_plan(plan)
        business_gap_state.update(
            {
                "recognitionStatus": "completed",
                "recognizedAt": recognized_at,
                "submittedForReview": False,
                "reviewConfirmed": False,
                "reviewedAt": "",
                "plan": plan,
                "planFile": str(plan.get("planFile") or ""),
                "integrity": self._business_gap_integrity(plan),
            }
        )
        plan["integrity"] = copy.deepcopy(business_gap_state["integrity"])
        project["updatedAt"] = recognized_at
        self._persist_project(project)
        return self._build_business_gap_payload(project, business_gap_state)

    def get_business_gap_filling(
        self,
        project_id: str,
        *,
        browser_base_url: str = "",
        onlyoffice_base_url: str = "",
    ) -> dict[str, Any]:
        project = self._require(project_id)
        if normalize_bid_type(str(project.get("bidType") or "")) != "商务标":
            raise ValueError("商务标缺口处理仅支持商务标项目。")
        business_gap_state = self._ensure_business_gap_state(project)
        if business_gap_state["recognitionStatus"] != "completed":
            return self._build_business_gap_payload(project, business_gap_state)
        if self._refresh_business_gap_template_candidates(project, business_gap_state):
            business_gap_state = self._ensure_business_gap_state(project)
        payload = self._build_business_gap_payload(project, business_gap_state)
        plan = payload.get("plan") if isinstance(payload.get("plan"), dict) else {}
        refresh_business_gap_artifact_urls(
            project_id,
            plan,
            browser_base_url=browser_base_url,
            onlyoffice_base_url=onlyoffice_base_url,
        )
        payload["plan"] = plan
        payload["businessGapPlan"] = plan
        return payload

    def list_business_gap_selectable_materials(self, project_id: str, *, keyword: str = "") -> dict[str, Any]:
        project = self._require(project_id)
        if normalize_bid_type(str(project.get("bidType") or "")) != "商务标":
            raise ValueError("商务标素材选择仅支持商务标项目。")
        business_gap_state = self._ensure_business_gap_state(project)
        if business_gap_state["recognitionStatus"] != "completed":
            raise ValueError("请先生成商务标缺口计划。")
        picker = build_business_gap_material_picker_index(project)
        templates = [item for item in picker.get("templateIndex") or [] if isinstance(item, dict)]
        materials = [item for item in picker.get("materialIndex") or [] if isinstance(item, dict)]
        segments = [item for item in picker.get("evidenceSegments") or [] if isinstance(item, dict)]
        material_by_id = {
            str(item.get("id") or item.get("materialId") or ""): item
            for item in materials
            if str(item.get("id") or item.get("materialId") or "")
        }
        normalized_keyword = str(keyword or "").strip().lower()

        def matches(text: str) -> bool:
            return not normalized_keyword or normalized_keyword in str(text or "").lower()

        selectable_templates: list[dict[str, Any]] = []
        for template in templates:
            haystack = " ".join(
                [
                    str(template.get("templateName") or ""),
                    str(template.get("fileName") or ""),
                    str(template.get("sourceLabel") or ""),
                    str(template.get("sourceMode") or ""),
                    str(template.get("filePath") or ""),
                ]
            )
            if not matches(haystack):
                continue
            selectable_templates.append(
                {
                    "id": str(template.get("templateId") or template.get("filePath") or ""),
                    "kind": "template",
                    "templateId": str(template.get("templateId") or ""),
                    "templateName": str(template.get("templateName") or template.get("fileName") or ""),
                    "fileName": str(template.get("fileName") or ""),
                    "filePath": str(template.get("filePath") or ""),
                    "sourceMode": str(template.get("sourceMode") or ""),
                    "sourceLabel": str(template.get("sourceLabel") or ""),
                    "templateScope": str(template.get("templateScope") or ""),
                    "assemblyMode": "template_fill_docx",
                    "materialUsage": "fill_template",
                    "mimeType": str(template.get("mimeType") or ""),
                    "score": float(template.get("score") or 0),
                    "reason": str(template.get("reason") or ""),
                }
            )

        selectable_segments: list[dict[str, Any]] = []
        for segment in segments:
            material_id = str(segment.get("material_id") or segment.get("materialId") or "").strip()
            material = material_by_id.get(material_id, {})
            if not material_id or not material:
                continue
            name = str(material.get("name") or material.get("fileName") or segment.get("title") or material_id)
            folder_path = str(material.get("folderPath") or segment.get("path") or "")
            haystack = " ".join(
                [
                    name,
                    folder_path,
                    str(segment.get("title") or ""),
                    str(segment.get("summary") or ""),
                    " ".join(str(item) for item in segment.get("keywords") or []),
                    str(segment.get("business_category") or ""),
                    str(segment.get("document_type") or ""),
                ]
            )
            if material_id and not matches(haystack):
                continue
            selectable_segments.append(
                {
                    "id": str(segment.get("segment_id") or ""),
                    "kind": "segment",
                    "materialId": material_id,
                    "materialName": name,
                    "folderPath": folder_path,
                    "materialTier": str(material.get("materialTier") or segment.get("material_tier") or ""),
                    "cleanStatus": str(material.get("cleanStatus") or ""),
                    "hasCleanedWord": bool(material.get("hasCleanedWord")),
                    "cleanedFileName": str(material.get("cleanedFileName") or ""),
                    "evidenceSegmentId": str(segment.get("segment_id") or ""),
                    "evidenceSegmentTitle": str(segment.get("title") or ""),
                    "evidenceSegmentType": str(segment.get("segment_type") or ""),
                    "evidenceSourcePages": str(segment.get("source_pages") or ""),
                    "evidenceSummary": str(segment.get("summary") or ""),
                    "wikiCardId": str(segment.get("card_id") or ""),
                    "wikiUsageMode": str(segment.get("usage_mode") or ""),
                    "wikiEvidence": {
                        "validityStatus": str(segment.get("validity_status") or ""),
                        "expiryDate": str(segment.get("expiry_date") or ""),
                        "riskNotes": str(segment.get("risk_notes") or ""),
                        "ocrStatus": str(segment.get("ocr_status") or ""),
                        "ocrConfidence": str(segment.get("ocr_confidence") or ""),
                    },
                    "keywords": [str(item) for item in segment.get("keywords") or [] if str(item).strip()][:16],
                    "updatedAt": str(material.get("updatedAt") or ""),
                }
            )

        material_ids_with_segments = {str(item.get("materialId") or "") for item in selectable_segments if str(item.get("materialId") or "")}
        selectable_materials: list[dict[str, Any]] = []
        for material in materials:
            material_id = str(material.get("id") or material.get("materialId") or "").strip()
            if not material_id:
                continue
            name = str(material.get("name") or material.get("fileName") or material_id)
            folder_path = str(material.get("folderPath") or "")
            haystack = " ".join([name, folder_path, str(material.get("cleanedFileName") or ""), str(material.get("turbineModelLabel") or "")])
            if not matches(haystack):
                continue
            selectable_materials.append(
                {
                    "id": material_id,
                    "kind": "material",
                    "materialId": material_id,
                    "materialName": name,
                    "folderPath": folder_path,
                    "materialTier": str(material.get("materialTier") or ""),
                    "cleanStatus": str(material.get("cleanStatus") or ""),
                    "hasCleanedWord": bool(material.get("hasCleanedWord")),
                    "cleanedFileName": str(material.get("cleanedFileName") or ""),
                    "size": int(material.get("size") or 0),
                    "turbineModelLabel": str(material.get("turbineModelLabel") or ""),
                    "updatedAt": str(material.get("updatedAt") or ""),
                    "segmentCount": sum(1 for item in selectable_segments if str(item.get("materialId") or "") == material_id),
                    "hasSegments": material_id in material_ids_with_segments,
                }
            )

        return {
            "schemaVersion": "bid-business-selectable-materials-v1",
            "projectId": project_id,
            "bidType": "商务标",
            "keyword": str(keyword or ""),
            "templates": selectable_templates,
            "items": selectable_materials,
            "segments": selectable_segments,
            "summary": {
                "templateCount": len(selectable_templates),
                "materialCount": len(selectable_materials),
                "segmentCount": len(selectable_segments),
                "wikiEvidenceSegmentCount": int((picker.get("businessWikiIndexSummary") or {}).get("evidenceSegmentCount") or 0),
            },
            "materialScope": copy.deepcopy(picker.get("materialScope") or {}),
            "selectedBusinessTurbineModel": copy.deepcopy(picker.get("selectedBusinessTurbineModel") or {}),
        }

    def get_business_gap_fact_table(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        if normalize_bid_type(str(project.get("bidType") or "")) != "商务标":
            raise ValueError("商务标项目事实表仅支持商务标项目。")
        business_gap_state = self._ensure_business_gap_state(project)
        table = business_gap_state.get("projectFactTable") if isinstance(business_gap_state.get("projectFactTable"), dict) else {}
        if table.get("schemaVersion") == PROJECT_FACT_TABLE_SCHEMA_VERSION:
            return copy.deepcopy(table)
        return self._empty_project_fact_table(project_id)

    def build_business_gap_fact_table(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        if normalize_bid_type(str(project.get("bidType") or "")) != "商务标":
            raise ValueError("商务标项目事实表仅支持商务标项目。")
        business_gap_state = self._ensure_business_gap_state(project)
        if business_gap_state["recognitionStatus"] != "completed":
            raise ValueError("请先生成商务标缺口计划，再维护项目事实表。")
        table = self._build_project_fact_table(project, business_gap_state)
        business_gap_state["projectFactTable"] = table
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        return copy.deepcopy(table)

    def save_business_gap_fact_table(self, project_id: str, data: dict[str, Any]) -> dict[str, Any]:
        project = self._require(project_id)
        if normalize_bid_type(str(project.get("bidType") or "")) != "商务标":
            raise ValueError("商务标项目事实表仅支持商务标项目。")
        business_gap_state = self._ensure_business_gap_state(project)
        if business_gap_state["recognitionStatus"] != "completed":
            raise ValueError("请先生成商务标缺口计划，再维护项目事实表。")
        current = business_gap_state.get("projectFactTable")
        if not isinstance(current, dict) or current.get("schemaVersion") != PROJECT_FACT_TABLE_SCHEMA_VERSION:
            current = self._build_project_fact_table(project, business_gap_state)
        incoming_fields = data.get("fields") if isinstance(data.get("fields"), list) else current.get("fields") or []
        confirm = bool(data.get("confirm") or data.get("confirmed"))
        operator = str(data.get("operator") or "当前用户")
        saved_at = now_iso()
        fields = [
            self._normalize_project_fact_field(field, index=index, confirm=confirm, operator=operator, saved_at=saved_at)
            for index, field in enumerate(incoming_fields, start=1)
            if isinstance(field, dict)
        ]
        table = {
            "schemaVersion": PROJECT_FACT_TABLE_SCHEMA_VERSION,
            "projectId": project_id,
            "status": "confirmed" if confirm else "draft",
            "builtAt": str(current.get("builtAt") or saved_at),
            "updatedAt": saved_at,
            "confirmedAt": saved_at if confirm else str(current.get("confirmedAt") or ""),
            "confirmedBy": operator if confirm else str(current.get("confirmedBy") or ""),
            "fields": fields,
            "summary": self._summarize_project_fact_fields(fields),
        }
        business_gap_state["projectFactTable"] = table
        project["updatedAt"] = saved_at
        self._persist_project(project)
        return copy.deepcopy(table)

    def update_business_gap_task(self, project_id: str, task_id: str, data: dict[str, Any]) -> dict[str, Any]:
        project = self._require(project_id)
        business_gap_state = self._ensure_business_gap_state(project)
        plan = business_gap_state.get("plan") if isinstance(business_gap_state.get("plan"), dict) else {}
        task = self._find_business_gap_task(plan, task_id)
        allowed = {
            "status",
            "decision",
            "selectedMaterialRefs",
            "notes",
            "confirmed",
            "riskFlags",
            "assemblyMode",
            "materialUsage",
            "fillPlan",
            "selectedEvidenceSegments",
        }
        for key in allowed:
            if key in data:
                task[key] = copy.deepcopy(data[key])
        intent_changed = "assemblyMode" in data or "materialUsage" in data or "selectedEvidenceSegments" in data
        explicit_status_change = "status" in data or "decision" in data
        if intent_changed:
            if "assemblyMode" in data and "materialUsage" not in data:
                task["materialUsage"] = self._business_material_usage_for_assembly_mode(str(task.get("assemblyMode") or ""))
            elif not task.get("materialUsage"):
                task["materialUsage"] = self._business_material_usage_for_assembly_mode(str(task.get("assemblyMode") or ""))
            task["fillPlan"] = self._business_task_fill_plan(task)
            if not explicit_status_change:
                self._recompute_business_gap_task_after_artifact_change(task)
            if str(task.get("assemblyMode") or "") == "template_fill_docx":
                self._refresh_business_gap_template_candidates(project, business_gap_state)
                plan = business_gap_state.get("plan") if isinstance(business_gap_state.get("plan"), dict) else plan
                task = self._find_business_gap_task(plan, task_id)
        task["updatedAt"] = now_iso()
        self._finalize_business_gap_plan_update(project, business_gap_state, plan, updated_at=task["updatedAt"])
        return {"task": copy.deepcopy(task), "plan": copy.deepcopy(plan), "integrity": copy.deepcopy(business_gap_state["integrity"])}

    def create_business_gap_manual_task(
        self,
        project_id: str,
        toc_node_id: str,
        data: dict[str, Any],
    ) -> dict[str, Any]:
        project = self._require(project_id)
        if normalize_bid_type(str(project.get("bidType") or "")) != "商务标":
            raise ValueError("商务标缺口处理仅支持商务标项目。")
        business_gap_state = self._ensure_business_gap_state(project)
        if business_gap_state["recognitionStatus"] != "completed":
            raise ValueError("请先生成商务标缺口计划。")
        plan = business_gap_state.get("plan") if isinstance(business_gap_state.get("plan"), dict) else {}
        toc_ref = self._find_business_gap_toc_ref(plan, toc_node_id)
        tasks = plan.setdefault("tasks", [])
        if not isinstance(tasks, list):
            plan["tasks"] = tasks = []
        created_at = now_iso()
        existing_ids = {str(task.get("id") or "") for task in tasks if isinstance(task, dict)}
        base_index = len(tasks) + 1
        task_id = f"BTASK-MANUAL-{base_index:04d}"
        while task_id in existing_ids:
            base_index += 1
            task_id = f"BTASK-MANUAL-{base_index:04d}"
        title = str(data.get("title") or "").strip() or f"{toc_ref.get('title') or '本章节'}补充材料"
        task = {
            "id": task_id,
            "taskKey": f"manual_{safe_filename(str(toc_ref.get('nodeId') or toc_node_id), 'toc')}_{base_index}",
            "title": title,
            "titleAlias": [],
            "moduleKey": str(data.get("moduleKey") or "commitments_and_notes"),
            "taskType": str(data.get("taskType") or "attachment"),
            "decision": "material_required",
            "status": "needs_input",
            "sourceType": "manual_user",
            "sourceRequirement": {
                "triggerText": str(data.get("requirement") or f"操作人针对目录章节「{toc_ref.get('title') or toc_node_id}」手动补充材料。"),
                "triggerContext": "",
                "normalizedTopic": title,
                "fromSection": str(toc_ref.get("number") or ""),
                "extractionMethod": "manual_user",
            },
            "sourceEvidenceRefs": [],
            "tocTarget": {
                "nodeId": str(toc_ref.get("nodeId") or toc_node_id),
                "number": str(toc_ref.get("number") or ""),
                "title": str(toc_ref.get("title") or ""),
                "required": True,
            },
            "candidateMaterials": [],
            "selectedMaterialRefs": [],
            "resolvedArtifacts": [],
            "assemblyMode": self._default_business_task_assembly_mode(
                str(data.get("moduleKey") or "commitments_and_notes"),
                str(data.get("taskType") or "attachment"),
                "material_required",
                title,
            ),
            "materialUsage": "",
            "fillPlan": {},
            "selectedEvidenceSegments": [],
            "riskFlags": ["manual_upload_required"],
            "requirementLevel": "required",
            "assigneeMode": "manual_select",
            "displayOrder": 9000 + base_index,
            "fingerprint": f"manual:{project_id}:{toc_node_id}:{base_index}",
            "updatedAt": created_at,
        }
        task["materialUsage"] = self._business_material_usage_for_assembly_mode(str(task.get("assemblyMode") or ""))
        task["fillPlan"] = self._business_task_fill_plan(task)
        tasks.append(task)
        task_ids = toc_ref.setdefault("taskIds", [])
        if task_id not in task_ids:
            task_ids.append(task_id)
        self._update_business_gap_toc_ref_statuses(plan)
        self._finalize_business_gap_plan_update(project, business_gap_state, plan, updated_at=created_at)
        return {
            "message": "已为当前目录章节创建人工补料任务。",
            "task": copy.deepcopy(task),
            "plan": copy.deepcopy(plan),
            "integrity": copy.deepcopy(business_gap_state["integrity"]),
        }

    def confirm_business_gap_artifact(self, project_id: str, task_id: str, data: dict[str, Any]) -> dict[str, Any]:
        project = self._require(project_id)
        business_gap_state = self._ensure_business_gap_state(project)
        plan = business_gap_state.get("plan") if isinstance(business_gap_state.get("plan"), dict) else {}
        task = self._find_business_gap_task(plan, task_id)
        artifact_id = str(data.get("artifactId") or "").strip()
        confirmed = bool(data.get("confirmed", True))
        artifacts = task.get("resolvedArtifacts") if isinstance(task.get("resolvedArtifacts"), list) else []
        artifact = None
        for item in artifacts:
            if not isinstance(item, dict):
                continue
            if not artifact_id or artifact_id == str(item.get("artifactId") or ""):
                artifact = item
                break
        if artifact is None:
            raise KeyError(artifact_id or task_id)
        artifact["confirmed"] = confirmed
        artifact["reviewStatus"] = "approved" if confirmed else "pending_review"
        artifact["confirmedAt"] = now_iso() if confirmed else ""
        if confirmed:
            task["decision"] = "ready"
            task["status"] = "ready"
            task["riskFlags"] = [
                flag
                for flag in (task.get("riskFlags") if isinstance(task.get("riskFlags"), list) else [])
                if flag not in {"missing_material", "parser_generated_unconfirmed"}
            ]
        else:
            task["decision"] = "review_required"
            task["status"] = "review_required"
        task["updatedAt"] = now_iso()
        self._finalize_business_gap_plan_update(project, business_gap_state, plan, updated_at=task["updatedAt"])
        return {"task": copy.deepcopy(task), "artifact": copy.deepcopy(artifact), "plan": copy.deepcopy(plan)}

    def remove_business_gap_artifact(
        self,
        project_id: str,
        task_id: str,
        artifact_id: str,
        *,
        browser_base_url: str = "",
        onlyoffice_base_url: str = "",
    ) -> dict[str, Any]:
        project = self._require(project_id)
        if normalize_bid_type(str(project.get("bidType") or "")) != "商务标":
            raise ValueError("商务标缺口处理仅支持商务标项目。")
        business_gap_state = self._ensure_business_gap_state(project)
        if business_gap_state["recognitionStatus"] != "completed":
            raise ValueError("请先生成商务标缺口计划。")
        plan = business_gap_state.get("plan") if isinstance(business_gap_state.get("plan"), dict) else {}
        task = self._find_business_gap_task(plan, task_id)
        artifacts = task.get("resolvedArtifacts") if isinstance(task.get("resolvedArtifacts"), list) else []
        target = str(artifact_id or "").strip()
        kept: list[dict[str, Any]] = []
        removed: dict[str, Any] | None = None
        for artifact in artifacts:
            if isinstance(artifact, dict) and str(artifact.get("artifactId") or "") == target:
                removed = artifact
                continue
            if isinstance(artifact, dict):
                kept.append(artifact)
        if removed is None:
            raise KeyError(artifact_id)
        if str(removed.get("sourceMode") or "") not in {"uploaded_in_business_s3", "selected_from_business_material_library"}:
            raise ValueError("解析生成产物不能在 S3 页面直接取消，请在解析产物审核处处理。")
        if str(removed.get("materialSyncStatus") or "") == "synced_to_project_material":
            raise ValueError("该补料已同步到项目素材库，不能直接取消；如需删除，请在素材库中处理。")

        file_path = Path(str(removed.get("filePath") or ""))
        if file_path.exists() and file_path.is_file():
            try:
                file_path.unlink()
            except OSError:
                removed["deleteWarning"] = "文件删除失败，仅从当前任务中移除记录。"

        task["resolvedArtifacts"] = kept
        removed_material_id = str(removed.get("materialId") or "")
        if removed_material_id and isinstance(task.get("selectedMaterialRefs"), list):
            task["selectedMaterialRefs"] = [
                ref
                for ref in task["selectedMaterialRefs"]
                if not isinstance(ref, dict) or str(ref.get("materialId") or ref.get("id") or "") != removed_material_id
            ]
        self._recompute_business_gap_task_after_artifact_change(task)
        now = now_iso()
        task["updatedAt"] = now
        self._finalize_business_gap_plan_update(project, business_gap_state, plan, updated_at=now)
        self._refresh_business_gap_urls_for_result(
            project_id,
            plan,
            browser_base_url=browser_base_url,
            onlyoffice_base_url=onlyoffice_base_url,
        )
        return {
            "message": "已取消该补料。",
            "task": copy.deepcopy(task),
            "artifact": copy.deepcopy(removed),
            "plan": copy.deepcopy(plan),
            "integrity": copy.deepcopy(business_gap_state["integrity"]),
        }

    def upload_business_gap_artifact(
        self,
        project_id: str,
        task_id: str,
        data: dict[str, Any],
        *,
        browser_base_url: str = "",
        onlyoffice_base_url: str = "",
    ) -> dict[str, Any]:
        project = self._require(project_id)
        if normalize_bid_type(str(project.get("bidType") or "")) != "商务标":
            raise ValueError("商务标缺口处理仅支持商务标项目。")
        business_gap_state = self._ensure_business_gap_state(project)
        if business_gap_state["recognitionStatus"] != "completed":
            raise ValueError("请先生成商务标缺口计划。")
        plan = business_gap_state.get("plan") if isinstance(business_gap_state.get("plan"), dict) else {}
        task = self._find_business_gap_task(plan, task_id)
        files = [entry for entry in data.get("files") or [] if isinstance(entry, dict)]
        if not files:
            raise ValueError("至少需要上传一个文件。")

        upload_records: list[dict[str, Any]] = []
        for index, file in enumerate(files, start=1):
            source_name = str(file.get("name") or file.get("fileName") or f"{task_id}-{index}.bin")
            content = str(
                file.get("data")
                or file.get("dataUrl")
                or file.get("content")
                or file.get("base64")
                or file.get("base64Content")
                or ""
            )
            raw_bytes, mime_type = self._decode_business_gap_upload_content(
                content,
                fallback_mime=str(file.get("mimeType") or ""),
            )
            upload_records.append(
                {
                    "name": source_name,
                    "mimeType": mime_type or str(file.get("mimeType") or ""),
                    "rawBytes": raw_bytes,
                }
            )
        return self._register_business_gap_upload_records(
            project,
            business_gap_state,
            plan,
            task,
            upload_records,
            operator=str(data.get("operator") or "当前用户"),
            browser_base_url=browser_base_url,
            onlyoffice_base_url=onlyoffice_base_url,
        )

    def upload_business_gap_artifact_bytes(
        self,
        project_id: str,
        task_id: str,
        files: list[dict[str, Any]],
        *,
        operator: str = "当前用户",
        browser_base_url: str = "",
        onlyoffice_base_url: str = "",
    ) -> dict[str, Any]:
        project = self._require(project_id)
        if normalize_bid_type(str(project.get("bidType") or "")) != "商务标":
            raise ValueError("商务标缺口处理仅支持商务标项目。")
        business_gap_state = self._ensure_business_gap_state(project)
        if business_gap_state["recognitionStatus"] != "completed":
            raise ValueError("请先生成商务标缺口计划。")
        plan = business_gap_state.get("plan") if isinstance(business_gap_state.get("plan"), dict) else {}
        task = self._find_business_gap_task(plan, task_id)
        return self._register_business_gap_upload_records(
            project,
            business_gap_state,
            plan,
            task,
            files,
            operator=operator,
            browser_base_url=browser_base_url,
            onlyoffice_base_url=onlyoffice_base_url,
        )

    def _register_business_gap_upload_records(
        self,
        project: dict[str, Any],
        business_gap_state: dict[str, Any],
        plan: dict[str, Any],
        task: dict[str, Any],
        files: list[dict[str, Any]],
        *,
        operator: str,
        browser_base_url: str = "",
        onlyoffice_base_url: str = "",
    ) -> dict[str, Any]:
        project_id = str(project.get("id") or "")
        task_id = str(task.get("id") or task.get("taskKey") or "task")
        records = [entry for entry in files if isinstance(entry, dict)]
        if not records:
            raise ValueError("至少需要上传一个文件。")
        work_dir = business_workspace_dir(project_id) / "gaps" / "uploads" / safe_filename(task_id, "task")
        work_dir.mkdir(parents=True, exist_ok=True)
        created_at = now_iso()
        artifacts: list[dict[str, Any]] = []
        existing_count = len(task.get("resolvedArtifacts") if isinstance(task.get("resolvedArtifacts"), list) else [])
        for index, file in enumerate(records, start=1):
            source_name = str(file.get("name") or file.get("fileName") or f"{task_id}-{index}.bin")
            file_name = safe_filename(source_name, f"{task_id}-{index}.bin")
            raw_bytes = file.get("rawBytes") or file.get("bytes") or b""
            if isinstance(raw_bytes, str):
                raw_bytes, mime_type = self._decode_business_gap_upload_content(
                    raw_bytes,
                    fallback_mime=str(file.get("mimeType") or ""),
                )
            else:
                raw_bytes = bytes(raw_bytes or b"")
                mime_type = str(file.get("mimeType") or "")
            if not raw_bytes:
                raise ValueError(f"上传文件内容为空：{file_name}")
            target_path = self._unique_business_gap_path(work_dir, file_name)
            target_path.write_bytes(raw_bytes)
            artifact_id = f"BART-{safe_filename(task_id, 'TASK')}-UPLOAD-{existing_count + index}"
            artifact = {
                "artifactId": artifact_id,
                "artifactType": "manual_upload",
                "fileName": target_path.name,
                "filePath": str(target_path),
                "sourceMode": "uploaded_in_business_s3",
                "assemblyMode": self._assembly_mode_for_business_artifact(task, {"fileName": target_path.name, "mimeType": mime_type}),
                "materialUsage": "",
                "materialSyncStatus": "not_synced",
                "materialSyncPolicy": "manual_project_only",
                "materialTargetPath": self._default_business_gap_material_target_path(project_id, task),
                "wikiSyncStatus": "not_synced",
                "version": 1,
                "previewable": True,
                "confirmed": True,
                "reviewStatus": "approved",
                "confirmedAt": created_at,
                "uploadedAt": created_at,
                "operator": operator or "当前用户",
                "mimeType": mime_type or mimetypes.guess_type(target_path.name)[0] or "application/octet-stream",
            }
            artifact["materialUsage"] = self._business_material_usage_for_assembly_mode(str(artifact.get("assemblyMode") or ""))
            artifacts.append(artifact)

        task.setdefault("resolvedArtifacts", []).extend(artifacts)
        self._apply_business_task_artifact_intent(task, artifacts)
        task["decision"] = "ready"
        task["status"] = "ready"
        task["updatedAt"] = created_at
        task["resolvedAt"] = created_at
        task["resolvedSource"] = artifacts[0]["fileName"]
        task["riskFlags"] = [
            flag
            for flag in (task.get("riskFlags") if isinstance(task.get("riskFlags"), list) else [])
            if flag not in {"missing_material", "manual_upload_required", "ai_draft_required"}
        ]
        self._finalize_business_gap_plan_update(project, business_gap_state, plan, updated_at=created_at)
        self._refresh_business_gap_urls_for_result(
            project_id,
            plan,
            browser_base_url=browser_base_url,
            onlyoffice_base_url=onlyoffice_base_url,
        )
        return {
            "task": copy.deepcopy(task),
            "artifact": copy.deepcopy(artifacts[0]),
            "artifacts": copy.deepcopy(artifacts),
            "plan": copy.deepcopy(plan),
            "integrity": copy.deepcopy(business_gap_state["integrity"]),
        }

    async def select_business_gap_material(
        self,
        project_id: str,
        task_id: str,
        data: dict[str, Any],
        *,
        browser_base_url: str = "",
        onlyoffice_base_url: str = "",
    ) -> dict[str, Any]:
        project = self._require(project_id)
        if normalize_bid_type(str(project.get("bidType") or "")) != "商务标":
            raise ValueError("商务标缺口处理仅支持商务标项目。")
        business_gap_state = self._ensure_business_gap_state(project)
        if business_gap_state["recognitionStatus"] != "completed":
            raise ValueError("请先生成商务标缺口计划。")
        plan = business_gap_state.get("plan") if isinstance(business_gap_state.get("plan"), dict) else {}
        task = self._find_business_gap_task(plan, task_id)
        selected = self._selected_business_material_entries(data)
        if not selected:
            raise ValueError("至少需要选择一份素材。")

        work_dir = business_workspace_dir(project_id) / "gaps" / "selected-materials" / safe_filename(task_id, "task")
        work_dir.mkdir(parents=True, exist_ok=True)
        created_at = now_iso()
        existing_count = len(task.get("resolvedArtifacts") if isinstance(task.get("resolvedArtifacts"), list) else [])
        material_refs: list[dict[str, Any]] = []
        artifacts: list[dict[str, Any]] = []
        for index, material in enumerate(selected, start=1):
            material_id = str(material.get("id") or material.get("materialId") or "").strip()
            if not material_id:
                continue
            payload, source_kind = await self._business_material_download_payload(material_id)
            file_name = safe_filename(
                str(material.get("cleanedFileName") or payload.get("fileName") or material.get("materialName") or material.get("name") or f"{material_id}.bin"),
                f"{material_id}.bin",
            )
            target_path = self._unique_business_gap_path(work_dir, f"{index:02d}-{file_name}")
            minio_client.download_file(str(payload["bucket"]), str(payload["key"]), target_path)
            ref = {
                "materialId": material_id,
                "materialName": str(material.get("materialName") or material.get("name") or payload.get("fileName") or material_id),
                "folderPath": str(material.get("folderPath") or material.get("path") or ""),
                "materialTier": str(material.get("materialTier") or material.get("libraryScope") or ""),
                "sourceKind": source_kind,
                "selectedAt": created_at,
            }
            artifact_id = f"BART-{safe_filename(task_id, 'TASK')}-MAT-{existing_count + index}"
            artifact = {
                "artifactId": artifact_id,
                "artifactType": "selected_material",
                "fileName": target_path.name,
                "filePath": str(target_path),
                "sourceMode": "selected_from_business_material_library",
                "assemblyMode": self._assembly_mode_for_business_artifact(task, {**material, "fileName": target_path.name, "mimeType": payload.get("mimeType")}),
                "materialUsage": str(material.get("wikiUsageMode") or ""),
                "materialId": material_id,
                "materialName": ref["materialName"],
                "folderPath": ref["folderPath"],
                "materialTier": ref["materialTier"],
                "sourceKind": source_kind,
                "version": 1,
                "previewable": True,
                "confirmed": True,
                "reviewStatus": "approved",
                "confirmedAt": created_at,
                "selectedAt": created_at,
                "operator": str(data.get("operator") or "当前用户"),
                "mimeType": str(payload.get("mimeType") or mimetypes.guess_type(target_path.name)[0] or "application/octet-stream"),
                "wikiCardId": str(material.get("wikiCardId") or ""),
                "wikiUsageMode": str(material.get("wikiUsageMode") or ""),
                "evidenceSegmentId": str(material.get("evidenceSegmentId") or ""),
                "evidenceSegmentTitle": str(material.get("evidenceSegmentTitle") or ""),
                "evidenceSegmentType": str(material.get("evidenceSegmentType") or ""),
                "evidenceSourcePages": str(material.get("evidenceSourcePages") or ""),
                "evidenceSummary": str(material.get("evidenceSummary") or ""),
                "wikiEvidence": copy.deepcopy(material.get("wikiEvidence") or {}),
            }
            if not artifact["materialUsage"]:
                artifact["materialUsage"] = self._business_material_usage_for_assembly_mode(str(artifact.get("assemblyMode") or ""))
            material_refs.append(ref)
            artifacts.append(artifact)

        if not artifacts:
            raise ValueError("至少需要选择一份有效素材。")
        existing_refs = task.get("selectedMaterialRefs") if isinstance(task.get("selectedMaterialRefs"), list) else []
        task["selectedMaterialRefs"] = self._merge_business_material_refs(existing_refs, material_refs)
        task.setdefault("resolvedArtifacts", []).extend(artifacts)
        self._apply_business_task_artifact_intent(task, artifacts)
        self._record_business_material_feedback(
            business_gap_state,
            task,
            artifacts,
            operator=str(data.get("operator") or "当前用户"),
            selected_at=created_at,
        )
        task["decision"] = "ready"
        task["status"] = "ready"
        task["updatedAt"] = created_at
        task["resolvedAt"] = created_at
        task["resolvedSource"] = artifacts[0]["fileName"]
        task["riskFlags"] = [
            flag
            for flag in (task.get("riskFlags") if isinstance(task.get("riskFlags"), list) else [])
            if flag not in {"missing_material"}
        ]
        self._finalize_business_gap_plan_update(project, business_gap_state, plan, updated_at=created_at)
        self._refresh_business_gap_urls_for_result(
            project_id,
            plan,
            browser_base_url=browser_base_url,
            onlyoffice_base_url=onlyoffice_base_url,
        )
        return {
            "task": copy.deepcopy(task),
            "artifact": copy.deepcopy(artifacts[0]),
            "artifacts": copy.deepcopy(artifacts),
            "selectedMaterialRefs": copy.deepcopy(task["selectedMaterialRefs"]),
            "plan": copy.deepcopy(plan),
            "integrity": copy.deepcopy(business_gap_state["integrity"]),
        }

    def select_business_gap_template(
        self,
        project_id: str,
        task_id: str,
        data: dict[str, Any],
        *,
        browser_base_url: str = "",
        onlyoffice_base_url: str = "",
    ) -> dict[str, Any]:
        project = self._require(project_id)
        if normalize_bid_type(str(project.get("bidType") or "")) != "商务标":
            raise ValueError("商务标模板选择仅支持商务标项目。")
        business_gap_state = self._ensure_business_gap_state(project)
        if business_gap_state["recognitionStatus"] != "completed":
            raise ValueError("请先生成商务标缺口计划。")
        plan = business_gap_state.get("plan") if isinstance(business_gap_state.get("plan"), dict) else {}
        task = self._find_business_gap_task(plan, task_id)
        selected = self._selected_business_template_entry(task, data)
        source_path = Path(str(selected.get("filePath") or selected.get("path") or "")).expanduser()
        if not source_path.exists() or not source_path.is_file():
            raise ValueError("候选模板文件不存在，请重新生成商务标缺口计划或重新上传模板。")
        if source_path.suffix.lower() != ".docx":
            raise ValueError("模板填充 Word 目前仅支持 DOCX 模板。")

        work_dir = business_workspace_dir(project_id) / "gaps" / "selected-templates" / safe_filename(task_id, "task")
        work_dir.mkdir(parents=True, exist_ok=True)
        created_at = now_iso()
        existing_count = len(task.get("resolvedArtifacts") if isinstance(task.get("resolvedArtifacts"), list) else [])
        target_path = self._unique_business_gap_path(work_dir, source_path.name)
        shutil.copy2(source_path, target_path)
        artifact = {
            "artifactId": f"BART-{safe_filename(task_id, 'TASK')}-TPL-{existing_count + 1}",
            "artifactType": "selected_bid_template",
            "fileName": target_path.name,
            "filePath": str(target_path),
            "sourceMode": str(selected.get("sourceMode") or "selected_from_bid_template"),
            "templateId": str(selected.get("templateId") or ""),
            "templateName": str(selected.get("templateName") or selected.get("fileName") or target_path.name),
            "templateScope": str(selected.get("templateScope") or ""),
            "sourceLabel": str(selected.get("sourceLabel") or ""),
            "assemblyMode": "template_fill_docx",
            "materialUsage": "fill_template",
            "version": 1,
            "previewable": True,
            "confirmed": True,
            "reviewStatus": "approved",
            "confirmedAt": created_at,
            "selectedAt": created_at,
            "operator": str(data.get("operator") or "当前用户"),
            "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
        task.setdefault("resolvedArtifacts", []).append(artifact)
        self._apply_business_task_artifact_intent(task, [artifact])
        self._recompute_business_gap_task_after_artifact_change(task)
        task["updatedAt"] = created_at
        task["resolvedAt"] = created_at
        task["resolvedSource"] = artifact["fileName"]
        self._finalize_business_gap_plan_update(project, business_gap_state, plan, updated_at=created_at)
        self._refresh_business_gap_urls_for_result(
            project_id,
            plan,
            browser_base_url=browser_base_url,
            onlyoffice_base_url=onlyoffice_base_url,
        )
        return {
            "message": "已选择模板并快照到商务 S3 工作区。",
            "task": copy.deepcopy(task),
            "artifact": copy.deepcopy(artifact),
            "plan": copy.deepcopy(plan),
            "integrity": copy.deepcopy(business_gap_state["integrity"]),
        }

    def run_business_gap_ai_draft(
        self,
        project_id: str,
        task_id: str,
        data: dict[str, Any],
        *,
        browser_base_url: str = "",
        onlyoffice_base_url: str = "",
    ) -> dict[str, Any]:
        project = self._require(project_id)
        if normalize_bid_type(str(project.get("bidType") or "")) != "商务标":
            raise ValueError("商务标 AI 起草仅支持商务标项目。")
        business_gap_state = self._ensure_business_gap_state(project)
        if business_gap_state["recognitionStatus"] != "completed":
            raise ValueError("请先生成商务标缺口计划。")
        plan = business_gap_state.get("plan") if isinstance(business_gap_state.get("plan"), dict) else {}
        task = self._find_business_gap_task(plan, task_id)
        if not self._business_task_can_ai_draft(task):
            raise ValueError("该任务不适合 AI 起草，请选择候选素材或人工上传补料。")

        fact_table = business_gap_state.get("projectFactTable")
        if not isinstance(fact_table, dict) or fact_table.get("schemaVersion") != PROJECT_FACT_TABLE_SCHEMA_VERSION:
            fact_table = self._build_project_fact_table(project, business_gap_state)
            business_gap_state["projectFactTable"] = fact_table
        facts = self._fact_table_value_map(fact_table)
        created_at = now_iso()
        work_dir = business_workspace_dir(project_id) / "gaps" / "ai-drafts" / safe_filename(task_id, "task")
        work_dir.mkdir(parents=True, exist_ok=True)
        existing_count = len(task.get("resolvedArtifacts") if isinstance(task.get("resolvedArtifacts"), list) else [])
        title = str(task.get("title") or "商务响应文件")
        output_path = self._unique_business_gap_path(work_dir, f"{safe_filename(title, '商务响应文件')}-AI起草.docx")
        self._write_business_ai_draft_docx(output_path, project, task, facts, data)

        artifact_id = f"BART-{safe_filename(task_id, 'TASK')}-AI-{existing_count + 1}"
        artifact = {
            "artifactId": artifact_id,
            "artifactType": "ai_draft",
            "fileName": output_path.name,
            "filePath": str(output_path),
            "sourceMode": "generated_by_business_s3_ai_draft",
            "version": 1,
            "previewable": True,
            "confirmed": True,
            "reviewStatus": "approved",
            "confirmedAt": created_at,
            "createdAt": created_at,
            "operator": str(data.get("operator") or "当前用户"),
            "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "assemblyMode": "ai_draft",
            "materialUsage": "generate_draft",
            "draftMode": "controlled_business_ai_draft",
            "draftPolicy": "only_for_s3_ai_draft_task",
            "factTableStatus": str(fact_table.get("status") or "draft"),
        }
        task.setdefault("resolvedArtifacts", []).append(artifact)
        self._apply_business_task_artifact_intent(task, [artifact])
        task["decision"] = "ready"
        task["status"] = "ready"
        task["assigneeMode"] = "ai_draft"
        task["updatedAt"] = created_at
        task["resolvedAt"] = created_at
        task["resolvedSource"] = artifact["fileName"]
        task["riskFlags"] = [
            flag
            for flag in (task.get("riskFlags") if isinstance(task.get("riskFlags"), list) else [])
            if flag not in {"missing_material", "manual_upload_required", "ai_draft_required"}
        ]
        if str(fact_table.get("status") or "") != "confirmed":
            flags = task.setdefault("riskFlags", [])
            if isinstance(flags, list) and "fact_table_unconfirmed" not in flags:
                flags.append("fact_table_unconfirmed")
        self._finalize_business_gap_plan_update(project, business_gap_state, plan, updated_at=created_at)
        self._refresh_business_gap_urls_for_result(
            project_id,
            plan,
            browser_base_url=browser_base_url,
            onlyoffice_base_url=onlyoffice_base_url,
        )
        return {
            "message": "已生成商务响应文件 AI 起草稿。",
            "task": copy.deepcopy(task),
            "artifact": copy.deepcopy(artifact),
            "plan": copy.deepcopy(plan),
            "integrity": copy.deepcopy(business_gap_state["integrity"]),
            "projectFactTable": copy.deepcopy(fact_table),
        }

    def sync_business_gap_artifact_to_material_library(
        self,
        project_id: str,
        task_id: str,
        data: dict[str, Any],
        *,
        browser_base_url: str = "",
        onlyoffice_base_url: str = "",
    ) -> dict[str, Any]:
        project = self._require(project_id)
        if normalize_bid_type(str(project.get("bidType") or "")) != "商务标":
            raise ValueError("商务标缺口处理仅支持商务标项目。")
        business_gap_state = self._ensure_business_gap_state(project)
        if business_gap_state["recognitionStatus"] != "completed":
            raise ValueError("请先生成商务标缺口计划。")
        plan = business_gap_state.get("plan") if isinstance(business_gap_state.get("plan"), dict) else {}
        task = self._find_business_gap_task(plan, task_id)
        artifact_id = str(data.get("artifactId") or "").strip()
        if not artifact_id:
            raise ValueError("artifactId 不能为空。")
        artifact = self._find_business_gap_artifact_in_task(task, artifact_id)
        if str(artifact.get("sourceMode") or "") != "uploaded_in_business_s3":
            raise ValueError("仅支持将人工上传的 S3 补料同步到商务标项目素材库。")
        source_path = Path(str(artifact.get("filePath") or ""))
        if not source_path.exists() or not source_path.is_file():
            raise ValueError("补料文件不存在，无法同步。")

        target_path = str(data.get("targetPath") or "").strip().strip("/")
        if not target_path:
            target_path = self._default_business_gap_material_target_path(project_id, task)
        if not target_path.startswith(f"商务标/项目素材/{project_id}/"):
            raise ValueError("S3 补料默认只允许同步到当前商务标项目素材库。")

        raw_bytes = source_path.read_bytes()
        upload_result = _run_async_material_upload(
            target_path=target_path,
            project_id=project_id,
            project_code=str(project.get("projectCode") or project_id),
            project_name=str(project.get("name") or project_id),
            bid_type="商务标",
            material_tier="project",
            customer_id=str(project.get("customerId") or ""),
            customer_name=str(project.get("customerCanonicalName") or project.get("customerName") or project.get("owner") or ""),
            on_conflict="version",
            files=[
                {
                    "name": str(artifact.get("fileName") or source_path.name),
                    "mimeType": str(artifact.get("mimeType") or mimetypes.guess_type(source_path.name)[0] or "application/octet-stream"),
                    "data": base64.b64encode(raw_bytes).decode("ascii"),
                }
            ],
        )
        synced_items = [item for item in upload_result.get("items") or [] if isinstance(item, dict)]
        if not synced_items:
            raise ValueError("补料同步失败：素材库未返回上传文件。")
        synced = synced_items[0]
        now = now_iso()
        artifact["materialSyncStatus"] = "synced_to_project_material"
        artifact["materialSyncedAt"] = now
        artifact["materialTargetPath"] = target_path
        artifact["materialId"] = str(synced.get("id") or "")
        artifact["materialName"] = str(synced.get("name") or artifact.get("fileName") or "")
        artifact["materialTier"] = str(synced.get("materialTier") or "project")
        artifact["folderPath"] = str(synced.get("folderPath") or target_path)
        artifact["cleanStatus"] = str(synced.get("cleanStatus") or "")
        artifact["cleanMessage"] = str(synced.get("cleanMessage") or "")
        artifact["wikiSyncStatus"] = "wiki_rebuild_required"
        task["updatedAt"] = now
        add_note = f"补料已同步到项目素材库：{target_path}"
        task["notes"] = "\n".join(part for part in [str(task.get("notes") or "").strip(), add_note] if part)
        business_gap_state["wikiRebuildRequired"] = True
        business_gap_state["lastMaterialSyncAt"] = now
        business_gap_state["lastMaterialSyncTargetPath"] = target_path
        self._finalize_business_gap_plan_update(project, business_gap_state, plan, updated_at=now)
        self._refresh_business_gap_urls_for_result(
            project_id,
            plan,
            browser_base_url=browser_base_url,
            onlyoffice_base_url=onlyoffice_base_url,
        )
        return {
            "message": "补料已同步到商务标项目素材库，商务 Wiki 需要重新生成/更新。",
            "task": copy.deepcopy(task),
            "artifact": copy.deepcopy(artifact),
            "material": copy.deepcopy(synced),
            "materialUpload": copy.deepcopy(upload_result),
            "wikiRebuildRequired": True,
            "plan": copy.deepcopy(plan),
            "integrity": copy.deepcopy(business_gap_state["integrity"]),
        }

    def get_business_gap_artifact(self, project_id: str, artifact_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        business_gap_state = self._ensure_business_gap_state(project)
        plan = business_gap_state.get("plan") if isinstance(business_gap_state.get("plan"), dict) else {}
        for task in plan.get("tasks") or []:
            if not isinstance(task, dict):
                continue
            for artifact in task.get("resolvedArtifacts") or []:
                if isinstance(artifact, dict) and str(artifact.get("artifactId") or "") == artifact_id:
                    return copy.deepcopy(artifact)
        raise KeyError(artifact_id)

    def get_gap_artifact(self, project_id: str, artifact_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        gap_state = self._ensure_gap_state(project)
        plan = gap_state.get("plan") if isinstance(gap_state.get("plan"), dict) else {}
        for item in plan.get("items") or []:
            for artifact in item.get("resolvedArtifacts") or []:
                if str(artifact.get("id") or "") == artifact_id:
                    return copy.deepcopy(artifact)
        raise KeyError(artifact_id)

    def list_gap_submissions(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        gap_state = self._ensure_gap_state(project)
        submissions = copy.deepcopy(gap_state["submissions"])
        return {"items": submissions, "total": len(submissions)}

    def submit_gap_material(self, project_id: str, data: dict[str, Any]) -> dict[str, Any]:
        project = self._require(project_id)
        gap_state = self._ensure_gap_state(project)
        if gap_state["recognitionStatus"] != "completed":
            raise ValueError("请先完成缺口识别。")

        missing_id = str(data.get("missingId") or "").strip()
        files = list(data.get("files") or [])
        if not missing_id:
            raise ValueError("missingId 不能为空。")
        if not files:
            raise ValueError("至少需要提交一个文件。")

        item = self._find_gap_item(gap_state, missing_id)
        receipts: list[dict[str, Any]] = []
        timestamp = now_iso()
        for index, file in enumerate(files, start=1):
            file_name = str(file.get("name") or f"{missing_id}-{index}.docx").strip() or f"{missing_id}-{index}.docx"
            receipt = {
                "receiptId": f"mr-{project_id}-{len(gap_state['submissions']) + index}",
                "projectId": project_id,
                "missingId": missing_id,
                "fileId": f"raw-{project_id}-{len(gap_state['submissions']) + index}",
                "fileName": file_name,
                "storedPath": f"技术标/项目素材/{project_id}",
                "action": "upload",
                "operator": str(data.get('operator') or "当前用户"),
                "submittedAt": timestamp,
                "traceId": f"mock-{project_id}-{len(gap_state['submissions']) + index}",
                "auditId": f"audit-{project_id}-{len(gap_state['submissions']) + index}",
            }
            receipts.append(receipt)

        gap_state["submissions"] = receipts + list(gap_state["submissions"])
        item["latestUploadAt"] = timestamp
        item["latestSubmissionId"] = receipts[0]["receiptId"]
        if item["status"] != "resolved":
            item["status"] = "checking"
        plan_item = self._find_gap_plan_item(gap_state, missing_id)
        if plan_item is not None:
            plan_item["status"] = "filling"
            plan_item["latestUploadAt"] = timestamp
            plan_item["latestSubmissionId"] = receipts[0]["receiptId"]
            plan_item.setdefault("resolvedArtifacts", []).extend(
                {
                    "id": receipt["receiptId"],
                    "source": "manual_upload",
                    "fileName": receipt["fileName"],
                    "path": receipt["storedPath"],
                    "createdAt": receipt["submittedAt"],
                    "s7Ready": False,
                }
                for receipt in receipts
            )
            if isinstance(gap_state.get("plan"), dict):
                gap_state["plan"]["summary"] = summarize_gap_plan(gap_state["plan"])
                gap_state["integrity"] = check_gap_integrity(gap_state["plan"])
                gap_state["plan"]["integrity"] = gap_state["integrity"]
        gap_state["submittedForReview"] = False
        gap_state["reviewConfirmed"] = False
        gap_state["reviewedAt"] = ""
        project["review_document_state"] = self._default_review_document_state(project)
        project["updatedAt"] = timestamp
        self._persist_project(project)
        return {
            "message": f"补料提交成功，共 {len(receipts)} 个文件。",
            "item": copy.deepcopy(item),
            "receipts": receipts,
            "payload": self.get_gap_filling(project_id),
            "traceId": receipts[0]["traceId"],
        }

    def update_gap_item(self, project_id: str, gap_id: str, data: dict[str, Any]) -> dict[str, Any]:
        project = self._require(project_id)
        gap_state = self._ensure_gap_state(project)
        if gap_state["recognitionStatus"] != "completed":
            raise ValueError("请先完成缺口识别。")

        item = self._find_gap_item(gap_state, gap_id)
        plan_item = self._find_gap_plan_item(gap_state, gap_id)
        action = str(data.get("action") or data.get("status") or "").strip()
        if action in {"skip", "skipped"}:
            item["status"] = "skipped"
            item["skipReason"] = str(data.get("reason") or item.get("skipReason") or "未填写原因")
            item["resolvedSource"] = ""
            item["resolvedAt"] = ""
            if plan_item is not None:
                plan_item["status"] = "ignored"
                plan_item["skipReason"] = item["skipReason"]
                plan_item["reviewNotes"] = list(plan_item.get("reviewNotes") or []) + [f"人工忽略：{item['skipReason']}"]
        elif action in {"resolve", "resolved"}:
            source = data.get("source") or {}
            if isinstance(source, dict):
                source_name = str(source.get("name") or "")
            else:
                source_name = str(source)
            item["status"] = "resolved"
            item["resolvedSource"] = source_name.strip() or str(data.get("resolvedSource") or item.get("resolvedSource") or "已补录")
            item["skipReason"] = ""
            item["resolvedAt"] = now_iso()
            if plan_item is not None:
                plan_item["status"] = "resolved"
                plan_item["resolvedSource"] = item["resolvedSource"]
                plan_item["resolvedAt"] = item["resolvedAt"]
                plan_item.setdefault("resolvedArtifacts", []).append(
                    {
                        "id": f"ART-{gap_id}-{len(plan_item.get('resolvedArtifacts') or []) + 1}",
                        "source": "manual",
                        "fileName": item["resolvedSource"],
                        "createdAt": item["resolvedAt"],
                        "s7Ready": True,
                    }
                )
        elif action in {"checking", "pending"}:
            item["status"] = action
            if plan_item is not None:
                plan_item["status"] = "filling" if action == "checking" else "needs_input"
        else:
            raise ValueError("不支持的缺口状态更新。")

        gap_state["submittedForReview"] = False
        gap_state["reviewConfirmed"] = False
        gap_state["reviewedAt"] = ""
        if isinstance(gap_state.get("plan"), dict):
            gap_state["plan"]["summary"] = summarize_gap_plan(gap_state["plan"])
            gap_state["integrity"] = check_gap_integrity(gap_state["plan"])
            gap_state["plan"]["integrity"] = gap_state["integrity"]
        project["review_document_state"] = self._default_review_document_state(project)
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        return {
            "message": "缺口状态已更新",
            "item": copy.deepcopy(item),
            "payload": self.get_gap_filling(project_id),
        }

    def patch_missing_material(self, project_id: str, missing_id: str, data: dict[str, Any]) -> dict[str, Any]:
        return self.update_gap_item(project_id, missing_id, data)

    def submit_gap_review(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        gap_state = self._ensure_gap_state(project)
        if gap_state["recognitionStatus"] != "completed":
            raise ValueError("请先完成缺口识别后再提交确认。")

        integrity = check_gap_integrity(gap_state.get("plan") or {})
        gap_state["integrity"] = integrity
        if integrity["status"] != "passed":
            raise ValueError(f"仍有 {integrity['blockingCount']} 项缺口未解决，暂不可提交审核。")

        gap_state["submittedForReview"] = True
        gap_state["reviewConfirmed"] = False
        gap_state["reviewedAt"] = ""
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        return {
            "message": "缺口处理已通过完整性校验并提交审核。",
            "payload": self.get_gap_filling(project_id),
        }

    def get_review_items(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        gap_state = self._ensure_gap_state(project)
        return self._build_review_payload(project, gap_state)

    def prepare_review_document(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        gap_state = self._ensure_gap_state(project)
        if not gap_state["submittedForReview"]:
            raise ValueError("请先在缺口处理页提交确认后再生成预览文档。")

        pending = [item for item in gap_state["items"] if item["status"] not in {"resolved", "skipped"}]
        if pending:
            raise ValueError(f"仍有 {len(pending)} 项素材未处理，暂不可生成预览文档。")

        review_state = self._ensure_review_document_state(project)
        parsed_at = now_iso()
        review_state.update(
            {
                "parseStatus": "completed",
                "parsedAt": parsed_at,
                "sourceFileName": self._review_source_file_name(gap_state),
                "fileName": f"{project['name']}_缺口处理确认预览.docx",
                "fileType": "docx",
                "content": self._build_review_document_content(project, gap_state),
                "lastSavedAt": parsed_at,
                "version": 1,
                "documentKey": f"{project_id}-s6-v1",
            }
        )
        project["updatedAt"] = parsed_at
        self._persist_project(project)
        return {
            "message": "缺口处理确认预览已生成，可继续生成标书。",
            "payload": copy.deepcopy(review_state),
        }

    def get_review_document_state(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        return copy.deepcopy(self._ensure_review_document_state(project))

    def save_review_document_content(self, project_id: str, content: str) -> dict[str, Any]:
        project = self._require(project_id)
        review_state = self._ensure_review_document_state(project)
        next_version = int(review_state.get("version") or 1) + 1
        review_state.update(
            {
                "parseStatus": "completed",
                "content": content,
                "version": next_version,
                "lastSavedAt": now_iso(),
                "documentKey": f"{project_id}-s6-v{next_version}",
            }
        )
        project["updatedAt"] = review_state["lastSavedAt"]
        self._persist_project(project)
        return copy.deepcopy(review_state)

    def force_save_review_document(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        review_state = self._ensure_review_document_state(project)
        next_version = int(review_state.get("version") or 1) + 1
        review_state.update(
            {
                "parseStatus": "completed",
                "version": next_version,
                "lastSavedAt": now_iso(),
                "documentKey": f"{project_id}-s6-v{next_version}",
            }
        )
        project["updatedAt"] = review_state["lastSavedAt"]
        self._persist_project(project)
        return copy.deepcopy(review_state)

    def confirm_review(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        gap_state = self._ensure_gap_state(project)
        if not gap_state["submittedForReview"]:
            raise ValueError("请先在缺口处理页提交确认后再执行确认。")

        integrity = check_gap_integrity(gap_state.get("plan") or {})
        gap_state["integrity"] = integrity
        if integrity["status"] != "passed":
            raise ValueError(f"仍有 {integrity['blockingCount']} 项缺口未解决，暂不可确认审核。")

        gap_state["reviewConfirmed"] = True
        gap_state["reviewedAt"] = now_iso()
        project["updatedAt"] = gap_state["reviewedAt"]
        self._persist_project(project)
        return {
            "message": "缺口处理已确认，可进入标书生成。",
            "reviewStatus": "confirmed",
            "payload": self._build_review_payload(project, gap_state),
        }

    def get_fill_state(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        return copy.deepcopy(project["fill_state"])

    def start_fill_generation(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        document_label = fill_document_label(project)
        payload = {
            "status": "running",
            "percentage": 5,
            "filledAt": "",
            "runDurationSec": 0,
            "runDuration": "",
            "summary": f"已开始拼装{document_label}，正在准备 S2 目录、Wiki 与素材库。",
            "output": None,
            "sections": [],
            "opencodeOutput": build_directory_opencode_output(),
            "events": [
                build_directory_event(
                    f"已开始{document_label}拼装任务，正在准备 S2 目录、Wiki 与素材库。",
                    step="bootstrap",
                ),
            ],
            "tasks": [
                {"id": "task-1", "label": "准备 S2 目录、Wiki 与素材库", "status": "running"},
                {"id": "task-2", "label": fill_task_label(project), "status": "pending"},
                {"id": "task-3", "label": "写入并规范化 Word 正文", "status": "pending"},
            ],
        }
        project["fill_state"] = payload
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        return copy.deepcopy(payload)

    def update_fill_generation_state(
        self,
        project_id: str,
        *,
        percentage: int | None = None,
        summary: str | None = None,
        tasks: list[dict[str, Any]] | None = None,
        status: str | None = None,
        event_message: str | None = None,
        event_level: str = "info",
        event_step: str = "general",
        opencode_output: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        project = self._require(project_id)
        current = copy.deepcopy(project.get("fill_state") or default_fill_state(project))
        if percentage is not None:
            current["percentage"] = max(0, min(100, int(percentage)))
        if summary is not None:
            current["summary"] = summary
        if tasks is not None:
            current["tasks"] = copy.deepcopy(tasks)
        if status is not None:
            current["status"] = status
        if opencode_output is not None:
            merged_output = build_directory_opencode_output()
            merged_output.update(copy.deepcopy(current.get("opencodeOutput") or {}))
            merged_output.update(copy.deepcopy(opencode_output))
            merged_output["parts"] = copy.deepcopy(merged_output.get("parts") or [])[-20:]
            current["opencodeOutput"] = merged_output
        if event_message:
            events = list(current.get("events") or [])
            events.append(
                build_directory_event(
                    event_message,
                    level=event_level,
                    step=event_step,
                )
            )
            current["events"] = events[-20:]
        project["fill_state"] = current
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        return copy.deepcopy(current)

    def fail_fill_generation(
        self,
        project_id: str,
        message: str,
        tasks: list[dict[str, Any]] | None = None,
    ) -> dict[str, Any]:
        project = self._require(project_id)
        current = copy.deepcopy(project.get("fill_state") or default_fill_state(project))
        current["status"] = "failed"
        current["summary"] = message
        current_output = build_directory_opencode_output()
        current_output.update(copy.deepcopy(current.get("opencodeOutput") or {}))
        if (
            current_output.get("status") != "idle"
            or current_output.get("sessionId")
            or current_output.get("providerId")
            or current_output.get("modelId")
            or current_output.get("parts")
        ):
            current_output["status"] = "failed"
        current["opencodeOutput"] = current_output
        if tasks is not None:
            current["tasks"] = copy.deepcopy(tasks)
        events = list(current.get("events") or [])
        events.append(build_directory_event(message, level="error", step="failed"))
        current["events"] = events[-20:]
        project["fill_state"] = current
        project["updatedAt"] = now_iso()
        self._persist_project(project)
        return copy.deepcopy(current)

    def complete_fill_generation(self, project_id: str, data: dict[str, Any]) -> dict[str, Any]:
        project = self._require(project_id)
        filled_at = now_iso()
        sections = [
            {
                "nodeId": "OL-1",
                "title": "项目概况",
                "generationMode": "generated",
                "evidenceRefs": [],
                "riskFlags": [],
            },
            {
                "nodeId": "OL-2",
                "title": "企业业绩",
                "generationMode": "placeholder",
                "evidenceRefs": [],
                "riskFlags": ["FACT_REQUIRED"],
            },
        ]
        document_label = fill_document_label(project)
        project["fill_state"] = {
            "status": "completed",
            "percentage": 100,
            "filledAt": filled_at,
            "runDurationSec": 79,
            "runDuration": "1分19秒",
            "summary": f"{document_label}拼装完成。",
            "output": {
                "fileName": f"{project['name']}_正文.docx",
                "fileType": "docx",
                "size": "2.8 MB",
                "fileUrl": f"/api/projects/{project_id}/document/file",
            },
            "sections": sections,
            "opencodeOutput": build_directory_opencode_output(),
            "events": [
                build_directory_event(f"{document_label}拼装完成。", level="success", step="done", at=filled_at),
            ],
            "tasks": [
                {"id": "task-1", "label": "准备 S2 目录、Wiki 与素材库", "status": "done"},
                {"id": "task-2", "label": fill_task_label(project), "status": "done"},
                {"id": "task-3", "label": "写入并规范化 Word 正文", "status": "done"},
            ],
        }
        project["document_state"]["sourceFileName"] = f"{project['name']}_正文.docx"
        project["document_state"]["fileName"] = f"{project['name']}_正文.docx"
        project["document_state"]["fallback"]["content"] = (
            f"# {project['name']}\n\n## 项目概况\n本稿为 MVP 占位正文。\n\n## 企业业绩\n【此处待补充真实业绩信息】"
        )
        project["updatedAt"] = filled_at
        self._persist_project(project)
        return copy.deepcopy(project["fill_state"])

    def save_fill_generation_result(
        self,
        project_id: str,
        *,
        summary: str,
        sections: list[dict[str, Any]],
        content: str,
        filled_at: str,
        run_duration_sec: int,
        file_size_bytes: int,
        opencode_output: dict[str, Any] | None = None,
        file_name: str | None = None,
        coverage: dict[str, Any] | None = None,
        assembly: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        project = self._require(project_id)
        file_name = file_name or f"{project['name']}_正文.docx"
        current_state = copy.deepcopy(project.get("fill_state") or {})
        current_events = list(copy.deepcopy(current_state.get("events") or []))
        current_output = build_directory_opencode_output()
        current_output.update(copy.deepcopy(current_state.get("opencodeOutput") or {}))
        if opencode_output is not None:
            current_output.update(copy.deepcopy(opencode_output))
        current_output["parts"] = copy.deepcopy(current_output.get("parts") or [])[-20:]
        document_label = fill_document_label(project)
        current_events.append(
            build_directory_event(
                f"{document_label}拼装完成，已输出 {len(sections)} 个目录章节。",
                level="success",
                step="done",
                at=filled_at,
            )
        )
        project["fill_state"] = {
            "status": "completed",
            "percentage": 100,
            "filledAt": filled_at,
            "runDurationSec": int(run_duration_sec),
            "runDuration": self._format_duration(run_duration_sec),
            "summary": summary,
            "output": {
                "fileName": file_name,
                "fileType": "docx",
                "size": self._format_file_size(file_size_bytes),
                "fileUrl": f"/api/projects/{project_id}/document/file",
            },
            "sections": copy.deepcopy(sections),
            "coverage": copy.deepcopy(coverage or {}),
            "assembly": copy.deepcopy(assembly or {}),
            "opencodeOutput": current_output,
            "events": current_events[-20:],
            "tasks": [
                {"id": "task-1", "label": "准备 S2 目录、Wiki 与素材库", "status": "done"},
                {"id": "task-2", "label": fill_task_label(project), "status": "done"},
                {"id": "task-3", "label": "写入并规范化 Word 正文", "status": "done"},
            ],
        }

        document_state = project["document_state"]
        next_version = int(document_state.get("version") or 1)
        if document_state.get("lastSavedAt"):
            next_version += 1
        document_state["status"] = "ready"
        document_state["sourceFileName"] = file_name
        document_state["fileName"] = file_name
        document_state["fileType"] = "docx"
        document_state["version"] = next_version
        document_state["lastSavedAt"] = filled_at
        document_state["fallback"]["content"] = content
        document_state["onlyoffice"]["title"] = file_name
        document_state["onlyoffice"]["documentKey"] = f"{project_id}-v{next_version}"

        project["updatedAt"] = filled_at
        self._persist_project(project)
        return copy.deepcopy(project["fill_state"])

    def get_coverage(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        assembly_coverage = (project.get("fill_state") or {}).get("coverage")
        if isinstance(assembly_coverage, dict) and assembly_coverage:
            return copy.deepcopy(assembly_coverage)

        outline_nodes = list((project.get("outline_state") or {}).get("nodes") or [])
        if not outline_nodes:
            return {
                "percentage": 100,
                "fullCover": 0,
                "partialCover": 0,
                "noCover": 0,
                "tree": [],
                "partialItems": [],
                "noCoverItems": [],
            }

        sections = list((project.get("fill_state") or {}).get("sections") or [])
        section_status_map = {
            str(section.get("nodeId") or "").strip(): self._mode_to_coverage_status(section.get("generationMode"))
            for section in sections
            if str(section.get("nodeId") or "").strip()
        }

        tree = self._build_coverage_tree(outline_nodes, section_status_map)
        partial_items, no_cover_items = self._build_coverage_issue_lists(tree)
        full_cover, partial_cover, no_cover = self._summarize_coverage(tree)
        total = full_cover + partial_cover + no_cover
        percentage = 100 if total == 0 else round(((full_cover * 1.0) + (partial_cover * 0.5)) / total * 100)
        return {
            "percentage": percentage,
            "fullCover": full_cover,
            "partialCover": partial_cover,
            "noCover": no_cover,
            "tree": tree,
            "partialItems": partial_items,
            "noCoverItems": no_cover_items,
        }

    def get_document_state(self, project_id: str) -> dict[str, Any]:
        return copy.deepcopy(self._require(project_id)["document_state"])

    def save_document_content(self, project_id: str, content: str) -> dict[str, Any]:
        project = self._require(project_id)
        state = project["document_state"]
        next_version = int(state["version"] or 1) + 1
        state["version"] = next_version
        state["lastSavedAt"] = now_iso()
        state["onlyoffice"]["documentKey"] = f"{project_id}-v{next_version}"
        state["fallback"]["content"] = content
        project["updatedAt"] = state["lastSavedAt"]
        self._persist_project(project)
        return copy.deepcopy(state)

    def force_save_document(self, project_id: str) -> dict[str, Any]:
        project = self._require(project_id)
        state = project["document_state"]
        next_version = int(state["version"] or 1) + 1
        state["version"] = next_version
        state["lastSavedAt"] = now_iso()
        state["onlyoffice"]["documentKey"] = f"{project_id}-v{next_version}"
        project["updatedAt"] = state["lastSavedAt"]
        self._persist_project(project)
        return copy.deepcopy(state)

    def get_final_document(self, project_id: str) -> dict[str, Any]:
        state = self._require(project_id)["document_state"]
        return {
            "ready": True,
            "fileName": state["fileName"],
            "fileType": state["fileType"],
            "fileUrl": f"/api/projects/{project_id}/final-document/file",
            "lastSavedAt": state["lastSavedAt"],
            "version": state["version"],
        }

    @staticmethod
    def _fact_label_key(label: Any) -> str:
        return re.sub(r"\s+", "", AppStore._canonical_fact_label(label)).lower()

    @staticmethod
    def _canonical_fact_label(label: Any) -> str:
        raw = str(label or "").strip()
        if not raw:
            return ""
        text = re.sub(r"\s+", "", raw)
        text = re.sub(r"[（(]\s*(?:MW|kW|m|m2/kW|m²/kW|%|h|MWh/y|MWh/a|台)\s*[）)]", "", text, flags=re.I)
        text = text.strip("：:；;，,、")
        aliases = {
            "方案": "投标方案",
            "项目方案": "投标方案",
            "机型": "投标方案",
            "建设容量": "总装机容量",
            "标段规模": "总装机容量",
            "机组数量": "机组台数",
            "风机数量": "机组台数",
            "台数": "机组台数",
            "总容量": "总装机容量",
            "容量": "总装机容量",
            "单机容量": "单机容量",
            "机组额定功率": "单机容量",
            "项目编号": "招标编号",
            "招标文件编号": "招标编号",
            "项目单位": "招标人",
            "建设单位": "招标人",
            "业主": "招标人",
            "交货期": "交货周期",
            "质量保证期": "质保期",
            "投标截止时间": "投标截止日期",
            "轮毂中心高度": "轮毂高度",
            "轮毂高度": "轮毂高度",
            "风轮直径": "叶轮直径",
            "叶轮直径": "叶轮直径",
            "发电小时数承诺": "保证有效小时数",
            "保证有效小时": "保证有效小时数",
            "风电机组设备年平均可利用率保证值": "全场可利用率",
            "适用等级": "安全等级",
        }
        if text in aliases:
            return aliases[text]
        if "总装机容量" in text or text.startswith("总容量"):
            return "总装机容量"
        if (
            "年平均风速" in text
            or "代表年风速" in text
            or ("平均风速" in text and ("机位" in text or "尾流" in text or "轮毂" in text))
        ):
            return "年平均风速"
        if "轮毂" in text and "高度" in text:
            return "轮毂高度"
        if "叶轮直径" in text or "风轮直径" in text:
            return "叶轮直径"
        if ("机组" in text or "风机" in text) and ("台数" in text or "数量" in text):
            return "机组台数"
        if "单机容量" in text or "额定功率" in text:
            return "单机容量"
        if "安全等级" in text or ("安全" in text and "等级" in text):
            return "安全等级"
        if "设计寿命" in text:
            return "设计寿命"
        if "单位千瓦扫风面积" in text:
            return "单位千瓦扫风面积"
        if "空气密度" in text and not re.search(r"参数|系数", text):
            return "空气密度"
        if "湍流强度" in text:
            return "湍流强度"
        if "极端风速" in text or "极大风速" in text:
            return "极端风速"
        if "风剪切" in text or "风切变" in text or "风剪切指数" in text:
            return "风剪切"
        if "功率曲线" in text and ("保证" in text or "保证率" in text):
            return "功率曲线保证率"
        if "单台" in text and "可利用率" in text:
            return "单台可利用率"
        if ("全场" in text or "风电场" in text or "年平均" in text) and "可利用率" in text:
            return "全场可利用率"
        if "发电量" in text and ("保证" in text or "承诺" in text):
            return "保证发电量"
        if "有效小时" in text or "发电小时" in text or "等效利用小时" in text:
            return "保证有效小时数"
        return text

    @staticmethod
    def _empty_project_fact_table(project_id: str) -> dict[str, Any]:
        return {
            "schemaVersion": PROJECT_FACT_TABLE_SCHEMA_VERSION,
            "projectId": project_id,
            "status": "empty",
            "builtAt": "",
            "updatedAt": "",
            "confirmedAt": "",
            "confirmedBy": "",
            "fields": [],
            "summary": {
                "totalCount": 0,
                "requiredCount": 0,
                "confirmedCount": 0,
                "candidateCount": 0,
                "missingCount": 0,
                "conflictCount": 0,
            },
        }

    @staticmethod
    def _summarize_project_fact_fields(fields: list[dict[str, Any]]) -> dict[str, int]:
        return {
            "totalCount": len(fields),
            "requiredCount": sum(1 for field in fields if field.get("required", True)),
            "confirmedCount": sum(1 for field in fields if str(field.get("status") or "") == "confirmed"),
            "candidateCount": sum(1 for field in fields if str(field.get("status") or "") == "candidate"),
            "missingCount": sum(1 for field in fields if str(field.get("status") or "") == "missing"),
            "conflictCount": sum(1 for field in fields if str(field.get("status") or "") == "conflict"),
        }

    @staticmethod
    def _fact_source_ref_priority(ref: dict[str, Any]) -> int:
        source_type = str(ref.get("type") or "").strip()
        if source_type in {"project", "projectIdentity", "projectTurbineModel", "derived"}:
            return 320
        if source_type in {"materialFact", "derivedMaterialFact"}:
            tier = str(ref.get("materialTier") or "").strip() or "standard"
            return FACT_MATERIAL_SOURCE_PRIORITIES.get(tier, 50)
        return 0

    @classmethod
    def _normalize_fact_source_refs(cls, refs: Any) -> list[dict[str, Any]]:
        normalized: list[tuple[int, int, dict[str, Any]]] = []
        seen: set[str] = set()
        for index, ref in enumerate(refs if isinstance(refs, list) else []):
            if not isinstance(ref, dict):
                continue
            item = copy.deepcopy(ref)
            key = json.dumps(item, ensure_ascii=False, sort_keys=True)
            if key in seen:
                continue
            seen.add(key)
            normalized.append((cls._fact_source_ref_priority(item), index, item))
        normalized.sort(key=lambda item: (-item[0], item[1]))
        return [item for _, _, item in normalized]

    @staticmethod
    def _normalize_project_fact_field(
        field: dict[str, Any],
        *,
        index: int,
        confirm: bool,
        operator: str,
        saved_at: str,
    ) -> dict[str, Any]:
        value = str(field.get("value") or "").strip()
        status = str(field.get("status") or ("candidate" if value else "missing")).strip()
        if confirm:
            status = "confirmed" if value else "missing"
        source_refs = AppStore._normalize_fact_source_refs(field.get("sourceRefs"))
        source_priority = int(field.get("sourcePriority") or 0)
        if source_refs:
            source_priority = max(source_priority, AppStore._fact_source_ref_priority(source_refs[0]))
        normalized = {
            "id": str(field.get("id") or f"FACT-{index:04d}"),
            "key": str(field.get("key") or AppStore._fact_label_key(field.get("label")) or f"fact-{index}"),
            "label": str(field.get("label") or ""),
            "category": str(field.get("category") or "项目事实"),
            "value": value,
            "unit": str(field.get("unit") or ""),
            "required": bool(field.get("required", True)),
            "status": status if status in {"candidate", "confirmed", "missing", "conflict"} else "candidate",
            "confidence": float(field.get("confidence") or 0),
            "sourcePriority": source_priority,
            "sourceRefs": source_refs,
            "alternatives": copy.deepcopy(field.get("alternatives") if isinstance(field.get("alternatives"), list) else []),
            "notes": str(field.get("notes") or ""),
            "updatedAt": saved_at,
            "updatedBy": operator,
        }
        if normalized["status"] == "confirmed":
            normalized["confirmedAt"] = saved_at
            normalized["confirmedBy"] = operator
        else:
            normalized["confirmedAt"] = str(field.get("confirmedAt") or "")
            normalized["confirmedBy"] = str(field.get("confirmedBy") or "")
        return normalized

    def _build_project_fact_table(self, project: dict[str, Any], gap_state: dict[str, Any]) -> dict[str, Any]:
        built_at = now_iso()
        existing_table = gap_state.get("projectFactTable") if isinstance(gap_state.get("projectFactTable"), dict) else {}
        existing_by_key = {
            self._fact_label_key(field.get("label")): field
            for field in (existing_table.get("fields") if isinstance(existing_table.get("fields"), list) else [])
            if isinstance(field, dict) and self._fact_label_key(field.get("label"))
        }
        fields_by_key: dict[str, dict[str, Any]] = {}

        def is_material_fact_ref(ref: dict[str, Any]) -> bool:
            return str(ref.get("type") or "") in {"materialFact", "derivedMaterialFact"}

        def blank_source_paths() -> set[str]:
            paths: set[str] = set()
            plan = gap_state.get("plan") if isinstance(gap_state.get("plan"), dict) else {}
            for item in plan.get("items") or []:
                if not isinstance(item, dict):
                    continue
                for task in item.get("fillTasks") or []:
                    if not isinstance(task, dict):
                        continue
                    blank = task.get("blankSource") if isinstance(task.get("blankSource"), dict) else {}
                    for key in ("docxPath", "path", "workspacePath"):
                        value = str(blank.get(key) or "").strip()
                        if value:
                            paths.add(str(Path(value).resolve()))
            return paths

        def add_candidate(
            label: str,
            value: Any,
            *,
            category: str,
            source_ref: dict[str, Any],
            confidence: float = 0.8,
            required: bool = True,
            unit: str = "",
            source_priority: int = 0,
        ) -> None:
            label_text = self._canonical_fact_label(label)
            if not label_text:
                return
            key = self._fact_label_key(label_text)
            value_text = str(value or "").strip()
            existing = existing_by_key.get(key)
            preserve_existing = bool(
                existing
                and str(existing.get("status") or "") == "confirmed"
                and str(existing.get("value") or "").strip()
            )
            if preserve_existing:
                value_text = str(existing.get("value") or "").strip()
            field = fields_by_key.get(key)
            incoming_priority = int(source_priority or 0)
            if not field:
                field = {
                    "id": str((existing or {}).get("id") or f"FACT-{len(fields_by_key) + 1:04d}"),
                    "key": key,
                    "label": label_text,
                    "category": category,
                    "value": value_text,
                    "unit": str(((existing or {}).get("unit") if preserve_existing else unit) or ""),
                    "required": bool((existing or {}).get("required", required)),
                    "status": "confirmed" if preserve_existing else ("candidate" if value_text else "missing"),
                    "confidence": float(((existing or {}).get("confidence") if preserve_existing else None) or (confidence if value_text else 0) or 0),
                    "sourcePriority": int((existing or {}).get("sourcePriority") if preserve_existing else (incoming_priority if value_text else 0)),
                    "sourceRefs": [],
                    "alternatives": copy.deepcopy(
                        (existing or {}).get("alternatives")
                        if preserve_existing and isinstance((existing or {}).get("alternatives"), list)
                        else []
                    ),
                    "notes": str((existing or {}).get("notes") if preserve_existing else ""),
                    "updatedAt": str((existing or {}).get("updatedAt") if preserve_existing else built_at),
                    "updatedBy": str((existing or {}).get("updatedBy") if preserve_existing else ""),
                    "confirmedAt": str((existing or {}).get("confirmedAt") if preserve_existing else ""),
                    "confirmedBy": str((existing or {}).get("confirmedBy") if preserve_existing else ""),
                }
                fields_by_key[key] = field
            elif value_text and field.get("value") and value_text != field["value"]:
                alternatives = field.setdefault("alternatives", [])
                existing_rank = (int(field.get("sourcePriority") or 0), float(field.get("confidence") or 0))
                incoming_rank = (incoming_priority, float(confidence or 0))
                if incoming_rank > existing_rank and str(field.get("status") or "") != "confirmed":
                    old_value = str(field.get("value") or "")
                    if old_value and old_value not in [str(item.get("value") or "") for item in alternatives if isinstance(item, dict)]:
                        alternatives.append({"value": old_value, "source": (field.get("sourceRefs") or [{}])[0]})
                    field["value"] = value_text
                    field["unit"] = str(unit or field.get("unit") or "")
                    field["category"] = category
                    field["status"] = "candidate"
                    field["confidence"] = float(confidence or 0)
                    field["sourcePriority"] = incoming_priority
                    if source_ref:
                        field["sourceRefs"] = [source_ref] + list(field.get("sourceRefs") or [])
                        source_ref = {}
                elif incoming_rank == existing_rank and str(field.get("status") or "") != "confirmed":
                    existing_material = any(
                        is_material_fact_ref(ref)
                        for ref in (field.get("sourceRefs") if isinstance(field.get("sourceRefs"), list) else [])
                        if isinstance(ref, dict)
                    )
                    if not (existing_material and is_material_fact_ref(source_ref)):
                        field["status"] = "conflict"
                    if value_text not in [str(item.get("value") or "") for item in alternatives if isinstance(item, dict)]:
                        alternatives.append({"value": value_text, "source": source_ref})
                elif value_text not in [str(item.get("value") or "") for item in alternatives if isinstance(item, dict)]:
                    alternatives.append({"value": value_text, "source": source_ref})
            elif value_text and field.get("value") and value_text == field.get("value"):
                existing_rank = (int(field.get("sourcePriority") or 0), float(field.get("confidence") or 0))
                incoming_rank = (incoming_priority, float(confidence or 0))
                if incoming_rank > existing_rank and str(field.get("status") or "") != "confirmed":
                    field["unit"] = str(unit or field.get("unit") or "")
                    field["category"] = category
                    field["confidence"] = float(confidence or 0)
                    field["sourcePriority"] = incoming_priority
                    if source_ref:
                        field["sourceRefs"] = [source_ref] + list(field.get("sourceRefs") or [])
                        source_ref = {}
            elif value_text and not field.get("value"):
                field["value"] = value_text
                field["status"] = "candidate"
                field["unit"] = str(unit or field.get("unit") or "")
                field["confidence"] = max(float(field.get("confidence") or 0), float(confidence or 0))
                field["sourcePriority"] = incoming_priority
                field["category"] = category
                if source_ref:
                    field["sourceRefs"] = [source_ref] + list(field.get("sourceRefs") or [])
                    source_ref = {}
            if source_ref:
                field.setdefault("sourceRefs", []).append(source_ref)
            if preserve_existing and value_text:
                field["status"] = "confirmed"

        def preserve_manual_fields() -> None:
            for existing in existing_by_key.values():
                if not isinstance(existing, dict):
                    continue
                label_text = self._canonical_fact_label(existing.get("label"))
                key = self._fact_label_key(label_text)
                if not key or key in fields_by_key:
                    continue
                source_refs = [
                    ref
                    for ref in (existing.get("sourceRefs") if isinstance(existing.get("sourceRefs"), list) else [])
                    if isinstance(ref, dict)
                ]
                is_manual = any(str(ref.get("type") or "") == "manualFact" for ref in source_refs)
                if not is_manual:
                    continue
                has_value = bool(str(existing.get("value") or "").strip())
                field = copy.deepcopy(existing)
                field["label"] = label_text
                field["key"] = key
                field["category"] = str(field.get("category") or "人工补充事实")
                field["status"] = str(field.get("status") or ("candidate" if has_value else "missing"))
                field["sourceRefs"] = source_refs or [{"type": "manualFact", "title": "人工新增", "field": label_text}]
                fields_by_key[key] = field

        trusted_parse_facts = self._trusted_parse_fact_fields(project.get("parse_result"))
        first_parse_value = {
            self._fact_label_key(fact.get("label")): fact.get("value")
            for fact in trusted_parse_facts
            if fact.get("value")
        }
        identity = project.get("identity") if isinstance(project.get("identity"), dict) else {}
        owner = identity.get("owner") or identity.get("customerCanonicalName") or identity.get("customerName") or project.get("owner") or project.get("customerName")
        project_name = first_parse_value.get(self._fact_label_key("项目名称")) or project.get("name")
        add_candidate("项目名称", project_name, category="项目基础信息", source_ref={"type": "project", "field": "name", "title": "项目名称"}, confidence=0.86, source_priority=320)
        add_candidate("招标方", owner, category="项目基础信息", source_ref={"type": "projectIdentity", "field": "owner", "title": "招标方"}, confidence=0.92, source_priority=320)
        add_candidate("招标人", owner, category="项目基础信息", source_ref={"type": "projectIdentity", "field": "owner", "title": "招标人"}, confidence=0.92, source_priority=320)
        add_candidate("客户名称", project.get("customerName"), category="项目基础信息", source_ref={"type": "project", "field": "customerName", "title": "客户名称"}, confidence=0.9, source_priority=320)
        add_candidate("日期", datetime.now(UTC).strftime("%Y年%m月%d日"), category="系统字段", source_ref={"type": "system", "field": "currentDate", "title": "当前日期"}, confidence=0.62)

        turbine = project_turbine_model(project)
        model = turbine.get("model") or turbine.get("turbineModel")
        hub_height = turbine.get("hubHeightM")
        add_candidate("投标机型", model, category="机型参数", source_ref={"type": "projectTurbineModel", "field": "model", "title": "投标机型"}, confidence=0.98, source_priority=320)
        rated_kw = turbine.get("ratedPowerKw")
        rated_mw = ""
        if isinstance(rated_kw, (int, float)):
            rated_mw = f"{rated_kw / 1000:g}"
        add_candidate("单机容量", rated_mw or rated_kw, category="机型参数", source_ref={"type": "projectTurbineModel", "field": "ratedPowerKw", "title": "单机容量"}, confidence=0.9, unit="MW" if rated_mw else "", source_priority=320)
        add_candidate("叶轮直径", turbine.get("rotorDiameterM"), category="机型参数", source_ref={"type": "projectTurbineModel", "field": "rotorDiameterM", "title": "叶轮直径"}, confidence=0.9, unit="m", source_priority=320)
        add_candidate("轮毂高度", hub_height, category="机型参数", source_ref={"type": "projectTurbineModel", "field": "hubHeightM", "title": "轮毂高度"}, confidence=0.86, unit="m", source_priority=320)
        if model and hub_height:
            add_candidate("投标方案", f"{model}-{hub_height}m", category="方案口径", source_ref={"type": "derived", "field": "modelHubHeight", "title": "投标方案"}, confidence=0.78, source_priority=320)
            add_candidate("方案", f"{model}-{hub_height}m", category="方案口径", source_ref={"type": "derived", "field": "modelHubHeight", "title": "方案"}, confidence=0.78, source_priority=320)
        elif model:
            add_candidate("投标方案", model, category="方案口径", source_ref={"type": "derived", "field": "model", "title": "投标方案"}, confidence=0.64, source_priority=80)

        for fact in trusted_parse_facts:
            add_candidate(
                str(fact.get("label") or ""),
                fact.get("value"),
                category=str(fact.get("category") or "招标解析字段"),
                source_ref=copy.deepcopy(fact.get("sourceRef") or {}),
                confidence=float(fact.get("confidence") or 0.82),
                required=bool(fact.get("required", False)),
                unit=str(fact.get("unit") or ""),
                source_priority=260,
            )

        for fact in self._project_material_fact_fields(project, gap_state, excluded_paths=blank_source_paths()):
            if fact.get("internal"):
                continue
            add_candidate(
                str(fact.get("label") or ""),
                fact.get("value"),
                category=str(fact.get("category") or "素材库事实"),
                source_ref=copy.deepcopy(fact.get("sourceRef") or {}),
                confidence=float(fact.get("confidence") or 0.78),
                required=bool(fact.get("required", False)),
                unit=str(fact.get("unit") or ""),
                source_priority=int(fact.get("sourcePriority") or 0),
            )

        plan = gap_state.get("plan") if isinstance(gap_state.get("plan"), dict) else {}
        for item in plan.get("items") or []:
            if not isinstance(item, dict):
                continue
            for task in item.get("fillTasks") or []:
                if not isinstance(task, dict):
                    continue
                blank = task.get("blankSource") if isinstance(task.get("blankSource"), dict) else {}
                for label in blank.get("placeholderLabels") or []:
                    add_candidate(
                        str(label),
                        "",
                        category="待填写Word字段",
                        source_ref={
                            "type": "gapPlaceholder",
                            "gapId": str(item.get("id") or ""),
                            "title": str(item.get("title") or ""),
                            "field": str(label),
                        },
                        confidence=0.0,
                    )
                for label in self._fillable_table_labels_from_blank_source(blank):
                    add_candidate(
                        label,
                        "",
                        category="待填写表格字段",
                        source_ref={
                            "type": "gapTableField",
                            "gapId": str(item.get("id") or ""),
                            "title": str(item.get("title") or ""),
                            "field": label,
                            "blankSourceId": str(blank.get("id") or ""),
                        },
                        confidence=0.0,
                    )

        for task in plan.get("tasks") or []:
            if not isinstance(task, dict):
                continue
            for label in self._business_fact_labels_from_task(task):
                add_candidate(
                    label,
                    "",
                    category="商务待填写字段",
                    source_ref={
                        "type": "businessGapTask",
                        "taskId": str(task.get("id") or ""),
                        "title": str(task.get("title") or ""),
                        "field": label,
                    },
                    confidence=0.0,
                )

        preserve_manual_fields()

        fields = list(fields_by_key.values())
        for field in fields:
            source_refs = self._normalize_fact_source_refs(field.get("sourceRefs"))
            field["sourceRefs"] = source_refs
            if source_refs:
                field["sourcePriority"] = max(
                    int(field.get("sourcePriority") or 0),
                    self._fact_source_ref_priority(source_refs[0]),
                )
        category_order = {
            "项目基础信息": 0,
            "机型参数": 1,
            "方案口径": 2,
            "系统字段": 3,
            "招标解析字段": 4,
            "素材库事实": 5,
            "性能保证": 6,
            "待填写Word字段": 7,
            "待填写表格字段": 8,
            "商务待填写字段": 9,
        }
        fields.sort(
            key=lambda field: (
                category_order.get(str(field.get("category") or ""), 9),
                0 if field.get("required") else 1,
                field.get("label") or "",
            )
        )
        for index, field in enumerate(fields, start=1):
            field["id"] = field.get("id") or f"FACT-{index:04d}"
        return {
            "schemaVersion": PROJECT_FACT_TABLE_SCHEMA_VERSION,
            "projectId": str(project.get("id") or ""),
            "status": "draft",
            "builtAt": built_at,
            "updatedAt": built_at,
            "confirmedAt": "",
            "confirmedBy": "",
            "fields": fields,
            "summary": self._summarize_project_fact_fields(fields),
        }

    @classmethod
    def _business_fact_labels_from_task(cls, task: dict[str, Any]) -> list[str]:
        title = str(task.get("title") or "")
        task_key = str(task.get("taskKey") or "")
        text = f"{title} {task_key}"
        labels: list[str] = []
        if re.search(r"投标函|授权|廉洁|专用章|保证金|保函|履约|其他说明|承诺|声明", text):
            labels.extend(["项目名称", "招标编号", "招标人", "投标人", "日期"])
        if re.search(r"价格|报价|开标|投标价格", text):
            labels.extend(["项目名称", "招标编号", "投标报价", "币种"])
        if re.search(r"规格|货物|供货范围|供货清单", text):
            labels.extend(["投标机型", "单机容量", "总装机容量", "机组台数"])
        if re.search(r"商务偏差|条款偏差", text):
            labels.extend(["招标编号", "项目名称", "商务偏差说明"])
        return list(dict.fromkeys(labels))

    @classmethod
    def _fact_table_value_map(cls, fact_table: dict[str, Any]) -> dict[str, str]:
        values: dict[str, str] = {}
        for field in fact_table.get("fields") or []:
            if not isinstance(field, dict):
                continue
            label = cls._canonical_fact_label(str(field.get("label") or field.get("title") or ""))
            value = str(field.get("value") or "").strip()
            if label and value:
                values[label] = value
                values[cls._fact_label_key(label)] = value
        return values

    @classmethod
    def _iter_parse_fact_fields(cls, value: Any) -> list[dict[str, Any]]:
        fields: list[dict[str, Any]] = []

        def visit(node: Any) -> None:
            if len(fields) >= 200:
                return
            if isinstance(node, dict):
                has_label = any(key in node for key in ("label", "title", "key", "id"))
                has_value = any(key in node for key in ("value", "keyValue", "evidence"))
                if has_label and has_value:
                    fields.append(node)
                for child in node.values():
                    visit(child)
            elif isinstance(node, list):
                for child in node:
                    visit(child)

        visit(value)
        return fields

    @classmethod
    def _trusted_parse_fact_fields(cls, parse_result: Any) -> list[dict[str, Any]]:
        fields: dict[str, dict[str, Any]] = {}

        def add(
            label: str,
            value: Any,
            *,
            category: str,
            source_field: dict[str, Any],
            confidence: float,
            required: bool = False,
            unit: str = "",
        ) -> None:
            label_text = cls._canonical_fact_label(label)
            value_text = str(value or "").strip()
            if not label_text or not value_text:
                return
            key = cls._fact_label_key(label_text)
            current = fields.get(key)
            fact = {
                "label": label_text,
                "value": value_text,
                "category": category,
                "confidence": confidence,
                "required": required,
                "unit": unit,
                "sourceRef": {
                    "type": "parseField",
                    "field": str(source_field.get("id") or source_field.get("fieldKey") or source_field.get("title") or ""),
                    "fieldKey": str(source_field.get("fieldKey") or ""),
                    "title": str(source_field.get("title") or source_field.get("label") or label_text),
                    "sourceFile": str(source_field.get("sourceFile") or ""),
                },
            }
            if current is None or confidence > float(current.get("confidence") or 0):
                fields[key] = fact

        for field in cls._iter_parse_fact_fields(parse_result):
            field_key = str(field.get("fieldKey") or "").strip()
            label = str(field.get("title") or field.get("label") or field.get("key") or field.get("id") or "").strip()
            value = str(field.get("value") or field.get("keyValue") or "").strip()
            evidence = str(field.get("evidence") or "").strip()
            text = "。".join(part for part in (label, value, evidence) if part)

            if field_key == "projectName" or cls._fact_label_key(label) == cls._fact_label_key("项目名称"):
                if cls._looks_like_project_name(value):
                    add("项目名称", value, category="项目基础信息", source_field=field, confidence=0.95, required=True)
            elif field_key == "tenderNo" or cls._fact_label_key(label) == cls._fact_label_key("招标编号"):
                if cls._looks_like_tender_no(value):
                    add("招标编号", value, category="项目基础信息", source_field=field, confidence=0.94, required=True)
            elif field_key in {"tenderer", "owner", "customerName"} or cls._fact_label_key(label) in {
                cls._fact_label_key("招标人"),
                cls._fact_label_key("招标方"),
                cls._fact_label_key("客户名称"),
            }:
                if cls._looks_like_party_name(value):
                    add("招标人", value, category="项目基础信息", source_field=field, confidence=0.84, required=False)
            elif field_key == "managementUnit" or cls._fact_label_key(label) == cls._fact_label_key("管理单位"):
                if cls._looks_like_party_name(value):
                    add("管理单位", value, category="项目基础信息", source_field=field, confidence=0.82, required=False)
            elif field_key == "bidSectionScale" or cls._fact_label_key(label) in {
                cls._fact_label_key("标段规模"),
                cls._fact_label_key("招标规模"),
            }:
                add("标段规模", value, category="项目基础信息", source_field=field, confidence=0.78, required=False)
            elif field_key == "deliveryPeriod" or cls._fact_label_key(label) == cls._fact_label_key("交货周期"):
                add("交货周期", value, category="项目基础信息", source_field=field, confidence=0.78, required=False)
            elif field_key == "warrantyPeriod" or cls._fact_label_key(label) == cls._fact_label_key("质保期"):
                add("质保期", value, category="项目基础信息", source_field=field, confidence=0.78, required=False)
            elif field_key == "bidStartDate" or cls._fact_label_key(label) == cls._fact_label_key("投标起始日期"):
                add("投标起始日期", value, category="投标时间信息", source_field=field, confidence=0.78, required=False)
            elif field_key == "bidDeadline" or cls._fact_label_key(label) in {
                cls._fact_label_key("投标截止日期"),
                cls._fact_label_key("投标截止时间"),
            }:
                add("投标截止日期", value, category="投标时间信息", source_field=field, confidence=0.82, required=True)

            cls._add_performance_facts_from_parse_text(text, field, add)

        return list(fields.values())

    @staticmethod
    def _looks_like_project_name(value: Any) -> bool:
        text = str(value or "").strip()
        if not text or len(text) > 160:
            return False
        if re.search(r"投标人|招标人|应当|必须|不得|标准|规范|条款|认可|提供", text):
            return False
        return "项目" in text or "工程" in text

    @staticmethod
    def _looks_like_tender_no(value: Any) -> bool:
        text = str(value or "").strip()
        return bool(text and len(text) <= 80 and re.fullmatch(r"[A-Za-z0-9][A-Za-z0-9_\-./]+", text))

    @staticmethod
    def _looks_like_party_name(value: Any) -> bool:
        text = str(value or "").strip()
        if not text or len(text) > 80:
            return False
        if re.search(r"[。；;]|投标人|应|必须|不得|标准|规范|条款|认可|提供|要求|报告|测试|审查", text):
            return False
        return bool(re.search(r"公司|集团|有限|招标|业主|电力|能源|华能|国电|大唐|华电", text))

    @classmethod
    def _add_performance_facts_from_parse_text(
        cls,
        text: str,
        source_field: dict[str, Any],
        add: Any,
    ) -> None:
        normalized = re.sub(r"\s+", "", str(text or ""))
        if not normalized:
            return

        patterns = [
            (r"功率曲线[^。；;]{0,24}(?:不低于|≥|>=)(?:保证值的)?([0-9]+(?:\.[0-9]+)?%)", "功率曲线保证率"),
            (r"风电场机组年平均可利用率(?:≥|>=|不低于)([0-9]+(?:\.[0-9]+)?%)", "全场可利用率"),
            (r"(?:全部机组|全场).*?平均可利用率(?:≥|>=|不低于)([0-9]+(?:\.[0-9]+)?%)", "全场可利用率"),
            (r"单台机组年平均可利用率(?:≥|>=|不低于)([0-9]+(?:\.[0-9]+)?%)", "单台可利用率"),
            (r"主要部件更换率(?:低于|不高于|≤|<=)([0-9]+(?:\.[0-9]+)?%)", "主要部件更换率"),
        ]
        for pattern, label in patterns:
            match = re.search(pattern, normalized)
            if match:
                add(
                    label,
                    match.group(1),
                    category="性能保证",
                    source_field=source_field,
                    confidence=0.86,
                    required=False,
                    unit="%",
                )

    @classmethod
    def _fillable_table_labels_from_blank_source(cls, blank: dict[str, Any]) -> list[str]:
        path = cls._blank_source_docx_path(blank)
        if path is None:
            return []
        try:
            document = Document(str(path))
        except Exception:
            return []

        labels: list[str] = []
        seen: set[str] = set()
        for table in document.tables:
            for row in table.rows:
                cells = [cls._clean_table_cell_text(cell.text) for cell in row.cells]
                label = cls._table_field_label_from_row(cells)
                if not label:
                    continue
                key = cls._fact_label_key(label)
                if not key or key in seen:
                    continue
                seen.add(key)
                labels.append(label)
        return labels

    @staticmethod
    def _blank_source_docx_path(blank: dict[str, Any]) -> Path | None:
        for key in ("docxPath", "path", "workspacePath"):
            value = str(blank.get(key) or "").strip()
            if not value:
                continue
            path = Path(value)
            if path.exists():
                return path
        return None

    @staticmethod
    def _clean_table_cell_text(value: Any) -> str:
        return re.sub(r"\s+", "", str(value or "")).strip()

    @classmethod
    def _table_field_label_from_row(cls, cells: list[str]) -> str:
        if not cells:
            return ""
        fill_positions = [
            index
            for index, text in enumerate(cells)
            if not text or re.search(r"待(?:人工)?(?:补充|填写|解析)|未填写|待确认", text)
        ]
        if not fill_positions:
            return ""
        candidates = cells[: fill_positions[0]]
        for candidate in reversed(candidates):
            label = cls._canonical_fact_label(candidate)
            if cls._looks_like_table_field_label(label):
                return label
        return ""

    @staticmethod
    def _looks_like_table_field_label(label: str) -> bool:
        text = str(label or "").strip()
        if not text or len(text) < 2 or len(text) > 80:
            return False
        if text in FACT_TABLE_HEADER_WORDS:
            return False
        if re.fullmatch(r"[\d一二三四五六七八九十]+[.、]?", text):
            return False
        if re.search(r"待(?:人工)?(?:补充|填写|解析)|未填写|授权人签名|日期", text):
            return False
        if re.search(r"同等质量|知名品牌|件套|厂家|品牌|Fluke|FLUKE|SKYLOTEC|DEHN|ABB|西门子|施耐德", text, flags=re.I):
            return False
        if re.search(r"参数|方法|折减|系数", text):
            return False
        if re.fullmatch(r"[A-Z]{1,8}[-A-Z0-9（）()\"'.—]+", text):
            return False
        if re.match(r"^\d", text) and not re.search(r"风速|年|容量|功率|高度|直径|小时|电量|温度", text):
            return False
        if text in COMMON_PROJECT_FACT_LABELS:
            return True
        return bool(
            re.search(
                r"投标机型|机组类型|机组台数|风机台数|单机容量|总装机容量|叶轮直径|风轮直径|轮毂.*高度|"
                r"扫风面积|比功率|安全等级|设计寿命|功率曲线|可利用率|保证电量|保证发电量|发电小时|"
                r"有效小时|等效利用小时|平均风速|空气密度|湍流|风切变|风剪切|极端风速|极大风速|"
                r"低温|高温|海拔|覆冰|盐雾|沙尘|雷电",
                text,
            )
        )

    def _project_material_fact_fields(
        self,
        project: dict[str, Any],
        gap_state: dict[str, Any],
        *,
        excluded_paths: set[str] | None = None,
    ) -> list[dict[str, Any]]:
        materials = self._project_fact_material_index(project, gap_state)
        if not materials:
            return []
        prepared = self._prepare_project_fact_materials(project, materials)
        facts: list[dict[str, Any]] = []
        for material in prepared:
            if not isinstance(material, dict):
                continue
            facts.extend(self._facts_from_material_name(material))
            path_text = str(material.get("path") or material.get("docx") or "").strip()
            if not path_text:
                continue
            path = Path(path_text)
            if not path.exists():
                continue
            if str(path.resolve()) in (excluded_paths or set()):
                continue
            suffix = path.suffix.lower()
            if suffix in {".docx", ".doc"}:
                facts.extend(self._facts_from_docx_material(path, material))
            elif suffix in {".xlsx", ".xlsm"}:
                facts.extend(self._facts_from_xlsx_material(path, material, project))
        facts.extend(self._derived_material_fact_fields(project, facts))
        return facts

    def _project_fact_material_index(self, project: dict[str, Any], gap_state: dict[str, Any]) -> list[dict[str, Any]]:
        plan = gap_state.get("plan") if isinstance(gap_state.get("plan"), dict) else {}
        materials = [
            dict(item)
            for item in (plan.get("materialIndex") if isinstance(plan.get("materialIndex"), list) else [])
            if isinstance(item, dict)
        ]
        if not materials and isinstance(plan.get("tasks"), list):
            seen: set[str] = set()
            for task in plan.get("tasks") or []:
                if not isinstance(task, dict):
                    continue
                for raw in [*(task.get("candidateMaterials") or []), *(task.get("selectedMaterialRefs") or [])]:
                    if not isinstance(raw, dict):
                        continue
                    material_id = str(raw.get("id") or raw.get("materialId") or "").strip()
                    if not material_id or material_id in seen:
                        continue
                    seen.add(material_id)
                    materials.append(
                        {
                            "id": material_id,
                            "name": str(raw.get("name") or raw.get("materialName") or raw.get("fileName") or material_id),
                            "folderPath": str(raw.get("folderPath") or raw.get("path") or ""),
                            "materialTier": str(raw.get("materialTier") or raw.get("libraryScope") or ""),
                            "cleanedFileName": str(raw.get("cleanedFileName") or ""),
                            "hasCleanedWord": bool(raw.get("hasCleanedWord")),
                            "turbineModelLabel": str(raw.get("turbineModelLabel") or ""),
                        }
                    )
        if not materials:
            try:
                materials = _allowed_material_index(build_project_material_scope(project), project_turbine_model(project))
            except Exception:
                materials = []
        return [item for item in materials if self._material_is_fact_relevant(item)]

    @staticmethod
    def _material_is_fact_relevant(material: dict[str, Any]) -> bool:
        tier = str(material.get("materialTier") or "").strip()
        if tier == "project":
            return True
        text = " ".join(
            str(material.get(key) or "")
            for key in ("name", "cleanedFileName", "folderPath", "path")
        )
        return bool(
            re.search(
                r"参数|机型|功率曲线|风资源|发电量|报价|容量|安全|场址|载荷|工程量|技术承诺|投标关键数据",
                text,
            )
        )

    @staticmethod
    def _prepare_project_fact_materials(project: dict[str, Any], materials: list[dict[str, Any]]) -> list[dict[str, Any]]:
        path_materials = [item for item in materials if item.get("path")]
        if path_materials and len(path_materials) == len(materials) and all(
            Path(str(item.get("path") or "")).exists() for item in path_materials
        ):
            return materials
        project_id = str(project.get("id") or "project")
        bid_type = normalize_bid_type(str(project.get("bidType") or "")) or "技术标"
        workspace_dir = "business-workspace" if bid_type == "商务标" else "technical-workspace"
        work_dir = settings.documents_dir / project_id / workspace_dir / "gaps" / "fact_table_materials"
        work_dir.mkdir(parents=True, exist_ok=True)
        try:
            return _prepare_material_index_files(materials, work_dir, limit=120)
        except Exception:
            return materials

    @classmethod
    def _facts_from_material_name(cls, material: dict[str, Any]) -> list[dict[str, Any]]:
        text = str(material.get("name") or material.get("cleanedFileName") or "")
        facts: list[dict[str, Any]] = []
        for pattern, label, unit in [
            (r"空气密度\s*([0-9]+(?:\.[0-9]+)?)", "空气密度", "kg/m3"),
            (r"湍流强度\s*([0-9]+(?:\.[0-9]+)?)", "湍流强度", ""),
            (r"风(?:剪切|切变)(?:指数)?\s*([0-9]+(?:\.[0-9]+)?)", "风剪切", ""),
        ]:
            match = re.search(pattern, text, flags=re.I)
            if match:
                facts.append(cls._material_fact(label, match.group(1), material, unit=unit, confidence=0.9))
        return facts

    @classmethod
    def _facts_from_docx_material(cls, path: Path, material: dict[str, Any]) -> list[dict[str, Any]]:
        facts: list[dict[str, Any]] = []
        try:
            document = Document(str(path))
        except Exception:
            return facts
        text_parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for para_idx, paragraph in enumerate(document.paragraphs, start=1):
            text = cls._clean_fact_text(paragraph.text)
            if not text or len(text) > 180:
                continue
            match = re.match(r"^([^:：]{2,40})[:：]\s*(.{1,100})$", text)
            if match:
                fact = cls._material_fact_from_label_value(match.group(1), match.group(2), material, location=f"P{para_idx}", confidence=0.78)
                if fact:
                    facts.append(fact)
        for table_idx, table in enumerate(document.tables, start=1):
            facts.extend(cls._facts_from_guarantee_table(table, material, table_idx=table_idx))
            for row_idx, row in enumerate(table.rows, start=1):
                cells = [cls._clean_fact_text(cell.text) for cell in row.cells]
                text_parts.append(" | ".join(cell for cell in cells if cell))
                facts.extend(cls._facts_from_table_cells(cells, material, location=f"T{table_idx}/R{row_idx}"))
        facts.extend(cls._facts_from_free_text("\n".join(text_parts), material))
        return facts

    @classmethod
    def _facts_from_guarantee_table(cls, table: Any, material: dict[str, Any], *, table_idx: int) -> list[dict[str, Any]]:
        if not getattr(table, "rows", None) or len(table.rows) < 2:
            return []
        header = " ".join(cls._clean_fact_text(cell.text) for cell in table.rows[0].cells)
        if not ("年平均风速" in header and "保证年上网电量" in header and "满负荷小时" in header):
            return []
        facts: list[dict[str, Any]] = []
        for row_idx, row in enumerate(table.rows[1:], start=2):
            cells = [cls._clean_fact_text(cell.text) for cell in row.cells]
            if len(cells) < 3:
                continue
            wind_speed = cls._clean_fact_value("年平均风速", cells[0])
            energy = cls._clean_fact_value("保证发电量", cells[1])
            hours = cls._clean_fact_value("保证有效小时数", cells[2])
            if not (wind_speed and energy and hours):
                continue
            matrix_fact = cls._material_fact(
                "__guaranteeMatrixRow",
                {"windSpeed": wind_speed, "energyMwh": energy, "hours": hours},
                material,
                location=f"T{table_idx}/R{row_idx}",
                confidence=0.82,
            )
            matrix_fact["internal"] = True
            facts.append(matrix_fact)
        return facts

    @classmethod
    def _facts_from_xlsx_material(cls, path: Path, material: dict[str, Any], project: dict[str, Any]) -> list[dict[str, Any]]:
        facts: list[dict[str, Any]] = []
        try:
            workbook = load_workbook(path, data_only=True, read_only=True)
        except Exception:
            return facts
        project_model = str((project_turbine_model(project) or {}).get("model") or "")
        model_key = re.sub(r"(上置|下置|内置|外置|塔上|塔下)", "", project_model)
        for worksheet in workbook.worksheets:
            selected_col = cls._xlsx_model_column(worksheet, model_key)
            for row_idx, row in enumerate(worksheet.iter_rows(values_only=True), start=1):
                cells = [cls._clean_fact_text(cell) for cell in row]
                if selected_col is not None and selected_col < len(cells):
                    for label_idx in (2, 1, 0):
                        if label_idx < len(cells):
                            fact = cls._material_fact_from_label_value(
                                cells[label_idx],
                                cells[selected_col],
                                material,
                                unit=cells[3] if len(cells) > 3 else "",
                                location=f"{worksheet.title}!R{row_idx}",
                                confidence=0.82,
                            )
                            if fact:
                                facts.append(fact)
                                break
                facts.extend(cls._facts_from_table_cells(cells, material, location=f"{worksheet.title}!R{row_idx}"))
                if len(facts) >= 800:
                    return facts
        return facts

    @staticmethod
    def _xlsx_model_column(worksheet: Any, model_key: str) -> int | None:
        if not model_key:
            return None
        normalized_model = re.sub(r"\s+", "", model_key)
        for row in worksheet.iter_rows(min_row=1, max_row=min(12, worksheet.max_row), values_only=True):
            for index, value in enumerate(row):
                text = re.sub(r"\s+", "", str(value or ""))
                if normalized_model and normalized_model in text:
                    return index
        return None

    @classmethod
    def _facts_from_table_cells(cls, cells: list[str], material: dict[str, Any], *, location: str) -> list[dict[str, Any]]:
        facts: list[dict[str, Any]] = []
        nonempty = [(index, value) for index, value in enumerate(cells) if value]
        if len(nonempty) < 2:
            return facts
        # Common project tables often use: category | label | unit | value.
        if len(cells) >= 4:
            fact = cls._material_fact_from_label_value(cells[1], cells[3], material, unit=cells[2], location=location, confidence=0.88)
            if fact:
                facts.append(fact)
        for first, second in zip(nonempty, nonempty[1:]):
            fact = cls._material_fact_from_label_value(first[1], second[1], material, location=location, confidence=0.76)
            if fact:
                facts.append(fact)
                break
        # Wind-resource matrix rows use: label | unit | tower1 | tower2...
        if len(cells) >= 3 and cells[0]:
            wind_fact = cls._material_fact_from_label_value(cells[0], cells[2], material, unit=cells[1], location=location, confidence=0.84)
            if wind_fact:
                facts.append(wind_fact)
        return facts

    @classmethod
    def _material_fact_from_label_value(
        cls,
        label: Any,
        value: Any,
        material: dict[str, Any],
        *,
        unit: str = "",
        location: str = "",
        confidence: float = 0.78,
    ) -> dict[str, Any] | None:
        label_text = cls._canonical_fact_label(label)
        value_text = cls._clean_fact_value(label_text, value)
        if not label_text or not value_text:
            return None
        if label_text not in COMMON_PROJECT_FACT_LABELS and not cls._looks_like_table_field_label(label_text):
            return None
        unit_text = cls._clean_fact_unit(unit)
        raw_label = str(label or "")
        raw_value = str(value or "")
        if not unit_text:
            raw_context = f"{raw_label}{raw_value}"
            if label_text in {"轮毂高度", "叶轮直径"} and re.search(r"(?:m|米)", raw_context, flags=re.I):
                unit_text = "m"
            elif label_text in {"极端风速", "年平均风速"} and re.search(r"m/?s|米/秒", raw_context, flags=re.I):
                unit_text = "m/s"
            elif label_text == "空气密度" and re.search(r"kg/?m|kg/m3|kg/m³", raw_context, flags=re.I):
                unit_text = "kg/m3"
            elif label_text == "机组台数" and "台" in raw_value:
                unit_text = "台"
        if not unit_text:
            if label_text in {"轮毂高度", "叶轮直径"}:
                unit_text = "m"
            elif label_text in {"极端风速", "年平均风速"}:
                unit_text = "m/s"
            elif label_text == "空气密度":
                unit_text = "kg/m3"
            elif label_text == "机组台数":
                unit_text = "台"
        return cls._material_fact(label_text, value_text, material, unit=unit_text, location=location, confidence=confidence)

    @classmethod
    def _facts_from_free_text(cls, text: str, material: dict[str, Any]) -> list[dict[str, Any]]:
        compact = cls._clean_fact_text(text)
        facts: list[dict[str, Any]] = []
        patterns = [
            (r"(?:总装机容量|建设容量|标段规模|总容量)[^0-9]{0,12}([0-9]+(?:\.[0-9]+)?\s*(?:MW|万千瓦|kW)?)", "总装机容量", ""),
            (r"(?:机组台数|机组数量|风机台数|风机数量|安装)[^0-9]{0,12}([0-9]+)\s*台", "机组台数", "台"),
            (r"轮毂(?:中心)?高度[^0-9]{0,12}([0-9]+(?:\.[0-9]+)?\s*m?)", "轮毂高度", "m"),
            (r"(?:安全等级|适用等级|设计等级)[^A-Za-z0-9]{0,12}((?:IEC\s*)?[A-Z0-9][A-Z0-9/ .-]{0,20})", "安全等级", ""),
            (r"空气密度[^0-9]{0,12}([0-9]+(?:\.[0-9]+)?)", "空气密度", "kg/m3"),
            (r"湍流强度[^0-9]{0,12}([0-9]+(?:\.[0-9]+)?)", "湍流强度", ""),
            (r"(?:极端风速|极大风速|Ve50)[^0-9]{0,12}([0-9]+(?:\.[0-9]+)?\s*m/s?)", "极端风速", "m/s"),
            (r"年平均风速[^0-9]{0,12}([0-9]+(?:\.[0-9]+)?\s*m/s?)", "年平均风速", "m/s"),
        ]
        for pattern, label, unit in patterns:
            match = re.search(pattern, compact, flags=re.I)
            if match:
                value = cls._clean_fact_value(label, match.group(1))
                if value:
                    facts.append(cls._material_fact(label, value, material, unit=unit, confidence=0.78))
        return facts

    @staticmethod
    def _clean_fact_text(value: Any) -> str:
        return re.sub(r"\s+", "", str(value or "")).strip()

    @staticmethod
    def _clean_fact_unit(value: Any) -> str:
        text = re.sub(r"\s+", "", str(value or "")).strip()
        text = text.strip("：:；;，,、")
        if text in {"", "-", "/", "—", "NA", "N/A", "字段", "值", "年份", "参数内容", "结果", "说明", "备注", "机型", "型号"}:
            return ""
        text = text.replace("m³", "m3")
        text = re.sub(r"kg/?m3", "kg/m3", text, flags=re.I)
        text = re.sub(r"m/?s", "m/s", text, flags=re.I)
        return text

    @classmethod
    def _clean_fact_value(cls, label: str, value: Any) -> str:
        text = str(value or "").strip()
        text = re.sub(r"\s+", "", text)
        text = text.strip("：:；;，,、")
        if not text or len(text) > 120:
            return ""
        if any(token in text for token in ("待填写", "待人工", "未填写")):
            return ""
        if text in {"-", "/", "—", "无", "暂无", "值", "结果", "参数内容", "单位", "年份"}:
            return ""
        numeric_ranges = {
            "机组台数": (1, 1000),
            "轮毂高度": (40, 250),
            "叶轮直径": (50, 350),
            "空气密度": (0.7, 1.5),
            "湍流强度": (0, 1),
            "风剪切": (0, 1),
            "极端风速": (20, 100),
            "年平均风速": (2, 15),
        }
        numeric_noise = (
            "年份",
            "各年",
            "版本",
            "编制",
            "校核",
            "审核",
            "批准",
            "日期",
            "参数内容",
            "结果结果",
            "场址空气密度下",
        )
        if label in numeric_ranges and (
            any(token in text for token in numeric_noise)
            or re.search(r"\d{4}[-/年]\d{1,2}", text)
        ):
            return ""
        if label in {"总装机容量", "单机容量"}:
            match = re.search(r"([0-9]+(?:\.[0-9]+)?)(万千瓦|MW|kW)?", text, flags=re.I)
            if match:
                return f"{match.group(1)}{match.group(2) or ''}".strip()
        if label in {"机组台数"}:
            match = re.search(r"([0-9]+)", text)
            if not match:
                return ""
            number = float(match.group(1))
            low, high = numeric_ranges[label]
            return match.group(1) if low <= number <= high else ""
        if label in {"保证发电量", "保证有效小时数"}:
            if re.search(r"风电场|保证年上网电量|满负荷小时|字段|单位", text):
                return ""
            match = re.search(r"([0-9]+(?:\.[0-9]+)?)", text)
            if not match:
                return ""
            number = float(match.group(1))
            if label == "保证发电量" and not (1 <= number <= 10000000):
                return ""
            if label == "保证有效小时数" and not (1 <= number <= 8760):
                return ""
            return match.group(1)
        if label in {"极端风速", "年平均风速"}:
            match = re.search(r"([0-9]+(?:\.[0-9]+)?)", text)
            if match:
                number = float(match.group(1))
                low, high = numeric_ranges[label]
                if not (low <= number <= high):
                    return ""
                return f"{match.group(1)}m/s" if re.search(r"m/?s|米/秒", text, flags=re.I) else match.group(1)
        if label in {"轮毂高度", "叶轮直径", "空气密度", "湍流强度", "风剪切"}:
            match = re.search(r"([0-9]+(?:\.[0-9]+)?)", text)
            if match:
                number = float(match.group(1))
                low, high = numeric_ranges[label]
                if not (low <= number <= high):
                    return ""
                return match.group(1)
        if label in {"投标方案", "投标机型", "机组类型"}:
            if text in {"机型", "投标机型", "方案", "投标方案"}:
                return ""
            model_match = re.search(r"([A-Z]{1,6}\d+(?:\.\d+)?[-—]\d+(?:[-—]\d+)?)", text, flags=re.I)
            if model_match:
                return model_match.group(1).replace("—", "-")
        if label == "安全等级":
            text = re.sub(r"^IEC\s*", "IEC ", text, flags=re.I).strip()
        return text

    @classmethod
    def _material_fact(
        cls,
        label: str,
        value: Any,
        material: dict[str, Any],
        *,
        unit: str = "",
        location: str = "",
        confidence: float = 0.78,
    ) -> dict[str, Any]:
        tier = str(material.get("materialTier") or "").strip() or "standard"
        return {
            "label": cls._canonical_fact_label(label),
            "value": value,
            "category": "素材库事实",
            "unit": cls._clean_fact_unit(unit),
            "confidence": confidence,
            "sourcePriority": FACT_MATERIAL_SOURCE_PRIORITIES.get(tier, 50),
            "sourceRef": {
                "type": "materialFact",
                "materialId": str(material.get("id") or material.get("materialId") or ""),
                "materialTier": tier,
                "name": str(material.get("name") or material.get("fileName") or material.get("cleanedFileName") or ""),
                "folderPath": str(material.get("folderPath") or ""),
                "path": str(material.get("path") or ""),
                "location": location,
            },
        }

    @classmethod
    def _derived_material_fact_fields(cls, project: dict[str, Any], facts: list[dict[str, Any]]) -> list[dict[str, Any]]:
        by_label: dict[str, dict[str, Any]] = {}
        guarantee_matrix_rows: list[dict[str, Any]] = []
        for fact in facts:
            if fact.get("label") == "__guaranteeMatrixRow":
                guarantee_matrix_rows.append(fact)
                continue
            label = cls._canonical_fact_label(fact.get("label"))
            if not label or not fact.get("value"):
                continue
            current = by_label.get(label)
            rank = (int(fact.get("sourcePriority") or 0), float(fact.get("confidence") or 0))
            if current is None or rank > (int(current.get("sourcePriority") or 0), float(current.get("confidence") or 0)):
                by_label[label] = fact
        rated_kw = project_turbine_model(project).get("ratedPowerKw")
        rated_mw = float(rated_kw) / 1000 if isinstance(rated_kw, (int, float)) and rated_kw else 0
        result: list[dict[str, Any]] = []
        total = cls._number_from_fact(by_label.get("总装机容量"))
        count = cls._number_from_fact(by_label.get("机组台数"))
        source = by_label.get("总装机容量") or by_label.get("机组台数") or {}
        material_ref = copy.deepcopy(source.get("sourceRef") if isinstance(source.get("sourceRef"), dict) else {})
        if total and rated_mw and not count:
            derived_count = total / rated_mw
            rounded = round(derived_count)
            if abs(derived_count - rounded) < 0.01:
                result.append(
                    {
                        "label": "机组台数",
                        "value": str(rounded),
                        "category": "素材库事实",
                        "unit": "台",
                        "confidence": 0.86,
                        "sourcePriority": int(source.get("sourcePriority") or 0),
                        "sourceRef": {**material_ref, "type": "derivedMaterialFact", "field": "总装机容量/单机容量"},
                    }
                )
        if count and rated_mw and not total:
            result.append(
                {
                    "label": "总装机容量",
                    "value": f"{count * rated_mw:g}MW",
                    "category": "素材库事实",
                    "unit": "MW",
                    "confidence": 0.82,
                    "sourcePriority": int(source.get("sourcePriority") or 0),
                    "sourceRef": {**material_ref, "type": "derivedMaterialFact", "field": "机组台数/单机容量"},
                }
            )
        year_avg = cls._number_from_fact(by_label.get("年平均风速"))
        if year_avg and guarantee_matrix_rows:
            def matrix_distance(row: dict[str, Any]) -> float:
                value = row.get("value") if isinstance(row.get("value"), dict) else {}
                wind = cls._number_from_fact({"value": value.get("windSpeed")})
                return abs(float(wind or 0) - year_avg) if wind else 9999

            selected = min(guarantee_matrix_rows, key=matrix_distance)
            if matrix_distance(selected) <= 0.08:
                selected_value = selected.get("value") if isinstance(selected.get("value"), dict) else {}
                selected_ref = copy.deepcopy(selected.get("sourceRef") if isinstance(selected.get("sourceRef"), dict) else {})
                wind_speed = str(selected_value.get("windSpeed") or "")
                if selected_value.get("energyMwh"):
                    result.append(
                        {
                            "label": "保证发电量",
                            "value": str(selected_value.get("energyMwh")),
                            "category": "性能保证",
                            "unit": "MWh",
                            "confidence": 0.86,
                            "sourcePriority": int(selected.get("sourcePriority") or 0),
                            "sourceRef": {**selected_ref, "type": "derivedMaterialFact", "field": f"发电量保证矩阵/{wind_speed}m/s"},
                        }
                    )
                if selected_value.get("hours"):
                    result.append(
                        {
                            "label": "保证有效小时数",
                            "value": str(selected_value.get("hours")),
                            "category": "性能保证",
                            "unit": "h",
                            "confidence": 0.86,
                            "sourcePriority": int(selected.get("sourcePriority") or 0),
                            "sourceRef": {**selected_ref, "type": "derivedMaterialFact", "field": f"发电量保证矩阵/{wind_speed}m/s"},
                        }
                    )
        return result

    @staticmethod
    def _number_from_fact(fact: dict[str, Any] | None) -> float | None:
        if not isinstance(fact, dict):
            return None
        match = re.search(r"([0-9]+(?:\.[0-9]+)?)", str(fact.get("value") or ""))
        return float(match.group(1)) if match else None

    @staticmethod
    def _require_confirmed_project_fact_table(gap_state: dict[str, Any]) -> dict[str, Any]:
        table = gap_state.get("projectFactTable") if isinstance(gap_state.get("projectFactTable"), dict) else {}
        if table.get("status") not in PROJECT_FACT_CONFIRMED_STATUSES:
            raise ValueError("请先维护并确认项目事实表，再执行 AI 填写。")
        return table

    @staticmethod
    def _aggregate_gap_fill_quality(results: list[dict[str, Any]], errors: list[dict[str, str]]) -> dict[str, Any]:
        reports = [result.get("qualityReport") for result in results if isinstance(result.get("qualityReport"), dict)]
        if not reports:
            return {
                "status": "failed" if errors else "empty",
                "coverageRate": 0.0,
                "correctnessRate": 0.0,
                "completenessRate": 0.0,
                "thresholds": {"coverageRate": 0.85, "correctnessRate": 0.85, "completenessRate": 0.85},
            }
        expected = sum(int(report.get("expectedFieldCount") or 0) for report in reports)
        filled = sum(int(report.get("filledFieldCount") or 0) for report in reports)
        unfilled = sum(int(report.get("unfilledFieldCount") or 0) for report in reports)
        evidence = sum(int(report.get("evidenceRefCount") or 0) for report in reports)
        if expected > 0:
            coverage = filled / expected
            correctness = min(1.0, evidence / max(1, filled)) if filled else 0.0
            completeness = max(0.0, (expected - unfilled) / expected)
        else:
            coverage = sum(float(report.get("coverageRate") or 0) for report in reports) / len(reports)
            correctness = sum(float(report.get("correctnessRate") or 0) for report in reports) / len(reports)
            completeness = sum(float(report.get("completenessRate") or 0) for report in reports) / len(reports)
        thresholds = {"coverageRate": 0.85, "correctnessRate": 0.85, "completenessRate": 0.85}
        passed = (
            not errors
            and coverage >= thresholds["coverageRate"]
            and correctness >= thresholds["correctnessRate"]
            and completeness >= thresholds["completenessRate"]
        )
        return {
            "status": "passed" if passed else "needs_review",
            "coverageRate": round(coverage, 4),
            "correctnessRate": round(correctness, 4),
            "completenessRate": round(completeness, 4),
            "expectedFieldCount": expected,
            "filledFieldCount": filled,
            "unfilledFieldCount": unfilled,
            "evidenceRefCount": evidence,
            "taskCount": len(reports),
            "passedTaskCount": sum(1 for report in reports if report.get("status") == "passed"),
            "needsReviewTaskCount": sum(1 for report in reports if report.get("status") != "passed"),
            "thresholds": thresholds,
        }

    def _ensure_gap_state(self, project: dict[str, Any]) -> dict[str, Any]:
        gap_state = project.get("gap_state")
        if not isinstance(gap_state, dict):
            gap_state = {}
            project["gap_state"] = gap_state
        gap_state.setdefault("recognitionStatus", "idle")
        gap_state.setdefault("recognizedAt", "")
        gap_state.setdefault("submittedForReview", False)
        gap_state.setdefault("reviewConfirmed", False)
        gap_state.setdefault("reviewedAt", "")
        gap_state.setdefault("items", [])
        gap_state.setdefault("submissions", [])
        gap_state.setdefault("plan", {})
        gap_state.setdefault("planFile", "")
        gap_state.setdefault("integrity", {})
        gap_state.setdefault("projectFactTable", {})
        return gap_state

    def _ensure_business_gap_state(self, project: dict[str, Any]) -> dict[str, Any]:
        state = project.get("business_gap_state")
        if not isinstance(state, dict):
            state = {}
            project["business_gap_state"] = state
        state.setdefault("recognitionStatus", "idle")
        state.setdefault("recognizedAt", "")
        state.setdefault("submittedForReview", False)
        state.setdefault("reviewConfirmed", False)
        state.setdefault("reviewedAt", "")
        state.setdefault("plan", {})
        state.setdefault("planFile", "")
        state.setdefault("integrity", {})
        state.setdefault("projectFactTable", {})
        return state

    @staticmethod
    def _business_gap_integrity(plan: dict[str, Any]) -> dict[str, Any]:
        tasks = [
            task
            for task in (plan.get("tasks") if isinstance(plan.get("tasks"), list) else [])
            if isinstance(task, dict)
        ]
        blocking = [
            {
                "id": str(task.get("id") or ""),
                "title": str(task.get("title") or ""),
                "status": str(task.get("status") or ""),
                "decision": str(task.get("decision") or ""),
            }
            for task in tasks
            if str(task.get("status") or "") in {"needs_input", "filling", "review_required"}
        ]
        return {
            "status": "passed" if not blocking else "blocked",
            "checkedAt": now_iso(),
            "blockingCount": len(blocking),
            "blockingTasks": blocking,
            "summary": summarize_business_gap_plan(plan),
        }

    @staticmethod
    def _find_business_gap_task(plan: dict[str, Any], task_id: str) -> dict[str, Any]:
        target = str(task_id or "")
        for task in plan.get("tasks") or []:
            if not isinstance(task, dict):
                continue
            if target in {str(task.get("id") or ""), str(task.get("taskKey") or "")}:
                return task
        raise KeyError(task_id)

    @staticmethod
    def _find_business_gap_toc_ref(plan: dict[str, Any], node_id: str) -> dict[str, Any]:
        target = str(node_id or "")
        for ref in plan.get("tocRefs") or []:
            if isinstance(ref, dict) and str(ref.get("nodeId") or "") == target:
                return ref
        raise KeyError(node_id)

    @staticmethod
    def _update_business_gap_toc_ref_statuses(plan: dict[str, Any]) -> None:
        task_by_id = {
            str(task.get("id") or ""): task
            for task in plan.get("tasks") or []
            if isinstance(task, dict) and str(task.get("id") or "")
        }
        for ref in plan.get("tocRefs") or []:
            if not isinstance(ref, dict):
                continue
            bound = [task_by_id[task_id] for task_id in [str(item) for item in ref.get("taskIds") or []] if task_id in task_by_id]
            if not bound:
                ref["status"] = "empty"
            elif all(str(task.get("status") or "") in {"ready", "resolved", "ignored"} for task in bound):
                ref["status"] = "ready"
            elif any(str(task.get("status") or "") == "review_required" for task in bound):
                ref["status"] = "review_required"
            else:
                ref["status"] = "partial"

    @staticmethod
    def _find_business_gap_artifact_in_task(task: dict[str, Any], artifact_id: str) -> dict[str, Any]:
        target = str(artifact_id or "")
        for artifact in task.get("resolvedArtifacts") or []:
            if isinstance(artifact, dict) and str(artifact.get("artifactId") or "") == target:
                return artifact
        raise KeyError(artifact_id)

    @staticmethod
    def _business_gap_artifact_matches_assembly_mode(task: dict[str, Any], artifact: dict[str, Any]) -> bool:
        mode = str(task.get("assemblyMode") or "")
        if not mode:
            return True
        artifact_mode = str(artifact.get("assemblyMode") or "")
        usage = str(artifact.get("wikiUsageMode") or artifact.get("materialUsage") or "")
        file_name = str(artifact.get("fileName") or artifact.get("materialName") or artifact.get("filePath") or "")
        suffix = Path(file_name).suffix.lower()
        source_mode = str(artifact.get("sourceMode") or "")
        artifact_type = str(artifact.get("artifactType") or "")
        if mode == "template_fill_docx":
            if artifact_mode == "template_fill_docx" or usage == "fill_template":
                return True
            if artifact_type == "parse_appendix_template" or source_mode == "parsed_from_tender_attachment_template":
                return True
            return suffix in {".doc", ".docx"} and bool(re.search(r"模板|格式|底稿|投标函|授权|廉洁|承诺|声明|说明|履约", file_name))
        if mode == "table_fill_from_material":
            return artifact_mode == mode or usage == "fill_table" or suffix in {".xls", ".xlsx", ".xlsm"}
        if mode == "embed_scan_or_image":
            return artifact_mode == mode or usage == "embed_scan" or suffix in {".pdf", ".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
        if mode == "attach_whole_file":
            return artifact_mode == mode or usage == "attach_whole"
        if mode == "extract_segment":
            return artifact_mode == mode or usage == "extract_segment" or bool(artifact.get("evidenceSegmentId"))
        if mode == "ai_draft":
            return artifact_mode == mode or usage == "generate_draft" or source_mode == "generated_by_business_s3_ai_draft"
        if mode == "manual_upload":
            return artifact_mode == mode or usage == "manual_upload" or source_mode == "uploaded_in_business_s3"
        return True

    @staticmethod
    def _business_gap_candidate_matches_assembly_mode(task: dict[str, Any], candidate: dict[str, Any]) -> bool:
        mode = str(task.get("assemblyMode") or "")
        if not mode:
            return True
        usage = str(candidate.get("wikiUsageMode") or candidate.get("materialUsage") or "")
        file_name = str(candidate.get("materialName") or candidate.get("name") or candidate.get("fileName") or "")
        folder_path = str(candidate.get("folderPath") or candidate.get("path") or "")
        text = f"{file_name} {folder_path}"
        suffix = Path(file_name).suffix.lower()
        if mode == "template_fill_docx":
            return usage == "fill_template" or (
                suffix in {"", ".doc", ".docx"} and bool(re.search(r"模板|格式|底稿|投标函|授权|廉洁|承诺|声明|说明|履约", text))
            )
        if mode == "table_fill_from_material":
            return usage == "fill_table" or suffix in {".xls", ".xlsx", ".xlsm"} or bool(re.search(r"表|清单|报价|价格|规格|偏差", text))
        if mode == "embed_scan_or_image":
            return usage == "embed_scan" or suffix in {".pdf", ".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"} or bool(re.search(r"证书|扫描|图片|回单|保函", text))
        if mode == "attach_whole_file":
            return usage == "attach_whole"
        if mode == "extract_segment":
            return usage == "extract_segment" or bool(candidate.get("evidenceSegmentId"))
        if mode == "ai_draft":
            return usage == "generate_draft"
        if mode == "manual_upload":
            return usage == "manual_upload"
        return True

    @classmethod
    def _business_task_wants_template_candidates(cls, task: dict[str, Any]) -> bool:
        mode = str(task.get("assemblyMode") or "")
        decision = str(task.get("decision") or "")
        title = str(task.get("title") or "")
        if mode == "template_fill_docx":
            return True
        if decision == "fill_required":
            return True
        return cls._default_business_task_assembly_mode(
            str(task.get("moduleKey") or ""),
            str(task.get("taskType") or ""),
            decision,
            title,
        ) == "template_fill_docx"

    @staticmethod
    def _business_task_can_ai_draft(task: dict[str, Any]) -> bool:
        task_type = str(task.get("taskType") or "")
        decision = str(task.get("decision") or "")
        assignee = str(task.get("assigneeMode") or "")
        title = str(task.get("title") or "")
        if task_type in {"certificate", "attachment", "bundle"}:
            return False
        if decision == "fill_required" or assignee in {"generate_or_fill", "ai_draft"}:
            return True
        return bool(re.search(r"投标函|授权|廉洁|承诺|声明|说明|专用章|履约", title))

    @staticmethod
    def _default_business_task_assembly_mode(module_key: str, task_type: str, decision: str, title: str) -> str:
        normalized = re.sub(r"[\s:：，,。.！!？?\-_/'\"“”‘’·（）()【】\[\]/]+", "", str(title or ""))
        if decision == "ai_draft_required":
            return "ai_draft"
        if task_type == "table":
            return "table_fill_from_material"
        if task_type == "certificate":
            return "embed_scan_or_image"
        if task_type == "attachment":
            if any(token in normalized for token in ["扫描", "回单", "保函", "证书", "图片", "pdf"]):
                return "embed_scan_or_image"
            return "attach_whole_file"
        if task_type == "bundle":
            if module_key in {"enterprise_finance_supply", "performance_cooperation_support"}:
                return "extract_segment"
            return "attach_whole_file"
        if decision == "fill_required":
            return "template_fill_docx"
        if any(token in normalized for token in ["投标函", "授权", "廉洁", "承诺", "声明", "说明", "履约"]):
            return "template_fill_docx"
        return "manual_upload"

    @staticmethod
    def _business_material_usage_for_assembly_mode(assembly_mode: str) -> str:
        return {
            "template_fill_docx": "fill_template",
            "table_fill_from_material": "fill_table",
            "attach_whole_file": "attach_whole",
            "embed_scan_or_image": "embed_scan",
            "extract_segment": "extract_segment",
            "ai_draft": "generate_draft",
            "manual_upload": "manual_upload",
        }.get(str(assembly_mode or ""), "")

    @classmethod
    def _business_task_fill_plan(cls, task: dict[str, Any]) -> dict[str, Any]:
        mode = str(task.get("assemblyMode") or cls._default_business_task_assembly_mode(
            str(task.get("moduleKey") or ""),
            str(task.get("taskType") or ""),
            str(task.get("decision") or ""),
            str(task.get("title") or ""),
        ))
        plan = {
            "mode": mode,
            "materialUsage": str(task.get("materialUsage") or cls._business_material_usage_for_assembly_mode(mode)),
            "requiresProjectFacts": mode in {"template_fill_docx", "table_fill_from_material", "ai_draft"},
            "requiresMaterialEvidence": mode in {"table_fill_from_material", "extract_segment", "attach_whole_file", "embed_scan_or_image"},
            "requiresHumanReview": True,
            "outputArtifactType": {
                "template_fill_docx": "filled_docx",
                "table_fill_from_material": "filled_table_docx",
                "attach_whole_file": "whole_file_attachment",
                "embed_scan_or_image": "embedded_scan",
                "extract_segment": "extracted_segment_docx",
                "ai_draft": "ai_draft_docx",
            }.get(mode, "manual_artifact"),
        }
        if task.get("selectedEvidenceSegments"):
            plan["evidenceSegmentCount"] = len(task.get("selectedEvidenceSegments") or [])
        return plan

    @classmethod
    def _assembly_mode_for_business_artifact(cls, task: dict[str, Any], artifact: dict[str, Any]) -> str:
        usage = str(artifact.get("wikiUsageMode") or artifact.get("materialUsage") or "").strip().lower().replace("-", "_")
        usage_modes = {
            "attach_whole": "attach_whole_file",
            "attach_whole_file": "attach_whole_file",
            "embed_scan": "embed_scan_or_image",
            "embed_scan_or_image": "embed_scan_or_image",
            "extract_segment": "extract_segment",
            "fill_template": "template_fill_docx",
            "template_fill_docx": "template_fill_docx",
            "fill_table": "table_fill_from_material",
            "table_fill_from_material": "table_fill_from_material",
        }
        if usage in usage_modes:
            return usage_modes[usage]
        suffix = Path(str(artifact.get("fileName") or artifact.get("filePath") or "")).suffix.lower()
        mime_type = str(artifact.get("mimeType") or "").lower()
        task_mode = str(task.get("assemblyMode") or "")
        if suffix in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff", ".pdf"} or "image/" in mime_type or "pdf" in mime_type:
            if task_mode in {"attach_whole_file", "extract_segment"}:
                return task_mode
            return "embed_scan_or_image"
        if task_mode:
            return task_mode
        return cls._default_business_task_assembly_mode(
            str(task.get("moduleKey") or ""),
            str(task.get("taskType") or ""),
            str(task.get("decision") or ""),
            str(task.get("title") or ""),
        )

    @classmethod
    def _apply_business_task_artifact_intent(cls, task: dict[str, Any], artifacts: list[dict[str, Any]]) -> None:
        selected = next((artifact for artifact in artifacts if isinstance(artifact, dict)), None)
        if selected:
            task["assemblyMode"] = str(selected.get("assemblyMode") or task.get("assemblyMode") or "")
            task["materialUsage"] = str(selected.get("materialUsage") or task.get("materialUsage") or "")
        if not task.get("assemblyMode"):
            task["assemblyMode"] = cls._default_business_task_assembly_mode(
                str(task.get("moduleKey") or ""),
                str(task.get("taskType") or ""),
                str(task.get("decision") or ""),
                str(task.get("title") or ""),
            )
        if not task.get("materialUsage"):
            task["materialUsage"] = cls._business_material_usage_for_assembly_mode(str(task.get("assemblyMode") or ""))
        segments: list[dict[str, Any]] = []
        seen: set[str] = set()
        for artifact in artifacts:
            if not isinstance(artifact, dict):
                continue
            segment_id = str(artifact.get("evidenceSegmentId") or "").strip()
            if not segment_id or segment_id in seen:
                continue
            seen.add(segment_id)
            segments.append(
                {
                    "segmentId": segment_id,
                    "title": str(artifact.get("evidenceSegmentTitle") or ""),
                    "materialId": str(artifact.get("materialId") or ""),
                    "materialName": str(artifact.get("materialName") or ""),
                    "wikiCardId": str(artifact.get("wikiCardId") or ""),
                    "sourcePages": str(artifact.get("evidenceSourcePages") or ""),
                    "summary": str(artifact.get("evidenceSummary") or ""),
                    "usageMode": str(artifact.get("wikiUsageMode") or artifact.get("materialUsage") or ""),
                }
            )
        if segments:
            current = [item for item in task.get("selectedEvidenceSegments") or [] if isinstance(item, dict)]
            incoming_ids = {str(item.get("segmentId") or "") for item in segments if str(item.get("segmentId") or "")}
            by_id = {
                str(item.get("segmentId") or ""): item
                for item in current
                if str(item.get("segmentId") or "") and str(item.get("segmentId") or "") not in incoming_ids
            }
            for segment in segments:
                by_id[str(segment.get("segmentId") or "")] = segment
            task["selectedEvidenceSegments"] = [*segments, *by_id.values()][:12]
        task["fillPlan"] = cls._business_task_fill_plan(task)

    @classmethod
    def _record_business_material_feedback(
        cls,
        business_gap_state: dict[str, Any],
        task: dict[str, Any],
        artifacts: list[dict[str, Any]],
        *,
        operator: str,
        selected_at: str,
    ) -> None:
        entries = business_gap_state.setdefault("materialFeedback", [])
        if not isinstance(entries, list):
            entries = []
            business_gap_state["materialFeedback"] = entries
        existing_keys = {
            str(item.get("feedbackKey") or "")
            for item in entries
            if isinstance(item, dict) and str(item.get("feedbackKey") or "")
        }
        task_title = str(task.get("title") or "")
        toc_target = task.get("tocTarget") if isinstance(task.get("tocTarget"), dict) else {}
        for artifact in artifacts:
            if not isinstance(artifact, dict):
                continue
            material_id = str(artifact.get("materialId") or "").strip()
            segment_id = str(artifact.get("evidenceSegmentId") or "").strip()
            if not material_id:
                continue
            feedback_key = cls._business_material_feedback_key(task, material_id, segment_id)
            if feedback_key in existing_keys:
                continue
            existing_keys.add(feedback_key)
            entries.append(
                {
                    "schemaVersion": "bid-business-material-feedback-item-v1",
                    "feedbackKey": feedback_key,
                    "taskId": str(task.get("id") or ""),
                    "taskKey": str(task.get("taskKey") or ""),
                    "taskTitle": task_title,
                    "moduleKey": str(task.get("moduleKey") or ""),
                    "taskType": str(task.get("taskType") or ""),
                    "assemblyMode": str(task.get("assemblyMode") or artifact.get("assemblyMode") or ""),
                    "materialUsage": str(task.get("materialUsage") or artifact.get("materialUsage") or ""),
                    "tocNodeId": str(toc_target.get("nodeId") or ""),
                    "tocNumber": str(toc_target.get("number") or ""),
                    "tocTitle": str(toc_target.get("title") or ""),
                    "materialId": material_id,
                    "materialName": str(artifact.get("materialName") or artifact.get("fileName") or ""),
                    "folderPath": str(artifact.get("folderPath") or ""),
                    "materialTier": str(artifact.get("materialTier") or ""),
                    "wikiCardId": str(artifact.get("wikiCardId") or ""),
                    "evidenceSegmentId": segment_id,
                    "evidenceSegmentTitle": str(artifact.get("evidenceSegmentTitle") or ""),
                    "evidenceSummary": str(artifact.get("evidenceSummary") or ""),
                    "selectedAt": selected_at,
                    "operator": operator or "当前用户",
                    "source": "business_s3_manual_selection",
                }
            )
        business_gap_state["materialFeedback"] = entries[-300:]

    @staticmethod
    def _business_material_feedback_key(task: dict[str, Any], material_id: str, segment_id: str) -> str:
        raw = "|".join(
            [
                str(task.get("moduleKey") or ""),
                str(task.get("taskType") or ""),
                str(task.get("title") or ""),
                str(material_id or ""),
                str(segment_id or ""),
            ]
        )
        import hashlib

        return "biz-feedback-" + hashlib.sha1(raw.encode("utf-8")).hexdigest()[:12]

    @classmethod
    def _write_business_ai_draft_docx(
        cls,
        output_path: Path,
        project: dict[str, Any],
        task: dict[str, Any],
        facts: dict[str, str],
        data: dict[str, Any],
    ) -> None:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        title = str(task.get("title") or "商务响应文件")
        normalized = title.replace("函格式", "函").replace("格式", "").strip() or "商务响应文件"
        doc = Document()
        doc.add_heading(normalized, level=1)

        tenderer = cls._fact_value(facts, "招标人") or str(project.get("customerName") or project.get("owner") or "")
        bidder = cls._fact_value(facts, "投标人") or str(data.get("bidderName") or "投标人")
        project_name = cls._fact_value(facts, "项目名称") or str(project.get("name") or "")
        tender_no = cls._fact_value(facts, "招标编号") or str(project.get("projectCode") or project.get("id") or "")
        date_value = cls._fact_value(facts, "日期") or "    年    月    日"

        if tenderer:
            doc.add_paragraph(f"致：{tenderer}")
        body = cls._business_ai_draft_body(normalized, project_name, tender_no)
        for paragraph in body:
            doc.add_paragraph(paragraph)

        doc.add_paragraph(f"项目名称：{project_name or '[待填写：项目名称]'}")
        doc.add_paragraph(f"招标编号：{tender_no or '[待填写：招标编号]'}")
        if "报价" in normalized or "价格" in normalized:
            doc.add_paragraph(f"投标报价：{cls._fact_value(facts, '投标报价') or '[待填写：投标报价]'}")
            doc.add_paragraph(f"币种：{cls._fact_value(facts, '币种') or '人民币'}")
        if "规格" in normalized or "供货" in normalized:
            doc.add_paragraph(f"投标机型：{cls._fact_value(facts, '投标机型') or '[待填写：投标机型]'}")
            doc.add_paragraph(f"总装机容量：{cls._fact_value(facts, '总装机容量') or '[待填写：总装机容量]'}")
        doc.add_paragraph("")
        doc.add_paragraph(f"投标人：{bidder}")
        doc.add_paragraph("法定代表人或授权代表：________________")
        doc.add_paragraph("日期：" + date_value)
        doc.add_paragraph("")
        doc.add_paragraph("注：本稿由商务标 S3 根据目录任务和项目事实表受控起草，签字、盖章、报价、金额及承诺范围须人工复核。")
        doc.save(output_path)

    @staticmethod
    def _business_ai_draft_body(title: str, project_name: str, tender_no: str) -> list[str]:
        project_label = project_name or "本项目"
        if "授权" in title:
            return [
                f"我单位现授权本单位工作人员作为 {project_label}（招标编号：{tender_no or '待填写'}）投标事宜的合法代理人。",
                "代理人在本项目投标、澄清、签署相关文件及处理投标事务中的行为，我单位均予以认可并承担相应法律责任。",
            ]
        if "廉洁" in title:
            return [
                f"我单位参加 {project_label} 投标活动，承诺严格遵守国家法律法规、招标文件及廉洁从业要求。",
                "我单位不以任何不正当方式影响评审结果，不向相关人员输送利益，并接受招标人及监管机构监督。",
            ]
        if "承诺" in title or "声明" in title or "说明" in title:
            return [
                f"我单位已充分理解 {project_label} 招标文件要求，并承诺按招标文件、投标文件及后续合同约定履行相关义务。",
                "如本承诺与招标文件专用条款或澄清文件存在不一致，以经人工复核后的最终投标文件为准。",
            ]
        if "履约" in title:
            return [
                f"我单位承诺在 {project_label} 中标后，按招标文件和合同约定及时提交履约保证金或履约保函。",
                "若未按要求提交，我单位愿承担招标文件及合同约定的相关责任。",
            ]
        if "投标函" in title:
            return [
                f"我单位已认真阅读 {project_label} 招标文件及其补遗、澄清文件，愿按招标文件要求提交商务投标文件。",
                "我单位承诺投标文件所载内容真实、准确、完整，并在投标有效期内保持投标文件有效。",
            ]
        return [
            f"我单位针对 {project_label} 编制本商务响应文件。",
            "本文件内容依据招标文件目录任务和项目事实表生成，需结合最终招标要求人工复核完善。",
        ]

    @staticmethod
    def _fact_value(facts: dict[str, str], label: str) -> str:
        key = re.sub(r"[\s　：:（）()【】\\[\\]]+", "", str(label or ""))
        if label in facts:
            return facts[label]
        if key in facts:
            return facts[key]
        for existing, value in facts.items():
            if key and (key in existing or existing in key):
                return value
        return ""

    @staticmethod
    def _default_business_gap_material_target_path(project_id: str, task: dict[str, Any]) -> str:
        root = f"商务标/项目素材/{safe_filename(project_id, 'project')}"
        module_key = str(task.get("moduleKey") or "")
        task_type = str(task.get("taskType") or "")
        sub_module = str(task.get("subModuleKey") or "")
        title = str(task.get("title") or "")
        if module_key == "qualification_compliance_certificates" and (
            sub_module == "special_certificates" or task_type == "certificate" or "认证" in title or "证书" in title
        ):
            return f"{root}/02-商务响应文件/专题证书"
        if module_key in {"base_documents_guarantees", "structured_response_tables", "commitments_and_notes"}:
            return f"{root}/02-商务响应文件"
        if module_key == "performance_cooperation_support":
            return f"{root}/01-客户关系与专项证明"
        return f"{root}/03-模板底稿与过程文件"

    @staticmethod
    def _decode_business_gap_upload_content(content: str, *, fallback_mime: str = "") -> tuple[bytes, str]:
        text = str(content or "").strip()
        if not text:
            return b"", fallback_mime
        if text.startswith("data:"):
            header, separator, payload = text.partition(",")
            if not separator:
                return b"", fallback_mime
            mime_type = header[5:].split(";", 1)[0] or fallback_mime
            if ";base64" in header.lower():
                try:
                    return base64.b64decode(payload, validate=True), mime_type
                except (binascii.Error, ValueError) as exc:
                    raise ValueError("上传文件 Base64 内容无法解析。") from exc
            return payload.encode("utf-8"), mime_type
        try:
            return base64.b64decode(text, validate=True), fallback_mime
        except (binascii.Error, ValueError):
            return text.encode("utf-8"), fallback_mime or "text/plain"

    @staticmethod
    def _unique_business_gap_path(directory: Path, file_name: str) -> Path:
        target = directory / safe_filename(file_name, "artifact.bin")
        if not target.exists():
            return target
        stem = target.stem or "artifact"
        suffix = target.suffix
        for index in range(2, 1000):
            candidate = directory / f"{stem}-{index}{suffix}"
            if not candidate.exists():
                return candidate
        raise ValueError(f"无法生成不重复的文件名：{file_name}")

    @staticmethod
    def _selected_business_material_entries(data: dict[str, Any]) -> list[dict[str, Any]]:
        raw_entries = data.get("materials")
        if raw_entries is None:
            raw_entries = data.get("materialIds") or data.get("referenceMaterialIds") or []
        entries: list[dict[str, Any]] = []
        for entry in raw_entries if isinstance(raw_entries, list) else []:
            if isinstance(entry, str):
                entries.append({"materialId": entry})
            elif isinstance(entry, dict):
                entries.append(copy.deepcopy(entry))
        return entries

    @staticmethod
    def _selected_business_template_entry(task: dict[str, Any], data: dict[str, Any]) -> dict[str, Any]:
        raw = data.get("template") if isinstance(data.get("template"), dict) else data
        template_id = str(raw.get("templateId") or "").strip() if isinstance(raw, dict) else ""
        file_path = str(raw.get("filePath") or raw.get("path") or "").strip() if isinstance(raw, dict) else ""
        for candidate in task.get("templateCandidates") or []:
            if not isinstance(candidate, dict):
                continue
            if template_id and template_id == str(candidate.get("templateId") or ""):
                return copy.deepcopy(candidate)
            if file_path and file_path == str(candidate.get("filePath") or candidate.get("path") or ""):
                return copy.deepcopy(candidate)
        if isinstance(raw, dict) and file_path:
            return copy.deepcopy(raw)
        raise ValueError("请选择一份有效模板候选。")

    @staticmethod
    def _merge_business_material_refs(current: list[Any], incoming: list[dict[str, Any]]) -> list[dict[str, Any]]:
        merged: list[dict[str, Any]] = []
        seen: set[str] = set()
        for item in [*current, *incoming]:
            if not isinstance(item, dict):
                continue
            material_id = str(item.get("materialId") or item.get("id") or "").strip()
            if not material_id or material_id in seen:
                continue
            seen.add(material_id)
            merged.append(copy.deepcopy(item))
        return merged

    @classmethod
    def _recompute_business_gap_task_after_artifact_change(cls, task: dict[str, Any]) -> None:
        flags = task.setdefault("riskFlags", [])
        if not isinstance(flags, list):
            flags = []
            task["riskFlags"] = flags
        for stale_flag in (
            "missing_material",
            "manual_upload_required",
            "template_missing_for_fill",
            "candidate_template_unconfirmed",
        ):
            if stale_flag in flags:
                flags.remove(stale_flag)
        artifacts = [item for item in task.get("resolvedArtifacts") or [] if isinstance(item, dict)]
        mode_artifacts = [
            item
            for item in artifacts
            if cls._business_gap_artifact_matches_assembly_mode(task, item)
        ]
        confirmed_artifacts = [item for item in mode_artifacts if bool(item.get("confirmed"))]
        template_candidates = [item for item in task.get("templateCandidates") or [] if isinstance(item, dict)]
        candidates = [item for item in task.get("candidateMaterials") or [] if isinstance(item, dict)]
        mode_candidates = [
            item
            for item in candidates
            if cls._business_gap_candidate_matches_assembly_mode(task, item)
        ]
        if confirmed_artifacts:
            task["decision"] = "ready"
            task["status"] = "ready"
            return
        if mode_artifacts:
            task["decision"] = "review_required"
            task["status"] = "review_required"
            return
        if str(task.get("assemblyMode") or "") == "template_fill_docx" and template_candidates:
            task["decision"] = "review_required"
            task["status"] = "review_required"
            if "candidate_template_unconfirmed" not in flags:
                flags.append("candidate_template_unconfirmed")
            return
        if str(task.get("assemblyMode") or "") == "ai_draft":
            task["decision"] = "ai_draft_required"
            task["status"] = "needs_input"
            if "ai_draft_required" not in flags:
                flags.append("ai_draft_required")
            return
        if mode_candidates:
            task["decision"] = "review_required"
            task["status"] = "review_required"
            if str(task.get("assemblyMode") or "") == "template_fill_docx" and "candidate_template_unconfirmed" not in flags:
                flags.append("candidate_template_unconfirmed")
            return
        if str(task.get("assemblyMode") or "") == "template_fill_docx":
            task["decision"] = "fill_required"
            task["status"] = "needs_input"
            for flag in ("missing_material", "template_missing_for_fill"):
                if flag not in flags:
                    flags.append(flag)
            return
        if str(task.get("sourceType") or "") == "manual_user":
            task["decision"] = "material_required"
            task["status"] = "needs_input"
            for flag in ("missing_material", "manual_upload_required"):
                if flag not in flags:
                    flags.append(flag)
            return
        if str(task.get("taskType") or "") in {"attachment", "certificate", "bundle"} or str(task.get("decision") or "") == "material_required":
            task["decision"] = "material_required"
            task["status"] = "needs_input"
            if "missing_material" not in flags:
                flags.append("missing_material")
            return
        task["decision"] = "review_required"
        task["status"] = "review_required"

    def _refresh_business_gap_template_candidates(
        self,
        project: dict[str, Any],
        business_gap_state: dict[str, Any],
    ) -> bool:
        plan = business_gap_state.get("plan") if isinstance(business_gap_state.get("plan"), dict) else {}
        tasks = [task for task in plan.get("tasks") or [] if isinstance(task, dict)]
        if not tasks:
            return False

        project_id = str(project.get("id") or "")
        work_dir = business_workspace_dir(project_id) / "gaps"
        template_index = _business_template_index(project, work_dir)
        project_template_ids = {
            str(candidate.get("templateId") or candidate.get("filePath") or candidate.get("originalPath") or "")
            for candidate in template_index
            if isinstance(candidate, dict) and str(candidate.get("sourceMode") or "") == "project_uploaded_bid_template"
        }
        active_template_keys = {
            str(candidate.get("templateId") or candidate.get("filePath") or candidate.get("originalPath") or "")
            for candidate in template_index
            if isinstance(candidate, dict)
        }
        changed = False
        has_project_template = bool(project_template_ids)
        for task in tasks:
            if not self._business_task_wants_template_candidates(task):
                continue
            existing = [item for item in task.get("templateCandidates") or [] if isinstance(item, dict)]
            task_changed = False
            if active_template_keys:
                kept_existing: list[dict[str, Any]] = []
                for item in existing:
                    item_key = str(item.get("templateId") or item.get("filePath") or item.get("originalPath") or "")
                    source_mode = str(item.get("sourceMode") or "")
                    if item_key not in active_template_keys:
                        continue
                    if has_project_template and source_mode == "system_default_bid_template":
                        continue
                    kept_existing.append(item)
                if len(kept_existing) != len(existing):
                    existing = kept_existing
                    task_changed = True
                    changed = True
            elif existing:
                existing = []
                task_changed = True
                changed = True
            existing_keys = {
                str(item.get("templateId") or item.get("filePath") or item.get("originalPath") or "")
                for item in existing
                if str(item.get("templateId") or item.get("filePath") or item.get("originalPath") or "")
            }
            merged = [copy.deepcopy(item) for item in existing]
            for candidate in template_index:
                candidate_key = str(
                    candidate.get("templateId")
                    or candidate.get("filePath")
                    or candidate.get("originalPath")
                    or ""
                )
                if not candidate_key or candidate_key in existing_keys:
                    continue
                merged.append(copy.deepcopy(candidate))
                existing_keys.add(candidate_key)
                task_changed = True
                changed = True
            if not task_changed:
                continue
            task["templateCandidates"] = sorted(
                merged,
                key=lambda item: 0 if str(item.get("sourceMode") or "") == "project_uploaded_bid_template" else 1,
            )[:8]
            if task["templateCandidates"]:
                if str(task.get("assemblyMode") or "") != "template_fill_docx":
                    task["assemblyMode"] = "template_fill_docx"
                task["materialUsage"] = self._business_material_usage_for_assembly_mode("template_fill_docx")
                task["fillPlan"] = self._business_task_fill_plan(task)
            self._recompute_business_gap_task_after_artifact_change(task)

        if not changed:
            return False
        updated_at = now_iso()
        self._finalize_business_gap_plan_update(project, business_gap_state, plan, updated_at=updated_at)
        return True

    @staticmethod
    async def _business_material_download_payload(material_id: str) -> tuple[dict[str, Any], str]:
        try:
            payload = await material_store.raw_download_cleaned_content(material_id)
            return payload, "cleaned"
        except Exception:
            payload = await material_store.raw_download_content(material_id)
            return payload, "raw"

    def _finalize_business_gap_plan_update(
        self,
        project: dict[str, Any],
        business_gap_state: dict[str, Any],
        plan: dict[str, Any],
        *,
        updated_at: str,
    ) -> None:
        plan["updatedAt"] = updated_at
        self._update_business_gap_toc_ref_statuses(plan)
        plan["summary"] = summarize_business_gap_plan(plan)
        business_gap_state["plan"] = plan
        business_gap_state["submittedForReview"] = False
        business_gap_state["reviewConfirmed"] = False
        business_gap_state["reviewedAt"] = ""
        business_gap_state["integrity"] = self._business_gap_integrity(plan)
        plan["integrity"] = copy.deepcopy(business_gap_state["integrity"])
        project["updatedAt"] = updated_at
        self._persist_project(project)

    @staticmethod
    def _refresh_business_gap_urls_for_result(
        project_id: str,
        plan: dict[str, Any],
        *,
        browser_base_url: str = "",
        onlyoffice_base_url: str = "",
    ) -> None:
        refresh_business_gap_artifact_urls(
            project_id,
            plan,
            browser_base_url=browser_base_url,
            onlyoffice_base_url=onlyoffice_base_url,
        )

    def _build_business_gap_payload(self, project: dict[str, Any], business_gap_state: dict[str, Any]) -> dict[str, Any]:
        plan = copy.deepcopy(business_gap_state.get("plan") or {})
        summary = plan.get("summary") if isinstance(plan.get("summary"), dict) else summarize_business_gap_plan(plan)
        return {
            "status": business_gap_state["recognitionStatus"],
            "recognizedAt": business_gap_state["recognizedAt"],
            "submittedForReview": bool(business_gap_state["submittedForReview"]),
            "reviewConfirmed": bool(business_gap_state["reviewConfirmed"]),
            "summary": copy.deepcopy(summary),
            "plan": plan,
            "businessGapPlan": copy.deepcopy(plan),
            "tocRefs": copy.deepcopy(plan.get("tocRefs") or []),
            "tasks": copy.deepcopy(plan.get("tasks") or []),
            "moduleGroups": copy.deepcopy(plan.get("moduleGroups") or []),
            "integrity": copy.deepcopy(business_gap_state.get("integrity") or {}),
            "source": {
                "fromStage": "商务标缺口处理",
                "projectId": project["id"],
                "projectName": project["name"],
                "bidType": project.get("bidType") or "商务标",
            },
        }

    def _repair_gap_state_fill_task_skills(self, gap_state: dict[str, Any]) -> bool:
        plan = gap_state.get("plan") if isinstance(gap_state.get("plan"), dict) else {}
        repaired = normalize_gap_plan_fill_task_skills(plan)
        if not repaired:
            return False
        gap_state["plan"] = plan
        gap_state["items"] = self._legacy_gap_items_from_plan(plan)
        gap_state["integrity"] = check_gap_integrity(plan)
        plan["integrity"] = gap_state["integrity"]
        return True

    @staticmethod
    def _legacy_gap_items_from_plan(plan: dict[str, Any]) -> list[dict[str, Any]]:
        items: list[dict[str, Any]] = []
        for index, item in enumerate(plan.get("items") or [], start=1):
            status = str(item.get("status") or "")
            if status in {"matched", "structural"}:
                continue
            items.append(
                {
                    "id": str(item.get("id") or f"GAP-{index}"),
                    "section": str(item.get("section") or ""),
                    "title": str(item.get("title") or ""),
                    "desc": str(item.get("gapReason") or "请补充该目录项所需素材。"),
                    "priority": str(item.get("priority") or "medium"),
                    "bidType": "技术标",
                    "status": "resolved" if status == "resolved" else "skipped" if status == "ignored" else "pending",
                    "skipReason": str(item.get("skipReason") or ""),
                    "resolvedSource": str(item.get("resolvedSource") or ""),
                    "resolvedAt": str(item.get("resolvedAt") or ""),
                    "latestUploadAt": str(item.get("latestUploadAt") or ""),
                    "latestSubmissionId": str(item.get("latestSubmissionId") or ""),
                }
            )
        return items

    def _ensure_review_document_state(self, project: dict[str, Any]) -> dict[str, Any]:
        state = project.get("review_document_state")
        if not isinstance(state, dict):
            state = {}
            project["review_document_state"] = state
        state.setdefault("parseStatus", "idle")
        state.setdefault("parsedAt", "")
        state.setdefault("documentId", f"REV-DOC-{project['id']}")
        state.setdefault("sourceFileName", "")
        state.setdefault("fileName", f"{project['name']}_缺口处理确认预览.docx")
        state.setdefault("fileType", "docx")
        state.setdefault("content", "")
        state.setdefault("lastSavedAt", "")
        state.setdefault("version", 1)
        state.setdefault("documentKey", f"{project['id']}-s6-v1")
        return state

    def _default_review_document_state(self, project: dict[str, Any]) -> dict[str, Any]:
        return {
            "parseStatus": "idle",
            "parsedAt": "",
            "documentId": f"REV-DOC-{project['id']}",
            "sourceFileName": "",
            "fileName": f"{project['name']}_缺口处理确认预览.docx",
            "fileType": "docx",
            "content": "",
            "lastSavedAt": "",
            "version": 1,
            "documentKey": f"{project['id']}-s6-v1",
        }

    def _build_gap_items_from_outline(self, project: dict[str, Any]) -> list[dict[str, Any]]:
        outline_nodes = list((project.get("outline_state") or {}).get("nodes") or [])
        candidates = self._collect_outline_candidates(outline_nodes)
        if not candidates:
            candidates = [
                {"section": "项目概况", "title": "项目背景材料"},
                {"section": "技术方案", "title": "技术方案支撑材料"},
                {"section": "实施与保障", "title": "实施组织与资源证明"},
            ]

        items: list[dict[str, Any]] = []
        priorities = ["high", "high", "medium", "medium", "low", "low"]
        for index, candidate in enumerate(candidates[:6], start=1):
            items.append(
                {
                    "id": f"GAP-{index}",
                    "section": candidate["section"],
                    "title": candidate["title"],
                    "desc": f"请补充“{candidate['title']}”对应的佐证材料或说明。",
                    "priority": priorities[index - 1] if index - 1 < len(priorities) else "low",
                    "bidType": "技术标",
                    "status": "pending",
                    "skipReason": "",
                    "resolvedSource": "",
                    "resolvedAt": "",
                    "latestUploadAt": "",
                    "latestSubmissionId": "",
                }
            )
        return items

    def _collect_outline_candidates(
        self,
        nodes: list[dict[str, Any]],
        parents: list[str] | None = None,
    ) -> list[dict[str, str]]:
        result: list[dict[str, str]] = []
        parent_titles = parents or []
        for node in nodes:
            title = str(node.get("title") or "").strip()
            children = list(node.get("children") or [])
            current_parents = parent_titles + ([title] if title else [])
            if children:
                result.extend(self._collect_outline_candidates(children, current_parents))
                continue
            if not title:
                continue
            section = " / ".join(parent_titles[-2:] or [title])
            result.append({"section": section, "title": title})
        return result

    def _build_gap_detection_payload(self, project: dict[str, Any], gap_state: dict[str, Any]) -> dict[str, Any]:
        items = copy.deepcopy(gap_state["items"])
        gap_plan = copy.deepcopy(gap_state.get("plan") or {})
        plan_summary = gap_plan.get("summary") if isinstance(gap_plan.get("summary"), dict) else {}
        high_priority_count = sum(1 for item in items if item.get("priority") == "high")
        medium_priority_count = sum(1 for item in items if item.get("priority") == "medium")
        low_priority_count = max(0, len(items) - high_priority_count - medium_priority_count)
        return {
            "status": gap_state["recognitionStatus"],
            "recognizedAt": gap_state["recognizedAt"],
            "summary": {
                "totalMissing": len(items),
                "totalTocItems": int(plan_summary.get("totalTocItems") or 0),
                "matchedCount": int(plan_summary.get("matchedCount") or 0),
                "missingCount": int(plan_summary.get("missingCount") or len(items)),
                "resolvedCount": int(plan_summary.get("resolvedCount") or 0),
                "fillableTaskCount": int(plan_summary.get("fillableTaskCount") or 0),
                "highPriorityCount": high_priority_count,
                "mediumPriorityCount": medium_priority_count,
                "lowPriorityCount": low_priority_count,
            },
            "items": items,
            "gapPlan": gap_plan,
            "integrity": copy.deepcopy(gap_state.get("integrity") or {}),
            "source": {
                "fromStage": "缺口处理",
                "projectId": project["id"],
                "projectName": project["name"],
            },
        }

    @staticmethod
    def _refresh_gap_plan_artifact_urls(
        project_id: str,
        gap_plan: dict[str, Any],
        *,
        browser_base_url: str = "",
        onlyoffice_base_url: str = "",
    ) -> None:
        if not isinstance(gap_plan, dict):
            return
        for item in gap_plan.get("items") or []:
            if not isinstance(item, dict):
                continue
            for artifact in item.get("resolvedArtifacts") or []:
                if not isinstance(artifact, dict):
                    continue
                artifact_id = str(artifact.get("id") or "")
                file_name = str(artifact.get("fileName") or Path(str(artifact.get("path") or "")).name)
                if not artifact_id or not file_name:
                    continue
                artifact["onlyoffice"] = {
                    **(artifact.get("onlyoffice") if isinstance(artifact.get("onlyoffice"), dict) else {}),
                    **_artifact_onlyoffice_payload(
                        project_id=project_id,
                        artifact_id=artifact_id,
                        file_name=file_name,
                        browser_base_url=browser_base_url,
                        onlyoffice_base_url=onlyoffice_base_url,
                    ),
                }

    def _find_gap_item(self, gap_state: dict[str, Any], gap_id: str) -> dict[str, Any]:
        for item in gap_state["items"]:
            if item.get("id") == gap_id:
                return item
        raise KeyError(gap_id)

    @staticmethod
    def _find_gap_plan_item(gap_state: dict[str, Any], gap_id: str) -> dict[str, Any] | None:
        plan = gap_state.get("plan") if isinstance(gap_state.get("plan"), dict) else {}
        for item in plan.get("items") or []:
            if str(item.get("id") or "") == gap_id:
                return item
        return None

    def _review_summary(self, gap_state: dict[str, Any]) -> dict[str, int]:
        items = list(gap_state["items"])
        return {
            "total": len(items),
            "resolvedCount": sum(1 for item in items if item["status"] == "resolved"),
            "skippedCount": sum(1 for item in items if item["status"] == "skipped"),
            "pendingCount": sum(1 for item in items if item["status"] not in {"resolved", "skipped"}),
        }

    def _build_review_payload(self, project: dict[str, Any], gap_state: dict[str, Any]) -> dict[str, Any]:
        latest_submission_by_missing_id: dict[str, dict[str, Any]] = {}
        for submission in gap_state["submissions"]:
            missing_id = str(submission.get("missingId") or "")
            if missing_id and missing_id not in latest_submission_by_missing_id:
                latest_submission_by_missing_id[missing_id] = submission

        items = []
        for item in gap_state["items"]:
            items.append(
                {
                    "id": item["id"],
                    "section": item["section"],
                    "title": item["title"],
                    "bidType": item.get("bidType") or "技术标",
                    "status": item["status"],
                    "resolvedSource": item.get("resolvedSource") or "",
                    "skipReason": item.get("skipReason") or "",
                    "resolvedAt": item.get("resolvedAt") or "",
                    "priority": item.get("priority") or "low",
                    "submission": copy.deepcopy(latest_submission_by_missing_id.get(item["id"])),
                }
            )

        review_state = self._ensure_review_document_state(project)
        return {
            "status": "ready" if gap_state["submittedForReview"] else "idle",
            "confirmed": bool(gap_state["reviewConfirmed"]),
            "reviewedAt": gap_state["reviewedAt"] or "",
            "summary": self._review_summary(gap_state),
            "items": items,
            "source": {
                "fromStage": "缺口处理",
                "projectId": project["id"],
                "projectName": project["name"],
            },
            "parse": {
                "status": review_state["parseStatus"],
                "parsedAt": review_state["parsedAt"],
                "fileName": review_state["fileName"],
            },
        }

    def _review_source_file_name(self, gap_state: dict[str, Any]) -> str:
        for item in gap_state["items"]:
            if item.get("status") == "resolved" and item.get("resolvedSource"):
                return str(item["resolvedSource"])
        return "缺口处理结果.docx"

    def _build_review_document_content(self, project: dict[str, Any], gap_state: dict[str, Any]) -> str:
        resolved_items = [item for item in gap_state["items"] if item["status"] == "resolved"]
        skipped_items = [item for item in gap_state["items"] if item["status"] == "skipped"]
        lines = [
            f"# {project['name']}（缺口处理确认预览）",
            "",
            "该文档由缺口处理提交确认后自动生成，用于生成标书前预览。",
            "",
            f"- 已补录：{len(resolved_items)} 项",
            f"- 未补录：{len(skipped_items)} 项",
            "",
            "## 已补录素材",
        ]
        if not resolved_items:
            lines.append("- 无")
        else:
            for index, item in enumerate(resolved_items, start=1):
                lines.append(f"{index}. {item['title']}（{item['section']}）")
                lines.append(f"   - 来源：{item.get('resolvedSource') or '已补录'}")

        lines.extend(["", "## 未补录素材"])
        if not skipped_items:
            lines.append("- 无")
        else:
            for index, item in enumerate(skipped_items, start=1):
                lines.append(f"{index}. {item['title']}（{item['section']}）")
                lines.append(f"   - 原因：{item.get('skipReason') or '未填写'}")
        return "\n".join(lines)

    @staticmethod
    def _count_nodes(nodes: list[dict[str, Any]]) -> int:
        total = 0
        for node in nodes:
            total += 1
            total += AppStore._count_nodes(node.get("children") or [])
        return total

    @staticmethod
    def _format_duration(seconds: int) -> str:
        total = max(0, int(seconds or 0))
        minutes = total // 60
        remain_seconds = total % 60
        if minutes <= 0:
            return f"{remain_seconds}秒"
        return f"{minutes}分{remain_seconds}秒"

    @staticmethod
    def _format_file_size(size_bytes: int) -> str:
        safe_size = max(0, int(size_bytes or 0))
        if safe_size < 1024:
            return f"{safe_size} B"
        if safe_size < 1024 * 1024:
            return f"{safe_size / 1024:.1f} KB"
        return f"{safe_size / (1024 * 1024):.1f} MB"

    @staticmethod
    def _mode_to_coverage_status(mode: Any) -> str:
        mode_text = str(mode or "").strip()
        if mode_text == "generated":
            return "full"
        if mode_text == "generated_with_placeholder":
            return "partial"
        return "none"

    def _build_coverage_tree(
        self,
        nodes: list[dict[str, Any]],
        section_status_map: dict[str, str],
        inherited_status: str | None = None,
    ) -> list[dict[str, Any]]:
        tree: list[dict[str, Any]] = []
        for node in nodes:
            node_id = str(node.get("id") or "")
            node_title = str(node.get("title") or "").strip() or "未命名章节"
            current_status = section_status_map.get(node_id) or inherited_status
            children = self._build_coverage_tree(
                node.get("children") or [],
                section_status_map,
                current_status,
            )
            if children:
                avg = sum(int(child.get("coverage") or 0) for child in children) / len(children)
                tree.append(
                    {
                        "id": node_id,
                        "title": node_title,
                        "coverage": round(avg),
                        "children": children,
                    }
                )
                continue

            leaf_status = current_status or "none"
            tree.append(
                {
                    "id": node_id,
                    "title": node_title,
                    "status": leaf_status,
                    "coverage": {"full": 100, "partial": 50, "none": 0}.get(leaf_status, 0),
                    "children": [],
                }
            )
        return tree

    def _build_coverage_issue_lists(
        self,
        tree: list[dict[str, Any]],
        path: list[str] | None = None,
    ) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
        partial_items: list[dict[str, Any]] = []
        no_cover_items: list[dict[str, Any]] = []
        current_path = path or []
        for node in tree:
            next_path = [*current_path, str(node.get("title") or "未命名章节")]
            children = list(node.get("children") or [])
            if children:
                nested_partial, nested_none = self._build_coverage_issue_lists(children, next_path)
                partial_items.extend(nested_partial)
                no_cover_items.extend(nested_none)
                continue

            item = {
                "id": str(node.get("id") or ""),
                "title": str(node.get("title") or "未命名章节"),
                "nodeTitle": " / ".join(current_path) if current_path else str(node.get("title") or "未命名章节"),
                "status": str(node.get("status") or "none"),
            }
            if item["status"] == "partial":
                partial_items.append(item)
            elif item["status"] == "none":
                no_cover_items.append(item)
        return partial_items, no_cover_items

    def _summarize_coverage(self, tree: list[dict[str, Any]]) -> tuple[int, int, int]:
        full_cover = 0
        partial_cover = 0
        no_cover = 0

        def visit(node: dict[str, Any]) -> None:
            nonlocal full_cover, partial_cover, no_cover
            children = list(node.get("children") or [])
            if children:
                for child in children:
                    visit(child)
                return

            status = str(node.get("status") or "none")
            if status == "full":
                full_cover += 1
            elif status == "partial":
                partial_cover += 1
            else:
                no_cover += 1

        for node in tree:
            visit(node)
        return full_cover, partial_cover, no_cover


store = AppStore()
