from __future__ import annotations

import copy
import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document

from app.core.config import settings
from app.services.bid_parse_state import update_parse_result_state
from app.services.bid_type import BUSINESS_BID_TYPE
from app.services.business_material_store import business_material_store
from app.services.file_utils import safe_segment
from app.services.identity import build_project_identity
from app.services.material_folder_scope import project_material_root_path
from app.services.onlyoffice_documents import WORD_MEDIA_TYPE
from app.services.parse_profiles import BUSINESS_PARSE_PROFILE
from app.services.parsing import (
    materialize_parse_appendix_docx_assets,
    materialize_parse_business_commitment_letter_docx_assets,
)
from app.services.workspace_project_access import (
    get_workspace_project_runtime_state,
    persist_workspace_project_state,
    require_workspace_project_for_update,
)
from app.services.workspace_artifacts import workspace_parse_dir


BUSINESS_APPENDIX_MATERIAL_FOLDER = "02-商务响应文件"
BUSINESS_SCORING_MATERIAL_FOLDER = "03-模板底稿与过程文件"
BUSINESS_COMMITMENT_MATERIAL_FOLDER = "02-商务响应文件"
BUSINESS_PARSE_ASSET_SCHEMA_VERSION = "business-parse-assets-v1"


class BusinessParseAssetError(Exception):
    def __init__(self, status_code: int, detail: str) -> None:
        super().__init__(detail)
        self.status_code = status_code
        self.detail = detail


def _now_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def _safe_docx_material_name(title: str, fallback: str) -> str:
    name = safe_segment(title, fallback)
    if not name.lower().endswith(".docx"):
        name = f"{Path(name).stem or fallback}.docx"
    return name


def _read_file_bytes(path: Path) -> bytes:
    if not path.exists() or not path.is_file():
        raise BusinessParseAssetError(404, "解析产物文件不存在。")
    return path.read_bytes()


def _materialize_business_parse_result(project_id: str, parse_result: dict[str, Any]) -> dict[str, Any]:
    payload = materialize_parse_appendix_docx_assets(
        project_id,
        parse_result,
        bid_type=BUSINESS_PARSE_PROFILE.bid_type,
    )
    return materialize_parse_business_commitment_letter_docx_assets(
        project_id,
        payload,
        bid_type=BUSINESS_PARSE_PROFILE.bid_type,
    )


def _business_payload(project_id: str) -> tuple[dict[str, Any], dict[str, Any], dict[str, Any]]:
    project = get_workspace_project_runtime_state(
        project_id,
        bid_type=BUSINESS_PARSE_PROFILE.bid_type,
        not_found_error=KeyError,
        wrong_type_error=lambda _project_id: BusinessParseAssetError(400, "仅商务标解析产物支持该操作。"),
    )
    parse_result = copy.deepcopy(project.get("parse_result") if isinstance(project.get("parse_result"), dict) else {})
    if parse_result.get("status") != "completed":
        raise BusinessParseAssetError(400, "请先完成商务标解析。")
    payload = _materialize_business_parse_result(project_id, parse_result)
    structured = payload.get("structured") if isinstance(payload.get("structured"), dict) else {}
    if not isinstance(structured, dict):
        raise BusinessParseAssetError(400, "商务标解析结果缺少 structured 数据。")
    return project, payload, structured


def _update_structured_result_file(parse_storage: dict[str, Any], parse_result: dict[str, Any]) -> None:
    structured_path = Path(str(parse_storage.get("structuredResultPath") or ""))
    if not structured_path.exists() or not structured_path.is_file():
        return
    try:
        payload = json.loads(structured_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        payload = {}
    payload["items"] = copy.deepcopy(parse_result.get("items") or payload.get("items") or [])
    payload["structured"] = copy.deepcopy(parse_result.get("structured") or {})
    structured = payload.get("structured")
    if isinstance(structured, dict):
        payload["schemaVersion"] = str(structured.get("schemaVersion") or BUSINESS_PARSE_PROFILE.schema_version)
    else:
        payload["schemaVersion"] = BUSINESS_PARSE_PROFILE.schema_version
    structured_path.parent.mkdir(parents=True, exist_ok=True)
    structured_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def _persist_business_parse_result(project_id: str, parse_result: dict[str, Any]) -> dict[str, Any]:
    project = require_workspace_project_for_update(
        project_id,
        bid_type=BUSINESS_PARSE_PROFILE.bid_type,
        not_found_error=KeyError,
        wrong_type_error=lambda _project_id: BusinessParseAssetError(400, "仅商务标解析产物支持该操作。"),
    )
    parse_storage = copy.deepcopy(project.get("parse_storage") if isinstance(project.get("parse_storage"), dict) else {})
    if isinstance(parse_storage, dict):
        parse_storage["items"] = copy.deepcopy(parse_result.get("items") or parse_storage.get("items") or [])
        parse_storage["structured"] = copy.deepcopy(parse_result.get("structured") or {})
        _update_structured_result_file(parse_storage, parse_result)
    payload = update_parse_result_state(project, parse_result, parse_storage=parse_storage)
    persist_workspace_project_state(project)
    return payload


def _scoring_groups(structured: dict[str, Any]) -> dict[str, Any]:
    scoring = structured.get("scoringCriteria") if isinstance(structured, dict) else {}
    return scoring if isinstance(scoring, dict) else {}


def _business_scoring_row_count(scoring_groups: dict[str, Any]) -> int:
    total = 0
    for rows in scoring_groups.values():
        if isinstance(rows, list):
            total += len(rows)
    return total


def _write_business_scoring_docx(path: Path, scoring_groups: dict[str, Any]) -> int:
    path.parent.mkdir(parents=True, exist_ok=True)
    rows_total = 0
    doc = Document()
    doc.add_heading("商务评分标准", level=1)
    group_specs = (
        ("business", "商务评分标准"),
        ("price", "投标报价评分标准"),
        ("compliance", "符合性审查标准"),
    )
    handled = set()
    for key, title in group_specs:
        rows = scoring_groups.get(key) if isinstance(scoring_groups.get(key), list) else []
        handled.add(key)
        if not rows:
            continue
        rows_total += len(rows)
        doc.add_heading(title, level=2)
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = ["序号", "评分/审查项", "分值", "得分点/要求", "证明材料要求", "证据位置"]
        for index, header in enumerate(headers):
            table.cell(0, index).text = header
        for row_index, row in enumerate(rows, start=1):
            cells = table.add_row().cells
            cells[0].text = str(row.get("order") or row_index)
            cells[1].text = str(row.get("scoringItem") or "")
            cells[2].text = str(row.get("score") or "")
            cells[3].text = str(row.get("scorePoint") or row.get("evidence") or "")
            cells[4].text = str(row.get("proofRequirement") or "")
            cells[5].text = str(row.get("evidenceLocation") or "")

    for key, rows in scoring_groups.items():
        if key in handled or not isinstance(rows, list) or not rows:
            continue
        rows_total += len(rows)
        doc.add_heading(str(key), level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ["序号", "评分/审查项", "要求", "证据位置"]
        for index, header in enumerate(headers):
            table.cell(0, index).text = header
        for row_index, row in enumerate(rows, start=1):
            cells = table.add_row().cells
            cells[0].text = str(row.get("order") or row_index)
            cells[1].text = str(row.get("scoringItem") or row.get("title") or "")
            cells[2].text = str(row.get("scorePoint") or row.get("proofRequirement") or row.get("evidence") or "")
            cells[3].text = str(row.get("evidenceLocation") or "")

    if rows_total <= 0:
        raise BusinessParseAssetError(400, "未识别到可审核的商务评分标准。")
    doc.save(path)
    return rows_total


def _business_scoring_output_path(project_id: str) -> Path:
    return workspace_parse_dir(project_id, BUSINESS_PARSE_PROFILE) / "business-parse-assets" / "商务评分标准.docx"


def _appendix_material_file(appendix: dict[str, Any], path: Path) -> dict[str, Any]:
    title = str(appendix.get("title") or path.stem or "商务附件模板").strip()
    return {
        "name": _safe_docx_material_name(title, "商务附件模板"),
        "type": WORD_MEDIA_TYPE,
        "mimeType": WORD_MEDIA_TYPE,
        "data": _read_file_bytes(path),
        "relativePath": "",
    }


def _business_scoring_material_file(path: Path) -> dict[str, Any]:
    return {
        "name": "商务评分标准.docx",
        "type": WORD_MEDIA_TYPE,
        "mimeType": WORD_MEDIA_TYPE,
        "data": _read_file_bytes(path),
        "relativePath": "",
    }


def _commitment_material_file(letter: dict[str, Any], path: Path) -> dict[str, Any]:
    title = str(letter.get("title") or path.stem or "承诺函").strip()
    return {
        "name": _safe_docx_material_name(title, "承诺函"),
        "type": WORD_MEDIA_TYPE,
        "mimeType": WORD_MEDIA_TYPE,
        "data": _read_file_bytes(path),
        "relativePath": "",
    }


def _business_material_target(project: dict[str, Any], target_folder: str) -> tuple[dict[str, Any], str, str, str, str, str]:
    identity = build_project_identity(project)
    project_id = str(project.get("id") or "")
    material_project_id = safe_segment(str(identity.get("projectId") or project_id), project_id)
    if not material_project_id:
        raise BusinessParseAssetError(400, "项目素材 ID 为空，无法同步商务解析产物。")
    project_code = str(identity.get("projectCode") or material_project_id)
    project_name = str(identity.get("projectName") or project.get("name") or "")
    customer_name = str(identity.get("customerCanonicalName") or identity.get("customerName") or project.get("owner") or "")
    target_path = f"{project_material_root_path(BUSINESS_BID_TYPE, material_project_id)}/{target_folder}"
    return identity, target_path, material_project_id, project_code, project_name, customer_name


async def _upload_business_material_files(
    project: dict[str, Any],
    *,
    target_folder: str,
    files: list[dict[str, Any]],
    on_conflict: str = "version",
) -> dict[str, Any]:
    if not files:
        raise BusinessParseAssetError(400, "没有可同步的商务解析产物。")
    identity, target_path, material_project_id, project_code, project_name, customer_name = _business_material_target(
        project,
        target_folder,
    )
    try:
        return await business_material_store.raw_upload(
            target_path=target_path,
            project_id=material_project_id,
            project_code=project_code,
            project_name=project_name,
            material_tier="project",
            customer_id=str(identity.get("customerId") or ""),
            customer_name=customer_name,
            on_conflict=on_conflict,
            files=files,
        )
    except Exception as exc:
        detail = getattr(exc, "detail", None) or getattr(exc, "message", None) or str(exc)
        status_code = int(getattr(exc, "status_code", 500) or 500)
        raise BusinessParseAssetError(status_code, str(detail)) from exc


def approve_business_appendix_asset(project_id: str, appendix_id: str, *, approved: bool = True) -> dict[str, Any]:
    _project, parse_result, structured = _business_payload(project_id)
    appendices = structured.get("appendices") if isinstance(structured.get("appendices"), list) else []
    target: dict[str, Any] | None = None
    at = _now_iso()
    for appendix in appendices:
        if not isinstance(appendix, dict):
            continue
        if str(appendix.get("id") or "") != appendix_id:
            continue
        appendix["assetKind"] = "businessAppendixTemplate"
        appendix["assetReviewStatus"] = "approved" if approved else "pending_review"
        appendix["assetApprovedAt"] = at if approved else ""
        appendix["assetSchemaVersion"] = BUSINESS_PARSE_ASSET_SCHEMA_VERSION
        appendix["assetMaterialFolder"] = BUSINESS_APPENDIX_MATERIAL_FOLDER
        target = copy.deepcopy(appendix)
        break
    if target is None:
        raise BusinessParseAssetError(404, "未找到对应的商务附件模板。")
    parse_result["structured"] = structured
    persisted = _persist_business_parse_result(project_id, parse_result)
    message = "商务附件模板已审核通过；确认参与投标并完善项目信息后，将自动同步至项目素材库。"
    return {"message": message, "appendix": target, "parseResult": persisted}


def approve_all_business_appendix_assets(project_id: str, *, approved: bool = True) -> dict[str, Any]:
    _project, parse_result, structured = _business_payload(project_id)
    appendices = structured.get("appendices") if isinstance(structured.get("appendices"), list) else []
    at = _now_iso()
    count = 0
    for appendix in appendices:
        if not isinstance(appendix, dict):
            continue
        appendix["assetKind"] = "businessAppendixTemplate"
        appendix["assetReviewStatus"] = "approved" if approved else "pending_review"
        appendix["assetApprovedAt"] = at if approved else ""
        appendix["assetSchemaVersion"] = BUSINESS_PARSE_ASSET_SCHEMA_VERSION
        appendix["assetMaterialFolder"] = BUSINESS_APPENDIX_MATERIAL_FOLDER
        count += 1
    if count <= 0:
        raise BusinessParseAssetError(400, "未识别到可审核的商务附件模板。")
    parse_result["structured"] = structured
    persisted = _persist_business_parse_result(project_id, parse_result)
    return {
        "message": f"已审核通过 {count} 个商务附件模板；确认参与投标并完善项目信息后，将自动同步至项目素材库。",
        "approvedCount": count,
        "parseResult": persisted,
    }


def approve_business_scoring_asset(project_id: str, *, approved: bool = True) -> dict[str, Any]:
    _project, parse_result, structured = _business_payload(project_id)
    scoring = _scoring_groups(structured)
    row_count = _business_scoring_row_count(scoring)
    if row_count <= 0:
        raise BusinessParseAssetError(400, "未识别到可审核的商务评分标准。")
    at = _now_iso()
    asset = structured.get("businessScoringAsset") if isinstance(structured.get("businessScoringAsset"), dict) else {}
    asset.update(
        {
            "assetKind": "businessScoringCriteria",
            "assetSchemaVersion": BUSINESS_PARSE_ASSET_SCHEMA_VERSION,
            "status": "approved" if approved else "pending_review",
            "reviewStatus": "approved" if approved else "pending_review",
            "approvedAt": at if approved else "",
            "rowCount": row_count,
            "assetMaterialFolder": BUSINESS_SCORING_MATERIAL_FOLDER,
        }
    )
    structured["businessScoringAsset"] = asset
    parse_result["structured"] = structured
    persisted = _persist_business_parse_result(project_id, parse_result)
    return {
        "message": "商务评分标准已审核通过；确认参与投标并完善项目信息后，将自动同步至项目素材库。",
        "asset": copy.deepcopy(asset),
        "parseResult": persisted,
    }


def approve_business_commitment_letter_asset(project_id: str, letter_id: str, *, approved: bool = True) -> dict[str, Any]:
    _project, parse_result, structured = _business_payload(project_id)
    letters = structured.get("commitmentLetters") if isinstance(structured.get("commitmentLetters"), list) else []
    target: dict[str, Any] | None = None
    at = _now_iso()
    for letter in letters:
        if not isinstance(letter, dict):
            continue
        if str(letter.get("id") or "") != letter_id:
            continue
        letter["assetKind"] = "businessCommitmentLetter"
        letter["assetReviewStatus"] = "approved" if approved else "pending_review"
        letter["assetApprovedAt"] = at if approved else ""
        letter["assetSchemaVersion"] = BUSINESS_PARSE_ASSET_SCHEMA_VERSION
        letter["assetMaterialFolder"] = BUSINESS_COMMITMENT_MATERIAL_FOLDER
        target = copy.deepcopy(letter)
        break
    if target is None:
        raise BusinessParseAssetError(404, "未找到对应的承诺函。")
    parse_result["structured"] = structured
    persisted = _persist_business_parse_result(project_id, parse_result)
    message = "承诺函已审核通过；确认参与投标并完善项目信息后，将自动同步至项目素材库。"
    return {"message": message, "letter": target, "parseResult": persisted}


def approve_all_business_commitment_letter_assets(project_id: str, *, approved: bool = True) -> dict[str, Any]:
    _project, parse_result, structured = _business_payload(project_id)
    letters = structured.get("commitmentLetters") if isinstance(structured.get("commitmentLetters"), list) else []
    at = _now_iso()
    count = 0
    for letter in letters:
        if not isinstance(letter, dict):
            continue
        letter["assetKind"] = "businessCommitmentLetter"
        letter["assetReviewStatus"] = "approved" if approved else "pending_review"
        letter["assetApprovedAt"] = at if approved else ""
        letter["assetSchemaVersion"] = BUSINESS_PARSE_ASSET_SCHEMA_VERSION
        letter["assetMaterialFolder"] = BUSINESS_COMMITMENT_MATERIAL_FOLDER
        count += 1
    if count <= 0:
        raise BusinessParseAssetError(400, "未识别到可审核的承诺函。")
    parse_result["structured"] = structured
    persisted = _persist_business_parse_result(project_id, parse_result)
    return {
        "message": f"已审核通过 {count} 个承诺函；确认参与投标并完善项目信息后，将自动同步至项目素材库。",
        "approvedCount": count,
        "parseResult": persisted,
    }


def _mark_appendix_sync_failed(appendix: dict[str, Any], error: str) -> None:
    appendix["assetSyncStatus"] = "failed"
    appendix["assetSyncError"] = error
    appendix["assetSyncUpdatedAt"] = _now_iso()


def _mark_appendix_synced(appendix: dict[str, Any], uploaded_item: dict[str, Any]) -> None:
    appendix["assetSyncStatus"] = "synced"
    appendix["assetSyncedAt"] = _now_iso()
    appendix["assetSyncError"] = ""
    appendix["assetMaterialId"] = uploaded_item.get("id") or ""
    appendix["assetMaterialPath"] = "/".join(
        part for part in [str(uploaded_item.get("folderPath") or ""), str(uploaded_item.get("name") or "")] if part
    )
    appendix["assetMaterialItem"] = copy.deepcopy(uploaded_item)


def _mark_letter_sync_failed(letter: dict[str, Any], error: str) -> None:
    letter["assetSyncStatus"] = "failed"
    letter["assetSyncError"] = error
    letter["assetSyncUpdatedAt"] = _now_iso()


def _mark_letter_synced(letter: dict[str, Any], uploaded_item: dict[str, Any]) -> None:
    letter["assetSyncStatus"] = "synced"
    letter["assetSyncedAt"] = _now_iso()
    letter["assetSyncError"] = ""
    letter["assetMaterialId"] = uploaded_item.get("id") or ""
    letter["assetMaterialPath"] = "/".join(
        part for part in [str(uploaded_item.get("folderPath") or ""), str(uploaded_item.get("name") or "")] if part
    )
    letter["assetMaterialItem"] = copy.deepcopy(uploaded_item)


def _mark_scoring_sync_failed(asset: dict[str, Any], error: str) -> None:
    asset["syncStatus"] = "failed"
    asset["syncError"] = error
    asset["syncUpdatedAt"] = _now_iso()


def _mark_scoring_synced(asset: dict[str, Any], uploaded_item: dict[str, Any], *, output_path: Path) -> None:
    asset["syncStatus"] = "synced"
    asset["syncedAt"] = _now_iso()
    asset["syncError"] = ""
    asset["docxPath"] = str(output_path)
    asset["materialId"] = uploaded_item.get("id") or ""
    asset["materialPath"] = "/".join(
        part for part in [str(uploaded_item.get("folderPath") or ""), str(uploaded_item.get("name") or "")] if part
    )
    asset["materialItem"] = copy.deepcopy(uploaded_item)


async def sync_approved_business_parse_assets(project_id: str) -> dict[str, Any]:
    project, parse_result, structured = _business_payload(project_id)
    summary: dict[str, Any] = {
        "status": "skipped",
        "projectId": project_id,
        "syncedAppendixCount": 0,
        "failedAppendixCount": 0,
        "syncedCommitmentLetterCount": 0,
        "failedCommitmentLetterCount": 0,
        "syncedScoring": False,
        "messages": [],
    }

    appendices = structured.get("appendices") if isinstance(structured.get("appendices"), list) else []
    appendix_uploads: list[dict[str, Any]] = []
    appendix_refs: list[dict[str, Any]] = []
    for appendix in appendices:
        if not isinstance(appendix, dict):
            continue
        if str(appendix.get("assetReviewStatus") or "") != "approved":
            continue
        path = Path(str(appendix.get("docxPath") or ""))
        if not path.exists() or not path.is_file():
            _mark_appendix_sync_failed(appendix, "已审核附件模板的 Word 文件不存在。")
            summary["failedAppendixCount"] += 1
            continue
        try:
            appendix_uploads.append(_appendix_material_file(appendix, path))
            appendix_refs.append(appendix)
        except BusinessParseAssetError as exc:
            _mark_appendix_sync_failed(appendix, exc.detail)
            summary["failedAppendixCount"] += 1

    if appendix_uploads:
        try:
            result = await _upload_business_material_files(
                project,
                target_folder=BUSINESS_APPENDIX_MATERIAL_FOLDER,
                files=appendix_uploads,
                on_conflict="version",
            )
            uploaded_items = result.get("items") if isinstance(result.get("items"), list) else []
            for appendix, uploaded_item in zip(appendix_refs, uploaded_items, strict=False):
                _mark_appendix_synced(appendix, uploaded_item if isinstance(uploaded_item, dict) else {})
                summary["syncedAppendixCount"] += 1
            summary["messages"].append(result.get("message") or "商务附件模板已同步。")
        except BusinessParseAssetError as exc:
            for appendix in appendix_refs:
                _mark_appendix_sync_failed(appendix, exc.detail)
            summary["failedAppendixCount"] += len(appendix_refs)
            summary["messages"].append(f"商务附件模板同步失败：{exc.detail}")

    letters = structured.get("commitmentLetters") if isinstance(structured.get("commitmentLetters"), list) else []
    letter_uploads: list[dict[str, Any]] = []
    letter_refs: list[dict[str, Any]] = []
    for letter in letters:
        if not isinstance(letter, dict):
            continue
        if str(letter.get("assetReviewStatus") or "") != "approved":
            continue
        path = Path(str(letter.get("docxPath") or ""))
        if not path.exists() or not path.is_file():
            _mark_letter_sync_failed(letter, "已审核承诺函的 Word 文件不存在。")
            summary["failedCommitmentLetterCount"] += 1
            continue
        try:
            letter_uploads.append(_commitment_material_file(letter, path))
            letter_refs.append(letter)
        except BusinessParseAssetError as exc:
            _mark_letter_sync_failed(letter, exc.detail)
            summary["failedCommitmentLetterCount"] += 1

    if letter_uploads:
        try:
            result = await _upload_business_material_files(
                project,
                target_folder=BUSINESS_COMMITMENT_MATERIAL_FOLDER,
                files=letter_uploads,
                on_conflict="version",
            )
            uploaded_items = result.get("items") if isinstance(result.get("items"), list) else []
            for letter, uploaded_item in zip(letter_refs, uploaded_items, strict=False):
                _mark_letter_synced(letter, uploaded_item if isinstance(uploaded_item, dict) else {})
                summary["syncedCommitmentLetterCount"] += 1
            summary["messages"].append(result.get("message") or "承诺函已同步。")
        except BusinessParseAssetError as exc:
            for letter in letter_refs:
                _mark_letter_sync_failed(letter, exc.detail)
            summary["failedCommitmentLetterCount"] += len(letter_refs)
            summary["messages"].append(f"承诺函同步失败：{exc.detail}")

    scoring_asset = structured.get("businessScoringAsset") if isinstance(structured.get("businessScoringAsset"), dict) else {}
    if str(scoring_asset.get("reviewStatus") or scoring_asset.get("status") or "") == "approved":
        output_path = _business_scoring_output_path(project_id)
        try:
            row_count = _write_business_scoring_docx(output_path, _scoring_groups(structured))
            scoring_asset["rowCount"] = row_count
            result = await _upload_business_material_files(
                project,
                target_folder=BUSINESS_SCORING_MATERIAL_FOLDER,
                files=[_business_scoring_material_file(output_path)],
                on_conflict="version",
            )
            uploaded_items = result.get("items") if isinstance(result.get("items"), list) else []
            uploaded_item = uploaded_items[0] if uploaded_items and isinstance(uploaded_items[0], dict) else {}
            _mark_scoring_synced(scoring_asset, uploaded_item, output_path=output_path)
            structured["businessScoringAsset"] = scoring_asset
            summary["syncedScoring"] = True
            summary["messages"].append(result.get("message") or "商务评分标准已同步。")
        except BusinessParseAssetError as exc:
            _mark_scoring_sync_failed(scoring_asset, exc.detail)
            structured["businessScoringAsset"] = scoring_asset
            summary["messages"].append(f"商务评分标准同步失败：{exc.detail}")

    parse_result["structured"] = structured
    persisted = _persist_business_parse_result(project_id, parse_result)
    summary["parseResult"] = persisted
    has_synced = summary["syncedAppendixCount"] or summary["syncedCommitmentLetterCount"] or summary["syncedScoring"]
    has_failed = summary["failedAppendixCount"] or summary["failedCommitmentLetterCount"] or any("失败" in message for message in summary["messages"])
    if has_synced:
        summary["status"] = "partial" if has_failed else "synced"
    elif has_failed:
        summary["status"] = "failed"
    return summary
