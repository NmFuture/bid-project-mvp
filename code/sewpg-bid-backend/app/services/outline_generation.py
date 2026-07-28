from __future__ import annotations

import copy
import hashlib
import importlib.util
import json
import os
import re
import shutil
import sys
import threading
from collections import Counter
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable

from docx import Document

from app.core.config import settings
from app.services.bid_outline_state import save_generated_outline_state
from app.services.bid_project_state import project_parse_input_records
from app.services.bid_type import BUSINESS_BID_TYPE, require_bid_type
from app.services.opencode_client import OpencodeClient
from app.services.parsing import IMAGE_SUFFIXES, _ocr_fallback_text
from app.services.bid_runtime_state import build_directory_opencode_output, now_iso
from app.services.system_settings import system_settings_service
from app.services.template_store import is_valid_docx_file
from app.services.workspace_artifacts import workspace_dir
from app.services.workspace_project_access import (
    get_any_workspace_project_runtime_state,
    persist_workspace_project_state,
    require_any_workspace_project_for_update,
)
from app.services.turbine_models import project_turbine_model

OUTLINE_SKILL_NAME = "bid-tech-outline-generator"
BUSINESS_OUTLINE_SKILL_NAME = "bid-business-outline-generator"
BUSINESS_OUTLINE_SKILL_COMMAND = "business-outline"
TECH_OUTLINE_FINALIZE_COMMAND = "s2outline finalize"
TECH_OUTLINE_FINALIZE_EARLY_COMMAND = "s2outline-finalize"
TECH_OUTLINE_HANDOFF_DECISION_UNITS = 1
TECH_OUTLINE_CHAPTER_WORKERS = 6
_TECH_OUTLINE_REQUEST_SLOTS = threading.BoundedSemaphore(TECH_OUTLINE_CHAPTER_WORKERS)
PUBLIC_EVIDENCE_DECISION_LIMIT = 80


class _ChapterParallelUnsupported(RuntimeError):
    pass
TECHNICAL_SUGGESTION_ACTIONS = {"必要", "建议增加", "建议删除", "待确认"}

# 注入 S2 manifest 的事实表状态：已确认/已抽取/待人工确认的值可信可用；
# 未提取/缺来源/冲突/不适用于不注入。
MANIFEST_FACT_VALUE_STATUSES = {"confirmed", "extracted", "pending_confirmation"}


def project_facts_for_manifest(project: dict[str, Any]) -> dict[str, str]:
    """从 S3 项目事实表提取 label→value 映射，供 S2 manifest 注入。

    S2 通常早于 S3 事实表构建，gap_state 无事实表时返回空映射（调用方不写该键）。
    """
    gap_state = project.get("gap_state") if isinstance(project.get("gap_state"), dict) else {}
    fact_table = gap_state.get("projectFactTable") if isinstance(gap_state.get("projectFactTable"), dict) else {}
    facts: dict[str, str] = {}
    for field in fact_table.get("fields") or []:
        if not isinstance(field, dict):
            continue
        label = str(field.get("label") or "").strip()
        value = str(field.get("value") or "").strip()
        if label and value and str(field.get("status") or "") in MANIFEST_FACT_VALUE_STATUSES:
            facts.setdefault(label, value)
    return facts


def _load_technical_outline_runner() -> Any:
    module_name = "_sewpg_bid_technical_outline_runner"
    loaded = sys.modules.get(module_name)
    if loaded is not None:
        return loaded
    script_path = (
        Path(__file__).resolve().parents[2]
        / "opencode"
        / "skills"
        / OUTLINE_SKILL_NAME
        / "scripts"
        / "run_from_manifest.py"
    )
    spec = importlib.util.spec_from_file_location(module_name, script_path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"无法加载技术标目录 Skill：{script_path}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    return module


def _capture_trusted_technical_outline_input(manifest_path: Path) -> dict[str, Any]:
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    template_file = Path(str(manifest.get("templateFile") or "")).expanduser()
    output_file = Path(str(manifest.get("outputFile") or "")).expanduser()
    if not template_file.is_file():
        raise RuntimeError(f"技术标历史模板不存在：{template_file}")
    if not str(manifest.get("outputFile") or "").strip():
        raise RuntimeError("技术标目录 manifest 缺少 outputFile。")
    runner = _load_technical_outline_runner()
    try:
        structure = runner.extract_template_structure(template_file)
        raw_tender_files = manifest.get("tenderFiles")
        tender_files = copy.deepcopy(raw_tender_files) if isinstance(raw_tender_files, list) else []
        appendix_inventory = runner.extract_tender_appendix_inventory(tender_files)
        appendix_items = runner.review_workflow.decision_appendix_items_from_inventory(
            appendix_inventory
        )
        tender_inputs_digest = (
            runner.review_workflow.tender_input_fingerprint(tender_files)
            if tender_files
            else ""
        )
    except SystemExit as exc:
        raise RuntimeError(f"技术标可信输入提取失败：{exc}") from exc
    return {
        "templateFile": str(template_file.resolve()),
        "templateFileSha256": hashlib.sha256(template_file.read_bytes()).hexdigest(),
        "outputFile": str(output_file.resolve()),
        "templateStructure": copy.deepcopy(structure),
        "tenderFiles": tender_files,
        "tenderInputsDigest": tender_inputs_digest,
        "appendixItems": copy.deepcopy(appendix_items),
    }


def _finalize_current_technical_outline(manifest_path: Path) -> dict[str, Any]:
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    manifest["_runtimeRequireComposedOutline"] = True
    runner = _load_technical_outline_runner()
    try:
        return runner.finalize_manifest(manifest, manifest_path)
    except SystemExit as exc:
        raise RuntimeError(str(exc)) from exc


def _build_outline_handoff_prompt(manifest_path: Path, handoff_index: int) -> str:
    first_pass = handoff_index == 1
    startup = (
        "先执行一次 `s2outline prepare`，再将招标目录按 `next_cursor` 分页读到 complete=true。"
        if first_pass
        else (
            "步骤 1 和步骤 2 已由前序会话完整完成。继续使用现有决策状态；"
            "不要重复执行 prepare，不要执行 template-headings、headings、next-batch 或 review-batch，"
            "也不要重新读取全量模板或招标目录。直接从 decision-next 开始当前决策单元。"
        )
    )
    mandatory_start = (
        ""
        if first_pass
        else (
            "加载 Skill 后，第一条非 Skill 工具调用必须是 Bash，且 command 必须精确等于："
            f"`s2outline decision-next {manifest_path}`。"
            "禁止调用 Read、Glob、Grep，也不要执行 pwd、ls、cat 或搜索 manifest；"
            "manifest 路径已经给出，不需要探路。"
        )
    )
    return f"""
Use the {OUTLINE_SKILL_NAME} skill.

这是 S2 技术标目录的第 {handoff_index} 个受控接力会话。
manifest：{manifest_path}

{startup}
{mandatory_start}
本会话只做模板正文目录的自主判断：循环调用 `decision-next`，每个决策单元都按 Skill 自主使用 `search` 和 `section` 阅读相关招标原文，再提交“保留 / 建议增加 / 建议删除”。最多完成 {TECH_OUTLINE_HANDOFF_DECISION_UNITS} 个成功提交的决策单元；不足时做到 `decision-next complete=true` 为止。

本会话不得执行 appendix-next、review-complete、decisions、compose 或 finalize。到达本会话边界后立即停止，不要继续读后续章节；只返回一个简短 JSON：{{"workflowStage":"decision_checkpoint"}}。
""".strip()


def _build_outline_chapter_prompt(
    manifest_path: Path,
    chapter: dict[str, Any],
) -> str:
    return f"""
Use the {OUTLINE_SKILL_NAME} skill.

这是 S2 技术标目录的独立章节决策会话，只处理一级章：{chapter.get('number')} {chapter.get('title')}（chapter_id={chapter.get('chapter_id')}）。
manifest：{manifest_path}

准备产物已经由后端生成，不要执行 `s2outline prepare`。先用 `template-headings` 按 `next_cursor` 读完模板目录，再用 `headings` 按 `next_cursor` 读完整本招标目录，以便识别跨章等价项。
然后循环执行 `decision-next`。每个决策单元都按 Skill 自主使用 `search`、`section`、`read` 阅读相关招标原文，完成“保留 / 建议增加 / 建议删除”三类判断并执行 `decision-batch`。同一章包含多个决策单元时继续处理，直到 `decision-next complete=true`。
不要执行 `appendix-next`、`review-complete`、`decisions`、`compose` 或 `finalize`，也不要处理其他一级章。完成后只返回简短 JSON：{{"workflowStage":"chapter_complete"}}。
""".strip()


def _outline_chapter_base_urls() -> list[str]:
    configured = [
        item.strip().rstrip("/")
        for item in os.getenv("OPENCODE_CHAPTER_BASE_URLS", "").split(",")
        if item.strip()
    ]
    return configured or [str(settings.opencode_base_url).rstrip("/")]


def _prepare_outline_chapter_workspaces(
    manifest_path: Path,
    structure: dict[str, Any],
) -> tuple[list[dict[str, Any]], dict[str, Path], Path, dict[str, str]]:
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    work_dir = Path(str(manifest.get("workDir") or manifest_path.parent)).expanduser()
    runner = _load_technical_outline_runner()
    runner.write_template_structure(manifest, manifest_path)
    chapters = list(runner.decision_workflow.decision_chapters(structure)["chapters"])
    chapter_root = work_dir.parent / f".{work_dir.name}-chapter-decisions"
    if chapter_root.exists():
        shutil.rmtree(chapter_root)
    chapter_root.mkdir(parents=True)

    chapter_manifests: dict[str, Path] = {}
    baseline_files = [path for path in work_dir.iterdir() if path.is_file()]
    for index, chapter in enumerate(chapters, start=1):
        chapter_id = str(chapter["chapter_id"])
        chapter_dir = chapter_root / f"chapter-{index:02d}"
        chapter_dir.mkdir()
        for source in baseline_files:
            shutil.copy2(source, chapter_dir / source.name)
        chapter_manifest = copy.deepcopy(manifest)
        chapter_manifest["workDir"] = str(chapter_dir)
        chapter_manifest["outputFile"] = str(chapter_dir / "toc.json")
        chapter_manifest["_runtimeDecisionChapterId"] = chapter_id
        chapter_manifest_path = chapter_dir / "s2_input.json"
        chapter_manifest_path.write_text(
            json.dumps(chapter_manifest, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        chapter_manifests[chapter_id] = chapter_manifest_path

    cursor = 0
    while True:
        headings = runner.dispatch_command(
            "headings",
            manifest,
            manifest_path,
            ["--cursor", str(cursor), "--page-size", "200"],
        )
        if headings.get("requires_full_review"):
            shutil.rmtree(chapter_root, ignore_errors=True)
            raise _ChapterParallelUnsupported("招标文件缺少可分页目录结构")
        if headings.get("complete"):
            break
        cursor = int(headings["next_cursor"])
    workflow_binding = runner._strict_workflow_binding(manifest, work_dir) or {}
    return chapters, chapter_manifests, chapter_root, workflow_binding


def _run_parallel_outline_chapters(
    manifest_path: Path,
    structure: dict[str, Any],
    *,
    progress_callback: Callable[[str, dict[str, Any] | None], None] | None = None,
) -> list[str]:
    runner = _load_technical_outline_runner()
    chapters, chapter_manifests, chapter_root, workflow_binding = (
        _prepare_outline_chapter_workspaces(manifest_path, structure)
    )
    session_ids: dict[str, str] = {}
    chapter_base_urls = _outline_chapter_base_urls()
    model_config = system_settings_service.get_opencode_model_config_sync()
    chapter_indexes = {
        str(chapter["chapter_id"]): index for index, chapter in enumerate(chapters)
    }

    def run_chapter(chapter: dict[str, Any]) -> tuple[str, str]:
        chapter_id = str(chapter["chapter_id"])
        chapter_manifest_path = chapter_manifests[chapter_id]
        chapter_manifest = json.loads(chapter_manifest_path.read_text(encoding="utf-8"))
        chapter_work_dir = Path(str(chapter_manifest["workDir"]))

        def validate_complete() -> dict[str, Any]:
            binding = runner._strict_workflow_binding(chapter_manifest, chapter_work_dir) or {}
            return runner.decision_workflow.chapter_decision_progress(
                chapter_work_dir,
                structure,
                chapter_id,
                workflow_binding=binding,
            )

        def session_ready(details: dict[str, Any]) -> None:
            if progress_callback:
                progress_callback(
                    "outline_session_ready",
                    {**details, "chapterId": chapter_id, "chapterTitle": chapter.get("title")},
                )

        result = OpencodeClient(
            base_url=chapter_base_urls[chapter_indexes[chapter_id] % len(chapter_base_urls)],
            timeout_ms=int(settings.opencode_timeout_sec * 1000),
            model_config=model_config,
            request_slots=_TECH_OUTLINE_REQUEST_SLOTS,
        ).run_outline_decision_session(
            _build_outline_chapter_prompt(chapter_manifest_path, chapter),
            session_title=f"S2 目录决策·{chapter.get('number') or chapter_id}",
            completion_validator=validate_complete,
            session_ready_callback=session_ready,
            stream_callback=(
                (lambda details: progress_callback("outline_delta", details))
                if progress_callback
                else None
            ),
        )
        return chapter_id, str(result["sessionId"])

    try:
        with ThreadPoolExecutor(
            max_workers=min(
                TECH_OUTLINE_CHAPTER_WORKERS,
                len(chapter_base_urls),
                max(1, len(chapters)),
            )
        ) as executor:
            futures = {executor.submit(run_chapter, chapter): chapter for chapter in chapters}
            for future in as_completed(futures):
                chapter_id, session_id = future.result()
                session_ids[chapter_id] = session_id

        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        work_dir = Path(str(manifest.get("workDir") or manifest_path.parent)).expanduser()
        runner.decision_workflow.merge_chapter_decisions(
            work_dir,
            structure,
            {
                chapter_id: Path(
                    json.loads(path.read_text(encoding="utf-8"))["workDir"]
                )
                for chapter_id, path in chapter_manifests.items()
            },
            workflow_binding=workflow_binding,
        )
    except Exception:
        raise
    else:
        shutil.rmtree(chapter_root, ignore_errors=True)
    return [session_ids[str(chapter["chapter_id"])] for chapter in chapters]


def _technical_outline_handoff_state(
    manifest_path: Path,
    *,
    previous_decided_count: int,
) -> dict[str, Any]:
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    runner = _load_technical_outline_runner()
    try:
        progress = runner.dispatch_command("decision-next", manifest, manifest_path, [])
    except SystemExit as exc:
        raise RuntimeError(str(exc)) from exc
    decided_count = int(progress.get("decided_count") or 0)
    complete = bool(progress.get("complete"))
    if not complete and decided_count <= previous_decided_count:
        raise RuntimeError(
            "S2 目录接力会话没有提交新的目录判断，已停止以避免空转。"
        )
    return {
        "complete": complete,
        "decidedCount": decided_count,
        "remainingCount": int(progress.get("remaining_count") or 0),
    }


def _build_outline_finalize_prompt(manifest_path: Path) -> str:
    return f"""
Use the {OUTLINE_SKILL_NAME} skill.

这是 S2 技术标目录的最终收口会话。模板正文目录的三类判断已经由前序接力会话全部完成并持久化。

manifest：{manifest_path}

不要执行 `s2outline prepare`、`template-headings`、`decision-next`、`decision-batch`、`next-batch` 或 `review-batch`。不要直接读取或修改 `outline_decision_state.json`、`outline_authoring_decisions.json`、`toc.json` 或 Skill 脚本，只使用 `s2outline` 受控命令。

从 `s2outline appendix-next {manifest_path} --max-items 40` 开始，按 Skill 完成技术附表判断。随后只做一次全局复核：用 `s2outline headings --review` 按 `next_cursor` 从头分页读完整本招标目录，逐项查漏；疑似缺项必须用 `s2outline section` 详读原文，必要时用 `search` 定位后继续详读。发现遗漏或误判就用 `review-corrections` 写回并重新复核，不能只写在总结里。

确认无问题后依次执行 `s2outline review-complete`、`s2outline decisions`、`s2outline compose` 和 `{TECH_OUTLINE_FINALIZE_COMMAND} {manifest_path}`。最后原样返回 finalize 的严格 JSON，不加 Markdown 或解释。
""".strip()


def _is_business_bid(bid_type: Any) -> bool:
    return require_bid_type(
        bid_type,
        error_message="目录生成必须显式传入技术标或商务标。",
    ) == BUSINESS_BID_TYPE


def _outline_skill_name(bid_type: Any) -> str:
    return BUSINESS_OUTLINE_SKILL_NAME if _is_business_bid(bid_type) else OUTLINE_SKILL_NAME


def generate_outline_for_project(project_id: str, data: dict[str, Any] | None = None) -> dict[str, Any]:
    return generate_outline_for_project_with_progress(project_id, data)


def generate_outline_for_project_with_progress(
    project_id: str,
    data: dict[str, Any] | None = None,
    progress_callback: Callable[[str, dict[str, Any] | None], None] | None = None,
) -> dict[str, Any]:
    project = get_any_workspace_project_runtime_state(project_id, not_found_error=KeyError)
    parse_storage = copy.deepcopy(project.get("parse_storage") if isinstance(project.get("parse_storage"), dict) else {})
    tender_file_records, template_file_records = project_parse_input_records(project_id, project)
    combined_text_path = Path(str(parse_storage.get("combinedTextPath") or ""))
    if not combined_text_path.exists():
        raise ValueError("S1 解析结果不存在，请先完成解析。")

    combined_text = combined_text_path.read_text(encoding="utf-8").strip()
    if not combined_text:
        raise ValueError("S1 解析文本为空，暂时无法生成目录。")

    skill_workspace = _prepare_toc_skill_workspace(
        project_id=project_id,
        project=project,
        parse_storage=parse_storage,
        tender_file_records=tender_file_records,
        template_file_records=template_file_records,
    )
    if progress_callback:
        progress_callback(
            "inputs_ready",
            {
                "tenderFileCount": skill_workspace["tenderFileCount"],
                "templateFileCount": skill_workspace["templateFileCount"],
                "workDir": skill_workspace["workDir"],
                "bidType": skill_workspace["bidType"],
            },
        )

    toc_result = _run_outline_skill(
        Path(str(skill_workspace["canonicalManifestPath"])),
        bid_type=skill_workspace["bidType"],
        progress_callback=progress_callback,
    )
    publish_info = _publish_toc_skill_workspace(skill_workspace, toc_result)
    toc_result = publish_info["result"]
    nodes = _nodes_from_generation_result(
        toc_result,
        compact_technical=not _is_business_bid(skill_workspace["bidType"]),
    )
    summary = _summary_from_generation_result(toc_result)
    opencode_output = toc_result.get("opencodeOutput") if isinstance(toc_result.get("opencodeOutput"), dict) else {}
    if not opencode_output:
        opencode_output = build_directory_opencode_output(status="received")
    skill_name = _outline_skill_name(skill_workspace["bidType"])
    opencode_output.update(
        {
            "engine": skill_name,
            "skill": skill_name,
            "workDir": publish_info["workDir"],
            "manifestPath": publish_info["manifestPath"],
            "canonicalManifestPath": publish_info["canonicalManifestPath"],
            "stagingWorkDir": publish_info["stagingWorkDir"],
            "archiveRoot": publish_info["archiveRoot"],
            "tocJsonPath": str(toc_result.get("outputFile") or publish_info["outputFile"]),
        }
    )
    evidence_path = str(toc_result.get("evidenceFile") or publish_info.get("evidenceFile") or "").strip()
    if evidence_path:
        opencode_output["evidencePath"] = evidence_path
    if progress_callback:
        progress_callback(
            "normalizing_result",
            {
                "chapterCount": len(nodes),
            },
        )

    generated_at = now_iso()
    project_for_update = require_any_workspace_project_for_update(project_id, not_found_error=KeyError)
    payload = save_generated_outline_state(
        project_for_update,
        nodes=nodes,
        generated_at=generated_at,
        summary=summary,
        opencode_output=opencode_output,
        rule_evidence=toc_result.get("ruleEvidence") if isinstance(toc_result.get("ruleEvidence"), dict) else {},
    )
    persist_workspace_project_state(project_for_update)
    return payload


def _run_outline_skill(
    manifest_path: Path,
    *,
    bid_type: Any,
    progress_callback: Callable[[str, dict[str, Any] | None], None] | None = None,
) -> dict[str, Any]:
    if _is_business_bid(bid_type):
        return _run_business_outline_skill(manifest_path, progress_callback=progress_callback)

    prompt = _build_outline_finalize_prompt(manifest_path)
    try:
        trusted_input = _capture_trusted_technical_outline_input(manifest_path)
        chapter_session_ids: list[str] = []
        handoff_kwargs: dict[str, Any] = {}
        try:
            chapter_session_ids = _run_parallel_outline_chapters(
                manifest_path,
                trusted_input["templateStructure"],
                progress_callback=progress_callback,
            )
        except _ChapterParallelUnsupported:
            previous_decided_count = [-1]

            def handoff_state(handoff_index: int) -> dict[str, Any]:
                del handoff_index
                state = _technical_outline_handoff_state(
                    manifest_path,
                    previous_decided_count=previous_decided_count[0],
                )
                previous_decided_count[0] = int(state["decidedCount"])
                return state

            handoff_kwargs = {
                "handoff_prompt_factory": lambda index: _build_outline_handoff_prompt(
                    manifest_path,
                    index,
                ),
                "handoff_state_callback": handoff_state,
            }

        generated = OpencodeClient(
            timeout_ms=int(settings.opencode_timeout_sec * 1000)
        ).generate_outline_with_trace(
                prompt,
                session_ready_callback=(
                    (lambda details: progress_callback("outline_session_ready", details))
                    if progress_callback
                    else None
                ),
                stream_callback=(
                    (lambda details: progress_callback("outline_delta", details))
                    if progress_callback
                    else None
                ),
                early_tool_command=TECH_OUTLINE_FINALIZE_EARLY_COMMAND,
                terminal_validator=lambda: _finalize_current_technical_outline(manifest_path),
                **handoff_kwargs,
            )
        loaded = _load_outline_result(
            generated,
            manifest_path,
            expected_bid_type=bid_type,
            trusted_technical_input=trusted_input,
        )
        if chapter_session_ids:
            output_trace = loaded.setdefault("opencodeOutput", {})
            final_session_id = str(output_trace.get("sessionId") or "")
            output_trace["sessionIds"] = [
                *chapter_session_ids,
                *([final_session_id] if final_session_id else []),
            ]
            output_trace["chapterSessionCount"] = len(chapter_session_ids)
            output_trace["parallelChapterWorkers"] = min(
                TECH_OUTLINE_CHAPTER_WORKERS,
                len(chapter_session_ids),
            )
        return loaded
    except Exception as exc:
        if progress_callback:
            progress_callback(
                "outline_failed",
                {"error": str(exc), "manifestPath": str(manifest_path)},
            )
        raise RuntimeError(
            "技术标目录生成失败：目录生成需要 opencode 自主决策，"
            f"futurecode/opencode 执行失败：{exc}。"
        ) from exc


def _run_business_outline_skill(
    manifest_path: Path,
    *,
    progress_callback: Callable[[str, dict[str, Any] | None], None] | None = None,
) -> dict[str, Any]:
    prompt = _build_business_outline_prompt(manifest_path)
    try:
        result = OpencodeClient(timeout_ms=int(settings.opencode_timeout_sec * 1000)).generate_outline_with_trace(
            prompt,
            session_ready_callback=(
                (lambda details: progress_callback("outline_session_ready", details))
                if progress_callback
                else None
            ),
            stream_callback=(
                (lambda details: progress_callback("outline_delta", details))
                if progress_callback
                else None
            ),
            early_tool_command="",
        )
    except Exception as exc:
        if progress_callback:
            progress_callback(
                "outline_fallback",
                {"error": str(exc), "manifestPath": str(manifest_path)},
            )
        raise RuntimeError(
            "商务标目录生成失败："
            f"futurecode 执行失败：{exc}。"
            "本地 bid-business-outline-generator 只负责准备候选材料，不能兜底生成最终 outline.json。"
        ) from exc
    return _load_outline_result(result, manifest_path, expected_bid_type=BUSINESS_BID_TYPE)


def _run_local_outline_skill(manifest_path: Path) -> dict[str, Any]:
    raise RuntimeError(
        "目录生成需要 opencode 自主决策，本地 s2toc 兼容命令不再生成最终目录。"
    )


def _build_outline_prompt(manifest_path: Path, bid_type: Any) -> str:
    bid_type_text = require_bid_type(
        bid_type,
        error_message="目录生成必须显式传入技术标或商务标。",
    )
    if _is_business_bid(bid_type_text):
        return _build_business_outline_prompt(manifest_path)
    skill_name = _outline_skill_name(bid_type_text)
    return f"""
Use the {skill_name} skill.

生成 S2 {bid_type_text}目录。目录学习、招标新增项和适用性建议由 Opencode 按 Skill 自主判断，不得把未判断节点自动当成必要。

manifest：{manifest_path}

历史投标模板提供目录经验，当前招标文件提供本项目要求。完整学习模板一至三级目录，模板已有第三级目录统一进入结果供用户确认，但不预设任何模板节点必须保留。每个模板节点由 Opencode 自主选择保留或建议删除，并自主判断建议增加项；建议删除的节点仍保留供用户确认。最终目录最多三级，第四级及更深层级只作为对应第三级节点的内容参考，不把参数、条款或表格字段机械扩成目录，再结合招标文件逐项判断。

严格按 Skill 执行受控流程。`prepare` 只执行一次；完成后不要再执行同功能的 `s2outline template`，不要直接读取 `template_structure.json`，模板节点只通过 `decision-next` 按章获取。招标目录必须按 `next_cursor` 分页读到 `complete=true`；每章自主使用 `s2outline section` 阅读相关章节或小节，使用 `s2outline search` 跨章节查漏并继续详读原文。每个决策单元一次提交保留、建议增加、建议删除，不做章节复核。完成全部章节后用 `s2outline appendix-next {manifest_path} --max-items 40` 判断附表，再只做一次全局复核；发现遗漏或误判必须用 `review-corrections` 写入目录决策，不能只记在总结或留给后续阶段。确认无问题后执行 `review-complete`、`decisions` 和 `compose`。不得编写临时脚本批量拼装判断，不得自行写入 manifest.outputFile 或决策状态文件，也不得读取决策状态文件。最后执行：

首次执行 `s2outline prepare {manifest_path}` 时，Bash 必须显式设置 `timeout=300000`。若仍超时，只增大 timeout 后重试同一命令；不要检查脚本或包装器，不要绕过 `s2outline`。

{TECH_OUTLINE_FINALIZE_COMMAND} {manifest_path}

finalize 只校验结果，也是后端的完成信号。最后原样返回 finalize 的严格 JSON，不要 Markdown 或解释文字。
""".strip()


def _build_business_outline_prompt(manifest_path: Path) -> str:
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    work_dir = Path(str(manifest.get("workDir") or manifest_path.parent)).expanduser()
    business_outline_file = str(work_dir / "outline.json")
    decisions_file = str(work_dir / "outline_authoring_decisions.json")
    history_file = str(work_dir / "history_bid_outline_inputs.json")
    tender_file = str(work_dir / "tender_map_inputs.json")
    document_structure_index_file = str(work_dir / "document_structure_index.json")
    source_text_candidates_file = str(work_dir / "source_text_candidates.json")
    return f"""
Use the {BUSINESS_OUTLINE_SKILL_NAME} skill.

你现在在做 S2 商务标目录生成。必须完整执行 bid-business-outline-generator Skill，并严格遵循原有 Skill 的产物边界：准备脚本只生成输入材料，opencode 只输出语义选择、状态判断和保留/延后理由，最终 outline.json 必须通过固定的 outline_authoring_helper.py 机械写回。

manifest：{manifest_path}

集成准备动作：直接调用一次 Bash 工具执行下面命令，Bash 工具 timeout 必须设置为 1800000 毫秒或更高。不要先检查工作目录，不要先执行 pwd/ls/cat/read/glob，不要拆成多条命令，不要改写命令或路径：

{BUSINESS_OUTLINE_SKILL_COMMAND} {manifest_path}

强制工具顺序：加载 skill 之后，第一条非 skill 工具调用必须是 Bash，且 Bash command 必须完全等于上面这一行。禁止在这条 Bash 命令完成前调用 read、glob、list、ls、cat、head、tail、grep 或任何读取 manifest/JSON 的工具。不要读取 manifest 内容来“理解输入”；准备命令会读取 manifest 并产出后续判断所需材料。

Mandatory tool order: after the skill tool is loaded, your first non-skill tool call MUST be the Bash tool with exactly the `business-outline {manifest_path}` command above. Do not call read, glob, list, ls, cat, head, tail, grep, or any manifest/JSON inspection tool before that Bash command completes. Do not inspect the manifest first.

该命令只负责根据 manifest.templateFile 和 manifest.tenderFiles 生成：
- {history_file}
- {tender_file}
- {document_structure_index_file}
- {source_text_candidates_file}

该命令不得被视为最终目录生成；不得把它的 stdout、summary 或任何候选信息当作最终 outline.json。

AI 判断动作：命令完成后，继续按 bid-business-outline-generator Skill 的步骤 2-6 做 AI 判断：学习历史商务标目录结构，分析当前招标文件，读取并消费 document_structure_index.json 与 source_text_candidates.json，匹配每个目录项的 source_text，判断 required_status，补强必须提交材料，并把每个保留目录项的语义决策写入 outline_authoring_decisions.json。不得现场编写临时 Python 写回脚本。

必须使用后端 manifest.templateFile 作为历史商务标/商务模板来源，不扫描当前工作目录，不使用 user_confirmed_inputs.json。

后续判断只基于原始 Skill 输入产物：
- 历史商务标输入：{history_file}
- 招标文件输入：{tender_file}
- 文档结构索引：{document_structure_index_file}
- source_text 候选：{source_text_candidates_file}

source_text 选择必须先消费 source_text_candidates.json 的首选候选：若某目录项已有 candidates[0]，且候选不是目录页/目次页，也不是合计、总计、小计等汇总行，最终 section.source_text 应优先逐字采用该候选，并把候选的 scope、evidence_strength、evidence_category、match_reason 写入 section.evidence_scope、section.evidence_strength、section.evidence_category、section.reason。不要用同一章节内的表格汇总行替换强标题候选或强段落候选；若首选候选是目录项标题本身或明确提交材料名称，最终 source_text 必须保留该候选，不得改用“合计 | 100”这类汇总行。

历史继承策略：章节级、材料级目录应保留；具体项目业绩清单、具体证书扫描件、协议明细、过程材料明细、逐页附件、图片说明、合同逐项列表等细碎内容，应由 opencode 判断为“素材库组装项/正文素材”，在 outline_authoring_decisions.json 中显式写 action: "defer" 并说明理由，不能因为只有历史原文就默认以 history_fallback 全部保留进目录。

禁止调用 read 工具；不要使用 cat/head/tail/grep 直接打印 JSON 大文件。需要访问文件内容或写回结果时，只能调用 Bash 工具执行 python3 脚本读取上述原始产物、按 bid-business-outline-generator Skill 逻辑分析和写回。Python 脚本可以完整读取 JSON 文件到内存，但每次 stdout 只输出当前判断所需的简短检查结果，避免刷屏或截断。

必须先把 opencode 的语义判断写入固定决策文件：
{decisions_file}

outline_authoring_decisions.json 只表达 opencode 的判断，不负责机械拼装。至少包含：
{{
  "document_name": "商务标目录",
  "sections": [
    {{
      "id": "BIZ-FALLBACK-0001",
      "candidate_source_id": "hist-cand-001",
      "selected_candidate_id": "cand-001",
      "required_status": "必要",
      "reason": "结合当前招标文件证据与历史目录语义保留。"
    }}
  ],
  "review_items": []
}}

写好决策文件后，必须调用固定 helper 机械生成最终 outline.json，不得自己现场编写 Python 写回逻辑：
python scripts/outline_authoring_helper.py --history "{history_file}" --source-candidates "{source_text_candidates_file}" --decisions "{decisions_file}" --output "{business_outline_file}"

helper 只负责读取候选、保持 ID、组装/写回 outline.json、运行基础校验；它不判断章节是否必要，不写死商务标题。

最终原生产物必须写入：
{business_outline_file}

outline.json 必须满足：
{{
  "schema_version": "business_bid_outline.v1",
  "sections": [
    {{
      "id": "sec-001",
      "title": "目录标题",
      "number": null,
      "level": 1,
      "required_status": "待确认",
      "source_text": "逐字证据",
      "evidence_scope": "parent_context",
      "evidence_strength": "strong",
      "children": []
    }}
  ]
}}

每一个 sections[*] 以及所有子级 section 都必须显式包含 number 字段。有历史编号时保留字符串编号；历史无编号、空编号或无法可靠推断时写为 null 或空字符串，禁止由层级顺序强行生成 1、1.1、1.2 等编号。

不要自行生成或修改前端兼容 toc.json；后端会根据最终 outline.json 自动转换。

最后只返回严格 JSON，不要 Markdown，不要解释文字：
{{
  "schema_version": "business_bid_outline.v1",
  "businessOutlineFile": "{business_outline_file}",
  "historyBidOutlineInputsFile": "{history_file}",
  "tenderMapInputsFile": "{tender_file}",
  "sourceTextCandidatesFile": "{source_text_candidates_file}",
  "outlineAuthoringDecisionsFile": "{decisions_file}",
  "summary": {{"total_sections": 0}}
}}
""".strip()

def _load_outline_result(
    result: dict[str, Any],
    manifest_path: Path,
    *,
    expected_bid_type: Any | None = None,
    trusted_technical_input: dict[str, Any] | None = None,
) -> dict[str, Any]:
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    resolved_bid_type = expected_bid_type if expected_bid_type is not None else manifest.get("bidType")
    is_business_bid = _is_business_bid(resolved_bid_type)
    trusted_output = (
        trusted_technical_input.get("outputFile")
        if isinstance(trusted_technical_input, dict) and not is_business_bid
        else None
    )
    output_file = Path(
        str(trusted_output or result.get("outputFile") or manifest.get("outputFile") or "")
    ).expanduser()
    if is_business_bid:
        evidence_file = Path(str(result.get("evidenceFile") or manifest.get("evidenceFile") or "")).expanduser()
        return _load_business_outline_result(result, manifest, output_file, evidence_file)

    if not output_file.exists():
        raise RuntimeError(f"S2 目录 Skill 未生成 outputFile：{output_file}")
    validated_outline: dict[str, Any] | None = None
    if expected_bid_type is not None:
        validated_outline = _validate_technical_compose_report(
            manifest_path.parent,
            output_file,
            trusted_input=trusted_technical_input,
        )

    outline = validated_outline or json.loads(output_file.read_text(encoding="utf-8"))
    if (
        not isinstance(outline, dict)
        or outline.get("schema_version") != "technical-outline.v1"
        or not isinstance(outline.get("nodes"), list)
        or not outline["nodes"]
    ):
        raise RuntimeError("S2 目录 Skill 输出不是有效 technical-outline.v1。")

    outline["outputFile"] = str(output_file)
    if isinstance(result.get("opencodeOutput"), dict):
        outline["opencodeOutput"] = result["opencodeOutput"]
    outline["ruleEvidence"] = _technical_rule_evidence(outline["nodes"])
    summary = result.get("summary") if isinstance(result.get("summary"), dict) else {}
    outline["summary"] = summary or {
        "total_nodes": outline["ruleEvidence"]["nodeCount"],
        "action_counts": outline["ruleEvidence"]["actionCounts"],
    }
    return outline


def _validate_technical_compose_report(
    work_dir: Path,
    output_file: Path,
    *,
    trusted_input: dict[str, Any] | None,
) -> dict[str, Any]:
    report_path = work_dir / "outline_compose_report.json"
    if not report_path.exists():
        raise RuntimeError("技术标目录缺少 compose report，拒绝接收直接写入的 outputFile。")
    try:
        report = json.loads(report_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise RuntimeError("技术标目录 compose report 无法读取。") from exc
    if not isinstance(report, dict) or report.get("schema_version") != "technical-outline-compose-report.v1":
        raise RuntimeError("技术标目录 compose report Schema 无效。")
    required_fields = {
        "inputFingerprint",
        "decisionsDigest",
        "outputSha256",
        "outputFile",
    }
    if str((trusted_input or {}).get("tenderInputsDigest") or "").strip():
        required_fields.update(
            {"tenderInputsDigest", "headingsStateDigest", "decisionStateDigest"}
        )
    if any(not str(report.get(field) or "").strip() for field in required_fields):
        raise RuntimeError("技术标目录 compose report 字段不完整。")
    reported_output = Path(str(report.get("outputFile") or "")).expanduser()
    if reported_output.resolve() != output_file.resolve():
        raise RuntimeError("技术标目录 compose report 指向了不同的 outputFile。")
    try:
        output_bytes = output_file.read_bytes()
        actual_outline = json.loads(output_bytes.decode("utf-8"))
    except (OSError, UnicodeDecodeError, json.JSONDecodeError) as exc:
        raise RuntimeError("技术标目录 outputFile 无法读取。") from exc
    output_sha256 = hashlib.sha256(output_bytes).hexdigest()
    if report.get("outputSha256") != output_sha256:
        raise RuntimeError("技术标目录 outputFile 在 compose 后被修改。")
    if not isinstance(trusted_input, dict):
        raise RuntimeError("技术标目录缺少后端可信模板快照。")

    template_file = Path(str(trusted_input.get("templateFile") or "")).expanduser()
    expected_template_sha256 = str(trusted_input.get("templateFileSha256") or "")
    trusted_structure = trusted_input.get("templateStructure")
    if not template_file.is_file() or not expected_template_sha256 or not isinstance(trusted_structure, dict):
        raise RuntimeError("技术标目录后端可信模板快照无效。")
    current_template_sha256 = hashlib.sha256(template_file.read_bytes()).hexdigest()
    if current_template_sha256 != expected_template_sha256:
        raise RuntimeError("技术标历史模板在 Opencode 执行期间被修改。")

    runner = _load_technical_outline_runner()
    composer = runner.outline_composer
    try:
        trusted_fingerprint = composer.template_fingerprint(
            composer.annotate_template_structure(trusted_structure)
        )
        structure_path = work_dir / "template_structure.json"
        current_structure = json.loads(structure_path.read_text(encoding="utf-8"))
        current_fingerprint = composer.template_fingerprint(
            composer.annotate_template_structure(current_structure)
        )
        decisions = composer.load_decisions(work_dir, trusted_structure, required=True)
        workflow_proof: dict[str, str] = {}
        trusted_tender_files = trusted_input.get("tenderFiles")
        trusted_tender_digest = str(trusted_input.get("tenderInputsDigest") or "")
        if trusted_tender_digest:
            if not isinstance(trusted_tender_files, list) or not trusted_tender_files:
                raise RuntimeError("技术标目录后端可信招标文件快照无效。")
            current_tender_digest = runner.review_workflow.tender_input_fingerprint(
                trusted_tender_files
            )
            if current_tender_digest != trusted_tender_digest:
                raise RuntimeError("技术标招标文件在 Opencode 执行期间被修改。")
            trusted_appendix_items = trusted_input.get("appendixItems")
            if not isinstance(trusted_appendix_items, list) or any(
                not isinstance(item, dict) for item in trusted_appendix_items
            ):
                raise RuntimeError("技术标目录后端可信附表清单快照无效。")
            workspace_appendix_items = (
                runner.review_workflow.decision_appendix_items(work_dir)
            )
            _, trusted_appendix_digest = (
                runner.decision_workflow._normalized_appendix_inventory(
                    trusted_appendix_items
                )
            )
            _, workspace_appendix_digest = (
                runner.decision_workflow._normalized_appendix_inventory(
                    workspace_appendix_items
                )
            )
            if workspace_appendix_digest != trusted_appendix_digest:
                raise RuntimeError("技术标目录工作区附表清单与后端可信快照不一致。")
            workflow_proof = runner.review_workflow.require_headings_complete(
                work_dir,
                trusted_tender_files,
            )
            workflow_proof.update(
                runner.decision_workflow.validate_finalized_decisions(
                    work_dir,
                    trusted_structure,
                    decisions,
                    workflow_binding=workflow_proof,
                    appendix_items=trusted_appendix_items,
                )
            )
            composer.validate_compose_report(
                work_dir=work_dir,
                output_file=output_file,
                structure=trusted_structure,
                decisions=decisions,
                workflow_proof=workflow_proof,
            )
        expected_outline, context = composer.build_composition(trusted_structure, decisions)
    except RuntimeError:
        raise
    except (OSError, json.JSONDecodeError, ValueError, SystemExit) as exc:
        raise RuntimeError(f"技术标目录 compose 可信校验失败：{exc}") from exc

    if current_fingerprint != trusted_fingerprint:
        raise RuntimeError("技术标模板结构与后端可信快照不一致。")
    if report.get("inputFingerprint") != trusted_fingerprint:
        raise RuntimeError("技术标目录 compose report 的模板指纹不可信。")
    if report.get("decisionsDigest") != context["decisionsDigest"]:
        raise RuntimeError("技术标目录 compose report 的 decisions 摘要不一致。")
    if actual_outline != expected_outline:
        raise RuntimeError("技术标目录输出无法由可信模板与 decisions 确定性重组。")
    return copy.deepcopy(expected_outline)


def _load_business_outline_result(
    result: dict[str, Any],
    manifest: dict[str, Any],
    output_file: Path,
    evidence_file: Path,
) -> dict[str, Any]:
    work_dir = Path(str(manifest.get("workDir") or output_file.parent)).expanduser()
    business_outline_file = Path(str(result.get("businessOutlineFile") or work_dir / "outline.json")).expanduser()
    business_outline = _load_business_outline_json(business_outline_file)
    _write_business_toc_from_outline_payload(manifest, business_outline, business_outline_file, output_file, evidence_file)
    toc = json.loads(output_file.read_text(encoding="utf-8"))
    if not isinstance(toc, dict) or not isinstance(toc.get("items"), list):
        raise RuntimeError("商务标 outline.json 未能转换为有效 bid-toc-json-v1。")
    toc["items"] = _clean_toc_items(toc["items"])
    _rewrite_toc_file(output_file, toc, evidence_file)
    toc["outputFile"] = str(output_file)
    toc["evidenceFile"] = str(evidence_file)
    toc["businessOutlineFile"] = str(business_outline_file)
    if isinstance(result.get("opencodeOutput"), dict):
        toc["opencodeOutput"] = result["opencodeOutput"]
    toc["ruleEvidence"] = _public_rule_evidence_from_file(evidence_file)
    return toc


def _load_business_outline_json(path: Path) -> dict[str, Any]:
    if not path.exists():
        raise RuntimeError(f"商务标目录 Skill 未生成最终 outline.json：{path}")
    business_outline = _load_json_dict(path)
    if business_outline.get("schema_version") != "business_bid_outline.v1":
        raise RuntimeError("商务标 outline.json schema_version 必须是 business_bid_outline.v1。")
    sections = business_outline.get("sections")
    if not isinstance(sections, list) or not sections:
        raise RuntimeError("商务标 outline.json 必须包含非空 sections[]。")
    _validate_business_outline_section_numbers(sections)
    return business_outline


def _validate_business_outline_section_numbers(sections: list[Any], path: str = "sections") -> None:
    for index, section in enumerate(sections):
        section_path = f"{path}[{index}]"
        if not isinstance(section, dict):
            continue
        if "number" not in section:
            raise RuntimeError(f"商务标 outline.json {section_path}.number 缺失。")
        _business_section_number(section.get("number"))
        children = section.get("children")
        if isinstance(children, list):
            _validate_business_outline_section_numbers(children, f"{section_path}.children")


def _write_business_toc_from_outline(
    manifest: dict[str, Any],
    result: dict[str, Any],
    output_file: Path,
    evidence_file: Path,
) -> None:
    work_dir = Path(str(manifest.get("workDir") or output_file.parent)).expanduser()
    business_outline_file = Path(str(result.get("businessOutlineFile") or work_dir / "outline.json")).expanduser()
    business_outline = _load_business_outline_json(business_outline_file)
    _write_business_toc_from_outline_payload(manifest, business_outline, business_outline_file, output_file, evidence_file)


def _write_business_toc_from_outline_payload(
    manifest: dict[str, Any],
    business_outline: dict[str, Any],
    business_outline_file: Path,
    output_file: Path,
    evidence_file: Path,
) -> None:
    work_dir = Path(str(manifest.get("workDir") or output_file.parent)).expanduser()
    sections = business_outline.get("sections") if isinstance(business_outline.get("sections"), list) else []
    if not sections:
        raise RuntimeError("商务标 outline.json 必须包含非空 sections[]。")

    items = _clean_toc_items(_business_toc_items_from_sections(sections))
    counts = Counter(str(item.get("annotation") or "") for item in items)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    evidence_file.parent.mkdir(parents=True, exist_ok=True)
    source_files = {
        "tender": manifest.get("tenderFiles") if isinstance(manifest.get("tenderFiles"), list) else [],
        "template": str(manifest.get("templateFile") or ""),
        "output": str(output_file),
        "evidence": str(evidence_file),
        "businessOutline": str(business_outline_file),
        "tenderMapInputs": str(work_dir / "tender_map_inputs.json"),
        "historyBidOutlineInputs": str(work_dir / "history_bid_outline_inputs.json"),
    }
    toc = {
        "schema_version": "bid-toc-json-v1",
        "document_title": str(business_outline.get("document_name") or "商务标目录"),
        "project": {
            "projectId": str(manifest.get("projectId") or ""),
            "projectCode": str(manifest.get("projectCode") or ""),
            "projectName": str(manifest.get("projectName") or ""),
            "bidType": require_bid_type(
                manifest.get("bidType"),
                error_message="商务标目录生成必须显式传入商务标。",
            ),
        },
        "source_files": source_files,
        "summary": {
            "total_items": len(items),
            "annotation_counts": dict(counts),
        },
        "items": items,
        "outputFile": str(output_file),
        "evidenceFile": str(evidence_file),
        "businessOutlineFile": str(business_outline_file),
    }
    evidence = {
        "schema_version": "bid-toc-evidence-v1",
        "engine": "bid-business-outline-generator",
        "inputs": source_files,
        "businessOutlineFile": str(business_outline_file),
        "decisions": business_outline.get("review_items") if isinstance(business_outline.get("review_items"), list) else [],
    }
    output_file.write_text(json.dumps(toc, ensure_ascii=False, indent=2), encoding="utf-8")
    if not evidence_file.exists():
        evidence_file.write_text(json.dumps(evidence, ensure_ascii=False, indent=2), encoding="utf-8")


def _business_toc_items_from_sections(sections: list[Any]) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []

    def append_section(section: dict[str, Any], fallback_level: int) -> None:
        order = len(items) + 1
        required_status = _normalize_required_status(section.get("required_status") or section.get("requiredStatus"))
        source_text = str(section.get("source_text") or section.get("sourceText") or "").strip()
        source_refs = _business_source_refs_from_section(section, source_text)
        number = _business_section_number(section.get("number"))
        items.append(
            {
                "itemId": f"TOC-{order:04d}",
                "order": order,
                "number": number,
                "title": str(section.get("title") or section.get("name") or f"商务标目录项{order}").strip(),
                "level": _coerce_toc_level(section.get("level") or fallback_level),
                "annotation": _business_annotation_from_required_status(required_status),
                "required_status": required_status,
                "requiredStatus": required_status,
                "source_text": source_text,
                "sourceText": source_text,
                "source": "business_outline",
                "reason": str(section.get("reason") or "商务标目录项来自 futurecode 生成的 outline.json。").strip(),
                "source_refs": source_refs,
                "material_refs": [],
            }
        )
        children = section.get("children") if isinstance(section.get("children"), list) else []
        for child in children:
            if isinstance(child, dict):
                append_section(child, fallback_level + 1)

    for section in sections:
        if isinstance(section, dict):
            append_section(section, 1)
    return items


def _business_section_number(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    raise RuntimeError("商务标 outline.json sections[].number 必须是字符串或 null。")


def _normalize_required_status(value: Any) -> str:
    text = str(value or "").strip()
    if text in {"必要", "可选", "待确认"}:
        return text
    if text in {"必选", "必须", "应提交", "须提交", "保留"}:
        return "必要"
    if text in {"选填", "按需", "如适用", "适用时提交"}:
        return "可选"
    return text or "待确认"


def _business_annotation_from_required_status(required_status: str) -> str:
    if required_status == "待确认":
        return "待确认"
    if required_status == "可选":
        return "可选"
    return "保留"


def _business_source_refs_from_section(section: dict[str, Any], source_text: str) -> list[dict[str, Any]]:
    raw_refs = section.get("source_refs") if isinstance(section.get("source_refs"), list) else []
    if not raw_refs and isinstance(section.get("sourceRefs"), list):
        raw_refs = section["sourceRefs"]
    if not raw_refs and isinstance(section.get("source_ref"), dict):
        raw_refs = [section["source_ref"]]
    refs = [_clean_source_ref(ref) for ref in raw_refs if isinstance(ref, dict)]
    if refs:
        return refs
    if not source_text:
        return []
    return [
        _clean_source_ref(
            {
                "type": "tender",
                "role": "basis",
                "kind": "business_outline_section",
                "sectionId": str(section.get("id") or ""),
                "title": str(section.get("title") or ""),
                "raw_text": source_text,
                "rawText": source_text,
                "basisText": source_text,
                "searchText": source_text,
                "reason": "商务标 outline.json 目录项依据",
            }
        )
    ]



def _apply_agent_decisions(
    toc: dict[str, Any],
    agent_decisions: list[Any],
    evidence_file: Path,
) -> dict[str, Any]:
    evidence = _load_json_dict(evidence_file)
    candidates = {
        str(item.get("id") or item.get("candidateId") or ""): item
        for item in evidence.get("tenderCandidates", [])
        if isinstance(item, dict)
    }
    items = _clean_toc_items(toc.get("items") if isinstance(toc.get("items"), list) else [])
    item_index = _item_lookup(items)
    for raw_decision in agent_decisions:
        if not isinstance(raw_decision, dict):
            continue
        decision = str(raw_decision.get("decision") or raw_decision.get("action") or "").strip()
        candidate_id = str(raw_decision.get("candidateId") or raw_decision.get("id") or "")
        candidate = candidates.get(candidate_id)
        if candidate is None:
            continue
        if decision in {"attach_evidence", "attach", "covered"}:
            target = _find_decision_target(raw_decision, item_index, items)
            if target is None:
                continue
            target.setdefault("source_refs", []).append(_source_ref_from_agent_candidate(candidate, raw_decision))
        elif decision in {"append_item", "add", "add_appendix"}:
            items.append(_toc_item_from_agent_candidate(len(items) + 1, candidate, raw_decision))
        elif decision in {"exclude", "candidate", "ignore"}:
            continue
    toc["items"] = _clean_toc_items(items)
    return toc


def _rewrite_toc_file(
    output_file: Path,
    toc: dict[str, Any],
    evidence_file: Path,
    *,
    agent_decisions: list[Any] | None = None,
) -> None:
    raw_items = toc.get("items") if isinstance(toc.get("items"), list) else []
    items = _clean_toc_items(raw_items)
    counts = Counter(str(item.get("annotation") or "") for item in items)
    toc["items"] = items
    summary = toc.get("summary") if isinstance(toc.get("summary"), dict) else {}
    summary.update(
        {
            "total_items": len(items),
            "annotation_counts": dict(counts),
        }
    )
    toc["summary"] = summary
    toc["outputFile"] = str(output_file)
    toc["evidenceFile"] = str(evidence_file)
    output_file.write_text(json.dumps(toc, ensure_ascii=False, indent=2), encoding="utf-8")
    if evidence_file.exists():
        evidence = _load_json_dict(evidence_file)
        if agent_decisions is not None:
            evidence["agentDecisions"] = [item for item in agent_decisions if isinstance(item, dict)]
        evidence_file.write_text(json.dumps(evidence, ensure_ascii=False, indent=2), encoding="utf-8")

def _clean_toc_items(items: list[Any]) -> list[dict[str, Any]]:
    cleaned: list[dict[str, Any]] = []
    for index, item in enumerate(items, start=1):
        if not isinstance(item, dict):
            continue
        source_refs = item.get("source_refs")
        if not isinstance(source_refs, list):
            source_refs = item.get("sourceRefs") if isinstance(item.get("sourceRefs"), list) else []
        annotation = str(item.get("annotation") or "").strip() or "保留"
        required_status = str(item.get("required_status") or item.get("requiredStatus") or "").strip()
        if not required_status:
            required_status = _required_status_from_annotation(annotation)
        source_text = _toc_item_source_text(item)
        cleaned.append(
            {
                "itemId": str(item.get("itemId") or f"TOC-{index:04d}"),
                "order": index,
                "number": str(item.get("number") or "").strip(),
                "title": str(item.get("title") or item.get("name") or f"未命名章节{index}").strip(),
                "level": _coerce_toc_level(item.get("level")),
                "annotation": annotation,
                "required_status": required_status,
                "requiredStatus": required_status,
                "source_text": source_text,
                "sourceText": source_text,
                "source": str(item.get("source") or "").strip() or "template",
                "reason": str(item.get("reason") or "").strip(),
                "source_refs": [_clean_source_ref(ref) for ref in source_refs if isinstance(ref, dict)],
                "material_refs": [],
            }
        )
    return cleaned


def _required_status_from_annotation(annotation: str) -> str:
    text = str(annotation or "").strip()
    if text in {"待确认", "可选"}:
        return text
    if text in {"保留", "必要"}:
        return "必要"
    return text


def _clean_source_ref(ref: dict[str, Any]) -> dict[str, Any]:
    cleaned = dict(ref)
    if "raw_text" in cleaned and "rawText" not in cleaned:
        cleaned["rawText"] = cleaned.get("raw_text")
    if "rawText" in cleaned and "raw_text" not in cleaned:
        cleaned["raw_text"] = cleaned.get("rawText")
    basis_text = str(cleaned.get("basisText") or cleaned.get("rawText") or cleaned.get("raw_text") or "")
    if basis_text and not cleaned.get("basisText"):
        cleaned["basisText"] = basis_text
    if basis_text and not cleaned.get("searchText"):
        cleaned["searchText"] = basis_text
    return cleaned


def _item_lookup(items: list[dict[str, Any]]) -> dict[str, dict[str, Any]]:
    result: dict[str, dict[str, Any]] = {}
    for item in items:
        for key in (
            str(item.get("itemId") or ""),
            str(item.get("title") or ""),
            _title_key(str(item.get("title") or "")),
        ):
            if key:
                result[key] = item
    return result


def _find_decision_target(
    decision: dict[str, Any],
    item_index: dict[str, dict[str, Any]],
    items: list[dict[str, Any]],
) -> dict[str, Any] | None:
    for key in (
        str(decision.get("targetItemId") or ""),
        str(decision.get("targetTitle") or ""),
        _title_key(str(decision.get("targetTitle") or "")),
    ):
        if key and key in item_index:
            return item_index[key]
    return items[-1] if items else None


def _source_ref_from_agent_candidate(candidate: dict[str, Any], decision: dict[str, Any]) -> dict[str, Any]:
    basis_text = str(candidate.get("searchText") or candidate.get("basisText") or candidate.get("rawText") or "")
    return {
        "type": "tender",
        "role": "basis",
        "kind": "codex_semantic",
        "candidateKind": str(candidate.get("kind") or ""),
        "candidateId": str(candidate.get("id") or decision.get("candidateId") or ""),
        "relation": str(decision.get("relation") or "semantic_match"),
        "confidence": _coerce_confidence(decision.get("confidence")),
        "reason": str(decision.get("reason") or ""),
        "fileId": str(candidate.get("fileId") or ""),
        "fileName": str(candidate.get("fileName") or ""),
        "path": str(candidate.get("sourceFile") or candidate.get("path") or ""),
        "paragraphIndex": candidate.get("paragraphIndex"),
        "raw_text": str(candidate.get("rawText") or ""),
        "rawText": str(candidate.get("rawText") or ""),
        "basisText": basis_text,
        "searchText": basis_text,
        "title": str(candidate.get("title") or ""),
        "number": str(candidate.get("number") or ""),
        "contextTitle": str(candidate.get("contextTitle") or ""),
    }


def _toc_item_from_agent_candidate(order: int, candidate: dict[str, Any], decision: dict[str, Any]) -> dict[str, Any]:
    number = str(decision.get("number") or candidate.get("number") or "").strip()
    title = str(decision.get("title") or candidate.get("title") or "").strip()
    if number and title.startswith(number):
        title = title[len(number) :].strip(" ：:、.-")
    return {
        "itemId": f"TOC-{order:04d}",
        "order": order,
        "number": number,
        "title": title or f"未命名附表{order}",
        "level": _coerce_toc_level(decision.get("level") or 1),
        "annotation": str(decision.get("annotation") or "新增-副表"),
        "source": "tender",
        "reason": str(decision.get("reason") or "Agent 判断招标文件要求追加该目录项。"),
        "source_refs": [_source_ref_from_agent_candidate(candidate, decision)],
        "material_refs": [],
    }


def _coerce_confidence(value: Any) -> float:
    try:
        parsed = float(value)
    except (TypeError, ValueError):
        return 0.0
    return max(0.0, min(1.0, parsed))


def _load_json_dict(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}
    return payload if isinstance(payload, dict) else {}


def _public_rule_evidence_from_file(path: Path) -> dict[str, Any]:
    evidence = _load_json_dict(path)
    decisions = evidence.get("decisions") if isinstance(evidence.get("decisions"), list) else []
    if evidence.get("schema_version") == "bid-toc-evidence-v2":
        action_counts = Counter(
            str(item.get("action") or "")
            for item in decisions
            if isinstance(item, dict) and str(item.get("action") or "")
        )
        return {
            "schemaVersion": "bid-toc-evidence-v2",
            "engine": str(evidence.get("engine") or ""),
            "ruleVersion": str(evidence.get("ruleVersion") or evidence.get("rule_version") or ""),
            "decisionCount": len(decisions),
            "reviewCount": sum(
                1
                for item in decisions
                if isinstance(item, dict) and bool(item.get("review_required") or item.get("reviewRequired"))
            ),
            "actionCounts": dict(action_counts),
        }
    candidates = evidence.get("tenderCandidates") if isinstance(evidence.get("tenderCandidates"), list) else []
    template_outline = evidence.get("templateOutline") if isinstance(evidence.get("templateOutline"), list) else []
    item_sources = evidence.get("itemSources") if isinstance(evidence.get("itemSources"), list) else []
    return {
        "schemaVersion": str(evidence.get("schema_version") or ""),
        "engine": str(evidence.get("engine") or ""),
        "ruleVersion": str(evidence.get("ruleVersion") or evidence.get("rule_version") or ""),
        "templateOutlineCount": len(template_outline),
        "tenderCandidateCount": len(candidates),
        "itemSources": [
            dict(item)
            for item in item_sources
            if isinstance(item, dict)
        ][:PUBLIC_EVIDENCE_DECISION_LIMIT],
        "itemSourceCount": len(item_sources),
        "decisions": [
            dict(item)
            for item in decisions
            if isinstance(item, dict)
        ][:PUBLIC_EVIDENCE_DECISION_LIMIT],
        "decisionCount": len(decisions),
        "agentDecisions": [
            dict(item)
            for item in (evidence.get("agentDecisions") if isinstance(evidence.get("agentDecisions"), list) else [])
            if isinstance(item, dict)
        ][:PUBLIC_EVIDENCE_DECISION_LIMIT],
    }


def _title_key(value: str) -> str:
    text = re.sub(r"\s+", "", str(value or "").lower())
    text = re.sub(r"[，,。.:：;；、（）()\[\]【】《》<>\"'“”‘’\\/_-]+", "", text)
    return text


def _prepare_toc_skill_workspace(
    *,
    project_id: str,
    project: dict[str, Any],
    parse_storage: dict[str, Any],
    tender_file_records: list[dict[str, Any]],
    template_file_records: list[dict[str, Any]],
) -> dict[str, Any]:
    bid_type = _outline_bid_type(str(project.get("bidType") or ""))
    project_dir = workspace_dir(project_id, bid_type)
    project_dir.mkdir(parents=True, exist_ok=True)

    published_work_dir = project_dir / "s2_toc_workdir"
    staging_work_dir = project_dir / "s2_toc_workdir.new"
    archive_root = project_dir / "s2_toc_workdir.runs"
    _archive_workspace_if_exists(staging_work_dir, archive_root, "stale")
    staging_work_dir.mkdir(parents=True, exist_ok=True)
    _remove_manifest_alias(project_dir)

    tender_inputs = _copy_tender_inputs(tender_file_records, staging_work_dir, parse_storage)
    template_path, attach_path = _copy_template_inputs(template_file_records, staging_work_dir, project_id)
    if template_path is None:
        raise ValueError("投标模板不存在，请先上传可读取的投标模板文件。")
    output_file = staging_work_dir / _safe_file_name(settings.s2_toc_output_file_name, "toc.json")
    manifest_path = staging_work_dir / "s2_input.json"
    manifest = {
        "projectId": project_id,
        "projectCode": str(project.get("projectCode") or project_id),
        "projectName": str(project.get("name") or project_id),
        "bidType": bid_type,
        "workDir": str(staging_work_dir),
        "tenderFiles": tender_inputs,
        "templateFile": str(template_path) if template_path else "",
        "attachFile": str(attach_path) if attach_path else "",
        "outputFile": str(output_file),
    }
    # 下游对接：项目机型与 S3 事实表值注入 manifest（须在 _trustedManifest 快照之前）。
    # S2 通常早于 S3 事实表构建，缺失时不写这两个键，不报错。
    turbine_model = project_turbine_model(project)
    if turbine_model:
        manifest["turbineModel"] = turbine_model
    project_facts = project_facts_for_manifest(project)
    if project_facts:
        manifest["projectFacts"] = project_facts
    if _is_business_bid(bid_type):
        manifest["evidenceFile"] = str(
            staging_work_dir / _safe_file_name(settings.s2_toc_evidence_file_name, "toc_evidence.json")
        )
    else:
        manifest["requireComposedOutline"] = True
    manifest_text = json.dumps(manifest, ensure_ascii=False, indent=2)
    manifest_path.write_text(manifest_text, encoding="utf-8")
    return {
        **manifest,
        "_trustedManifest": copy.deepcopy(manifest),
        "manifestPath": str(manifest_path),
        "canonicalManifestPath": str(manifest_path),
        "publishedWorkDir": str(published_work_dir),
        "stagingWorkDir": str(staging_work_dir),
        "archiveRoot": str(archive_root),
        "tenderFileCount": len(tender_inputs),
        "templateFileCount": 1 if template_path else 0,
        "hasAttachFile": bool(attach_path),
    }


def _publish_toc_skill_workspace(skill_workspace: dict[str, Any], toc_result: dict[str, Any]) -> dict[str, Any]:
    staging_work_dir = Path(str(skill_workspace.get("stagingWorkDir") or skill_workspace.get("workDir") or "")).expanduser()
    published_work_dir = Path(str(skill_workspace.get("publishedWorkDir") or skill_workspace.get("workDir") or "")).expanduser()
    archive_root = Path(str(skill_workspace.get("archiveRoot") or published_work_dir.with_name("s2_toc_workdir.runs"))).expanduser()
    if not staging_work_dir.exists():
        raise RuntimeError(f"S2 staging 工作目录不存在：{staging_work_dir}")

    replacements = {str(staging_work_dir): str(published_work_dir)}
    staging_manifest_path = staging_work_dir / "s2_input.json"
    current_manifest = _load_json_dict(staging_manifest_path)
    is_business_bid = _is_business_bid(skill_workspace.get("bidType"))
    trusted_manifest = skill_workspace.get("_trustedManifest")
    if not is_business_bid:
        if not isinstance(trusted_manifest, dict) or current_manifest != trusted_manifest:
            raise RuntimeError("技术标目录 manifest 在 Opencode 执行期间被修改。")
        manifest = copy.deepcopy(trusted_manifest)
    else:
        manifest = current_manifest
    staging_output_file = Path(
        str(manifest.get("outputFile") or staging_work_dir / _safe_file_name(settings.s2_toc_output_file_name, "toc.json"))
    ).expanduser()
    staging_evidence_file = (
        Path(str(manifest.get("evidenceFile") or "")).expanduser()
        if is_business_bid
        else None
    )
    if is_business_bid:
        manifest = _remap_workspace_paths(manifest, replacements)
    else:
        manifest["workDir"] = str(published_work_dir)
        manifest["outputFile"] = str(published_work_dir / staging_output_file.name)
        for key in ("templateFile", "attachFile"):
            if str(manifest.get(key) or ""):
                manifest[key] = _remap_workspace_paths(manifest[key], replacements)
        for tender_file in manifest.get("tenderFiles") or []:
            if not isinstance(tender_file, dict):
                continue
            for key in ("path", "originalPath"):
                if str(tender_file.get(key) or ""):
                    tender_file[key] = _remap_workspace_paths(tender_file[key], replacements)
    manifest_path = published_work_dir / "s2_input.json"
    output_file = Path(
        str(manifest.get("outputFile") or published_work_dir / _safe_file_name(settings.s2_toc_output_file_name, "toc.json"))
    ).expanduser()
    evidence_file = (
        Path(
            str(
                manifest.get("evidenceFile")
                or published_work_dir / _safe_file_name(settings.s2_toc_evidence_file_name, "toc_evidence.json")
            )
        ).expanduser()
        if is_business_bid
        else None
    )
    business_outline_file = published_work_dir / "outline.json"
    tender_map_inputs_file = published_work_dir / "tender_map_inputs.json"
    history_bid_outline_inputs_file = published_work_dir / "history_bid_outline_inputs.json"
    manifest["workDir"] = str(published_work_dir)
    manifest["outputFile"] = str(output_file)
    if evidence_file is not None:
        manifest["evidenceFile"] = str(evidence_file)
    else:
        manifest.pop("evidenceFile", None)
    staging_manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    if is_business_bid:
        _remap_json_file(
            staging_output_file,
            replacements,
            ({"outputFile": str(output_file), "evidenceFile": str(evidence_file)} if evidence_file is not None else None),
        )
    if staging_evidence_file is not None:
        _remap_json_file(staging_evidence_file, replacements)
    _remap_json_file(staging_work_dir / "outline.json", replacements)
    _remap_json_file(staging_work_dir / "tender_map_inputs.json", replacements)
    _remap_json_file(staging_work_dir / "history_bid_outline_inputs.json", replacements)
    compose_report_path = staging_work_dir / "outline_compose_report.json"
    if is_business_bid:
        _remap_json_file(compose_report_path, replacements)
    else:
        _remap_json_file(compose_report_path, {}, {"outputFile": str(output_file)})

    result = (
        _remap_workspace_paths(toc_result, replacements)
        if is_business_bid
        else copy.deepcopy(toc_result)
    )
    result["outputFile"] = str(output_file)
    if evidence_file is not None:
        result["evidenceFile"] = str(evidence_file)
    else:
        result.pop("evidenceFile", None)
    if (staging_work_dir / "outline.json").exists():
        result["businessOutlineFile"] = str(business_outline_file)
    if (staging_work_dir / "tender_map_inputs.json").exists():
        result["tenderMapInputsFile"] = str(tender_map_inputs_file)
    if (staging_work_dir / "history_bid_outline_inputs.json").exists():
        result["historyBidOutlineInputsFile"] = str(history_bid_outline_inputs_file)
    if isinstance(result.get("opencodeOutput"), dict):
        result["opencodeOutput"]["workDir"] = str(published_work_dir)
        result["opencodeOutput"]["manifestPath"] = str(manifest_path)
        result["opencodeOutput"]["canonicalManifestPath"] = str(manifest_path)
        result["opencodeOutput"]["tocJsonPath"] = str(output_file)
        if evidence_file is not None:
            result["opencodeOutput"]["evidencePath"] = str(evidence_file)
        else:
            result["opencodeOutput"].pop("evidencePath", None)
        if (staging_work_dir / "outline.json").exists():
            result["opencodeOutput"]["businessOutlinePath"] = str(business_outline_file)
        if (staging_work_dir / "tender_map_inputs.json").exists():
            result["opencodeOutput"]["tenderMapInputsPath"] = str(tender_map_inputs_file)
        if (staging_work_dir / "history_bid_outline_inputs.json").exists():
            result["opencodeOutput"]["historyBidOutlineInputsPath"] = str(history_bid_outline_inputs_file)

    previous_archive = ""
    try:
        previous_archive = _archive_workspace_if_exists(published_work_dir, archive_root, "previous")
        published_work_dir.parent.mkdir(parents=True, exist_ok=True)
        shutil.move(str(staging_work_dir), str(published_work_dir))
    except Exception:
        previous_archive_path = Path(previous_archive) if previous_archive else None
        if previous_archive_path and previous_archive_path.exists() and not published_work_dir.exists():
            shutil.move(str(previous_archive_path), str(published_work_dir))
        raise

    publish_result = {
        "result": result,
        "workDir": str(published_work_dir),
        "stagingWorkDir": str(staging_work_dir),
        "archiveRoot": str(archive_root),
        "previousArchive": previous_archive,
        "manifestPath": str(manifest_path),
        "canonicalManifestPath": str(manifest_path),
        "outputFile": str(output_file),
    }
    if evidence_file is not None:
        publish_result["evidenceFile"] = str(evidence_file)
    return publish_result


def _archive_workspace_if_exists(target: Path, archive_root: Path, label: str) -> str:
    if not target.exists():
        return ""
    archive_root.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    archive_path = _unique_path(archive_root / f"{timestamp}-{label}-{target.name}")
    shutil.move(str(target), str(archive_path))
    return str(archive_path)


def _remove_manifest_alias(project_dir: Path) -> None:
    alias_path = project_dir / "s2.json"
    if alias_path.exists():
        alias_path.unlink()


def _remap_json_file(path: Path, replacements: dict[str, str], updates: dict[str, Any] | None = None) -> None:
    if not path.exists():
        return
    payload = _remap_workspace_paths(_load_json_dict(path), replacements)
    if updates:
        payload.update(updates)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def _remap_workspace_paths(value: Any, replacements: dict[str, str]) -> Any:
    if isinstance(value, str):
        result = value
        for old, new in replacements.items():
            result = result.replace(old, new)
        return result
    if isinstance(value, list):
        return [_remap_workspace_paths(item, replacements) for item in value]
    if isinstance(value, dict):
        return {key: _remap_workspace_paths(item, replacements) for key, item in value.items()}
    return value


def _heading_style_for_line(line: str) -> str | None:
    text = str(line or "").strip()
    if re.match(r"^第[一二三四五六七八九十百千万零〇两0-9]+章", text):
        return "Heading 1"
    match = re.match(r"^(?P<number>\d+(?:\.\d+)*)(?:[\s　:：、.-]+)", text)
    if not match:
        return None
    level = min(match.group("number").count(".") + 1, 4)
    return f"Heading {level}"


def _write_text_docx(path: Path, title: str, text: str) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_paragraph(title)
    for line in str(text or "").splitlines():
        if line.strip():
            style = _heading_style_for_line(line)
            document.add_paragraph(line.strip(), style=style) if style else document.add_paragraph(line.strip())
    document.save(path)
    return path


def _copy_tender_inputs(file_records: list[dict[str, Any]], work_dir: Path, parse_storage: dict[str, Any]) -> list[dict[str, str]]:
    copied: list[dict[str, str]] = []
    for index, record in enumerate(file_records, start=1):
        source = Path(str(record.get("path") or "")).expanduser()
        if not source.exists() or source.suffix.lower() != ".docx":
            continue
        name = _safe_file_name(str(record.get("name") or source.name), f"tender-{index}.docx")
        destination = _unique_path(work_dir / name)
        shutil.copy2(source, destination)
        copied.append(
            {
                "id": str(record.get("id") or f"tender-{index}"),
                "name": name,
                "path": str(destination),
                "originalPath": str(source),
            }
        )
    if not copied:
        combined_text_path = Path(str(parse_storage.get("combinedTextPath") or "")).expanduser()
        if combined_text_path.exists():
            text = combined_text_path.read_text(encoding="utf-8", errors="replace")
            generated_path = _write_text_docx(work_dir / "tender-from-s1-text.docx", "招标文件解析文本", text)
            copied.append(
                {
                    "id": "TEN-S1-TEXT",
                    "name": "招标文件解析文本.docx",
                    "path": str(generated_path),
                    "originalPath": str(combined_text_path),
                }
            )
    return copied


def _copy_template_inputs(file_records: list[dict[str, Any]], work_dir: Path, project_id: str) -> tuple[Path | None, Path | None]:
    docx_records = [
        record
        for record in file_records
        if Path(str(record.get("path") or "")).expanduser().exists()
        and Path(str(record.get("path") or "")).suffix.lower() == ".docx"
    ]
    attach_record = next((record for record in docx_records if _looks_like_attachment_template(record)), None)
    template_record = next(
        (
            record
            for record in docx_records
            if record is not attach_record
        ),
        None,
    )
    if template_record is None and docx_records:
        template_record = docx_records[0]

    template_path = None
    if template_record is not None:
        template_path = _copy_single_template(template_record, work_dir / "template-main.docx")
    elif file_records:
        template_path = _copy_visual_template_input(project_id, file_records[0], work_dir / "template-main.docx")
    attach_path = (
        _copy_single_template(attach_record, work_dir / "template-attachment.docx")
        if attach_record is not None and attach_record is not template_record
        else None
    )
    return template_path, attach_path


def _looks_like_attachment_template(record: dict[str, Any]) -> bool:
    role = str(record.get("role") or record.get("templateRole") or record.get("type") or "").lower()
    if role in {"attachment", "attachments", "appendix", "appendices", "attach"}:
        return True
    name = str(record.get("name") or Path(str(record.get("path") or "")).name)
    return bool(re.search(r"(附表|附件|appendix|attachment|attach)", name, re.IGNORECASE))


def _copy_single_template(record: dict[str, Any], destination: Path) -> Path:
    source = Path(str(record.get("path") or "")).expanduser()
    if source.suffix.lower() == ".docx" and not is_valid_docx_file(source):
        source_label = "系统默认模板" if str(record.get("source") or "") == "system-default" else "投标模板"
        raise ValueError(f"{source_label}不是有效 DOCX 文件，请重新上传或更换默认模板。")
    resolved_destination = destination.with_suffix(source.suffix.lower() or ".docx")
    shutil.copy2(source, resolved_destination)
    return resolved_destination


def _copy_visual_template_input(project_id: str, record: dict[str, Any], destination: Path) -> Path | None:
    source = Path(str(record.get("path") or "")).expanduser()
    if not source.exists() or source.suffix.lower() not in {".pdf", *IMAGE_SUFFIXES}:
        return None
    text, meta = _ocr_fallback_text(project_id, record, source)
    if not text:
        message = meta.get("message") if isinstance(meta, dict) else ""
        raise ValueError(f"投标模板为图片或 PDF，但视觉模型未能读取：{message or '未知错误'}")
    return _write_text_docx(destination, str(record.get("name") or source.name), text)


def _safe_file_name(value: str, fallback: str) -> str:
    text = re.sub(r"[\\/:*?\"<>|]+", "-", str(value or "").strip())
    text = re.sub(r"\s+", " ", text).strip(" .")
    return text or fallback


def _unique_path(path: Path) -> Path:
    if not path.exists():
        return path
    stem = path.stem
    suffix = path.suffix
    for index in range(2, 1000):
        candidate = path.with_name(f"{stem}-{index}{suffix}")
        if not candidate.exists():
            return candidate
    raise RuntimeError(f"无法为文件生成唯一路径：{path}")


def _outline_bid_type(value: str) -> str:
    return require_bid_type(
        value,
        error_message="目录工作区必须显式传入技术标或商务标。",
    )


def _nodes_from_generation_result(
    result: dict[str, Any],
    *,
    compact_technical: bool = False,
) -> list[dict[str, Any]]:
    if compact_technical and isinstance(result.get("nodes"), list):
        return _clean_technical_outline_nodes(result["nodes"])
    if isinstance(result.get("items"), list):
        return _nodes_from_toc_items(result["items"])
    raise ValueError("目录 JSON 缺少 nodes[] 或 items[]。")


def _summary_from_generation_result(result: dict[str, Any]) -> str:
    summary = result.get("summary")
    if isinstance(summary, str) and summary.strip():
        return summary.strip()
    if isinstance(summary, dict):
        total_items = summary.get("total_nodes") if "total_nodes" in summary else summary.get("total_items")
        status_counts = (
            summary.get("action_counts")
            or summary.get("required_status_counts")
            or summary.get("annotation_counts")
            or {}
        )
        if isinstance(status_counts, dict) and status_counts:
            counts = "，".join(
                f"{key}{value}"
                for key, value in status_counts.items()
                if value
            )
            if counts:
                return f"目录生成完成，共 {total_items or 0} 条目录项（{counts}）。"
        return f"目录生成完成，共 {total_items or 0} 条目录项。"
    return "目录生成完成。"


def _clean_technical_outline_nodes(
    nodes: list[Any],
    *,
    parent_id: str = "OL",
) -> list[dict[str, Any]]:
    cleaned: list[dict[str, Any]] = []
    for index, raw_node in enumerate(nodes, start=1):
        if not isinstance(raw_node, dict):
            continue
        node_id = f"{parent_id}-{index}"
        action = str(raw_node.get("suggestion_action") or raw_node.get("suggestionAction") or "待确认").strip()
        if action not in TECHNICAL_SUGGESTION_ACTIONS:
            action = "待确认"
        reason = str(raw_node.get("suggestion_reason") or raw_node.get("suggestionReason") or "").strip()
        if action != "必要" and not reason:
            reason = "该目录项需要人工确认。"
        basis = raw_node.get("tender_basis")
        if not isinstance(basis, dict):
            basis = raw_node.get("tenderBasis") if isinstance(raw_node.get("tenderBasis"), dict) else None
        clean_basis = _clean_tender_basis(basis)
        number = str(raw_node.get("number") or raw_node.get("tocNumber") or "").strip()
        raw_children = raw_node.get("children") if isinstance(raw_node.get("children"), list) else []
        node = {
            "id": node_id,
            "number": number,
            "tocNumber": number,
            "title": str(raw_node.get("title") or "未命名章节").strip(),
            "suggestionAction": action,
            "suggestionReason": reason,
            "children": _clean_technical_outline_nodes(raw_children, parent_id=node_id),
        }
        if clean_basis:
            node["tenderBasis"] = clean_basis
        cleaned.append(node)
    return cleaned


def _clean_tender_basis(value: dict[str, Any] | None) -> dict[str, str] | None:
    if not isinstance(value, dict):
        return None
    file_id = str(value.get("file_id") or value.get("fileId") or "").strip()
    search_text = str(value.get("search_text") or value.get("searchText") or "").strip()
    if not file_id or not search_text:
        return None
    result = {"fileId": file_id, "searchText": search_text}
    evidence_id = str(value.get("evidence_id") or value.get("evidenceId") or "").strip()
    if evidence_id:
        result["evidenceId"] = evidence_id
    return result


def _nodes_from_toc_items(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    roots: list[dict[str, Any]] = []
    stack: list[tuple[int, dict[str, Any]]] = []
    counters: list[int] = []

    for fallback_order, item in enumerate(items, start=1):
        if not isinstance(item, dict):
            continue
        level = _coerce_toc_level(item.get("level"))
        if not stack and level > 1:
            level = 1
        elif stack and level > stack[-1][0] + 1:
            level = stack[-1][0] + 1

        while stack and stack[-1][0] >= level:
            stack.pop()

        counters = counters[:level]
        if len(counters) < level:
            counters.extend([0] * (level - len(counters)))
        counters[level - 1] += 1
        node_id = "OL-" + "-".join(str(part) for part in counters[:level] if part)

        title = _toc_item_title(item, fallback_order)
        source_text = _toc_item_source_text(item)
        annotation = str(item.get("annotation") or "").strip()
        required_status = str(item.get("required_status") or item.get("requiredStatus") or "").strip()
        if not required_status:
            required_status = _required_status_from_annotation(annotation)
        node = {
            "id": node_id,
            "title": title,
            "children": [],
            "tocNumber": str(item.get("number") or "").strip(),
            "annotation": annotation,
            "required_status": required_status,
            "requiredStatus": required_status,
            "source_text": source_text,
            "sourceText": source_text,
            "source": str(item.get("source") or "").strip(),
            "reason": str(item.get("reason") or "").strip(),
        }
        if isinstance(item.get("source_refs"), list):
            node["sourceRefs"] = item["source_refs"]
        if isinstance(item.get("material_refs"), list):
            node["materialRefs"] = item["material_refs"]
        if stack:
            stack[-1][1].setdefault("children", []).append(node)
        else:
            roots.append(node)
        stack.append((level, node))

    return roots


def _coerce_toc_level(value: Any) -> int:
    try:
        level = int(value)
    except (TypeError, ValueError):
        level = 1
    return max(1, level)


def _toc_item_title(
    item: dict[str, Any],
    fallback_order: int,
) -> str:
    title = str(item.get("title") or "").strip()
    number = str(item.get("number") or "").strip()
    if title:
        if str(item.get("source") or "").strip() == "business_outline":
            return title
        if number and not re.fullmatch(r"\d+(?:\.\d+)*", number):
            return f"{number} {title}".strip()
        return title
    if number:
        return number
    return f"未命名章节{fallback_order}"

def _technical_rule_evidence(nodes: list[Any]) -> dict[str, Any]:
    action_counts: Counter[str] = Counter()
    node_count = 0

    def walk(items: list[Any]) -> None:
        nonlocal node_count
        for item in items:
            if not isinstance(item, dict):
                continue
            node_count += 1
            action = str(item.get("suggestion_action") or item.get("suggestionAction") or "待确认").strip()
            action_counts[action if action in TECHNICAL_SUGGESTION_ACTIONS else "待确认"] += 1
            children = item.get("children")
            if isinstance(children, list):
                walk(children)

    walk(nodes)
    return {
        "schemaVersion": "technical-outline.v1",
        "engine": OUTLINE_SKILL_NAME,
        "nodeCount": node_count,
        "actionCounts": dict(action_counts),
    }


def _toc_item_source_text(item: dict[str, Any]) -> str:
    explicit = str(item.get("source_text") or item.get("sourceText") or "").strip()
    if explicit:
        return explicit
    source_refs = item.get("source_refs")
    if not isinstance(source_refs, list):
        source_refs = item.get("sourceRefs") if isinstance(item.get("sourceRefs"), list) else []
    for ref in source_refs:
        if not isinstance(ref, dict):
            continue
        text = str(
            ref.get("searchText")
            or ref.get("basisText")
            or ref.get("rawText")
            or ref.get("raw_text")
            or ""
        ).strip()
        if text:
            return text
    return ""
