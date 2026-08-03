from __future__ import annotations

import copy
import json
import logging
import re
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook

from app.core.config import settings
from app.services.bid_type import TECHNICAL_BID_TYPE
from app.services.file_utils import run_awaitable_sync
from app.services.identity import build_project_material_scope
from app.services.project_fact_materials import prepare_project_fact_material_files
from app.services.technical_fact_field_specs import (
    SPEC_LABEL_ALIASES,
    fillable_specs,
    spec_category,
)
from app.services.technical_fact_spec_versions import fact_specs_ref, resolve_project_specs
from app.services.technical_material_store import technical_material_store
from app.services.turbine_models import project_turbine_model

logger = logging.getLogger(__name__)


PROJECT_FACT_TABLE_SCHEMA_VERSION = "bid-project-fact-table-v2"

# 字段状态模型（七态，对应任务文档 §7）：
# unextracted 未提取 → extracted 已自动提取 → pending_confirmation 待人工确认 → confirmed 已人工确认
#   ├─ missing_source 缺少来源（登记了缺口但无值来源）
#   ├─ conflict 存在冲突（多来源同优先级不同值，人工裁决后进 confirmed）
#   └─ not_applicable 不适用（字段对本项目不适用，notes 注明原因）
FACT_STATUS_UNEXTRACTED = "unextracted"
FACT_STATUS_EXTRACTED = "extracted"
FACT_STATUS_PENDING_CONFIRMATION = "pending_confirmation"
FACT_STATUS_CONFIRMED = "confirmed"
FACT_STATUS_MISSING_SOURCE = "missing_source"
FACT_STATUS_CONFLICT = "conflict"
FACT_STATUS_NOT_APPLICABLE = "not_applicable"

FACT_FIELD_STATUSES = {
    FACT_STATUS_UNEXTRACTED,
    FACT_STATUS_EXTRACTED,
    FACT_STATUS_PENDING_CONFIRMATION,
    FACT_STATUS_CONFIRMED,
    FACT_STATUS_MISSING_SOURCE,
    FACT_STATUS_CONFLICT,
    FACT_STATUS_NOT_APPLICABLE,
}

# v1 四态 → v2 七态迁移映射
LEGACY_FACT_STATUS_MAP = {
    "candidate": FACT_STATUS_EXTRACTED,
    "missing": FACT_STATUS_MISSING_SOURCE,
}


def normalize_fact_status(status: Any, *, has_value: bool) -> str:
    """归一字段状态：旧四态映射为七态，非法/空值按有无取值给默认态。"""
    text = str(status or "").strip()
    text = LEGACY_FACT_STATUS_MAP.get(text, text)
    if text in FACT_FIELD_STATUSES:
        return text
    return FACT_STATUS_EXTRACTED if has_value else FACT_STATUS_MISSING_SOURCE

FACT_TABLE_HEADER_WORDS = {
    "编号",
    "序号",
    "项目",
    "名称",
    "内容",
    "备注",
    "说明",
    "单位",
    "计量单位",
    "技术参数与规格",
    "主要项目",
    "投标机型1",
    "投标机型2",
    "保证值",
    "授权人签名",
}

COMMON_PROJECT_FACT_LABELS = {
    "项目名称",
    "招标编号",
    "招标人",
    "招标方",
    "客户名称",
    "投标方案",
    "投标机型",
    "机组类型",
    "机组台数",
    "总装机容量",
    "单机容量",
    "叶轮直径",
    "轮毂高度",
    "扫风面积",
    "比功率",
    "安全等级",
    "设计寿命",
    "空气密度",
    "湍流强度",
    "极端风速",
    "年平均风速",
    "风剪切",
    "保证发电量",
    "保证有效小时数",
    "功率曲线保证率",
    "全场可利用率",
    "单台可利用率",
    "主要部件更换率",
}

FACT_MATERIAL_SOURCE_PRIORITIES = {
    "project": 300,
    "customer": 200,
    "standard": 100,
}


def empty_fact_summary() -> dict[str, int]:
    return {
        "totalCount": 0,
        "requiredCount": 0,
        "confirmedCount": 0,
        "extractedCount": 0,
        "pendingConfirmationCount": 0,
        "unextractedCount": 0,
        "missingSourceCount": 0,
        "conflictCount": 0,
        "notApplicableCount": 0,
        "specTotal": 0,
        "specBuiltTotal": 0,
        # deprecated：旧口径"有值的 spec 行数"，保留兼容，前端展示改用下方四段确认进度
        "specMatched": 0,
        # 清单确认进度四段：互斥穷尽已构建的 spec 行，加总 == specBuiltTotal
        "specConfirmedCount": 0,
        "specPendingConfirmationCount": 0,
        "specUnfilledCount": 0,
        "specFilledUnconfirmedCount": 0,
        # v1 兼容别名：candidate=已自动提取，missing=缺少来源
        "candidateCount": 0,
        "missingCount": 0,
    }


def empty_project_fact_table(project_id: str) -> dict[str, Any]:
    return {
        "schemaVersion": PROJECT_FACT_TABLE_SCHEMA_VERSION,
        "projectId": project_id,
        "status": "empty",
        "builtAt": "",
        "updatedAt": "",
        "confirmedAt": "",
        "confirmedBy": "",
        "fields": [],
        "summary": empty_fact_summary(),
    }


def summarize_project_fact_fields(fields: list[dict[str, Any]], spec_total: int | None = None) -> dict[str, int]:
    def count(status: str) -> int:
        return sum(1 for field in fields if str(field.get("status") or "") == status)

    # 四段进度只统计规则骨架行；specSeq=0 也是合法序号，不能按真值过滤。
    spec_rows = [field for field in fields if field.get("specSeq") is not None]
    spec_confirmed = sum(1 for field in spec_rows if str(field.get("status") or "") == FACT_STATUS_CONFIRMED)
    spec_pending = sum(1 for field in spec_rows if str(field.get("status") or "") == FACT_STATUS_PENDING_CONFIRMATION)
    # 「未填」：值为空，或状态仍属 unextracted/missing_source；confirmed/pending 已各自成段，不重复计
    spec_unfilled = sum(
        1
        for field in spec_rows
        if str(field.get("status") or "") not in {FACT_STATUS_CONFIRMED, FACT_STATUS_PENDING_CONFIRMATION}
        and (
            not str(field.get("value") or "").strip()
            or str(field.get("status") or "") in {FACT_STATUS_UNEXTRACTED, FACT_STATUS_MISSING_SOURCE}
        )
    )
    summary = empty_fact_summary()
    summary.update(
        {
            "totalCount": len(fields),
            "requiredCount": sum(1 for field in fields if field.get("required", True)),
            "confirmedCount": count(FACT_STATUS_CONFIRMED),
            "extractedCount": count(FACT_STATUS_EXTRACTED),
            "pendingConfirmationCount": count(FACT_STATUS_PENDING_CONFIRMATION),
            "unextractedCount": count(FACT_STATUS_UNEXTRACTED),
            "missingSourceCount": count(FACT_STATUS_MISSING_SOURCE),
            "conflictCount": count(FACT_STATUS_CONFLICT),
            "notApplicableCount": count(FACT_STATUS_NOT_APPLICABLE),
            # Dev 口径保持不变：已绑定规则条数是稳定分母，不能随当前建表结果波动。
            "specTotal": len(spec_rows) if spec_total is None else max(0, int(spec_total)),
            "specBuiltTotal": len(spec_rows),
            "specMatched": sum(1 for field in spec_rows if str(field.get("value") or "").strip()),
            "specConfirmedCount": spec_confirmed,
            "specPendingConfirmationCount": spec_pending,
            "specUnfilledCount": spec_unfilled,
            # 剩余 spec 行（extracted/conflict 等有值但未确认）：四段加总 == specBuiltTotal
            "specFilledUnconfirmedCount": len(spec_rows) - spec_confirmed - spec_pending - spec_unfilled,
        }
    )
    summary["candidateCount"] = summary["extractedCount"]
    summary["missingCount"] = summary["missingSourceCount"]
    return summary


def canonical_fact_label(label: Any) -> str:
    raw = str(label or "").strip()
    if not raw:
        return ""
    text = re.sub(r"\s+", "", raw)
    text = re.sub(r"[（(]\s*(?:MW|kW|m|m2/kW|m²/kW|%|h|MWh/y|MWh/a|台)\s*[）)]", "", text, flags=re.I)
    text = text.strip("：:；;，,、")
    aliases = {
        "方案": "投标方案",
        "项目方案": "投标方案",
        "机型": "投标方案",
        "建设容量": "总装机容量",
        "标段规模": "总装机容量",
        "机组数量": "机组台数",
        "风机数量": "机组台数",
        "台数": "机组台数",
        "总容量": "总装机容量",
        "容量": "总装机容量",
        "单机容量": "单机容量",
        "机组额定功率": "单机容量",
        "项目编号": "招标编号",
        "招标文件编号": "招标编号",
        "项目单位": "招标人",
        "建设单位": "招标人",
        "业主": "招标人",
        "交货期": "交货周期",
        "质量保证期": "质保期",
        "投标截止时间": "投标截止日期",
        "轮毂中心高度": "轮毂高度",
        "轮毂高度": "轮毂高度",
        "风轮直径": "叶轮直径",
        "叶轮直径": "叶轮直径",
        "发电小时数承诺": "保证有效小时数",
        "保证有效小时": "保证有效小时数",
        "风电机组设备年平均可利用率保证值": "全场可利用率",
        "适用等级": "安全等级",
    }
    if text in aliases:
        return aliases[text]
    if "总装机容量" in text or text.startswith("总容量"):
        return "总装机容量"
    if (
        "年平均风速" in text
        or "代表年风速" in text
        or ("平均风速" in text and ("机位" in text or "尾流" in text or "轮毂" in text))
    ):
        return "年平均风速"
    if "轮毂" in text and "高度" in text:
        return "轮毂高度"
    if "叶轮直径" in text or "风轮直径" in text:
        return "叶轮直径"
    if ("机组" in text or "风机" in text) and ("台数" in text or "数量" in text):
        return "机组台数"
    if "单机容量" in text or "额定功率" in text:
        return "单机容量"
    if "安全等级" in text or ("安全" in text and "等级" in text):
        return "安全等级"
    if "设计寿命" in text:
        return "设计寿命"
    if "单位千瓦扫风面积" in text:
        return "单位千瓦扫风面积"
    if "空气密度" in text and not re.search(r"参数|系数", text):
        return "空气密度"
    if "湍流强度" in text:
        return "湍流强度"
    if "极端风速" in text or "极大风速" in text:
        return "极端风速"
    if "风剪切" in text or "风切变" in text or "风剪切指数" in text:
        return "风剪切"
    if "功率曲线" in text and ("保证" in text or "保证率" in text):
        return "功率曲线保证率"
    if "单台" in text and "可利用率" in text:
        return "单台可利用率"
    if ("全场" in text or "风电场" in text or "年平均" in text) and "可利用率" in text:
        return "全场可利用率"
    if "发电量" in text and ("保证" in text or "承诺" in text):
        return "保证发电量"
    if "有效小时" in text or "发电小时" in text or "等效利用小时" in text:
        return "保证有效小时数"
    return text


def fact_label_key(label: Any) -> str:
    return re.sub(r"\s+", "", canonical_fact_label(label)).lower()


def fact_source_ref_priority(ref: dict[str, Any]) -> int:
    source_type = str(ref.get("type") or "").strip()
    if source_type in {"project", "projectIdentity", "projectTurbineModel", "derived"}:
        return 320
    if source_type in {"materialFact", "derivedMaterialFact"}:
        tier = str(ref.get("materialTier") or "").strip() or "standard"
        return FACT_MATERIAL_SOURCE_PRIORITIES.get(tier, 50)
    return 0


def normalize_fact_source_refs(refs: Any) -> list[dict[str, Any]]:
    normalized: list[tuple[int, int, dict[str, Any]]] = []
    seen: set[str] = set()
    for index, ref in enumerate(refs if isinstance(refs, list) else []):
        if not isinstance(ref, dict):
            continue
        item = copy.deepcopy(ref)
        key = json.dumps(item, ensure_ascii=False, sort_keys=True)
        if key in seen:
            continue
        seen.add(key)
        normalized.append((fact_source_ref_priority(item), index, item))
    normalized.sort(key=lambda item: (-item[0], item[1]))
    return [item for _, _, item in normalized]


def normalize_project_fact_field(
    field: dict[str, Any],
    *,
    index: int,
    confirm: bool,
    operator: str,
    saved_at: str,
) -> dict[str, Any]:
    value = str(field.get("value") or "").strip()
    if confirm:
        # 整表确认时保留人工标记的"不适用"，其余有值→已人工确认、无值→缺少来源
        incoming = normalize_fact_status(field.get("status"), has_value=bool(value))
        status = (
            incoming
            if incoming == FACT_STATUS_NOT_APPLICABLE
            else (FACT_STATUS_CONFIRMED if value else FACT_STATUS_MISSING_SOURCE)
        )
    else:
        status = normalize_fact_status(field.get("status"), has_value=bool(value))
    source_refs = normalize_fact_source_refs(field.get("sourceRefs"))
    source_priority = int(field.get("sourcePriority") or 0)
    if source_refs:
        source_priority = max(source_priority, fact_source_ref_priority(source_refs[0]))
    normalized = {
        "id": str(field.get("id") or f"FACT-{index:04d}"),
        "key": str(field.get("key") or fact_label_key(field.get("label")) or f"fact-{index}"),
        "label": str(field.get("label") or ""),
        "category": str(field.get("category") or "项目事实"),
        "value": value,
        "unit": str(field.get("unit") or ""),
        "required": bool(field.get("required", True)),
        "status": status,
        "confidence": float(field.get("confidence") or 0),
        "sourcePriority": source_priority,
        "sourceRefs": source_refs,
        "alternatives": copy.deepcopy(field.get("alternatives") if isinstance(field.get("alternatives"), list) else []),
        "notes": str(field.get("notes") or ""),
        "updatedAt": saved_at,
        "updatedBy": operator,
    }
    # 清单 spec 元数据（有则保留，供前端展示"待确认"标记与复核口径）
    for meta_key in ("specSeq", "specKey", "reviewLabel", "needsConfirmation", "sourceKind", "sourceHint"):
        if field.get(meta_key) is not None:
            normalized[meta_key] = copy.deepcopy(field.get(meta_key))
    if field.get("outOfSpec"):
        normalized["outOfSpec"] = True
    if normalized["status"] == FACT_STATUS_CONFIRMED:
        normalized["confirmedAt"] = saved_at
        normalized["confirmedBy"] = operator
    else:
        normalized["confirmedAt"] = str(field.get("confirmedAt") or "")
        normalized["confirmedBy"] = str(field.get("confirmedBy") or "")
    return normalized


def fact_table_value_map(fact_table: dict[str, Any]) -> dict[str, str]:
    values: dict[str, str] = {}
    for field in fact_table.get("fields") or []:
        if not isinstance(field, dict):
            continue
        label = canonical_fact_label(str(field.get("label") or field.get("title") or ""))
        value = str(field.get("value") or "").strip()
        if label and value:
            values[label] = value
            values[fact_label_key(label)] = value
    return values


def _spec_match_keys(spec: dict[str, Any]) -> list[str]:
    """spec 的候选匹配键：label、reviewLabel、手工别名，统一走 canonical 归一。"""
    keys: list[str] = []
    for label in [spec.get("label"), spec.get("reviewLabel"), *SPEC_LABEL_ALIASES.get(str(spec.get("label") or ""), [])]:
        key = fact_label_key(label)
        if key and key not in keys:
            keys.append(key)
    return keys


def reconcile_fact_fields_with_specs(
    fields_by_key: dict[str, dict[str, Any]],
    existing_by_key: dict[str, dict[str, Any]] | None = None,
    specs: list[dict[str, Any]] | None = None,
) -> None:
    """以填值 spec 为骨架对齐抽取结果。

    - 匹配到的字段打上 spec 元数据；needsConfirmation 且已自动提取的转"待人工确认"。
    - 未匹配到的 spec 生成"未提取"骨架字段，保证清单字段在事实表中齐全。
    - 一个启发式字段只归属一个 spec（按清单序号顺序先到先得）。
    - 上一轮已有人工值/人工状态的 spec 字段在重建时保留（人工确认结果不丢）。
    - specs 为 None 时回退全局 fillable_specs()；项目构建链路显式传项目级清单
      （用户上传的事实表 Excel 解析结果）。
    """
    existing_by_key = existing_by_key or {}
    matched_field_keys: set[str] = set()
    for spec in (specs if specs is not None else fillable_specs()):
        match_keys = _spec_match_keys(spec)
        field: dict[str, Any] | None = None
        for key in match_keys:
            candidate = fields_by_key.get(key)
            if candidate is not None and key not in matched_field_keys:
                field = candidate
                matched_field_keys.add(key)
                break
        if field is not None:
            field.pop("outOfSpec", None)
            field["specSeq"] = int(spec.get("seq") or 0)
            field["specKey"] = str(spec.get("key") or "")
            field["reviewLabel"] = str(spec.get("reviewLabel") or "")
            field["needsConfirmation"] = bool(spec.get("needsConfirmation"))
            field["sourceKind"] = str(spec.get("sourceKind") or "")
            if spec.get("needsConfirmation") and str(field.get("status") or "") == FACT_STATUS_EXTRACTED:
                field["status"] = FACT_STATUS_PENDING_CONFIRMATION
            continue
        key = fact_label_key(spec.get("label")) or f"spec-{int(spec.get('seq') or 0):03d}"
        if key in fields_by_key:
            key = f"spec-{int(spec.get('seq') or 0):03d}"
        skeleton = {
            "id": "",
            "key": key,
            "label": str(spec.get("label") or ""),
            "category": spec_category(spec),
            "value": "",
            "unit": "",
            "required": True,
            "status": FACT_STATUS_UNEXTRACTED,
            "confidence": 0.0,
            "sourcePriority": 0,
            "sourceRefs": [],
            "alternatives": [],
            "notes": str(spec.get("note") or ""),
            "updatedAt": "",
            "updatedBy": "",
            "confirmedAt": "",
            "confirmedBy": "",
            "specSeq": int(spec.get("seq") or 0),
            "specKey": str(spec.get("key") or ""),
            "reviewLabel": str(spec.get("reviewLabel") or ""),
            "needsConfirmation": bool(spec.get("needsConfirmation")),
            "sourceKind": str(spec.get("sourceKind") or ""),
            "sourceHint": str(spec.get("referenceFile") or ""),
        }
        # 上一轮的人工结果（有值或人工置过的状态）随重建保留
        previous = next(
            (existing_by_key[k] for k in [key, *match_keys, f"spec-{int(spec.get('seq') or 0):03d}"] if k in existing_by_key),
            None,
        )
        if previous is not None:
            prev_value = str(previous.get("value") or "").strip()
            prev_status = normalize_fact_status(previous.get("status"), has_value=bool(prev_value))
            if prev_value or prev_status in {FACT_STATUS_CONFIRMED, FACT_STATUS_NOT_APPLICABLE, FACT_STATUS_PENDING_CONFIRMATION}:
                skeleton["value"] = prev_value
                skeleton["unit"] = str(previous.get("unit") or "")
                skeleton["status"] = prev_status
                skeleton["sourceRefs"] = copy.deepcopy(
                    previous.get("sourceRefs") if isinstance(previous.get("sourceRefs"), list) else []
                )
                skeleton["confirmedAt"] = str(previous.get("confirmedAt") or "")
                skeleton["confirmedBy"] = str(previous.get("confirmedBy") or "")
        fields_by_key[key] = skeleton
        # 骨架键同样占位，避免后续 spec 把别人的骨架当成匹配字段
        matched_field_keys.add(key)


def is_manual_fact_field(field: dict[str, Any]) -> bool:
    """人工新增字段：sourceRefs 含 manualFact 来源。无 specSeq，不计入清单统计。"""
    refs = field.get("sourceRefs") if isinstance(field.get("sourceRefs"), list) else []
    return any(isinstance(ref, dict) and str(ref.get("type") or "") == "manualFact" for ref in refs)


def build_project_fact_table(project: dict[str, Any], gap_state: dict[str, Any]) -> dict[str, Any]:
    built_at = _now_iso()
    existing_table = gap_state.get("projectFactTable") if isinstance(gap_state.get("projectFactTable"), dict) else {}
    existing_by_key = {
        fact_label_key(field.get("label")): field
        for field in (existing_table.get("fields") if isinstance(existing_table.get("fields"), list) else [])
        if isinstance(field, dict) and fact_label_key(field.get("label"))
    }
    fields_by_key: dict[str, dict[str, Any]] = {}
    # 任务启动时固化规则快照（R06-B04-02）：本项目绑定版本优先，无绑定回落系统默认清单
    project_specs, fact_specs_meta = resolve_project_specs(gap_state)

    def is_material_fact_ref(ref: dict[str, Any]) -> bool:
        return str(ref.get("type") or "") in {"materialFact", "derivedMaterialFact"}

    def blank_source_paths() -> set[str]:
        paths: set[str] = set()
        plan = gap_state.get("plan") if isinstance(gap_state.get("plan"), dict) else {}
        for item in plan.get("items") or []:
            if not isinstance(item, dict):
                continue
            for task in item.get("fillTasks") or []:
                if not isinstance(task, dict):
                    continue
                blank = task.get("blankSource") if isinstance(task.get("blankSource"), dict) else {}
                for key in ("docxPath", "path", "workspacePath"):
                    value = str(blank.get(key) or "").strip()
                    if value:
                        paths.add(str(Path(value).resolve()))
        return paths

    def add_candidate(
        label: str,
        value: Any,
        *,
        category: str,
        source_ref: dict[str, Any],
        confidence: float = 0.8,
        required: bool = True,
        unit: str = "",
        source_priority: int = 0,
    ) -> None:
        label_text = canonical_fact_label(label)
        if not label_text:
            return
        key = fact_label_key(label_text)
        value_text = str(value or "").strip()
        existing = existing_by_key.get(key)
        preserve_existing = bool(
            existing
            and str(existing.get("status") or "") == FACT_STATUS_CONFIRMED
            and str(existing.get("value") or "").strip()
        )
        if preserve_existing:
            value_text = str(existing.get("value") or "").strip()
        field = fields_by_key.get(key)
        incoming_priority = int(source_priority or 0)
        if not field:
            field = {
                "id": str((existing or {}).get("id") or f"FACT-{len(fields_by_key) + 1:04d}"),
                "key": key,
                "label": label_text,
                "category": category,
                "value": value_text,
                "unit": str(((existing or {}).get("unit") if preserve_existing else unit) or ""),
                "required": bool((existing or {}).get("required", required)),
                "status": FACT_STATUS_CONFIRMED if preserve_existing else (FACT_STATUS_EXTRACTED if value_text else FACT_STATUS_MISSING_SOURCE),
                "confidence": float(((existing or {}).get("confidence") if preserve_existing else None) or (confidence if value_text else 0) or 0),
                "sourcePriority": int((existing or {}).get("sourcePriority") if preserve_existing else (incoming_priority if value_text else 0)),
                "sourceRefs": [],
                "alternatives": copy.deepcopy(
                    (existing or {}).get("alternatives")
                    if preserve_existing and isinstance((existing or {}).get("alternatives"), list)
                    else []
                ),
                "notes": str((existing or {}).get("notes") if preserve_existing else ""),
                "updatedAt": str((existing or {}).get("updatedAt") if preserve_existing else built_at),
                "updatedBy": str((existing or {}).get("updatedBy") if preserve_existing else ""),
                "confirmedAt": str((existing or {}).get("confirmedAt") if preserve_existing else ""),
                "confirmedBy": str((existing or {}).get("confirmedBy") if preserve_existing else ""),
            }
            fields_by_key[key] = field
        elif value_text and field.get("value") and value_text != field["value"]:
            alternatives = field.setdefault("alternatives", [])
            existing_rank = (int(field.get("sourcePriority") or 0), float(field.get("confidence") or 0))
            incoming_rank = (incoming_priority, float(confidence or 0))
            if incoming_rank > existing_rank and str(field.get("status") or "") != FACT_STATUS_CONFIRMED:
                old_value = str(field.get("value") or "")
                if old_value and old_value not in [str(item.get("value") or "") for item in alternatives if isinstance(item, dict)]:
                    alternatives.append({"value": old_value, "source": (field.get("sourceRefs") or [{}])[0]})
                field["value"] = value_text
                field["unit"] = str(unit or field.get("unit") or "")
                field["category"] = category
                field["status"] = FACT_STATUS_EXTRACTED
                field["confidence"] = float(confidence or 0)
                field["sourcePriority"] = incoming_priority
                if source_ref:
                    field["sourceRefs"] = [source_ref] + list(field.get("sourceRefs") or [])
                    source_ref = {}
            elif incoming_rank == existing_rank and str(field.get("status") or "") != FACT_STATUS_CONFIRMED:
                existing_material = any(
                    is_material_fact_ref(ref)
                    for ref in (field.get("sourceRefs") if isinstance(field.get("sourceRefs"), list) else [])
                    if isinstance(ref, dict)
                )
                if not (existing_material and is_material_fact_ref(source_ref)):
                    field["status"] = FACT_STATUS_CONFLICT
                if value_text not in [str(item.get("value") or "") for item in alternatives if isinstance(item, dict)]:
                    alternatives.append({"value": value_text, "source": source_ref})
            elif value_text not in [str(item.get("value") or "") for item in alternatives if isinstance(item, dict)]:
                alternatives.append({"value": value_text, "source": source_ref})
        elif value_text and field.get("value") and value_text == field.get("value"):
            existing_rank = (int(field.get("sourcePriority") or 0), float(field.get("confidence") or 0))
            incoming_rank = (incoming_priority, float(confidence or 0))
            if incoming_rank > existing_rank and str(field.get("status") or "") != FACT_STATUS_CONFIRMED:
                field["unit"] = str(unit or field.get("unit") or "")
                field["category"] = category
                field["confidence"] = float(confidence or 0)
                field["sourcePriority"] = incoming_priority
                if source_ref:
                    field["sourceRefs"] = [source_ref] + list(field.get("sourceRefs") or [])
                    source_ref = {}
        elif value_text and not field.get("value"):
            field["value"] = value_text
            field["status"] = FACT_STATUS_EXTRACTED
            field["unit"] = str(unit or field.get("unit") or "")
            field["confidence"] = max(float(field.get("confidence") or 0), float(confidence or 0))
            field["sourcePriority"] = incoming_priority
            field["category"] = category
            if source_ref:
                field["sourceRefs"] = [source_ref] + list(field.get("sourceRefs") or [])
                source_ref = {}
        if source_ref:
            field.setdefault("sourceRefs", []).append(source_ref)
        if preserve_existing and value_text:
            field["status"] = FACT_STATUS_CONFIRMED

    def preserve_compatible_existing_fields() -> None:
        """保留人工新增字段及旧规则下已确认的字段，避免重建时丢失人工结论。

        已确认字段先移除旧 spec 元数据，再参与当前规则对齐；若没有命中当前规则，
        则以 outOfSpec 标记留在表尾，且不计入当前规则版本的进度。
        """
        for existing in existing_by_key.values():
            if not isinstance(existing, dict):
                continue
            label_text = canonical_fact_label(existing.get("label"))
            key = fact_label_key(label_text)
            if not key:
                continue
            source_refs = [
                ref
                for ref in (existing.get("sourceRefs") if isinstance(existing.get("sourceRefs"), list) else [])
                if isinstance(ref, dict)
            ]
            is_manual = any(str(ref.get("type") or "") == "manualFact" for ref in source_refs)
            is_confirmed = str(existing.get("status") or "") == FACT_STATUS_CONFIRMED
            if not (is_manual or is_confirmed):
                continue
            has_value = bool(str(existing.get("value") or "").strip())
            field = copy.deepcopy(existing)
            for meta_key in ("specSeq", "specKey", "reviewLabel", "needsConfirmation", "sourceKind", "sourceHint"):
                field.pop(meta_key, None)
            field["label"] = label_text
            field["key"] = key
            field["category"] = str(field.get("category") or ("人工补充事实" if is_manual else "清单外历史事实"))
            field["status"] = normalize_fact_status(field.get("status"), has_value=has_value)
            field["sourceRefs"] = source_refs or (
                [{"type": "manualFact", "title": "人工新增", "field": label_text}] if is_manual else []
            )
            if is_confirmed and not is_manual:
                field["outOfSpec"] = True
            fields_by_key[key] = field

    trusted_parse_facts = trusted_parse_fact_fields(project.get("parse_result"))
    first_parse_value = {
        fact_label_key(fact.get("label")): fact.get("value")
        for fact in trusted_parse_facts
        if fact.get("value")
    }
    identity = project.get("identity") if isinstance(project.get("identity"), dict) else {}
    owner = identity.get("owner") or identity.get("customerCanonicalName") or identity.get("customerName") or project.get("owner") or project.get("customerName")
    project_name = first_parse_value.get(fact_label_key("项目名称")) or project.get("name")
    add_candidate("项目名称", project_name, category="项目基础信息", source_ref={"type": "project", "field": "name", "title": "项目名称"}, confidence=0.86, source_priority=320)
    add_candidate("招标方", owner, category="项目基础信息", source_ref={"type": "projectIdentity", "field": "owner", "title": "招标方"}, confidence=0.92, source_priority=320)
    add_candidate("招标人", owner, category="项目基础信息", source_ref={"type": "projectIdentity", "field": "owner", "title": "招标人"}, confidence=0.92, source_priority=320)
    add_candidate("客户名称", project.get("customerName"), category="项目基础信息", source_ref={"type": "project", "field": "customerName", "title": "客户名称"}, confidence=0.9, source_priority=320)
    add_candidate("日期", datetime.now(UTC).strftime("%Y年%m月%d日"), category="系统字段", source_ref={"type": "system", "field": "currentDate", "title": "当前日期"}, confidence=0.62)

    turbine = project_turbine_model(project)
    model = turbine.get("model") or turbine.get("turbineModel")
    hub_height = turbine.get("hubHeightM")
    add_candidate("投标机型", model, category="机型参数", source_ref={"type": "projectTurbineModel", "field": "model", "title": "投标机型"}, confidence=0.98, source_priority=320)
    rated_kw = turbine.get("ratedPowerKw")
    rated_mw = ""
    if isinstance(rated_kw, (int, float)):
        rated_mw = f"{rated_kw / 1000:g}"
    add_candidate("单机容量", rated_mw or rated_kw, category="机型参数", source_ref={"type": "projectTurbineModel", "field": "ratedPowerKw", "title": "单机容量"}, confidence=0.9, unit="MW" if rated_mw else "", source_priority=320)
    add_candidate("叶轮直径", turbine.get("rotorDiameterM"), category="机型参数", source_ref={"type": "projectTurbineModel", "field": "rotorDiameterM", "title": "叶轮直径"}, confidence=0.9, unit="m", source_priority=320)
    add_candidate("轮毂高度", hub_height, category="机型参数", source_ref={"type": "projectTurbineModel", "field": "hubHeightM", "title": "轮毂高度"}, confidence=0.86, unit="m", source_priority=320)
    if model and hub_height:
        add_candidate("投标方案", f"{model}-{hub_height}m", category="方案口径", source_ref={"type": "derived", "field": "modelHubHeight", "title": "投标方案"}, confidence=0.78, source_priority=320)
        add_candidate("方案", f"{model}-{hub_height}m", category="方案口径", source_ref={"type": "derived", "field": "modelHubHeight", "title": "方案"}, confidence=0.78, source_priority=320)
    elif model:
        add_candidate("投标方案", model, category="方案口径", source_ref={"type": "derived", "field": "model", "title": "投标方案"}, confidence=0.64, source_priority=80)

    for fact in trusted_parse_facts:
        add_candidate(
            str(fact.get("label") or ""),
            fact.get("value"),
            category=str(fact.get("category") or "招标解析字段"),
            source_ref=copy.deepcopy(fact.get("sourceRef") or {}),
            confidence=float(fact.get("confidence") or 0.82),
            required=bool(fact.get("required", False)),
            unit=str(fact.get("unit") or ""),
            source_priority=260,
        )

    for fact in project_material_fact_fields(
        project, gap_state, excluded_paths=blank_source_paths(), specs=project_specs
    ):
        if fact.get("internal"):
            continue
        add_candidate(
            str(fact.get("label") or ""),
            fact.get("value"),
            category=str(fact.get("category") or "素材库事实"),
            source_ref=copy.deepcopy(fact.get("sourceRef") or {}),
            confidence=float(fact.get("confidence") or 0.78),
            required=bool(fact.get("required", False)),
            unit=str(fact.get("unit") or ""),
            source_priority=int(fact.get("sourcePriority") or 0),
        )

    plan = gap_state.get("plan") if isinstance(gap_state.get("plan"), dict) else {}
    for item in plan.get("items") or []:
        if not isinstance(item, dict):
            continue
        for task in item.get("fillTasks") or []:
            if not isinstance(task, dict):
                continue
            blank = task.get("blankSource") if isinstance(task.get("blankSource"), dict) else {}
            for label in blank.get("placeholderLabels") or []:
                add_candidate(
                    str(label),
                    "",
                    category="待填写Word字段",
                    source_ref={
                        "type": "gapPlaceholder",
                        "gapId": str(item.get("id") or ""),
                        "title": str(item.get("title") or ""),
                        "field": str(label),
                    },
                    confidence=0.0,
                )
            for label in fillable_table_labels_from_blank_source(blank):
                add_candidate(
                    label,
                    "",
                    category="待填写表格字段",
                    source_ref={
                        "type": "gapTableField",
                        "gapId": str(item.get("id") or ""),
                        "title": str(item.get("title") or ""),
                        "field": label,
                        "blankSourceId": str(blank.get("id") or ""),
                    },
                    confidence=0.0,
                )

    for task in plan.get("tasks") or []:
        if not isinstance(task, dict):
            continue
        for label in technical_fact_labels_from_task(task):
            add_candidate(
                label,
                "",
                category="技术待填写字段",
                source_ref={
                    "type": "technicalGapTask",
                    "taskId": str(task.get("id") or ""),
                    "title": str(task.get("title") or ""),
                    "field": label,
                },
                confidence=0.0,
            )

    preserve_compatible_existing_fields()
    # 字段骨架来自任务启动时固化的规则快照（项目绑定版本，无绑定回落系统默认清单）
    spec_mode = bool(project_specs)
    if not spec_mode:
        logger.warning("项目 %s 无可用事实表字段规则，按来源并集构建", project.get("id"))
    reconcile_fact_fields_with_specs(fields_by_key, existing_by_key, project_specs)

    fields = list(fields_by_key.values())
    if spec_mode:
        # 以清单为唯一字段骨架：匹配不到 spec 的来源字段不再单独成行，
        # 只保留 spec 行、人工新增字段，以及旧规则下已经人工确认的兼容字段。
        fields = [
            field
            for field in fields
            if field.get("specSeq") is not None or is_manual_fact_field(field) or bool(field.get("outOfSpec"))
        ]
    for field in fields:
        source_refs = normalize_fact_source_refs(field.get("sourceRefs"))
        field["sourceRefs"] = source_refs
        if source_refs:
            field["sourcePriority"] = max(
                int(field.get("sourcePriority") or 0),
                fact_source_ref_priority(source_refs[0]),
            )
    category_order = {
        "项目基础信息": 0,
        "机型参数": 1,
        "方案口径": 2,
        "系统字段": 3,
        "招标解析字段": 4,
        "素材库事实": 5,
        "性能保证": 6,
        "待填写Word字段": 7,
        "待填写表格字段": 8,
        "技术待填写字段": 9,
        "清单-招标文件": 10,
        "清单-项目定制材料": 11,
        "清单-认证证书": 12,
        "清单-平台输入": 13,
        "清单-自动生成": 14,
    }
    fields.sort(
        key=lambda field: (
            # 清单模式下，人工新增和清单外历史字段追加在当前规则骨架之后。
            1 if spec_mode and field.get("specSeq") is None else 0,
            category_order.get(str(field.get("category") or ""), 9),
            0 if field.get("required") else 1,
            int(field.get("specSeq")) if field.get("specSeq") is not None else 9999,
            field.get("label") or "",
        )
    )
    # 补 id 时避开已有 id：骨架行按排序序号补 FACT-XXXX，可能与清单外候选自带的
    # FACT-XXXX 撞号（清单模式下杂行被过滤后序号前移），撞号则顺延到未占用的序号
    used_ids: set[str] = set()
    for index, field in enumerate(fields, start=1):
        field_id = str(field.get("id") or "")
        if not field_id or field_id in used_ids:
            serial = index
            field_id = f"FACT-{serial:04d}"
            while field_id in used_ids:
                serial += 1
                field_id = f"FACT-{serial:04d}"
        field["id"] = field_id
        used_ids.add(field_id)
    return {
        "schemaVersion": PROJECT_FACT_TABLE_SCHEMA_VERSION,
        "projectId": str(project.get("id") or ""),
        "status": "draft",
        "builtAt": built_at,
        "updatedAt": built_at,
        "confirmedAt": "",
        "confirmedBy": "",
        "fields": fields,
        "summary": summarize_project_fact_fields(fields, spec_total=len(project_specs)),
        # 本次构建实际使用的规则版本快照（审计：正式标书用了哪版规则）
        "factSpecsRef": fact_specs_meta,
    }


def technical_fact_labels_from_task(task: dict[str, Any]) -> list[str]:
    title = str(task.get("title") or "")
    task_key = str(task.get("taskKey") or "")
    text = f"{title} {task_key}"
    labels: list[str] = []
    if re.search(r"技术响应|技术偏差|技术参数|技术方案|供货范围|机组|风机|塔筒|叶片|发电量|功率曲线|保证值|承诺|声明", text):
        labels.extend(["项目名称", "招标编号", "招标人", "投标人", "日期"])
    if re.search(r"规格|货物|供货范围|供货清单|设备清单|机组配置", text):
        labels.extend(["投标机型", "单机容量", "总装机容量", "机组台数"])
    if re.search(r"技术偏差|参数偏差|条款偏差", text):
        labels.extend(["招标编号", "项目名称", "技术偏差说明"])
    if re.search(r"发电量|有效小时|利用率|功率曲线|性能保证", text):
        labels.extend(["保证发电量", "保证有效小时数", "功率曲线保证率", "全场可利用率"])
    return list(dict.fromkeys(labels))


def trusted_parse_fact_fields(parse_result: Any) -> list[dict[str, Any]]:
    fields: dict[str, dict[str, Any]] = {}

    def add(
        label: str,
        value: Any,
        *,
        category: str,
        source_field: dict[str, Any],
        confidence: float,
        required: bool = False,
        unit: str = "",
    ) -> None:
        label_text = canonical_fact_label(label)
        value_text = str(value or "").strip()
        if not label_text or not value_text:
            return
        key = fact_label_key(label_text)
        current = fields.get(key)
        fact = {
            "label": label_text,
            "value": value_text,
            "category": category,
            "confidence": confidence,
            "required": required,
            "unit": unit,
            "sourceRef": {
                "type": "parseField",
                "field": str(source_field.get("id") or source_field.get("fieldKey") or source_field.get("title") or ""),
                "fieldKey": str(source_field.get("fieldKey") or ""),
                "title": str(source_field.get("title") or source_field.get("label") or label_text),
                "sourceFile": str(source_field.get("sourceFile") or ""),
            },
        }
        if current is None or confidence > float(current.get("confidence") or 0):
            fields[key] = fact

    for field in iter_parse_fact_fields(parse_result):
        field_key = str(field.get("fieldKey") or "").strip()
        label = str(field.get("title") or field.get("label") or field.get("key") or field.get("id") or "").strip()
        value = str(field.get("value") or field.get("keyValue") or "").strip()
        evidence = str(field.get("evidence") or "").strip()
        text = "。".join(part for part in (label, value, evidence) if part)

        if field_key == "projectName" or fact_label_key(label) == fact_label_key("项目名称"):
            if looks_like_project_name(value):
                add("项目名称", value, category="项目基础信息", source_field=field, confidence=0.95, required=True)
        elif field_key == "tenderNo" or fact_label_key(label) == fact_label_key("招标编号"):
            if looks_like_tender_no(value):
                add("招标编号", value, category="项目基础信息", source_field=field, confidence=0.94, required=True)
        elif field_key in {"tenderer", "owner", "customerName"} or fact_label_key(label) in {
            fact_label_key("招标人"),
            fact_label_key("招标方"),
            fact_label_key("客户名称"),
        }:
            if looks_like_party_name(value):
                add("招标人", value, category="项目基础信息", source_field=field, confidence=0.84, required=False)
        elif field_key == "managementUnit" or fact_label_key(label) == fact_label_key("管理单位"):
            if looks_like_party_name(value):
                add("管理单位", value, category="项目基础信息", source_field=field, confidence=0.82, required=False)
        elif field_key == "bidSectionScale" or fact_label_key(label) in {
            fact_label_key("标段规模"),
            fact_label_key("招标规模"),
        }:
            add("标段规模", value, category="项目基础信息", source_field=field, confidence=0.78, required=False)
        elif field_key == "deliveryPeriod" or fact_label_key(label) == fact_label_key("交货周期"):
            add("交货周期", value, category="项目基础信息", source_field=field, confidence=0.78, required=False)
        elif field_key == "warrantyPeriod" or fact_label_key(label) == fact_label_key("质保期"):
            add("质保期", value, category="项目基础信息", source_field=field, confidence=0.78, required=False)
        elif field_key == "bidStartDate" or fact_label_key(label) == fact_label_key("投标起始日期"):
            add("投标起始日期", value, category="投标时间信息", source_field=field, confidence=0.78, required=False)
        elif field_key == "bidDeadline" or fact_label_key(label) in {
            fact_label_key("投标截止日期"),
            fact_label_key("投标截止时间"),
        }:
            add("投标截止日期", value, category="投标时间信息", source_field=field, confidence=0.82, required=True)

        add_performance_facts_from_parse_text(text, field, add)

    return list(fields.values())


def iter_parse_fact_fields(value: Any) -> list[dict[str, Any]]:
    fields: list[dict[str, Any]] = []

    def visit(node: Any) -> None:
        if len(fields) >= 200:
            return
        if isinstance(node, dict):
            has_label = any(key in node for key in ("label", "title", "key", "id"))
            has_value = any(key in node for key in ("value", "keyValue", "evidence"))
            if has_label and has_value:
                fields.append(node)
            for child in node.values():
                visit(child)
        elif isinstance(node, list):
            for child in node:
                visit(child)

    visit(value)
    return fields


def looks_like_project_name(value: Any) -> bool:
    text = str(value or "").strip()
    if not text or len(text) > 160:
        return False
    if re.search(r"投标人|招标人|应当|必须|不得|标准|规范|条款|认可|提供", text):
        return False
    return "项目" in text or "工程" in text


def looks_like_tender_no(value: Any) -> bool:
    text = str(value or "").strip()
    return bool(text and len(text) <= 80 and re.fullmatch(r"[A-Za-z0-9][A-Za-z0-9_\-./]+", text))


def looks_like_party_name(value: Any) -> bool:
    text = str(value or "").strip()
    if not text or len(text) > 80:
        return False
    if re.search(r"[。；;]|投标人|应|必须|不得|标准|规范|条款|认可|提供|要求|报告|测试|审查", text):
        return False
    return bool(re.search(r"公司|集团|有限|招标|业主|电力|能源|华能|国电|大唐|华电", text))


def add_performance_facts_from_parse_text(text: str, source_field: dict[str, Any], add: Any) -> None:
    normalized = re.sub(r"\s+", "", str(text or ""))
    if not normalized:
        return

    patterns = [
        (r"功率曲线[^。；;]{0,24}(?:不低于|≥|>=)(?:保证值的)?([0-9]+(?:\.[0-9]+)?%)", "功率曲线保证率"),
        (r"风电场机组年平均可利用率(?:≥|>=|不低于)([0-9]+(?:\.[0-9]+)?%)", "全场可利用率"),
        (r"(?:全部机组|全场).*?平均可利用率(?:≥|>=|不低于)([0-9]+(?:\.[0-9]+)?%)", "全场可利用率"),
        (r"单台机组年平均可利用率(?:≥|>=|不低于)([0-9]+(?:\.[0-9]+)?%)", "单台可利用率"),
        (r"主要部件更换率(?:低于|不高于|≤|<=)([0-9]+(?:\.[0-9]+)?%)", "主要部件更换率"),
    ]
    for pattern, label in patterns:
        match = re.search(pattern, normalized)
        if match:
            add(
                label,
                match.group(1),
                category="性能保证",
                source_field=source_field,
                confidence=0.86,
                required=False,
                unit="%",
            )


def fillable_table_labels_from_blank_source(blank: dict[str, Any]) -> list[str]:
    path = blank_source_docx_path(blank)
    if path is None:
        return []
    try:
        document = Document(str(path))
    except Exception:
        return []

    labels: list[str] = []
    seen: set[str] = set()
    for table in document.tables:
        for row in table.rows:
            cells = [clean_table_cell_text(cell.text) for cell in row.cells]
            label = table_field_label_from_row(cells)
            if not label:
                continue
            key = fact_label_key(label)
            if not key or key in seen:
                continue
            seen.add(key)
            labels.append(label)
    return labels


def blank_source_docx_path(blank: dict[str, Any]) -> Path | None:
    for key in ("docxPath", "path", "workspacePath"):
        value = str(blank.get(key) or "").strip()
        if not value:
            continue
        path = Path(value)
        if path.exists():
            return path
    return None


def clean_table_cell_text(value: Any) -> str:
    return re.sub(r"\s+", "", str(value or "")).strip()


def table_field_label_from_row(cells: list[str]) -> str:
    if not cells:
        return ""
    fill_positions = [
        index
        for index, text in enumerate(cells)
        if not text or re.search(r"待(?:人工)?(?:补充|填写|解析)|未填写|待确认", text)
    ]
    if not fill_positions:
        return ""
    candidates = cells[: fill_positions[0]]
    for candidate in reversed(candidates):
        label = canonical_fact_label(candidate)
        if looks_like_table_field_label(label):
            return label
    return ""


def looks_like_table_field_label(label: str) -> bool:
    text = str(label or "").strip()
    if not text or len(text) < 2 or len(text) > 80:
        return False
    if text in FACT_TABLE_HEADER_WORDS:
        return False
    if re.fullmatch(r"[\d一二三四五六七八九十]+[.、]?", text):
        return False
    if re.search(r"待(?:人工)?(?:补充|填写|解析)|未填写|授权人签名|日期", text):
        return False
    if re.search(r"同等质量|知名品牌|件套|厂家|品牌|Fluke|FLUKE|SKYLOTEC|DEHN|ABB|西门子|施耐德", text, flags=re.I):
        return False
    if re.search(r"参数|方法|折减|系数", text):
        return False
    if re.fullmatch(r"[A-Z]{1,8}[-A-Z0-9（）()\"'.—]+", text):
        return False
    if re.match(r"^\d", text) and not re.search(r"风速|年|容量|功率|高度|直径|小时|电量|温度", text):
        return False
    if text in COMMON_PROJECT_FACT_LABELS:
        return True
    return bool(
        re.search(
            r"投标机型|机组类型|机组台数|风机台数|单机容量|总装机容量|叶轮直径|风轮直径|轮毂.*高度|"
            r"扫风面积|比功率|安全等级|设计寿命|功率曲线|可利用率|保证电量|保证发电量|发电小时|"
            r"有效小时|等效利用小时|平均风速|空气密度|湍流|风切变|风剪切|极端风速|极大风速|"
            r"低温|高温|海拔|覆冰|盐雾|沙尘|雷电",
            text,
        )
    )


def project_material_fact_fields(
    project: dict[str, Any],
    gap_state: dict[str, Any],
    *,
    excluded_paths: set[str] | None = None,
    specs: list[dict[str, Any]] | None = None,
) -> list[dict[str, Any]]:
    materials = project_fact_material_index(project, gap_state)
    if not materials:
        return []
    prepared = prepare_project_fact_materials(project, materials)
    # 延迟 import 避免循环：专项模块复用本模块的 material_fact/clean_fact_text
    from app.services.technical_fact_special_extractors import (
        facts_from_certificate_materials,
        run_special_extractor,
        special_extractor_for_material,
    )

    facts: list[dict[str, Any]] = []
    cert_materials: list[tuple[dict[str, Any], Path]] = []
    for material in prepared:
        if not isinstance(material, dict):
            continue
        facts.extend(facts_from_material_name(material))
        path_text = str(material.get("path") or material.get("docx") or "").strip()
        if not path_text:
            continue
        path = Path(path_text)
        if not path.exists():
            continue
        if str(path.resolve()) in (excluded_paths or set()):
            continue
        kind = special_extractor_for_material(material)
        if kind == "certificate":
            # 证书按"型式认证 > 设计认证"成组处理
            cert_materials.append((material, path))
            continue
        if kind:
            special_facts = run_special_extractor(kind, path, material, project, specs=specs)
            if special_facts is not None:
                facts.extend(special_facts)
                continue
        suffix = path.suffix.lower()
        if suffix in {".docx", ".doc"}:
            facts.extend(facts_from_docx_material(path, material))
        elif suffix in {".xlsx", ".xlsm"}:
            facts.extend(facts_from_xlsx_material(path, material, project))
    facts.extend(facts_from_certificate_materials(cert_materials, project, specs=specs))
    facts.extend(derived_material_fact_fields(project, facts))
    return facts


def project_fact_material_index(project: dict[str, Any], gap_state: dict[str, Any]) -> list[dict[str, Any]]:
    plan = gap_state.get("plan") if isinstance(gap_state.get("plan"), dict) else {}
    materials: list[dict[str, Any]] = []
    seen: set[str] = set()

    def material_key(item: dict[str, Any]) -> str:
        material_id = str(item.get("id") or item.get("materialId") or "").strip()
        if material_id:
            return f"id:{material_id}"
        path = str(item.get("path") or item.get("folderPath") or "").strip()
        name = str(item.get("name") or item.get("materialName") or item.get("fileName") or "").strip()
        return f"path:{path}/{name}" if path or name else ""

    def append_material(item: dict[str, Any]) -> None:
        key = material_key(item)
        if not key or key in seen:
            return
        seen.add(key)
        materials.append(item)

    for item in (plan.get("materialIndex") if isinstance(plan.get("materialIndex"), list) else []):
        if isinstance(item, dict):
            append_material(dict(item))

    if not materials and isinstance(plan.get("tasks"), list):
        for task in plan.get("tasks") or []:
            if not isinstance(task, dict):
                continue
            for raw in [*(task.get("candidateMaterials") or []), *(task.get("selectedMaterialRefs") or [])]:
                if not isinstance(raw, dict):
                    continue
                material_id = str(raw.get("id") or raw.get("materialId") or "").strip()
                if not material_id:
                    continue
                append_material(
                    {
                        "id": material_id,
                        "name": str(raw.get("name") or raw.get("materialName") or raw.get("fileName") or material_id),
                        "folderPath": str(raw.get("folderPath") or raw.get("path") or ""),
                        "materialTier": str(raw.get("materialTier") or raw.get("libraryScope") or ""),
                        "cleanedFileName": str(raw.get("cleanedFileName") or ""),
                        "hasCleanedWord": bool(raw.get("hasCleanedWord")),
                        "turbineModelLabel": str(raw.get("turbineModelLabel") or ""),
                    }
                )
    try:
        selected_model = project_turbine_model(project)

        def collect_scope_files(folder_path: str, material_tier: str) -> None:
            payload = run_async_material_files(
                folder_path=folder_path,
                bid_type=TECHNICAL_BID_TYPE,
                material_tier=material_tier,
                turbine_model=selected_model,
                recursive=True,
                page=1,
                page_size=1000,
            )
            for raw in payload.get("items") or []:
                if not isinstance(raw, dict):
                    continue
                material_id = str(raw.get("id") or "")
                if not material_id:
                    continue
                append_material(
                    {
                        "id": material_id,
                        "name": str(raw.get("name") or ""),
                        "folderPath": str(raw.get("folderPath") or ""),
                        "materialTier": str(raw.get("materialTier") or material_tier),
                        "hasCleanedWord": bool(raw.get("hasCleanedWord")),
                        "cleanedFileName": str(raw.get("cleanedFileName") or ""),
                        "turbineModelLabel": str(raw.get("turbineModelLabel") or ""),
                        "size": int(raw.get("size") or 0),
                    }
                )

        # 无现成索引时，沿用项目默认目录扫描作为回退。
        if not materials:
            material_scope = build_project_material_scope(project)
            for scope in material_scope.get("readableScopes") or []:
                if not isinstance(scope, dict):
                    continue
                # 事实表匹配只扫本项目「项目定制」目录（recursive 覆盖下一级子目录，
                # 相关项目素材由用户归置到该目录下），不扫标准文件/客户定制目录
                if str(scope.get("materialTier") or "") != "project":
                    continue
                folder_path = str(scope.get("path") or "").strip()
                if folder_path:
                    collect_scope_files(folder_path, str(scope.get("materialTier") or ""))

        # 用户显式选择的目录始终叠加到计划索引，并按素材 ID 去重。
        custom_paths = gap_state.get("factMaterialPaths") if isinstance(gap_state.get("factMaterialPaths"), list) else []
        for raw_path in custom_paths:
            folder_path = str(raw_path or "").strip().strip("/")
            if folder_path and not folder_path.startswith(f"{TECHNICAL_BID_TYPE}/"):
                folder_path = f"{TECHNICAL_BID_TYPE}/{folder_path}"
            if folder_path:
                # 显式目录不按 tier 过滤，相关性由 material_is_fact_relevant 把守。
                collect_scope_files(folder_path, "")
    except Exception:
        logger.exception("项目事实素材索引查询失败，保留已收集素材继续构建")
    return [item for item in materials if material_is_fact_relevant(item)]


# 业主待填目标表格模板的文件名前缀（「待填写-附表X….docx」「待填写、待用印-….docx」等，
# 前缀取自业主下发的空白附表命名约定）：它们是要填的目标，不是取数素材，不进事实表素材体系
FILL_TEMPLATE_NAME_PREFIX = "待填写"


def material_is_fill_template(material: dict[str, Any]) -> bool:
    """按文件名前缀识别待填目标表格模板（不看 folderPath，避免按目录名猜内容）。"""
    name = str(material.get("name") or material.get("cleanedFileName") or "").strip()
    return name.startswith(FILL_TEMPLATE_NAME_PREFIX)


def material_is_fact_relevant(material: dict[str, Any]) -> bool:
    if material_is_fill_template(material):
        return False
    tier = str(material.get("materialTier") or "").strip()
    if tier == "project":
        return True
    text = " ".join(
        str(material.get(key) or "")
        for key in ("name", "cleanedFileName", "folderPath", "path")
    )
    return bool(
        re.search(
            r"参数|机型|功率曲线|风资源|发电量|报价|容量|安全|场址|载荷|工程量|技术承诺|投标关键数据|"
            r"弯矩|认证|承诺函|生产制造基地",
            text,
        )
    )


def prepare_project_fact_materials(project: dict[str, Any], materials: list[dict[str, Any]]) -> list[dict[str, Any]]:
    path_materials = [item for item in materials if item.get("path")]
    if path_materials and len(path_materials) == len(materials) and all(
        Path(str(item.get("path") or "")).exists() for item in path_materials
    ):
        return materials
    project_id = str(project.get("id") or "project")
    bid_type = TECHNICAL_BID_TYPE
    workspace_dir = "technical-workspace"
    work_dir = settings.documents_dir / project_id / workspace_dir / "gaps" / "fact_table_materials"
    work_dir.mkdir(parents=True, exist_ok=True)
    try:
        return prepare_project_fact_material_files(materials, work_dir, bid_type=bid_type, limit=120)
    except Exception:
        return materials


def facts_from_material_name(material: dict[str, Any]) -> list[dict[str, Any]]:
    text = str(material.get("name") or material.get("cleanedFileName") or "")
    facts: list[dict[str, Any]] = []
    for pattern, label, unit in [
        (r"空气密度\s*([0-9]+(?:\.[0-9]+)?)", "空气密度", "kg/m3"),
        (r"湍流强度\s*([0-9]+(?:\.[0-9]+)?)", "湍流强度", ""),
        (r"风(?:剪切|切变)(?:指数)?\s*([0-9]+(?:\.[0-9]+)?)", "风剪切", ""),
    ]:
        match = re.search(pattern, text, flags=re.I)
        if match:
            facts.append(material_fact(label, match.group(1), material, unit=unit, confidence=0.9))
    return facts


def facts_from_docx_material(path: Path, material: dict[str, Any]) -> list[dict[str, Any]]:
    facts: list[dict[str, Any]] = []
    try:
        document = Document(str(path))
    except Exception:
        return facts
    text_parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for para_idx, paragraph in enumerate(document.paragraphs, start=1):
        text = clean_fact_text(paragraph.text)
        if not text or len(text) > 180:
            continue
        match = re.match(r"^([^:：]{2,40})[:：]\s*(.{1,100})$", text)
        if match:
            fact = material_fact_from_label_value(match.group(1), match.group(2), material, location=f"P{para_idx}", confidence=0.78)
            if fact:
                facts.append(fact)
    for table_idx, table in enumerate(document.tables, start=1):
        facts.extend(facts_from_guarantee_table(table, material, table_idx=table_idx))
        for row_idx, row in enumerate(table.rows, start=1):
            cells = [clean_fact_text(cell.text) for cell in row.cells]
            text_parts.append(" | ".join(cell for cell in cells if cell))
            facts.extend(facts_from_table_cells(cells, material, location=f"T{table_idx}/R{row_idx}"))
    facts.extend(facts_from_free_text("\n".join(text_parts), material))
    return facts


def facts_from_guarantee_table(table: Any, material: dict[str, Any], *, table_idx: int) -> list[dict[str, Any]]:
    if not getattr(table, "rows", None) or len(table.rows) < 2:
        return []
    header = " ".join(clean_fact_text(cell.text) for cell in table.rows[0].cells)
    if not ("年平均风速" in header and "保证年上网电量" in header and "满负荷小时" in header):
        return []
    facts: list[dict[str, Any]] = []
    for row_idx, row in enumerate(table.rows[1:], start=2):
        cells = [clean_fact_text(cell.text) for cell in row.cells]
        if len(cells) < 3:
            continue
        wind_speed = clean_fact_value("年平均风速", cells[0])
        energy = clean_fact_value("保证发电量", cells[1])
        hours = clean_fact_value("保证有效小时数", cells[2])
        if not (wind_speed and energy and hours):
            continue
        matrix_fact = material_fact(
            "__guaranteeMatrixRow",
            {"windSpeed": wind_speed, "energyMwh": energy, "hours": hours},
            material,
            location=f"T{table_idx}/R{row_idx}",
            confidence=0.82,
        )
        matrix_fact["internal"] = True
        facts.append(matrix_fact)
    return facts


def facts_from_xlsx_material(path: Path, material: dict[str, Any], project: dict[str, Any]) -> list[dict[str, Any]]:
    facts: list[dict[str, Any]] = []
    try:
        workbook = load_workbook(path, data_only=True, read_only=True)
    except Exception:
        return facts
    project_model = str((project_turbine_model(project) or {}).get("model") or "")
    model_key = re.sub(r"(上置|下置|内置|外置|塔上|塔下)", "", project_model)
    for worksheet in workbook.worksheets:
        selected_col = xlsx_model_column(worksheet, model_key)
        for row_idx, row in enumerate(worksheet.iter_rows(values_only=True), start=1):
            cells = [clean_fact_text(cell) for cell in row]
            if selected_col is not None and selected_col < len(cells):
                for label_idx in (2, 1, 0):
                    if label_idx < len(cells):
                        fact = material_fact_from_label_value(
                            cells[label_idx],
                            cells[selected_col],
                            material,
                            unit=cells[3] if len(cells) > 3 else "",
                            location=f"{worksheet.title}!R{row_idx}",
                            confidence=0.82,
                        )
                        if fact:
                            facts.append(fact)
                            break
            facts.extend(facts_from_table_cells(cells, material, location=f"{worksheet.title}!R{row_idx}"))
            if len(facts) >= 800:
                return facts
    return facts


def xlsx_model_column(worksheet: Any, model_key: str) -> int | None:
    if not model_key:
        return None
    normalized_model = re.sub(r"\s+", "", model_key)
    for row in worksheet.iter_rows(min_row=1, max_row=min(12, worksheet.max_row), values_only=True):
        for index, value in enumerate(row):
            text = re.sub(r"\s+", "", str(value or ""))
            if normalized_model and normalized_model in text:
                return index
    return None


def facts_from_table_cells(cells: list[str], material: dict[str, Any], *, location: str) -> list[dict[str, Any]]:
    facts: list[dict[str, Any]] = []
    nonempty = [(index, value) for index, value in enumerate(cells) if value]
    if len(nonempty) < 2:
        return facts
    if len(cells) >= 4:
        fact = material_fact_from_label_value(cells[1], cells[3], material, unit=cells[2], location=location, confidence=0.88)
        if fact:
            facts.append(fact)
    for first, second in zip(nonempty, nonempty[1:]):
        fact = material_fact_from_label_value(first[1], second[1], material, location=location, confidence=0.76)
        if fact:
            facts.append(fact)
            break
    if len(cells) >= 3 and cells[0]:
        wind_fact = material_fact_from_label_value(cells[0], cells[2], material, unit=cells[1], location=location, confidence=0.84)
        if wind_fact:
            facts.append(wind_fact)
    return facts


def material_fact_from_label_value(
    label: Any,
    value: Any,
    material: dict[str, Any],
    *,
    unit: str = "",
    location: str = "",
    confidence: float = 0.78,
) -> dict[str, Any] | None:
    label_text = canonical_fact_label(label)
    value_text = clean_fact_value(label_text, value)
    if not label_text or not value_text:
        return None
    if label_text not in COMMON_PROJECT_FACT_LABELS and not looks_like_table_field_label(label_text):
        return None
    unit_text = clean_fact_unit(unit)
    raw_label = str(label or "")
    raw_value = str(value or "")
    if not unit_text:
        raw_context = f"{raw_label}{raw_value}"
        if label_text in {"轮毂高度", "叶轮直径"} and re.search(r"(?:m|米)", raw_context, flags=re.I):
            unit_text = "m"
        elif label_text in {"极端风速", "年平均风速"} and re.search(r"m/?s|米/秒", raw_context, flags=re.I):
            unit_text = "m/s"
        elif label_text == "空气密度" and re.search(r"kg/?m|kg/m3|kg/m³", raw_context, flags=re.I):
            unit_text = "kg/m3"
        elif label_text == "机组台数" and "台" in raw_value:
            unit_text = "台"
    if not unit_text:
        if label_text in {"轮毂高度", "叶轮直径"}:
            unit_text = "m"
        elif label_text in {"极端风速", "年平均风速"}:
            unit_text = "m/s"
        elif label_text == "空气密度":
            unit_text = "kg/m3"
        elif label_text == "机组台数":
            unit_text = "台"
    return material_fact(label_text, value_text, material, unit=unit_text, location=location, confidence=confidence)


def facts_from_free_text(text: str, material: dict[str, Any]) -> list[dict[str, Any]]:
    compact = clean_fact_text(text)
    facts: list[dict[str, Any]] = []
    patterns = [
        (r"(?:总装机容量|建设容量|标段规模|总容量)[^0-9]{0,12}([0-9]+(?:\.[0-9]+)?\s*(?:MW|万千瓦|kW)?)", "总装机容量", ""),
        (r"(?:机组台数|机组数量|风机台数|风机数量|安装)[^0-9]{0,12}([0-9]+)\s*台", "机组台数", "台"),
        (r"轮毂(?:中心)?高度[^0-9]{0,12}([0-9]+(?:\.[0-9]+)?\s*m?)", "轮毂高度", "m"),
        (r"(?:安全等级|适用等级|设计等级)[^A-Za-z0-9]{0,12}((?:IEC\s*)?[A-Z0-9][A-Z0-9/ .-]{0,20})", "安全等级", ""),
        (r"空气密度[^0-9]{0,12}([0-9]+(?:\.[0-9]+)?)", "空气密度", "kg/m3"),
        (r"湍流强度[^0-9]{0,12}([0-9]+(?:\.[0-9]+)?)", "湍流强度", ""),
        (r"(?:极端风速|极大风速|Ve50)[^0-9]{0,12}([0-9]+(?:\.[0-9]+)?\s*m/s?)", "极端风速", "m/s"),
        (r"年平均风速[^0-9]{0,12}([0-9]+(?:\.[0-9]+)?\s*m/s?)", "年平均风速", "m/s"),
    ]
    for pattern, label, unit in patterns:
        match = re.search(pattern, compact, flags=re.I)
        if match:
            value = clean_fact_value(label, match.group(1))
            if value:
                facts.append(material_fact(label, value, material, unit=unit, confidence=0.78))
    return facts


def clean_fact_text(value: Any) -> str:
    return re.sub(r"\s+", "", str(value or "")).strip()


def clean_fact_unit(value: Any) -> str:
    text = re.sub(r"\s+", "", str(value or "")).strip()
    text = text.strip("：:；;，,、")
    if text in {"", "-", "/", "—", "NA", "N/A", "字段", "值", "年份", "参数内容", "结果", "说明", "备注", "机型", "型号"}:
        return ""
    text = text.replace("m³", "m3")
    text = re.sub(r"kg/?m3", "kg/m3", text, flags=re.I)
    text = re.sub(r"m/?s", "m/s", text, flags=re.I)
    return text


def clean_fact_value(label: str, value: Any) -> str:
    text = str(value or "").strip()
    text = re.sub(r"\s+", "", text)
    text = text.strip("：:；;，,、")
    if not text or len(text) > 120:
        return ""
    if any(token in text for token in ("待填写", "待人工", "未填写")):
        return ""
    if text in {"-", "/", "—", "无", "暂无", "值", "结果", "参数内容", "单位", "年份"}:
        return ""
    numeric_ranges = {
        "机组台数": (1, 1000),
        "轮毂高度": (40, 250),
        "叶轮直径": (50, 350),
        "空气密度": (0.7, 1.5),
        "湍流强度": (0, 1),
        "风剪切": (0, 1),
        "极端风速": (20, 100),
        "年平均风速": (2, 15),
    }
    numeric_noise = (
        "年份",
        "各年",
        "版本",
        "编制",
        "校核",
        "审核",
        "批准",
        "日期",
        "参数内容",
        "结果结果",
        "场址空气密度下",
    )
    if label in numeric_ranges and (
        any(token in text for token in numeric_noise)
        or re.search(r"\d{4}[-/年]\d{1,2}", text)
    ):
        return ""
    if label in {"总装机容量", "单机容量"}:
        match = re.search(r"([0-9]+(?:\.[0-9]+)?)(万千瓦|MW|kW)?", text, flags=re.I)
        if match:
            return f"{match.group(1)}{match.group(2) or ''}".strip()
    if label in {"机组台数"}:
        match = re.search(r"([0-9]+)", text)
        if not match:
            return ""
        number = float(match.group(1))
        low, high = numeric_ranges[label]
        return match.group(1) if low <= number <= high else ""
    if label in {"保证发电量", "保证有效小时数"}:
        if re.search(r"风电场|保证年上网电量|满负荷小时|字段|单位", text):
            return ""
        match = re.search(r"([0-9]+(?:\.[0-9]+)?)", text)
        if not match:
            return ""
        number = float(match.group(1))
        if label == "保证发电量" and not (1 <= number <= 10000000):
            return ""
        if label == "保证有效小时数" and not (1 <= number <= 8760):
            return ""
        return match.group(1)
    if label in {"极端风速", "年平均风速"}:
        match = re.search(r"([0-9]+(?:\.[0-9]+)?)", text)
        if match:
            number = float(match.group(1))
            low, high = numeric_ranges[label]
            if not (low <= number <= high):
                return ""
            return f"{match.group(1)}m/s" if re.search(r"m/?s|米/秒", text, flags=re.I) else match.group(1)
    if label in {"轮毂高度", "叶轮直径", "空气密度", "湍流强度", "风剪切"}:
        match = re.search(r"([0-9]+(?:\.[0-9]+)?)", text)
        if match:
            number = float(match.group(1))
            low, high = numeric_ranges[label]
            if not (low <= number <= high):
                return ""
            return match.group(1)
    if label in {"投标方案", "投标机型", "机组类型"}:
        if text in {"机型", "投标机型", "方案", "投标方案"}:
            return ""
        model_match = re.search(r"([A-Z]{1,6}\d+(?:\.\d+)?[-—]\d+(?:[-—]\d+)?)", text, flags=re.I)
        if model_match:
            return model_match.group(1).replace("—", "-")
    if label == "安全等级":
        text = re.sub(r"^IEC\s*", "IEC ", text, flags=re.I).strip()
    return text


def material_fact(
    label: str,
    value: Any,
    material: dict[str, Any],
    *,
    unit: str = "",
    location: str = "",
    confidence: float = 0.78,
) -> dict[str, Any]:
    tier = str(material.get("materialTier") or "").strip() or "standard"
    return {
        "label": canonical_fact_label(label),
        "value": value,
        "category": "素材库事实",
        "unit": clean_fact_unit(unit),
        "confidence": confidence,
        "sourcePriority": FACT_MATERIAL_SOURCE_PRIORITIES.get(tier, 50),
        "sourceRef": {
            "type": "materialFact",
            "materialId": str(material.get("id") or material.get("materialId") or ""),
            "materialTier": tier,
            "name": str(material.get("name") or material.get("fileName") or material.get("cleanedFileName") or ""),
            "folderPath": str(material.get("folderPath") or ""),
            "path": str(material.get("path") or ""),
            "location": location,
        },
    }


def derived_material_fact_fields(project: dict[str, Any], facts: list[dict[str, Any]]) -> list[dict[str, Any]]:
    by_label: dict[str, dict[str, Any]] = {}
    guarantee_matrix_rows: list[dict[str, Any]] = []
    for fact in facts:
        if fact.get("label") == "__guaranteeMatrixRow":
            guarantee_matrix_rows.append(fact)
            continue
        label = canonical_fact_label(fact.get("label"))
        if not label or not fact.get("value"):
            continue
        current = by_label.get(label)
        rank = (int(fact.get("sourcePriority") or 0), float(fact.get("confidence") or 0))
        if current is None or rank > (int(current.get("sourcePriority") or 0), float(current.get("confidence") or 0)):
            by_label[label] = fact
    rated_kw = project_turbine_model(project).get("ratedPowerKw")
    rated_mw = float(rated_kw) / 1000 if isinstance(rated_kw, (int, float)) and rated_kw else 0
    result: list[dict[str, Any]] = []
    total = number_from_fact(by_label.get("总装机容量"))
    count = number_from_fact(by_label.get("机组台数"))
    source = by_label.get("总装机容量") or by_label.get("机组台数") or {}
    material_ref = copy.deepcopy(source.get("sourceRef") if isinstance(source.get("sourceRef"), dict) else {})
    if total and rated_mw and not count:
        derived_count = total / rated_mw
        rounded = round(derived_count)
        if abs(derived_count - rounded) < 0.01:
            result.append(
                {
                    "label": "机组台数",
                    "value": str(rounded),
                    "category": "素材库事实",
                    "unit": "台",
                    "confidence": 0.86,
                    "sourcePriority": int(source.get("sourcePriority") or 0),
                    "sourceRef": {**material_ref, "type": "derivedMaterialFact", "field": "总装机容量/单机容量"},
                }
            )
    if count and rated_mw and not total:
        result.append(
            {
                "label": "总装机容量",
                "value": f"{count * rated_mw:g}MW",
                "category": "素材库事实",
                "unit": "MW",
                "confidence": 0.82,
                "sourcePriority": int(source.get("sourcePriority") or 0),
                "sourceRef": {**material_ref, "type": "derivedMaterialFact", "field": "机组台数/单机容量"},
            }
        )
    year_avg = number_from_fact(by_label.get("年平均风速"))
    if year_avg and guarantee_matrix_rows:
        def matrix_distance(row: dict[str, Any]) -> float:
            value = row.get("value") if isinstance(row.get("value"), dict) else {}
            wind = number_from_fact({"value": value.get("windSpeed")})
            return abs(float(wind or 0) - year_avg) if wind else 9999

        selected = min(guarantee_matrix_rows, key=matrix_distance)
        if matrix_distance(selected) <= 0.08:
            selected_value = selected.get("value") if isinstance(selected.get("value"), dict) else {}
            selected_ref = copy.deepcopy(selected.get("sourceRef") if isinstance(selected.get("sourceRef"), dict) else {})
            wind_speed = str(selected_value.get("windSpeed") or "")
            if selected_value.get("energyMwh"):
                result.append(
                    {
                        "label": "保证发电量",
                        "value": str(selected_value.get("energyMwh")),
                        "category": "性能保证",
                        "unit": "MWh",
                        "confidence": 0.86,
                        "sourcePriority": int(selected.get("sourcePriority") or 0),
                        "sourceRef": {**selected_ref, "type": "derivedMaterialFact", "field": f"发电量保证矩阵/{wind_speed}m/s"},
                    }
                )
            if selected_value.get("hours"):
                result.append(
                    {
                        "label": "保证有效小时数",
                        "value": str(selected_value.get("hours")),
                        "category": "性能保证",
                        "unit": "h",
                        "confidence": 0.86,
                        "sourcePriority": int(selected.get("sourcePriority") or 0),
                        "sourceRef": {**selected_ref, "type": "derivedMaterialFact", "field": f"发电量保证矩阵/{wind_speed}m/s"},
                    }
                )
    return result


def number_from_fact(fact: dict[str, Any] | None) -> float | None:
    if not isinstance(fact, dict):
        return None
    match = re.search(r"([0-9]+(?:\.[0-9]+)?)", str(fact.get("value") or ""))
    return float(match.group(1)) if match else None


def run_async_material_files(**kwargs: Any) -> dict[str, Any]:
    kwargs.pop("bid_type", None)
    result = run_awaitable_sync(technical_material_store.raw_files(**kwargs))
    return result if isinstance(result, dict) else {}


def _now_iso() -> str:
    return datetime.now(UTC).isoformat()
