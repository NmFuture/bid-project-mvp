"""进程内执行技术标正文组装。"""

from __future__ import annotations

import json
import re
import shutil
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document

from .build_assembly import apply_gap_plan, build_plan, rearrange_appendices
from .create_tech_master import apply_page_setup, apply_style_overrides, prune_unreferenced_media, strip_body
from .finalize import force_update_fields, insert_toc_field, reapply_heading_fonts, replace_header_text
from .merger import merge
from .numbering_fixer import (
    enforce_no_auto_numbering_on_numbered_headings,
    strip_numPr_from_body,
    strip_numPr_from_heading_styles,
)
from .parse_toc import parse_toc_docx, parse_toc_json
from .verify import scan_docx


RESOURCES_DIR = Path(__file__).resolve().parents[1] / "resources"
STYLE_SPEC_PATH = RESOURCES_DIR / "heading_style.json"


def _path(value: Any, *, required: bool = False, label: str = "path") -> Path | None:
    text = str(value or "").strip()
    if not text:
        if required:
            raise RuntimeError(f"{label} is required")
        return None
    path = Path(text)
    if required and not path.exists():
        raise RuntimeError(f"{label} does not exist: {path}")
    return path


def _safe_filename(value: str, fallback: str) -> str:
    text = re.sub(r'[\\/:*?"<>|]+', "-", str(value or "").strip())
    return re.sub(r"\s+", " ", text).strip(" .") or fallback


def _create_master(manifest: dict[str, Any], work_dir: Path) -> Path:
    target = work_dir / "templates" / "技术投标母版模板.docx"
    sample = _path(manifest.get("templateFile"))
    target.parent.mkdir(parents=True, exist_ok=True)
    if not sample or not sample.exists():
        Document().save(str(target))
        return target
    try:
        shutil.copy2(sample, target)
        style_cfg = json.loads(STYLE_SPEC_PATH.read_text(encoding="utf-8"))
        doc = Document(str(target))
        strip_body(doc)
        apply_page_setup(doc, style_cfg["page"])
        apply_style_overrides(doc, style_cfg)
        doc.save(str(target))
        prune_unreferenced_media(target)
    except Exception:
        Document().save(str(target))
    return target


def finalize_merged_output(
    input_file: str | Path,
    output_file: str | Path,
    project_params: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """为 formatter 失败场景生成可交付的降级稿。"""
    source = Path(input_file)
    target = Path(output_file)
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, target)
    style_cfg = json.loads(STYLE_SPEC_PATH.read_text(encoding="utf-8"))
    doc = Document(str(target))
    insert_toc_field(doc)
    reapply_heading_fonts(doc, style_cfg)
    strip_numPr_from_heading_styles(doc)
    strip_numPr_from_body(doc)
    enforce_no_auto_numbering_on_numbered_headings(doc)
    doc.save(str(target))
    project_name = str((project_params or {}).get("project_name") or "")
    if project_name:
        replace_header_text(target, project_name)
    force_update_fields(target)
    return scan_docx(target)


def _summary(plan: list[dict[str, Any]], merge_result: dict[str, Any], scan: dict[str, Any]) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    counts = Counter(str(item.get("status") or "") for item in plan)
    warnings = [dict(item) for item in merge_result.get("warnings") or [] if isinstance(item, dict)]
    for code, key, message in (
        ("PLACEHOLDER_REMAINS", "placeholders", "交付稿中仍有占位符"),
        ("EMPTY_SECTION", "empty_leaf_headings", "交付稿中存在空章节"),
        ("DUPLICATE_HEADING", "dup_alerts", "交付稿中存在重复标题"),
    ):
        count = len(scan.get(key) or [])
        if count:
            warnings.append({"code": code, "message": f"{message}：{count} 处", "count": count})
    summary = {
        "total": len(plan),
        "byStatus": dict(counts),
        "usedPathCount": len({str(path) for item in plan for path in item.get("paths") or []}),
        "assembledCount": int(merge_result.get("merged_materials") or 0),
        "warningCount": sum(int(item.get("count") or 0) for item in warnings),
    }
    return summary, warnings


def run_from_manifest(manifest_path: str | Path) -> dict[str, Any]:
    path = Path(manifest_path)
    manifest = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(manifest, dict):
        raise RuntimeError("manifest must be a JSON object")
    work_dir = _path(manifest.get("workDir")) or path.parent
    work_dir.mkdir(parents=True, exist_ok=True)
    toc_source = _path(manifest.get("tocJsonPath") or manifest.get("tocPath"), required=True, label="tocJsonPath")
    gap_plan = _path(manifest.get("gapPlanPath"), required=True, label="gapPlanPath")
    material_library = _path(manifest.get("materialLibraryDir"), required=True, label="materialLibraryDir")
    assert toc_source and gap_plan and material_library

    toc = parse_toc_json(toc_source) if toc_source.suffix.lower() == ".json" else parse_toc_docx(toc_source)
    toc_file = work_dir / "toc_entries.json"
    toc_file.write_text(json.dumps(toc, ensure_ascii=False, indent=2), encoding="utf-8")
    params = manifest.get("projectParams") if isinstance(manifest.get("projectParams"), dict) else {}
    params_file = _path(manifest.get("projectParamsPath")) or work_dir / "project_params.json"
    params_file.write_text(json.dumps(params, ensure_ascii=False, indent=2), encoding="utf-8")

    plan = rearrange_appendices(apply_gap_plan(build_plan(toc, [], params), gap_plan))
    plan_file = work_dir / "assembly_plan.json"
    plan_file.write_text(json.dumps(plan, ensure_ascii=False, indent=2), encoding="utf-8")
    merged_file = work_dir / "bid_merged.docx"
    merge_result = merge(_create_master(manifest, work_dir), plan, material_library, params, work_dir / "bid_prep", merged_file)

    requested_output = _path(manifest.get("outputFile"))
    finalize_output = manifest.get("finalizeOutput", True) is not False
    output_file = requested_output or merged_file
    if finalize_output:
        finalize_merged_output(merged_file, output_file, params)
    else:
        output_file = merged_file
    scan = scan_docx(output_file)
    summary, warnings = _summary(plan, merge_result, scan)
    return {
        "schema_version": "bid-tech-assembly-v1",
        "workDir": str(work_dir),
        "tocJson": str(toc_file),
        "planFile": str(plan_file),
        "projectParamsFile": str(params_file),
        "gapPlanFile": str(gap_plan),
        "outputFile": str(output_file),
        "summary": summary,
        "warnings": warnings,
    }
