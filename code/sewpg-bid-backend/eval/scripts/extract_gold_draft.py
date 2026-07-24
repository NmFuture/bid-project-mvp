from __future__ import annotations

import argparse
import csv
import re
from pathlib import Path

from docx import Document

from gold_schema import GOLD_COLUMNS


PLACEHOLDER_RE = re.compile(r"\[(?P<label>[^\[\]]{1,80}?)(?:[，,：:]\s*)?(?:待填写|待补充|待人工补充)[^\[\]]*?\]")


def extract_body_placeholders(docx_path: Path, project_id: str) -> list[dict[str, str]]:
    document = Document(docx_path)
    rows: list[dict[str, str]] = []
    heading_stack: list[str] = []
    seen: set[tuple[str, str]] = set()

    for para_index, paragraph in enumerate(document.paragraphs, start=1):
        text = _clean_text(paragraph.text)
        if not text:
            continue
        style_name = str(paragraph.style.name or "")
        heading_level = _heading_level(style_name)
        if heading_level:
            heading_stack = heading_stack[: heading_level - 1]
            heading_stack.append(text)
        for match in PLACEHOLDER_RE.finditer(text):
            label = _placeholder_label(match.group("label"))
            locator = "/".join(heading_stack) if heading_stack else f"P{para_index}"
            key = (locator, label)
            if key in seen:
                continue
            seen.add(key)
            rows.append(_gold_row(project_id, "body", locator, label, "paragraph", "T5"))

    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                text = _clean_text(cell.text)
                for match in PLACEHOLDER_RE.finditer(text):
                    label = _placeholder_label(match.group("label"))
                    locator = f"表{table_index}/行{row_index}/列{cell_index}"
                    key = (locator, label)
                    if key in seen:
                        continue
                    seen.add(key)
                    rows.append(_gold_row(project_id, "body", locator, label, "phrase", "T5"))

    return rows


def extract_appendix_table_fields(docx_path: Path, project_id: str) -> list[dict[str, str]]:
    document = Document(docx_path)
    rows: list[dict[str, str]] = []
    seen: set[tuple[str, str]] = set()
    current_title = docx_path.stem

    paragraphs = list(document.paragraphs)
    if paragraphs:
        for paragraph in paragraphs:
            text = _clean_text(paragraph.text)
            if text:
                current_title = text[:80]
                break

    for table_index, table in enumerate(document.tables, start=1):
        if not table.rows:
            continue
        for row_index, row in enumerate(table.rows, start=1):
            cells = [_clean_text(cell.text) for cell in row.cells]
            if not any(cells):
                continue
            field_name = _field_name_from_cells(cells)
            if not field_name:
                continue
            locator = f"{current_title}/表{table_index}/行{row_index}"
            key = (locator, field_name)
            if key in seen:
                continue
            seen.add(key)
            rows.append(_gold_row(project_id, "appendix", locator, field_name, "phrase", "T5"))

    return rows


def write_gold_csv(rows: list[dict[str, str]], output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with output_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=GOLD_COLUMNS)
        writer.writeheader()
        writer.writerows(rows)


def _gold_row(
    project_id: str,
    doc_type: str,
    locator: str,
    field_name: str,
    field_type: str,
    difficulty_tier: str,
) -> dict[str, str]:
    return {
        "project_id": project_id,
        "doc_type": doc_type,
        "locator": locator,
        "field_name": field_name,
        "human_answer": "",
        "field_type": field_type,
        "difficulty_tier": difficulty_tier,
        "evidence_source": "",
    }


def _clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _heading_level(style_name: str) -> int:
    match = re.search(r"(?:Heading|标题)\s*([1-6])", style_name, flags=re.I)
    return int(match.group(1)) if match else 0


def _placeholder_label(value: str) -> str:
    text = _clean_text(value)
    text = re.sub(r"[，,：:]\s*$", "", text).strip()
    return text or "待填写字段"


def _field_name_from_cells(cells: list[str]) -> str:
    skip_tokens = ("序号", "单位", "备注", "填写", "待填写", "投标人填写", "响应")
    candidates = [cell for cell in cells[:3] if cell and not any(token == cell for token in skip_tokens)]
    if not candidates:
        candidates = [cell for cell in cells if cell and not any(token in cell for token in ("待填写", "备注"))]
    if not candidates:
        return ""
    value = candidates[0]
    return value[:80]


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract a human-review draft gold CSV from a docx file.")
    parser.add_argument("--project-id", required=True)
    parser.add_argument("--doc-type", choices=["appendix", "body"], required=True)
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()

    if args.doc_type == "body":
        rows = extract_body_placeholders(args.input, args.project_id)
    else:
        rows = extract_appendix_table_fields(args.input, args.project_id)
    write_gold_csv(rows, args.output)
    print(f"wrote {len(rows)} draft gold rows to {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
