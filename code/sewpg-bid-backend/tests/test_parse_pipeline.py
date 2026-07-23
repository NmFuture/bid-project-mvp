from __future__ import annotations

import io
import json
import sqlite3
import subprocess
import sys
import tempfile
import time
import unittest
import zipfile
from pathlib import Path
from unittest.mock import AsyncMock, patch

from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from app.core.config import settings
from app.services import parsing as parsing_service
from app.services.file_utils import format_size_mb
from app.services.bid_parse_state import complete_parse_state
from app.services.bid_project_state import project_parse_input_records
from app.services.store import store


def build_docx_bytes(*lines: str) -> bytes:
    file_obj = io.BytesIO()
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    doc.save(file_obj)
    return file_obj.getvalue()


def parse_inputs_for_tests(project_id: str):
    project = store.get_project_runtime_state(project_id)
    return project_parse_input_records(project_id, project)


def complete_parse_for_tests(
    project_id: str,
    tender_files: list[dict],
    template_files: list[dict],
    *,
    summary: dict | None = None,
    parse_storage: dict | None = None,
) -> dict:
    project = store.require_project_for_update(project_id)
    payload = complete_parse_state(
        project,
        tender_files,
        template_files,
        summary=summary,
        parse_storage=parse_storage,
    )
    store.persist_project_state(project)
    return payload


def build_docx_blocks_bytes(*blocks: str | list[list[str]]) -> bytes:
    file_obj = io.BytesIO()
    doc = Document()
    for block in blocks:
        if isinstance(block, str):
            doc.add_paragraph(block)
            continue
        if not block:
            continue
        column_count = max(len(row) for row in block)
        table = doc.add_table(rows=len(block), cols=column_count)
        for row_index, row in enumerate(block):
            for col_index in range(column_count):
                table.cell(row_index, col_index).text = row[col_index] if col_index < len(row) else ""
    doc.save(file_obj)
    return file_obj.getvalue()


def build_appendix_docx_bytes() -> bytes:
    file_obj = io.BytesIO()
    doc = Document()
    doc.add_paragraph("附表1：供货范围空表")
    table = doc.add_table(rows=3, cols=3)
    values = [
        ["序号", "设备名称", "投标响应"],
        ["1", "风力发电机组", ""],
        ["2", "塔筒", ""],
    ]
    for row_index, row in enumerate(values):
        for col_index, value in enumerate(row):
            table.cell(row_index, col_index).text = value
    doc.save(file_obj)
    return file_obj.getvalue()


def build_appendix_with_merges_docx_bytes() -> bytes:
    """Build an RFP-like docx whose appendix table uses both horizontal and vertical
    cell merges, so we can assert that the parser preserves these structures end-to-end."""

    file_obj = io.BytesIO()
    doc = Document()
    doc.add_paragraph("附表D.1 标准及风电场空气密度功率曲线")
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    # Row 0 header: merge all 4 cells into one banner
    table.cell(0, 0).text = "机型：投标机型1"
    table.cell(0, 0).merge(table.cell(0, 3))
    # Row 1 sub-header: merge last two columns into "对比图"
    headers = ["风速区间(m/s)", "区间平均风速(m/s)", "标准空气密度下功率(kW)", "对比图"]
    for col_index, value in enumerate(headers):
        table.cell(1, col_index).text = value
    table.cell(1, 3).merge(table.cell(1, 3))
    # Row 2: data row, with cols 2-3 vertically merged into row 3 (vMerge)
    table.cell(2, 0).text = "0.00-0.50"
    table.cell(2, 1).text = "0"
    table.cell(2, 2).text = "/"
    table.cell(2, 3).text = "/"
    # Row 3: another data row; col 3 merged into row 2's col 3 (vertical)
    table.cell(3, 0).text = "0.50-1.00"
    table.cell(3, 1).text = "0.5"
    table.cell(3, 2).text = "/"
    table.cell(3, 3).text = "/"
    table.cell(2, 3).merge(table.cell(3, 3))
    doc.save(file_obj)
    return file_obj.getvalue()


def build_business_attachment_templates_docx_bytes() -> bytes:
    file_obj = io.BytesIO()
    doc = Document()
    doc.add_paragraph("第一章 招标公告")
    doc.add_paragraph("此处为普通正文，不应被提取为商务附件模板。")
    doc.add_paragraph("第六章 投标文件格式")
    doc.add_paragraph("附件1 投标函")
    doc.add_paragraph("致：华能集团")
    doc.add_paragraph("我方已仔细研究招标文件的全部内容，愿意参加本项目投标。")
    doc.add_paragraph("投标人（盖章）：____________")
    doc.add_paragraph("附件2 开标价格表")
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    table.cell(0, 0).text = "开标价格表"
    table.cell(0, 0).merge(table.cell(0, 2))
    values = [
        ["序号", "项目名称", "投标报价"],
        ["1", "", ""],
    ]
    for row_index, row in enumerate(values, start=1):
        for col_index, value in enumerate(row):
            table.cell(row_index, col_index).text = value
    doc.add_paragraph("附件3 法定代表人授权书")
    doc.add_paragraph("本人授权以下代表作为我方合法代理人参加本项目投标。")
    doc.add_paragraph("授权代表签字：____________")
    doc.save(file_obj)
    return file_obj.getvalue()


def build_business_section_tree_docx_bytes() -> bytes:
    file_obj = io.BytesIO()
    doc = Document()
    doc.add_paragraph("目录")
    doc.add_paragraph("第一章 招标公告 1")
    doc.add_paragraph("第二章 供应商须知 8")
    doc.add_paragraph("第三章 评审办法 30")
    doc.add_heading("第一章 招标公告", level=1)
    doc.add_heading("3. 供应商资格要求", level=2)
    doc.add_paragraph("3.1 供应商须为中华人民共和国境内合法注册的独立法人。")
    doc.add_heading("第二章 供应商须知", level=1)
    doc.add_heading("供应商须知前附表", level=2)
    table = doc.add_table(rows=2, cols=3)
    for col, text in enumerate(["条款号", "条款名称", "编列内容"]):
        table.cell(0, col).text = text
    for col, text in enumerate(["1.1.2", "采购人", "示例采购人"]):
        table.cell(1, col).text = text
    doc.add_heading("第三章 评审办法", level=1)
    doc.add_heading("商务评分标准", level=2)
    doc.add_paragraph("企业业绩评分标准。")
    doc.save(file_obj)
    return file_obj.getvalue()


def build_business_section_tree_toc_docx_bytes() -> bytes:
    file_obj = io.BytesIO()
    doc = Document()
    doc.add_paragraph("目录")
    doc.add_paragraph("第一章 招标公告 2")
    doc.add_paragraph("3. 供应商资格要求 5")
    doc.add_paragraph("第二章 供应商须知 8")
    doc.add_paragraph("第一章 招标公告")
    doc.add_paragraph("3. 供应商资格要求")
    doc.add_paragraph("3.1 供应商须为中华人民共和国境内合法注册的独立法人。")
    doc.add_paragraph("第二章 供应商须知")
    doc.add_paragraph("供应商须知前附表")
    doc.save(file_obj)
    return file_obj.getvalue()


def build_business_multilevel_template_cluster_docx_bytes() -> bytes:
    file_obj = io.BytesIO()
    doc = Document()
    doc.add_paragraph("第一章 招标公告")
    doc.add_paragraph("投标文件格式目录")
    doc.add_paragraph("附件2 投标价格表 ........ 12")
    doc.add_page_break()
    doc.add_paragraph("第六章 投标文件格式")
    doc.add_paragraph("附件2 投标价格表")
    doc.add_paragraph("A投标价格总表")
    doc.add_paragraph("表1 A-1  标段一")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "序号"
    table.cell(0, 1).text = "价格"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = ""
    doc.add_page_break()
    doc.add_paragraph("D 技术服务的分项报价")
    doc.add_paragraph("D-1除质保期服务外的技术指导")
    next_table = doc.add_table(rows=1, cols=2)
    next_table.cell(0, 0).text = "服务"
    next_table.cell(0, 1).text = "报价"
    doc.save(file_obj)
    return file_obj.getvalue()


def build_business_attachment_templates_with_toc_docx_bytes() -> bytes:
    file_obj = io.BytesIO()
    doc = Document()
    doc.add_paragraph("目录")
    doc.add_paragraph("第六章 投标文件格式 108")
    doc.add_paragraph("附件1 投标函 109")
    doc.add_paragraph("附件3 货物规格一览表 110")
    doc.add_paragraph("附件6 履约保证函格式承诺书和质量保函 111")
    doc.add_paragraph("第一章 招标公告")
    doc.add_paragraph("普通招标公告正文。")
    doc.add_paragraph("第六章 投标文件格式")
    doc.add_paragraph("附件1 投标函")
    doc.add_paragraph("致：华能集团")
    doc.add_paragraph("投标人（盖章）：____________")
    doc.add_paragraph("附件3 货物规格一览表")
    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    values = [
        ["序号", "货物名称", "规格型号", "数量"],
        ["1", "风力发电机组", "", ""],
        ["2", "塔筒", "", ""],
    ]
    for row_index, row in enumerate(values):
        for col_index, value in enumerate(row):
            table.cell(row_index, col_index).text = value
    doc.add_paragraph("附件6 履约保证函格式承诺书和质量保函")
    doc.add_paragraph("我方承诺按招标文件要求提交履约保证函并承担质量保函责任。")
    doc.add_paragraph("投标人（盖章）：____________")
    doc.save(file_obj)
    return file_obj.getvalue()


def build_business_fingerprint_only_tables_docx_bytes() -> bytes:
    file_obj = io.BytesIO()
    doc = Document()
    doc.add_paragraph("第一章 招标公告")
    doc.add_paragraph("本章为公告。")
    doc.add_heading("商务标格式", level=1)
    doc.add_paragraph("投标报价明细")
    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    values = [
        ["序号", "货物名称", "规格型号", "数量"],
        ["1", "风力发电机组", "", ""],
        ["2", "塔筒", "", ""],
    ]
    for row_index, row in enumerate(values):
        for col_index, value in enumerate(row):
            table.cell(row_index, col_index).text = value
    doc.add_heading("合同条款", level=1)
    doc.add_paragraph("本章不属于投标文件格式。")
    doc.save(file_obj)
    return file_obj.getvalue()


def build_business_commitment_template_alignment_docx_bytes() -> bytes:
    file_obj = io.BytesIO()
    doc = Document()
    doc.add_paragraph("第二章 投标人须知")
    doc.add_paragraph("投标人须无条件承诺在本采购项目第一台合同设备供货前取得本条a和b所述材料，需提供承诺书。")
    doc.add_heading("第六章 投标文件格式", level=1)
    doc.add_paragraph("附件1 材料取得承诺书")
    doc.add_paragraph("我方承诺在本采购项目第一台合同设备供货前取得本条a和b所述材料。")
    doc.add_paragraph("投标人（盖章）：____________")
    doc.save(file_obj)
    return file_obj.getvalue()


def field_by_key(items: list[dict], key: str) -> dict:
    return next(item for item in items if item["key"] == key)


def sample_evaluation_docx_bytes() -> bytes:
    return build_docx_blocks_bytes(
        "第三章 评标办法（综合评估法）",
        "附表2：技术评分标准表",
        [
            ["序号", "评分项", "分值", "得分点", "证明材料要求"],
            ["1", "技术方案", "30分", "总体方案完整、技术路线先进得满分。", "提供技术方案和技术承诺函。"],
            ["2", "供货保障", "10分", "供货计划合理、保障措施充分得满分。", "提供供货计划。"],
        ],
        "附表3：商务评分标准表",
        [
            ["序号", "评分项", "分值", "得分点", "证明材料要求"],
            ["1", "企业业绩", "20分", "近三年同类风电项目业绩满足要求得满分。", "提供合同或中标通知书。"],
            ["2", "财务状况", "10分", "财务状况良好得满分。", "提供审计报告。"],
        ],
        "附表4：投标报价评分标准",
        [
            ["序号", "评分项", "满分", "评分办法"],
            ["1", "投标报价", "100分", "以评标基准价为基础计算报价得分。"],
        ],
        "附表5：投标度电成本评分标准",
        [
            ["序号", "评分项", "满分", "评分办法"],
            ["1", "度电成本", "100分", "按度电成本由低到高计算得分。"],
        ],
        "附表1：符合性审查标准表",
        [
            ["序号", "审查项目", "审查标准", "证明材料要求"],
            ["1", "投标文件签署", "投标文件按招标文件要求签字盖章。", "提供签字盖章页。"],
        ],
    )


def sample_technical_spec_docx_bytes() -> bytes:
    return build_docx_blocks_bytes(
        "第二卷 技术规范书",
        "1.1.1 项目概况",
        [
            ["项目名称", "华能甘肃100MW风电项目"],
            ["招标编号", "HN-2026-001"],
            ["招标人", "华能集团"],
            ["管理单位", "华能甘肃公司"],
            ["标段规模", "100MW"],
            ["交货周期", "2026年10月1日至2027年3月31日"],
            ["质保期", "5年"],
            ["技术承诺", "投标人应承诺满足全部技术规范。"],
        ],
        "招标机型要求",
        [
            ["参数", "要求"],
            ["单机容量", "6.25MW"],
            ["叶轮直径", "200m"],
            ["轮毂高度", "120m"],
            ["叶片最低点距地", "20m"],
            ["塔筒型式", "钢混塔筒"],
            ["箱变型式", "华式箱变"],
            ["安全等级", "IEC IIB"],
            ["空气密度", "1.225kg/m3"],
            ["风速", "8.5m/s"],
            ["湍流强度", "0.14"],
        ],
        "性能保证指标",
        [
            ["指标", "要求"],
            ["功率曲线", "投标人应提供经认证功率曲线。"],
            ["可利用率", "97%"],
            ["发电量", "年上网电量不少于300GWh"],
            ["涉网性能", "满足高低电压穿越要求。"],
        ],
        "环境适应性要求",
        [
            ["要求", "说明"],
            ["抗低温", "满足-30℃低温运行。"],
            ["抗覆冰防凝露", "具备覆冰及防凝露措施。"],
            ["防潮湿", "适应高湿环境。"],
            ["防雷暴", "配置防雷保护。"],
            ["防风沙", "满足风沙环境防护。"],
            ["抗高温", "满足高温环境运行。"],
        ],
        "专题方案：应提供叶片专题、变桨系统专题、主轴专题、齿轮箱专题。",
        "供货范围：风力发电机组、塔筒、箱变及备品备件。",
        "考核条款：发电量考核、可利用率考核、功率曲线考核、部件考核、认证考核。",
    )


class ParsePipelineTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        base = Path(self.temp_dir.name)
        settings.uploads_dir = base / "uploads"
        settings.documents_dir = base / "documents"
        settings.parsed_dir = base / "parsed"
        settings.ensure_dirs()

        store.reset_for_tests()
        store._ensure_db()
        self.technical_appendix_sync_patcher = patch(
            "app.services.bid_project_service.sync_technical_parse_appendices",
            new=AsyncMock(return_value={"status": "synced", "syncedCount": 0}),
        )
        self.technical_appendix_sync_patcher.start()
        self.client = TestClient(app, base_url="http://127.0.0.1:8000")

    def tearDown(self) -> None:
        self.client.close()
        self.technical_appendix_sync_patcher.stop()
        self.temp_dir.cleanup()

    def test_long_running_parse_step_emits_progress_heartbeats_until_done(self) -> None:
        heartbeats = []

        def slow_step() -> str:
            time.sleep(0.05)
            return "done"

        result = parsing_service._run_with_progress_heartbeat(
            slow_step,
            heartbeat=lambda metadata: heartbeats.append(metadata),
            interval_seconds=0.01,
        )

        self.assertEqual(result, "done")
        self.assertGreaterEqual(len(heartbeats), 1)
        self.assertTrue(heartbeats[0]["heartbeat"])
        self.assertEqual(heartbeats[0]["heartbeatIndex"], 1)
        self.assertIn("elapsedSeconds", heartbeats[0])

    def test_long_running_parse_step_checks_cancel_before_heartbeat(self) -> None:
        heartbeats = []

        def slow_step() -> str:
            time.sleep(0.05)
            return "done"

        with self.assertRaises(parsing_service.ParseCancelledError):
            parsing_service._run_with_progress_heartbeat(
                slow_step,
                heartbeat=lambda metadata: heartbeats.append(metadata),
                interval_seconds=0.01,
                cancel_check=lambda: True,
            )

        self.assertEqual(heartbeats, [])

    def create_project(self) -> str:
        response = self.client.post(
            "/api/technical/projects",
            json={"name": "解析测试项目", "customerName": "测试业主"},
        )
        response.raise_for_status()
        return response.json()["id"]

    def create_business_project(self) -> str:
        response = self.client.post(
            "/api/business/projects",
            json={"name": "商务解析测试项目", "customerName": "测试业主"},
        )
        response.raise_for_status()
        return response.json()["id"]

    def project_url(self, project_id: str) -> str:
        project = store._require(project_id)
        if project.get("bidType") == "商务标":
            return f"/api/business/projects/{project_id}"
        return f"/api/technical/projects/{project_id}"

    def parse_results_url(self, project_id: str, suffix: str = "") -> str:
        return f"{self.project_url(project_id)}/parse-results{suffix}"

    def test_upload_and_parse_docx_extracts_text_and_preview(self) -> None:
        project_id = self.create_project()
        file_bytes = build_docx_bytes(
            "上海电气风电项目招标文件",
            "第一章 项目概况",
            "本项目建设地点位于江苏。",
        )

        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[
                (
                    "tenderFiles",
                    ("招标文件.docx", file_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                )
            ],
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["status"], "completed")
        self.assertEqual(payload["summary"]["fileCount"], 1)
        self.assertGreater(payload["summary"]["textLength"], 10)
        self.assertIn("上海电气风电项目招标文件", payload["summary"]["textPreview"])

    def test_upload_and_parse_markdown_extracts_text_and_preview(self) -> None:
        project_id = self.create_project()
        file_bytes = "# Markdown 招标说明\n\n本项目允许使用 Markdown 素材文件。".encode("utf-8")

        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[
                (
                    "tenderFiles",
                    ("招标说明.md", file_bytes, "text/markdown"),
                )
            ],
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["status"], "completed")
        self.assertEqual(payload["sourceFiles"][0]["type"], "MD")
        self.assertIn("Markdown 招标说明", payload["summary"]["textPreview"])

    def test_business_tender_parse_prompt_allows_agentic_navigation_workflow(self) -> None:
        prompt = parsing_service._build_tender_parse_prompt(
            Path("C:/tmp/s1_parse_manifest.json"),
            parsing_service.BUSINESS_PARSE_PROFILE,
        )

        self.assertIn("s1parse prepare", prompt)
        self.assertIn("s1parse overview", prompt)
        self.assertIn("--page-size 60", prompt)
        self.assertIn("s1parse search", prompt)
        self.assertIn("--limit 40", prompt)
        self.assertIn("s1parse read", prompt)
        self.assertIn("--max-chars 4000", prompt)
        self.assertIn("s1parse window", prompt)
        self.assertIn("s1parse table", prompt)
        self.assertIn("--rows 1-24", prompt)
        self.assertIn("--max-chars 8000", prompt)
        self.assertIn("s1parse submit", prompt)
        self.assertIn("s1parse validate", prompt)
        self.assertIn("s1parse status", prompt)
        self.assertIn("s1parse finalize", prompt)
        self.assertIn("opencode-agentic-navigation", prompt)
        self.assertIn("manifest.structuredResultPath", prompt)
        self.assertIn("解析中间产物的大 JSON", prompt)
        self.assertIn("s1parse 小输出导航命令", prompt)
        self.assertNotIn("candidate_package.json", prompt)
        self.assertNotIn("review_plan.json", prompt)
        self.assertNotIn("ai_tasks/**", prompt)
        self.assertNotIn("s1parse tasks", prompt)
        self.assertNotIn("s1parse task ", prompt)
        self.assertNotIn("s1parse decision-all", prompt)
        self.assertNotIn("s1parse decision-set", prompt)
        self.assertNotIn("s1parse validate-decision", prompt)
        self.assertNotIn("decisionPath", prompt)
        self.assertIn("opencode", prompt)
        self.assertIn("read", prompt)
        self.assertIn("Task/subagent", prompt)
        self.assertIn("evidenceId", prompt)
        self.assertIn("finalize", prompt)
        self.assertIn("表格类内容必须完整读取后再提交", prompt)
        self.assertIn("不得基于预览、summary 或局部行推断", prompt)
        self.assertNotIn("Bash 工具执行下面命令", prompt)
        self.assertLess(len(prompt), 2500)
        self.assertNotIn("必须覆盖这些目标", prompt)
        self.assertNotIn("完整性硬约束", prompt)
        self.assertNotIn("原文具备的条款必须逐条提交", prompt)
        self.assertNotIn("每条资格要求必须包含 applicableScope", prompt)
        self.assertNotIn("商务部分评分项目、分值、评分标准", prompt)

    def test_business_tender_parser_skill_declares_frontend_delivery_contract(self) -> None:
        skill_path = (
            Path(__file__).resolve().parents[1]
            / "opencode"
            / "skills"
            / "bid-business-tender-structured-parser"
            / "SKILL.md"
        )
        content = skill_path.read_text(encoding="utf-8")

        self.assertLess(len(content), 5000)
        self.assertIn("s1parse overview <manifest> --page 1 --page-size 60", content)
        self.assertIn('s1parse search <manifest> "<query>" --limit 40', content)
        self.assertIn("s1parse read <manifest> <evidenceId> --mode summary --max-chars 4000", content)
        self.assertIn("s1parse table <manifest> <tableId> --rows 1-24 --max-chars 8000", content)
        self.assertIn("表格类内容必须完整读取后再提交", content)
        self.assertIn("不得基于预览、summary 或局部行推断", content)
        self.assertIn("你是招投标专家，不是关键词匹配器", content)
        self.assertIn("只提交前端清单需要的业务字段", content)
        self.assertIn("资格要求和商务评分的序号不需要提交", content)
        self.assertIn("项目基础信息", content)
        self.assertIn("投标人资格要求", content)
        self.assertIn("投标人须知前附表", content)
        self.assertIn("商务废标项", content)
        self.assertIn("商务评分标准", content)
        self.assertIn("不要为了前端不用的字段额外提交证明材料要求", content)

    def test_business_template_skill_runs_before_structured_parser_and_passes_manifest(self) -> None:
        project_id = self.create_business_project()
        calls: list[str] = []
        seen_manifest: dict[str, object] = {}
        seen_cancel_check: dict[str, bool] = {}

        def fake_section_tree(documents: list[dict], project_dir: Path):
            calls.append("section_tree")
            tree_path = project_dir / "business_section_tree.json"
            payload = {
                "schemaVersion": "bid-business-section-tree-v1",
                "maxLevel": 3,
                "documents": [{"id": documents[0]["id"], "name": documents[0]["name"]}],
                "nodes": [],
                "toc": {"detected": False, "entries": []},
                "validation": {
                    "status": "not_applicable",
                    "tocEntryCount": 0,
                    "matchedTocEntryCount": 0,
                    "unmatchedTocTitles": [],
                },
                "summary": {
                    "documentCount": 1,
                    "nodeCount": 0,
                    "tocEntryCount": 0,
                    "validationStatus": "not_applicable",
                    "warnings": [],
                },
            }
            tree_path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            return tree_path, payload

        def fake_template_extractor(
            *,
            project_id: str,
            documents: list[dict],
            project_dir: Path,
            progress_callback=None,
            cancel_check,
        ):
            calls.append("template")
            seen_cancel_check["callable"] = callable(cancel_check)
            output_dir = project_dir / "business_template_extraction"
            template_dir = output_dir / "templates"
            template_dir.mkdir(parents=True, exist_ok=True)
            template_docx = template_dir / "TPL-0001.docx"
            Document().save(str(template_docx))
            payload = {
                "schemaVersion": "bid-business-template-extractor-v1",
                "skillName": "bid-business-template-extractor",
                "summary": {"templateCount": 1},
                "appendices": [
                    {
                        "id": "APPX-0001",
                        "title": "Bid Letter",
                        "artifactType": "business_attachment_template",
                        "templateType": "bid_letter",
                        "status": "generated",
                        "docxPath": str(template_docx),
                        "sourceDocumentId": documents[0]["id"],
                        "sourceDocumentName": documents[0]["name"],
                        "extractionMode": "business_template_extractor_skill",
                    }
                ],
            }
            extraction_path = output_dir / "business_template_extraction.json"
            extraction_path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            return payload["appendices"], payload, ""

        def fake_structured_parser(skill_manifest_path: Path, **kwargs):
            calls.append("structured")
            manifest = json.loads(skill_manifest_path.read_text(encoding="utf-8"))
            seen_manifest.update(manifest)
            extraction_path = Path(str(manifest.get("businessTemplateExtractionPath") or ""))
            appendices = []
            if extraction_path.is_file():
                extraction_payload = json.loads(extraction_path.read_text(encoding="utf-8"))
                appendices = extraction_payload.get("appendices") or []
            return {
                "items": [],
                "structured": {
                    "schemaVersion": "bid-business-tender-structured-v1",
                    "targetSkill": "bid-business-tender-structured-parser",
                    "mode": "opencode-skill",
                    "sourceDocuments": [],
                    "scoringCriteria": {"business": []},
                    "fieldGroups": {},
                    "requirementPresence": {},
                    "coverage": [],
                    "projectDates": {"endDate": ""},
                    "appendices": appendices,
                    "commitmentLetters": [],
                    "commitmentClues": [],
                    "projectFactFields": [],
                    "categoryCounts": {},
                },
            }, ""

        with patch("app.services.parsing.settings.s1_parse_opencode_enabled", True), patch(
            "app.services.parsing.write_business_section_tree",
            side_effect=fake_section_tree,
        ), patch(
            "app.services.parsing.run_business_template_extractor",
            side_effect=fake_template_extractor,
            create=True,
        ), patch(
            "app.services.parsing._run_parse_skill",
            side_effect=fake_structured_parser,
        ), patch(
            "app.services.parsing._needs_business_s1_finalize_guard",
            return_value=False,
        ):
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[
                    (
                        "tenderFiles",
                        (
                            "business-tender.docx",
                            build_business_attachment_templates_docx_bytes(),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        ),
                    )
                ],
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(calls, ["section_tree", "template", "structured"])
        self.assertTrue(seen_cancel_check["callable"])
        extraction_path = Path(str(seen_manifest.get("businessTemplateExtractionPath") or ""))
        self.assertTrue(extraction_path.is_file())
        self.assertEqual(seen_manifest.get("businessTemplateExtractionSummary"), {"templateCount": 1})
        self.assertTrue(Path(str(seen_manifest.get("businessSectionTreePath") or "")).is_file())
        appendices = response.json()["structured"]["appendices"]
        self.assertEqual(len(appendices), 1)
        self.assertEqual(appendices[0]["extractionMode"], "business_template_extractor_skill")

    def test_business_pdf_uses_docling_document_nav_without_lightweight_fallback(self) -> None:
        project_id = self.create_business_project()
        pdf_path = settings.uploads_dir / project_id / "business.pdf"
        pdf_path.parent.mkdir(parents=True, exist_ok=True)
        pdf_path.write_bytes(b"%PDF-1.4\n")
        nav_payload = {
            "schemaVersion": "business-document-nav-v1",
            "sourceEngine": "docling",
            "documents": [{"id": "DOC-1", "sourcePath": str(pdf_path)}],
            "pages": [{"pageNo": 1, "textDensity": 0.8}],
            "blocks": [
                {
                    "id": "DOC-1:B000001",
                    "type": "heading",
                    "text": "第六章 投标文件格式",
                    "sourceEngine": "docling",
                },
                {
                    "id": "DOC-1:B000002",
                    "type": "paragraph",
                    "text": "本章包含商务偏差表，请投标人填写。",
                    "sourceEngine": "docling",
                }
            ],
            "tables": [],
            "images": [],
            "evidence": [],
            "quality": {"engine": "docling", "status": "completed", "fallbackUsed": False},
        }

        def fake_parse_pdf(self, *, project_id: str, document: dict, output_dir: Path):
            nav_path = output_dir / "document_nav.json"
            quality_path = output_dir / "parse_quality.json"
            nav_path.write_text(json.dumps(nav_payload, ensure_ascii=False), encoding="utf-8")
            quality_path.write_text(
                json.dumps(
                    {
                        "engine": "docling",
                        "status": "completed",
                        "fallbackUsed": False,
                        "doclingMode": "local-text-layer",
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            return {
                "documentParseEngine": "docling",
                "status": "completed",
                "documentNavPath": str(nav_path),
                "parseQualityPath": str(quality_path),
                "doclingMode": "local-text-layer",
            }

        def fake_structured_parser(skill_manifest_path: Path, **kwargs):
            manifest = json.loads(skill_manifest_path.read_text(encoding="utf-8"))
            document = manifest["documents"][0]
            self.assertEqual(document["documentParseEngine"], "docling")
            self.assertEqual(document["doclingMode"], "local-text-layer")
            self.assertTrue(Path(document["documentNavPath"]).is_file())
            self.assertTrue(Path(document["parseQualityPath"]).is_file())
            text_path = Path(document["textPath"])
            self.assertIn("第六章 投标文件格式", text_path.read_text(encoding="utf-8"))
            return {
                "items": [],
                "structured": {
                    "schemaVersion": "bid-business-tender-structured-v1",
                    "targetSkill": "bid-business-tender-structured-parser",
                    "mode": "opencode-skill",
                    "sourceDocuments": [],
                    "scoringCriteria": {"business": []},
                    "fieldGroups": {},
                    "requirementPresence": {},
                    "coverage": [],
                    "projectDates": {"endDate": ""},
                    "appendices": [],
                    "commitmentLetters": [],
                    "commitmentClues": [],
                    "projectFactFields": [],
                    "categoryCounts": {},
                },
            }, ""

        with patch("app.services.parsing.settings.s1_parse_opencode_enabled", True), patch(
            "app.services.parsing.settings.business_pdf_parse_engine",
            "docling",
            create=True,
        ), patch(
            "app.services.parsing.settings.business_pdf_engine_fallback",
            "none",
            create=True,
        ), patch(
            "app.services.parsing.DoclingParseEngine.parse_pdf",
            new=fake_parse_pdf,
        ), patch(
            "app.services.parsing.extract_pdf_text",
            side_effect=AssertionError("Docling business PDF path must not call lightweight extract_pdf_text fallback"),
        ), patch(
            "app.services.parsing.run_business_template_extractor",
            return_value=([], {"schemaVersion": "bid-business-template-extractor-v1", "summary": {"templateCount": 0}}, ""),
        ), patch(
            "app.services.parsing._run_parse_skill",
            side_effect=fake_structured_parser,
        ), patch(
            "app.services.parsing._needs_business_s1_finalize_guard",
            return_value=False,
        ):
            summary, storage = parsing_service.parse_tender_documents(
                project_id,
                [
                    {
                        "id": "DOC-1",
                        "name": "business.pdf",
                        "path": str(pdf_path),
                        "content_type": "application/pdf",
                    }
                ],
                bid_type="商务标",
            )

        self.assertIn("第六章 投标文件格式", summary["textPreview"])
        self.assertEqual(storage["documents"][0]["documentParseEngine"], "docling")
        self.assertEqual(storage["documents"][0]["textLength"], len("第六章 投标文件格式\n\n本章包含商务偏差表，请投标人填写。"))
        self.assertTrue(Path(storage["documents"][0]["documentNavPath"]).is_file())
        self.assertIn("第六章 投标文件格式", Path(storage["documents"][0]["textPath"]).read_text(encoding="utf-8"))
        quality = json.loads(Path(storage["documents"][0]["parseQualityPath"]).read_text(encoding="utf-8"))
        self.assertEqual(quality["engine"], "docling")
        self.assertEqual(quality["sourceEngine"], "docling")
        self.assertEqual(quality["status"], "completed")
        self.assertEqual(quality["qualityStatus"], "needs_review")
        self.assertEqual(quality["doclingMode"], "local-text-layer")
        self.assertFalse(quality["fallbackUsed"])
        self.assertTrue(quality["reviewRequired"])

    def test_business_pdf_low_quality_pages_append_ocr_blocks(self) -> None:
        project_id = self.create_business_project()
        pdf_path = settings.uploads_dir / project_id / "business-low-quality.pdf"
        pdf_path.parent.mkdir(parents=True, exist_ok=True)
        pdf_path.write_bytes(b"%PDF-1.4\n")
        nav_payload = {
            "schemaVersion": "business-document-nav-v1",
            "sourceEngine": "docling",
            "documents": [{"id": "DOC-1", "sourcePath": str(pdf_path)}],
            "pages": [{"pageNo": 1, "textDensity": 0.01}],
            "blocks": [{"id": "DOC-1:B000001", "type": "paragraph", "text": "Docling 原始文本", "pageNo": 1}],
            "tables": [],
            "images": [],
            "evidence": [],
            "quality": {"engine": "docling", "status": "completed", "fallbackUsed": False},
        }

        def fake_parse_pdf(self, *, project_id: str, document: dict, output_dir: Path):
            nav_path = output_dir / "document_nav.json"
            quality_path = output_dir / "parse_quality.json"
            nav_path.write_text(json.dumps(nav_payload, ensure_ascii=False), encoding="utf-8")
            quality_path.write_text(
                json.dumps({"engine": "docling", "status": "completed", "fallbackUsed": False}, ensure_ascii=False),
                encoding="utf-8",
            )
            return {
                "documentParseEngine": "docling",
                "status": "completed",
                "documentNavPath": str(nav_path),
                "parseQualityPath": str(quality_path),
            }

        def fake_ocr_pages(*, project_id: str, document: dict, file_path: Path, page_numbers: list[int]):
            self.assertEqual(page_numbers, [1])
            return {1: {"text": "OCR 补充文本", "meta": {"status": "completed", "pageCount": 1}}}

        with patch("app.services.parsing.settings.s1_parse_opencode_enabled", True), patch(
            "app.services.parsing.settings.business_pdf_parse_engine",
            "docling",
            create=True,
        ), patch(
            "app.services.parsing.DoclingParseEngine.parse_pdf",
            new=fake_parse_pdf,
        ), patch(
            "app.services.parsing._ocr_business_pdf_pages",
            side_effect=fake_ocr_pages,
        ), patch(
            "app.services.parsing.run_business_template_extractor",
            return_value=([], {"schemaVersion": "bid-business-template-extractor-v1", "summary": {"templateCount": 0}}, ""),
        ), patch(
            "app.services.parsing._run_parse_skill",
            return_value=(
                {
                    "items": [],
                    "structured": {
                        "schemaVersion": "bid-business-tender-structured-v1",
                        "targetSkill": "bid-business-tender-structured-parser",
                        "mode": "opencode-skill",
                        "sourceDocuments": [],
                        "scoringCriteria": {"business": []},
                        "fieldGroups": {},
                        "requirementPresence": {},
                        "coverage": [],
                        "projectDates": {"endDate": ""},
                        "appendices": [],
                        "commitmentLetters": [],
                        "commitmentClues": [],
                        "projectFactFields": [],
                        "categoryCounts": {},
                    },
                },
                "",
            ),
        ), patch(
            "app.services.parsing._needs_business_s1_finalize_guard",
            return_value=False,
        ):
            _, storage = parsing_service.parse_tender_documents(
                project_id,
                [
                    {
                        "id": "DOC-1",
                        "name": "business-low-quality.pdf",
                        "path": str(pdf_path),
                        "content_type": "application/pdf",
                    }
                ],
                bid_type="商务标",
            )

        document = storage["documents"][0]
        nav = json.loads(Path(document["documentNavPath"]).read_text(encoding="utf-8"))
        texts = [block["text"] for block in nav["blocks"]]
        quality = json.loads(Path(document["parseQualityPath"]).read_text(encoding="utf-8"))
        self.assertIn("Docling 原始文本", texts)
        self.assertIn("OCR 补充文本", texts)
        self.assertTrue(any(block["type"] == "ocr_text" for block in nav["blocks"]))
        self.assertEqual(document["pageOcr"]["appliedPages"], [1])
        self.assertEqual(quality["engine"], "docling")
        self.assertFalse(quality["fallbackUsed"])
        self.assertEqual(quality["ocrAppliedPages"], [1])

    def test_parse_tender_documents_stops_when_cancel_requested_before_extracting(self) -> None:
        project_id = self.create_project()
        tender_path = settings.uploads_dir / project_id / "cancel-source.md"
        tender_path.parent.mkdir(parents=True, exist_ok=True)
        tender_path.write_text("取消测试文件", encoding="utf-8")

        with self.assertRaisesRegex(RuntimeError, "解析已取消"):
            parsing_service.parse_tender_documents(
                project_id,
                [
                    {
                        "id": "DOC-1",
                        "name": "cancel-source.md",
                        "path": str(tender_path),
                        "content_type": "text/markdown",
                    }
                ],
                bid_type="技术标",
                cancel_check=lambda: True,
            )

    def test_business_section_tree_is_ready_before_structured_parser(self) -> None:
        project_id = self.create_business_project()
        calls: list[str] = []
        seen_manifest: dict[str, object] = {}

        def fake_template_extractor(
            *,
            project_id: str,
            documents: list[dict],
            project_dir: Path,
            progress_callback=None,
            cancel_check=None,
        ):
            calls.append("template")
            return [], {"schemaVersion": "bid-business-template-extractor-v1", "summary": {"templateCount": 0}}, ""

        def fake_structured_parser(skill_manifest_path: Path, **kwargs):
            calls.append("structured")
            manifest = json.loads(skill_manifest_path.read_text(encoding="utf-8"))
            seen_manifest.update(manifest)
            tree_path = Path(str(manifest.get("businessSectionTreePath") or ""))
            self.assertTrue(tree_path.is_file(), manifest)
            tree_payload = json.loads(tree_path.read_text(encoding="utf-8"))
            self.assertEqual(tree_payload["schemaVersion"], "bid-business-section-tree-v1")
            titles = [node["title"] for node in tree_payload["nodes"]]
            self.assertIn("3. 供应商资格要求", titles)
            self.assertIn("供应商须知前附表", titles)
            self.assertIn("商务评分标准", titles)
            qualification_node = next(node for node in tree_payload["nodes"] if node["title"] == "3. 供应商资格要求")
            self.assertLessEqual(qualification_node["contentStartLine"], qualification_node["endLine"])
            self.assertEqual(qualification_node["documentId"], manifest["documents"][0]["id"])
            return {
                "items": [],
                "structured": {
                    "schemaVersion": "bid-business-tender-structured-v1",
                    "targetSkill": "bid-business-tender-structured-parser",
                    "mode": "opencode-skill",
                    "sourceDocuments": [],
                    "scoringCriteria": {"business": []},
                    "fieldGroups": {},
                    "requirementPresence": {},
                    "coverage": [],
                    "projectDates": {"endDate": ""},
                    "appendices": [],
                    "commitmentLetters": [],
                    "commitmentClues": [],
                    "projectFactFields": [],
                    "categoryCounts": {},
                },
            }, ""

        with patch("app.services.parsing.settings.s1_parse_opencode_enabled", True), patch(
            "app.services.parsing.run_business_template_extractor",
            side_effect=fake_template_extractor,
            create=True,
        ), patch(
            "app.services.parsing._run_parse_skill",
            side_effect=fake_structured_parser,
        ), patch(
            "app.services.parsing._needs_business_s1_finalize_guard",
            return_value=False,
        ):
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[
                    (
                        "tenderFiles",
                        (
                            "business-section-tree.docx",
                            build_business_section_tree_docx_bytes(),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        ),
                    )
                ],
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(calls, ["template", "structured"])
        self.assertTrue(Path(str(seen_manifest.get("businessSectionTreePath") or "")).is_file())

    def test_business_section_tree_keeps_plain_toc_lines_out_of_nodes(self) -> None:
        project_id = self.create_business_project()
        seen_tree: dict[str, object] = {}

        def fake_template_extractor(
            *,
            project_id: str,
            documents: list[dict],
            project_dir: Path,
            progress_callback=None,
            cancel_check=None,
        ):
            return [], {"schemaVersion": "bid-business-template-extractor-v1", "summary": {"templateCount": 0}}, ""

        def fake_structured_parser(skill_manifest_path: Path, **kwargs):
            manifest = json.loads(skill_manifest_path.read_text(encoding="utf-8"))
            tree_path = Path(str(manifest.get("businessSectionTreePath") or ""))
            seen_tree.update(json.loads(tree_path.read_text(encoding="utf-8")))
            return {
                "items": [],
                "structured": {
                    "schemaVersion": "bid-business-tender-structured-v1",
                    "targetSkill": "bid-business-tender-structured-parser",
                    "mode": "opencode-skill",
                    "sourceDocuments": [],
                    "scoringCriteria": {"business": []},
                    "fieldGroups": {},
                    "requirementPresence": {},
                    "coverage": [],
                    "projectDates": {"endDate": ""},
                    "appendices": [],
                    "commitmentLetters": [],
                    "commitmentClues": [],
                    "projectFactFields": [],
                    "categoryCounts": {},
                },
            }, ""

        with patch("app.services.parsing.settings.s1_parse_opencode_enabled", True), patch(
            "app.services.parsing.run_business_template_extractor",
            side_effect=fake_template_extractor,
            create=True,
        ), patch(
            "app.services.parsing._run_parse_skill",
            side_effect=fake_structured_parser,
        ), patch(
            "app.services.parsing._needs_business_s1_finalize_guard",
            return_value=False,
        ):
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[
                    (
                        "tenderFiles",
                        (
                            "business-section-tree-toc.docx",
                            build_business_section_tree_toc_docx_bytes(),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        ),
                    )
                ],
            )

        self.assertEqual(response.status_code, 200)
        titles = [node["title"] for node in seen_tree["nodes"]]
        self.assertNotIn("第一章 招标公告 2", titles)
        self.assertNotIn("3. 供应商资格要求 5", titles)
        self.assertIn("3. 供应商资格要求", titles)
        self.assertEqual(seen_tree["validation"]["status"], "passed")

    def test_business_template_extractor_appendices_are_kept_when_structured_parser_returns_empty(self) -> None:
        project_id = self.create_business_project()
        template_docx = settings.parsed_dir / project_id / "business_template_extraction" / "templates" / "TPL-0001.docx"

        def fake_template_extractor(
            *,
            project_id: str,
            documents: list[dict],
            project_dir: Path,
            progress_callback=None,
            cancel_check=None,
        ):
            template_docx.parent.mkdir(parents=True, exist_ok=True)
            Document().save(str(template_docx))
            payload = {
                "schemaVersion": "bid-business-template-extractor-v1",
                "skillName": "bid-business-template-extractor",
                "summary": {"templateCount": 1},
                "appendices": [
                    {
                        "id": "APPX-0001",
                        "title": "Single Block Placeholder",
                        "artifactType": "business_attachment_template",
                        "templateType": "attachment_placeholder",
                        "status": "generated",
                        "docxPath": str(template_docx),
                        "sourceDocumentId": documents[0]["id"],
                        "sourceDocumentName": documents[0]["name"],
                        "extractionMode": "business_template_extractor_skill",
                    }
                ],
            }
            output_path = project_dir / "business_template_extraction" / "business_template_extraction.json"
            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            return payload["appendices"], payload, ""

        def fake_structured_parser(skill_manifest_path: Path, **kwargs):
            return {
                "items": [],
                "structured": {
                    "schemaVersion": "bid-business-tender-structured-v1",
                    "targetSkill": "bid-business-tender-structured-parser",
                    "mode": "opencode-skill",
                    "sourceDocuments": [],
                    "scoringCriteria": {"business": []},
                    "fieldGroups": {},
                    "requirementPresence": {},
                    "coverage": [],
                    "projectDates": {"endDate": ""},
                    "appendices": [],
                    "commitmentLetters": [],
                    "commitmentClues": [],
                    "projectFactFields": [],
                    "categoryCounts": {},
                },
            }, ""

        with patch("app.services.parsing.settings.s1_parse_opencode_enabled", True), patch(
            "app.services.parsing.run_business_template_extractor",
            side_effect=fake_template_extractor,
            create=True,
        ), patch(
            "app.services.parsing._run_parse_skill",
            side_effect=fake_structured_parser,
        ), patch(
            "app.services.parsing._needs_business_s1_finalize_guard",
            return_value=False,
        ):
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[
                    (
                        "tenderFiles",
                        (
                            "business-tender.docx",
                            build_business_attachment_templates_docx_bytes(),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        ),
                    )
                ],
            )

        self.assertEqual(response.status_code, 200)
        appendices = response.json()["structured"]["appendices"]
        self.assertEqual(len(appendices), 1)
        self.assertEqual(appendices[0]["title"], "Single Block Placeholder")
        self.assertEqual(appendices[0]["extractionMode"], "business_template_extractor_skill")

    def test_business_finalized_skill_workflow_is_not_rewritten_by_local_transform(self) -> None:
        project_id = self.create_business_project()
        validation_report_path = settings.parsed_dir / project_id / "agentic_validation_report.json"
        validation_report_path.parent.mkdir(parents=True, exist_ok=True)
        validation_report_path.write_text(
            json.dumps({"schemaVersion": "bid-business-agentic-validation-v1", "status": "passed"}, ensure_ascii=False),
            encoding="utf-8",
        )
        tender = "\n".join(
            [
                "# 商务招标文件",
                "项目名称：后端覆盖回归测试项目",
                "第一章 招标公告",
                "3. 投标人资格要求",
                "3.1 本地旧转换不应覆盖 finalized skill 结果。",
            ]
        ).encode("utf-8")

        skill_payload = {
            "items": [
                {
                    "id": "AI-ITEM-0001",
                    "category": "资格要求",
                    "content": "AI 已接收的资格要求",
                    "sourceFile": "商务招标文件.md",
                    "sourceDocumentId": "DOC-AI",
                    "section": "第一章 招标公告 > 3. 投标人资格要求",
                    "evidence": "AI 已接收的资格要求",
                    "evidenceLocation": "DOC-AI:L4",
                }
            ],
            "structured": {
                "schemaVersion": "bid-business-tender-structured-v1",
                "targetSkill": "bid-business-tender-structured-parser",
                "mode": "opencode-skill",
                "workflow": {
                    "stage": "finalized",
                    "aiReviewTrusted": True,
                    "mode": "opencode-agentic-navigation",
                    "validationReportPath": str(validation_report_path),
                },
                "sourceDocuments": [],
                "scoringCriteria": {"business": [], "price": [], "compliance": [], "lcoe": []},
                "fieldGroups": {
                    "projectBasics": [],
                    "businessResponse": [],
                    "qualificationSupport": [],
                    "qualificationRequirements": [
                        {
                            "id": "QUAL-AI-0001",
                            "content": "AI 已接收的资格要求",
                            "applicableScope": "全部标段",
                            "sourceText": "商务招标文件.md：第一章招标公告第3条",
                            "sourceFile": "商务招标文件.md",
                            "sourceDocumentId": "DOC-AI",
                            "section": "第一章 招标公告 > 3. 投标人资格要求",
                            "evidence": "AI 已接收的资格要求",
                            "evidenceLocation": "DOC-AI:L4",
                            "evidenceIds": ["DOC-AI:L4"],
                        }
                    ],
                    "bidderInstructions": [],
                    "commercialRejectionClauses": [],
                    "commitmentRequirements": [],
                },
                "requirementPresence": {},
                "coverage": [],
                "projectDates": {"startDate": "", "endDate": ""},
                "appendices": [],
                "commitmentLetters": [],
                "commitmentClues": [],
                "projectFactFields": [],
                "categoryCounts": {},
            },
        }

        with patch("app.services.parsing.settings.s1_parse_opencode_enabled", True), patch(
            "app.services.parsing._run_parse_skill",
            return_value=(skill_payload, ""),
        ):
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[("tenderFiles", ("商务招标文件.md", tender, "text/markdown"))],
            )

        self.assertEqual(response.status_code, 200)
        structured = response.json()["structured"]
        self.assertEqual(structured["workflow"]["stage"], "finalized")
        self.assertTrue(structured["workflow"]["aiReviewTrusted"])
        qualification_text = "\n".join(
            row["content"] for row in structured["fieldGroups"]["qualificationRequirements"]
        )
        self.assertIn("AI 已接收的资格要求", qualification_text)

    def test_business_parse_manifest_does_not_include_script_ai_review_config(self) -> None:
        project_id = self.create_business_project()
        tender = "\n".join(
            [
                "# 商务招标文件",
                "项目名称：AI 审查配置注入测试项目",
                "第一章 招标公告",
                "3. 投标人资格要求",
                "3.1 投标人须为境内合法注册的独立法人。",
            ]
        ).encode("utf-8")
        captured_manifest: dict[str, Any] = {}

        def fake_run_parse_skill(skill_manifest_path: Path, **_kwargs: Any):
            captured_manifest.update(json.loads(skill_manifest_path.read_text(encoding="utf-8")))
            return _kwargs["local_result"], "unit-test stops before opencode"

        with patch("app.services.parsing.settings.s1_parse_opencode_enabled", True), patch(
            "app.services.system_settings.system_settings_service.get_opencode_model_config_sync",
            return_value={
                "enabled": True,
                "baseUrl": "https://llm.example.com/v1",
                "apiKey": "llm-secret",
                "model": "deepseek-v4-pro",
                "modelId": "deepseek-v4-pro",
                "timeoutMs": 45000,
                "maxTokens": 12000,
            },
        ), patch("app.services.parsing._run_parse_skill", side_effect=fake_run_parse_skill):
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[("tenderFiles", ("商务招标文件.md", tender, "text/markdown"))],
            )

        self.assertEqual(response.status_code, 200)
        self.assertNotIn("aiReviewMode", captured_manifest)
        self.assertNotIn("aiReviewBaseUrl", captured_manifest)
        self.assertNotIn("aiReviewApiKey", captured_manifest)
        self.assertNotIn("aiReviewModel", captured_manifest)
        self.assertNotIn("aiReviewTimeoutSec", captured_manifest)
        self.assertNotIn("aiReviewMaxTokens", captured_manifest)

    def test_business_skill_result_is_authoritative_without_local_semantic_backfill(self) -> None:
        project_id = self.create_business_project()
        validation_report_path = settings.parsed_dir / project_id / "agentic_validation_report.json"
        validation_report_path.parent.mkdir(parents=True, exist_ok=True)
        validation_report_path.write_text(
            json.dumps({"schemaVersion": "bid-business-agentic-validation-v1", "status": "passed"}, ensure_ascii=False),
            encoding="utf-8",
        )
        tender = "\n".join(
            [
                "# Business tender",
                "Project name: backend must not backfill this",
                "Bid deadline: 2026-03-18 09:30",
                "Qualification: bidder must be an independent legal person.",
            ]
        ).encode("utf-8")
        skill_payload = {
            "items": [],
            "structured": {
                "schemaVersion": "bid-business-tender-structured-v1",
                "targetSkill": "bid-business-tender-structured-parser",
                "mode": "opencode-skill",
                "workflow": {
                    "stage": "finalized",
                    "aiReviewTrusted": True,
                    "mode": "opencode-agentic-navigation",
                    "validationReportPath": str(validation_report_path),
                },
                "sourceDocuments": [],
                "scoringCriteria": {},
                "fieldGroups": {},
                "coverage": [],
                "projectDates": {},
                "projectFactFields": [],
            },
        }

        with patch("app.services.parsing.settings.s1_parse_opencode_enabled", True), patch(
            "app.services.parsing._run_parse_skill",
            return_value=(skill_payload, ""),
        ), patch(
            "app.services.parsing._transform_to_business_contract",
            wraps=parsing_service._transform_to_business_contract,
        ) as local_transform:
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[("tenderFiles", ("business-tender.md", tender, "text/markdown"))],
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(local_transform.call_count, 0)
        structured = response.json()["structured"]
        self.assertEqual(structured["fieldGroups"], {})
        self.assertEqual(structured["projectDates"], {})

    def test_technical_skill_result_does_not_backfill_project_basics_from_local_parser(self) -> None:
        project_id = self.create_project()
        tender = "\n".join(
            [
                "华能蒙东新能源公司赤峰市200万千瓦自建调峰能力风光储多能互补一体化+荒漠治理基地项目（翁牛特旗120万千瓦风电项目区）",
                "风力发电机组（不含塔架）及附属设备采购",
                "招 标 文 件",
                "招标编号：HNZB2025-12-1-382-01",
                "招标人：华能内蒙古东部能源有限公司",
                "项目单位：华能翁牛特旗新能源有限公司",
                "招标代理机构：中国华能集团有限公司北京睿采数动科技分公司",
                "投标文件递交截止时间：2026年01月26日15时00分",
                "本项目招标范围为整套风力发电机组及塔筒内所有必要设备。",
            ]
        ).encode("utf-8")

        def fake_run_parse_skill(_skill_manifest_path: Path, **kwargs):
            skill_result = json.loads(json.dumps(kwargs["local_result"], ensure_ascii=False))
            project_basics = skill_result["structured"].setdefault("fieldGroups", {}).setdefault("projectBasics", [])
            for row in project_basics:
                row["value"] = ""
                row["status"] = "missing"
                row["sourceFile"] = ""
                row["sourceDocumentId"] = ""
                row["section"] = ""
                row["evidence"] = ""
                row["evidenceLocation"] = ""
            skill_result["structured"].pop("projectDates", None)
            skill_result["structured"]["projectFactFields"] = project_basics
            skill_result["structured"]["workflow"] = {
                "stage": "finalized",
                "mode": "opencode-agentic-navigation",
                "submittedTargetCount": 2,
                "missingTargets": [],
                "validationErrors": [],
            }
            return skill_result, ""

        with patch("app.services.parsing.settings.s1_parse_opencode_enabled", True), patch(
            "app.services.parsing._run_parse_skill",
            side_effect=fake_run_parse_skill,
        ):
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[("tenderFiles", ("招标文件-技术规范.md", tender, "text/markdown"))],
            )

        self.assertEqual(response.status_code, 200)
        structured = response.json()["structured"]
        project_basics = structured["fieldGroups"]["projectBasics"]
        self.assertEqual(field_by_key(project_basics, "projectName")["value"], "")
        self.assertEqual(field_by_key(project_basics, "tenderNo")["value"], "")
        self.assertEqual(field_by_key(project_basics, "projectUnit")["value"], "")
        self.assertEqual(field_by_key(project_basics, "tenderer")["value"], "")
        self.assertEqual(field_by_key(project_basics, "tenderAgency")["value"], "")
        self.assertEqual(field_by_key(project_basics, "bidDeadline")["value"], "")
        self.assertNotIn("projectDates", structured)
        self.assertEqual(structured["projectFactFields"], project_basics)
        self.assertNotIn("localProjectBasicsMerged", structured["workflow"])

    def test_technical_skill_result_updates_project_deadline_from_project_basics_without_project_dates(self) -> None:
        project_id = self.create_project()
        tender = "\n".join(
            [
                "都匀市盛黔风电场风力发电机组及附属设备采购项目",
                "招标编号：PC-0307-26J1-FG0002",
                "招标人：都匀盛黔新能源有限公司",
                "投标文件递交截止时间：2026年05月06日10时00分",
                "本项目招标范围为整套风力发电机组及塔筒内所有必要设备。",
            ]
        ).encode("utf-8")

        def fake_run_parse_skill(_skill_manifest_path: Path, **kwargs):
            skill_result = json.loads(json.dumps(kwargs["local_result"], ensure_ascii=False))
            project_basics = [
                {"key": "projectName", "fieldKey": "projectName", "label": "项目名称", "value": "都匀市盛黔风电场风力发电机组及附属设备采购项目"},
                {"key": "tenderNo", "fieldKey": "tenderNo", "label": "招标编号", "value": "PC-0307-26J1-FG0002"},
                {"key": "projectUnit", "fieldKey": "projectUnit", "label": "项目单位", "value": ""},
                {"key": "tenderer", "fieldKey": "tenderer", "label": "招标人", "value": "都匀盛黔新能源有限公司"},
                {"key": "tenderAgency", "fieldKey": "tenderAgency", "label": "招标代理机构", "value": ""},
                {"key": "bidDeadline", "fieldKey": "bidDeadline", "label": "递交截止时间", "status": "found", "value": "2026-05-06 10:00"},
            ]
            structured = skill_result.setdefault("structured", {})
            structured.setdefault("fieldGroups", {})["projectBasics"] = project_basics
            structured.pop("projectDates", None)
            structured["projectFactFields"] = project_basics
            structured["workflow"] = {
                "stage": "finalized",
                "mode": "opencode-agentic-navigation",
                "submittedTargetCount": 2,
                "missingTargets": [],
                "validationErrors": [],
            }
            return skill_result, ""

        with patch("app.services.parsing.settings.s1_parse_opencode_enabled", True), patch(
            "app.services.parsing._run_parse_skill",
            side_effect=fake_run_parse_skill,
        ):
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[("tenderFiles", ("招标文件-技术规范.md", tender, "text/markdown"))],
            )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        structured = payload["structured"]
        project_basics = structured["fieldGroups"]["projectBasics"]
        self.assertEqual(field_by_key(project_basics, "bidDeadline")["value"], "2026-05-06 10:00")
        self.assertNotIn("projectDates", structured)
        self.assertNotIn("projectDates", payload["summary"])

        project = store._require(project_id)
        self.assertEqual(project["endDate"], "2026-05-06 10:00")
        self.assertEqual(project["deadline"], "2026-05-06 10:00")

    def test_technical_parse_results_preserves_agentic_project_basics_without_text_repair(self) -> None:
        project_id = self.create_project()
        parse_dir = settings.parsed_dir / project_id
        parse_dir.mkdir(parents=True, exist_ok=True)
        text_path = parse_dir / "TEN-1.txt"
        text_path.write_text(
            "\n".join(
                [
                    "华能蒙东新能源公司赤峰市200万千瓦自建调峰能力风光储多能互补一体化+荒漠治理基地项目（翁牛特旗120万千瓦风电项目区）",
                    "风力发电机组（不含塔架）及附属设备采购",
                    "招 标 文 件",
                    "招标编号：HNZB2025-12-1-382-01",
                    "招标人",
                    "：",
                    "华能内蒙古东部能源有限公司",
                    "管理单位",
                    "：",
                    "华能翁牛特旗新能源有限公司",
                    "招标代理机构",
                    "：",
                    "中国华能集团有限公司北京睿采数动科技分公司",
                ]
            ),
            encoding="utf-8",
        )
        structured_path = parse_dir / "s1_structured_result.json"
        empty_basics = [
            {"key": key, "label": label, "value": "", "status": "missing", "fieldKey": key}
            for key, label in (
                ("projectName", "项目名称"),
                ("tenderNo", "招标编号"),
                ("projectUnit", "项目单位"),
                ("tenderer", "招标人"),
                ("tenderAgency", "招标代理机构"),
                ("bidDeadline", "递交截止时间"),
            )
        ]
        structured_path.write_text(
            json.dumps(
                {
                    "items": [],
                    "structured": {
                        "schemaVersion": "bid-tender-structured-v1",
                        "workflow": {"stage": "finalized", "mode": "opencode-agentic-navigation"},
                        "fieldGroups": {"projectBasics": empty_basics},
                        "projectDates": {"startDate": "", "endDate": ""},
                        "projectFactFields": empty_basics,
                        "technicalInterpretation": {"items": [], "summary": {"total": 0}},
                    },
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        complete_parse_for_tests(
            project_id,
            [{"id": "TEN-1", "name": "招标文件-技术规范.md", "size_label": "1 KB"}],
            [],
            summary={"fileCount": 1, "extractedCount": 0, "textLength": 0, "warnings": []},
            parse_storage={
                "structuredResultPath": str(structured_path),
                "documents": [
                    {
                        "id": "TEN-1",
                        "name": "招标文件-技术规范.md",
                        "textPath": str(text_path),
                    }
                ],
                "items": [],
                "structured": {
                    "schemaVersion": "bid-tender-structured-v1",
                    "workflow": {"stage": "finalized", "mode": "opencode-agentic-navigation"},
                    "fieldGroups": {"projectBasics": empty_basics},
                    "projectDates": {"startDate": "", "endDate": ""},
                    "projectFactFields": empty_basics,
                    "technicalInterpretation": {"items": [], "summary": {"total": 0}},
                },
            },
        )

        response = self.client.get(self.parse_results_url(project_id))

        self.assertEqual(response.status_code, 200)
        structured = response.json()["structured"]
        project_basics = structured["fieldGroups"]["projectBasics"]
        self.assertEqual(field_by_key(project_basics, "projectName")["value"], "")
        self.assertEqual(field_by_key(project_basics, "tenderer")["value"], "")
        self.assertEqual(field_by_key(project_basics, "projectUnit")["value"], "")
        self.assertNotIn("localProjectBasicsMerged", structured["workflow"])
        persisted = json.loads(structured_path.read_text(encoding="utf-8"))
        self.assertEqual(
            field_by_key(persisted["structured"]["fieldGroups"]["projectBasics"], "tenderAgency")["value"],
            "",
        )

    def test_business_skill_failure_fails_parse_without_local_fallback_result(self) -> None:
        project_id = self.create_business_project()
        tender = "\n".join(
            [
                "# Business tender",
                "Project name: local fallback must not become completed result",
                "Bid deadline: 2026-03-18 09:30",
            ]
        ).encode("utf-8")

        with patch("app.services.parsing.settings.s1_parse_opencode_enabled", True), patch(
            "app.services.parsing._run_parse_skill",
            side_effect=RuntimeError("unit-test skill failure"),
        ):
            client = TestClient(app, base_url="http://127.0.0.1:8000", raise_server_exceptions=False)
            try:
                response = client.post(
                    self.parse_results_url(project_id, "/upload-and-run"),
                    files=[("tenderFiles", ("business-tender.md", tender, "text/markdown"))],
                )
            finally:
                client.close()

        self.assertGreaterEqual(response.status_code, 500)
        progress = self.client.get(self.parse_results_url(project_id, "/progress")).json()
        self.assertEqual(progress["status"], "failed")
        self.assertIn("unit-test skill failure", progress["summary"])

    def test_business_agentic_prepared_workflow_triggers_backend_finalize_guard(self) -> None:
        project_id = self.create_business_project()
        tender = "\n".join(
            [
                "# Business tender",
                "Project name: prepared workflow guard project",
                "Qualification: bidder must be an independent legal person.",
            ]
        ).encode("utf-8")

        def fake_run_parse_skill(skill_manifest_path: Path, **kwargs):
            parse_dir = skill_manifest_path.parent
            prepared = json.loads(json.dumps(kwargs["local_result"], ensure_ascii=False))
            prepared["structured"]["mode"] = "opencode-skill"
            prepared["structured"]["workflow"] = {
                "stage": "prepared",
                "mode": "opencode-agentic-navigation",
                "navStorePath": str(parse_dir / "s1_nav.sqlite"),
                "documentMapPath": str(parse_dir / "document_map.json"),
                "submissionPath": str(parse_dir / "agentic_submissions.json"),
                "validationReportPath": str(parse_dir / "validation_report.json"),
                "submittedTargetCount": 0,
                "missingTargets": ["qualificationRequirements"],
                "validationErrors": [],
            }
            return prepared, ""

        def fake_finalize(skill_manifest_path: Path, structured_result: dict, profile):
            finalized = json.loads(json.dumps(structured_result, ensure_ascii=False))
            validation_report_path = skill_manifest_path.parent / "validation_report.json"
            validation_report_path.write_text(
                json.dumps({"schemaVersion": "bid-business-agentic-validation-v1", "status": "passed"}, ensure_ascii=False),
                encoding="utf-8",
            )
            finalized["structured"]["mode"] = "opencode-skill"
            finalized["structured"]["workflow"] = {
                "stage": "finalized",
                "mode": "opencode-agentic-navigation",
                "navStorePath": str(skill_manifest_path.parent / "s1_nav.sqlite"),
                "documentMapPath": str(skill_manifest_path.parent / "document_map.json"),
                "submissionPath": str(skill_manifest_path.parent / "agentic_submissions.json"),
                "validationReportPath": str(validation_report_path),
                "submittedTargetCount": 1,
                "missingTargets": [],
                "validationErrors": [],
            }
            finalized["structured"].setdefault("fieldGroups", {})["qualificationRequirements"] = [
                {"content": "finalize guard qualification", "evidenceIds": ["DOC-1:B000001"]}
            ]
            return finalized, ""

        with patch("app.services.parsing.settings.s1_parse_opencode_enabled", True), patch(
            "app.services.parsing._run_parse_skill",
            side_effect=fake_run_parse_skill,
        ), patch("app.services.parsing._finalize_business_s1_result", side_effect=fake_finalize) as finalize_guard:
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[("tenderFiles", ("business-tender.md", tender, "text/markdown"))],
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(finalize_guard.call_count, 1)
        structured = response.json()["structured"]
        self.assertEqual(structured["workflow"]["stage"], "finalized")
        self.assertEqual(structured["workflow"]["mode"], "opencode-agentic-navigation")
        qualification_text = "\n".join(row["content"] for row in structured["fieldGroups"]["qualificationRequirements"])
        self.assertIn("finalize guard qualification", qualification_text)

    def test_business_finalized_agentic_workflow_without_validation_report_triggers_finalize_guard(self) -> None:
        project_id = self.create_business_project()
        tender = "\n".join(
            [
                "# Business tender",
                "Project name: validation report guard project",
                "Qualification: bidder must be an independent legal person.",
            ]
        ).encode("utf-8")

        def fake_run_parse_skill(skill_manifest_path: Path, **kwargs):
            parse_dir = skill_manifest_path.parent
            validation_report_path = parse_dir / "validation_report.json"
            validation_report_path.unlink(missing_ok=True)
            finalized = json.loads(json.dumps(kwargs["local_result"], ensure_ascii=False))
            finalized["structured"]["mode"] = "opencode-skill"
            finalized["structured"]["workflow"] = {
                "stage": "finalized",
                "mode": "opencode-agentic-navigation",
                "navStorePath": str(parse_dir / "s1_nav.sqlite"),
                "documentMapPath": str(parse_dir / "document_map.json"),
                "submissionPath": str(parse_dir / "agentic_submissions.json"),
                "validationReportPath": str(validation_report_path),
                "submittedTargetCount": 5,
                "missingTargets": [],
                "validationErrors": [],
            }
            return finalized, ""

        def fake_finalize(skill_manifest_path: Path, structured_result: dict, profile):
            finalized = json.loads(json.dumps(structured_result, ensure_ascii=False))
            validation_report_path = skill_manifest_path.parent / "validation_report.json"
            validation_report_path.write_text(
                json.dumps({"schemaVersion": "bid-business-agentic-validation-v1", "status": "passed"}, ensure_ascii=False),
                encoding="utf-8",
            )
            finalized["structured"]["workflow"]["validationReportPath"] = str(validation_report_path)
            return finalized, ""

        with patch("app.services.parsing.settings.s1_parse_opencode_enabled", True), patch(
            "app.services.parsing._run_parse_skill",
            side_effect=fake_run_parse_skill,
        ), patch("app.services.parsing._finalize_business_s1_result", side_effect=fake_finalize) as finalize_guard:
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[("tenderFiles", ("business-tender.md", tender, "text/markdown"))],
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(finalize_guard.call_count, 1)
        workflow = response.json()["structured"]["workflow"]
        self.assertEqual(workflow["stage"], "finalized")
        self.assertEqual(workflow["mode"], "opencode-agentic-navigation")
        self.assertTrue(Path(workflow["validationReportPath"]).is_file())

    def test_business_finalize_guard_preserves_opencode_trace_and_separates_backend_stdout(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            parse_dir = Path(tmp)
            structured_path = parse_dir / "s1_structured_result.json"
            manifest_path = parse_dir / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps({"structuredResultPath": str(structured_path)}, ensure_ascii=False),
                encoding="utf-8",
            )
            original_trace = {
                "status": "received",
                "sessionId": "ses-prj0017",
                "parts": [{"type": "tool", "text": "read /data/parsed/PRJ-0017/document_map.json running"}],
            }
            prepared_result = {
                "items": [],
                "structured": {
                    "schemaVersion": "bid-business-tender-structured-v1",
                    "targetSkill": "bid-business-tender-structured-parser",
                    "mode": "opencode-skill",
                    "workflow": {
                        "stage": "prepared",
                        "mode": "opencode-agentic-navigation",
                        "navStorePath": str(parse_dir / "s1_nav.sqlite"),
                        "documentMapPath": str(parse_dir / "document_map.json"),
                        "submissionPath": str(parse_dir / "agentic_submissions.json"),
                        "validationReportPath": str(parse_dir / "validation_report.json"),
                    },
                    "sourceDocuments": [],
                    "fieldGroups": {},
                    "scoringCriteria": {"business": []},
                    "coverage": [],
                    "projectDates": {"startDate": "", "endDate": ""},
                    "appendices": [],
                    "opencodeOutput": original_trace,
                },
            }
            finalized_result = json.loads(json.dumps(prepared_result, ensure_ascii=False))
            finalized_result["structured"]["workflow"]["stage"] = "finalized"
            finalized_result["structured"]["workflow"]["submittedTargetCount"] = 0
            structured_path.write_text(json.dumps(finalized_result, ensure_ascii=False), encoding="utf-8")

            class Completed:
                returncode = 0
                stdout = '{"summary":{"workflowStage":"finalized","submittedTargetCount":0}}'
                stderr = ""

            with patch("app.services.parsing.subprocess.run", return_value=Completed()):
                result, warning = parsing_service._finalize_business_s1_result(
                    manifest_path,
                    prepared_result,
                    parsing_service.BUSINESS_PARSE_PROFILE,
                )

        self.assertEqual(warning, "")
        structured = result["structured"]
        self.assertEqual(structured["opencodeOutput"], original_trace)
        self.assertEqual(
            structured["backendFinalizeOutput"]["stdout"],
            '{"summary":{"workflowStage":"finalized","submittedTargetCount":0}}',
        )
        self.assertTrue(structured["workflow"]["backendFinalizeGuardApplied"])

    def test_business_run_parse_skill_raises_when_opencode_never_produces_skill_result(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            manifest_path = Path(tmp) / "s1_parse_manifest.json"
            manifest_path.write_text("{}", encoding="utf-8")
            local_result = {
                "items": [{"id": "LOCAL-1", "title": "local fallback should not be returned"}],
                "structured": {
                    "schemaVersion": "bid-business-tender-structured-v1",
                    "targetSkill": "bid-business-tender-structured-parser",
                    "mode": "local-structured-parser",
                    "sourceDocuments": [],
                    "fieldGroups": {},
                    "scoringCriteria": {},
                    "coverage": [],
                    "projectDates": {"startDate": "", "endDate": ""},
                },
            }
            error = RuntimeError("unit-test opencode failed before finalize")
            error.opencode_trace = {
                "status": "stalled",
                "sessionId": "ses-business-no-finalize",
                "agentStatus": "stalled",
                "lastTool": "bash",
                "lastToolStatus": "running",
                "failureReason": "s1parse finalize did not complete",
            }

            with patch("app.services.parsing.settings.s1_parse_opencode_enabled", True), patch(
                "app.services.parsing.OpencodeClient.generate_tender_parse_with_trace",
                side_effect=error,
            ):
                with self.assertRaisesRegex(RuntimeError, "S1 商务解析 Skill 调用失败"):
                    parsing_service._run_parse_skill(
                        manifest_path,
                        local_result=local_result,
                        profile=parsing_service.BUSINESS_PARSE_PROFILE,
                    )

    def test_technical_run_parse_skill_preserves_stalled_opencode_trace_on_fallback(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            manifest_path = Path(tmp) / "s1_parse_manifest.json"
            manifest_path.write_text("{}", encoding="utf-8")
            local_result = {
                "items": [],
                "structured": {
                    "schemaVersion": "bid-tender-structured-v1",
                    "targetSkill": "bid-tech-tender-structured-parser",
                    "mode": "local-structured-parser",
                    "sourceDocuments": [],
                    "fieldGroups": {},
                    "scoringCriteria": {},
                    "coverage": [],
                    "projectDates": {"startDate": "", "endDate": ""},
                },
            }
            error = RuntimeError("opencode incomplete/stalled: sessionId=ses-prj0017")
            error.opencode_trace = {
                "status": "stalled",
                "sessionId": "ses-prj0017",
                "agentStatus": "stalled",
                "lastTool": "read",
                "lastToolStatus": "running",
                "lastToolInput": {"filePath": "/data/parsed/PRJ-0017/document_map.json"},
                "failureReason": "read document_map.json running",
            }

            with patch("app.services.parsing.settings.s1_parse_opencode_enabled", True), patch(
                "app.services.parsing.OpencodeClient.generate_tender_parse_with_trace",
                side_effect=error,
            ):
                progress_events = []
                result, warning = parsing_service._run_parse_skill(
                    manifest_path,
                    local_result=local_result,
                    profile=parsing_service.TECHNICAL_PARSE_PROFILE,
                    progress_callback=lambda event, details: progress_events.append((event, details)),
                )

        structured = result["structured"]
        self.assertIn("Skill", warning)
        self.assertIn("opencode incomplete/stalled", warning)
        self.assertEqual(structured["opencodeOutput"]["sessionId"], "ses-prj0017")
        self.assertEqual(structured["workflow"]["opencodeSessionId"], "ses-prj0017")
        self.assertEqual(structured["workflow"]["opencodeAgentStatus"], "stalled")
        self.assertEqual(structured["workflow"]["opencodeLastTool"], "read")
        self.assertEqual(structured["workflow"]["opencodeLastToolStatus"], "running")
        self.assertIn("document_map.json", json.dumps(structured["opencodeOutput"]["lastToolInput"], ensure_ascii=False))
        self.assertEqual(progress_events[-1][0], "opencode_delta")
        self.assertEqual(progress_events[-1][1]["sessionId"], "ses-prj0017")

    def test_run_parse_skill_retries_once_after_incomplete_opencode_session(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            manifest_path = Path(tmp) / "s1_parse_manifest.json"
            structured_path = Path(tmp) / "s1_structured_result.json"
            manifest_path.write_text(json.dumps({"structuredResultPath": str(structured_path)}), encoding="utf-8")
            local_result = {
                "items": [],
                "structured": {
                    "schemaVersion": "bid-business-tender-structured-v1",
                    "targetSkill": "bid-business-tender-structured-parser",
                    "mode": "local-structured-parser",
                    "sourceDocuments": [],
                    "fieldGroups": {},
                    "scoringCriteria": {},
                    "coverage": [],
                    "projectDates": {"startDate": "", "endDate": ""},
                },
            }
            final_result = {
                "items": [],
                "structured": {
                    "schemaVersion": "bid-business-tender-structured-v1",
                    "targetSkill": "bid-business-tender-structured-parser",
                    "mode": "opencode-skill",
                    "workflow": {
                        "stage": "finalized",
                        "mode": "opencode-agentic-navigation",
                    },
                    "fieldGroups": {"qualificationRequirements": [{"content": "retry success"}]},
                    "scoringCriteria": {"business": []},
                    "coverage": [],
                    "projectDates": {},
                },
            }
            structured_path.write_text(json.dumps(final_result, ensure_ascii=False), encoding="utf-8")
            error = RuntimeError("opencode incomplete/stalled: sessionId=ses-first")
            error.opencode_trace = {
                "status": "stalled",
                "sessionId": "ses-first",
                "agentStatus": "stalled",
                "lastTool": "bash",
                "lastToolStatus": "completed",
                "failureReason": "did not complete s1parse finalize",
            }

            calls: list[str] = []

            def fake_generate(prompt: str, **_kwargs: Any) -> dict[str, Any]:
                calls.append(prompt)
                if len(calls) == 1:
                    raise error
                return {"outputFile": str(structured_path)}

            with patch("app.services.parsing.settings.s1_parse_opencode_enabled", True), patch(
                "app.services.parsing.OpencodeClient.generate_tender_parse_with_trace",
                side_effect=fake_generate,
            ):
                result, warning = parsing_service._run_parse_skill(
                    manifest_path,
                    local_result=local_result,
                    profile=parsing_service.BUSINESS_PARSE_PROFILE,
                )

        self.assertEqual(warning, "")
        self.assertEqual(len(calls), 2)
        self.assertIn("s1parse status", calls[1])
        self.assertIn("s1parse finalize", calls[1])
        structured = result["structured"]
        self.assertEqual(structured["mode"], "opencode-skill")
        self.assertEqual(structured["fieldGroups"]["qualificationRequirements"][0]["content"], "retry success")
        attempts = structured["workflow"]["opencodeAttempts"]
        self.assertEqual(len(attempts), 2)
        self.assertEqual(attempts[0]["sessionId"], "ses-first")
        self.assertEqual(attempts[0]["status"], "stalled")
        self.assertEqual(attempts[1]["status"], "succeeded")
        self.assertEqual(attempts[1]["attempt"], 2)

    def test_upload_and_parse_image_uses_visual_recognition_without_manual_ocr_flow(self) -> None:
        project_id = self.create_project()

        fake_response = {
            "choices": [
                {
                    "message": {
                        "content": "项目名称：图片型招标文件\n招标编号：IMG-2026-001\n投标截止日期：2026年8月20日"
                    }
                }
            ]
        }

        async def fake_post(*_args, **_kwargs):
            class Response:
                status_code = 200

                @staticmethod
                def json():
                    return fake_response

            return Response()

        with patch("app.services.system_settings.system_settings_service.get_model_secret_config") as config, patch(
            "httpx.AsyncClient.post",
            side_effect=fake_post,
        ):
            config.return_value = {
                "enabled": True,
                "baseUrl": "https://ocr.example.com/v1",
                "apiKey": "ocr-secret-key",
                "model": "deepseek-ai/DeepSeek-OCR",
                "timeoutMs": 60000,
                "maxTokens": 2048,
            }
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[
                    (
                        "tenderFiles",
                        ("图片型招标文件.png", b"\x89PNG\r\n\x1a\nfake", "image/png"),
                    )
                ],
            )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["status"], "completed")
        self.assertIn("图片型招标文件", payload["summary"]["textPreview"])
        self.assertIn("图片文件已通过 OCR/视觉模型转为可解析文本。", payload["summary"]["warnings"])
        self.assertNotIn("ocrConfirmedFields", payload["structured"])
        self.assertEqual(payload["structured"]["projectDates"]["endDate"], "2026-08-20")

    def test_upload_and_parse_multiple_tenders_extracts_structured_requirements_and_dates(self) -> None:
        project_id = self.create_project()
        main_tender = "\n".join(
            [
                "# 总发包招标文件",
                "项目名称：华能甘肃100MW风电项目",
                "招标编号：HN-2026-001",
                "招标人：华能集团",
                "招标文件获取时间：2026年6月1日至2026年6月10日",
                "投标截止日期：2026年9月30日",
                "评分细则：技术方案30分，供货保障10分。",
                "交货周期：2026年10月1日至2027年3月31日",
            ]
        ).encode("utf-8")
        child_tender = "\n".join(
            [
                "# 子项目招标文件",
                "单机容量：6.25MW",
                "叶轮直径：200m",
                "轮毂高度：120m",
                "可利用率：97%",
                "功率曲线保证率：95%",
                "环境适应性要求：低温-30℃、覆冰、防雷暴。",
                "专题方案要求：叶片专题方案、变桨系统专题方案。",
            ]
        ).encode("utf-8")

        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[
                ("tenderFiles", ("总发包招标文件.md", main_tender, "text/markdown")),
                ("tenderFiles", ("子项目招标文件.md", child_tender, "text/markdown")),
            ],
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["summary"]["fileCount"], 2)
        self.assertGreaterEqual(payload["summary"]["extractedCount"], 10)

        category_labels = {category["label"] for category in payload["structured"]["categories"]}
        self.assertIn("评分细则", category_labels)
        self.assertIn("项目基础信息", category_labels)
        self.assertIn("风机核心参数", category_labels)
        self.assertIn("性能保证指标", category_labels)
        self.assertIn("环境适应性要求", category_labels)
        self.assertIn("专题方案要求", category_labels)

        item_types = {item["type"] for item in payload["items"]}
        self.assertIn("评分细则", item_types)
        self.assertIn("项目基础信息", item_types)
        self.assertIn("风机核心参数", item_types)
        self.assertIn("性能保证指标", item_types)
        self.assertIn("环境适应性要求", item_types)
        self.assertIn("专题方案要求", item_types)

        source_files = {item["sourceFile"] for item in payload["items"]}
        self.assertIn("总发包招标文件.md", source_files)
        self.assertIn("子项目招标文件.md", source_files)
        self.assertTrue(all(item.get("evidence") for item in payload["items"]))
        self.assertTrue(all(item.get("evidenceLocation") for item in payload["items"]))

        parsed_dates = payload["structured"]["projectDates"]
        self.assertEqual(parsed_dates["startDate"], "2026-06-01")
        self.assertEqual(parsed_dates["endDate"], "2026-09-30")

        project = store._require(project_id)
        self.assertEqual(project["startDate"], "2026-06-01")
        self.assertEqual(project["endDate"], "2026-09-30")
        self.assertEqual(project["deadline"], "2026-09-30")

    def test_bid_dates_ignore_supply_delivery_ranges(self) -> None:
        project_id = self.create_project()
        tender = "\n".join(
            [
                "# 招标文件",
                "项目名称：供货日期不应污染投标日期",
                "供货范围及交货进度",
                "主机设备2026年4月10日前开始供货，截止2026年8月30日前完成全部供货。",
                "安装调试服务期：2026年9月1日至2026年10月30日。",
            ]
        ).encode("utf-8")

        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[("tenderFiles", ("招标文件.md", tender, "text/markdown"))],
        )

        self.assertEqual(response.status_code, 200)
        parsed_dates = response.json()["structured"]["projectDates"]
        self.assertEqual(parsed_dates["startDate"], "")
        self.assertEqual(parsed_dates["endDate"], "")

        project = store._require(project_id)
        self.assertEqual(project["startDate"], "")
        self.assertEqual(project["endDate"], "")
        self.assertEqual(project["deadline"], "")

    def test_bid_dates_parse_bid_submission_and_opening_dates(self) -> None:
        project_id = self.create_project()
        tender = "\n".join(
            [
                "# 招标公告",
                "招标文件获取时间：2026年5月8日至2026年5月15日。",
                "投标文件递交截止时间：2026年6月20日09时30分。",
                "开标时间：2026年6月20日09时30分。",
                "交货周期：2026年10月1日至2027年3月31日。",
            ]
        ).encode("utf-8")

        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[("tenderFiles", ("招标文件.md", tender, "text/markdown"))],
        )

        self.assertEqual(response.status_code, 200)
        parsed_dates = response.json()["structured"]["projectDates"]
        self.assertEqual(parsed_dates["startDate"], "2026-05-08")
        self.assertEqual(parsed_dates["endDate"], "2026-06-20")

        project = store._require(project_id)
        self.assertEqual(project["startDate"], "2026-05-08")
        self.assertEqual(project["endDate"], "2026-06-20")
        self.assertEqual(project["deadline"], "2026-06-20")

    def test_upload_and_parse_multifile_docx_tables_builds_structured_contract(self) -> None:
        project_id = self.create_project()

        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[
                (
                    "tenderFiles",
                    (
                        "评标办法.docx",
                        sample_evaluation_docx_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
                (
                    "tenderFiles",
                    (
                        "技术规范书.docx",
                        sample_technical_spec_docx_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
            ],
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        structured = payload["structured"]

        source_documents = structured["sourceDocuments"]
        self.assertEqual(len(source_documents), 2)
        self.assertEqual(source_documents[0]["role"], "evaluation")
        self.assertEqual(source_documents[1]["role"], "technical_spec")

        scoring = structured["scoringCriteria"]
        self.assertEqual(len(scoring["technical"]), 2)
        self.assertEqual(len(scoring["business"]), 2)
        self.assertEqual(len(scoring["price"]), 1)
        self.assertEqual(len(scoring["lcoe"]), 1)
        self.assertEqual(len(scoring["compliance"]), 1)
        self.assertEqual(scoring["technical"][0]["scoringItem"], "技术方案")
        self.assertEqual(scoring["technical"][0]["score"], "30分")
        self.assertIn("技术承诺函", scoring["technical"][0]["proofRequirement"])
        self.assertEqual(scoring["business"][0]["scoringItem"], "企业业绩")
        self.assertIn("合同", scoring["business"][0]["proofRequirement"])

        for bucket in scoring.values():
            for row in bucket:
                self.assertTrue(row["sourceFile"])
                self.assertTrue(row["sourceDocumentId"])
                self.assertTrue(row["section"])
                self.assertTrue(row["evidence"])
                self.assertTrue(row["evidenceLocation"])

        field_groups = structured["fieldGroups"]
        project_basics = field_groups["projectBasics"]
        self.assertEqual(
            [field["key"] for field in project_basics],
            ["projectName", "tenderNo", "projectUnit", "tenderer", "tenderAgency", "bidDeadline"],
        )
        self.assertTrue(all(field["value"] == "" for field in project_basics))
        self.assertEqual(field_by_key(field_groups["turbineCoreParameters"], "singleCapacity")["value"], "6.25MW")
        self.assertEqual(field_by_key(field_groups["turbineCoreParameters"], "bladeTipClearance")["value"], "20m")
        self.assertIn("认证功率曲线", field_by_key(field_groups["performanceGuarantees"], "powerCurve")["value"])
        self.assertIn("防凝露", field_by_key(field_groups["environmentAdaptation"], "icingCondensation")["value"])

        for group in field_groups.values():
            if isinstance(group, list):
                for field in group:
                    if field["status"] == "found":
                        self.assertTrue(field["sourceFile"])
                        self.assertTrue(field["sourceDocumentId"])
                        self.assertTrue(field["evidence"])
                        self.assertTrue(field["evidenceLocation"])

        presence = structured["requirementPresence"]
        self.assertEqual(presence["topicPlans"]["status"], "present")
        self.assertEqual(presence["supplyScope"]["status"], "present")
        self.assertEqual(presence["assessmentTerms"]["status"], "present")

    def test_s1parse_skill_script_outputs_same_multifile_structured_contract(self) -> None:
        project_dir = Path(self.temp_dir.name) / "skill-script"
        project_dir.mkdir()
        evaluation_path = project_dir / "评标办法.docx"
        technical_path = project_dir / "技术规范书.docx"
        evaluation_path.write_bytes(sample_evaluation_docx_bytes())
        technical_path.write_bytes(sample_technical_spec_docx_bytes())
        evaluation_text = project_dir / "evaluation.txt"
        technical_text = project_dir / "technical.txt"
        evaluation_text.write_text("第三章 评标办法（综合评估法）\n附表2：技术评分标准表\n", encoding="utf-8")
        technical_text.write_text(
            "第二卷 技术规范书\n专题方案：应提供叶片专题、变桨系统专题、主轴专题、齿轮箱专题。\n",
            encoding="utf-8",
        )
        output_path = project_dir / "s1_structured_result.json"
        manifest_path = project_dir / "s1_parse_manifest.json"
        manifest_path.write_text(
            json.dumps(
                {
                    "projectId": "PRJ-SKILL",
                    "structuredResultPath": str(output_path),
                    "documents": [
                        {
                            "id": "DOC-1",
                            "name": "评标办法.docx",
                            "sourcePath": str(evaluation_path),
                            "textPath": str(evaluation_text),
                        },
                        {
                            "id": "DOC-2",
                            "name": "技术规范书.docx",
                            "sourcePath": str(technical_path),
                            "textPath": str(technical_text),
                        },
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        script_path = (
            Path(__file__).resolve().parents[1]
            / "opencode"
            / "skills"
            / "bid-tech-tender-structured-parser"
            / "scripts"
            / "run_from_manifest.py"
        )

        completed = subprocess.run(
            [sys.executable, str(script_path), str(manifest_path)],
            check=True,
            capture_output=True,
            text=True,
        )

        summary = json.loads(completed.stdout)
        self.assertEqual(summary["schemaVersion"], "bid-tender-structured-v1")
        payload = json.loads(output_path.read_text(encoding="utf-8"))
        self.assertEqual(len(payload["structured"]["scoringCriteria"]["technical"]), 2)
        self.assertEqual(len(payload["structured"]["scoringCriteria"]["business"]), 2)
        project_basics = payload["structured"]["fieldGroups"]["projectBasics"]
        self.assertEqual(
            [field["key"] for field in project_basics],
            ["projectName", "tenderNo", "projectUnit", "tenderer", "tenderAgency", "bidDeadline"],
        )
        self.assertTrue(all(field["value"] == "" for field in project_basics))

    def test_parse_result_exposes_fixed_fields_presence_and_appendix_docx_assets(self) -> None:
        project_id = self.create_project()
        tender = "\n".join(
            [
                "# 招标文件",
                "项目名称：华能甘肃100MW风电项目",
                "招标编号：HN-2026-001",
                "招标人：华能集团",
                "管理单位：华能甘肃公司",
                "标段规模：100MW",
                "交货周期：2026年10月1日至2027年3月31日",
                "质保期：5年",
                "技术承诺：投标人应承诺满足全部技术规范。",
                "评分细则：技术方案30分，需提供技术响应表和证明材料；供货保障10分，需提供供货计划。",
                "单机容量：6.25MW",
                "叶轮直径：200m",
                "轮毂高度：120m",
                "叶片最低点距地：20m",
                "塔筒型式：钢混塔筒",
                "箱变型式：华式箱变",
                "安全等级：IEC IIB",
                "空气密度：1.225kg/m3",
                "风速：8.5m/s",
                "湍流强度：0.14",
                "功率曲线：投标人应提供经认证功率曲线。",
                "可利用率：97%",
                "发电量：年上网电量不少于300GWh",
                "涉网性能：满足高低电压穿越要求。",
                "环境适应性：抗低温、抗覆冰防凝露、防潮湿、防雷暴、防风沙、抗高温。",
                "专题方案：应提供叶片专题、变桨系统专题、主轴专题、齿轮箱专题。",
                "供货范围：风力发电机组、塔筒、箱变及备品备件。",
                "考核条款：发电量考核、可利用率考核、功率曲线考核、部件考核、认证考核。",
                "附表1：技术参数响应表",
                "| 序号 | 参数 | 投标响应 |",
                "| --- | --- | --- |",
                "| 1 | 单机容量 | |",
                "| 2 | 叶轮直径 | |",
            ]
        ).encode("utf-8")

        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[("tenderFiles", ("招标文件.md", tender, "text/markdown"))],
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        structured = payload["structured"]
        field_groups = structured["fieldGroups"]

        project_basics = field_groups["projectBasics"]
        self.assertEqual(
            [field["key"] for field in project_basics],
            ["projectName", "tenderNo", "projectUnit", "tenderer", "tenderAgency", "bidDeadline"],
        )
        self.assertTrue(all(field["value"] == "" for field in project_basics))

        turbine = field_groups["turbineCoreParameters"]
        self.assertEqual(field_by_key(turbine, "singleCapacity")["value"], "6.25MW")
        self.assertEqual(field_by_key(turbine, "rotorDiameter")["value"], "200m")
        self.assertEqual(field_by_key(turbine, "hubHeight")["value"], "120m")
        self.assertEqual(field_by_key(turbine, "bladeTipClearance")["value"], "20m")
        self.assertEqual(field_by_key(turbine, "towerType")["value"], "钢混塔筒")
        self.assertEqual(field_by_key(turbine, "boxTransformerType")["value"], "华式箱变")
        self.assertEqual(field_by_key(turbine, "safetyClass")["value"], "IEC IIB")
        self.assertEqual(field_by_key(turbine, "airDensity")["value"], "1.225kg/m3")
        self.assertEqual(field_by_key(turbine, "windSpeed")["value"], "8.5m/s")
        self.assertEqual(field_by_key(turbine, "turbulenceIntensity")["value"], "0.14")

        performance = field_groups["performanceGuarantees"]
        self.assertIn("认证功率曲线", field_by_key(performance, "powerCurve")["value"])
        self.assertEqual(field_by_key(performance, "availability")["value"], "97%")
        self.assertEqual(field_by_key(performance, "generation")["value"], "年上网电量不少于300GWh")
        self.assertIn("电压穿越", field_by_key(performance, "gridPerformance")["value"])

        scoring = field_groups["scoringCriteria"]
        self.assertEqual(scoring[0]["scoringItem"], "技术方案")
        self.assertEqual(scoring[0]["score"], "30分")
        self.assertIn("证明材料", scoring[0]["proofRequirement"])
        self.assertEqual(scoring[1]["scoringItem"], "供货保障")
        self.assertEqual(scoring[1]["score"], "10分")
        self.assertIn("供货计划", scoring[1]["proofRequirement"])

        presence = structured["requirementPresence"]
        self.assertEqual(presence["topicPlans"]["status"], "present")
        self.assertIn("叶片专题", presence["topicPlans"]["summary"])
        self.assertEqual(presence["supplyScope"]["status"], "present")
        self.assertIn("风力发电机组", presence["supplyScope"]["summary"])
        self.assertEqual(presence["assessmentTerms"]["status"], "present")
        self.assertIn("发电量考核", presence["assessmentTerms"]["summary"])

        appendices = structured["appendices"]
        self.assertEqual(len(appendices), 1)
        self.assertEqual(appendices[0]["title"], "附表1：技术参数响应表")
        self.assertEqual(appendices[0]["status"], "generated")
        self.assertEqual(appendices[0]["rowCount"], 3)
        appendix_path = Path(appendices[0]["docxPath"])
        self.assertTrue(appendix_path.exists())
        self.assertIn(str(settings.parsed_dir / project_id / "s1_appendices"), str(appendix_path))
        appendix_doc = Document(str(appendix_path))
        self.assertEqual(len(appendix_doc.tables), 1)
        self.assertEqual(appendix_doc.tables[0].cell(0, 1).text, "参数")
        self.assertEqual(appendix_doc.tables[0].cell(1, 2).text, "")

    def test_parse_docx_appendix_table_generates_workspace_docx(self) -> None:
        project_id = self.create_project()
        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[
                (
                    "tenderFiles",
                    (
                        "含附表招标文件.docx",
                        build_appendix_docx_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
            ],
        )

        self.assertEqual(response.status_code, 200)
        appendices = response.json()["structured"]["appendices"]
        self.assertEqual(len(appendices), 1)
        self.assertEqual(appendices[0]["status"], "generated")
        self.assertEqual(appendices[0]["rowCount"], 3)
        appendix_doc = Document(appendices[0]["docxPath"])
        self.assertEqual(appendix_doc.tables[0].cell(0, 1).text, "设备名称")
        self.assertEqual(appendix_doc.tables[0].cell(1, 2).text, "")

    def test_parse_pdf_text_appendix_generates_workspace_docx(self) -> None:
        project_id = self.create_project()
        tender_text = "\n".join(
            [
                "招标文件技术规范",
                "附表G.3.1 关键设计方法 ..................................................206",
                "附表G.2.4 场址投标机组疲劳载荷与认证机组疲劳载荷对比（m=10）206",
                "附表G.5.2 招标项目场址机组等效疲劳载荷与认证机组等效疲劳载荷对",
                "比 ................................................................................................................. 209",
                "附表G.3.2 塔筒极限强度设计安全余量",
                "序号 截面位置 安全余量 投标响应",
                "1 塔底 不低于5% ",
                "2 门洞 不低于5% ",
                "附表G.3.3 塔筒屈曲稳定性安全余量",
                "序号 截面位置 安全余量 投标响应",
                "1 塔底 不低于5% ",
            ]
        )

        def fake_extract_pdf_text(_path: Path):
            return tender_text, {"pageCount": 3, "warnings": [], "requiresOcr": False}

        with patch("app.services.parsing.settings.business_pdf_parse_engine", "lightweight", create=True), patch(
            "app.services.parsing.extract_pdf_text",
            side_effect=fake_extract_pdf_text,
        ):
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[("tenderFiles", ("招标文件-技术规范.pdf", b"%PDF-1.4\n", "application/pdf"))],
            )

        self.assertEqual(response.status_code, 200)
        appendices = response.json()["structured"]["appendices"]
        self.assertEqual([item["title"] for item in appendices], ["附表G.3.2 塔筒极限强度设计安全余量", "附表G.3.3 塔筒屈曲稳定性安全余量"])
        self.assertEqual([item["rowCount"] for item in appendices], [3, 2])
        appendix_doc = Document(appendices[0]["docxPath"])
        self.assertEqual(len(appendix_doc.tables), 1)
        self.assertEqual(appendix_doc.tables[0].cell(0, 1).text, "截面位置")
        self.assertEqual(appendix_doc.tables[0].cell(1, 3).text, "")

    def test_parse_pdf_document_nav_appendix_preserves_table_blocks(self) -> None:
        project_id = self.create_project()
        nav_payload = {
            "schemaVersion": "business-document-nav-v1",
            "sourceEngine": "docling",
            "documents": [{"id": "DOC-1", "sourcePath": "technical.pdf"}],
            "pages": [{"pageNo": page_no, "textDensity": 0.8} for page_no in range(1, 556)],
            "blocks": [
                {"id": "DOC-1:B000001", "type": "paragraph", "text": "招标文件技术规范", "pageNo": 1},
                {"id": "DOC-1:B000002", "type": "heading", "text": "附表A.1 投标机型总方案信息表", "pageNo": 1, "bbox": [50, 80, 500, 110]},
                {"id": "DOC-1:B000003", "type": "paragraph", "text": "投标人应按下表填写。", "pageNo": 1, "bbox": [50, 120, 500, 145]},
                {"id": "DOC-1:B000004", "type": "table", "text": "", "tableId": "T1", "pageNo": 1, "bbox": [50, 150, 540, 260]},
                {"id": "DOC-1:B000005", "type": "heading", "text": "附表A.2 机型配置品牌表", "pageNo": 2, "bbox": [50, 80, 500, 110]},
                {"id": "DOC-1:B000006", "type": "table", "text": "", "tableId": "T2", "pageNo": 2, "bbox": [50, 150, 540, 230]},
            ],
            "tables": [
                {
                    "id": "T1",
                    "rows": [
                        ["编号", "项目", "", "备注"],
                        ["12", "塔筒重量（t）", "TG1", "说明"],
                        ["", "", "TG2", ""],
                    ],
                    "cells": [
                        {"rowStart": 0, "rowEnd": 1, "colStart": 0, "colEnd": 1, "rowSpan": 1, "colSpan": 1, "text": "编号", "bbox": []},
                        {"rowStart": 0, "rowEnd": 1, "colStart": 1, "colEnd": 3, "rowSpan": 1, "colSpan": 2, "text": "项目", "bbox": []},
                        {"rowStart": 0, "rowEnd": 1, "colStart": 3, "colEnd": 4, "rowSpan": 1, "colSpan": 1, "text": "备注", "bbox": []},
                        {"rowStart": 1, "rowEnd": 3, "colStart": 0, "colEnd": 1, "rowSpan": 2, "colSpan": 1, "text": "12", "bbox": []},
                        {"rowStart": 1, "rowEnd": 3, "colStart": 1, "colEnd": 2, "rowSpan": 2, "colSpan": 1, "text": "塔筒重量（t）", "bbox": []},
                        {"rowStart": 1, "rowEnd": 2, "colStart": 2, "colEnd": 3, "rowSpan": 1, "colSpan": 1, "text": "TG1", "bbox": []},
                        {"rowStart": 1, "rowEnd": 3, "colStart": 3, "colEnd": 4, "rowSpan": 2, "colSpan": 1, "text": "说明", "bbox": []},
                        {"rowStart": 2, "rowEnd": 3, "colStart": 2, "colEnd": 3, "rowSpan": 1, "colSpan": 1, "text": "TG2", "bbox": []},
                    ],
                    "sourceEngine": "docling",
                },
                {
                    "id": "T2",
                    "rows": [
                        ["序号", "部件", "品牌"],
                        ["1", "叶片", ""],
                    ],
                    "sourceEngine": "docling",
                },
            ],
            "images": [],
            "evidence": [],
            "quality": {"engine": "docling", "status": "completed", "fallbackUsed": False},
        }

        def fake_parse_pdf(self, *, project_id: str, document: dict, output_dir: Path):
            nav_path = output_dir / "DOC-1_document_nav.json"
            quality_path = output_dir / "document_parse" / "docling" / "DOC-1" / "parse_quality.json"
            quality_path.parent.mkdir(parents=True, exist_ok=True)
            nav_path.write_text(json.dumps(nav_payload, ensure_ascii=False), encoding="utf-8")
            quality_path.write_text(
                json.dumps({"engine": "docling", "status": "completed", "fallbackUsed": False}, ensure_ascii=False),
                encoding="utf-8",
            )
            return {
                "documentParseEngine": "docling",
                "status": "completed",
                "documentNavPath": str(nav_path),
                "parseQualityPath": str(quality_path),
            }

        with patch("app.services.parsing.settings.business_pdf_parse_engine", "docling", create=True), patch(
            "app.services.parsing.DoclingParseEngine.parse_pdf",
            new=fake_parse_pdf,
        ), patch(
            "app.services.parsing.extract_pdf_text",
            side_effect=AssertionError("技术标 DocumentNav 已包含全文，不应重复扫描 PDF"),
        ):
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[("tenderFiles", ("招标文件-技术规范.pdf", b"%PDF-1.4\n", "application/pdf"))],
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["sourceFiles"][0]["pageCount"], 555)
        self.assertEqual(response.json()["sourceFiles"][0]["textLength"], len(parsing_service.nav_to_text(nav_payload)))
        appendices = response.json()["structured"]["appendices"]
        self.assertEqual(
            [item["title"] for item in appendices],
            ["附表A.1 投标机型总方案信息表", "附表A.2 机型配置品牌表"],
        )
        self.assertEqual([item["extractionMode"] for item in appendices], ["pdf_document_nav_slice", "pdf_document_nav_slice"])
        self.assertEqual([item["rowCount"] for item in appendices], [3, 2])
        self.assertEqual(appendices[0]["sourceEngine"], "docling")
        self.assertEqual(appendices[0]["sourceStart"], "P1")
        self.assertEqual(appendices[0]["contentBlocks"][1]["cells"][1]["colSpan"], 2)
        self.assertEqual(appendices[0]["contentBlocks"][1]["cells"][3]["rowSpan"], 2)
        appendix_doc = Document(appendices[0]["docxPath"])
        self.assertEqual(len(appendix_doc.tables), 1)
        self.assertEqual(appendix_doc.tables[0].cell(0, 1).text, "项目")
        self.assertEqual(appendix_doc.tables[0].cell(1, 1).text, "塔筒重量（t）")
        with zipfile.ZipFile(appendices[0]["docxPath"]) as docx_zip:
            document_xml = docx_zip.read("word/document.xml").decode("utf-8")
        self.assertIn("gridSpan", document_xml)
        self.assertIn("vMerge", document_xml)

    def test_technical_pdf_docling_failure_does_not_use_local_table_nav_fallback(self) -> None:
        project_id = self.create_project()
        pdf_path = settings.uploads_dir / project_id / "technical.pdf"
        pdf_path.parent.mkdir(parents=True, exist_ok=True)
        pdf_path.write_bytes(b"%PDF-1.4\n")

        def fake_run_docling_conversion(_pdf_path: Path, _output_dir: Path):
            raise RuntimeError("rapidocr model missing")

        with patch("app.services.parsing.settings.s1_parse_opencode_enabled", False), patch(
            "app.services.parsing.settings.business_pdf_parse_engine",
            "docling",
            create=True,
        ), patch(
            "app.services.parsing.settings.business_pdf_engine_fallback",
            "none",
            create=True,
        ), patch(
            "app.services.docling_engine.run_docling_conversion",
            side_effect=fake_run_docling_conversion,
        ), patch(
            "app.services.docling_engine.run_docling_local_text_layer_conversion",
            side_effect=AssertionError("technical PDF must not use local-text-layer fallback"),
        ), patch(
            "app.services.parsing.extract_pdf_text",
            return_value=("技术标全文正文", {"pageCount": 1, "warnings": [], "requiresOcr": False}),
        ):
            _summary, storage = parsing_service.parse_tender_documents(
                project_id,
                [
                    {
                        "id": "DOC-1",
                        "name": "招标文件-技术规范.pdf",
                        "path": str(pdf_path),
                        "content_type": "application/pdf",
                    }
                ],
                bid_type="技术标",
            )

        self.assertEqual(storage["structured"]["appendices"], [])
        self.assertEqual(storage["documents"][0]["documentParseStatus"], "failed")
        self.assertIn("rapidocr model missing", storage["documents"][0]["fallbackReason"])

    def test_parse_docx_appendix_preserves_cell_merges_via_source_slicing(self) -> None:
        """Ensure the appendix docx generated from a docx-source RFP keeps the
        original <w:vMerge>/<w:gridSpan> structures rather than being rebuilt
        from a flattened rows list. This is the format-preservation guarantee."""
        import zipfile

        project_id = self.create_project()
        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[
                (
                    "tenderFiles",
                    (
                        "含合并附表招标文件.docx",
                        build_appendix_with_merges_docx_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
            ],
        )

        self.assertEqual(response.status_code, 200)
        appendices = response.json()["structured"]["appendices"]
        self.assertEqual(len(appendices), 1)

        appendix_path = Path(appendices[0]["docxPath"])
        self.assertTrue(appendix_path.exists(), f"appendix docx missing: {appendix_path}")

        # 1. The generated appendix table must still carry cell-merge XML markers.
        with zipfile.ZipFile(appendix_path) as zf:
            doc_xml = zf.read("word/document.xml").decode("utf-8")
        self.assertIn(
            "gridSpan",
            doc_xml,
            "horizontal merge (<w:gridSpan>) was lost: the table was rebuilt from rows instead of sliced from source",
        )
        self.assertIn(
            "vMerge",
            doc_xml,
            "vertical merge (<w:vMerge>) was lost: the table was rebuilt from rows instead of sliced from source",
        )

        # 2. The appendix should ship its own styles.xml (proof we copied the source
        #    docx, not regenerated from scratch). A regenerated python-docx file has
        #    a default styles.xml that is significantly smaller than what the source
        #    carries because the source was authored with full Word style definitions.
        with zipfile.ZipFile(appendix_path) as zf:
            self.assertIn("word/styles.xml", zf.namelist())

        # 3. The body should only retain the appendix heading + the table; everything
        #    else from the source RFP body must be removed.
        appendix_doc = Document(str(appendix_path))
        self.assertEqual(len(appendix_doc.tables), 1, "exactly one table expected in the sliced appendix")
        non_empty_paragraphs = [
            paragraph.text.strip()
            for paragraph in appendix_doc.paragraphs
            if paragraph.text.strip()
        ]
        self.assertEqual(
            non_empty_paragraphs,
            ["附表D.1 标准及风电场空气密度功率曲线"],
            "only the appendix heading paragraph should remain in body; got %r" % (non_empty_paragraphs,),
        )

    def test_technical_docx_plain_tables_skip_source_slice_but_merged_tables_keep_it(self) -> None:
        plain_path = settings.uploads_dir / "plain-appendix.docx"
        merged_path = settings.uploads_dir / "merged-appendix.docx"
        plain_path.write_bytes(
            build_docx_blocks_bytes(
                "附表A.1 投标机型总方案信息表",
                [
                    ["序号", "项目", "投标响应"],
                    ["1", "机型总方案", ""],
                ],
            )
        )
        merged_path.write_bytes(build_appendix_with_merges_docx_bytes())

        captured: list[dict] = []

        def fake_materialize(project_id: str, appendix: dict, *, profile):
            captured.append(appendix)
            return dict(appendix)

        with patch("app.services.parsing.materialize_appendix_docx", side_effect=fake_materialize):
            parsing_service._extract_docx_appendices(
                "PRJ-SLICE-POLICY",
                [
                    {"name": "plain-appendix.docx", "sourcePath": str(plain_path)},
                    {"name": "merged-appendix.docx", "sourcePath": str(merged_path)},
                ],
            )

        self.assertEqual(len(captured), 2)
        self.assertNotIn("_slice", captured[0])
        self.assertIn("_slice", captured[1])

    def test_large_docx_slice_stores_document_xml_without_expensive_deflate(self) -> None:
        source_path = settings.uploads_dir / "large-slice-source.docx"
        target_path = settings.parsed_dir / "large-slice-target.docx"
        source_path.parent.mkdir(parents=True, exist_ok=True)
        target_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        doc.add_paragraph("附表A.1 大型测试表")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "序号"
        table.cell(0, 1).text = "内容"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "大段内容" * 20
        doc.save(source_path)

        source_state = parsing_service._build_appendix_slice_state(source_path)
        with patch("app.services.parsing.DOCX_SLICE_STORED_XML_THRESHOLD_BYTES", 1, create=True):
            sliced = parsing_service._slice_appendix_from_source(
                source_path,
                target_path,
                0,
                1,
                source_state=source_state,
            )

        self.assertTrue(sliced)
        with zipfile.ZipFile(target_path) as zf:
            self.assertEqual(zf.getinfo("word/document.xml").compress_type, zipfile.ZIP_STORED)

    def test_large_rebuilt_appendix_table_writes_within_budget(self) -> None:
        target_path = settings.parsed_dir / "large-rebuilt-appendix-table.docx"
        rows = [
            [f"R{row_index}C{column_index} " * 2 for column_index in range(5)]
            for row_index in range(160)
        ]

        started_at = time.perf_counter()
        parsing_service._write_appendix_docx(
            target_path,
            "技术附表I 技术条款偏差表",
            rows,
            [{"type": "table", "rows": rows}],
        )
        elapsed = time.perf_counter() - started_at

        self.assertLess(elapsed, 4.0)
        doc = Document(str(target_path))
        self.assertEqual(len(doc.tables), 1)
        self.assertEqual(len(doc.tables[0].rows), len(rows))
        self.assertEqual(doc.tables[0].cell(159, 4).text, rows[159][4])

    def test_parse_docx_appendix_slicing_handles_large_body_within_budget(self) -> None:
        """Regression guard: a real RFP body can have thousands of paragraphs and
        dozens of appendices. The slice path must stay O(N) overall, not O(N^2)
        per appendix. Earlier the implementation called ``body.remove(child)``
        per non-keeper, which froze the request thread for huge documents."""

        import time

        project_id = self.create_project()

        # Build an RFP-shaped docx that's representative of a real bid file:
        # ~5000 narrative paragraphs + 80 back-to-back appendices. Without
        # source-tree caching, each appendix would re-parse a multi-MB docx
        # via python-docx (several seconds per appendix => minutes total),
        # which is what froze the upload thread in production.
        narrative_blocks: list[str | list[list[str]]] = [
            f"第{i // 50 + 1}章 章节{i + 1} 这一段是正文铺垫文本，内容足够长以便撑出体积。"
            for i in range(5000)
        ]
        appendix_blocks: list[str | list[list[str]]] = []
        for i in range(1, 81):
            appendix_blocks.append(f"附表X.{i} 测试附表{i}")
            appendix_blocks.append(
                [
                    ["序号", "字段", "值"],
                    ["1", f"项目{i}", ""],
                ]
            )

        rfp_bytes = build_docx_blocks_bytes(*narrative_blocks, *appendix_blocks)

        deadline = time.monotonic()
        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[
                (
                    "tenderFiles",
                    (
                        "大体量招标文件.docx",
                        rfp_bytes,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
            ],
        )
        elapsed = time.monotonic() - deadline

        self.assertEqual(response.status_code, 200)
        appendices = response.json()["structured"]["appendices"]
        self.assertEqual(
            len(appendices),
            80,
            f"expected all 80 appendices to be discovered; got {len(appendices)}",
        )
        # 60s is a generous ceiling — the goal is to fail loudly if the
        # algorithm regresses (e.g. re-parsing the 20+ MB source per appendix
        # the way the first cut did, which took minutes per upload).
        self.assertLess(
            elapsed,
            60.0,
            f"appendix slicing took {elapsed:.1f}s for 5080-block / 80-appendix docx; "
            "suspect a regression in _slice_appendix_from_source caching",
        )

    def test_parse_docx_appendices_ignores_toc_titles_with_page_numbers(self) -> None:
        project_id = self.create_project()
        file_bytes = build_docx_blocks_bytes(
            "目录",
            "附表A.1 投标机型总方案信息表169",
            "附表B.1.2 机型配置品牌表1173",
            "附表B.9.1 双馈型风电机组179",
            "正文",
            "附表A.1 投标机型总方案信息表",
            [
                ["序号", "项目", "投标响应"],
                ["1", "总方案", ""],
            ],
            "附表B.1.2 机型配置品牌表1",
            [
                ["序号", "部件", "品牌"],
                ["1", "叶片", ""],
            ],
        )

        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[
                (
                    "tenderFiles",
                    (
                        "含目录附表招标文件.docx",
                        file_bytes,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
            ],
        )

        self.assertEqual(response.status_code, 200)
        appendices = response.json()["structured"]["appendices"]
        self.assertEqual(
            [appendix["title"] for appendix in appendices],
            [
                "附表A.1 投标机型总方案信息表",
                "附表B.1.2 机型配置品牌表1",
            ],
        )
        self.assertEqual([appendix["id"] for appendix in appendices], ["APPX-0001", "APPX-0002"])
        self.assertEqual([appendix["rowCount"] for appendix in appendices], [2, 2])

    def test_markdown_appendix_heading_without_table_generates_workspace_docx(self) -> None:
        project_id = self.create_project()
        tender = "\n".join(
            [
                "# 招标文件",
                "项目名称：附表空表测试项目",
                "附表2：投标偏离表",
                "请投标人按招标文件要求填写。",
            ]
        ).encode("utf-8")

        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[("tenderFiles", ("招标文件.md", tender, "text/markdown"))],
        )

        self.assertEqual(response.status_code, 200)
        appendices = response.json()["structured"]["appendices"]
        self.assertEqual(len(appendices), 1)
        self.assertEqual(appendices[0]["status"], "generated")
        self.assertEqual(appendices[0]["rowCount"], 0)
        appendix_path = Path(appendices[0]["docxPath"])
        self.assertTrue(appendix_path.exists())
        self.assertEqual(appendices[0]["workspacePath"], f"s1_appendices/{appendix_path.name}")
        appendix_doc = Document(str(appendix_path))
        self.assertEqual(len(appendix_doc.tables), 0)
        self.assertIn("附表2：投标偏离表", [paragraph.text for paragraph in appendix_doc.paragraphs])

    def test_participating_promotes_parse_json_and_appendices_to_workspace(self) -> None:
        project_id = self.create_project()
        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[
                (
                    "tenderFiles",
                    (
                        "含附表招标文件.docx",
                        build_appendix_docx_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
            ],
        )
        self.assertEqual(response.status_code, 200)
        temp_appendix_path = Path(response.json()["structured"]["appendices"][0]["docxPath"])
        self.assertIn(str(settings.parsed_dir / project_id / "s1_appendices"), str(temp_appendix_path))
        temp_project_dir = settings.parsed_dir / project_id
        self.assertTrue(temp_project_dir.exists())
        stale_path = settings.documents_dir / project_id / "technical-workspace" / "appendices" / "stale.docx"
        stale_path.parent.mkdir(parents=True, exist_ok=True)
        stale_path.write_bytes(b"old")

        updated = self.client.put(
            self.project_url(project_id),
            json={
                "name": "参与后归档项目",
                "customerName": "测试业主",
                "manager": "项目经理",
                "startDate": "2026-01-01",
                "endDate": "2026-02-01",
                "bidType": "技术标",
                "reviewDecision": "participate",
            },
        )

        self.assertEqual(updated.status_code, 200)
        self.assertFalse(stale_path.exists())
        workspace_parse_dir = settings.documents_dir / project_id / "technical-workspace" / "parse"
        workspace_appendix_dir = settings.documents_dir / project_id / "technical-workspace" / "appendices"
        self.assertTrue((workspace_parse_dir / "s1_structured_result.json").exists())
        self.assertTrue((workspace_parse_dir / "parse-result.workspace.json").exists())
        workspace_appendices = sorted(workspace_appendix_dir.glob("*.docx"))
        self.assertEqual(len(workspace_appendices), 1)
        self.assertFalse(temp_project_dir.exists())

        promoted_payload = self.client.get(self.parse_results_url(project_id))
        self.assertEqual(promoted_payload.status_code, 200)
        appendix = promoted_payload.json()["structured"]["appendices"][0]
        self.assertIn(str(workspace_appendix_dir), appendix["docxPath"])
        self.assertEqual(appendix["workspacePath"], f"technical-workspace/appendices/{workspace_appendices[0].name}")

        project = store._require(project_id)
        parse_storage = project["parse_storage"]
        self.assertEqual(Path(parse_storage["projectDir"]), settings.documents_dir / project_id / "technical-workspace")
        self.assertEqual(Path(parse_storage["parseDir"]), workspace_parse_dir)
        self.assertEqual(Path(parse_storage["combinedTextPath"]), workspace_parse_dir / "combined.txt")
        self.assertEqual(Path(parse_storage["structuredResultPath"]), workspace_parse_dir / "s1_structured_result.json")
        self.assertEqual(Path(parse_storage["manifestPath"]), workspace_parse_dir / "manifest.json")
        self.assertEqual(Path(parse_storage["skillManifestPath"]), workspace_parse_dir / "s1_parse_manifest.json")
        self.assertTrue(all(str(workspace_parse_dir) in item["textPath"] for item in parse_storage["documents"]))

        preview = self.client.get(self.parse_results_url(project_id, "/appendices/APPX-0001/preview"))
        self.assertEqual(preview.status_code, 200)
        self.assertIn(str(workspace_appendix_dir), preview.json()["docxPath"])
        self.assertFalse(temp_project_dir.exists())

    def test_participating_project_reparse_refreshes_workspace_parse_artifacts(self) -> None:
        project_id = self.create_project()

        def fake_parse_with_appendix(title: str):
            def fake_parse(project_id_arg, tender_files, *, bid_type, progress_callback=None, cancel_check=None):
                parse_dir = settings.parsed_dir / project_id_arg
                parse_dir.mkdir(parents=True, exist_ok=True)
                combined_path = parse_dir / "combined.txt"
                manifest_path = parse_dir / "manifest.json"
                skill_manifest_path = parse_dir / "s1_parse_manifest.json"
                structured_path = parse_dir / "s1_structured_result.json"
                rows = [["序号", "项目", "投标响应"], ["1", title, ""]]
                structured = {
                    "schemaVersion": "bid-tender-structured-v1",
                    "projectDates": {"startDate": "", "endDate": ""},
                    "appendices": [
                        {
                            "id": "APPX-0001",
                            "title": title,
                            "status": "generated",
                            "sourceFile": tender_files[0]["name"],
                            "rows": rows,
                            "contentBlocks": [{"type": "table", "rows": rows}],
                            "rowCount": len(rows),
                            "docxPath": "",
                            "extractionMode": "pdf_document_nav_slice",
                        }
                    ],
                }
                combined_path.write_text(title, encoding="utf-8")
                manifest = {
                    "documents": [
                        {
                            "id": "TEN-1",
                            "name": tender_files[0]["name"],
                            "textPath": str(combined_path),
                            "pageCount": 1,
                            "textLength": len(title),
                        }
                    ]
                }
                manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
                skill_manifest_path.write_text(json.dumps({"structuredResultPath": str(structured_path)}), encoding="utf-8")
                structured_path.write_text(
                    json.dumps({"items": [], "structured": structured}, ensure_ascii=False, indent=2),
                    encoding="utf-8",
                )
                return (
                    {"fileCount": 1, "extractedCount": 0, "textLength": len(title), "textPreview": title, "warnings": []},
                    {
                        "documents": manifest["documents"],
                        "items": [],
                        "structured": structured,
                        "combinedTextPath": str(combined_path),
                        "manifestPath": str(manifest_path),
                        "skillManifestPath": str(skill_manifest_path),
                        "structuredResultPath": str(structured_path),
                        "projectUpdates": {},
                    },
                )

            return fake_parse

        with patch("app.services.bid_parse_service.parse_tender_documents", side_effect=fake_parse_with_appendix("附表1 初版")):
            first = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[("tenderFiles", ("first.pdf", b"%PDF-1.4\n", "application/pdf"))],
            )
        self.assertEqual(first.status_code, 200)

        updated = self.client.put(
            self.project_url(project_id),
            json={
                "name": "参与后重复解析项目",
                "customerName": "测试业主",
                "bidType": "技术标",
                "reviewDecision": "participate",
            },
        )
        self.assertEqual(updated.status_code, 200)

        workspace_parse_result = settings.documents_dir / project_id / "technical-workspace" / "parse" / "parse-result.workspace.json"
        self.assertTrue(workspace_parse_result.exists())
        self.assertEqual(
            json.loads(workspace_parse_result.read_text(encoding="utf-8"))["structured"]["appendices"][0]["title"],
            "附表1 初版",
        )

        with patch("app.services.bid_parse_service.parse_tender_documents", side_effect=fake_parse_with_appendix("附表2 重解析新版")):
            second = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[("tenderFiles", ("second.pdf", b"%PDF-1.4\n", "application/pdf"))],
            )
        self.assertEqual(second.status_code, 200)

        refreshed = json.loads(workspace_parse_result.read_text(encoding="utf-8"))
        appendix = refreshed["structured"]["appendices"][0]
        self.assertEqual(appendix["title"], "附表2 重解析新版")
        self.assertEqual(appendix["extractionMode"], "pdf_document_nav_slice")
        self.assertIn("technical-workspace/appendices/", appendix["workspacePath"])
        self.assertTrue(Path(appendix["docxPath"]).exists())

    def test_business_bid_parse_returns_business_contract_without_technical_groups(self) -> None:
        project_id = self.create_business_project()
        tender = "\n".join(
            [
                "# 商务招标文件",
                "项目名称：华能甘肃100MW风电项目",
                "招标编号：HN-BUS-2026-001",
                "招标人：华能集团",
                "交货周期：2026年10月1日至2027年3月31日",
                "质保期：5年",
                "附表3：商务评分标准表",
                "| 序号 | 评分项 | 分值 | 得分点 | 证明材料要求 |",
                "| --- | --- | --- | --- | --- |",
                "| 1 | 企业业绩 | 20分 | 近三年同类项目业绩满足要求得满分。 | 提供合同或中标通知书。 |",
                "投标函：按招标文件格式填写并签字盖章。",
                "法定代表人授权委托书：须加盖公章。",
                "商务偏差表：投标人应逐项响应。",
                "投标保证金：须提供电汇回单或保函。",
                "投标人证明其是合格投标人并有资格履行合同的证明文件。",
                "投标人不得存在下列情形之一。",
                "投标人需要说明的其他内容。",
            ]
        ).encode("utf-8")

        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[("tenderFiles", ("商务招标文件.md", tender, "text/markdown"))],
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        structured = payload["structured"]
        self.assertEqual(structured["schemaVersion"], "bid-business-tender-structured-v1")
        self.assertEqual(payload["summary"]["targetSkill"], "bid-business-tender-structured-parser")
        self.assertIn("commitmentLetterCount", payload["summary"])
        fact_by_key = {field["fieldKey"]: field for field in structured["projectFactFields"]}
        self.assertEqual(fact_by_key["projectName"]["value"], "华能甘肃100MW风电项目")
        self.assertEqual(fact_by_key["tenderNo"]["value"], "HN-BUS-2026-001")
        self.assertEqual(fact_by_key["tenderer"]["value"], "华能集团")

        field_groups = structured["fieldGroups"]
        self.assertIn("projectBasics", field_groups)
        self.assertIn("businessResponse", field_groups)
        self.assertIn("qualificationSupport", field_groups)
        self.assertIn("qualificationRequirements", field_groups)
        self.assertIn("bidderInstructions", field_groups)
        self.assertIn("commercialRejectionClauses", field_groups)
        self.assertIn("commitmentRequirements", field_groups)
        self.assertIn("tenderAgency", {field["key"] for field in field_groups["projectBasics"]})
        self.assertIn("bidDeadline", {field["key"] for field in field_groups["projectBasics"]})
        self.assertNotIn("turbineCoreParameters", field_groups)
        self.assertNotIn("performanceGuarantees", field_groups)
        self.assertNotIn("environmentAdaptation", field_groups)

        scoring = structured["scoringCriteria"]
        self.assertEqual(set(scoring.keys()), {"business", "price", "compliance"})
        self.assertGreaterEqual(len(scoring["business"]), 1)
        self.assertEqual(scoring["business"][0]["scoringItem"], "企业业绩")

        self.assertEqual(structured["requirementPresence"]["bidSecurity"]["status"], "present")
        self.assertEqual(structured["requirementPresence"]["qualificationDocuments"]["status"], "present")
        self.assertEqual(structured["requirementPresence"]["disqualificationClauses"]["status"], "present")

        commitment_fields = field_groups["commitmentRequirements"]
        self.assertEqual(field_by_key(commitment_fields, "generalCommitmentCount")["value"], "1")
        self.assertEqual(field_by_key(commitment_fields, "generatedCommitmentCount")["value"], "1")
        self.assertEqual(field_by_key(commitment_fields, "pendingCommitmentCount")["value"], "0")
        self.assertEqual(field_by_key(commitment_fields, "disqualificationCommitmentRequired")["status"], "found")

        letters = structured["commitmentLetters"]
        self.assertEqual(len(letters), 1)
        self.assertEqual(letters[0]["title"], "投标人不存在下列情形之一承诺函")
        self.assertEqual(letters[0]["commitmentType"], "disqualification")
        self.assertEqual(letters[0]["status"], "generated")
        self.assertTrue(letters[0]["docxPath"])
        self.assertTrue(Path(letters[0]["docxPath"]).exists())

        preview = self.client.get(
            self.parse_results_url(project_id, f"/commitment-letters/{letters[0]['id']}/preview")
        )
        self.assertEqual(preview.status_code, 200)
        preview_payload = preview.json()
        self.assertEqual(preview_payload["id"], letters[0]["id"])
        self.assertIn("/parse-results/commitment-letters/", preview_payload["onlyoffice"]["browserFileUrl"])

        approved = self.client.post(
            self.parse_results_url(project_id, f"/commitment-letters/{letters[0]['id']}/approve"),
            json={"approved": True},
        )
        self.assertEqual(approved.status_code, 200)
        approved_letter = approved.json()["letter"]
        self.assertEqual(approved_letter["assetReviewStatus"], "approved")
        self.assertEqual(approved_letter["assetMaterialFolder"], "资格审查与商务响应成册")

    def test_business_bid_deadline_preserves_minutes_and_ignores_opening_time(self) -> None:
        project_id = self.create_business_project()
        tender = "\n".join(
            [
                "# 商务招标文件",
                "项目名称：京能风电设备采购项目",
                "招标编号：ZBA272600801",
                "招标人：山西漳山发电有限责任公司",
                "投标文件的递交",
                "投标截止时间",
                "递交截止时间：2026年03月18日 09时30分",
                "开标时间及地点",
                "开标时间：2026年03月18日 10时30分",
                "交货周期：2026年10月1日至2027年3月31日。",
            ]
        ).encode("utf-8")

        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[("tenderFiles", ("商务招标文件.md", tender, "text/markdown"))],
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        project_basics = payload["structured"]["fieldGroups"]["projectBasics"]
        bid_deadline = field_by_key(project_basics, "bidDeadline")
        self.assertEqual(bid_deadline["value"], "2026-03-18 09:30")
        self.assertIn("递交截止时间", bid_deadline["evidence"])
        self.assertNotIn("开标时间", bid_deadline["evidence"])
        self.assertEqual(payload["structured"]["projectDates"]["endDate"], "2026-03-18 09:30")
        fact_by_key = {field["fieldKey"]: field for field in payload["structured"]["projectFactFields"]}
        self.assertEqual(fact_by_key["bidDeadline"]["value"], "2026-03-18 09:30")

        project = store._require(project_id)
        self.assertEqual(project["endDate"], "2026-03-18 09:30")
        self.assertEqual(project["deadline"], "2026-03-18 09:30")

    def test_business_bid_deadline_normalizer_prefers_datetime_candidate(self) -> None:
        self.assertEqual(
            parsing_service._normalize_bid_deadline(
                "2026-03-18 4.2.1 | 投标截止时间 | 2026年03月18日 09时30分"
            ),
            "2026-03-18 09:30",
        )

    def test_business_bid_deadline_requires_real_date_candidate(self) -> None:
        project_id = self.create_business_project()
        tender = "\n".join(
            [
                "# 商务招标文件",
                "项目名称：华能风电项目",
                "招标编号：HNZB2025-12-1-382",
                "投标保证金到账截止时间为投标截止时间；请投标人确保投标截止时间前到账。",
                "5.1 递交截止时间：2026年1月26日15时00分",
                "开标时间：2026年1月26日16时00分",
            ]
        ).encode("utf-8")

        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[("tenderFiles", ("商务招标文件.md", tender, "text/markdown"))],
        )

        self.assertEqual(response.status_code, 200)
        project_basics = response.json()["structured"]["fieldGroups"]["projectBasics"]
        bid_deadline = field_by_key(project_basics, "bidDeadline")
        self.assertEqual(bid_deadline["value"], "2026-01-26 15:00")
        self.assertIn("递交截止时间", bid_deadline["evidence"])
        self.assertNotIn("投标保证金到账截止时间", bid_deadline["evidence"])

    def test_business_local_transform_discards_preface_reference_project_basics(self) -> None:
        payload = {
            "items": [
                {
                    "fieldKey": "projectName",
                    "title": "招标项目名称",
                    "keyEntity": "招标项目名称",
                    "value": "见招标公告",
                    "sourceFile": "商务招标文件.md",
                    "sourceDocumentId": "DOC-1",
                    "section": "投标人须知前附表",
                    "evidence": "1.1.4 | 招标项目名称 | 见招标公告",
                    "evidenceLocation": "L12",
                    "confidence": 0.9,
                },
                {
                    "fieldKey": "projectName",
                    "title": "招标项目名称",
                    "keyEntity": "招标项目名称",
                    "value": "公告真实项目",
                    "sourceFile": "商务招标文件.md",
                    "sourceDocumentId": "DOC-1",
                    "section": "第一章 招标公告",
                    "evidence": "招标项目名称：公告真实项目",
                    "evidenceLocation": "L4",
                    "confidence": 0.86,
                },
                {
                    "fieldKey": "tenderer",
                    "title": "招标人",
                    "keyEntity": "招标人",
                    "value": "见招标公告",
                    "sourceFile": "商务招标文件.md",
                    "sourceDocumentId": "DOC-1",
                    "section": "投标人须知前附表",
                    "evidence": "1.1.2 | 招标人 | 见招标公告",
                    "evidenceLocation": "L10",
                    "confidence": 0.9,
                },
                {
                    "fieldKey": "tenderer",
                    "title": "招标人",
                    "keyEntity": "招标人",
                    "value": "公告真实招标单位",
                    "sourceFile": "商务招标文件.md",
                    "sourceDocumentId": "DOC-1",
                    "section": "第一章 招标公告",
                    "evidence": "招标人：公告真实招标单位",
                    "evidenceLocation": "L5",
                    "confidence": 0.86,
                },
            ],
            "structured": {"projectDates": {"startDate": "", "endDate": ""}, "appendices": []},
        }
        result = parsing_service._transform_to_business_contract(
            "PRJ-LOCAL-REFERENCE",
            payload,
            profile=parsing_service.BUSINESS_PARSE_PROFILE,
            documents=[],
            texts_by_id={},
            run_semantic_review=False,
        )

        project_basics = result["structured"]["fieldGroups"]["projectBasics"]
        self.assertEqual(field_by_key(project_basics, "projectName")["value"], "公告真实项目")
        self.assertEqual(field_by_key(project_basics, "tenderer")["value"], "公告真实招标单位")
        self.assertNotEqual(field_by_key(project_basics, "projectName")["value"], "见招标公告")
        self.assertNotEqual(field_by_key(project_basics, "tenderer")["value"], "见招标公告")

    def test_business_bid_participating_promotes_parse_json_to_business_workspace(self) -> None:
        project_id = self.create_business_project()
        tender = "\n".join(
            [
                "# 商务招标文件",
                "项目名称：商务归档测试项目",
                "招标编号：BUS-2026-002",
                "投标保证金：须提交保函。",
                "投标人不得存在下列情形之一。",
                "附表1：商务偏差表",
                "| 序号 | 条款 | 偏差说明 |",
                "| --- | --- | --- |",
                "| 1 | 付款条件 | |",
            ]
        ).encode("utf-8")

        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[("tenderFiles", ("商务招标文件.md", tender, "text/markdown"))],
        )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["structured"]["schemaVersion"], "bid-business-tender-structured-v1")
        approve_response = self.client.post(
            self.parse_results_url(project_id, "/commitment-letters/approve"),
            json={"approved": True},
        )
        self.assertEqual(approve_response.status_code, 200)
        self.assertEqual(approve_response.json()["approvedCount"], 1)

        updated = self.client.put(
            self.project_url(project_id),
            json={
                "name": "商务归档测试项目",
                "customerName": "测试业主",
                "bidType": "商务标",
                "reviewDecision": "participate",
            },
        )
        self.assertEqual(updated.status_code, 200)

        workspace_parse_dir = settings.documents_dir / project_id / "business-workspace" / "parse"
        workspace_appendix_dir = settings.documents_dir / project_id / "business-workspace" / "appendices"
        workspace_commitment_dir = settings.documents_dir / project_id / "business-workspace" / "commitment-letters"
        self.assertTrue((workspace_parse_dir / "s1_structured_result.json").exists())
        self.assertTrue((workspace_parse_dir / "parse-result.workspace.json").exists())
        self.assertTrue(workspace_appendix_dir.exists())
        self.assertTrue(workspace_commitment_dir.exists())

        project = store._require(project_id)
        parse_storage = project["parse_storage"]
        self.assertEqual(Path(parse_storage["projectDir"]), settings.documents_dir / project_id / "business-workspace")
        self.assertEqual(Path(parse_storage["parseDir"]), workspace_parse_dir)
        s1_handoff = project["stageArtifacts"]["s1"]
        self.assertEqual(s1_handoff["schemaVersion"], "business-s1-handoff-v1")
        self.assertEqual(s1_handoff["status"], "published")
        self.assertEqual(s1_handoff["parseProfile"], "business")
        self.assertEqual(Path(s1_handoff["paths"]["structuredResultPath"]), workspace_parse_dir / "s1_structured_result.json")
        self.assertEqual(Path(s1_handoff["paths"]["appendicesDir"]), workspace_appendix_dir)
        self.assertEqual(Path(s1_handoff["paths"]["commitmentLettersDir"]), workspace_commitment_dir)

        structured_result = json.loads((workspace_parse_dir / "s1_structured_result.json").read_text(encoding="utf-8"))
        self.assertEqual(structured_result["schemaVersion"], "bid-business-tender-structured-v1")
        self.assertEqual(structured_result["structured"]["schemaVersion"], "bid-business-tender-structured-v1")
        commitment_letters = structured_result["structured"]["commitmentLetters"]
        self.assertEqual(len(commitment_letters), 1)
        self.assertIn(str(workspace_commitment_dir), commitment_letters[0]["docxPath"])
        self.assertEqual(
            commitment_letters[0]["workspacePath"],
            f"business-workspace/commitment-letters/{Path(commitment_letters[0]['docxPath']).name}",
        )

        preview = self.client.get(
            self.parse_results_url(project_id, f"/commitment-letters/{commitment_letters[0]['id']}/preview")
        )
        self.assertEqual(preview.status_code, 200)
        self.assertIn(str(workspace_commitment_dir), preview.json()["docxPath"])

    def test_business_parse_results_recovers_completed_structured_file_after_idle_state(self) -> None:
        project_id = self.create_business_project()
        parse_dir = settings.parsed_dir / project_id
        parse_dir.mkdir(parents=True, exist_ok=True)
        structured_payload = {
            "items": [
                {
                    "id": "REQ-1",
                    "fieldKey": "projectName",
                    "title": "项目名称",
                    "value": "后台\x00恢复测试项目",
                    "sourceFile": "商务招标文件.pdf",
                }
            ],
            "structured": {
                "schemaVersion": "bid-business-tender-structured-v1",
                "fieldGroups": {
                    "projectBasics": [
                        {
                            "key": "projectName",
                            "label": "项目名称",
                            "value": "后台\x00恢复测试项目",
                        }
                    ]
                },
                "appendices": [],
                "commitmentLetters": [],
                "workflow": {
                    "documentParseEngine": "docling",
                    "documentParseStatus": "completed",
                    "fallbackUsed": False,
                },
            },
            "summary": {
                "fileCount": 1,
                "extractedCount": 1,
                "textLength": 12,
                "textPreview": "后台\x00恢复测试项目",
                "warnings": [],
            },
        }
        (parse_dir / "s1_structured_result.json").write_text(
            json.dumps(structured_payload, ensure_ascii=False),
            encoding="utf-8",
        )
        (parse_dir / "manifest.json").write_text(
            json.dumps(
                {
                    "documents": [
                        {
                            "id": "TEN-1",
                            "name": "商务招标文件.pdf",
                            "documentParseEngine": "docling",
                            "documentParseStatus": "completed",
                            "fallbackUsed": False,
                        }
                    ]
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        (parse_dir / "combined.txt").write_text("后台\x00恢复测试项目", encoding="utf-8")

        project = store._require(project_id)
        project["parse_result"] = {
            "status": "idle",
            "parsedAt": "",
            "sourceFiles": [],
            "items": [],
            "structured": {},
            "summary": {"fileCount": 0, "extractedCount": 0, "textLength": 0, "textPreview": "", "warnings": []},
        }
        project["parse_storage"] = {
            "projectDir": "",
            "parseDir": "",
            "combinedTextPath": "",
            "manifestPath": "",
            "documents": [],
        }
        project["parse_progress"] = {
            "status": "completed",
            "percentage": 100,
            "summary": "解析完成",
            "startedAt": "",
            "completedAt": "",
            "events": [],
        }
        store.persist_project_state(project)

        response = self.client.get(self.parse_results_url(project_id))

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["status"], "completed")
        self.assertEqual(len(payload["items"]), 1)
        self.assertEqual(payload["items"][0]["value"], "后台恢复测试项目")
        self.assertNotIn("\x00", json.dumps(payload, ensure_ascii=False))
        self.assertEqual(payload["structured"]["workflow"]["documentParseEngine"], "docling")
        self.assertFalse(payload["structured"]["workflow"]["fallbackUsed"])

        project = store._require(project_id)
        self.assertEqual(project["parse_result"]["status"], "completed")
        self.assertEqual(len(project["parse_result"]["items"]), 1)
        self.assertNotIn("\x00", json.dumps(project["parse_result"], ensure_ascii=False))
        self.assertNotIn("\x00", json.dumps(project["parse_storage"], ensure_ascii=False))
        self.assertEqual(Path(project["parse_storage"]["structuredResultPath"]), parse_dir / "s1_structured_result.json")
        self.assertEqual(Path(project["parse_storage"]["combinedTextPath"]), parse_dir / "combined.txt")

    def test_business_complete_parse_strips_nul_chars_before_persisting(self) -> None:
        project_id = self.create_business_project()
        parse_dir = settings.parsed_dir / project_id
        parse_dir.mkdir(parents=True, exist_ok=True)
        structured_path = parse_dir / "s1_structured_result.json"
        structured_path.write_text("{}", encoding="utf-8")

        from app.services.bid_parse_service import business_parse_service

        parse_result = business_parse_service.complete_parse(
            project_id,
            [
                {
                    "id": "TEN-1",
                    "name": "商务招标文件.pdf",
                    "size_label": "1 MB",
                    "path": str(settings.uploads_dir / project_id / "商务招标文件.pdf"),
                }
            ],
            [],
            summary={
                "fileCount": 1,
                "extractedCount": 1,
                "textLength": 10,
                "textPreview": "预览\x00文本",
                "warnings": ["警告\x00内容"],
            },
            parse_storage={
                "projectDir": str(parse_dir),
                "combinedTextPath": str(parse_dir / "combined.txt"),
                "manifestPath": str(parse_dir / "manifest.json"),
                "structuredResultPath": str(structured_path),
                "documents": [{"id": "TEN-1", "name": "商务\x00招标文件.pdf"}],
                "items": [{"id": "REQ-1", "value": "字段\x00值"}],
                "structured": {"fieldGroups": {"projectBasics": [{"key": "projectName", "value": "项目\x00名称"}]}},
            },
        )

        self.assertEqual(parse_result["items"][0]["value"], "字段值")
        self.assertEqual(parse_result["structured"]["fieldGroups"]["projectBasics"][0]["value"], "项目名称")
        project = store._require(project_id)
        self.assertNotIn("\x00", json.dumps(project["parse_result"], ensure_ascii=False))
        self.assertNotIn("\x00", json.dumps(project["parse_storage"], ensure_ascii=False))

    def test_business_bid_text_attachment_template_docx_keeps_template_body(self) -> None:
        project_id = self.create_business_project()
        tender = "\n".join(
            [
                "第一章 投标文件格式",
                "附件1 投标函",
                "致：华能集团",
                "我方已仔细研究招标文件的全部内容，愿意按招标文件要求参加投标。",
                "投标人（盖章）：____________",
                "二、法定代表人授权书",
                "本人授权以下代表作为我方合法代理人参加本项目投标。",
                "授权代表签字：____________",
            ]
        ).encode("utf-8")

        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[("tenderFiles", ("商务招标文件.md", tender, "text/markdown"))],
        )

        self.assertEqual(response.status_code, 200)
        appendices = response.json()["structured"]["appendices"]
        self.assertGreaterEqual(len(appendices), 2)
        bid_letter = next(item for item in appendices if "投标函" in item["title"])
        self.assertEqual(bid_letter["rowCount"], 0)
        doc = Document(str(Path(bid_letter["docxPath"])))
        paragraph_texts = [paragraph.text for paragraph in doc.paragraphs]
        self.assertIn("致：华能集团", paragraph_texts)
        self.assertTrue(any("愿意按招标文件要求参加投标" in text for text in paragraph_texts))

    def test_business_bid_extracts_sixth_chapter_attachment_templates(self) -> None:
        project_id = self.create_business_project()
        tender = "\n".join(
            [
                "第二章 投标人须知",
                "投标人须无条件承诺在本采购项目第一台合同设备供货前取得本条a和b所述材料，需提供承诺书。",
                "第六章 投标文件格式",
                "附件1 投标函",
                "致：华能集团",
                "投标人（盖章）：____________",
                "附件2 法定代表人（单位负责人）身份证明",
                "姓名：____________ 身份证号：____________",
                "附件3 业绩情况表",
                "| 序号 | 项目名称 | 合同容量 | 投运时间 |",
                "| --- | --- | --- | --- |",
                "| 1 |  |  |  |",
            ]
        ).encode("utf-8")

        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[("tenderFiles", ("商务招标文件.md", tender, "text/markdown"))],
        )

        self.assertEqual(response.status_code, 200)
        structured = response.json()["structured"]
        appendices = structured["appendices"]
        titles = [item["title"] for item in appendices]
        self.assertTrue(any("附件1 投标函" in title for title in titles))
        self.assertTrue(any("法定代表人" in title and "身份证明" in title for title in titles))
        self.assertTrue(any("业绩情况表" in title for title in titles))
        bid_letter = next(item for item in appendices if "投标函" in item["title"])
        doc = Document(str(Path(bid_letter["docxPath"])))
        self.assertIn("致：华能集团", [paragraph.text for paragraph in doc.paragraphs])
        self.assertNotIn("投标人须无条件承诺", "".join(titles))

    def test_business_bid_docx_attachment_templates_do_not_use_legacy_slice_when_agent_missing(self) -> None:
        project_id = self.create_business_project()
        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[
                (
                    "tenderFiles",
                    (
                        "商务附件模板招标文件.docx",
                        build_business_attachment_templates_docx_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
            ],
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["structured"]["appendices"], [])
        self.assertEqual(payload["summary"]["appendixCount"], 0)
        self.assertTrue(payload["summary"].get("warnings"))

    def test_business_bid_uses_template_extractor_and_keeps_header_cluster(self) -> None:
        project_id = self.create_business_project()
        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[
                (
                    "tenderFiles",
                    (
                        "商务招标文件.docx",
                        build_business_multilevel_template_cluster_docx_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
            ],
        )
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        structured = payload["structured"]
        appendices = structured["appendices"]
        self.assertEqual(appendices, [])
        parse_dir = settings.parsed_dir / project_id
        extraction_path = parse_dir / "business_template_extraction" / "business_template_extraction.json"
        self.assertTrue(extraction_path.is_file())
        extraction_payload = json.loads(extraction_path.read_text(encoding="utf-8"))
        self.assertEqual(extraction_payload["summary"]["templateCount"], 0)
        self.assertEqual(extraction_payload.get("schemaVersion"), "bid-business-template-extractor-v1")
        self.assertFalse((extraction_path.parent / "DOC-1" / "candidate_templates.json").exists())
        self.assertFalse((extraction_path.parent / "DOC-1" / "llm_boundary_decisions.json").exists())
        skill_manifest = json.loads((parse_dir / "s1_parse_manifest.json").read_text(encoding="utf-8"))
        self.assertEqual(skill_manifest["businessTemplateExtractionPath"], str(extraction_path))
        self.assertEqual(skill_manifest["businessTemplateExtractionSummary"]["templateCount"], 0)

    def test_business_template_extractor_appendices_survive_skill_result_merge(self) -> None:
        project_id = self.create_business_project()
        template_docx = settings.parsed_dir / project_id / "business_template_extraction" / "templates" / "TPL-0001.docx"

        def fake_template_extractor(
            *,
            project_id: str,
            documents: list[dict],
            project_dir: Path,
            progress_callback=None,
            cancel_check=None,
        ):
            template_docx.parent.mkdir(parents=True, exist_ok=True)
            Document().save(str(template_docx))
            appendix = {
                "id": "APPX-0001",
                "title": "A投标价格总表",
                "artifactType": "business_attachment_template",
                "templateType": "price_table",
                "status": "generated",
                "docxPath": str(template_docx),
                "sourceDocumentId": documents[0]["id"],
                "sourceDocumentName": documents[0]["name"],
                "extractionMode": "business_template_extractor_skill",
            }
            payload = {
                "schemaVersion": "bid-business-template-extractor-v1",
                "skillName": "bid-business-template-extractor",
                "summary": {"templateCount": 1},
                "appendices": [appendix],
            }
            output_path = project_dir / "business_template_extraction" / "business_template_extraction.json"
            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            return [appendix], payload, ""

        with patch("app.services.parsing.run_business_template_extractor", side_effect=fake_template_extractor):
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[
                    (
                        "tenderFiles",
                        (
                            "商务招标文件.docx",
                            build_business_multilevel_template_cluster_docx_bytes(),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        ),
                    )
                ],
            )

        self.assertEqual(response.status_code, 200)
        appendices = response.json()["structured"]["appendices"]
        self.assertTrue(all(item["extractionMode"] == "business_template_extractor_skill" for item in appendices))
        self.assertTrue(any("A投标价格总表" in item["title"] for item in appendices))

    def test_business_template_extractor_empty_result_does_not_fallback_to_legacy(self) -> None:
        project_id = self.create_business_project()
        with patch(
            "app.services.parsing.run_business_template_extractor",
            return_value=([], {"summary": {"templateCount": 0}, "appendices": []}, "template skill empty"),
        ), patch(
            "app.services.parsing._extract_docx_appendices",
            wraps=parsing_service._extract_docx_appendices,
        ) as docx_appendix_extractor, patch(
            "app.services.parsing._extract_markdown_appendices",
            wraps=parsing_service._extract_markdown_appendices,
        ) as markdown_appendix_extractor, patch(
            "app.services.parsing._extract_text_business_appendices",
            wraps=parsing_service._extract_text_business_appendices,
        ) as text_appendix_extractor:
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[
                    (
                        "tenderFiles",
                        (
                            "商务招标文件.docx",
                            build_business_attachment_templates_docx_bytes(),
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        ),
                    )
                ],
            )
        self.assertEqual(response.status_code, 200)
        appendices = response.json()["structured"]["appendices"]
        self.assertEqual(appendices, [])
        self.assertEqual(docx_appendix_extractor.call_count, 0)
        self.assertEqual(markdown_appendix_extractor.call_count, 0)
        self.assertEqual(text_appendix_extractor.call_count, 0)

    def test_business_template_agent_failure_warning_does_not_allow_preview_fallback(self) -> None:
        warning = "商务模板提取 Agent 未完成，未启用脚本兜底：opencode incomplete/stalled"

        self.assertFalse(parsing_service._business_template_extractor_allows_preview_fallback(warning))

    def test_business_template_missing_agent_decision_warning_does_not_allow_preview_fallback(self) -> None:
        warning = "模板边界 Agent 裁决未完成，未启用脚本兜底：缺少 Agent 裁决文件：llm_boundary_decisions.json"

        self.assertFalse(parsing_service._business_template_extractor_allows_preview_fallback(warning))

    def test_business_template_empty_skill_result_does_not_allow_preview_fallback(self) -> None:
        warning = "商务模板提取 skill 未识别到模板。"

        self.assertFalse(parsing_service._business_template_extractor_allows_preview_fallback(warning))

    def test_business_bid_docx_attachment_templates_with_toc_do_not_use_legacy_slice_when_agent_missing(self) -> None:
        project_id = self.create_business_project()
        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[
                (
                    "tenderFiles",
                    (
                        "商务附件模板含目录招标文件.docx",
                        build_business_attachment_templates_with_toc_docx_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
            ],
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["structured"]["appendices"], [])
        self.assertEqual(payload["summary"]["appendixCount"], 0)
        self.assertTrue(payload["summary"].get("warnings"))

    def test_business_bid_docx_table_fingerprint_does_not_use_legacy_slice_when_agent_missing(self) -> None:
        project_id = self.create_business_project()
        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[
                (
                    "tenderFiles",
                    (
                        "商务表格指纹招标文件.docx",
                        build_business_fingerprint_only_tables_docx_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
            ],
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["structured"]["appendices"], [])
        self.assertEqual(payload["summary"]["appendixCount"], 0)
        self.assertTrue(payload["summary"].get("warnings"))

    def test_business_bid_title_only_attachment_template_is_not_materialized(self) -> None:
        project_id = self.create_business_project()
        tender = build_docx_blocks_bytes(
            "第六章 投标文件格式",
            "附件1 投标函",
            "附件2 开标价格表",
            [
                ["序号", "项目名称", "投标报价"],
                ["1", "", ""],
            ],
        )

        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[
                (
                    "tenderFiles",
                    (
                        "商务招标文件.docx",
                        tender,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
            ],
        )

        self.assertEqual(response.status_code, 200)
        appendices = response.json()["structured"]["appendices"]
        titles = [item["title"] for item in appendices]
        self.assertNotIn("附件1 投标函", titles)
        self.assertEqual(titles, [])

    def test_business_bid_probably_incomplete_attachment_template_is_not_materialized(self) -> None:
        project_id = self.create_business_project()
        tender = build_docx_blocks_bytes(
            "第六章 投标文件格式",
            "附件1 其他说明",
            "本附件用于说明投标人认为需要说明的其他事项，具体内容由投标人结合项目实际情况自行说明",
            "附件2 开标价格表",
            [
                ["序号", "项目名称", "投标报价"],
                ["1", "", ""],
            ],
        )

        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[
                (
                    "tenderFiles",
                    (
                        "商务招标文件.docx",
                        tender,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
            ],
        )

        self.assertEqual(response.status_code, 200)
        appendices = response.json()["structured"]["appendices"]
        titles = [item["title"] for item in appendices]
        self.assertNotIn("附件1 其他说明", titles)
        self.assertEqual(titles, [])

    def test_business_bid_existing_commitment_template_suppresses_generated_duplicate(self) -> None:
        project_id = self.create_business_project()
        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[
                (
                    "tenderFiles",
                    (
                        "商务承诺模板招标文件.docx",
                        build_business_commitment_template_alignment_docx_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
            ],
        )

        self.assertEqual(response.status_code, 200)
        structured = response.json()["structured"]
        appendices = structured["appendices"]
        self.assertEqual(appendices, [])
        self.assertEqual([item["title"] for item in structured["commitmentLetters"]], ["材料取得承诺书"])
        self.assertEqual(structured["commitmentTemplateAlignments"], [])

    def test_business_bid_unrelated_commitment_template_does_not_suppress_required_letter(self) -> None:
        project_id = self.create_business_project()
        tender = "\n".join(
            [
                "# 商务招标文件",
                "投标人不得存在下列情形之一。",
                "# 第六章 投标文件格式",
                "附件6 履约保证函格式",
                "我方承诺按招标文件要求提交履约保证函。",
                "投标人（盖章）：____________",
            ]
        ).encode("utf-8")

        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[("tenderFiles", ("商务招标文件.md", tender, "text/markdown"))],
        )

        self.assertEqual(response.status_code, 200)
        structured = response.json()["structured"]
        self.assertEqual([item["title"] for item in structured["appendices"]], ["附件6 履约保证函格式"])
        titles = [item["title"] for item in structured["commitmentLetters"]]
        self.assertIn("投标人不存在下列情形之一承诺函", titles)
        self.assertEqual(structured["commitmentTemplateAlignments"], [])

    def test_business_bid_generates_semantic_business_commitment_and_filters_technical_commitments(self) -> None:
        project_id = self.create_business_project()
        tender = "\n".join(
            [
                "# 商务招标文件",
                "项目名称：商务承诺优化项目",
                "招标编号：BUS-2026-COM-001",
                "招标人：测试招标人",
                "投标人需同时承诺两个保证年等效满负荷小时数值（需提供书面承诺函）。",
                "投标人保证年等效满负荷小时数（需提供书面承诺函）。",
                "投标人须无条件承诺在本采购项目第一台合同设备供货前取得本条a和b所述材料，需提供承诺书。",
                "投标人不得存在下列情形之一。",
            ]
        ).encode("utf-8")

        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[("tenderFiles", ("商务招标文件.md", tender, "text/markdown"))],
        )

        self.assertEqual(response.status_code, 200)
        letters = response.json()["structured"]["commitmentLetters"]
        titles = [item["title"] for item in letters]
        all_text = "".join(titles + [str(item.get("triggerContext") or "") for item in letters])
        self.assertIn("材料取得承诺书", titles)
        self.assertIn("投标人不存在下列情形之一承诺函", titles)
        self.assertNotIn("等效满负荷小时", all_text)
        self.assertNotIn("发电量", all_text)

    def test_business_bid_commitment_semantic_review_can_override_strong_rule_candidate(self) -> None:
        project_id = self.create_business_project()
        tender = "\n".join(
            [
                "# 商务招标文件",
                "项目名称：商务承诺 AI 复核项目",
                "招标编号：BUS-2026-AI-001",
                "招标人：测试招标人",
                "投标人须无条件承诺在本采购项目第一台合同设备供货前取得本条a和b所述材料，需提供承诺书。",
            ]
        ).encode("utf-8")

        with patch.object(
            parsing_service.OpencodeClient,
            "review_business_commitments_with_trace",
            return_value={
                "decisions": [
                    {
                        "id": "RAW-0001",
                        "action": "ignore",
                        "topicKey": "certificate_obtainment",
                        "preferredTitle": "",
                        "reason": "AI 判断该项不需要自动生成商务承诺书。",
                    }
                ]
            },
        ) as mocked_review:
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[("tenderFiles", ("商务招标文件.md", tender, "text/markdown"))],
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(mocked_review.call_count, 1)
        self.assertEqual(response.json()["structured"]["commitmentLetters"], [])

    def test_business_bid_commitment_strong_rule_fallback_when_semantic_review_unavailable(self) -> None:
        project_id = self.create_business_project()
        tender = "\n".join(
            [
                "# 商务招标文件",
                "项目名称：商务承诺兜底项目",
                "招标编号：BUS-2026-FB-001",
                "招标人：测试招标人",
                "投标人须无条件承诺在本采购项目第一台合同设备供货前取得本条a和b所述材料，需提供承诺书。",
            ]
        ).encode("utf-8")

        with patch.object(
            parsing_service.OpencodeClient,
            "review_business_commitments_with_trace",
            side_effect=RuntimeError("semantic review unavailable"),
        ) as mocked_review:
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[("tenderFiles", ("商务招标文件.md", tender, "text/markdown"))],
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(mocked_review.call_count, 1)
        letters = response.json()["structured"]["commitmentLetters"]
        self.assertEqual([item["title"] for item in letters], ["材料取得承诺书"])
        self.assertIn("rule_fallback_generated", letters[0]["riskFlags"])

    def test_business_bid_commitment_titles_use_semantic_review_and_dedupe_same_topic(self) -> None:
        project_id = self.create_business_project()
        tender = "\n".join(
            [
                "# 商务招标文件",
                "项目名称：商务承诺测试项目",
                "招标编号：BUS-2026-SEM-001",
                "招标人：测试招标人",
                "第八章 投标人需要说明的其他内容",
                "保密承诺书",
                "本项仅为目录标题，不构成单独提交要求。",
                "另附保密承诺书。",
                "投标人须提供保密承诺书。",
                "保密承诺书",
                "投标人应提供保密承诺书。",
                "发电量承诺书另附。",
                "投标人不得存在下列情形之一。",
            ]
        ).encode("utf-8")

        with patch.object(
            parsing_service.OpencodeClient,
            "review_business_commitments_with_trace",
            return_value={
                "decisions": [
                    {
                        "id": "RAW-0001",
                        "action": "ignore",
                        "topicKey": "confidentiality",
                        "preferredTitle": "",
                        "reason": "仅标题，无明确单独提交要求。",
                    }
                ]
            },
        ) as mocked_review:
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[("tenderFiles", ("商务招标文件.md", tender, "text/markdown"))],
            )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        structured = payload["structured"]
        letters = structured["commitmentLetters"]
        titles = [item["title"] for item in letters]

        self.assertEqual(mocked_review.call_count, 1)
        self.assertEqual(len(letters), 2)
        self.assertEqual(titles.count("保密承诺书"), 1)
        self.assertIn("投标人不存在下列情形之一承诺函", titles)
        self.assertNotIn("发电量承诺书", "".join(titles))
        self.assertEqual(
            field_by_key(structured["fieldGroups"]["commitmentRequirements"], "generatedCommitmentCount")["value"],
            "2",
        )

    def test_business_bid_commitment_title_only_candidate_becomes_clue_when_semantic_review_uncertain(self) -> None:
        project_id = self.create_business_project()
        tender = "\n".join(
            [
                "# 商务招标文件",
                "项目名称：商务承诺线索测试项目",
                "招标编号：BUS-2026-SEM-002",
                "招标人：测试招标人",
                "第八章 投标人需要说明的其他内容",
                "保密承诺书",
                "本节为模板目录，具体内容见附件。",
            ]
        ).encode("utf-8")

        with patch.object(
            parsing_service.OpencodeClient,
            "review_business_commitments_with_trace",
            return_value={
                "decisions": [
                    {
                        "id": "RAW-0001",
                        "action": "clue",
                        "topicKey": "confidentiality",
                        "preferredTitle": "",
                        "reason": "仅识别到标题，建议人工确认是否需要单独成文。",
                    }
                ]
            },
        ) as mocked_review:
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[("tenderFiles", ("商务招标文件.md", tender, "text/markdown"))],
            )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        structured = payload["structured"]
        self.assertEqual(mocked_review.call_count, 1)
        self.assertEqual(structured["commitmentLetters"], [])
        self.assertEqual(len(structured["commitmentClues"]), 1)
        self.assertEqual(structured["commitmentClues"][0]["topicKey"], "confidentiality")
        self.assertIn("人工确认", structured["commitmentClues"][0]["recommendedAction"])
        self.assertEqual(
            field_by_key(structured["fieldGroups"]["commitmentRequirements"], "pendingCommitmentCount")["value"],
            "1",
        )

    def test_delete_project_cleans_parse_temp_workspace(self) -> None:
        project_id = self.create_project()
        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[
                (
                    "tenderFiles",
                    (
                        "含附表招标文件.docx",
                        build_appendix_docx_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
            ],
        )
        self.assertEqual(response.status_code, 200)
        temp_project_dir = settings.parsed_dir / project_id
        self.assertTrue(temp_project_dir.exists())

        deleted = self.client.delete(self.project_url(project_id))

        self.assertEqual(deleted.status_code, 200)
        self.assertFalse(temp_project_dir.exists())

    def test_delete_business_project_cleans_project_material_folder(self) -> None:
        project_id = self.create_business_project()
        deleted_calls: list[tuple[str, str]] = []

        def fake_delete_folder(path: str, *, expected_project_id: str = "") -> dict[str, object]:
            deleted_calls.append((path, expected_project_id))
            return {"message": "deleted", "folderPath": path, "deletedFileCount": 2}

        with patch("app.services.bid_project_state.run_workspace_material_folder_delete", side_effect=fake_delete_folder):
            deleted = self.client.delete(self.project_url(project_id))

        self.assertEqual(deleted.status_code, 200)
        self.assertEqual(deleted_calls, [(f"商务标/项目素材/{project_id}", project_id)])

    def test_parse_results_materializes_legacy_required_appendix_preview_docx(self) -> None:
        project_id = self.create_project()
        complete_parse_for_tests(
            project_id,
            [{"id": "TEN-1", "name": "招标文件.docx", "size_label": "1.0 MB"}],
            [],
            summary={"fileCount": 1, "extractedCount": 0, "textLength": 0, "warnings": []},
            parse_storage={
                "items": [],
                "structured": {
                    "appendices": [
                        {
                            "id": "APPX-0007",
                            "title": "附表7：技术资料递交表",
                            "status": "required_no_template",
                            "sourceFile": "招标文件.docx",
                            "evidence": "附表7：技术资料递交表",
                            "evidenceLocation": "L88",
                            "rows": [],
                            "rowCount": 0,
                            "docxPath": "",
                        }
                    ]
                },
            },
        )

        response = self.client.get(self.parse_results_url(project_id))
        self.assertEqual(response.status_code, 200)
        appendix = response.json()["structured"]["appendices"][0]
        self.assertEqual(appendix["status"], "generated")
        self.assertEqual(appendix["rowCount"], 0)
        appendix_path = Path(appendix["docxPath"])
        self.assertTrue(appendix_path.exists())

        preview = self.client.get(self.parse_results_url(project_id, "/appendices/APPX-0007/preview"))
        self.assertEqual(preview.status_code, 200)
        preview_payload = preview.json()
        self.assertEqual(preview_payload["id"], "APPX-0007")
        self.assertEqual(preview_payload["onlyoffice"]["documentType"], "word")
        self.assertTrue(preview_payload["onlyoffice"]["documentKey"])

        file_response = self.client.get(
            self.parse_results_url(project_id, f"/appendices/APPX-0007/file/{appendix_path.name}")
        )
        self.assertEqual(file_response.status_code, 200)
        self.assertGreater(len(file_response.content), 0)

    def test_business_parse_results_refreshes_finalized_structured_file_before_returning(self) -> None:
        project_id = self.create_business_project()
        structured_path = settings.parsed_dir / project_id / "s1_structured_result.json"
        structured_path.parent.mkdir(parents=True, exist_ok=True)
        structured_path.write_text(
            json.dumps(
                {
                    "items": [],
                    "structured": {
                        "schemaVersion": "bid-business-tender-structured-v1",
                        "workflow": {"stage": "finalized", "mode": "opencode-agentic-navigation"},
                        "fieldGroups": {
                            "qualificationRequirements": [
                                {
                                    "content": "供应商须为入围供应商。",
                                    "evidenceIds": ["TEN-1:B000089"],
                                }
                            ],
                            "bidderInstructions": [],
                            "commercialRejectionClauses": [],
                            "projectBasics": [],
                        },
                        "scoringCriteria": {"business": []},
                    },
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        complete_parse_for_tests(
            project_id,
            [{"id": "TEN-1", "name": "商务招标文件.docx", "size_label": "1.0 MB"}],
            [],
            summary={"fileCount": 1, "extractedCount": 0, "textLength": 0, "warnings": []},
            parse_storage={
                "items": [],
                "structuredResultPath": str(structured_path),
                "structured": {
                    "schemaVersion": "bid-business-tender-structured-v1",
                    "workflow": {"stage": "finalized", "mode": "opencode-agentic-navigation"},
                    "fieldGroups": {
                        "qualificationRequirements": [
                            {
                                "content": "供应商须为入围供应商。",
                                "__evidenceIds": ["TEN-1:B000089"],
                            }
                        ],
                        "bidderInstructions": [],
                        "commercialRejectionClauses": [],
                        "projectBasics": [],
                    },
                    "scoringCriteria": {"business": []},
                },
            },
        )

        response = self.client.get(self.parse_results_url(project_id))

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        text = json.dumps(payload, ensure_ascii=False)
        self.assertNotIn("__evidenceIds", text)
        qualification = payload["structured"]["fieldGroups"]["qualificationRequirements"][0]
        self.assertEqual(qualification["evidenceIds"], ["TEN-1:B000089"])

        project = store._require(project_id)
        persisted_text = json.dumps(project["parse_result"], ensure_ascii=False)
        self.assertNotIn("__evidenceIds", persisted_text)
        self.assertEqual(
            project["parse_storage"]["structured"]["fieldGroups"]["qualificationRequirements"][0]["evidenceIds"],
            ["TEN-1:B000089"],
        )

    def test_business_parse_results_hydrates_template_appendices_when_structured_file_omits_them(self) -> None:
        project_id = self.create_business_project()
        parse_dir = settings.parsed_dir / project_id
        structured_path = parse_dir / "s1_structured_result.json"
        extraction_path = parse_dir / "business_template_extraction" / "business_template_extraction.json"
        template_docx = extraction_path.parent / "templates" / "TPL-0001.docx"
        template_docx.parent.mkdir(parents=True, exist_ok=True)
        Document().save(str(template_docx))
        structured_path.parent.mkdir(parents=True, exist_ok=True)
        structured_path.write_text(
            json.dumps(
                {
                    "items": [],
                    "structured": {
                        "schemaVersion": "bid-business-tender-structured-v1",
                        "workflow": {"stage": "finalized", "mode": "opencode-agentic-navigation"},
                        "fieldGroups": {"projectBasics": []},
                        "scoringCriteria": {"business": []},
                    },
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        extraction_path.parent.mkdir(parents=True, exist_ok=True)
        extraction_path.write_text(
            json.dumps(
                {
                    "schemaVersion": "bid-business-template-extractor-v1",
                    "skillName": "bid-business-template-extractor",
                    "summary": {"templateCount": 1},
                    "appendices": [
                        {
                            "id": "APPX-0001",
                            "title": "Bid Letter Template",
                            "evidence": "Bid Letter Template",
                            "artifactType": "business_attachment_template",
                            "templateType": "bid_letter",
                            "status": "generated",
                            "docxPath": str(template_docx),
                            "sourceDocumentId": "TEN-1",
                            "sourceDocumentName": "business-tender.docx",
                            "extractionMode": "business_template_extractor_skill",
                        }
                    ],
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        complete_parse_for_tests(
            project_id,
            [{"id": "TEN-1", "name": "business-tender.docx", "size_label": "1.0 MB"}],
            [],
            summary={"fileCount": 1, "extractedCount": 0, "textLength": 0, "warnings": [], "appendixCount": 1},
            parse_storage={
                "items": [],
                "structuredResultPath": str(structured_path),
                "businessTemplateExtractionPath": str(extraction_path),
                "structured": {
                    "schemaVersion": "bid-business-tender-structured-v1",
                    "workflow": {"stage": "finalized", "mode": "opencode-agentic-navigation"},
                    "fieldGroups": {"projectBasics": []},
                    "scoringCriteria": {"business": []},
                    "appendices": [
                        {
                            "id": "APPX-0001",
                            "title": "Bid Letter Template",
                            "status": "generated",
                            "docxPath": str(template_docx),
                            "extractionMode": "business_template_extractor_skill",
                        }
                    ],
                },
            },
        )

        response = self.client.get(self.parse_results_url(project_id))

        self.assertEqual(response.status_code, 200)
        appendices = response.json()["structured"]["appendices"]
        self.assertEqual(len(appendices), 1)
        self.assertEqual(appendices[0]["title"], "Bid Letter Template")
        self.assertEqual(appendices[0]["extractionMode"], "business_template_extractor_skill")
        project = store._require(project_id)
        self.assertEqual(len(project["parse_result"]["structured"]["appendices"]), 1)
        self.assertEqual(len(project["parse_storage"]["structured"]["appendices"]), 1)

    def test_business_parse_results_materializes_evidence_ids_to_readable_sources(self) -> None:
        project_id = self.create_business_project()
        parse_dir = settings.parsed_dir / project_id
        parse_dir.mkdir(parents=True, exist_ok=True)
        nav_path = parse_dir / "s1_nav.sqlite"
        conn = sqlite3.connect(nav_path)
        try:
            conn.executescript(
                """
                CREATE TABLE evidence (
                    id TEXT PRIMARY KEY,
                    document_id TEXT NOT NULL,
                    kind TEXT NOT NULL,
                    body_index INTEGER NOT NULL,
                    table_id TEXT NOT NULL DEFAULT '',
                    row_index INTEGER,
                    col_index INTEGER,
                    text TEXT NOT NULL
                );
                CREATE TABLE blocks (
                    id TEXT PRIMARY KEY,
                    document_id TEXT NOT NULL,
                    body_index INTEGER NOT NULL,
                    block_type TEXT NOT NULL,
                    text TEXT NOT NULL,
                    heading_path TEXT NOT NULL DEFAULT ''
                );
                CREATE TABLE tables (
                    id TEXT PRIMARY KEY,
                    document_id TEXT NOT NULL,
                    body_index INTEGER NOT NULL,
                    title TEXT NOT NULL DEFAULT '',
                    heading_path TEXT NOT NULL DEFAULT ''
                );
                """
            )
            conn.execute(
                "INSERT INTO tables(id, document_id, body_index, title, heading_path) VALUES (?, ?, ?, ?, ?)",
                ("TEN-1:T0003", "TEN-1", 311, "投标人须知前附表", "第二章 投标人须知 > 投标人须知前附表"),
            )
            conn.execute(
                "INSERT INTO blocks(id, document_id, body_index, block_type, text, heading_path) VALUES (?, ?, ?, ?, ?, ?)",
                (
                    "TEN-1:B000214",
                    "TEN-1",
                    214,
                    "paragraph",
                    "3.1.1 投标人为中华人民共和国境内合法注册的独立法人或其他组织。",
                    "第一章 招标公告 > 3.1 通用资格条件 > 3.1.1 投标人为中华人民共和国境内合法注册的独立法人或其他组织。",
                ),
            )
            conn.execute(
                "INSERT INTO blocks(id, document_id, body_index, block_type, text, heading_path) VALUES (?, ?, ?, ?, ?, ?)",
                (
                    "TEN-1:T0003",
                    "TEN-1",
                    311,
                    "table",
                    "投标人须知前附表",
                    "第二章 投标人须知 > 投标人须知前附表",
                ),
            )
            conn.execute(
                "INSERT INTO blocks(id, document_id, body_index, block_type, text, heading_path) VALUES (?, ?, ?, ?, ?, ?)",
                (
                    "TEN-1:B000224",
                    "TEN-1",
                    224,
                    "paragraph",
                    "标段一和标段二：投标人须提供近3年风电机组通过试运行业绩。",
                    "第一章 招标公告 > (6) 法定代表人或单位负责人为同一人的两个及两个以上法人，母公司、全资子公司及其控股公司，不得在同一标段同时投标。",
                ),
            )
            conn.execute(
                "INSERT INTO evidence(id, document_id, kind, body_index, table_id, row_index, col_index, text) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                (
                    "TEN-1:B000214",
                    "TEN-1",
                    "paragraph",
                    214,
                    "",
                    None,
                    None,
                    "3.1.1 投标人为中华人民共和国境内合法注册的独立法人或其他组织。",
                ),
            )
            conn.execute(
                "INSERT INTO evidence(id, document_id, kind, body_index, table_id, row_index, col_index, text) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                (
                    "TEN-1:T0003:R0004",
                    "TEN-1",
                    "table_row",
                    311,
                    "TEN-1:T0003",
                    4,
                    None,
                    "1.1.4 | 招标项目名称 | 华能赤峰风电项目",
                ),
            )
            conn.execute(
                "INSERT INTO evidence(id, document_id, kind, body_index, table_id, row_index, col_index, text) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                (
                    "TEN-1:B000224",
                    "TEN-1",
                    "paragraph",
                    224,
                    "",
                    None,
                    None,
                    "标段一和标段二：投标人须提供近3年风电机组通过试运行业绩。",
                ),
            )
            conn.commit()
        finally:
            conn.close()

        structured_path = parse_dir / "s1_structured_result.json"
        structured = {
            "schemaVersion": "bid-business-tender-structured-v1",
            "sourceDocuments": [
                {
                    "id": "TEN-1",
                    "name": "招标文件-华能赤峰风电项目招标文件.docx",
                    "sourcePath": "/data/uploads/PRJ/tender.docx",
                    "textPath": "/data/parsed/PRJ/TEN-1.txt",
                }
            ],
            "workflow": {
                "stage": "finalized",
                "mode": "opencode-agentic-navigation",
                "navStorePath": str(nav_path),
            },
            "fieldGroups": {
                "projectBasics": [
                    {
                        "key": "projectName",
                        "label": "项目名称",
                        "value": "华能赤峰风电项目",
                        "evidenceIds": ["TEN-1:T0003:R0004"],
                        "evidenceLocation": "表格第4行",
                        "sourceText": "招标文件-华能赤峰风电项目招标文件.docx / 第二章 投标人须知 > 投标人须知前附表 / 表格第4行",
                    }
                ],
                "qualificationRequirements": [
                    {
                        "content": "投标人为中华人民共和国境内合法注册的独立法人或其他组织。",
                        "evidenceIds": ["TEN-1:B000214"],
                        "evidenceLocation": "正文第214段",
                        "sourceText": "招标文件-华能赤峰风电项目招标文件.docx / 第一章 招标公告 > 3.1 通用资格条件 / 正文第214段",
                    },
                    {
                        "content": "标段一和标段二：投标人须提供近3年风电机组通过试运行业绩。",
                        "evidenceIds": ["TEN-1:B000224"],
                    }
                ],
                "bidderInstructions": [],
                "commercialRejectionClauses": [],
            },
            "scoringCriteria": {"business": []},
        }
        structured_path.write_text(
            json.dumps({"items": [], "structured": structured}, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        complete_parse_for_tests(
            project_id,
            [{"id": "TEN-1", "name": "招标文件-华能赤峰风电项目招标文件.docx", "size_label": "1.0 MB"}],
            [],
            summary={"fileCount": 1, "extractedCount": 0, "textLength": 0, "warnings": []},
            parse_storage={
                "items": [],
                "structuredResultPath": str(structured_path),
                "structured": structured,
            },
        )

        response = self.client.get(self.parse_results_url(project_id))

        self.assertEqual(response.status_code, 200)
        field_groups = response.json()["structured"]["fieldGroups"]
        project_name = field_by_key(field_groups["projectBasics"], "projectName")
        self.assertEqual(project_name["sourceFile"], "招标文件-华能赤峰风电项目招标文件.docx")
        self.assertEqual(project_name["section"], "第二章 投标人须知 > 投标人须知前附表")
        self.assertEqual(project_name["evidenceLocation"], "1.1.4 招标项目名称")
        self.assertEqual(
            project_name["sourceText"],
            "招标文件-华能赤峰风电项目招标文件.docx / 第二章 投标人须知 > 投标人须知前附表 / 1.1.4 招标项目名称",
        )
        self.assertIn("招标项目名称", project_name["evidence"])
        self.assertNotIn("TEN-1:", project_name["evidenceLocation"])
        self.assertNotIn("表格第", project_name["sourceText"])

        qualification = field_groups["qualificationRequirements"][0]
        self.assertEqual(qualification["sourceFile"], "招标文件-华能赤峰风电项目招标文件.docx")
        self.assertEqual(qualification["section"], "第一章 招标公告 > 3.1 通用资格条件")
        self.assertEqual(qualification["evidenceLocation"], "3.1.1 投标人为中华人民共和国境内合法注册的独立法人或其他组织")
        self.assertIn("合法注册", qualification["evidence"])
        self.assertIn("招标文件-华能赤峰风电项目招标文件.docx", qualification["sourceText"])
        self.assertNotIn("TEN-1:", qualification["sourceText"])
        self.assertNotIn("正文第", qualification["sourceText"])
        self.assertEqual(qualification["evidenceIds"], ["TEN-1:B000214"])

        noisy_heading_qualification = field_groups["qualificationRequirements"][1]
        self.assertEqual(noisy_heading_qualification["section"], "第一章 招标公告")
        self.assertNotIn("法定代表人", noisy_heading_qualification["sourceText"])

    def test_business_upload_and_parse_returns_readable_sources_after_auto_redirect(self) -> None:
        project_id = self.create_business_project()
        parse_dir = settings.parsed_dir / project_id
        parse_dir.mkdir(parents=True, exist_ok=True)
        nav_path = parse_dir / "s1_nav.sqlite"
        conn = sqlite3.connect(nav_path)
        try:
            conn.executescript(
                """
                CREATE TABLE evidence (
                    id TEXT PRIMARY KEY,
                    document_id TEXT NOT NULL,
                    kind TEXT NOT NULL,
                    body_index INTEGER NOT NULL,
                    table_id TEXT NOT NULL DEFAULT '',
                    row_index INTEGER,
                    col_index INTEGER,
                    text TEXT NOT NULL
                );
                CREATE TABLE blocks (
                    id TEXT PRIMARY KEY,
                    document_id TEXT NOT NULL,
                    body_index INTEGER NOT NULL,
                    block_type TEXT NOT NULL,
                    text TEXT NOT NULL,
                    heading_path TEXT NOT NULL DEFAULT ''
                );
                CREATE TABLE tables (
                    id TEXT PRIMARY KEY,
                    document_id TEXT NOT NULL,
                    body_index INTEGER NOT NULL,
                    title TEXT NOT NULL DEFAULT '',
                    heading_path TEXT NOT NULL DEFAULT ''
                );
                """
            )
            conn.execute(
                "INSERT INTO tables(id, document_id, body_index, title, heading_path) VALUES (?, ?, ?, ?, ?)",
                ("TEN-1:T0001", "TEN-1", 10, "Bidder Instructions Table", "Chapter 2 > Bidder Instructions"),
            )
            conn.execute(
                "INSERT INTO blocks(id, document_id, body_index, block_type, text, heading_path) VALUES (?, ?, ?, ?, ?, ?)",
                (
                    "TEN-1:B0002",
                    "TEN-1",
                    20,
                    "paragraph",
                    "3.1.1 Bidder must be a legal entity registered in China.",
                    "Chapter 1 > 3.1 Qualification > 3.1.1 Bidder must be a legal entity registered in China.",
                ),
            )
            conn.execute(
                "INSERT INTO evidence(id, document_id, kind, body_index, table_id, row_index, col_index, text) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                ("TEN-1:T0001:R0004", "TEN-1", "table_row", 10, "TEN-1:T0001", 4, None, "1.1.4 | Project name | Liangshan wind project"),
            )
            conn.execute(
                "INSERT INTO evidence(id, document_id, kind, body_index, table_id, row_index, col_index, text) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                ("TEN-1:B0002", "TEN-1", "paragraph", 20, "", None, None, "3.1.1 Bidder must be a legal entity registered in China."),
            )
            conn.execute(
                "INSERT INTO evidence(id, document_id, kind, body_index, table_id, row_index, col_index, text) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                ("TEN-1:B0003", "TEN-1", "paragraph", 30, "", None, None, "Other bidder instruction."),
            )
            conn.commit()
        finally:
            conn.close()

        structured_path = parse_dir / "s1_structured_result.json"
        structured = {
            "schemaVersion": "bid-business-tender-structured-v1",
            "sourceDocuments": [{"id": "TEN-1", "name": "tender.md"}],
            "workflow": {"stage": "finalized", "mode": "opencode-agentic-navigation", "navStorePath": str(nav_path)},
            "fieldGroups": {
                "projectBasics": [
                    {
                        "key": "projectName",
                        "label": "Project name",
                        "value": "Liangshan wind project",
                        "evidenceIds": ["TEN-1:T0001:R0004"],
                    }
                ],
                "qualificationRequirements": [
                    {
                        "content": "Bidder must be a legal entity registered in China.",
                        "evidenceIds": ["TEN-1:B0002"],
                    }
                ],
                "bidderInstructions": [
                    {
                        "content": "Other bidder instruction.",
                        "evidenceIds": ["TEN-1:B0003"],
                    }
                ],
            },
            "scoringCriteria": {"business": []},
            "appendices": [],
        }
        structured_path.write_text(
            json.dumps({"items": [], "structured": structured}, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

        def fake_parse(project_id, tender_files, *, bid_type, progress_callback=None, cancel_check=None):
            return (
                {"fileCount": 1, "extractedCount": 0, "textLength": 10, "textPreview": "", "warnings": []},
                {
                    "documents": [{"name": "tender.md", "pageCount": 1, "textLength": 10}],
                    "items": [],
                    "structured": structured,
                    "structuredResultPath": str(structured_path),
                    "projectUpdates": {},
                },
            )

        with patch("app.services.bid_parse_service.parse_tender_documents", side_effect=fake_parse):
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[("tenderFiles", ("tender.md", b"business source test", "text/markdown"))],
            )

        self.assertEqual(response.status_code, 200)
        field_groups = response.json()["structured"]["fieldGroups"]
        project_name = field_groups["projectBasics"][0]
        self.assertEqual(project_name["sourceFile"], "tender.md")
        self.assertEqual(project_name["section"], "Chapter 2 > Bidder Instructions")
        self.assertEqual(project_name["evidenceLocation"], "1.1.4 Project name")
        self.assertEqual(project_name["sourceText"], "tender.md / Chapter 2 > Bidder Instructions / 1.1.4 Project name")

        qualification = field_groups["qualificationRequirements"][0]
        self.assertEqual(qualification["sourceFile"], "tender.md")
        self.assertEqual(qualification["section"], "Chapter 1 > 3.1 Qualification")
        self.assertEqual(qualification["evidenceLocation"], "3.1.1 Bidder must be a legal entity registered in China.")
        self.assertIn("tender.md / Chapter 1 > 3.1 Qualification", qualification["sourceText"])
        self.assertNotIn("TEN-1:", qualification["sourceText"])

        self.assertNotIn("sourceText", field_groups["bidderInstructions"][0])

    def test_parse_progress_records_real_steps_and_completion(self) -> None:
        project_id = self.create_project()
        tender = "项目名称：进度测试项目\n单机容量：6.25MW\n".encode("utf-8")

        before = self.client.get(self.parse_results_url(project_id, "/progress"))
        self.assertEqual(before.status_code, 200)
        self.assertEqual(before.json()["status"], "idle")

        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[("tenderFiles", ("招标文件.md", tender, "text/markdown"))],
        )
        self.assertEqual(response.status_code, 200)

        progress = self.client.get(self.parse_results_url(project_id, "/progress"))
        self.assertEqual(progress.status_code, 200)
        payload = progress.json()
        self.assertEqual(payload["status"], "completed")
        self.assertEqual(payload["percentage"], 100)
        steps = [event["step"] for event in payload["events"]]
        self.assertIn("upload", steps)
        self.assertIn("extract", steps)
        self.assertIn("skill", steps)
        self.assertIn("appendix", steps)
        self.assertIn("complete", steps)

    def test_parse_progress_completion_replaces_streaming_opencode_output(self) -> None:
        project_id = self.create_project()
        tender = "progress opencode output closeout test\n".encode("utf-8")

        structured = {
            "schemaVersion": "bid-tender-structured-v1",
            "projectDates": {"startDate": "", "endDate": ""},
            "appendices": [],
            "opencodeOutput": {
                "status": "received",
                "sessionId": "ses-s1",
                "parts": [{"type": "text", "text": "s1parse finalize completed"}],
                "earlyCompletion": True,
            },
        }

        with patch(
            "app.services.bid_parse_service.parse_tender_documents",
            return_value=(
                {"fileCount": 1, "extractedCount": 1, "textLength": 10, "textPreview": "", "warnings": []},
                {
                    "documents": [{"name": "tender.md", "pageCount": 1, "textLength": 10}],
                    "items": [{"id": "REQ-1", "title": "parsed"}],
                    "structured": structured,
                    "projectUpdates": {},
                },
            ),
        ):
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[("tenderFiles", ("tender.md", tender, "text/markdown"))],
            )

        self.assertEqual(response.status_code, 200)
        progress = self.client.get(self.parse_results_url(project_id, "/progress")).json()
        self.assertEqual(progress["status"], "completed")
        self.assertEqual(progress["opencodeOutput"]["status"], "received")
        self.assertNotEqual(progress["opencodeOutput"]["status"], "streaming")

    def test_parse_progress_completion_closes_stale_streaming_without_final_trace(self) -> None:
        project_id = self.create_project()
        tender = "progress stale streaming closeout test\n".encode("utf-8")

        def fake_parse(project_id, tender_files, *, bid_type, progress_callback=None, cancel_check=None):
            if progress_callback:
                progress_callback(
                    "opencode_delta",
                    {
                        "status": "streaming",
                        "sessionId": "ses-stale",
                        "parts": [{"type": "text", "text": "agent is still writing"}],
                    },
                )
            return (
                {"fileCount": 1, "extractedCount": 1, "textLength": 10, "textPreview": "", "warnings": []},
                {
                    "documents": [{"name": "tender.md", "pageCount": 1, "textLength": 10}],
                    "items": [{"id": "REQ-1", "title": "parsed"}],
                    "structured": {
                        "schemaVersion": "bid-tender-structured-v1",
                        "projectDates": {"startDate": "", "endDate": ""},
                        "appendices": [],
                    },
                    "projectUpdates": {},
                },
            )

        with patch("app.services.bid_parse_service.parse_tender_documents", side_effect=fake_parse):
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[("tenderFiles", ("tender.md", tender, "text/markdown"))],
            )

        self.assertEqual(response.status_code, 200)
        progress = self.client.get(self.parse_results_url(project_id, "/progress")).json()
        self.assertEqual(progress["status"], "completed")
        self.assertEqual(progress["opencodeOutput"]["status"], "received")
        self.assertEqual(progress["opencodeOutput"]["sessionId"], "ses-stale")

    def test_cancel_parse_endpoint_marks_progress_and_aborts_opencode_session(self) -> None:
        project_id = self.create_project()
        project = store.require_project_for_update(project_id)
        project["parse_progress"] = {
            "status": "running",
            "percentage": 80,
            "summary": "opencode 正在返回解析输出。",
            "startedAt": "2026-07-04T09:00:00Z",
            "completedAt": "",
            "events": [],
            "opencodeOutput": {
                "status": "streaming",
                "sessionId": "ses_cancel_parse_probe",
                "parts": [{"type": "text", "text": "s1parse 正在执行"}],
            },
        }
        store.persist_project_state(project)

        with patch(
            "app.services.bid_parse_service.OpencodeClient.abort_session",
            return_value=True,
            create=True,
        ) as abort_session:
            response = self.client.post(self.parse_results_url(project_id, "/cancel"))

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["status"], "cancelled")
        self.assertTrue(payload["cancelRequested"])
        self.assertEqual(payload["opencodeOutput"]["status"], "cancelled")
        self.assertEqual(payload["opencodeAbort"]["sessionId"], "ses_cancel_parse_probe")
        self.assertTrue(payload["opencodeAbort"]["aborted"])
        abort_session.assert_called_once_with("ses_cancel_parse_probe")

        progress = self.client.get(self.parse_results_url(project_id, "/progress")).json()
        self.assertEqual(progress["status"], "cancelled")
        self.assertTrue(progress["cancelRequested"])

    def test_upload_parse_returns_cancelled_when_cancel_requested_before_completion(self) -> None:
        project_id = self.create_project()
        tender = "cancel before complete\n".encode("utf-8")

        def fake_parse(project_id_arg, tender_files, *, bid_type, progress_callback=None, cancel_check=None):
            project = store.require_project_for_update(project_id_arg)
            progress = project.get("parse_progress")
            self.assertIsInstance(progress, dict)
            progress["status"] = "cancelled"
            progress["cancelRequested"] = True
            progress["summary"] = "已请求停止后端解析和 Opencode 任务。"
            project["parse_progress"] = progress
            store.persist_project_state(project)
            if cancel_check:
                self.assertTrue(cancel_check())
            return (
                {"fileCount": 1, "extractedCount": 1, "textLength": 10, "textPreview": "", "warnings": []},
                {
                    "documents": [{"name": "tender.md", "pageCount": 1, "textLength": 10}],
                    "items": [{"id": "REQ-1", "title": "should not be committed"}],
                    "structured": {
                        "schemaVersion": "bid-tender-structured-v1",
                        "appendices": [],
                    },
                    "projectUpdates": {},
                },
            )

        with patch("app.services.bid_parse_service.parse_tender_documents", side_effect=fake_parse):
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[("tenderFiles", ("tender.md", tender, "text/markdown"))],
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["status"], "cancelled")
        project = store._require(project_id)
        self.assertNotEqual(project["parse_result"]["status"], "completed")
        self.assertEqual(project["parse_progress"]["status"], "cancelled")

    def test_business_template_extraction_progress_is_visible(self) -> None:
        project_id = self.create_business_project()
        tender = "business template progress test\n".encode("utf-8")

        def fake_parse(project_id, tender_files, *, bid_type, progress_callback=None, cancel_check=None):
            if progress_callback:
                progress_callback(
                    "business_template_extraction_started",
                    {"documentCount": 1},
                )
                progress_callback(
                    "business_template_extraction_agent",
                    {
                        "status": "streaming",
                        "sessionId": "ses-template",
                        "parts": [{"type": "text", "text": "btplnav extracting templates"}],
                    },
                )
            return (
                {"fileCount": 1, "extractedCount": 1, "textLength": 10, "textPreview": "", "warnings": []},
                {
                    "documents": [{"name": "tender.md", "pageCount": 1, "textLength": 10}],
                    "items": [{"id": "REQ-1", "title": "parsed"}],
                    "structured": {
                        "schemaVersion": "bid-business-tender-structured-v1",
                        "fieldGroups": {},
                        "appendices": [],
                    },
                    "projectUpdates": {},
                },
            )

        with patch("app.services.bid_parse_service.parse_tender_documents", side_effect=fake_parse):
            response = self.client.post(
                self.parse_results_url(project_id, "/upload-and-run"),
                files=[("tenderFiles", ("tender.md", tender, "text/markdown"))],
            )

        self.assertEqual(response.status_code, 200)
        progress = self.client.get(self.parse_results_url(project_id, "/progress")).json()
        steps = [event["step"] for event in progress["events"]]
        self.assertIn("template", steps)
        self.assertEqual(progress["opencodeOutput"]["sessionId"], "ses-template")

    def test_project_create_and_update_support_start_and_end_dates(self) -> None:
        response = self.client.post(
            "/api/technical/projects",
            json={
                "name": "日期测试项目",
                "customerName": "测试业主",
                "startDate": "2026-05-10",
                "endDate": "2026-08-20",
                "deadline": "2026-08-20",
            },
        )
        self.assertEqual(response.status_code, 200)
        created = response.json()
        self.assertEqual(created["startDate"], "2026-05-10")
        self.assertEqual(created["endDate"], "2026-08-20")
        self.assertEqual(created["deadline"], "2026-08-20")

        update_response = self.client.put(
            f"/api/technical/projects/{created['id']}",
            json={
                "startDate": "2026-05-15",
                "endDate": "2026-09-01",
                "deadline": "2026-09-01",
            },
        )
        self.assertEqual(update_response.status_code, 200)
        updated = update_response.json()
        self.assertEqual(updated["startDate"], "2026-05-15")
        self.assertEqual(updated["endDate"], "2026-09-01")
        self.assertEqual(updated["deadline"], "2026-09-01")

        list_response = self.client.get("/api/technical/projects")
        self.assertEqual(list_response.status_code, 200)
        listed = list_response.json()["items"][0]
        self.assertEqual(listed["startDate"], "2026-05-15")
        self.assertEqual(listed["endDate"], "2026-09-01")

    def test_upload_and_parse_persists_text_to_disk_artifact(self) -> None:
        project_id = self.create_project()
        file_bytes = build_docx_bytes(
            "招标文件正文",
            "第二章 技术方案",
            "这里是一段比较长的测试内容，用于验证解析结果不会直接塞进项目状态数据库，而是落到磁盘文件。",
        )

        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[
                (
                    "tenderFiles",
                    ("招标文件.docx", file_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                )
            ],
        )

        self.assertEqual(response.status_code, 200)
        project = store._require(project_id)
        parse_storage = project["parse_storage"]
        combined_text_path = Path(parse_storage["combinedTextPath"])
        self.assertTrue(combined_text_path.exists())
        content = combined_text_path.read_text(encoding="utf-8")
        self.assertIn("第二章 技术方案", content)
        self.assertGreater(parse_storage["documents"][0]["textLength"], 10)

    def test_template_only_reparse_works_after_tender_uploaded(self) -> None:
        project_id = self.create_project()
        tender_bytes = build_docx_bytes("招标文件正文", "项目概况", "这是第一次上传的招标文件。")
        template_bytes = build_docx_bytes("投标模板", "封面", "这是后补上传的模板文件。")

        first_response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[
                (
                    "tenderFiles",
                    ("招标文件.docx", tender_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                )
            ],
        )
        self.assertEqual(first_response.status_code, 200)

        second_response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[
                (
                    "templateFiles",
                    ("投标模板.docx", template_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                )
            ],
        )

        self.assertEqual(second_response.status_code, 200)
        payload = second_response.json()
        self.assertEqual(payload["summary"]["fileCount"], 1)
        self.assertEqual(len(payload["project"]["templateFiles"]), 1)
        self.assertEqual(payload["project"]["templateFiles"][0]["name"], "投标模板.docx")

    def test_parse_inputs_do_not_use_legacy_template_when_project_has_no_template(self) -> None:
        project_id = self.create_project()
        tender_bytes = build_docx_bytes("招标文件正文", "项目概况", "项目没有单独上传投标模板。")
        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[
                (
                    "tenderFiles",
                    ("招标文件.docx", tender_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                )
            ],
        )
        self.assertEqual(response.status_code, 200)

        with patch("app.services.template_store.resolve_system_default_bid_template_file", return_value=None):
            _, template_files = parse_inputs_for_tests(project_id)

        self.assertEqual(template_files, [])

    def test_parse_inputs_use_settings_default_template_when_project_has_no_template(self) -> None:
        project_id = self.create_project()
        tender_bytes = build_docx_bytes("招标文件正文", "项目概况", "项目没有单独上传投标模板。")
        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[
                (
                    "tenderFiles",
                    ("招标文件.docx", tender_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                )
            ],
        )
        self.assertEqual(response.status_code, 200)

        default_path = settings.uploads_dir / project_id / "system-default-template" / "默认技术标模板.docx"
        default_path.parent.mkdir(parents=True, exist_ok=True)
        default_path.write_bytes(build_docx_bytes("默认技术标模板", "第一章 模板章节"))
        default_record = {
            "id": "TPL-0001",
            "name": "默认技术标模板.docx",
            "stored_name": "默认技术标模板.docx",
            "size_bytes": default_path.stat().st_size,
            "size_label": format_size_mb(default_path.stat().st_size),
            "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "path": str(default_path),
            "source": "system-default",
            "isFallback": True,
            "templateType": "technical",
            "templateTypeLabel": "技术标",
            "minioBucket": "bid-templates",
            "minioKey": "templates/default/technical/default-template.docx",
        }

        with patch("app.services.template_store.resolve_system_default_bid_template_file", return_value=default_record):
            _, template_files = parse_inputs_for_tests(project_id)

        self.assertEqual(len(template_files), 1)
        self.assertEqual(template_files[0]["name"], "默认技术标模板.docx")
        self.assertEqual(template_files[0]["source"], "system-default")
        self.assertEqual(template_files[0]["templateType"], "technical")

    def test_project_template_overrides_fallback_template(self) -> None:
        project_id = self.create_project()
        tender_bytes = build_docx_bytes("招标文件正文", "项目概况", "项目后续上传自己的投标模板。")
        template_bytes = build_docx_bytes("项目投标模板", "第一章 项目模板章节")
        response = self.client.post(
            self.parse_results_url(project_id, "/upload-and-run"),
            files=[
                (
                    "tenderFiles",
                    ("招标文件.docx", tender_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                ),
                (
                    "templateFiles",
                    ("项目模板.docx", template_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
                ),
            ],
        )
        self.assertEqual(response.status_code, 200)

        default_record = {
            "id": "TPL-0001",
            "name": "默认技术标模板.docx",
            "path": "/tmp/default-template.docx",
            "source": "system-default",
            "isFallback": True,
        }
        with patch("app.services.template_store.resolve_system_default_bid_template_file", return_value=default_record):
            _, template_files = parse_inputs_for_tests(project_id)

        self.assertEqual(len(template_files), 1)
        self.assertEqual(template_files[0]["name"], "项目模板.docx")
        self.assertNotEqual(template_files[0].get("source"), "system-default")


if __name__ == "__main__":
    unittest.main()
