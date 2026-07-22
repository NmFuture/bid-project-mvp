from __future__ import annotations

import importlib.util
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


BACKEND_ROOT = Path(__file__).resolve().parents[1]
NUMBERING_FIXER_PATH = (
    BACKEND_ROOT
    / "opencode"
    / "skills"
    / "bid-tech-assembler"
    / "scripts"
    / "numbering_fixer.py"
)


def _load_numbering_fixer():
    spec = importlib.util.spec_from_file_location(
        "numbering_fixer_regression",
        NUMBERING_FIXER_PATH,
    )
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def _set_num_id(element, value: int) -> None:
    p_pr = element.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(value))
    num_pr.append(num_id)
    p_pr.append(num_pr)


def _direct_num_id(paragraph) -> str | None:
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return None
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return None
    num_id = num_pr.find(qn("w:numId"))
    return num_id.get(qn("w:val")) if num_id is not None else None


def _custom_heading_document():
    doc = Document()
    custom_style = doc.styles.add_style("标题6-标书", WD_STYLE_TYPE.PARAGRAPH)
    custom_style.base_style = doc.styles["Heading 6"]
    _set_num_id(custom_style.element, 1)
    return doc, custom_style


def test_heading_injection_preserves_numid_zero_suppression() -> None:
    numbering_fixer = _load_numbering_fixer()
    doc, custom_style = _custom_heading_document()
    heading = doc.add_paragraph("低风速段出力优势", style=custom_style)
    _set_num_id(heading._p, 0)

    numbering_fixer.inject_prefix_to_headings(doc, "1.7")

    assert heading.text == "1.7.1  低风速段出力优势"
    assert _direct_num_id(heading) == "0"


def test_heading_style_numbering_cleanup_follows_based_on_chain() -> None:
    numbering_fixer = _load_numbering_fixer()
    doc, custom_style = _custom_heading_document()
    doc.add_paragraph("低风速段出力优势", style=custom_style)

    assert numbering_fixer.strip_numPr_from_heading_styles(doc) == 1

    custom_p_pr = custom_style.element.find(qn("w:pPr"))
    assert custom_p_pr is not None
    assert custom_p_pr.find(qn("w:numPr")) is None


def test_heading_paragraph_cleanup_keeps_suppression_but_removes_active_numbering() -> None:
    numbering_fixer = _load_numbering_fixer()
    doc, custom_style = _custom_heading_document()
    suppressed = doc.add_paragraph("低风速段出力优势", style=custom_style)
    active = doc.add_paragraph("降载优势", style=custom_style)
    body_list = doc.add_paragraph("正文列表")
    _set_num_id(suppressed._p, 0)
    _set_num_id(active._p, 7)
    _set_num_id(body_list._p, 8)

    assert numbering_fixer.strip_numPr_from_body(doc) == 1

    assert _direct_num_id(suppressed) == "0"
    assert _direct_num_id(active) is None
    assert _direct_num_id(body_list) == "8"
