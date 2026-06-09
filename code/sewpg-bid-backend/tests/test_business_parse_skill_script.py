from __future__ import annotations

import importlib
import json
import shutil
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


def field_by_key(fields: list[dict], key: str) -> dict:
    return next(field for field in fields if field["key"] == key)


class BusinessParseSkillScriptTests(unittest.TestCase):
    def runner_path(self) -> Path:
        backend_root = Path(__file__).resolve().parents[1]
        return (
            backend_root
            / "opencode"
            / "skill"
            / "bid-business-tender-structured-parser"
            / "scripts"
            / "run_from_manifest.py"
        )

    def skill_path(self) -> Path:
        return self.runner_path().parents[1] / "SKILL.md"

    def test_business_skill_runs_without_sibling_technical_skill(self) -> None:
        source_skill_dir = self.runner_path().parents[1]

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            isolated_skill_dir = tmp_path / source_skill_dir.name
            shutil.copytree(source_skill_dir, isolated_skill_dir)
            script_path = isolated_skill_dir / "scripts" / "run_from_manifest.py"

            source_path = tmp_path / "business_tender.md"
            source_path.write_text("# business tender\nno structured Chinese clues\n", encoding="utf-8")
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-ISOLATED-BUSINESS-SKILL",
                        "bidType": "business",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "documents": [
                            {
                                "id": "DOC-1",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                                "textPath": str(source_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            completed = subprocess.run(
                [sys.executable, str(script_path), "offline-fallback", str(manifest_path)],
                capture_output=True,
                text=True,
            )

            self.assertEqual(completed.returncode, 0, completed.stderr)
            self.assertNotIn("parser_core", completed.stderr)
            payload = json.loads(output_path.read_text(encoding="utf-8"))
            self.assertEqual(payload["structured"]["targetSkill"], "bid-business-tender-structured-parser")

    def test_business_skill_instructions_forbid_read_and_subagent_shortcuts(self) -> None:
        instructions = self.skill_path().read_text(encoding="utf-8")

        self.assertIn("s1parse tasks <manifest>", instructions)
        self.assertIn("s1parse task <manifest> <taskId>", instructions)
        self.assertIn("s1parse decision-all <manifest>", instructions)
        self.assertIn("s1parse decision-set <manifest>", instructions)
        self.assertIn("s1parse qualification-item <manifest>", instructions)
        self.assertIn("禁止用 Bash heredoc", instructions)
        self.assertIn("s1parse validate-decision <manifest> <taskId>", instructions)
        self.assertIn("s1parse status <manifest>", instructions)
        self.assertIn("禁止用 read 工具打开 `review_plan.json`", instructions)
        self.assertIn("禁止使用 opencode 的 Task/subagent/子代理/任务委派工具", instructions)
        self.assertIn("不得调用 Task 工具", instructions)

    def write_plan_decision(
        self,
        task_path: Path,
        decision_path: Path,
        *,
        accepted_candidate_ids: set[str] | None = None,
    ) -> None:
        task = json.loads(task_path.read_text(encoding="utf-8"))
        accepted_candidate_ids = accepted_candidate_ids or set()
        decision = {
            "schemaVersion": "bid-business-ai-decision-v1",
            "task": task["task"],
            "taskId": task["taskId"],
            "adapter": "unit-test-opencode-agent",
            "accepted": [],
            "rejected": [],
            "needsReview": [],
            "reason": "单元测试按 review_plan 为每个小任务生成决策。",
            "evidenceIds": [],
        }
        if task["task"] == "qualification_review":
            decision = {
                "schemaVersion": "bid-business-ai-decision-v1",
                "task": task["task"],
                "taskId": task["taskId"],
                "adapter": "unit-test-opencode-agent",
                "qualificationItems": [],
                "rejectedEvidenceIds": [],
                "reason": "单元测试按资格整节切片生成 qualificationItems。",
                "evidenceIds": [],
            }
        for candidate in task.get("candidates") or []:
            bucket = "accepted" if candidate["candidateId"] in accepted_candidate_ids else "rejected"
            if task["task"] == "qualification_review":
                if bucket == "accepted":
                    for line in candidate.get("lines") or []:
                        text = str(line.get("text") or "").strip()
                        if not text or not any(cue in text for cue in ("投标人", "联合体", "须", "应", "不得", "不接受", "具有", "具备")):
                            continue
                        if any(skip in text for skip in ("投标人资格要求", "通用资格条件", "专用资格条件", "资格审查资料", "见投标人须知前附表")):
                            continue
                        evidence_id = str(line.get("evidenceId") or "")
                        item = {
                            "content": text,
                            "applicableScope": str(line.get("applicableScopeHint") or "全部标段"),
                            "sourceText": str(line.get("sourceText") or candidate.get("sourceText") or "投标人资格要求"),
                            "evidenceIds": [evidence_id] if evidence_id else candidate.get("evidenceIds") or [],
                        }
                        decision["qualificationItems"].append(item)
                        decision["evidenceIds"].extend(item["evidenceIds"])
                else:
                    decision["rejectedEvidenceIds"].extend(candidate.get("evidenceIds") or [])
                    decision["evidenceIds"].extend(candidate.get("evidenceIds") or [])
                continue
            item = {
                "candidateId": candidate["candidateId"],
                "decision": bucket,
                "fieldType": "qualification_requirement" if task["task"] == "qualification_review" else str(candidate.get("candidateType") or "review_candidate"),
                "content": candidate["content"],
                "applicableScope": "全部标段",
                "sourceText": candidate.get("sourceText") or "招标文件 > 候选来源",
                "reason": "测试决策。",
                "evidenceIds": candidate.get("evidenceIds") or [],
            }
            decision[bucket].append(item)
            decision["evidenceIds"].extend(item["evidenceIds"])
        if task["task"] == "qualification_review":
            decision["rejectedEvidenceIds"] = sorted(set(decision["rejectedEvidenceIds"]))
        decision["evidenceIds"] = sorted(set(decision["evidenceIds"]))
        decision_path.parent.mkdir(parents=True, exist_ok=True)
        decision_path.write_text(json.dumps(decision, ensure_ascii=False, indent=2), encoding="utf-8")

    def write_bid_deadline_ai_decision(self, tmp_path: Path, review_plan: dict[str, Any], *, content: str) -> None:
        for task_ref in review_plan["tasks"]:
            task_path = tmp_path / task_ref["taskPath"]
            decision_path = tmp_path / task_ref["decisionPath"]
            task_payload = json.loads(task_path.read_text(encoding="utf-8"))
            if task_ref["task"] == "project_basics_reference_review":
                candidate = next(
                    item
                    for item in task_payload["candidates"]
                    if item.get("fieldKey") == "bidDeadline"
                )
                decision_path.parent.mkdir(parents=True, exist_ok=True)
                decision_path.write_text(
                    json.dumps(
                        {
                            "schemaVersion": "bid-business-ai-decision-v1",
                            "task": task_ref["task"],
                            "taskId": task_ref["taskId"],
                            "adapter": "unit-test-opencode-agent",
                            "accepted": [
                                {
                                    "candidateId": candidate["candidateId"],
                                    "decision": "accepted",
                                    "fieldType": "bidDeadline",
                                    "content": content,
                                    "applicableScope": "全部标段",
                                    "sourceText": candidate["sourceText"],
                                    "reason": "候选窗口明确给出了递交截止时间。",
                                    "evidenceIds": candidate["evidenceIds"],
                                }
                            ],
                            "rejected": [],
                            "needsReview": [],
                            "reason": "单元测试确认 AI 提取递交截止时间。",
                            "evidenceIds": candidate["evidenceIds"],
                        },
                        ensure_ascii=False,
                        indent=2,
                    ),
                    encoding="utf-8",
                )
                continue
            self.write_plan_decision(task_path, decision_path, accepted_candidate_ids=set())

    def assert_rejection_display_fields(self, rows: list[dict]) -> None:
        self.assertTrue(rows)
        for row in rows:
            self.assertIn(row.get("riskLevel"), {"high", "medium"})
            self.assertTrue(row.get("matchedKeywords"), row)

    def test_business_skill_script_outputs_business_contract(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "商务招标文件.docx"
            doc = Document()
            doc.add_paragraph("脱敏风电设备采购项目")
            cover_table = doc.add_table(rows=3, cols=3)
            cover_table.cell(0, 0).text = "招标人"
            cover_table.cell(0, 1).text = "："
            cover_table.cell(0, 2).text = "示例招标单位"
            cover_table.cell(1, 0).text = "项目单位"
            cover_table.cell(1, 1).text = "："
            cover_table.cell(1, 2).text = "示例项目单位"
            cover_table.cell(2, 0).text = "招标代理机构"
            cover_table.cell(2, 1).text = "："
            cover_table.cell(2, 2).text = "示例代理机构"
            doc.add_paragraph("招标编号：BUS-GEN-2026-001")
            doc.add_paragraph("投标人须知前附表")
            instruction_table = doc.add_table(rows=5, cols=3)
            for col, text in enumerate(["条款号", "条款名称", "编列内容"]):
                instruction_table.cell(0, col).text = text
            instruction_rows = [
                ("1.1.2", "招标人", "招标人：示例招标单位 地 址：示例地址 联 系 人：联系人甲 电 话：010-00000000"),
                ("1.1.3", "招标代理机构", "招标代理机构：示例代理机构 地址：示例代理地址 联系人：联系人乙 电话：400-000-0000"),
                ("1.1.4", "招标项目名称", "脱敏风电设备采购项目"),
                ("4.2.1", "投标截止时间", "2026年1月26日09时00分"),
            ]
            for row_index, values in enumerate(instruction_rows, start=1):
                for col, text in enumerate(values):
                    instruction_table.cell(row_index, col).text = text
            doc.add_paragraph("附表1：符合性审查标准表")
            compliance_table = doc.add_table(rows=2, cols=3)
            for col, text in enumerate(["序号", "审查项目", "审查标准"]):
                compliance_table.cell(0, col).text = text
            for col, text in enumerate(["1", "投标保证金", "按照招标文件要求提供投标保证金且无瑕疵。"]):
                compliance_table.cell(1, col).text = text
            doc.add_paragraph("附表3：商务评分标准表")
            business_table = doc.add_table(rows=2, cols=5)
            for col, text in enumerate(["序号", "评分项", "分值", "得分点", "证明材料要求"]):
                business_table.cell(0, col).text = text
            for col, text in enumerate(["1", "企业业绩", "20分", "近三年同类风电项目业绩满足要求得满分。", "提供合同或中标通知书。"]):
                business_table.cell(1, col).text = text
            doc.add_paragraph("附表4：投标报价评分标准")
            price_table = doc.add_table(rows=2, cols=4)
            for col, text in enumerate(["序号", "评分项", "分值", "得分点"]):
                price_table.cell(0, col).text = text
            for col, text in enumerate(["1", "评标价", "100分", "评标价等于评标基准价时得100分。"]):
                price_table.cell(1, col).text = text
            doc.add_paragraph("投标函：按招标文件格式填写并签字盖章。")
            doc.add_paragraph("投标保证金：须提供电汇回单或保函。")
            doc.add_paragraph("投标人证明其是合格投标人并有资格履行合同的证明文件。")
            doc.add_paragraph("投标人须提供供货周期承诺书。")
            doc.add_paragraph("保密承诺书")
            doc.add_paragraph("投标人须提供保密承诺书。")
            doc.add_paragraph("投标人应提供保密承诺书。")
            doc.add_paragraph("发电量承诺书另附。")
            doc.add_paragraph("技术承诺：详见技术部分。")
            doc.add_paragraph("投标文件应当对招标文件的实质性要求作出响应，否则投标将被否决。")
            doc.add_paragraph("电子投标文件逾期上传或者未成功上传指定信息平台，招标人不予受理。")
            doc.add_paragraph("投标人不得存在下列情形之一。")
            doc.add_paragraph("投标人需要说明的其他内容。")
            doc.save(source_path)
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-BUSINESS-SKILL",
                        "bidType": "商务标",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "documents": [
                            {
                                "id": "DOC-1",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            completed = subprocess.run(
                [sys.executable, str(script_path), "offline-fallback", str(manifest_path)],
                check=True,
                capture_output=True,
                text=True,
            )

            summary = json.loads(completed.stdout)
            payload = json.loads(output_path.read_text(encoding="utf-8"))
            structured = payload["structured"]

            self.assertEqual(summary["schemaVersion"], "bid-business-tender-structured-v1")
            self.assertEqual(summary["targetSkill"], "bid-business-tender-structured-parser")
            self.assertEqual(structured["schemaVersion"], "bid-business-tender-structured-v1")
            self.assertEqual(structured["targetSkill"], "bid-business-tender-structured-parser")
            field_groups = structured["fieldGroups"]
            self.assertEqual(
                set(field_groups.keys()),
                {"projectBasics", "qualificationRequirements", "bidderInstructions", "commercialRejectionClauses"},
            )
            self.assertEqual(set(structured["scoringCriteria"].keys()), {"business"})
            self.assertEqual(len(structured["scoringCriteria"]["business"]), 1)
            self.assertEqual(structured["scoringCriteria"]["business"][0]["scoringItem"], "企业业绩")
            project_basics = field_groups["projectBasics"]
            self.assertEqual(field_by_key(project_basics, "projectName")["value"], "脱敏风电设备采购项目")
            self.assertEqual(field_by_key(project_basics, "tenderNo")["value"], "BUS-GEN-2026-001")
            self.assertEqual(field_by_key(project_basics, "tenderer")["value"], "示例招标单位")
            self.assertEqual(field_by_key(project_basics, "tenderAgency")["value"], "示例代理机构")
            self.assertEqual(field_by_key(project_basics, "bidDeadline")["value"], "2026-01-26 09:00")
            self.assertEqual(field_groups["qualificationRequirements"], [])
            self.assertEqual(field_groups["bidderInstructions"][0]["clauseNo"], "1.1.2")
            self.assertEqual(field_groups["bidderInstructions"][1]["clauseName"], "招标代理机构")
            rejections = field_groups["commercialRejectionClauses"]
            self.assertGreaterEqual(len(rejections), 2)
            self.assertTrue(any("否决" in row["content"] for row in rejections))
            self.assertTrue(any("不予受理" in row["content"] for row in rejections))
            self.assert_rejection_display_fields(rejections)
            denied_row = next(row for row in rejections if "否决" in row["content"])
            rejected_upload_row = next(row for row in rejections if "不予受理" in row["content"])
            self.assertEqual(denied_row["riskLevel"], "high")
            self.assertIn("否决", denied_row["matchedKeywords"])
            self.assertEqual(rejected_upload_row["riskLevel"], "high")
            self.assertIn("不予受理", rejected_upload_row["matchedKeywords"])
            self.assertEqual(structured["appendices"], [])
            self.assertEqual(structured["commitmentLetters"], [])
            self.assertEqual(structured["commitmentClues"], [])
            fact_by_key = {field["fieldKey"]: field for field in structured["projectFactFields"]}
            self.assertEqual(fact_by_key["projectName"]["value"], "脱敏风电设备采购项目")
            self.assertEqual(fact_by_key["tenderNo"]["value"], "BUS-GEN-2026-001")
            self.assertEqual(fact_by_key["tenderer"]["value"], "示例招标单位")

    def test_qualification_requirements_are_section_based_and_filtered(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "脱敏资格要求片段.docx"
            doc = Document()
            doc.add_paragraph("目录")
            doc.add_paragraph("3.5 资格审查资料\t23")
            doc.add_paragraph("第一章 招标公告")
            doc.add_paragraph("3. 投标人资格要求")
            doc.add_paragraph("3.1 通用资格条件")
            doc.add_paragraph("3.1.1 投标人为中华人民共和国境内合法注册的独立法人或其他组织，具有独立承担民事责任能力，具有独立订立合同的权利。")
            doc.add_paragraph("3.1.2 投标人财务、信誉等方面应具备下列条件：")
            doc.add_paragraph("（1）没有处于行政主管部门或采购人系统内单位确认的禁止投标范围和处罚期内。")
            doc.add_paragraph("3.2 专用资格条件")
            doc.add_paragraph("3.2.1 业绩要求：")
            doc.add_paragraph("标段一至标段四（需同时满足）：")
            doc.add_paragraph("（1）投标人须提供近3年有6.25兆瓦或以上容量风电机组通过试运行业绩。")
            doc.add_paragraph("（2）投标人须提供近3年超过100台6.25兆瓦或以上容量等级风电机组合同业绩。")
            doc.add_paragraph("标段五（需同时满足）：")
            doc.add_paragraph("（1）投标人须提供近3年单机容量8兆瓦或以上容量等级海上风电机组通过试运行业绩。")
            doc.add_paragraph("3.2.2 资格能力要求：")
            doc.add_paragraph("标段一至标段四（需同时满足）：")
            doc.add_paragraph("（1）投标人需提供任意6.25兆瓦级别风力发电机组完整型式认证一项。")
            doc.add_paragraph("（2）投标机型已取得对应各项目安全等级要求的设计认证。")
            doc.add_paragraph("（3）投标机型已取得整机、叶片及大部件完整型式试验证书及附页材料；若尚未取得，须无条件承诺在本采购项目第一台合同设备供货前取得，需提供承诺书。")
            doc.add_paragraph("标段五（需同时满足）：")
            doc.add_paragraph("（1）投标人需提供任意10兆瓦或以上级别海上风力发电机组完整型式认证一项。")
            doc.add_paragraph("（2）投标机型已取得海上风力发电机组完整型式试验证书及附页材料；若尚未取得，须无条件承诺在本采购项目第一台合同设备供货前取得，需提供承诺书。")
            doc.add_paragraph("3.2.3 本项目不允许联合体投标。")
            doc.add_paragraph("第二章 投标人须知")
            doc.add_paragraph("1.4 投标人资格要求")
            doc.add_paragraph("1.4.1 投标人应具备承担本招标项目资质条件、能力和信誉：见投标人须知前附表。")
            doc.add_paragraph("3.5 资格审查资料")
            doc.add_paragraph("除投标人须知前附表另有规定外，投标人应按下列规定提供资格审查资料，以证明其满足本章第 1.4 款规定的资质、财务、业绩、信誉等要求。")
            doc.add_paragraph("3.5.1 投标人基本情况表应附营业执照复印件。")
            doc.add_paragraph("第三章 评标办法")
            doc.add_paragraph("附表3：商务评分标准表")
            score_table = doc.add_table(rows=2, cols=4)
            for col, text in enumerate(["序号", "评审因素", "分值", "评审标准"]):
                score_table.cell(0, col).text = text
            for col, text in enumerate(["1", "类似合同业绩", "20", "满足最低资格要求的合同业绩数量者得基础分12分，每增加100台加1分。"]):
                score_table.cell(1, col).text = text
            doc.add_paragraph("投标文件应当对招标文件的实质性要求作出响应，否则投标将被否决。")
            doc.save(source_path)

            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-QUAL-SECTION",
                        "bidType": "商务标",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "documents": [
                            {
                                "id": "DOC-QUAL",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            completed = subprocess.run(
                [sys.executable, str(script_path), "offline-fallback", str(manifest_path)],
                check=True,
                capture_output=True,
                text=True,
            )

            self.assertEqual(json.loads(completed.stdout)["schemaVersion"], "bid-business-tender-structured-v1")
            payload = json.loads(output_path.read_text(encoding="utf-8"))
            rows = payload["structured"]["fieldGroups"]["qualificationRequirements"]
            contents = [row["content"] for row in rows]

            self.assertGreaterEqual(len(rows), 8)
            self.assertTrue(any("合法注册的独立法人" in text for text in contents))
            self.assertTrue(any("6.25兆瓦或以上容量风电机组通过试运行业绩" in text for text in contents))
            self.assertTrue(any("超过100台6.25兆瓦" in text for text in contents))
            self.assertTrue(any("8兆瓦或以上容量等级海上风电机组" in text for text in contents))
            self.assertTrue(any("完整型式认证" in text for text in contents))
            self.assertTrue(any("设计认证" in text for text in contents))
            self.assertTrue(any("完整型式试验证书及附页材料" in text for text in contents))
            self.assertTrue(any("不允许联合体投标" in text for text in contents))

            joined = "\n".join(contents)
            self.assertNotIn("投标人财务、信誉等方面应具备下列条件", joined)
            self.assertNotIn("满足最低资格要求的合同业绩数量者得基础分", joined)
            self.assertNotIn("资格审查资料\t23", joined)
            self.assertNotIn("营业执照复印件", joined)
            self.assertNotIn("投标将被否决", joined)
            self.assertNotIn("见投标人须知前附表", joined)

            scoped = [row for row in rows if "6.25兆瓦" in row["content"]]
            self.assertTrue(scoped)
            self.assertTrue(all(row["applicableScope"] == "标段一至标段四" for row in scoped))
            offshore = [row for row in rows if "8兆瓦或以上容量等级海上风电机组" in row["content"]]
            self.assertTrue(offshore)
            self.assertEqual(offshore[0]["applicableScope"], "标段五")

            for row in rows:
                self.assertIn("sourceText", row)
                self.assertNotRegex(row["sourceText"], r"^L\d+$")
                self.assertNotRegex(row["sourceText"], r"B\d+/R\d+")
                self.assertTrue(row["sourceText"].startswith("脱敏资格要求片段.docx："))
                self.assertIn("sourceFile", row)
                self.assertIn("section", row)
                self.assertIn("evidence", row)
                self.assertIn("evidenceLocation", row)

    def test_business_parser_preserves_template_extractor_appendices(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            temp_dir = Path(temp)
            structured_path = temp_dir / "s1_structured_result.json"
            manifest_path = temp_dir / "manifest.json"
            structured_path.write_text(
                json.dumps(
                    {
                        "structured": {
                            "appendices": [
                                {
                                    "id": "APPX-0001",
                                    "title": "附件2 投标价格表\nA投标价格总表\n表1 A-1  标段一",
                                    "artifactType": "business_attachment_template",
                                    "extractionMode": "business_template_extractor_skill",
                                    "docxPath": "C:/tmp/TPL-0001.docx",
                                }
                            ]
                        }
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            combined_path = temp_dir / "combined.txt"
            combined_path.write_text("第六章 投标文件格式\n商务评分 企业业绩 5分", encoding="utf-8")
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "proj-1",
                        "parseProfile": "business",
                        "combinedTextPath": str(combined_path),
                        "structuredResultPath": str(structured_path),
                        "documents": [],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            completed = subprocess.run(
                [sys.executable, str(self.runner_path()), "offline-fallback", str(manifest_path)],
                text=True,
                capture_output=True,
                check=False,
            )

            self.assertEqual(completed.returncode, 0, completed.stderr)
            payload = json.loads(completed.stdout)
            self.assertEqual(payload["targetSkill"], "bid-business-tender-structured-parser")
            result = json.loads(structured_path.read_text(encoding="utf-8"))
            appendices = result["structured"]["appendices"]
            self.assertEqual(len(appendices), 1)
            self.assertEqual(appendices[0]["id"], "APPX-0001")
            self.assertEqual(appendices[0]["extractionMode"], "business_template_extractor_skill")
            self.assertEqual(
                set(result["structured"]["fieldGroups"].keys()),
                {"projectBasics", "qualificationRequirements", "bidderInstructions", "commercialRejectionClauses"},
            )

    def test_business_parser_imports_template_extractor_appendices_from_manifest_path(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            temp_dir = Path(temp)
            template_docx = temp_dir / "TPL-0001.docx"
            Document().save(str(template_docx))
            extraction_path = temp_dir / "business_template_extraction.json"
            extraction_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-template-extractor-v1",
                        "skillName": "bid-business-template-extractor",
                        "summary": {"templateCount": 1},
                        "appendices": [
                            {
                                "id": "APPX-0009",
                                "title": "Bid Letter",
                                "artifactType": "business_attachment_template",
                                "templateType": "bid_letter",
                                "status": "generated",
                                "docxPath": str(template_docx),
                                "sourceDocumentId": "DOC-1",
                                "sourceDocumentName": "business_tender.md",
                                "extractionMode": "business_template_extractor_skill",
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            source_path = temp_dir / "business_tender.md"
            source_path.write_text(
                "项目名称：模板 skill 传递测试\n递交截止时间：2026年1月26日15时00分\n",
                encoding="utf-8",
            )
            structured_path = temp_dir / "s1_structured_result.json"
            manifest_path = temp_dir / "manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-TEMPLATE-SKILL",
                        "parseProfile": "business",
                        "combinedTextPath": str(source_path),
                        "structuredResultPath": str(structured_path),
                        "businessTemplateExtractionPath": str(extraction_path),
                        "documents": [
                            {
                                "id": "DOC-1",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                                "textPath": str(source_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            completed = subprocess.run(
                [sys.executable, str(self.runner_path()), "offline-fallback", str(manifest_path)],
                text=True,
                capture_output=True,
                check=False,
            )

            self.assertEqual(completed.returncode, 0, completed.stderr)
            result = json.loads(structured_path.read_text(encoding="utf-8"))
            appendices = result["structured"]["appendices"]
            self.assertEqual(len(appendices), 1)
            self.assertEqual(appendices[0]["id"], "APPX-0009")
            self.assertEqual(appendices[0]["extractionMode"], "business_template_extractor_skill")
            self.assertEqual(appendices[0]["docxPath"], str(template_docx))

    def test_business_parser_without_real_ai_marks_offline_review_as_untrusted_fallback(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "脱敏商务招标片段.md"
            source_path.write_text(
                "\n".join(
                    [
                        "# 商务招标文件",
                        "项目名称：脱敏风电设备采购项目",
                        "招标编号：BUS-GEN-2026-001",
                        "招标人：示例招标单位",
                        "第一章 招标公告",
                        "3. 投标人资格要求",
                        "3.1 投标人须为中华人民共和国境内合法注册的独立法人或其他组织。",
                        "3.2 投标人须提供近三年同类设备供货业绩。",
                        "3.3 资格审查资料的特殊要求：须附营业执照复印件。",
                        "第二章 投标人须知",
                        "投标文件逾期上传的，招标人不予受理。",
                        "异议材料未按要求提交的，招标人可以不予受理。",
                        "第三章 评标办法",
                        "附表1：商务评分标准表",
                        "| 序号 | 评分项 | 分值 | 得分点 | 证明材料要求 |",
                        "| --- | --- | --- | --- | --- |",
                        "| 1 | 企业业绩 | 20分 | 满足同类业绩要求得分。 | 提供合同。 |",
                        "附表2：技术评分标准表",
                        "| 序号 | 评分项 | 分值 | 得分点 | 证明材料要求 |",
                        "| --- | --- | --- | --- | --- |",
                        "| 1 | 技术方案 | 30分 | 技术方案完整得分。 | 提供技术方案。 |",
                        "附表3：度电成本评分标准表",
                        "| 序号 | 评分项 | 分值 | 得分点 | 证明材料要求 |",
                        "| --- | --- | --- | --- | --- |",
                        "| 1 | 度电成本 | 10分 | 按度电成本排序计分。 | 提供计算表。 |",
                        "第六章 投标文件格式",
                        "投标人须单独提供供货能力承诺函。",
                        "合同条款：卖方承诺按合同约定履行全部义务。",
                    ]
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-BUSINESS-WORKFLOW",
                        "bidType": "商务标",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "documents": [
                            {
                                "id": "DOC-1",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                                "textPath": str(source_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            completed = subprocess.run(
                [sys.executable, str(script_path), "offline-fallback", str(manifest_path)],
                check=True,
                capture_output=True,
                text=True,
            )

            summary = json.loads(completed.stdout)
            self.assertEqual(summary["summary"]["workflowStage"], "fallback")

            candidate_path = tmp_path / "candidate_package.json"
            validation_path = tmp_path / "validation_report.json"
            self.assertTrue(candidate_path.exists())
            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            self.assertTrue(any(item["task"] == "qualification_review" for item in review_plan["tasks"]))
            self.assertTrue(any(item["task"] == "rejection_clause_review" for item in review_plan["tasks"]))
            self.assertFalse(any(item["task"] == "scoring_review" for item in review_plan["tasks"]))
            self.assertFalse(any(item["task"] == "scoring_table_review" for item in review_plan["tasks"]))
            self.assertIn("scoring_table_review", review_plan["skippedAiModules"])
            qualification_ref = next(item for item in review_plan["tasks"] if item["task"] == "qualification_review")
            self.assertTrue((tmp_path / qualification_ref["taskPath"]).exists())
            self.assertTrue((tmp_path / qualification_ref["decisionPath"]).exists())
            self.assertTrue(validation_path.exists())

            candidate_package = json.loads(candidate_path.read_text(encoding="utf-8"))
            self.assertEqual(candidate_package["schemaVersion"], "bid-business-candidate-package-v1")
            self.assertIn("documents", candidate_package)
            self.assertIn("sections", candidate_package)
            self.assertIn("blocks", candidate_package)
            self.assertIn("tables", candidate_package)
            self.assertIn("qualification", candidate_package["candidates"])
            self.assertIn("rejection", candidate_package["candidates"])
            self.assertIn("scoring", candidate_package["candidates"])
            self.assertIn("scoringTableReview", candidate_package["candidates"])
            self.assertTrue(all(block["id"].startswith("DOC-1:") for block in candidate_package["blocks"]))
            self.assertTrue(all("evidenceIds" in item for values in candidate_package["candidates"].values() for item in values))

            decision = json.loads((tmp_path / qualification_ref["decisionPath"]).read_text(encoding="utf-8"))
            self.assertEqual(decision["task"], "qualification_review")
            self.assertIn("accepted", decision)
            self.assertIn("rejected", decision)
            self.assertIn("needsReview", decision)
            self.assertTrue(all(item.get("evidenceIds") for item in decision["accepted"]))
            for bucket in ("accepted", "rejected", "needsReview"):
                for item in decision[bucket]:
                    self.assertTrue(item.get("fieldType"))
                    self.assertIn("content", item)
                    self.assertTrue(item.get("applicableScope"))
                    self.assertTrue(item.get("sourceText"))
                    self.assertTrue(item.get("reason"))
            deterministic = candidate_package["deterministicExtracts"]
            deterministic_text = json.dumps(deterministic["scoringTables"], ensure_ascii=False)
            self.assertEqual([row["scoringItem"] for row in deterministic["scoringTables"]["business"]], ["企业业绩"])
            self.assertNotIn("技术评分", deterministic_text)

            validation = json.loads(validation_path.read_text(encoding="utf-8"))
            self.assertEqual(validation["schemaVersion"], "bid-business-validation-report-v1")
            self.assertEqual(validation["status"], "passed")
            self.assertTrue(all(check["status"] == "passed" for check in validation["checks"]))

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            structured = payload["structured"]
            self.assertEqual(structured["workflow"]["stage"], "fallback")
            self.assertFalse(structured["workflow"]["aiReviewTrusted"])
            self.assertEqual(structured["workflow"]["semanticReviewMode"], "offline-fallback")
            self.assertTrue(structured["workflow"]["offlineAdapterUsed"])
            self.assertEqual(structured["workflow"]["candidatePackagePath"], str(candidate_path))
            self.assertEqual(structured["workflow"]["validationReportPath"], str(validation_path))
            self.assertEqual(set(structured["scoringCriteria"].keys()), {"business"})

            qualifications = structured["fieldGroups"]["qualificationRequirements"]
            qualification_text = "\n".join(row["content"] for row in qualifications)
            self.assertIn("合法注册的独立法人", qualification_text)
            self.assertIn("近三年同类设备供货业绩", qualification_text)
            self.assertNotIn("营业执照复印件", qualification_text)
            self.assertTrue(all(row.get("evidenceIds") for row in qualifications))

            rejections = structured["fieldGroups"]["commercialRejectionClauses"]
            rejection_text = "\n".join(row["content"] for row in rejections)
            self.assertIn("投标文件逾期上传", rejection_text)
            self.assertNotIn("异议材料", rejection_text)

            self.assertEqual(len(structured["scoringCriteria"]["business"]), 1)
            self.assertNotIn("technical", structured["scoringCriteria"])
            self.assertNotIn("lcoe", structured["scoringCriteria"])
            self.assertEqual(structured["commitmentLetters"], [])
            self.assertEqual(structured["commitmentClues"], [])

    def test_business_parser_default_prepare_stage_writes_candidate_package_without_finalizing(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "候选包商务片段.md"
            source_path.write_text(
                "\n".join(
                    [
                        "# 商务招标文件",
                        "项目名称：脱敏设备采购项目",
                        "第一章 招标公告",
                        "3. 投标人资格要求",
                        "3.1 投标人须为中华人民共和国境内合法注册的独立法人或其他组织。",
                    ]
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-PREPARE-ONLY",
                        "bidType": "商务标",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "documents": [
                            {
                                "id": "DOC-1",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                                "textPath": str(source_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            completed = subprocess.run(
                [sys.executable, str(script_path), str(manifest_path)],
                check=True,
                capture_output=True,
                text=True,
            )

            summary = json.loads(completed.stdout)
            self.assertEqual(summary["summary"]["workflowStage"], "prepared")
            self.assertTrue(output_path.exists())
            payload = json.loads(output_path.read_text(encoding="utf-8"))
            self.assertEqual(payload["structured"]["workflow"]["stage"], "prepared")
            self.assertTrue((tmp_path / "candidate_package.json").exists())
            review_plan_path = tmp_path / "review_plan.json"
            self.assertTrue(review_plan_path.exists())
            review_plan = json.loads(review_plan_path.read_text(encoding="utf-8"))
            self.assertEqual(review_plan["schemaVersion"], "bid-business-review-plan-v1")
            self.assertEqual(review_plan["status"], "pending")
            self.assertEqual(review_plan["taskCount"], 1)
            self.assertEqual(review_plan["requiredTaskCount"], 1)
            self.assertEqual([item["task"] for item in review_plan["tasks"]], ["qualification_review"])
            qualification_refs = [item for item in review_plan["tasks"] if item["task"] == "qualification_review"]
            self.assertTrue(qualification_refs)
            self.assertTrue((tmp_path / qualification_refs[0]["taskPath"]).exists())
            self.assertEqual(Path(qualification_refs[0]["taskPath"]).parts[0], "ai_tasks")
            self.assertEqual(Path(qualification_refs[0]["taskPath"]).parts[1], "qualification_review")
            self.assertEqual(Path(qualification_refs[0]["decisionPath"]).parts[0], "ai_decisions")
            self.assertEqual(Path(qualification_refs[0]["decisionPath"]).parts[1], "qualification_review")
            self.assertTrue((tmp_path / qualification_refs[0]["decisionPath"]).parent.is_dir())
            workflow = payload["structured"]["workflow"]
            self.assertEqual(workflow["reviewPlanPath"], str(review_plan_path))
            self.assertEqual(workflow["requiredDecisionTaskCount"], review_plan["taskCount"])
            self.assertIn("deterministicModules", workflow)
            self.assertIn("projectBasics", workflow["deterministicModules"])
            self.assertIn("bidderInstructions", workflow["deterministicModules"])
            self.assertIn("businessScoringTables", workflow["deterministicModules"])
            self.assertEqual(workflow["aiReviewModules"], ["qualification"])
            self.assertEqual(workflow["skippedAiModules"], ["commercialRejectionClauses", "scoring_table_review"])
            self.assertFalse(any((tmp_path / "ai_decisions").rglob("*.json")))
            self.assertFalse((tmp_path / "validation_report.json").exists())

    def test_business_parser_task_helper_commands_are_bounded_and_validate_decision(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "辅助命令商务片段.md"
            source_path.write_text(
                "\n".join(
                    [
                        "# 商务招标文件",
                        "第一章 招标公告",
                        "3. 投标人资格要求",
                        "3.1 投标人须为中华人民共和国境内合法注册的独立法人或其他组织。",
                    ]
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-TASK-HELPERS",
                        "bidType": "商务标",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "documents": [
                            {
                                "id": "DOC-1",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                                "textPath": str(source_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)

            tasks_completed = subprocess.run(
                [sys.executable, str(script_path), "tasks", str(manifest_path)],
                check=True,
                capture_output=True,
                text=True,
            )
            task_listing = json.loads(tasks_completed.stdout)
            self.assertEqual(task_listing["schemaVersion"], "bid-business-task-list-v1")
            self.assertEqual(task_listing["requiredTaskCount"], 1)
            self.assertNotIn("candidates", json.dumps(task_listing, ensure_ascii=False))
            task_ref = next(item for item in task_listing["tasks"] if item["task"] == "qualification_review")

            task_completed = subprocess.run(
                [sys.executable, str(script_path), "task", str(manifest_path), task_ref["taskId"]],
                check=True,
                capture_output=True,
                text=True,
            )
            task_payload = json.loads(task_completed.stdout)
            self.assertEqual(task_payload["taskId"], task_ref["taskId"])
            self.assertTrue(task_payload["candidates"])
            self.assertLess(len(task_completed.stdout), 60000)

            self.write_plan_decision(
                tmp_path / task_ref["taskPath"],
                tmp_path / task_ref["decisionPath"],
            )
            validation_completed = subprocess.run(
                [sys.executable, str(script_path), "validate-decision", str(manifest_path), task_ref["taskId"]],
                check=True,
                capture_output=True,
                text=True,
            )
            validation = json.loads(validation_completed.stdout)
            self.assertEqual(validation["schemaVersion"], "bid-business-decision-validation-v1")
            self.assertEqual(validation["taskId"], task_ref["taskId"])
            self.assertEqual(validation["status"], "passed")
            self.assertEqual(validation["issueCount"], 0)

            status_completed = subprocess.run(
                [sys.executable, str(script_path), "status", str(manifest_path)],
                check=True,
                capture_output=True,
                text=True,
            )
            coverage = json.loads(status_completed.stdout)
            self.assertEqual(coverage["schemaVersion"], "bid-business-review-status-v1")
            self.assertEqual(coverage["presentDecisionTaskCount"], 1)
            self.assertIn(task_ref["taskId"], coverage["presentDecisionTasks"])

    def test_prepare_uses_deterministic_extracts_and_only_needed_ai_tasks(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "脱敏确定性商务片段.docx"
            doc = Document()
            doc.add_paragraph("第一章 招标公告")
            doc.add_paragraph("招标编号：DET-2026-001")
            doc.add_paragraph("投标文件递交截止时间：2026年3月18日09时30分")
            doc.add_paragraph("第二章 投标人须知")
            doc.add_paragraph("投标人须知前附表")
            instruction_table = doc.add_table(rows=5, cols=3)
            for col, text in enumerate(["条款号", "条款名称", "编列内容"]):
                instruction_table.cell(0, col).text = text
            instruction_rows = [
                ("1.1.2", "招标人", "名称：示例招标单位 地址：示例地址"),
                ("1.1.3", "招标代理机构", "名称：示例代理机构 地址：代理地址"),
                ("1.1.4", "招标项目名称", "脱敏确定性采购项目"),
                ("4.2.1", "投标截止时间", "2026年4月1日10时00分"),
            ]
            for row_index, values in enumerate(instruction_rows, start=1):
                for col, text in enumerate(values):
                    instruction_table.cell(row_index, col).text = text
            doc.add_paragraph("评标办法前附表")
            mixed_table = doc.add_table(rows=4, cols=5)
            for col, text in enumerate(["条款号", "评分因素", "分值", "评分标准", "证明材料要求"]):
                mixed_table.cell(0, col).text = text
            mixed_rows = [
                ("2.2.4(1)", "商务评分标准（20分）", "20分", "企业业绩满足要求得20分。", "提供合同。"),
                ("", "技术评分标准（30分）", "30分", "技术方案完整得30分。", "提供技术方案。"),
                ("", "报价评分标准（50分）", "50分", "按评标价计算。", "提供报价表。"),
            ]
            for row_index, values in enumerate(mixed_rows, start=1):
                for col, text in enumerate(values):
                    mixed_table.cell(row_index, col).text = text
            doc.add_paragraph("第三章 评标办法")
            doc.add_paragraph("综合评分表")
            ambiguous_table = doc.add_table(rows=2, cols=4)
            for col, text in enumerate(["序号", "评审因素", "分值", "评分标准"]):
                ambiguous_table.cell(0, col).text = text
            for col, text in enumerate(["1", "服务响应", "5分", "响应完整得分。"]):
                ambiguous_table.cell(1, col).text = text
            doc.save(source_path)

            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-DETERMINISTIC",
                        "bidType": "商务标",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "documents": [
                            {
                                "id": "DOC-1",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            structured = payload["structured"]
            workflow = structured["workflow"]
            self.assertEqual(workflow["stage"], "prepared")
            self.assertEqual(workflow["deterministicModules"], ["projectBasics", "bidderInstructions", "businessScoringTables"])
            self.assertEqual(workflow["aiReviewModules"], [])
            self.assertIn("qualification", workflow["skippedAiModules"])
            self.assertIn("scoring_table_review", workflow["skippedAiModules"])

            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            deterministic = candidate_package["deterministicExtracts"]
            self.assertEqual(field_by_key(deterministic["projectBasics"], "projectName")["value"], "脱敏确定性采购项目")
            self.assertEqual(field_by_key(deterministic["projectBasics"], "tenderNo")["value"], "DET-2026-001")
            self.assertEqual(field_by_key(deterministic["projectBasics"], "tenderer")["value"], "示例招标单位")
            self.assertEqual(field_by_key(deterministic["projectBasics"], "tenderAgency")["value"], "示例代理机构")
            self.assertEqual(field_by_key(deterministic["projectBasics"], "bidDeadline")["value"], "2026-04-01 10:00")
            self.assertEqual(len(deterministic["bidderInstructions"]), 4)
            self.assertEqual([row["scoringItem"] for row in deterministic["scoringTables"]["business"]], ["商务评分标准（20分）"])
            self.assertNotIn("price", deterministic["scoringTables"])
            rejected_text = json.dumps(deterministic["scoringTables"], ensure_ascii=False)
            self.assertNotIn("技术评分标准", rejected_text)
            self.assertNotIn("综合评分表", rejected_text)

            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            self.assertFalse(any(task["task"] == "scoring_table_review" for task in review_plan["tasks"]))

    def test_bidder_instruction_locator_stops_at_first_target_table(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "prj0012-like.docx"
            doc = Document()
            doc.add_paragraph("目录")
            doc.add_paragraph("投标人须知前附表 11")
            doc.add_paragraph("评标办法前附表 33")
            doc.add_paragraph("第二章 投标人须知")
            doc.add_paragraph("投标人须知前附表")
            instruction_table = doc.add_table(rows=4, cols=3)
            for col, text in enumerate(["条款号", "条款名称", "编列内容"]):
                instruction_table.cell(0, col).text = text
            instruction_rows = [
                ("1.1.2", "招标人", "示例招标人"),
                ("1.1.3", "招标代理机构", "示例代理机构"),
                ("4.2.1", "投标截止时间", "2026年4月1日09时00分"),
            ]
            for row_index, values in enumerate(instruction_rows, start=1):
                for col, text in enumerate(values):
                    instruction_table.cell(row_index, col).text = text
            doc.add_paragraph("11. 需要补充的其他内容")
            doc.add_paragraph("见投标人须知前附表。")
            doc.add_paragraph("第三章 评标办法(综合评估法)")
            doc.add_paragraph("评标办法前附表")
            eval_table = doc.add_table(rows=3, cols=3)
            for col, text in enumerate(["条款号", "条款内容", "编列内容"]):
                eval_table.cell(0, col).text = text
            eval_rows = [
                ("1.2", "中标候选人推荐原则", "采用综合评估法。"),
                ("2.2.1", "分值构成", "商务部分10分。"),
            ]
            for row_index, values in enumerate(eval_rows, start=1):
                for col, text in enumerate(values):
                    eval_table.cell(row_index, col).text = text
            doc.save(source_path)

            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-BIDDER-INSTRUCTION-ONLY",
                        "bidType": "商务标",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "documents": [
                            {
                                "id": "DOC-1",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)
            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            self.write_bid_deadline_ai_decision(tmp_path, review_plan, content="2026年5月6日09时30分")
            subprocess.run([sys.executable, str(script_path), "finalize", str(manifest_path)], check=True, capture_output=True, text=True)

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            rows = payload["structured"]["fieldGroups"]["bidderInstructions"]
            self.assertEqual(len(rows), 3)
            self.assertEqual({row["section"] for row in rows}, {"投标人须知前附表"})
            self.assertNotIn("中标候选人推荐原则", "\n".join(row["content"] for row in rows))
            self.assertEqual(rows[0]["headers"], ["条款号", "条款名称", "编列内容"])
            self.assertEqual(rows[0]["cells"], ["1.1.2", "招标人", "示例招标人"])

            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            deterministic_rows = candidate_package["deterministicExtracts"]["bidderInstructions"]
            self.assertEqual(len(deterministic_rows), 3)
            self.assertTrue(all(row["tableTitle"] == "投标人须知前附表" for row in deterministic_rows))

    def test_bidder_instruction_fallback_uses_nearby_title_and_dynamic_headers(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "dynamic-bidder-instructions.md"
            source_path.write_text(
                "\n".join(
                    [
                        "# 目录",
                        "投标人须知前附表 12",
                        "",
                        "# 第二章 投标人须知",
                        "投标人须知前附表",
                        "说明：下表为本项目专用条款。",
                        "单位：人民币",
                        "| 序号 | 事项 | 要求 | 备注 |",
                        "| --- | --- | --- | --- |",
                        "| 1 | 招标人 | 示例招标人 | 以公告为准 |",
                        "| 2 | 递交截止时间 | 2026年4月1日09时00分 | 电子平台递交 |",
                        "",
                        "# 第三章 评标办法",
                        "评标办法前附表",
                        "| 条款号 | 条款内容 | 编列内容 |",
                        "| --- | --- | --- |",
                        "| 1.2 | 中标候选人推荐原则 | 采用综合评估法 |",
                    ]
                ),
                encoding="utf-8",
            )

            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-DYNAMIC-BIDDER-INSTRUCTION-HEADER",
                        "bidType": "商务标",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "documents": [
                            {
                                "id": "DOC-1",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                                "textPath": str(source_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)
            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            self.write_bid_deadline_ai_decision(tmp_path, review_plan, content="2026年5月6日09时30分")
            subprocess.run([sys.executable, str(script_path), "finalize", str(manifest_path)], check=True, capture_output=True, text=True)

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            rows = payload["structured"]["fieldGroups"]["bidderInstructions"]
            self.assertEqual(len(rows), 2)
            self.assertEqual(rows[0]["headers"], ["序号", "事项", "要求", "备注"])
            self.assertEqual(rows[0]["cells"], ["1", "招标人", "示例招标人", "以公告为准"])
            self.assertEqual(rows[0]["clauseNo"], "1")
            self.assertEqual(rows[0]["clauseName"], "招标人")
            self.assertEqual(rows[0]["content"], "示例招标人；以公告为准")
            self.assertNotIn("中标候选人推荐原则", "\n".join(row["content"] for row in rows))

    def test_exact_business_scoring_table_is_deterministic_and_skips_scoring_ai(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "exact-business-scoring.docx"
            doc = Document()
            doc.add_paragraph("第三章 评标办法")
            doc.add_paragraph("评标办法前附表")
            preface_table = doc.add_table(rows=4, cols=3)
            for col, text in enumerate(["条款号", "条款内容", "编列内容"]):
                preface_table.cell(0, col).text = text
            preface_rows = [
                ("1.2", "中标候选人推荐原则", "采用综合评估法，其中商务权重为10%，按最终得分推荐中标候选人。"),
                ("2.2.3", "商务评审标准", "续表3：商务部分评审评分标准表"),
                ("3.2.2.1", "确定参与报价评审的条件", "技术和商务得分均不小于80分，参加报价评审。"),
            ]
            for row_index, values in enumerate(preface_rows, start=1):
                for col, text in enumerate(values):
                    preface_table.cell(row_index, col).text = text

            doc.add_paragraph("附表3：商务评分标准表")
            scoring_table = doc.add_table(rows=3, cols=4)
            for col, text in enumerate(["序号", "评审因素", "分值", "评审标准"]):
                scoring_table.cell(0, col).text = text
            scoring_rows = [
                ("1", "企业综合实力", "15", "根据企业资质、综合竞争力等情况综合评审，最高得15分。"),
                ("2", "投标文件完整性", "4", "投标文件编制条理清楚、完整性强者得2-4分。"),
            ]
            for row_index, values in enumerate(scoring_rows, start=1):
                for col, text in enumerate(values):
                    scoring_table.cell(row_index, col).text = text
            doc.save(source_path)

            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-EXACT-BUSINESS-SCORING",
                        "bidType": "商务标",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "documents": [
                            {
                                "id": "DOC-1",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)

            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            deterministic_rows = candidate_package["deterministicExtracts"]["scoringTables"]["business"]
            self.assertEqual([row["scoringItem"] for row in deterministic_rows], ["企业综合实力", "投标文件完整性"])
            self.assertEqual(candidate_package["candidates"]["scoringTableReview"], [])

            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            self.assertFalse(any(task["task"] == "scoring_table_review" for task in review_plan["tasks"]))
            self.assertIn("scoring_table_review", review_plan["skippedAiModules"])

    def test_business_part_review_detail_table_is_deterministic_and_skips_scoring_ai(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "business-part-review-detail.docx"
            doc = Document()
            doc.add_paragraph("第三章 评标办法")
            doc.add_paragraph("表 2：商务部分评审细则（满分 10 分）")
            scoring_table = doc.add_table(rows=5, cols=4)
            for col, text in enumerate(["序号", "评审项目", "分 值", "评分细则"]):
                scoring_table.cell(0, col).text = text
            scoring_rows = [
                ("一", "售后服务、技术服务方案", "2", "供应商提供的售后服务承诺及措施完整，方案合理并具有可操作性，最优得满分，最低不得分。"),
                ("二", "综合费用合理性", "2", "对运行维护费用、备品备件、塔筒造价及其他特殊情况增加费用等进行综合评比，最优得满分，最低不得分。"),
                ("三", "同机型设备业绩", "2", "近三年响应机型批量并网和订货业绩清晰有效，本项最优得满分，最低不得分。"),
                ("四", "战略合作", "4", "按照战略合作项目数量和规模两方面评分，本项最高得4分，最低不得分。"),
            ]
            for row_index, values in enumerate(scoring_rows, start=1):
                for col, text in enumerate(values):
                    scoring_table.cell(row_index, col).text = text
            doc.add_paragraph("表 3 技术部分评分细则（满分 35 分）")
            technical_table = doc.add_table(rows=2, cols=4)
            for col, text in enumerate(["序号", "评审项目", "分 值", "评分细则"]):
                technical_table.cell(0, col).text = text
            for col, text in enumerate(["一", "响应机型成熟度", "3", "根据响应机型成熟度评分。"]):
                technical_table.cell(1, col).text = text
            doc.save(source_path)

            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-BUSINESS-PART-REVIEW-DETAIL",
                        "bidType": "商务标",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "documents": [
                            {
                                "id": "DOC-1",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)

            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            deterministic_rows = candidate_package["deterministicExtracts"]["scoringTables"]["business"]
            self.assertEqual(
                [row["scoringItem"] for row in deterministic_rows],
                ["售后服务、技术服务方案", "综合费用合理性", "同机型设备业绩", "战略合作"],
            )
            self.assertEqual([row["score"] for row in deterministic_rows], ["2", "2", "2", "4"])
            self.assertEqual(candidate_package["candidates"]["scoringTableReview"], [])

            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            self.assertFalse(any(task["task"] == "scoring_table_review" for task in review_plan["tasks"]))
            self.assertIn("scoring_table_review", review_plan["skippedAiModules"])

    def test_exact_business_scoring_embedded_table_without_score_header_goes_to_ai_row_block(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "embedded-business-scoring-no-score-header.docx"
            doc = Document()
            doc.add_paragraph("第三章 评标办法")
            doc.add_paragraph("评标办法前附表")
            table = doc.add_table(rows=5, cols=4)
            for col, text in enumerate(["条款号", "条款号", "评分因素", "评分标准"]):
                table.cell(0, col).text = text
            rows = [
                ("2.2.4（1）", "商务评分标准（10分）", "交货期保证（2分）", "满分2分，交货期满足招标文件要求得2分。"),
                ("2.2.4（1）", "商务评分标准（10分）", "企业财务状况（3分）", "满分3分，根据近三年财务状况得0至3分。"),
                ("2.2.4（1）", "商务评分标准（10分）", "供货业绩（5分）", "满分5分，每提供一项有效业绩得1分。"),
                ("2.2.4（2）", "技术评分标准（30分）", "技术支持（4分）", "满分4分，技术支持方案完整得4分。"),
            ]
            for row_index, values in enumerate(rows, start=1):
                for col, text in enumerate(values):
                    table.cell(row_index, col).text = text
            doc.save(source_path)

            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-EMBEDDED-BUSINESS-SCORING-NO-SCORE-HEADER",
                        "bidType": "商务标",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "documents": [
                            {
                                "id": "DOC-1",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)

            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            self.assertEqual(candidate_package["deterministicExtracts"]["scoringTables"]["business"], [])
            row_block_candidates = candidate_package["candidates"]["scoringTableReview"]
            self.assertEqual(len(row_block_candidates), 1)
            self.assertEqual(row_block_candidates[0]["candidateType"], "business_scoring_row_block_review")
            self.assertEqual(len(row_block_candidates[0]["evidenceIds"]), 3)
            self.assertTrue(all("/R" in evidence_id for evidence_id in row_block_candidates[0]["evidenceIds"]))
            self.assertIn("交货期保证（2分）", row_block_candidates[0]["content"])
            self.assertNotIn("技术支持", row_block_candidates[0]["content"])

            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            scoring_ref = next(task for task in review_plan["tasks"] if task["task"] == "scoring_table_review")
            subprocess.run(
                [
                    sys.executable,
                    str(script_path),
                    "decision-all",
                    str(manifest_path),
                    scoring_ref["taskId"],
                    "accepted",
                    "business",
                    "unit test accepts exact-anchor embedded business scoring rows",
                ],
                check=True,
                capture_output=True,
                text=True,
            )
            subprocess.run([sys.executable, str(script_path), "finalize", str(manifest_path)], check=True, capture_output=True, text=True)

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            business_rows = payload["structured"]["scoringCriteria"]["business"]
            self.assertEqual([row["scoringItem"] for row in business_rows], ["交货期保证", "企业财务状况", "供货业绩"])
            self.assertEqual([row["score"] for row in business_rows], ["2分", "3分", "5分"])
            self.assertTrue(all("商务评分标准" not in row["scoringItem"] for row in business_rows))
            self.assertFalse(any("技术评分标准" in row["evidence"] for row in business_rows))

    def test_mixed_preface_business_scoring_block_excludes_score_composition_and_price_rows(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "mixed-preface-business-scoring.docx"
            doc = Document()
            doc.add_paragraph("第三章 评标办法")
            doc.add_paragraph("评标办法前附表")
            table = doc.add_table(rows=11, cols=4)
            for col, text in enumerate(["条款号", "评分因素（偏差率）", "评分项", "评分标准"]):
                table.cell(0, col).text = text
            rows = [
                ("2.2.1", "分值构成（总分100分）", "", "投标报价：45分；商务部分：10分；技术部分：45分。"),
                ("2.2.2", "评标基准价计算方法", "", "按有效投标报价算术平均值计算。"),
                ("2.2.3", "投标报价的偏差率计算公式", "", "偏差率=（评标价－评标基准价）/评标基准价×100%。"),
                ("2.2.4（1）", "投标报价评分标准", "45分", "评标价等于评标基准价，得基础分35分。"),
                ("2.2.4（2）", "商务评分标准", "业绩（5分）", "在满足资格要求的基础上，每增加1000MW加1分，最多得5分。"),
                ("2.2.4（2）", "商务评分标准", "企业综合实力及财务状况（2分）", "根据企业实力及近三年财务状况横向对比，优秀得2分。"),
                ("2.2.4（2）", "商务评分标准", "主要商务条款响应程度（2分）", "主要商务条款全部响应招标文件要求得2分，否则酌情扣分。"),
                ("2.2.4（2）", "商务评分标准", "投标文件编制质量（1分）", "投标文件清晰完整、按招标文件要求格式编制得1分。"),
                ("2.2.4（3）", "技术评分标准", "投标方案先进性（10分）", "根据投标方案先进性横向对比评分。"),
                ("2.2.4（4）", "其他因素评分标准", "/", "无。"),
            ]
            for row_index, values in enumerate(rows, start=1):
                for col, text in enumerate(values):
                    table.cell(row_index, col).text = text
            doc.save(source_path)

            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-MIXED-PREFACE-BUSINESS-SCORING",
                        "bidType": "商务标",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "documents": [
                            {
                                "id": "DOC-1",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)

            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            self.assertEqual(candidate_package["deterministicExtracts"]["scoringTables"]["business"], [])
            row_block_candidates = candidate_package["candidates"]["scoringTableReview"]
            self.assertEqual(len(row_block_candidates), 1)
            content = row_block_candidates[0]["content"]
            self.assertIn("业绩（5分）", content)
            self.assertIn("投标文件编制质量（1分）", content)
            self.assertNotIn("分值构成", content)
            self.assertNotIn("投标报价评分标准", content)
            self.assertNotIn("技术评分标准", content)
            self.assertEqual(len(row_block_candidates[0]["evidenceIds"]), 4)

    def test_preface_scoring_noise_is_not_appended_without_concrete_scores(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "preface-scoring-noise.docx"
            doc = Document()
            doc.add_paragraph("第三章 评标办法")
            doc.add_paragraph("评标办法前附表")
            table = doc.add_table(rows=4, cols=3)
            for col, text in enumerate(["条款号", "条款内容", "编列内容"]):
                table.cell(0, col).text = text
            rows = [
                ("1.2", "中标候选人推荐原则", "采用综合评估法，其中价格、技术和商务权重分别为50%、40%、10%，按最终得分推荐中标候选人。"),
                ("2.2.1", "分值构成", "商务部分B满分100分，权重K2 10%；综合得分=A+B+C。"),
                ("3.2.2.1", "确定参与报价评审的条件", "技术和商务得分均不小于80分，参加报价评审。"),
            ]
            for row_index, values in enumerate(rows, start=1):
                for col, text in enumerate(values):
                    table.cell(row_index, col).text = text
            doc.save(source_path)

            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-PREFACE-SCORING-NOISE",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)

            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            self.assertFalse(any(task["task"] == "scoring_table_review" for task in review_plan["tasks"]))
            for task_ref in review_plan["tasks"]:
                self.write_plan_decision(tmp_path / task_ref["taskPath"], tmp_path / task_ref["decisionPath"])

            subprocess.run([sys.executable, str(script_path), "finalize", str(manifest_path)], check=True, capture_output=True, text=True)

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            scoring = payload["structured"]["scoringCriteria"]
            self.assertEqual(scoring["business"], [])
            self.assertNotIn("中标候选人推荐原则", json.dumps(scoring, ensure_ascii=False))
            self.assertNotIn("确定参与报价评审的条件", json.dumps(scoring, ensure_ascii=False))

    def test_business_parser_finalize_without_ai_decisions_reports_untrusted_fallback(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "缺少审查商务片段.md"
            source_path.write_text(
                "\n".join(
                    [
                        "# 商务招标文件",
                        "第一章 招标公告",
                        "3. 投标人资格要求",
                        "3.1 投标人须为中华人民共和国境内合法注册的独立法人或其他组织。",
                    ]
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-MISSING-AI-DECISIONS",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                        "textPath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)
            completed = subprocess.run(
                [sys.executable, str(script_path), "finalize", str(manifest_path)],
                check=True,
                capture_output=True,
                text=True,
            )

            summary = json.loads(completed.stdout)
            self.assertEqual(summary["summary"]["workflowStage"], "fallback")
            self.assertFalse(any((tmp_path / "ai_decisions").rglob("*.json")))
            payload = json.loads(output_path.read_text(encoding="utf-8"))
            workflow = payload["structured"]["workflow"]
            self.assertEqual(workflow["stage"], "fallback")
            self.assertFalse(workflow["aiReviewTrusted"])
            self.assertEqual(workflow["semanticReviewMode"], "opencode-agent")
            self.assertEqual(workflow["requiredDecisionTaskCount"], 1)
            self.assertEqual(workflow["presentDecisionTaskCount"], 0)
            self.assertTrue(any(task.startswith("qualification_review/part-") for task in workflow["missingDecisionTasks"]))
            validation = json.loads((tmp_path / "validation_report.json").read_text(encoding="utf-8"))
            coverage_check = next(check for check in validation["checks"] if check["name"] == "ai_decision_coverage")
            self.assertEqual(coverage_check["status"], "failed")
            self.assertEqual(validation["aiDecisionCoverage"]["status"], "failed")
            self.assertEqual(validation["aiDecisionCoverage"]["requiredDecisionTaskCount"], workflow["requiredDecisionTaskCount"])
            self.assertEqual(validation["aiDecisionCoverage"]["presentDecisionTaskCount"], 0)
            self.assertEqual(validation["aiDecisionCoverage"]["missingDecisionTasks"], workflow["missingDecisionTasks"])

    def test_business_parser_decision_all_helper_rejects_qualification_auto_split(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "business-source.md"
            source_path.write_text(
                "\n".join(
                    [
                        "# \u5546\u52a1\u62db\u6807\u6587\u4ef6",
                        "\u7b2c\u4e00\u7ae0 \u62db\u6807\u516c\u544a",
                        "3. \u6295\u6807\u4eba\u8d44\u683c\u8981\u6c42",
                        "3.1 \u6295\u6807\u4eba\u987b\u4e3a\u4e2d\u534e\u4eba\u6c11\u5171\u548c\u56fd\u5883\u5185\u5408\u6cd5\u6ce8\u518c\u7684\u72ec\u7acb\u6cd5\u4eba\u3002",
                    ]
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-DECISION-ALL",
                "bidType": "\u5546\u52a1\u6807",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                        "textPath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)
            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            task_id = review_plan["tasks"][0]["taskId"]
            decision_path = tmp_path / review_plan["tasks"][0]["decisionPath"]

            completed = subprocess.run(
                [
                    sys.executable,
                    str(script_path),
                    "decision-all",
                    str(manifest_path),
                    task_id,
                    "rejected",
                    "process_note",
                    "unit test rejects every candidate without shell heredoc",
                ],
                check=True,
                capture_output=True,
                text=True,
            )

            write_summary = json.loads(completed.stdout)
            self.assertEqual(write_summary["status"], "failed")
            self.assertEqual(write_summary["decisionPath"], str(decision_path))
            self.assertEqual(write_summary["issues"][0]["code"], "qualification_requires_raw_ai_items")
            self.assertFalse(decision_path.exists())

    def test_business_parser_decision_set_helper_rejects_qualification_auto_split(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "business-source.md"
            source_path.write_text(
                "\n".join(
                    [
                        "# \u5546\u52a1\u62db\u6807\u6587\u4ef6",
                        "\u7b2c\u4e00\u7ae0 \u62db\u6807\u516c\u544a",
                        "3. \u6295\u6807\u4eba\u8d44\u683c\u8981\u6c42",
                        "3.1 \u6295\u6807\u4eba\u987b\u4e3a\u4e2d\u534e\u4eba\u6c11\u5171\u548c\u56fd\u5883\u5185\u5408\u6cd5\u6ce8\u518c\u7684\u72ec\u7acb\u6cd5\u4eba\u3002",
                        "3.2 \u6295\u6807\u4eba\u987b\u5177\u6709\u8fd1\u4e09\u5e74\u540c\u7c7b\u4e1a\u7ee9\u3002",
                    ]
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-DECISION-SET",
                "bidType": "\u5546\u52a1\u6807",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                        "textPath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)
            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            task_id = review_plan["tasks"][0]["taskId"]
            task_payload = json.loads((tmp_path / review_plan["tasks"][0]["taskPath"]).read_text(encoding="utf-8"))
            accepted_id = task_payload["candidates"][0]["candidateId"]

            completed = subprocess.run(
                [
                    sys.executable,
                    str(script_path),
                    "decision-set",
                    str(manifest_path),
                    task_id,
                    accepted_id,
                    "",
                    "",
                    "rejected",
                    "qualification_requirement",
                    "unit test accepts selected candidate id only",
                ],
                check=True,
                capture_output=True,
                text=True,
            )

            write_result = json.loads(completed.stdout)
            self.assertEqual(write_result["status"], "failed")
            self.assertEqual(write_result["issues"][0]["code"], "qualification_requires_raw_ai_items")
            self.assertFalse((tmp_path / review_plan["tasks"][0]["decisionPath"]).exists())

    def test_business_parser_qualification_item_helper_preserves_ai_raw_content(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "business-source.md"
            source_path.write_text(
                "\n".join(
                    [
                        "# 商务招标文件",
                        "第一章 招标公告",
                        "3. 投标人资格要求",
                        "3.1 投标人须为中华人民共和国境内合法注册的独立法人。",
                    ]
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-QUALIFICATION-RAW-AI",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                        "textPath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)
            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            task_ref = next(task for task in review_plan["tasks"] if task["task"] == "qualification_review")
            task_payload = json.loads((tmp_path / task_ref["taskPath"]).read_text(encoding="utf-8"))
            source_line = next(line for line in task_payload["candidates"][0]["lines"] if "独立法人" in line["text"])
            ai_content = "3.1.1 AI原文条款：投标人须为中华人民共和国境内合法注册的独立法人。"

            completed = subprocess.run(
                [
                    sys.executable,
                    str(script_path),
                    "qualification-item",
                    str(manifest_path),
                    task_ref["taskId"],
                    ai_content,
                    "全部标段",
                    source_line["evidenceId"],
                    source_line["sourceText"],
                    "AI 原始拆分结果。",
                ],
                check=True,
                capture_output=True,
                text=True,
            )

            self.assertEqual(json.loads(completed.stdout)["status"], "passed")
            decision = json.loads((tmp_path / task_ref["decisionPath"]).read_text(encoding="utf-8"))
            self.assertEqual(decision["qualificationItems"][0]["content"], ai_content)

            subprocess.run([sys.executable, str(script_path), "finalize", str(manifest_path)], check=True, capture_output=True, text=True)
            payload = json.loads(output_path.read_text(encoding="utf-8"))
            qualifications = payload["structured"]["fieldGroups"]["qualificationRequirements"]
            self.assertEqual(qualifications[0]["content"], ai_content)

    def test_business_parser_scoring_table_review_decision_materializes_rows(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "ambiguous-scoring.docx"
            doc = Document()
            doc.add_paragraph("\u7b2c\u4e09\u7ae0 \u8bc4\u6807\u529e\u6cd5")
            doc.add_paragraph("\u8bc4\u6807\u529e\u6cd5\u524d\u9644\u8868")
            table = doc.add_table(rows=4, cols=4)
            for col, text in enumerate(["\u6761\u6b3e\u53f7", "\u8bc4\u5ba1\u56e0\u7d20", "\u5206\u503c", "\u8bc4\u5ba1\u6807\u51c6"]):
                table.cell(0, col).text = text
            rows = [
                ("2.2.4\uff081\uff09", "\u5546\u52a1\u8bc4\u5ba1\u6807\u51c6", "10", "\u4ea4\u8d27\u671f\u548c\u552e\u540e\u670d\u52a1\u54cd\u5e94\u6e05\u6670\u3002"),
                ("", "\u4ed8\u6b3e\u54cd\u5e94", "5", "\u4ed8\u6b3e\u6761\u4ef6\u54cd\u5e94\u5b8c\u6574\u3002"),
                ("2.2.4\uff082\uff09", "\u6280\u672f\u8bc4\u5ba1\u6807\u51c6", "30", "\u6280\u672f\u65b9\u6848\u5b8c\u6574\u3002"),
            ]
            for row_index, values in enumerate(rows, start=1):
                for col, text in enumerate(values):
                    table.cell(row_index, col).text = text
            doc.save(source_path)

            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-SCORING-TABLE-REVIEW",
                        "bidType": "\u5546\u52a1\u6807",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "documents": [
                            {
                                "id": "DOC-1",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)
            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            self.assertEqual(review_plan["taskCount"], 1)
            self.assertEqual(review_plan["tasks"][0]["task"], "scoring_table_review")
            task_id = review_plan["tasks"][0]["taskId"]
            task_payload = json.loads((tmp_path / review_plan["tasks"][0]["taskPath"]).read_text(encoding="utf-8"))
            self.assertIn("\u5177\u4f53\u5206\u503c", task_payload["instruction"])
            self.assertEqual(len(task_payload["candidates"]), 1)
            self.assertEqual(len(task_payload["candidates"][0]["evidenceIds"]), 2)
            self.assertIn("\u4ed8\u6b3e\u54cd\u5e94", task_payload["candidates"][0]["content"])
            self.assertNotIn("\u6280\u672f\u8bc4\u5ba1\u6807\u51c6", task_payload["candidates"][0]["content"])

            subprocess.run(
                [
                    sys.executable,
                    str(script_path),
                    "decision-all",
                    str(manifest_path),
                    task_id,
                    "accepted",
                    "business",
                    "unit test classifies the ambiguous scoring table as business scoring",
                ],
                check=True,
                capture_output=True,
                text=True,
            )
            completed = subprocess.run(
                [sys.executable, str(script_path), "finalize", str(manifest_path)],
                check=True,
                capture_output=True,
                text=True,
            )

            self.assertEqual(json.loads(completed.stdout)["summary"]["workflowStage"], "finalized")
            payload = json.loads(output_path.read_text(encoding="utf-8"))
            business_rows = payload["structured"]["scoringCriteria"]["business"]
            self.assertEqual([row["scoringItem"] for row in business_rows], ["\u5546\u52a1\u8bc4\u5ba1\u6807\u51c6", "\u4ed8\u6b3e\u54cd\u5e94"])
            self.assertFalse(any("\u6280\u672f" in row["evidence"] for row in business_rows))

    def test_project_basics_are_from_bidder_instruction_table_not_body_candidates(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "project-basics.docx"
            doc = Document()
            doc.add_paragraph("第一章 招标公告")
            doc.add_paragraph("招标编号：BODY-2026-999")
            doc.add_paragraph("招标人不接受未购买招标文件的投标人参加投标。")
            doc.add_paragraph("投标文件递交截止时间：2026年4月20日 09时30分")
            doc.add_paragraph("第二章 投标人须知")
            doc.add_paragraph("投标人须知前附表")
            table = doc.add_table(rows=5, cols=3)
            for col, text in enumerate(["条款号", "条款名称", "编列内容"]):
                table.cell(0, col).text = text
            rows = [
                ("1.1.2", "招标人", "名称：前附表招标单位 地址：示例地址"),
                ("1.1.3", "招标代理机构", "名称：前附表代理机构 地址：代理地址"),
                ("1.1.4", "招标项目名称", "前附表结构化项目"),
                ("1.1.5", "招标编号", "TABLE-2026-001"),
            ]
            for row_index, values in enumerate(rows, start=1):
                for col, text in enumerate(values):
                    table.cell(row_index, col).text = text
            doc.add_paragraph("第三章 评标办法")
            doc.add_paragraph("商务评分标准")
            score_table = doc.add_table(rows=2, cols=4)
            for col, text in enumerate(["序号", "评分项", "分值", "评分标准"]):
                score_table.cell(0, col).text = text
            for col, text in enumerate(["1", "企业业绩", "20分", "满足业绩要求得分。"]):
                score_table.cell(1, col).text = text
            doc.save(source_path)

            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-PROJECT-BASICS-TABLE",
                        "bidType": "商务标",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "documents": [
                            {
                                "id": "DOC-1",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            subprocess.run(
                [sys.executable, str(script_path), "offline-fallback", str(manifest_path)],
                check=True,
                capture_output=True,
                text=True,
            )

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            project_basics = payload["structured"]["fieldGroups"]["projectBasics"]
            self.assertEqual(field_by_key(project_basics, "projectName")["value"], "前附表结构化项目")
            self.assertEqual(field_by_key(project_basics, "tenderNo")["value"], "TABLE-2026-001")
            self.assertEqual(field_by_key(project_basics, "tenderer")["value"], "前附表招标单位")
            self.assertEqual(field_by_key(project_basics, "tenderAgency")["value"], "前附表代理机构")
            self.assertEqual(field_by_key(project_basics, "bidDeadline")["value"], "2026-04-20 09:30")
            self.assertIn("投标文件递交截止时间", field_by_key(project_basics, "bidDeadline")["evidence"])

            fact_by_key = {field["fieldKey"]: field for field in payload["structured"]["projectFactFields"]}
            self.assertEqual(fact_by_key["tenderer"]["value"], "前附表招标单位")
            self.assertEqual(fact_by_key["bidDeadline"]["value"], "2026-04-20 09:30")
            self.assertTrue(field_by_key(project_basics, "bidDeadline").get("evidenceIds"))

            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            self.assertEqual(field_by_key(candidate_package["deterministicExtracts"]["projectBasics"], "tenderer")["value"], "前附表招标单位")
            self.assertEqual(len(candidate_package["deterministicExtracts"]["bidderInstructions"]), 4)

    def test_project_basics_resolve_preface_references_to_announcement_section_tree(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "preface-reference-announcement.md"
            lines = [
                "# 商务招标文件",
                "第一章 招标公告",
                "1. 招标条件",
                "招标项目名称：公告真实项目",
                "招标人：公告真实招标单位",
                "招标代理机构：公告真实代理机构",
                "投标文件递交截止时间：2026年5月6日09时30分",
                "第二章 投标人须知",
                "投标人须知前附表",
                "| 条款号 | 条款名称 | 编列内容 |",
                "| --- | --- | --- |",
                "| 1.1.2 | 招标人 | 见招标公告 |",
                "| 1.1.3 | 招标代理机构 | 详见招标公告 |",
                "| 1.1.4 | 招标项目名称 | 见招标公告 |",
                "| 4.2.1 | 投标截止时间 | 详见招标公告 |",
                "第三章 评标办法",
            ]
            source_path.write_text("\n".join(lines), encoding="utf-8")
            section_tree_path = tmp_path / "business_section_tree.json"
            section_tree_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-section-tree-v1",
                        "documents": [{"id": "DOC-1", "name": source_path.name}],
                        "nodes": [
                            {
                                "id": "DOC-1-S0001",
                                "documentId": "DOC-1",
                                "level": 1,
                                "number": "第一章",
                                "title": "第一章 招标公告",
                                "path": ["第一章 招标公告"],
                                "source": "heading",
                                "confidence": 0.95,
                                "startBlockIndex": 2,
                                "contentStartBlockIndex": 3,
                                "endBlockIndex": 7,
                                "startLine": 2,
                                "contentStartLine": 3,
                                "endLine": 7,
                            },
                            {
                                "id": "DOC-1-S0002",
                                "documentId": "DOC-1",
                                "level": 2,
                                "number": "",
                                "title": "投标人须知前附表",
                                "path": ["第二章 投标人须知", "投标人须知前附表"],
                                "source": "heading",
                                "confidence": 0.95,
                                "startBlockIndex": 9,
                                "contentStartBlockIndex": 10,
                                "endBlockIndex": 15,
                                "startLine": 9,
                                "contentStartLine": 10,
                                "endLine": 15,
                            },
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-PROJECT-BASICS-REFERENCE",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "businessSectionTreePath": str(section_tree_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                        "textPath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)
            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            self.write_bid_deadline_ai_decision(tmp_path, review_plan, content="2026年5月6日09时30分")
            subprocess.run([sys.executable, str(script_path), "finalize", str(manifest_path)], check=True, capture_output=True, text=True)

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            project_basics = payload["structured"]["fieldGroups"]["projectBasics"]
            self.assertEqual(field_by_key(project_basics, "projectName")["value"], "公告真实项目")
            self.assertEqual(field_by_key(project_basics, "tenderer")["value"], "公告真实招标单位")
            self.assertEqual(field_by_key(project_basics, "tenderAgency")["value"], "公告真实代理机构")
            self.assertEqual(field_by_key(project_basics, "bidDeadline")["value"], "2026-05-06 09:30")
            self.assertEqual(payload["structured"]["projectDates"]["endDate"], "2026-05-06 09:30")
            self.assertTrue(all("见招标公告" not in field["value"] for field in project_basics))
            self.assertTrue(all("招标公告" in field["section"] for field in project_basics if field["status"] == "found"))

    def test_project_basics_prefer_cover_procurement_labels_and_map_to_tender_fields(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "procurement-cover.docx"
            doc = Document()
            doc.add_paragraph("脱敏采购项目")
            cover_table = doc.add_table(rows=4, cols=3)
            rows = [
                ("项目名称", "：", "封面采购项目"),
                ("采购编号", "：", "CG-2026-001"),
                ("采购人", "：", "封面采购单位"),
                ("采购代理机构", "：", "封面采购代理"),
            ]
            for row_index, values in enumerate(rows):
                for col, text in enumerate(values):
                    cover_table.cell(row_index, col).text = text
            doc.add_paragraph("第一章 采购公告")
            doc.add_paragraph("采购人：公告采购单位")
            doc.add_paragraph("采购代理机构：公告采购代理")
            doc.add_paragraph("响应文件提交截止时间：2026年6月7日10时00分")
            doc.save(source_path)

            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-PROCUREMENT-COVER",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)
            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            self.write_bid_deadline_ai_decision(tmp_path, review_plan, content="2026年5月6日09时30分")
            subprocess.run([sys.executable, str(script_path), "finalize", str(manifest_path)], check=True, capture_output=True, text=True)

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            project_basics = payload["structured"]["fieldGroups"]["projectBasics"]
            self.assertEqual(field_by_key(project_basics, "projectName")["value"], "封面采购项目")
            self.assertEqual(field_by_key(project_basics, "tenderNo")["value"], "CG-2026-001")
            self.assertEqual(field_by_key(project_basics, "tenderer")["value"], "封面采购单位")
            self.assertEqual(field_by_key(project_basics, "tenderAgency")["value"], "封面采购代理")

    def test_project_basics_reference_ai_candidates_use_one_line_context_window(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "project-basics-ai-fallback.md"
            lines = [
                "# 商务招标文件",
                "第一章 招标公告",
                "本项目已具备招标条件。",
                "项目业主为脱敏能源有限公司，现委托脱敏代理有限公司进行公开招标。",
                "建设地点详见技术规范书。",
                "第二章 投标人须知",
                "投标人须知前附表",
                "| 条款号 | 条款名称 | 编列内容 |",
                "| --- | --- | --- |",
                "| 1.1.2 | 招标人 | 见招标公告 |",
            ]
            source_path.write_text("\n".join(lines), encoding="utf-8")
            section_tree_path = tmp_path / "business_section_tree.json"
            section_tree_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-section-tree-v1",
                        "documents": [{"id": "DOC-1", "name": source_path.name}],
                        "nodes": [
                            {
                                "id": "DOC-1-S0001",
                                "documentId": "DOC-1",
                                "level": 1,
                                "number": "第一章",
                                "title": "第一章 招标公告",
                                "path": ["第一章 招标公告"],
                                "source": "heading",
                                "confidence": 0.95,
                                "startBlockIndex": 2,
                                "contentStartBlockIndex": 3,
                                "endBlockIndex": 5,
                                "startLine": 2,
                                "contentStartLine": 3,
                                "endLine": 5,
                            },
                            {
                                "id": "DOC-1-S0002",
                                "documentId": "DOC-1",
                                "level": 2,
                                "number": "",
                                "title": "投标人须知前附表",
                                "path": ["第二章 投标人须知", "投标人须知前附表"],
                                "source": "heading",
                                "confidence": 0.95,
                                "startBlockIndex": 7,
                                "contentStartBlockIndex": 8,
                                "endBlockIndex": 10,
                                "startLine": 7,
                                "contentStartLine": 8,
                                "endLine": 10,
                            },
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-PROJECT-BASICS-AI-FALLBACK",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "businessSectionTreePath": str(section_tree_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                        "textPath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)

            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            fallback_candidates = candidate_package["candidates"]["projectBasicsReferenceReview"]
            self.assertEqual(len(fallback_candidates), 1)
            candidate = fallback_candidates[0]
            self.assertEqual(candidate["fieldKey"], "tenderer")
            self.assertEqual(candidate["referenceTarget"], "招标公告")
            self.assertEqual(candidate["contextLines"], lines[2:5])

            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            self.assertTrue(any(task["task"] == "project_basics_reference_review" for task in review_plan["tasks"]))

    def test_project_basics_reference_ai_decision_updates_final_project_basics(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "project-basics-ai-finalize.md"
            lines = [
                "# 商务招标文件",
                "第一章 招标公告",
                "本项目已具备招标条件。",
                "项目业主为脱敏能源有限公司，现委托脱敏代理有限公司进行公开招标。",
                "建设地点详见技术规范书。",
                "第二章 投标人须知",
                "投标人须知前附表",
                "| 条款号 | 条款名称 | 编列内容 |",
                "| --- | --- | --- |",
                "| 1.1.2 | 招标人 | 见招标公告 |",
            ]
            source_path.write_text("\n".join(lines), encoding="utf-8")
            section_tree_path = tmp_path / "business_section_tree.json"
            section_tree_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-section-tree-v1",
                        "documents": [{"id": "DOC-1", "name": source_path.name}],
                        "nodes": [
                            {
                                "id": "DOC-1-S0001",
                                "documentId": "DOC-1",
                                "level": 1,
                                "number": "第一章",
                                "title": "第一章 招标公告",
                                "path": ["第一章 招标公告"],
                                "source": "heading",
                                "confidence": 0.95,
                                "startBlockIndex": 2,
                                "contentStartBlockIndex": 3,
                                "endBlockIndex": 5,
                                "startLine": 2,
                                "contentStartLine": 3,
                                "endLine": 5,
                            },
                            {
                                "id": "DOC-1-S0002",
                                "documentId": "DOC-1",
                                "level": 2,
                                "number": "",
                                "title": "投标人须知前附表",
                                "path": ["第二章 投标人须知", "投标人须知前附表"],
                                "source": "heading",
                                "confidence": 0.95,
                                "startBlockIndex": 7,
                                "contentStartBlockIndex": 8,
                                "endBlockIndex": 10,
                                "startLine": 7,
                                "contentStartLine": 8,
                                "endLine": 10,
                            },
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-PROJECT-BASICS-AI-FINALIZE",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "businessSectionTreePath": str(section_tree_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                        "textPath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)
            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            for task_ref in review_plan["tasks"]:
                task_path = tmp_path / task_ref["taskPath"]
                decision_path = tmp_path / task_ref["decisionPath"]
                task_payload = json.loads(task_path.read_text(encoding="utf-8"))
                if task_ref["task"] == "project_basics_reference_review":
                    candidate = task_payload["candidates"][0]
                    decision_path.parent.mkdir(parents=True, exist_ok=True)
                    decision_path.write_text(
                        json.dumps(
                            {
                                "schemaVersion": "bid-business-ai-decision-v1",
                                "task": task_ref["task"],
                                "taskId": task_ref["taskId"],
                                "adapter": "unit-test-opencode-agent",
                                "accepted": [
                                    {
                                        "candidateId": candidate["candidateId"],
                                        "decision": "accepted",
                                        "fieldType": "tenderer",
                                        "content": "脱敏能源有限公司",
                                        "applicableScope": "全部标段",
                                        "sourceText": candidate["sourceText"],
                                        "reason": "候选窗口中项目业主即招标人。",
                                        "evidenceIds": candidate["evidenceIds"],
                                    }
                                ],
                                "rejected": [],
                                "needsReview": [],
                                "reason": "单元测试确认 AI 兜底可回填项目基础信息。",
                                "evidenceIds": candidate["evidenceIds"],
                            },
                            ensure_ascii=False,
                            indent=2,
                        ),
                        encoding="utf-8",
                    )
                    continue
                self.write_plan_decision(task_path, decision_path, accepted_candidate_ids=set())

            subprocess.run([sys.executable, str(script_path), "finalize", str(manifest_path)], check=True, capture_output=True, text=True)

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            tenderer = field_by_key(payload["structured"]["fieldGroups"]["projectBasics"], "tenderer")
            self.assertEqual(tenderer["value"], "脱敏能源有限公司")
            self.assertEqual(tenderer["sourcePriority"], "ai_reference_section")
            self.assertNotEqual(tenderer["value"], "见招标公告")

    def test_bid_deadline_reference_ai_candidates_use_keyword_and_date_context(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "bid-deadline-ai-candidates.md"
            lines = [
                "# 商务招标文件",
                "第一章 招标公告",
                "本项目已具备招标条件。",
                "投标文件应于 2026 年 3 月 23 日 10:00 之前递交到电子平台。",
                "逾期递交的投标文件将被拒绝。",
                "投标截止时间前完成 CA 绑定。",
                "请投标人合理安排上传时间。",
                "第二章 投标人须知",
                "投标人须知前附表",
                "| 条款号 | 条款名称 | 编列内容 |",
                "| --- | --- | --- |",
                "| 4.2.1 | 投标截止时间 | 详见招标公告 |",
            ]
            source_path.write_text("\n".join(lines), encoding="utf-8")
            section_tree_path = tmp_path / "business_section_tree.json"
            section_tree_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-section-tree-v1",
                        "documents": [{"id": "DOC-1", "name": source_path.name}],
                        "nodes": [
                            {
                                "id": "DOC-1-S0001",
                                "documentId": "DOC-1",
                                "level": 1,
                                "number": "第一章",
                                "title": "第一章 招标公告",
                                "path": ["第一章 招标公告"],
                                "source": "heading",
                                "confidence": 0.95,
                                "startBlockIndex": 2,
                                "contentStartBlockIndex": 3,
                                "endBlockIndex": 7,
                                "startLine": 2,
                                "contentStartLine": 3,
                                "endLine": 7,
                            },
                            {
                                "id": "DOC-1-S0002",
                                "documentId": "DOC-1",
                                "level": 2,
                                "number": "",
                                "title": "投标人须知前附表",
                                "path": ["第二章 投标人须知", "投标人须知前附表"],
                                "source": "heading",
                                "confidence": 0.95,
                                "startBlockIndex": 9,
                                "contentStartBlockIndex": 10,
                                "endBlockIndex": 12,
                                "startLine": 9,
                                "contentStartLine": 10,
                                "endLine": 12,
                            },
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-BID-DEADLINE-AI-CANDIDATES",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "businessSectionTreePath": str(section_tree_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                        "textPath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)

            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            deadline_candidates = [
                candidate
                for candidate in candidate_package["candidates"]["projectBasicsReferenceReview"]
                if candidate["fieldKey"] == "bidDeadline"
            ]
            self.assertEqual(len(deadline_candidates), 1)
            candidate = deadline_candidates[0]
            self.assertEqual(candidate["content"], lines[3])
            self.assertEqual(candidate["contextLines"], lines[2:5])
            self.assertEqual(candidate["referenceTarget"], "招标公告")

    def test_bid_deadline_ai_candidate_date_context_accepts_common_date_formats(self) -> None:
        scripts_dir = self.runner_path().parent
        sys.path.insert(0, str(scripts_dir))
        try:
            workflow = importlib.import_module("business_workflow")
            accepted_lines = [
                "投标文件应于 2026 年 03 月 23 日 10:00 前递交。",
                "响应文件应于2026年3月23日10:00前提交。",
                "应答文件应于2026/03/23 10:00前提交。",
                "截止时间：2026/3/23 10:00。",
                "截止时间：2026-03-23 10:00。",
                "截止时间：2026-3-23 10:00。",
                "截止时间：2026.03.23 10:00。",
                "截止时间：2026.3.23 10:00。",
            ]
            for line in accepted_lines:
                with self.subTest(line=line):
                    self.assertTrue(workflow._is_bid_deadline_ai_candidate_line(line, [line]))
            self.assertFalse(workflow._is_bid_deadline_ai_candidate_line("投标截止时间前完成 CA 绑定。", ["投标截止时间前完成 CA 绑定。"]))
        finally:
            sys.path.remove(str(scripts_dir))

    def test_bid_deadline_reference_ai_decision_updates_final_project_basics(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "bid-deadline-ai-finalize.md"
            lines = [
                "# 商务招标文件",
                "第一章 招标公告",
                "本项目已具备招标条件。",
                "响应文件应于2026/3/23 10:00之前提交到电子平台。",
                "逾期提交的响应文件将被拒绝。",
                "第二章 投标人须知",
                "投标人须知前附表",
                "| 条款号 | 条款名称 | 编列内容 |",
                "| --- | --- | --- |",
                "| 4.2.1 | 投标截止时间 | 详见招标公告 |",
            ]
            source_path.write_text("\n".join(lines), encoding="utf-8")
            section_tree_path = tmp_path / "business_section_tree.json"
            section_tree_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-section-tree-v1",
                        "documents": [{"id": "DOC-1", "name": source_path.name}],
                        "nodes": [
                            {
                                "id": "DOC-1-S0001",
                                "documentId": "DOC-1",
                                "level": 1,
                                "number": "第一章",
                                "title": "第一章 招标公告",
                                "path": ["第一章 招标公告"],
                                "source": "heading",
                                "confidence": 0.95,
                                "startBlockIndex": 2,
                                "contentStartBlockIndex": 3,
                                "endBlockIndex": 5,
                                "startLine": 2,
                                "contentStartLine": 3,
                                "endLine": 5,
                            },
                            {
                                "id": "DOC-1-S0002",
                                "documentId": "DOC-1",
                                "level": 2,
                                "number": "",
                                "title": "投标人须知前附表",
                                "path": ["第二章 投标人须知", "投标人须知前附表"],
                                "source": "heading",
                                "confidence": 0.95,
                                "startBlockIndex": 7,
                                "contentStartBlockIndex": 8,
                                "endBlockIndex": 10,
                                "startLine": 7,
                                "contentStartLine": 8,
                                "endLine": 10,
                            },
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-BID-DEADLINE-AI-FINALIZE",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "businessSectionTreePath": str(section_tree_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                        "textPath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)
            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            for task_ref in review_plan["tasks"]:
                task_path = tmp_path / task_ref["taskPath"]
                decision_path = tmp_path / task_ref["decisionPath"]
                task_payload = json.loads(task_path.read_text(encoding="utf-8"))
                if task_ref["task"] == "project_basics_reference_review":
                    candidate = next(item for item in task_payload["candidates"] if item["fieldKey"] == "bidDeadline")
                    decision_path.parent.mkdir(parents=True, exist_ok=True)
                    decision_path.write_text(
                        json.dumps(
                            {
                                "schemaVersion": "bid-business-ai-decision-v1",
                                "task": task_ref["task"],
                                "taskId": task_ref["taskId"],
                                "adapter": "unit-test-opencode-agent",
                                "accepted": [
                                    {
                                        "candidateId": candidate["candidateId"],
                                        "decision": "accepted",
                                        "fieldType": "projectBasics",
                                        "content": "2026/3/23 10:00",
                                        "applicableScope": "全部标段",
                                        "sourceText": candidate["sourceText"],
                                        "reason": "候选窗口明确说明响应文件应于该时间之前提交。",
                                        "evidenceIds": candidate["evidenceIds"],
                                    }
                                ],
                                "rejected": [],
                                "needsReview": [],
                                "reason": "单元测试确认 AI 提取递交截止时间。",
                                "evidenceIds": candidate["evidenceIds"],
                            },
                            ensure_ascii=False,
                            indent=2,
                        ),
                        encoding="utf-8",
                    )
                    continue
                self.write_plan_decision(task_path, decision_path, accepted_candidate_ids=set())

            subprocess.run([sys.executable, str(script_path), "finalize", str(manifest_path)], check=True, capture_output=True, text=True)

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            bid_deadline = field_by_key(payload["structured"]["fieldGroups"]["projectBasics"], "bidDeadline")
            self.assertEqual(bid_deadline["value"], "2026-03-23 10:00")
            self.assertEqual(bid_deadline["sourcePriority"], "ai_reference_section")
            self.assertEqual(payload["structured"]["projectDates"]["endDate"], "2026-03-23 10:00")

    def test_project_basics_reference_ai_does_not_replace_concrete_preface_value(self) -> None:
        scripts_dir = self.runner_path().parent
        sys.path.insert(0, str(scripts_dir))
        try:
            workflow = importlib.import_module("business_workflow")
            project_basics = [
                {
                    "key": "projectName",
                    "label": "项目名称",
                    "value": "前附表真实项目名称",
                    "status": "found",
                    "sourcePriority": "bidder_instruction",
                    "evidenceIds": ["DOC-1:B1/R1"],
                },
                {"key": "tenderNo", "label": "招标编号", "value": "", "status": "missing"},
                {"key": "tenderer", "label": "招标人", "value": "", "status": "missing"},
                {"key": "tenderAgency", "label": "招标代理机构", "value": "", "status": "missing"},
                {"key": "bidDeadline", "label": "递交截止时间", "value": "", "status": "missing"},
            ]
            candidates = [
                {
                    "id": "PROJECT-BASIC-REF-DOC-1-projectName-0001",
                    "candidateId": "PROJECT-BASIC-REF-DOC-1-projectName-0001",
                    "fieldKey": "projectName",
                    "content": "公告段落里的长项目描述",
                    "sourceFile": "招标文件.docx",
                    "sourceDocumentId": "DOC-1",
                    "section": "第一章 招标公告",
                    "evidence": "公告段落里的长项目描述",
                    "evidenceIds": ["DOC-1:L1"],
                    "referenceTarget": "招标公告",
                }
            ]
            decision = {
                "schemaVersion": "bid-business-ai-decision-v1",
                "task": "project_basics_reference_review",
                "taskId": "project_basics_reference_review/part-001",
                "accepted": [
                    {
                        "candidateId": "PROJECT-BASIC-REF-DOC-1-projectName-0001",
                        "decision": "accepted",
                        "fieldType": "projectBasics",
                        "content": "公告段落里的长项目描述",
                        "applicableScope": "全部标段",
                        "sourceText": "招标文件.docx：第一章 招标公告",
                        "reason": "AI 认为公告段落可作为项目名称候选。",
                        "evidenceIds": ["DOC-1:L1"],
                    }
                ],
                "rejected": [],
                "needsReview": [],
            }

            updated = workflow._apply_project_basic_ai_decisions(project_basics, candidates, decision)

            self.assertEqual(field_by_key(updated, "projectName")["value"], "前附表真实项目名称")
            self.assertEqual(field_by_key(updated, "projectName")["sourcePriority"], "bidder_instruction")
        finally:
            sys.path.remove(str(scripts_dir))

    def test_bid_deadline_ignores_preface_reference_and_opening_time_preserves_minutes(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "bid-deadline-reference-and-opening-time.docx"
            doc = Document()
            doc.add_paragraph("第一章 招标公告")
            doc.add_paragraph("5.1 递交截止时间：2026年1月26日15时00分")
            doc.add_paragraph("开标时间：2026年1月26日16时00分")
            doc.add_paragraph("第二章 投标人须知")
            doc.add_paragraph("投标人须知前附表")
            table = doc.add_table(rows=6, cols=3)
            for col, text in enumerate(["条款号", "条款名称", "编列内容"]):
                table.cell(0, col).text = text
            rows = [
                ("1.1.2", "招标人", "名称：前附表招标单位 地址：示例地址"),
                ("1.1.3", "招标代理机构", "名称：前附表代理机构 地址：代理地址"),
                ("1.1.4", "招标项目名称", "前附表结构化项目"),
                ("4.2.1", "投标截止时间", "详见招标公告"),
                ("5.1", "开标时间和地点", "开标时间：同投标截止时间 开标地点：同递交投标文件地点"),
            ]
            for row_index, values in enumerate(rows, start=1):
                for col, text in enumerate(values):
                    table.cell(row_index, col).text = text
            doc.save(source_path)

            text_path = tmp_path / "bid-deadline-reference-and-opening-time.txt"
            text_path.write_text(
                "\n".join(
                    [
                        "第一章 招标公告",
                        "5.1 递交截止时间：2026年1月26日15时00分",
                        "开标时间：2026年1月26日16时00分",
                        "第二章 投标人须知",
                        "投标人须知前附表",
                    ]
                ),
                encoding="utf-8",
            )
            section_tree_path = tmp_path / "business_section_tree.json"
            section_tree_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-section-tree-v1",
                        "documents": [{"id": "DOC-1", "name": source_path.name}],
                        "nodes": [
                            {
                                "id": "DOC-1-S0001",
                                "documentId": "DOC-1",
                                "level": 1,
                                "number": "第一章",
                                "title": "第一章 招标公告",
                                "path": ["第一章 招标公告"],
                                "source": "heading",
                                "confidence": 0.95,
                                "startBlockIndex": 1,
                                "contentStartBlockIndex": 2,
                                "endBlockIndex": 3,
                                "startLine": 1,
                                "contentStartLine": 2,
                                "endLine": 3,
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-BID-DEADLINE-REFERENCE-OPENING",
                        "bidType": "商务标",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "businessSectionTreePath": str(section_tree_path),
                        "documents": [
                            {
                                "id": "DOC-1",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                                "textPath": str(text_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)
            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            self.write_bid_deadline_ai_decision(tmp_path, review_plan, content="2026年1月26日15时00分")
            subprocess.run([sys.executable, str(script_path), "finalize", str(manifest_path)], check=True, capture_output=True, text=True)

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            project_basics = payload["structured"]["fieldGroups"]["projectBasics"]
            bid_deadline = field_by_key(project_basics, "bidDeadline")
            self.assertEqual(bid_deadline["value"], "2026-01-26 15:00")
            self.assertIn("递交截止时间", bid_deadline["evidence"])
            self.assertIn("开标时间", bid_deadline["evidence"])
            self.assertEqual(bid_deadline["sourcePriority"], "ai_reference_section")
            self.assertEqual(payload["structured"]["projectDates"]["endDate"], "2026-01-26 15:00")

            self.assertEqual(field_by_key(project_basics, "projectName")["value"], "前附表结构化项目")
            self.assertEqual(field_by_key(project_basics, "tenderer")["value"], "前附表招标单位")
            self.assertEqual(field_by_key(project_basics, "tenderAgency")["value"], "前附表代理机构")

            fact_by_key = {field["fieldKey"]: field for field in payload["structured"]["projectFactFields"]}
            self.assertEqual(fact_by_key["bidDeadline"]["value"], "2026-01-26 15:00")

    def test_project_basics_prefer_exact_preface_clauses_and_exclude_relative_deadlines(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "project-basics-clause-priority.docx"
            doc = Document()
            doc.add_paragraph("第二章 投标人须知")
            doc.add_paragraph("投标人须知前附表")
            table = doc.add_table(rows=5, cols=3)
            for col, text in enumerate(["条款号", "条款名称", "编列内容"]):
                table.cell(0, col).text = text
            rows = [
                ("3.2.5", "投标报价的其他要求", "招标人不接受投标人任何形式的价格调整声明；投标截止时间10日前不得修改报价。"),
                ("2.2.1", "澄清招标文件的时间", "投标截止时间10日前。"),
                ("1.1.2", "招标人", "名称：山西漳山发电有限责任公司 地址：示例地址"),
                ("4.2.1", "投标文件递交截止时间", "2026年7月1日09时30分"),
            ]
            for row_index, values in enumerate(rows, start=1):
                for col, text in enumerate(values):
                    table.cell(row_index, col).text = text
            doc.save(source_path)

            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-PROJECT-BASICS-CLAUSE-PRIORITY",
                        "bidType": "商务标",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "documents": [
                            {
                                "id": "DOC-1",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            subprocess.run(
                [sys.executable, str(script_path), "offline-fallback", str(manifest_path)],
                check=True,
                capture_output=True,
                text=True,
            )

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            project_basics = payload["structured"]["fieldGroups"]["projectBasics"]
            tenderer = field_by_key(project_basics, "tenderer")
            bid_deadline = field_by_key(project_basics, "bidDeadline")
            self.assertEqual(tenderer["value"], "山西漳山发电有限责任公司")
            self.assertIn("1.1.2", tenderer["evidence"])
            self.assertEqual(bid_deadline["value"], "2026-07-01 09:30")
            self.assertIn("投标文件递交截止时间", bid_deadline["evidence"])
            self.assertTrue(bid_deadline.get("sourceFile"))
            self.assertTrue(bid_deadline.get("section"))
            self.assertTrue(bid_deadline.get("evidenceLocation"))
            self.assertTrue(bid_deadline.get("evidenceIds"))

    def test_markdown_bidder_instruction_preface_table_is_extracted(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "markdown-preface.md"
            source_path.write_text(
                "\n".join(
                    [
                        "# 商务招标文件",
                        "第二章 投标人须知",
                        "投标人须知前附表",
                        "| 条款号 | 条款名称 | 编列内容 |",
                        "| --- | --- | --- |",
                        "| 1.1.2 | 招标人 | 示例招标人 |",
                        "| 1.1.3 | 招标代理机构 | 示例代理机构 |",
                        "| 4.2.1 | 投标截止时间 | 2026年7月1日09时30分 |",
                        "第三章 评标办法",
                        "商务评分细则",
                        "| 评分项 | 分值 | 评分标准 |",
                        "| --- | --- | --- |",
                        "| 企业信誉 | 5 | 信誉良好得5分。 |",
                    ]
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-MARKDOWN-PREFACE",
                        "bidType": "商务标",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "documents": [
                            {
                                "id": "DOC-1",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                                "textPath": str(source_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            subprocess.run([sys.executable, str(script_path), "offline-fallback", str(manifest_path)], check=True, capture_output=True, text=True)

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            rows = payload["structured"]["fieldGroups"]["bidderInstructions"]
            self.assertEqual([row["clauseNo"] for row in rows], ["1.1.2", "1.1.3", "4.2.1"])
            self.assertEqual(rows[0]["clauseName"], "招标人")
            self.assertEqual(rows[0]["content"], "示例招标人")
            self.assertEqual(rows[0]["evidenceLocation"], "L6")
            self.assertTrue(rows[0].get("evidenceIds"))

            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            self.assertEqual(len(candidate_package["deterministicExtracts"]["bidderInstructions"]), 3)

    def test_business_finalize_registers_table_row_evidence(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "table-row-evidence.docx"
            doc = Document()
            doc.add_paragraph("第一章 招标公告")
            doc.add_paragraph("3. 投标人资格要求")
            doc.add_paragraph("3.1 投标人须为中华人民共和国境内依法注册的独立法人。")
            doc.add_paragraph("第二章 投标人须知")
            doc.add_paragraph("投标人须知前附表")
            instruction_table = doc.add_table(rows=5, cols=3)
            for col, text in enumerate(["条款号", "条款名称", "编列内容"]):
                instruction_table.cell(0, col).text = text
            instruction_rows = [
                ("1.1.2", "招标人", "名称：山西漳山发电有限责任公司"),
                ("1.1.3", "招标代理机构", "名称：中招代理有限公司"),
                ("1.1.4", "招标项目名称", "表格行证据回溯项目"),
                ("4.2.1", "投标截止时间", "2026年7月1日09时30分"),
            ]
            for row_index, values in enumerate(instruction_rows, start=1):
                for col, text in enumerate(values):
                    instruction_table.cell(row_index, col).text = text
            doc.add_paragraph("第三章 评标办法")
            doc.add_paragraph("评标办法前附表")
            scoring_table = doc.add_table(rows=3, cols=5)
            for col, text in enumerate(["条款号", "评分因素", "分值", "评分标准", "证明材料要求"]):
                scoring_table.cell(0, col).text = text
            scoring_rows = [
                ("2.2.4（1）", "商务评分标准", "10分", "交货期满足招标要求得10分。", "提供承诺函。"),
                ("", "售后服务", "5分", "售后服务方案完整得5分。", "提供服务方案。"),
            ]
            for row_index, values in enumerate(scoring_rows, start=1):
                for col, text in enumerate(values):
                    scoring_table.cell(row_index, col).text = text
            doc.add_paragraph("投标文件应当对招标文件的实质性要求作出响应，否则投标将被否决。")
            doc.save(source_path)

            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-TABLE-ROW-EVIDENCE",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "documents": [
                    {
                        "id": "DOC-TABLE",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)
            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            for task_ref in review_plan["tasks"]:
                task_payload = json.loads((tmp_path / task_ref["taskPath"]).read_text(encoding="utf-8"))
                self.write_plan_decision(
                    tmp_path / task_ref["taskPath"],
                    tmp_path / task_ref["decisionPath"],
                    accepted_candidate_ids={candidate["candidateId"] for candidate in task_payload.get("candidates") or []},
                )

            subprocess.run([sys.executable, str(script_path), "finalize", str(manifest_path)], check=True, capture_output=True, text=True)

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            structured = payload["structured"]
            validation_report = json.loads((tmp_path / "validation_report.json").read_text(encoding="utf-8"))
            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            evidence_index = candidate_package["evidenceIndex"]

            self.assertEqual(validation_report["status"], "passed", validation_report)
            self.assertEqual(structured["workflow"]["validationStatus"], "passed")
            self.assertEqual(structured["workflow"]["stage"], "finalized")
            self.assertFalse(
                any(
                    check["name"] == "evidence_references" and check["status"] == "failed"
                    for check in validation_report["checks"]
                )
            )
            final_records = []
            for key in ("projectBasics", "qualificationRequirements", "bidderInstructions", "commercialRejectionClauses"):
                final_records.extend(structured["fieldGroups"].get(key) or [])
            final_records.extend(structured["scoringCriteria"].get("business") or [])
            final_records.extend(structured.get("projectFactFields") or [])
            for record in final_records:
                if record.get("evidence"):
                    self.assertTrue(
                        any(eid in evidence_index for eid in record.get("evidenceIds", [])),
                        record,
                    )

    def test_qualification_review_uses_main_announcement_section_payloads(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "qualification-main-section.md"
            source_path.write_text(
                "\n".join(
                    [
                        "# 商务招标文件",
                        "第一章 招标公告",
                        "3. 投标人资格要求",
                        "3.1 通用资格条件",
                        "3.1.1 投标人须为中华人民共和国境内依法注册的独立法人。",
                        "3.1.2 投标人财务、信誉等方面应具备下列条件：",
                        "(1) 没有处于行政主管部门禁止投标处罚期内。",
                        "3.2 专用资格条件",
                        "3.2.1 业绩要求：",
                        "标段一至标段二：",
                        "(1) 投标人须提供近三年同类设备供货业绩。",
                        "1、投标机型应取得完整型式认证证书。",
                        "第二章 投标人须知",
                        "1.4 投标人资格要求",
                        "1.4.1 投标人应具备承担本招标项目资质条件、能力和信誉：见投标人须知前附表。",
                        "3.5 资格审查资料",
                        "投标人应附营业执照复印件。",
                    ]
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-MAIN-QUALIFICATION",
                        "bidType": "商务标",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "documents": [
                            {
                                "id": "DOC-1",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                                "textPath": str(source_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            subprocess.run([sys.executable, str(script_path), "offline-fallback", str(manifest_path)], check=True, capture_output=True, text=True)

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            qualifications = payload["structured"]["fieldGroups"]["qualificationRequirements"]
            joined = "\n".join(row["content"] for row in qualifications)
            self.assertIn("依法注册的独立法人", joined)
            self.assertIn("禁止投标处罚期内", joined)
            self.assertIn("同类设备供货业绩", joined)
            self.assertIn("完整型式认证证书", joined)
            self.assertNotIn("见投标人须知前附表", joined)
            self.assertNotIn("营业执照复印件", joined)
            self.assertTrue(all(row["section"].startswith("3. 投标人资格要求") for row in qualifications))
            self.assertTrue(all(row["sourceText"].startswith("qualification-main-section.md：") for row in qualifications))
            self.assertTrue(all(row.get("evidenceIds") for row in qualifications))

            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            qualification_tasks = [task for task in review_plan["tasks"] if task["task"] == "qualification_review"]
            self.assertEqual(len(qualification_tasks), 1)
            task_payload = json.loads((tmp_path / qualification_tasks[0]["taskPath"]).read_text(encoding="utf-8"))
            self.assertEqual(len(task_payload["candidates"]), 1)
            self.assertTrue(all(candidate["candidateType"] == "qualification_section_slice" for candidate in task_payload["candidates"]))
            self.assertTrue(all("1.4 投标人资格要求" not in candidate["content"] for candidate in task_payload["candidates"]))

    def test_qualification_review_ignores_toc_anchor_and_reopens_main_section(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "qualification-toc.md"
            source_path.write_text(
                "\n".join(
                    [
                        "# 商务招标文件",
                        "目录",
                        "1. 招标公告 1",
                        "2. 项目概况与招标范围 3",
                        "3. 投标人资格要求 5",
                        "4. 招标文件的获取 6",
                        "5. 投标文件的递交 7",
                        "第一章 招标公告",
                        "3. 投标人资格要求",
                        "3.1 通用资格条件",
                        "3.1.1 投标人为中华人民共和国境内合法注册的独立法人或其他组织，具有独立承担民事责任能力，具有独立订立合同的权利；",
                        "3.1.2 投标人财务、信誉等方面应具备下列条件：",
                        "(1) 没有处于行政主管部门禁止投标处罚期内。",
                        "3.2 专用资格条件",
                        "3.2.1 业绩要求：",
                        "(1) 投标人须提供近三年同类设备供货业绩。",
                        "第二章 投标人须知",
                        "1.4 投标人资格要求",
                        "1.4.1 投标人应具备承担本招标项目资质条件、能力和信誉：见投标人须知前附表。",
                        "1.4.2 投标人须具有第二章重复资格条件。",
                    ]
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-QUALIFICATION-TOC",
                        "bidType": "商务标",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "documents": [
                            {
                                "id": "DOC-1",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                                "textPath": str(source_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            subprocess.run([sys.executable, str(script_path), "offline-fallback", str(manifest_path)], check=True, capture_output=True, text=True)

            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            qualification_candidates = candidate_package["candidates"]["qualification"]
            self.assertEqual(len(qualification_candidates), 1)
            qualification_slice = qualification_candidates[0]
            self.assertEqual(qualification_slice["candidateType"], "qualification_section_slice")
            self.assertEqual(qualification_slice["section"], "3. 投标人资格要求")
            self.assertTrue(qualification_slice.get("lines"))
            self.assertTrue(all(line.get("evidenceId") for line in qualification_slice["lines"]))
            candidate_text = qualification_slice["content"]
            self.assertIn("合法注册的独立法人", candidate_text)
            self.assertIn("禁止投标处罚期内", candidate_text)
            self.assertIn("同类设备供货业绩", candidate_text)
            self.assertNotIn("投标人资格要求 5", candidate_text)
            self.assertNotIn("第二章重复资格条件", candidate_text)
            candidate_sources = "\n".join(line["sourceText"] for line in qualification_slice["lines"])
            self.assertIn("3. 投标人资格要求 > 3.1 通用资格条件 > 3.1.1 投标人为中华人民共和国境内合法注册的独立法人", candidate_sources)
            self.assertIn("3. 投标人资格要求 > 3.1 通用资格条件 > 3.1.2 投标人财务、信誉等方面应具备下列条件： > (1)", candidate_sources)

            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            qualification_ref = next(task for task in review_plan["tasks"] if task["task"] == "qualification_review")
            task_payload = json.loads((tmp_path / qualification_ref["taskPath"]).read_text(encoding="utf-8"))
            self.assertIn("qualificationItems", task_payload["decisionContract"])
            self.assertEqual(len(task_payload["candidates"]), 1)
            self.assertEqual(task_payload["candidates"][0]["candidateType"], "qualification_section_slice")
            self.assertTrue(task_payload["candidates"][0].get("lines"))
            task_text = task_payload["candidates"][0]["content"]
            self.assertIn("合法注册的独立法人", task_text)
            self.assertNotIn("投标人资格要求 5", task_text)
            self.assertNotIn("第二章重复资格条件", task_text)
            decision = json.loads((tmp_path / qualification_ref["decisionPath"]).read_text(encoding="utf-8"))
            self.assertIn("qualificationItems", decision)
            self.assertTrue(decision["qualificationItems"])
            self.assertTrue(all(item.get("content") for item in decision["qualificationItems"]))
            self.assertTrue(all(item.get("applicableScope") for item in decision["qualificationItems"]))
            self.assertTrue(all(item.get("sourceText") for item in decision["qualificationItems"]))
            self.assertTrue(all(item.get("evidenceIds") for item in decision["qualificationItems"]))

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            qualifications = payload["structured"]["fieldGroups"]["qualificationRequirements"]
            joined = "\n".join(row["content"] for row in qualifications)
            self.assertIn("合法注册的独立法人", joined)
            self.assertIn("禁止投标处罚期内", joined)
            self.assertIn("同类设备供货业绩", joined)
            self.assertNotIn("见投标人须知前附表", joined)
            self.assertNotIn("第二章重复资格条件", joined)
            self.assertTrue(all(row["section"].startswith("3. 投标人资格要求") for row in qualifications))
            qualification_sources = "\n".join(row["sourceText"] for row in qualifications)
            self.assertIn("3. 投标人资格要求 > 3.1 通用资格条件 > 3.1.1 投标人为中华人民共和国境内合法注册的独立法人", qualification_sources)
            self.assertIn("3. 投标人资格要求 > 3.1 通用资格条件 > 3.1.2 投标人财务、信誉等方面应具备下列条件： > (1)", qualification_sources)

    def test_rejection_noise_is_excluded_and_scoring_embedded_rows_stop_at_next_module(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "rejection-and-scoring.docx"
            doc = Document()
            doc.add_paragraph("第一章 招标公告")
            doc.add_paragraph("投标文件逾期上传或者未成功上传指定平台的，招标人不予受理。")
            doc.add_paragraph("投标文件未对商务实质性要求响应的，其投标将被否决。")
            doc.add_paragraph("异议材料未按要求提交的，招标人可以不予受理。")
            doc.add_paragraph("投诉材料未按要求提交的，不予受理。")
            doc.add_paragraph("合同履行过程中发生争议的，按合同约定处理。")
            doc.add_paragraph("投标保证金不退还的情形按招标文件规定执行。")
            doc.add_paragraph("第三章 评标办法")
            doc.add_paragraph("评标办法前附表")
            table = doc.add_table(rows=6, cols=5)
            for col, text in enumerate(["条款号", "评分因素", "分值", "评分标准", "证明材料要求"]):
                table.cell(0, col).text = text
            rows = [
                ("2.2.4(1)", "商务评分标准（20分）", "20分", "企业业绩满足要求得20分。", "提供合同。"),
                ("", "服务响应", "5分", "响应完整得5分。", "提供承诺。"),
                ("", "技术评分标准（30分）", "30分", "技术方案完整得30分。", "提供技术方案。"),
                ("", "报价评分标准（50分）", "50分", "按评标价计算。", "提供报价表。"),
                ("", "串通投标认定", "", "不同投标人文件异常一致时按否决处理。", ""),
            ]
            for row_index, values in enumerate(rows, start=1):
                for col, text in enumerate(values):
                    table.cell(row_index, col).text = text
            doc.save(source_path)

            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-REJECTION-SCORING",
                        "bidType": "商务标",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "documents": [
                            {
                                "id": "DOC-1",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            subprocess.run([sys.executable, str(script_path), "offline-fallback", str(manifest_path)], check=True, capture_output=True, text=True)

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            rejection_text = "\n".join(row["content"] for row in payload["structured"]["fieldGroups"]["commercialRejectionClauses"])
            self.assertIn("逾期上传", rejection_text)
            self.assertIn("商务实质性要求响应", rejection_text)
            self.assertNotIn("异议材料", rejection_text)
            self.assertNotIn("投诉材料", rejection_text)
            self.assertNotIn("合同履行", rejection_text)
            self.assertNotIn("投标保证金不退还", rejection_text)

            scoring = payload["structured"]["scoringCriteria"]
            business_items = [row["scoringItem"] for row in scoring["business"]]
            self.assertEqual(business_items, ["商务评分标准（20分）", "服务响应"])
            self.assertTrue(all("串通投标" not in json.dumps(row, ensure_ascii=False) for row in scoring["business"]))
            self.assertNotIn("price", scoring)
            self.assertNotIn("compliance", scoring)
            self.assertTrue(all("技术评分" not in json.dumps(row, ensure_ascii=False) for row in scoring["business"]))
            self.assertEqual(scoring.get("technical"), None)

            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            self.assertTrue(any(task["task"] == "rejection_clause_review" for task in review_plan["tasks"]))
            self.assertFalse(any(task["task"] == "scoring_table_review" for task in review_plan["tasks"]))

    def test_scoring_table_review_uses_row_level_business_boundaries(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "scoring-row-boundary.docx"
            doc = Document()
            doc.add_paragraph("第三章 评标办法")
            doc.add_paragraph("评标办法前附表")
            table = doc.add_table(rows=8, cols=5)
            for col, text in enumerate(["条款号", "评分因素", "分值", "评分标准", "证明材料要求"]):
                table.cell(0, col).text = text
            rows = [
                ("2.2.4（1）", "商务评分标准", "10分", "交货期满足要求得10分。", "提供承诺。"),
                ("", "售后服务", "5分", "售后服务方案完整得5分。", "提供方案。"),
                ("", "付款响应", "5分", "付款条件响应得5分。", "提供响应表。"),
                ("2.2.4（2）", "技术评分标准", "30分", "技术方案完整得30分。", "提供技术方案。"),
                ("", "施工组织", "10分", "施工组织合理得10分。", "提供施工组织。"),
                ("2.2.4（3）", "投标报价评分标准", "50分", "按评标价计算。", "提供报价表。"),
                ("", "报价偏差", "5分", "报价偏差满足要求得5分。", "提供说明。"),
            ]
            for row_index, values in enumerate(rows, start=1):
                for col, text in enumerate(values):
                    table.cell(row_index, col).text = text
            doc.save(source_path)

            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-SCORING-ROW-BOUNDARY",
                        "bidType": "商务标",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "documents": [
                            {
                                "id": "DOC-1",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            subprocess.run([sys.executable, str(script_path), "offline-fallback", str(manifest_path)], check=True, capture_output=True, text=True)

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            scoring = payload["structured"]["scoringCriteria"]
            self.assertEqual(len(scoring["business"]), 3)
            self.assertFalse(any("技术" in row["evidence"] for row in scoring["business"]))
            self.assertFalse(any("报价" in row["evidence"] for row in scoring["business"]))
            self.assertNotIn("price", scoring)
            self.assertNotIn("compliance", scoring)

    def test_rejection_clause_review_expands_parent_heading_children(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "rejection-clause-block.md"
            source_path.write_text(
                "\n".join(
                    [
                        "# 商务招标文件",
                        "第二章 投标人须知",
                        "1.4.3 投标人不得存在下列情形之一：",
                        "（1）为招标人不具有独立法人资格的附属机构；",
                        "（2）为本招标项目前期准备提供设计或咨询服务；",
                        "1.4.4 其他条款",
                    ]
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-REJECTION-CLAUSE-BLOCK",
                        "bidType": "商务标",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "documents": [
                            {
                                "id": "DOC-1",
                                "name": source_path.name,
                                "sourcePath": str(source_path),
                                "textPath": str(source_path),
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            subprocess.run([sys.executable, str(script_path), "offline-fallback", str(manifest_path)], check=True, capture_output=True, text=True)

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            rejection_text = "\n".join(
                row["content"] for row in payload["structured"]["fieldGroups"]["commercialRejectionClauses"]
            )
            self.assertIn("不具有独立法人资格的附属机构", rejection_text)
            self.assertIn("前期准备提供设计或咨询服务", rejection_text)
            self.assertNotIn("1.4.4 其他条款", rejection_text)

            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            rejection_ref = next(task for task in review_plan["tasks"] if task["task"] == "rejection_clause_review")
            task_payload = json.loads((tmp_path / rejection_ref["taskPath"]).read_text(encoding="utf-8"))
            candidate_text = "\n".join(candidate["content"] for candidate in task_payload["candidates"])
            self.assertIn("投标人不得存在下列情形之一", candidate_text)
            self.assertIn("前期准备提供设计或咨询服务", candidate_text)

    def test_business_parser_finalize_with_complete_review_plan_decisions_is_trusted(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "完整审查商务片段.md"
            source_path.write_text(
                "\n".join(
                    [
                        "# 商务招标文件",
                        "第一章 招标公告",
                        "3. 投标人资格要求",
                        "3.1 投标人须为中华人民共和国境内合法注册的独立法人或其他组织。",
                        "3.2 投标人须具有近三年同类设备供货业绩。",
                        "第二章 投标人须知",
                        "投标文件逾期上传的，招标人不予受理。",
                    ]
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-COMPLETE-PLAN-AI",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                        "textPath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)

            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            accepted_candidate = candidate_package["candidates"]["qualification"][0]
            accepted_rejection = candidate_package["candidates"]["rejection"][0]
            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            for task_ref in review_plan["tasks"]:
                accepted_ids = set()
                if task_ref["task"] == "qualification_review":
                    accepted_ids = {accepted_candidate["id"]}
                elif task_ref["task"] == "rejection_clause_review":
                    accepted_ids = {accepted_rejection["id"]}
                self.write_plan_decision(
                    tmp_path / task_ref["taskPath"],
                    tmp_path / task_ref["decisionPath"],
                    accepted_candidate_ids=accepted_ids,
                )

            completed = subprocess.run(
                [sys.executable, str(script_path), "finalize", str(manifest_path)],
                check=True,
                capture_output=True,
                text=True,
            )

            self.assertEqual(json.loads(completed.stdout)["summary"]["workflowStage"], "finalized")
            payload = json.loads(output_path.read_text(encoding="utf-8"))
            workflow = payload["structured"]["workflow"]
            self.assertEqual(workflow["stage"], "finalized")
            self.assertTrue(workflow["aiReviewTrusted"])
            self.assertEqual(workflow["semanticReviewMode"], "opencode-agent")
            self.assertEqual(workflow["requiredDecisionTaskCount"], review_plan["taskCount"])
            self.assertEqual(workflow["presentDecisionTaskCount"], review_plan["taskCount"])
            self.assertEqual(workflow["missingDecisionTasks"], [])
            self.assertEqual(workflow["reviewPlanPath"], str(tmp_path / "review_plan.json"))
            validation = json.loads((tmp_path / "validation_report.json").read_text(encoding="utf-8"))
            coverage_check = next(check for check in validation["checks"] if check["name"] == "ai_decision_coverage")
            self.assertEqual(coverage_check["status"], "passed")
            qualification_text = "\n".join(
                row["content"] for row in payload["structured"]["fieldGroups"]["qualificationRequirements"]
            )
            self.assertIn("合法注册的独立法人", qualification_text)
            rejections = payload["structured"]["fieldGroups"]["commercialRejectionClauses"]
            self.assertEqual(len(rejections), 1)
            self.assertIn("不予受理", rejections[0]["content"])
            self.assert_rejection_display_fields(rejections)
            self.assertEqual(rejections[0]["riskLevel"], "high")
            self.assertIn("不予受理", rejections[0]["matchedKeywords"])

    def test_business_finalize_keeps_all_accepted_qualification_candidates(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "qualification-six-candidates.md"
            source_path.write_text(
                "\n".join(
                    [
                        "# 商务招标文件",
                        "第一章 招标公告",
                        "3. 投标人资格要求",
                        "3.1 投标人须为中华人民共和国境内合法注册的独立法人或其他组织。",
                        "3.2 投标人须具有有效的营业执照。",
                        "3.3 投标人须具有近三年同类设备供货业绩。",
                        "3.4 投标人不得处于行政主管部门禁止投标处罚期内。",
                        "3.5 本项目不接受联合体投标。",
                        "3.6 投标人须提供银行资信证明。",
                    ]
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-SIX-QUALIFICATIONS",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                        "textPath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)
            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            qualification_ref = next(task for task in review_plan["tasks"] if task["task"] == "qualification_review")
            task_payload = json.loads((tmp_path / qualification_ref["taskPath"]).read_text(encoding="utf-8"))
            accepted_ids = {candidate["candidateId"] for candidate in task_payload.get("candidates") or []}
            self.assertEqual(len(accepted_ids), 1)
            self.write_plan_decision(
                tmp_path / qualification_ref["taskPath"],
                tmp_path / qualification_ref["decisionPath"],
                accepted_candidate_ids=accepted_ids,
            )
            for task_ref in review_plan["tasks"]:
                if task_ref["task"] == "qualification_review":
                    continue
                self.write_plan_decision(
                    tmp_path / task_ref["taskPath"],
                    tmp_path / task_ref["decisionPath"],
                    accepted_candidate_ids=set(),
                )

            subprocess.run([sys.executable, str(script_path), "finalize", str(manifest_path)], check=True, capture_output=True, text=True)

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            field_groups = payload["structured"]["fieldGroups"]
            workflow = payload["structured"]["workflow"]
            self.assertGreaterEqual(len(field_groups["qualificationRequirements"]), 6)
            self.assertTrue(all(row.get("evidenceIds") for row in field_groups["qualificationRequirements"]))
            self.assertEqual(workflow["validationStatus"], "passed")

    def test_business_parser_uses_section_tree_for_qualification_scope(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "章节树资格范围.md"
            lines = [
                "# 商务招标文件",
                "目录",
                "3. 供应商资格要求 ........ 3",
                "第一章 招标公告",
                "3. 供应商资格要求",
                "3.1 供应商须为中华人民共和国境内合法注册的独立法人。",
                "3.2 本项目不接受联合体投标。",
                "第二章 投标人须知",
                "1.4 投标人资格要求",
                "1.4.1 见投标人须知前附表，本行不应进入资格候选范围。",
                "3.5 资格审查资料",
                "投标人须提供营业执照复印件，本行属于资料要求而非资格要求。",
                "第三章 评标办法",
                "商务评分标准",
                "类似合同业绩每增加一项加2分，本行不应进入资格候选范围。",
            ]
            source_path.write_text("\n".join(lines), encoding="utf-8")
            section_tree_path = tmp_path / "business_section_tree.json"
            section_tree_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-section-tree-v1",
                        "documents": [{"id": "DOC-1", "name": source_path.name}],
                        "nodes": [
                            {
                                "id": "DOC-1-S0001",
                                "documentId": "DOC-1",
                                "level": 2,
                                "number": "3.",
                                "title": "3. 供应商资格要求",
                                "path": ["第一章 招标公告", "3. 供应商资格要求"],
                                "source": "heading",
                                "confidence": 0.95,
                                "startBlockIndex": 5,
                                "contentStartBlockIndex": 6,
                                "endBlockIndex": 7,
                                "startLine": 5,
                                "contentStartLine": 6,
                                "endLine": 7,
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-SECTION-TREE-QUAL",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "businessSectionTreePath": str(section_tree_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                        "textPath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)

            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            qualification_candidates = candidate_package["candidates"]["qualification"]
            self.assertEqual(len(qualification_candidates), 1)
            qualification_slice = qualification_candidates[0]
            slice_text = "\n".join(line["text"] for line in qualification_slice["lines"])
            self.assertIn("合法注册的独立法人", slice_text)
            self.assertIn("不接受联合体", slice_text)
            self.assertNotIn("见投标人须知前附表", slice_text)
            self.assertNotIn("营业执照复印件", slice_text)
            self.assertNotIn("类似合同业绩", slice_text)
            self.assertEqual(qualification_slice["startLine"], 5)
            self.assertEqual(qualification_slice["endLine"], 7)
            self.assertEqual(qualification_slice["section"], "3. 供应商资格要求")

    def test_business_parser_prefers_main_qualification_node_when_section_tree_has_bidder_instruction_duplicate(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "章节树重复资格节点.md"
            lines = [
                "# 商务招标文件",
                "第一章 招标公告",
                "3. 投标人资格要求",
                "3.1 投标人须为中华人民共和国境内合法注册的独立法人。",
                "3.2 投标人须具有近三年同类设备供货业绩。",
                "第二章 投标人须知",
                "1. 总则",
                "1.4 投标人资格要求",
                "1.4.1 投标人应具备承担本项目的资格条件：见投标人须知前附表。",
                "1.4.2 投标人须提供营业执照复印件，本行属于资料要求而非资格要求。",
            ]
            source_path.write_text("\n".join(lines), encoding="utf-8")
            section_tree_path = tmp_path / "business_section_tree.json"
            section_tree_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-section-tree-v1",
                        "documents": [{"id": "DOC-1", "name": source_path.name}],
                        "nodes": [
                            {
                                "id": "DOC-1-S0001",
                                "documentId": "DOC-1",
                                "level": 2,
                                "number": "3.",
                                "title": "3. 投标人资格要求",
                                "path": ["第一章 招标公告", "3. 投标人资格要求"],
                                "source": "heading",
                                "confidence": 0.95,
                                "startBlockIndex": 3,
                                "contentStartBlockIndex": 4,
                                "endBlockIndex": 5,
                                "startLine": 3,
                                "contentStartLine": 4,
                                "endLine": 5,
                            },
                            {
                                "id": "DOC-1-S0002",
                                "documentId": "DOC-1",
                                "level": 3,
                                "number": "1.4",
                                "title": "1.4 投标人资格要求",
                                "path": ["第二章 投标人须知", "1. 总则", "1.4 投标人资格要求"],
                                "source": "heading",
                                "confidence": 0.95,
                                "startBlockIndex": 8,
                                "contentStartBlockIndex": 9,
                                "endBlockIndex": 10,
                                "startLine": 8,
                                "contentStartLine": 9,
                                "endLine": 10,
                            },
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-SECTION-TREE-QUAL-DUPLICATE",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "businessSectionTreePath": str(section_tree_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                        "textPath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run(
                [sys.executable, str(script_path), "offline-fallback", str(manifest_path)],
                check=True,
                capture_output=True,
                text=True,
            )

            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            qualification_candidates = candidate_package["candidates"]["qualification"]
            self.assertEqual(len(qualification_candidates), 1)
            self.assertEqual(qualification_candidates[0]["section"], "3. 投标人资格要求")
            candidate_text = qualification_candidates[0]["content"]
            self.assertIn("合法注册的独立法人", candidate_text)
            self.assertIn("同类设备供货业绩", candidate_text)
            self.assertNotIn("见投标人须知前附表", candidate_text)
            self.assertNotIn("营业执照复印件", candidate_text)

            rows = json.loads(output_path.read_text(encoding="utf-8"))["structured"]["fieldGroups"]["qualificationRequirements"]
            joined = "\n".join(row["content"] for row in rows)
            self.assertEqual({row["section"] for row in rows}, {"3. 投标人资格要求"})
            self.assertIn("合法注册的独立法人", joined)
            self.assertIn("同类设备供货业绩", joined)
            self.assertNotIn("见投标人须知前附表", joined)
            self.assertNotIn("营业执照复印件", joined)

    def test_business_contract_uses_shared_selector_for_duplicate_qualification_nodes(self) -> None:
        scripts_dir = self.runner_path().parent
        sys.path.insert(0, str(scripts_dir))
        try:
            contract = importlib.import_module("business_contract")
            section_tree = {
                "schemaVersion": "bid-business-section-tree-v1",
                "nodes": [
                    {
                        "id": "DOC-1-S0001",
                        "documentId": "DOC-1",
                        "level": 2,
                        "number": "3.",
                        "title": "3. 投标人资格要求",
                        "path": ["第一章 招标公告", "3. 投标人资格要求"],
                        "startLine": 3,
                        "endLine": 5,
                    },
                    {
                        "id": "DOC-1-S0002",
                        "documentId": "DOC-1",
                        "level": 3,
                        "number": "1.4",
                        "title": "1.4 投标人资格要求",
                        "path": ["第二章 投标人须知", "1. 总则", "1.4 投标人资格要求"],
                        "startLine": 8,
                        "endLine": 10,
                    },
                ],
            }

            nodes = contract._qualification_section_tree_nodes(section_tree, "DOC-1")

            self.assertEqual([node["title"] for node in nodes], ["3. 投标人资格要求"])
        finally:
            sys.path.remove(str(scripts_dir))

    def test_business_parser_does_not_duplicate_qualification_child_nodes_from_section_tree_path(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "章节树资格子节点.md"
            lines = [
                "# 商务采购文件",
                "第一章 采购公告",
                "3. 供应商资格要求",
                "3.1 通用资格条件",
                "供应商须为中华人民共和国境内合法注册的独立法人。",
                "3.2 专用资格条件",
                "本项目不接受联合体投标。",
                "第二章 供应商须知",
            ]
            source_path.write_text("\n".join(lines), encoding="utf-8")
            section_tree_path = tmp_path / "business_section_tree.json"
            section_tree_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-section-tree-v1",
                        "documents": [{"id": "DOC-1", "name": source_path.name}],
                        "nodes": [
                            {
                                "id": "DOC-1-S0001",
                                "documentId": "DOC-1",
                                "level": 2,
                                "number": "3.",
                                "title": "3. 供应商资格要求",
                                "path": ["第一章 采购公告", "3. 供应商资格要求"],
                                "source": "heading",
                                "confidence": 0.95,
                                "startBlockIndex": 3,
                                "contentStartBlockIndex": 4,
                                "endBlockIndex": 7,
                                "startLine": 3,
                                "contentStartLine": 4,
                                "endLine": 7,
                            },
                            {
                                "id": "DOC-1-S0002",
                                "documentId": "DOC-1",
                                "level": 3,
                                "number": "3.1",
                                "title": "3.1 通用资格条件",
                                "path": ["第一章 采购公告", "3. 供应商资格要求", "3.1 通用资格条件"],
                                "source": "heading",
                                "confidence": 0.93,
                                "startBlockIndex": 4,
                                "contentStartBlockIndex": 5,
                                "endBlockIndex": 5,
                                "startLine": 4,
                                "contentStartLine": 5,
                                "endLine": 5,
                            },
                            {
                                "id": "DOC-1-S0003",
                                "documentId": "DOC-1",
                                "level": 3,
                                "number": "3.2",
                                "title": "3.2 专用资格条件",
                                "path": ["第一章 采购公告", "3. 供应商资格要求", "3.2 专用资格条件"],
                                "source": "heading",
                                "confidence": 0.93,
                                "startBlockIndex": 6,
                                "contentStartBlockIndex": 7,
                                "endBlockIndex": 7,
                                "startLine": 6,
                                "contentStartLine": 7,
                                "endLine": 7,
                            },
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-SECTION-TREE-QUAL-CHILDREN",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "businessSectionTreePath": str(section_tree_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                        "textPath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)

            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            qualification_candidates = candidate_package["candidates"]["qualification"]
            self.assertEqual(len(qualification_candidates), 1)
            self.assertEqual(qualification_candidates[0]["section"], "3. 供应商资格要求")

    def test_business_parser_final_qualification_rows_use_section_tree_scope(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "章节树最终资格范围.md"
            lines = [
                "# 商务采购文件",
                "目录",
                "3. 供应商资格要求 ........ 3",
                "第一章 采购公告",
                "3. 供应商资格要求",
                "3.1 供应商须为中华人民共和国境内合法注册的独立法人。",
                "3.2 本项目不接受联合体投标。",
                "（二）投标人资格要求相关证明材料",
                "投标人须提供近三年业绩证明材料。",
                "第二章 投标人须知",
                "1.4 投标人资格要求",
                "1.4.1 投标人应具备承担本项目的资格条件：见投标人须知前附表。",
                "3.5 资格审查资料",
                "投标人须提供营业执照复印件，本行属于资料要求而非资格要求。",
                "第三章 评审办法",
                "商务评分标准",
                "类似合同业绩每增加一项加2分，本行不应进入资格结果。",
            ]
            source_path.write_text("\n".join(lines), encoding="utf-8")
            section_tree_path = tmp_path / "business_section_tree.json"
            section_tree_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-section-tree-v1",
                        "documents": [{"id": "DOC-1", "name": source_path.name}],
                        "nodes": [
                            {
                                "id": "DOC-1-S0001",
                                "documentId": "DOC-1",
                                "level": 2,
                                "number": "3.",
                                "title": "3. 供应商资格要求",
                                "path": ["第一章 采购公告", "3. 供应商资格要求"],
                                "source": "heading",
                                "confidence": 0.95,
                                "startBlockIndex": 5,
                                "contentStartBlockIndex": 6,
                                "endBlockIndex": 7,
                                "startLine": 5,
                                "contentStartLine": 6,
                                "endLine": 7,
                            },
                            {
                                "id": "DOC-1-S0002",
                                "documentId": "DOC-1",
                                "level": 2,
                                "number": "（二）",
                                "title": "（二）投标人资格要求相关证明材料",
                                "path": ["第一章 采购公告", "（二）投标人资格要求相关证明材料"],
                                "source": "heading",
                                "confidence": 0.9,
                                "startBlockIndex": 8,
                                "contentStartBlockIndex": 9,
                                "endBlockIndex": 9,
                                "startLine": 8,
                                "contentStartLine": 9,
                                "endLine": 9,
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-SECTION-TREE-FINAL-QUAL",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "businessSectionTreePath": str(section_tree_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                        "textPath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run(
                [sys.executable, str(script_path), "offline-fallback", str(manifest_path)],
                check=True,
                capture_output=True,
                text=True,
            )

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            rows = payload["structured"]["fieldGroups"]["qualificationRequirements"]
            contents = [row["content"] for row in rows]
            joined = "\n".join(contents)

            self.assertEqual(len(rows), 2)
            self.assertTrue(any("合法注册的独立法人" in text for text in contents))
            self.assertTrue(any("不接受联合体投标" in text for text in contents))
            self.assertNotIn("业绩证明材料", joined)
            self.assertNotIn("见投标人须知前附表", joined)
            self.assertNotIn("营业执照复印件", joined)
            self.assertNotIn("类似合同业绩", joined)
            self.assertEqual({row["section"] for row in rows}, {"3. 供应商资格要求"})

    def test_business_parser_uses_section_tree_for_supplier_instruction_table(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "供应商须知前附表.md"
            lines = [
                "# 商务采购文件",
                "第一章 采购公告",
                "第二章 供应商须知",
                "供应商须知前附表",
                "| 条款号 | 条款名称 | 编列内容 |",
                "| --- | --- | --- |",
                "| 1.1.2 | 采购人 | 示例采购人 |",
                "| 4.2.1 | 响应文件递交截止时间 | 2026年3月8日09时00分 |",
                "第三章 评审办法",
            ]
            source_path.write_text("\n".join(lines), encoding="utf-8")
            section_tree_path = tmp_path / "business_section_tree.json"
            section_tree_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-section-tree-v1",
                        "documents": [{"id": "DOC-1", "name": source_path.name}],
                        "nodes": [
                            {
                                "id": "DOC-1-S0001",
                                "documentId": "DOC-1",
                                "level": 2,
                                "number": "",
                                "title": "供应商须知前附表",
                                "path": ["第二章 供应商须知", "供应商须知前附表"],
                                "source": "heading",
                                "confidence": 0.96,
                                "startBlockIndex": 4,
                                "contentStartBlockIndex": 5,
                                "endBlockIndex": 8,
                                "startLine": 4,
                                "contentStartLine": 5,
                                "endLine": 8,
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-SECTION-TREE-SUPPLIER-INSTRUCTIONS",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "businessSectionTreePath": str(section_tree_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                        "textPath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), "offline-fallback", str(manifest_path)], check=True, capture_output=True, text=True)

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            rows = payload["structured"]["fieldGroups"]["bidderInstructions"]
            self.assertEqual(len(rows), 2)
            self.assertEqual({row["section"] for row in rows}, {"供应商须知前附表"})
            self.assertEqual(rows[0]["clauseName"], "采购人")
            self.assertEqual(rows[1]["clauseNo"], "4.2.1")
            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            deterministic_rows = candidate_package["deterministicExtracts"]["bidderInstructions"]
            self.assertEqual({row["tableTitle"] for row in deterministic_rows}, {"供应商须知前附表"})

    def test_business_parser_binds_docx_scoring_table_section_from_section_tree(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "章节树商务评分表.docx"
            doc = Document()
            doc.add_paragraph("第三章 评审办法")
            doc.add_paragraph("商务评分标准")
            doc.add_paragraph("以下表格按项目实际情况进行评分。")
            scoring_table = doc.add_table(rows=2, cols=4)
            for col, text in enumerate(["序号", "评审因素", "分值", "评审标准"]):
                scoring_table.cell(0, col).text = text
            for col, text in enumerate(["1", "企业业绩", "20", "近三年类似项目业绩满足要求得20分。"]):
                scoring_table.cell(1, col).text = text
            doc.save(source_path)
            section_tree_path = tmp_path / "business_section_tree.json"
            section_tree_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-section-tree-v1",
                        "documents": [{"id": "DOC-1", "name": source_path.name}],
                        "nodes": [
                            {
                                "id": "DOC-1-S0001",
                                "documentId": "DOC-1",
                                "level": 2,
                                "number": "",
                                "title": "商务评分标准",
                                "path": ["第三章 评审办法", "商务评分标准"],
                                "source": "heading",
                                "confidence": 0.96,
                                "startBlockIndex": 2,
                                "contentStartBlockIndex": 3,
                                "endBlockIndex": 4,
                                "startLine": 2,
                                "contentStartLine": 3,
                                "endLine": 4,
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-SECTION-TREE-SCORING-TABLE",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "businessSectionTreePath": str(section_tree_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)

            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            table = next(item for item in candidate_package["tables"] if item["headers"] == ["序号", "评审因素", "分值", "评审标准"])
            self.assertEqual(table["section"], "商务评分标准")
            self.assertEqual(table["tableType"], "business")
            scoring_rows = candidate_package["deterministicExtracts"]["scoringTables"]["business"]
            self.assertEqual(len(scoring_rows), 1)
            self.assertEqual(scoring_rows[0]["section"], "商务评分标准")

    def test_business_parser_uses_external_ai_decision_source(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "外部审查商务片段.md"
            source_path.write_text(
                "\n".join(
                    [
                        "# 商务招标文件",
                        "第一章 招标公告",
                        "3. 投标人资格要求",
                        "3.1 投标人须为中华人民共和国境内合法注册的独立法人或其他组织。",
                        "3.2 投标人须具有近三年同类设备供货业绩。",
                    ]
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-EXTERNAL-AI",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                        "textPath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)

            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            qualification_candidates = candidate_package["candidates"]["qualification"]
            qualification_slice = qualification_candidates[0]
            accepted_line = next(line for line in qualification_slice["lines"] if "独立法人" in line["text"])
            rejected_line = next(line for line in qualification_slice["lines"] if "供货业绩" in line["text"])
            external_dir = tmp_path / "external_ai_decisions"
            external_dir.mkdir()
            for task_name in (
                "project_facts_review",
                "bidder_instructions_review",
                "rejection_clause_review",
                "scoring_review",
                "business_response_review",
                "commitment_review",
                "appendix_review",
            ):
                (external_dir / f"{task_name}.json").write_text(
                    json.dumps(
                        {
                            "schemaVersion": "bid-business-ai-decision-v1",
                            "task": task_name,
                            "adapter": "unit-test-external-ai",
                            "accepted": [],
                            "rejected": [],
                            "needsReview": [],
                            "reason": "外部 AI 本轮无采纳项。",
                            "evidenceIds": [],
                        },
                        ensure_ascii=False,
                        indent=2,
                    ),
                    encoding="utf-8",
                )
            (external_dir / "qualification_review.json").write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-ai-decision-v1",
                        "task": "qualification_review",
                        "adapter": "unit-test-external-ai",
                        "qualificationItems": [
                            {
                                "content": accepted_line["text"],
                                "applicableScope": "全部标段",
                                "sourceText": accepted_line["sourceText"],
                                "evidenceIds": [accepted_line["evidenceId"]],
                            }
                        ],
                        "rejectedEvidenceIds": [rejected_line["evidenceId"]],
                        "reason": "使用外部 AI 决策源覆盖默认离线审查。",
                        "evidenceIds": [accepted_line["evidenceId"], rejected_line["evidenceId"]],
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )

            manifest_with_external = {
                **manifest,
                "aiDecisionSourceDir": str(external_dir),
                "aiDecisionsDir": str(tmp_path / "normalized_ai_decisions"),
                "validationReportPath": str(tmp_path / "external_validation_report.json"),
            }
            manifest_path.write_text(json.dumps(manifest_with_external, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), "finalize", str(manifest_path)], check=True, capture_output=True, text=True)

            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            qualification_ref = next(item for item in review_plan["tasks"] if item["task"] == "qualification_review")
            decision = json.loads((tmp_path / qualification_ref["decisionPath"]).read_text(encoding="utf-8"))
            self.assertEqual(decision["adapter"], "unit-test-external-ai")
            payload = json.loads(output_path.read_text(encoding="utf-8"))
            workflow = payload["structured"]["workflow"]
            self.assertEqual(workflow["stage"], "finalized")
            self.assertTrue(workflow["aiReviewTrusted"])
            self.assertEqual(workflow["semanticReviewMode"], "opencode-agent")
            self.assertFalse(workflow["offlineAdapterUsed"])
            qualification_text = "\n".join(
                row["content"] for row in payload["structured"]["fieldGroups"]["qualificationRequirements"]
            )
            self.assertIn("合法注册的独立法人", qualification_text)
            self.assertNotIn("同类设备供货业绩", qualification_text)
            validation = json.loads((tmp_path / "external_validation_report.json").read_text(encoding="utf-8"))
            decision_check = next(check for check in validation["checks"] if check["name"] == "ai_decision_contract")
            self.assertEqual(decision_check["status"], "passed")

    def test_business_parser_rejects_ai_decisions_without_candidate_evidence(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "非法审查商务片段.md"
            source_path.write_text(
                "\n".join(
                    [
                        "# 商务招标文件",
                        "第一章 招标公告",
                        "3. 投标人资格要求",
                        "3.1 投标人须为中华人民共和国境内合法注册的独立法人或其他组织。",
                    ]
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-INVALID-AI",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                        "textPath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)

            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            accepted_candidate = candidate_package["candidates"]["qualification"][0]
            accepted_line = next(line for line in accepted_candidate["lines"] if "独立法人" in line["text"])
            external_dir = tmp_path / "invalid_ai_decisions"
            external_dir.mkdir()
            (external_dir / "qualification_review.json").write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-ai-decision-v1",
                        "task": "qualification_review",
                        "adapter": "unit-test-invalid-ai",
                        "qualificationItems": [
                            {
                                "content": accepted_line["text"],
                                "applicableScope": "全部标段",
                                "sourceText": accepted_line["sourceText"],
                                "evidenceIds": [],
                            },
                            {
                                "content": accepted_line["text"],
                                "applicableScope": "全部标段",
                                "sourceText": accepted_line["sourceText"],
                                "evidenceIds": ["DOC-1:L999"],
                            },
                            {
                                "content": accepted_line["text"],
                                "applicableScope": "全部标段",
                                "sourceText": "L4",
                                "evidenceIds": [accepted_line["evidenceId"]],
                            },
                        ],
                        "rejectedEvidenceIds": [],
                        "reason": "这些接受项都应被验真层剔除。",
                        "evidenceIds": [accepted_line["evidenceId"], "DOC-1:L999"],
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )

            manifest_with_external = {
                **manifest,
                "aiDecisionSourceDir": str(external_dir),
                "aiDecisionsDir": str(tmp_path / "normalized_invalid_ai_decisions"),
                "validationReportPath": str(tmp_path / "invalid_validation_report.json"),
            }
            manifest_path.write_text(json.dumps(manifest_with_external, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), "finalize", str(manifest_path)], check=True, capture_output=True, text=True)

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            qualification_text = "\n".join(
                row["content"] for row in payload["structured"]["fieldGroups"]["qualificationRequirements"]
            )
            self.assertNotIn("合法注册的独立法人", qualification_text)
            validation = json.loads((tmp_path / "invalid_validation_report.json").read_text(encoding="utf-8"))
            decision_check = next(check for check in validation["checks"] if check["name"] == "ai_decision_contract")
            self.assertEqual(decision_check["status"], "failed")
            self.assertGreaterEqual(decision_check["count"], 3)
            self.assertEqual(payload["structured"]["workflow"]["validationStatus"], "failed")

    def test_business_parser_rejects_ai_decision_items_missing_required_contract_fields(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "缺字段审查商务片段.md"
            source_path.write_text(
                "\n".join(
                    [
                        "# 商务招标文件",
                        "第一章 招标公告",
                        "3. 投标人资格要求",
                        "3.1 投标人须为中华人民共和国境内合法注册的独立法人或其他组织。",
                    ]
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-MISSING-CONTRACT-FIELDS",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                        "textPath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)

            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            accepted_candidate = candidate_package["candidates"]["qualification"][0]
            accepted_line = next(line for line in accepted_candidate["lines"] if "独立法人" in line["text"])
            external_dir = tmp_path / "missing_contract_fields_ai_decisions"
            external_dir.mkdir()
            (external_dir / "qualification_review.json").write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-ai-decision-v1",
                        "task": "qualification_review",
                        "adapter": "unit-test-missing-contract-fields",
                        "qualificationItems": [
                            {
                                "content": accepted_line["text"],
                                "evidenceIds": [accepted_line["evidenceId"]],
                            }
                        ],
                        "reason": "缺字段决策项应被验真层剔除。",
                        "evidenceIds": [accepted_line["evidenceId"]],
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )

            manifest_with_external = {
                **manifest,
                "aiDecisionSourceDir": str(external_dir),
                "aiDecisionsDir": str(tmp_path / "normalized_missing_contract_fields_ai_decisions"),
                "validationReportPath": str(tmp_path / "missing_contract_fields_validation_report.json"),
            }
            manifest_path.write_text(json.dumps(manifest_with_external, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), "finalize", str(manifest_path)], check=True, capture_output=True, text=True)

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            qualification_text = "\n".join(
                row["content"] for row in payload["structured"]["fieldGroups"]["qualificationRequirements"]
            )
            self.assertNotIn("合法注册的独立法人", qualification_text)
            validation = json.loads((tmp_path / "missing_contract_fields_validation_report.json").read_text(encoding="utf-8"))
            decision_check = next(check for check in validation["checks"] if check["name"] == "ai_decision_contract")
            self.assertEqual(decision_check["status"], "failed")
            self.assertEqual(validation["aiDecisionIssues"][0]["code"], "missing_required_decision_fields")
            self.assertEqual(payload["structured"]["workflow"]["validationStatus"], "failed")

    def test_business_parser_finalize_does_not_call_openai_compatible_ai_adapter(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "模型审查商务片段.md"
            source_path.write_text(
                "\n".join(
                    [
                        "# 商务招标文件",
                        "第一章 招标公告",
                        "3. 投标人资格要求",
                        "3.1 投标人须为中华人民共和国境内合法注册的独立法人或其他组织。",
                        "3.2 投标人须具有近三年同类设备供货业绩。",
                    ]
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            seed_manifest = {
                "projectId": "PRJ-OPENAI-ADAPTER",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                        "textPath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(seed_manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)

            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            accepted_candidate = candidate_package["candidates"]["qualification"][0]
            review_plan = json.loads((tmp_path / "review_plan.json").read_text(encoding="utf-8"))
            decisions_dir = tmp_path / "agent_ai_decisions"
            for task_ref in review_plan["tasks"]:
                task_path = tmp_path / task_ref["taskPath"]
                decision_path = decisions_dir / Path(task_ref["decisionPath"]).relative_to("ai_decisions")
                self.write_plan_decision(
                    task_path,
                    decision_path,
                    accepted_candidate_ids={accepted_candidate["id"]},
                )

            manifest = {
                **seed_manifest,
                "aiReviewMode": "openai-compatible",
                "aiReviewBaseUrl": "http://127.0.0.1:9",
                "aiReviewApiKey": "unit-test-key",
                "aiReviewModel": "unit-test-model",
                "aiDecisionsDir": str(decisions_dir),
                "validationReportPath": str(tmp_path / "agent_validation_report.json"),
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run(
                [sys.executable, str(script_path), "finalize", str(manifest_path)],
                check=True,
                capture_output=True,
                text=True,
            )

            qualification_ref = next(item for item in review_plan["tasks"] if item["task"] == "qualification_review")
            decision = json.loads((decisions_dir / Path(qualification_ref["decisionPath"]).relative_to("ai_decisions")).read_text(encoding="utf-8"))
            self.assertEqual(decision["adapter"], "unit-test-opencode-agent")
            payload = json.loads(output_path.read_text(encoding="utf-8"))
            workflow = payload["structured"]["workflow"]
            self.assertEqual(workflow["stage"], "finalized")
            self.assertTrue(workflow["aiReviewTrusted"])
            self.assertEqual(workflow["semanticReviewMode"], "opencode-agent")
            self.assertFalse(workflow["offlineAdapterUsed"])
            qualification_text = "\n".join(
                row["content"] for row in payload["structured"]["fieldGroups"]["qualificationRequirements"]
            )
            self.assertIn("合法注册的独立法人", qualification_text)

    def test_business_parser_rejects_ai_decisions_with_final_structured_fields(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "越权审查商务片段.md"
            source_path.write_text(
                "\n".join(
                    [
                        "# 商务招标文件",
                        "第一章 招标公告",
                        "3. 投标人资格要求",
                        "3.1 投标人须为中华人民共和国境内合法注册的独立法人或其他组织。",
                    ]
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-FORBIDDEN-AI",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                        "textPath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)

            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            accepted_candidate = candidate_package["candidates"]["qualification"][0]
            accepted_line = next(line for line in accepted_candidate["lines"] if "独立法人" in line["text"])
            external_dir = tmp_path / "unexpected_ai_decisions"
            external_dir.mkdir()
            (external_dir / "qualification_review.json").write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-ai-decision-v1",
                        "task": "qualification_review",
                        "adapter": "unit-test-unexpected-ai",
                        "structured": {"fieldGroups": {"qualificationRequirements": [{"content": "模型越权最终结果"}]}},
                        "qualificationItems": [
                            {
                                "content": accepted_line["text"],
                                "applicableScope": "全部标段",
                                "sourceText": accepted_line["sourceText"],
                                "evidenceIds": [accepted_line["evidenceId"]],
                            }
                        ],
                        "reason": "应被脚本拒绝。",
                        "evidenceIds": [accepted_line["evidenceId"]],
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )

            manifest_with_external = {
                **manifest,
                "aiDecisionSourceDir": str(external_dir),
                "aiDecisionsDir": str(tmp_path / "normalized_unexpected_ai_decisions"),
                "validationReportPath": str(tmp_path / "unexpected_validation_report.json"),
            }
            manifest_path.write_text(json.dumps(manifest_with_external, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), "finalize", str(manifest_path)], check=True, capture_output=True, text=True)

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            qualification_text = "\n".join(
                row["content"] for row in payload["structured"]["fieldGroups"]["qualificationRequirements"]
            )
            self.assertNotIn("合法注册的独立法人", qualification_text)
            self.assertNotIn("模型越权最终结果", json.dumps(payload, ensure_ascii=False))
            validation = json.loads((tmp_path / "unexpected_validation_report.json").read_text(encoding="utf-8"))
            decision_check = next(check for check in validation["checks"] if check["name"] == "ai_decision_contract")
            self.assertEqual(decision_check["status"], "failed")
            self.assertEqual(validation["aiDecisionIssues"][0]["code"], "unexpected_top_level_fields")

    def test_business_parser_rejects_ai_decision_with_invalid_schema_version(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "无效契约商务片段.md"
            source_path.write_text(
                "\n".join(
                    [
                        "# 商务招标文件",
                        "第一章 招标公告",
                        "3. 投标人资格要求",
                        "3.1 投标人须为中华人民共和国境内合法注册的独立法人或其他组织。",
                    ]
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            manifest = {
                "projectId": "PRJ-INVALID-SCHEMA-AI",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                        "textPath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), "finalize", str(manifest_path)], check=True, capture_output=True, text=True)

            candidate_package = json.loads((tmp_path / "candidate_package.json").read_text(encoding="utf-8"))
            accepted_candidate = candidate_package["candidates"]["qualification"][0]
            accepted_line = next(line for line in accepted_candidate["lines"] if "独立法人" in line["text"])
            external_dir = tmp_path / "invalid_schema_ai_decisions"
            external_dir.mkdir()
            (external_dir / "qualification_review.json").write_text(
                json.dumps(
                    {
                        "schemaVersion": "legacy-decision-schema",
                        "task": "qualification_review",
                        "adapter": "unit-test-invalid-schema-ai",
                        "qualificationItems": [
                            {
                                "content": accepted_line["text"],
                                "applicableScope": "全部标段",
                                "sourceText": accepted_line["sourceText"],
                                "evidenceIds": [accepted_line["evidenceId"]],
                            }
                        ],
                        "reason": "应被脚本拒绝。",
                        "evidenceIds": [accepted_line["evidenceId"]],
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )

            manifest_with_external = {
                **manifest,
                "aiDecisionSourceDir": str(external_dir),
                "aiDecisionsDir": str(tmp_path / "normalized_invalid_schema_ai_decisions"),
                "validationReportPath": str(tmp_path / "invalid_schema_validation_report.json"),
            }
            manifest_path.write_text(json.dumps(manifest_with_external, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), "finalize", str(manifest_path)], check=True, capture_output=True, text=True)

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            qualification_text = "\n".join(
                row["content"] for row in payload["structured"]["fieldGroups"]["qualificationRequirements"]
            )
            self.assertNotIn("合法注册的独立法人", qualification_text)
            validation = json.loads((tmp_path / "invalid_schema_validation_report.json").read_text(encoding="utf-8"))
            decision_check = next(check for check in validation["checks"] if check["name"] == "ai_decision_contract")
            self.assertEqual(decision_check["status"], "failed")
            self.assertEqual(validation["aiDecisionIssues"][0]["code"], "invalid_schema_version")

    def test_business_parser_marks_invalid_full_ai_review_as_untrusted_fallback(self) -> None:
        script_path = self.runner_path()

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source_path = tmp_path / "全量非法审查商务片段.md"
            source_path.write_text(
                "\n".join(
                    [
                        "# 商务招标文件",
                        "第一章 招标公告",
                        "3. 投标人资格要求",
                        "3.1 投标人须为中华人民共和国境内合法注册的独立法人或其他组织。",
                    ]
                ),
                encoding="utf-8",
            )
            output_path = tmp_path / "s1_structured_result.json"
            manifest_path = tmp_path / "s1_parse_manifest.json"
            seed_manifest = {
                "projectId": "PRJ-FULL-INVALID-AI",
                "bidType": "商务标",
                "parseProfile": "business",
                "structuredResultPath": str(output_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source_path.name,
                        "sourcePath": str(source_path),
                        "textPath": str(source_path),
                    }
                ],
            }
            manifest_path.write_text(json.dumps(seed_manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), str(manifest_path)], check=True, capture_output=True, text=True)

            external_dir = tmp_path / "full_invalid_ai_decisions"
            external_dir.mkdir()
            for task_name in (
                "project_facts_review",
                "bidder_instructions_review",
                "qualification_review",
                "rejection_clause_review",
                "scoring_review",
                "business_response_review",
                "commitment_review",
                "appendix_review",
            ):
                (external_dir / f"{task_name}.json").write_text(
                    json.dumps(
                        {
                            "schemaVersion": "legacy-decision-schema",
                            "task": task_name,
                            "adapter": "unit-test-invalid-full-ai",
                            "accepted": [],
                            "rejected": [],
                            "needsReview": [],
                            "reason": "schemaVersion 非法时不能标记为可信 finalized。",
                            "evidenceIds": [],
                        },
                        ensure_ascii=False,
                        indent=2,
                    ),
                    encoding="utf-8",
                )

            manifest = {
                **seed_manifest,
                "aiDecisionSourceDir": str(external_dir),
                "aiDecisionsDir": str(tmp_path / "full_invalid_normalized_ai_decisions"),
                "validationReportPath": str(tmp_path / "full_invalid_validation_report.json"),
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            subprocess.run([sys.executable, str(script_path), "finalize", str(manifest_path)], check=True, capture_output=True, text=True)

            payload = json.loads(output_path.read_text(encoding="utf-8"))
            workflow = payload["structured"]["workflow"]
            self.assertEqual(workflow["stage"], "fallback")
            self.assertFalse(workflow["aiReviewTrusted"])
            self.assertEqual(workflow["validationStatus"], "failed")
