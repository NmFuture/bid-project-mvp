"""待插入素材备料单测：扫描占位符 → 按素材范围检索 → 写入 manifest.embedSources。

素材检索要联网查库，按「skill 只依据 manifest 工作」的既有约定放在后端；filler 只依据
这里给出的本地路径与 status 决定嵌入还是标黄。
fixture 为脱敏合成数据（通用领域词面，无真实项目数据）。
"""
from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from app.services import technical_gap_ai_fill as ai_fill


def _write_docx(path: Path, paragraphs: list[str], cell_text: str = "") -> None:
    from docx import Document

    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if cell_text:
        table = document.add_table(rows=1, cols=1)
        table.rows[0].cells[0].text = cell_text
    document.save(str(path))


class ScanEmbedPlaceholdersTests(unittest.TestCase):
    def test_scans_paragraph_embed_placeholders_and_dedupes(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            path = Path(raw) / "待填写-方案.docx"
            _write_docx(
                path,
                [
                    "本项目安全等级为[安全等级，待填写]。",
                    "[设备清单，待插入]",
                    "【设备清单, 待插入】",
                    "详见[附件材料，待插入]后附。",
                ],
                cell_text="[表内素材，待插入]",
            )

            labels = ai_fill._scan_embed_placeholders(path)

        # 待填写不进来；同名占位符归一化后去重；表格单元格不备料（filler 一律标黄）
        self.assertEqual(labels, ["设备清单", "附件材料"])

    def test_document_without_embed_placeholders_returns_empty(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            path = Path(raw) / "待填写-方案.docx"
            _write_docx(path, ["本项目安全等级为[安全等级，待填写]。"])

            self.assertEqual(ai_fill._scan_embed_placeholders(path), [])


class MaterialSpecificityTests(unittest.TestCase):
    def test_more_specific_tier_wins(self) -> None:
        picked, ambiguous = ai_fill._pick_most_specific_material(
            [
                {"id": "RAW-1", "name": "设备清单.docx", "materialTier": "standard"},
                {"id": "RAW-2", "name": "设备清单.docx", "materialTier": "project"},
                {"id": "RAW-3", "name": "设备清单.docx", "materialTier": "customer"},
            ]
        )

        self.assertEqual(picked["id"], "RAW-2")
        self.assertFalse(ambiguous)

    def test_same_tier_collision_is_ambiguous(self) -> None:
        _picked, ambiguous = ai_fill._pick_most_specific_material(
            [
                {"id": "RAW-1", "name": "设备清单.docx", "materialTier": "project"},
                {"id": "RAW-2", "name": "设备清单.docx", "materialTier": "project"},
            ]
        )

        self.assertTrue(ambiguous)


class EmbedSourcesTests(unittest.TestCase):
    def setUp(self) -> None:
        self.project = {"id": "PRJ-0001", "name": "示例项目", "bidType": "技术标"}

    def _run(self, tmp: Path, materials: list[dict], paragraphs: list[str]) -> list[dict]:
        blank = tmp / "待填写-方案.docx"
        _write_docx(blank, paragraphs)
        with (
            patch.object(ai_fill, "build_project_material_scope", return_value={"readableScopes": []}),
            patch.object(ai_fill, "project_turbine_model", return_value={}),
            patch.object(ai_fill, "_allowed_technical_material_index", return_value=materials),
        ):
            return ai_fill._embed_sources_for_fill(self.project, blank, tmp)

    def test_excel_material_is_refused_by_original_suffix(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            sources = self._run(
                Path(raw),
                [{"id": "RAW-9", "name": "价格表.xlsx", "materialTier": "project", "folderPath": "技术标/项目定制"}],
                ["[价格表，待插入]"],
            )

        self.assertEqual(len(sources), 1)
        # 按原始后缀判断，不看能不能取到 docx：xlsx 被 Wiki 预览转换过时会拿到有损清洗稿，
        # 悄悄嵌进投标材料就是静默降级
        self.assertEqual(sources[0]["status"], "unsupported_format")
        self.assertIn("另存为 Word", sources[0]["statusMessage"])
        self.assertNotIn("docxPath", sources[0])

    def test_missing_material_is_reported_not_silently_skipped(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            sources = self._run(Path(raw), [], ["[设备清单，待插入]"])

        self.assertEqual(len(sources), 1)
        self.assertEqual(sources[0]["status"], "not_found")
        self.assertIn("设备清单", sources[0]["statusMessage"])

    def test_same_tier_collision_is_marked_ambiguous(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            sources = self._run(
                Path(raw),
                [
                    {"id": "RAW-1", "name": "设备清单.docx", "materialTier": "project"},
                    {"id": "RAW-2", "name": "设备清单.docx", "materialTier": "project"},
                ],
                ["[设备清单，待插入]"],
            )

        self.assertEqual(sources[0]["status"], "ambiguous")
        self.assertEqual(sources[0]["candidateCount"], 2)

    def test_download_failure_degrades_to_manual_not_exception(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            blank = Path(raw) / "待填写-方案.docx"
            _write_docx(blank, ["[设备清单，待插入]"])
            with (
                patch.object(ai_fill, "build_project_material_scope", return_value={"readableScopes": []}),
                patch.object(ai_fill, "project_turbine_model", return_value={}),
                patch.object(
                    ai_fill,
                    "_allowed_technical_material_index",
                    return_value=[{"id": "RAW-1", "name": "设备清单.docx", "materialTier": "project"}],
                ),
                patch.object(ai_fill, "_run_async", side_effect=RuntimeError("minio down")),
            ):
                sources = ai_fill._embed_sources_for_fill(self.project, blank, Path(raw))

        # 一份素材取不到不能中断整份文件的填写
        self.assertEqual(sources[0]["status"], "download_failed")
        self.assertIn("minio down", sources[0]["statusMessage"])

    def test_ready_source_carries_local_path_for_the_skill(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            tmp = Path(raw)
            blank = tmp / "待填写-方案.docx"
            _write_docx(blank, ["[设备清单，待插入]"])
            payload = {"bucket": "materials", "key": "cleaned/设备清单.docx", "fileName": "设备清单.docx"}

            def _run_async_stub(awaitable: object) -> tuple[dict, str]:
                # 真实 _run_async 会 await 掉协程；mock 不 await，这里显式关闭免得留下未等待警告
                if hasattr(awaitable, "close"):
                    awaitable.close()
                return payload, "cleaned"

            with (
                patch.object(ai_fill, "build_project_material_scope", return_value={"readableScopes": []}),
                patch.object(ai_fill, "project_turbine_model", return_value={}),
                patch.object(
                    ai_fill,
                    "_allowed_technical_material_index",
                    return_value=[{"id": "RAW-1", "name": "设备清单.docx", "materialTier": "project"}],
                ),
                patch.object(ai_fill, "_run_async", side_effect=_run_async_stub),
                patch.object(ai_fill.minio_client, "download_file") as download,
            ):
                sources = ai_fill._embed_sources_for_fill(self.project, blank, tmp)

        self.assertEqual(sources[0]["status"], "ready")
        self.assertEqual(sources[0]["materialTier"], "project")
        self.assertTrue(sources[0]["docxPath"].endswith(".docx"))
        self.assertIn("embed_sources", sources[0]["docxPath"])
        download.assert_called_once()

    def test_no_embed_placeholder_skips_material_lookup_entirely(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            tmp = Path(raw)
            blank = tmp / "待填写-方案.docx"
            _write_docx(blank, ["本项目安全等级为[安全等级，待填写]。"])
            with patch.object(ai_fill, "_allowed_technical_material_index") as lookup:
                sources = ai_fill._embed_sources_for_fill(self.project, blank, tmp)

        self.assertEqual(sources, [])
        lookup.assert_not_called()


if __name__ == "__main__":
    unittest.main()
