from __future__ import annotations

"""技术标事实表专项抽取器（technical_fact_special_extractors）的测试。

样本用 python-docx/openpyxl 在临时目录现造，断言 label/value/location，
并覆盖 reconcile 集成、路由器与 material_is_fact_relevant 的关键词补充。
"""

import tempfile
import unittest
from pathlib import Path

from docx import Document
from openpyxl import Workbook

from app.services.technical_fact_field_specs import fillable_specs
from app.services.technical_fact_special_extractors import (
    facts_from_certificate_materials,
    facts_from_foundation_moment_xlsx,
    facts_from_hours_commitment_docx,
    facts_from_production_base_docx,
    facts_from_tower_quantity_docx,
    facts_from_wind_resource_docx,
    parse_certificate_wind_params,
    special_extractor_for_material,
)
from app.services.technical_gap_fact_table import (
    FACT_STATUS_UNEXTRACTED,
    fact_label_key,
    material_is_fact_relevant,
    reconcile_fact_fields_with_specs,
)


def _material(name: str, folder: str = "项目定制") -> dict:
    return {"id": f"test-{name}", "name": name, "folderPath": folder, "materialTier": "project"}


def _facts_by_label(facts: list[dict]) -> dict[str, dict]:
    return {str(fact.get("label")): fact for fact in facts}


def _write_docx(path: Path, builder) -> Path:
    document = Document()
    builder(document)
    document.save(path)
    return path


class TestTowerQuantityExtractor(unittest.TestCase):
    def test_steel_tower_sections_and_renames(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "塔架与基础工程量.docx"

            def build(document: Document) -> None:
                table = document.add_table(rows=0, cols=3)
                rows = [
                    ("塔筒段数（段）", "塔筒段数（段）", "5"),
                    ("第2段 塔节Q355NE", "长度（m）", "20.72"),
                    ("第2段 塔节Q355NE", "底部直径（m）", "4.998"),
                    ("第1段(底) 塔节Q355NE", "重量（kg）", "129122"),
                    ("钢材型号Q355NE的筒节质量（kg）", "钢材型号Q355NE的筒节质量（kg）", "452317"),
                    ("塔架总重（筒壁+法兰+内附件）（kg）", "塔架总重（筒壁+法兰+内附件）（kg）", "~ 527632"),
                    ("挖方（m3）", "挖方（m3）", "2628.934"),
                ]
                for row in rows:
                    cells = table.add_row().cells
                    for idx, value in enumerate(row):
                        cells[idx].text = value

            _write_docx(path, build)
            facts = facts_from_tower_quantity_docx(path, _material(path.name), {})
        by_label = _facts_by_label(facts)
        # 段名状态机：col0 段名 + col1 子项拼 spec label
        self.assertEqual(by_label["第2段塔节长度"]["value"], "20.72")
        self.assertEqual(by_label["第2段塔节底部直径"]["value"], "4.998")
        self.assertEqual(by_label["第1段（底）塔节重量（kg）"]["value"], "129122")
        # 命名差异映射 + "~" 前缀剥离
        self.assertEqual(by_label["Q355NE筒节质量（kg）"]["value"], "452317")
        self.assertEqual(by_label["塔架总重（筒壁+法兰+内附件，kg）"]["value"], "527632")
        self.assertEqual(by_label["基础挖方量（m³）"]["value"], "2628.934")
        self.assertEqual(by_label["塔筒段数（段）"]["value"], "5")
        # location 与来源结构
        fact = by_label["第2段塔节长度"]
        self.assertRegex(str(fact["sourceRef"]["location"]), r"^T1/R\d+$")
        self.assertEqual(fact["sourceRef"]["type"], "materialFact")
        # 不在 spec 清单内的 label 不得产出
        self.assertNotIn("机型", by_label)

    def test_mixed_tower_foundation_rows(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "混塔工程量.docx"

            def build(document: Document) -> None:
                table = document.add_table(rows=0, cols=2)
                rows = [
                    ("三、基础参数", ""),
                    ("开挖量（m3）", "1850"),
                    ("垫层混凝土等级", "C15"),
                    ("钢筋标号", "HRB400"),
                    ("钢筋用量（t）", "68"),
                ]
                for label, value in rows:
                    cells = table.add_row().cells
                    cells[0].text = label
                    cells[1].text = value

            _write_docx(path, build)
            facts = facts_from_tower_quantity_docx(path, _material(path.name), {})
        by_label = _facts_by_label(facts)
        self.assertEqual(by_label["基础挖方量（m³）"]["value"], "1850")
        self.assertEqual(by_label["垫层混凝土型号"]["value"], "C15")
        self.assertEqual(by_label["基础钢筋型号"]["value"], "HRB400")
        # t → kg 换算
        self.assertEqual(by_label["基础钢筋用量（kg）"]["value"], "68000")

    def test_unsupported_suffix_returns_none(self) -> None:
        facts = facts_from_tower_quantity_docx(Path("工程量.pdf"), _material("工程量.pdf"), {})
        self.assertIsNone(facts)


class TestFoundationMomentExtractor(unittest.TestCase):
    def _write_moment_xlsx(self, path: Path) -> Path:
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Sheet1"
        sheet.append([None, "EW6.25-220-160"])
        header = [None, "载荷工况", "Mx", "My", "Mxy", "Mz", "Fx", "Fy", "Fxy", "Fz", "Safety factor"]
        sheet.append(header)
        sheet.append([None, None, "kNm", "kNm", "kNm", "kNm", "kN", "kN", "kN", "kN", "-"])
        sheet.append([None, "正常运行载荷工况", 46791.3, 144611, 145463, -12527.2, 1269.02, -331.649, 1272.65, -8154.59, 1])
        sheet.append([None, "极端载荷工况", 175523, 159704, 184172, -19490.8, 1304.71, -1500.44, 1517.06, -8166.27, 1])
        workbook.save(path)
        return path

    def test_moment_table_all_18_labels(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = self._write_moment_xlsx(Path(tmp) / "基础弯矩表.xlsx")
            facts = facts_from_foundation_moment_xlsx(path, _material(path.name), {})
        by_label = _facts_by_label(facts)
        self.assertEqual(len(facts), 18)
        self.assertEqual(by_label["正常工况-Mx（kNm）"]["value"], "46791.3")
        self.assertEqual(by_label["极端工况-Mxy（kNm）"]["value"], "184172")
        self.assertEqual(by_label["正常工况-Fy（kN）"]["value"], "-331.65")
        self.assertEqual(by_label["极端工况-Fz（kN）"]["value"], "-8166.27")
        # 安全系数无单位
        self.assertEqual(by_label["正常工况-安全系数"]["value"], "1")
        self.assertIn("安全系数", [label.split("-")[1] for label in by_label if label.startswith("极端")])
        self.assertRegex(
            str(by_label["极端工况-Mx（kNm）"]["sourceRef"]["location"]), r"^Sheet1!R\d+$"
        )

    def test_moment_reconcile_integration(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = self._write_moment_xlsx(Path(tmp) / "基础弯矩表.xlsx")
            facts = facts_from_foundation_moment_xlsx(path, _material(path.name), {})
        fields_by_key = {
            fact_label_key(fact["label"]): {
                "label": fact["label"],
                "value": fact["value"],
                "status": "extracted",
                "sourceRefs": [fact["sourceRef"]],
            }
            for fact in facts
        }
        reconcile_fact_fields_with_specs(fields_by_key)
        field = fields_by_key[fact_label_key("极端工况-Mx（kNm）")]
        spec = next(s for s in fillable_specs() if s.get("label") == "极端工况-Mx（kNm）")
        self.assertEqual(field.get("specKey"), spec.get("key"))
        self.assertEqual(field.get("specSeq"), 51)
        self.assertEqual(field.get("sourceKind"), "material")
        self.assertEqual(field.get("value"), "175523")
        # 未抽到的 spec 生成"未提取"骨架
        skeleton = fields_by_key[fact_label_key("塔筒段数（段）")]
        self.assertEqual(skeleton.get("status"), FACT_STATUS_UNEXTRACTED)


class TestWindResourceExtractor(unittest.TestCase):
    def _write_wind_docx(self, path: Path) -> Path:
        def build(document: Document) -> None:
            document.add_paragraph("风资源评估报告")
            commitment = document.add_table(rows=0, cols=4)
            rows = [
                ("项目", "单位", "承诺考核值", "承诺保证值"),
                ("机型", "-", "EW10.0-220-125", "EW10.0-220-125"),
                ("容量", "MW", "600", "600"),
                ("折减系数", "%", "77", "75"),
                ("全场净上网电量", "MWh/y", "1750086", "1704629"),
                ("等效上网小时数", "h", "2917", "2841"),
            ]
            for row in rows:
                cells = commitment.add_row().cells
                for idx, value in enumerate(row):
                    cells[idx].text = value
            rep_year = document.add_table(rows=0, cols=4)
            for row in [
                ("测风塔", "轮毂高度(m)", "测量风速(m/s)", "代表年风速(m/s)"),
                ("26304#", "125", "7.36", "7.36"),
                ("210302#", "125", "6.86", "6.86"),
            ]:
                cells = rep_year.add_row().cells
                for idx, value in enumerate(row):
                    cells[idx].text = value
            config = document.add_table(rows=0, cols=2)
            for row in [("轮毂高度(m)", "125"), ("适用等级", "IEC S")]:
                cells = config.add_row().cells
                cells[0].text = row[0]
                cells[1].text = row[1]

        return _write_docx(path, build)

    def test_commitment_double_column_table(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = self._write_wind_docx(Path(tmp) / "风资源评估报告.docx")
            facts = facts_from_wind_resource_docx(path, _material(path.name), {})
        by_label = _facts_by_label(facts)
        # 第1列=承诺考核值、第2列=承诺保证值
        self.assertEqual(by_label["承诺方式（第1列）"]["value"], "承诺考核值")
        self.assertEqual(by_label["承诺方式（第2列）"]["value"], "承诺保证值")
        self.assertEqual(by_label["折减系数（考核值，%）"]["value"], "77")
        self.assertEqual(by_label["折减系数（保证值，%）"]["value"], "75")
        self.assertEqual(by_label["折减值（考核值）"]["value"], "77")
        self.assertEqual(by_label["全场净上网电量（考核值，MWh/y）"]["value"], "1750086")
        self.assertEqual(by_label["全场净上网电量（保证值，MWh/y）"]["value"], "1704629")
        self.assertEqual(by_label["等效上网小时数（考核值，h）"]["value"], "2917")
        self.assertEqual(by_label["等效上网小时数（保证值，h）"]["value"], "2841")
        self.assertEqual(by_label["年等效满负荷小时数（保证值，h）"]["value"], "2841")
        self.assertEqual(by_label["投标总容量"]["value"], "600")

    def test_rep_year_and_config_tables(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = self._write_wind_docx(Path(tmp) / "风资源评估报告.docx")
            facts = facts_from_wind_resource_docx(path, _material(path.name), {})
        by_label = _facts_by_label(facts)
        # 代表年风速表（label 经 canonical 归一为"年平均风速"）
        self.assertEqual(by_label["年平均风速"]["value"], "7.36/6.86")
        self.assertEqual(by_label["测风塔参考高度"]["value"], "125")
        self.assertIn("26304#", by_label["指定测风塔/测风条件"]["value"])
        # 配置表（轮毂高度 canonical 归一；安全等级规范化为 "IEC S"）
        self.assertEqual(by_label["轮毂高度"]["value"], "125")
        self.assertEqual(by_label["安全等级"]["value"], "IEC S")
        # 叶尖高度 = 轮毂 125 + 叶轮 220/2（机型串解析）
        self.assertEqual(by_label["投标叶尖高度"]["value"], "235")


class TestCertificateExtractor(unittest.TestCase):
    def test_parse_certificate_wind_params(self) -> None:
        text = (
            "型式认证证书：参考风速 50 m/s，湍流强度期望值 0.15，"
            "50年一遇极端风速 59.5 m/s，安全等级 IEC S，满足 IEC 61400-22 要求。"
        )
        params = parse_certificate_wind_params(text)
        self.assertEqual(params.get("vref"), "50")
        self.assertEqual(params.get("turbulence"), "0.15")
        self.assertEqual(params.get("extreme50"), "59.5")
        self.assertEqual(params.get("iecClass"), "IEC S")
        self.assertEqual(parse_certificate_wind_params(""), {})
        self.assertEqual(parse_certificate_wind_params("供招投标使用"), {})

    def _write_cert_docx(self, path: Path, text: str) -> Path:
        return _write_docx(path, lambda doc: doc.add_paragraph(text))

    def test_type_cert_priority_over_design_cert(self) -> None:
        filler = "认证证书正文。" * 30  # 凑过文本层最小长度
        with tempfile.TemporaryDirectory() as tmp:
            type_path = self._write_cert_docx(
                Path(tmp) / "EW6.7-220-125型式认证B.docx",
                f"型式认证证书 参考风速 42.5 m/s，湍流强度 0.14，安全等级 IEC IA。{filler}",
            )
            design_path = self._write_cert_docx(
                Path(tmp) / "EW6.7-220-125设计认证A.docx",
                f"设计认证证书 参考风速 50 m/s，湍流强度 0.15，安全等级 IEC S。{filler}",
            )
            materials = [
                (_material(design_path.name), design_path),
                (_material(type_path.name), type_path),
            ]
            facts = facts_from_certificate_materials(materials, {})
        by_label = _facts_by_label(facts)
        # 型式认证优先：同字段取值式认证的值（湍流强度 label 经 canonical 归一）
        self.assertEqual(by_label["机型认证10分钟平均极限风速（m/s）"]["value"], "42.5")
        self.assertEqual(by_label["湍流强度"]["value"], "0.14")
        self.assertIn("型式认证", by_label["机型认证10分钟平均极限风速（m/s）"]["sourceRef"]["name"])
        # 设计认证核心风资源参数：摘要串
        summary = by_label["设计认证核心风资源参数"]["value"]
        self.assertIn("参考风速42.5m/s", summary)
        self.assertIn("湍流强度0.14", summary)

    def test_scanned_cert_yields_nothing(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = self._write_cert_docx(Path(tmp) / "设计认证D.docx", "供招投标使用")
            facts = facts_from_certificate_materials([(_material(path.name), path)], {})
        self.assertEqual(facts, [])


class TestHoursCommitmentAndProductionBase(unittest.TestCase):
    def test_commitment_version_label(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "发电小时数承诺函（承诺保证值）.docx"
            _write_docx(path, lambda doc: doc.add_paragraph("电量承诺书（承诺保证值）"))
            facts = facts_from_hours_commitment_docx(path, _material(path.name), {})
        by_label = _facts_by_label(facts)
        # spec label canonical 坍缩，产出别名"电量承诺函版本"（SPEC_LABEL_ALIASES 归位）
        self.assertEqual(by_label["电量承诺函版本"]["value"], "承诺保证值")
        # reconcile 能经别名归位到 spec seq=7
        fields_by_key = {
            fact_label_key(fact["label"]): {"label": fact["label"], "value": fact["value"], "status": "extracted"}
            for fact in facts
        }
        reconcile_fact_fields_with_specs(fields_by_key)
        field = fields_by_key[fact_label_key("电量承诺函版本")]
        spec = next(s for s in fillable_specs() if s.get("label") == "发电小时数/电量承诺函版本")
        self.assertEqual(field.get("specKey"), spec.get("key"))
        self.assertEqual(field.get("specSeq"), 7)

    def test_production_base_name_and_intro(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "本项目叶片供货制造基地.docx"

            def build(document: Document) -> None:
                document.add_paragraph("本项目叶片供货制造基地-中材叶片连云港基地")
                document.add_paragraph("生产基地概况")
                document.add_paragraph("中复连众（连云港）风电叶片有限公司成立于2023年6月30日。")

            _write_docx(path, build)
            facts = facts_from_production_base_docx(path, _material(path.name), {})
        by_label = _facts_by_label(facts)
        self.assertEqual(by_label["叶片供货制造基地名称"]["value"], "中材叶片连云港基地")
        self.assertIn("中复连众", by_label["叶片供货制造基地介绍"]["value"])
        self.assertNotIn("主机供货制造基地名称", by_label)


class TestSpecialExtractorRouter(unittest.TestCase):
    def test_router_keyword_mapping(self) -> None:
        cases = [
            ("风资源评估报告.docx", "项目定制/风资源评估报告", "wind_resource"),
            ("塔架与基础工程量.docx", "项目定制/塔架与基础工程量", "tower_quantity"),
            ("基础弯矩表.xlsx", "项目定制/基础弯矩表", "foundation_moment"),
            ("EW6.7-220-125型式认证B.pdf", "标准文件/EW6.7-220/认证证书", "certificate"),
            ("EW10.0-220设计认证D.pdf", "标准文件/EW10.0-220/认证证书", "certificate"),
            ("发电小时数承诺函（承诺保证值）.docx", "项目定制/发电小时数承诺函", "hours_commitment"),
            ("生产制造基地专题_锡盟基地.docx", "项目定制/项目生产制造基地专题", "production_base"),
            ("本项目叶片供货制造基地.docx", "项目定制/项目生产制造基地专题", "production_base"),
            ("采购部短名单.xlsx", "项目定制/短名单", None),
            ("低电压穿越评估证书.pdf", "标准文件/EW6.7-220/认证证书", None),
        ]
        for name, folder, expected in cases:
            with self.subTest(name=name):
                self.assertEqual(
                    special_extractor_for_material({"name": name, "folderPath": folder}), expected
                )

    def test_material_is_fact_relevant_new_keywords(self) -> None:
        for name in [
            "基础弯矩表.xlsx",
            "EW6.7-220-125型式认证B.pdf",
            "发电小时数承诺函（承诺考核值）.docx",
            "生产制造基地专题_锡盟基地.docx",
        ]:
            with self.subTest(name=name):
                self.assertTrue(material_is_fact_relevant({"name": name, "materialTier": "standard"}))
        self.assertFalse(material_is_fact_relevant({"name": "采购部短名单.xlsx", "materialTier": "standard"}))


if __name__ == "__main__":
    unittest.main()
