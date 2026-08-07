from __future__ import annotations

import importlib.util
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


WORD_CLEANER_PATH = (
    Path(__file__).resolve().parents[1]
    / "opencode"
    / "skills"
    / "bid-material-format-cleaner"
    / "scripts"
    / "word_cleaner.py"
)


def _load_word_cleaner():
    module_name = "test_bid_material_format_cleaner_word_cleaner"
    spec = importlib.util.spec_from_file_location(module_name, WORD_CLEANER_PATH)
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    return module


class MaterialCleanerBlankHeadingTests(unittest.TestCase):
    def test_normalize_preserves_trailing_blank_heading_with_line_break(self) -> None:
        word_cleaner = _load_word_cleaner()

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "trailing-blank-heading.docx"
            document = Document()
            document.add_paragraph("正文")
            blank_heading = document.add_paragraph(style="Heading 2")
            blank_heading.add_run().add_break(WD_BREAK.LINE)
            document.save(path)

            word_cleaner.cmd_normalize(path)
            word_cleaner.cmd_normalize(path)

            cleaned = Document(path)
            line_break_paragraphs = [
                paragraph
                for paragraph in cleaned.paragraphs
                if any(
                    br.get(qn("w:type"), "") != "page"
                    for br in paragraph._p.iter(qn("w:br"))
                )
            ]

        self.assertEqual(len(cleaned.paragraphs), 2)
        self.assertEqual(len(line_break_paragraphs), 1)
        self.assertEqual(line_break_paragraphs[0].style.name, "Normal")

    def test_normalize_clears_inherited_heading_metadata_without_removing_page_break(self) -> None:
        word_cleaner = _load_word_cleaner()

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "blank-heading.docx"
            document = Document()
            custom_style = document.styles.add_style("空白标题样式", WD_STYLE_TYPE.PARAGRAPH)
            custom_style.base_style = document.styles["Heading 2"]
            document.add_paragraph("前文")
            blank_heading = document.add_paragraph(style=custom_style)
            blank_heading.add_run().add_break(WD_BREAK.PAGE)
            document.add_paragraph("后文")
            document.save(path)

            word_cleaner.cmd_normalize(path)

            cleaned = Document(path)
            page_break_paragraphs = [
                paragraph
                for paragraph in cleaned.paragraphs
                if any(
                    br.get(qn("w:type"), "") == "page"
                    for br in paragraph._p.iter(qn("w:br"))
                )
            ]

        self.assertEqual(len(page_break_paragraphs), 1)
        paragraph = page_break_paragraphs[0]
        self.assertEqual(paragraph.style.name, "Normal")
        p_pr = paragraph._p.find(qn("w:pPr"))
        self.assertIsNotNone(p_pr)
        self.assertIsNone(p_pr.find(qn("w:outlineLvl")))

    def test_normalize_clears_direct_outline_level_without_removing_line_break(self) -> None:
        word_cleaner = _load_word_cleaner()

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "blank-outline.docx"
            document = Document()
            document.add_paragraph("前文")
            blank_outline = document.add_paragraph()
            blank_outline.add_run().add_break(WD_BREAK.LINE)
            outline = OxmlElement("w:outlineLvl")
            outline.set(qn("w:val"), "3")
            blank_outline._p.get_or_add_pPr().append(outline)
            document.add_paragraph("后文")
            document.save(path)

            word_cleaner.cmd_normalize(path)

            cleaned = Document(path)
            line_break_paragraphs = [
                paragraph
                for paragraph in cleaned.paragraphs
                if any(
                    br.get(qn("w:type"), "") != "page"
                    for br in paragraph._p.iter(qn("w:br"))
                )
            ]

        self.assertEqual(len(line_break_paragraphs), 1)
        paragraph = line_break_paragraphs[0]
        self.assertEqual(paragraph.style.name, "Normal")
        p_pr = paragraph._p.find(qn("w:pPr"))
        self.assertIsNotNone(p_pr)
        self.assertIsNone(p_pr.find(qn("w:outlineLvl")))

    def test_normalize_preserves_trailing_blank_heading_with_carriage_return(self) -> None:
        word_cleaner = _load_word_cleaner()

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "trailing-blank-heading-w-cr.docx"
            document = Document()
            document.add_paragraph("正文")
            blank_heading = document.add_paragraph(style="Heading 2")
            blank_heading.add_run()._r.append(OxmlElement("w:cr"))
            document.save(path)

            word_cleaner.cmd_normalize(path)
            word_cleaner.cmd_normalize(path)

            cleaned = Document(path)
            carriage_return_paragraphs = [
                paragraph
                for paragraph in cleaned.paragraphs
                if next(paragraph._p.iter(qn("w:cr")), None) is not None
            ]

        self.assertEqual(len(cleaned.paragraphs), 2)
        self.assertEqual(len(carriage_return_paragraphs), 1)
        self.assertEqual(carriage_return_paragraphs[0].style.name, "Normal")

    def test_normalize_preserves_carriage_returns_in_body_and_table_cell(self) -> None:
        word_cleaner = _load_word_cleaner()

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "body-and-table-w-cr.docx"
            document = Document()
            body_paragraph = document.add_paragraph()
            body_paragraph.add_run("正文上")
            body_paragraph.add_run()._r.append(OxmlElement("w:cr"))
            body_paragraph.add_run("正文下")

            cell_paragraph = document.add_table(rows=1, cols=1).cell(0, 0).paragraphs[0]
            cell_paragraph.add_run("单元格上")
            cell_paragraph.add_run()._r.append(OxmlElement("w:cr"))
            cell_paragraph.add_run("单元格下")
            document.save(path)

            word_cleaner.cmd_normalize(path)

            cleaned = Document(path)
            cleaned_body = cleaned.paragraphs[0]
            cleaned_cell = cleaned.tables[0].cell(0, 0).paragraphs[0]

        self.assertEqual(len(list(cleaned_body._p.iter(qn("w:cr")))), 1)
        self.assertEqual(len(list(cleaned_cell._p.iter(qn("w:cr")))), 1)

    def test_normalize_collapses_blank_headings_without_manual_breaks(self) -> None:
        word_cleaner = _load_word_cleaner()

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "blank-heading-run.docx"
            document = Document()
            document.add_paragraph("before")
            for _ in range(30):
                document.add_paragraph(style="Heading 2")
            document.add_paragraph("after")
            document.save(path)

            word_cleaner.cmd_normalize(path)

            cleaned = Document(path)
            preserved_blank_headings = [
                bookmark
                for paragraph in cleaned.paragraphs
                for bookmark in paragraph._p.iter(qn("w:bookmarkStart"))
                if str(bookmark.get(qn("w:name")) or "").startswith(
                    word_cleaner._BLANK_HEADING_BOOKMARK_PREFIX
                )
            ]
            page_break_paragraphs = [
                paragraph
                for paragraph in cleaned.paragraphs
                if any(
                    br.get(qn("w:type"), "") == "page"
                    for br in paragraph._p.iter(qn("w:br"))
                )
            ]

        self.assertEqual(preserved_blank_headings, [])
        self.assertEqual(len(page_break_paragraphs), 1)
        self.assertEqual(len(cleaned.paragraphs), 3)

    def test_normalize_removes_legacy_blank_heading_markers_without_breaks(self) -> None:
        word_cleaner = _load_word_cleaner()

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "legacy-blank-heading-run.docx"
            document = Document()
            document.add_paragraph("before")
            for _ in range(30):
                paragraph = document.add_paragraph()
                word_cleaner._mark_preserved_blank_heading(paragraph._p)
            document.add_paragraph("after")
            document.save(path)

            word_cleaner.cmd_normalize(path)

            cleaned = Document(path)
            preserved_blank_headings = [
                bookmark
                for paragraph in cleaned.paragraphs
                for bookmark in paragraph._p.iter(qn("w:bookmarkStart"))
                if str(bookmark.get(qn("w:name")) or "").startswith(
                    word_cleaner._BLANK_HEADING_BOOKMARK_PREFIX
                )
            ]

        self.assertEqual(preserved_blank_headings, [])
        self.assertEqual(len(cleaned.paragraphs), 3)

    def test_numbered_outline_only_heading_preserves_normal_style(self) -> None:
        word_cleaner = _load_word_cleaner()

        document = Document()
        heading = document.add_paragraph("1.7 项目业绩")
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), "3")
        heading._p.get_or_add_pPr().append(outline)

        normalized = word_cleaner._strip_numbered_heading_prefixes(document)

        self.assertEqual(normalized, 1)
        self.assertEqual(heading.text, "项目业绩")
        self.assertEqual(heading.style.name, "Normal")
        preserved_outline = heading._p.get_or_add_pPr().find(qn("w:outlineLvl"))
        self.assertIsNotNone(preserved_outline)
        self.assertEqual(preserved_outline.get(qn("w:val")), "3")


if __name__ == "__main__":
    unittest.main()
