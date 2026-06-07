from __future__ import annotations

import unittest
import zipfile
from io import BytesIO
from unittest.mock import AsyncMock, patch

from docx import Document
from docx.shared import Inches

from app.api.routes import business as business_routes
from app.models.materials import RawFile
from app.services.business_material_store import business_material_store
from app.services.business_material_splitter import confirm_business_material_split, preview_business_material_split, _send_openai_compatible_prompt
from app.services.peripheral import PeripheralError


def _sample_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("第六章 投标文件格式", level=1)
    doc.add_paragraph("附件1 投标函")
    doc.add_paragraph("致：招标人。我方愿意参加本项目投标。")
    doc.add_paragraph("附件2 法定代表人授权书")
    doc.add_paragraph("兹授权以下人员代表我方办理投标事宜。")
    doc.add_paragraph("EW10.0-220 设计认证证书")
    doc.add_paragraph("证书编号：CERT-001。有效期至2028年12月31日。")
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _order_summary_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("订单汇总表", level=1)
    table = doc.add_table(rows=3, cols=6)
    headers = ["序号", "项目名称", "合同编号", "客户名称", "机型", "签订日期"]
    row1 = ["1", "华能某风电项目", "HT-2024-001", "华能集团", "EW10.0-220", "2024年6月"]
    row2 = ["2", "国家能源某风电项目", "HT-2025-002", "国家能源集团", "EW8.5-230", "2025年3月"]
    for col, value in enumerate(headers):
        table.cell(0, col).text = value
    for col, value in enumerate(row1):
        table.cell(1, col).text = value
    for col, value in enumerate(row2):
        table.cell(2, col).text = value
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _plain_order_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("历史订单资料", level=1)
    doc.add_paragraph("华能山东某风电项目 HT-2024-001 上海电气 EW10.0-220 10台 2024年6月签订")
    doc.add_paragraph("本订单包含风机设备、塔筒接口及相关服务。")
    doc.add_paragraph("国家能源内蒙古某风电项目 HT-2025-002 上海电气 EW8.5-230 12台 2025年3月签订")
    doc.add_paragraph("本订单已完成交货并通过阶段验收。")
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _single_paragraph_many_orders_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("订单汇总资料", level=1)
    doc.add_paragraph(
        "华能山东某风电项目 HT-2024-001 上海电气 EW10.0-220 10台 2024年6月签订；"
        "国家能源内蒙古某风电项目 HT-2025-002 上海电气 EW8.5-230 12台 2025年3月签订；"
        "大唐江苏某风电项目 HT-2025-003 上海电气 EW9.0-220 8台 2025年5月签订。"
    )
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _docx_with_embedded_image_bytes() -> bytes:
    image_bytes = bytes.fromhex(
        "89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c4890000000d49444154789c63fccfc0500f000583027f96ecc3e20000000049454e44ae426082"
    )
    image_stream = BytesIO(image_bytes)
    doc = Document()
    doc.add_paragraph("EW10.0-220 设计认证证书")
    doc.add_picture(image_stream, width=Inches(1))
    doc.add_paragraph("证书编号：CERT-IMG-001。有效期至2028年12月31日。")
    doc.add_paragraph("附件1 投标函")
    doc.add_paragraph("致：招标人。我方愿意参加本项目投标。")
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _disable_ai_enhance_patch():
    return patch("app.services.business_material_splitter._should_ai_enhance_split", return_value=False)


def _long_single_paragraph_many_orders_docx_bytes(count: int = 12) -> bytes:
    doc = Document()
    doc.add_heading("历史订单台账", level=1)
    chunks = []
    for index in range(1, count + 1):
        chunks.append(
            f"华能{index:02d}号风电项目 HT-2025-{index:03d} 上海电气 EW{6 + index % 5}.0-220 "
            f"{index + 3}台 2025年{index:02d}月签订"
        )
    doc.add_paragraph("；".join(chunks))
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


class BusinessMaterialSplitterTests(unittest.IsolatedAsyncioTestCase):
    def test_raw_file_payload_exposes_split_traceability_fields(self) -> None:
        raw = RawFile(
            id=7,
            name="投标函.docx",
            size_bytes=1024,
            minio_key="raw/test.docx",
            ext_fields={
                "bidType": "商务标",
                "splitParentMaterialId": "RAW-0001",
                "splitParentFileName": "商务附件合集.docx",
                "splitFragmentId": "frag-001",
                "splitFragmentTitle": "投标函",
                "splitMaterialType": "商务附件模板",
                "splitConfidence": 0.9,
                "splitRiskTips": ["需人工复核。"],
                "reviewStatus": "pending_review",
            },
        )

        payload = raw.to_dict()

        self.assertEqual(payload["splitParentMaterialId"], "RAW-0001")
        self.assertEqual(payload["splitFragmentTitle"], "投标函")
        self.assertEqual(payload["splitMaterialType"], "商务附件模板")
        self.assertEqual(payload["splitConfidence"], 0.9)
        self.assertEqual(payload["reviewStatus"], "pending_review")

    async def test_preview_business_material_split_builds_docx_fragments(self) -> None:
        with patch(
            "app.services.business_material_splitter._load_business_source_file",
            new=AsyncMock(
                return_value={
                    "id": "RAW-0001",
                    "name": "商务附件合集.docx",
                    "folderPath": "商务标/通用素材/通用模板底稿库",
                    "materialTier": "standard",
                    "content": _sample_docx_bytes(),
                }
            ),
        ), _disable_ai_enhance_patch():
            plan = await preview_business_material_split("RAW-0001")

        titles = [item["title"] for item in plan["fragments"]]
        self.assertIn("投标函", titles)
        self.assertIn("法定代表人授权书", titles)
        self.assertTrue(any("设计认证" in title for title in titles))
        self.assertTrue(any(item["materialType"] == "证书" for item in plan["fragments"]))
        self.assertFalse(any(title.startswith("证书编号") for title in titles))

    async def test_preview_business_material_split_builds_order_row_fragments(self) -> None:
        with patch(
            "app.services.business_material_splitter._load_business_source_file",
            new=AsyncMock(
                return_value={
                    "id": "RAW-0002",
                    "name": "订单汇总.docx",
                    "folderPath": "商务标/通用素材/企业能力与综合实力",
                    "materialTier": "standard",
                    "content": _order_summary_docx_bytes(),
                }
            ),
        ), _disable_ai_enhance_patch():
            plan = await preview_business_material_split("RAW-0002")

        order_fragments = [item for item in plan["fragments"] if item["materialType"] == "业绩订单" and (item.get("sourceLocation") or {}).get("mode") == "tableRow"]
        self.assertEqual(len(order_fragments), 2)
        self.assertTrue(any("华能某风电项目" in item["title"] for item in order_fragments))
        self.assertTrue(any("国家能源某风电项目" in item["title"] for item in order_fragments))
        self.assertTrue(all(item["suggestedPath"] == "" for item in order_fragments))
        self.assertTrue(all(item["selected"] is False for item in order_fragments))
        self.assertTrue(all(any("业绩库导入" in tip for tip in item["riskTips"]) for item in order_fragments))

    async def test_preview_business_material_split_falls_back_to_local_semantic_fragments(self) -> None:
        with patch(
            "app.services.business_material_splitter._load_business_source_file",
            new=AsyncMock(
                return_value={
                    "id": "RAW-0003",
                    "name": "历史订单资料.docx",
                    "folderPath": "商务标/通用素材/企业能力与综合实力",
                    "materialTier": "standard",
                    "content": _plain_order_docx_bytes(),
                }
            ),
        ), _disable_ai_enhance_patch():
            plan = await preview_business_material_split("RAW-0003")

        semantic_fragments = [item for item in plan["fragments"] if (item.get("sourceLocation") or {}).get("mode") == "localSemantic"]
        self.assertEqual(len(semantic_fragments), 2)
        self.assertTrue(any("华能山东" in item["title"] for item in semantic_fragments))
        self.assertTrue(any("国家能源内蒙古" in item["title"] for item in semantic_fragments))

    async def test_confirm_business_material_split_accepts_semantic_client_fragment(self) -> None:
        uploaded_calls = []

        async def fake_raw_upload(**kwargs):
            uploaded_calls.append(kwargs)
            return {
                "items": [
                    {
                        "id": "RAW-0101",
                        "name": kwargs["files"][0]["name"],
                        "splitParentMaterialId": kwargs["files"][0]["extFields"]["splitParentMaterialId"],
                    }
                ]
            }

        with patch(
            "app.services.business_material_splitter._load_business_source_file",
            new=AsyncMock(
                return_value={
                    "id": "RAW-0004",
                    "name": "普通文本资料.docx",
                    "folderPath": "商务标/通用素材/企业能力与综合实力",
                    "materialTier": "standard",
                    "content": _plain_order_docx_bytes(),
                }
            ),
        ), patch(
            "app.services.business_material_splitter._upload_business_split_files",
            side_effect=fake_raw_upload,
        ):
            with self.assertRaises(PeripheralError) as context:
                await confirm_business_material_split(
                    "RAW-0004",
                    fragments=[
                        {
                            "id": "frag-ai-001",
                            "selected": True,
                            "title": "AI识别订单片段",
                            "materialType": "业绩订单",
                            "fileName": "AI识别订单片段.docx",
                            "targetPath": "",
                            "sourceLocation": {"mode": "aiSemantic", "blockStart": 1, "blockEnd": 3},
                        }
                    ],
                )

        self.assertEqual(context.exception.code, "BUSINESS_SPLIT_PERFORMANCE_REQUIRES_LIBRARY")
        self.assertEqual(uploaded_calls, [])

    async def test_preview_business_material_split_can_use_ai_quote_fragments_inside_one_block(self) -> None:
        fake_reply = {
            "reply": (
                '{"fragments": ['
                '{"title": "华能山东某风电项目", "materialType": "业绩订单", "blockStart": 1, "blockEnd": 2, "quote": "华能山东某风电项目 HT-2024-001 上海电气 EW10.0-220 10台 2024年6月签订", "reason": "同一段落内的独立订单"},'
                '{"title": "国家能源内蒙古某风电项目", "materialType": "业绩订单", "blockStart": 1, "blockEnd": 2, "quote": "国家能源内蒙古某风电项目 HT-2025-002 上海电气 EW8.5-230 12台 2025年3月签订", "reason": "同一段落内的独立订单"},'
                '{"title": "大唐江苏某风电项目", "materialType": "业绩订单", "blockStart": 1, "blockEnd": 2, "quote": "大唐江苏某风电项目 HT-2025-003 上海电气 EW9.0-220 8台 2025年5月签订", "reason": "同一段落内的独立订单"}'
                ']}'
            )
        }

        with patch(
            "app.services.business_material_splitter._load_business_source_file",
            new=AsyncMock(
                return_value={
                    "id": "RAW-0005",
                    "name": "订单汇总资料.docx",
                    "folderPath": "商务标/通用素材/企业能力与综合实力",
                    "materialTier": "standard",
                    "content": _single_paragraph_many_orders_docx_bytes(),
                }
            ),
        ), patch("app.services.business_material_splitter._send_business_split_ai_prompt", return_value=fake_reply):
            plan = await preview_business_material_split("RAW-0005", ai_mode="force")

        fragments = plan["fragments"]
        self.assertEqual(len(fragments), 3)
        self.assertEqual(plan["diagnostics"]["strategy"], "ai_enhanced")
        self.assertTrue(all((item.get("sourceLocation") or {}).get("quote") for item in fragments))

    async def test_preview_business_material_split_chunks_long_blocks_for_ai(self) -> None:
        fake_reply = {
            "reply": (
                '{"fragments": ['
                '{"title": "华能01号风电项目", "materialType": "业绩订单", "blockStart": "1#1/2", "blockEnd": "2", "quote": "华能01号风电项目 HT-2025-001 上海电气 EW7.0-220 4台 2025年01月签订"},'
                '{"title": "华能02号风电项目", "materialType": "业绩订单", "blockStart": "1#1/2", "blockEnd": "2", "quote": "华能02号风电项目 HT-2025-002 上海电气 EW8.0-220 5台 2025年02月签订"},'
                '{"title": "华能21号风电项目", "materialType": "业绩订单", "blockStart": "1#2/2", "blockEnd": "2", "quote": "华能21号风电项目 HT-2025-021 上海电气 EW7.0-220 24台 2025年21月签订"},'
                '{"title": "华能22号风电项目", "materialType": "业绩订单", "blockStart": "1#2/2", "blockEnd": "2", "quote": "华能22号风电项目 HT-2025-022 上海电气 EW8.0-220 25台 2025年22月签订"}'
                ']}'
            )
        }

        with patch(
            "app.services.business_material_splitter._load_business_source_file",
            new=AsyncMock(
                return_value={
                    "id": "RAW-0007",
                    "name": "历史订单台账.docx",
                    "folderPath": "商务标/通用素材/企业能力与综合实力",
                    "materialTier": "standard",
                    "content": _long_single_paragraph_many_orders_docx_bytes(count=24),
                }
            ),
        ), patch("app.services.business_material_splitter._send_business_split_ai_prompt", return_value=fake_reply) as send_ai:
            plan = await preview_business_material_split("RAW-0007", ai_mode="force")

        prompt = send_ai.call_args.args[0]
        self.assertIn("[1#1/", prompt)
        self.assertIn("[1#2/", prompt)
        self.assertEqual(len(plan["fragments"]), 4)
        self.assertEqual({(item["sourceLocation"] or {})["blockStart"] for item in plan["fragments"]}, {1})
        self.assertTrue(all((item.get("sourceLocation") or {}).get("quote") for item in plan["fragments"]))

    async def test_confirm_business_material_split_renders_quote_fragment(self) -> None:
        uploaded_calls = []

        async def fake_raw_upload(**kwargs):
            uploaded_calls.append(kwargs)
            return {
                "items": [
                    {
                        "id": "RAW-0102",
                        "name": kwargs["files"][0]["name"],
                        "splitParentMaterialId": kwargs["files"][0]["extFields"]["splitParentMaterialId"],
                    }
                ]
            }

        quote = "华能山东某风电项目 HT-2024-001 上海电气 EW10.0-220 10台 2024年6月签订"
        with patch(
            "app.services.business_material_splitter._load_business_source_file",
            new=AsyncMock(
                return_value={
                    "id": "RAW-0006",
                    "name": "订单汇总资料.docx",
                    "folderPath": "商务标/通用素材/企业能力与综合实力",
                    "materialTier": "standard",
                    "content": _single_paragraph_many_orders_docx_bytes(),
                }
            ),
        ), patch(
            "app.services.business_material_splitter._upload_business_split_files",
            side_effect=fake_raw_upload,
        ):
            with self.assertRaises(PeripheralError) as context:
                await confirm_business_material_split(
                    "RAW-0006",
                    fragments=[
                        {
                            "id": "frag-ai-001",
                            "selected": True,
                            "title": "华能山东某风电项目",
                            "materialType": "业绩订单",
                            "fileName": "华能山东某风电项目.docx",
                            "targetPath": "",
                            "sourceLocation": {"mode": "aiSemantic", "blockStart": 1, "blockEnd": 2, "quote": quote},
                        }
                    ],
                )

        self.assertEqual(context.exception.code, "BUSINESS_SPLIT_PERFORMANCE_REQUIRES_LIBRARY")
        self.assertEqual(uploaded_calls, [])

    async def test_confirm_business_material_split_uploads_selected_fragments_with_metadata(self) -> None:
        uploaded_calls = []

        async def fake_raw_upload(**kwargs):
            uploaded_calls.append(kwargs)
            return {
                "items": [
                    {
                        "id": "RAW-0100",
                        "name": kwargs["files"][0]["name"],
                        "splitParentMaterialId": kwargs["files"][0]["extFields"]["splitParentMaterialId"],
                    }
                ]
            }

        with patch(
            "app.services.business_material_splitter._load_business_source_file",
            new=AsyncMock(
                return_value={
                    "id": "RAW-0001",
                    "name": "商务附件合集.docx",
                    "folderPath": "商务标/通用素材/通用模板底稿库",
                    "materialTier": "standard",
                    "content": _sample_docx_bytes(),
                }
            ),
        ), patch(
            "app.services.business_material_splitter._upload_business_split_files",
            side_effect=fake_raw_upload,
        ):
            result = await confirm_business_material_split(
                "RAW-0001",
                fragments=[
                    {
                        "id": "frag-001",
                        "selected": True,
                        "title": "投标函",
                        "fileName": "投标函.docx",
                        "targetPath": "商务标/通用素材/通用模板底稿库",
                    }
                ],
            )

        self.assertEqual(len(uploaded_calls), 1)
        self.assertEqual(uploaded_calls[0]["target_path"], "商务标/通用素材/通用模板底稿库")
        self.assertNotIn("bid_type", uploaded_calls[0])
        self.assertEqual(uploaded_calls[0]["files"][0]["extFields"]["splitParentMaterialId"], "RAW-0001")
        self.assertEqual(result["items"][0]["splitParentMaterialId"], "RAW-0001")

    async def test_confirm_business_material_split_preserves_embedded_images(self) -> None:
        uploaded_calls = []

        async def fake_raw_upload(**kwargs):
            uploaded_calls.append(kwargs)
            return {"items": [{"id": "RAW-0200", "name": kwargs["files"][0]["name"]}]}

        with patch(
            "app.services.business_material_splitter._load_business_source_file",
            new=AsyncMock(
                return_value={
                    "id": "RAW-0008",
                    "name": "证书合集.docx",
                    "folderPath": "商务标/通用素材/专题证书库/机型认证证书",
                    "materialTier": "standard",
                    "content": _docx_with_embedded_image_bytes(),
                }
            ),
        ), patch(
            "app.services.business_material_splitter._upload_business_split_files",
            side_effect=fake_raw_upload,
        ):
            await confirm_business_material_split(
                "RAW-0008",
                fragments=[
                    {
                        "id": "frag-001",
                        "selected": True,
                        "title": "EW10.0-220 设计认证证书",
                        "fileName": "EW10.0-220 设计认证证书.docx",
                        "targetPath": "商务标/通用素材/专题证书库/机型认证证书",
                    }
                ],
            )

        docx_bytes = uploaded_calls[0]["files"][0]["data"]
        output_doc = Document(BytesIO(docx_bytes))
        with zipfile.ZipFile(BytesIO(docx_bytes), "r") as zf:
            media_files = [name for name in zf.namelist() if name.startswith("word/media/")]
        self.assertGreaterEqual(len(output_doc.part._package.image_parts), 1)
        self.assertGreaterEqual(len(media_files), 1)

    async def test_business_split_confirm_route_accepts_frontend_fragments_payload(self) -> None:
        fragment = {
            "id": "frag-001",
            "selected": True,
            "title": "投标函",
            "targetPath": "商务标/通用素材/通用模板底稿库",
        }

        with patch.object(
            business_routes.business_material_store,
            "confirm_business_split",
            new=AsyncMock(return_value={"ok": True}),
        ) as confirm_business_split:
            result = await business_routes.business_raw_confirm_split(
                "RAW-0001",
                {
                    "fragments": [fragment],
                    "targetPath": "商务标/通用素材/通用模板底稿库",
                    "onConflict": "rename",
                },
            )

        self.assertEqual(result, {"ok": True})
        confirm_business_split.assert_awaited_once()
        self.assertEqual(confirm_business_split.call_args.kwargs["fragments"], [fragment])
        self.assertEqual(confirm_business_split.call_args.kwargs["target_path"], "商务标/通用素材/通用模板底稿库")
        self.assertEqual(confirm_business_split.call_args.kwargs["on_conflict"], "rename")

    async def test_business_material_store_forwards_split_contract_to_service(self) -> None:
        fragment = {
            "id": "frag-001",
            "selected": True,
            "title": "投标函",
            "targetPath": "商务标/通用素材/通用模板底稿库",
        }

        with patch.object(business_material_store, "ensure_raw_file", new=AsyncMock()) as ensure_raw_file, patch(
            "app.services.business_material_store.confirm_business_material_split",
            new=AsyncMock(return_value={"ok": True}),
        ) as confirm_material_split:
            result = await business_material_store.confirm_business_split(
                "RAW-0001",
                fragments=[fragment],
                target_path="商务标/通用素材/通用模板底稿库",
                on_conflict="rename",
            )

        self.assertEqual(result, {"ok": True})
        ensure_raw_file.assert_awaited_once_with("RAW-0001")
        confirm_material_split.assert_awaited_once()
        self.assertEqual(confirm_material_split.call_args.kwargs["fragments"], [fragment])
        self.assertEqual(
            confirm_material_split.call_args.kwargs["default_target_path"],
            "商务标/通用素材/通用模板底稿库",
        )
        self.assertEqual(confirm_material_split.call_args.kwargs["on_conflict"], "rename")

    def test_openai_compatible_prompt_clamps_extreme_max_tokens(self) -> None:
        captured = {}

        class FakeResponse:
            status_code = 200

            def json(self):
                return {"choices": [{"message": {"content": '{"fragments":[]}'}, "finish_reason": "stop"}]}

        class FakeClient:
            def __init__(self, *args, **kwargs):
                pass

            def __enter__(self):
                return self

            def __exit__(self, *args):
                return False

            def post(self, url, headers=None, json=None):
                captured["payload"] = json
                return FakeResponse()

        with patch("app.services.business_material_splitter.httpx.Client", FakeClient):
            result = _send_openai_compatible_prompt(
                {
                    "enabled": True,
                    "baseUrl": "https://api.example.com",
                    "apiKey": "test",
                    "model": "deepseek-v4-pro",
                    "maxTokens": 9999999999,
                    "timeoutMs": 1,
                },
                '请返回 {"fragments":[]}',
            )

        self.assertEqual(result["reply"], '{"fragments":[]}')
        self.assertEqual(captured["payload"]["max_tokens"], 8192)


if __name__ == "__main__":
    unittest.main()
