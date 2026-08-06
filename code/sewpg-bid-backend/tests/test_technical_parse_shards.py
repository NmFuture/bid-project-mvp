from __future__ import annotations

import json
import subprocess
import sys
import tempfile
import unittest
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

from docx import Document


SKILL_SCRIPTS = (
    Path(__file__).resolve().parents[1]
    / "opencode"
    / "skills"
    / "bid-tech-tender-structured-parser"
    / "scripts"
)

if str(SKILL_SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SKILL_SCRIPTS))

from agentic.checklist import (  # noqa: E402
    checklist_for_shard,
    load_checklist,
    load_shards,
    shard_keys,
    shard_of_row,
)
from agentic.paths import load_manifest  # noqa: E402
from agentic.submission_store import load as load_submissions  # noqa: E402
from agentic.submission_store import shard_progress, submit  # noqa: E402


RUNNER = SKILL_SCRIPTS / "run_from_manifest.py"


class ChecklistShardConfigTests(unittest.TestCase):
    def test_shards_cover_every_checklist_row_exactly_once(self) -> None:
        checklist_rows = {int(item["rowNo"]) for item in load_checklist()}
        covered: list[int] = []
        for shard in load_shards():
            covered.extend(int(row_no) for row_no in shard["rowNos"])
        self.assertEqual(len(covered), len(set(covered)), "分片行号出现重叠")
        self.assertEqual(set(covered), checklist_rows, "分片未完整覆盖清单行")
        self.assertEqual(len(covered), 58)

    def test_shard_row_counts_stay_balanced(self) -> None:
        # 长板分片决定并发总耗时；行数一旦失衡（比如某片被塞进 20 行），并发就白做了。
        counts = [len(shard["rowNos"]) for shard in load_shards()]
        self.assertLessEqual(max(counts), 12)
        self.assertGreaterEqual(min(counts), 6)

    def test_shard_of_row_matches_checklist_for_shard(self) -> None:
        for key in shard_keys():
            for row in checklist_for_shard(key):
                self.assertEqual(shard_of_row(int(row["rowNo"])), key)


class ShardSubmissionTests(unittest.TestCase):
    def setUp(self) -> None:
        self._tmp = tempfile.TemporaryDirectory()
        self.root = Path(self._tmp.name)
        self.manifest_path = self.root / "s1_parse_manifest.json"
        self.manifest_path.write_text(
            json.dumps(
                {
                    "projectId": "PRJ-SHARD-UNIT",
                    "bidType": "技术标",
                    "parseProfile": "technical",
                    "structuredResultPath": str(self.root / "s1_structured_result.json"),
                    "documents": [],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        self.manifest = load_manifest(self.manifest_path)

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def _rows_for(self, shard_key: str) -> list[dict[str, object]]:
        return [
            {
                "rowNo": int(row["rowNo"]),
                "status": "missing",
                "conclusion": f"{shard_key} 结论 {row['rowNo']}",
                "evidenceSummary": "",
                "evidenceIds": [],
            }
            for row in checklist_for_shard(shard_key)
        ]

    def test_shard_submissions_merge_instead_of_overwrite(self) -> None:
        first, second = shard_keys()[0], shard_keys()[1]
        submit(self.manifest_path, self.manifest, "technicalInterpretation", self._rows_for(first), shard=first)
        submit(self.manifest_path, self.manifest, "technicalInterpretation", self._rows_for(second), shard=second)

        rows = load_submissions(self.manifest_path, self.manifest)["targets"]["technicalInterpretation"]
        submitted = {int(row["rowNo"]) for row in rows}
        expected = {int(r["rowNo"]) for r in checklist_for_shard(first)} | {
            int(r["rowNo"]) for r in checklist_for_shard(second)
        }
        self.assertEqual(submitted, expected)

    def test_resubmitting_a_shard_replaces_only_its_own_rows(self) -> None:
        first, second = shard_keys()[0], shard_keys()[1]
        submit(self.manifest_path, self.manifest, "technicalInterpretation", self._rows_for(first), shard=first)
        submit(self.manifest_path, self.manifest, "technicalInterpretation", self._rows_for(second), shard=second)

        revised = self._rows_for(first)
        for row in revised:
            row["conclusion"] = "revised"
        submit(self.manifest_path, self.manifest, "technicalInterpretation", revised, shard=first)

        rows = load_submissions(self.manifest_path, self.manifest)["targets"]["technicalInterpretation"]
        by_row_no = {int(row["rowNo"]): row for row in rows}
        for row in checklist_for_shard(first):
            self.assertEqual(by_row_no[int(row["rowNo"])]["conclusion"], "revised")
        for row in checklist_for_shard(second):
            self.assertNotEqual(by_row_no[int(row["rowNo"])]["conclusion"], "revised")

    def test_shard_cannot_submit_rows_outside_its_range(self) -> None:
        first, second = shard_keys()[0], shard_keys()[1]
        foreign = self._rows_for(second)
        with self.assertRaises(RuntimeError) as ctx:
            submit(self.manifest_path, self.manifest, "technicalInterpretation", foreign, shard=first)
        self.assertIn("outside its range", str(ctx.exception))

    def test_unknown_shard_is_rejected(self) -> None:
        with self.assertRaises(RuntimeError):
            submit(self.manifest_path, self.manifest, "technicalInterpretation", [], shard="not-a-shard")

    def test_project_basics_does_not_accept_shard(self) -> None:
        with self.assertRaises(RuntimeError):
            submit(self.manifest_path, self.manifest, "projectBasics", [], shard=shard_keys()[0])

    def test_concurrent_shard_submissions_do_not_lose_rows(self) -> None:
        """并发分片提交必须靠文件锁串行化；没有锁时这里会丢行。"""
        keys = shard_keys()
        with ThreadPoolExecutor(max_workers=len(keys)) as pool:
            futures = [
                pool.submit(
                    submit,
                    self.manifest_path,
                    self.manifest,
                    "technicalInterpretation",
                    self._rows_for(key),
                    shard=key,
                )
                for key in keys
            ]
            for future in futures:
                future.result()

        rows = load_submissions(self.manifest_path, self.manifest)["targets"]["technicalInterpretation"]
        submitted = {int(row["rowNo"]) for row in rows}
        self.assertEqual(submitted, {int(item["rowNo"]) for item in load_checklist()})
        self.assertEqual(len(rows), 58)

        progress = shard_progress(self.manifest_path, self.manifest)
        self.assertEqual(progress["pendingShards"], [])
        self.assertEqual(sorted(progress["submittedShards"]), sorted(keys))

    def test_concurrent_shard_submissions_across_processes(self) -> None:
        """分片会话是独立进程（opencode bash），跨进程锁必须同样成立。"""
        keys = shard_keys()

        def run(shard_key: str) -> subprocess.CompletedProcess[str]:
            return subprocess.run(
                [
                    sys.executable,
                    str(RUNNER),
                    "submit",
                    str(self.manifest_path),
                    "technicalInterpretation",
                    json.dumps(self._rows_for(shard_key), ensure_ascii=False),
                    "--shard",
                    shard_key,
                ],
                capture_output=True,
                text=True,
                encoding="utf-8",
            )

        with ThreadPoolExecutor(max_workers=len(keys)) as pool:
            results = list(pool.map(run, keys))
        for result in results:
            self.assertEqual(result.returncode, 0, result.stderr)

        rows = load_submissions(self.manifest_path, self.manifest)["targets"]["technicalInterpretation"]
        self.assertEqual({int(row["rowNo"]) for row in rows}, {int(i["rowNo"]) for i in load_checklist()})


class ShardChecklistCommandTests(unittest.TestCase):
    def setUp(self) -> None:
        self._tmp = tempfile.TemporaryDirectory()
        self.root = Path(self._tmp.name)
        source_path = self.root / "technical_tender.docx"
        doc = Document()
        doc.add_paragraph("第一章 招标公告")
        doc.add_paragraph("本项目供货范围为整套风力发电机组及塔筒内所有必要设备，含主控柜、通讯电缆。")
        doc.add_paragraph("塔筒型式为常规钢塔，塔筒底部直径应满足运输条件。")
        doc.add_paragraph("箱变采用上置于塔筒内布置，环网柜含箱变保护测控。")
        doc.save(source_path)
        self.manifest_path = self.root / "s1_parse_manifest.json"
        self.manifest_path.write_text(
            json.dumps(
                {
                    "projectId": "PRJ-SHARD-CLI",
                    "bidType": "技术标",
                    "parseProfile": "technical",
                    "structuredResultPath": str(self.root / "s1_structured_result.json"),
                    "documents": [{"id": "DOC-1", "name": source_path.name, "sourcePath": str(source_path)}],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        self._run("prepare", str(self.manifest_path))

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def _run(self, *args: str) -> dict:
        result = subprocess.run(
            [sys.executable, str(RUNNER), *args],
            capture_output=True,
            text=True,
            encoding="utf-8",
        )
        self.assertEqual(result.returncode, 0, result.stderr)
        return json.loads(result.stdout)

    def test_checklist_command_returns_only_shard_rows_with_hints(self) -> None:
        shard_key = "selection_supply"
        payload = self._run("checklist", str(self.manifest_path), "--shard", shard_key)
        self.assertEqual(payload["shard"], shard_key)
        expected = [int(row["rowNo"]) for row in checklist_for_shard(shard_key)]
        self.assertEqual([int(row["rowNo"]) for row in payload["rows"]], expected)
        self.assertTrue(payload["hintsIncluded"])
        # 预检索必须真的命中素材，否则等于没省掉模型的试探性检索。
        self.assertTrue(any(row.get("hints") for row in payload["rows"]))

    def test_checklist_command_can_skip_hints(self) -> None:
        payload = self._run("checklist", str(self.manifest_path), "--shard", "grid_documents", "--no-hints")
        self.assertFalse(payload["hintsIncluded"])
        self.assertTrue(all("hints" not in row for row in payload["rows"]))

    def test_search_accepts_multiple_queries_in_one_call(self) -> None:
        payload = self._run("search", str(self.manifest_path), "供货范围", "塔筒", "环网柜", "--limit", "5")
        self.assertEqual(payload["queryCount"], 3)
        self.assertEqual([item["query"] for item in payload["results"]], ["供货范围", "塔筒", "环网柜"])
        self.assertGreater(payload["matchCount"], 0)

    def test_search_many_reuses_one_database_connection(self) -> None:
        from unittest.mock import patch

        from agentic import navigator_cli

        manifest = load_manifest(self.manifest_path)
        original_connect = navigator_cli._connect
        with patch.object(navigator_cli, "_connect", wraps=original_connect) as mocked_connect:
            payload = navigator_cli.search_many(
                self.manifest_path,
                manifest,
                ["供货范围", "塔筒", "环网柜"],
                limit=5,
            )

        self.assertEqual(payload["queryCount"], 3)
        self.assertEqual(mocked_connect.call_count, 1)

    def test_single_query_search_keeps_legacy_shape(self) -> None:
        payload = self._run("search", str(self.manifest_path), "供货范围", "--limit", "5")
        self.assertIn("matches", payload)
        self.assertNotIn("results", payload)


class ShardProgressAggregatorTests(unittest.TestCase):
    def _aggregator(self):
        from app.services.parsing import _ShardProgressAggregator

        events: list[dict] = []
        agg = _ShardProgressAggregator(["a", "b", "c", "d"], lambda _e, payload: events.append(payload))
        return agg, events

    def test_one_finished_shard_does_not_push_overall_to_full(self) -> None:
        """单个分片完成时整体进度不能跳到 100%，否则其余分片还在跑进度条已经满了。"""
        agg, events = self._aggregator()
        agg.on_finished("a")
        self.assertEqual(events[-1]["completedShards"], 1)
        self.assertEqual(events[-1]["shardProgress"], 25)
        self.assertEqual(events[-1]["status"], "running")

    def test_overall_progress_never_reports_completed_status(self) -> None:
        agg, events = self._aggregator()
        for key in ("a", "b", "c", "d"):
            agg.on_finished(key)
        # 全部完成也只到 99，收口由 opencode_finished 事件负责
        self.assertEqual(events[-1]["shardProgress"], 99)
        self.assertEqual(events[-1]["completedShards"], 4)
        self.assertTrue(all(item["status"] == "running" for item in events))

    def test_progress_is_monotonic_per_shard(self) -> None:
        """会话内进度可能回跳（trace 重置），整体进度不能因此倒退。"""
        agg, events = self._aggregator()
        agg.on_stream("a", {"parts": [1, 2, 3, 4, 5]})
        high = events[-1]["shardProgress"]
        agg.on_stream("a", {"parts": []})
        self.assertGreaterEqual(events[-1]["shardProgress"], high)


class ShardedOrchestrationTests(unittest.TestCase):
    """端到端验证后端编排：prepare(一次) → 并发分片 submit → finalize。"""

    def setUp(self) -> None:
        self._tmp = tempfile.TemporaryDirectory()
        self.root = Path(self._tmp.name)
        source_path = self.root / "technical_tender.docx"
        doc = Document()
        doc.add_paragraph("某某风电场风力发电机组采购项目")
        doc.add_paragraph("招标编号：PC-TEST-0001")
        doc.add_paragraph("招标人：某某新能源有限公司")
        doc.add_paragraph("投标文件递交截止时间：2026年05月06日10时00分")
        doc.add_paragraph("本项目供货范围为整套风力发电机组及塔筒内所有必要设备。")
        doc.add_paragraph("塔筒型式为常规钢塔；箱变上置于塔筒内。")
        doc.save(source_path)
        self.output_path = self.root / "s1_structured_result.json"
        self.manifest_path = self.root / "s1_parse_manifest.json"
        self.manifest_path.write_text(
            json.dumps(
                {
                    "projectId": "PRJ-SHARD-ORCH",
                    "bidType": "技术标",
                    "parseProfile": "technical",
                    "structuredResultPath": str(self.output_path),
                    "documents": [{"id": "DOC-1", "name": source_path.name, "sourcePath": str(source_path)}],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def _fake_client_class(
        self,
        failing_shards: set[str] | None = None,
        silent_shards: set[str] | None = None,
        partial_shards: set[str] | None = None,
    ):
        """模拟 opencode 会话：按分片提示词真实调用 s1parse CLI 提交。

        failing_shards 模拟会话报错；silent_shards 模拟会话正常结束但没调 submit。
        """
        import re
        import threading

        manifest_path = self.manifest_path
        failing = failing_shards or set()
        silent = silent_shards or set()
        partial = partial_shards or set()
        seen: list[str] = []
        seen_lock = threading.Lock()

        def cli(*args: str) -> None:
            result = subprocess.run(
                [sys.executable, str(RUNNER), *args],
                capture_output=True,
                text=True,
                encoding="utf-8",
            )
            if result.returncode != 0:
                raise RuntimeError(result.stderr)

        class FakeClient:
            def __init__(self, *_args, **_kwargs) -> None:
                pass

            def run_tender_parse_shard_with_trace(self, prompt: str, **_kwargs) -> dict:
                match = re.search(r"--shard (\S+)", prompt)
                key = match.group(1) if match else "projectBasics"
                with seen_lock:
                    seen.append(key)
                if key in failing:
                    raise RuntimeError(f"injected failure for {key}")
                if key in silent:
                    # 会话「正常结束」，但什么都没提交
                    return {"opencodeOutput": {"sessionId": f"ses-{key}", "status": "completed"}}
                if key == "projectBasics":
                    payload = [
                        {"key": field_key, "label": field_key, "status": "missing",
                         "value": "当前文件未提及，建议补充上传对应文件"}
                        for field_key in (
                            "projectName", "tenderNo", "projectUnit", "tenderer", "tenderAgency", "bidDeadline"
                        )
                    ]
                    if key in partial:
                        payload = payload[:1]
                    cli("submit", str(manifest_path), "projectBasics", json.dumps(payload, ensure_ascii=False))
                else:
                    rows = [
                        {"rowNo": int(row["rowNo"]), "status": "missing", "conclusion": f"{key} 结论",
                         "evidenceSummary": "", "evidenceIds": []}
                        for row in checklist_for_shard(key)
                    ]
                    if key in partial:
                        rows = rows[:1]
                    cli("submit", str(manifest_path), "technicalInterpretation",
                        json.dumps(rows, ensure_ascii=False), "--shard", key)
                return {"opencodeOutput": {"sessionId": f"ses-{key}", "status": "completed"}}

        return FakeClient, seen

    def _run(
        self,
        failing_shards: set[str] | None = None,
        silent_shards: set[str] | None = None,
        partial_shards: set[str] | None = None,
    ):
        from unittest.mock import patch

        from app.services import parsing
        from app.services.parse_profiles import TECHNICAL_PARSE_PROFILE

        fake_cls, seen = self._fake_client_class(failing_shards, silent_shards, partial_shards)
        local_result = {"items": [], "structured": {}}
        with patch.object(parsing, "OpencodeClient", fake_cls):
            resolved, message = parsing._run_technical_sharded_parse_skill(
                self.manifest_path,
                local_result=local_result,
                profile=TECHNICAL_PARSE_PROFILE,
            )
        return resolved, message, seen

    def test_all_shards_run_once_and_finalize_covers_full_checklist(self) -> None:
        resolved, message, seen = self._run()

        self.assertEqual(message, "")
        # 每个分片和 projectBasics 各跑一次，不重复、不遗漏
        self.assertEqual(sorted(seen), sorted(["projectBasics", *shard_keys()]))

        self.assertEqual(len(resolved["items"]), 58)
        workflow = resolved["structured"]["workflow"]
        self.assertEqual(workflow["mode"], "opencode-agentic-navigation-sharded")
        self.assertEqual(workflow["failedShards"], [])
        self.assertTrue(self.output_path.is_file())

        submissions = load_submissions(self.manifest_path, load_manifest(self.manifest_path))
        rows = submissions["targets"]["technicalInterpretation"]
        self.assertEqual({int(r["rowNo"]) for r in rows}, {int(i["rowNo"]) for i in load_checklist()})

    def test_model_config_is_loaded_once_and_shared_by_all_sessions(self) -> None:
        from unittest.mock import patch

        from app.services import parsing
        from app.services.parse_profiles import TECHNICAL_PARSE_PROFILE

        model_config = {
            "enabled": True,
            "baseUrl": "https://llm.example.com/v1",
            "modelId": "test-model",
        }
        fake_cls, seen = self._fake_client_class()
        initialized_configs: list[dict | None] = []
        original_init = fake_cls.__init__

        def capture_init(client, *args, **kwargs) -> None:
            initialized_configs.append(kwargs.get("model_config"))
            original_init(client, *args, **kwargs)

        fake_cls.__init__ = capture_init
        with patch.object(parsing, "OpencodeClient", fake_cls), patch.object(
            parsing.system_settings_service,
            "get_opencode_model_config_sync",
            return_value=model_config,
        ) as load_model_config:
            parsing._run_technical_sharded_parse_skill(
                self.manifest_path,
                local_result={"items": [], "structured": {}},
                profile=TECHNICAL_PARSE_PROFILE,
            )

        load_model_config.assert_called_once_with()
        self.assertEqual(sorted(seen), sorted(["projectBasics", *shard_keys()]))
        self.assertEqual(len(initialized_configs), 7)
        self.assertTrue(all(config is model_config for config in initialized_configs))

    def test_failed_shard_is_retried_then_surfaced_without_killing_the_run(self) -> None:
        broken = shard_keys()[0]
        resolved, message, seen = self._run(failing_shards={broken})

        # 首轮 + 重试轮，失败分片被跑了两次
        self.assertEqual(seen.count(broken), 2)
        # 失败必须显式暴露，不能静默当成成功
        self.assertIn(broken, resolved["structured"]["workflow"]["failedShards"])
        self.assertIn("未完成", message)
        # 其余分片照常收口，finalize 仍输出完整 58 条（失败分片按 missing）
        self.assertEqual(len(resolved["items"]), 58)

    def test_session_that_finishes_without_submitting_counts_as_failure(self) -> None:
        """会话「正常结束」但没调 submit，必须判失败并暴露，不能静默当成解析完成。"""
        silent = shard_keys()[1]
        resolved, message, seen = self._run(silent_shards={silent})

        self.assertEqual(seen.count(silent), 2, "没提交的分片必须被重试")
        self.assertIn(silent, resolved["structured"]["workflow"]["failedShards"])
        self.assertIn("没有提交结果", message)

    def test_session_that_submits_only_part_of_its_rows_counts_as_failure(self) -> None:
        """只提交分片部分行不能冒充成功，否则剩余行会被 finalize 静默写成 missing。"""
        partial = shard_keys()[2]
        resolved, message, seen = self._run(partial_shards={partial})

        self.assertEqual(seen.count(partial), 2, "提交不完整的分片必须被重试")
        self.assertIn(partial, resolved["structured"]["workflow"]["failedShards"])
        self.assertIn("提交不完整", message)

    def test_project_basics_must_submit_all_six_standard_fields(self) -> None:
        resolved, message, seen = self._run(partial_shards={"projectBasics"})

        self.assertEqual(seen.count("projectBasics"), 2, "基础信息提交不完整时必须重试")
        self.assertIn("projectBasics", resolved["structured"]["workflow"]["failedShards"])
        self.assertIn("提交不完整", message)

    def test_new_run_does_not_accept_stale_shard_submission(self) -> None:
        """同一 manifest 重跑时，上一轮提交不能冒充本轮会话产出。"""
        stale = shard_keys()[0]
        manifest = load_manifest(self.manifest_path)
        rows = [
            {
                "rowNo": int(row["rowNo"]),
                "status": "missing",
                "conclusion": "上一轮残留结果",
                "evidenceSummary": "",
                "evidenceIds": [],
            }
            for row in checklist_for_shard(stale)
        ]
        submit(self.manifest_path, manifest, "technicalInterpretation", rows, shard=stale)

        resolved, message, seen = self._run(silent_shards={stale})

        self.assertEqual(seen.count(stale), 2, "本轮未提交的分片必须重试")
        self.assertIn(stale, resolved["structured"]["workflow"]["failedShards"])
        self.assertIn("没有提交结果", message)


class DeterministicCliTests(unittest.TestCase):
    def test_s1parse_subprocess_timeout_is_reported_as_runtime_error(self) -> None:
        from unittest.mock import patch

        from app.services import parsing

        with tempfile.TemporaryDirectory() as tmp:
            manifest_path = Path(tmp) / "s1_parse_manifest.json"
            with patch.object(
                parsing.subprocess,
                "run",
                side_effect=subprocess.TimeoutExpired(cmd="s1parse prepare", timeout=600),
            ):
                with self.assertRaisesRegex(RuntimeError, "超时"):
                    parsing._run_s1parse_cli("prepare", manifest_path)


if __name__ == "__main__":
    unittest.main()
