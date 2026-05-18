from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path

from docx import Document

from app.core.config import settings
from app.services.business_document_editing import apply_controlled_business_rewrite
from app.services.onlyoffice_documents import document_path


class BusinessDocumentEditingTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        self.original_documents_dir = settings.documents_dir
        settings.documents_dir = Path(self.temp_dir.name) / "documents"
        settings.documents_dir.mkdir(parents=True, exist_ok=True)

    def tearDown(self) -> None:
        settings.documents_dir = self.original_documents_dir
        self.temp_dir.cleanup()

    def test_apply_controlled_rewrite_replaces_unique_paragraph_and_writes_history(self) -> None:
        project_id = "PRJ-BIZ-REWRITE"
        path = document_path(project_id)
        doc = Document()
        doc.add_paragraph("投标人承诺按招标文件要求履行合同。")
        doc.add_paragraph("本段不应被修改。")
        doc.save(path)

        result = apply_controlled_business_rewrite(
            project_id,
            original_text="投标人承诺按招标文件要求履行合同。",
            replacement_text="投标人承诺严格按照招标文件及合同约定履行全部商务义务。",
            operator="测试用户",
        )

        updated = Document(str(path))
        text = "\n".join(paragraph.text for paragraph in updated.paragraphs)
        self.assertIn("投标人承诺严格按照招标文件及合同约定履行全部商务义务。", text)
        self.assertIn("本段不应被修改。", text)
        self.assertTrue(Path(result["backupFile"]).exists())

        history_path = Path(result["historyFile"])
        self.assertTrue(history_path.exists())
        history = [json.loads(line) for line in history_path.read_text(encoding="utf-8").splitlines()]
        self.assertEqual(history[-1]["operator"], "测试用户")
        self.assertEqual(history[-1]["matchMode"], "exact_paragraph")

    def test_apply_controlled_rewrite_rejects_ambiguous_matches(self) -> None:
        project_id = "PRJ-BIZ-REWRITE-DUP"
        path = document_path(project_id)
        doc = Document()
        doc.add_paragraph("请严格响应招标文件要求。")
        doc.add_paragraph("请严格响应招标文件要求。")
        doc.save(path)

        with self.assertRaises(ValueError):
            apply_controlled_business_rewrite(
                project_id,
                original_text="请严格响应招标文件要求。",
                replacement_text="请严格响应招标文件全部商务要求。",
                operator="测试用户",
            )

        updated = Document(str(path))
        self.assertEqual(
            [paragraph.text for paragraph in updated.paragraphs],
            ["请严格响应招标文件要求。", "请严格响应招标文件要求。"],
        )


if __name__ == "__main__":
    unittest.main()
