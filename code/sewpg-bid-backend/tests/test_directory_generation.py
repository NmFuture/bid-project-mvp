from __future__ import annotations

import copy
import hashlib
import importlib.util
import json
import re
import shutil
import sys
import tempfile
import unittest
from concurrent.futures import ThreadPoolExecutor as RealThreadPoolExecutor
from pathlib import Path
from unittest.mock import MagicMock, patch

from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from app.core.config import settings
from app.services.bid_parse_state import complete_parse_state
from app.services.bid_project_state import update_template_fallback_state
from app.services.bid_outline_state import (
    complete_directory_generation_state,
    directory_state_with_rule_evidence,
    regenerate_outline_state,
    save_outline_state,
    start_directory_generation_state,
)
from app.services.store import store
from app.services.workspace_artifacts import business_workspace_dir, technical_workspace_dir


BACKEND_ROOT = Path(__file__).resolve().parents[1]


def skill_script_path(skill_name: str, script_name: str) -> Path:
    for root_name in ("skills", "skill"):
        candidate = BACKEND_ROOT / "opencode" / root_name / skill_name / "scripts" / script_name
        if candidate.exists():
            return candidate
    return BACKEND_ROOT / "opencode" / "skills" / skill_name / "scripts" / script_name


def load_technical_outline_runner():
    module_name = "directory_generation_technical_outline_runner"
    loaded = sys.modules.get(module_name)
    if loaded is not None:
        return loaded
    script_path = skill_script_path("bid-tech-outline-generator", "run_from_manifest.py")
    spec = importlib.util.spec_from_file_location(module_name, script_path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"无法加载技术标目录 Skill：{script_path}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    return module


class DirectoryGenerationTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        base = Path(self.temp_dir.name)
        settings.uploads_dir = base / "uploads"
        settings.documents_dir = base / "documents"
        settings.parsed_dir = base / "parsed"
        settings.ensure_dirs()

        store.reset_for_tests()
        self.client = TestClient(app, base_url="http://127.0.0.1:8000")
        self.parallel_outline_patcher = patch(
            "app.services.outline_generation._run_parallel_outline_chapters",
            return_value=[],
        )
        self.parallel_outline_patcher.start()

    def tearDown(self) -> None:
        self.parallel_outline_patcher.stop()
        self.client.close()
        self.temp_dir.cleanup()

    def _write_docx(self, path: Path, paragraphs: list[tuple[str, str | None]]) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        for text, style in paragraphs:
            doc.add_paragraph(text, style=style) if style else doc.add_paragraph(text)
        doc.save(path)

    def _parse_runtime_state(self, project_id: str) -> tuple[dict, dict]:
        project = store.get_project_runtime_state(project_id)
        return copy.deepcopy(project["parse_result"]), copy.deepcopy(project["parse_storage"])

    def _directory_state_for_tests(self, project_id: str) -> dict:
        project = store.require_project_for_update(project_id)
        return directory_state_with_rule_evidence(project)

    def _outline_state_for_tests(self, project_id: str) -> dict:
        project = store.require_project_for_update(project_id)
        return copy.deepcopy(project["outline_state"])

    def _update_template_fallback_for_tests(self, project_id: str, data: dict) -> dict:
        project = store.require_project_for_update(project_id)
        payload = update_template_fallback_state(project, data)
        store.persist_project_state(project)
        return payload

    def _start_directory_generation_for_tests(self, project_id: str) -> dict:
        project = store.require_project_for_update(project_id)
        payload = start_directory_generation_state(project)
        store.persist_project_state(project)
        return payload

    def _save_outline_for_tests(self, project_id: str, nodes: list[dict]) -> dict:
        project = store.require_project_for_update(project_id)
        payload = save_outline_state(project, nodes)
        store.persist_project_state(project)
        return payload

    def _regenerate_outline_for_tests(self, project_id: str) -> dict:
        project = store.require_project_for_update(project_id)
        payload = regenerate_outline_state(project)
        store.persist_project_state(project)
        return payload

    def _complete_parse_for_tests(
        self,
        project_id: str,
        tender_files: list[dict],
        template_files: list[dict],
        *,
        summary: dict | None = None,
        parse_storage: dict | None = None,
    ) -> dict:
        project = store.require_project_for_update(project_id)
        payload = complete_parse_state(
            project,
            tender_files,
            template_files,
            summary=summary,
            parse_storage=parse_storage,
        )
        store.persist_project_state(project)
        return payload

    def _prepare_project_with_parse_result(self, bid_type: str = "技术标") -> str:
        project = store.create_project(
            {
                "name": "目录生成联调项目",
                "customerName": "测试业主",
                "bidType": bid_type,
            }
        )
        project_id = project["id"]

        project_dir = business_workspace_dir(project_id) if bid_type == "商务标" else technical_workspace_dir(project_id)
        project_dir.mkdir(parents=True, exist_ok=True)

        tender_path = settings.uploads_dir / project_id / "tender" / "招标文件.docx"
        self._write_docx(
            tender_path,
            [
                ("第一章 采购需求", "Heading 1"),
                ("1.1 项目范围", "Heading 2"),
                ("投标人应提供实施方案。", None),
                ("投标人须提交服务团队安排。", None),
                ("第二章 评审办法", "Heading 1"),
                ("评分项：实施方案，分值30分，证明材料要求：提供项目计划。", None),
            ],
        )

        template_path = settings.uploads_dir / project_id / "template" / "投标文件-正文.docx"
        self._write_docx(
            template_path,
            [
                ("第一章 投标响应概述", "Heading 1"),
                ("1.1 项目理解", "Heading 2"),
                ("第二章 实施方案", "Heading 1"),
                ("2.1 工作计划", "Heading 2"),
            ],
        )

        combined_text_path = project_dir / "combined.txt"
        combined_text_path.write_text(
            "\n".join(
                [
                    "# 文件：招标文件.docx",
                    "",
                    "第一章 采购需求",
                    "1.1 项目范围",
                    "投标人应提供实施方案。",
                    "投标人须提交服务团队安排。",
                    "第二章 评审办法",
                    "评分项：实施方案，分值30分，证明材料要求：提供项目计划。",
                ]
            ),
            encoding="utf-8",
        )

        self._complete_parse_for_tests(
            project_id,
            tender_files=[
                {
                    "id": "TEN-1",
                    "name": "招标文件.docx",
                    "path": str(tender_path),
                    "size_label": "1.0 MB",
                }
            ],
            template_files=[
                {
                    "id": "TPL-1",
                    "name": "投标文件-正文.docx",
                    "path": str(template_path),
                    "size_label": "1.0 MB",
                }
            ],
            summary={
                "fileCount": 1,
                "extractedCount": 2,
                "textLength": 120,
                "textPreview": "",
                "warnings": [],
            },
            parse_storage={
                "projectDir": str(project_dir),
                "combinedTextPath": str(combined_text_path),
                "manifestPath": "",
                "documents": [],
            },
        )
        return project_id

    def _technical_manifest_from_prompt(self, prompt: str, kwargs: dict) -> Path:
        self.assertIn("Use the bid-tech-outline-generator skill", prompt)
        self.assertIn("s2outline appendix-next", prompt)
        self.assertIn("--max-items 40", prompt)
        self.assertIn("s2outline headings", prompt)
        self.assertIn("s2outline section", prompt)
        self.assertIn("review-corrections", prompt)
        self.assertIn("review-complete", prompt)
        self.assertIn("decisions", prompt)
        self.assertIn("compose", prompt)
        self.assertIn("不要执行 `s2outline prepare`", prompt)
        self.assertIn("模板正文目录的三类判断已经由前序接力会话全部完成", prompt)
        self.assertIn("s2outline finalize", prompt)
        self.assertNotIn("s2toc ", prompt)
        self.assertEqual(kwargs.get("early_tool_command"), "s2outline-finalize")
        self.assertTrue(callable(kwargs.get("terminal_validator")))
        self.assertIsNone(kwargs.get("handoff_prompt_factory"))
        self.assertIsNone(kwargs.get("handoff_state_callback"))
        match = re.search(r"manifest[：:]\s*(?P<path>.+)", prompt)
        self.assertIsNotNone(match)
        manifest_path = Path(str(match.group("path")).strip())
        self.assertTrue(manifest_path.exists())
        return manifest_path

    def _write_mock_technical_toc_outputs(
        self,
        manifest_path: Path,
        *,
        first_root_title: str = "投标响应概述",
    ) -> dict:
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        runner = load_technical_outline_runner()
        runner.write_template_structure(manifest, manifest_path)
        structure = json.loads(
            (manifest_path.parent / "template_structure.json").read_text(encoding="utf-8")
        )
        items = structure["items"]
        roots = [item for item in items if item.get("parent_id") is None]
        self.assertGreaterEqual(len(roots), 2)
        headings = runner.dispatch_command("headings", manifest, manifest_path, [])
        self.assertTrue(headings["complete"])
        section_id = next(
            item["section_id"]
            for file_item in headings["files"]
            for item in file_item["items"]
            if item.get("section_id")
        )
        section = runner.dispatch_command(
            "section", manifest, manifest_path, [section_id]
        )
        evidence_id = next(
            record["evidence_id"]
            for record in section["records"]
            if record.get("evidence_id")
        )
        appendix_candidates = runner.review_workflow.decision_appendix_items(
            manifest_path.parent
        )

        additions = []
        if first_root_title != "投标响应概述":
            additions.append(
                {
                    "node_id": "ADD-MOCK-SEMANTIC-PATH",
                    "parent_id": None,
                    "number": "第三章",
                    "title": first_root_title,
                    "reason": "验证发布过程不会改写目录语义文本。",
                }
            )
        appendix_number = "第四章" if additions else "第三章"
        tender_document = Document(str(manifest["tenderFiles"][0]["path"]))
        tender_text = "\n".join(paragraph.text for paragraph in tender_document.paragraphs)
        appendix_basis = (
            {
                "file_id": "TEN-1",
                "search_text": "附表D.7 性能及考核承诺保证表",
            }
            if "附表D.7 性能及考核承诺保证表" in tender_text
            else None
        )
        appendix_addition = {
            "node_id": "ADD-MOCK-APPENDIX-D7",
            "parent_id": "ADD-MOCK-APPENDIX",
            "number": "附表D.7",
            "title": "性能及考核承诺保证表",
            "reason": "招标文件新增需填写技术附表。",
        }
        if appendix_basis:
            appendix_addition["tender_basis"] = appendix_basis
        if not appendix_candidates:
            additions.extend(
                [
                    {
                        "node_id": "ADD-MOCK-APPENDIX",
                        "parent_id": None,
                        "number": appendix_number,
                        "title": "技术附表",
                        "reason": "招标文件包含需单独确认的技术附表。",
                    },
                    appendix_addition,
                ]
            )
        for addition in additions:
            addition["evidence_id"] = evidence_id
        pending_additions = additions
        while True:
            batch = runner.dispatch_command(
                "decision-next",
                manifest,
                manifest_path,
                [],
            )
            if batch["complete"]:
                break
            runner.dispatch_command(
                "section",
                manifest,
                manifest_path,
                [section_id],
            )
            runner.dispatch_command(
                "decision-batch",
                manifest,
                manifest_path,
                [
                    json.dumps(
                        {
                            "batch_token": batch["batch_token"],
                            "items": [
                                {
                                    "target_id": item["target_id"],
                                    "decision": "retain",
                                    "reason": "历史模板中的专业目录仍适用于本项目",
                                }
                                for item in batch["items"]
                            ],
                            "additions": pending_additions,
                        },
                        ensure_ascii=False,
                    )
                ],
            )
            pending_additions = []
        if appendix_candidates:
            appendix_batch = runner.dispatch_command(
                "appendix-next", manifest, manifest_path, []
            )
            runner.dispatch_command(
                "appendix-decision-batch",
                manifest,
                manifest_path,
                [
                    json.dumps(
                        {
                            "batch_token": appendix_batch["batch_token"],
                            "root_addition": {
                                "node_id": "ADD-MOCK-APPENDIX",
                                "parent_id": None,
                                "number": appendix_number,
                                "title": "技术附表",
                                "reason": "招标文件包含需单独确认的技术附表。",
                            },
                            "items": [
                                {
                                    "appendix_id": item["appendix_id"],
                                    "decision": "include",
                                    "node_id": f"ADD-MOCK-APPENDIX-{index}",
                                    "parent_id": "ADD-MOCK-APPENDIX",
                                    "reason": "招标文件要求独立填写。",
                                }
                                for index, item in enumerate(
                                    appendix_batch["items"], start=1
                                )
                            ],
                        },
                        ensure_ascii=False,
                    )
                ],
            )
        runner.dispatch_command("section", manifest, manifest_path, [section_id])
        runner.dispatch_command(
            "review-complete",
            manifest,
            manifest_path,
            [json.dumps({"review_summary": "测试全局复核。", "issues": []}, ensure_ascii=False)],
        )
        runner.dispatch_command("decisions", manifest, manifest_path, [])
        return runner.compose_manifest(manifest, manifest_path)

    def _mock_futurecode_outline(self, prompt: str, *args, **kwargs) -> dict:
        manifest_path = self._technical_manifest_from_prompt(prompt, kwargs)
        result = self._write_mock_technical_toc_outputs(manifest_path)
        result["opencodeOutput"] = {
            "status": "received",
            "sessionId": "test-s2-session",
            "providerId": "opencode",
            "modelId": "big-pickle",
            "receivedAt": "2026-05-02T00:00:00Z",
            "parts": [{"type": "text", "text": "{}"}],
        }
        return result

    def _mock_business_futurecode_outline(self, prompt: str, *args, **kwargs) -> dict:
        self.assertIn("Use the bid-business-outline-generator skill", prompt)
        self.assertNotIn('"agentDecisions"', prompt)
        self.assertIn("完整执行 bid-business-outline-generator Skill", prompt)
        self.assertIn("准备脚本只生成输入材料", prompt)
        self.assertIn("不要自行生成或修改前端兼容 toc.json", prompt)
        self.assertIn("manifest.templateFile", prompt)
        self.assertIn("不扫描当前工作目录", prompt)
        self.assertIn("不使用 user_confirmed_inputs.json", prompt)
        self.assertIn("第一条非 skill 工具调用必须是 Bash", prompt)
        self.assertIn("禁止在这条 Bash 命令完成前调用 read", prompt)
        self.assertIn("不要读取 manifest 内容来“理解输入”", prompt)
        self.assertIn("Mandatory tool order", prompt)
        self.assertIn("Do not inspect the manifest first", prompt)
        self.assertIn("source_text_candidates.json 的首选候选", prompt)
        self.assertIn("outline_authoring_decisions.json", prompt)
        self.assertIn("outline_authoring_helper.py", prompt)
        self.assertIn("opencode 只输出语义选择、状态判断和保留/延后理由", prompt)
        self.assertIn("evidence_scope", prompt)
        self.assertIn("evidence_strength", prompt)
        self.assertIn("合计 | 100", prompt)
        self.assertIn('"schema_version": "business_bid_outline.v1"', prompt)
        self.assertTrue(any(line.strip().startswith("business-outline ") for line in prompt.splitlines()))
        self.assertEqual(kwargs.get("early_tool_command"), "")
        manifest_line = next(line for line in prompt.splitlines() if line.strip().startswith("business-outline "))
        manifest_path = Path(manifest_line.strip().split(" ", 1)[1])
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        work_dir = Path(manifest["workDir"])
        business_outline_file = work_dir / "outline.json"
        (work_dir / "history_bid_outline_inputs.json").write_text(
            json.dumps({"document_name": "template-main.docx", "outline_candidates": []}, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        (work_dir / "tender_map_inputs.json").write_text(
            json.dumps({"document_name": "招标文件.docx", "blocks": [], "tables": []}, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        business_outline_file.write_text(
            json.dumps(
                {
                    "schema_version": "business_bid_outline.v1",
                    "document_name": "商务标目录",
                    "sections": [
                        {
                            "id": "BIZ-1",
                            "title": "商务响应文件",
                            "number": "一、",
                            "level": 1,
                            "required_status": "必要",
                            "source_text": "投标人须提交服务团队安排。",
                            "children": [],
                        }
                    ],
                    "review_items": [],
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        return {
            "schema_version": "business_bid_outline.v1",
            "businessOutlineFile": str(business_outline_file),
            "historyBidOutlineInputsFile": str(work_dir / "history_bid_outline_inputs.json"),
            "tenderMapInputsFile": str(work_dir / "tender_map_inputs.json"),
            "summary": {"total_sections": 1},
            "agentDecisions": [],
            "opencodeOutput": {
                "status": "received",
                "sessionId": "test-business-outline-session",
                "providerId": "opencode",
                "modelId": "big-pickle",
                "receivedAt": "2026-05-04T00:00:00Z",
                "parts": [{"type": "text", "text": "{}"}],
            },
        }

    def _mock_business_futurecode_outline_with_result_overrides(self, prompt: str, *args, **kwargs) -> dict:
        result = self._mock_business_futurecode_outline(prompt, *args, **kwargs)
        manifest_line = next(line for line in prompt.splitlines() if line.strip().startswith("business-outline "))
        manifest_path = Path(manifest_line.strip().split(" ", 1)[1])
        work_dir = Path(json.loads(manifest_path.read_text(encoding="utf-8"))["workDir"])
        business_outline_file = work_dir / "outline.json"
        business_outline_file.write_text(
            json.dumps(
                {
                    "schema_version": "business_bid_outline.v1",
                    "sections": [
                        {
                            "id": "BIZ-1",
                            "title": "Business complete parent",
                            "number": "A.",
                            "level": 1,
                            "required_status": "necessary",
                            "source_text": "Tender requires the complete business parent item.",
                            "children": [
                                {
                                    "id": "BIZ-1-1",
                                    "title": "Business complete child",
                                    "number": "",
                                    "level": 2,
                                    "required_status": "necessary",
                                    "source_text": "Tender requires the complete business child item.",
                                    "children": [],
                                }
                            ],
                        }
                    ],
                    "review_items": [],
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        result["items"] = [
            {
                "itemId": "TOC-MODEL-0001",
                "order": 1,
                "title": "Incomplete model result must not win",
                "level": 1,
                "annotation": "adapted",
                "source_refs": [],
                "material_refs": [],
            }
        ]
        result["agentDecisions"] = [
            {
                "itemId": "TOC-MODEL-0001",
                "title": "Explanatory model decision must not rewrite business toc",
                "decision": "improved_source",
                "new_source": "Explanation text, not a directory patch.",
            }
        ]
        return result

    def _mock_business_futurecode_outline_missing_outline(self, prompt: str, *args, **kwargs) -> dict:
        manifest_line = next(line for line in prompt.splitlines() if line.strip().startswith("business-outline "))
        manifest_path = Path(manifest_line.strip().split(" ", 1)[1])
        work_dir = Path(json.loads(manifest_path.read_text(encoding="utf-8"))["workDir"])
        (work_dir / "history_bid_outline_inputs.json").write_text("{}", encoding="utf-8")
        (work_dir / "tender_map_inputs.json").write_text("{}", encoding="utf-8")
        return {
            "schema_version": "business_bid_outline.v1",
            "businessOutlineFile": str(work_dir / "outline.json"),
            "summary": {"total_sections": 0},
            "opencodeOutput": {"status": "received", "parts": []},
        }

    def _mock_business_futurecode_outline_with_outline_payload(self, payload: dict) -> callable:
        def _mock(prompt: str, *args, **kwargs) -> dict:
            manifest_line = next(line for line in prompt.splitlines() if line.strip().startswith("business-outline "))
            manifest_path = Path(manifest_line.strip().split(" ", 1)[1])
            work_dir = Path(json.loads(manifest_path.read_text(encoding="utf-8"))["workDir"])
            (work_dir / "history_bid_outline_inputs.json").write_text("{}", encoding="utf-8")
            (work_dir / "tender_map_inputs.json").write_text("{}", encoding="utf-8")
            business_outline_file = work_dir / "outline.json"
            business_outline_file.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
            return {
                "schema_version": "business_bid_outline.v1",
                "businessOutlineFile": str(business_outline_file),
                "summary": {"total_sections": len(payload.get("sections") or [])},
                "opencodeOutput": {"status": "received", "parts": []},
            }

        return _mock

    def test_generate_business_outline_uses_business_workspace_and_skill(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project_id = self._prepare_project_with_parse_result("商务标")

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=self._mock_business_futurecode_outline,
        ):
            payload = generate_outline_for_project(project_id, {"outlineStrategy": "strict"})

        work_dir = Path(payload["opencodeOutput"]["workDir"])
        manifest = json.loads(Path(payload["opencodeOutput"]["canonicalManifestPath"]).read_text(encoding="utf-8"))
        self.assertEqual(payload["status"], "completed")
        self.assertIn("business-workspace", str(work_dir))
        self.assertNotIn("technical-workspace", str(work_dir))
        self.assertEqual(work_dir, business_workspace_dir(project_id) / "s2_toc_workdir")
        self.assertEqual(payload["opencodeOutput"]["engine"], "bid-business-outline-generator")
        self.assertEqual(payload["opencodeOutput"]["skill"], "bid-business-outline-generator")
        self.assertEqual(manifest["bidType"], "商务标")
        self.assertIn("business-workspace", manifest["workDir"])
        self.assertEqual(Path(manifest["templateFile"]).parent, work_dir)
        self.assertEqual(Path(manifest["templateFile"]).name, "template-main.docx")
        toc = json.loads(Path(payload["opencodeOutput"]["tocJsonPath"]).read_text(encoding="utf-8"))
        business_outline_path = business_workspace_dir(project_id) / "s2_toc_workdir" / "outline.json"
        history_inputs_path = business_workspace_dir(project_id) / "s2_toc_workdir" / "history_bid_outline_inputs.json"
        tender_map_inputs_path = business_workspace_dir(project_id) / "s2_toc_workdir" / "tender_map_inputs.json"
        self.assertEqual(toc["schema_version"], "bid-toc-json-v1")
        self.assertTrue(business_outline_path.exists())
        self.assertTrue(history_inputs_path.exists())
        self.assertTrue(tender_map_inputs_path.exists())
        self.assertEqual(toc["businessOutlineFile"], str(business_outline_path))
        self.assertEqual(payload["opencodeOutput"]["businessOutlinePath"], str(business_outline_path))
        self.assertEqual(payload["opencodeOutput"]["historyBidOutlineInputsPath"], str(history_inputs_path))
        self.assertEqual(payload["opencodeOutput"]["tenderMapInputsPath"], str(tender_map_inputs_path))
        self.assertTrue(all(item.get("source") == "business_outline" for item in toc["items"]))
        self.assertTrue(all(item.get("required_status") and item.get("requiredStatus") for item in toc["items"]))
        self.assertTrue(all(item.get("source_text") and item.get("sourceText") for item in toc["items"]))
        searchable_items = [item for item in toc["items"] if item.get("source_refs")]
        self.assertTrue(searchable_items)
        self.assertTrue(
            all(
                ref.get("searchText") and ref.get("basisText") and ref.get("rawText")
                for item in searchable_items
                for ref in item.get("source_refs", [])
            )
        )
        self.assertTrue((business_workspace_dir(project_id) / "s2_toc_workdir" / "s2_input.json").exists())
        self.assertFalse((technical_workspace_dir(project_id) / "s2_toc_workdir").exists())

    def test_generate_business_outline_uses_outline_json_over_model_result_items(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project_id = self._prepare_project_with_parse_result("商务标")

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=self._mock_business_futurecode_outline_with_result_overrides,
        ):
            payload = generate_outline_for_project(project_id, {"outlineStrategy": "strict"})

        toc = json.loads(Path(payload["opencodeOutput"]["tocJsonPath"]).read_text(encoding="utf-8"))
        self.assertEqual(toc["schema_version"], "bid-toc-json-v1")
        self.assertEqual([item["title"] for item in toc["items"]], ["Business complete parent", "Business complete child"])
        self.assertTrue(all(item.get("source") == "business_outline" for item in toc["items"]))
        self.assertTrue(all(item.get("source_text") for item in toc["items"]))
        self.assertFalse(toc.get("ruleEvidence", {}).get("agentDecisions"))

    def test_generate_business_outline_preserves_skill_number_values(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project_id = self._prepare_project_with_parse_result("商务标")
        outline_payload = {
            "schema_version": "business_bid_outline.v1",
            "sections": [
                {
                    "id": "BIZ-1",
                    "title": "投标函及授权文件",
                    "number": "一、",
                    "level": 1,
                    "required_status": "必要",
                    "source_text": "投标文件应包括投标函、法定代表人身份证明或授权委托书。",
                    "children": [
                        {
                            "id": "BIZ-1-1",
                            "title": "投标函",
                            "number": "1.1",
                            "level": 2,
                            "required_status": "必要",
                            "source_text": "投标函格式",
                            "children": [],
                        },
                        {
                            "id": "BIZ-1-2",
                            "title": "商务评分索引表",
                            "number": "",
                            "level": 2,
                            "required_status": "待确认",
                            "source_text": "商务评分索引表",
                            "children": [],
                        },
                        {
                            "id": "BIZ-1-3",
                            "title": "供货保障专题",
                            "number": None,
                            "level": 2,
                            "required_status": "待确认",
                            "source_text": "供货保障专题",
                            "children": [],
                        },
                    ],
                }
            ],
            "review_items": [],
        }

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=self._mock_business_futurecode_outline_with_outline_payload(outline_payload),
        ):
            payload = generate_outline_for_project(project_id, {"outlineStrategy": "strict"})

        toc = json.loads(Path(payload["opencodeOutput"]["tocJsonPath"]).read_text(encoding="utf-8"))
        self.assertEqual(
            [(item["title"], item["number"]) for item in toc["items"]],
            [
                ("投标函及授权文件", "一、"),
                ("投标函", "1.1"),
                ("商务评分索引表", ""),
                ("供货保障专题", ""),
            ],
        )
        outline = self._outline_state_for_tests(project_id)
        root = outline["nodes"][0]
        self.assertEqual(root["title"], "投标函及授权文件")
        self.assertEqual(root["tocNumber"], "一、")
        self.assertEqual(root["children"][0]["tocNumber"], "1.1")
        self.assertEqual(root["children"][1]["tocNumber"], "")
        self.assertEqual(root["children"][2]["tocNumber"], "")

    def test_generate_business_outline_requires_final_outline_json(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project_id = self._prepare_project_with_parse_result("商务标")

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=self._mock_business_futurecode_outline_missing_outline,
        ):
            with self.assertRaisesRegex(RuntimeError, "outline.json"):
                generate_outline_for_project(project_id, {"outlineStrategy": "strict"})

    def test_business_outline_uses_task_timeout_for_long_opencode_run(self) -> None:
        from app.services.outline_generation import _run_business_outline_skill

        work_dir = settings.documents_dir / "business-timeout-work"
        work_dir.mkdir(parents=True, exist_ok=True)
        manifest_path = work_dir / "s2_input.json"
        manifest_path.write_text(
            json.dumps(
                {
                    "projectId": "PRJ-TIMEOUT",
                    "projectCode": "PRJ-TIMEOUT",
                    "projectName": "商务目录长任务",
                    "bidType": "商务标",
                    "workDir": str(work_dir),
                    "tenderFiles": [],
                    "templateFile": str(work_dir / "template.docx"),
                    "outputFile": str(work_dir / "toc.json"),
                    "evidenceFile": str(work_dir / "toc_evidence.json"),
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        business_outline_file = work_dir / "outline.json"

        def _generate(*args, **kwargs) -> dict:
            business_outline_file.write_text(
                json.dumps(
                    {
                        "schema_version": "business_bid_outline.v1",
                        "sections": [
                            {
                                "id": "BIZ-1",
                                "title": "商务响应文件",
                                "number": None,
                                "level": 1,
                                "required_status": "必要",
                                "source_text": "招标文件要求提交商务响应文件。",
                                "children": [],
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            return {
                "schema_version": "business_bid_outline.v1",
                "businessOutlineFile": str(business_outline_file),
                "summary": {"total_sections": 1},
                "opencodeOutput": {"status": "received", "parts": []},
            }

        with patch("app.services.outline_generation.OpencodeClient") as client_cls:
            client_cls.return_value.generate_outline_with_trace.side_effect = _generate

            _run_business_outline_skill(manifest_path)

        self.assertEqual(
            client_cls.call_args.kwargs["timeout_ms"],
            int(settings.opencode_timeout_sec * 1000),
        )

    def test_generate_business_outline_rejects_invalid_outline_schema(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project_id = self._prepare_project_with_parse_result("商务标")

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=self._mock_business_futurecode_outline_with_outline_payload(
                {"schema_version": "bid-toc-json-v1", "sections": [{"title": "商务响应文件"}]}
            ),
        ):
            with self.assertRaisesRegex(RuntimeError, "business_bid_outline.v1"):
                generate_outline_for_project(project_id, {"outlineStrategy": "strict"})

    def test_generate_business_outline_rejects_empty_outline_sections(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project_id = self._prepare_project_with_parse_result("商务标")

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=self._mock_business_futurecode_outline_with_outline_payload(
                {"schema_version": "business_bid_outline.v1", "sections": []}
            ),
        ):
            with self.assertRaisesRegex(RuntimeError, "sections"):
                generate_outline_for_project(project_id, {"outlineStrategy": "strict"})

    def test_generate_business_outline_does_not_fallback_to_local_business_runner(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project_id = self._prepare_project_with_parse_result("商务标")

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=RuntimeError("offline business outline"),
        ):
            with self.assertRaisesRegex(RuntimeError, "只负责准备候选材料"):
                generate_outline_for_project(project_id, {"outlineStrategy": "strict"})

        work_dir = business_workspace_dir(project_id) / "s2_toc_workdir"
        self.assertFalse((work_dir / "outline.json").exists())
        self.assertFalse((work_dir / "toc.json").exists())

    def test_generate_business_outline_rejects_missing_section_number(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project_id = self._prepare_project_with_parse_result("商务标")

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=self._mock_business_futurecode_outline_with_outline_payload(
                {
                    "schema_version": "business_bid_outline.v1",
                    "sections": [
                        {
                            "id": "BIZ-1",
                            "title": "商务响应文件",
                            "level": 1,
                            "required_status": "必要",
                            "source_text": "投标人须提交商务响应文件。",
                            "children": [],
                        }
                    ],
                    "review_items": [],
                }
            ),
        ):
            with self.assertRaisesRegex(RuntimeError, "sections\\[0\\]\\.number"):
                generate_outline_for_project(project_id, {"outlineStrategy": "strict"})

    def test_business_outline_runner_prepares_inputs_without_final_outline(self) -> None:
        import importlib.util
        import sys

        script_path = skill_script_path("bid-business-outline-generator", "run_from_manifest.py")
        spec = importlib.util.spec_from_file_location("business_outline_runner_input_only_test", script_path)
        runner = importlib.util.module_from_spec(spec)
        assert spec and spec.loader
        sys.modules["business_outline_runner_input_only_test"] = runner
        spec.loader.exec_module(runner)

        project_id = self._prepare_project_with_parse_result("商务标")
        work_dir = business_workspace_dir(project_id) / "runner-input-only"
        work_dir.mkdir(parents=True, exist_ok=True)
        template_file = settings.uploads_dir / project_id / "template" / "投标文件-正文.docx"
        tender_file = settings.uploads_dir / project_id / "tender" / "招标文件.docx"
        manifest_path = work_dir / "s2_input.json"
        manifest = {
            "projectId": project_id,
            "projectName": "目录生成联调项目",
            "bidType": "商务标",
            "workDir": str(work_dir),
            "templateFile": str(template_file),
            "tenderFiles": [{"id": "TEN-1", "name": "招标文件.docx", "path": str(tender_file)}],
            "outputFile": str(work_dir / "toc.json"),
            "evidenceFile": str(work_dir / "toc_evidence.json"),
        }
        manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

        result = runner.run_manifest(manifest, manifest_path)

        self.assertTrue((work_dir / "history_bid_outline_inputs.json").exists())
        self.assertTrue((work_dir / "tender_map_inputs.json").exists())
        self.assertTrue((work_dir / "document_structure_index.json").exists())
        self.assertTrue((work_dir / "source_text_candidates.json").exists())
        outline_path = work_dir / "outline.json"
        self.assertFalse(outline_path.exists())
        self.assertFalse((work_dir / "toc.json").exists())
        self.assertFalse((work_dir / "toc_evidence.json").exists())
        self.assertFalse((work_dir / "agent_review_input.json").exists())
        self.assertEqual(result["summary"]["schema_version"], "business-outline-inputs-v1")
        self.assertNotIn("businessOutlineFile", result["summary"])
        self.assertEqual(result["summary"]["historyBidOutlineInputsFile"], str(work_dir / "history_bid_outline_inputs.json"))
        self.assertEqual(result["summary"]["tenderMapInputsFile"], str(work_dir / "tender_map_inputs.json"))
        self.assertEqual(result["summary"]["documentStructureIndexFile"], str(work_dir / "document_structure_index.json"))
        self.assertEqual(result["summary"]["sourceTextCandidatesFile"], str(work_dir / "source_text_candidates.json"))

    def test_business_outline_regenerate_uses_generated_business_toc_not_technical_defaults(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project_id = self._prepare_project_with_parse_result("商务标")
        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=self._mock_business_futurecode_outline,
        ):
            generate_outline_for_project(project_id, {"outlineStrategy": "strict"})

        self._save_outline_for_tests(project_id, [{"id": "OL-X", "title": "人工改坏目录", "children": []}])
        payload = self._regenerate_outline_for_tests(project_id)
        titles = [node["title"] for node in payload["nodes"]]

        self.assertEqual(titles, ["商务响应文件"])
        self.assertNotIn("技术方案", titles)
        self.assertNotIn("项目概况", titles)
        self.assertEqual(payload["reviewStatus"], "draft")

    def test_generate_outline_for_project_uses_opencode_direct_decision(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project_id = self._prepare_project_with_parse_result()

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=self._mock_futurecode_outline,
        ) as mock_generate:
            payload = generate_outline_for_project(project_id, {"outlineStrategy": "strict"})

        self.assertEqual(payload["status"], "completed")
        self.assertEqual(payload["output"]["chapterCount"], 3)
        self.assertTrue(payload["events"])
        self.assertEqual(payload["events"][-1]["level"], "success")
        self.assertTrue(mock_generate.called)
        self.assertEqual(mock_generate.call_args.kwargs["early_tool_command"], "s2outline-finalize")
        self.assertEqual(payload["opencodeOutput"]["status"], "received")

    def test_technical_outline_prompt_delegates_decisions_to_skill_without_duplicate_contract(self) -> None:
        from app.services.outline_generation import _build_outline_prompt

        prompt = _build_outline_prompt(Path("C:/workspace/s2_input.json"), "技术标")

        self.assertIn("Use the bid-tech-outline-generator skill", prompt)
        self.assertIn("next_cursor", prompt)
        self.assertIn("decision-next", prompt)
        self.assertIn("review-corrections", prompt)
        self.assertIn("review-complete", prompt)
        self.assertIn("decisions", prompt)
        self.assertIn("compose", prompt)
        self.assertIn("自主判断", prompt)
        self.assertIn("完整学习模板一至三级目录", prompt)
        self.assertIn("模板已有第三级目录统一进入结果供用户确认", prompt)
        self.assertIn("最终目录最多三级", prompt)
        self.assertIn("第四级及更深层级只作为对应第三级节点的内容参考", prompt)
        self.assertIn("再结合招标文件", prompt)
        self.assertIn("不得把未判断节点自动当成必要", prompt)
        self.assertIn("不得自行写入", prompt)
        self.assertIn("s2outline finalize", prompt)
        self.assertIn("timeout=300000", prompt)
        self.assertIn("不要检查脚本或包装器", prompt)
        self.assertNotIn("--max-items 50", prompt)
        self.assertNotIn("decision-context", prompt)
        self.assertNotIn("历史投标模板是主骨架", prompt)
        self.assertNotIn("仅因招标未提及", prompt)
        self.assertNotIn("没有不适用证据", prompt)
        self.assertNotIn("required_status", prompt)
        self.assertNotIn("粒度收敛和增删建议必须由 Opencode", prompt)
        self.assertNotIn("source_refs", prompt)
        self.assertLess(len(prompt), 1800)

    def test_technical_outline_fast_path_closes_without_llm_finalize_session(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project_id = self._prepare_project_with_parse_result()
        runner = load_technical_outline_runner()

        def fake_parallel_chapters(manifest_path, structure, *, progress_callback=None):
            # 模拟并行章节会话合并后的主工作区状态：机械完成全部模板章节决策
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            runner.write_template_structure(manifest, manifest_path)
            headings = runner.dispatch_command("headings", manifest, manifest_path, [])
            assert headings["complete"]
            while True:
                batch = runner.dispatch_command("decision-next", manifest, manifest_path, [])
                if batch["complete"]:
                    break
                runner.dispatch_command(
                    "decision-batch",
                    manifest,
                    manifest_path,
                    [
                        json.dumps(
                            {
                                "batch_token": batch["batch_token"],
                                "items": [
                                    {
                                        "target_id": item["target_id"],
                                        "decision": "retain",
                                        "reason": "历史模板中的专业目录仍适用于本项目",
                                    }
                                    for item in batch["items"]
                                ],
                                "additions": [],
                            },
                            ensure_ascii=False,
                        )
                    ],
                )
            return {
                "chapterSessionIds": ["ses-chapter-1"],
                "appendixPredecided": True,
                "appendixResult": {
                    "sessionId": "ses-appendix-1",
                    "opencodeOutput": {"sessionId": "ses-appendix-1"},
                },
            }

        with patch(
            "app.services.outline_generation._run_parallel_outline_chapters",
            side_effect=fake_parallel_chapters,
        ), patch(
            "app.services.outline_generation._run_outline_appendix_session",
        ) as mock_appendix, patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
        ) as mock_generate:
            payload = generate_outline_for_project(project_id, {"outlineStrategy": "strict"})

        # 快路径：不再启动串行 LLM 收口会话（含全局复核），由后端直跑受控收口
        mock_generate.assert_not_called()
        mock_appendix.assert_not_called()
        self.assertEqual(payload["status"], "completed")
        self.assertEqual(payload["output"]["chapterCount"], 2)
        self.assertEqual(
            payload["opencodeOutput"]["sessionIds"],
            ["ses-chapter-1", "ses-appendix-1"],
        )
        self.assertEqual(payload["opencodeOutput"]["chapterSessionCount"], 1)
        self.assertEqual(payload["opencodeOutput"]["parallelAppendixSessionCount"], 1)
        self.assertEqual(payload["opencodeOutput"]["parallelDecisionWorkers"], 2)

    def test_technical_outline_chapter_prompt_explains_sparse_chapter_mode(self) -> None:
        from app.services.outline_generation import _build_outline_chapter_prompt

        prompt = _build_outline_chapter_prompt(
            Path("C:/workspace/s2_input.json"),
            {"chapter_id": "TPL-0001", "number": "2", "title": "技术规范响应"},
        )

        self.assertIn("authoring_mode=sparse_chapter", prompt)
        self.assertIn("连续通读", prompt)
        self.assertIn("search 全会话最多 2 次", prompt)
        self.assertIn("一次提交全部新增", prompt)

    def test_technical_outline_finalize_prompt_resumes_after_template_decisions(self) -> None:
        from app.services.outline_generation import _build_outline_finalize_prompt

        prompt = _build_outline_finalize_prompt(Path("C:/workspace/s2_input.json"))

        self.assertIn("Use the bid-tech-outline-generator skill", prompt)
        self.assertIn("appendix-next", prompt)
        self.assertIn("s2outline headings", prompt)
        self.assertIn("s2outline section", prompt)
        self.assertIn("review-corrections", prompt)
        self.assertIn("review-complete", prompt)
        self.assertIn("s2outline decisions", prompt)
        self.assertIn("s2outline compose", prompt)
        self.assertIn("s2outline finalize", prompt)
        self.assertIn("不要执行 `s2outline prepare`", prompt)
        self.assertIn("`template-headings`", prompt)
        self.assertIn("`decision-next`", prompt)
        self.assertIn("outline_decision_state.json", prompt)
        self.assertIn("outline_authoring_decisions.json", prompt)
        self.assertIn("toc.json", prompt)
        self.assertIn("不要直接读取或修改", prompt)
        self.assertIn("只使用 `s2outline` 受控命令", prompt)
        self.assertNotIn("首次执行", prompt)
        self.assertLess(len(prompt), 1000)

    def test_technical_outline_prepares_isolated_workspaces_per_root_chapter(self) -> None:
        from app.services.outline_generation import (
            _capture_trusted_technical_outline_input,
            _prepare_outline_chapter_workspaces,
        )

        root = Path(self.temp_dir.name) / "parallel-chapters"
        root.mkdir(parents=True)
        template = root / "template.docx"
        tender = root / "tender.docx"
        manifest_path = root / "s2_input.json"
        self._write_docx(
            template,
            [
                ("第一章 总体方案", "Heading 1"),
                ("1.1 实施组织", "Heading 2"),
                ("第二章 质量安全", "Heading 1"),
                ("2.1 质量保证", "Heading 2"),
            ],
        )
        self._write_docx(
            tender,
            [
                ("1 技术要求", "Heading 1"),
                ("投标人应提交总体方案。", None),
            ],
        )
        manifest_path.write_text(
            json.dumps(
                {
                    "workDir": str(root),
                    "templateFile": str(template),
                    "tenderFiles": [{"id": "TEN-1", "name": tender.name, "path": str(tender)}],
                    "outputFile": str(root / "toc.json"),
                    "requireComposedOutline": True,
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        trusted = _capture_trusted_technical_outline_input(manifest_path)
        chapters, chapter_manifests, chapter_root, binding = (
            _prepare_outline_chapter_workspaces(
                manifest_path,
                trusted["templateStructure"],
            )
        )
        try:
            child_payloads = [
                json.loads(path.read_text(encoding="utf-8"))
                for path in chapter_manifests.values()
            ]
            self.assertEqual(len(chapters), 2)
            self.assertEqual(len({item["workDir"] for item in child_payloads}), 2)
            self.assertEqual(
                {item["_runtimeDecisionChapterId"] for item in child_payloads},
                set(chapter_manifests),
            )
            self.assertTrue(binding["headingsStateDigest"])
        finally:
            shutil.rmtree(chapter_root, ignore_errors=True)

    def test_technical_outline_uses_configured_opencode_instances_for_chapters(self) -> None:
        from app.services.outline_generation import (
            TECH_OUTLINE_CHAPTER_WORKERS,
            _outline_chapter_base_urls,
        )

        self.assertEqual(TECH_OUTLINE_CHAPTER_WORKERS, 6)

        with patch.dict(
            "os.environ",
            {
                "OPENCODE_CHAPTER_BASE_URLS": (
                    "http://opencode:4096, http://opencode-2:4096,"
                    "http://opencode-3:4096,http://opencode-4:4096,"
                    "http://opencode-5:4096,http://opencode-6:4096"
                )
            },
        ):
            self.assertEqual(
                _outline_chapter_base_urls(),
                [
                    "http://opencode:4096",
                    "http://opencode-2:4096",
                    "http://opencode-3:4096",
                    "http://opencode-4:4096",
                    "http://opencode-5:4096",
                    "http://opencode-6:4096",
                ],
            )

    def test_parallel_outline_chapters_include_appendix_as_seventh_worker(self) -> None:
        self.parallel_outline_patcher.stop()
        from app.services.outline_generation import (
            _TECH_OUTLINE_REQUEST_SLOTS,
            _run_parallel_outline_chapters,
        )

        root = Path(self.temp_dir.name) / "parallel-model-config"
        root.mkdir(parents=True)
        manifest_path = root / "s2_input.json"
        manifest_path.write_text("{}", encoding="utf-8")
        chapter_root = root / "chapters"
        chapter_root.mkdir()
        chapters = [
            {
                "chapter_id": f"TPL-{index:04d}",
                "number": f"第{index}章",
                "title": f"第{index}章",
            }
            for index in range(1, 7)
        ]
        chapter_manifests: dict[str, Path] = {}
        for chapter in chapters:
            chapter_dir = chapter_root / str(chapter["chapter_id"])
            chapter_dir.mkdir()
            path = chapter_dir / "s2_input.json"
            path.write_text(
                json.dumps({"workDir": str(chapter_dir)}, ensure_ascii=False),
                encoding="utf-8",
            )
            chapter_manifests[str(chapter["chapter_id"])] = path

        model_config = {
            "providerId": "configured-provider",
            "modelId": "configured-model",
            "timeoutMs": 45000,
        }
        runner = MagicMock()
        runner.decision_workflow.chapter_decision_progress.return_value = {
            "complete": True
        }
        runner.review_workflow.decision_appendix_items.return_value = [
            {
                "appendix_id": "APP-0001",
                "file_id": "TEN-1",
                "number": "附表A.1",
                "title": "技术参数表",
                "raw_text": "附表A.1 技术参数表",
                "following_table_count": 1,
                "source_status": "present",
            }
        ]
        runner.decision_workflow.appendix_decision_progress.return_value = {
            "complete": True,
            "decidedCount": 1,
            "remainingCount": 0,
        }
        worker_counts: list[int] = []

        def recording_executor(*, max_workers: int) -> RealThreadPoolExecutor:
            worker_counts.append(max_workers)
            return RealThreadPoolExecutor(max_workers=max_workers)

        with (
            patch(
                "app.services.outline_generation._prepare_outline_chapter_workspaces",
                return_value=(chapters, chapter_manifests, chapter_root, {}),
            ),
            patch(
                "app.services.outline_generation._load_technical_outline_runner",
                return_value=runner,
            ),
            patch(
                "app.services.outline_generation.system_settings_service.get_opencode_model_config_sync",
                return_value=model_config,
            ) as load_config,
            patch(
                "app.services.outline_generation._outline_chapter_base_urls",
                return_value=["http://opencode:4096"],
            ),
            patch(
                "app.services.outline_generation.ThreadPoolExecutor",
                side_effect=recording_executor,
            ),
            patch("app.services.outline_generation.OpencodeClient") as client_class,
        ):
            client_class.return_value.run_outline_decision_session.side_effect = [
                {"sessionId": f"ses-{index}", "opencodeOutput": {}}
                for index in range(1, 8)
            ]
            result = _run_parallel_outline_chapters(manifest_path, {})

        load_config.assert_called_once_with()
        self.assertEqual(worker_counts, [7])
        self.assertEqual(len(result["chapterSessionIds"]), 6)
        self.assertTrue(result["appendixPredecided"])
        self.assertTrue(result["appendixResult"]["sessionId"])
        self.assertEqual(client_class.call_count, 7)
        for call in client_class.call_args_list:
            self.assertIs(call.kwargs["model_config"], model_config)
            self.assertIs(call.kwargs["request_slots"], _TECH_OUTLINE_REQUEST_SLOTS)
        runner.decision_workflow.materialize_appendix_predecisions.assert_called_once()

    def test_technical_outline_handoff_stops_when_a_session_makes_no_decision_progress(self) -> None:
        from app.services.outline_generation import (
            _build_outline_handoff_prompt,
            _technical_outline_handoff_state,
        )

        manifest_path = Path("C:/workspace/s2_input.json")
        prompt = _build_outline_handoff_prompt(manifest_path, 1)
        self.assertIn("最多完成 1 个", prompt)
        resumed_prompt = _build_outline_handoff_prompt(manifest_path, 2)
        self.assertIn("第一条非 Skill 工具调用必须是 Bash", resumed_prompt)
        self.assertIn(
            f"s2outline decision-next {manifest_path}",
            resumed_prompt,
        )
        self.assertIn("禁止调用 Read、Glob、Grep", resumed_prompt)

        root = Path(self.temp_dir.name) / "handoff-no-progress"
        root.mkdir(parents=True, exist_ok=True)
        manifest_path = root / "s2_input.json"
        manifest_path.write_text(
            json.dumps({"workDir": str(root)}, ensure_ascii=False),
            encoding="utf-8",
        )
        runner = MagicMock()
        runner.dispatch_command.return_value = {
            "complete": False,
            "decided_count": 12,
            "remaining_count": 20,
        }

        with patch(
            "app.services.outline_generation._load_technical_outline_runner",
            return_value=runner,
        ):
            with self.assertRaisesRegex(RuntimeError, "没有提交新的目录判断"):
                _technical_outline_handoff_state(
                    manifest_path,
                    previous_decided_count=12,
                )

    def test_technical_outline_loader_does_not_trust_agent_modified_manifest_gate(self) -> None:
        from app.services.outline_generation import _load_outline_result

        root = Path(self.temp_dir.name) / "compose-gate"
        root.mkdir(parents=True, exist_ok=True)
        output = root / "toc.json"
        manifest_path = root / "s2_input.json"
        output.write_text(
            json.dumps(
                {
                    "schema_version": "technical-outline.v1",
                    "nodes": [
                        {
                            "number": "第1章",
                            "title": "技术方案",
                            "suggestion_action": "必要",
                            "suggestion_reason": "",
                            "children": [],
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        manifest_path.write_text(
            json.dumps(
                {
                    "bidType": "商务标",
                    "workDir": str(root),
                    "outputFile": str(output),
                    "requireComposedOutline": False,
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        with self.assertRaisesRegex(RuntimeError, "compose report"):
            _load_outline_result(
                {"outputFile": str(output)},
                manifest_path,
                expected_bid_type="技术标",
            )

        (root / "outline_compose_report.json").write_text(
            json.dumps(
                {
                    "schema_version": "technical-outline-compose-report.v1",
                    "outputFile": str(output),
                    "outputSha256": hashlib.sha256(output.read_bytes()).hexdigest(),
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        with self.assertRaisesRegex(RuntimeError, "compose report"):
            _load_outline_result(
                {"outputFile": str(output)},
                manifest_path,
                expected_bid_type="技术标",
            )

    def test_technical_outline_rejects_agent_rebased_artifacts_after_template_tamper(self) -> None:
        from app.services.outline_generation import _run_outline_skill

        root = Path(self.temp_dir.name) / "compose-template-tamper"
        root.mkdir(parents=True, exist_ok=True)
        template = root / "template.docx"
        tender = root / "tender.docx"
        output = root / "toc.json"
        manifest_path = root / "s2_input.json"
        self._write_docx(template, [("第一章 原始技术方案", "Heading 1")])
        self._write_docx(tender, [("第一章 招标技术要求", "Heading 1")])
        manifest = {
            "bidType": "技术标",
            "workDir": str(root),
            "templateFile": str(template),
            "tenderFiles": [{"id": "TEN-1", "name": tender.name, "path": str(tender)}],
            "outputFile": str(output),
            "requireComposedOutline": True,
        }
        manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
        runner = load_technical_outline_runner()

        def write_rebased_artifacts(_prompt: str, *args, **kwargs) -> dict:
            self._write_docx(template, [("第一章 被篡改的技术方案", "Heading 1")])
            runner.write_template_structure(manifest, manifest_path)
            headings = runner.dispatch_command("headings", manifest, manifest_path, [])
            self.assertTrue(headings["complete"])
            batch = runner.dispatch_command("decision-next", manifest, manifest_path, [])
            runner.dispatch_command(
                "section",
                manifest,
                manifest_path,
                [headings["files"][0]["items"][0]["section_id"]],
            )
            runner.dispatch_command(
                "decision-batch",
                manifest,
                manifest_path,
                [
                    json.dumps(
                        {
                            "batch_token": batch["batch_token"],
                            "items": [
                                {
                                    "target_id": item["target_id"],
                                    "decision": "retain",
                                    "reason": "历史模板中的专业目录仍适用于本项目",
                                }
                                for item in batch["items"]
                            ],
                            "additions": [],
                        },
                        ensure_ascii=False,
                    )
                ],
            )
            runner.dispatch_command(
                "section",
                manifest,
                manifest_path,
                [headings["files"][0]["items"][0]["section_id"]],
            )
            runner.dispatch_command(
                "review-complete",
                manifest,
                manifest_path,
                [json.dumps({"review_summary": "测试全局复核。", "issues": []}, ensure_ascii=False)],
            )
            runner.dispatch_command("decisions", manifest, manifest_path, [])
            return runner.compose_manifest(manifest, manifest_path)

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=write_rebased_artifacts,
        ):
            with self.assertRaisesRegex(RuntimeError, "模板|template"):
                _run_outline_skill(manifest_path, bid_type="技术标")

    def test_technical_outline_rejects_rebased_artifacts_after_tender_tamper(self) -> None:
        from app.services.outline_generation import _run_outline_skill

        root = Path(self.temp_dir.name) / "compose-tender-tamper"
        root.mkdir(parents=True, exist_ok=True)
        template = root / "template.docx"
        tender = root / "tender.docx"
        output = root / "toc.json"
        manifest_path = root / "s2_input.json"
        self._write_docx(template, [("第一章 技术方案", "Heading 1")])
        self._write_docx(tender, [("第一章 原始招标要求", "Heading 1")])
        manifest = {
            "bidType": "技术标",
            "workDir": str(root),
            "templateFile": str(template),
            "tenderFiles": [{"id": "TEN-1", "name": tender.name, "path": str(tender)}],
            "outputFile": str(output),
            "requireComposedOutline": True,
        }
        manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
        runner = load_technical_outline_runner()

        def write_rebased_artifacts(_prompt: str, *args, **kwargs) -> dict:
            self._write_docx(tender, [("第一章 被篡改的招标要求", "Heading 1")])
            runner.write_template_structure(manifest, manifest_path)
            headings = runner.dispatch_command("headings", manifest, manifest_path, [])
            batch = runner.dispatch_command("decision-next", manifest, manifest_path, [])
            runner.dispatch_command(
                "section",
                manifest,
                manifest_path,
                [headings["files"][0]["items"][0]["section_id"]],
            )
            runner.dispatch_command(
                "decision-batch",
                manifest,
                manifest_path,
                [
                    json.dumps(
                        {
                            "batch_token": batch["batch_token"],
                            "items": [
                                {
                                    "target_id": item["target_id"],
                                    "decision": "retain",
                                    "reason": "历史模板中的专业目录仍适用于本项目",
                                }
                                for item in batch["items"]
                            ],
                            "additions": [],
                        },
                        ensure_ascii=False,
                    )
                ],
            )
            runner.dispatch_command(
                "section",
                manifest,
                manifest_path,
                [headings["files"][0]["items"][0]["section_id"]],
            )
            runner.dispatch_command(
                "review-complete",
                manifest,
                manifest_path,
                [json.dumps({"review_summary": "测试全局复核。", "issues": []}, ensure_ascii=False)],
            )
            runner.dispatch_command("decisions", manifest, manifest_path, [])
            return runner.compose_manifest(manifest, manifest_path)

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=write_rebased_artifacts,
        ):
            with self.assertRaisesRegex(RuntimeError, "招标文件.*修改"):
                _run_outline_skill(manifest_path, bid_type="技术标")

    def test_technical_compose_trusted_validation_reuses_appendix_inventory(self) -> None:
        from app.services.outline_generation import (
            _capture_trusted_technical_outline_input,
            _run_outline_skill,
            _validate_technical_compose_report,
        )

        root = Path(self.temp_dir.name) / "compose-controlled-appendix"
        root.mkdir(parents=True, exist_ok=True)
        template = root / "template.docx"
        tender = root / "tender.docx"
        output = root / "toc.json"
        manifest_path = root / "s2_input.json"
        self._write_docx(template, [("第1章 技术方案", "Heading 1")])
        tender_doc = Document()
        tender_doc.add_paragraph("附表A.1 技术参数表")
        table = tender_doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "参数"
        table.cell(0, 1).text = "要求"
        table.cell(1, 0).text = "示例"
        table.cell(1, 1).text = "投标人填写"
        tender_doc.add_paragraph("附表A.2 供货范围表")
        table = tender_doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "设备"
        table.cell(0, 1).text = "数量"
        table.cell(1, 0).text = "示例"
        table.cell(1, 1).text = "投标人填写"
        tender_doc.save(tender)
        manifest = {
            "bidType": "技术标",
            "workDir": str(root),
            "templateFile": str(template),
            "tenderFiles": [{"id": "TEN-1", "name": tender.name, "path": str(tender)}],
            "outputFile": str(output),
            "requireComposedOutline": True,
        }
        manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
        runner = load_technical_outline_runner()
        trusted_input = _capture_trusted_technical_outline_input(manifest_path)

        def write_controlled_appendix_artifacts(_prompt: str, *args, **kwargs) -> dict:
            runner.write_template_structure(manifest, manifest_path)
            cursor = 0
            while True:
                headings = runner.dispatch_command(
                    "headings", manifest, manifest_path, ["--cursor", str(cursor)]
                )
                if headings["complete"]:
                    break
                cursor = int(headings["next_cursor"])

            while True:
                batch = runner.dispatch_command("decision-next", manifest, manifest_path, [])
                if batch["complete"]:
                    break
                appendix_evidence_id = runner.review_workflow.decision_appendix_items(root)[0][
                    "evidence_id"
                ]
                runner.dispatch_command(
                    "read",
                    manifest,
                    manifest_path,
                    [appendix_evidence_id],
                )
                runner.dispatch_command(
                    "decision-batch",
                    manifest,
                    manifest_path,
                    [
                        json.dumps(
                            {
                                "batch_token": batch["batch_token"],
                                "items": [
                                    {
                                        "target_id": item["target_id"],
                                        "decision": "retain",
                                        "reason": "历史模板中的专业目录仍适用于本项目",
                                    }
                                    for item in batch["items"]
                                ],
                                "additions": [],
                            },
                            ensure_ascii=False,
                        )
                    ],
                )

            appendix_batch = runner.dispatch_command(
                "appendix-next", manifest, manifest_path, []
            )
            self.assertEqual(len(appendix_batch["items"]), 2)
            runner.dispatch_command(
                "appendix-decision-batch",
                manifest,
                manifest_path,
                [
                    json.dumps(
                        {
                            "batch_token": appendix_batch["batch_token"],
                            "root_addition": {
                                "node_id": "ADD-TECH-APPENDIX",
                                "parent_id": None,
                                "number": "第2章",
                                "title": "技术附表",
                                "reason": "招标文件包含受控附表。",
                            },
                            "items": [
                                {
                                    "appendix_id": item["appendix_id"],
                                    "decision": "include",
                                    "reason": "招标文件要求投标人填写。",
                                    "node_id": f"ADD-APPENDIX-{index}",
                                    "parent_id": "ADD-TECH-APPENDIX",
                                }
                                for index, item in enumerate(
                                    appendix_batch["items"], start=1
                                )
                            ],
                        },
                        ensure_ascii=False,
                    )
                ],
            )
            runner.dispatch_command(
                "read",
                manifest,
                manifest_path,
                [appendix_batch["items"][0]["evidence_id"]],
            )
            runner.dispatch_command(
                "review-complete",
                manifest,
                manifest_path,
                [json.dumps({"review_summary": "测试全局复核。", "issues": []}, ensure_ascii=False)],
            )
            runner.dispatch_command("decisions", manifest, manifest_path, [])
            return runner.compose_manifest(manifest, manifest_path)

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=write_controlled_appendix_artifacts,
        ):
            result = _run_outline_skill(manifest_path, bid_type="技术标")

        self.assertEqual(result["nodes"][-1]["title"], "技术附表")
        self.assertEqual(len(result["nodes"][-1]["children"]), 2)

        inventory_path = root / "tender_appendix_inventory.json"
        inventory = json.loads(inventory_path.read_text(encoding="utf-8"))
        inventory["items"] = inventory["items"][:1]
        inventory_path.write_text(json.dumps(inventory, ensure_ascii=False), encoding="utf-8")
        chunks = json.loads((root / "tender_review_chunks.json").read_text(encoding="utf-8"))
        files_by_id, _ = runner.review_workflow._collect_heading_files(chunks)
        headings_state_path = root / "tender_headings_state.json"
        headings_state = json.loads(headings_state_path.read_text(encoding="utf-8"))
        headings_state["headings_catalog_digest"] = runner.review_workflow._heading_catalog_digest(
            files_by_id,
            inventory["items"],
        )
        headings_state["appendix_count"] = 1
        headings_state_path.write_text(
            json.dumps(headings_state, ensure_ascii=False), encoding="utf-8"
        )
        workflow_binding = runner.review_workflow.require_headings_complete(
            root,
            runner.tender_inputs(manifest),
        )
        state_path = root / "outline_decision_state.json"
        state = json.loads(state_path.read_text(encoding="utf-8"))
        current_appendix_items = runner.review_workflow.decision_appendix_items(root)
        state["appendix_inventory_digest"] = runner.decision_workflow._normalized_appendix_inventory(
            current_appendix_items
        )[1]
        state["appendix_decisions"] = {
            appendix_id: decision
            for appendix_id, decision in state["appendix_decisions"].items()
            if appendix_id == current_appendix_items[0]["appendix_id"]
        }
        state["additions"] = [
            addition
            for addition in state["additions"]
            if addition.get("node_id") in {"ADD-TECH-APPENDIX", "ADD-APPENDIX-1"}
        ]
        state["headings_state_digest"] = workflow_binding["headingsStateDigest"]
        state["finalized_decisions_digest"] = ""
        state_path.write_text(json.dumps(state, ensure_ascii=False), encoding="utf-8")
        runner.dispatch_command(
            "review-complete",
            manifest,
            manifest_path,
            [json.dumps({"review_summary": "测试全局复核。", "issues": []}, ensure_ascii=False)],
        )
        runner.dispatch_command("decisions", manifest, manifest_path, [])
        runner.compose_manifest(manifest, manifest_path)

        with self.assertRaisesRegex(RuntimeError, "附表清单.*可信快照"):
            _validate_technical_compose_report(
                root,
                output,
                trusted_input=trusted_input,
            )

    def test_generate_outline_rejects_agent_modified_manifest_before_publish(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project_id = self._prepare_project_with_parse_result()

        def write_valid_artifacts_then_mutate_manifest(prompt: str, *args, **kwargs) -> dict:
            manifest_path = self._technical_manifest_from_prompt(prompt, kwargs)
            result = self._write_mock_technical_toc_outputs(manifest_path)
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            manifest["outputFile"] = str(manifest_path.parent / "agent-selected-output.json")
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
            return result

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=write_valid_artifacts_then_mutate_manifest,
        ):
            with self.assertRaisesRegex(RuntimeError, "manifest.*修改"):
                generate_outline_for_project(project_id, {"outlineStrategy": "strict"})

    def test_publish_preserves_verified_technical_semantics_and_output_hash(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project_id = self._prepare_project_with_parse_result()
        staging_paths: list[str] = []

        def write_semantic_path_artifacts(prompt: str, *args, **kwargs) -> dict:
            manifest_path = self._technical_manifest_from_prompt(prompt, kwargs)
            staging_path = str(manifest_path.parent)
            staging_paths.append(staging_path)
            return self._write_mock_technical_toc_outputs(
                manifest_path,
                first_root_title=f"路径原文 {staging_path}",
            )

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=write_semantic_path_artifacts,
        ):
            payload = generate_outline_for_project(project_id, {"outlineStrategy": "strict"})

        expected_title = f"路径原文 {staging_paths[0]}"
        outline = self._outline_state_for_tests(project_id)
        self.assertIn(expected_title, [node["title"] for node in outline["nodes"]])
        published_work_dir = Path(payload["opencodeOutput"]["workDir"])
        output_file = Path(payload["opencodeOutput"]["tocJsonPath"])
        output_payload = json.loads(output_file.read_text(encoding="utf-8"))
        report = json.loads(
            (published_work_dir / "outline_compose_report.json").read_text(encoding="utf-8")
        )
        self.assertIn(expected_title, [node["title"] for node in output_payload["nodes"]])
        self.assertEqual(report["outputFile"], str(output_file))
        self.assertEqual(report["outputSha256"], hashlib.sha256(output_file.read_bytes()).hexdigest())
        self.assertTrue(report["tenderInputsDigest"])
        self.assertTrue(report["headingsStateDigest"])
        self.assertTrue(report["decisionStateDigest"])

    def test_clean_technical_outline_nodes_uses_frontend_contract(self) -> None:
        from app.services.outline_generation import _clean_technical_outline_nodes

        nodes = _clean_technical_outline_nodes(
            [
                {
                    "number": "5.8.1",
                    "title": "叶片专题",
                    "suggestion_action": "待确认",
                    "suggestion_reason": "确认专题归属。",
                    "tender_basis": {
                        "evidence_id": "TEN-1:B000123",
                        "file_id": "TEN-1",
                        "search_text": "投标人应编制叶片专题。",
                    },
                    "children": [],
                }
            ]
        )

        self.assertEqual(
            nodes,
            [
                {
                    "id": "OL-1",
                    "number": "5.8.1",
                    "tocNumber": "5.8.1",
                    "title": "叶片专题",
                    "suggestionAction": "待确认",
                    "suggestionReason": "确认专题归属。",
                    "tenderBasis": {
                        "evidenceId": "TEN-1:B000123",
                        "fileId": "TEN-1",
                        "searchText": "投标人应编制叶片专题。",
                    },
                    "children": [],
                }
            ],
        )

    def test_technical_nodes_keep_nested_number_title_and_actions(self) -> None:
        from app.services.outline_generation import _nodes_from_generation_result

        nodes = _nodes_from_generation_result(
            {
                "nodes": [
                    {
                    "number": "第一章",
                    "title": "标前概述",
                    "suggestion_action": "必要",
                    "suggestion_reason": "",
                    "children": [
                        {
                            "number": "1.1",
                            "title": "评分索引",
                            "suggestion_action": "建议增加",
                            "suggestion_reason": "招标要求独立提供评分索引。",
                            "children": [],
                        }
                    ],
                    }
                ]
            },
            compact_technical=True,
        )

        self.assertEqual(nodes[0]["tocNumber"], "第一章")
        self.assertEqual(nodes[0]["suggestionAction"], "必要")
        self.assertEqual(nodes[0]["children"][0]["id"], "OL-1-1")
        self.assertEqual(nodes[0]["children"][0]["suggestionAction"], "建议增加")

    def test_technical_generation_summary_uses_suggestion_action_counts(self) -> None:
        from app.services.outline_generation import _summary_from_generation_result

        summary = _summary_from_generation_result(
            {
                "summary": {
                    "total_nodes": 3,
                    "action_counts": {"必要": 2, "建议增加": 1},
                }
            }
        )

        self.assertEqual(summary, "目录生成完成，共 3 条目录项（必要2，建议增加1）。")

    def test_runtime_toc_recovery_understands_compact_technical_contract(self) -> None:
        from app.services.bid_runtime_state import outline_nodes_from_toc_items

        nodes = outline_nodes_from_toc_items(
            [
                {
                    "number": "附表G.2.3",
                    "title": "附表G.2.3 场址载荷对比",
                    "level": 1,
                    "required_status": "必要",
                    "review_required": False,
                    "review_note": "",
                    "source_refs": [],
                }
            ],
            compact_technical=True,
        )

        self.assertEqual(
            nodes,
            [
                {
                    "id": "OL-1",
                    "title": "场址载荷对比",
                    "children": [],
                    "tocNumber": "附表G.2.3",
                    "requiredStatus": "必要",
                    "reviewRequired": False,
                    "reviewNote": "",
                    "sourceRefs": [],
                }
            ],
        )

    def test_generate_outline_uses_s1_text_and_visual_template_for_non_docx_inputs(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project = store.create_project(
            {
                "name": "图片模板目录项目",
                "customerName": "测试业主",
                "bidType": "技术标",
            }
        )
        project_id = project["id"]
        project_dir = technical_workspace_dir(project_id)
        project_dir.mkdir(parents=True, exist_ok=True)

        tender_path = settings.uploads_dir / project_id / "tender" / "招标文件.png"
        tender_path.parent.mkdir(parents=True, exist_ok=True)
        tender_path.write_bytes(b"\x89PNG\r\n\x1a\nfake")
        template_path = settings.uploads_dir / project_id / "template" / "投标模板.png"
        template_path.parent.mkdir(parents=True, exist_ok=True)
        template_path.write_bytes(b"\x89PNG\r\n\x1a\nfake-template")

        combined_text_path = project_dir / "combined.txt"
        combined_text_path.write_text(
            "\n".join(
                [
                    "# 文件：招标文件.png",
                    "",
                    "第一章 采购需求",
                    "投标人应提供实施方案。",
                    "投标人须提交服务团队安排。",
                ]
            ),
            encoding="utf-8",
        )
        self._complete_parse_for_tests(
            project_id,
            tender_files=[
                {
                    "id": "TEN-1",
                    "name": "招标文件.png",
                    "path": str(tender_path),
                    "size_label": "1.0 MB",
                }
            ],
            template_files=[
                {
                    "id": "TPL-1",
                    "name": "投标模板.png",
                    "path": str(template_path),
                    "size_label": "1.0 MB",
                }
            ],
            summary={
                "fileCount": 1,
                "extractedCount": 1,
                "textLength": 80,
                "textPreview": "",
                "warnings": [],
            },
            parse_storage={
                "projectDir": str(project_dir),
                "combinedTextPath": str(combined_text_path),
                "manifestPath": "",
                "documents": [],
            },
        )

        with patch(
            "app.services.outline_generation._ocr_fallback_text",
            return_value=("第一章 投标响应概述\n1.1 项目理解\n第二章 实施方案\n2.1 工作计划", {"status": "completed"}),
        ), patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=self._mock_futurecode_outline,
        ):
            payload = generate_outline_for_project(project_id, {"outlineStrategy": "strict"})

        self.assertEqual(payload["status"], "completed")
        self.assertEqual(payload["output"]["chapterCount"], 3)
        workspace = self._directory_state_for_tests(project_id)["opencodeOutput"]["workDir"]
        manifest = json.loads((Path(workspace) / "s2_input.json").read_text(encoding="utf-8"))
        self.assertEqual(Path(manifest["templateFile"]).suffix, ".docx")
        self.assertEqual(Path(manifest["tenderFiles"][0]["path"]).suffix, ".docx")
        self.assertIs(manifest["requireComposedOutline"], True)
        compose_report = json.loads(
            (Path(workspace) / "outline_compose_report.json").read_text(encoding="utf-8")
        )
        self.assertEqual(compose_report["outputFile"], manifest["outputFile"])
        self.assertEqual(payload["opencodeOutput"]["engine"], "bid-tech-outline-generator")
        self.assertEqual(payload["opencodeOutput"]["skill"], "bid-tech-outline-generator")
        self.assertTrue(payload["opencodeOutput"]["parts"])
        self.assertTrue(Path(payload["opencodeOutput"]["tocJsonPath"]).exists())
        self.assertNotIn("evidencePath", payload["opencodeOutput"])
        self.assertEqual(
            payload["opencodeOutput"]["manifestPath"],
            payload["opencodeOutput"]["canonicalManifestPath"],
        )
        self.assertTrue(Path(payload["opencodeOutput"]["canonicalManifestPath"]).exists())
        self.assertFalse((settings.parsed_dir / project_id / "s2.json").exists())
        self.assertEqual(Path(payload["opencodeOutput"]["workDir"]).name, "s2_toc_workdir")

        outline = self._outline_state_for_tests(project_id)
        self.assertEqual(outline["reviewStatus"], "draft")
        self.assertEqual(len(outline["nodes"]), 3)
        self.assertEqual(outline["nodes"][0]["title"], "投标响应概述")
        self.assertEqual(outline["nodes"][1]["title"], "实施方案")
        self.assertEqual(outline["summary"]["totalNodeCount"], 6)

    def test_generate_outline_archives_previous_successful_workspace_on_success(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project_id = self._prepare_project_with_parse_result()

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=self._mock_futurecode_outline,
        ):
            first_payload = generate_outline_for_project(project_id, {"outlineStrategy": "strict"})

        project_dir = technical_workspace_dir(project_id)
        work_dir = project_dir / "s2_toc_workdir"
        marker = work_dir / "previous-marker.txt"
        marker.write_text("previous run", encoding="utf-8")

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=self._mock_futurecode_outline,
        ):
            second_payload = generate_outline_for_project(project_id, {"outlineStrategy": "strict"})

        self.assertTrue(Path(first_payload["opencodeOutput"]["tocJsonPath"]).exists())
        self.assertTrue(Path(second_payload["opencodeOutput"]["tocJsonPath"]).exists())
        self.assertFalse(marker.exists())
        archive_root = project_dir / "s2_toc_workdir.runs"
        archived_markers = list(archive_root.glob("*/previous-marker.txt"))
        self.assertEqual(len(archived_markers), 1)

        manifest = json.loads(Path(second_payload["opencodeOutput"]["canonicalManifestPath"]).read_text(encoding="utf-8"))
        self.assertEqual(manifest["workDir"], str(work_dir))
        self.assertFalse(any(".new" in str(value) for value in manifest.values() if isinstance(value, str)))

    def test_generate_outline_failure_preserves_previous_successful_workspace(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project_id = self._prepare_project_with_parse_result()

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=self._mock_futurecode_outline,
        ):
            first_payload = generate_outline_for_project(project_id, {"outlineStrategy": "strict"})

        work_dir = technical_workspace_dir(project_id) / "s2_toc_workdir"
        previous_toc = Path(first_payload["opencodeOutput"]["tocJsonPath"])
        marker = work_dir / "previous-marker.txt"
        marker.write_text("keep me", encoding="utf-8")

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=RuntimeError("futurecode down"),
        ), patch("app.services.outline_generation._run_local_outline_skill") as local_fallback:
            with self.assertRaisesRegex(RuntimeError, "目录生成需要 opencode 自主决策"):
                generate_outline_for_project(project_id, {"outlineStrategy": "strict"})
        local_fallback.assert_not_called()

        self.assertTrue(work_dir.exists())
        self.assertTrue(previous_toc.exists())
        self.assertEqual(marker.read_text(encoding="utf-8"), "keep me")
        self.assertTrue((technical_workspace_dir(project_id) / "s2_toc_workdir.new").exists())

    def test_publish_failure_preserves_previous_successful_workspace(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project_id = self._prepare_project_with_parse_result()

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=self._mock_futurecode_outline,
        ):
            first_payload = generate_outline_for_project(project_id, {"outlineStrategy": "strict"})

        project_dir = technical_workspace_dir(project_id)
        work_dir = project_dir / "s2_toc_workdir"
        previous_toc = Path(first_payload["opencodeOutput"]["tocJsonPath"])
        marker = work_dir / "previous-marker.txt"
        marker.write_text("keep me", encoding="utf-8")

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=self._mock_futurecode_outline,
        ), patch(
            "app.services.outline_generation._remap_json_file",
            side_effect=RuntimeError("remap failed"),
        ):
            with self.assertRaises(RuntimeError):
                generate_outline_for_project(project_id, {"outlineStrategy": "strict"})

        self.assertTrue(work_dir.exists())
        self.assertTrue(previous_toc.exists())
        self.assertEqual(marker.read_text(encoding="utf-8"), "keep me")
        self.assertTrue((project_dir / "s2_toc_workdir.new").exists())

    def test_generate_outline_outputs_compact_technical_contract(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project_id = self._prepare_project_with_parse_result()

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=self._mock_futurecode_outline,
        ):
            payload = generate_outline_for_project(project_id, {"outlineStrategy": "strict"})
        toc_path = Path(payload["opencodeOutput"]["tocJsonPath"])
        toc = __import__("json").loads(toc_path.read_text(encoding="utf-8"))

        self.assertEqual(toc["schema_version"], "technical-outline.v1")
        self.assertTrue(toc["nodes"])
        allowed_keys = {
            "number",
            "title",
            "suggestion_action",
            "suggestion_reason",
            "tender_basis",
            "children",
        }

        def assert_node_contract(nodes: list[dict]) -> None:
            for node in nodes:
                self.assertLessEqual(set(node), allowed_keys)
                self.assertTrue({"number", "title", "suggestion_action", "suggestion_reason", "children"} <= set(node))
                assert_node_contract(node["children"])

        assert_node_contract(toc["nodes"])

    def test_generate_outline_records_action_summary_without_evidence_file(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project_id = self._prepare_project_with_parse_result()
        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=self._mock_futurecode_outline,
        ):
            payload = generate_outline_for_project(project_id, {"outlineStrategy": "strict"})
        self.assertNotIn("evidencePath", payload["opencodeOutput"])
        self.assertEqual(
            payload["ruleEvidence"],
            {
                "schemaVersion": "technical-outline.v1",
                "engine": "bid-tech-outline-generator",
                "nodeCount": 6,
                "actionCounts": {"必要": 4, "建议增加": 2},
            },
        )

    def test_public_technical_evidence_v2_exposes_summary_only(self) -> None:
        from app.services.outline_generation import _public_rule_evidence_from_file

        evidence_path = Path(self.temp_dir.name) / "toc_evidence.json"
        evidence_path.write_text(
            json.dumps(
                {
                    "schema_version": "bid-toc-evidence-v2",
                    "engine": "bid-tech-outline-generator",
                    "ruleVersion": "opencode-autonomous-v2",
                    "decisions": [
                        {"action": "include", "review_required": False},
                        {"action": "merge", "review_required": True},
                        {"action": "exclude", "review_required": False},
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        public = _public_rule_evidence_from_file(evidence_path)

        self.assertEqual(
            public,
            {
                "schemaVersion": "bid-toc-evidence-v2",
                "engine": "bid-tech-outline-generator",
                "ruleVersion": "opencode-autonomous-v2",
                "decisionCount": 3,
                "reviewCount": 1,
                "actionCounts": {"include": 1, "merge": 1, "exclude": 1},
            },
        )

    def test_generate_outline_links_minimal_tender_basis_and_adds_appendix_items(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project_id = self._prepare_project_with_parse_result()
        tender_path = settings.uploads_dir / project_id / "tender" / "招标文件.docx"
        self._write_docx(
            tender_path,
            [
                ("第一章 采购需求", "Heading 1"),
                ("投标人应提供实施方案。", None),
                ("附表D.7 性能及考核承诺保证表", None),
            ],
        )
        tender_doc = Document(tender_path)
        appendix_table = tender_doc.add_table(rows=2, cols=2)
        appendix_table.cell(0, 0).text = "承诺事项"
        appendix_table.cell(0, 1).text = "投标响应"
        tender_doc.save(tender_path)

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=self._mock_futurecode_outline,
        ):
            payload = generate_outline_for_project(project_id, {"outlineStrategy": "strict"})
        toc = __import__("json").loads(Path(payload["opencodeOutput"]["tocJsonPath"]).read_text(encoding="utf-8"))

        implementation = next(item for item in toc["nodes"] if "实施方案" in item["title"])
        self.assertNotIn("tender_basis", implementation)
        self.assertEqual(implementation["suggestion_action"], "必要")

        appendix = toc["nodes"][-1]
        self.assertEqual(appendix["title"], "技术附表")
        self.assertEqual(appendix["children"][0]["number"], "附表D.7")
        self.assertEqual(
            appendix["children"][0]["tender_basis"],
            {
                "evidence_id": "TEN-1:B000003",
                "file_id": "TEN-1",
                "search_text": "附表D.7 性能及考核承诺保证表",
            },
        )
        self.assertEqual(appendix["children"][0]["children"], [])

    def test_generate_outline_fails_without_opencode_output_file(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project_id = self._prepare_project_with_parse_result()

        def _mock_missing_output(prompt: str, *args, **kwargs) -> dict:
            manifest_path = self._technical_manifest_from_prompt(prompt, kwargs)
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            return {
                "schema_version": "technical-outline.v1",
                "outputFile": manifest["outputFile"],
                "summary": {"total_nodes": 0},
                "opencodeOutput": {"status": "received", "parts": []},
            }

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=_mock_missing_output,
        ), patch("app.services.outline_generation._run_local_outline_skill") as local_fallback:
            with self.assertRaisesRegex(RuntimeError, "目录生成需要 opencode 自主决策"):
                generate_outline_for_project(project_id, {"outlineStrategy": "strict"})

        local_fallback.assert_not_called()

    def test_get_directory_state_keeps_generated_action_summary(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project_id = self._prepare_project_with_parse_result()
        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=self._mock_futurecode_outline,
        ):
            payload = generate_outline_for_project(project_id, {"outlineStrategy": "strict"})
        state = self._directory_state_for_tests(project_id)

        self.assertEqual(
            state["ruleEvidence"],
            {
                "schemaVersion": "technical-outline.v1",
                "engine": "bid-tech-outline-generator",
                "nodeCount": 6,
                "actionCounts": {"必要": 4, "建议增加": 2},
            },
        )
        self.assertNotIn("evidencePath", payload["opencodeOutput"])

    def test_run_directory_generation_returns_running_state_immediately(self) -> None:
        project_id = self._prepare_project_with_parse_result()

        with patch("app.services.bid_directory_flow._schedule_directory_generation_job"):
            response = self.client.post(
                f"/api/technical/projects/{project_id}/directory-generation/run",
                json={"outlineStrategy": "strict"},
            )

        self.assertEqual(response.status_code, 202)
        payload = response.json()
        self.assertEqual(payload["status"], "running")
        self.assertEqual(payload["tasks"][0]["status"], "running")
        self.assertEqual(payload["tasks"][1]["status"], "pending")

    def test_generate_outline_fails_when_template_is_missing(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project_id = self._prepare_project_with_parse_result()
        self._update_template_fallback_for_tests(project_id, {"enabled": False})
        parse_result, parse_storage = self._parse_runtime_state(project_id)
        template_upload_dir = settings.uploads_dir / project_id / "template"
        for path in template_upload_dir.glob("*"):
            if path.is_file():
                path.unlink()
        self._complete_parse_for_tests(
            project_id,
            tender_files=[
                {
                    "id": "TEN-1",
                    "name": "招标文件.docx",
                    "path": str(settings.uploads_dir / project_id / "tender" / "招标文件.docx"),
                    "size_label": "1.0 MB",
                }
            ],
            template_files=[],
            summary=parse_result["summary"],
            parse_storage=parse_storage,
        )

        with self.assertRaisesRegex(ValueError, "投标模板"):
            generate_outline_for_project(project_id, {"outlineStrategy": "strict"})

    def test_generate_outline_rejects_invalid_project_template_docx(self) -> None:
        from app.services.outline_generation import generate_outline_for_project

        project_id = self._prepare_project_with_parse_result()
        parse_result, parse_storage = self._parse_runtime_state(project_id)
        bad_template = settings.uploads_dir / project_id / "template" / "bad-template.docx"
        bad_template.parent.mkdir(parents=True, exist_ok=True)
        bad_template.write_bytes(b"bad docx")
        self._complete_parse_for_tests(
            project_id,
            tender_files=[
                {
                    "id": "TEN-1",
                    "name": "招标文件.docx",
                    "path": str(settings.uploads_dir / project_id / "tender" / "招标文件.docx"),
                    "size_label": "1.0 MB",
                }
            ],
            template_files=[
                {
                    "id": "TPL-1",
                    "name": "坏模板.docx",
                    "path": str(bad_template),
                    "size_label": "7 B",
                }
            ],
            summary=parse_result["summary"],
            parse_storage=parse_storage,
        )

        with self.assertRaisesRegex(ValueError, "不是有效 DOCX"):
            generate_outline_for_project(project_id, {"outlineStrategy": "strict"})

    def test_futurecode_progress_updates_before_completion(self) -> None:
        from app.services.bid_directory_flow import _handle_directory_progress

        project_id = self._prepare_project_with_parse_result()
        self._start_directory_generation_for_tests(project_id)

        _handle_directory_progress(
            project_id,
            "outline_session_ready",
            {
                "sessionId": "SESSION-1",
                "providerId": "opencode",
                "modelId": "big-pickle",
            },
        )

        state = self._directory_state_for_tests(project_id)
        self.assertEqual(state["status"], "running")
        self.assertEqual(state["percentage"], 8)
        self.assertEqual(state["opencodeOutput"]["status"], "waiting")
        self.assertEqual(state["opencodeOutput"]["sessionId"], "SESSION-1")
        self.assertEqual(state["events"][-1]["step"], "futurecode_session")
        self.assertIn("futurecode", state["events"][-1]["message"])

    def test_decision_progress_drives_percentage_counts_and_monotonicity(self) -> None:
        from app.services.bid_directory_flow import _handle_directory_progress

        project_id = self._prepare_project_with_parse_result()
        self._start_directory_generation_for_tests(project_id)

        started_state = self._directory_state_for_tests(project_id)
        self.assertTrue(started_state["startedAt"])
        self.assertEqual(started_state["percentage"], 2)

        _handle_directory_progress(
            project_id,
            "decision_progress",
            {"phase": "chapters", "decided": 0, "total": 240, "chaptersDone": 0, "chaptersTotal": 12},
        )
        state = self._directory_state_for_tests(project_id)
        self.assertEqual(state["percentage"], 5)
        self.assertEqual(state["decisionProgress"]["decided"], 0)
        self.assertEqual(state["decisionProgress"]["total"], 240)
        self.assertEqual(state["decisionProgress"]["chaptersTotal"], 12)

        _handle_directory_progress(
            project_id,
            "decision_progress",
            {"phase": "chapters", "decided": 120, "total": 240, "chaptersDone": 5, "chaptersTotal": 12},
        )
        state = self._directory_state_for_tests(project_id)
        self.assertEqual(state["percentage"], 42)
        self.assertEqual(state["decisionProgress"]["decided"], 120)
        self.assertIn("120/240", state["summary"])

        # 乱序到达的旧计数不能让百分比回退
        _handle_directory_progress(
            project_id,
            "decision_progress",
            {"phase": "chapters", "decided": 60, "total": 240},
        )
        state = self._directory_state_for_tests(project_id)
        self.assertEqual(state["percentage"], 42)

        _handle_directory_progress(
            project_id,
            "decision_progress",
            {"phase": "appendix", "decided": 6, "total": 12},
        )
        state = self._directory_state_for_tests(project_id)
        self.assertEqual(state["percentage"], 83)
        self.assertEqual(state["decisionProgress"]["phase"], "appendix")
        self.assertIn("附表", state["summary"])

    def test_concurrent_progress_updates_do_not_lose_events(self) -> None:
        """并行章节会话同时上报时，目录 state 的读改写不能互相覆盖。

        实测 PRJ-0002：6 个章节会话在同一秒上报「session 已建立」，events 只落了 1 条。
        目录 state 的写是「读快照 → 改 → 整份回写」，接 postgres 时读还要走一次库往返；
        这里在快照与回写之间注入等待，把那段真实窗口放大成确定性交错——去掉写锁本用例必失败。
        """
        import threading

        import app.services.bid_outline_state as outline_state
        from app.services.bid_directory_flow import _handle_directory_progress

        project_id = self._prepare_project_with_parse_result()
        self._start_directory_generation_for_tests(project_id)

        session_count = 6
        original_build_event = outline_state.build_directory_event
        arrival_lock = threading.Lock()
        arrived_count = [0]
        all_arrived = threading.Event()

        def delayed_build_event(*args: object, **kwargs: object) -> dict:
            # 本调用点位于「已读快照、尚未回写」之间：无锁时 6 个线程会在此集合，
            # 各自持有同一份旧快照，随后互相覆盖，只剩最后一次写入。
            # 有锁时写被串行化，凑不齐并发线程，各自等待很短的超时后继续。
            with arrival_lock:
                arrived_count[0] += 1
                if arrived_count[0] >= session_count:
                    all_arrived.set()
            all_arrived.wait(timeout=0.3)
            return original_build_event(*args, **kwargs)

        errors: list[BaseException] = []

        def report_session(index: int) -> None:
            try:
                _handle_directory_progress(
                    project_id,
                    "outline_session_ready",
                    {"sessionId": f"SESSION-{index}", "providerId": "opencode", "modelId": "big-pickle"},
                )
            except BaseException as exc:  # pragma: no cover - 只在回归时触发
                errors.append(exc)

        with patch.object(outline_state, "build_directory_event", delayed_build_event):
            threads = [threading.Thread(target=report_session, args=(index,)) for index in range(session_count)]
            for thread in threads:
                thread.start()
            for thread in threads:
                thread.join(timeout=30)

        self.assertEqual(errors, [])
        state = self._directory_state_for_tests(project_id)
        session_events = [
            event for event in state["events"] if event.get("step") == "futurecode_session"
        ]
        self.assertEqual(len(session_events), session_count)

    def test_chapter_stream_delta_suppresses_percentage_anchor(self) -> None:
        from app.services.bid_directory_flow import _handle_directory_progress

        project_id = self._prepare_project_with_parse_result()
        self._start_directory_generation_for_tests(project_id)

        _handle_directory_progress(
            project_id,
            "decision_progress",
            {"phase": "chapters", "decided": 30, "total": 240},
        )
        _handle_directory_progress(
            project_id,
            "outline_delta",
            {"status": "streaming", "parts": [{"type": "text", "text": "delta"}], "suppressPercentage": True},
        )

        state = self._directory_state_for_tests(project_id)
        self.assertEqual(state["percentage"], 14)
        self.assertNotIn("suppressPercentage", state["opencodeOutput"])
        self.assertIn("30/240", state["summary"])

    def test_background_job_updates_running_state_then_completes(self) -> None:
        from app.services.bid_directory_flow import _handle_directory_progress, _run_directory_generation_job

        project_id = self._prepare_project_with_parse_result()
        self._start_directory_generation_for_tests(project_id)

        _handle_directory_progress(
            project_id,
            "inputs_ready",
            {"tenderFileCount": 1, "templateFileCount": 1},
        )
        running_state = self._directory_state_for_tests(project_id)
        self.assertEqual(running_state["status"], "running")
        self.assertEqual(running_state["percentage"], 5)
        self.assertEqual(running_state["tasks"][1]["status"], "running")
        self.assertEqual(running_state["events"][-1]["step"], "hint_ready")
        self.assertEqual(running_state["opencodeOutput"]["status"], "idle")

        _handle_directory_progress(
            project_id,
            "outline_session_ready",
            {
                "sessionId": "SESSION-1",
                "providerId": "opencode",
                "modelId": "big-pickle",
            },
        )

        with patch(
            "app.services.opencode_client.OpencodeClient.generate_outline_with_trace",
            side_effect=self._mock_futurecode_outline,
        ):
            _run_directory_generation_job(project_id, {"outlineStrategy": "strict"})

        completed_state = self._directory_state_for_tests(project_id)
        self.assertEqual(completed_state["status"], "completed")
        self.assertEqual(completed_state["percentage"], 100)
        self.assertEqual(completed_state["tasks"][2]["status"], "done")
        self.assertEqual(completed_state["events"][-1]["level"], "success")
        self.assertEqual(completed_state["opencodeOutput"]["status"], "received")
        self.assertTrue(completed_state["opencodeOutput"]["parts"])

    def test_directory_generation_stream_returns_event_stream_payload(self) -> None:
        project_id = self._prepare_project_with_parse_result()
        project = store.require_project_for_update(project_id)
        complete_directory_generation_state(project, {})
        store.persist_project_state(project)

        with self.client.stream(
            "GET",
            f"/api/technical/projects/{project_id}/directory-generation/stream",
        ) as response:
            chunks = response.iter_text()
            first_chunk = next(chunks)

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.headers["content-type"].split(";")[0], "text/event-stream")
        self.assertIn('"status": "completed"', first_chunk)
        self.assertIn('"summary": "目录生成完成。"', first_chunk)


if __name__ == "__main__":
    unittest.main()
