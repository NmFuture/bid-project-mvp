from __future__ import annotations

import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from app.services.business_template_extractor import (
    build_business_template_navigation_prompt,
    run_business_template_extractor,
)


ROOT = Path(__file__).resolve().parents[1]
RUNNER = ROOT / "opencode" / "skills" / "bid-business-template-extractor" / "scripts" / "run_from_manifest.py"
SKILL_MD = ROOT / "opencode" / "skills" / "bid-business-template-extractor" / "SKILL.md"


def build_agentic_template_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("第一章 招标公告")
    doc.add_paragraph("这里不是模板。")
    doc.add_page_break()
    doc.add_paragraph("第六章 投标文件格式")
    doc.add_paragraph("投标函")
    doc.add_paragraph("致：招标人")
    doc.add_paragraph("投标人（盖章）：")
    doc.add_page_break()
    doc.add_paragraph("授权委托书")
    doc.add_paragraph("委托代理人姓名：")
    doc.add_paragraph("身份证号码：")
    doc.save(path)


def build_unfamiliar_template_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("商务响应文件编制样式")
    doc.add_paragraph("星河能源承诺页")
    doc.add_paragraph("投标人应在此填写承诺内容并盖章。")
    doc.save(path)


def write_manifest(path: Path, *, source: Path, output_dir: Path) -> None:
    path.write_text(
        json.dumps(
            {
                "schemaVersion": "bid-business-template-extractor-v1",
                "skillName": "bid-business-template-extractor",
                "projectId": "PRJ-AGENTIC",
                "outputDir": str(output_dir),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source.name,
                        "sourcePath": str(source),
                    }
                ],
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )


def run_btplnav(*args: object) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, str(RUNNER), *(str(arg) for arg in args)],
        cwd=str(ROOT),
        text=True,
        encoding="utf-8",
        errors="replace",
        capture_output=True,
        check=False,
    )


def stdout_json(completed: subprocess.CompletedProcess[str]) -> dict:
    if completed.returncode != 0:
        raise AssertionError(completed.stderr or completed.stdout)
    return json.loads(completed.stdout)


def find_block_id(blocks: list[dict], text: str) -> int:
    return next(int(block["blockId"]) for block in blocks if text in str(block.get("text") or ""))


class BusinessTemplateExtractorAgenticTests(unittest.TestCase):
    def test_btplnav_slices_ai_submitted_ranges_without_script_candidates(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "agentic.docx"
            output_dir = root / "out"
            manifest = root / "manifest.json"
            build_agentic_template_docx(source)
            write_manifest(manifest, source=source, output_dir=output_dir)

            prepared = stdout_json(run_btplnav("prepare", manifest))
            document_output = Path(prepared["documents"][0]["outputDir"])
            self.assertFalse((document_output / "candidate_templates.json").exists())
            self.assertFalse((document_output / "regions.json").exists())

            overview = stdout_json(run_btplnav("overview", manifest, "--page", "1", "--page-size", "40"))
            blocks = overview["blocks"]
            bid_letter_start = find_block_id(blocks, "投标函")
            bid_letter_end = find_block_id(blocks, "投标人（盖章）")
            auth_start = find_block_id(blocks, "授权委托书")
            auth_end = find_block_id(blocks, "身份证号码")

            submit = {
                "templates": [
                    {
                        "sourceDocumentId": "DOC-1",
                        "title": "投标函",
                        "templateType": "bid_letter",
                        "startBlockId": bid_letter_start,
                        "endBlockId": bid_letter_end,
                        "confidence": 0.96,
                        "reason": "AI 判断为需要投标人填写并盖章的投标函模板。",
                    },
                    {
                        "sourceDocumentId": "DOC-1",
                        "title": "授权委托书",
                        "templateType": "authorization_letter",
                        "startBlockId": auth_start,
                        "endBlockId": auth_end,
                        "confidence": 0.94,
                        "reason": "AI 判断为需要填写委托代理人信息的模板。",
                    },
                ]
            }
            stdout_json(run_btplnav("submit", manifest, "templates", json.dumps(submit, ensure_ascii=False)))
            validation = stdout_json(run_btplnav("validate", manifest))
            self.assertEqual(validation["status"], "passed")

            finalized = stdout_json(run_btplnav("finalize", manifest))
            self.assertEqual(finalized["summary"]["templateCount"], 2)
            payload = json.loads((output_dir / "business_template_extraction.json").read_text(encoding="utf-8"))
            self.assertEqual([item["title"] for item in payload["appendices"]], ["投标函", "授权委托书"])
            for appendix in payload["appendices"]:
                self.assertTrue(Path(appendix["docxPath"]).is_file())

    def test_btplnav_validator_is_structural_not_section_name_limited(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "unfamiliar.docx"
            output_dir = root / "out"
            manifest = root / "manifest.json"
            build_unfamiliar_template_docx(source)
            write_manifest(manifest, source=source, output_dir=output_dir)

            overview = stdout_json(run_btplnav("prepare", manifest))
            blocks = overview["documents"][0]["blocksPreview"]
            start = find_block_id(blocks, "星河能源承诺页")
            end = find_block_id(blocks, "投标人应在此填写")

            stdout_json(
                run_btplnav(
                    "submit",
                    manifest,
                    "templates",
                    json.dumps(
                        {
                            "templates": [
                                {
                                    "sourceDocumentId": "DOC-1",
                                    "title": "星河能源承诺页",
                                    "startBlockId": start,
                                    "endBlockId": end,
                                    "confidence": 0.88,
                                    "reason": "AI 语义判断该范围需要投标人填写并盖章。",
                                }
                            ]
                        },
                        ensure_ascii=False,
                    ),
                )
            )
            validation = stdout_json(run_btplnav("validate", manifest))

            self.assertEqual(validation["status"], "passed")
            self.assertEqual(validation["validationErrors"], [])

    def test_skill_document_keeps_agentic_flow_and_drops_candidate_batches(self) -> None:
        text = SKILL_MD.read_text(encoding="utf-8")

        self.assertIn("DOCX 导航阅读器", text)
        self.assertIn("Word 切片器", text)
        self.assertIn("结构校验器", text)
        self.assertIn("AI", text)
        self.assertIn("btplnav prepare", text)
        self.assertIn("btplnav submit", text)
        for forbidden in ("candidate-batch", "boundary-batch", "candidate_templates.json", "高召回候选"):
            self.assertNotIn(forbidden, text)

    def test_backend_prompt_delegates_navigation_to_ai_without_btplbound_batches(self) -> None:
        prompt = build_business_template_navigation_prompt(
            project_id="PRJ-AGENTIC",
            manifest_path=Path("/data/parsed/PRJ-AGENTIC/business_template_extraction_manifest.json"),
        )

        self.assertIn("btplnav prepare", prompt)
        self.assertIn("btplnav submit", prompt)
        self.assertIn("自主", prompt)
        self.assertIn("遵循 skill 中的执行流程", prompt)
        self.assertIn("逐个粗章节内部下钻", prompt)
        self.assertIn("独立编制任务", prompt)
        self.assertIn("父级集合回查", prompt)
        self.assertIn("只提取商务部分", prompt)
        self.assertIn("技术部分", prompt)
        self.assertIn("validate 只做结构校验", prompt)
        self.assertIn("不要额外套用未写在 skill 里的排除清单", prompt)
        self.assertNotIn("合同附件", prompt)
        self.assertNotIn("履约保证金格式", prompt)
        self.assertNotIn("目录页、目录清单", prompt)
        self.assertNotIn("普通说明", prompt)
        self.assertNotIn("第一个有意义标题", prompt)
        self.assertNotIn("归入父标题", prompt)
        self.assertNotIn("标段一", prompt)
        self.assertNotIn("标段二", prompt)
        self.assertNotIn("candidate-batch", prompt)
        self.assertNotIn("boundary-batch", prompt)
        self.assertNotIn("btplbound", prompt)

    def test_backend_uses_single_agentic_navigation_run_and_never_drives_batches(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            project_dir = Path(tmp)
            source = project_dir / "source.docx"
            build_agentic_template_docx(source)
            output_dir = project_dir / "business_template_extraction"

            def fake_agent(_client, prompt: str):  # type: ignore[no-untyped-def]
                self.assertIn("btplnav finalize", prompt)
                output_dir.mkdir(parents=True, exist_ok=True)
                (output_dir / "business_template_extraction.json").write_text(
                    json.dumps(
                        {
                            "schemaVersion": "bid-business-template-extractor-v1",
                            "skillName": "bid-business-template-extractor",
                            "projectId": "PRJ-AGENTIC",
                            "outputDir": str(output_dir),
                            "appendices": [
                                {
                                    "id": "APPX-0001",
                                    "title": "投标函",
                                    "artifactType": "business_attachment_template",
                                    "templateType": "bid_letter",
                                    "docxPath": str(output_dir / "DOC-1" / "templates" / "TPL-0001-投标函.docx"),
                                    "sourceDocumentId": "DOC-1",
                                    "sourceDocumentName": source.name,
                                }
                            ],
                            "warnings": [],
                            "quality": {"scriptFallbackUsed": False, "acceptedTemplateCount": 1},
                            "summary": {"templateCount": 1},
                        },
                        ensure_ascii=False,
                    ),
                    encoding="utf-8",
                )
                return {
                    "schemaVersion": "bid-business-template-extractor-v1",
                    "outputFile": str(output_dir / "business_template_extraction.json"),
                    "summary": {"templateCount": 1},
                    "opencodeOutput": {"sessionId": "ses-btplnav", "completionSource": "btplnav-finalize"},
                }

            with (
                patch("app.services.business_template_extractor.OpencodeClient.extract_business_templates_with_trace", new=fake_agent),
            ):
                appendices, payload, warning = run_business_template_extractor(
                    project_id="PRJ-AGENTIC",
                    documents=[{"id": "DOC-1", "name": source.name, "sourcePath": str(source)}],
                    project_dir=project_dir,
                )

            self.assertEqual(warning, "")
            self.assertEqual(len(appendices), 1)
            self.assertEqual(appendices[0]["title"], "投标函")
            self.assertEqual((payload or {})["opencodeOutput"]["completionSource"], "btplnav-finalize")


if __name__ == "__main__":
    unittest.main()
