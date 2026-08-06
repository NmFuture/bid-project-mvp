from __future__ import annotations

import copy
import base64
import json
import tempfile
import unittest
from pathlib import Path
from urllib.parse import quote
from unittest.mock import patch

from docx import Document
from fastapi.testclient import TestClient
import openpyxl

from app.main import app
from app.core.config import settings
from app.services import technical_gap_ai_fill
from app.services.bid_outline_state import confirm_outline_state, save_generated_outline_state
from app.services.technical_gap_domain import aggregate_technical_gap_fill_quality
from app.services.bid_runtime_state import count_outline_nodes, now_iso, outline_nodes_from_toc_items
from app.services.store import store
from app.services.technical_fact_field_specs import fillable_specs
from app.services.technical_fact_spec_import import EXPECTED_HEADER
from app.services.workspace_artifacts import technical_workspace_dir


def _test_fact_specs() -> dict:
    """build_facts 门控 seed：测试绕过实时表上传，直接注入全局字段清单作为项目 specs。"""
    return {
        "fileName": "测试实时表.xlsx",
        "uploadedAt": "2026-07-27T00:00:00",
        "specs": copy.deepcopy(fillable_specs()),
    }


def _seed_fact_specs(project_id: str) -> None:
    """gaps-detection 等链路重建的 gap_state 没有 factSpecs，补 seed 后再 build。"""
    project = store._require(project_id)
    project.setdefault("gap_state", {})["factSpecs"] = _test_fact_specs()
    store._persist_project(project)


def minimal_gap_plan_from_manifest(manifest: dict) -> dict:
    toc = json.loads(Path(manifest["tocJsonPath"]).read_text(encoding="utf-8"))
    items = [
        {
            "id": f"GAP-{index:04d}",
            "number": str(item.get("number") or ""),
            "title": str(item.get("title") or f"目录项-{index}"),
            "level": int(item.get("level") or 1),
            "status": "missing",
            "decision": "material_required",
            "matchedMaterials": [],
            "candidateMaterials": [],
            "fillTasks": [],
            "resolvedArtifacts": [],
        }
        for index, item in enumerate(toc.get("items") or [], start=1)
        if isinstance(item, dict)
    ]
    count = len(items)
    return {
        "schemaVersion": "bid-tech-gap-plan-v1",
        "projectId": str(manifest.get("projectId") or ""),
        "status": "ready",
        "summary": {
            "totalTocItems": count,
            "matchedCount": 0,
            "missingCount": count,
            "resolvedCount": 0,
            "ignoredCount": 0,
            "structuralCount": 0,
            "fillableTaskCount": 0,
            "blockingCount": count,
        },
        "items": items,
    }


def _save_generated_outline_for_tests(
    project_id: str,
    *,
    nodes: list[dict],
    generated_at: str,
    summary: str,
) -> dict:
    project = store.require_project_for_update(project_id)
    payload = save_generated_outline_state(
        project,
        nodes=nodes,
        generated_at=generated_at,
        summary=summary,
    )
    store.persist_project_state(project)
    return payload


def _confirm_outline_for_tests(project_id: str) -> dict:
    project = store.require_project_for_update(project_id)
    payload = confirm_outline_state(project)
    store.persist_project_state(project)
    return payload


def _replace_confirmed_outline_from_toc(project_id: str, toc: dict) -> None:
    project = store.require_project_for_update(project_id)
    nodes = outline_nodes_from_toc_items(toc.get("items") or [])
    project["outline_state"]["nodes"] = nodes
    project["outline_state"]["summary"] = {"totalNodeCount": count_outline_nodes(nodes)}
    project["outline_state"]["reviewStatus"] = "confirmed"
    store.persist_project_state(project)


class GapReviewFlowTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        base = Path(self.temp_dir.name)
        settings.uploads_dir = base / "uploads"
        settings.documents_dir = base / "documents"
        settings.parsed_dir = base / "parsed"
        settings.ensure_dirs()

        store.reset_for_tests()
        self.client = TestClient(app, base_url="http://127.0.0.1:8000")
        # 缺口识别已改为直跑本地脚本（不再经 OpenCode），原先「打桩逼出本地 fallback」
        # 的 OpencodeClient 打桩随之失效且不再需要。

    def tearDown(self) -> None:
        self.client.close()
        self.temp_dir.cleanup()

    def _create_project_with_confirmed_outline(self) -> str:
        response = self.client.post(
            "/api/technical/projects",
            json={
                "name": "S4-S6联调项目",
                "customerName": "测试业主",
            },
        )
        response.raise_for_status()
        project_id = response.json()["id"]

        _save_generated_outline_for_tests(
            project_id=project_id,
            nodes=[
                {
                    "id": "OL-1",
                    "title": "第1章 标前概述",
                    "children": [
                        {"id": "OL-1-1", "title": "技术评分标准索引表", "children": []},
                        {"id": "OL-1-2", "title": "投标方案优势说明", "children": []},
                    ],
                },
                {
                    "id": "OL-2",
                    "title": "第2章 技术标准",
                    "children": [
                        {"id": "OL-2-1", "title": "性能保证", "children": []},
                    ],
                },
                {
                    "id": "OL-3",
                    "title": "第3章 风资源评估与机位排布方案",
                    "children": [],
                },
            ],
            generated_at=now_iso(),
            summary="目录已生成。",
        )
        _confirm_outline_for_tests(project_id)
        return project_id

    def _create_project_with_confirmed_directory_json(self) -> str:
        project_id = self._create_project_with_confirmed_outline()
        project_dir = technical_workspace_dir(project_id)
        work_dir = project_dir / "s2_toc_workdir"
        work_dir.mkdir(parents=True, exist_ok=True)
        toc_json = work_dir / "投标文件-总目录.json"
        toc_json.write_text(
            json.dumps(
                {
                    "schema_version": "bid-toc-json-v1",
                    "document_title": "S4-S6联调项目投标文件总目录",
                    "project": {
                        "owner": "测试业主",
                        "name": "S4-S6联调项目",
                        "code": project_id,
                    },
                    "items": [
                        {
                            "order": 1,
                            "number": "1",
                            "title": "标前概述",
                            "level": 1,
                            "annotation": "保留",
                            "source": "template",
                            "reason": "",
                            "material_refs": [],
                        },
                        {
                            "order": 2,
                            "number": "1.1",
                            "title": "技术评分标准索引表",
                            "level": 2,
                            "annotation": "适配",
                            "source": "wiki",
                            "reason": "匹配素材库：评分索引",
                            "material_refs": [
                                {
                                    "id": "RAW-0001",
                                    "docx": "技术标/通用素材/技术评分标准索引表.docx",
                                    "usage": "both",
                                }
                            ],
                        },
                        {
                            "order": 3,
                            "number": "2.1",
                            "title": "性能保证",
                            "level": 2,
                            "annotation": "新增-招标要求",
                            "source": "tender_special",
                            "reason": "招标文件要求填写性能保证附表",
                            "material_refs": [],
                        },
                    ],
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        _save_generated_outline_for_tests(
            project_id=project_id,
            nodes=[
                {
                    "id": "OL-1",
                    "tocNumber": "1",
                    "title": "标前概述",
                    "source": "template",
                    "children": [
                        {
                            "id": "OL-1-1",
                            "tocNumber": "1.1",
                            "title": "技术评分标准索引表",
                            "annotation": "适配",
                            "source": "wiki",
                            "reason": "匹配素材库：评分索引",
                            "materialRefs": [
                                {
                                    "id": "RAW-0001",
                                    "docx": "技术标/通用素材/技术评分标准索引表.docx",
                                    "usage": "both",
                                }
                            ],
                            "children": [],
                        },
                        {
                            "id": "OL-1-2",
                            "tocNumber": "2.1",
                            "title": "性能保证",
                            "annotation": "新增-招标要求",
                            "source": "tender_special",
                            "reason": "招标文件要求填写性能保证附表",
                            "children": [],
                        },
                    ],
                }
            ],
            generated_at=now_iso(),
            summary="目录已生成，共 3 条目录项。",
        )
        _confirm_outline_for_tests(project_id)
        parse_storage = copy.deepcopy(store.get_project_runtime_state(project_id)["parse_storage"])
        parse_storage["projectDir"] = str(project_dir)
        project = store._require(project_id)
        project["parse_storage"] = parse_storage
        project["parse_result"]["structured"] = {
            "appendices": [
                {
                    "id": "APP-PERF",
                    "title": "性能保证附表",
                    "sourceFile": "招标文件.docx",
                    "blankTable": True,
                }
            ],
            "projectDates": {"startDate": "2026-06-01", "endDate": "2026-09-30"},
        }
        store._persist_project(project)
        return project_id

    def test_gap_detection_uses_confirmed_outline_without_deduplicating_numbers(self) -> None:
        project_id = self._create_project_with_confirmed_outline()
        project = store.require_project_for_update(project_id)
        project["outline_state"]["nodes"] = [
            {
                "id": "OL-B8-1",
                "tocNumber": "附表B.8",
                "title": "附表B.8 出质保后备品备件服务",
                "source": "tender",
                "children": [],
            },
            {
                "id": "OL-B8-2",
                "tocNumber": "附表B.8",
                "title": "附表B.8 出质保后备品备件服务 无",
                "source": "tender",
                "children": [],
            },
        ]
        project["outline_state"]["summary"] = {"totalNodeCount": 2}
        stale_toc = technical_workspace_dir(project_id) / "s2_toc_workdir" / "toc.json"
        stale_toc.parent.mkdir(parents=True, exist_ok=True)
        stale_toc.write_text(
            json.dumps(
                {
                    "schema_version": "bid-toc-json-v1",
                    "items": [
                        {
                            "number": "附表B.8",
                            "title": "出质保后备品备件服务",
                            "level": 1,
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        project["directory_state"]["opencodeOutput"]["tocJsonPath"] = str(stale_toc)
        store.persist_project_state(project)

        captured_toc: list[dict] = []

        def fake_gap_planner(manifest_path):
            manifest = json.loads(Path(manifest_path).read_text(encoding="utf-8"))
            toc = json.loads(Path(manifest["tocJsonPath"]).read_text(encoding="utf-8"))
            captured_toc.extend(toc["items"])
            output_file = Path(manifest["outputFile"])
            plan = minimal_gap_plan_from_manifest(manifest)
            output_file.write_text(json.dumps(plan, ensure_ascii=False), encoding="utf-8")
            return {"schema_version": "bid-tech-gap-plan-v1", "outputFile": str(output_file)}

        with patch("app.services.technical_gap_planner.run_technical_gap_planner_skill", side_effect=fake_gap_planner):
            response = self.client.post(f"/api/technical/projects/{project_id}/gaps-detection/run")

        self.assertEqual(response.status_code, 200, response.text)
        self.assertEqual(response.json()["summary"]["totalTocItems"], 2)
        self.assertEqual([item["number"] for item in captured_toc], ["附表B.8", "附表B.8"])
        self.assertEqual(
            [item["title"] for item in captured_toc],
            ["出质保后备品备件服务", "出质保后备品备件服务 无"],
        )
        self.assertEqual(len(json.loads(stale_toc.read_text(encoding="utf-8"))["items"]), 1)

    def _confirm_project_fact_table(self, project_id: str, extra_values: dict[str, str] | None = None) -> dict:
        _seed_fact_specs(project_id)
        build_response = self.client.post(f"/api/technical/projects/{project_id}/gaps/facts/build")
        self.assertEqual(build_response.status_code, 200, build_response.text)
        fields = build_response.json()["fields"]
        for field in fields:
            label = str(field.get("label") or "")
            if extra_values and label in extra_values:
                field["value"] = extra_values[label]
            if str(field.get("value") or "").strip():
                field["status"] = "confirmed"
        confirm_response = self.client.put(
            f"/api/technical/projects/{project_id}/gaps/facts",
            json={"fields": fields, "confirm": True, "operator": "测试用户"},
        )
        self.assertEqual(confirm_response.status_code, 200, confirm_response.text)
        self.assertEqual(confirm_response.json()["status"], "confirmed")
        return confirm_response.json()

    def test_gap_detection_creates_real_gap_plan_from_directory_material_refs_and_parse_appendices(self) -> None:
        project_id = self._create_project_with_confirmed_directory_json()

        response = self.client.post(f"/api/technical/projects/{project_id}/gaps-detection/run")

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["status"], "completed")
        self.assertIn("gapPlan", payload)
        self.assertEqual(payload["gapPlan"]["schemaVersion"], "bid-tech-gap-plan-v1")
        self.assertEqual(payload["summary"]["totalTocItems"], 3)
        self.assertEqual(payload["summary"]["matchedCount"], 1)
        self.assertEqual(payload["summary"]["missingCount"], 1)
        self.assertGreaterEqual(payload["summary"]["fillableTaskCount"], 1)
        matched = next(item for item in payload["gapPlan"]["items"] if item["status"] == "matched")
        self.assertEqual(matched["matchedMaterials"][0]["id"], "RAW-0001")
        missing = next(item for item in payload["gapPlan"]["items"] if item["status"] == "needs_input")
        self.assertEqual(missing["number"], "2.1")
        self.assertEqual(missing["fillTasks"][0]["skill"], "bid-tech-table-filler")
        self.assertTrue(payload["gapPlan"]["planFile"].endswith("gap_plan.json"))

    def test_gap_detection_rejects_shrunken_skill_output(self) -> None:
        project_id = self._create_project_with_confirmed_directory_json()

        def fake_gap_planner(manifest_path):
            manifest = json.loads(Path(manifest_path).read_text(encoding="utf-8"))
            output_file = Path(manifest["outputFile"])
            output_file.parent.mkdir(parents=True, exist_ok=True)
            plan = {
                "schemaVersion": "bid-tech-gap-plan-v1",
                "projectId": project_id,
                "status": "ready",
                "summary": {
                    "totalTocItems": 1,
                    "matchedCount": 1,
                    "missingCount": 0,
                    "resolvedCount": 0,
                    "ignoredCount": 0,
                    "structuralCount": 0,
                    "fillableTaskCount": 0,
                    "blockingCount": 0,
                },
                "items": [
                    {
                        "id": "GAP-0001",
                        "number": "1",
                        "title": "标前概述",
                        "level": 1,
                        "status": "matched",
                        "decision": "ready",
                        "matchedMaterials": [],
                        "candidateMaterials": [],
                        "fillTasks": [],
                        "resolvedArtifacts": [],
                    }
                ],
            }
            output_file.write_text(json.dumps(plan, ensure_ascii=False, indent=2), encoding="utf-8")
            return {"schema_version": "bid-tech-gap-plan-v1", "outputFile": str(output_file)}

        with patch("app.services.technical_gap_planner.run_technical_gap_planner_skill", side_effect=fake_gap_planner):
            response = self.client.post(f"/api/technical/projects/{project_id}/gaps-detection/run")

        self.assertEqual(response.status_code, 400, response.text)
        self.assertIn("缺口识别结果不完整", response.json()["detail"])

    def test_gap_detection_rerun_returns_clean_first_step_plan(self) -> None:
        project_id = self._create_project_with_confirmed_directory_json()
        detection_response = self.client.post(f"/api/technical/projects/{project_id}/gaps-detection/run")
        self.assertEqual(detection_response.status_code, 200)
        gap_plan = detection_response.json()["gapPlan"]
        fill_item = next(item for item in gap_plan["items"] if item["fillTasks"])
        fill_item["fillTasks"][0]["status"] = "completed"
        fill_item["fillTasks"][0]["outputArtifactId"] = "ART-OLD"
        fill_item["resolvedArtifacts"] = [
            {
                "id": "ART-OLD",
                "source": "ai_fill",
                "skill": "bid-tech-table-filler",
                "fileName": "旧填写产物.docx",
                "s7Ready": True,
            }
        ]
        project = store._require(project_id)
        project["gap_state"]["plan"] = gap_plan
        project["gap_state"]["integrity"] = {
            "status": "passed",
        }
        project["gap_state"]["submissions"] = [{"id": "SUB-OLD"}]
        store._persist_project(project)

        response = self.client.post(f"/api/technical/projects/{project_id}/gaps-detection/run")

        self.assertEqual(response.status_code, 200, response.text)
        payload = response.json()
        next_plan = payload["gapPlan"]
        self.assertEqual(next_plan.get("phase"), "gap_detection")
        self.assertEqual(next_plan.get("integrity"), {})
        next_fill_item = next(item for item in next_plan["items"] if item["id"] == fill_item["id"])
        self.assertEqual(next_fill_item["resolvedArtifacts"], [])
        self.assertEqual(next_fill_item["fillTasks"][0]["status"], "pending")
        self.assertNotIn("outputArtifactId", next_fill_item["fillTasks"][0])
        self.assertEqual(payload["integrity"], {})

    def test_gap_ai_fill_calls_opencode_skill_and_registers_resolved_artifact(self) -> None:
        project_id = self._create_project_with_confirmed_directory_json()
        detection_response = self.client.post(f"/api/technical/projects/{project_id}/gaps-detection/run")
        self.assertEqual(detection_response.status_code, 200)
        gap_plan = detection_response.json()["gapPlan"]
        fill_item = next(item for item in gap_plan["items"] if item["fillTasks"])
        gap_id = fill_item["id"]
        fill_task_id = fill_item["fillTasks"][0]["id"]
        self._confirm_project_fact_table(project_id)

        def fake_run_table_filler(manifest_path, progress_callback=None):
            manifest = json.loads(Path(manifest_path).read_text(encoding="utf-8"))
            output_file = Path(manifest["outputFile"])
            output_file.parent.mkdir(parents=True, exist_ok=True)
            doc = Document()
            doc.add_paragraph("性能保证附表")
            doc.add_paragraph("已根据参考素材和解析字段填写。")
            doc.save(output_file)
            return {
                "schema_version": "bid-tech-table-fill-v1",
                "outputFile": str(output_file),
                "unfilledFields": ["保证值来源页码"],
                "evidenceRefs": [{"source": "招标文件.docx", "field": "性能保证"}],
                "opencodeOutput": {
                    "status": "received",
                    "sessionId": str(manifest_path),
                    "providerId": "local-skill",
                    "modelId": "bid-tech-table-filler",
                    "receivedAt": "2026-05-02T00:00:00Z",
                    "parts": [{"type": "text", "text": "{\"schema_version\":\"bid-tech-table-fill-v1\"}"}],
                },
            }

        with patch(
            "app.services.technical_gap_ai_fill.run_technical_table_filler_skill",
            side_effect=fake_run_table_filler,
        ):
            response = self.client.post(
                f"/api/technical/projects/{project_id}/gaps/{gap_id}/ai-fill",
                json={
                    "fillTaskId": fill_task_id,
                    "referenceMaterialIds": ["RAW-0001"],
                    "parseFieldIds": ["APP-PERF"],
                    "operator": "测试用户",
                },
            )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["item"]["status"], "needs_input")
        self.assertEqual(payload["item"]["decision"], "review_required")
        self.assertEqual(payload["artifact"]["source"], "ai_fill")
        self.assertEqual(payload["artifact"]["skill"], "bid-tech-table-filler")
        self.assertFalse(payload["artifact"]["s7Ready"])
        self.assertEqual(payload["artifact"]["referenceMaterialIds"], ["RAW-0001"])
        self.assertIn("AI 填写仍有未填字段：1 项", payload["item"]["reviewNotes"])
        self.assertTrue(Path(payload["artifact"]["path"]).exists())
        self.assertIn("onlyoffice", payload["artifact"])
        encoded_name = quote(payload["artifact"]["fileName"])
        self.assertIn(encoded_name, payload["artifact"]["onlyoffice"]["fileUrl"])
        self.assertIn(encoded_name, payload["artifact"]["onlyoffice"]["browserFileUrl"])
        self.assertIn(encoded_name, payload["artifact"]["onlyoffice"]["documentServerFileUrl"])
        self.assertNotIn(payload["artifact"]["fileName"], payload["artifact"]["onlyoffice"]["fileUrl"])
        self.assertEqual(
            payload["artifact"]["onlyoffice"]["fileUrl"],
            payload["artifact"]["onlyoffice"]["documentServerFileUrl"],
        )
        self.assertTrue(payload["artifact"]["onlyoffice"]["browserFileUrl"].startswith("http://127.0.0.1:8000/"))
        updated_gap = self.client.get(f"/api/technical/projects/{project_id}/gaps").json()["gapPlan"]
        updated_item = next(item for item in updated_gap["items"] if item["id"] == gap_id)
        self.assertEqual(updated_item["resolvedArtifacts"][0]["source"], "ai_fill")
        self.assertEqual(updated_item["resolvedArtifacts"][0]["referenceMaterialIds"], ["RAW-0001"])
        self.assertEqual(updated_item["resolvedArtifacts"][0]["referenceMaterials"][0]["id"], "RAW-0001")
        self.assertEqual(updated_item["fillTasks"][0]["status"], "completed")

        confirm_response = self.client.post(
            f"/api/technical/projects/{project_id}/gaps/{gap_id}/artifacts/{payload['artifact']['id']}/confirm",
            json={"operator": "复核用户"},
        )
        self.assertEqual(confirm_response.status_code, 200, confirm_response.text)
        confirmed = confirm_response.json()
        self.assertEqual(confirmed["item"]["status"], "resolved")
        self.assertEqual(confirmed["item"]["decision"], "ready")
        self.assertTrue(confirmed["artifact"]["s7Ready"])
        self.assertEqual(confirmed["artifact"]["qualityGate"], "human_confirmed")
        self.assertEqual(confirmed["artifact"]["confirmedBy"], "复核用户")

    def test_gap_ai_fill_confirm_scopes_human_confirmed_until_all_fill_tasks_done(self) -> None:
        """R10-B07-02：多 fillTask 目录项只复核一个任务，不得把整项置为 human_confirmed。"""
        project_id = self._create_project_with_confirmed_directory_json()
        project = store._require(project_id)
        table_template = technical_workspace_dir(project_id) / "table-template.docx"
        doc = Document()
        doc.add_paragraph("机组参数：[保证值，待填写]")
        doc.save(table_template)
        table_template_2 = technical_workspace_dir(project_id) / "table-template-2.docx"
        doc = Document()
        doc.add_paragraph("塔筒参数：[保证值，待填写]")
        doc.save(table_template_2)
        project["gap_state"] = {
            "recognitionStatus": "completed",
            "recognizedAt": now_iso(),
            "submittedForReview": False,
            "reviewConfirmed": False,
            "reviewedAt": "",
            "items": [],
            "submissions": [],
            "plan": {
                "schemaVersion": "bid-tech-gap-plan-v1",
                "projectId": project_id,
                "status": "ready",
                "items": [
                    {
                        "id": "GAP-TABLE",
                        "number": "附表A.1",
                        "title": "投标关键数据一览表",
                        "status": "needs_input",
                        "decision": "fill_required",
                        "usage": "appendix_fill",
                        "matchedMaterials": [],
                        "candidateMaterials": [],
                        "appendixTasks": [],
                        "fillTasks": [
                            {
                                "id": "FILL-TABLE",
                                "skill": "bid-tech-table-filler",
                                "status": "pending",
                                "blankSource": {
                                    "id": "APP-TABLE",
                                    "docxPath": str(table_template),
                                    "placeholderLabels": ["保证值"],
                                },
                            },
                            {
                                "id": "FILL-TABLE-2",
                                "skill": "bid-tech-table-filler",
                                "status": "pending",
                                "blankSource": {
                                    "id": "APP-TABLE-2",
                                    "docxPath": str(table_template_2),
                                    "placeholderLabels": ["保证值"],
                                },
                            },
                        ],
                        "resolvedArtifacts": [],
                    },
                ],
            },
            "planFile": "",
            "integrity": {},
        }
        store._persist_project(project)
        self._confirm_project_fact_table(project_id, {"保证值": "满足招标要求"})

        def fake_run_table_filler(manifest_path, progress_callback=None):
            manifest = json.loads(Path(manifest_path).read_text(encoding="utf-8"))
            output_file = Path(manifest["outputFile"])
            output_file.parent.mkdir(parents=True, exist_ok=True)
            doc = Document()
            doc.add_paragraph("性能保证：满足招标要求")
            doc.save(output_file)
            return {
                "schema_version": "bid-tech-table-fill-v1",
                "outputFile": str(output_file),
                "unfilledFields": [],
                "evidenceRefs": [{"field": "保证值"}],
                "fillReport": {"filledFieldCount": 1, "unfilledFieldCount": 0},
            }

        with patch(
            "app.services.technical_gap_ai_fill.run_technical_table_filler_skill",
            side_effect=fake_run_table_filler,
        ):
            first_fill = self.client.post(
                f"/api/technical/projects/{project_id}/gaps/GAP-TABLE/ai-fill",
                json={"fillTaskId": "FILL-TABLE", "operator": "测试用户"},
            )
            self.assertEqual(first_fill.status_code, 200, first_fill.text)
            first_artifact_id = first_fill.json()["artifact"]["id"]
            first_confirm = self.client.post(
                f"/api/technical/projects/{project_id}/gaps/GAP-TABLE/artifacts/{first_artifact_id}/confirm",
                json={"operator": "复核用户"},
            )
            self.assertEqual(first_confirm.status_code, 200, first_confirm.text)

        first_item = first_confirm.json()["item"]
        # 只复核第一个任务：整项不收口 human_confirmed，决策仍是 fill_required。
        self.assertNotEqual(str(first_item.get("qualityStatus") or ""), "human_confirmed")
        self.assertEqual(first_item["decision"], "fill_required")
        self.assertEqual(first_item["status"], "needs_input")
        tasks_by_id = {task["id"]: task for task in first_item["fillTasks"]}
        self.assertEqual(tasks_by_id["FILL-TABLE"]["status"], "completed")
        self.assertEqual(tasks_by_id["FILL-TABLE-2"]["status"], "pending")

        with patch(
            "app.services.technical_gap_ai_fill.run_technical_table_filler_skill",
            side_effect=fake_run_table_filler,
        ):
            second_fill = self.client.post(
                f"/api/technical/projects/{project_id}/gaps/GAP-TABLE/ai-fill",
                json={"fillTaskId": "FILL-TABLE-2", "operator": "测试用户"},
            )
            self.assertEqual(second_fill.status_code, 200, second_fill.text)
            second_artifact_id = second_fill.json()["artifact"]["id"]
            second_confirm = self.client.post(
                f"/api/technical/projects/{project_id}/gaps/GAP-TABLE/artifacts/{second_artifact_id}/confirm",
                json={"operator": "复核用户"},
            )
            self.assertEqual(second_confirm.status_code, 200, second_confirm.text)

        second_item = second_confirm.json()["item"]
        # 全部任务完成且各自产物均复核通过后，整项才允许收口为已就绪。
        self.assertEqual(second_item["qualityStatus"], "human_confirmed")
        self.assertEqual(second_item["decision"], "ready")
        self.assertEqual(second_item["status"], "resolved")

    def test_gap_ai_fill_manifest_carries_appendix_context_and_recommended_materials(self) -> None:
        project_id = self._create_project_with_confirmed_directory_json()
        project = store._require(project_id)
        project["parse_result"]["structured"]["appendices"][0]["availableParseFields"] = [
            {"id": "FIELD-POWER", "label": "单机容量", "value": "10MW", "sourceFile": "招标文件.docx"},
            {"id": "FIELD-ROTOR", "label": "叶轮直径", "value": "220m", "sourceFile": "招标文件.docx"},
        ]
        store._persist_project(project)
        detection_response = self.client.post(f"/api/technical/projects/{project_id}/gaps-detection/run")
        self.assertEqual(detection_response.status_code, 200)
        gap_plan = detection_response.json()["gapPlan"]
        fill_item = next(item for item in gap_plan["items"] if item["fillTasks"])
        gap_id = fill_item["id"]
        fill_task_id = fill_item["fillTasks"][0]["id"]
        fill_item["appendixTasks"][0]["recommendedMaterials"] = [
            {
                "id": "RAW-0001",
                "name": "性能保证基准素材.docx",
                "folderPath": "技术标/通用素材",
                "materialTier": "standard",
                "usage": "table_source",
            }
        ]
        fill_item["appendixTasks"][0]["sourceRouting"] = {
            "status": "matched",
            "source": "appendix_source_matrix",
            "ruleId": "Sheet1!R12",
            "customer": "测试业主",
            "tableTitle": "性能保证附表",
            "projectSources": [],
            "standardSources": ["性能保证基准素材"],
            "otherSources": [],
            "matchedMaterials": [{"id": "RAW-0001", "name": "性能保证基准素材.docx"}],
            "manualRequired": False,
            "useTenderParseFields": False,
        }
        project = store._require(project_id)
        project["gap_state"]["plan"] = gap_plan
        store._persist_project(project)
        self._confirm_project_fact_table(project_id)
        manifests: list[dict] = []

        def fake_run_table_filler(manifest_path, progress_callback=None):
            manifest = json.loads(Path(manifest_path).read_text(encoding="utf-8"))
            manifests.append(manifest)
            output_file = Path(manifest["outputFile"])
            doc = Document()
            doc.add_paragraph("性能保证附表")
            doc.save(output_file)
            return {
                "schema_version": "bid-tech-table-fill-v1",
                "outputFile": str(output_file),
                "unfilledFields": [],
                "evidenceRefs": [{"type": "material", "id": "RAW-0001"}],
                "fillReport": {"filledFieldCount": 2, "referenceMaterialCount": 1},
            }

        with patch(
            "app.services.technical_gap_ai_fill.run_technical_table_filler_skill",
            side_effect=fake_run_table_filler,
        ):
            response = self.client.post(
                f"/api/technical/projects/{project_id}/gaps/{gap_id}/ai-fill",
                json={
                    "fillTaskId": fill_task_id,
                    "referenceMaterialIds": ["RAW-0001"],
                    "parseFieldIds": ["APP-PERF", "FIELD-POWER"],
                    "operator": "测试用户",
                },
            )

        self.assertEqual(response.status_code, 200, response.text)
        manifest = manifests[0]
        self.assertEqual(manifest["gapItem"]["id"], gap_id)
        self.assertEqual(manifest["appendixTask"]["id"], "APP-PERF")
        self.assertEqual(manifest["appendixTask"]["sourceRouting"]["source"], "appendix_source_matrix")
        self.assertEqual(manifest["appendixTask"]["sourceRouting"]["standardSources"], ["性能保证基准素材"])
        self.assertEqual(manifest["recommendedMaterials"][0]["id"], "RAW-0001")
        self.assertEqual(manifest["referenceMaterials"][0]["id"], "RAW-0001")
        self.assertEqual([item["id"] for item in manifest["materialIndex"]], ["RAW-0001"])
        self.assertEqual(manifest["parseFields"][0]["id"], "FIELD-POWER")
        self.assertEqual(manifest["blankSource"]["id"], "APP-PERF")
        artifact = response.json()["artifact"]
        self.assertEqual(artifact["fillReport"]["filledFieldCount"], 2)
        self.assertEqual(artifact["referenceMaterials"][0]["id"], "RAW-0001")

    def test_gap_ai_fill_registers_each_batch_table_output_as_previewable_artifact(self) -> None:
        project_id = self._create_project_with_confirmed_directory_json()
        detection_response = self.client.post(f"/api/technical/projects/{project_id}/gaps-detection/run")
        self.assertEqual(detection_response.status_code, 200, detection_response.text)
        gap_plan = detection_response.json()["gapPlan"]
        fill_item = next(item for item in gap_plan["items"] if item["fillTasks"])
        gap_id = fill_item["id"]
        fill_task_id = fill_item["fillTasks"][0]["id"]
        self._confirm_project_fact_table(project_id)

        def fake_run_batch_table_filler(manifest_path, progress_callback=None):
            manifest = json.loads(Path(manifest_path).read_text(encoding="utf-8"))
            output_root = Path(manifest["outputFile"]).parent
            output_root.mkdir(parents=True, exist_ok=True)
            first_output = output_root / "001-APPX-A1_AI填写.docx"
            second_output = output_root / "002-APPX-B2_AI填写.docx"
            for path, text in (
                (first_output, "附表A.1：EW10.0-220上置"),
                (second_output, "附表B.2：满足招标要求"),
            ):
                doc = Document()
                doc.add_paragraph(text)
                doc.save(path)
            batch_report = output_root / "batch_fill_report.json"
            result = {
                "schema_version": "bid-tech-table-fill-v1",
                "outputFile": str(batch_report),
                "outputFiles": [str(first_output), str(second_output)],
                "unfilledFields": [],
                "evidenceRefs": [{"field": "投标机型"}, {"field": "供货范围"}],
                "targetResults": [
                    {
                        "schema_version": "bid-tech-table-fill-v1",
                        "outputFile": str(first_output),
                        "unfilledFields": [],
                        "evidenceRefs": [{"field": "投标机型"}],
                        "fillReport": {
                            "title": "附表A.1 投标机型总方案信息表",
                            "appendixId": "APPX-A1",
                            "filledFieldCount": 1,
                            "unfilledFieldCount": 0,
                            "targetFieldCount": 1,
                        },
                        "filledAt": "2026-05-02T00:00:00Z",
                    },
                    {
                        "schema_version": "bid-tech-table-fill-v1",
                        "outputFile": str(second_output),
                        "unfilledFields": [],
                        "evidenceRefs": [{"field": "供货范围"}],
                        "fillReport": {
                            "title": "附表B.2 供货范围响应表",
                            "appendixId": "APPX-B2",
                            "filledFieldCount": 1,
                            "unfilledFieldCount": 0,
                            "targetFieldCount": 1,
                        },
                        "filledAt": "2026-05-02T00:00:00Z",
                    },
                ],
                "fillReport": {
                    "batch": True,
                    "targetCount": 2,
                    "successfulTargetCount": 2,
                    "failedTargetCount": 0,
                    "filledFieldCount": 2,
                    "unfilledFieldCount": 0,
                    "targetFieldCount": 2,
                },
            }
            batch_report.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
            return result

        with patch("app.services.technical_gap_ai_fill.run_technical_table_filler_skill", side_effect=fake_run_batch_table_filler):
            response = self.client.post(
                f"/api/technical/projects/{project_id}/gaps/{gap_id}/ai-fill",
                json={"fillTaskId": fill_task_id, "operator": "测试用户"},
            )

        self.assertEqual(response.status_code, 200, response.text)
        payload = response.json()
        self.assertEqual(len(payload["artifacts"]), 2)
        self.assertEqual(payload["item"]["fillTasks"][0]["outputArtifactIds"], [item["id"] for item in payload["artifacts"]])
        self.assertEqual(payload["item"]["resolvedSource"], "2 份AI填写产物")
        self.assertEqual(payload["item"]["qualityReport"]["status"], "passed")
        for index, artifact in enumerate(payload["artifacts"], start=1):
            self.assertTrue(artifact["fileName"].endswith(".docx"))
            self.assertEqual(artifact["batchTargetIndex"], index)
            self.assertEqual(artifact["batchTargetCount"], 2)
            self.assertEqual(artifact["qualityReport"]["status"], "passed")
            self.assertIn(quote(artifact["fileName"]), artifact["onlyoffice"]["browserFileUrl"])

        updated_gap = self.client.get(f"/api/technical/projects/{project_id}/gaps").json()["gapPlan"]
        updated_item = next(item for item in updated_gap["items"] if item["id"] == gap_id)
        self.assertEqual(len(updated_item["resolvedArtifacts"]), 2)
        first_artifact = updated_item["resolvedArtifacts"][0]
        download = self.client.get(
            f"/api/technical/projects/{project_id}/gaps/artifacts/{first_artifact['id']}/content/{quote(first_artifact['fileName'])}"
        )
        self.assertEqual(download.status_code, 200)
        self.assertGreater(len(download.content), 0)

    def test_project_fact_table_builds_required_fields_from_project_and_gap_placeholders(self) -> None:
        project_id = self._create_project_with_confirmed_directory_json()
        project = store._require(project_id)
        project["owner"] = "华能集团"
        project["customerName"] = "华能集团"
        project["identity"] = {"owner": "华能集团", "customerName": "华能集团"}
        project["turbineModel"] = {
            "model": "EW10.0-220下置",
            "ratedPowerKw": 10000,
            "rotorDiameterM": 220,
            "hubHeightM": 125,
        }
        store._persist_project(project)
        detection_response = self.client.post(f"/api/technical/projects/{project_id}/gaps-detection/run")
        self.assertEqual(detection_response.status_code, 200, detection_response.text)
        gap_plan = detection_response.json()["gapPlan"]
        fill_item = next(item for item in gap_plan["items"] if item["fillTasks"])
        fill_item["fillTasks"][0]["blankSource"]["placeholderLabels"] = ["投标方案", "招标方", "未知保证值"]
        project = store._require(project_id)
        project["gap_state"]["plan"] = gap_plan
        project["gap_state"]["factSpecs"] = _test_fact_specs()
        store._persist_project(project)

        response = self.client.post(f"/api/technical/projects/{project_id}/gaps/facts/build")

        self.assertEqual(response.status_code, 200, response.text)
        payload = response.json()
        self.assertEqual(payload["status"], "draft")
        labels = [field["label"] for field in payload["fields"]]
        # 以清单为唯一骨架：匹配到 spec 的（投标机型、投标方案经 spec 100 别名）成行；
        # 匹配不到的（招标方、缺口占位"未知保证值"）不再单独成行
        self.assertIn("投标机型", labels)
        self.assertIn("投标方案", labels)
        self.assertNotIn("招标方", labels)
        self.assertNotIn("未知保证值", labels)
        model = next(field for field in payload["fields"] if field["label"] == "投标机型")
        self.assertEqual(model["value"], "EW10.0-220下置")
        self.assertEqual(model["status"], "extracted")
        self.assertTrue(model["sourceRefs"])
        # 未提取骨架计 unextracted，missingCount 只统计 missing_source
        self.assertEqual(payload["summary"]["missingCount"], 0)
        self.assertEqual(payload["summary"]["specTotal"], 128)

        confirmed = self._confirm_project_fact_table(project_id, {"承诺函致函对象全称": "按招标文件要求执行"})

        self.assertEqual(
            confirmed["summary"]["confirmedCount"],
            sum(1 for field in confirmed["fields"] if str(field.get("value") or "").strip()),
        )
        self.assertTrue(confirmed["confirmedAt"])
        confirmed_unknown = next(field for field in confirmed["fields"] if field["label"] == "承诺函致函对象全称")
        self.assertEqual(confirmed_unknown["status"], "confirmed")
        self.assertEqual(confirmed_unknown["value"], "按招标文件要求执行")

    def test_project_fact_table_preserves_manual_fields_when_rebuilt(self) -> None:
        project_id = self._create_project_with_confirmed_directory_json()
        detection_response = self.client.post(f"/api/technical/projects/{project_id}/gaps-detection/run")
        self.assertEqual(detection_response.status_code, 200, detection_response.text)
        _seed_fact_specs(project_id)

        build_response = self.client.post(f"/api/technical/projects/{project_id}/gaps/facts/build")
        self.assertEqual(build_response.status_code, 200, build_response.text)
        fields = build_response.json()["fields"]
        fields.append(
            {
                "id": "FACT-MANUAL-1",
                "label": "现场特殊要求",
                "category": "人工补充事实",
                "value": "满足低温施工窗口要求",
                "required": False,
                "status": "candidate",
                "confidence": 1,
                "sourcePriority": 360,
                "sourceRefs": [{"type": "manualFact", "title": "人工新增", "field": "现场特殊要求"}],
            }
        )
        save_response = self.client.put(
            f"/api/technical/projects/{project_id}/gaps/facts",
            json={"fields": fields, "confirm": True, "operator": "测试用户"},
        )
        self.assertEqual(save_response.status_code, 200, save_response.text)

        rebuild_response = self.client.post(f"/api/technical/projects/{project_id}/gaps/facts/build")
        self.assertEqual(rebuild_response.status_code, 200, rebuild_response.text)
        by_label = {field["label"]: field for field in rebuild_response.json()["fields"]}
        self.assertEqual(by_label["现场特殊要求"]["value"], "满足低温施工窗口要求")
        self.assertEqual(by_label["现场特殊要求"]["category"], "人工补充事实")
        self.assertEqual(by_label["现场特殊要求"]["sourceRefs"][0]["type"], "manualFact")

    def test_fact_field_patch_unknown_id_returns_404_without_creating(self) -> None:
        """R06-B07-08：普通 PATCH 不应隐式创建字段，未持久化的临时 id 必须 404 且不落库。"""
        project_id = self._create_project_with_confirmed_directory_json()
        detection_response = self.client.post(f"/api/technical/projects/{project_id}/gaps-detection/run")
        self.assertEqual(detection_response.status_code, 200, detection_response.text)
        _seed_fact_specs(project_id)
        build_response = self.client.post(f"/api/technical/projects/{project_id}/gaps/facts/build")
        self.assertEqual(build_response.status_code, 200, build_response.text)
        field_total = len(build_response.json()["fields"])

        patch_response = self.client.patch(
            f"/api/technical/projects/{project_id}/gaps/facts/FACT-MANUAL-999",
            json={"value": "张三", "confirm": True, "operator": "测试用户"},
        )
        self.assertEqual(patch_response.status_code, 404, patch_response.text)

        facts_response = self.client.get(f"/api/technical/projects/{project_id}/gaps/facts")
        self.assertEqual(facts_response.status_code, 200, facts_response.text)
        self.assertEqual(len(facts_response.json()["fields"]), field_total)

    def test_manual_fact_field_saved_then_patch_confirm_succeeds(self) -> None:
        """R06-B07-08：人工新增字段整表持久化后保留 id，逐字段确认幂等且不产生重复字段。"""
        project_id = self._create_project_with_confirmed_directory_json()
        detection_response = self.client.post(f"/api/technical/projects/{project_id}/gaps-detection/run")
        self.assertEqual(detection_response.status_code, 200, detection_response.text)
        _seed_fact_specs(project_id)
        build_response = self.client.post(f"/api/technical/projects/{project_id}/gaps/facts/build")
        self.assertEqual(build_response.status_code, 200, build_response.text)
        fields = build_response.json()["fields"]
        field_total = len(fields)
        fields.append(
            {
                "id": "FACT-MANUAL-1",
                "label": "项目联系人",
                "category": "人工补充事实",
                "value": "张三",
                "required": False,
                "status": "extracted",
                "confidence": 1,
                "sourcePriority": 360,
                "sourceRefs": [{"type": "manualFact", "title": "人工新增", "field": "项目联系人"}],
            }
        )
        # 模拟前端"先整表持久化"：confirm 缺省为草稿保存
        save_response = self.client.put(
            f"/api/technical/projects/{project_id}/gaps/facts",
            json={"fields": fields, "operator": "测试用户"},
        )
        self.assertEqual(save_response.status_code, 200, save_response.text)
        saved_ids = [field["id"] for field in save_response.json()["fields"]]
        self.assertIn("FACT-MANUAL-1", saved_ids)
        self.assertEqual(len(saved_ids), field_total + 1)

        for _ in range(2):
            patch_response = self.client.patch(
                f"/api/technical/projects/{project_id}/gaps/facts/FACT-MANUAL-1",
                json={"value": "张三", "status": "extracted", "confirm": True, "operator": "测试用户"},
            )
            self.assertEqual(patch_response.status_code, 200, patch_response.text)
            self.assertEqual(patch_response.json()["field"]["status"], "confirmed")

        facts_response = self.client.get(f"/api/technical/projects/{project_id}/gaps/facts")
        self.assertEqual(facts_response.status_code, 200, facts_response.text)
        manual_fields = [field for field in facts_response.json()["fields"] if field["id"] == "FACT-MANUAL-1"]
        self.assertEqual(len(manual_fields), 1)
        self.assertEqual(manual_fields[0]["value"], "张三")
        self.assertEqual(manual_fields[0]["status"], "confirmed")

    def test_fact_specs_upload_accepts_xlsx(self) -> None:
        """R06-B07-09：项目级实时表上传 .xlsx 成功用例。"""
        project_id = self._create_project_with_confirmed_outline()
        xlsx_path = Path(self.temp_dir.name) / "实时表.xlsx"
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.append(EXPECTED_HEADER)
        for index in range(1, 4):
            sheet.append([index, "招标文件-技术规范书", "第一章 1.1", f"上传字段{index}", "说明", "", "招标文件/技术规范书"])
        workbook.save(xlsx_path)

        with xlsx_path.open("rb") as handle:
            upload_response = self.client.post(
                f"/api/technical/projects/{project_id}/gaps/facts/specs-upload",
                files={"file": ("实时表.xlsx", handle, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
            )
        self.assertEqual(upload_response.status_code, 200, upload_response.text)
        self.assertEqual(upload_response.json()["specTotal"], 3)

        facts_response = self.client.get(f"/api/technical/projects/{project_id}/gaps/facts")
        self.assertEqual(facts_response.status_code, 200, facts_response.text)
        self.assertTrue(facts_response.json()["specsImported"])
        self.assertEqual(facts_response.json()["specTotal"], 3)

    def test_fact_specs_upload_rejects_xls(self) -> None:
        """R06-B07-09：.xls 与前后端规则一致地被 400 拒绝，且不写入项目清单。"""
        project_id = self._create_project_with_confirmed_outline()
        upload_response = self.client.post(
            f"/api/technical/projects/{project_id}/gaps/facts/specs-upload",
            files={"file": ("实时表.xls", b"not a real xls", "application/vnd.ms-excel")},
        )
        self.assertEqual(upload_response.status_code, 400, upload_response.text)
        self.assertIn(".xlsx", upload_response.json()["detail"])

        facts_response = self.client.get(f"/api/technical/projects/{project_id}/gaps/facts")
        self.assertEqual(facts_response.status_code, 200, facts_response.text)
        self.assertFalse(facts_response.json()["specsImported"])

    def test_project_fact_table_filters_noisy_parse_items_and_extracts_table_fields(self) -> None:
        project_id = self._create_project_with_confirmed_directory_json()
        project = store._require(project_id)
        project["identity"] = {"owner": "测试业主", "customerName": "测试业主"}
        project["turbineModel"] = {
            "model": "EW10.0-220上置",
            "ratedPowerKw": 10000,
            "rotorDiameterM": 220,
        }
        project["parse_result"]["items"] = [
            {
                "id": "REQ-PROJECT",
                "title": "项目名称",
                "value": "华能真实项目名称",
                "keyValue": "华能真实项目名称",
                "fieldKey": "projectName",
                "fieldGroup": "projectBasics",
                "sourceFile": "招标文件.docx",
            },
            {
                "id": "REQ-NO",
                "title": "招标编号",
                "value": "HNZB2025-12-1-382-01",
                "keyValue": "HNZB2025-12-1-382-01",
                "fieldKey": "tenderNo",
                "fieldGroup": "projectBasics",
                "sourceFile": "招标文件.docx",
            },
            {
                "id": "REQ-NOISE-OWNER",
                "title": "招标人",
                "value": "提供其采用标准或规范的中文版本。只有当其采用的标准或规范不低于本技术规范的要求时，投标人采用的标准或规范才能为招标人认可",
                "keyValue": "提供其采用标准或规范的中文版本。只有当其采用的标准或规范不低于本技术规范的要求时，投标人采用的标准或规范才能为招标人认可",
                "fieldKey": "tenderer",
                "fieldGroup": "projectBasics",
                "sourceFile": "招标文件.docx",
            },
            {
                "id": "REQ-NOISE-COMMIT",
                "title": "技术承诺",
                "value": "项目单位应按项目具体要求填写专用部分的非固化内容，投标人至少填写专用部分的项目技术承诺要求、专题方案要求、其他要求以及附表。",
                "keyValue": "项目单位应按项目具体要求填写专用部分的非固化内容，投标人至少填写专用部分的项目技术承诺要求、专题方案要求、其他要求以及附表。",
                "fieldKey": "technicalCommitment",
                "fieldGroup": "projectBasics",
                "sourceFile": "招标文件.docx",
            },
        ]
        table_docx = technical_workspace_dir(project_id) / "appendix-a1.docx"
        doc = Document()
        table = doc.add_table(rows=4, cols=5)
        for row, values in enumerate(
            [
                ["编号", "项目", "项目", "投标机型1", "投标机型2"],
                ["1", "投标机型", "投标机型", "", ""],
                ["2", "机组台数", "机组台数", "", ""],
                ["3", "总容量（MW）", "总容量（MW）", "", ""],
            ]
        ):
            for col, value in enumerate(values):
                table.cell(row, col).text = value
        doc.save(table_docx)
        project["gap_state"] = {
            "recognitionStatus": "completed",
            "recognizedAt": now_iso(),
            "submittedForReview": False,
            "reviewConfirmed": False,
            "reviewedAt": "",
            "items": [],
            "submissions": [],
            "plan": {
                "schemaVersion": "bid-tech-gap-plan-v1",
                "projectId": project_id,
                "items": [
                    {
                        "id": "GAP-TABLE",
                        "number": "附表A.1",
                        "title": "投标机型总方案信息表",
                        "status": "needs_input",
                        "decision": "fill_required",
                        "usage": "appendix_fill",
                        "matchedMaterials": [],
                        "candidateMaterials": [],
                        "appendixTasks": [],
                        "fillTasks": [
                            {
                                "id": "FILL-TABLE",
                                "skill": "bid-tech-table-filler",
                                "status": "pending",
                                "blankSource": {"id": "APPX-A1", "docxPath": str(table_docx)},
                            }
                        ],
                        "resolvedArtifacts": [],
                    }
                ],
            },
            "planFile": "",
            "integrity": {},
            "factSpecs": _test_fact_specs(),
        }
        store._persist_project(project)

        response = self.client.post(f"/api/technical/projects/{project_id}/gaps/facts/build")

        self.assertEqual(response.status_code, 200, response.text)
        fields = response.json()["fields"]
        by_label = {field["label"]: field for field in fields}
        self.assertEqual(by_label["项目名称"]["value"], "华能真实项目名称")
        self.assertIn("机组台数", by_label)
        self.assertIn("总装机容量", by_label)
        self.assertEqual(by_label["机组台数"]["category"], "待填写表格字段")
        # 清单之外的字段不再成行：招标编号/招标人无匹配 spec，技术承诺仍是噪声
        self.assertNotIn("招标编号", by_label)
        self.assertNotIn("招标人", by_label)
        self.assertNotIn("技术承诺", by_label)

    def test_project_fact_table_fills_core_facts_by_project_customer_standard_priority(self) -> None:
        project_id = self._create_project_with_confirmed_directory_json()
        project = store._require(project_id)
        project["identity"] = {"owner": "测试业主", "customerName": "测试业主"}
        project["turbineModel"] = {
            "model": "EW10.0-220上置",
            "ratedPowerKw": 10000,
            "rotorDiameterM": 220,
        }
        work_dir = technical_workspace_dir(project_id)

        def write_fact_docx(file_name: str, rows: list[tuple[str, str]]) -> Path:
            path = work_dir / file_name
            doc = Document()
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "字段"
            table.cell(0, 1).text = "值"
            for label, value in rows:
                cells = table.add_row().cells
                cells[0].text = label
                cells[1].text = value
            doc.save(path)
            return path

        standard_path = write_fact_docx(
            "standard-facts.docx",
            [
                ("总装机容量", "400MW"),
                ("机组台数", "40"),
                ("轮毂高度", "110m"),
                ("安全等级", "IEC IIA"),
                ("空气密度", "1.15kg/m3"),
                ("设计寿命", "20年"),
            ],
        )
        customer_path = write_fact_docx(
            "customer-facts.docx",
            [
                ("总装机容量", "500MW"),
                ("机组台数", "50"),
                ("轮毂高度", "120m"),
                ("安全等级", "IEC IIB"),
                ("湍流强度", "0.12"),
                ("设计寿命", "20年"),
            ],
        )
        project_path = write_fact_docx(
            "project-facts.docx",
            [
                ("总装机容量", "600MW"),
                ("机组台数", "60"),
                ("轮毂高度", "125m"),
                ("安全等级", "IEC S"),
                ("空气密度", "1.16kg/m3"),
                ("湍流强度", "0.10"),
                ("极端风速", "52.5m/s"),
                ("设计寿命", "20年"),
            ],
        )
        project["gap_state"] = {
            "recognitionStatus": "completed",
            "recognizedAt": now_iso(),
            "submittedForReview": False,
            "reviewConfirmed": False,
            "reviewedAt": "",
            "items": [],
            "submissions": [],
            "plan": {
                "schemaVersion": "bid-tech-gap-plan-v1",
                "projectId": project_id,
                "materialIndex": [
                    {
                        "id": "RAW-STANDARD",
                        "name": standard_path.name,
                        "path": str(standard_path),
                        "materialTier": "standard",
                        "folderPath": "技术标/通用素材",
                    },
                    {
                        "id": "RAW-CUSTOMER",
                        "name": customer_path.name,
                        "path": str(customer_path),
                        "materialTier": "customer",
                        "folderPath": "技术标/客户素材/测试业主",
                    },
                    {
                        "id": "RAW-PROJECT",
                        "name": project_path.name,
                        "path": str(project_path),
                        "materialTier": "project",
                        "folderPath": "技术标/项目素材/PRJ-TEST",
                    },
                ],
                "items": [],
            },
            "planFile": "",
            "integrity": {},
            "factSpecs": _test_fact_specs(),
        }
        store._persist_project(project)

        response = self.client.post(f"/api/technical/projects/{project_id}/gaps/facts/build")

        self.assertEqual(response.status_code, 200, response.text)
        by_label = {field["label"]: field for field in response.json()["fields"]}
        self.assertEqual(by_label["总装机容量"]["value"], "600MW")
        self.assertEqual(by_label["机组台数"]["value"], "60")
        self.assertEqual(by_label["轮毂高度"]["value"], "125")
        self.assertEqual(by_label["轮毂高度"]["unit"], "m")
        self.assertEqual(by_label["安全等级"]["value"], "IEC S")
        self.assertEqual(by_label["湍流强度"]["value"], "0.10")
        self.assertEqual(by_label["总装机容量"]["status"], "extracted")
        self.assertEqual(by_label["总装机容量"]["sourceRefs"][0]["materialTier"], "project")
        # 空气密度/极端风速/设计寿命匹配不到 spec，清单模式下不再成行
        self.assertNotIn("空气密度", by_label)
        self.assertNotIn("极端风速", by_label)
        self.assertNotIn("设计寿命", by_label)

    def test_project_fact_table_rejects_header_and_author_noise_from_materials(self) -> None:
        project_id = self._create_project_with_confirmed_directory_json()
        project = store._require(project_id)
        project["identity"] = {"owner": "测试业主", "customerName": "测试业主"}
        project["turbineModel"] = {
            "model": "EW10.0-220上置",
            "ratedPowerKw": 10000,
            "rotorDiameterM": 220,
        }
        work_dir = technical_workspace_dir(project_id)
        noisy_path = work_dir / "project-facts-with-noise.docx"
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "字段"
        table.cell(0, 1).text = "值"
        for label, value in [
            ("轮毂高度", "池建昌 2025-08-07 00:00:00"),
            ("年平均风速", "各年平均风速(m/s) 年份2025"),
            ("空气密度", "场址空气密度下"),
            ("轮毂高度", "125m"),
            ("年平均风速", "7.20m/s"),
            ("空气密度", "1.16kg/m3"),
        ]:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value
        doc.save(noisy_path)
        project["gap_state"] = {
            "recognitionStatus": "completed",
            "recognizedAt": now_iso(),
            "submittedForReview": False,
            "reviewConfirmed": False,
            "reviewedAt": "",
            "items": [],
            "submissions": [],
            "plan": {
                "schemaVersion": "bid-tech-gap-plan-v1",
                "projectId": project_id,
                "materialIndex": [
                    {
                        "id": "RAW-NOISY-PROJECT",
                        "name": noisy_path.name,
                        "path": str(noisy_path),
                        "materialTier": "project",
                        "folderPath": "技术标/项目素材/PRJ-TEST",
                    },
                ],
                "items": [],
            },
            "planFile": "",
            "integrity": {},
            "factSpecs": _test_fact_specs(),
        }
        store._persist_project(project)

        response = self.client.post(f"/api/technical/projects/{project_id}/gaps/facts/build")

        self.assertEqual(response.status_code, 200, response.text)
        by_label = {field["label"]: field for field in response.json()["fields"]}
        self.assertEqual(by_label["轮毂高度"]["value"], "125")
        self.assertEqual(by_label["年平均风速"]["value"], "7.20m/s")
        self.assertEqual(by_label["轮毂高度"]["status"], "extracted")
        self.assertEqual(by_label["年平均风速"]["status"], "extracted")
        # 空气密度匹配不到 spec，清单模式下不再成行
        self.assertNotIn("空气密度", by_label)

    def test_project_fact_table_derives_guarantee_values_from_wind_speed_matrix(self) -> None:
        project_id = self._create_project_with_confirmed_directory_json()
        project = store._require(project_id)
        project["identity"] = {"owner": "测试业主", "customerName": "测试业主"}
        project["turbineModel"] = {"model": "EW10.0-220上置", "ratedPowerKw": 10000}
        work_dir = technical_workspace_dir(project_id)
        guarantee_path = work_dir / "project-guarantee-matrix.docx"
        doc = Document()
        summary = doc.add_table(rows=1, cols=4)
        summary.cell(0, 0).text = "发电量"
        summary.cell(0, 1).text = "机位点尾流后平均风速"
        summary.cell(0, 2).text = "m/s"
        summary.cell(0, 3).text = "7.22"
        matrix = doc.add_table(rows=1, cols=3)
        matrix.cell(0, 0).text = "指定测风塔125m高度处的年平均风速(m/s)"
        matrix.cell(0, 1).text = "风电场保证年上网电量(MWh)"
        matrix.cell(0, 2).text = "风电场保证年等效满负荷小时数(h)"
        for wind_speed, energy, hours in [
            ("7.1", "1667083", "2778"),
            ("7.2", "1701601", "2836"),
            ("7.3", "1735525", "2893"),
        ]:
            cells = matrix.add_row().cells
            cells[0].text = wind_speed
            cells[1].text = energy
            cells[2].text = hours
        doc.save(guarantee_path)
        project["gap_state"] = {
            "recognitionStatus": "completed",
            "recognizedAt": now_iso(),
            "submittedForReview": False,
            "reviewConfirmed": False,
            "reviewedAt": "",
            "items": [],
            "submissions": [],
            "plan": {
                "schemaVersion": "bid-tech-gap-plan-v1",
                "projectId": project_id,
                "materialIndex": [
                    {
                        "id": "RAW-GUARANTEE",
                        "name": guarantee_path.name,
                        "path": str(guarantee_path),
                        "materialTier": "project",
                        "folderPath": "技术标/项目素材/PRJ-TEST",
                    },
                ],
                "items": [],
            },
            "planFile": "",
            "integrity": {},
            "factSpecs": _test_fact_specs(),
        }
        store._persist_project(project)

        response = self.client.post(f"/api/technical/projects/{project_id}/gaps/facts/build")

        self.assertEqual(response.status_code, 200, response.text)
        by_label = {field["label"]: field for field in response.json()["fields"]}
        self.assertEqual(by_label["年平均风速"]["value"], "7.22")
        self.assertEqual(by_label["年平均风速"]["unit"], "m/s")
        self.assertEqual(by_label["保证有效小时数"]["value"], "2836")
        self.assertEqual(by_label["保证有效小时数"]["unit"], "h")
        self.assertEqual(by_label["保证有效小时数"]["sourceRefs"][0]["materialTier"], "project")
        # 保证发电量匹配不到 spec，清单模式下不再成行
        self.assertNotIn("保证发电量", by_label)

    def test_gap_ai_fill_requires_confirmed_fact_table_and_manifest_carries_it(self) -> None:
        project_id = self._create_project_with_confirmed_directory_json()
        detection_response = self.client.post(f"/api/technical/projects/{project_id}/gaps-detection/run")
        self.assertEqual(detection_response.status_code, 200, detection_response.text)
        gap_plan = detection_response.json()["gapPlan"]
        fill_item = next(item for item in gap_plan["items"] if item["fillTasks"])
        gap_id = fill_item["id"]
        fill_task_id = fill_item["fillTasks"][0]["id"]

        blocked = self.client.post(
            f"/api/technical/projects/{project_id}/gaps/{gap_id}/ai-fill",
            json={"fillTaskId": fill_task_id, "operator": "测试用户"},
        )
        self.assertEqual(blocked.status_code, 400)
        self.assertIn("项目事实表", blocked.json()["detail"])

        self._confirm_project_fact_table(project_id)
        manifests: list[dict] = []

        def fake_run_table_filler(manifest_path, progress_callback=None):
            manifest = json.loads(Path(manifest_path).read_text(encoding="utf-8"))
            manifests.append(manifest)
            output_file = Path(manifest["outputFile"])
            doc = Document()
            doc.add_paragraph("性能保证附表")
            doc.save(output_file)
            return {
                "schema_version": "bid-tech-table-fill-v1",
                "outputFile": str(output_file),
                "unfilledFields": [],
                "evidenceRefs": [{"field": "单机容量"}, {"field": "招标方"}],
                "fillReport": {"filledFieldCount": 2, "unfilledFieldCount": 0},
            }

        with patch("app.services.technical_gap_ai_fill.run_technical_table_filler_skill", side_effect=fake_run_table_filler):
            response = self.client.post(
                f"/api/technical/projects/{project_id}/gaps/{gap_id}/ai-fill",
                json={"fillTaskId": fill_task_id, "operator": "测试用户"},
            )

        self.assertEqual(response.status_code, 200, response.text)
        self.assertEqual(manifests[0]["projectFactTable"]["status"], "confirmed")
        # 清单模式下事实表只含 spec 骨架行（招标方等清单外字段不再成行），用 spec 11 投标机型断言随单
        self.assertTrue([field for field in manifests[0]["projectFactTable"]["fields"] if field["label"] == "投标机型"])
        artifact = response.json()["artifact"]
        self.assertEqual(artifact["qualityReport"]["status"], "passed")
        self.assertGreaterEqual(artifact["qualityReport"]["coverageRate"], 0.85)
        self.assertGreaterEqual(artifact["qualityReport"]["correctnessRate"], 0.85)
        self.assertGreaterEqual(artifact["qualityReport"]["completenessRate"], 0.85)

    def test_gap_ai_fill_all_runs_word_tasks_before_table_tasks_and_returns_quality_summary(self) -> None:
        project_id = self._create_project_with_confirmed_directory_json()
        project = store._require(project_id)
        word_template = technical_workspace_dir(project_id) / "word-template.docx"
        doc = Document()
        doc.add_paragraph("招标方：[招标方，待填写]")
        doc.save(word_template)
        table_template = technical_workspace_dir(project_id) / "table-template.docx"
        doc = Document()
        doc.add_paragraph("性能保证：[保证值，待填写]")
        doc.save(table_template)
        table_template_2 = technical_workspace_dir(project_id) / "table-template-2.docx"
        doc = Document()
        doc.add_paragraph("性能保证2：[保证值，待填写]")
        doc.save(table_template_2)
        project["gap_state"] = {
            "recognitionStatus": "completed",
            "recognizedAt": now_iso(),
            "submittedForReview": False,
            "reviewConfirmed": False,
            "reviewedAt": "",
            "items": [],
            "submissions": [],
            "plan": {
                "schemaVersion": "bid-tech-gap-plan-v1",
                "projectId": project_id,
                "status": "ready",
                "items": [
                    {
                        "id": "GAP-TABLE",
                        "number": "1.6",
                        "title": "投标关键数据一览表",
                        "status": "needs_input",
                        "decision": "fill_required",
                        "usage": "appendix_fill",
                        "matchedMaterials": [],
                        "candidateMaterials": [],
                        "appendixTasks": [],
                        "fillTasks": [
                            {
                                "id": "FILL-TABLE",
                                "skill": "bid-tech-table-filler",
                                "status": "pending",
                                "blankSource": {"id": "APP-TABLE", "docxPath": str(table_template), "placeholderLabels": ["保证值"]},
                            },
                            {
                                "id": "FILL-TABLE-2",
                                "skill": "bid-tech-table-filler",
                                "status": "pending",
                                "blankSource": {"id": "APP-TABLE-2", "docxPath": str(table_template_2), "placeholderLabels": ["保证值"]},
                            }
                        ],
                        "resolvedArtifacts": [],
                    },
                    {
                        "id": "GAP-WORD",
                        "number": "1.4",
                        "title": "风电机组自主可控推广应用的承诺",
                        "status": "needs_input",
                        "decision": "fill_required",
                        "usage": "section_fill",
                        "matchedMaterials": [],
                        "candidateMaterials": [],
                        "appendixTasks": [],
                        "fillTasks": [
                            {
                                "id": "FILL-WORD",
                                "skill": "bid-tech-word-placeholder-filler",
                                "status": "pending",
                                "blankSource": {
                                    "id": "RAW-WORD",
                                    "sourceType": "material_fill_template",
                                    "docxPath": str(word_template),
                                    "placeholderLabels": ["招标方"],
                                },
                            }
                        ],
                        "resolvedArtifacts": [],
                    },
                ],
            },
            "planFile": "",
            "integrity": {},
        }
        project["owner"] = "华能集团"
        project["customerName"] = "华能集团"
        project["identity"] = {"owner": "华能集团", "customerName": "华能集团"}
        store._persist_project(project)
        self._confirm_project_fact_table(project_id, {"保证值": "满足招标要求"})
        calls: list[str] = []

        def fake_run_word_filler(manifest_path, progress_callback=None):
            calls.append("word")
            manifest = json.loads(Path(manifest_path).read_text(encoding="utf-8"))
            output_file = Path(manifest["outputFile"])
            doc = Document()
            doc.add_paragraph("招标方：华能集团")
            doc.save(output_file)
            return {
                "schema_version": "bid-tech-word-placeholder-fill-v1",
                "outputFile": str(output_file),
                "unfilledFields": [],
                "evidenceRefs": [{"field": "招标方"}],
                "fillReport": {"filledPlaceholderCount": 1, "unfilledPlaceholderCount": 0},
            }

        def fake_run_table_filler(manifest_path, progress_callback=None):
            calls.append("table")
            manifest = json.loads(Path(manifest_path).read_text(encoding="utf-8"))
            output_file = Path(manifest["outputFile"])
            doc = Document()
            doc.add_paragraph("性能保证：满足招标要求")
            doc.save(output_file)
            return {
                "schema_version": "bid-tech-table-fill-v1",
                "outputFile": str(output_file),
                "unfilledFields": [],
                "evidenceRefs": [{"field": "保证值"}],
                "fillReport": {"filledFieldCount": 1, "unfilledFieldCount": 0},
            }

        with patch("app.services.technical_gap_ai_fill.run_technical_word_placeholder_filler_skill", side_effect=fake_run_word_filler), patch(
            "app.services.technical_gap_ai_fill.run_technical_table_filler_skill",
            side_effect=fake_run_table_filler,
        ):
            response = self.client.post(
                f"/api/technical/projects/{project_id}/gaps/ai-fill-all",
                json={"operator": "测试用户"},
            )

        self.assertEqual(response.status_code, 200, response.text)
        self.assertEqual(calls, ["word", "table", "table"])
        payload = response.json()
        self.assertEqual(payload["summary"]["total"], 3)
        self.assertEqual(payload["summary"]["passed"], 3)
        self.assertEqual(payload["qualityReport"]["status"], "passed")
        self.assertEqual([result["gapId"] for result in payload["results"]], ["GAP-WORD", "GAP-TABLE", "GAP-TABLE"])
        table_files = [result["fileName"] for result in payload["results"] if result["gapId"] == "GAP-TABLE"]
        self.assertEqual(len(table_files), len(set(table_files)))

    def test_gap_ai_fill_all_quality_summary_uses_weighted_field_counts(self) -> None:
        aggregate = aggregate_technical_gap_fill_quality(
            [
                {
                    "qualityReport": {
                        "status": "needs_review",
                        "coverageRate": 0.0,
                        "correctnessRate": 0.0,
                        "completenessRate": 0.0,
                        "expectedFieldCount": 1,
                        "filledFieldCount": 0,
                        "unfilledFieldCount": 1,
                        "evidenceRefCount": 0,
                    }
                },
                {
                    "qualityReport": {
                        "status": "passed",
                        "coverageRate": 1.0,
                        "correctnessRate": 1.0,
                        "completenessRate": 1.0,
                        "expectedFieldCount": 99,
                        "filledFieldCount": 99,
                        "unfilledFieldCount": 0,
                        "evidenceRefCount": 99,
                    }
                },
            ],
            [],
        )

        self.assertEqual(aggregate["status"], "passed")
        self.assertEqual(aggregate["coverageRate"], 0.99)
        self.assertEqual(aggregate["correctnessRate"], 1.0)
        self.assertEqual(aggregate["completenessRate"], 0.99)

    def test_fill_quality_does_not_pass_when_target_fields_are_unknown(self) -> None:
        report = technical_gap_ai_fill._build_fill_quality_report(
            {
                "outputFile": "filled.docx",
                "unfilledFields": [],
                "evidenceRefs": [{"field": "来源"}],
                "fillReport": {"targetFieldCount": 0, "filledFieldCount": 0, "unfilledFieldCount": 0},
            },
            output_exists=True,
        )

        self.assertEqual(report["status"], "needs_review")
        self.assertEqual(report["coverageRate"], 0.0)
        self.assertEqual(report["expectedFieldCount"], 0)

    def test_fill_quality_uses_semantic_validation_for_correctness(self) -> None:
        report = technical_gap_ai_fill._build_fill_quality_report(
            {
                "outputFile": "filled.docx",
                "unfilledFields": [],
                "evidenceRefs": [{"field": "来源1"}, {"field": "来源2"}, {"field": "来源3"}],
                "fillReport": {
                    "placeholderCount": 3,
                    "filledPlaceholderCount": 3,
                    "unfilledPlaceholderCount": 0,
                    "semanticCheckCount": 3,
                    "semanticPassedCount": 1,
                    "semanticFailedCount": 2,
                    "semanticValidationRate": 0.3333,
                },
            },
            output_exists=True,
        )

        self.assertEqual(report["status"], "needs_review")
        self.assertEqual(report["coverageRate"], 1.0)
        self.assertEqual(report["evidenceChainRate"], 1.0)
        self.assertEqual(report["correctnessRate"], 0.3333)
        self.assertEqual(report["semanticFailedCount"], 2)

    def test_stale_material_word_fill_task_is_repaired_to_word_skill_before_ai_fill(self) -> None:
        project_id = self._create_project_with_confirmed_directory_json()
        project = store._require(project_id)
        template = technical_workspace_dir(project_id) / "word-template.docx"
        doc = Document()
        doc.add_paragraph("招标方：[招标方，待填写]")
        doc.save(template)
        plan = {
            "schemaVersion": "bid-tech-gap-plan-v1",
            "projectId": project_id,
            "status": "ready",
            "items": [
                {
                    "id": "GAP-WORD",
                    "number": "1.4",
                    "title": "风电机组自主可控推广应用的承诺",
                    "status": "needs_input",
                    "decision": "fill_required",
                    "usage": "section_fill",
                    "matchedMaterials": [],
                    "candidateMaterials": [],
                    "appendixTasks": [],
                    "fillTasks": [
                        {
                            "id": "FILL-GAP-WORD-RAW-WORD",
                            "skill": "bid-tech-table-filler",
                            "status": "pending",
                            "blankSource": {
                                "id": "RAW-WORD",
                                "materialId": "RAW-WORD",
                                "sourceType": "material_fill_template",
                                "title": "待填写-承诺函.docx",
                                "docxPath": str(template),
                                "placeholderCount": 1,
                                "placeholderLabels": ["招标方"],
                            },
                        }
                    ],
                    "resolvedArtifacts": [],
                }
            ],
        }
        project["gap_state"] = {
            "recognitionStatus": "completed",
            "recognizedAt": now_iso(),
            "submittedForReview": False,
            "reviewConfirmed": False,
            "reviewedAt": "",
            "items": [],
            "submissions": [],
            "plan": plan,
            "planFile": "",
            "integrity": {},
        }
        project["owner"] = "华能集团"
        project["customerName"] = "华能集团"
        project["identity"] = {"owner": "华能集团", "customerName": "华能集团"}
        store._persist_project(project)

        repaired = self.client.get(f"/api/technical/projects/{project_id}/gaps-detection").json()["gapPlan"]
        task = repaired["items"][0]["fillTasks"][0]
        self.assertEqual(task["skill"], "bid-tech-word-placeholder-filler")
        self._confirm_project_fact_table(project_id)

        manifests: list[dict] = []

        def fake_run_word_filler(manifest_path, progress_callback=None):
            manifest = json.loads(Path(manifest_path).read_text(encoding="utf-8"))
            manifests.append(manifest)
            output_file = Path(manifest["outputFile"])
            output_doc = Document()
            output_doc.add_paragraph("招标方：华能集团")
            output_doc.save(output_file)
            return {
                "schema_version": "bid-tech-word-placeholder-fill-v1",
                "outputFile": str(output_file),
                "unfilledFields": [],
                "evidenceRefs": [{"type": "projectIdentity", "field": "招标方"}],
                "fillReport": {"filledPlaceholderCount": 1, "unfilledPlaceholderCount": 0},
            }

        with patch("app.services.technical_gap_ai_fill.run_technical_table_filler_skill", side_effect=AssertionError("table filler should not run")), patch(
            "app.services.technical_gap_ai_fill.run_technical_word_placeholder_filler_skill",
            side_effect=fake_run_word_filler,
        ):
            response = self.client.post(
                f"/api/technical/projects/{project_id}/gaps/GAP-WORD/ai-fill",
                json={"fillTaskId": "FILL-GAP-WORD-RAW-WORD", "operator": "测试用户"},
            )

        self.assertEqual(response.status_code, 200, response.text)
        self.assertEqual(manifests[0]["schemaVersion"], "bid-tech-word-placeholder-fill-v1")
        self.assertEqual(manifests[0]["projectIdentity"]["owner"], "华能集团")
        payload = response.json()
        self.assertEqual(payload["artifact"]["skill"], "bid-tech-word-placeholder-filler")
        self.assertEqual(payload["item"]["fillTasks"][0]["status"], "completed")

    def test_fill_material_index_is_locked_to_selected_references(self) -> None:
        index = technical_gap_ai_fill._material_index_for_fill(
            {"id": "PRJ-TEST"},
            {
                "materialIndex": [
                    {"id": "RAW-SELECTED", "name": "已选素材.docx"},
                    {"id": "RAW-UNSELECTED", "name": "未选素材.docx"},
                ]
            },
            {
                "candidateMaterials": [{"id": "RAW-CANDIDATE", "name": "候选素材.docx"}],
                "matchedMaterials": [],
                "appendixTasks": [],
            },
            ["RAW-SELECTED"],
            [{"id": "RAW-SELECTED", "name": "已选素材.docx"}],
        )

        self.assertEqual([item["id"] for item in index], ["RAW-SELECTED"])

    def test_ai_fill_reference_selection_uses_all_routed_materials_and_respects_explicit_empty(self) -> None:
        item = {
            "sourceRouting": {"source": "appendix_source_matrix"},
            "sourceRoutedMaterials": [{"id": "RAW-ITEM-1"}, {"id": "RAW-ITEM-2"}],
            "matchedMaterials": [{"id": "RAW-MATCHED"}],
        }
        appendix_task = {
            "sourceRouting": {"source": "appendix_source_matrix"},
            "recommendedMaterials": [{"id": "RAW-TASK-1"}, {"id": "RAW-TASK-2"}],
        }

        self.assertEqual(
            technical_gap_ai_fill._selected_reference_material_ids(item, appendix_task, {}),
            ["RAW-TASK-1", "RAW-TASK-2", "RAW-ITEM-1", "RAW-ITEM-2"],
        )
        self.assertEqual(
            technical_gap_ai_fill._selected_reference_material_ids(
                item,
                appendix_task,
                {"referenceMaterialIds": []},
            ),
            [],
        )

    def test_gap_upload_registers_real_project_artifact_for_s7(self) -> None:
        project_id = self._create_project_with_confirmed_directory_json()
        detection_response = self.client.post(f"/api/technical/projects/{project_id}/gaps-detection/run")
        self.assertEqual(detection_response.status_code, 200)
        gap_plan = detection_response.json()["gapPlan"]
        gap_id = next(item for item in gap_plan["items"] if item["status"] == "needs_input")["id"]

        response = self.client.post(
            f"/api/technical/projects/{project_id}/gaps/{gap_id}/upload",
            json={
                "bidType": "技术标",
                "files": [
                    {
                        "name": "客户补充性能保证.docx",
                        "type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        "data": "客户补充性能保证内容",
                    }
                ],
            },
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["item"]["status"], "resolved")
        artifact = payload["artifact"]
        self.assertEqual(artifact["source"], "manual_upload")
        self.assertEqual(artifact["skill"], "")
        self.assertTrue(Path(artifact["path"]).exists())
        self.assertTrue(artifact["s7Ready"])
        gap_payload = self.client.get(f"/api/technical/projects/{project_id}/gaps").json()
        updated_item = next(item for item in gap_payload["gapPlan"]["items"] if item["id"] == gap_id)
        self.assertEqual(updated_item["resolvedArtifacts"][0]["source"], "manual_upload")
        self.assertEqual(updated_item["status"], "resolved")

    def test_gap_upload_preserves_browser_docx_data_url_for_s7(self) -> None:
        project_id = self._create_project_with_confirmed_directory_json()
        detection_response = self.client.post(f"/api/technical/projects/{project_id}/gaps-detection/run")
        self.assertEqual(detection_response.status_code, 200)
        gap_plan = detection_response.json()["gapPlan"]
        gap_id = next(item for item in gap_plan["items"] if item["status"] == "needs_input")["id"]

        upload_docx = Path(self.temp_dir.name) / "客户补充性能保证.docx"
        doc = Document()
        doc.add_paragraph("客户原始上传 Word 内容")
        doc.save(upload_docx)
        data_url = (
            "data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,"
            + base64.b64encode(upload_docx.read_bytes()).decode("ascii")
        )

        response = self.client.post(
            f"/api/technical/projects/{project_id}/gaps/{gap_id}/upload",
            json={
                "bidType": "技术标",
                "files": [
                    {
                        "name": "客户补充性能保证.docx",
                        "type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        "data": data_url,
                    }
                ],
            },
        )

        self.assertEqual(response.status_code, 200)
        artifact_path = Path(response.json()["artifact"]["path"])
        self.assertTrue(artifact_path.exists())
        uploaded_doc = Document(str(artifact_path))
        self.assertIn(
            "客户原始上传 Word 内容",
            "\n".join(paragraph.text for paragraph in uploaded_doc.paragraphs),
        )

    def test_gap_select_existing_material_registers_real_artifact_for_s7(self) -> None:
        project_id = self._create_project_with_confirmed_directory_json()
        detection_response = self.client.post(f"/api/technical/projects/{project_id}/gaps-detection/run")
        self.assertEqual(detection_response.status_code, 200)
        gap_plan = detection_response.json()["gapPlan"]
        gap_id = next(item for item in gap_plan["items"] if item["status"] == "needs_input")["id"]

        prepared_docx = Path(self.temp_dir.name) / "素材库既有性能保证.docx"
        doc = Document()
        doc.add_paragraph("素材库既有 Word 内容")
        doc.save(prepared_docx)

        async def fake_prepare_existing_gap_material_files(project, selected_gap_id, data):
            self.assertEqual(selected_gap_id, gap_id)
            self.assertEqual(data["materials"][0]["id"], "RAW-0099")
            return [
                {
                    "materialId": "RAW-0099",
                    "materialName": "素材库既有性能保证",
                    "fileName": prepared_docx.name,
                    "path": str(prepared_docx),
                    "folderPath": "技术标/通用素材",
                    "materialTier": "standard",
                    "sourceKind": "cleaned",
                }
            ]

        with patch(
            "app.services.technical_gap_service.prepare_technical_existing_gap_material_files",
            side_effect=fake_prepare_existing_gap_material_files,
        ):
            response = self.client.post(
                f"/api/technical/projects/{project_id}/gaps/{gap_id}/select-material",
                json={
                    "materials": [
                        {
                            "id": "RAW-0099",
                            "name": "素材库既有性能保证",
                            "folderPath": "技术标/通用素材",
                            "materialTier": "standard",
                        }
                    ],
                },
            )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["item"]["status"], "resolved")
        artifact = payload["artifact"]
        self.assertEqual(artifact["source"], "material_library")
        self.assertEqual(artifact["materialId"], "RAW-0099")
        self.assertEqual(artifact["path"], str(prepared_docx))
        self.assertTrue(artifact["s7Ready"])
        self.assertIn("onlyoffice", artifact)
        updated_gap = self.client.get(f"/api/technical/projects/{project_id}/gaps").json()["gapPlan"]
        updated_item = next(item for item in updated_gap["items"] if item["id"] == gap_id)
        self.assertEqual(updated_item["resolvedArtifacts"][0]["source"], "material_library")
        self.assertEqual(updated_item["resolvedArtifacts"][0]["path"], str(prepared_docx))

    def test_gap_detection_matches_material_from_s2_wiki_cards_when_toc_has_no_refs(self) -> None:
        project_id = self._create_project_with_confirmed_directory_json()
        project_dir = technical_workspace_dir(project_id)
        toc_json = project_dir / "s2_toc_workdir" / "投标文件-总目录.json"
        toc_data = json.loads(toc_json.read_text(encoding="utf-8"))
        for item in toc_data["items"]:
            item["material_refs"] = []
        toc_json.write_text(json.dumps(toc_data, ensure_ascii=False, indent=2), encoding="utf-8")
        _replace_confirmed_outline_from_toc(project_id, toc_data)
        wiki_cards = project_dir / "s2_toc_workdir" / "wiki" / "卡片"
        wiki_cards.mkdir(parents=True, exist_ok=True)
        (wiki_cards / "技术评分标准索引表.md").write_text(
            "---\n"
            "name: 技术评分标准索引表\n"
            "path: 技术标/通用素材/技术评分标准索引表.docx\n"
            "scope: 通用\n"
            "category: 技术标\n"
            "material_id: RAW-0001\n"
            "skeleton_section: \"1.1\"\n"
            "deprecated: false\n"
            "---\n",
            encoding="utf-8",
        )

        response = self.client.post(f"/api/technical/projects/{project_id}/gaps-detection/run")

        self.assertEqual(response.status_code, 200)
        plan = response.json()["gapPlan"]
        matched = next(item for item in plan["items"] if item["number"] == "1.1")
        self.assertEqual(matched["status"], "matched")
        self.assertEqual(matched["matchedMaterials"][0]["id"], "RAW-0001")
        self.assertEqual(matched["matchedMaterials"][0]["source"], "wiki")

    def test_gap_detection_uses_one_parent_chapter_material_for_wind_resource_chapter(self) -> None:
        project_id = self._create_project_with_confirmed_directory_json()
        project_dir = technical_workspace_dir(project_id)
        toc_json = project_dir / "s2_toc_workdir" / "投标文件-总目录.json"
        wind_candidates = [
            {
                "id": "RAW-0471",
                "name": "定制-项目风资源评估与机组选型排布及发电量计算.docx",
                "docx": "技术标/项目素材/MAT-HN-CHIFENG-001/技术标-专题方案要求/定制-项目风资源评估与机组选型排布及发电量计算.docx",
                "folderPath": "技术标/项目素材/MAT-HN-CHIFENG-001/技术标-专题方案要求",
                "materialTier": "project",
            },
            {
                "id": "RAW-0473",
                "name": "定制-风资源评估与机位排布方案.docx",
                "docx": "技术标/项目素材/MAT-HN-CHIFENG-001/技术标-风资源评估与机位排布方案/定制-风资源评估与机位排布方案.docx",
                "folderPath": "技术标/项目素材/MAT-HN-CHIFENG-001/技术标-风资源评估与机位排布方案",
                "materialTier": "project",
            },
            {
                "id": "RAW-0478",
                "name": "风资源评估报告.docx",
                "docx": "技术标/项目素材/MAT-HN-CHIFENG-001/风资源评估报告（解决方案部_风资源处：风资源流程传递）/风资源评估报告.docx",
                "folderPath": "技术标/项目素材/MAT-HN-CHIFENG-001/风资源评估报告（解决方案部_风资源处：风资源流程传递）",
                "materialTier": "project",
            },
        ]
        toc_data = {
            "schema_version": "bid-toc-json-v1",
            "items": [
                        {
                            "order": 1,
                            "number": "第3章",
                            "title": "风资源评估与机位排布方案",
                            "level": 1,
                            "annotation": "保留",
                            "source": "template",
                            "material_refs": wind_candidates,
                        },
                        *[
                            {
                                "order": index + 1,
                                "number": f"3.{index}",
                                "title": title,
                                "level": 2,
                                "annotation": "保留",
                                "source": "template",
                                "material_refs": wind_candidates if index in {4, 7} else [],
                            }
                            for index, title in enumerate(
                                [
                                    "总体方案概览",
                                    "项目概况",
                                    "风资源分析",
                                    "机组选型",
                                    "风机适应性分析",
                                    "方案及发电量结果",
                                    "不确定性分析",
                                ],
                                start=1,
                            )
                        ],
                        {
                            "order": 9,
                            "number": "附表E.1",
                            "title": "投标人风资源评估与机位排布方案",
                            "level": 1,
                            "annotation": "新增-附表空表",
                            "source": "tender_appendix",
                            "material_refs": [],
                        },
            ],
        }
        toc_json.write_text(json.dumps(toc_data, ensure_ascii=False, indent=2), encoding="utf-8")
        _replace_confirmed_outline_from_toc(project_id, toc_data)
        project = store._require(project_id)
        project["materialProjectId"] = "MAT-HN-CHIFENG-001"
        project["materialProjectCode"] = "MAT-HN-CHIFENG-001"
        project["materialProjectMode"] = "library"
        project["materialCustomerName"] = "华能集团"
        project["parse_result"]["structured"] = {
            "appendices": [
                {
                    "id": "APPX-0033",
                    "title": "附表E.1 投标人风资源评估与机位排布方案",
                    "sourceFile": "招标文件.docx",
                    "workspacePath": "/data/documents/PRJ-0003/parse/appendices/APPX-0033.docx",
                    "rowCount": 10,
                    "blankTable": True,
                }
            ]
        }
        store._persist_project(project)

        response = self.client.post(f"/api/technical/projects/{project_id}/gaps-detection/run")

        self.assertEqual(response.status_code, 200, response.text)
        plan = response.json()["gapPlan"]
        self.assertFalse([item for item in plan["items"] if len(item.get("matchedMaterials") or []) > 1])
        parent = next(item for item in plan["items"] if item["number"] == "第3章")
        self.assertEqual(parent["decision"], "ready")
        self.assertEqual(parent["status"], "matched")
        self.assertEqual(parent["coverageRole"], "chapter_master")
        self.assertEqual(parent["usage"], "chapter_master")
        self.assertEqual(parent["matchedMaterials"][0]["id"], "RAW-0473")
        self.assertEqual(parent["matchedMaterials"][0]["usage"], "chapter_master")
        self.assertEqual(parent["fillTasks"], [])
        self.assertEqual(parent["appendixTasks"], [])
        for number in [f"3.{index}" for index in range(1, 8)]:
            child = next(item for item in plan["items"] if item["number"] == number)
            self.assertEqual(child["decision"], "ready")
            self.assertEqual(child["status"], "matched")
            self.assertEqual(child["coverageRole"], "covered_by_parent")
            self.assertEqual(child["coveredByParent"], parent["id"])
            self.assertEqual(child["matchedMaterials"], [])
            self.assertEqual(child["appendixTasks"], [])
            self.assertEqual(child["fillTasks"], [])
        appendix_item = next(item for item in plan["items"] if item["number"] == "附表E.1")
        self.assertEqual(appendix_item["decision"], "fill_required")
        self.assertEqual(appendix_item["status"], "needs_input")
        self.assertEqual(appendix_item["matchedMaterials"], [])
        self.assertEqual([task["id"] for task in appendix_item["appendixTasks"]], ["APPX-0033"])
        self.assertIn(
            "RAW-0473",
            [item["id"] for item in appendix_item["appendixTasks"][0]["recommendedMaterials"]],
        )
        self.assertEqual(appendix_item["appendixTasks"][0]["recommendedMaterials"][0]["id"], "RAW-0473")

    def test_gap_detection_manifest_contains_project_scoped_material_index(self) -> None:
        project_id = self._create_project_with_confirmed_directory_json()
        project = store._require(project_id)
        project["customerName"] = "华能集团"
        project["materialCustomerName"] = "华能集团"
        project["materialProjectId"] = "MAT-HN-CHIFENG-001"
        project["materialProjectCode"] = "MAT-HN-CHIFENG-001"
        project["materialProjectMode"] = "library"
        store._persist_project(project)
        manifests: list[dict] = []

        async def fake_raw_files(**kwargs):
            items = []
            if (
                kwargs.get("material_tier") == "project"
                and kwargs.get("project_id") == "MAT-HN-CHIFENG-001"
            ):
                items.append(
                    {
                        "id": "RAW-0473",
                        "name": "定制-风资源评估与机位排布方案.docx",
                        "folderPath": "技术标/项目定制/S4-S6联调项目/技术标-风资源评估与机位排布方案",
                        "materialTier": "project",
                        "hasCleanedWord": True,
                        "cleanedFileName": "定制-风资源评估与机位排布方案.docx",
                    }
                )
            return {"items": items, "total": len(items), "page": 1, "pageSize": 1000}

        def fake_gap_planner(manifest_path):
            manifest = json.loads(Path(manifest_path).read_text(encoding="utf-8"))
            manifests.append(manifest)
            output_file = Path(manifest["outputFile"])
            plan = minimal_gap_plan_from_manifest(manifest)
            output_file.write_text(json.dumps(plan, ensure_ascii=False), encoding="utf-8")
            return {"schema_version": "bid-tech-gap-plan-v1", "outputFile": str(output_file)}

        with patch("app.services.technical_gap_planner.technical_material_store.raw_files", side_effect=fake_raw_files), \
            patch("app.services.technical_gap_planner.run_technical_gap_planner_skill", side_effect=fake_gap_planner):
            response = self.client.post(f"/api/technical/projects/{project_id}/gaps-detection/run")

        self.assertEqual(response.status_code, 200, response.text)
        manifest = manifests[0]
        self.assertEqual(manifest["customerName"], "华能集团")
        self.assertEqual(
            manifest["materialScope"]["paths"],
            [
                "技术标/标准文件",
                "技术标/客户定制/华能集团",
                "技术标/项目定制/S4-S6联调项目",
            ],
        )
        self.assertEqual(manifest["materialIndex"][0]["id"], "RAW-0473")
        self.assertTrue(
            all(
                item["folderPath"].startswith(tuple(manifest["materialScope"]["paths"]))
                for item in manifest["materialIndex"]
            )
        )

    def test_gap_review_mock_flow_runs_from_s4_to_s6(self) -> None:
        project_id = self._create_project_with_confirmed_outline()

        detection_response = self.client.post(f"/api/technical/projects/{project_id}/gaps-detection/run")
        self.assertEqual(detection_response.status_code, 200)
        detection_payload = detection_response.json()
        self.assertEqual(detection_payload["status"], "completed")
        self.assertGreater(len(detection_payload["items"]), 0)

        gaps_response = self.client.get(f"/api/technical/projects/{project_id}/gaps")
        self.assertEqual(gaps_response.status_code, 200)
        gaps_payload = gaps_response.json()
        self.assertEqual(gaps_payload["status"], "ready")
        items = gaps_payload["items"]
        self.assertGreater(len(items), 0)

        for index, item in enumerate(items):
            submission_response = self.client.post(
                f"/api/technical/projects/{project_id}/materials/submissions",
                json={
                    "missingId": item["id"],
                    "bidType": item.get("bidType") or "技术标",
                    "files": [
                        {
                            "name": f"{item['id']}.docx",
                            "size": 1024,
                            "type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        }
                    ],
                },
            )
            self.assertEqual(submission_response.status_code, 200)

            if index % 2 == 0:
                update_response = self.client.put(
                    f"/api/technical/projects/{project_id}/gaps/{item['id']}",
                    json={"action": "resolve", "source": {"name": f"{item['id']}.docx"}},
                )
            else:
                update_response = self.client.patch(
                    f"/api/technical/projects/{project_id}/materials/missing/{item['id']}",
                    json={"status": "skipped", "reason": "MVP阶段先跳过"},
                )
            self.assertEqual(update_response.status_code, 200)

        submit_review_response = self.client.post(f"/api/technical/projects/{project_id}/gaps/submit-review")
        self.assertEqual(submit_review_response.status_code, 200)
        submit_payload = submit_review_response.json()["payload"]
        self.assertTrue(submit_payload["submittedForReview"])
        self.assertGreater(len(submit_payload["items"]), 0)

    def test_submit_review_blocks_until_gap_plan_is_resolved(self) -> None:
        project_id = self._create_project_with_confirmed_outline()

        detection_response = self.client.post(f"/api/technical/projects/{project_id}/gaps-detection/run")
        self.assertEqual(detection_response.status_code, 200)

        submit_review_response = self.client.post(f"/api/technical/projects/{project_id}/gaps/submit-review")
        self.assertEqual(submit_review_response.status_code, 400)
        self.assertIn("缺口未解决", submit_review_response.json()["detail"])

        gaps_payload = self.client.get(f"/api/technical/projects/{project_id}/gaps").json()
        for item in gaps_payload["items"]:
            update_response = self.client.patch(
                f"/api/technical/projects/{project_id}/materials/missing/{item['id']}",
                json={"status": "skipped", "reason": "测试中人工确认忽略"},
            )
            self.assertEqual(update_response.status_code, 200)

        recheck_response = self.client.post(f"/api/technical/projects/{project_id}/gaps/recheck")
        self.assertEqual(recheck_response.status_code, 200)
        self.assertEqual(recheck_response.json()["integrity"]["status"], "passed")

        submit_review_response = self.client.post(f"/api/technical/projects/{project_id}/gaps/submit-review")
        self.assertEqual(submit_review_response.status_code, 200)
        submit_payload = submit_review_response.json()["payload"]
        self.assertTrue(submit_payload["submittedForReview"])
        self.assertGreater(len(submit_payload["items"]), 0)
        self.assertTrue(all(item["status"] == "skipped" for item in submit_payload["items"]))

        final_gap_payload = self.client.get(f"/api/technical/projects/{project_id}/gaps").json()
        self.assertTrue(final_gap_payload["submittedForReview"])
        self.assertTrue(all(item["status"] == "skipped" for item in submit_payload["items"]))

    def test_gap_recheck_passes_when_all_items_are_resolved_or_ignored_with_s4_ready_artifacts(self) -> None:
        project_id = self._create_project_with_confirmed_directory_json()
        detection_response = self.client.post(f"/api/technical/projects/{project_id}/gaps-detection/run")
        self.assertEqual(detection_response.status_code, 200)
        gap_plan = detection_response.json()["gapPlan"]
        for item in gap_plan["items"]:
            status = item.get("status")
            if status in {"matched", "structural"}:
                continue
            if item.get("fillTasks"):
                item["status"] = "resolved"
                for task in item["fillTasks"]:
                    task["status"] = "completed"
                item["resolvedArtifacts"] = [
                    {
                        "id": f"ART-{item['id']}",
                        "source": "ai_fill",
                        "skill": "bid-tech-table-filler",
                        "fileName": f"{item['id']}.docx",
                        "s7Ready": True,
                        "qualityReport": {"status": "passed"},
                        "unfilledFields": [],
                        "fillReport": {"filledFieldCount": 2, "unfilledFieldCount": 0},
                        "evidenceRefs": [{"type": "material", "id": "RAW-0001"}],
                    }
                ]
            else:
                item["status"] = "ignored"
                item["skipReason"] = "测试中人工确认不适用"
        project = store._require(project_id)
        project["gap_state"]["plan"] = gap_plan
        store._persist_project(project)

        response = self.client.post(f"/api/technical/projects/{project_id}/gaps/recheck")

        self.assertEqual(response.status_code, 200, response.text)
        integrity = response.json()["integrity"]
        self.assertEqual(integrity["status"], "passed")
        self.assertEqual(integrity["blockingCount"], 0)
        self.assertEqual(integrity["blockingItems"], [])


if __name__ == "__main__":
    unittest.main()
