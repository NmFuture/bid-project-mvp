from __future__ import annotations

import json
import unittest
from io import BytesIO
from pathlib import Path
from subprocess import CalledProcessError, CompletedProcess
from types import SimpleNamespace
from unittest.mock import AsyncMock, patch
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches

from app.services.performance_package_service import (
    ITEM_CONTRACT_FORMAT_VERSION,
    PerformancePackageService,
    ensure_contract_docx_file_name,
    match_contract_chunks_to_items,
    parse_performance_summary_docx,
    render_contract_item_docx,
    split_performance_contract_docx,
)
from app.services.peripheral import PeripheralError


def _summary_docx_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("陆上11MW业绩")
    doc.add_paragraph("合同业绩台数共计69台，已通过240小时试运行台数共计0台，投运容量共计0MW。")
    table = doc.add_table(rows=3, cols=9)
    headers = [
        "序号",
        "型号",
        "项目名称/合同名称",
        "买方名称",
        "合同台数",
        "试运行台数",
        "投运容量(MW)",
        "交货期/投运时间",
        "联系人及 电话",
    ]
    rows = [
        ["1", "11-230", "华电新疆喀什 2×66 万千瓦", "华电喀什能源有限公司哈密分公司", "59", "/", "/", "/", "翟文东0998-5817059"],
        ["2", "11-230", "内蒙古能源察右前旗50万千瓦风光实验实证项目", "内蒙古电力勘测设计院有限责任公司", "10", "/", "/", "/", "安文凭13644843963"],
    ]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    for row_index, row in enumerate(rows, start=1):
        for column_index, value in enumerate(row):
            table.cell(row_index, column_index).text = value
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _offshore_summary_docx_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("海上8MW及以上机型合同业绩表")
    doc.add_paragraph("投标人8MW及以上机型国内合同业绩台数共计628台，投运容量共计241台。")
    table = doc.add_table(rows=4, cols=10)
    headers = [
        "序号",
        "型号和规格",
        "项目名称",
        "总容量（MW)",
        "台数",
        "买方名称",
        "合同时间",
        "投运时间",
        "投运容量（MW)",
        "联系人及电话",
    ]
    rows = [
        ["1", "SG8.0-167", "长乐外海Ｃ区第一批项目", "200", "25", "福建省福能海峡发电有限公司", "2020年", "2021年", "200", "黄长凤137745142615"],
        ["2", "EW8.0-230 EW8.5-230", "山能渤中海上风电基地B1项目", "100", "4 8", "中国电建集团中南勘测设计研究院有限公司", "2024年", "2025年", "/", "程总17673048550"],
        ["3", "EW8.5-230", "国华山东B2场址项目", "501.5", "59", "国家能源投资集团（济南）新能源有限责任公司", "2022年", "/", "/", "刘欢18615226720"],
    ]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    for row_index, row in enumerate(rows, start=1):
        for column_index, value in enumerate(row):
            table.cell(row_index, column_index).text = value
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _tiny_png_bytes() -> bytes:
    return (
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01"
        b"\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde"
        b"\x00\x00\x00\x0cIDATx\x9cc\xf8\xff\xff?\x00\x05"
        b"\xfe\x02\xfeA\x86\xa7\x9b\x00\x00\x00\x00IEND\xaeB`\x82"
    )


def _contract_bundle_docx_bytes() -> bytes:
    image_stream = BytesIO(_tiny_png_bytes())
    doc = Document()
    doc.add_paragraph("华电新疆喀什 2×66 万千瓦风力发电机组采购合同")
    doc.paragraphs[0].runs[0].font.name = "方正兰亭超细黑简体"
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "封面页"
    table.cell(0, 0).paragraphs[0].runs[0].font.name = "宋体"
    table.cell(0, 1).paragraphs[0].add_run().add_picture(image_stream, width=Inches(0.2))
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("内蒙古能源察右前旗50万千瓦风光实验实证项目风力发电机组合同")
    image_stream = BytesIO(_tiny_png_bytes())
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "合同范围"
    table.cell(0, 1).paragraphs[0].add_run().add_picture(image_stream, width=Inches(0.2))
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _contract_with_table_layout_docx_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("山东华电枣庄山亭19万千瓦风电项目风力发电机组合同")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "合同首页"
    table.cell(0, 1).text = "主要商务条款页"
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "技术参数表第一页"
    table.cell(0, 1).text = "技术参数表第二页"
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _contract_with_paired_page_table_docx_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("山东华电枣庄山亭19万千瓦风电项目风力发电机组合同")
    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "合同首页"
    table.cell(0, 1).text = "签字盖章页"
    table.cell(1, 0).paragraphs[0].add_run().add_picture(BytesIO(_tiny_png_bytes()), width=Inches(0.2))
    table.cell(1, 1).paragraphs[0].add_run().add_picture(BytesIO(_tiny_png_bytes()), width=Inches(0.2))
    table.cell(2, 0).text = "主要技术参数页"
    table.cell(2, 1).text = "主要技术参数页"
    table.cell(3, 0).paragraphs[0].add_run().add_picture(BytesIO(_tiny_png_bytes()), width=Inches(0.2))
    table.cell(3, 1).paragraphs[0].add_run().add_picture(BytesIO(_tiny_png_bytes()), width=Inches(0.2))
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def _docx_declared_fonts(content: bytes) -> set[str]:
    fonts: set[str] = set()
    font_attrs = {
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii",
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi",
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cs",
    }
    with ZipFile(BytesIO(content)) as zf:
        for name in zf.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            try:
                root = ET.fromstring(zf.read(name))
            except ET.ParseError:
                continue
            for element in root.iter():
                local_name = element.tag.rsplit("}", 1)[-1]
                if local_name == "rFonts":
                    fonts.update(value for attr, value in element.attrib.items() if attr in font_attrs and value)
                if local_name == "font":
                    font_name = element.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name")
                    if font_name:
                        fonts.add(font_name)
                    theme_font_name = element.attrib.get("typeface")
                    if theme_font_name:
                        fonts.add(theme_font_name)
                if local_name in {"latin", "ea", "cs"}:
                    theme_font_name = element.attrib.get("typeface")
                    if theme_font_name:
                        fonts.add(theme_font_name)
    return fonts


def _docx_rfont_non_name_refs(content: bytes) -> set[str]:
    refs: set[str] = set()
    ref_attrs = {
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}asciiTheme",
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsiTheme",
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsiaTheme",
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cstheme",
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hint",
    }
    with ZipFile(BytesIO(content)) as zf:
        for name in zf.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            try:
                root = ET.fromstring(zf.read(name))
            except ET.ParseError:
                continue
            for element in root.iter():
                if element.tag.rsplit("}", 1)[-1] == "rFonts":
                    refs.update(value for attr, value in element.attrib.items() if attr in ref_attrs and value)
    return refs


def _docx_settings_count(content: bytes, local_name: str) -> int:
    with ZipFile(BytesIO(content), "r") as zf:
        settings_xml = zf.read("word/settings.xml")
    root = ET.fromstring(settings_xml)
    return sum(1 for element in root.iter() if element.tag.rsplit("}", 1)[-1] == local_name)


def _docx_local_name_count(content: bytes, local_name: str) -> int:
    with ZipFile(BytesIO(content), "r") as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    return sum(1 for element in root.iter() if element.tag.rsplit("}", 1)[-1] == local_name)


def _docx_effect_extents(content: bytes) -> list[dict[str, str]]:
    with ZipFile(BytesIO(content), "r") as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    return [
        {attr: element.attrib.get(attr, "") for attr in ("l", "t", "r", "b")}
        for element in root.iter()
        if element.tag.rsplit("}", 1)[-1] == "effectExtent"
    ]


def _docx_picture_ids(content: bytes) -> list[str]:
    with ZipFile(BytesIO(content), "r") as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    return [
        element.attrib.get("id", "")
        for element in root.iter()
        if element.tag.rsplit("}", 1)[-1] == "docPr"
    ]


def _docx_body_layout(content: bytes) -> list[dict[str, object]]:
    with ZipFile(BytesIO(content)) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    body = root.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}body")
    assert body is not None
    layout = []
    for child in list(body):
        local_name = child.tag.rsplit("}", 1)[-1]
        layout.append(
            {
                "kind": local_name,
                "text": "".join(child.itertext()).strip(),
                "pageBreak": any(
                    node.tag.rsplit("}", 1)[-1] == "br"
                    and node.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type") == "page"
                    for node in child.iter()
                ),
            }
        )
    return layout


def _docx_table_row_pagination(content: bytes) -> list[dict[str, object]]:
    with ZipFile(BytesIO(content)) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    rows = []
    for row in root.iter(f"{namespace}tr"):
        trpr = row.find(f"{namespace}trPr")
        rows.append(
            {
                "text": "".join(row.itertext()).strip(),
                "cantSplit": trpr is not None and trpr.find(f"{namespace}cantSplit") is not None,
                "keepNext": any(
                    (ppr := paragraph.find(f"{namespace}pPr")) is not None
                    and ppr.find(f"{namespace}keepNext") is not None
                    for paragraph in row.iter(f"{namespace}p")
                ),
            }
        )
    return rows


def test_parse_performance_summary_docx_extracts_category_and_rows() -> None:
    parsed = parse_performance_summary_docx(_summary_docx_bytes(), file_name="陆上11MW业绩_汇总表.docx")

    assert parsed["categoryName"] == "陆上11MW业绩"
    assert parsed["scene"] == "陆上"
    assert parsed["powerRating"] == "11MW"
    assert parsed["rowCount"] == 2
    assert [field["label"] for field in parsed["fieldSchema"]] == [
        "序号",
        "型号",
        "项目名称/合同名称",
        "买方名称",
        "合同台数",
        "试运行台数",
        "投运容量(MW)",
        "交货期/投运时间",
        "联系人及 电话",
    ]
    assert parsed["rows"][0]["projectName"] == "华电新疆喀什 2×66 万千瓦"
    assert parsed["rows"][0]["customerName"] == "华电喀什能源有限公司哈密分公司"
    assert parsed["rows"][0]["contractQuantity"] == "59"
    assert parsed["rows"][0]["trialOperationQuantity"] == ""
    assert parsed["rows"][0]["turbineModels"] == ["11-230"]


def test_parse_performance_summary_docx_normalizes_models_and_time_fields() -> None:
    parsed = parse_performance_summary_docx(_offshore_summary_docx_bytes(), file_name="海上8MW及以上业绩_汇总表.docx")

    assert parsed["categoryName"] == "海上8MW及以上机型合同业绩表"
    assert parsed["scene"] == "海上"
    assert parsed["powerRating"] == "8MW及以上"
    assert parsed["rowCount"] == 3

    first = parsed["rows"][0]
    assert first["turbineModel"] == "SG8.0-167"
    assert first["turbineModels"] == ["SG8.0-167"]
    assert first["deliveryOrOperationTime"] == "2021年"
    assert first["contractYear"] == 2020
    assert first["operationYear"] == 2021
    assert first["timeFacts"]["contractTimeRaw"] == "2020年"
    assert first["timeFacts"]["operationTimeRaw"] == "2021年"

    multi_model = parsed["rows"][1]
    assert multi_model["turbineModel"] == "EW8.0-230 EW8.5-230"
    assert multi_model["turbineModels"] == ["EW8.0-230", "EW8.5-230"]
    assert multi_model["contractYear"] == 2024
    assert multi_model["operationYear"] == 2025

    not_commissioned = parsed["rows"][2]
    assert not_commissioned["contractYear"] == 2022
    assert not_commissioned["operationYear"] is None
    assert not_commissioned["deliveryOrOperationTime"] == ""


def test_split_performance_contract_docx_preserves_item_documents_and_images() -> None:
    chunks = split_performance_contract_docx(_contract_bundle_docx_bytes(), file_name="陆上11MW业绩_合同.docx")

    assert len(chunks) == 2
    assert chunks[0]["title"] == "华电新疆喀什 2×66 万千瓦风力发电机组采购合同"
    assert chunks[1]["title"] == "内蒙古能源察右前旗50万千瓦风光实验实证项目风力发电机组合同"

    first_doc = Document(BytesIO(chunks[0]["content"]))
    second_doc = Document(BytesIO(chunks[1]["content"]))
    assert first_doc.paragraphs[0].text == "华电新疆喀什 2×66 万千瓦风力发电机组"
    assert second_doc.paragraphs[0].text == "内蒙古能源察右前旗50万千瓦风光实验实证项目风力发电机组"
    assert len(first_doc.tables) == 1
    assert len(second_doc.tables) == 1
    assert first_doc.tables[0].cell(0, 0).text == "封面页"
    assert second_doc.tables[0].cell(0, 0).text == "合同范围"
    assert first_doc.element.xpath('.//*[local-name()="drawing" or local-name()="pict"]')
    assert second_doc.element.xpath('.//*[local-name()="drawing" or local-name()="pict"]')


def test_split_performance_contract_docx_preserves_internal_table_layout_breaks() -> None:
    chunks = split_performance_contract_docx(_contract_with_table_layout_docx_bytes(), file_name="枣庄山亭19万千瓦业绩_合同.docx")

    assert len(chunks) == 1
    content = render_contract_item_docx(chunks[0], output_title="山东华电枣庄山亭19万千瓦风电项目")
    rendered = Document(BytesIO(content))
    assert len(rendered.tables) == 2

    layout = _docx_body_layout(content)
    first_table_index = next(index for index, item in enumerate(layout) if item["kind"] == "tbl")
    second_table_index = next(index for index, item in enumerate(layout[first_table_index + 1 :], start=first_table_index + 1) if item["kind"] == "tbl")
    between_tables = layout[first_table_index + 1 : second_table_index]

    assert sum(1 for item in between_tables if item["kind"] == "p" and not item["text"]) >= 3
    assert any(item["pageBreak"] for item in between_tables)


def test_render_contract_item_docx_keeps_table_caption_rows_with_image_rows() -> None:
    chunks = split_performance_contract_docx(_contract_with_paired_page_table_docx_bytes(), file_name="枣庄山亭19万千瓦业绩_合同.docx")

    content = render_contract_item_docx(chunks[0], output_title="山东华电枣庄山亭19万千瓦风电项目")
    rendered = Document(BytesIO(content))
    rows = _docx_table_row_pagination(content)

    assert len(rendered.tables) == 1
    assert len(rows) == 4
    assert "合同首页" in rows[0]["text"]
    assert "主要技术参数页" in rows[2]["text"]
    assert _docx_local_name_count(content, "drawing") == 4


def test_render_contract_item_docx_uses_short_title_and_table_spacing() -> None:
    chunks = split_performance_contract_docx(_contract_bundle_docx_bytes(), file_name="陆上11MW业绩_合同.docx")
    content = render_contract_item_docx(chunks[0], output_title="华电新疆喀什 2×66 万千瓦")

    rendered = Document(BytesIO(content))
    assert rendered.paragraphs[0].text == "华电新疆喀什 2×66 万千瓦"
    assert all(paragraph.text != "华电新疆喀什 2×66 万千瓦风力发电机组采购合同" for paragraph in rendered.paragraphs)
    assert rendered.tables[0].cell(0, 0).text == "封面页"
    assert ITEM_CONTRACT_FORMAT_VERSION == 15


def test_render_contract_item_docx_sanitizes_fonts_for_text_documents() -> None:
    chunks = split_performance_contract_docx(_contract_with_table_layout_docx_bytes(), file_name="枣庄山亭19万千瓦业绩_合同.docx")
    content = render_contract_item_docx(chunks[0], output_title="山东华电枣庄山亭19万千瓦风电项目")

    fonts = _docx_declared_fonts(content)
    assert "宋体" in fonts
    assert "Times New Roman" in fonts
    assert "方正兰亭超细黑简体" not in fonts
    assert "Songti SC" not in fonts
    assert "Noto Serif CJK SC" not in fonts
    assert "Liberation Serif" not in fonts
    assert "Angsana New" not in fonts
    assert "Mangal" not in fonts
    assert not any("Theme" in font or font.startswith("major") or font.startswith("minor") for font in fonts)
    assert _docx_rfont_non_name_refs(content) == set()
    assert _docx_settings_count(content, "compat") == 0
    assert _docx_settings_count(content, "rsids") == 0


def test_render_contract_item_docx_preserves_drawing_documents_for_onlyoffice() -> None:
    chunks = split_performance_contract_docx(_contract_bundle_docx_bytes(), file_name="陆上11MW业绩_合同.docx")
    content = render_contract_item_docx(chunks[0], output_title="华电新疆喀什 2×66 万千瓦")

    assert _docx_local_name_count(content, "drawing") == 1
    assert _docx_local_name_count(content, "useLocalDpi") == 0
    assert _docx_picture_ids(content) == ["1"]
    assert len(Document(BytesIO(content)).tables) == 1
    Document(BytesIO(content))


def test_render_contract_item_docx_uses_stable_ooxml_prefixes() -> None:
    chunks = split_performance_contract_docx(_contract_with_paired_page_table_docx_bytes(), file_name="陆上7MW业绩_合同.docx")
    content = render_contract_item_docx(chunks[0], output_title="吉林公主岭风电乡村振兴项目")

    with ZipFile(BytesIO(content), "r") as zf:
        document_xml = zf.read("word/document.xml").decode("utf-8")

    assert "<w:document" in document_xml
    assert "<ns0:document" not in document_xml
    assert "mc:Ignorable=\"w14 wp14 w15\"" not in document_xml
    Document(BytesIO(content))


def test_render_contract_item_docx_normalizes_with_soffice() -> None:
    chunks = split_performance_contract_docx(_contract_with_table_layout_docx_bytes(), file_name="陆上11MW业绩_合同.docx")
    calls: list[list[str]] = []

    def fake_run(args: list[str], **_: object) -> CompletedProcess[str]:
        calls.append(args)
        source_path = Path(args[-1])
        outdir = Path(args[args.index("--outdir") + 1])
        outdir.mkdir(parents=True, exist_ok=True)
        (outdir / "source.docx").write_bytes(source_path.read_bytes())
        return CompletedProcess(args, 0, stdout="ok", stderr="")

    with (
        patch("app.services.performance_package_service.shutil.which", return_value="/usr/bin/soffice"),
        patch("app.services.performance_package_service.subprocess.run", side_effect=fake_run),
    ):
        content = render_contract_item_docx(chunks[0], output_title="山东华电枣庄山亭19万千瓦风电项目")

    Document(BytesIO(content))
    assert calls
    assert "--convert-to" in calls[0]


def test_render_contract_item_docx_keeps_original_when_soffice_fails() -> None:
    chunks = split_performance_contract_docx(_contract_with_table_layout_docx_bytes(), file_name="陆上11MW业绩_合同.docx")

    with (
        patch("app.services.performance_package_service.shutil.which", return_value="/usr/bin/soffice"),
        patch(
            "app.services.performance_package_service.subprocess.run",
            side_effect=CalledProcessError(1, ["soffice"], stderr="boom"),
        ),
    ):
        content = render_contract_item_docx(chunks[0], output_title="山东华电枣庄山亭19万千瓦风电项目")

    Document(BytesIO(content))
    assert _docx_settings_count(content, "compat") == 0


def test_ensure_contract_docx_file_name_accepts_docx() -> None:
    ensure_contract_docx_file_name("陆上11MW业绩_合同.docx")
    ensure_contract_docx_file_name("A/B 合同.DOCX")


def test_ensure_contract_docx_file_name_rejects_legacy_doc() -> None:
    with unittest.TestCase().assertRaises(PeripheralError) as ctx:
        ensure_contract_docx_file_name("陆上11MW业绩_合同.doc")

    assert ctx.exception.status_code == 400
    assert ctx.exception.code == "PERFORMANCE_CONTRACT_DOC_NOT_SUPPORTED"
    assert ".docx" in ctx.exception.detail


def test_match_contract_chunks_to_items_prefers_project_name_then_row_order() -> None:
    parsed = parse_performance_summary_docx(_summary_docx_bytes(), file_name="陆上11MW业绩_汇总表.docx")
    chunks = split_performance_contract_docx(_contract_bundle_docx_bytes(), file_name="陆上11MW业绩_合同.docx")
    items = [
        {
            "id": f"PERITEM-{index:04d}",
            "categoryId": "PERCAT-0001",
            **row,
        }
        for index, row in enumerate(parsed["rows"], start=1)
    ]

    matches = match_contract_chunks_to_items(chunks, items)

    assert len(matches) == 2
    assert [match["item"]["projectName"] for match in matches] == [
        "华电新疆喀什 2×66 万千瓦",
        "内蒙古能源察右前旗50万千瓦风光实验实证项目",
    ]
    assert [match["chunk"]["title"] for match in matches] == [
        "华电新疆喀什 2×66 万千瓦风力发电机组采购合同",
        "内蒙古能源察右前旗50万千瓦风光实验实证项目风力发电机组合同",
    ]
    assert all(match["method"] == "project_name" for match in matches)


def _partner_summary_docx_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("陆上5MW业绩")
    doc.add_paragraph("合同业绩台数共计2台，已通过240小时试运行台数共计2台，投运容量共计10MW。")
    table = doc.add_table(rows=2, cols=6)
    headers = [
        "序号",
        "型号",
        "项目名称",
        "买方名称",
        "项目合作方单位",
        "合同台数",
    ]
    rows = [
        ["1", "5-200", "张北某风电项目", "某能源集团有限公司", "某勘测设计院有限公司", "2"],
    ]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    for row_index, row in enumerate(rows, start=1):
        for column_index, value in enumerate(row):
            table.cell(row_index, column_index).text = value
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def test_parse_performance_summary_docx_maps_partner_name() -> None:
    parsed = parse_performance_summary_docx(_partner_summary_docx_bytes(), file_name="陆上5MW业绩_汇总表.docx")

    assert parsed["rowCount"] == 1
    assert parsed["rows"][0]["projectName"] == "张北某风电项目"
    assert parsed["rows"][0]["customerName"] == "某能源集团有限公司"
    assert parsed["rows"][0]["partnerName"] == "某勘测设计院有限公司"


def test_core_values_partner_name_aliases() -> None:
    from app.services.performance_package_service import _core_values

    assert _core_values({"项目合作方单位": "甲公司"})["partnerName"] == "甲公司"
    assert _core_values({"合作方单位": "乙公司"})["partnerName"] == "乙公司"
    assert _core_values({"合作方": "丙公司"})["partnerName"] == "丙公司"
    assert _core_values({"买方名称": "丁公司"})["partnerName"] == ""


class _FakePerformanceResult:
    def __init__(self, row: dict | None = None, *, rows: list[dict] | None = None, scalar: int | None = None) -> None:
        self.row = row
        self.rows = rows
        self.scalar = scalar

    def first(self):
        if self.row is None:
            return None
        return SimpleNamespace(_mapping=self.row)

    def scalar_one(self):
        return self.scalar if self.scalar is not None else 0

    def __iter__(self):
        return iter([SimpleNamespace(_mapping=row) for row in (self.rows if self.rows is not None else [])])


class _FakePerformanceSession:
    def __init__(
        self,
        row: dict | None = None,
        *,
        rows: list[dict] | None = None,
        scalar: int | None = None,
        results: list[_FakePerformanceResult] | None = None,
    ) -> None:
        self.row = row
        self.rows = rows
        self.scalar = scalar
        self.results = list(results or [])
        self.statements: list[str] = []
        self.params: list[dict] = []
        self.committed = False

    async def __aenter__(self):
        return self

    async def __aexit__(self, exc_type, exc, tb):
        return False

    async def execute(self, statement, params=None):
        self.statements.append(str(statement))
        self.params.append(dict(params or {}))
        if self.results:
            return self.results.pop(0)
        return _FakePerformanceResult(self.row, rows=self.rows, scalar=self.scalar)

    async def commit(self):
        self.committed = True


class PerformanceItemUpdateTests(unittest.IsolatedAsyncioTestCase):
    def _item_row(self, **overrides) -> dict:
        row = {
            "id": 2,
            "category_id": 1,
            "row_index": 1,
            "project_name": "原项目",
            "customer_name": "原买方",
            "partner_name": "原合作方",
            "turbine_model": "EW5.0-200",
            "turbine_models": ["EW5.0-200"],
            "contract_quantity": "10",
            "trial_operation_quantity": "8",
            "commissioned_capacity_mw": "50",
            "delivery_or_operation_time": "2021年",
            "contract_year": 2020,
            "delivery_year": 2021,
            "operation_year": 2021,
            "time_facts": {
                "contractTimeRaw": "2020年",
                "deliveryOrOperationTimeRaw": "2021年",
                "contractYear": 2020,
                "deliveryYear": 2021,
                "operationYear": 2021,
                "years": [2020, 2021],
            },
            "contact_info": "张三13800000000",
            "row_values": {"项目名称": "原项目"},
            "created_at": None,
            "updated_at": None,
        }
        row.update(overrides)
        return row

    async def test_update_item_updates_fields_and_derived_payloads(self) -> None:
        service = PerformancePackageService()
        existing_values = {
            "项目名称/合同名称": "原项目",
            "型号和规格": "EW5.0-200",
            "合同时间": "2020年",
            "交货期/投运时间": "2021年",
            "自定义备注": "保留",
        }
        updated_values = {
            "项目名称/合同名称": "新项目",
            "型号和规格": "EW6.25-202",
            "合同时间": "2022年",
            "交货期/投运时间": "2021年",
            "自定义备注": "保留",
            "项目合作方单位": "新合作方",
        }
        updated_time_facts = {
            "contractTimeRaw": "2022年",
            "deliveryTimeRaw": "2021年",
            "operationTimeRaw": "2021年",
            "deliveryOrOperationTimeRaw": "2021年",
            "contractYear": 2022,
            "deliveryYear": 2021,
            "operationYear": 2021,
            "years": [2022, 2021],
        }
        session = _FakePerformanceSession(
            results=[
                _FakePerformanceResult(row=self._item_row(row_values=existing_values)),
                _FakePerformanceResult(
                    row=self._item_row(
                        project_name="新项目",
                        partner_name="新合作方",
                        turbine_model="EW6.25-202",
                        turbine_models=["EW6.25-202"],
                        contract_year=2022,
                        row_values=updated_values,
                        time_facts=updated_time_facts,
                    )
                ),
            ]
        )
        with patch("app.services.performance_package_service.async_session", return_value=session), patch(
            "app.services.performance_package_service.ensure_material_runtime_tables", new=AsyncMock()
        ):
            result = await service.update_item(
                "PERCAT-0001",
                "PERITEM-0002",
                {
                    "projectName": " 新项目 ",
                    "partnerName": " 新合作方 ",
                    "turbineModel": "EW6.25-202",
                    "contractYear": "2022年",
                },
            )

        self.assertIn("SELECT *", session.statements[0])
        self.assertIn("UPDATE performance_items", session.statements[1])
        self.assertIn("row_values = CAST(:row_values AS JSONB)", session.statements[1])
        self.assertIn("RETURNING *", session.statements[1])
        self.assertIn("UPDATE performance_categories SET updated_at = NOW()", session.statements[2])
        update_params = session.params[1]
        self.assertEqual(update_params["project_name"], "新项目")
        self.assertEqual(update_params["partner_name"], "新合作方")
        self.assertEqual(update_params["turbine_model"], "EW6.25-202")
        self.assertEqual(update_params["contract_year"], 2022)
        self.assertEqual(json.loads(update_params["turbine_models"]), ["EW6.25-202"])
        row_values = json.loads(update_params["row_values"])
        self.assertEqual(row_values, updated_values)
        self.assertNotIn("项目名称", row_values)
        self.assertNotIn("型号", row_values)
        self.assertNotIn("原项目", update_params["row_values"])
        self.assertNotIn("EW5.0-200", update_params["row_values"])
        self.assertNotIn("2020", update_params["row_values"])
        time_facts = json.loads(update_params["time_facts"])
        self.assertEqual(time_facts, updated_time_facts)
        self.assertTrue(session.committed)
        self.assertEqual(result["message"], "业绩明细已更新")
        self.assertEqual(result["item"]["id"], "PERITEM-0002")
        self.assertEqual(result["item"]["partnerName"], "新合作方")
        self.assertEqual(result["item"]["contractYear"], 2022)
        self.assertEqual(result["item"]["values"], updated_values)

    async def test_update_item_repairs_existing_row_value_drift(self) -> None:
        service = PerformancePackageService()
        existing_values = {
            "项目名称/合同名称": "旧项目",
            "交货期/投运时间": "2021年",
            "备注": "保留",
        }
        updated_values = {
            "项目名称/合同名称": "新项目",
            "交货期/投运时间": "2021年",
            "备注": "保留",
        }
        existing = self._item_row(
            project_name="新项目",
            row_values=existing_values,
        )
        session = _FakePerformanceSession(
            results=[
                _FakePerformanceResult(row=existing),
                _FakePerformanceResult(
                    row=self._item_row(
                        project_name="新项目",
                        row_values=updated_values,
                    )
                ),
            ]
        )
        with patch("app.services.performance_package_service.async_session", return_value=session), patch(
            "app.services.performance_package_service.ensure_material_runtime_tables", new=AsyncMock()
        ):
            result = await service.update_item(
                "PERCAT-0001",
                "PERITEM-0002",
                {"projectName": "新项目"},
            )

        self.assertIn("row_values = CAST(:row_values AS JSONB)", session.statements[1])
        self.assertEqual(json.loads(session.params[1]["row_values"]), updated_values)
        self.assertNotIn("旧项目", session.params[1]["row_values"])
        self.assertNotIn("time_facts", session.params[1])
        self.assertEqual(result["item"]["values"], updated_values)

    async def test_update_item_full_payload_does_not_copy_total_capacity_into_commissioned_capacity(self) -> None:
        service = PerformancePackageService()
        existing_values = {
            "型号和规格": "EW8.5-230",
            "项目名称": "国华山东B2场址项目",
            "总容量（MW)": "501.5",
            "台数": "59",
            "买方名称": "国家能源投资集团（济南）新能源有限责任公司",
            "合同时间": "2022年",
            "投运时间": "/",
            "投运容量（MW)": "/",
            "联系人及电话": "刘欢18615226720",
        }
        updated_values = {**existing_values, "项目名称": "更正后的国华山东B2场址项目"}
        existing = self._item_row(
            project_name="国华山东B2场址项目",
            customer_name="国家能源投资集团（济南）新能源有限责任公司",
            partner_name="",
            turbine_model="EW8.5-230",
            turbine_models=["EW8.5-230"],
            contract_quantity="59",
            trial_operation_quantity="",
            commissioned_capacity_mw="501.5",
            delivery_or_operation_time="/",
            contract_year=2022,
            delivery_year=None,
            operation_year=None,
            contact_info="刘欢18615226720",
            row_values=existing_values,
        )
        session = _FakePerformanceSession(
            results=[
                _FakePerformanceResult(row=existing),
                _FakePerformanceResult(
                    row={**existing, "project_name": "更正后的国华山东B2场址项目", "row_values": updated_values}
                ),
            ]
        )
        with patch("app.services.performance_package_service.async_session", return_value=session), patch(
            "app.services.performance_package_service.ensure_material_runtime_tables", new=AsyncMock()
        ):
            await service.update_item(
                "PERCAT-0001",
                "PERITEM-0002",
                {
                    "projectName": "更正后的国华山东B2场址项目",
                    "customerName": existing["customer_name"],
                    "partnerName": existing["partner_name"],
                    "turbineModel": existing["turbine_model"],
                    "contractQuantity": existing["contract_quantity"],
                    "trialOperationQuantity": existing["trial_operation_quantity"],
                    "commissionedCapacityMw": existing["commissioned_capacity_mw"],
                    "deliveryOrOperationTime": existing["delivery_or_operation_time"],
                    "contractYear": existing["contract_year"],
                    "deliveryYear": existing["delivery_year"],
                    "operationYear": existing["operation_year"],
                    "contactInfo": existing["contact_info"],
                },
            )

        row_values = json.loads(session.params[1]["row_values"])
        self.assertEqual(row_values, updated_values)
        self.assertEqual(row_values["总容量（MW)"], "501.5")
        self.assertEqual(row_values["投运容量（MW)"], "/")

    async def test_update_item_does_not_overwrite_total_capacity(self) -> None:
        service = PerformancePackageService()
        existing_values = {"总容量（MW)": "200", "投运容量（MW)": "200", "备注": "保留"}
        updated_values = {"总容量（MW)": "200", "投运容量（MW)": "180", "备注": "保留"}
        session = _FakePerformanceSession(
            results=[
                _FakePerformanceResult(row=self._item_row(commissioned_capacity_mw="200", row_values=existing_values)),
                _FakePerformanceResult(row=self._item_row(commissioned_capacity_mw="180", row_values=updated_values)),
            ]
        )
        with patch("app.services.performance_package_service.async_session", return_value=session), patch(
            "app.services.performance_package_service.ensure_material_runtime_tables", new=AsyncMock()
        ):
            await service.update_item("PERCAT-0001", "PERITEM-0002", {"commissionedCapacityMw": "180"})

        self.assertEqual(json.loads(session.params[1]["row_values"]), updated_values)

    async def test_update_item_clears_time_row_value_and_stale_year(self) -> None:
        service = PerformancePackageService()
        existing_values = {"交货时间": "预计2021年", "备注": "保留"}
        updated_values = {"交货时间": "", "备注": "保留"}
        empty_time_facts = {
            "contractTimeRaw": "",
            "deliveryTimeRaw": "",
            "operationTimeRaw": "",
            "deliveryOrOperationTimeRaw": "",
            "contractYear": None,
            "deliveryYear": None,
            "operationYear": None,
            "years": [],
        }
        existing = self._item_row(
            contract_year=None,
            delivery_or_operation_time="预计2021年",
            delivery_year=2021,
            operation_year=None,
            row_values=existing_values,
            time_facts={
                **empty_time_facts,
                "deliveryTimeRaw": "预计2021年",
                "deliveryOrOperationTimeRaw": "预计2021年",
                "deliveryYear": 2021,
                "years": [2021],
            },
        )
        session = _FakePerformanceSession(
            results=[
                _FakePerformanceResult(row=existing),
                _FakePerformanceResult(
                    row=self._item_row(
                        contract_year=None,
                        delivery_or_operation_time="",
                        delivery_year=None,
                        operation_year=None,
                        row_values=updated_values,
                        time_facts=empty_time_facts,
                    )
                ),
            ]
        )
        with patch("app.services.performance_package_service.async_session", return_value=session), patch(
            "app.services.performance_package_service.ensure_material_runtime_tables", new=AsyncMock()
        ):
            await service.update_item(
                "PERCAT-0001",
                "PERITEM-0002",
                {"deliveryOrOperationTime": "", "deliveryYear": ""},
            )

        update_params = session.params[1]
        self.assertEqual(json.loads(update_params["row_values"]), updated_values)
        self.assertEqual(json.loads(update_params["time_facts"]), empty_time_facts)
        self.assertNotIn("2021", update_params["row_values"])
        self.assertNotIn("2021", update_params["time_facts"])
        self.assertNotIn("turbine_models", update_params)

    async def test_update_item_rejects_invalid_year(self) -> None:
        service = PerformancePackageService()
        with self.assertRaises(PeripheralError) as context:
            await service.update_item("PERCAT-0001", "PERITEM-0002", {"contractYear": "abc"})
        self.assertEqual(context.exception.status_code, 400)
        self.assertEqual(context.exception.code, "PERFORMANCE_ITEM_YEAR_INVALID")

    async def test_update_item_rejects_unknown_or_empty_fields(self) -> None:
        service = PerformancePackageService()
        with self.assertRaises(PeripheralError) as unknown_context:
            await service.update_item("PERCAT-0001", "PERITEM-0002", {"status": "enabled"})
        self.assertEqual(unknown_context.exception.status_code, 400)
        self.assertEqual(unknown_context.exception.code, "PERFORMANCE_ITEM_FIELDS_INVALID")
        with self.assertRaises(PeripheralError) as empty_context:
            await service.update_item("PERCAT-0001", "PERITEM-0002", {})
        self.assertEqual(empty_context.exception.status_code, 400)
        self.assertEqual(empty_context.exception.code, "PERFORMANCE_ITEM_UPDATE_EMPTY")

    async def test_update_item_raises_not_found_for_missing_row(self) -> None:
        service = PerformancePackageService()
        session = _FakePerformanceSession(results=[_FakePerformanceResult(row=None)])
        with patch("app.services.performance_package_service.async_session", return_value=session), patch(
            "app.services.performance_package_service.ensure_material_runtime_tables", new=AsyncMock()
        ):
            with self.assertRaises(PeripheralError) as context:
                await service.update_item("PERCAT-0001", "PERITEM-9999", {"partnerName": "新合作方"})
        self.assertEqual(context.exception.status_code, 404)
        self.assertEqual(context.exception.code, "PERFORMANCE_ITEM_NOT_FOUND")


class PerformanceItemCreateTests(unittest.IsolatedAsyncioTestCase):
    def _created_row(self, **overrides) -> dict:
        row = {
            "id": 7,
            "category_id": 1,
            "row_index": 4,
            "project_name": "新项目",
            "customer_name": "",
            "partner_name": "",
            "turbine_model": "EW6.25-202",
            "turbine_models": ["EW6.25-202"],
            "contract_quantity": "",
            "trial_operation_quantity": "",
            "commissioned_capacity_mw": "",
            "delivery_or_operation_time": "",
            "contract_year": 2022,
            "delivery_year": 2021,
            "operation_year": None,
            "time_facts": {},
            "contact_info": "",
            "row_values": {},
            "created_at": None,
            "updated_at": None,
        }
        row.update(overrides)
        return row

    async def test_create_item_inserts_with_derived_payloads(self) -> None:
        service = PerformancePackageService()
        expected_row_values = {
            "项目名称": "新项目",
            "型号": "EW6.25-202",
            "合同时间": "2022年",
            "交货时间": "2021年",
        }
        expected_time_facts = {
            "contractTimeRaw": "2022年",
            "deliveryTimeRaw": "2021年",
            "operationTimeRaw": "",
            "deliveryOrOperationTimeRaw": "2021年",
            "contractYear": 2022,
            "deliveryYear": 2021,
            "operationYear": None,
            "years": [2022, 2021],
        }
        session = _FakePerformanceSession(
            results=[
                _FakePerformanceResult(row={"id": 1}),
                _FakePerformanceResult(scalar=4),
                _FakePerformanceResult(row=self._created_row()),
                _FakePerformanceResult(),
            ]
        )
        with patch("app.services.performance_package_service.async_session", return_value=session), patch(
            "app.services.performance_package_service.ensure_material_runtime_tables", new=AsyncMock()
        ):
            result = await service.create_item(
                "PERCAT-0001",
                {
                    "projectName": " 新项目 ",
                    "turbineModel": "EW6.25-202",
                    "contractYear": "2022年",
                    "deliveryYear": 2021,
                },
            )

        self.assertIn("FROM performance_categories", session.statements[0])
        self.assertIn("MAX(row_index)", session.statements[1])
        self.assertIn("INSERT INTO performance_items", session.statements[2])
        self.assertIn("RETURNING *", session.statements[2])
        self.assertIn("UPDATE performance_categories SET updated_at = NOW()", session.statements[3])
        insert_params = session.params[2]
        self.assertEqual(insert_params["category_id"], 1)
        self.assertEqual(insert_params["row_index"], 4)
        self.assertEqual(insert_params["project_name"], "新项目")
        self.assertEqual(insert_params["turbine_model"], "EW6.25-202")
        self.assertEqual(insert_params["contract_year"], 2022)
        self.assertEqual(insert_params["delivery_year"], 2021)
        self.assertIsNone(insert_params["operation_year"])
        self.assertEqual(json.loads(insert_params["turbine_models"]), ["EW6.25-202"])
        self.assertEqual(json.loads(insert_params["row_values"]), expected_row_values)
        self.assertEqual(json.loads(insert_params["time_facts"]), expected_time_facts)
        self.assertTrue(session.committed)
        self.assertEqual(result["message"], "业绩明细已新增")
        self.assertEqual(result["item"]["id"], "PERITEM-0007")
        self.assertEqual(result["item"]["rowIndex"], 4)
        self.assertEqual(result["item"]["projectName"], "新项目")

    async def test_create_item_rejects_unknown_fields(self) -> None:
        service = PerformancePackageService()
        with self.assertRaises(PeripheralError) as context:
            await service.create_item("PERCAT-0001", {"status": "enabled"})
        self.assertEqual(context.exception.status_code, 400)
        self.assertEqual(context.exception.code, "PERFORMANCE_ITEM_FIELDS_INVALID")

    async def test_create_item_rejects_empty_content(self) -> None:
        service = PerformancePackageService()
        with self.assertRaises(PeripheralError) as context:
            await service.create_item("PERCAT-0001", {})
        self.assertEqual(context.exception.status_code, 400)
        self.assertEqual(context.exception.code, "PERFORMANCE_ITEM_CREATE_EMPTY")

    async def test_create_item_rejects_invalid_year(self) -> None:
        service = PerformancePackageService()
        with self.assertRaises(PeripheralError) as context:
            await service.create_item("PERCAT-0001", {"contractYear": "abc"})
        self.assertEqual(context.exception.status_code, 400)
        self.assertEqual(context.exception.code, "PERFORMANCE_ITEM_YEAR_INVALID")

    async def test_create_item_raises_not_found_for_missing_category(self) -> None:
        service = PerformancePackageService()
        session = _FakePerformanceSession(results=[_FakePerformanceResult(row=None)])
        with patch("app.services.performance_package_service.async_session", return_value=session), patch(
            "app.services.performance_package_service.ensure_material_runtime_tables", new=AsyncMock()
        ):
            with self.assertRaises(PeripheralError) as context:
                await service.create_item("PERCAT-9999", {"projectName": "新项目"})
        self.assertEqual(context.exception.status_code, 404)
        self.assertEqual(context.exception.code, "PERFORMANCE_CATEGORY_NOT_FOUND")


class PerformanceItemDeleteTests(unittest.IsolatedAsyncioTestCase):
    async def test_delete_item_removes_row_attachments_and_minio_objects(self) -> None:
        service = PerformancePackageService()
        session = _FakePerformanceSession(
            results=[
                _FakePerformanceResult(rows=[{"minio_bucket": "mat-bucket", "minio_key": "perf/item-2.docx"}]),
                _FakePerformanceResult(row={"id": 2}),
                _FakePerformanceResult(),
            ]
        )
        with patch("app.services.performance_package_service.async_session", return_value=session), patch(
            "app.services.performance_package_service.ensure_material_runtime_tables", new=AsyncMock()
        ), patch("app.services.performance_package_service.minio_client") as mock_minio:
            result = await service.delete_item("PERCAT-0001", "PERITEM-0002")

        self.assertIn("DELETE FROM performance_item_attachments", session.statements[0])
        self.assertIn("RETURNING minio_bucket, minio_key", session.statements[0])
        self.assertIn("DELETE FROM performance_items", session.statements[1])
        self.assertIn("UPDATE performance_categories SET updated_at = NOW()", session.statements[2])
        self.assertEqual(session.params[0], {"item_id": 2, "category_id": 1})
        mock_minio.remove_object.assert_called_once_with("mat-bucket", "perf/item-2.docx")
        self.assertTrue(session.committed)
        self.assertEqual(result["message"], "业绩明细已删除")
        self.assertEqual(result["id"], "PERITEM-0002")

    async def test_delete_item_raises_not_found_for_missing_row(self) -> None:
        service = PerformancePackageService()
        session = _FakePerformanceSession(
            results=[
                _FakePerformanceResult(rows=[]),
                _FakePerformanceResult(row=None),
            ]
        )
        with patch("app.services.performance_package_service.async_session", return_value=session), patch(
            "app.services.performance_package_service.ensure_material_runtime_tables", new=AsyncMock()
        ), patch("app.services.performance_package_service.minio_client") as mock_minio:
            with self.assertRaises(PeripheralError) as context:
                await service.delete_item("PERCAT-0001", "PERITEM-9999")
        self.assertEqual(context.exception.status_code, 404)
        self.assertEqual(context.exception.code, "PERFORMANCE_ITEM_NOT_FOUND")
        mock_minio.remove_object.assert_not_called()
