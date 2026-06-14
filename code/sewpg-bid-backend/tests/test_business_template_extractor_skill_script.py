import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from app.services import business_template_extractor
from app.services.business_template_extractor import (
    build_business_template_extractor_manifest,
    convert_extractor_appendices,
    run_business_template_extractor,
)


ROOT = Path(__file__).resolve().parents[1]
RUNNER = ROOT / "opencode" / "skill" / "bid-business-template-extractor" / "scripts" / "run_from_manifest.py"
SKILL_DIR = ROOT / "opencode" / "skill" / "bid-business-template-extractor"
SKILL_MD = SKILL_DIR / "SKILL.md"


def build_business_format_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("第一章 招标公告")
    doc.add_paragraph("这里不是模板。")
    doc.add_page_break()
    doc.add_paragraph("第六章 投标文件格式")
    doc.add_paragraph("投标函")
    doc.add_paragraph("致：招标人")
    doc.add_paragraph("投标人（盖章）：")
    doc.add_page_break()
    doc.add_paragraph("法定代表人授权委托书")
    doc.add_paragraph("委托代理人姓名：")
    doc.add_paragraph("身份证号码：")
    doc.save(path)


def write_manifest(path: Path, *, source: Path, output_dir: Path) -> None:
    path.write_text(
        json.dumps(
            {
                "schemaVersion": "bid-business-template-extractor-v1",
                "skillName": "bid-business-template-extractor",
                "projectId": "PRJ-TPL",
                "outputDir": str(output_dir),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": source.name,
                        "sourcePath": str(source),
                    }
                ],
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )


def run_btplnav(*args: object) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, str(RUNNER), *(str(arg) for arg in args)],
        cwd=str(ROOT),
        text=True,
        encoding="utf-8",
        errors="replace",
        capture_output=True,
        check=False,
    )


def stdout_json(completed: subprocess.CompletedProcess[str]) -> dict:
    if completed.returncode != 0:
        raise AssertionError(completed.stderr or completed.stdout)
    return json.loads(completed.stdout)


def find_block_id(blocks: list[dict], text: str) -> int:
    return next(int(block["blockId"]) for block in blocks if text in str(block.get("text") or ""))


class BusinessTemplateExtractorSkillScriptTests(unittest.TestCase):
    def test_runner_requires_btplnav_subcommand_and_blocks_legacy_pipeline_entry(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "source.docx"
            manifest = root / "manifest.json"
            output_dir = root / "out"
            build_business_format_docx(source)
            write_manifest(manifest, source=source, output_dir=output_dir)

            completed = run_btplnav(manifest)

            self.assertEqual(completed.returncode, 64)
            self.assertIn("usage: run_from_manifest.py", completed.stderr)
            self.assertFalse((output_dir / "candidate_templates.json").exists())

    def test_btplnav_prepare_outputs_navigation_index_only(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "source.docx"
            manifest = root / "manifest.json"
            output_dir = root / "out"
            build_business_format_docx(source)
            write_manifest(manifest, source=source, output_dir=output_dir)

            prepared = stdout_json(run_btplnav("prepare", manifest))
            document_output = Path(prepared["documents"][0]["outputDir"])

            self.assertEqual(prepared["stage"], "prepared")
            self.assertTrue((output_dir / "template_nav.json").is_file())
            self.assertTrue((document_output / "blocks.json").is_file())
            self.assertFalse((document_output / "candidate_templates.json").exists())
            self.assertFalse((document_output / "regions.json").exists())

    def test_btplnav_submit_validate_and_finalize_slice_ai_ranges(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "source.docx"
            manifest = root / "manifest.json"
            output_dir = root / "out"
            build_business_format_docx(source)
            write_manifest(manifest, source=source, output_dir=output_dir)

            stdout_json(run_btplnav("prepare", manifest))
            overview = stdout_json(run_btplnav("overview", manifest, "--page", "1", "--page-size", "80"))
            blocks = overview["blocks"]
            bid_start = find_block_id(blocks, "投标函")
            bid_end = find_block_id(blocks, "投标人（盖章）")
            auth_start = find_block_id(blocks, "法定代表人授权委托书")
            auth_end = find_block_id(blocks, "身份证号码")

            stdout_json(
                run_btplnav(
                    "submit",
                    manifest,
                    "templates",
                    json.dumps(
                        {
                            "templates": [
                                {
                                    "sourceDocumentId": "DOC-1",
                                    "title": "投标函",
                                    "templateType": "bid_letter",
                                    "startBlockId": bid_start,
                                    "endBlockId": bid_end,
                                    "confidence": 0.96,
                                    "reason": "AI 判断为需要填写并盖章的投标函模板。",
                                },
                                {
                                    "sourceDocumentId": "DOC-1",
                                    "title": "法定代表人授权委托书",
                                    "templateType": "authorization_letter",
                                    "startBlockId": auth_start,
                                    "endBlockId": auth_end,
                                    "confidence": 0.94,
                                    "reason": "AI 判断为需要填写委托代理人身份信息的授权模板。",
                                },
                            ]
                        },
                        ensure_ascii=False,
                    ),
                )
            )
            validation = stdout_json(run_btplnav("validate", manifest))
            finalized = stdout_json(run_btplnav("finalize", manifest))

            self.assertEqual(validation["status"], "passed")
            self.assertEqual(finalized["summary"]["templateCount"], 2)
            payload = json.loads((output_dir / "business_template_extraction.json").read_text(encoding="utf-8"))
            self.assertEqual([item["title"] for item in payload["appendices"]], ["投标函", "法定代表人授权委托书"])
            self.assertFalse(payload["quality"]["scriptFallbackUsed"])
            self.assertEqual(payload["quality"]["agentSubmittedTemplateCount"], 2)
            self.assertEqual(payload["quality"]["structurallyValidatedTemplateCount"], 2)
            self.assertEqual(payload["quality"]["slicedTemplateCount"], 2)
            for legacy_key in (
                "formatRegionCount",
                "candidateAnchorCount",
                "candidateTemplateCount",
                "boundaryReferenceCount",
                "sectionContainerCount",
                "boundaryOnlyCount",
            ):
                self.assertNotIn(legacy_key, payload["quality"])
            for appendix in payload["appendices"]:
                self.assertTrue(Path(appendix["docxPath"]).is_file())

    def test_btplnav_validator_rejects_structurally_invalid_ranges_without_section_rules(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source = root / "source.docx"
            manifest = root / "manifest.json"
            output_dir = root / "out"
            build_business_format_docx(source)
            write_manifest(manifest, source=source, output_dir=output_dir)

            stdout_json(run_btplnav("prepare", manifest))
            stdout_json(
                run_btplnav(
                    "submit",
                    manifest,
                    "templates",
                    json.dumps(
                        {
                            "templates": [
                                {
                                    "sourceDocumentId": "DOC-1",
                                    "title": "任意商务承诺页",
                                    "startBlockId": 9999,
                                    "endBlockId": 10000,
                                    "confidence": 0.75,
                                    "reason": "测试结构校验，不依赖章节名白名单。",
                                }
                            ]
                        },
                        ensure_ascii=False,
                    ),
                )
            )
            validation = stdout_json(run_btplnav("validate", manifest))

            self.assertEqual(validation["status"], "failed")
            self.assertEqual(validation["validationErrors"][0]["code"], "unknown_block_id")


class BusinessTemplateExtractorWrapperTests(unittest.TestCase):
    def test_build_manifest_keeps_only_docx_sources_and_output_dir(self) -> None:
        output_dir = Path("C:/tmp/business-template-output")
        manifest = build_business_template_extractor_manifest(
            project_id="proj-1",
            documents=[
                {"id": "DOC-1", "name": "招标.docx", "sourcePath": "C:/tmp/招标.docx"},
                {"id": "DOC-2", "name": "说明.txt", "sourcePath": "C:/tmp/说明.txt"},
            ],
            output_dir=output_dir,
        )

        self.assertEqual(manifest["projectId"], "proj-1")
        self.assertEqual(manifest["outputDir"], str(output_dir))
        self.assertEqual(manifest["stage"], "prepare")
        self.assertNotIn("fallbackMode", manifest)
        self.assertEqual(len(manifest["documents"]), 1)
        self.assertEqual(manifest["documents"][0]["id"], "DOC-1")

    def test_backend_service_does_not_expose_legacy_btplbound_orchestration(self) -> None:
        removed_names = {
            "btplbound_runner_path",
            "_run_btplbound_command",
            "_run_btplbound_agent_phase",
            "_run_btplbound_backend_agent_decisions",
            "build_business_template_batch_decision_prompt",
            "build_business_template_boundary_decision_prompt",
            "_missing_decision_paths",
        }

        exposed_names = set(dir(business_template_extractor))

        self.assertTrue(removed_names.isdisjoint(exposed_names))
        self.assertTrue(hasattr(business_template_extractor, "build_business_template_navigation_prompt"))
        self.assertTrue(hasattr(business_template_extractor, "run_business_template_extractor"))

    def test_template_extractor_skill_document_is_principle_driven_not_case_rule_driven(self) -> None:
        skill_text = SKILL_MD.read_text(encoding="utf-8")

        self.assertIn("招投标专家", skill_text)
        self.assertIn("AI 负责业务判断", skill_text)
        self.assertIn("文档浏览器", skill_text)
        self.assertIn("Word 切片器", skill_text)
        self.assertIn("结构校验器", skill_text)
        self.assertIn("需要填写、粘贴材料或签章的完整模板单元", skill_text)
        self.assertIn("总体原则", skill_text)
        self.assertIn("封面", skill_text)
        self.assertIn("第一个有意义标题", skill_text)
        self.assertIn("btplnav prepare", skill_text)
        self.assertIn("btplnav submit", skill_text)
        self.assertIn("btplnav finalize", skill_text)
        self.assertIn("父标题若承载一组需要整体编制或提交的子表、材料或附件，优先作为一个模板", skill_text)
        self.assertIn("子项只有在脱离父级也必须单独填写、签章或交付时，才拆成独立模板", skill_text)
        self.assertIn("目录页、目录清单、普通说明、合同附件、履约保证金格式、纯噪声不作为模板输出", skill_text)
        for overfit_rule in (
            "candidate-batch",
            "boundary-batch",
            "candidate_templates.json",
            "sub_table_code + near_following_table",
            "表2 E",
            "表3 A",
            "表1 A-1",
            "7D-1表",
            "承诺书/声明函/保密承诺书/保证函格式",
            "裁决规则",
            "template_start",
            "section_container",
        ):
            self.assertNotIn(overfit_rule, skill_text)

    def test_run_extractor_uses_agentic_btplnav_decisions_without_script_fallback(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            project_dir = Path(tmp)
            source = project_dir / "agent-success.docx"
            source.write_bytes(b"fake-docx")
            calls: list[dict[str, object]] = []
            agent_prompts: list[str] = []

            def completed(stdout: str = "{}", returncode: int = 0, stderr: str = ""):
                class Completed:
                    pass

                item = Completed()
                item.returncode = returncode
                item.stdout = stdout
                item.stderr = stderr
                return item

            def fake_subprocess_run(args, **kwargs):  # type: ignore[no-untyped-def]
                command = str(args[2])
                manifest_path = Path(args[3])
                calls.append({"command": command})
                self.assertEqual(command, "prepare")
                manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
                output_dir = Path(manifest["outputDir"])
                document_output = output_dir / "DOC-1"
                document_output.mkdir(parents=True, exist_ok=True)
                (output_dir / "business_template_extraction.json").write_text(
                    json.dumps(
                        {
                            "stage": "prepare",
                            "schemaVersion": "bid-business-template-extractor-v1",
                            "outputDir": str(output_dir),
                            "documents": [{"id": "DOC-1", "outputDir": str(document_output)}],
                            "appendices": [],
                            "quality": {"scriptFallbackUsed": False},
                        },
                        ensure_ascii=False,
                    ),
                    encoding="utf-8",
                )
                return completed()

            def fake_agentic_template_extractor(_client, prompt: str):  # type: ignore[no-untyped-def]
                agent_prompts.append(prompt)
                self.assertIn("btplnav prepare", prompt)
                self.assertIn("btplnav submit", prompt)
                self.assertIn("btplnav finalize", prompt)
                self.assertNotIn("candidate-batch", prompt)
                self.assertNotIn("boundary-batch", prompt)
                output_dir = project_dir / "business_template_extraction"
                document_output = output_dir / "DOC-1"
                document_output.mkdir(parents=True, exist_ok=True)
                (output_dir / "business_template_extraction.json").write_text(
                    json.dumps(
                        {
                            "stage": "finalize",
                            "schemaVersion": "bid-business-template-extractor-v1",
                            "outputDir": str(output_dir),
                            "appendices": [
                                {
                                    "id": "APPX-0001",
                                    "title": "Bid Letter",
                                    "artifactType": "business_attachment_template",
                                    "docxPath": str(document_output / "templates" / "TPL-0001.docx"),
                                    "sourceDocumentId": "DOC-1",
                                }
                            ],
                            "quality": {"scriptFallbackUsed": False},
                        },
                        ensure_ascii=False,
                    ),
                    encoding="utf-8",
                )
                return {
                    "schemaVersion": "bid-business-template-extractor-v1",
                    "outputFile": str(output_dir / "business_template_extraction.json"),
                    "summary": {"templateCount": 1},
                    "opencodeOutput": {"sessionId": "ses-btplnav", "completionSource": "btplnav-finalize"},
                }

            with (
                patch("app.services.business_template_extractor.subprocess.run", side_effect=fake_subprocess_run),
                patch(
                    "app.services.business_template_extractor.OpencodeClient.extract_business_templates_with_trace",
                    new=fake_agentic_template_extractor,
                ),
            ):
                appendices, payload, warning = run_business_template_extractor(
                    project_id="PRJ-1",
                    documents=[{"id": "DOC-1", "name": "agent-success.docx", "sourcePath": str(source)}],
                    project_dir=project_dir,
                )

        self.assertEqual(warning, "")
        self.assertEqual([call["command"] for call in calls], ["prepare"])
        self.assertEqual(len(agent_prompts), 1)
        self.assertEqual(len(appendices), 1)
        self.assertEqual(appendices[0]["extractionMode"], "business_template_extractor_skill")
        self.assertFalse((payload or {})["quality"]["scriptFallbackUsed"])
        self.assertEqual((payload or {})["opencodeOutput"]["completionSource"], "btplnav-finalize")

    def test_run_extractor_records_agent_failure_without_script_fallback(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            project_dir = Path(tmp)
            source = project_dir / "agent-failure.docx"
            source.write_bytes(b"fake-docx")
            calls: list[str] = []

            def fake_subprocess_run(args, **kwargs):  # type: ignore[no-untyped-def]
                command = str(args[2])
                manifest_path = Path(args[3])
                calls.append(command)
                manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
                output_dir = Path(manifest["outputDir"])
                document_output = output_dir / "DOC-1"
                document_output.mkdir(parents=True, exist_ok=True)
                (output_dir / "business_template_extraction.json").write_text(
                    json.dumps(
                        {
                            "stage": "prepare",
                            "schemaVersion": "bid-business-template-extractor-v1",
                            "outputDir": str(output_dir),
                            "documents": [{"id": "DOC-1", "outputDir": str(document_output)}],
                            "appendices": [],
                            "warnings": [],
                            "quality": {"scriptFallbackUsed": False},
                        },
                        ensure_ascii=False,
                    ),
                    encoding="utf-8",
                )

                class Completed:
                    returncode = 0
                    stdout = "{}"
                    stderr = ""

                return Completed()

            with (
                patch("app.services.business_template_extractor.subprocess.run", side_effect=fake_subprocess_run),
                patch(
                    "app.services.business_template_extractor.OpencodeClient.extract_business_templates_with_trace",
                    side_effect=RuntimeError("agent stopped before finalize"),
                ),
            ):
                appendices, payload, warning = run_business_template_extractor(
                    project_id="PRJ-1",
                    documents=[{"id": "DOC-1", "name": "agent-failure.docx", "sourcePath": str(source)}],
                    project_dir=project_dir,
        )

        self.assertEqual(calls, ["prepare"])
        self.assertEqual(appendices, [])
        self.assertIn("商务模板提取 Agent 未完成", warning)
        self.assertIn("agent stopped before finalize", warning)
        self.assertFalse((payload or {})["quality"]["scriptFallbackUsed"])
        self.assertIn("agent stopped before finalize", (payload or {})["quality"]["agentFallbackReason"])
        self.assertTrue(any(item["code"] == "business_template_agent_failed" for item in (payload or {})["warnings"]))

    def test_convert_extractor_appendices_preserves_docx_path_for_prepare_outputs(self) -> None:
        payload = {
            "appendices": [
                {
                    "id": "APPX-0007",
                    "title": "附件2 投标价格表\nA投标价格总表\n表1 A-1  标段一",
                    "artifactType": "business_attachment_template",
                    "templateType": "business_template",
                    "templateSectionTitle": "第六章 投标文件格式",
                    "status": "generated",
                    "docxPath": "C:/tmp/TPL-0001.docx",
                    "sourceDocumentId": "DOC-1",
                    "sourceDocumentName": "招标.docx",
                }
            ]
        }

        appendices = convert_extractor_appendices(payload)

        self.assertEqual(len(appendices), 1)
        self.assertEqual(appendices[0]["id"], "APPX-0007")
        self.assertEqual(appendices[0]["title"], "附件2 投标价格表\nA投标价格总表\n表1 A-1  标段一")
        self.assertEqual(appendices[0]["artifactType"], "business_attachment_template")
        self.assertEqual(appendices[0]["extractionMode"], "business_template_extractor_skill")
        self.assertEqual(appendices[0]["docxPath"], "C:/tmp/TPL-0001.docx")
        self.assertEqual(appendices[0]["workspacePath"], "")
        self.assertEqual(appendices[0]["sourceFile"], "招标.docx")


if __name__ == "__main__":
    unittest.main()
