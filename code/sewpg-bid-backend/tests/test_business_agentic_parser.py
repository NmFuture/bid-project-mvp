from __future__ import annotations

import json
import os
import sqlite3
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


class BusinessAgenticParserTests(unittest.TestCase):
    def runner_path(self) -> Path:
        return (
            Path(__file__).resolve().parents[1]
            / "opencode"
            / "skill"
            / "bid-business-tender-structured-parser"
            / "scripts"
            / "run_from_manifest.py"
        )

    def write_sample_docx(self, root: Path) -> tuple[Path, Path, Path]:
        source_path = root / "business_tender.docx"
        output_path = root / "s1_structured_result.json"
        manifest_path = root / "s1_parse_manifest.json"
        doc = Document()
        doc.add_paragraph("项目编号：PC-0307-26J1-FG0002")
        doc.add_paragraph("都匀市盛黔风电场风力发电机组及附属设备采购项目")
        doc.add_paragraph("谈判采购供应商须知前附表")
        instruction_table = doc.add_table(rows=3, cols=3)
        for col, text in enumerate(["条款号", "条款名称", "编列内容"]):
            instruction_table.cell(0, col).text = text
        for col, text in enumerate(["3.3", "报价", "最高限价：最终报价不得超过人民币8250万元。"]):
            instruction_table.cell(1, col).text = text
        for col, text in enumerate(["3.7", "响应文件份数", "响应文件开启后30分钟内发送可编辑版本。"]):
            instruction_table.cell(2, col).text = text
        doc.add_paragraph("响应文件递交的截止时间为2026年05月06日10:00。")
        doc.add_paragraph("供应商资格要求：已成为中国电建集团2025年度风力发电机组框架入围集中采购项目的入围供应商。")
        doc.add_paragraph("响应文件符合性审查，有以下情形之一的，为重大偏差，其响应文件将被视为无效。")
        doc.add_paragraph("报价高于采购文件设定的最高限价的。")
        doc.add_paragraph("表2：商务部分评审细则（满分10分）")
        scoring_table = doc.add_table(rows=2, cols=4)
        for col, text in enumerate(["序号", "评审项目", "分值", "评分细则"]):
            scoring_table.cell(0, col).text = text
        for col, text in enumerate(["一", "售后服务、技术服务方案", "2", "方案合理并具有可操作性，最优得满分。"]):
            scoring_table.cell(1, col).text = text
        doc.add_paragraph("采购人：都匀盛黔新能源有限公司")
        doc.add_paragraph("采购代理机构：中电建集中采购平台")
        doc.save(source_path)
        manifest_path.write_text(
            json.dumps(
                {
                    "projectId": "PRJ-AGENTIC-UNIT",
                    "bidType": "商务标",
                    "parseProfile": "business",
                    "structuredResultPath": str(output_path),
                    "documents": [
                        {
                            "id": "DOC-1",
                            "name": source_path.name,
                            "sourcePath": str(source_path),
                        }
                    ],
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        return source_path, manifest_path, output_path

    def run_s1parse(self, *args: str, check: bool = True) -> subprocess.CompletedProcess[str]:
        env = os.environ.copy()
        env["PYTHONIOENCODING"] = "utf-8"
        return subprocess.run(
            [sys.executable, str(self.runner_path()), *args],
            check=check,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            env=env,
        )

    def test_prepare_preserves_docx_body_order_and_builds_sqlite_index(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path, _ = self.write_sample_docx(root)

            completed = self.run_s1parse("prepare", str(manifest_path))

            summary = json.loads(completed.stdout)
            self.assertEqual(summary["stage"], "prepared")
            nav_path = Path(summary["navStorePath"])
            self.assertTrue(nav_path.is_file())
            conn = sqlite3.connect(nav_path)
            rows = conn.execute(
                "SELECT body_index, block_type, text, table_id FROM blocks ORDER BY body_index"
            ).fetchall()
            self.assertEqual([row[1] for row in rows[:5]], ["paragraph", "paragraph", "paragraph", "table", "paragraph"])
            self.assertEqual(rows[3][3], "DOC-1:T0001")
            self.assertIn("响应文件递交的截止时间", rows[4][2])
            evidence_count = conn.execute("SELECT COUNT(*) FROM evidence").fetchone()[0]
            table_row_count = conn.execute("SELECT COUNT(*) FROM table_rows").fetchone()[0]
            cell_count = conn.execute("SELECT COUNT(*) FROM table_cells").fetchone()[0]
            conn.close()
            self.assertGreaterEqual(evidence_count, 10)
            self.assertEqual(table_row_count, 5)
            self.assertGreaterEqual(cell_count, 14)

    def test_navigation_commands_return_small_outputs(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path, _ = self.write_sample_docx(root)
            self.run_s1parse("prepare", str(manifest_path))

            overview = json.loads(self.run_s1parse("overview", str(manifest_path), "--page", "1", "--page-size", "4").stdout)
            search = json.loads(self.run_s1parse("search", str(manifest_path), "递交截止", "--limit", "5").stdout)
            read = json.loads(self.run_s1parse("read", str(manifest_path), "DOC-1:B000005", "--max-chars", "200").stdout)
            window = json.loads(self.run_s1parse("window", str(manifest_path), "DOC-1:B000005", "--before", "1", "--after", "1").stdout)
            table = json.loads(self.run_s1parse("table", str(manifest_path), "DOC-1:T0001", "--rows", "1-3").stdout)

            self.assertLess(len(json.dumps(overview, ensure_ascii=False)), 2500)
            self.assertEqual(len(overview["blocks"]), 4)
            self.assertEqual(search["matchCount"], 1)
            self.assertIn("2026年05月06日10:00", read["record"]["text"])
            self.assertEqual([row["bodyIndex"] for row in window["blocks"]], [4, 5, 6])
            self.assertEqual(table["table"]["rowCount"], 3)
            self.assertEqual(len(table["rows"]), 3)

    def test_navigation_stdout_is_utf8_even_when_console_code_page_is_legacy(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source_path = root / "yen.docx"
            output_path = root / "s1_structured_result.json"
            manifest_path = root / "s1_parse_manifest.json"
            doc = Document()
            doc.add_paragraph("报价币种：¥人民币")
            doc.save(source_path)
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-UTF8-STDOUT",
                        "bidType": "business",
                        "parseProfile": "business",
                        "structuredResultPath": str(output_path),
                        "documents": [{"id": "DOC-1", "name": source_path.name, "sourcePath": str(source_path)}],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            self.run_s1parse("prepare", str(manifest_path))

            completed = self.run_s1parse("search", str(manifest_path), "人民币", "--limit", "5")

            payload = json.loads(completed.stdout)
            self.assertEqual(payload["matchCount"], 1)
            self.assertIn("¥人民币", payload["matches"][0]["text"])

    def test_submit_validate_finalize_writes_compatible_result(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path, output_path = self.write_sample_docx(root)
            self.run_s1parse("prepare", str(manifest_path))
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "projectBasics",
                json.dumps(
                    [
                        {
                            "key": "projectName",
                            "label": "项目名称",
                            "value": "都匀市盛黔风电场风力发电机组及附属设备采购项目",
                            "evidenceIds": ["DOC-1:B000002"],
                        },
                        {
                            "key": "tenderNo",
                            "label": "招标编号",
                            "value": "PC-0307-26J1-FG0002",
                            "evidenceIds": ["DOC-1:B000001"],
                        },
                        {
                            "key": "tenderer",
                            "label": "招标人",
                            "value": "都匀盛黔新能源有限公司",
                            "evidenceIds": ["DOC-1:B000011"],
                        },
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "qualificationRequirements",
                json.dumps(
                    [
                        {
                            "content": "已成为中国电建集团2025年度风力发电机组框架入围集中采购项目的入围供应商。",
                            "evidenceIds": ["DOC-1:B000006"],
                        }
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "bidderInstructions",
                json.dumps(
                    [
                        {
                            "clauseNo": "3.3",
                            "clauseName": "报价",
                            "content": "最高限价：最终报价不得超过人民币8250万元。",
                            "evidenceIds": ["DOC-1:T0001:R0002"],
                        }
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "commercialRejectionClauses",
                json.dumps(
                    [
                        {
                            "riskLevel": "high",
                            "content": "报价高于采购文件设定的最高限价的。",
                            "evidenceIds": ["DOC-1:B000008"],
                        }
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "businessScoringCriteria",
                json.dumps(
                    [
                        {
                            "scoringItem": "售后服务、技术服务方案",
                            "score": "2",
                            "scoringStandard": "方案合理并具有可操作性，最优得满分。",
                            "evidenceIds": ["DOC-1:T0002:R0002"],
                        }
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "projectDates",
                json.dumps({"endDate": "2026-05-06 10:00", "evidenceIds": ["DOC-1:B000005"]}, ensure_ascii=False),
            )

            validation = json.loads(self.run_s1parse("validate", str(manifest_path)).stdout)
            finalize = json.loads(self.run_s1parse("finalize", str(manifest_path)).stdout)
            payload = json.loads(output_path.read_text(encoding="utf-8"))

            self.assertEqual(validation["status"], "passed")
            self.assertEqual(finalize["summary"]["workflowStage"], "finalized")
            structured = payload["structured"]
            self.assertEqual(structured["workflow"]["mode"], "opencode-agentic-navigation")
            self.assertNotIn("candidatePackagePath", structured["workflow"])
            self.assertNotIn("reviewPlanPath", structured["workflow"])
            self.assertNotIn("aiTasksDir", structured["workflow"])
            self.assertEqual(structured["fieldGroups"]["projectBasics"][0]["key"], "projectName")
            self.assertEqual(structured["fieldGroups"]["qualificationRequirements"][0]["content"], "已成为中国电建集团2025年度风力发电机组框架入围集中采购项目的入围供应商。")
            self.assertEqual(structured["fieldGroups"]["qualificationRequirements"][0]["applicableScope"], "全部标段")
            self.assertEqual(structured["scoringCriteria"]["business"][0]["scoringItem"], "售后服务、技术服务方案")
            self.assertEqual(structured["projectDates"]["endDate"], "2026-05-06 10:00")

    def test_finalize_normalizes_agentic_dict_payload_for_frontend_tables(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path, output_path = self.write_sample_docx(root)
            self.run_s1parse("prepare", str(manifest_path))
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "projectBasics",
                json.dumps(
                    {
                        "projectName": {
                            "key": "projectName",
                            "label": "Project name",
                            "value": "都匀市盛黔风电场风力发电机组及附属设备采购项目",
                            "evidenceIds": ["DOC-1:B000002"],
                        },
                        "tenderNumber": {
                            "key": "tenderNumber",
                            "label": "Tender number",
                            "value": "HNZB-2026-001",
                        },
                        "projectUnit": {
                            "key": "projectUnit",
                            "label": "Project unit",
                            "value": "Example project unit",
                        },
                        "tenderer": {
                            "key": "tenderer",
                            "label": "Tenderer",
                            "value": "都匀盛黔新能源有限公司",
                            "evidenceIds": ["DOC-1:B000011"],
                        },
                        "agency": {
                            "key": "agency",
                            "label": "Agency",
                            "value": "Example agency",
                        },
                    },
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "qualificationRequirements",
                json.dumps(
                    [{"content": "Bidder must satisfy qualification requirements.", "evidenceIds": ["DOC-1:B000006"]}],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "bidderInstructions",
                json.dumps(
                    {
                        "tableTitle": "Bidder instruction preface table",
                        "tableId": "DOC-1:T0001",
                        "rowCount": 3,
                        "keyItems": [
                            {
                                "clause": "1.1.2",
                                "name": "Tenderer",
                                "content": "Example tenderer",
                                "rowId": "DOC-1:T0001:R0002",
                                "evidenceIds": ["DOC-1:T0001:R0002"],
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "commercialRejectionClauses",
                json.dumps(
                    [
                        {
                            "riskLevel": "high",
                            "label": "Price exceeds ceiling",
                            "content": "Response is invalid if price exceeds ceiling.",
                            "evidenceIds": ["DOC-1:B000008"],
                        }
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "businessScoringCriteria",
                json.dumps(
                    [
                        {
                            "scoringItem": "Service plan",
                            "score": "2",
                            "scoringStandard": "Best reasonable plan gets full score.",
                            "evidenceIds": ["DOC-1:T0002:R0002"],
                        }
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "projectDates",
                json.dumps({"endDate": "2026-05-06 10:00", "evidenceIds": ["DOC-1:B000005"]}, ensure_ascii=False),
            )

            validation = json.loads(self.run_s1parse("validate", str(manifest_path)).stdout)
            self.run_s1parse("finalize", str(manifest_path))
            payload = json.loads(output_path.read_text(encoding="utf-8"))
            structured = payload["structured"]
            basics = structured["fieldGroups"]["projectBasics"]
            by_key = {row["key"]: row for row in basics}

            self.assertEqual(validation["status"], "passed")
            self.assertEqual(by_key["projectName"]["value"], "都匀市盛黔风电场风力发电机组及附属设备采购项目")
            self.assertEqual(by_key["tenderNo"]["value"], "HNZB-2026-001")
            self.assertEqual(by_key["projectUnit"]["value"], "Example project unit")
            self.assertEqual(by_key["tenderer"]["value"], "都匀盛黔新能源有限公司")
            self.assertEqual(by_key["tenderAgency"]["value"], "Example agency")
            self.assertEqual(by_key["bidDeadline"]["value"], "2026-05-06 10:00")
            self.assertEqual(structured["fieldGroups"]["bidderInstructions"][0]["clauseNo"], "1.1.2")
            self.assertEqual(structured["fieldGroups"]["bidderInstructions"][0]["clauseName"], "Tenderer")
            self.assertEqual(structured["fieldGroups"]["commercialRejectionClauses"][0]["riskLevel"], "high")
            self.assertEqual(structured["fieldGroups"]["commercialRejectionClauses"][0]["matchedKeywords"], "Price exceeds ceiling")
            self.assertEqual(structured["scoringCriteria"]["business"][0]["scorePoint"], "Best reasonable plan gets full score.")

    def test_frontend_contract_normalizes_legacy_aliases_orders_and_deadline_list(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path, output_path = self.write_sample_docx(root)
            self.run_s1parse("prepare", str(manifest_path))
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "projectBasics",
                json.dumps(
                    [
                        {"label": "项目名称", "value": "都匀市盛黔风电场风力发电机组及附属设备采购项目", "evidenceIds": ["DOC-1:B000002"]},
                        {"label": "采购编号", "value": "PC-0307-26J1-FG0002", "evidenceIds": ["DOC-1:B000001"]},
                        {"label": "采购人", "value": "都匀盛黔新能源有限公司", "evidenceIds": ["DOC-1:B000011"]},
                        {"label": "采购代理机构", "value": "中电建集中采购平台", "evidenceIds": ["DOC-1:B000012"]},
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "qualificationRequirements",
                json.dumps(
                    [
                        {"content": "入围供应商。", "evidenceIds": ["DOC-1:B000006"]},
                        {"content": "不接受联合体。", "applicableScope": "本项目", "evidenceIds": ["DOC-1:B000006"]},
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "bidderInstructions",
                json.dumps(
                    [
                        {
                            "clauseName": "2.2 供应商提出异议/澄清截止时间",
                            "content": "2026年04月24日17时00分前",
                            "evidenceIds": ["DOC-1:T0001:R0002"],
                        }
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "commercialRejectionClauses",
                json.dumps(
                    [
                        {"riskLevel": "high", "matchedKeywords": "否决", "content": "报价高于采购文件设定的最高限价的，响应文件将被视为无效。", "evidenceIds": ["DOC-1:B000008"]},
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "businessScoringCriteria",
                json.dumps(
                    [
                        {"scoringItem": "售后服务、技术服务方案", "score": "2", "scoringStandard": "方案合理并具有可操作性，最优得满分。", "evidenceIds": ["DOC-1:T0002:R0002"]},
                        {"scoringItem": "商务偏差", "score": "1", "scorePoint": "无商务偏差得1分。", "evidenceIds": ["DOC-1:T0002:R0002"]},
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "projectDates",
                json.dumps(
                    [{"endDate": "2026-05-06 10:00", "source": "响应文件递交截止时间", "evidenceIds": ["DOC-1:B000005"]}],
                    ensure_ascii=False,
                ),
            )

            validation = json.loads(self.run_s1parse("validate", str(manifest_path)).stdout)
            self.run_s1parse("finalize", str(manifest_path))
            payload = json.loads(output_path.read_text(encoding="utf-8"))
            structured = payload["structured"]
            basics_by_key = {row["key"]: row for row in structured["fieldGroups"]["projectBasics"]}

            self.assertEqual(validation["status"], "passed")
            self.assertEqual(basics_by_key["tenderNo"]["value"], "PC-0307-26J1-FG0002")
            self.assertEqual(basics_by_key["tenderer"]["value"], "都匀盛黔新能源有限公司")
            self.assertEqual(basics_by_key["tenderAgency"]["value"], "中电建集中采购平台")
            self.assertEqual(basics_by_key["bidDeadline"]["value"], "2026-05-06 10:00")
            self.assertEqual(structured["projectDates"]["endDate"], "2026-05-06 10:00")
            self.assertEqual([row["order"] for row in structured["fieldGroups"]["qualificationRequirements"]], [1, 2])
            self.assertEqual([row["order"] for row in structured["scoringCriteria"]["business"]], [1, 2])
            self.assertEqual(structured["fieldGroups"]["bidderInstructions"][0]["clauseNo"], "2.2")
            self.assertEqual(structured["fieldGroups"]["bidderInstructions"][0]["clauseName"], "供应商提出异议/澄清截止时间")
            self.assertEqual(structured["fieldGroups"]["bidderInstructions"][0]["content"], "2026年04月24日17时00分前")
            self.assertEqual(len(structured["fieldGroups"]["commercialRejectionClauses"]), 1)
            self.assertNotIn("proofRequirement", structured["scoringCriteria"]["business"][0])

    def test_validate_fails_when_frontend_required_fields_are_not_displayable(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path, _ = self.write_sample_docx(root)
            self.run_s1parse("prepare", str(manifest_path))
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "projectBasics",
                json.dumps(
                    [{"label": "项目名称", "value": "都匀市盛黔风电场风力发电机组及附属设备采购项目", "evidenceIds": ["DOC-1:B000002"]}],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "qualificationRequirements",
                json.dumps([{"content": "入围供应商。", "evidenceIds": ["DOC-1:B000006"]}], ensure_ascii=False),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "bidderInstructions",
                json.dumps([{"clauseName": "只有名称", "content": "缺少条款号", "evidenceIds": ["DOC-1:T0001:R0002"]}], ensure_ascii=False),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "commercialRejectionClauses",
                json.dumps([{"content": "供应商保证金不予退还。", "evidenceIds": ["DOC-1:B000008"]}], ensure_ascii=False),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "businessScoringCriteria",
                json.dumps([{"scoringItem": "售后服务", "score": "2", "evidenceIds": ["DOC-1:T0002:R0002"]}], ensure_ascii=False),
            )

            validation = json.loads(self.run_s1parse("validate", str(manifest_path), check=False).stdout)
            codes = {item["code"] for item in validation["validationErrors"]}

            self.assertEqual(validation["status"], "failed")
            self.assertIn("missing_displayable_project_basic", codes)
            self.assertIn("invalid_bidder_instruction_row", codes)
            self.assertIn("invalid_business_scoring_row", codes)
            self.assertIn("non_rejection_deposit_clause", codes)

    def test_validate_fails_when_project_basic_value_is_not_supported_by_evidence(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path, _ = self.write_sample_docx(root)
            self.run_s1parse("prepare", str(manifest_path))
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "projectBasics",
                json.dumps(
                    {
                        "projectName": "都匀市盛黔风电场风力发电机组及附属设备采购项目",
                        "tenderer": "中国电力建设股份有限公司",
                        "bidDeadline": "2026年05月06日10:00",
                        "sourceEvidence": [
                            {"field": "projectName", "evidenceId": "DOC-1:B000002"},
                            {"field": "tenderer", "evidenceId": "DOC-1:B000002"},
                            {"field": "bidDeadline", "evidenceId": "DOC-1:B000005"},
                        ],
                    },
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "qualificationRequirements",
                json.dumps([{"content": "入围供应商。", "evidenceIds": ["DOC-1:B000006"]}], ensure_ascii=False),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "bidderInstructions",
                json.dumps(
                    [
                        {
                            "clauseNo": "3.3",
                            "clauseName": "报价",
                            "content": "最高限价：最终报价不得超过人民币8250万元。",
                            "evidenceIds": ["DOC-1:T0001:R0002"],
                        }
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "commercialRejectionClauses",
                json.dumps([{"riskLevel": "high", "content": "报价高于采购文件设定的最高限价的，响应文件将被视为无效。", "evidenceIds": ["DOC-1:B000008"]}], ensure_ascii=False),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "businessScoringCriteria",
                json.dumps(
                    [
                        {
                            "scoringItem": "售后服务、技术服务方案",
                            "score": "2",
                            "scorePoint": "方案合理并具有可操作性，最优得满分。",
                            "evidenceIds": ["DOC-1:T0002:R0002"],
                        }
                    ],
                    ensure_ascii=False,
                ),
            )

            validation = json.loads(self.run_s1parse("validate", str(manifest_path), check=False).stdout)
            errors = validation["validationErrors"]

            self.assertEqual(validation["status"], "failed")
            self.assertIn("project_basic_value_not_supported_by_evidence", {item["code"] for item in errors})
            tenderer_errors = [item for item in errors if item.get("fieldKey") == "tenderer"]
            self.assertEqual(len(tenderer_errors), 1)
            self.assertIn("DOC-1:B000002", tenderer_errors[0]["evidenceIds"])
            self.assertFalse([item for item in errors if item.get("fieldKey") == "bidDeadline"])

    def test_validate_fails_when_display_project_basics_have_no_evidence(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path, _ = self.write_sample_docx(root)
            self.run_s1parse("prepare", str(manifest_path))
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "projectBasics",
                json.dumps(
                    {
                        "projectName": "都匀市盛黔风电场风力发电机组及附属设备采购项目",
                        "tenderer": "都匀盛黔新能源有限公司",
                        "bidDeadline": "2026-05-06 10:00",
                    },
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "qualificationRequirements",
                json.dumps([{"content": "入围供应商。", "evidenceIds": ["DOC-1:B000006"]}], ensure_ascii=False),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "bidderInstructions",
                json.dumps([{"clauseNo": "3.3", "clauseName": "报价", "content": "最高限价：最终报价不得超过人民币8250万元。", "evidenceIds": ["DOC-1:T0001:R0002"]}], ensure_ascii=False),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "commercialRejectionClauses",
                json.dumps([{"riskLevel": "high", "content": "报价高于采购文件设定的最高限价的，响应文件将被视为无效。", "evidenceIds": ["DOC-1:B000008"]}], ensure_ascii=False),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "businessScoringCriteria",
                json.dumps([{"scoringItem": "售后服务、技术服务方案", "score": "2", "scorePoint": "方案合理并具有可操作性，最优得满分。", "evidenceIds": ["DOC-1:T0002:R0002"]}], ensure_ascii=False),
            )

            validation = json.loads(self.run_s1parse("validate", str(manifest_path), check=False).stdout)
            codes = {item["code"] for item in validation["validationErrors"]}

            self.assertEqual(validation["status"], "failed")
            self.assertIn("missing_project_basic_evidence", codes)

    def test_validate_fails_when_rejection_clause_risk_level_is_not_enum_value(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path, _ = self.write_sample_docx(root)
            self.run_s1parse("prepare", str(manifest_path))
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "projectBasics",
                json.dumps(
                    {
                        "projectName": "都匀市盛黔风电场风力发电机组及附属设备采购项目",
                        "tenderer": "都匀盛黔新能源有限公司",
                        "bidDeadline": "2026-05-06 10:00",
                        "sourceEvidence": [
                            {"field": "projectName", "evidenceId": "DOC-1:B000002"},
                            {"field": "tenderer", "evidenceId": "DOC-1:B000011"},
                            {"field": "bidDeadline", "evidenceId": "DOC-1:B000005"},
                        ],
                    },
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "qualificationRequirements",
                json.dumps([{"content": "入围供应商。", "evidenceIds": ["DOC-1:B000006"]}], ensure_ascii=False),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "bidderInstructions",
                json.dumps([{"clauseNo": "3.3", "clauseName": "报价", "content": "最高限价：最终报价不得超过人民币8250万元。", "evidenceIds": ["DOC-1:T0001:R0002"]}], ensure_ascii=False),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "commercialRejectionClauses",
                json.dumps(
                    [
                        {
                            "riskLevel": "否决投标",
                            "matchedKeywords": "无效",
                            "content": "报价高于采购文件设定的最高限价的，响应文件将被视为无效。",
                            "evidenceIds": ["DOC-1:B000008"],
                        }
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "businessScoringCriteria",
                json.dumps([{"scoringItem": "售后服务、技术服务方案", "score": "2", "scorePoint": "方案合理并具有可操作性，最优得满分。", "evidenceIds": ["DOC-1:T0002:R0002"]}], ensure_ascii=False),
            )

            validation = json.loads(self.run_s1parse("validate", str(manifest_path), check=False).stdout)
            errors = validation["validationErrors"]
            risk_errors = [item for item in errors if item["code"] == "invalid_rejection_clause_risk_level"]

            self.assertEqual(validation["status"], "failed")
            self.assertEqual(len(risk_errors), 1)
            self.assertEqual(risk_errors[0]["rowIndex"], 1)
            self.assertEqual(risk_errors[0]["value"], "否决投标")

    def test_frontend_contract_accepts_evidence_objects_and_backfills_project_dates_from_deadline(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path, output_path = self.write_sample_docx(root)
            self.run_s1parse("prepare", str(manifest_path))
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "projectBasics",
                json.dumps(
                    {
                        "projectName": "都匀市盛黔风电场风力发电机组及附属设备采购项目",
                        "tenderer": "都匀市盛黔风电场风力发电机组及附属设备采购项目",
                        "bidDeadline": "2026年05月06日10:00",
                        "evidence": [
                            {"id": "DOC-1:B000002", "field": "projectName", "text": "都匀市盛黔风电场风力发电机组及附属设备采购项目"},
                            {"id": "DOC-1:B000002", "field": "tenderer", "text": "都匀市盛黔风电场风力发电机组及附属设备采购项目"},
                            {"id": "DOC-1:B000005", "field": "bidDeadline", "text": "响应文件递交的截止时间为2026年05月06日10:00。"},
                        ],
                    },
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "qualificationRequirements",
                json.dumps([{"content": "入围供应商。", "evidence": [{"id": "DOC-1:B000006"}]}], ensure_ascii=False),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "bidderInstructions",
                json.dumps(
                    [
                        {
                            "clauseNo": "3.3",
                            "clauseName": "报价",
                            "content": "最高限价：最终报价不得超过人民币8250万元。",
                            "evidence": [{"id": "DOC-1:T0001:R0002"}],
                        }
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "commercialRejectionClauses",
                json.dumps([{"riskLevel": "high", "content": "报价高于采购文件设定的最高限价的，响应文件将被视为无效。", "evidence": [{"id": "DOC-1:B000008"}]}], ensure_ascii=False),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "businessScoringCriteria",
                json.dumps(
                    [
                        {
                            "scoringItem": "售后服务、技术服务方案",
                            "score": "2",
                            "scorePoint": "方案合理并具有可操作性，最优得满分。",
                            "evidence": [{"id": "DOC-1:T0002:R0002"}],
                        }
                    ],
                    ensure_ascii=False,
                ),
            )

            validation = json.loads(self.run_s1parse("validate", str(manifest_path)).stdout)
            self.run_s1parse("finalize", str(manifest_path))
            payload = json.loads(output_path.read_text(encoding="utf-8"))
            basics_by_key = {row["key"]: row for row in payload["structured"]["fieldGroups"]["projectBasics"]}

            self.assertEqual(validation["status"], "passed")
            self.assertGreaterEqual(validation["evidenceCount"], 5)
            self.assertEqual(basics_by_key["bidDeadline"]["value"], "2026-05-06 10:00")
            self.assertEqual(payload["structured"]["projectDates"]["endDate"], "2026-05-06 10:00")

    def test_frontend_contract_normalizes_private_evidence_ids_alias(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path, output_path = self.write_sample_docx(root)
            self.run_s1parse("prepare", str(manifest_path))
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "projectBasics",
                json.dumps(
                    [
                        {"key": "projectName", "value": "都匀市盛黔风电场风力发电机组及附属设备采购项目", "evidenceIds": ["DOC-1:B000002"]},
                        {"key": "tenderer", "value": "都匀盛黔新能源有限公司", "evidenceIds": ["DOC-1:B000011"]},
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "qualificationRequirements",
                json.dumps([{"content": "入围供应商。", "__evidenceIds": ["DOC-1:B000006"]}], ensure_ascii=False),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "bidderInstructions",
                json.dumps([{"clauseNo": "3.3", "clauseName": "报价", "content": "最高限价：最终报价不得超过人民币8250万元。", "__evidenceIds": ["DOC-1:T0001:R0002"]}], ensure_ascii=False),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "commercialRejectionClauses",
                json.dumps([{"riskLevel": "high", "content": "报价高于采购文件设定的最高限价的，响应文件将被视为无效。", "__evidenceIds": ["DOC-1:B000008"]}], ensure_ascii=False),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "businessScoringCriteria",
                json.dumps([{"scoringItem": "售后服务、技术服务方案", "score": "2", "scorePoint": "方案合理并具有可操作性，最优得满分。", "__evidenceIds": ["DOC-1:T0002:R0002"]}], ensure_ascii=False),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "projectDates",
                json.dumps({"endDate": "2026-05-06 10:00", "evidenceIds": ["DOC-1:B000005"]}, ensure_ascii=False),
            )

            validation = json.loads(self.run_s1parse("validate", str(manifest_path)).stdout)
            self.run_s1parse("finalize", str(manifest_path))
            payload = json.loads(output_path.read_text(encoding="utf-8"))
            structured = payload["structured"]

            self.assertEqual(validation["status"], "passed")
            self.assertGreaterEqual(validation["evidenceCount"], 7)
            self.assertIn("evidenceIds", structured["fieldGroups"]["qualificationRequirements"][0])
            self.assertNotIn("__evidenceIds", structured["fieldGroups"]["qualificationRequirements"][0])
            self.assertIn("evidenceIds", structured["fieldGroups"]["bidderInstructions"][0])
            self.assertIn("evidenceIds", structured["fieldGroups"]["commercialRejectionClauses"][0])
            self.assertIn("evidenceIds", structured["scoringCriteria"]["business"][0])

    def test_template_extractor_directory_is_not_modified_by_agentic_parser(self) -> None:
        backend_root = Path(__file__).resolve().parents[1]
        extractor_dir = backend_root / "opencode" / "skill" / "bid-business-template-extractor"
        before_files = sorted(path.relative_to(extractor_dir).as_posix() for path in extractor_dir.rglob("*") if path.is_file())
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path, _ = self.write_sample_docx(root)
            self.run_s1parse("prepare", str(manifest_path))
            self.run_s1parse("validate", str(manifest_path))

        after_files = sorted(path.relative_to(extractor_dir).as_posix() for path in extractor_dir.rglob("*") if path.is_file())
        self.assertEqual(before_files, after_files)
