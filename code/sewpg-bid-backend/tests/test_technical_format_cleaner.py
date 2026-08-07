from __future__ import annotations

import ast
import json
import tempfile
import unittest
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from app.document_processing.technical_document.formatting import cleaner


BACKEND_ROOT = Path(__file__).resolve().parents[1]
TECH_ASSEMBLY_PATH = BACKEND_ROOT / "app" / "services" / "tech_assembly.py"


def _run_application_cleaner(manifest_path: Path):
    return cleaner.run_manifest(manifest_path)


def _write_style_spec(path: Path, *, insert_toc: bool) -> None:
    path.write_text(
        json.dumps(
            {
                "body": {
                    "zh_font": "宋体",
                    "en_font": "Arial",
                    "size_pt": 16,
                    "bold": False,
                    "align": "both",
                    "space_before_pt": 0,
                    "space_after_pt": 0,
                    "line_spacing": 1.5,
                    "first_line_indent_chars": 2,
                },
                "toc": {"insert_when_missing": insert_toc, "page_break_after": True},
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )


def _write_empty_outline(path: Path) -> None:
    path.write_text(json.dumps({"schema_version": "tech_bid_outline.v1", "sections": []}), encoding="utf-8")


def _write_manifest(
    path: Path,
    *,
    input_file: Path,
    outline_file: Path,
    output_file: Path,
    style_file: Path,
    force_canonical_toc: bool = False,
) -> None:
    path.write_text(
        json.dumps(
            {
                "inputFile": str(input_file),
                "outlineFile": str(outline_file),
                "outputFile": str(output_file),
                "projectName": "脱敏项目",
                "styleSpecPath": str(style_file),
                "forceCanonicalToc": force_canonical_toc,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )


def _set_distinct_format(paragraph, index: int) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.left_indent = Cm(index)
    paragraph.paragraph_format.space_before = Pt(index)
    paragraph.paragraph_format.space_after = Pt(index + 1)
    paragraph.paragraph_format.line_spacing = 1 + index / 10
    paragraph.runs[0].font.name = "Calibri"
    paragraph.runs[0].font.size = Pt(8 + index)


def _write_toc_format_document(path: Path) -> dict[str, int]:
    doc = Document()
    toc_style = doc.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
    toc_style.base_style = doc.styles["Normal"]
    inherited_style = doc.styles.add_style("自定义目录条目", WD_STYLE_TYPE.PARAGRAPH)
    inherited_style.base_style = toc_style
    chinese_style = doc.styles.add_style("目录 2", WD_STYLE_TYPE.PARAGRAPH)
    chinese_style.base_style = doc.styles["Normal"]

    entries = (
        ("TOC 1 条目", "TOC 1", 1),
        ("继承目录条目", "自定义目录条目", 2),
        ("中文目录条目", "目录 2", 3),
    )
    expected: dict[str, int] = {}
    for text, style_name, index in entries:
        paragraph = doc.add_paragraph(text, style=style_name)
        _set_distinct_format(paragraph, index)
        expected[text] = index
    doc.add_paragraph("普通正文")
    doc.save(path)
    return expected


def _write_existing_toc_document(path: Path) -> None:
    doc = Document()
    doc.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
    doc.styles.add_style("目录 2", WD_STYLE_TYPE.PARAGRAPH)
    field = doc.add_paragraph()
    instruction_run = OxmlElement("w:r")
    instruction = OxmlElement("w:instrText")
    instruction.text = ' TOC \\o "1-3" '
    instruction_run.append(instruction)
    field._element.append(instruction_run)
    doc.add_paragraph("旧英文目录结果", style="TOC 1")
    doc.add_paragraph("旧中文目录结果", style="目录 2")
    doc.add_paragraph("保留正文")
    doc.save(path)


def _write_custom_style_toc_document(path: Path) -> None:
    doc = Document()
    toc_style = doc.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
    toc_style.base_style = doc.styles["Normal"]
    inherited_style = doc.styles.add_style("Inherited Toc Entry", WD_STYLE_TYPE.PARAGRAPH)
    inherited_style.base_style = toc_style
    custom_style = doc.styles.add_style("Project Directory Entry", WD_STYLE_TYPE.PARAGRAPH)
    custom_style.base_style = doc.styles["Normal"]

    field = doc.add_paragraph()
    instruction_run = OxmlElement("w:r")
    instruction = OxmlElement("w:instrText")
    instruction.text = ' TOC \\o "1-3" '
    instruction_run.append(instruction)
    field._element.append(instruction_run)

    doc.add_paragraph("OLD_INHERITED_TOC_ENTRY", style=inherited_style)
    custom_entry = doc.add_paragraph(style=custom_style)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), "_Toc123456789")
    hyperlink_run = OxmlElement("w:r")
    hyperlink_text = OxmlElement("w:t")
    hyperlink_text.text = "OLD_CUSTOM_TOC_ENTRY"
    hyperlink_run.append(hyperlink_text)
    hyperlink.append(hyperlink_run)
    custom_entry._element.append(hyperlink)

    page_break = doc.add_paragraph()
    page_break_run = OxmlElement("w:r")
    page_break_element = OxmlElement("w:br")
    page_break_element.set(qn("w:type"), "page")
    page_break_run.append(page_break_element)
    page_break._element.append(page_break_run)
    doc.add_paragraph("BODY_MUST_REMAIN", style="Heading 1")
    doc.save(path)


class TechnicalFormatCleanerTests(unittest.TestCase):
    def test_run_manifest_returns_warnings_without_writing_markdown_report(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            input_file = root / "assembled.docx"
            outline_file = root / "outline.json"
            output_file = root / "assembled.formatted.docx"
            manifest_file = root / "manifest.json"

            doc = Document()
            doc.add_paragraph("技术方案", style="Heading 1")
            doc.add_paragraph("脱敏正文。")
            doc.add_table(rows=1, cols=1).cell(0, 0).text = "参数值"
            doc.save(input_file)
            outline_file.write_text(
                json.dumps(
                    {
                        "schema_version": "tech_bid_outline.v1",
                        "sections": [{"id": "S-1", "title": "技术方案", "level": 1}],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            manifest_file.write_text(
                json.dumps(
                    {
                        "inputFile": str(input_file),
                        "outlineFile": str(outline_file),
                        "outputFile": str(output_file),
                        "projectName": "脱敏项目",
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = _run_application_cleaner(manifest_file)

            self.assertTrue(output_file.exists())
            self.assertEqual(result["reportFile"], "")
            self.assertIsInstance(result["summary"], dict)
            self.assertEqual(
                result["summary"]["fontFamilies"],
                ["等线 Light", "Times New Roman", "等线", "宋体"],
            )
            with ZipFile(output_file) as archive:
                document_xml = "\n".join(
                    archive.read(name).decode("utf-8", errors="ignore")
                    for name in archive.namelist()
                    if name.startswith("word/") and name.endswith(".xml")
                )
            for family in ("等线", "等线 Light", "宋体", "Times New Roman"):
                self.assertIn(family, document_xml)
            for family in ("Noto Sans CJK SC", "Noto Serif CJK SC", "Liberation Serif"):
                self.assertNotIn(family, document_xml)
            self.assertIsInstance(result["warnings"], list)
            self.assertEqual(result["warnings"], result["summary"]["warnings"])
            self.assertTrue(all(set(item) == {"code", "message", "count"} for item in result["warnings"]))
            self.assertFalse((root / "tech_format_clean_report.md").exists())

    def test_production_cleaner_preserves_toc_styles_and_formats_body(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            input_file = root / "toc-input.docx"
            outline_file = root / "outline.json"
            output_file = root / "toc-output.docx"
            style_file = root / "style.json"
            manifest_file = root / "manifest.json"
            expected = _write_toc_format_document(input_file)
            _write_empty_outline(outline_file)
            _write_style_spec(style_file, insert_toc=False)
            _write_manifest(
                manifest_file,
                input_file=input_file,
                outline_file=outline_file,
                output_file=output_file,
                style_file=style_file,
            )

            _run_application_cleaner(manifest_file)

            output = Document(str(output_file))
            paragraphs = {paragraph.text: paragraph for paragraph in output.paragraphs}
            for text, index in expected.items():
                paragraph = paragraphs[text]
                self.assertEqual(paragraph.alignment, WD_ALIGN_PARAGRAPH.RIGHT)
                self.assertAlmostEqual(paragraph.paragraph_format.left_indent.cm, index, places=3)
                self.assertAlmostEqual(paragraph.paragraph_format.space_before.pt, index, places=3)
                self.assertAlmostEqual(paragraph.paragraph_format.space_after.pt, index + 1, places=3)
                self.assertAlmostEqual(paragraph.paragraph_format.line_spacing, 1 + index / 10, places=3)
                self.assertEqual(paragraph.runs[0].font.name, "Calibri")
                self.assertAlmostEqual(paragraph.runs[0].font.size.pt, 8 + index)

            body = paragraphs["普通正文"]
            self.assertEqual(body.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
            self.assertAlmostEqual(body.runs[0].font.size.pt, 16)
            with ZipFile(output_file) as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")
            self.assertIn('w:ascii="Arial"', document_xml)
            self.assertIn('w:eastAsia="宋体"', document_xml)

    def test_force_canonical_toc_removes_english_and_chinese_results(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            input_file = root / "existing-toc.docx"
            outline_file = root / "outline.json"
            output_file = root / "canonical-toc.docx"
            style_file = root / "style.json"
            manifest_file = root / "manifest.json"
            _write_existing_toc_document(input_file)
            _write_empty_outline(outline_file)
            _write_style_spec(style_file, insert_toc=True)
            _write_manifest(
                manifest_file,
                input_file=input_file,
                outline_file=outline_file,
                output_file=output_file,
                style_file=style_file,
                force_canonical_toc=True,
            )

            _run_application_cleaner(manifest_file)

            output = Document(str(output_file))
            texts = [paragraph.text for paragraph in output.paragraphs]
            self.assertNotIn("旧英文目录结果", texts)
            self.assertNotIn("旧中文目录结果", texts)
            self.assertIn("保留正文", texts)
            self.assertTrue(
                any("TOC" in (node.text or "").upper() for node in output.element.iter(qn("w:instrText")))
            )

    def test_force_canonical_toc_removes_custom_and_inherited_results_idempotently(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            input_file = root / "custom-existing-toc.docx"
            outline_file = root / "outline.json"
            first_output = root / "canonical-toc-first.docx"
            second_output = root / "canonical-toc-second.docx"
            style_file = root / "style.json"
            first_manifest = root / "manifest-first.json"
            second_manifest = root / "manifest-second.json"
            _write_custom_style_toc_document(input_file)
            _write_empty_outline(outline_file)
            _write_style_spec(style_file, insert_toc=True)

            for source, target, manifest in (
                (input_file, first_output, first_manifest),
                (first_output, second_output, second_manifest),
            ):
                _write_manifest(
                    manifest,
                    input_file=source,
                    outline_file=outline_file,
                    output_file=target,
                    style_file=style_file,
                    force_canonical_toc=True,
                )
                _run_application_cleaner(manifest)

            output = Document(str(second_output))
            all_text = "".join(node.text or "" for node in output.element.iter(qn("w:t")))
            self.assertNotIn("OLD_INHERITED_TOC_ENTRY", all_text)
            self.assertNotIn("OLD_CUSTOM_TOC_ENTRY", all_text)
            self.assertIn("BODY_MUST_REMAIN", all_text)
            toc_instructions = [
                node for node in output.element.iter(qn("w:instrText")) if "TOC" in (node.text or "").upper()
            ]
            self.assertEqual(len(toc_instructions), 1)
            toc_breaks = [
                node
                for node in output.element.iter(qn("w:bookmarkStart"))
                if node.get(qn("w:name")) == cleaner.TOC_BREAK_BOOKMARK
            ]
            self.assertEqual(len(toc_breaks), 1)
            self.assertEqual(sum(paragraph.text == "目录" for paragraph in output.paragraphs), 1)
            page_breaks = [
                node
                for node in output.element.iter(qn("w:br"))
                if (node.get(qn("w:type")) or "page") == "page"
            ]
            self.assertEqual(len(page_breaks), 2)

    def test_toc_style_identifier_accepts_only_supported_levels(self) -> None:
        for value in ("TOC 1", "toc1", "目录 1", "目录9"):
            self.assertTrue(cleaner._is_toc_style_identifier(value))
        for value in ("TOC 0", "TOC 10", "TOC Notes", "目录", "目录 10"):
            self.assertFalse(cleaner._is_toc_style_identifier(value))

    def test_toc_field_instruction_excludes_pageref_bookmarks(self) -> None:
        self.assertTrue(cleaner._is_toc_field_instruction(' TOC \\o "1-3" '))
        self.assertFalse(cleaner._is_toc_field_instruction(" PAGEREF _Toc123456789 \\h "))

    def test_tech_assembly_delegates_to_application_cleaner(self) -> None:
        tree = ast.parse(TECH_ASSEMBLY_PATH.read_text(encoding="utf-8"), filename=str(TECH_ASSEMBLY_PATH))
        imported = any(
            isinstance(node, ast.ImportFrom)
            and node.module == "app.document_processing.technical_document.formatting"
            and any(alias.name == "run_manifest" and alias.asname == "run_format_manifest" for alias in node.names)
            for node in tree.body
        )
        self.assertTrue(imported)

        function = next(
            node
            for node in tree.body
            if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef))
            and node.name == "_run_local_tech_format_cleaner"
        )
        self.assertTrue(
            any(
                isinstance(node, ast.Call)
                and isinstance(node.func, ast.Name)
                and node.func.id == "run_format_manifest"
                for node in ast.walk(function)
            )
        )


if __name__ == "__main__":
    unittest.main()
