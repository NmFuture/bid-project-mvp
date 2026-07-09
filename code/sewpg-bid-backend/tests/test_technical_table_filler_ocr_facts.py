"""table-filler PDF/OCR sidecar 事实抽取单测（F 系列认证表 0% 攻坚）。

金标反评定位：F 系列认证表素材是扫描版 PDF 证书，选源层能选中（kind=pdf）
但 collect_facts 无 pdf 分支，选了也不抽取。本组测试覆盖新增的
OCR sidecar 消费路径：sidecar 发现、键值/表格/机构名/领域字段抽取、
无 sidecar 显式报错（不静默）。
"""
from __future__ import annotations

import importlib.util
import sys
import tempfile
import unittest
from pathlib import Path

_SRC = (
    Path(__file__).resolve().parents[1]
    / "opencode"
    / "skills"
    / "bid-tech-table-filler"
    / "scripts"
    / "run_from_manifest.py"
)
_SPEC = importlib.util.spec_from_file_location("tech_table_filler_ocr_under_test", _SRC)
filler = importlib.util.module_from_spec(_SPEC)
assert _SPEC.loader is not None
sys.modules["tech_table_filler_ocr_under_test"] = filler
_SPEC.loader.exec_module(filler)


SAMPLE_OCR_TEXT = """\
型式认证证书
证书编号：TC-WT-2025-001
北京鉴衡认证中心
认证依据：IEC 61400-1:2019
| 项目 | 内容 |
| 认证等级 | IEC S |
| 发证日期 | 2025年3月18日 |
有效期至 2030年3月17日
"""


class _SourceFactory:
    def __init__(self, base_dir: Path) -> None:
        self.base_dir = base_dir

    def pdf_source(self, *, sidecar_text: str | None, name: str = "型式认证证书.pdf") -> object:
        pdf_path = self.base_dir / name
        pdf_path.write_bytes(b"%PDF-1.4 fake")
        if sidecar_text is not None:
            pdf_path.with_suffix(".ocr.txt").write_text(sidecar_text, encoding="utf-8")
        return filler.Source(name=name, path=pdf_path, kind="pdf", priority=80, route="test")


class ExtractPdfOcrFactsTests(unittest.TestCase):
    def setUp(self) -> None:
        self._tmp = tempfile.TemporaryDirectory()
        self.factory = _SourceFactory(Path(self._tmp.name))

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def test_missing_sidecar_raises_explicitly(self) -> None:
        source = self.factory.pdf_source(sidecar_text=None)
        with self.assertRaises(RuntimeError):
            filler.extract_pdf_ocr_facts(source)

    def test_colon_key_value_extracted(self) -> None:
        source = self.factory.pdf_source(sidecar_text=SAMPLE_OCR_TEXT)
        facts = filler.extract_pdf_ocr_facts(source)
        by_label = {fact["label"]: fact for fact in facts}
        self.assertIn("认证依据", by_label)
        self.assertEqual(by_label["认证依据"]["value"], "IEC 61400-1:2019")

    def test_markdown_table_row_extracted(self) -> None:
        source = self.factory.pdf_source(sidecar_text=SAMPLE_OCR_TEXT)
        facts = filler.extract_pdf_ocr_facts(source)
        by_label = {fact["label"]: fact["value"] for fact in facts}
        self.assertEqual(by_label.get("认证等级"), "IEC S")

    def test_issuer_standalone_line_recognized(self) -> None:
        source = self.factory.pdf_source(sidecar_text=SAMPLE_OCR_TEXT)
        facts = filler.extract_pdf_ocr_facts(source)
        issuer_values = [fact["value"] for fact in facts if fact["label"] == "认证机构"]
        self.assertIn("北京鉴衡认证中心", issuer_values)

    def test_domain_patterns_catch_no_colon_validity(self) -> None:
        source = self.factory.pdf_source(sidecar_text=SAMPLE_OCR_TEXT)
        facts = filler.extract_pdf_ocr_facts(source)
        validity = [fact for fact in facts if fact["label"] == "有效期"]
        self.assertTrue(validity)
        self.assertIn("2030", validity[0]["value"])

    def test_all_ocr_facts_carry_review_risk(self) -> None:
        source = self.factory.pdf_source(sidecar_text=SAMPLE_OCR_TEXT)
        facts = filler.extract_pdf_ocr_facts(source)
        self.assertTrue(facts)
        for fact in facts:
            self.assertTrue(fact["risk"], f"OCR 事实缺 risk 标记：{fact['label']}")


class CollectFactsPdfBranchTests(unittest.TestCase):
    def setUp(self) -> None:
        self._tmp = tempfile.TemporaryDirectory()
        self.factory = _SourceFactory(Path(self._tmp.name))

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def test_pdf_without_sidecar_surfaces_error_in_meta(self) -> None:
        source = self.factory.pdf_source(sidecar_text=None)
        meta, facts = filler.collect_facts([source], {}, {})
        self.assertIn("error", meta.get(source.name, {}))
        self.assertIn("OCR", meta[source.name]["error"])

    def test_pdf_with_sidecar_contributes_facts(self) -> None:
        source = self.factory.pdf_source(sidecar_text=SAMPLE_OCR_TEXT)
        _meta, facts = filler.collect_facts([source], {}, {})
        pdf_facts = [fact for fact in facts if fact["sourceKind"] == "pdf"]
        self.assertTrue(pdf_facts)


class SourceSidecarDiscoveryTests(unittest.TestCase):
    def setUp(self) -> None:
        self._tmp = tempfile.TemporaryDirectory()
        self.base_dir = Path(self._tmp.name)

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def test_add_source_resolves_sidecar_by_convention(self) -> None:
        pdf_path = self.base_dir / "MAT-1-设计认证证书.pdf"
        pdf_path.write_bytes(b"%PDF-1.4 fake")
        pdf_path.with_suffix(".ocr.txt").write_text("证书编号：X-1", encoding="utf-8")
        sources: list = []
        filler.add_source_from_material(
            sources,
            {"id": "MAT-1", "name": "设计认证证书", "path": str(pdf_path)},
            self.base_dir,
            priority=80,
            route="test",
        )
        self.assertEqual(len(sources), 1)
        self.assertEqual(sources[0].kind, "pdf")
        self.assertIsNotNone(sources[0].ocr_text_path)

    def test_add_source_prefers_manifest_ocr_text_path(self) -> None:
        pdf_path = self.base_dir / "MAT-2-并网认证.pdf"
        pdf_path.write_bytes(b"%PDF-1.4 fake")
        sidecar = self.base_dir / "elsewhere.ocr.txt"
        sidecar.write_text("证书编号：X-2", encoding="utf-8")
        sources: list = []
        filler.add_source_from_material(
            sources,
            {"id": "MAT-2", "name": "并网认证", "path": str(pdf_path), "ocrTextPath": str(sidecar)},
            self.base_dir,
            priority=80,
            route="test",
        )
        self.assertEqual(sources[0].ocr_text_path, sidecar.resolve())


class CertKeywordAndScoreTests(unittest.TestCase):
    def test_f_series_keywords_present(self) -> None:
        for prefix in ("F1", "F2", "F3", "F4", "F5"):
            self.assertTrue(filler.component_keywords_for(prefix), f"{prefix} 无选源关键词")

    def test_ocr_fact_maps_to_cert_field_via_generic_score(self) -> None:
        fact = {
            "label": "证书编号",
            "value": "TC-WT-2025-001",
            "unit": "",
            "sourceKind": "pdf",
            "sourcePriority": 80,
            "concepts": [],
            "baseConfidence": 0.74,
        }
        field = {"field": "证书编号", "unit": "", "concepts": [], "generic": True}
        self.assertGreaterEqual(filler.score(field, fact, "table"), 0.62)


class MapFieldsOcrPartialTests(unittest.TestCase):
    """OCR 事实必须穿过 map_fields 的 usable 线（0.62）和 risk 拦截线（0.76），
    以 partial（填入+高亮待人工）落地——此前 risk 拦截会把 0.76 以下全部置空，
    证书表退回全空。"""

    def setUp(self) -> None:
        self._tmp = tempfile.TemporaryDirectory()
        self.factory = _SourceFactory(Path(self._tmp.name))

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def _spec(self) -> object:
        return filler.AppendixSpec(
            appendix_id="APPX-F21",
            prefix="F2",
            title="附表F.2.1 投标机组设计认证",
            source=Path(self._tmp.name) / "blank.docx",
            table_index=0,
            header_row=0,
            field_col=0,
            value_col=1,
            unit_col=None,
            remark_col=None,
        )

    def test_ocr_fact_survives_risk_gate_as_partial(self) -> None:
        source = self.factory.pdf_source(sidecar_text=SAMPLE_OCR_TEXT)
        facts = filler.extract_pdf_ocr_facts(source)
        fields = [
            {
                "id": "F2-R02",
                "rowIndex": 2,
                "tableIndex": 0,
                "valueCol": 1,
                "unitCol": None,
                "group": "",
                "field": "证书编号",
                "unit": "",
                "remark": "",
                "requirementValue": "",
                "concepts": [],
                "generic": True,
            }
        ]
        mapping = filler.map_fields(self._spec(), fields, facts, "table")
        decision = next(d for d in mapping["decisions"] if d["field"] == "证书编号")
        self.assertEqual(decision["action"], "partial")
        self.assertEqual(decision["value"], "TC-WT-2025-001")


class CertMirrorColumnTests(unittest.TestCase):
    """F 系列"认证机型N/投标机型N"成对列：投标机型即认证机型，fill_doc 两列同值
    （金标反评 F.2.1：中标人 16 格逐格复制，差异列不动）。"""

    def setUp(self) -> None:
        self._tmp = tempfile.TemporaryDirectory()
        self.base = Path(self._tmp.name)

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def _build_blank(self) -> Path:
        from docx import Document

        doc = Document()
        table = doc.add_table(rows=3, cols=5)
        headers = ["项目名称", "项目名称", "认证机型1", "投标机型1", "认证机型1与投标机型1差异"]
        for idx, text in enumerate(headers):
            table.rows[0].cells[idx].text = text
        table.rows[1].cells[0].text = "基本信息"
        table.rows[1].cells[1].text = "认证机构"
        table.rows[2].cells[0].text = "基本信息"
        table.rows[2].cells[1].text = "机组型号"
        path = self.base / "blank_f21.docx"
        doc.save(str(path))
        return path

    def _spec(self, blank: Path) -> object:
        return filler.AppendixSpec(
            appendix_id="APPX-F21",
            prefix="F2",
            title="附表F.2.1 投标机组设计认证",
            source=blank,
            table_index=0,
            header_row=0,
            field_col=1,
            value_col=2,
            unit_col=None,
            remark_col=None,
        )

    def test_value_mirrored_to_bid_model_column(self) -> None:
        from docx import Document

        blank = self._build_blank()
        spec = self._spec(blank)
        mapping = {
            "decisions": [
                {"rowIndex": 1, "tableIndex": 0, "valueCol": 2, "unitCol": None, "field": "认证机构", "action": "partial", "value": "中国质量认证中心", "unit": ""},
                {"rowIndex": 2, "tableIndex": 0, "valueCol": 2, "unitCol": None, "field": "机组型号", "action": "manual", "value": "[待人工补充：机组型号]", "unit": ""},
            ]
        }
        out = self.base / "out.docx"
        filler.fill_doc(spec, mapping, out)
        table = Document(str(out)).tables[0]
        self.assertEqual(table.rows[1].cells[2].text, "中国质量认证中心")
        self.assertEqual(table.rows[1].cells[3].text, "中国质量认证中心")
        self.assertEqual(table.rows[1].cells[4].text, "")
        # manual 占位不镜像
        self.assertEqual(table.rows[2].cells[3].text, "")

    def test_no_mirror_without_paired_headers(self) -> None:
        from docx import Document

        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        for idx, text in enumerate(["项目", "内容", "备注"]):
            table.rows[0].cells[idx].text = text
        table.rows[1].cells[0].text = "证书编号"
        blank = self.base / "blank_plain.docx"
        doc.save(str(blank))
        spec = filler.AppendixSpec(
            appendix_id="APPX-F21",
            prefix="F2",
            title="附表F.2.1 投标机组设计认证",
            source=blank,
            table_index=0,
            header_row=0,
            field_col=0,
            value_col=1,
            unit_col=None,
            remark_col=None,
        )
        mapping = {"decisions": [{"rowIndex": 1, "tableIndex": 0, "valueCol": 1, "unitCol": None, "field": "证书编号", "action": "fill", "value": "X-1", "unit": ""}]}
        out = self.base / "out_plain.docx"
        filler.fill_doc(spec, mapping, out)
        table = Document(str(out)).tables[0]
        self.assertEqual(table.rows[1].cells[1].text, "X-1")
        self.assertEqual(table.rows[1].cells[2].text, "")


if __name__ == "__main__":
    unittest.main()
