"""附表 AI 填写 LLM 化 PR-1 单测：--prepare 简报与 --apply 计划校验/执行。

覆盖：
1. prepare 产出的 fill_brief.json 结构与现有检测一致（targetFields 同
   extract_target_fields，素材按 route 标注，事实表只透传 confirmed/extracted）。
2. apply 校验器五类用例：坐标越界 / excerpt 不在来源文件 / 单位换算 /
   manual 标记 / 清单型一行多列。
3. 端到端：手写正确 plan → apply → 输出 docx 格子值 + fill_report 带
   fillMode="llm-plan"，stdout 摘要与 compact_summary 同构。
"""
from __future__ import annotations

import importlib.util
import json
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

_SRC = (
    Path(__file__).resolve().parents[1]
    / "opencode"
    / "skills"
    / "bid-tech-table-filler"
    / "scripts"
    / "run_from_manifest.py"
)
_SPEC = importlib.util.spec_from_file_location("tech_table_filler_llm_plan_under_test", _SRC)
filler = importlib.util.module_from_spec(_SPEC)
assert _SPEC.loader is not None
sys.modules["tech_table_filler_llm_plan_under_test"] = filler
_SPEC.loader.exec_module(filler)


def _build_docx(path: Path, heading: str, rows: list[list[str]]) -> Path:
    doc = Document()
    doc.add_paragraph(heading)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for ri, row in enumerate(rows):
        for ci, value in enumerate(row):
            table.rows[ri].cells[ci].text = value
    doc.save(str(path))
    return path


def _build_text_docx(path: Path, paragraphs: list[str]) -> Path:
    doc = Document()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(str(path))
    return path


# 参数表：编号 + 项目 + 主要项目 + 技术参数与规格（响应列）+ 计量单位 + 备注
_PARAM_ROWS = [
    ["编号", "项目", "主要项目", "技术参数与规格", "计量单位", "备注"],
    ["1", "机型总体参数", "投标机型", "", "", ""],
    ["2", "机型总体参数", "单机容量", "", "MW", ""],
    ["3", "机型总体参数", "叶轮直径", "", "m", ""],
]

# 供货范围清单：序号空、货物名称有值、品牌/型号/原产地/数量待填
_SUPPLY_ROWS = [
    ["序号", "货物名称", "品牌或制造商名称", "型号和规格", "原产地", "数量", "备注"],
    ["", "主控系统", "", "", "", "", ""],
    ["", "发电机", "", "", "", "", ""],
]

_MATERIAL_PARAGRAPHS = [
    "机型参数说明",
    "投标机型为 EW6.25-220 机组。",
    "单机容量 6250 kW，叶轮直径 220 m。",
]


class LlmPlanTestBase(unittest.TestCase):
    def setUp(self) -> None:
        self._tmp = tempfile.TemporaryDirectory()
        self.base = Path(self._tmp.name)
        self.blank = _build_docx(self.base / "blank.docx", "附表C.1 机型总体参数与规格", _PARAM_ROWS)
        self.material = _build_text_docx(self.base / "material.docx", _MATERIAL_PARAGRAPHS)
        self.output = self.base / "out" / "filled.docx"

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def _write_manifest(self, *, blank: Path | None = None, title: str = "附表C.1 机型总体参数与规格") -> Path:
        manifest = {
            "blankSource": {"docxPath": str(blank or self.blank), "title": title, "id": "APPX-TEST"},
            "outputFile": str(self.output),
            "referenceMaterials": [{"id": "RAW-1", "name": "机型参数材料.docx", "path": str(self.material)}],
            "projectFactTable": {
                "status": "confirmed",
                "fields": [
                    {"label": "单机容量", "value": "6.25", "unit": "MW", "status": "confirmed"},
                    {"label": "噪音", "value": "104", "unit": "dB(A)", "status": "pending_confirmation"},
                ],
            },
            "parseFields": [{"label": "叶轮直径", "value": "220", "unit": "m"}],
            "projectTurbineModel": {"model": "EW6.25-220", "ratedPowerKw": 6250},
        }
        path = self.base / "manifest.json"
        path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
        return path

    def _write_plan(self, fills: list[dict]) -> Path:
        plan = {"schemaVersion": "bid-tech-table-fill-plan-v1", "fills": fills}
        path = self.base / "fill_plan.json"
        path.write_text(json.dumps(plan, ensure_ascii=False), encoding="utf-8")
        return path

    def _fill(self, target_field_id: str, field: str, value: str, excerpt: str, **extra) -> dict:
        fill = {
            "targetFieldId": target_field_id,
            "field": field,
            "action": "fill",
            "value": value,
            "unit": "",
            "confidence": 0.9,
            "evidence": {"sourceRoute": "referenceMaterial", "sourcePath": str(self.material), "excerpt": excerpt},
            "reason": "测试取值",
        }
        fill.update(extra)
        return fill


class PrepareBriefTests(LlmPlanTestBase):
    def test_brief_structure_and_fields_match_detection(self) -> None:
        manifest_path = self._write_manifest()
        payload = filler.run_prepare(manifest_path)
        brief_file = self.base / "fill_brief.json"
        self.assertEqual(payload["briefFile"], str(brief_file))
        self.assertEqual(payload["targetFieldCount"], 3)
        self.assertEqual(payload["materialCount"], 1)
        brief = json.loads(brief_file.read_text(encoding="utf-8"))

        self.assertEqual(brief["schemaVersion"], "bid-tech-table-fill-brief-v1")
        self.assertEqual(brief["blankDocxPath"], str(self.blank.resolve()))
        self.assertEqual(brief["outputFile"], str(self.output))
        self.assertTrue(brief["planFile"].endswith("fill_plan.json"))

        # 字段清单与现有检测逐字一致（同一 extract_target_fields 输出）
        spec = filler.detect_appendix_spec(self.blank.resolve(), json.loads(manifest_path.read_text(encoding="utf-8")))
        detected = filler.extract_target_fields(spec)
        self.assertEqual([f["targetFieldId"] for f in brief["targetFields"]], [f["id"] for f in detected])
        self.assertEqual([f["field"] for f in brief["targetFields"]], ["投标机型", "单机容量", "叶轮直径"])
        for field in brief["targetFields"]:
            for key in ("tableIndex", "rowIndex", "valueCol", "unitCol", "requirementValue", "unit", "group", "rowLabel", "columnLabel"):
                self.assertIn(key, field)

        # 素材按 route 标注，路径解析沿用现有定位键
        self.assertEqual(brief["materials"][0]["route"], "referenceMaterial")
        self.assertEqual(brief["materials"][0]["path"], str(self.material.resolve()))

        # 事实表只透传 confirmed/extracted；parseFields/机型原样透传
        self.assertEqual([f["label"] for f in brief["factTableFields"]], ["单机容量"])
        self.assertEqual(brief["parseFields"], [{"label": "叶轮直径", "value": "220", "unit": "m"}])
        self.assertEqual(brief["projectTurbineModel"]["model"], "EW6.25-220")

        # 填写铁律：非空字符串列表，含证据与待人工契约
        self.assertTrue(brief["rules"])
        self.assertTrue(any("excerpt" in rule for rule in brief["rules"]))
        self.assertTrue(any("待人工补充" in rule for rule in brief["rules"]))


class ApplyValidationTests(LlmPlanTestBase):
    def test_out_of_bounds_coordinate_fails_without_writing(self) -> None:
        manifest_path = self._write_manifest()
        fill = self._fill("C1-R01", "投标机型", "EW6.25-220", "EW6.25-220", tableIndex=99)
        self._write_plan([fill])
        with self.assertRaises(filler.PlanValidationError) as ctx:
            filler.run_apply(manifest_path)
        self.assertIn("越界", ctx.exception.errors[0]["error"])
        self.assertFalse(self.output.exists())
        self.assertFalse((self.base / "out" / "filled.fill_report.json").exists())

    def test_unknown_target_field_id_fails(self) -> None:
        manifest_path = self._write_manifest()
        self._write_plan([self._fill("C1-R99", "不存在的字段", "x", "x")])
        with self.assertRaises(filler.PlanValidationError) as ctx:
            filler.run_apply(manifest_path)
        self.assertIn("未知 targetFieldId", ctx.exception.errors[0]["error"])
        self.assertFalse(self.output.exists())

    def test_excerpt_not_found_downgrades_to_manual(self) -> None:
        manifest_path = self._write_manifest()
        self._write_plan([self._fill("C1-R03", "叶轮直径", "250", "叶轮直径 250 m")])
        result = filler.run_apply(manifest_path)
        decision = result["mapping"]["decisions"][2]
        self.assertEqual(decision["action"], "manual")
        self.assertEqual(decision["value"], "[待人工补充：叶轮直径]")
        self.assertIn("证据未命中", decision["reason"])
        row = Document(str(self.output)).tables[0].rows[3]
        self.assertEqual(row.cells[3].text.strip(), "[待人工补充：叶轮直径]")

    def test_unit_conversion_applied(self) -> None:
        manifest_path = self._write_manifest()
        # 模板单位列是 MW，计划值带 kW：走 normalize_value_for_field 的同族换算
        fill = self._fill("C1-R02", "单机容量", "6250", "单机容量 6250 kW", unit="kW")
        self._write_plan([fill])
        result = filler.run_apply(manifest_path)
        decision = result["mapping"]["decisions"][1]
        self.assertEqual(decision["action"], "fill")
        self.assertEqual(decision["value"], "6.25")
        row = Document(str(self.output)).tables[0].rows[2]
        self.assertEqual(row.cells[3].text.strip(), "6.25")

    def test_manual_action_forces_marker(self) -> None:
        manifest_path = self._write_manifest()
        self._write_plan([{"targetFieldId": "C1-R01", "field": "投标机型", "action": "manual", "value": "随便填的"}])
        result = filler.run_apply(manifest_path)
        decision = result["mapping"]["decisions"][0]
        self.assertEqual(decision["action"], "manual")
        self.assertEqual(decision["value"], "[待人工补充：投标机型]")
        row = Document(str(self.output)).tables[0].rows[1]
        self.assertEqual(row.cells[3].text.strip(), "[待人工补充：投标机型]")


class ApplyListTableTests(LlmPlanTestBase):
    """清单型附表一行多格：按列区分坐标，同一行可落多格。"""

    def test_same_row_multiple_columns_all_written(self) -> None:
        supply = _build_docx(self.base / "supply.docx", "附表B.1.1 供货范围清单", _SUPPLY_ROWS)
        material = _build_text_docx(self.base / "supply_material.docx", ["主控系统品牌为上海电气，数量共 60 台。"])
        self.material = material
        manifest_path = self._write_manifest(blank=supply, title="附表B.1.1 供货范围清单")
        filler.run_prepare(manifest_path)
        brief = json.loads((self.base / "fill_brief.json").read_text(encoding="utf-8"))
        row1 = [f for f in brief["targetFields"] if f["rowLabel"] == "主控系统"]
        by_column = {f["columnLabel"]: f for f in row1}
        fills = [
            self._fill(by_column["品牌或制造商名称"]["targetFieldId"], "主控系统 品牌或制造商名称", "上海电气", "上海电气"),
            self._fill(by_column["数量"]["targetFieldId"], "主控系统 数量", "60", "60 台"),
        ]
        self._write_plan(fills)
        result = filler.run_apply(manifest_path)
        filled = [d for d in result["mapping"]["decisions"] if d["action"] == "fill"]
        self.assertEqual(len(filled), 2)
        row = Document(str(self.output)).tables[0].rows[1]
        self.assertEqual(row.cells[2].text.strip(), "上海电气")
        self.assertEqual(row.cells[5].text.strip(), "60")


class ApplyEndToEndTests(LlmPlanTestBase):
    def test_correct_plan_fills_and_reports_llm_mode(self) -> None:
        manifest_path = self._write_manifest()
        fills = [
            self._fill("C1-R01", "投标机型", "EW6.25-220", "EW6.25-220"),
            self._fill("C1-R02", "单机容量", "6250", "单机容量 6250 kW", unit="kW"),
            # 叶轮直径不在计划里：按待人工处理
        ]
        self._write_plan(fills)
        result = filler.run_apply(manifest_path)

        table = Document(str(self.output)).tables[0]
        self.assertEqual(table.rows[1].cells[3].text.strip(), "EW6.25-220")
        self.assertEqual(table.rows[2].cells[3].text.strip(), "6.25")
        self.assertEqual(table.rows[3].cells[3].text.strip(), "[待人工补充：叶轮直径]")

        report_path = self.base / "out" / "filled.fill_report.json"
        self.assertTrue(report_path.exists())
        self.assertTrue((self.base / "out" / "filled.fill_report.md").exists())
        report = json.loads(report_path.read_text(encoding="utf-8"))
        self.assertEqual(report["fillReport"]["fillMode"], "llm-plan")
        self.assertEqual(report["fillReport"]["filledFieldCount"], 2)
        self.assertEqual(report["fillReport"]["unfilledFieldCount"], 1)

        # 每格证据进 evidenceRefs（含 excerpt），结构与现有 selected_fact 一致
        evidence = [e for e in report["evidenceRefs"] if e.get("type") == "selected_fact"]
        self.assertEqual(len(evidence), 2)
        self.assertEqual(evidence[0]["excerpt"], "EW6.25-220")
        self.assertTrue(evidence[0]["sourcePath"].endswith("material.docx"))

        # stdout 摘要与 compact_summary 同构
        summary = filler.compact_summary(result)
        for key in ("schema_version", "outputFile", "unfilledFields", "evidenceRefs", "fillReport", "filledAt"):
            self.assertIn(key, summary)
        self.assertEqual(summary["unfilledFields"], ["叶轮直径"])


if __name__ == "__main__":
    unittest.main()
