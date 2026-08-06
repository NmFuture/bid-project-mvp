"""组装验证的残留占位符扫描与 gap-planner 的待填写判定。

两处此前都只认「待填写」系列的部分写法：
- 组装验证只匹配前缀式 `[待填写：xxx]`，漏掉素材库占主流的后缀式 `[xxx，待填写]`，
  以及填写链路留下的 `[待人工补充：xxx]` / `[待人工插入：xxx]`；
- gap-planner 的文本判据不含「待插入」，只剩待插入占位符的模板拿不到 AI 填写任务。
fixture 为脱敏合成数据（通用领域词面，无真实项目数据）。
"""
from __future__ import annotations

import importlib.util
import sys
import tempfile
import unittest
from pathlib import Path


_BACKEND_VERIFY = (
    Path(__file__).resolve().parents[1]
    / "app"
    / "document_processing"
    / "technical_document"
    / "assembly"
    / "verify.py"
)
_SKILL_VERIFY = (
    Path(__file__).resolve().parents[1]
    / "opencode"
    / "skills"
    / "bid-tech-assembler"
    / "scripts"
    / "verify.py"
)
_PLANNER = (
    Path(__file__).resolve().parents[1]
    / "opencode"
    / "skills"
    / "bid-tech-gap-planner"
    / "scripts"
    / "run_from_manifest.py"
)


def _load(path: Path, name: str):
    spec = importlib.util.spec_from_file_location(name, path)
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    sys.modules[name] = module
    spec.loader.exec_module(module)
    return module


class PlaceholderPatternTests(unittest.TestCase):
    def setUp(self) -> None:
        self.verify = _load(_BACKEND_VERIFY, "assembly_verify_under_test")

    def _hits(self, text: str) -> list[str]:
        return [m.group(0) for pat in self.verify.PLACEHOLDER_PATTERNS for m in pat.finditer(text)]

    def test_prefix_form_still_matches(self) -> None:
        self.assertEqual(self._hits("[待填写：安全等级]"), ["[待填写：安全等级]"])
        self.assertEqual(self._hits("[缺失：发电小时数承诺函]"), ["[缺失：发电小时数承诺函]"])

    def test_suffix_form_is_matched(self) -> None:
        # 素材库里占主流的写法，修复前完全漏掉
        self.assertEqual(self._hits("[安全等级，待填写]"), ["[安全等级，待填写]"])
        self.assertEqual(self._hits("【安全等级, 待补充】"), ["【安全等级, 待补充】"])
        self.assertEqual(self._hits("[质保期，待确认]"), ["[质保期，待确认]"])

    def test_embed_placeholder_is_matched(self) -> None:
        self.assertEqual(self._hits("[设备清单，待插入]"), ["[设备清单，待插入]"])
        self.assertEqual(self._hits("详见[设备清单，待插入]后附"), ["[设备清单，待插入]"])

    def test_manual_markers_from_fill_pipeline_are_matched(self) -> None:
        self.assertEqual(self._hits("[待人工补充：投标机型]"), ["[待人工补充：投标机型]"])
        self.assertEqual(self._hits("[待人工插入：设备清单]"), ["[待人工插入：设备清单]"])

    def test_ordinary_bracketed_text_is_not_flagged(self) -> None:
        self.assertEqual(self._hits("参见[附录三]的说明"), [])
        self.assertEqual(self._hits("本机组[EW10.0-220-125]为投标机型"), [])

    def test_patterns_do_not_double_count(self) -> None:
        # 前缀式与后缀式两条规则若重叠，匹配结果会累进同一列表导致计数翻倍
        for text in ("[待填写：安全等级]", "[安全等级，待填写]", "[待人工插入：设备清单]"):
            self.assertEqual(len(self._hits(text)), 1, text)

    def test_skill_copy_stays_in_sync_with_backend(self) -> None:
        # 两份副本必须一致：backend 走 S4 组装，skill 副本在 opencode 容器里执行
        self.assertEqual(
            _BACKEND_VERIFY.read_text(encoding="utf-8"),
            _SKILL_VERIFY.read_text(encoding="utf-8"),
        )


class ScanDocxTests(unittest.TestCase):
    def setUp(self) -> None:
        self.verify = _load(_SKILL_VERIFY, "assembler_verify_under_test")

    def test_scan_reports_suffix_and_embed_placeholders_in_body_and_table(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as raw:
            path = Path(raw) / "assembled.docx"
            document = Document()
            document.add_paragraph("本项目安全等级为[安全等级，待填写]。")
            document.add_paragraph("[设备清单，待插入]")
            document.add_paragraph("[待人工插入：价格表]")
            table = document.add_table(rows=1, cols=1)
            table.rows[0].cells[0].text = "[质保期，待确认]"
            document.save(str(path))

            scan = self.verify.scan_docx(path)

        self.assertEqual(
            sorted(scan["placeholders"]),
            sorted(
                [
                    "[安全等级，待填写]",
                    "[设备清单，待插入]",
                    "[待人工插入：价格表]",
                    "[质保期，待确认]",
                ]
            ),
        )

    def test_clean_document_reports_no_placeholder(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as raw:
            path = Path(raw) / "clean.docx"
            document = Document()
            document.add_paragraph("本项目安全等级为一级，详见[附录三]。")
            document.save(str(path))

            scan = self.verify.scan_docx(path)

        self.assertEqual(scan["placeholders"], [])


class PlannerRequiresFillTests(unittest.TestCase):
    def setUp(self) -> None:
        self.planner = _load(_PLANNER, "gap_planner_under_test")

    def test_embed_only_template_requires_fill(self) -> None:
        # 文件名不带「待填写-」前缀时，此前只靠文件名巧合命中，改名即静默失效
        material = {"name": "塔架与基础工程量说明.docx", "matchReason": "正文含[塔架与基础工程量，待插入]"}

        self.assertTrue(self.planner.material_requires_fill(material))

    def test_fill_template_still_requires_fill(self) -> None:
        self.assertTrue(self.planner.material_requires_fill({"name": "待填写-塔筒设计方案专题报告.docx"}))
        self.assertTrue(self.planner.material_requires_fill({"name": "方案.docx", "matchReason": "含待补充占位符"}))

    def test_ready_material_does_not_require_fill(self) -> None:
        self.assertFalse(self.planner.material_requires_fill({"name": "塔架与基础工程量.docx"}))
        self.assertFalse(self.planner.material_requires_fill({"name": "基础弯矩表.docx", "folderPath": "项目定制/示例"}))

    def test_none_material_does_not_require_fill(self) -> None:
        self.assertFalse(self.planner.material_requires_fill(None))


if __name__ == "__main__":
    unittest.main()
