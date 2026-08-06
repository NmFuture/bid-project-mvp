"""素材后台深度解析（material_deep_parse）测试。

覆盖：PDF extract 全文提取（文字层优先、OCR 兜底、页数护栏、失败计数）、
XLSX 后台转 Word（仅技术标闸口 allowConvert 触发）、超大 docx 后台画像解析、
商务标闸口维持终态跳过、任务队列分发与本地兜底去重。
"""

from __future__ import annotations

import asyncio
import io
import unittest
from types import SimpleNamespace
from unittest.mock import AsyncMock, patch

from docx import Document

from app.services import material_deep_parse
from app.services.job_queue import KNOWN_JOB_TYPES
from app.services.material_deep_parse import (
    deep_parse_material_file,
    deep_parse_profile_for,
    deep_parse_status_allows_enqueue,
    enqueue_deep_parse_job,
    raw_file_deep_parse_kind,
)
from app.services.peripheral import PeripheralError
from app.services.technical_wiki_preview_generation import (
    PREVIEW_SCHEMA_VERSION,
    _build_preview_plans,
    _preview_signature,
    _profile_for_raw_file,
)
from app.services.wiki_blueprint_common import MAX_SYNC_DOCX_BYTES

OVER_LIMIT = MAX_SYNC_DOCX_BYTES + 1024


def _make_docx_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("总体方案", style="Heading 1")
    doc.add_paragraph("本机组适用于低温环境，额定功率 5.0MW。")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _raw_file(name: str, *, size: int = 1024, ext_fields: dict | None = None) -> SimpleNamespace:
    return SimpleNamespace(
        id=1,
        name=name,
        size_bytes=size,
        minio_bucket="bid-materials",
        minio_key="raw/1/file",
        ext_fields=dict(ext_fields or {}),
        folder=SimpleNamespace(path="技术标/标准文件/EW5.0", tier="", bid_type=""),
    )


class DeepParseKindTests(unittest.TestCase):
    def test_pdf_goes_extract_xlsx_goes_convert(self) -> None:
        self.assertEqual(raw_file_deep_parse_kind("报告.pdf", {}), "extract")
        self.assertEqual(raw_file_deep_parse_kind("台账.xlsx", {}), "convert")
        self.assertEqual(raw_file_deep_parse_kind("台账.xlsm", {}), "convert")

    def test_pdf_extract_disabled_falls_back_to_convert(self) -> None:
        with patch.object(material_deep_parse, "PDF_EXTRACT_ENABLED", False):
            self.assertEqual(raw_file_deep_parse_kind("报告.pdf", {}), "convert")

    def test_oversize_docx_needs_parse(self) -> None:
        self.assertEqual(raw_file_deep_parse_kind("方案.docx", {}), "parse")

    def test_pdf_with_legacy_cleaned_output_still_goes_extract(self) -> None:
        # 历史转换链产物（图片 docx）不被 extract 采信，PDF 始终走全文提取；
        # XLSX 有清洗稿且超上限时走清洗稿画像解析。
        self.assertEqual(
            raw_file_deep_parse_kind("扫描件.pdf", {"cleanedMinioKey": "legacy", "cleanedSize": OVER_LIMIT}),
            "extract",
        )
        for name in ["台账.xlsx", "台账.xls"]:
            self.assertEqual(
                raw_file_deep_parse_kind(name, {"cleanedMinioKey": "k", "cleanedSize": OVER_LIMIT}),
                "parse",
                name,
            )

    def test_small_or_converted_material_needs_nothing(self) -> None:
        self.assertEqual(raw_file_deep_parse_kind("图片.png", {}), "")

    def test_profile_only_used_when_source_key_matches(self) -> None:
        ext = {"deepParseProfile": {"sourceKey": "k1", "profile": {"headings": []}}}
        self.assertEqual(deep_parse_profile_for(ext, "k1"), {"headings": []})
        self.assertIsNone(deep_parse_profile_for(ext, "k2"))
        self.assertIsNone(deep_parse_profile_for({}, "k1"))

    def test_running_status_blocks_enqueue_until_stale(self) -> None:
        from datetime import UTC, datetime, timedelta

        fresh = {"deepParseStatus": "running", "deepParseUpdatedAt": datetime.now(UTC).isoformat()}
        stale = {
            "deepParseStatus": "running",
            "deepParseUpdatedAt": (datetime.now(UTC) - timedelta(hours=1)).isoformat(),
        }
        self.assertFalse(deep_parse_status_allows_enqueue(fresh))
        self.assertTrue(deep_parse_status_allows_enqueue(stale))
        self.assertTrue(deep_parse_status_allows_enqueue({"deepParseStatus": "failed"}))
        self.assertTrue(deep_parse_status_allows_enqueue({}))


class TechnicalProfileGateTests(unittest.TestCase):
    def test_pdf_without_extract_profile_marks_extract_pending(self) -> None:
        ext, profile = _profile_for_raw_file(_raw_file("检测报告.pdf"))

        self.assertEqual(ext, "pdf")
        self.assertTrue(profile["deepParsePending"])
        self.assertIn("全文提取", profile["parseError"])

    def test_pdf_extract_disabled_falls_back_to_convert_pending(self) -> None:
        # extract 关闭时 PDF 回落上游 convert 链：无清洗稿 → 排队转换。
        with patch("app.services.technical_wiki_preview_generation.PDF_EXTRACT_ENABLED", False):
            ext, profile = _profile_for_raw_file(_raw_file("检测报告.pdf"))

        self.assertEqual(ext, "pdf")
        self.assertTrue(profile["deepParsePending"])

    def test_pdf_ignores_legacy_cleaned_output(self) -> None:
        # extract 开启时 PDF 不采信 convert 链的图片 docx，始终走全文提取。
        ext, profile = _profile_for_raw_file(
            _raw_file(
                "检测报告.pdf",
                ext_fields={"cleanedMinioKey": "legacy.pdf.docx", "cleanedSize": OVER_LIMIT},
            )
        )

        self.assertEqual(ext, "pdf")
        self.assertTrue(profile["deepParsePending"])

    def test_xlsx_with_cleaned_output_is_upgraded_to_docx(self) -> None:
        # 转换产出的 cleaned docx 超同步上限：按 docx 走后台画像解析（上游 convert 链）。
        ext, profile = _profile_for_raw_file(
            _raw_file(
                "台账.xlsx",
                ext_fields={"cleanedMinioKey": "legacy.xlsx.docx", "cleanedSize": OVER_LIMIT},
            )
        )

        self.assertEqual(ext, "docx")
        self.assertTrue(profile["deepParsePending"])
        self.assertIn("30MB", profile["parseError"])

    def test_pdf_extract_profile_short_circuits_gate(self) -> None:
        profile_data = {
            "headings": [{"level": 1, "title": "型式认证证书"}],
            "paragraphs": ["认证编号 CGC2025461310034"],
            "tableCount": 0,
            "source": "pdfExtract",
            "fulltextKey": "parsed/RAW-0001/v1/fulltext.md",
        }
        item = _raw_file(
            "检测报告.pdf",
            ext_fields={"deepParseProfile": {"sourceKey": "raw/1/file", "profile": profile_data}},
        )

        ext, profile = _profile_for_raw_file(item)

        self.assertEqual(ext, "pdf")
        self.assertEqual(profile["headings"][0]["title"], "型式认证证书")
        self.assertEqual(profile["parseError"], "")
        self.assertEqual(profile["fulltextKey"], "parsed/RAW-0001/v1/fulltext.md")

    def test_pdf_extract_profile_expires_when_source_replaced(self) -> None:
        item = _raw_file(
            "检测报告.pdf",
            ext_fields={"deepParseProfile": {"sourceKey": "raw/1/old", "profile": {"headings": []}}},
        )

        _ext, profile = _profile_for_raw_file(item)

        self.assertTrue(profile["deepParsePending"])

    def test_oversize_docx_marks_deep_parse_pending(self) -> None:
        ext, profile = _profile_for_raw_file(_raw_file("方案.docx", size=OVER_LIMIT))

        self.assertEqual(ext, "docx")
        self.assertTrue(profile["deepParsePending"])
        self.assertIn("30MB", profile["parseError"])

    def test_deep_parse_profile_short_circuits_sync_gate(self) -> None:
        profile_data = {"headings": [{"level": 1, "title": "总体方案"}], "paragraphs": ["摘录"], "tableCount": 0}
        item = _raw_file(
            "方案.docx",
            size=OVER_LIMIT,
            ext_fields={"deepParseProfile": {"sourceKey": "raw/1/file", "profile": profile_data}},
        )

        ext, profile = _profile_for_raw_file(item)

        self.assertEqual(ext, "docx")
        self.assertEqual(profile["headings"][0]["title"], "总体方案")
        self.assertEqual(profile["parseError"], "")

    def test_non_convertible_file_keeps_empty_profile(self) -> None:
        ext, profile = _profile_for_raw_file(_raw_file("现场照片.png"))

        self.assertEqual(ext, "png")
        self.assertNotIn("deepParsePending", profile)


class _SingleExecuteSession:
    def __init__(self, items: list[object]) -> None:
        self._items = items

    async def __aenter__(self) -> "_SingleExecuteSession":
        return self

    async def __aexit__(self, *_args: object) -> None:
        return None

    async def execute(self, _statement: object) -> object:
        items = self._items

        class _Result:
            def scalars(self) -> "_Result":
                return self

            def all(self) -> list[object]:
                return list(items)

        return _Result()


class BuildPreviewPlansDeepParseTests(unittest.IsolatedAsyncioTestCase):
    async def test_pending_material_enqueues_and_stays_retryable(self) -> None:
        raw = _raw_file("检测报告.pdf")
        empty_profile = {
            "headings": [],
            "bodyHeadings": [],
            "paragraphs": [],
            "tables": [],
            "tableCount": 0,
            "parseError": "",
            "deepParsePending": True,
        }
        with (
            patch("app.models.async_session", return_value=_SingleExecuteSession([raw])),
            patch(
                "app.services.technical_wiki_preview_generation._profile_for_raw_file",
                return_value=("pdf", empty_profile),
            ),
            patch(
                "app.services.technical_wiki_preview_generation.enqueue_deep_parse_job",
            ) as enqueue_mock,
        ):
            plans, stats = await _build_preview_plans(
                [{"fileId": "RAW-0001", "tierCode": "standard", "tierLabel": "标准文件", "file": {"id": "RAW-0001"}}]
            )

        self.assertEqual(stats["skipped"], 1)
        payload = plans[0]["payload"]
        self.assertEqual(payload["status"], "fallback")
        self.assertTrue(payload["retryable"])
        self.assertIn("后台全文提取", payload["skipReason"])
        enqueue_mock.assert_called_once_with("RAW-0001", {"allowConvert": True, "bidType": "技术标"})

    async def test_oversize_docx_pending_keeps_deep_parse_message(self) -> None:
        raw = _raw_file("方案.docx")
        empty_profile = {
            "headings": [],
            "bodyHeadings": [],
            "paragraphs": [],
            "tables": [],
            "tableCount": 0,
            "parseError": "",
            "deepParsePending": True,
        }
        with (
            patch("app.models.async_session", return_value=_SingleExecuteSession([raw])),
            patch(
                "app.services.technical_wiki_preview_generation._profile_for_raw_file",
                return_value=("docx", empty_profile),
            ),
            patch(
                "app.services.technical_wiki_preview_generation.enqueue_deep_parse_job",
            ) as enqueue_mock,
        ):
            plans, _stats = await _build_preview_plans(
                [{"fileId": "RAW-0001", "tierCode": "standard", "tierLabel": "标准文件", "file": {"id": "RAW-0001"}}]
            )

        payload = plans[0]["payload"]
        self.assertTrue(payload["retryable"])
        self.assertIn("后台深度解析", payload["skipReason"])
        enqueue_mock.assert_called_once_with("RAW-0001", {"allowConvert": True, "bidType": "技术标"})

    async def test_pdf_extract_failures_over_limit_turn_terminal(self) -> None:
        raw = _raw_file(
            "检测报告.pdf",
            ext_fields={
                "deepParseStatus": "failed",
                "deepParseMessage": "后台全文提取失败：OCR 未配置",
                "deepParseFailCount": 3,
            },
        )
        empty_profile = {
            "headings": [],
            "bodyHeadings": [],
            "paragraphs": [],
            "tables": [],
            "tableCount": 0,
            "parseError": "",
            "deepParsePending": True,
        }
        with (
            patch("app.models.async_session", return_value=_SingleExecuteSession([raw])),
            patch(
                "app.services.technical_wiki_preview_generation._profile_for_raw_file",
                return_value=("pdf", empty_profile),
            ),
            patch(
                "app.services.technical_wiki_preview_generation.enqueue_deep_parse_job",
            ) as enqueue_mock,
        ):
            plans, _stats = await _build_preview_plans(
                [{"fileId": "RAW-0001", "tierCode": "standard", "tierLabel": "标准文件", "file": {"id": "RAW-0001"}}]
            )

        payload = plans[0]["payload"]
        self.assertFalse(payload["retryable"])
        self.assertIn("OCR 未配置", payload["skipReason"])
        enqueue_mock.assert_not_called()

    async def test_pdf_extract_profile_payload_carries_fulltext_info(self) -> None:
        raw = _raw_file("检测报告.pdf")
        extract_profile = {
            "headings": [{"level": 1, "title": "型式认证证书"}],
            "bodyHeadings": [],
            "paragraphs": ["认证编号 CGC2025461310034，产品型号 EW-CNVDL1146250"],
            "tables": [],
            "tableCount": 0,
            "parseError": "",
            "source": "pdfExtract",
            "pageCount": 12,
            "processedPages": 12,
            "truncated": False,
            "charCount": 3842,
            "fulltextKey": "parsed/RAW-0001/v1/fulltext.md",
        }
        with (
            patch("app.models.async_session", return_value=_SingleExecuteSession([raw])),
            patch(
                "app.services.technical_wiki_preview_generation._profile_for_raw_file",
                return_value=("pdf", extract_profile),
            ),
        ):
            plans, _stats = await _build_preview_plans(
                [{"fileId": "RAW-0001", "tierCode": "standard", "tierLabel": "标准文件", "file": {"id": "RAW-0001"}}]
            )

        plan = plans[0]
        self.assertFalse(plan.get("hit"))
        self.assertNotIn("payload", plan)
        self.assertEqual(plan["base"]["pdfFulltext"]["pageCount"], 12)
        self.assertEqual(plan["base"]["pdfFulltext"]["charCount"], 3842)

    async def test_pdf_stale_terminal_cache_does_not_block_extract(self) -> None:
        # 历史版本对 PDF 写的是终态 fallback 缓存（"非 docx"，retryable=False）。
        # 若签名不因 deepParsePending 变化，旧缓存会永远命中，extract 永远排不上。
        empty_profile = {
            "headings": [],
            "bodyHeadings": [],
            "paragraphs": [],
            "tables": [],
            "tableCount": 0,
            "parseError": "",
        }
        stale_signature = _preview_signature("检测报告.pdf", empty_profile, "技术标/标准文件/EW5.0")
        raw = _raw_file(
            "检测报告.pdf",
            ext_fields={
                "techWikiPreview": {
                    "schemaVersion": PREVIEW_SCHEMA_VERSION,
                    "signature": stale_signature,
                    "status": "fallback",
                    "retryable": False,
                    "skipReason": "非 docx，无可解析正文",
                    "preview": {"lead": "旧本地摘要", "source": "local"},
                }
            },
        )
        with (
            patch("app.models.async_session", return_value=_SingleExecuteSession([raw])),
            patch(
                "app.services.technical_wiki_preview_generation._profile_for_raw_file",
                return_value=("pdf", {**empty_profile, "deepParsePending": True}),
            ),
            patch(
                "app.services.technical_wiki_preview_generation.enqueue_deep_parse_job",
            ) as enqueue_mock,
        ):
            plans, stats = await _build_preview_plans(
                [{"fileId": "RAW-0001", "tierCode": "standard", "tierLabel": "标准文件", "file": {"id": "RAW-0001"}}]
            )

        self.assertEqual(stats["cached"], 0)
        payload = plans[0]["payload"]
        self.assertTrue(payload["retryable"])
        self.assertIn("后台全文提取", payload["skipReason"])
        enqueue_mock.assert_called_once_with("RAW-0001", {"allowConvert": True, "bidType": "技术标"})

    async def test_non_pending_skip_stays_terminal(self) -> None:
        raw = _raw_file("现场照片.png")
        empty_profile = {
            "headings": [],
            "bodyHeadings": [],
            "paragraphs": [],
            "tables": [],
            "tableCount": 0,
            "parseError": "",
        }
        with (
            patch("app.models.async_session", return_value=_SingleExecuteSession([raw])),
            patch(
                "app.services.technical_wiki_preview_generation._profile_for_raw_file",
                return_value=("png", empty_profile),
            ),
            patch(
                "app.services.technical_wiki_preview_generation.enqueue_deep_parse_job",
            ) as enqueue_mock,
        ):
            plans, _stats = await _build_preview_plans(
                [{"fileId": "RAW-0001", "tierCode": "standard", "tierLabel": "标准文件", "file": {"id": "RAW-0001"}}]
            )

        self.assertFalse(plans[0]["payload"]["retryable"])
        enqueue_mock.assert_not_called()


class BusinessProfileGateTests(unittest.TestCase):
    def _profile(self, item: SimpleNamespace) -> dict:
        from app.services.business_wiki_generation import _profile_raw_file

        return _profile_raw_file(item)

    def test_oversize_docx_enqueues_deep_parse(self) -> None:
        with patch("app.services.business_wiki_generation.enqueue_deep_parse_job") as enqueue_mock:
            profile = self._profile(_raw_file("合同.docx", size=OVER_LIMIT))

        self.assertTrue(profile["deepParsePending"])
        self.assertIn("后台深度解析", profile["parseError"])
        enqueue_mock.assert_called_once_with("RAW-0001", {"bidType": "商务标"})

    def test_xlsx_does_not_enqueue_background_convert(self) -> None:
        with patch("app.services.business_wiki_generation.enqueue_deep_parse_job") as enqueue_mock:
            profile = self._profile(_raw_file("价格表.xlsx"))

        self.assertNotIn("deepParsePending", profile)
        enqueue_mock.assert_not_called()

    def test_xlsx_ignores_legacy_cleaned_output(self) -> None:
        item = _raw_file(
            "价格表.xlsx",
            ext_fields={"cleanedMinioKey": "legacy.xlsx.docx", "cleanedSize": OVER_LIMIT},
        )
        with patch("app.services.business_wiki_generation.enqueue_deep_parse_job") as enqueue_mock:
            profile = self._profile(item)

        self.assertEqual(profile["ext"], "xlsx")
        self.assertFalse(profile["hasCleanedWord"])
        self.assertNotIn("deepParsePending", profile)
        enqueue_mock.assert_not_called()

    def test_deep_parse_profile_used_directly(self) -> None:
        profile_data = {"headings": [{"level": 1, "title": "资质"}], "paragraphs": [], "tableCount": 0}
        item = _raw_file(
            "合同.docx",
            size=OVER_LIMIT,
            ext_fields={"deepParseProfile": {"sourceKey": "raw/1/file", "profile": profile_data}},
        )
        with patch("app.services.business_wiki_generation.enqueue_deep_parse_job") as enqueue_mock:
            profile = self._profile(item)

        self.assertEqual(profile["headings"][0]["title"], "资质")
        self.assertEqual(profile["parseError"], "")
        enqueue_mock.assert_not_called()


class BusinessOcrGateTests(unittest.IsolatedAsyncioTestCase):
    async def test_pdf_source_ignores_fresh_legacy_ocr_cache(self) -> None:
        from app.services import business_wiki_generation

        signature = {
            "version": 1,
            "sourceMinioKey": "raw/1/report.pdf",
            "sourceSizeBytes": 1024,
            "cleanedMinioKey": "",
            "cleanedSize": 0,
        }
        item = SimpleNamespace(
            id=1,
            name="检测报告.pdf",
            version=1,
            size_bytes=1024,
            minio_bucket="bid-materials",
            minio_key="raw/1/report.pdf",
            ext_fields={
                "businessWikiOcr": {
                    "schemaVersion": business_wiki_generation.BUSINESS_WIKI_OCR_VERSION,
                    "signature": signature,
                    "status": "completed",
                    "sourceType": "source_file",
                    "text": "历史 PDF OCR 文本",
                }
            },
        )

        payload = await business_wiki_generation._ensure_business_wiki_ocr_cache(
            item,
            {"sourceExt": "pdf", "ext": "pdf"},
        )

        self.assertEqual(payload["status"], "not_required")
        self.assertEqual(payload["text"], "")

    async def test_pdf_source_does_not_trigger_business_wiki_ocr(self) -> None:
        from app.services import business_wiki_generation

        item = SimpleNamespace(
            id=1,
            name="检测报告.pdf",
            version=1,
            size_bytes=1024,
            minio_bucket="bid-materials",
            minio_key="raw/1/report.pdf",
            ext_fields={},
        )
        with patch.object(
            business_wiki_generation,
            "_recognize_source_file_for_wiki",
            new=AsyncMock(return_value={"status": "completed", "text": "不应读取"}),
        ) as recognize_mock:
            payload = await business_wiki_generation._ensure_business_wiki_ocr_cache(
                item,
                {"sourceExt": "pdf", "ext": "pdf"},
            )

        self.assertEqual(payload["status"], "not_required")
        recognize_mock.assert_not_awaited()


class _DeepParseSession:
    def __init__(self, item: SimpleNamespace) -> None:
        self._item = item

    async def __aenter__(self) -> "_DeepParseSession":
        return self

    async def __aexit__(self, *_args: object) -> None:
        return None

    async def execute(self, _statement: object) -> object:
        item = self._item
        return SimpleNamespace(scalar_one_or_none=lambda: item)

    async def get(self, _model: object, _pk: object) -> SimpleNamespace:
        return self._item

    async def commit(self) -> None:
        return None


def _make_pdf_bytes(page_texts: list[str]) -> bytes:
    import fitz

    doc = fitz.open()
    for text in page_texts:
        page = doc.new_page()
        if text:
            page.insert_text((72, 72), text)
    return doc.tobytes()


TEXT_PAGE = "CGC Type Certification Certificate for converter EW-CNVDL1146250 issued on 2025-04-07."


class DeepParseMaterialFileTests(unittest.IsolatedAsyncioTestCase):
    async def test_status_write_rejects_replaced_source(self) -> None:
        item = _raw_file("检测报告.pdf")
        item.version = 2
        before = dict(item.ext_fields)
        with patch(
            "app.services.material_deep_parse.async_session",
            return_value=_DeepParseSession(item),
        ):
            persisted = await material_deep_parse._write_deep_parse_status(
                1,
                "parsed",
                "old job",
                expected_source_version=1,
                expected_source_key="raw/1/file",
            )

        self.assertFalse(persisted)
        self.assertEqual(item.ext_fields, before)

    async def test_oversize_docx_writes_profile(self) -> None:
        item = _raw_file("方案.docx", size=OVER_LIMIT)
        docx_bytes = _make_docx_bytes()
        with (
            patch("app.services.material_deep_parse.async_session", return_value=_DeepParseSession(item)),
            patch("app.services.material_deep_parse.minio_client") as minio_mock,
        ):
            minio_mock.get_object.return_value = docx_bytes
            result = await deep_parse_material_file("RAW-0001")

        self.assertEqual(result["deepParseStatus"], "parsed")
        profile = item.ext_fields["deepParseProfile"]
        self.assertEqual(profile["sourceKey"], "raw/1/file")
        self.assertEqual(profile["profile"]["headings"][0]["title"], "总体方案")
        self.assertEqual(item.ext_fields["deepParseStatus"], "parsed")

    # ---- convert 分支（上游 #130）：XLSX 与 extract 关闭时的 PDF ----

    async def test_xlsx_without_allow_convert_is_not_converted(self) -> None:
        # 商务标闸口排队不带 allowConvert 标记：非 Word 素材维持终态，不触发转换
        item = _raw_file("台账.xlsx")
        with (
            patch("app.services.material_deep_parse.async_session", return_value=_DeepParseSession(item)),
            patch("app.services.material_deep_parse.clean_material_file") as clean_mock,
        ):
            result = await deep_parse_material_file("RAW-0001")

        self.assertEqual(result["deepParseStatus"], "failed")
        self.assertIn("不支持", result["deepParseMessage"])
        clean_mock.assert_not_called()

    async def test_pdf_extract_disabled_falls_back_to_convert_branch(self) -> None:
        # extract 关闭时 PDF 回落 convert 链：无 allowConvert 标记维持终态，不触发转换
        item = _raw_file("检测报告.pdf")
        with (
            patch.object(material_deep_parse, "PDF_EXTRACT_ENABLED", False),
            patch("app.services.material_deep_parse.async_session", return_value=_DeepParseSession(item)),
            patch("app.services.material_deep_parse.clean_material_file") as clean_mock,
        ):
            result = await deep_parse_material_file("RAW-0001")

        self.assertEqual(result["deepParseStatus"], "failed")
        self.assertIn("不支持", result["deepParseMessage"])
        clean_mock.assert_not_called()

    async def test_xlsx_converts_then_ready_when_cleaned_is_small(self) -> None:
        item = _raw_file("台账.xlsx")

        async def fake_clean(file_id: str, data: dict | None = None, *, allow_convert: bool = False) -> dict:
            assert allow_convert is True
            item.ext_fields = {
                "cleanedMinioKey": "cleaned/1.docx",
                "cleanedMinioBucket": "bid-materials",
                "cleanedSize": 2048,
            }
            return {"cleanStatus": "cleaned"}

        with (
            patch("app.services.material_deep_parse.async_session", return_value=_DeepParseSession(item)),
            patch("app.services.material_deep_parse.clean_material_file", side_effect=fake_clean),
        ):
            result = await deep_parse_material_file("RAW-0001", {"allowConvert": True})

        self.assertEqual(result["deepParseStatus"], "ready")
        self.assertEqual(item.ext_fields["deepParseStatus"], "ready")

    async def test_xlsx_convert_without_output_fails_with_message(self) -> None:
        item = _raw_file("台账.xlsx")

        async def fake_clean(file_id: str, data: dict | None = None, *, allow_convert: bool = False) -> dict:
            return {"cleanStatus": "failed", "cleanMessage": "转换失败：无可用转换工具。"}

        with (
            patch("app.services.material_deep_parse.async_session", return_value=_DeepParseSession(item)),
            patch("app.services.material_deep_parse.clean_material_file", side_effect=fake_clean),
        ):
            result = await deep_parse_material_file("RAW-0001", {"allowConvert": True})

        self.assertEqual(result["deepParseStatus"], "failed")
        self.assertIn("转换失败", result["deepParseMessage"])
        self.assertEqual(item.ext_fields["deepParseStatus"], "failed")

    async def test_xlsx_with_cleaned_output_parses_cleaned_docx(self) -> None:
        # 转换产出的 cleaned docx 超同步上限：直接对清洗稿做后台画像解析
        item = _raw_file(
            "台账.xlsx",
            ext_fields={
                "cleanedMinioKey": "legacy.xlsx.docx",
                "cleanedMinioBucket": "bid-materials",
                "cleanedSize": OVER_LIMIT,
            },
        )
        with (
            patch("app.services.material_deep_parse.async_session", return_value=_DeepParseSession(item)),
            patch("app.services.material_deep_parse.minio_client") as minio_mock,
        ):
            minio_mock.get_object.return_value = _make_docx_bytes()
            result = await deep_parse_material_file("RAW-0001")

        self.assertEqual(result["deepParseStatus"], "parsed")
        minio_mock.get_object.assert_called_once_with("bid-materials", "legacy.xlsx.docx")
        self.assertEqual(item.ext_fields["deepParseProfile"]["sourceKey"], "legacy.xlsx.docx")

    # ---- PDF extract 全文提取 ----

    async def test_pdf_text_layer_extracts_without_ocr(self) -> None:
        item = _raw_file("检测报告.pdf")
        pdf_bytes = _make_pdf_bytes([TEXT_PAGE, TEXT_PAGE])
        with (
            patch("app.services.material_deep_parse.async_session", return_value=_DeepParseSession(item)),
            patch("app.services.material_deep_parse.minio_client") as minio_mock,
            patch(
                "app.services.ocr_service.ocr_service.recognize_text_for_parse",
                new=AsyncMock(return_value=("不应触发", {})),
            ) as ocr_mock,
        ):
            minio_mock.get_object.return_value = pdf_bytes
            result = await deep_parse_material_file("RAW-0001")

        self.assertEqual(result["deepParseStatus"], "parsed")
        ocr_mock.assert_not_awaited()
        envelope = item.ext_fields["deepParseProfile"]
        self.assertEqual(envelope["sourceKey"], "raw/1/file")
        profile = envelope["profile"]
        self.assertEqual(profile["source"], "pdfExtract")
        self.assertEqual(profile["pageCount"], 2)
        self.assertFalse(profile["truncated"])
        self.assertEqual(profile["ocrPages"], 0)
        self.assertTrue(profile["paragraphs"])
        self.assertEqual(profile["fulltextKey"], "parsed/RAW-0001/v1/fulltext.md")
        # 全文落 MinIO，内容含页标记与文字层文本。
        put_call = minio_mock.put_object.call_args
        self.assertEqual(put_call.args[1], "parsed/RAW-0001/v1/fulltext.md")
        fulltext = put_call.args[2].decode("utf-8")
        self.assertIn("<!-- 第 1 页 -->", fulltext)
        self.assertIn("EW-CNVDL1146250", fulltext)
        self.assertEqual(item.ext_fields["deepParseFailCount"], 0)

    async def test_pdf_extract_discards_fulltext_when_source_changes_before_finalize(self) -> None:
        item = _raw_file("检测报告.pdf")
        pdf_bytes = _make_pdf_bytes([TEXT_PAGE])

        async def write_status(_numeric_id: int, status: str, _message: str, **_kwargs: object) -> bool:
            if status == "running":
                return True
            item.version = 2
            item.minio_key = "raw/1/replaced.pdf"
            return False

        with (
            patch("app.services.material_deep_parse.async_session", return_value=_DeepParseSession(item)),
            patch.object(material_deep_parse, "_write_deep_parse_status", new=write_status),
            patch("app.services.material_deep_parse.minio_client") as minio_mock,
        ):
            minio_mock.get_object.return_value = pdf_bytes
            result = await deep_parse_material_file("RAW-0001")

        self.assertEqual(result["deepParseStatus"], "stale")
        minio_mock.put_object.assert_called_once()
        minio_mock.remove_object.assert_called_once_with(
            "bid-materials",
            "parsed/RAW-0001/v1/fulltext.md",
        )

    async def test_pdf_extract_preserves_fulltext_when_source_key_changes_without_new_version(self) -> None:
        item = _raw_file("检测报告.pdf")
        pdf_bytes = _make_pdf_bytes([TEXT_PAGE])

        async def write_status(_numeric_id: int, status: str, _message: str, **_kwargs: object) -> bool:
            if status == "running":
                return True
            item.minio_key = "raw/1/moved.pdf"
            return False

        with (
            patch("app.services.material_deep_parse.async_session", return_value=_DeepParseSession(item)),
            patch.object(material_deep_parse, "_write_deep_parse_status", new=write_status),
            patch("app.services.material_deep_parse.minio_client") as minio_mock,
        ):
            minio_mock.get_object.return_value = pdf_bytes
            result = await deep_parse_material_file("RAW-0001")

        self.assertEqual(result["deepParseStatus"], "stale")
        minio_mock.put_object.assert_called_once()
        minio_mock.remove_object.assert_not_called()

    async def test_pdf_extract_preserves_current_fulltext_when_finalize_raises(self) -> None:
        item = _raw_file("检测报告.pdf")
        pdf_bytes = _make_pdf_bytes([TEXT_PAGE])
        write_status = AsyncMock(
            side_effect=[
                True,
                RuntimeError("database unavailable"),
                RuntimeError("database unavailable"),
            ]
        )
        with (
            patch("app.services.material_deep_parse.async_session", return_value=_DeepParseSession(item)),
            patch.object(material_deep_parse, "_write_deep_parse_status", new=write_status),
            patch("app.services.material_deep_parse.minio_client") as minio_mock,
            self.assertRaisesRegex(RuntimeError, "database unavailable"),
        ):
            minio_mock.get_object.return_value = pdf_bytes
            await deep_parse_material_file("RAW-0001")

        minio_mock.put_object.assert_called_once()
        minio_mock.remove_object.assert_not_called()
        self.assertEqual(write_status.await_count, 3)

    async def test_pdf_extract_preserves_current_fulltext_when_finalize_is_cancelled(self) -> None:
        item = _raw_file("检测报告.pdf")
        pdf_bytes = _make_pdf_bytes([TEXT_PAGE])
        write_status = AsyncMock(side_effect=[True, asyncio.CancelledError()])
        with (
            patch("app.services.material_deep_parse.async_session", return_value=_DeepParseSession(item)),
            patch.object(material_deep_parse, "_write_deep_parse_status", new=write_status),
            patch("app.services.material_deep_parse.minio_client") as minio_mock,
            self.assertRaises(asyncio.CancelledError),
        ):
            minio_mock.get_object.return_value = pdf_bytes
            await deep_parse_material_file("RAW-0001")

        minio_mock.put_object.assert_called_once()
        minio_mock.remove_object.assert_not_called()
        self.assertEqual(write_status.await_count, 2)

    async def test_pdf_extract_keeps_fulltext_when_finalize_commit_result_is_uncertain(self) -> None:
        item = _raw_file("检测报告.pdf")
        pdf_bytes = _make_pdf_bytes([TEXT_PAGE])
        statuses: list[str] = []

        async def write_status(_numeric_id: int, status: str, _message: str, **_kwargs: object) -> bool:
            statuses.append(status)
            if status == "running":
                return True
            if status == "parsed":
                item.ext_fields = {
                    "deepParseProfile": {
                        "sourceKey": "raw/1/file",
                        "profile": {
                            "fulltextBucket": "bid-materials",
                            "fulltextKey": "parsed/RAW-0001/v1/fulltext.md",
                        },
                    }
                }
            raise RuntimeError("database response lost")

        with (
            patch("app.services.material_deep_parse.async_session", return_value=_DeepParseSession(item)),
            patch.object(material_deep_parse, "_write_deep_parse_status", new=write_status),
            patch("app.services.material_deep_parse.minio_client") as minio_mock,
            self.assertRaisesRegex(RuntimeError, "database response lost"),
        ):
            minio_mock.get_object.return_value = pdf_bytes
            await deep_parse_material_file("RAW-0001")

        minio_mock.put_object.assert_called_once()
        minio_mock.remove_object.assert_not_called()
        self.assertEqual(statuses, ["running", "parsed", "failed"])

    async def test_pdf_scanned_pages_fall_back_to_ocr(self) -> None:
        item = _raw_file("扫描证书.pdf")
        pdf_bytes = _make_pdf_bytes(["", ""])
        with (
            patch("app.services.material_deep_parse.async_session", return_value=_DeepParseSession(item)),
            patch("app.services.material_deep_parse.minio_client") as minio_mock,
            patch(
                "app.services.ocr_service.ocr_service.recognize_text_for_parse",
                new=AsyncMock(return_value=("扫描页 OCR 文本内容", {"pageCount": 1})),
            ) as ocr_mock,
        ):
            minio_mock.get_object.return_value = pdf_bytes
            result = await deep_parse_material_file("RAW-0001")

        self.assertEqual(result["deepParseStatus"], "parsed")
        self.assertEqual(ocr_mock.await_count, 2)
        # 按页送 OCR：文件名带页码、内容是 PNG 字节。
        ocr_kwargs = ocr_mock.await_args_list[0].kwargs
        self.assertTrue(ocr_kwargs["file_name"].endswith("-p1.png"))
        self.assertEqual(ocr_kwargs["mime_type"], "image/png")
        profile = item.ext_fields["deepParseProfile"]["profile"]
        self.assertEqual(profile["ocrPages"], 2)
        self.assertEqual(profile["ocrFailedPages"], 0)
        fulltext = minio_mock.put_object.call_args.args[2].decode("utf-8")
        self.assertIn("扫描页 OCR 文本内容", fulltext)

    async def test_pdf_extract_respects_max_pages_guardrail(self) -> None:
        item = _raw_file("长报告.pdf")
        pdf_bytes = _make_pdf_bytes([TEXT_PAGE, TEXT_PAGE, TEXT_PAGE])
        with (
            patch("app.services.material_deep_parse.async_session", return_value=_DeepParseSession(item)),
            patch("app.services.material_deep_parse.minio_client") as minio_mock,
            patch.object(material_deep_parse, "PDF_EXTRACT_MAX_PAGES", 1),
        ):
            minio_mock.get_object.return_value = pdf_bytes
            result = await deep_parse_material_file("RAW-0001")

        self.assertEqual(result["deepParseStatus"], "parsed")
        self.assertIn("已截断", result["deepParseMessage"])
        profile = item.ext_fields["deepParseProfile"]["profile"]
        self.assertTrue(profile["truncated"])
        self.assertEqual(profile["pageCount"], 3)
        self.assertEqual(profile["processedPages"], 1)

    async def test_pdf_extract_all_empty_marks_failed_and_counts(self) -> None:
        item = _raw_file("扫描件.pdf")
        pdf_bytes = _make_pdf_bytes([""])
        with (
            patch("app.services.material_deep_parse.async_session", return_value=_DeepParseSession(item)),
            patch("app.services.material_deep_parse.minio_client") as minio_mock,
            patch(
                "app.services.ocr_service.ocr_service.recognize_text_for_parse",
                new=AsyncMock(side_effect=RuntimeError("OCR 未配置")),
            ),
        ):
            minio_mock.get_object.return_value = pdf_bytes
            result = await deep_parse_material_file("RAW-0001")

        self.assertEqual(result["deepParseStatus"], "failed")
        self.assertIn("文字层为空", result["deepParseMessage"])
        self.assertEqual(item.ext_fields["deepParseStatus"], "failed")
        self.assertEqual(item.ext_fields["deepParseFailCount"], 1)
        minio_mock.put_object.assert_not_called()

    async def test_unsupported_type_fails_explicitly(self) -> None:
        item = _raw_file("现场照片.png")
        with patch("app.services.material_deep_parse.async_session", return_value=_DeepParseSession(item)):
            result = await deep_parse_material_file("RAW-0001")

        self.assertEqual(result["deepParseStatus"], "failed")
        self.assertEqual(item.ext_fields["deepParseStatus"], "failed")


class PdfFulltextServiceTests(unittest.IsolatedAsyncioTestCase):
    def _item_with_extract(self) -> SimpleNamespace:
        return _raw_file(
            "检测报告.pdf",
            ext_fields={
                "deepParseProfile": {
                    "schemaVersion": 1,
                    "sourceKey": "raw/1/file",
                    "profile": {
                        "source": "pdfExtract",
                        "pageCount": 12,
                        "processedPages": 12,
                        "truncated": False,
                        "charCount": 100,
                        "fulltextBucket": "bid-materials",
                        "fulltextKey": "parsed/RAW-0001/v1/fulltext.md",
                    },
                }
            },
        )

    async def test_fulltext_returns_minio_text(self) -> None:
        from app.services.material_deep_parse import pdf_fulltext_for_raw_file

        item = self._item_with_extract()
        with (
            patch("app.services.material_deep_parse.async_session", return_value=_DeepParseSession(item)),
            patch("app.services.material_deep_parse.minio_client") as minio_mock,
        ):
            minio_mock.get_object.return_value = "<!-- 第 1 页 -->\n\n正文".encode("utf-8")
            payload = await pdf_fulltext_for_raw_file("RAW-0001")

        self.assertEqual(payload["fileId"], "RAW-0001")
        self.assertIn("正文", payload["text"])
        self.assertEqual(payload["pageCount"], 12)
        self.assertFalse(payload["truncated"])
        minio_mock.get_object.assert_called_once_with("bid-materials", "parsed/RAW-0001/v1/fulltext.md")

    async def test_fulltext_not_ready_raises_404(self) -> None:
        from app.services.material_deep_parse import pdf_fulltext_for_raw_file

        item = _raw_file("检测报告.pdf")
        with patch("app.services.material_deep_parse.async_session", return_value=_DeepParseSession(item)):
            with self.assertRaises(PeripheralError) as ctx:
                await pdf_fulltext_for_raw_file("RAW-0001")

        self.assertEqual(ctx.exception.status_code, 404)
        self.assertEqual(ctx.exception.code, "FULLTEXT_NOT_READY")

    async def test_fulltext_ignores_tampered_bucket_and_key(self) -> None:
        # 元数据 fulltextBucket/fulltextKey 可被 extFields 外部写入，不得作为读路径依据：
        # 读取必须固定走 materials 桶 + 按素材 id/版本本地重算的规范 key。
        from app.services.material_deep_parse import pdf_fulltext_for_raw_file

        item = self._item_with_extract()
        profile = item.ext_fields["deepParseProfile"]["profile"]
        profile["fulltextBucket"] = "other-bucket"
        profile["fulltextKey"] = "parsed/RAW-9999/v1/fulltext.md"
        with (
            patch("app.services.material_deep_parse.async_session", return_value=_DeepParseSession(item)),
            patch("app.services.material_deep_parse.minio_client") as minio_mock,
        ):
            minio_mock.get_object.return_value = "<!-- 第 1 页 -->\n\n正文".encode("utf-8")
            payload = await pdf_fulltext_for_raw_file("RAW-0001")

        self.assertIn("正文", payload["text"])
        minio_mock.get_object.assert_called_once_with("bid-materials", "parsed/RAW-0001/v1/fulltext.md")

    async def test_fulltext_rejects_non_pdf(self) -> None:
        from app.services.material_deep_parse import pdf_fulltext_for_raw_file

        item = _raw_file("方案.docx")
        with patch("app.services.material_deep_parse.async_session", return_value=_DeepParseSession(item)):
            with self.assertRaises(PeripheralError) as ctx:
                await pdf_fulltext_for_raw_file("RAW-0001")

        self.assertEqual(ctx.exception.code, "FULLTEXT_FILE_TYPE_INVALID")


class ExtractExcerptTests(unittest.TestCase):
    def test_excerpt_respects_limits(self) -> None:
        lines = [f"第 {index} 段正文内容，包含足够的字数用于摘录采样。" for index in range(100)]

        excerpt = material_deep_parse._excerpt_paragraphs(lines)

        self.assertLessEqual(len(excerpt), material_deep_parse.EXCERPT_PARAGRAPH_LIMIT)
        self.assertLessEqual(sum(len(item) for item in excerpt), material_deep_parse.EXCERPT_TOTAL_CHARS)
        self.assertTrue(all(len(item) <= material_deep_parse.EXCERPT_PARAGRAPH_CHARS for item in excerpt))
        self.assertEqual(excerpt[0][:8], lines[0][:8])
        self.assertIn(lines[-1][: material_deep_parse.EXCERPT_PARAGRAPH_CHARS], excerpt[-1])

    def test_heading_heuristic_finds_numbered_lines(self) -> None:
        lines = ["1.1 总体技术方案", "本段是正文，句号结尾。", "附录 A 规范性附录"]

        headings = material_deep_parse._guess_pdf_headings(lines, "报告")

        self.assertEqual(headings[0]["title"], "1.1 总体技术方案")

    def test_heading_falls_back_to_filename(self) -> None:
        headings = material_deep_parse._guess_pdf_headings(["这一段正文很长且没有编号。"], "CGC 证书")

        self.assertEqual(headings, [{"level": 1, "title": "CGC 证书"}])


class EnqueueAndWorkerTests(unittest.TestCase):
    def test_job_type_registered(self) -> None:
        self.assertIn("material_deep_parse", KNOWN_JOB_TYPES)

    def test_local_fallback_when_redis_unavailable_and_deduped(self) -> None:
        material_deep_parse._local_inflight.clear()
        with (
            patch(
                "app.services.material_deep_parse.enqueue_generation_job",
                side_effect=RuntimeError("redis down"),
            ),
            patch("app.services.material_deep_parse.submit_local_job") as local_mock,
        ):
            first = enqueue_deep_parse_job("RAW-0001", {"bidType": "技术标"})
            second = enqueue_deep_parse_job("RAW-0001", {"bidType": "技术标"})

        self.assertTrue(first["queued"])
        self.assertTrue(first["local"])
        self.assertTrue(second["deduped"])
        local_mock.assert_called_once()
        material_deep_parse._local_inflight.clear()

    def test_worker_dispatches_deep_parse_job(self) -> None:
        from app.workers import redis_worker

        with patch(
            "app.services.material_deep_parse.deep_parse_material_file_sync",
            return_value={"deepParseStatus": "parsed", "deepParseMessage": "ok"},
        ) as runner_mock, patch(
            "app.workers.redis_worker.renew_generation_lock",
            return_value=True,
        ):
            redis_worker._run_job(
                {"id": "job-1", "type": "material_deep_parse", "projectId": "RAW-0001", "data": {"fileId": "RAW-0001"}}
            )

        runner_mock.assert_called_once_with("RAW-0001", {"fileId": "RAW-0001"})

    def test_enqueue_deep_parse_job_requires_bid_type(self) -> None:
        # 深度解析任务必须携带并校验 bidType（R10-B04-01）。
        with self.assertRaises(ValueError):
            enqueue_deep_parse_job("RAW-0001")
        with self.assertRaises(ValueError):
            enqueue_deep_parse_job("RAW-0001", {"bidType": "标书"})

    def test_worker_passes_bid_type_to_deep_parse_hook(self) -> None:
        from app.workers import redis_worker

        with patch(
            "app.services.material_deep_parse.deep_parse_material_file_sync",
            return_value={"deepParseStatus": "parsed", "deepParseMessage": "ok"},
        ), patch(
            "app.workers.redis_worker.renew_generation_lock",
            return_value=True,
        ), patch(
            "app.services.material_wiki_auto.on_material_deep_parse_job_finished",
        ) as hook_mock:
            redis_worker._run_job(
                {
                    "id": "job-2",
                    "type": "material_deep_parse",
                    "projectId": "RAW-0001",
                    "data": {"fileId": "RAW-0001", "bidType": "商务标"},
                }
            )

        hook_mock.assert_called_once_with("RAW-0001", current_job_id="job-2", bid_type="商务标")

    def test_worker_passes_bid_type_to_cleaning_hook(self) -> None:
        from app.workers import redis_worker

        with patch(
            "app.services.material_cleaning.clean_material_file_sync",
            return_value={"cleanStatus": "cleaned", "cleanMessage": "ok"},
        ), patch(
            "app.workers.redis_worker.renew_generation_lock",
            return_value=True,
        ), patch(
            "app.services.material_wiki_auto.on_material_cleaning_job_finished",
        ) as hook_mock:
            redis_worker._run_job(
                {
                    "id": "job-3",
                    "type": "material_cleaning",
                    "projectId": "RAW-0001",
                    "data": {"fileId": "RAW-0001", "bidType": "商务标"},
                }
            )

        hook_mock.assert_called_once_with(current_job_id="job-3", bid_type="商务标")


if __name__ == "__main__":
    unittest.main()
