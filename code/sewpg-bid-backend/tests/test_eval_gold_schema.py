from __future__ import annotations

import importlib.util
import sys
import tempfile
import unittest
from pathlib import Path


EVAL_SCRIPT_DIR = Path(__file__).resolve().parents[1] / "eval" / "scripts"


def load_eval_script(name: str):
    module_name = f"eval_{name}"
    spec = importlib.util.spec_from_file_location(module_name, EVAL_SCRIPT_DIR / f"{name}.py")
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    sys.path.insert(0, str(EVAL_SCRIPT_DIR))
    try:
        sys.modules[module_name] = module
        spec.loader.exec_module(module)
        return module
    finally:
        sys.path.remove(str(EVAL_SCRIPT_DIR))


class EvalGoldSchemaTests(unittest.TestCase):
    def test_gold_schema_loads_valid_rows(self) -> None:
        gold_schema = load_eval_script("gold_schema")

        row = gold_schema.GoldRow.from_csv_row(
            {
                "project_id": "PRJ-0003",
                "doc_type": "appendix",
                "locator": "附表C/表1/行3",
                "field_name": "投标机型",
                "human_answer": "SE-15530",
                "field_type": "phrase",
                "difficulty_tier": "T2",
                "evidence_source": "投标机型参数表",
            }
        )

        self.assertEqual(row.key.project_id, "PRJ-0003")
        self.assertEqual(row.doc_type.value, "appendix")
        self.assertEqual(row.field_type.value, "phrase")

    def test_prediction_schema_loads_valid_rows(self) -> None:
        gold_schema = load_eval_script("gold_schema")

        row = gold_schema.PredictionRow.from_csv_row(
            {
                "project_id": "PRJ-0003",
                "doc_type": "body",
                "locator": "第1章/技术评分标准索引表/[投标机型]",
                "field_name": "投标机型",
                "predicted_answer": "SE-15530",
                "mark_yellow": "false",
                "evidence_refs": "投标机型参数表",
                "skill_name": "bid-tech-word-placeholder-filler",
                "fill_path": "kb_lookup",
                "confidence": "0.92",
            }
        )

        self.assertEqual(row.key.project_id, "PRJ-0003")
        self.assertFalse(row.mark_yellow)
        self.assertEqual(row.confidence, 0.92)

    def test_prediction_schema_rejects_bad_confidence(self) -> None:
        gold_schema = load_eval_script("gold_schema")

        with self.assertRaises(ValueError):
            gold_schema.PredictionRow.from_csv_row(
                {
                    "project_id": "PRJ-0003",
                    "doc_type": "body",
                    "locator": "第1章/技术评分标准索引表/[投标机型]",
                    "field_name": "投标机型",
                    "predicted_answer": "SE-15530",
                    "mark_yellow": "false",
                    "evidence_refs": "",
                    "skill_name": "bid-tech-word-placeholder-filler",
                    "fill_path": "kb_lookup",
                    "confidence": "1.5",
                }
            )

    def test_gold_schema_rejects_missing_non_t5_answer(self) -> None:
        gold_schema = load_eval_script("gold_schema")

        with self.assertRaises(ValueError):
            gold_schema.GoldRow.from_csv_row(
                {
                    "project_id": "PRJ-0003",
                    "doc_type": "appendix",
                    "locator": "附表C/表1/行3",
                    "field_name": "投标机型",
                    "human_answer": "",
                    "field_type": "phrase",
                    "difficulty_tier": "T2",
                    "evidence_source": "",
                }
            )

    def test_validate_gold_detects_duplicate_keys(self) -> None:
        validate_gold = load_eval_script("validate_gold")

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "gold.csv"
            path.write_text(
                "\n".join(
                    [
                        "project_id,doc_type,locator,field_name,human_answer,field_type,difficulty_tier,evidence_source",
                        "PRJ-0003,appendix,附表C/表1/行3,投标机型,SE-15530,phrase,T2,投标机型参数表",
                        "PRJ-0003,appendix,附表C/表1/行3,投标机型,SE-15530,phrase,T2,投标机型参数表",
                    ]
                ),
                encoding="utf-8",
            )

            errors = validate_gold.validate_gold_paths([path])

        self.assertEqual(len(errors), 1)
        self.assertIn("duplicate key", errors[0])

    def test_extract_gold_draft_writes_body_placeholders(self) -> None:
        extract_gold_draft = load_eval_script("extract_gold_draft")

        with tempfile.TemporaryDirectory() as tmp:
            from docx import Document

            root = Path(tmp)
            docx_path = root / "body.docx"
            output_path = root / "body.csv"
            document = Document()
            document.add_heading("技术评分标准索引表", level=1)
            document.add_paragraph("投标机型为[投标机型，待填写]。")
            document.save(docx_path)

            rows = extract_gold_draft.extract_body_placeholders(docx_path, "PRJ-0003")
            extract_gold_draft.write_gold_csv(rows, output_path)

            csv_text = output_path.read_text(encoding="utf-8-sig")

        self.assertEqual(len(rows), 1)
        self.assertIn("投标机型", csv_text)
        self.assertIn("difficulty_tier", csv_text)

    def test_compute_metrics_scores_exact_predictions(self) -> None:
        gold_schema = load_eval_script("gold_schema")
        compute_metrics = load_eval_script("compute_metrics")

        gold = [
            gold_schema.GoldRow.from_csv_row(
                {
                    "project_id": "PRJ-0003",
                    "doc_type": "appendix",
                    "locator": "附表C/表1/行3",
                    "field_name": "投标机型",
                    "human_answer": "SE-15530",
                    "field_type": "phrase",
                    "difficulty_tier": "T2",
                    "evidence_source": "",
                }
            )
        ]
        predictions = [
            gold_schema.PredictionRow.from_csv_row(
                {
                    "project_id": "PRJ-0003",
                    "doc_type": "appendix",
                    "locator": "附表C/表1/行3",
                    "field_name": "投标机型",
                    "predicted_answer": "SE-15530",
                    "mark_yellow": "false",
                    "evidence_refs": "参数表",
                    "skill_name": "bid-tech-table-filler",
                    "fill_path": "kb_lookup",
                    "confidence": "0.9",
                }
            )
        ]

        metrics = compute_metrics.compute_metrics(gold, predictions)

        self.assertEqual(metrics["fillingRate"], 1.0)
        self.assertEqual(metrics["accuracy"], 1.0)
        self.assertEqual(metrics["evidenceChainRate"], 1.0)

    def test_failure_breakdown_detects_wrong_yellow(self) -> None:
        gold_schema = load_eval_script("gold_schema")
        failure_breakdown = load_eval_script("failure_breakdown")

        gold = [
            gold_schema.GoldRow.from_csv_row(
                {
                    "project_id": "PRJ-0003",
                    "doc_type": "body",
                    "locator": "章节/[投标机型]",
                    "field_name": "投标机型",
                    "human_answer": "SE-15530",
                    "field_type": "phrase",
                    "difficulty_tier": "T2",
                    "evidence_source": "",
                }
            )
        ]
        predictions = [
            gold_schema.PredictionRow.from_csv_row(
                {
                    "project_id": "PRJ-0003",
                    "doc_type": "body",
                    "locator": "章节/[投标机型]",
                    "field_name": "投标机型",
                    "predicted_answer": "",
                    "mark_yellow": "true",
                    "evidence_refs": "",
                    "skill_name": "bid-tech-word-placeholder-filler",
                    "fill_path": "mark_yellow",
                    "confidence": "",
                }
            )
        ]

        report = failure_breakdown.build_failure_breakdown(gold, predictions)

        self.assertEqual(report["byFailureType"], {"错标黄": 1})

    def test_run_eval_writes_all_outputs(self) -> None:
        run_eval = load_eval_script("run_eval")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            gold_path = root / "gold.csv"
            predictions_path = root / "predictions.csv"
            run_dir = root / "run"
            gold_path.write_text(
                "\n".join(
                    [
                        "project_id,doc_type,locator,field_name,human_answer,field_type,difficulty_tier,evidence_source",
                        "PRJ-0003,appendix,附表C/表1/行3,投标机型,SE-15530,phrase,T2,参数表",
                    ]
                ),
                encoding="utf-8",
            )
            predictions_path.write_text(
                "\n".join(
                    [
                        "project_id,doc_type,locator,field_name,predicted_answer,mark_yellow,evidence_refs,skill_name,fill_path,confidence",
                        "PRJ-0003,appendix,附表C/表1/行3,投标机型,SE-15530,false,参数表,bid-tech-table-filler,kb_lookup,0.9",
                    ]
                ),
                encoding="utf-8",
            )

            outputs = run_eval.run_eval(
                gold_paths=[gold_path],
                prediction_paths=[predictions_path],
                run_dir=run_dir,
            )

            self.assertTrue(outputs["metrics"].exists())
            self.assertTrue(outputs["failure_breakdown"].exists())
            self.assertTrue(outputs["evidence_audit"].exists())
            self.assertTrue(outputs["report"].exists())
