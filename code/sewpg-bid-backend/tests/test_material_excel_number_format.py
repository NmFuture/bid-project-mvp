from __future__ import annotations

import importlib.util
import sys
import tempfile
import unittest
from pathlib import Path


CONVERTER_PATH = (
    Path(__file__).resolve().parents[1]
    / "opencode"
    / "skills"
    / "bid-material-format-cleaner"
    / "scripts"
    / "excel_to_word.py"
)


def _load_converter():
    module_name = "test_bid_material_excel_to_word"
    spec = importlib.util.spec_from_file_location(module_name, CONVERTER_PATH)
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    return module


class ExcelNumberFormatTests(unittest.TestCase):
    def setUp(self) -> None:
        self.converter = _load_converter()

    def test_integer_format_rounds_underlying_float(self) -> None:
        # 基础弯矩表把荷载列设成整数显示，底层却是 12 位小数的浮点。
        self.assertEqual(self.converter._format_number(-149696.921875, "0"), "-149697")
        self.assertEqual(self.converter._format_number(162974.625, "0"), "162975")

    def test_fixed_decimal_format_follows_declared_precision(self) -> None:
        self.assertEqual(self.converter._format_number(1269.02, "0.0"), "1269.0")
        self.assertEqual(self.converter._format_number(-331.649, "0.0"), "-331.6")
        self.assertEqual(self.converter._format_number(145463, "0.0"), "145463.0")
        self.assertEqual(self.converter._format_number(12.3456, "0.00"), "12.35")

    def test_general_format_keeps_value_without_float_tail(self) -> None:
        self.assertEqual(self.converter._format_number(144611.0, "General"), "144611")
        self.assertEqual(self.converter._format_number(46791.3, "General"), "46791.3")
        self.assertEqual(self.converter._format_number(1, ""), "1")

    def test_percent_and_thousands_formats(self) -> None:
        self.assertEqual(self.converter._format_number(0.856, "0.0%"), "85.6%")
        self.assertEqual(self.converter._format_number(1234567.891, "#,##0"), "1,234,568")
        self.assertEqual(self.converter._format_number(1234567.891, "#,##0.00"), "1,234,567.89")

    def test_format_spec_strips_literals_colors_and_negative_section(self) -> None:
        # 字面量与颜色段不参与小数位判断，负数段不影响正数渲染
        self.assertEqual(self.converter._format_number(12.5, '0.00"吨"'), "12.50")
        self.assertEqual(self.converter._format_number(5.0, "0;[红色]0.00"), "5")


class ExcelToWordConversionTests(unittest.TestCase):
    def setUp(self) -> None:
        self.converter = _load_converter()

    def _write_workbook(self, path: Path) -> None:
        import openpyxl

        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet["B2"] = "载荷工况"
        sheet["C2"] = "Mx"
        sheet["D2"] = "占比"
        sheet["B3"] = "极端载荷工况"
        sheet["C3"] = -149696.921875
        sheet["C3"].number_format = "0"
        sheet["D3"] = 0.856
        sheet["D3"].number_format = "0.0%"
        workbook.save(path)

    def test_converted_table_uses_displayed_values(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "基础弯矩表.xlsx"
            output = Path(tmp) / "基础弯矩表.docx"
            self._write_workbook(source)

            self.converter.convert_excel_to_word(source, output)

            document = Document(str(output))

        self.assertEqual(len(document.tables), 1)
        rows = [[cell.text.strip() for cell in row.cells] for row in document.tables[0].rows]
        # 前导全空行/列被裁掉，表头行在前，数据按显示格式渲染
        self.assertEqual(rows[0], ["载荷工况", "Mx", "占比"])
        self.assertEqual(rows[1], ["极端载荷工况", "-149697", "85.6%"])


if __name__ == "__main__":
    unittest.main()
