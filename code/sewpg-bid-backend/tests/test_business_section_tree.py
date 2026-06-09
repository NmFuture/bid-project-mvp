from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.services.business_section_tree import build_business_section_tree
from app.services.parsing import extract_docx_text


def _node_by_title(payload: dict, title: str, *, after_line: int = 0) -> dict:
    return next(
        node
        for node in payload["nodes"]
        if node["title"] == title and int(node.get("startLine") or 0) > after_line
    )


def _line_number(lines: list[str], text: str, *, occurrence: int = 1) -> int:
    matches = [index for index, line in enumerate(lines, start=1) if text in line]
    return matches[occurrence - 1]


def _content_control_with_text(text: str):
    sdt = OxmlElement("w:sdt")
    sdt_content = OxmlElement("w:sdtContent")
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    paragraph.append(run)
    sdt_content.append(paragraph)
    sdt.append(sdt_content)
    return sdt


def _content_control_with_paragraphs(texts: list[str]):
    sdt = OxmlElement("w:sdt")
    sdt_content = OxmlElement("w:sdtContent")
    for text in texts:
        paragraph = OxmlElement("w:p")
        run = OxmlElement("w:r")
        text_node = OxmlElement("w:t")
        text_node.set(qn("xml:space"), "preserve")
        text_node.text = text
        run.append(text_node)
        paragraph.append(run)
        sdt_content.append(paragraph)
    sdt.append(sdt_content)
    return sdt


def test_business_section_tree_skips_cover_date_and_label_value_lines(tmp_path: Path) -> None:
    source_path = tmp_path / "cover-noise.docx"
    doc = Document()
    doc.add_paragraph("2026 年 4 月")
    doc.add_paragraph("招标项目所在地区：云南省,文山壮族苗族自治州")
    doc.add_paragraph("项目名称：示例风电项目")
    doc.add_paragraph("第一章 招标公告")
    doc.add_paragraph("3. 投标人资格要求")
    doc.add_paragraph("投标人须为中华人民共和国境内合法注册的独立法人。")
    doc.save(source_path)

    payload = build_business_section_tree([{"id": "DOC-1", "name": source_path.name, "sourcePath": str(source_path)}])

    titles = [node["title"] for node in payload["nodes"]]
    assert "2026 年 4 月" not in titles
    assert "招标项目所在地区：云南省,文山壮族苗族自治州" not in titles
    assert "项目名称：示例风电项目" not in titles
    assert "第一章 招标公告" in titles
    assert "3. 投标人资格要求" in titles


def test_business_section_tree_reads_detailed_toc_inside_word_containers(tmp_path: Path) -> None:
    source_path = tmp_path / "container-toc.docx"
    doc = Document()
    doc.add_paragraph("封面标题")
    toc_control = _content_control_with_paragraphs(
        [
            "目    录",
            "\t目    录\t1",
            "\t第一章  采购公告\t11",
            "\t1 采购条件 \t11",
            "\t2 项目概况与采购范围 \t11",
            "\t3 框架供应商资格要求 \t12",
            "\t第二章  框架供应商须知\t17",
            "\t框架供应商须知前附表 \t17",
            "\t1  总则 \t25",
            "\t1.1  采购项目概况 \t25",
            "\t1.4  框架供应商资格要求 \t25",
            "\t第六章    应答文件格式\t118",
        ]
    )
    doc._element.body.insert(len(doc._element.body) - 1, toc_control)
    doc.add_paragraph("第一章  采购公告")
    doc.add_paragraph("1 采购条件")
    doc.add_paragraph("采购条件正文。")
    doc.add_paragraph("2 项目概况与采购范围")
    doc.add_paragraph("项目概况正文。")
    doc.add_paragraph("3 框架供应商资格要求")
    doc.add_paragraph("资格要求正文。")
    doc.add_paragraph("第二章  框架供应商须知")
    doc.add_paragraph("框架供应商须知前附表")
    doc.add_paragraph("1  总则")
    doc.add_paragraph("1.1  采购项目概况")
    doc.add_paragraph("1.4  框架供应商资格要求")
    doc.add_paragraph("第六章    应答文件格式")
    doc.add_paragraph("格式正文。")
    doc.save(source_path)

    payload = build_business_section_tree([{"id": "DOC-1", "name": source_path.name, "sourcePath": str(source_path)}])
    lines = extract_docx_text(source_path).splitlines()
    titles = [node["title"] for node in payload["nodes"]]

    assert payload["summary"]["sourceModes"]["DOC-1"] == "detailed_toc"
    assert titles == [
        "第一章 采购公告",
        "1 采购条件",
        "2 项目概况与采购范围",
        "3 框架供应商资格要求",
        "第二章 框架供应商须知",
        "框架供应商须知前附表",
        "1 总则",
        "1.1 采购项目概况",
        "1.4 框架供应商资格要求",
        "第六章 应答文件格式",
    ]
    qualification_node = _node_by_title(payload, "3 框架供应商资格要求")
    body_qualification_lines = [
        index
        for index, line in enumerate(lines, start=1)
        if _line_number(lines, "第六章    应答文件格式") < index and "3 框架供应商资格要求" in line
    ]
    assert qualification_node["startLine"] == body_qualification_lines[0]
    assert qualification_node["endLine"] == _line_number(lines, "资格要求正文。")


def test_business_section_tree_locates_detailed_toc_entries_after_local_toc(tmp_path: Path) -> None:
    source_path = tmp_path / "detailed-toc-with-local-toc.docx"
    doc = Document()
    doc.add_paragraph("目录")
    doc.add_paragraph("第六章 投标文件格式\t20")
    doc.add_paragraph("商务投标文件\t21")
    doc.add_paragraph("一、投标函\t22")
    doc.add_paragraph("二、投标保证金\t23")
    doc.add_paragraph("技术投标文件\t24")
    doc.add_paragraph("一、技术偏差表\t25")

    doc.add_paragraph("第六章 投标文件格式")
    doc.add_paragraph("商务投标文件")
    doc.add_paragraph("目录")
    doc.add_paragraph("一、投标函")
    doc.add_paragraph("二、投标保证金")
    doc.add_paragraph("一、投标函")
    doc.add_paragraph("投标函正文。")
    doc.add_paragraph("二、投标保证金")
    doc.add_paragraph("保证金正文。")
    doc.add_paragraph("技术投标文件")
    doc.add_paragraph("一、技术偏差表")
    doc.add_paragraph("技术偏差正文。")
    doc.save(source_path)

    payload = build_business_section_tree([{"id": "DOC-1", "name": source_path.name, "sourcePath": str(source_path)}])
    lines = extract_docx_text(source_path).splitlines()

    assert payload["summary"]["sourceModes"]["DOC-1"] == "detailed_toc"
    format_node = _node_by_title(payload, "第六章 投标文件格式")
    assert format_node["endLine"] == _line_number(lines, "技术偏差正文。")

    bid_letter = _node_by_title(payload, "一、投标函")
    assert bid_letter["startLine"] == _line_number(lines, "一、投标函", occurrence=3)
    assert bid_letter["endLine"] == _line_number(lines, "投标函正文。")

    guarantee = _node_by_title(payload, "二、投标保证金")
    assert guarantee["startLine"] == _line_number(lines, "二、投标保证金", occurrence=3)
    assert guarantee["endLine"] == _line_number(lines, "保证金正文。")


def test_business_section_tree_lines_match_extracted_docx_text_with_tables_and_late_sections(tmp_path: Path) -> None:
    source_path = tmp_path / "line-coordinate.docx"
    doc = Document()
    doc.add_paragraph("第一章 招标公告", style="Heading 1")
    doc.add_paragraph("2. 项目概况与招标范围", style="Heading 2")
    doc.add_paragraph("项目简介：示例项目。")
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "序号"
    table.cell(0, 1).text = "项目"
    table.cell(0, 2).text = "内容"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "总装机容量"
    capacity_cell = table.cell(1, 2)
    capacity_cell.text = "336.6MW"
    capacity_cell.add_paragraph("可通过限制一台风机的输出功率，实现336.6MW。")
    capacity_cell.add_paragraph("90MW")
    capacity_cell.add_paragraph("可通过限制一台风机的输出功率，实现90MW。")

    doc.add_paragraph("3. 投标人资格要求", style="Heading 2")
    doc.add_paragraph("3.1 投标人须为境内合法注册法人。")
    qualification_with_break = doc.add_paragraph()
    qualification_with_break.add_run("3.2 投标人须为所投产品制造商；")
    qualification_with_break.add_run().add_break()
    qualification_with_break.add_run("并提供制造商声明。")
    doc.add_paragraph("3.3 投标人须提供类似业绩证明。")
    doc.add_paragraph("4. 招标文件的获取", style="Heading 2")
    doc.add_paragraph("4.1 招标文件售价：200元。")

    doc.add_paragraph("第五章 技术规范书", style="Heading 1")
    for index in range(10):
        doc.add_paragraph(f"技术规范正文第{index + 1}段。")

    doc.add_paragraph("第六章 投标文件格式", style="Heading 1")
    doc.add_paragraph("目录")
    doc.add_paragraph("六、资格证明文件")
    doc.add_paragraph("七、投标人其他证明文件")
    doc.add_paragraph("六、资格证明文件", style="Heading 2")
    doc.add_paragraph("（一）基本情况表", style="Heading 3")
    doc.add_paragraph("投标人基本情况表内容。")
    doc.add_paragraph("（二）投标人资格要求相关证明材料", style="Heading 3")
    doc.add_paragraph("1. 提供营业执照。")
    doc.add_paragraph("2. 提供制造商声明。")
    proof_with_break = doc.add_paragraph()
    proof_with_break.add_run("3. 提供业绩合同首页；")
    proof_with_break.add_run().add_break()
    proof_with_break.add_run("以及签字盖章页。")
    doc.add_paragraph("（三）其他资料", style="Heading 3")
    doc.add_paragraph("其他资料内容。")
    doc.add_paragraph("七、投标人其他证明文件", style="Heading 2")
    doc.add_paragraph("其他证明文件内容。")
    doc.save(source_path)

    payload = build_business_section_tree([{"id": "DOC-1", "name": source_path.name, "sourcePath": str(source_path)}])
    lines = extract_docx_text(source_path).splitlines()

    qualification_node = _node_by_title(payload, "3. 投标人资格要求")
    qualification_start = _line_number(lines, "3. 投标人资格要求")
    next_section_start = _line_number(lines, "4. 招标文件的获取")
    assert qualification_node["startLine"] == qualification_start
    assert qualification_node["contentStartLine"] == _line_number(lines, "3.1 投标人须为境内合法注册法人。")
    assert qualification_node["endLine"] == next_section_start - 1
    qualification_slice = lines[qualification_node["startLine"] - 1 : qualification_node["endLine"]]
    assert qualification_slice[0].strip() == "3. 投标人资格要求"
    assert "并提供制造商声明。" in qualification_slice
    assert not any("336.6MW" in line or "90MW" in line for line in qualification_slice)
    assert not any("4. 招标文件的获取" in line for line in qualification_slice)

    format_node = _node_by_title(payload, "第六章 投标文件格式")
    assert format_node["startLine"] == _line_number(lines, "第六章 投标文件格式")

    proof_parent = _node_by_title(payload, "六、资格证明文件", after_line=format_node["startLine"])
    format_toc_line = _line_number(lines[format_node["startLine"] :], "七、投标人其他证明文件") + format_node["startLine"]
    next_parent_start = _line_number(lines[proof_parent["startLine"] :], "七、投标人其他证明文件") + proof_parent["startLine"]
    assert proof_parent["startLine"] == _line_number(lines[format_toc_line:], "六、资格证明文件") + format_toc_line
    assert proof_parent["endLine"] == next_parent_start - 1
    proof_parent_slice = lines[proof_parent["startLine"] - 1 : proof_parent["endLine"]]
    assert any("（二）投标人资格要求相关证明材料" in line for line in proof_parent_slice)
    assert not any("七、投标人其他证明文件" in line for line in proof_parent_slice)

    proof_node = _node_by_title(payload, "（二）投标人资格要求相关证明材料", after_line=proof_parent["startLine"])
    next_proof_start = _line_number(lines[proof_node["startLine"] :], "（三）其他资料") + proof_node["startLine"]
    assert proof_node["startLine"] == _line_number(lines, "（二）投标人资格要求相关证明材料")
    assert proof_node["endLine"] == next_proof_start - 1
    proof_slice = lines[proof_node["startLine"] - 1 : proof_node["endLine"]]
    assert proof_slice[0].strip() == "（二）投标人资格要求相关证明材料"
    assert "以及签字盖章页。" in proof_slice
    assert not any("（三）其他资料" in line for line in proof_slice)


def test_business_section_tree_preserves_extract_docx_line_offsets_for_non_top_level_body_content(tmp_path: Path) -> None:
    source_path = tmp_path / "content-control-offset.docx"
    doc = Document()
    doc._element.body.insert(0, _content_control_with_text("内容控件提示行"))
    doc.add_paragraph("第一章 招标公告", style="Heading 1")
    doc.add_paragraph("投标人须知前附表", style="Heading 2")
    doc.add_paragraph("")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "条款号"
    table.cell(0, 1).text = "条款名称"
    table.cell(1, 0).text = "1.1"
    table.cell(1, 1).text = "项目名称"
    doc.add_paragraph("第二章 评标办法", style="Heading 1")
    doc.save(source_path)

    payload = build_business_section_tree([{"id": "DOC-1", "name": source_path.name, "sourcePath": str(source_path)}])
    lines = extract_docx_text(source_path).splitlines()

    chapter = _node_by_title(payload, "第一章 招标公告")
    assert chapter["startLine"] == _line_number(lines, "第一章 招标公告")

    preface = _node_by_title(payload, "投标人须知前附表")
    assert preface["startLine"] == _line_number(lines, "投标人须知前附表")
    assert preface["contentStartLine"] == _line_number(lines, "条款号")
    assert preface["endLine"] == _line_number(lines, "项目名称")


def test_business_section_tree_uses_detailed_toc_as_exclusive_title_source(tmp_path: Path) -> None:
    source_path = tmp_path / "detailed-toc-exclusive.docx"
    doc = Document()
    doc.add_paragraph("目录")
    doc.add_paragraph("第一章 采购公告\t1")
    doc.add_paragraph("第二章 框架供应商须知\t2")
    doc.add_paragraph("1. 总则\t3")
    doc.add_paragraph("1.1 项目概况\t4")
    doc.add_paragraph("1.4.2 框架供应商须知前附表规定接受联合体应答的，\t5")
    doc.add_paragraph("第六章 应答文件格式\t30")

    doc.add_paragraph("第一章 采购公告", style="Heading 1")
    doc.add_paragraph("采购公告正文。")
    doc.add_paragraph("第二章 框架供应商须知", style="Heading 1")
    doc.add_paragraph("1. 总则", style="Heading 2")
    doc.add_paragraph("1.1  项目概况", style="Heading 3")
    doc.add_paragraph("项目概况正文。")
    doc.add_paragraph("1.4.2 框架供应商须知前附表规定接受联合体应答的，", style="Heading 3")
    doc.add_paragraph("联合体应答正文。")
    doc.add_paragraph("3.5 正文里存在但目录没有的结构标题", style="Heading 2")
    doc.add_paragraph("这个标题不应进入章节树。")
    doc.add_paragraph("第六章 应答文件格式", style="Heading 1")
    doc.add_paragraph("格式正文。")
    doc.save(source_path)

    payload = build_business_section_tree([{"id": "DOC-1", "name": source_path.name, "sourcePath": str(source_path)}])
    lines = extract_docx_text(source_path).splitlines()
    titles = [node["title"] for node in payload["nodes"]]

    assert titles == [
        "第一章 采购公告",
        "第二章 框架供应商须知",
        "1. 总则",
        "1.1 项目概况",
        "1.4.2 框架供应商须知前附表规定接受联合体应答的，",
        "第六章 应答文件格式",
    ]
    assert all(node["source"] == "toc" for node in payload["nodes"])
    assert "3.5 正文里存在但目录没有的结构标题" not in titles

    project_node = _node_by_title(payload, "1.1 项目概况")
    assert project_node["startLine"] == _line_number(lines, "1.1  项目概况")
    assert project_node["endLine"] == _line_number(lines, "项目概况正文。")


def test_business_section_tree_uses_body_structure_without_regex_when_toc_is_not_detailed(tmp_path: Path) -> None:
    source_path = tmp_path / "body-structure-exclusive.docx"
    doc = Document()
    doc.add_paragraph("目录")
    doc.add_paragraph("第一章 招标公告\t3")
    doc.add_paragraph("第二章 投标人须知\t8")
    doc.add_paragraph("第三章 评标办法\t30")

    doc.add_paragraph("招标项目所在地区：云南省,文山壮族苗族自治州", style="Heading 2")
    doc.add_paragraph("第一章 招标公告", style="Heading 1")
    doc.add_paragraph("招标公告正文。")
    doc.add_paragraph("第二章 投标人须知", style="Heading 1")
    doc.add_paragraph("投标人须知前附表", style="Heading 2")
    doc.add_paragraph("1.4.2投标人须知前附表规定接受联合体投标的，联合体除应符合本章第1.4.1项和投标人须知前附表的要求外，还应遵守以下规定：")
    doc.add_paragraph("5.2.1主持人按下列程序进行开标：")
    doc.add_paragraph("开标正文。")
    doc.add_paragraph("第三章 评标办法", style="Heading 1")
    doc.add_paragraph("评标办法正文。")
    doc.save(source_path)

    payload = build_business_section_tree([{"id": "DOC-1", "name": source_path.name, "sourcePath": str(source_path)}])
    titles = [node["title"] for node in payload["nodes"]]

    assert titles == [
        "第一章 招标公告",
        "第二章 投标人须知",
        "投标人须知前附表",
        "第三章 评标办法",
    ]
    assert "1.4.2投标人须知前附表规定接受联合体投标的，联合体除应符合本章第1.4.1项和投标人须知前附表的要求外，还应遵守以下规定：" not in titles
    assert "5.2.1主持人按下列程序进行开标：" not in titles
    assert all(node["source"] in {"heading-style", "style-outline", "outline"} for node in payload["nodes"])


def test_business_section_tree_merges_missing_top_level_regex_heading_at_body_anchor_only(tmp_path: Path) -> None:
    source_path = tmp_path / "body-structure-regex-top-level-anchor.docx"
    doc = Document()
    doc.add_paragraph("第一章 招标公告")
    doc.add_paragraph("封面提示。")
    doc.add_paragraph("第一章 招标公告")
    doc.add_paragraph("招标公告正文。")
    doc.add_paragraph("第二章 投标人须知", style="Heading 1")
    doc.add_paragraph("投标人须知前附表", style="Heading 2")
    doc.add_paragraph("1. 总则", style="Heading 2")
    doc.add_paragraph("1.1 项目概况", style="Heading 3")
    doc.add_paragraph("项目概况正文。")
    doc.add_paragraph("第三章 评标办法", style="Heading 1")
    doc.add_paragraph("评标办法正文。")
    doc.save(source_path)

    payload = build_business_section_tree([{"id": "DOC-1", "name": source_path.name, "sourcePath": str(source_path)}])
    lines = extract_docx_text(source_path).splitlines()
    titles = [node["title"] for node in payload["nodes"]]

    assert titles.count("第一章 招标公告") == 1
    assert payload["summary"]["sourceModes"]["DOC-1"] == "body_structure"
    announcement = _node_by_title(payload, "第一章 招标公告")
    assert announcement["startLine"] == _line_number(lines[1:], "第一章 招标公告") + 1
    assert announcement["endLine"] == _line_number(lines, "招标公告正文。")


def test_business_section_tree_regex_fallback_is_conservative_and_limited_to_three_levels(tmp_path: Path) -> None:
    source_path = tmp_path / "regex-fallback-conservative.docx"
    doc = Document()
    doc.add_paragraph("第一章 招标公告")
    doc.add_paragraph("1. 项目概况")
    doc.add_paragraph("1.1 招标范围")
    doc.add_paragraph("1.1.1 供货范围")
    doc.add_paragraph("1.1.1.1 不应进入第四级")
    doc.add_paragraph("1.4.2投标人须知前附表规定接受联合体投标的，联合体除应符合本章第1.4.1项和投标人须知前附表的要求外，还应遵守以下规定：")
    doc.add_paragraph("5.2.1主持人按下列程序进行开标：")
    doc.add_paragraph("第二章 投标人须知")
    doc.save(source_path)

    payload = build_business_section_tree([{"id": "DOC-1", "name": source_path.name, "sourcePath": str(source_path)}])
    titles = [node["title"] for node in payload["nodes"]]

    assert "第一章 招标公告" in titles
    assert "1. 项目概况" in titles
    assert "1.1 招标范围" in titles
    assert "1.1.1 供货范围" in titles
    assert "第二章 投标人须知" in titles
    assert "1.1.1.1 不应进入第四级" not in titles
    assert "1.4.2投标人须知前附表规定接受联合体投标的，联合体除应符合本章第1.4.1项和投标人须知前附表的要求外，还应遵守以下规定：" not in titles
    assert "5.2.1主持人按下列程序进行开标：" not in titles
    assert max(node["level"] for node in payload["nodes"]) <= 3


def test_business_section_tree_ignores_mid_document_local_toc_as_global_title_source(tmp_path: Path) -> None:
    source_path = tmp_path / "mid-document-local-toc.docx"
    doc = Document()
    doc.add_paragraph("第一章 招标公告", style="Heading 1")
    doc.add_paragraph("招标公告正文。")
    doc.add_paragraph("第二章 合同条款", style="Heading 1")
    doc.add_paragraph("目录")
    doc.add_paragraph("第一部分 合同协议书\t5")
    doc.add_paragraph("第二部分 通用合同条款\t8")
    doc.add_paragraph("1. 一般约定\t8")
    doc.add_paragraph("1.1 词语定义\t8")
    doc.add_paragraph("1.2 语言文字\t9")
    doc.add_paragraph("1.3 合同生效\t10")
    doc.add_paragraph("1.4 联络\t11")
    doc.add_paragraph("2. 合同范围\t12")
    doc.add_paragraph("2.1 供货范围\t13")
    doc.add_paragraph("2.2 服务范围\t14")
    doc.add_paragraph("第一部分 合同协议书", style="Heading 2")
    doc.add_paragraph("合同协议正文。")
    doc.add_paragraph("第二部分 通用合同条款", style="Heading 2")
    doc.add_paragraph("通用合同正文。")
    doc.save(source_path)

    payload = build_business_section_tree([{"id": "DOC-1", "name": source_path.name, "sourcePath": str(source_path)}])

    titles = [node["title"] for node in payload["nodes"]]
    assert titles[:2] == ["第一章 招标公告", "第二章 合同条款"]
    assert payload["summary"]["sourceModes"]["DOC-1"] == "body_structure"
    assert all(node["source"] != "toc" for node in payload["nodes"])
