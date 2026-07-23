from __future__ import annotations

import importlib.util
import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


BACKEND_ROOT = Path(__file__).resolve().parents[1]
ASSEMBLER_SCRIPTS = (
    BACKEND_ROOT / "opencode" / "skills" / "bid-tech-assembler" / "scripts"
)


def load_assembler_script(name: str):
    sys.path.insert(0, str(ASSEMBLER_SCRIPTS))
    try:
        spec = importlib.util.spec_from_file_location(
            f"test_technical_final_assembly_{name}",
            ASSEMBLER_SCRIPTS / f"{name}.py",
        )
        if spec is None or spec.loader is None:
            raise RuntimeError(f"无法加载 assembler 脚本: {name}")
        module = importlib.util.module_from_spec(spec)
        previous = sys.modules.get(spec.name)
        sys.modules[spec.name] = module
        try:
            spec.loader.exec_module(module)
        finally:
            if previous is None:
                sys.modules.pop(spec.name, None)
            else:
                sys.modules[spec.name] = previous
        return module
    finally:
        sys.path.pop(0)


class TechnicalFinalAssemblyTests(unittest.TestCase):
    def test_outline_title_keeps_business_number_without_separator(self) -> None:
        from app.services.technical_gap_domain import technical_outline_number_and_title

        self.assertEqual(
            technical_outline_number_and_title(
                {"tocNumber": "1", "title": "1号机组技术参数"},
                "3",
            ),
            ("1", "1号机组技术参数"),
        )

    def test_prune_unused_styles_keeps_references_dependencies_and_pipeline_styles(self) -> None:
        pruner = load_assembler_script("docx_style_pruner")
        doc = Document()
        base_style = doc.styles.add_style("Referenced Base", WD_STYLE_TYPE.PARAGRAPH)
        used_style = doc.styles.add_style("Referenced Child", WD_STYLE_TYPE.PARAGRAPH)
        used_style.base_style = base_style
        doc.add_paragraph("正文", style=used_style)
        for index in range(200):
            doc.styles.add_style(f"Unused Style {index}", WD_STYLE_TYPE.PARAGRAPH)

        before = len(doc.styles)
        result = pruner.prune_unused_styles(doc)
        remaining_by_id = {
            style.get(qn("w:styleId")): style
            for style in doc.styles.element.findall(qn("w:style"))
        }
        remaining_names = {
            name.get(qn("w:val"))
            for style in remaining_by_id.values()
            if (name := style.find(qn("w:name"))) is not None
        }

        self.assertGreaterEqual(result["removed"], 200)
        self.assertLess(len(doc.styles), before - 190)
        self.assertIn(used_style.style_id, remaining_by_id)
        self.assertIn(base_style.style_id, remaining_by_id)
        self.assertIn("Normal", remaining_names)
        self.assertIn("heading 6", {name.lower() for name in remaining_names})

    def test_preprocess_removes_unused_style_baggage_from_saved_document(self) -> None:
        preprocess = load_assembler_script("preprocess")
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source_path = root / "source.docx"
            output_path = root / "output.docx"
            source = Document()
            source.add_paragraph("正文")
            for index in range(200):
                source.styles.add_style(f"Unused Style {index}", WD_STYLE_TYPE.PARAGRAPH)
            source.save(source_path)

            stats = preprocess.preprocess(source_path, output_path)
            output = Document(output_path)

        self.assertGreaterEqual(stats["styles_pruned"], 200)
        self.assertLess(len(output.styles), 30)

    def test_batch_composer_defers_global_renumbering_until_finalize(self) -> None:
        merger = load_assembler_script("merger")
        master = Document()
        first = Document()
        first.add_paragraph("first")
        second = Document()
        second.add_paragraph("second")

        with (
            patch.object(merger.Composer, "renumber_bookmarks", autospec=True) as bookmarks,
            patch.object(merger.Composer, "renumber_docpr_ids", autospec=True) as docpr_ids,
            patch.object(merger.Composer, "renumber_nvpicpr_ids", autospec=True) as nvpicpr_ids,
        ):
            composer = merger.BatchComposer(master)
            composer.append(first)
            composer.append(second)

            self.assertEqual(bookmarks.call_count, 0)
            self.assertEqual(docpr_ids.call_count, 0)
            self.assertEqual(nvpicpr_ids.call_count, 0)

            composer.finalize_global_ids()

        self.assertEqual(bookmarks.call_count, 1)
        self.assertEqual(docpr_ids.call_count, 1)
        self.assertEqual(nvpicpr_ids.call_count, 1)

    def test_prepare_toc_json_preserves_confirmed_outline_numbers_and_stable_fallbacks(self) -> None:
        from app.services import tech_assembly

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            opencode_toc = root / "opencode-toc.json"
            opencode_toc.write_text(
                json.dumps(
                    {
                        "items": [
                            {"level": 1, "number": "9", "title": "旧第一章"},
                            {"level": 2, "number": "9.8", "title": "旧子节"},
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            work_dir = root / "assembly"
            work_dir.mkdir()
            outline_state = {
                "nodes": [
                    {
                        "id": "OL-1",
                        "title": "第一章 技术方案",
                        "tocNumber": "第一章",
                        "children": [
                            {
                                "id": "OL-1-1",
                                "title": "总体设计",
                                "tocNumber": "   ",
                                "toc_number": "1.1",
                                "children": [],
                            },
                            {"id": "OL-1-2", "title": "设备选型", "number": "1.2", "children": []},
                            {"id": "OL-1-3", "title": "缺号子项", "children": []},
                        ],
                    },
                    {
                        "id": "OL-2",
                        "title": "缺号章节",
                        "children": [
                            {"id": "OL-2-1", "title": "缺号孙项", "children": []},
                        ],
                    },
                ]
            }
            project = {
                "id": "PRJ-TEST",
                "name": "测试项目",
                "directory_state": {"opencodeOutput": {"tocJsonPath": str(opencode_toc)}},
            }

            output_path = tech_assembly._prepare_toc_json(
                "PRJ-TEST",
                project,
                outline_state,
                {},
                work_dir,
            )
            output = json.loads(output_path.read_text(encoding="utf-8"))

            self.assertEqual(
                [(item["number"], item["title"]) for item in output["items"]],
                [
                    ("第一章", "技术方案"),
                    ("1.1", "总体设计"),
                    ("1.2", "设备选型"),
                    ("1.3", "缺号子项"),
                    ("2", "缺号章节"),
                    ("2.1", "缺号孙项"),
                ],
            )

    def test_chinese_chapter_number_keeps_selected_materials_through_gap_plan(self) -> None:
        from app.services import tech_assembly, technical_gap_planner

        parse_toc = load_assembler_script("parse_toc")
        with patch.dict(sys.modules, {"yaml": object()}):
            build_assembly = load_assembler_script("build_assembly")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            work_dir = root / "assembly"
            work_dir.mkdir()
            outline_nodes = [
                {
                    "id": "OL-1",
                    "title": "第一章技术方案",
                    "tocNumber": "第一章",
                    "children": [
                        {
                            "id": "OL-1-1",
                            "title": "一、总体设计",
                            "tocNumber": "一、",
                            "children": [
                                {
                                    "id": "OL-1-1-1",
                                    "title": "（一）设备选型",
                                    "tocNumber": "（一）",
                                    "children": [],
                                }
                            ],
                        },
                        {
                            "id": "OL-1-2",
                            "title": "第一节实施方案",
                            "tocNumber": "第一节",
                            "children": [],
                        },
                    ],
                }
            ]
            s4_toc_items = technical_gap_planner._outline_nodes_to_toc_items(outline_nodes)
            toc_path = tech_assembly._prepare_toc_json(
                "PRJ-CHINESE-NUMBER",
                {"id": "PRJ-CHINESE-NUMBER", "name": "中文章节号测试"},
                {"nodes": outline_nodes},
                {},
                work_dir,
            )
            s7_toc_items = json.loads(toc_path.read_text(encoding="utf-8"))["items"]
            self.assertEqual(
                [item["number"] for item in s4_toc_items],
                ["第一章", "一、", "（一）", "第一节"],
            )
            self.assertEqual(
                [item["title"] for item in s4_toc_items],
                ["技术方案", "总体设计", "设备选型", "实施方案"],
            )
            self.assertEqual(
                [(item["number"], item["title"]) for item in s7_toc_items],
                [(item["number"], item["title"]) for item in s4_toc_items],
            )
            self.assertEqual(
                [item["chapter_no_flat"] for item in s7_toc_items],
                ["1", "1.1", "1.1.1", "1.2"],
            )
            gap_plan_path = root / "gap-plan.json"
            gap_plan_path.write_text(
                json.dumps(
                    {
                        "items": [
                            {
                                **item,
                                "id": f"GAP-{index}",
                                "matchedMaterials": [{"path": f"material-{index}.docx"}],
                            }
                            for index, item in enumerate(s4_toc_items, start=1)
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            toc_entries = parse_toc.parse_toc_json(toc_path)
            result = build_assembly.apply_gap_plan(
                build_assembly.build_plan(toc_entries, [], {}),
                gap_plan_path,
            )
            chapter_master_path = root / "chapter-master-gap-plan.json"
            chapter_master_path.write_text(
                json.dumps(
                    {
                        "items": [
                            {
                                **item,
                                **({"coverageRole": "chapter_master"} if item["number"] == "一、" else {}),
                                "id": f"MASTER-{index}",
                                "matchedMaterials": [{"path": f"material-{index}.docx"}],
                            }
                            for index, item in enumerate(s4_toc_items, start=1)
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            chapter_master_result = build_assembly.apply_gap_plan(
                build_assembly.build_plan(toc_entries, [], {}),
                chapter_master_path,
            )

        self.assertEqual(
            [item["chapter_no"] for item in result],
            ["第一章", "一、", "（一）", "第一节"],
        )
        self.assertEqual(
            [item["title"] for item in result],
            ["技术方案", "总体设计", "设备选型", "实施方案"],
        )
        self.assertEqual(
            [item["chapter_no_flat"] for item in result],
            ["1", "1.1", "1.1.1", "1.2"],
        )
        self.assertEqual(
            [item["paths"] for item in result],
            [[f"material-{index}.docx"] for index in range(1, 5)],
        )
        self.assertEqual([item["status"] for item in result], ["MATCHED"] * 4)
        self.assertEqual(
            [item["chapter_no"] for item in chapter_master_result],
            ["第一章", "一、", "第一节"],
        )

    def test_chapter_master_replaces_descendants_and_numbers_material_headings(self) -> None:
        merger = load_assembler_script("merger")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            library = root / "library"
            library.mkdir()
            source_path = library / "source-1.7.docx"
            source = Document()
            source.add_paragraph("投标方案优势说明", style="Heading 2")
            source.add_paragraph("投标机组技术路线优势", style="Heading 5")
            source.add_paragraph("正文一")
            source.add_paragraph("低风速全碳叶片机组优势", style="Heading 5")
            source.add_paragraph("碳纤维叶片介绍", style="Heading 5")
            source.save(source_path)

            template_path = root / "template.docx"
            Document().save(template_path)
            output_path = root / "assembled.docx"
            plan = [
                {
                    "status": "MATCHED",
                    "level": 2,
                    "chapter_no": "1.7",
                    "chapter_no_flat": "1.7",
                    "title": "投标方案优势说明",
                    "paths": [source_path.name],
                    "is_preface": False,
                    "is_appendix": False,
                    "coverage_role": "chapter_master",
                },
                {
                    "status": "NEEDS_REVIEW",
                    "level": 3,
                    "chapter_no": "1.7.1",
                    "chapter_no_flat": "1.7.1",
                    "title": "投标方案整体优势",
                    "paths": [],
                    "is_preface": False,
                    "is_appendix": False,
                },
                {
                    "status": "NEEDS_REVIEW",
                    "level": 3,
                    "chapter_no": "1.7.2",
                    "chapter_no_flat": "1.7.2",
                    "title": "技术路线优势",
                    "paths": [],
                    "is_preface": False,
                    "is_appendix": False,
                },
            ]

            stats = merger.merge(template_path, plan, library, {}, root / "prep", output_path)
            assembled = Document(output_path)
            headings = [
                paragraph.text.strip().replace("  ", " ")
                for paragraph in assembled.paragraphs
                if (paragraph.style.name or "").startswith("Heading")
            ]

            self.assertEqual(
                headings,
                [
                    "1.7 投标方案优势说明",
                    "1.7.1 投标机组技术路线优势",
                    "1.7.2 低风速全碳叶片机组优势",
                    "1.7.3 碳纤维叶片介绍",
                ],
            )
            self.assertEqual(stats["superseded"], 2)

    def test_apply_gap_plan_carries_existing_chapter_coverage_into_assembly_plan(self) -> None:
        with patch.dict(sys.modules, {"yaml": object()}):
            build_assembly = load_assembler_script("build_assembly")

        with tempfile.TemporaryDirectory() as tmp:
            gap_plan_path = Path(tmp) / "gap-plan.json"
            gap_plan_path.write_text(
                json.dumps(
                    {
                        "items": [
                            {
                                "id": "GAP-PARENT",
                                "number": "1.7",
                                "title": "投标方案优势说明",
                                "coverageRole": "chapter_master",
                                "matchedMaterials": [{"path": "source-1.7.docx"}],
                            },
                            {
                                "id": "GAP-CHILD",
                                "number": "1.7.1",
                                "title": "投标方案整体优势",
                                "coverageRole": "covered_by_parent",
                                "coveredByParent": "GAP-PARENT",
                            },
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            plan = [
                {
                    "chapter_no": "1.7",
                    "chapter_no_flat": "1.7",
                    "title": "投标方案优势说明",
                    "status": "UNMATCHED",
                    "paths": ["wiki-guessed-parent.docx"],
                },
                {
                    "chapter_no": "1.7.1",
                    "chapter_no_flat": "1.7.1",
                    "title": "投标方案整体优势",
                    "status": "NEEDS_REVIEW",
                    "paths": ["wiki-guessed-child.docx"],
                },
            ]

            result = build_assembly.apply_gap_plan(plan, gap_plan_path)

            self.assertEqual(len(result), 1)
            self.assertEqual(result[0]["coverage_role"], "chapter_master")
            self.assertEqual(result[0]["paths"], ["source-1.7.docx"])

    def test_apply_gap_plan_clears_wiki_materials_without_selected_match(self) -> None:
        with patch.dict(sys.modules, {"yaml": object()}):
            build_assembly = load_assembler_script("build_assembly")

        with tempfile.TemporaryDirectory() as tmp:
            gap_plan_path = Path(tmp) / "gap-plan.json"
            gap_plan_path.write_text(
                json.dumps(
                    {
                        "items": [
                            {
                                "id": "GAP-6-6",
                                "number": "6.6",
                                "title": "技术附表",
                                "coverageRole": "covered_by_parent",
                                "coveredByParent": "GAP-6",
                                "matchedMaterials": [],
                            }
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            plan = [
                {
                    "chapter_no": "6.6",
                    "chapter_no_flat": "6.6",
                    "title": "技术附表",
                    "status": "MATCHED",
                    "paths": ["appendix-b.docx", "appendix-c.docx", "appendix-i.docx"],
                    "shifts": [0, 0, 0],
                    "attach_modes": ["normal", "normal", "normal"],
                }
            ]

            result = build_assembly.apply_gap_plan(plan, gap_plan_path)

        self.assertEqual(result[0]["paths"], [])
        self.assertEqual(result[0]["coverage_role"], "covered_by_parent")

    def test_gap_plan_paths_uses_selected_material_when_fill_has_no_artifact(self) -> None:
        with patch.dict(sys.modules, {"yaml": object()}):
            build_assembly = load_assembler_script("build_assembly")

        item = {
            "fillTasks": [{"id": "FILL-1", "status": "pending"}],
            "matchedMaterials": [{"path": "selected-material.docx"}],
            "resolvedArtifacts": [],
        }

        self.assertEqual(build_assembly._gap_plan_paths(item), ["selected-material.docx"])

    def test_gap_plan_paths_does_not_fallback_from_unreviewed_ai_artifact(self) -> None:
        with patch.dict(sys.modules, {"yaml": object()}):
            build_assembly = load_assembler_script("build_assembly")

        item = {
            "matchedMaterials": [{"path": "selected-template.docx"}],
            "resolvedArtifacts": [
                {
                    "source": "ai_fill",
                    "path": "unreviewed-ai.docx",
                    "s7Ready": False,
                    "qualityReport": {"status": "needs_review"},
                }
            ],
        }

        self.assertEqual(build_assembly._gap_plan_paths(item), [])

    def test_init_params_accepts_unified_turbine_fields_and_unknown_extensions(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_path = root / "toc.json"
            params_path = root / "project_params.json"
            toc_path.write_text("{}", encoding="utf-8")
            params_path.write_text(
                json.dumps(
                    {
                        "project_name": "测试项目",
                        "project_short": "测试项目",
                        "client_name": "测试业主",
                        "tender_no": "TEST-001",
                        "turbine_model": "EW6.25-220",
                        "turbine_platform": "",
                        "rated_power_kw": "",
                        "rotor_diameter_m": "",
                        "turbine_layout": "",
                        "future_extension": "",
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            completed = subprocess.run(
                [
                    sys.executable,
                    str(ASSEMBLER_SCRIPTS / "init_params.py"),
                    "--toc",
                    str(toc_path),
                    "--out",
                    str(params_path),
                ],
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
            )

            self.assertEqual(completed.returncode, 0, completed.stderr)
            params = json.loads(params_path.read_text(encoding="utf-8"))
            self.assertEqual(params["turbine_model"], "EW6.25-220")
            self.assertIn("待填写", params["turbine_platform"])
            self.assertEqual(params["future_extension"], "")

    def test_preprocess_replaces_turbine_placeholders_with_unified_field_names(self) -> None:
        preprocess = load_assembler_script("preprocess")
        doc = Document()
        doc.add_paragraph("机型：[机型号]；额定功率：[额定功率]")

        replaced = preprocess.replace_placeholders(
            doc,
            {"turbine_model": "EW6.25-220", "rated_power_kw": "6250"},
        )

        self.assertEqual(replaced, 1)
        self.assertEqual(doc.paragraphs[0].text, "机型：EW6.25-220；额定功率：6250")

    def test_runner_returns_structured_contract_without_markdown_reports(self) -> None:
        runner = load_assembler_script("run_from_manifest")
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_file = root / "toc.json"
            gap_plan_file = root / "gap-plan.json"
            wiki_dir = root / "wiki"
            library_dir = root / "library"
            template_file = root / "template.docx"
            output_file = root / "output.docx"
            manifest_file = root / "manifest.json"
            wiki_dir.mkdir()
            library_dir.mkdir()
            toc_file.write_text("{}", encoding="utf-8")
            gap_plan_file.write_text('{"items": []}', encoding="utf-8")
            Document().save(template_file)
            manifest_file.write_text(
                json.dumps(
                    {
                        "workDir": str(root),
                        "tocJsonPath": str(toc_file),
                        "gapPlanPath": str(gap_plan_file),
                        "wikiDir": str(wiki_dir),
                        "materialLibraryDir": str(library_dir),
                        "templateFile": str(template_file),
                        "outputFile": str(output_file),
                    }
                ),
                encoding="utf-8",
            )

            def fake_run_capture(command: list[str]) -> str:
                script = Path(command[1]).name
                if script == "parse_toc.py":
                    Path(command[command.index("--out") + 1]).write_text("[]", encoding="utf-8")
                elif script == "init_params.py":
                    Path(command[command.index("--out") + 1]).write_text("{}", encoding="utf-8")
                elif script == "build_assembly.py":
                    self.assertEqual(Path(command[command.index("--gap-plan") + 1]), gap_plan_file)
                    Path(command[command.index("--out") + 1]).write_text(
                        json.dumps([{"status": "STRUCTURAL", "level": 1, "title": "技术方案", "paths": []}]),
                        encoding="utf-8",
                    )
                elif script == "merger.py":
                    Document().save(Path(command[command.index("--out") + 1]))
                    Path(command[command.index("--result") + 1]).write_text(
                        json.dumps({"merged_materials": 0, "warnings": []}), encoding="utf-8"
                    )
                elif script == "finalize.py":
                    Document().save(Path(command[command.index("--out") + 1]))
                elif script == "verify.py":
                    self.assertNotIn("--report", command)
                    self.assertNotIn("--review", command)
                    Path(command[command.index("--result") + 1]).write_text(
                        json.dumps(
                            {
                                "placeholders": [],
                                "empty_leaf_headings": [],
                                "dup_alerts": [],
                                "ghost_chapters": [],
                                "invalid_h1": [],
                                "invalid_prefix": [],
                            }
                        ),
                        encoding="utf-8",
                    )
                return ""

            with patch.object(runner, "run_capture", side_effect=fake_run_capture):
                result = runner.run_from_manifest(manifest_file)

            self.assertTrue(output_file.exists())
            self.assertTrue(Path(result["planFile"]).exists())
            self.assertEqual(result["assemblyReport"], "")
            self.assertEqual(result["needsReview"], "")
            self.assertIsInstance(result["summary"], dict)
            self.assertIsInstance(result["warnings"], list)
            self.assertFalse((root / "assembly_report.md").exists())
            self.assertFalse((root / "needs_review.md").exists())

    def test_runner_rejects_manifest_without_gap_plan(self) -> None:
        runner = load_assembler_script("run_from_manifest")
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_file = root / "toc.json"
            wiki_dir = root / "wiki"
            library_dir = root / "library"
            toc_file.write_text("{}", encoding="utf-8")
            wiki_dir.mkdir()
            library_dir.mkdir()
            manifest_file = root / "manifest.json"
            manifest_file.write_text(
                json.dumps(
                    {
                        "workDir": str(root),
                        "tocJsonPath": str(toc_file),
                        "wikiDir": str(wiki_dir),
                        "materialLibraryDir": str(library_dir),
                    }
                ),
                encoding="utf-8",
            )

            with self.assertRaisesRegex(RuntimeError, "gapPlanPath"):
                runner.run_from_manifest(manifest_file)

    def test_verify_returns_compact_json_without_markdown_reports(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            docx_path = root / "assembled.docx"
            plan_path = root / "assembly_plan.json"
            params_path = root / "project_params.json"
            result_path = root / "assembly_verify_result.json"

            doc = Document()
            doc.add_paragraph("技术方案", style="Heading 1")
            doc.add_paragraph("脱敏正文。")
            doc.save(docx_path)
            plan_path.write_text("[]", encoding="utf-8")
            params_path.write_text("{}", encoding="utf-8")

            completed = subprocess.run(
                [
                    sys.executable,
                    str(ASSEMBLER_SCRIPTS / "verify.py"),
                    "--docx",
                    str(docx_path),
                    "--plan",
                    str(plan_path),
                    "--params",
                    str(params_path),
                    "--result",
                    str(result_path),
                ],
                check=True,
                capture_output=True,
                text=True,
                encoding="utf-8",
            )
            payload = json.loads(completed.stdout)

            self.assertEqual(payload, json.loads(result_path.read_text(encoding="utf-8")))
            self.assertIn("heading_counts", payload)
            self.assertFalse((root / "assembly_report.md").exists())
            self.assertFalse((root / "needs_review.md").exists())

    def test_merger_deduplicates_heading_from_first_usable_material(self) -> None:
        merger = load_assembler_script("merger")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            library = root / "library"
            corrupt = library / "corrupt.docx"
            valid = library / "valid.docx"
            output = root / "assembled.docx"

            library.mkdir()
            Document().save(template)
            corrupt.write_bytes(b"not-a-docx")
            valid_doc = Document()
            valid_doc.add_paragraph("正常素材", style="Heading 2")
            valid_doc.add_paragraph("这是首份可用素材的正文。")
            valid_doc.save(valid)
            plan = [
                {
                    "status": "MATCHED",
                    "level": 2,
                    "title": "正常素材",
                    "chapter_no": "1.1",
                    "chapter_no_flat": "1.1",
                    "paths": ["missing.docx", corrupt.name, valid.name],
                }
            ]

            stats = merger.merge(template, plan, library, {}, root / "prep", output)
            result = Document(str(output))
            matching_paragraphs = [
                paragraph.text
                for paragraph in result.paragraphs
                if "正常素材" in paragraph.text
            ]

        self.assertEqual(matching_paragraphs, ["1.1  正常素材"])
        self.assertEqual(stats["merged_materials"], 1)

    def test_failed_compose_does_not_suppress_following_unmatched_child(self) -> None:
        merger = load_assembler_script("merger")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            library = root / "library"
            source = library / "parent.docx"
            output = root / "assembled.docx"

            library.mkdir()
            Document().save(template)
            source_doc = Document()
            source_doc.add_paragraph("父章节", style="Heading 1")
            source_doc.add_paragraph("待补子节", style="Heading 2")
            source_doc.add_paragraph("尚未成功合并的正文。")
            source_doc.save(source)
            plan = [
                {
                    "status": "MATCHED",
                    "level": 1,
                    "title": "父章节",
                    "chapter_no": "第一章",
                    "chapter_no_flat": "1",
                    "paths": [source.name],
                },
                {
                    "status": "UNMATCHED",
                    "level": 2,
                    "title": "待补子节",
                    "chapter_no": "1.1",
                    "chapter_no_flat": "1.1",
                    "paths": [],
                },
            ]

            with patch.object(merger.Composer, "append", side_effect=RuntimeError("compose failed")):
                stats = merger.merge(template, plan, library, {}, root / "prep", output)
            result = Document(str(output))
            text = "\n".join(paragraph.text for paragraph in result.paragraphs)

        self.assertIn("1.1  待补子节", text)
        self.assertIn("[缺失：待补子节——wiki 无匹配卡片，请人工处理]", text)
        self.assertEqual(stats.get("superseded", 0), 0)

    def test_merger_keeps_going_when_materials_are_missing_or_corrupt(self) -> None:
        merger = load_assembler_script("merger")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            library = root / "library"
            valid = library / "valid.docx"
            corrupt = library / "corrupt.docx"
            output = root / "assembled.docx"

            library.mkdir()
            Document().save(template)
            valid_doc = Document()
            valid_doc.add_paragraph("正常素材", style="Heading 2")
            valid_doc.add_paragraph("这是可用的脱敏正文。")
            valid_doc.save(valid)
            corrupt.write_bytes(b"not-a-docx")

            plan = [
                {
                    "status": "MATCHED",
                    "level": 2,
                    "title": "正常素材",
                    "chapter_no": "1.1",
                    "chapter_no_flat": "1.1",
                    "paths": [valid.name],
                },
                {
                    "status": "MATCHED",
                    "level": 2,
                    "title": "缺失素材",
                    "chapter_no": "1.2",
                    "chapter_no_flat": "1.2",
                    "paths": ["missing.docx"],
                },
                {
                    "status": "MATCHED",
                    "level": 2,
                    "title": "损坏素材",
                    "chapter_no": "1.3",
                    "chapter_no_flat": "1.3",
                    "paths": [corrupt.name],
                },
            ]

            stats = merger.merge(template, plan, library, {}, root / "prep", output)
            result = Document(str(output))
            text = "\n".join(paragraph.text for paragraph in result.paragraphs)

        self.assertIn("这是可用的脱敏正文。", text)
        self.assertIn("1.2  缺失素材", text)
        self.assertIn("1.3  损坏素材", text)
        self.assertIn("[缺失：缺失素材——没有可用素材，请补充后重试]", text)
        self.assertIn("[缺失：损坏素材——没有可用素材，请补充后重试]", text)
        self.assertEqual(stats["merged_materials"], 1)
        self.assertEqual(
            stats["warnings"],
            [
                {
                    "code": "MATERIAL_MISSING",
                    "message": "1 份素材不存在，已跳过",
                    "count": 1,
                },
                {
                    "code": "MATERIAL_MERGE_FAILED",
                    "message": "1 份素材处理或合并失败，已跳过",
                    "count": 1,
                },
                {
                    "code": "DIRECTORY_WITHOUT_MATERIAL",
                    "message": "2 个目录节点没有可用素材，已保留标题并插入占位提示",
                    "count": 2,
                },
            ],
        )

    def test_runner_builds_stable_warning_summary_from_plan_and_verify_scan(self) -> None:
        runner = load_assembler_script("run_from_manifest")
        plan = [
            {"status": "MATCHED", "paths": ["valid.docx"]},
            {"status": "UNMATCHED", "paths": []},
            {"status": "NEEDS_REVIEW", "paths": []},
            {"status": "STRUCTURAL", "paths": []},
        ]
        merger_result = {
            "merged_materials": 1,
            "warnings": [
                {"code": "MATERIAL_MISSING", "message": "素材不存在", "count": 1},
                {"code": "MATERIAL_MERGE_FAILED", "message": "素材合并失败", "count": 1},
            ],
        }
        scan = {
            "placeholders": ["[待填写：参数]", "[缺失：章节]"],
            "empty_leaf_headings": ["L2 1.2 空章节"],
            "dup_alerts": ["L2 相邻重复：1.3 标题"],
            "ghost_chapters": [],
            "invalid_h1": ["错误一级标题"],
            "invalid_prefix": [],
        }

        summary, warnings = runner.build_summary_and_warnings(plan, merger_result, scan)

        self.assertEqual(summary["assembledCount"], 1)
        self.assertEqual(summary["unmatchedCount"], 1)
        self.assertEqual(summary["needsReviewCount"], 1)
        self.assertEqual(summary["structuralCount"], 1)
        self.assertEqual(summary["verification"]["placeholderCount"], 2)
        self.assertEqual(summary["verification"]["emptySectionCount"], 1)
        self.assertEqual(summary["verification"]["duplicateHeadingCount"], 1)
        self.assertEqual(summary["verification"]["abnormalHeadingCount"], 1)
        self.assertEqual(summary["warningCount"], sum(item["count"] for item in warnings))
        self.assertEqual(
            {item["code"] for item in warnings},
            {
                "MATERIAL_MISSING",
                "MATERIAL_MERGE_FAILED",
                "DIRECTORY_UNMATCHED",
                "PLACEHOLDER_REMAINS",
                "FORMAT_RISK",
            },
        )
        self.assertTrue(all(set(item) == {"code", "message", "count"} for item in warnings))

    def test_runner_ignores_invalid_intermediate_counts(self) -> None:
        runner = load_assembler_script("run_from_manifest")
        merger_result = {
            "merged_materials": "bad",
            "warnings": [
                {"code": "BAD_TEXT", "message": "非法文本", "count": "bad"},
                {"code": "BAD_NONE", "message": "空值", "count": None},
                {"code": "BAD_BOOL", "message": "布尔值", "count": True},
                {"code": "BAD_NEGATIVE", "message": "负数", "count": -1},
                {"code": "VALID", "message": "合法告警", "count": 2},
            ],
        }
        scan = {
            "placeholders": [],
            "empty_leaf_headings": [],
            "dup_alerts": [],
            "ghost_chapters": [],
            "invalid_h1": [],
            "invalid_prefix": [],
        }

        summary, warnings = runner.build_summary_and_warnings([], merger_result, scan)

        self.assertEqual(summary["assembledCount"], 0)
        self.assertEqual(summary["warningCount"], 2)
        self.assertEqual(warnings, [{"code": "VALID", "message": "合法告警", "count": 2}])


if __name__ == "__main__":
    unittest.main()
