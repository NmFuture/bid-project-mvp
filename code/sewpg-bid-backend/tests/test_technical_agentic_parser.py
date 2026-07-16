from __future__ import annotations

import asyncio
import json
import os
import sqlite3
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from app.services.bid_parse_service import technical_parse_service
from app.services.store import store


class TechnicalAgenticParserTests(unittest.TestCase):
    def setUp(self) -> None:
        store.reset_for_tests(clear_persistent=False)

    def runner_path(self) -> Path:
        return (
            Path(__file__).resolve().parents[1]
            / "opencode"
            / "skills"
            / "bid-tech-tender-structured-parser"
            / "scripts"
            / "run_from_manifest.py"
        )

    def skill_path(self) -> Path:
        return self.runner_path().parents[1] / "SKILL.md"

    def write_sample_docx(self, root: Path) -> tuple[Path, Path, Path]:
        source_path = root / "technical_tender.docx"
        output_path = root / "s1_structured_result.json"
        manifest_path = root / "s1_parse_manifest.json"
        doc = Document()
        doc.add_paragraph("华能赤峰市翁牛特旗等6个风电项目共计1998兆瓦风力发电机组及其附属设备集中采购预招标")
        doc.add_paragraph("第一章 招标公告")
        doc.add_paragraph("本项目招标范围为整套风力发电机组及塔筒内所有必要设备，包含主控柜、通讯电缆、电力电缆等配套。")
        doc.add_paragraph("投标人应按第三卷 技术规范书和技术规范专用部分要求提供完整技术方案。")
        doc.add_paragraph("投标机型应满足高电压穿越、低电压穿越、一次调频、仿真建模及频率、电压、功率调节能力。")
        doc.save(source_path)
        manifest_path.write_text(
            json.dumps(
                {
                    "projectId": "PRJ-TECH-AGENTIC-UNIT",
                    "bidType": "技术标",
                    "parseProfile": "technical",
                    "structuredResultPath": str(output_path),
                    "documents": [
                        {
                            "id": "DOC-1",
                            "name": source_path.name,
                            "sourcePath": str(source_path),
                        }
                    ],
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        return source_path, manifest_path, output_path

    def write_project_basic_docx(self, root: Path) -> tuple[Path, Path, Path]:
        source_path = root / "technical_project_basic_tender.docx"
        output_path = root / "s1_structured_result.json"
        manifest_path = root / "s1_parse_manifest.json"
        doc = Document()
        doc.add_paragraph("都匀市盛黔风电场风力发电机组及附属设备采购项目")
        doc.add_paragraph("招标编号：PC-0307-26J1-FG0002")
        doc.add_paragraph("项目单位：贵州风电项目公司")
        doc.add_paragraph("招标人：都匀盛黔新能源有限公司")
        doc.add_paragraph("招标代理机构：中电建集中采购平台")
        doc.add_paragraph("投标文件递交截止时间：2026年05月06日10时00分")
        doc.add_paragraph("本项目招标范围为整套风力发电机组及塔筒内所有必要设备。")
        doc.save(source_path)
        manifest_path.write_text(
            json.dumps(
                {
                    "projectId": "PRJ-TECH-BASICS-UNIT",
                    "bidType": "技术标",
                    "parseProfile": "technical",
                    "structuredResultPath": str(output_path),
                    "documents": [
                        {
                            "id": "DOC-1",
                            "name": source_path.name,
                            "sourcePath": str(source_path),
                        }
                    ],
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        return source_path, manifest_path, output_path

    def run_s1parse(self, *args: str, check: bool = True) -> subprocess.CompletedProcess[str]:
        env = os.environ.copy()
        env["PYTHONIOENCODING"] = "utf-8"
        return subprocess.run(
            [sys.executable, str(self.runner_path()), *args],
            check=check,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            env=env,
        )

    def test_skill_md_contains_the_fixed_excel_checklist(self) -> None:
        content = self.skill_path().read_text(encoding="utf-8")

        self.assertIn("技术标解读清单", content)
        self.assertIn("设备选型适配", content)
        self.assertIn("CMS振动监测系统", content)
        self.assertIn("二次安防系统", content)
        checklist_rows = [
            line
            for line in content.splitlines()
            if line.startswith("| ") and line.count("|") >= 4 and line.split("|")[1].strip().isdigit()
        ]
        self.assertEqual(len(checklist_rows), 58)

    def test_skill_md_presents_only_project_basics_and_technical_interpretation_targets(self) -> None:
        content = self.skill_path().read_text(encoding="utf-8")

        self.assertIn("你是风力发电设备领域的招投标技术标解读专家", content)
        self.assertIn("本 Skill 只有两个提交目标", content)
        self.assertIn("技术解读清单只约束 `technicalInterpretation`，不替代基础信息字段，也不是额外 Excel 输入", content)
        self.assertNotIn("s1_parse_manifest.json", content)
        self.assertNotIn("历史文件名", content)
        self.assertNotIn("58 条技术标解读清单就是解析范围", content)
        self.assertNotIn("58 条清单只限定技术解读目标，基础信息不受该清单限制", content)
        self.assertIn("## 输出目标一：基础信息 projectBasics", content)
        self.assertIn("## 输出目标二：技术解读 technicalInterpretation", content)
        self.assertIn("### 基础信息规则", content)
        self.assertIn("### 技术解读规则", content)
        self.assertLess(
            content.index("## 输出目标一：基础信息 projectBasics"),
            content.index("## 输出目标二：技术解读 technicalInterpretation"),
        )
        self.assertLess(content.index("### 基础信息规则"), content.index("### 技术解读规则"))
        self.assertIn("s1parse submit <manifest> projectBasics '<json>'", content)
        self.assertIn("s1parse submit <manifest> technicalInterpretation '<json>'", content)
        self.assertNotIn("s1parse submit <manifest> projectDates '<json>'", content)
        self.assertNotIn("projectDates", content)
        self.assertIn("`structured.fieldGroups.projectBasics`", content)
        self.assertIn("`structured.technicalInterpretation`", content)

    def test_prepare_builds_navigation_index_and_exposes_checklist_count(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path, _ = self.write_sample_docx(root)

            completed = self.run_s1parse("prepare", str(manifest_path))

            summary = json.loads(completed.stdout)
            self.assertEqual(summary["stage"], "prepared")
            self.assertEqual(summary["checklistCount"], 58)
            nav_path = Path(summary["navStorePath"])
            self.assertTrue(nav_path.is_file())
            conn = sqlite3.connect(nav_path)
            evidence_count = conn.execute("SELECT COUNT(*) FROM evidence").fetchone()[0]
            conn.close()
            self.assertGreaterEqual(evidence_count, 5)

    def test_prepare_does_not_extract_pdf_source_when_docling_text_is_missing(self) -> None:
        scripts_dir = self.runner_path().parent
        sys.path.insert(0, str(scripts_dir))
        try:
            from agentic import docx_indexer
        finally:
            try:
                sys.path.remove(str(scripts_dir))
            except ValueError:
                pass

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            pdf_path = root / "technical.pdf"
            pdf_path.write_bytes(b"%PDF-1.4\n")
            missing_text_path = root / "missing-docling-text.txt"
            manifest_path = root / "s1_parse_manifest.json"
            output_path = root / "s1_structured_result.json"
            manifest = {
                "projectId": "PRJ-TECH-PDF-DOCLING-FAILED",
                "bidType": "技术标",
                "parseProfile": "technical",
                "structuredResultPath": str(output_path),
                "documents": [
                    {
                        "id": "DOC-1",
                        "name": pdf_path.name,
                        "sourcePath": str(pdf_path),
                        "textPath": str(missing_text_path),
                        "documentParseEngine": "docling",
                        "documentParseStatus": "failed",
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            with patch.object(
                docx_indexer,
                "_extract_pdf_text",
                side_effect=AssertionError("技术标 PDF 必须只使用 Docling 结果，不允许 Skill 自行抽 PDF 文本"),
            ):
                summary = docx_indexer.build_index(manifest_path, manifest)

            self.assertEqual(summary["blockCount"], 0)
            document_map = json.loads((root / "document_map.json").read_text(encoding="utf-8"))
            self.assertEqual(document_map["documents"][0]["blockCount"], 0)
            self.assertEqual(document_map["documents"][0]["tableCount"], 0)

    def test_submit_validate_finalize_writes_technical_interpretation(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path, output_path = self.write_sample_docx(root)
            self.run_s1parse("prepare", str(manifest_path))
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "projectBasics",
                json.dumps(
                    [
                        {
                            "key": "projectName",
                            "label": "项目名称",
                            "value": "华能赤峰市翁牛特旗等6个风电项目共计1998兆瓦风力发电机组及其附属设备集中采购预招标",
                            "evidenceIds": ["DOC-1:B000001"],
                        }
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "technicalInterpretation",
                json.dumps(
                    [
                        {
                            "rowNo": 6,
                            "status": "found",
                            "conclusion": "招标文件明确要求整套风力发电机组及塔筒内必要设备，并列明主控柜、通讯电缆、电力电缆等配套。",
                            "evidenceSummary": "整套风力发电机组及塔筒内所有必要设备。",
                            "evidenceIds": ["DOC-1:B000003"],
                        },
                        {
                            "rowNo": 21,
                            "status": "partial",
                            "conclusion": "现有文件明确涉网控制能力要求，但型式试验报告等细节需继续核对技术规范专用部分。",
                            "evidenceSummary": "高低电压穿越、一次调频、仿真建模等能力。",
                            "neededSourceName": "第三卷 技术规范书和技术规范专用部分",
                            "evidenceIds": ["DOC-1:B000005"],
                        },
                        {
                            "rowNo": 49,
                            "status": "needs_spec",
                            "conclusion": "当前第一卷/第三卷正文未列明 CMS 测点数量，需要按原文提到的技术规范专用部分核对。",
                            "evidenceSummary": "招标文件指向技术规范书和技术规范专用部分。",
                            "neededSourceName": "第三卷 技术规范书和技术规范专用部分",
                            "evidenceIds": ["DOC-1:B000004"],
                        },
                    ],
                    ensure_ascii=False,
                ),
            )

            validation = json.loads(self.run_s1parse("validate", str(manifest_path)).stdout)
            finalize = json.loads(self.run_s1parse("finalize", str(manifest_path)).stdout)
            payload = json.loads(output_path.read_text(encoding="utf-8"))

            self.assertEqual(validation["status"], "passed")
            self.assertEqual(validation["checklistCount"], 58)
            self.assertGreaterEqual(validation["evidenceCount"], 3)
            self.assertEqual(finalize["summary"]["workflowStage"], "finalized")
            structured = payload["structured"]
            interpretation = structured["technicalInterpretation"]
            self.assertEqual(interpretation["checklistVersion"], "excel-technical-2026-06-16")
            self.assertEqual(interpretation["summary"]["total"], 58)
            self.assertEqual(len(interpretation["items"]), 58)
            by_row = {item["rowNo"]: item for item in interpretation["items"]}
            self.assertEqual(by_row[6]["displayGroup"], "供货范围界定")
            self.assertEqual(by_row[6]["status"], "found")
            self.assertEqual(by_row[49]["status"], "needs_spec")
            self.assertEqual(by_row[49]["neededSourceName"], "第三卷 技术规范书和技术规范专用部分")
            self.assertTrue(by_row[6]["evidenceRefs"])
            self.assertIn("整套风力发电机组", by_row[6]["evidenceRefs"][0]["text"])
            self.assertEqual(structured["workflow"]["mode"], "opencode-agentic-navigation")

    def test_finalize_writes_technical_project_basics_with_fixed_frontend_fields(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path, output_path = self.write_project_basic_docx(root)
            self.run_s1parse("prepare", str(manifest_path))
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "projectBasics",
                json.dumps(
                    [
                        {"key": "projectName", "label": "项目名称", "value": "都匀市盛黔风电场风力发电机组及附属设备采购项目", "evidenceIds": ["DOC-1:B000001"]},
                        {"key": "tenderNo", "label": "招标编号", "value": "PC-0307-26J1-FG0002", "evidenceIds": ["DOC-1:B000002"]},
                        {"key": "projectUnit", "label": "项目单位", "value": "贵州风电项目公司", "evidenceIds": ["DOC-1:B000003"]},
                        {"key": "tenderer", "label": "招标人", "value": "都匀盛黔新能源有限公司", "evidenceIds": ["DOC-1:B000004"]},
                        {"key": "tenderAgency", "label": "招标代理机构", "value": "中电建集中采购平台", "evidenceIds": ["DOC-1:B000005"]},
                        {"key": "bidDeadline", "label": "递交截止时间", "value": "2026年05月06日10时00分", "evidenceIds": ["DOC-1:B000006"]},
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "technicalInterpretation",
                json.dumps(
                    [{"rowNo": 6, "status": "found", "conclusion": "已识别供货范围。", "evidenceIds": ["DOC-1:B000007"]}],
                    ensure_ascii=False,
                ),
            )

            validation = json.loads(self.run_s1parse("validate", str(manifest_path)).stdout)
            self.run_s1parse("finalize", str(manifest_path))
            payload = json.loads(output_path.read_text(encoding="utf-8"))
            structured = payload["structured"]
            basics = structured["fieldGroups"]["projectBasics"]
            basics_by_key = {row["key"]: row for row in basics}

            self.assertEqual(validation["status"], "passed")
            self.assertEqual([row["key"] for row in basics], ["projectName", "tenderNo", "projectUnit", "tenderer", "tenderAgency", "bidDeadline"])
            self.assertEqual(basics_by_key["projectName"]["value"], "都匀市盛黔风电场风力发电机组及附属设备采购项目")
            self.assertEqual(basics_by_key["tenderNo"]["value"], "PC-0307-26J1-FG0002")
            self.assertEqual(basics_by_key["projectUnit"]["value"], "贵州风电项目公司")
            self.assertEqual(basics_by_key["tenderer"]["value"], "都匀盛黔新能源有限公司")
            self.assertEqual(basics_by_key["tenderAgency"]["value"], "中电建集中采购平台")
            self.assertEqual(basics_by_key["bidDeadline"]["value"], "2026-05-06 10:00")
            self.assertNotIn("projectDates", structured)
            self.assertEqual(structured["projectFactFields"], basics)
            self.assertEqual(structured["technicalInterpretation"]["summary"]["total"], 58)

    def test_frontend_contract_requires_canonical_project_basic_keys_without_project_dates_fallback(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path, output_path = self.write_project_basic_docx(root)
            self.run_s1parse("prepare", str(manifest_path))
            rejected = self.run_s1parse(
                "submit",
                str(manifest_path),
                "projectDates",
                json.dumps({"endDate": "2026-05-06 10:00", "evidenceIds": ["DOC-1:B000006"]}, ensure_ascii=False),
                check=False,
            )
            self.assertNotEqual(rejected.returncode, 0)
            self.assertIn("unsupported targetKey: projectDates", rejected.stderr)
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "projectBasics",
                json.dumps(
                    [
                        {"label": "项目名称", "value": "都匀市盛黔风电场风力发电机组及附属设备采购项目", "evidenceIds": ["DOC-1:B000001"]},
                        {"label": "采购编号", "value": "PC-0307-26J1-FG0002", "evidenceIds": ["DOC-1:B000002"]},
                        {"label": "建设单位", "value": "贵州风电项目公司", "evidenceIds": ["DOC-1:B000003"]},
                        {"label": "采购人", "value": "都匀盛黔新能源有限公司", "evidenceIds": ["DOC-1:B000004"]},
                        {"label": "采购代理机构", "value": "中电建集中采购平台", "evidenceIds": ["DOC-1:B000005"]},
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "technicalInterpretation",
                json.dumps([], ensure_ascii=False),
            )

            validation = json.loads(self.run_s1parse("validate", str(manifest_path)).stdout)
            self.run_s1parse("finalize", str(manifest_path))
            payload = json.loads(output_path.read_text(encoding="utf-8"))
            basics_by_key = {row["key"]: row for row in payload["structured"]["fieldGroups"]["projectBasics"]}
            codes = {item["code"] for item in validation["validationErrors"]}

            self.assertEqual(validation["status"], "failed")
            self.assertIn("missing_displayable_project_basic", codes)
            self.assertEqual(basics_by_key["projectName"]["value"], "")
            self.assertEqual(basics_by_key["tenderNo"]["value"], "")
            self.assertEqual(basics_by_key["projectUnit"]["value"], "")
            self.assertEqual(basics_by_key["tenderer"]["value"], "")
            self.assertEqual(basics_by_key["tenderAgency"]["value"], "")
            self.assertEqual(basics_by_key["bidDeadline"]["value"], "")
            self.assertNotIn("projectDates", payload["structured"])

    def test_local_parser_does_not_extract_technical_project_basics(self) -> None:
        script_dir = self.runner_path().parent
        if str(script_dir) not in sys.path:
            sys.path.insert(0, str(script_dir))
        import parser_core  # noqa: PLC0415

        text = "\n".join(
            [
                "华能蒙东新能源公司赤峰市200万千瓦自建调峰能力风光储多能互补一体化+荒漠治理基地项目（翁牛特旗120万千瓦风电项目区）",
                "风力发电机组（不含塔架）及附属设备采购",
                "招 标 文 件",
                "招标编号：HNZB2025-12-1-382-01",
                "招标人",
                "：",
                "华能内蒙古东部能源有限公司",
                "管理单位",
                "：",
                "华能翁牛特旗新能源有限公司",
                "招标代理机构",
                "：",
                "中国华能集团有限公司北京睿采数动科技分公司",
                "目 录",
                "6.6.1招标人的权利、义务与责任 116",
            ]
        )

        result = parser_core.parse_documents(
            [{"id": "TEN-1", "name": "招标文件-技术规范.md"}],
            {"TEN-1": text},
        )
        basics_by_key = {row["key"]: row for row in result["structured"]["fieldGroups"]["projectBasics"]}

        self.assertEqual([row["key"] for row in result["structured"]["fieldGroups"]["projectBasics"]], ["projectName", "tenderNo", "projectUnit", "tenderer", "tenderAgency", "bidDeadline"])
        self.assertTrue(all(row["value"] == "" for row in basics_by_key.values()))

    def test_validate_fails_when_technical_project_basic_evidence_is_missing_or_unsupported(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path, _ = self.write_project_basic_docx(root)
            self.run_s1parse("prepare", str(manifest_path))
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "projectBasics",
                json.dumps(
                    [
                        {"key": "projectName", "label": "项目名称", "value": "都匀市盛黔风电场风力发电机组及附属设备采购项目"},
                        {"key": "tenderer", "label": "招标人", "value": "错误招标人", "evidenceIds": ["DOC-1:B000004"]},
                        {"key": "bidDeadline", "label": "递交截止时间", "value": "2026-05-06 10:00", "evidenceIds": ["DOC-1:B000006"]},
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "technicalInterpretation",
                json.dumps([], ensure_ascii=False),
            )

            validation = json.loads(self.run_s1parse("validate", str(manifest_path), check=False).stdout)
            codes = {item["code"] for item in validation["validationErrors"]}
            field_keys = {item.get("fieldKey") for item in validation["validationErrors"]}

            self.assertEqual(validation["status"], "failed")
            self.assertIn("missing_project_basic_evidence", codes)
            self.assertIn("project_basic_value_not_supported_by_evidence", codes)
            self.assertIn("projectName", field_keys)
            self.assertIn("tenderer", field_keys)

    def test_validate_allows_missing_project_basic_notice_without_evidence(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path, output_path = self.write_project_basic_docx(root)
            self.run_s1parse("prepare", str(manifest_path))
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "projectBasics",
                json.dumps(
                    [
                        {"key": "projectName", "label": "项目名称", "value": "都匀市盛黔风电场风力发电机组及附属设备采购项目", "evidenceIds": ["DOC-1:B000001"]},
                        {"key": "tenderer", "label": "招标人", "value": "都匀盛黔新能源有限公司", "evidenceIds": ["DOC-1:B000004"]},
                        {
                            "key": "bidDeadline",
                            "label": "递交截止时间",
                            "status": "needs_spec",
                            "value": "技术规范文件未提及递交截止时间，建议补充上传第一卷 投标人须知或招标公告",
                            "neededSourceName": "第一卷 投标人须知或招标公告",
                        },
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "technicalInterpretation",
                json.dumps([], ensure_ascii=False),
            )

            validation = json.loads(self.run_s1parse("validate", str(manifest_path)).stdout)
            self.run_s1parse("finalize", str(manifest_path))
            payload = json.loads(output_path.read_text(encoding="utf-8"))
            bid_deadline = {
                row["key"]: row
                for row in payload["structured"]["fieldGroups"]["projectBasics"]
            }["bidDeadline"]

            self.assertEqual(validation["status"], "passed")
            self.assertEqual(bid_deadline["status"], "needs_spec")
            self.assertEqual(
                bid_deadline["value"],
                "技术规范文件未提及递交截止时间，建议补充上传第一卷 投标人须知或招标公告",
            )
            self.assertNotIn("projectDates", payload["structured"])

    def test_validate_rejects_non_datetime_bid_deadline_even_when_evidence_matches(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source_path = root / "technical_volume_only.docx"
            output_path = root / "s1_structured_result.json"
            manifest_path = root / "s1_parse_manifest.json"
            doc = Document()
            doc.add_paragraph("都匀市盛黔风电场风力发电机组及附属设备采购项目")
            doc.add_paragraph("招标人：都匀盛黔新能源有限公司")
            doc.add_paragraph("（第二卷 技术规范）")
            doc.add_paragraph("供货期：截止2026年8月30日前完成全部供货")
            doc.save(source_path)
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-TECH-VOLUME-ONLY",
                        "bidType": "技术标",
                        "parseProfile": "technical",
                        "structuredResultPath": str(output_path),
                        "documents": [{"id": "DOC-1", "name": source_path.name, "sourcePath": str(source_path)}],
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )
            self.run_s1parse("prepare", str(manifest_path))
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "projectBasics",
                json.dumps(
                    [
                        {"key": "projectName", "label": "项目名称", "value": "都匀市盛黔风电场风力发电机组及附属设备采购项目", "evidenceIds": ["DOC-1:B000001"]},
                        {"key": "tenderer", "label": "招标人", "value": "都匀盛黔新能源有限公司", "evidenceIds": ["DOC-1:B000002"]},
                        {"key": "bidDeadline", "label": "递交截止时间", "value": "（第二卷 技术规范）", "evidenceIds": ["DOC-1:B000003"]},
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "technicalInterpretation",
                json.dumps([], ensure_ascii=False),
            )

            validation = json.loads(self.run_s1parse("validate", str(manifest_path), check=False).stdout)
            codes = {item["code"] for item in validation["validationErrors"]}
            field_keys = {item.get("fieldKey") for item in validation["validationErrors"]}

            self.assertEqual(validation["status"], "failed")
            self.assertIn("invalid_bid_deadline_datetime", codes)
            self.assertIn("bidDeadline", field_keys)

    def test_validate_fails_when_display_project_basics_are_missing(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path, _ = self.write_project_basic_docx(root)
            self.run_s1parse("prepare", str(manifest_path))
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "technicalInterpretation",
                json.dumps(
                    [{"rowNo": 6, "status": "found", "conclusion": "已识别供货范围。", "evidenceIds": ["DOC-1:B000007"]}],
                    ensure_ascii=False,
                ),
            )

            validation = json.loads(self.run_s1parse("validate", str(manifest_path), check=False).stdout)
            codes = {item["code"] for item in validation["validationErrors"]}
            field_keys = {item.get("fieldKey") for item in validation["validationErrors"]}

            self.assertEqual(validation["status"], "failed")
            self.assertIn("missing_displayable_project_basic", codes)
            self.assertIn("projectBasics", validation["missingTargets"])
            self.assertIn("projectBasics", field_keys)

    def test_validate_fails_for_found_item_without_evidence(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path, _ = self.write_sample_docx(root)
            self.run_s1parse("prepare", str(manifest_path))
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "technicalInterpretation",
                json.dumps(
                    [
                        {
                            "rowNo": 6,
                            "status": "found",
                            "conclusion": "有明确供货范围。",
                        }
                    ],
                    ensure_ascii=False,
                ),
            )

            validation = json.loads(self.run_s1parse("validate", str(manifest_path), check=False).stdout)
            codes = {item["code"] for item in validation["validationErrors"]}

            self.assertEqual(validation["status"], "failed")
            self.assertIn("missing_evidence_for_positive_status", codes)

    def test_validate_fails_for_needs_spec_without_source_name(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path, _ = self.write_sample_docx(root)
            self.run_s1parse("prepare", str(manifest_path))
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "technicalInterpretation",
                json.dumps(
                    [
                        {
                            "rowNo": 49,
                            "status": "needs_spec",
                            "conclusion": "需要进一步核对。",
                            "evidenceIds": ["DOC-1:B000004"],
                        }
                    ],
                    ensure_ascii=False,
                ),
            )

            validation = json.loads(self.run_s1parse("validate", str(manifest_path), check=False).stdout)
            codes = {item["code"] for item in validation["validationErrors"]}

            self.assertEqual(validation["status"], "failed")
            self.assertIn("missing_needed_source_name", codes)

    def test_results_materializes_technical_evidence_refs_from_structured_file(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path, output_path = self.write_sample_docx(root)
            self.run_s1parse("prepare", str(manifest_path))
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "technicalInterpretation",
                json.dumps(
                    [
                        {
                            "rowNo": 6,
                            "status": "found",
                            "conclusion": "招标文件要求整套风力发电机组及配套设备。",
                            "evidenceIds": ["DOC-1:B000003"],
                        }
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse("finalize", str(manifest_path))
            payload = json.loads(output_path.read_text(encoding="utf-8"))
            first_item = payload["structured"]["technicalInterpretation"]["items"][0]
            first_item["evidenceRefs"] = [{"id": "DOC-1:B000003"}]
            payload["items"][0]["evidenceRefs"] = [{"id": "DOC-1:B000003"}]
            output_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

            project = store.create_project({"name": "技术标证据物化测试", "bidType": "技术标"})
            project_id = project["id"]
            technical_parse_service.complete_parse(
                project_id,
                [{"id": "TEN-1", "name": "technical_tender.docx", "size_label": "1 KB"}],
                [],
                summary={"fileCount": 1, "extractedCount": 1, "textLength": 0, "warnings": []},
                parse_storage={
                    "structuredResultPath": str(output_path),
                    "items": [],
                    "structured": {},
                },
            )

            result = asyncio.run(technical_parse_service.results(project_id))

            item = result["structured"]["technicalInterpretation"]["items"][0]
            self.assertEqual(item["evidenceRefs"][0]["id"], "DOC-1:B000003")
            self.assertIn("风力发电机组", item["evidenceRefs"][0]["text"])
            self.assertEqual(result["items"][0]["evidenceRefs"][0]["text"], item["evidenceRefs"][0]["text"])

    def test_results_materializes_project_basic_sources_from_evidence_ids(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path, output_path = self.write_project_basic_docx(root)
            self.run_s1parse("prepare", str(manifest_path))
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "projectBasics",
                json.dumps(
                    [
                        {"key": "projectName", "label": "项目名称", "value": "都匀市盛黔风电场风力发电机组及附属设备采购项目", "evidenceIds": ["DOC-1:B000001"]},
                        {"key": "tenderNo", "label": "招标编号", "value": "PC-0307-26J1-FG0002", "evidenceIds": ["DOC-1:B000002"]},
                    ],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse(
                "submit",
                str(manifest_path),
                "technicalInterpretation",
                json.dumps(
                    [{"rowNo": 6, "status": "found", "conclusion": "已识别供货范围。", "evidenceIds": ["DOC-1:B000007"]}],
                    ensure_ascii=False,
                ),
            )
            self.run_s1parse("finalize", str(manifest_path))
            payload = json.loads(output_path.read_text(encoding="utf-8"))
            basics = payload["structured"]["fieldGroups"]["projectBasics"]
            self.assertNotIn("sourceFile", basics[0])
            output_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

            project = store.create_project({"name": "技术标基础信息来源物化测试", "bidType": "技术标"})
            project_id = project["id"]
            technical_parse_service.complete_parse(
                project_id,
                [{"id": "TEN-1", "name": "technical_project_basic_tender.docx", "size_label": "1 KB"}],
                [],
                summary={"fileCount": 1, "extractedCount": 1, "textLength": 0, "warnings": []},
                parse_storage={
                    "structuredResultPath": str(output_path),
                    "items": [],
                    "structured": {},
                },
            )

            result = asyncio.run(technical_parse_service.results(project_id))

            basics_by_key = {row["key"]: row for row in result["structured"]["fieldGroups"]["projectBasics"]}
            project_name = basics_by_key["projectName"]
            self.assertEqual(project_name["sourceFile"], "technical_project_basic_tender.docx")
            self.assertEqual(project_name["sourceDocumentId"], "DOC-1")
            self.assertIn("都匀市盛黔风电场", project_name["evidenceLocation"])
            self.assertIn("都匀市盛黔风电场", project_name["evidence"])
            self.assertIn("technical_project_basic_tender.docx", project_name["sourceText"])
            self.assertIn(project_name["evidenceLocation"], project_name["sourceText"])
            self.assertNotIn("正文第", project_name["sourceText"])
            self.assertNotIn("表格第", project_name["sourceText"])
            self.assertNotIn("原文", project_name["sourceText"])
            self.assertEqual(project_name["evidenceRefs"][0]["id"], "DOC-1:B000001")
            self.assertIn("都匀市盛黔风电场", project_name["evidenceRefs"][0]["text"])

    def test_results_builds_readable_project_basic_source_from_existing_evidence(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            structured_path = root / "s1_structured_result.json"
            structured = {
                "schemaVersion": "bid-tender-structured-v1",
                "workflow": {"stage": "finalized", "mode": "opencode-agentic-navigation"},
                "sourceDocuments": [{"id": "DOC-1", "name": "technical_project_basic_tender.docx"}],
                "fieldGroups": {
                    "projectBasics": [
                        {
                            "key": "tenderNo",
                            "fieldKey": "tenderNo",
                            "label": "招标编号",
                            "value": "PC-0307-26J1-FG0002",
                            "sourceFile": "technical_project_basic_tender.docx",
                            "section": "封面",
                            "evidenceLocation": "正文第3段",
                            "evidence": "招标编号：PC-0307-26J1-FG0002",
                        }
                    ]
                },
                "projectFactFields": [],
                "technicalInterpretation": {"items": [], "summary": {"total": 0}},
            }
            structured_path.write_text(
                json.dumps({"items": [], "structured": structured}, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            project = store.create_project({"name": "技术标已有证据来源可读化测试", "bidType": "技术标"})
            project_id = project["id"]
            technical_parse_service.complete_parse(
                project_id,
                [{"id": "TEN-1", "name": "technical_project_basic_tender.docx", "size_label": "1 KB"}],
                [],
                summary={"fileCount": 1, "extractedCount": 1, "textLength": 0, "warnings": []},
                parse_storage={
                    "structuredResultPath": str(structured_path),
                    "items": [],
                    "structured": {},
                },
            )

            result = asyncio.run(technical_parse_service.results(project_id))

            tender_no = result["structured"]["fieldGroups"]["projectBasics"][0]
            self.assertEqual(tender_no["evidenceLocation"], "招标编号")
            self.assertIn("封面", tender_no["sourceText"])
            self.assertIn("technical_project_basic_tender.docx", tender_no["sourceText"])
            self.assertIn("招标编号", tender_no["sourceText"])
            self.assertNotIn("正文第3段", tender_no["sourceText"])
            self.assertNotIn("表格第", tender_no["sourceText"])
            self.assertNotIn("原文", tender_no["sourceText"])
