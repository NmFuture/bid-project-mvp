from __future__ import annotations

import hashlib
import importlib.util
import io
import json
import re
import sys
import tempfile
import unittest
import zipfile
from contextlib import redirect_stdout
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook


BACKEND_ROOT = Path(__file__).resolve().parents[1]


def skill_script_dir(skill_name: str) -> Path:
    for root_name in ("skills", "skill"):
        candidate = BACKEND_ROOT / "opencode" / root_name / skill_name / "scripts"
        if candidate.exists():
            return candidate
    return BACKEND_ROOT / "opencode" / "skills" / skill_name / "scripts"


ASSEMBLER_SCRIPT_DIR = skill_script_dir("bid-tech-assembler")
OUTLINE_SCRIPT_DIR = skill_script_dir("bid-tech-outline-generator")
WIKI_SCRIPT_DIR = skill_script_dir("bid-tech-wiki-material-builder")
GAP_PLANNER_SCRIPT_DIR = skill_script_dir("bid-tech-gap-planner")
TABLE_FILLER_SCRIPT_DIR = skill_script_dir("bid-tech-table-filler")
WORD_FILLER_SCRIPT_DIR = skill_script_dir("bid-tech-word-placeholder-filler")
MATERIAL_CLEANER_SCRIPT_DIR = skill_script_dir("bid-material-format-cleaner")


def load_assembler_script(name: str):
    spec = importlib.util.spec_from_file_location(name, ASSEMBLER_SCRIPT_DIR / f"{name}.py")
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    sys.modules[name] = module
    spec.loader.exec_module(module)
    return module


def load_outline_script(name: str):
    module_name = f"outline_{name}"
    spec = importlib.util.spec_from_file_location(module_name, OUTLINE_SCRIPT_DIR / f"{name}.py")
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    return module


def load_wiki_script(name: str):
    module_name = f"wiki_{name}"
    spec = importlib.util.spec_from_file_location(module_name, WIKI_SCRIPT_DIR / f"{name}.py")
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    return module


def load_gap_planner_script(name: str):
    module_name = f"gap_planner_{name}"
    spec = importlib.util.spec_from_file_location(module_name, GAP_PLANNER_SCRIPT_DIR / f"{name}.py")
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    return module


def load_table_filler_script(name: str):
    module_name = f"table_filler_{name}"
    spec = importlib.util.spec_from_file_location(module_name, TABLE_FILLER_SCRIPT_DIR / f"{name}.py")
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    return module


def load_word_filler_script(name: str):
    module_name = f"word_filler_{name}"
    spec = importlib.util.spec_from_file_location(module_name, WORD_FILLER_SCRIPT_DIR / f"{name}.py")
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    return module


def load_material_cleaner_script(name: str):
    module_name = f"material_cleaner_{name}"
    spec = importlib.util.spec_from_file_location(module_name, MATERIAL_CLEANER_SCRIPT_DIR / f"{name}.py")
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    return module


def json_load(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def complete_outline_review(outline_runner, manifest: dict, manifest_path: Path) -> dict:
    return outline_runner.dispatch_command(
        "review-complete",
        manifest,
        manifest_path,
        [json.dumps({"review_summary": "已完成全局复核。", "issues": []}, ensure_ascii=False)],
    )


def write_decision_context_fixture(
    root: Path,
    *,
    heading_count: int = 80,
    heading_text: str | None = None,
    template_count: int = 1,
    template_title: str = "技术方案",
) -> tuple[dict, Path]:
    manifest_path = root / "s2_input.json"
    manifest = {"workDir": str(root), "outputFile": str(root / "toc.json")}
    manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
    (root / "template_structure.json").write_text(
        json.dumps(
            {
                "schema_version": "template-structure.v1",
                "items": [
                    {
                        "number": f"第{index}章",
                        "title": f"{template_title}{index}",
                        "level": 1,
                    }
                    for index in range(1, template_count + 1)
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    blocks = [
        {
            "type": "paragraph",
            "evidence_id": f"TEN-1:B{index:06d}",
            "body_index": index,
            "text": heading_text or f"{index} 招标技术要求 " + ("详细条款" * 70),
            "toc_level": 1,
            "heading_level": 0,
            "structural_title_level": 1,
        }
        for index in range(1, heading_count + 1)
    ]
    (root / "tender_review_chunks.json").write_text(
        json.dumps(
            {
                "schema_version": "tender-review-chunks.v1",
                "chunks": [
                    {
                        "file_id": "TEN-1",
                        "file_name": "招标文件.docx",
                        "blocks": blocks,
                    }
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    (root / "tender_headings_state.json").write_text(
        json.dumps(
            {
                "schema_version": "tender-headings-state.v1",
                "next_cursor": 0,
                "complete": True,
            }
        ),
        encoding="utf-8",
    )
    return manifest, manifest_path


def submit_outline_changes(outline_runner, manifest: dict, manifest_path: Path, changes: list[dict]):
    structure = json_load(Path(manifest["workDir"]) / "template_structure.json")
    annotated = outline_runner.outline_composer.annotate_template_structure(structure)
    fingerprint = annotated["input_fingerprint"]
    deleted = {
        str(change.get("target_id") or ""): change
        for change in changes
        if change.get("operation") == "suggest_delete"
    }
    decision_max_level = outline_runner.outline_composer.DECISION_MAX_LEVEL
    template_decisions = []
    for item in annotated["items"]:
        if int(item.get("level") or 1) > decision_max_level:
            continue
        target_id = item["template_id"]
        change = deleted.get(target_id)
        if change:
            decision = {
                "target_id": target_id,
                "decision": "suggest_delete",
                "reason": change["reason"],
            }
            if "tender_basis" in change:
                decision["tender_basis"] = change["tender_basis"]
        else:
            decision = {"target_id": target_id, "decision": "retain"}
        template_decisions.append(decision)
    return outline_runner.submit_outline_decisions(
        manifest,
        manifest_path,
        {
            "schema_version": "technical-outline-decisions.v1",
            "input_fingerprint": fingerprint,
            "template_decisions": template_decisions,
            "changes": changes,
        },
    )


class TocSkillScriptTests(unittest.TestCase):
    def _prepare_finalized_appendix_outline(
        self,
        root: Path,
        decisions: list[str],
        *,
        mutate_state=None,
    ) -> tuple[object, dict, Path]:
        outline_runner = load_outline_script("run_from_manifest")
        template = root / "template.docx"
        tender = root / "tender.docx"
        output = root / "toc.json"
        manifest_path = root / "s2_input.json"

        template_doc = Document()
        template_doc.add_paragraph("第1章 技术方案", style="Heading 1")
        template_doc.save(template)
        tender_doc = Document()
        for number, title in (("附表A.1", "技术参数表"), ("附表A.2", "供货范围表")):
            tender_doc.add_paragraph(f"{number} {title}")
            table = tender_doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "参数"
            table.cell(0, 1).text = "要求"
            table.cell(1, 0).text = "示例"
            table.cell(1, 1).text = "投标人填写"
        tender_doc.save(tender)
        manifest = {
            "workDir": str(root),
            "templateFile": str(template),
            "tenderFiles": [{"id": "TEN-1", "name": tender.name, "path": str(tender)}],
            "outputFile": str(output),
            "requireComposedOutline": True,
        }
        manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
        outline_runner.write_template_structure(manifest, manifest_path)
        outline_runner.dispatch_command("headings", manifest, manifest_path, [])
        batch = outline_runner.dispatch_command("decision-next", manifest, manifest_path, [])
        appendix_heading = outline_runner.dispatch_command(
            "search", manifest, manifest_path, ["附表A.1"]
        )
        outline_runner.dispatch_command(
            "read",
            manifest,
            manifest_path,
            [appendix_heading["results"][0]["evidence_id"]],
        )
        outline_runner.dispatch_command(
            "decision-batch",
            manifest,
            manifest_path,
            [
                json.dumps(
                    {
                        "batch_token": batch["batch_token"],
                        "items": [
                            {
                                "target_id": item["target_id"],
                                "decision": "retain",
                                "reason": "历史模板专家经验保留。",
                            }
                            for item in batch["items"]
                        ],
                        "additions": [],
                    },
                    ensure_ascii=False,
                )
            ],
        )
        appendix_batch = outline_runner.dispatch_command(
            "appendix-next", manifest, manifest_path, []
        )
        self.assertEqual(len(appendix_batch["items"]), len(decisions))
        appendix_payload = []
        for index, (item, decision) in enumerate(zip(appendix_batch["items"], decisions)):
            payload = {
                "appendix_id": item["appendix_id"],
                "decision": decision,
                "reason": "逐项核验招标附表。",
            }
            if decision == "include":
                payload.update(
                    {
                        "node_id": f"ADD-APPENDIX-{index + 1}",
                        "parent_id": "ADD-APPENDIX",
                    }
                )
            appendix_payload.append(payload)
        appendix_request = {
            "batch_token": appendix_batch["batch_token"],
            "items": appendix_payload,
        }
        if "include" in decisions:
            appendix_request["root_addition"] = {
                "node_id": "ADD-APPENDIX",
                "parent_id": None,
                "number": "第2章",
                "title": "技术附表",
                "reason": "招标文件包含独立附表。",
            }
        outline_runner.dispatch_command(
            "appendix-decision-batch",
            manifest,
            manifest_path,
            [
                json.dumps(
                    appendix_request,
                    ensure_ascii=False,
                )
            ],
        )
        evidence_id = str(appendix_batch["items"][0].get("evidence_id") or "")
        if evidence_id:
            outline_runner.dispatch_command(
                "read", manifest, manifest_path, [evidence_id]
            )
        if mutate_state is not None:
            state_path = root / "outline_decision_state.json"
            state = json_load(state_path)
            mutate_state(state, appendix_batch["items"])
            state_path.write_text(json.dumps(state, ensure_ascii=False), encoding="utf-8")
        complete_outline_review(outline_runner, manifest, manifest_path)
        outline_runner.dispatch_command("decisions", manifest, manifest_path, [])
        outline_runner.compose_manifest(manifest, manifest_path)
        return outline_runner, manifest, manifest_path

    def test_bid_gap_planner_routes_fill_template_material_to_ai_fill(self) -> None:
        gap_runner = load_gap_planner_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_path = root / "toc.json"
            parse_path = root / "parse.json"
            toc_path.write_text(
                json.dumps(
                    {
                        "items": [
                            {"number": "第1章", "title": "标前概述", "level": 1},
                            {"number": "1.1", "title": "技术评分标准索引表", "level": 2},
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            parse_path.write_text(json.dumps({"structured": {}}, ensure_ascii=False), encoding="utf-8")
            manifest = {
                "projectId": "PRJ-TEST",
                "projectName": "测试项目",
                "bidType": "技术标",
                "tocJsonPath": str(toc_path),
                "parseResultPath": str(parse_path),
                "materialScope": {"paths": ["技术标/客户素材/华能集团"]},
                "projectTurbineModel": {"model": "EW10.0-220上置"},
                "materialIndex": [
                    {
                        "id": "RAW-0453",
                        "name": "待填写-技术评分标准索引表.docx",
                        "folderPath": "技术标/客户素材/华能集团/技术标-标前概述",
                        "materialTier": "customer",
                        "hasCleanedWord": True,
                        "cleanedFileName": "待填写-技术评分标准索引表.docx",
                        "requiresFill": True,
                        "placeholderCount": 3,
                        "placeholderLabels": ["投标机型", "投标方案", "章节索引"],
                    }
                ],
            }

            plan = gap_runner.build_gap_plan(manifest)

        item = next(entry for entry in plan["items"] if entry["number"] == "1.1")
        self.assertEqual(item["decision"], "fill_required")
        self.assertEqual(item["status"], "needs_input")
        self.assertEqual(item["usage"], "section_fill")
        self.assertEqual(item["matchedMaterials"], [])
        self.assertEqual(item["fillTasks"][0]["skill"], "bid-tech-word-placeholder-filler")
        self.assertEqual(item["fillTasks"][0]["blankSource"]["id"], "RAW-0453")
        self.assertEqual(item["fillTasks"][0]["blankSource"]["placeholderCount"], 3)
        self.assertEqual(item["candidateMaterials"][0]["id"], "RAW-0453")

    def test_bid_gap_planner_requires_one_result_per_confirmed_toc_item(self) -> None:
        gap_runner = load_gap_planner_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_path = root / "toc.json"
            parse_path = root / "parse.json"
            toc_items = [
                {"number": "第1章", "title": "标前概述", "level": 1},
                {"number": "1.1", "title": "技术评分标准索引表", "level": 2},
                {"number": "1.2", "title": "与华能集团签署的战略合作协议", "level": 2},
                {"number": "1.3", "title": "缺失专题", "level": 2},
            ]
            toc_path.write_text(json.dumps({"items": toc_items}, ensure_ascii=False), encoding="utf-8")
            parse_path.write_text(json.dumps({"structured": {}}, ensure_ascii=False), encoding="utf-8")
            manifest = {
                "projectId": "PRJ-TEST",
                "projectName": "测试项目",
                "bidType": "技术标",
                "tocJsonPath": str(toc_path),
                "parseResultPath": str(parse_path),
                "materialScope": {"paths": ["技术标/客户素材/华能集团"]},
                "projectTurbineModel": {"model": "EW10.0-220上置"},
                "materialIndex": [
                    {
                        "id": "RAW-0453",
                        "name": "待填写-技术评分标准索引表.docx",
                        "folderPath": "技术标/客户素材/华能集团/技术标-标前概述",
                        "materialTier": "customer",
                        "requiresFill": True,
                        "placeholderCount": 2,
                    },
                    {
                        "id": "RAW-0454",
                        "name": "固定-与华能集团签署的战略合作协议.docx",
                        "folderPath": "技术标/客户素材/华能集团/技术标-标前概述",
                        "materialTier": "customer",
                    },
                ],
            }

            plan = gap_runner.build_gap_plan(manifest)

        self.assertEqual(plan["summary"]["totalTocItems"], len(toc_items))
        self.assertEqual(len(plan["items"]), len(toc_items))
        self.assertEqual([item["number"] for item in plan["items"]], [item["number"] for item in toc_items])
        self.assertEqual(plan["integrity"]["coverageStatus"], "passed")
        self.assertEqual(plan["items"][1]["decision"], "fill_required")
        self.assertEqual(plan["items"][2]["decision"], "ready")
        self.assertEqual(plan["items"][3]["decision"], "material_required")

    def test_bid_gap_planner_carries_parse_fields_to_appendix_tasks(self) -> None:
        gap_runner = load_gap_planner_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_path = root / "toc.json"
            parse_path = root / "parse.json"
            appendix_doc = root / "APPX-A1.docx"
            Document().save(appendix_doc)
            toc_path.write_text(
                json.dumps({"items": [{"number": "附表A.1", "title": "附表A.1 投标机型总方案信息表", "level": 2}]}, ensure_ascii=False),
                encoding="utf-8",
            )
            parse_path.write_text(
                json.dumps(
                    {
                        "items": [
                            *[
                                {
                                    "id": f"REQ-NOISE-{idx}",
                                    "type": "项目基础信息",
                                    "title": "招标人",
                                    "value": f"噪声字段{idx}",
                                    "sourceFile": "招标文件.docx",
                                }
                                for idx in range(180)
                            ],
                            {
                                "id": "REQ-SCALE",
                                "type": "项目基础信息",
                                "title": "标段规模",
                                "value": "600MW",
                                "sourceFile": "招标文件.docx",
                            }
                        ],
                        "structured": {
                            "appendices": [
                                {
                                    "id": "APPX-A1",
                                    "title": "附表A.1 投标机型总方案信息表",
                                    "docxPath": str(appendix_doc),
                                }
                            ]
                        },
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            plan = gap_runner.build_gap_plan(
                {
                    "projectId": "PRJ-TEST",
                    "projectName": "测试项目",
                    "tocJsonPath": str(toc_path),
                    "parseResultPath": str(parse_path),
                    "materialIndex": [],
                }
            )

        task = plan["items"][0]["appendixTasks"][0]
        self.assertEqual(task["availableParseFields"][0]["id"], "REQ-SCALE")
        self.assertEqual(task["availableParseFields"][0]["label"], "标段规模")
        self.assertGreater(len(task["availableParseFields"]), 160)

    def test_bid_gap_planner_summary_stdout_exposes_coverage_counts(self) -> None:
        gap_runner = load_gap_planner_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_path = root / "toc.json"
            parse_path = root / "parse.json"
            manifest_path = root / "s4_gap_input.json"
            output_path = root / "gap_plan.json"
            toc_path.write_text(
                json.dumps(
                    {
                        "items": [
                            {"number": "第1章", "title": "标前概述", "level": 1},
                            {"number": "1.1", "title": "技术评分标准索引表", "level": 2},
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            parse_path.write_text(json.dumps({"structured": {}}, ensure_ascii=False), encoding="utf-8")
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-TEST",
                        "projectName": "测试项目",
                        "bidType": "技术标",
                        "tocJsonPath": str(toc_path),
                        "parseResultPath": str(parse_path),
                        "materialScope": {"paths": ["技术标/客户素材/华能集团"]},
                        "projectTurbineModel": {"model": "EW10.0-220上置"},
                        "materialIndex": [],
                        "outputFile": str(output_path),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            with patch.object(sys, "argv", ["run_from_manifest.py", "--manifest", str(manifest_path)]), \
                patch("builtins.print") as mocked_print:
                gap_runner.main()

            response = json.loads(mocked_print.call_args.args[0])

        self.assertEqual(response["tocItemCount"], 2)
        self.assertEqual(response["itemCount"], 2)
        self.assertEqual(response["coverageStatus"], "passed")

    def test_bid_gap_planner_uses_project_commitment_chapter_word_for_chapter_four_children(self) -> None:
        gap_runner = load_gap_planner_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_path = root / "toc.json"
            parse_path = root / "parse.json"
            toc_path.write_text(
                json.dumps(
                    {
                        "items": [
                            {"number": "第4章", "title": "项目技术承诺函", "level": 1},
                            {"number": "4.1", "title": "发电小时数承诺函", "level": 2},
                            {"number": "4.2", "title": "投标机组可利用率承诺", "level": 2},
                            {"number": "4.3", "title": "投标机组功率曲线保证率承诺", "level": 2},
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            parse_path.write_text(json.dumps({"structured": {}}, ensure_ascii=False), encoding="utf-8")
            manifest = {
                "projectId": "PRJ-TEST",
                "projectName": "测试项目",
                "bidType": "技术标",
                "tocJsonPath": str(toc_path),
                "parseResultPath": str(parse_path),
                "materialScope": {"paths": ["技术标/项目素材/MAT-HN-CHIFENG-001"]},
                "projectTurbineModel": {"model": "EW10.0-220上置"},
                "materialIndex": [
                    {
                        "id": "RAW-0472",
                        "name": "定制-项目技术承诺函.docx",
                        "folderPath": "技术标/项目素材/MAT-HN-CHIFENG-001/技术标-项目技术承诺函",
                        "materialTier": "project",
                        "hasCleanedWord": True,
                        "cleanedFileName": "定制-项目技术承诺函.docx",
                        "requiresFill": True,
                        "placeholderCount": 1,
                        "placeholderLabels": ["项目承诺函"],
                    }
                ],
            }

            plan = gap_runner.build_gap_plan(manifest)

        parent = next(item for item in plan["items"] if item["number"] == "第4章")
        self.assertEqual(parent["decision"], "fill_required")
        self.assertEqual(parent["status"], "needs_input")
        self.assertEqual(parent["usage"], "chapter_fill")
        self.assertEqual(parent["coverageRole"], "chapter_master")
        self.assertEqual(parent["fillTasks"][0]["blankSource"]["id"], "RAW-0472")
        for number in ["4.1", "4.2", "4.3"]:
            child = next(item for item in plan["items"] if item["number"] == number)
            self.assertEqual(child["decision"], "fill_required")
            self.assertEqual(child["status"], "needs_input")
            self.assertEqual(child["usage"], "covered_by_parent")
            self.assertEqual(child["coverageRole"], "covered_by_parent")
            self.assertEqual(child["coveredByParent"], parent["id"])
            self.assertEqual(child["matchedMaterials"], [])
            self.assertEqual(child["fillTasks"], [])

    def test_bid_gap_planner_uses_standard_delivery_acceptance_word_for_chapter_six_children(self) -> None:
        gap_runner = load_gap_planner_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_path = root / "toc.json"
            parse_path = root / "parse.json"
            toc_path.write_text(
                json.dumps(
                    {
                        "items": [
                            {"number": "第6章", "title": "产品交付、考核及验收", "level": 1},
                            {"number": "6.1", "title": "技术资料和交付进度", "level": 2},
                            {"number": "6.2", "title": "试验、检验和监造", "level": 2},
                            {"number": "6.3", "title": "设备安装、调试与试运行", "level": 2},
                            {"number": "6.4", "title": "考核指标", "level": 2},
                            {"number": "6.5", "title": "项目验收", "level": 2},
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            parse_path.write_text(json.dumps({"structured": {}}, ensure_ascii=False), encoding="utf-8")
            manifest = {
                "projectId": "PRJ-TEST",
                "projectName": "测试项目",
                "bidType": "技术标",
                "tocJsonPath": str(toc_path),
                "parseResultPath": str(parse_path),
                "materialScope": {"paths": ["技术标/通用素材"]},
                "projectTurbineModel": {"model": "EW10.0-220上置"},
                "materialIndex": [
                    {
                        "id": "RAW-0429",
                        "name": "待填写-产品交付、考核及验收.docx",
                        "folderPath": "技术标/通用素材/技术标-产品交付、考核及验收",
                        "materialTier": "standard",
                        "hasCleanedWord": True,
                        "cleanedFileName": "待填写-产品交付、考核及验收.docx",
                        "requiresFill": True,
                        "placeholderCount": 1,
                        "placeholderLabels": ["招标要求"],
                    }
                ],
            }

            plan = gap_runner.build_gap_plan(manifest)

        parent = next(item for item in plan["items"] if item["number"] == "第6章")
        self.assertEqual(parent["decision"], "fill_required")
        self.assertEqual(parent["status"], "needs_input")
        self.assertEqual(parent["usage"], "chapter_fill")
        self.assertEqual(parent["coverageRole"], "chapter_master")
        self.assertEqual(parent["fillTasks"][0]["blankSource"]["id"], "RAW-0429")
        for number in ["6.1", "6.2", "6.3", "6.4", "6.5"]:
            child = next(item for item in plan["items"] if item["number"] == number)
            self.assertEqual(child["decision"], "fill_required")
            self.assertEqual(child["status"], "needs_input")
            self.assertEqual(child["usage"], "covered_by_parent")
            self.assertEqual(child["coverageRole"], "covered_by_parent")
            self.assertEqual(child["coveredByParent"], parent["id"])
            self.assertEqual(child["matchedMaterials"], [])
            self.assertEqual(child["fillTasks"], [])

    def test_bid_gap_planner_generically_uses_parent_chapter_word_for_any_chapter_children(self) -> None:
        gap_runner = load_gap_planner_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_path = root / "toc.json"
            parse_path = root / "parse.json"
            toc_path.write_text(
                json.dumps(
                    {
                        "items": [
                            {"number": "第8章", "title": "运输安装与验收方案", "level": 1},
                            {"number": "8.1", "title": "运输组织方案", "level": 2},
                            {"number": "8.2", "title": "安装调试计划", "level": 2},
                            {"number": "8.3", "title": "验收交付安排", "level": 2},
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            parse_path.write_text(json.dumps({"structured": {}}, ensure_ascii=False), encoding="utf-8")
            manifest = {
                "projectId": "PRJ-TEST",
                "projectName": "测试项目",
                "bidType": "技术标",
                "tocJsonPath": str(toc_path),
                "parseResultPath": str(parse_path),
                "materialScope": {"paths": ["技术标/项目素材/MAT-TEST"]},
                "projectTurbineModel": {"model": "EW10.0-220上置"},
                "materialIndex": [
                    {
                        "id": "RAW-0800",
                        "name": "定制-运输安装与验收方案.docx",
                        "folderPath": "技术标/项目素材/MAT-TEST/技术标-运输安装与验收方案",
                        "materialTier": "project",
                        "hasCleanedWord": True,
                        "cleanedFileName": "定制-运输安装与验收方案.docx",
                    },
                    {
                        "id": "RAW-0801",
                        "name": "固定-运输组织方案.docx",
                        "folderPath": "技术标/项目素材/MAT-TEST/技术标-运输安装与验收方案/子节",
                        "materialTier": "project",
                    },
                ],
            }

            plan = gap_runner.build_gap_plan(manifest)

        parent = next(item for item in plan["items"] if item["number"] == "第8章")
        self.assertEqual(parent["decision"], "ready")
        self.assertEqual(parent["status"], "matched")
        self.assertEqual(parent["usage"], "chapter_master")
        self.assertEqual(parent["coverageRole"], "chapter_master")
        self.assertEqual(parent["matchedMaterials"][0]["id"], "RAW-0800")
        # 金标反评 R3b：整章覆盖是默认值不是锁——8.1 自身存在剥修饰同名素材
        # （固定-运输组织方案.docx），保留子节自主匹配，不被父章吞并。
        own_child = next(item for item in plan["items"] if item["number"] == "8.1")
        self.assertEqual(own_child["decision"], "ready")
        self.assertEqual(own_child["status"], "matched")
        self.assertEqual(own_child["matchedMaterials"][0]["id"], "RAW-0801")
        self.assertEqual(own_child["fillTasks"], [])
        for number in ["8.2", "8.3"]:
            child = next(item for item in plan["items"] if item["number"] == number)
            self.assertEqual(child["decision"], "ready")
            self.assertEqual(child["status"], "matched")
            self.assertEqual(child["usage"], "covered_by_parent")
            self.assertEqual(child["coverageRole"], "covered_by_parent")
            self.assertEqual(child["coveredByParent"], parent["id"])
            self.assertEqual(child["matchedMaterials"], [])
            self.assertEqual(child["fillTasks"], [])

    def test_bid_gap_planner_uses_generic_chapter_title_match_for_wind_resource_chapter(self) -> None:
        gap_runner = load_gap_planner_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_path = root / "toc.json"
            parse_path = root / "parse.json"
            toc_path.write_text(
                json.dumps(
                    {
                        "items": [
                            {"number": "第3章", "title": "风资源评估与机位排布方案", "level": 1},
                            {"number": "3.1", "title": "项目概况", "level": 2},
                            {"number": "3.2", "title": "风资源分析", "level": 2},
                            {"number": "3.3", "title": "机组选型", "level": 2},
                            {"number": "3.4", "title": "方案及发电量结果", "level": 2},
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            parse_path.write_text(json.dumps({"structured": {}}, ensure_ascii=False), encoding="utf-8")
            manifest = {
                "projectId": "PRJ-TEST",
                "projectName": "测试项目",
                "bidType": "技术标",
                "tocJsonPath": str(toc_path),
                "parseResultPath": str(parse_path),
                "materialScope": {"paths": ["技术标/项目素材/MAT-HN-CHIFENG-001"]},
                "projectTurbineModel": {"model": "EW10.0-220上置"},
                "materialIndex": [
                    {
                        "id": "RAW-0471",
                        "name": "定制-项目风资源评估与机组选型排布及发电量计算.docx",
                        "folderPath": "技术标/项目素材/MAT-HN-CHIFENG-001/技术标-专题方案要求",
                        "materialTier": "project",
                    },
                    {
                        "id": "RAW-0473",
                        "name": "定制-风资源评估与机位排布方案.docx",
                        "folderPath": "技术标/项目素材/MAT-HN-CHIFENG-001/技术标-风资源评估与机位排布方案",
                        "materialTier": "project",
                        "hasCleanedWord": True,
                        "cleanedFileName": "定制-风资源评估与机位排布方案.docx",
                    },
                    {
                        "id": "RAW-0478",
                        "name": "风资源评估报告.docx",
                        "folderPath": "技术标/项目素材/MAT-HN-CHIFENG-001/风资源评估报告",
                        "materialTier": "project",
                    },
                ],
            }

            plan = gap_runner.build_gap_plan(manifest)

        parent = next(item for item in plan["items"] if item["number"] == "第3章")
        self.assertEqual(parent["decision"], "ready")
        self.assertEqual(parent["status"], "matched")
        self.assertEqual(parent["coverageRole"], "chapter_master")
        self.assertEqual(parent["matchedMaterials"][0]["id"], "RAW-0473")
        self.assertNotIn("RAW-0478", [item["id"] for item in parent["matchedMaterials"]])
        for number in ["3.1", "3.2", "3.3", "3.4"]:
            child = next(item for item in plan["items"] if item["number"] == number)
            self.assertEqual(child["coverageRole"], "covered_by_parent")
            self.assertEqual(child["coveredByParent"], parent["id"])
            self.assertEqual(child["matchedMaterials"], [])

    def test_bid_gap_planner_filters_toc_refs_by_allowed_material_index(self) -> None:
        gap_runner = load_gap_planner_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_path = root / "toc.json"
            parse_path = root / "parse.json"
            toc_path.write_text(
                json.dumps(
                    {
                        "items": [
                            {
                                "number": "1.2",
                                "title": "与华能集团签署的战略合作协议",
                                "level": 2,
                                "material_refs": [
                                    {
                                        "id": "RAW-OUTSIDE",
                                        "docx": "技术标/客户素材/其他客户/战略合作协议.docx",
                                    }
                                ],
                            }
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            parse_path.write_text(json.dumps({"structured": {}}, ensure_ascii=False), encoding="utf-8")
            manifest = {
                "projectId": "PRJ-TEST",
                "projectName": "测试项目",
                "bidType": "技术标",
                "tocJsonPath": str(toc_path),
                "parseResultPath": str(parse_path),
                "materialScope": {"paths": ["技术标/客户素材/华能集团"]},
                "projectTurbineModel": {"model": "EW10.0-220上置"},
                "materialIndex": [
                    {
                        "id": "RAW-ALLOWED",
                        "name": "固定-其他资料.docx",
                        "folderPath": "技术标/客户素材/华能集团/其他",
                        "materialTier": "customer",
                    }
                ],
            }

            plan = gap_runner.build_gap_plan(manifest)

        item = plan["items"][0]
        self.assertEqual(item["decision"], "material_required")
        self.assertEqual(item["status"], "missing")
        self.assertEqual(item["matchedMaterials"], [])

    def test_bid_wiki_builder_writes_full_blueprint_and_returns_small_summary(self) -> None:
        wiki_runner = load_wiki_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest_path = root / "wiki_manifest.json"
            output_file = root / "wiki_blueprint.json"
            manifest = {
                "bidType": "技术标",
                "rootTitle": "技术标Wiki（自动生成）",
                "workDir": str(root),
                "outputFile": str(output_file),
                "stats": {"fileCount": 1},
                "tiers": [
                    {
                        "name": "标准文件",
                        "tier": "standard",
                        "path": "技术标/标准文件",
                        "fileCount": 1,
                        "folders": [
                            {
                                "name": "EW5.0",
                                "path": "技术标/标准文件/EW5.0",
                                "fileCount": 1,
                                "files": [
                                    {
                                        "id": "M1",
                                        "name": "总体方案.docx",
                                        "path": "技术标/标准文件/EW5.0/总体方案.docx",
                                        "ext": "docx",
                                    }
                                ],
                            }
                        ],
                    }
                ],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            response = wiki_runner.run_manifest(manifest, manifest_path)

            self.assertEqual(response["outputFile"], str(output_file))
            self.assertIn("nodeTitles", response)
            self.assertLess(len(json.dumps(response, ensure_ascii=False)), 1000)
            blueprint = json_load(output_file)
            self.assertEqual(blueprint["rootTitle"], "技术标Wiki（自动生成）")
            self.assertEqual(
                [node["title"] for node in blueprint["nodes"]],
                ["标准文件"],
            )
            self.assertEqual(blueprint["nodes"][0]["children"][0]["children"][0]["title"], "总体方案.docx")

    def test_bid_outline_skill_is_compact_generic_and_autonomous(self) -> None:
        skill_path = OUTLINE_SCRIPT_DIR.parent / "SKILL.md"
        content = skill_path.read_text(encoding="utf-8")

        self.assertLessEqual(len(content.splitlines()), 200)
        for case_specific_text in (
            "PRJ-0119",
            "5.8.1",
            "附表 G.2.3",
            "附表G.2.3",
            "华能",
            "上海电气",
            "projectId",
            "projectName",
            "projectCode",
        ):
            self.assertNotIn(case_specific_text, content)

        for principle in (
            "历史模板提供成熟投标经验",
            "当前招标文件决定本项目约束",
            "决策只到二级",
            "跟随其二级父节点",
            "父节点保留则整个子树保留",
            "父节点建议删除则整个子树建议删除",
            "一至三级目录",
            "s2outline prepare",
            "s2outline template-headings",
            "s2outline headings",
            "s2outline search",
            "s2outline section",
            "next_cursor",
            "requires_full_review=true",
            "s2outline next-batch",
            "s2outline read",
            "s2outline window",
            "s2outline table",
            "s2outline review-batch",
            "s2outline decision-next",
            "s2outline decision-batch",
            "一个完整决策单元",
            "一个一级章的章根加它下面的全部二级节点",
            "不做章节复核",
            "retain",
            "suggest_delete",
            "招标目录没有同名标题，不等于该节点应删除",
            "独立编制、提交、评审或评分",
            "`additions` 即使为空也必须写 `[]`",
            "s2outline appendix-next",
            "s2outline appendix-decision-batch",
            "--max-items 40",
            "严格保持返回顺序",
            "`source_status=missing` 必须 `exclude`",
            "root_addition",
            "`present`",
            "`missing`",
            "只选 `include` 或 `exclude`",
            "只做一次全局复核",
            "s2outline review-corrections",
            "s2outline review-complete",
            "s2outline decisions",
            "s2outline compose",
            "不要自行写 `manifest.outputFile`",
            "s2outline finalize",
            "timeout=300000",
            "不要检查脚本或包装器",
            "不要执行同功能的 `template`",
            "不要直接读取 `template_structure.json`",
            "内容有投标表达价值，不等于必须独立成章",
            "宽泛父节点不当然覆盖",
            "企业通用能力介绍与本项目专项响应",
            "不得改用 `reason` 规避校验",
            "从招标侧检查遗漏",
            "从模板侧检查不适用、重复或可合并节点",
            "一个短关键词或短语",
            "不能把多个无关关键词拼成一次查询",
            "疑似独立成果必须逐项读原文",
            "附表只覆盖表格填写",
            "不当然覆盖正文方案、说明、报告或承诺",
            "逐项重扫完整招标目录",
            "完整掌握模板结构",
            "先识别响应单元，再比较目录节点",
            "语义等价且粒度相当",
            "单独表达能够让评审人更清楚地看到",
            "父章节能够容纳内容，不等于目录已经覆盖",
            "仅有营销属性不是删除理由",
            "招标目录只用于定位，不能据标题判定覆盖",
            "把它的整个三级子树当成一个整体",
            "不新增三级节点",
            "并行章节会话中，`parent_id` 必须引用当前章根",
            "给了 `evidence_id`，前端就能点击跳转招标原文",
            "会把每个二级决策下沉到它的三级子树",
        ):
            self.assertIn(principle, content)

        self.assertNotIn('"operation":"collapse"', content)
        self.assertNotIn("由父节点统一承载", content)
        self.assertNotIn("模板是主骨架", content)
        self.assertNotIn("仅因招标未提及不能建议删除", content)
        self.assertNotIn("确认没有不适用证据后选择 `retain`", content)
        # 已删除的程序性门禁不应再出现在指令里。
        self.assertNotIn("新的受控正文阅读", content)
        self.assertNotIn("超过 50 个节点的超大章", content)
        self.assertNotIn("--max-items 50", content)

        for redundant_contract in (
            "required_status",
            "review_required",
            "review_note",
            "source_refs",
            "itemId",
            "toc_evidence.json",
            "当前环境不提供 `read` 工具",
            "用 `bash` 调用 `python-docx`",
            "全文审阅是完成条件",
            "不得跳过任何分块",
            "未读完的表格",
        ):
            self.assertNotIn(redundant_contract, content)

    def test_bid_outline_finalize_allows_three_levels_and_rejects_fourth_level(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            output = root / "toc.json"
            manifest_path = root / "s2_input.json"
            manifest = {"workDir": str(root), "outputFile": str(output)}
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            payload = {
                "schema_version": "technical-outline.v1",
                "nodes": [
                    {
                        "number": "1",
                        "title": "总体技术方案",
                        "suggestion_action": "必要",
                        "suggestion_reason": "",
                        "children": [
                            {
                                "number": "1.1",
                                "title": "机组设计",
                                "suggestion_action": "必要",
                                "suggestion_reason": "",
                                "children": [
                                    {
                                        "number": "1.1.1",
                                        "title": "关键部件设计",
                                        "suggestion_action": "必要",
                                        "suggestion_reason": "",
                                        "children": [],
                                    }
                                ],
                            }
                        ],
                    }
                ],
            }
            output.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

            result = outline_runner.finalize_manifest(manifest, manifest_path)
            self.assertEqual(result["summary"]["total_nodes"], 3)

            payload["nodes"][0]["children"][0]["children"][0]["children"].append(
                {
                    "number": "1.1.1.1",
                    "title": "叶片设计参数",
                    "suggestion_action": "必要",
                    "suggestion_reason": "",
                    "children": [],
                }
            )
            output.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

            with self.assertRaisesRegex(SystemExit, "技术标目录最多三级"):
                outline_runner.finalize_manifest(manifest, manifest_path)

    def test_bid_outline_docker_command_exposes_review_navigation(self) -> None:
        dockerfile = (BACKEND_ROOT / "opencode" / "Dockerfile").read_text(encoding="utf-8")

        commands = (
            "prepare|template|template-headings|headings|search|section|next-batch|read|window|table|tables|"
            "review-batch|decision-next|decision-batch|decision-reopen|review-corrections|"
            "appendix-next|appendix-decision-batch|review-complete|decisions|compose|"
            "validate|status|finalize"
        )
        self.assertIn(f"  {commands}) ;;", dockerfile)
        self.assertIn(f"usage: s2outline [{commands}] <manifest> [...]", dockerfile)
        self.assertIn('run_from_manifest.py --require-compose "$@"', dockerfile)

    def test_bid_outline_template_headings_pages_complete_template_structure(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest, manifest_path = write_decision_context_fixture(
                root,
                template_count=50,
                heading_count=0,
            )

            first = outline_runner.dispatch_command(
                "template-headings", manifest, manifest_path, ["--cursor", "0", "--page-size", "20"]
            )
            second = outline_runner.dispatch_command(
                "template-headings",
                manifest,
                manifest_path,
                ["--cursor", first["next_cursor"], "--page-size", "20"],
            )
            third = outline_runner.dispatch_command(
                "template-headings",
                manifest,
                manifest_path,
                ["--cursor", second["next_cursor"], "--page-size", "20"],
            )

        self.assertEqual(len(first["items"]), 20)
        self.assertEqual(len(second["items"]), 20)
        self.assertEqual(len(third["items"]), 10)
        self.assertEqual(first["item_count"], 50)
        self.assertFalse(first["complete"])
        self.assertTrue(third["complete"])
        self.assertEqual(third["next_cursor"], "")
        self.assertEqual(
            [item["target_id"] for page in (first, second, third) for item in page["items"]],
            [f"TPL-{index:04d}" for index in range(1, 51)],
        )

    def test_bid_outline_agent_navigation_outputs_stay_below_hard_limit(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest, manifest_path = write_decision_context_fixture(
                root,
                heading_count=300,
                heading_text="招标技术要求",
            )
            (root / "tender_appendix_inventory.json").write_text(
                json.dumps(
                    {
                        "schema_version": "tender-appendix-inventory.v1",
                        "items": [
                            {
                                "file_id": "TEN-1",
                                "number": f"附表A.{index}",
                                "title": f"技术响应表{index}",
                                "following_table_count": 1,
                            }
                            for index in range(1, 65)
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            chunks = json_load(root / "tender_review_chunks.json")
            chunks["input_fingerprint"] = "test-navigation-input"
            (root / "tender_review_chunks.json").write_text(
                json.dumps(chunks, ensure_ascii=False), encoding="utf-8"
            )
            files_by_id, _ = outline_runner.review_workflow._collect_heading_files(chunks)
            inventory = json_load(root / "tender_appendix_inventory.json")
            headings_state = json_load(root / "tender_headings_state.json")
            headings_state["input_fingerprint"] = "test-navigation-input"
            headings_state["headings_catalog_digest"] = (
                outline_runner.review_workflow._heading_catalog_digest(
                    files_by_id, inventory["items"]
                )
            )
            (root / "tender_headings_state.json").write_text(
                json.dumps(headings_state, ensure_ascii=False), encoding="utf-8"
            )

            stdout_sizes: dict[str, int] = {}

            def invoke(command: str, *args: str) -> dict:
                stdout = io.StringIO()
                with patch.object(
                    sys,
                    "argv",
                    ["run_from_manifest.py", command, str(manifest_path), *args],
                ), redirect_stdout(stdout):
                    outline_runner.main()
                output = stdout.getvalue()
                stdout_sizes[command] = len(output)
                self.assertLess(stdout_sizes[command], 45000, stdout_sizes)
                self.assertNotIn("output truncated", output.lower())
                return json.loads(output)

            invoke("headings", "--page-size", "80")
            decision_batch = invoke("decision-next")
            self.assertNotIn("comparison_context", decision_batch)
            outline_runner.dispatch_command(
                "decision-batch",
                manifest,
                manifest_path,
                [
                    json.dumps(
                        {
                            "batch_token": decision_batch["batch_token"],
                            "items": [
                                {
                                    "target_id": decision_batch["items"][0]["target_id"],
                                    "decision": "retain",
                                    "reason": "历史模板专家经验保留。",
                                }
                            ],
                            "additions": [],
                        }
                    )
                ],
            )
            appendix_batch = invoke("appendix-next", "--max-items", "20")

        self.assertEqual(len(appendix_batch["items"]), 20)
        self.assertEqual(
            set(stdout_sizes),
            {"headings", "decision-next", "appendix-next"},
        )

    def test_bid_outline_headings_page_auto_shrinks_below_output_limit(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path = write_decision_context_fixture(
                root,
                heading_count=2,
                heading_text="超长目录标题" * 3400,
            )
            chunks = json_load(root / "tender_review_chunks.json")
            chunks["input_fingerprint"] = "headings-rollback-input"
            (root / "tender_review_chunks.json").write_text(
                json.dumps(chunks, ensure_ascii=False), encoding="utf-8"
            )
            files_by_id, _ = outline_runner.review_workflow._collect_heading_files(chunks)
            state_path = root / "tender_headings_state.json"
            headings_state = json_load(state_path)
            headings_state["input_fingerprint"] = "headings-rollback-input"
            headings_state["headings_catalog_digest"] = (
                outline_runner.review_workflow._heading_catalog_digest(files_by_id, [])
            )
            state_path.write_text(
                json.dumps(headings_state, ensure_ascii=False), encoding="utf-8"
            )
            # 两条 2 万字符级超长标题 + page-size 2：脚本按字符预算自动收缩为单条返回，
            # 不再触发 45000 字符硬限回滚，按 next_cursor 连续读完。
            first_stdout = io.StringIO()
            with patch.object(
                sys,
                "argv",
                [
                    "run_from_manifest.py",
                    "headings",
                    str(manifest_path),
                    "--page-size",
                    "2",
                ],
            ), redirect_stdout(first_stdout):
                outline_runner.main()
            first = json.loads(first_stdout.getvalue())
            self.assertEqual(first["cursor"], "0")
            self.assertEqual(first["next_cursor"], "1")
            self.assertEqual(first["returned_heading_count"], 1)

            second_stdout = io.StringIO()
            with patch.object(
                sys,
                "argv",
                [
                    "run_from_manifest.py",
                    "headings",
                    str(manifest_path),
                    "--cursor",
                    "1",
                    "--page-size",
                    "2",
                ],
            ), redirect_stdout(second_stdout):
                outline_runner.main()
            second = json.loads(second_stdout.getvalue())

        self.assertEqual(second["cursor"], "1")
        self.assertEqual(second["next_cursor"], "")
        self.assertTrue(second["complete"])
        self.assertEqual(second["returned_heading_count"], 1)

    def test_bid_outline_section_default_budget_stays_below_char_limit_on_chinese_text(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            manifest_path = root / "s2_input.json"
            template_doc = Document()
            template_doc.add_paragraph("第1章 技术方案", style="Heading 1")
            template_doc.save(template)
            # 100 段密集中文正文：默认 --max-chars 12000 在字符预算下应整页承载，
            # 每页真实携带 1.2 万字，不再被字节预算压缩成更小的页。
            tender_doc = Document()
            tender_doc.add_paragraph("1 技术要求", style="Heading 1")
            for index in range(100):
                tender_doc.add_paragraph(f"第{index}条 " + "招标技术条款正文" * 60)
            tender_doc.save(tender)
            manifest = {
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": tender.name, "path": str(tender)}],
                "outputFile": str(root / "toc.json"),
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
            outline_runner.write_template_structure(manifest, manifest_path)

            headings = outline_runner.dispatch_command("headings", manifest, manifest_path, [])
            section_id = headings["files"][0]["items"][0]["section_id"]
            with self.assertRaisesRegex(SystemExit, r"items\[\]\.section_id"):
                outline_runner.dispatch_command(
                    "section", manifest, manifest_path, ["TEN-1:B000001"]
                )

            cursor = 0
            rounds = 0
            seen_records = 0
            while True:
                stdout = io.StringIO()
                with patch.object(
                    sys,
                    "argv",
                    [
                        "run_from_manifest.py",
                        "section",
                        str(manifest_path),
                        section_id,
                        "--cursor",
                        str(cursor),
                    ],
                ), redirect_stdout(stdout):
                    outline_runner.main()
                output = stdout.getvalue()
                self.assertLess(len(output), 45000)
                page = json.loads(output)
                seen_records += len(page["records"])
                rounds += 1
                if page["complete"]:
                    break
                cursor = int(page["next_cursor"])
                self.assertLess(rounds, 60)

        self.assertGreaterEqual(seen_records, 100)
        # 约 4.9 万字正文按每页 1.2 万字（--max-chars 默认值）读完，页数不应超过 6；
        # 若回退为字节预算，每页只能装约 6600 字，页数会膨胀到 8 页以上。
        self.assertLessEqual(rounds, 6)

    def test_bid_outline_appendix_output_limit_rolls_back_state_and_can_retry(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, manifest_path = write_decision_context_fixture(root, heading_count=0)

            decision_stdout = io.StringIO()
            with patch.object(
                sys,
                "argv",
                [
                    "run_from_manifest.py",
                    "decision-next",
                    str(manifest_path),
                ],
            ), redirect_stdout(decision_stdout):
                outline_runner.main()
            decision_batch = json.loads(decision_stdout.getvalue())

            decision_payload = json.dumps(
                {
                    "batch_token": decision_batch["batch_token"],
                    "items": [
                        {
                            "target_id": decision_batch["items"][0]["target_id"],
                            "decision": "retain",
                        }
                    ],
                    "additions": [],
                }
            )
            with patch.object(
                sys,
                "argv",
                [
                    "run_from_manifest.py",
                    "decision-batch",
                    str(manifest_path),
                    decision_payload,
                ],
            ), redirect_stdout(io.StringIO()):
                outline_runner.main()

            long_title = "超长技术响应附表" * 3000
            (root / "tender_appendix_inventory.json").write_text(
                json.dumps(
                    {
                        "schema_version": "tender-appendix-inventory.v1",
                        "items": [
                            {
                                "file_id": "TEN-1",
                                "number": f"附表A.{index}",
                                "title": long_title,
                                "following_table_count": 1,
                            }
                            for index in range(1, 3)
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            state_path = root / outline_runner.decision_workflow.STATE_FILE_NAME
            state_before = state_path.read_bytes()

            stdout = io.StringIO()
            with patch.object(
                sys,
                "argv",
                [
                    "run_from_manifest.py",
                    "appendix-next",
                    str(manifest_path),
                    "--max-items",
                    "2",
                ],
            ), redirect_stdout(stdout), self.assertRaisesRegex(
                SystemExit, r"command=appendix-next, actual_chars=\d+"
            ) as raised:
                outline_runner.main()

            self.assertEqual(stdout.getvalue(), "")
            self.assertEqual(state_path.read_bytes(), state_before)
            self.assertIn("--max-items", str(raised.exception))

            retry_stdout = io.StringIO()
            with patch.object(
                sys,
                "argv",
                [
                    "run_from_manifest.py",
                    "appendix-next",
                    str(manifest_path),
                    "--max-items",
                    "1",
                ],
            ), redirect_stdout(retry_stdout):
                outline_runner.main()
            retry = json.loads(retry_stdout.getvalue())

        self.assertEqual(len(retry["items"]), 1)
        self.assertEqual(retry["remaining_count"], 2)

    def test_bid_outline_navigation_output_limit_counts_linux_lf_boundary(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            manifest_path = Path(tmp) / "s2_input.json"
            manifest_path.write_text("{}", encoding="utf-8")

            below_limit = {"payload": "x" * 44984}
            stdout = io.StringIO()
            with patch.object(
                sys,
                "argv",
                ["run_from_manifest.py", "headings", str(manifest_path)],
            ), patch.object(
                outline_runner, "dispatch_command", return_value=below_limit
            ), redirect_stdout(stdout):
                outline_runner.main()
            self.assertEqual(len(stdout.getvalue()), 44999)
            self.assertTrue(stdout.getvalue().endswith("\n"))
            self.assertFalse(stdout.getvalue().endswith("\r\n"))

            at_limit = {"payload": "x" * 44985}
            decision_state_path = (
                manifest_path.parent / outline_runner.decision_workflow.STATE_FILE_NAME
            )

            def write_new_state(*_args, **_kwargs):
                decision_state_path.write_bytes(b"new navigation state")
                return at_limit

            for command in ("decision-next",):
                with self.subTest(command=command), patch.object(
                    sys,
                    "argv",
                    ["run_from_manifest.py", command, str(manifest_path)],
                ), patch.object(
                    outline_runner, "dispatch_command", side_effect=write_new_state
                ), redirect_stdout(io.StringIO()), self.assertRaisesRegex(
                    SystemExit, rf"command={command}, actual_chars=45000"
                ):
                    outline_runner.main()
                self.assertFalse(decision_state_path.exists())

    def test_bid_outline_cli_runtime_compose_gate_cannot_be_disabled_by_manifest(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            output = root / "toc.json"
            manifest_path = root / "s2_input.json"
            manifest = {
                "workDir": str(root),
                "outputFile": str(output),
                "requireComposedOutline": False,
            }
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            (root / "template_structure.json").write_text(
                json.dumps(
                    {
                        "schema_version": "template-structure.v1",
                        "items": [{"number": "第1章", "title": "技术方案", "level": 1}],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            submit_outline_changes(outline_runner, manifest, manifest_path, [])
            output.write_text(
                json.dumps(
                    {
                        "schema_version": "technical-outline.v1",
                        "nodes": [
                            {
                                "number": "第1章",
                                "title": "技术方案",
                                "suggestion_action": "必要",
                                "suggestion_reason": "",
                                "children": [],
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            with patch.object(
                sys,
                "argv",
                ["run_from_manifest.py", "--require-compose", "finalize", str(manifest_path)],
            ), self.assertRaisesRegex(SystemExit, "尚未执行 s2outline compose"):
                outline_runner.main()

    def test_bid_outline_compose_keeps_all_explicitly_retained_template_level_three(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            output = root / "toc.json"
            manifest_path = root / "s2_input.json"
            manifest = {"workDir": str(root), "outputFile": str(output)}
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            (root / "template_structure.json").write_text(
                json.dumps(
                    {
                        "schema_version": "template-structure.v1",
                        "items": [
                            {"number": "第1章", "title": "技术方案", "level": 1},
                            {"number": "1.1", "title": "总体设计", "level": 2},
                            {"number": "1.1.1", "title": "设计原则", "level": 3},
                            {"number": "1.1.2", "title": "设计边界", "level": 3},
                            {"number": "1.2", "title": "供货方案", "level": 2},
                            {"number": "1.2.1", "title": "供货范围", "level": 3},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            submit_outline_changes(outline_runner, manifest, manifest_path, [])
            result = outline_runner.compose_manifest(manifest, manifest_path)
            outline = json_load(output)

        first_root = outline["nodes"][0]
        self.assertEqual(
            [child["title"] for parent in first_root["children"] for child in parent["children"]],
            ["设计原则", "设计边界", "供货范围"],
        )
        self.assertEqual(result["summary"]["templateLevel3"]["templateCount"], 3)
        self.assertEqual(result["summary"]["templateLevel3"]["retainedCount"], 3)
        self.assertEqual(result["summary"]["templateLevel3"]["unexplainedMissingCount"], 0)

    def test_bid_outline_compose_preserves_realistic_6_48_150_template_structure(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest_path = root / "s2_input.json"
            output = root / "toc.json"
            manifest = {"workDir": str(root), "outputFile": str(output)}
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            items = []
            expected_children_by_parent: dict[str, list[tuple[str, str]]] = {}
            parent_index = 0
            for chapter_index in range(1, 7):
                items.append(
                    {
                        "number": f"第{chapter_index}章",
                        "title": f"技术方案{chapter_index}",
                        "level": 1,
                    }
                )
                for section_index in range(1, 9):
                    parent_index += 1
                    parent_number = f"{chapter_index}.{section_index}"
                    items.append(
                        {
                            "number": parent_number,
                            "title": f"二级专题{parent_index}",
                            "level": 2,
                        }
                    )
                    child_count = 5 if parent_index <= 14 else 4 if parent_index <= 34 else 0
                    expected_children_by_parent[parent_number] = []
                    for child_index in range(1, child_count + 1):
                        child_number = f"{parent_number}.{child_index}"
                        child_title = f"三级专题{parent_index}-{child_index}"
                        items.append(
                            {
                                "number": child_number,
                                "title": child_title,
                                "level": 3,
                            }
                        )
                        expected_children_by_parent[parent_number].append(
                            (child_number, child_title)
                        )
            (root / "template_structure.json").write_text(
                json.dumps(
                    {"schema_version": "template-structure.v1", "items": items},
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            submit_outline_changes(outline_runner, manifest, manifest_path, [])
            result = outline_runner.compose_manifest(manifest, manifest_path)
            outline = json_load(output)

        parents = [parent for root_node in outline["nodes"] for parent in root_node["children"]]
        coverage = result["summary"]["templateLevel3"]
        coverage_by_parent = {item["number"]: item for item in coverage["parents"]}
        self.assertEqual(len(outline["nodes"]), 6)
        self.assertEqual(len(parents), 48)
        self.assertEqual(sum(len(parent["children"]) for parent in parents), 150)
        self.assertEqual(coverage["templateCount"], 150)
        self.assertEqual(coverage["retainedCount"], 150)
        self.assertEqual(coverage["unexplainedMissingCount"], 0)
        self.assertEqual(len(coverage["parents"]), 48)
        for parent in parents:
            expected_children = expected_children_by_parent[parent["number"]]
            self.assertEqual(
                [(child["number"], child["title"]) for child in parent["children"]],
                expected_children,
            )
            self.assertTrue(
                all(
                    child["number"].startswith(f'{parent["number"]}.')
                    for child in parent["children"]
                )
            )
            self.assertEqual(
                coverage_by_parent[parent["number"]]["finalCount"],
                len(expected_children),
            )

    def test_bid_outline_compose_applies_explicit_collapse_and_addition(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            output = root / "toc.json"
            manifest_path = root / "s2_input.json"
            manifest = {"workDir": str(root), "outputFile": str(output)}
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            (root / "template_structure.json").write_text(
                json.dumps(
                    {
                        "schema_version": "template-structure.v1",
                        "items": [
                            {"number": "第1章", "title": "技术方案", "level": 1},
                            {"number": "1.1", "title": "总体设计", "level": 2},
                            {"number": "1.1.1", "title": "设计原则", "level": 3},
                            {"number": "1.1.2", "title": "设计边界", "level": 3},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            decisions = {
                "schema_version": "technical-outline-decisions.v1",
                "changes": [
                    {
                        "operation": "collapse",
                        "target_id": "TPL-0003",
                        "reason": "内容由总体设计统一承载，不需要独立分工和审核。",
                    },
                    {
                        "operation": "add",
                        "node_id": "ADD-0001",
                        "parent_id": "TPL-0002",
                        "after_id": "TPL-0004",
                        "number": "1.1.3",
                        "title": "专项计算报告",
                        "suggestion_action": "建议增加",
                        "suggestion_reason": "招标明确要求独立提交专项计算报告。",
                    },
                ],
            }

            submit_outline_changes(outline_runner, manifest, manifest_path, decisions["changes"])
            result = outline_runner.compose_manifest(manifest, manifest_path)
            outline = json_load(output)

        children = outline["nodes"][0]["children"][0]["children"]
        self.assertEqual([child["title"] for child in children], ["设计边界", "专项计算报告"])
        self.assertEqual(result["summary"]["templateLevel3"]["collapsedCount"], 1)
        self.assertEqual(result["summary"]["templateLevel3"]["addedCount"], 1)
        self.assertEqual(result["summary"]["templateLevel3"]["unexplainedMissingCount"], 0)

    def test_bid_outline_finalize_reports_missing_template_level_three_without_blocking(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            output = root / "toc.json"
            manifest_path = root / "s2_input.json"
            manifest = {"workDir": str(root), "outputFile": str(output)}
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            (root / "template_structure.json").write_text(
                json.dumps(
                    {
                        "schema_version": "template-structure.v1",
                        "items": [
                            {"number": "第1章", "title": "技术方案", "level": 1},
                            {"number": "1.1", "title": "总体设计", "level": 2},
                            {"number": "1.1.1", "title": "设计原则", "level": 3},
                            {"number": "1.1.2", "title": "设计边界", "level": 3},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            output.write_text(
                json.dumps(
                    {
                        "schema_version": "technical-outline.v1",
                        "nodes": [
                            {
                                "number": "第1章",
                                "title": "技术方案",
                                "suggestion_action": "必要",
                                "suggestion_reason": "",
                                "children": [
                                    {
                                        "number": "1.1",
                                        "title": "总体设计",
                                        "suggestion_action": "必要",
                                        "suggestion_reason": "",
                                        "children": [
                                            {
                                                "number": "1.1.1",
                                                "title": "设计原则",
                                                "suggestion_action": "必要",
                                                "suggestion_reason": "",
                                                "children": [],
                                            }
                                        ],
                                    }
                                ],
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = outline_runner.finalize_manifest(manifest, manifest_path)

        coverage = result["summary"]["templateLevel3"]
        self.assertEqual(coverage["templateCount"], 2)
        self.assertEqual(coverage["retainedCount"], 1)
        self.assertEqual(coverage["unexplainedMissingCount"], 1)
        self.assertEqual(coverage["parentsWithUnexplainedMissing"], 1)

    def test_bid_outline_finalize_requires_unchanged_composed_output_when_enabled(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            output = root / "toc.json"
            manifest_path = root / "s2_input.json"
            manifest = {
                "workDir": str(root),
                "outputFile": str(output),
                "requireComposedOutline": True,
            }
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            (root / "template_structure.json").write_text(
                json.dumps(
                    {
                        "schema_version": "template-structure.v1",
                        "items": [{"number": "第1章", "title": "技术方案", "level": 1}],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            submit_outline_changes(outline_runner, manifest, manifest_path, [])
            outline_runner.compose_manifest(manifest, manifest_path)
            result = outline_runner.finalize_manifest(manifest, manifest_path)
            payload = json_load(output)
            payload["nodes"][0]["title"] = "被临时脚本改写的目录"
            output.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

            with self.assertRaisesRegex(SystemExit, "compose 后被修改"):
                outline_runner.finalize_manifest(manifest, manifest_path)

        self.assertEqual(result["summary"]["workflowStage"], "finalized")

    def test_bid_outline_compose_keeps_suggested_delete_and_tracks_level_three_move(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            output = root / "toc.json"
            manifest_path = root / "s2_input.json"
            manifest = {"workDir": str(root), "outputFile": str(output)}
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            (root / "template_structure.json").write_text(
                json.dumps(
                    {
                        "schema_version": "template-structure.v1",
                        "items": [
                            {"number": "第1章", "title": "技术方案", "level": 1},
                            {"number": "1.1", "title": "总体设计", "level": 2},
                            {"number": "1.1.1", "title": "原设计专题", "level": 3},
                            {"number": "1.2", "title": "专项方案", "level": 2},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            decisions = {
                "schema_version": "technical-outline-decisions.v1",
                "changes": [
                    {
                        "operation": "suggest_delete",
                        "target_id": "TPL-0002",
                        "reason": "当前项目供货边界不包含该总体设计工作。",
                    },
                    {
                        "operation": "update",
                        "target_id": "TPL-0003",
                        "parent_id": "TPL-0004",
                        "number": "1.2.1",
                        "title": "调整后的设计专题",
                        "reason": "该专题应归入专项方案并独立编制。",
                    },
                ],
            }

            submit_outline_changes(outline_runner, manifest, manifest_path, decisions["changes"])
            result = outline_runner.compose_manifest(manifest, manifest_path)
            outline = json_load(output)

        first_parent, second_parent = outline["nodes"][0]["children"]
        self.assertEqual(first_parent["suggestion_action"], "建议删除")
        self.assertEqual(first_parent["suggestion_reason"], "当前项目供货边界不包含该总体设计工作。")
        self.assertEqual(first_parent["children"], [])
        self.assertEqual(second_parent["children"][0]["number"], "1.2.1")
        self.assertEqual(second_parent["children"][0]["title"], "调整后的设计专题")
        parents = {
            item["parentId"]: item for item in result["summary"]["templateLevel3"]["parents"]
        }
        self.assertEqual(parents["TPL-0002"]["retainedCount"], 0)
        self.assertEqual(parents["TPL-0002"]["movedOutCount"], 1)
        self.assertEqual(parents["TPL-0004"]["movedInCount"], 1)
        self.assertEqual(parents["TPL-0004"]["finalCount"], 1)

    def test_bid_outline_level_three_report_includes_added_parent(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest_path = root / "s2_input.json"
            manifest = {"workDir": str(root), "outputFile": str(root / "toc.json")}
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            (root / "template_structure.json").write_text(
                json.dumps(
                    {
                        "schema_version": "template-structure.v1",
                        "items": [{"number": "第1章", "title": "技术方案", "level": 1}],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            changes = [
                {
                    "operation": "add",
                    "node_id": "ADD-L2",
                    "parent_id": "TPL-0001",
                    "number": "1.1",
                    "title": "新增专项",
                    "suggestion_action": "建议增加",
                    "suggestion_reason": "招标要求新增专项。",
                },
                {
                    "operation": "add",
                    "node_id": "ADD-L3",
                    "parent_id": "ADD-L2",
                    "number": "1.1.1",
                    "title": "新增专项报告",
                    "suggestion_action": "建议增加",
                    "suggestion_reason": "招标要求独立提交报告。",
                },
            ]

            submit_outline_changes(outline_runner, manifest, manifest_path, changes)
            result = outline_runner.compose_manifest(manifest, manifest_path)

        parents = {
            item["parentId"]: item for item in result["summary"]["templateLevel3"]["parents"]
        }
        self.assertEqual(parents["ADD-L2"]["templateCount"], 0)
        self.assertEqual(parents["ADD-L2"]["addedCount"], 1)
        self.assertEqual(parents["ADD-L2"]["finalCount"], 1)

    def test_bid_outline_level_three_report_distinguishes_cross_level_move_from_addition(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest_path = root / "s2_input.json"
            manifest = {"workDir": str(root), "outputFile": str(root / "toc.json")}
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            (root / "template_structure.json").write_text(
                json.dumps(
                    {
                        "schema_version": "template-structure.v1",
                        "items": [
                            {"number": "第1章", "title": "技术方案", "level": 1},
                            {"number": "1.1", "title": "总体设计", "level": 2},
                            {"number": "1.1.1", "title": "原三级专题", "level": 3},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            changes = [
                {
                    "operation": "update",
                    "target_id": "TPL-0003",
                    "parent_id": None,
                    "number": "第2章",
                    "reason": "该专题调整为独立根章节。",
                },
                {
                    "operation": "add",
                    "node_id": "ADD-L3",
                    "parent_id": "TPL-0002",
                    "number": "1.1.2",
                    "title": "新增三级报告",
                    "suggestion_action": "建议增加",
                    "suggestion_reason": "招标要求独立提交。",
                },
            ]

            submit_outline_changes(outline_runner, manifest, manifest_path, changes)
            result = outline_runner.compose_manifest(manifest, manifest_path)

        coverage = result["summary"]["templateLevel3"]
        self.assertEqual(coverage["templateCount"], 1)
        self.assertEqual(coverage["outputCount"], 1)
        self.assertEqual(coverage["retainedCount"], 0)
        self.assertEqual(coverage["movedOutOfLevel3Count"], 1)
        self.assertEqual(coverage["addedCount"], 1)
        self.assertEqual(coverage["unexplainedMissingCount"], 0)

    def test_bid_outline_level_three_report_counts_template_node_moved_into_level_three(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest_path = root / "s2_input.json"
            manifest = {"workDir": str(root), "outputFile": str(root / "toc.json")}
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            (root / "template_structure.json").write_text(
                json.dumps(
                    {
                        "schema_version": "template-structure.v1",
                        "items": [
                            {"number": "第1章", "title": "技术方案", "level": 1},
                            {"number": "1.1", "title": "原二级叶节点", "level": 2},
                            {"number": "1.2", "title": "目标二级节点", "level": 2},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            changes = [
                {
                    "operation": "update",
                    "target_id": "TPL-0002",
                    "parent_id": "TPL-0003",
                    "number": "1.2.1",
                    "reason": "调整为目标二级节点下的独立三级专题。",
                }
            ]

            submit_outline_changes(outline_runner, manifest, manifest_path, changes)
            result = outline_runner.compose_manifest(manifest, manifest_path)

        coverage = result["summary"]["templateLevel3"]
        parents = {item["parentId"]: item for item in coverage["parents"]}
        self.assertEqual(coverage["outputCount"], 1)
        self.assertEqual(coverage["retainedCount"], 0)
        self.assertEqual(coverage["addedCount"], 0)
        self.assertEqual(coverage["movedIntoLevel3Count"], 1)
        self.assertEqual(parents["TPL-0003"]["movedInCount"], 1)
        self.assertEqual(parents["TPL-0003"]["finalCount"], 1)

    def test_bid_outline_finalize_rejects_missing_or_stale_compose_receipt(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            output = root / "toc.json"
            manifest_path = root / "s2_input.json"
            manifest = {
                "workDir": str(root),
                "outputFile": str(output),
                "requireComposedOutline": True,
            }
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            (root / "template_structure.json").write_text(
                json.dumps(
                    {
                        "schema_version": "template-structure.v1",
                        "items": [{"number": "第1章", "title": "技术方案", "level": 1}],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            output.write_text(
                json.dumps(
                    {
                        "schema_version": "technical-outline.v1",
                        "nodes": [
                            {
                                "number": "第1章",
                                "title": "技术方案",
                                "suggestion_action": "必要",
                                "suggestion_reason": "",
                                "children": [],
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            submit_outline_changes(outline_runner, manifest, manifest_path, [])
            with self.assertRaisesRegex(SystemExit, "尚未执行 s2outline compose"):
                outline_runner.finalize_manifest(manifest, manifest_path)

            outline_runner.compose_manifest(manifest, manifest_path)
            decisions_path = root / "outline_authoring_decisions.json"
            changed_decisions = json_load(decisions_path)
            changed_decisions["changes"] = [
                {
                    "operation": "add",
                    "node_id": "ADD-STALE",
                    "parent_id": None,
                    "number": "第2章",
                    "title": "新增章节",
                    "suggestion_action": "建议增加",
                    "suggestion_reason": "用于验证决策变更后回执失效。",
                }
            ]
            decisions_path.write_text(json.dumps(changed_decisions, ensure_ascii=False), encoding="utf-8")
            with self.assertRaisesRegex(SystemExit, "decisions.*compose 后被修改"):
                outline_runner.finalize_manifest(manifest, manifest_path)

    def test_bid_outline_compose_requires_current_explicit_decisions(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest_path = root / "s2_input.json"
            manifest = {"workDir": str(root), "outputFile": str(root / "toc.json")}
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            structure_path = root / "template_structure.json"
            structure_path.write_text(
                json.dumps(
                    {
                        "schema_version": "template-structure.v1",
                        "items": [{"number": "第1章", "title": "技术方案", "level": 1}],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            with self.assertRaisesRegex(SystemExit, "必须先执行 s2outline decisions"):
                outline_runner.compose_manifest(manifest, manifest_path)

            with self.assertRaisesRegex(SystemExit, "input_fingerprint.*required"):
                outline_runner.submit_outline_decisions(
                    manifest,
                    manifest_path,
                    {"schema_version": "technical-outline-decisions.v1", "changes": []},
                )

            structure = json_load(structure_path)
            fingerprint = outline_runner.outline_composer.annotate_template_structure(structure)[
                "input_fingerprint"
            ]
            with self.assertRaisesRegex(SystemExit, "does not match"):
                outline_runner.submit_outline_decisions(
                    manifest,
                    manifest_path,
                    {
                        "schema_version": "technical-outline-decisions.v1",
                        "input_fingerprint": "stale-template-fingerprint",
                        "template_decisions": [],
                        "changes": [],
                    },
                )

            outline_runner.submit_outline_decisions(
                manifest,
                manifest_path,
                {
                    "schema_version": "technical-outline-decisions.v1",
                    "input_fingerprint": fingerprint,
                    "template_decisions": [
                        {"target_id": "TPL-0001", "decision": "retain"}
                    ],
                    "changes": [],
                },
            )
            result = outline_runner.compose_manifest(manifest, manifest_path)

        self.assertEqual(result["summary"]["total_nodes"], 1)

    def test_bid_outline_decisions_require_an_explicit_choice_for_every_template_node(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest_path = root / "s2_input.json"
            manifest = {"workDir": str(root), "outputFile": str(root / "toc.json")}
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            structure = {
                "schema_version": "template-structure.v1",
                "items": [
                    {"number": "Chapter 1", "title": "Overview", "level": 1},
                    {"number": "1.1", "title": "Legacy topic", "level": 2},
                ],
            }
            (root / "template_structure.json").write_text(
                json.dumps(structure), encoding="utf-8"
            )
            annotated = outline_runner.outline_composer.annotate_template_structure(structure)
            incomplete = {
                "schema_version": "technical-outline-decisions.v1",
                "input_fingerprint": annotated["input_fingerprint"],
                "template_decisions": [
                    {"target_id": "TPL-0001", "decision": "retain"},
                ],
                "changes": [],
            }

            with self.assertRaisesRegex(SystemExit, "TPL-0002"):
                outline_runner.submit_outline_decisions(manifest, manifest_path, incomplete)

            complete = {
                **incomplete,
                "template_decisions": [
                    {"target_id": "TPL-0001", "decision": "retain"},
                    {
                        "target_id": "TPL-0002",
                        "decision": "suggest_delete",
                        "reason": "The tender explicitly excludes this scope.",
                    },
                ],
            }
            outline_runner.submit_outline_decisions(manifest, manifest_path, complete)
            outline_runner.compose_manifest(manifest, manifest_path)
            outline = json_load(root / "toc.json")

        self.assertEqual(outline["nodes"][0]["suggestion_action"], "必要")
        self.assertEqual(
            outline["nodes"][0]["children"][0]["suggestion_action"],
            "建议删除",
        )

    def test_bid_outline_decision_next_returns_only_current_template_chapter(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest, manifest_path = write_decision_context_fixture(
                root,
                template_count=50,
                template_title="超长中文技术方案章节" * 30,
            )

            batch = outline_runner.dispatch_command(
                "decision-next", manifest, manifest_path, []
            )

        compact = json.dumps(batch, ensure_ascii=False, separators=(",", ":"))
        self.assertLess(len(compact), 45000)
        self.assertEqual(len(batch["items"]), 1)
        self.assertEqual(batch["remaining_count"], 50)
        self.assertNotIn("comparison_context", batch)
        self.assertEqual(batch["decision_level"], 2)
        # 该夹具的章只有章根（无二级节点），属稀疏章：切换为连续通读构建纪律
        self.assertEqual(batch["authoring_mode"], "sparse_chapter")
        self.assertEqual(
            batch["decision_steps"],
            [
                "本章模板二级节点稀疏，需要从招标原文构建子目录：先通读完整招标目录，再自主圈定与本章主题对应的少数上级章节",
                "对圈定章节用 section --max-chars 30000 连续通读原文，不逐段跳读；search 全会话最多 2 次，仅用于跨章定位",
                "从已读原文自主提炼响应单元并确定标题与粒度，一次性提交全部新增；每个新增 reason + evidence_id",
                "章根与既有二级节点仍按 retain / suggest_delete 表态",
            ],
        )
        self.assertEqual(
            batch["submission_contract"],
            {
                "required_fields": ["batch_token", "items", "additions"],
                "items_must_match_batch": True,
                "additions_must_be_explicit": True,
                "decision_covers_subtree": "对二级节点的判断适用于其下全部三级节点",
                "addition_levels": "章节会话只能在当前一级章下新增二级节点，parent_id 必须引用当前章根；不新增一级章或三级节点",
            },
        )

    def test_bid_outline_decision_next_uses_one_complete_root_chapter(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest, manifest_path = write_decision_context_fixture(root, heading_count=0)
            (root / "template_structure.json").write_text(
                json.dumps(
                    {
                        "schema_version": "template-structure.v1",
                        "items": [
                            {"number": "第1章", "title": "总体方案", "level": 1},
                            {"number": "1.1", "title": "设计依据", "level": 2},
                            {"number": "1.1.1", "title": "采用标准", "level": 3},
                            {"number": "第2章", "title": "供货方案", "level": 1},
                            {"number": "2.1", "title": "供货范围", "level": 2},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            first = outline_runner.dispatch_command(
                "decision-next", manifest, manifest_path, []
            )
            outline_runner.dispatch_command(
                "decision-batch",
                manifest,
                manifest_path,
                [
                    json.dumps(
                        {
                            "batch_token": first["batch_token"],
                            "items": [
                                {"target_id": item["target_id"], "decision": "retain"}
                                for item in first["items"]
                            ],
                            "additions": [],
                        },
                        ensure_ascii=False,
                    )
                ],
            )
            second = outline_runner.dispatch_command(
                "decision-next", manifest, manifest_path, []
            )

        self.assertEqual(first["chapter_id"], "TPL-0001")
        # TPL-0003 是三级节点，跟随二级父节点 TPL-0002，不进入决策批次。
        self.assertEqual([item["target_id"] for item in first["items"]], ["TPL-0001", "TPL-0002"])
        self.assertEqual(second["chapter_id"], "TPL-0004")
        self.assertEqual([item["target_id"] for item in second["items"]], ["TPL-0004", "TPL-0005"])

    def test_bid_outline_decision_next_keeps_one_batch_per_chapter_at_level_two(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest, manifest_path = write_decision_context_fixture(root, heading_count=0)
            items = [{"number": "第1章", "title": "专题方案", "level": 1}]
            for section in range(1, 4):
                items.append(
                    {"number": f"1.{section}", "title": f"专题{section}", "level": 2}
                )
                items.extend(
                    {
                        "number": f"1.{section}.{child}",
                        "title": f"专题{section}响应内容{child}",
                        "level": 3,
                    }
                    for child in range(1, 21)
                )
            (root / "template_structure.json").write_text(
                json.dumps(
                    {"schema_version": "template-structure.v1", "items": items},
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            chapter = outline_runner.dispatch_command(
                "decision-next", manifest, manifest_path, []
            )
            outline_runner.dispatch_command(
                "decision-batch",
                manifest,
                manifest_path,
                [
                    json.dumps(
                        {
                            "batch_token": chapter["batch_token"],
                            "items": [
                                {
                                    "target_id": item["target_id"],
                                    "decision": "retain",
                                    "reason": "保留完整专题方案章节。",
                                }
                                for item in chapter["items"]
                            ],
                            "additions": [],
                        },
                        ensure_ascii=False,
                    )
                ],
            )
            after = outline_runner.dispatch_command(
                "decision-next", manifest, manifest_path, []
            )

        # 60 个三级节点跟随各自的二级父节点，整章只剩章根 + 3 个二级节点一批决完。
        self.assertEqual(chapter["chapter_id"], "TPL-0001")
        self.assertEqual(
            [item["number"] for item in chapter["items"]],
            ["第1章", "1.1", "1.2", "1.3"],
        )
        self.assertTrue(after["complete"])

    def test_bid_outline_decision_next_uses_actual_first_page_size_for_large_chapter(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest, manifest_path = write_decision_context_fixture(
                root,
                heading_count=160,
                heading_text="招标技术要求与独立响应事项" * 4,
            )
            items = [{"number": "1", "title": "总体技术方案", "level": 1}]
            items.extend(
                {
                    "number": f"1.{index}",
                    "title": "需要结合招标目录逐项判断的专业技术专题" * 6,
                    "level": 2,
                }
                for index in range(1, 30)
            )
            (root / "template_structure.json").write_text(
                json.dumps({"schema_version": "template-structure.v1", "items": items}, ensure_ascii=False),
                encoding="utf-8",
            )

            batch = outline_runner.dispatch_command("decision-next", manifest, manifest_path, [])

        compact = json.dumps(batch, ensure_ascii=False, separators=(",", ":"))
        self.assertEqual(len(batch["items"]), 30)
        self.assertNotIn("comparison_context", batch)
        self.assertLess(len(compact), 45000)

    def test_bid_outline_decision_batch_requires_explicit_additions(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest, manifest_path = write_decision_context_fixture(root, heading_count=0)
            batch = outline_runner.dispatch_command(
                "decision-next", manifest, manifest_path, []
            )

            with self.assertRaisesRegex(SystemExit, "additions.*required"):
                outline_runner.dispatch_command(
                    "decision-batch",
                    manifest,
                    manifest_path,
                    [
                        json.dumps(
                            {
                                "batch_token": batch["batch_token"],
                                "items": [
                                    {"target_id": item["target_id"], "decision": "retain"}
                                    for item in batch["items"]
                                ],
                            }
                        )
                    ],
                )

    def test_bid_outline_decision_batch_requires_reason_or_read_tender_evidence(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            manifest_path = root / "s2_input.json"
            template_doc = Document()
            template_doc.add_paragraph("第1章 技术方案", style="Heading 1")
            template_doc.save(template)
            tender_doc = Document()
            tender_doc.add_paragraph("1 专项方案", style="Heading 1")
            tender_doc.add_paragraph("投标人必须提交海上运输安全专项方案。")
            tender_doc.save(tender)
            manifest = {
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": tender.name, "path": str(tender)}],
                "outputFile": str(root / "toc.json"),
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
            outline_runner.write_template_structure(manifest, manifest_path)
            headings = outline_runner.dispatch_command("headings", manifest, manifest_path, [])
            batch = outline_runner.dispatch_command("decision-next", manifest, manifest_path, [])

            self.assertNotIn("comparison_context", batch)
            with self.assertRaisesRegex(SystemExit, "evidence_id.*reason"):
                outline_runner.dispatch_command(
                    "decision-batch",
                    manifest,
                    manifest_path,
                    [
                        json.dumps(
                            {
                                "batch_token": batch["batch_token"],
                                "items": [
                                    {"target_id": batch["items"][0]["target_id"], "decision": "retain"}
                                ],
                                "additions": [],
                            },
                            ensure_ascii=False,
                        )
                    ],
                )

            search = outline_runner.dispatch_command(
                "search", manifest, manifest_path, ["运输安全"]
            )
            evidence_id = search["results"][0]["evidence_id"]
            payload = {
                "batch_token": batch["batch_token"],
                "items": [
                    {
                        "target_id": batch["items"][0]["target_id"],
                        "decision": "retain",
                        "evidence_id": evidence_id,
                        "reason": "招标要求本章承接整体技术方案。",
                    }
                ],
                "additions": [
                    {
                        "node_id": "AI-0001",
                        "parent_id": batch["items"][0]["target_id"],
                        "number": "1.1",
                        "title": "海上运输安全专项方案",
                        "reason": "招标文件明确要求独立提交。",
                        "evidence_id": evidence_id,
                    }
                ],
            }
            with self.assertRaisesRegex(SystemExit, "尚未通过受控阅读"):
                outline_runner.dispatch_command(
                    "decision-batch",
                    manifest,
                    manifest_path,
                    [json.dumps(payload, ensure_ascii=False)],
                )
            outline_runner.dispatch_command(
                "section",
                manifest,
                manifest_path,
                [headings["files"][0]["items"][0]["section_id"]],
            )
            result = outline_runner.dispatch_command(
                "decision-batch",
                manifest,
                manifest_path,
                [json.dumps(payload, ensure_ascii=False)],
            )
            state = json_load(root / "outline_decision_state.json")

        self.assertEqual(result["decided_count"], 1)
        # evidence_id 与 reason 是前端展示契约：有依据就能跳转原文，reason 同时作为说明保留。
        retained = state["template_decisions"][batch["items"][0]["target_id"]]
        self.assertEqual(retained["tender_basis"]["evidence_id"], evidence_id)
        self.assertEqual(retained["reason"], "招标要求本章承接整体技术方案。")
        self.assertEqual(state["additions"][0]["tender_basis"]["evidence_id"], evidence_id)

    def test_bid_outline_appendix_candidates_keep_missing_source_status(self) -> None:
        review_workflow = load_outline_script("review_workflow")

        items = review_workflow.decision_appendix_items_from_inventory(
            {
                "schema_version": "tender-appendix-inventory.v1",
                "items": [
                    {"file_id": "TEN-1", "number": "技术附表E", "title": "技术要求", "following_table_count": 0},
                    {"file_id": "TEN-1", "number": "E.3", "title": "偏差表", "following_table_count": 1},
                    {"file_id": "TEN-1", "number": "E.4", "title": "承诺表", "following_table_count": 0},
                ],
            }
        )

        self.assertEqual([item["number"] for item in items], ["E.3", "E.4"])
        self.assertEqual([item["source_status"] for item in items], ["present", "missing"])

    def test_bid_outline_global_review_is_required_and_invalidated_by_reopen(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest, manifest_path = write_decision_context_fixture(root, heading_count=0)
            batch = outline_runner.dispatch_command("decision-next", manifest, manifest_path, [])
            outline_runner.dispatch_command(
                "decision-batch",
                manifest,
                manifest_path,
                [
                    json.dumps(
                        {
                            "batch_token": batch["batch_token"],
                            "items": [{"target_id": batch["items"][0]["target_id"], "decision": "retain"}],
                            "additions": [],
                        }
                    )
                ],
            )

            with self.assertRaisesRegex(SystemExit, "全局复核"):
                outline_runner.dispatch_command("decisions", manifest, manifest_path, [])
            reviewed = outline_runner.dispatch_command(
                "review-complete",
                manifest,
                manifest_path,
                [json.dumps({"review_summary": "已对照招标目录完成全局复核。", "issues": []}, ensure_ascii=False)],
            )
            outline_runner.dispatch_command("decisions", manifest, manifest_path, [])
            outline_runner.dispatch_command(
                "decision-reopen", manifest, manifest_path, [batch["chapter_id"]]
            )
            reopened = outline_runner.dispatch_command("decision-next", manifest, manifest_path, [])
            outline_runner.dispatch_command(
                "decision-batch",
                manifest,
                manifest_path,
                [
                    json.dumps(
                        {
                            "batch_token": reopened["batch_token"],
                            "items": [{"target_id": reopened["items"][0]["target_id"], "decision": "retain"}],
                            "additions": [],
                        }
                    )
                ],
            )
            with self.assertRaisesRegex(SystemExit, "全局复核"):
                outline_runner.dispatch_command("decisions", manifest, manifest_path, [])

        self.assertTrue(reviewed["review_digest"])

    def test_bid_outline_global_review_completes_without_forced_extra_read(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            manifest_path = root / "s2_input.json"
            template_doc = Document()
            template_doc.add_paragraph("第1章 技术方案", style="Heading 1")
            template_doc.save(template)
            tender_doc = Document()
            tender_doc.add_paragraph("1 技术要求", style="Heading 1")
            tender_doc.add_paragraph("投标人应提交完整技术方案。")
            tender_doc.save(tender)
            manifest = {
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": tender.name, "path": str(tender)}],
                "outputFile": str(root / "toc.json"),
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
            outline_runner.write_template_structure(manifest, manifest_path)
            headings = outline_runner.dispatch_command("headings", manifest, manifest_path, [])
            section_id = headings["files"][0]["items"][0]["section_id"]
            section = outline_runner.dispatch_command(
                "section", manifest, manifest_path, [section_id]
            )
            evidence_id = section["records"][1]["evidence_id"]
            batch = outline_runner.dispatch_command("decision-next", manifest, manifest_path, [])
            outline_runner.dispatch_command(
                "read", manifest, manifest_path, [evidence_id]
            )
            outline_runner.dispatch_command(
                "decision-batch",
                manifest,
                manifest_path,
                [
                    json.dumps(
                        {
                            "batch_token": batch["batch_token"],
                            "items": [
                                {
                                    "target_id": batch["items"][0]["target_id"],
                                    "decision": "retain",
                                    "evidence_id": evidence_id,
                                }
                            ],
                            "additions": [],
                        },
                        ensure_ascii=False,
                    )
                ],
            )
            payload = json.dumps(
                {"review_summary": "已从招标侧完成全局查漏。", "issues": []},
                ensure_ascii=False,
            )
            # 查漏由 opencode 自主安排，不再用阅读次数基线当门禁；只有摘要缺失才拒绝。
            with self.assertRaisesRegex(SystemExit, "review_summary is required"):
                outline_runner.dispatch_command(
                    "review-complete",
                    manifest,
                    manifest_path,
                    [json.dumps({"review_summary": "", "issues": []})],
                )
            completed = outline_runner.dispatch_command(
                "review-complete", manifest, manifest_path, [payload]
            )

        self.assertTrue(completed["review_complete"])

    def test_bid_outline_global_review_corrections_apply_omissions_without_reopening_chapter(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            manifest_path = root / "s2_input.json"
            template_doc = Document()
            template_doc.add_paragraph("第1章 技术方案", style="Heading 1")
            template_doc.save(template)
            tender_doc = Document()
            tender_doc.add_paragraph("1 专项要求", style="Heading 1")
            tender_doc.add_paragraph("投标人应独立提交项目专项承诺。")
            tender_doc.save(tender)
            manifest = {
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": tender.name, "path": str(tender)}],
                "outputFile": str(root / "toc.json"),
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
            outline_runner.write_template_structure(manifest, manifest_path)
            headings = outline_runner.dispatch_command("headings", manifest, manifest_path, [])
            section = outline_runner.dispatch_command(
                "section",
                manifest,
                manifest_path,
                [headings["files"][0]["items"][0]["section_id"]],
            )
            evidence_id = section["records"][1]["evidence_id"]
            batch = outline_runner.dispatch_command("decision-next", manifest, manifest_path, [])
            target_id = batch["items"][0]["target_id"]
            outline_runner.dispatch_command("read", manifest, manifest_path, [evidence_id])
            outline_runner.dispatch_command(
                "decision-batch",
                manifest,
                manifest_path,
                [
                    json.dumps(
                        {
                            "batch_token": batch["batch_token"],
                            "items": [
                                {
                                    "target_id": target_id,
                                    "decision": "retain",
                                    "evidence_id": evidence_id,
                                }
                            ],
                            "additions": [],
                        }
                    )
                ],
            )
            outline_runner.dispatch_command("read", manifest, manifest_path, [evidence_id])

            corrected = outline_runner.dispatch_command(
                "review-corrections",
                manifest,
                manifest_path,
                [
                    json.dumps(
                        {
                            "items": [
                                {
                                    "target_id": target_id,
                                    "decision": "suggest_delete",
                                    "reason": "全局复核发现该节点与新增专项重复。",
                                }
                            ],
                            "additions": [
                                {
                                    "node_id": "ADD-REVIEW-1",
                                    "parent_id": target_id,
                                    "number": "1.1",
                                    "title": "项目专项承诺",
                                    "reason": "全局复核发现招标要求独立提交。",
                                    "evidence_id": evidence_id,
                                }
                            ],
                        },
                        ensure_ascii=False,
                    )
                ],
            )
            with self.assertRaisesRegex(SystemExit, "全局复核"):
                outline_runner.dispatch_command("decisions", manifest, manifest_path, [])
            outline_runner.dispatch_command("read", manifest, manifest_path, [evidence_id])
            complete_outline_review(outline_runner, manifest, manifest_path)
            decisions = outline_runner.dispatch_command(
                "decisions", manifest, manifest_path, []
            )
            decision_payload = json_load(Path(decisions["decisionsFile"]))

        self.assertEqual(corrected["corrected_item_count"], 1)
        self.assertEqual(corrected["added_count"], 1)
        self.assertEqual(decision_payload["changes"][0]["node_id"], "ADD-REVIEW-1")
        self.assertEqual(
            decision_payload["template_decisions"][0],
            {
                "target_id": target_id,
                "decision": "suggest_delete",
                "reason": "全局复核发现该节点与新增专项重复。",
            },
        )

    def test_bid_outline_controlled_decision_batches_build_final_decisions(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest_path = root / "s2_input.json"
            manifest = {"workDir": str(root), "outputFile": str(root / "toc.json")}
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            (root / "template_structure.json").write_text(
                json.dumps(
                    {
                        "schema_version": "template-structure.v1",
                        "items": [
                            {"number": "Chapter 1", "title": "Overview", "level": 1},
                            {"number": "1.1", "title": "Design", "level": 2},
                            {"number": "1.1.1", "title": "Legacy case", "level": 3},
                        ],
                    }
                ),
                encoding="utf-8",
            )

            first = outline_runner.dispatch_command("decision-next", manifest, manifest_path, [])
            first_result = outline_runner.dispatch_command(
                "decision-batch",
                manifest,
                manifest_path,
                [
                    json.dumps(
                        {
                            "batch_token": first["batch_token"],
                            "items": [
                                {"target_id": item["target_id"], "decision": "retain"}
                                for item in first["items"]
                            ],
                            "additions": [],
                        }
                    )
                ],
            )
            complete_outline_review(outline_runner, manifest, manifest_path)
            finalized = outline_runner.dispatch_command(
                "decisions", manifest, manifest_path, []
            )

        # 一级章 + 二级节点各一次决策；三级 Legacy case 跟随二级父节点。
        self.assertEqual(first_result["decided_count"], 2)
        self.assertEqual(finalized["templateDecisionCount"], 2)
        self.assertEqual(finalized["remainingTemplateDecisionCount"], 0)

    def test_bid_outline_level_two_decision_cascades_to_level_three_subtree(self) -> None:
        outline_composer = load_outline_script("outline_composer")
        structure = {
            "schema_version": "template-structure.v1",
            "items": [
                {"number": "第1章", "title": "总体方案", "level": 1},
                {"number": "1.1", "title": "设计依据", "level": 2},
                {"number": "1.1.1", "title": "采用标准", "level": 3},
                {"number": "1.1.2", "title": "设计输入", "level": 3},
                {"number": "1.2", "title": "重复小节", "level": 2},
                {"number": "1.2.1", "title": "重复细项", "level": 3},
                {"number": "第2章", "title": "其他分册内容", "level": 1},
            ],
        }
        basis = {"evidence_id": "TEN-1:B000123", "file_id": "TEN-1", "search_text": "设计依据"}
        decisions = {
            "schema_version": "technical-outline-decisions.v1",
            "input_fingerprint": outline_composer.annotate_template_structure(structure)[
                "input_fingerprint"
            ],
            "template_decisions": [
                {"target_id": "TPL-0001", "decision": "retain", "reason": "承接整体技术论证。"},
                {"target_id": "TPL-0002", "decision": "retain", "tender_basis": basis},
                {"target_id": "TPL-0005", "decision": "suggest_delete", "reason": "与 1.1 语义重复。"},
                {"target_id": "TPL-0007", "decision": "suggest_delete", "reason": "属于商务分册。"},
            ],
            "changes": [],
        }

        outline, _ = outline_composer.build_composition(structure, decisions)

        chapter, other = outline["nodes"]
        design, duplicate = chapter["children"]
        # 保留的二级节点：整个三级子树跟随标签与招标依据。
        self.assertEqual([child["title"] for child in design["children"]], ["采用标准", "设计输入"])
        for child in design["children"]:
            self.assertEqual(child["suggestion_action"], "必要")
            self.assertEqual(child["tender_basis"], basis)
        # 建议删除的二级节点：子树全部建议删除，并复用父节点理由。
        self.assertEqual(duplicate["suggestion_action"], "建议删除")
        self.assertEqual(
            [
                (child["suggestion_action"], child["suggestion_reason"])
                for child in duplicate["children"]
            ],
            [("建议删除", "与 1.1 语义重复。")],
        )
        # 整章删除同样成立，且一级章自身可以独立决策。
        self.assertEqual(other["suggestion_action"], "建议删除")
        self.assertEqual(other["suggestion_reason"], "属于商务分册。")

    def test_bid_outline_additions_are_limited_to_chapter_and_section_levels(self) -> None:
        decision_workflow = load_outline_script("run_from_manifest").decision_workflow
        structure = {
            "schema_version": "template-structure.v1",
            "items": [
                {"number": "第1章", "title": "总体方案", "level": 1},
                {"number": "1.1", "title": "设计依据", "level": 2},
                {"number": "1.1.1", "title": "采用标准", "level": 3},
            ],
        }

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            batch = decision_workflow.next_decision_batch(root, structure)
            self.assertEqual(
                [item["target_id"] for item in batch["items"]], ["TPL-0001", "TPL-0002"]
            )

            def submit(additions: list[dict]) -> dict:
                return decision_workflow.submit_decision_batch(
                    root,
                    structure,
                    {
                        "batch_token": batch["batch_token"],
                        "items": [
                            {"target_id": item["target_id"], "decision": "retain"}
                            for item in batch["items"]
                        ],
                        "additions": additions,
                    },
                )

            for parent_id in ("TPL-0002", "TPL-0003"):
                with self.subTest(parent_id=parent_id), self.assertRaisesRegex(
                    SystemExit, "只能新增一级章或二级节点"
                ):
                    submit(
                        [
                            {
                                "node_id": "ADD-0001",
                                "parent_id": parent_id,
                                "number": "1.1.2",
                                "title": "专项细项",
                                "reason": "招标要求补充。",
                            }
                        ]
                    )

            result = submit(
                [
                    {
                        "node_id": "ADD-CHAPTER",
                        "parent_id": None,
                        "number": "第2章",
                        "title": "专项试验方案",
                        "reason": "招标要求独立成章。",
                    },
                    {
                        "node_id": "ADD-SECTION",
                        "parent_id": "TPL-0001",
                        "number": "1.2",
                        "title": "海上运输安全专项方案",
                        "reason": "招标要求独立提交。",
                    },
                ]
            )

        self.assertEqual(result["addition_count"], 2)

    def test_bid_outline_added_chapter_keeps_technical_appendix_last(self) -> None:
        outline_composer = load_outline_script("outline_composer")
        structure = {
            "schema_version": "template-structure.v1",
            "items": [{"number": "第1章", "title": "总体方案", "level": 1}],
        }
        decisions = {
            "schema_version": "technical-outline-decisions.v1",
            "input_fingerprint": outline_composer.annotate_template_structure(structure)[
                "input_fingerprint"
            ],
            "template_decisions": [
                {"target_id": "TPL-0001", "decision": "retain", "reason": "承接整体技术论证。"}
            ],
            "changes": [
                {
                    "operation": "add",
                    "node_id": "ADD-APPENDIX",
                    "parent_id": None,
                    "number": "附录",
                    "title": "技术附表",
                    "suggestion_action": "建议增加",
                    "suggestion_reason": "招标文件包含独立附表。",
                },
                {
                    "operation": "add",
                    "node_id": "ADD-CHAPTER",
                    "parent_id": None,
                    "number": "第2章",
                    "title": "专项试验方案",
                    "suggestion_action": "建议增加",
                    "suggestion_reason": "全局查漏发现的独立成章要求。",
                },
            ],
        }

        outline, _ = outline_composer.build_composition(structure, decisions)

        self.assertEqual(
            [node["title"] for node in outline["nodes"]],
            ["总体方案", "专项试验方案", "技术附表"],
        )

    def test_bid_outline_decision_batch_rejects_appendix_additions_outside_queue(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest_path = root / "s2_input.json"
            manifest = {"workDir": str(root), "outputFile": str(root / "toc.json")}
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            (root / "template_structure.json").write_text(
                json.dumps(
                    {
                        "schema_version": "template-structure.v1",
                        "items": [{"number": "第1章", "title": "技术方案", "level": 1}],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            (root / "tender_appendix_inventory.json").write_text(
                json.dumps(
                    {
                        "schema_version": "tender-appendix-inventory.v1",
                        "items": [
                            {
                                "file_id": "TEN-1",
                                "file_name": "招标文件.docx",
                                "number": "附表A.1",
                                "title": "投标机型总方案信息表",
                                "raw_text": "附表A.1 投标机型总方案信息表",
                                "following_table_count": 1,
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            batch = outline_runner.dispatch_command(
                "decision-next", manifest, manifest_path, []
            )
            with self.assertRaisesRegex(
                SystemExit, "技术附表必须通过 appendix-decision-batch 决策"
            ):
                outline_runner.dispatch_command(
                    "decision-batch",
                    manifest,
                    manifest_path,
                    [
                        json.dumps(
                            {
                                "batch_token": batch["batch_token"],
                                "items": [
                                    {
                                        "target_id": batch["items"][0]["target_id"],
                                        "decision": "retain",
                                    }
                                ],
                                "additions": [
                                    {
                                        "node_id": "ADD-TECH-APPENDIX",
                                        "parent_id": None,
                                        "number": "第2章",
                                        "title": "技术附表",
                                        "reason": "招标包含实际技术附表。",
                                    },
                                    {
                                        "node_id": "ADD-APPENDIX-A1",
                                        "parent_id": "ADD-TECH-APPENDIX",
                                        "appendix_id": "APP-0001",
                                        "reason": "招标结构化清单中的实际表单。",
                                    },
                                ],
                            },
                            ensure_ascii=False,
                        )
                    ],
                )
            state = json_load(root / "outline_decision_state.json")

        self.assertEqual(state["additions"], [])
        self.assertEqual(state["appendix_decisions"], {})
        self.assertEqual(state["active_batch"]["target_ids"], ["TPL-0001"])

    def test_bid_outline_appendix_batches_require_explicit_include_or_exclude(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest, manifest_path = write_decision_context_fixture(root, heading_count=0)
            template_batch = outline_runner.dispatch_command(
                "decision-next", manifest, manifest_path, []
            )
            with self.assertRaisesRegex(SystemExit, "先完成模板"):
                outline_runner.dispatch_command("appendix-next", manifest, manifest_path, [])
            outline_runner.dispatch_command(
                "decision-batch",
                manifest,
                manifest_path,
                [
                    json.dumps(
                        {
                            "batch_token": template_batch["batch_token"],
                            "items": [
                                {
                                    "target_id": template_batch["items"][0]["target_id"],
                                    "decision": "retain",
                                }
                            ],
                            "additions": [],
                        }
                    )
                ],
            )
            inventory_path = root / "tender_appendix_inventory.json"
            inventory_path.write_text(
                json.dumps(
                    {
                        "schema_version": "tender-appendix-inventory.v1",
                        "items": [
                            {
                                "file_id": "TEN-1",
                                "number": f"Appendix A.{index}",
                                "title": f"Controlled form {index}",
                                "following_table_count": 0 if index == 2 else 1,
                            }
                            for index in range(1, 4)
                        ],
                    }
                ),
                encoding="utf-8",
            )

            for invalid_max_items in (0, 41):
                with self.subTest(max_items=invalid_max_items):
                    with self.assertRaisesRegex(SystemExit, "between 1 and 40"):
                        outline_runner.dispatch_command(
                            "appendix-next",
                            manifest,
                            manifest_path,
                            ["--max-items", str(invalid_max_items)],
                        )

            for invalid_args in (
                ["--unknown"],
                ["--max-items"],
                ["--max-items", "2", "extra"],
            ):
                with self.subTest(invalid_args=invalid_args):
                    with self.assertRaisesRegex(SystemExit, "appendix-next usage"):
                        outline_runner.dispatch_command(
                            "appendix-next", manifest, manifest_path, invalid_args
                        )

            batch = outline_runner.dispatch_command(
                "appendix-next", manifest, manifest_path, ["--max-items", "2"]
            )
            self.assertEqual(
                [item["appendix_id"] for item in batch["items"]],
                ["APP-0001", "APP-0002"],
            )
            self.assertEqual(batch["decided_count"], 0)
            self.assertEqual(batch["remaining_count"], 3)
            self.assertFalse(batch["complete"])
            self.assertEqual(
                batch["submission_contract"],
                {
                    "items_must_match_batch": True,
                    "items_must_keep_returned_order": True,
                    "exclude_fields": ["appendix_id", "decision", "reason"],
                    "include_fields": [
                        "appendix_id",
                        "decision",
                        "node_id",
                        "parent_id",
                        "reason",
                    ],
                    "include_parent_id": "必须引用本批 root_addition.node_id 或已有唯一技术附表根节点",
                    "reason_required": "include 与 exclude 都必须提交 reason",
                    "missing_rule": "source_status=missing 必须 exclude；只有 source_status=present 才自主判断 include 或 exclude",
                    "root_addition": {
                        "required_when": "首次 include 且尚无唯一的技术附表根节点",
                        "omit_when": "根节点已建立后的所有后续批次禁止再提交 root_addition",
                        "fields": ["node_id", "reason"],
                        "generated_fields": {
                            "parent_id": None,
                            "number": "附录",
                            "title": "技术附表",
                        },
                    },
                },
            )

            with self.assertRaisesRegex(
                SystemExit, "appendix-decision-batch requires exactly one JSON payload"
            ):
                outline_runner.dispatch_command(
                    "appendix-decision-batch",
                    manifest,
                    manifest_path,
                    [
                        json.dumps(
                            {
                                "batch_token": batch["batch_token"],
                                "items": [
                                    {
                                        "appendix_id": item["appendix_id"],
                                        "decision": "exclude",
                                        "reason": "Not required.",
                                    }
                                    for item in batch["items"]
                                ],
                            }
                        ),
                        "extra",
                    ],
                )

            invalid_item_lists = [
                [
                    {"appendix_id": "APP-0001", "decision": "exclude", "reason": "no"},
                    {"appendix_id": "APP-0001", "decision": "exclude", "reason": "no"},
                ],
                [
                    {"appendix_id": "APP-0001", "decision": "exclude", "reason": "no"},
                    {"appendix_id": "APP-9999", "decision": "exclude", "reason": "no"},
                ],
                [
                    {"appendix_id": "APP-0002", "decision": "exclude", "reason": "no"},
                    {"appendix_id": "APP-0001", "decision": "exclude", "reason": "no"},
                ],
                [
                    {"appendix_id": "APP-0001", "decision": "exclude", "reason": "no"},
                ],
            ]
            for invalid_items in invalid_item_lists:
                with self.subTest(invalid_items=invalid_items):
                    with self.assertRaisesRegex(SystemExit, "exactly match"):
                        outline_runner.dispatch_command(
                            "appendix-decision-batch",
                            manifest,
                            manifest_path,
                            [json.dumps({"batch_token": batch["batch_token"], "items": invalid_items})],
                        )

            with self.assertRaisesRegex(SystemExit, "source_status=missing.*must be exclude"):
                outline_runner.dispatch_command(
                    "appendix-decision-batch",
                    manifest,
                    manifest_path,
                    [
                        json.dumps(
                            {
                                "batch_token": batch["batch_token"],
                                "items": [
                                    {
                                        "appendix_id": "APP-0001",
                                        "decision": "exclude",
                                        "reason": "Not required.",
                                    },
                                    {
                                        "appendix_id": "APP-0002",
                                        "decision": "include",
                                        "node_id": "ADD-APP-2",
                                        "parent_id": "ADD-APP-ROOT",
                                        "reason": "Include it.",
                                    },
                                ],
                            }
                        )
                    ],
                )
            state_after_missing = json_load(root / "outline_decision_state.json")
            self.assertEqual(state_after_missing["appendix_decisions"], {})
            self.assertEqual(
                state_after_missing["active_appendix_batch"]["appendix_ids"],
                ["APP-0001", "APP-0002"],
            )

            outline_runner.dispatch_command(
                "appendix-decision-batch",
                manifest,
                manifest_path,
                [
                    json.dumps(
                        {
                            "batch_token": batch["batch_token"],
                            "items": [
                                {
                                    "appendix_id": item["appendix_id"],
                                    "decision": "exclude",
                                    "reason": "Not required by this tender.",
                                }
                                for item in batch["items"]
                            ],
                        }
                    )
                ],
            )
            second = outline_runner.dispatch_command(
                "appendix-next", manifest, manifest_path, ["--max-items", "2"]
            )
            self.assertEqual(
                [item["appendix_id"] for item in second["items"]], ["APP-0003"]
            )

            inventory = json_load(inventory_path)
            inventory["items"][2]["title"] = "Changed after token issuance"
            inventory_path.write_text(json.dumps(inventory), encoding="utf-8")
            with self.assertRaisesRegex(SystemExit, "inventory.*changed"):
                outline_runner.dispatch_command("appendix-next", manifest, manifest_path, [])
            with self.assertRaisesRegex(SystemExit, "inventory.*changed"):
                outline_runner.dispatch_command(
                    "appendix-decision-batch",
                    manifest,
                    manifest_path,
                    [
                        json.dumps(
                            {
                                "batch_token": second["batch_token"],
                                "items": [
                                    {
                                        "appendix_id": "APP-0003",
                                        "decision": "exclude",
                                        "reason": "Not required.",
                                    }
                                ],
                            }
                        )
                    ],
                )
            with self.assertRaisesRegex(SystemExit, "inventory.*changed"):
                outline_runner.dispatch_command("decisions", manifest, manifest_path, [])

    def test_bid_outline_chapter_decisions_are_isolated_and_merge_completely(self) -> None:
        decision_workflow = load_outline_script("run_from_manifest").decision_workflow
        structure = {
            "schema_version": "template-structure.v1",
            "items": [
                {"number": "1", "title": "总体方案", "level": 1},
                {"number": "1.1", "title": "实施组织", "level": 2},
                {"number": "2", "title": "质量安全", "level": 1},
                {"number": "2.1", "title": "质量保证", "level": 2},
            ],
        }

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            chapter_catalog = decision_workflow.decision_chapters(structure)
            chapter_dirs = {
                item["chapter_id"]: root / item["chapter_id"]
                for item in chapter_catalog["chapters"]
            }
            chapter_titles = {}
            chapter_parent_ids = {}
            for chapter_id, chapter_dir in chapter_dirs.items():
                chapter_dir.mkdir()
                while True:
                    batch = decision_workflow.next_decision_batch(
                        chapter_dir,
                        structure,
                        chapter_id=chapter_id,
                    )
                    if batch["complete"]:
                        break
                    chapter_titles[chapter_id] = tuple(item["title"] for item in batch["items"])
                    chapter_parent_ids[chapter_id] = batch["items"][0]["target_id"]
                    other_parent_id = next(
                        item["template_id"]
                        for item in decision_workflow._annotated_items(structure)[1]
                        if item["template_id"] not in {
                            batch_item["target_id"] for batch_item in batch["items"]
                        }
                    )
                    with self.assertRaisesRegex(SystemExit, "outside the active chapter"):
                        decision_workflow.submit_decision_batch(
                            chapter_dir,
                            structure,
                            {
                                "batch_token": batch["batch_token"],
                                "items": [
                                    {
                                        "target_id": item["target_id"],
                                        "decision": "retain",
                                        "reason": "历史模板专家经验仍适用。",
                                    }
                                    for item in batch["items"]
                                ],
                                "additions": [
                                    {
                                        "node_id": "ADD-0001",
                                        "parent_id": other_parent_id,
                                        "number": "1.9",
                                        "title": "跨章新增",
                                        "reason": "不允许跨章挂载。",
                                    }
                                ],
                            },
                            chapter_id=chapter_id,
                        )
                    with self.assertRaisesRegex(SystemExit, "must stay under the active chapter"):
                        decision_workflow.submit_decision_batch(
                            chapter_dir,
                            structure,
                            {
                                "batch_token": batch["batch_token"],
                                "items": [
                                    {
                                        "target_id": item["target_id"],
                                        "decision": "retain",
                                        "reason": "历史模板专家经验仍适用。",
                                    }
                                    for item in batch["items"]
                                ],
                                "additions": [
                                    {
                                        "node_id": "ADD-ROOT",
                                        "parent_id": None,
                                        "number": "9",
                                        "title": "跨章一级新增",
                                        "reason": "章节会话不得创建无归属的一级章。",
                                    }
                                ],
                            },
                            chapter_id=chapter_id,
                        )
                    decision_workflow.submit_decision_batch(
                        chapter_dir,
                        structure,
                        {
                            "batch_token": batch["batch_token"],
                            "items": [
                                {
                                    "target_id": item["target_id"],
                                    "decision": "retain",
                                    "reason": "历史模板专家经验仍适用。",
                                }
                                for item in batch["items"]
                            ],
                            "additions": [
                                {
                                    "node_id": "ADD-0001",
                                    "parent_id": batch["items"][0]["target_id"],
                                    "number": "1.9",
                                    "title": "本章专项响应",
                                    "reason": "本章需要独立表达。",
                                }
                            ],
                        },
                        chapter_id=chapter_id,
                    )

            merged = decision_workflow.merge_chapter_decisions(
                root,
                structure,
                chapter_dirs,
            )
            state = json_load(root / "outline_decision_state.json")

        self.assertEqual(chapter_catalog["chapter_count"], 2)
        self.assertEqual(set(chapter_titles.values()), {
            ("总体方案", "实施组织"),
            ("质量安全", "质量保证"),
        })
        self.assertTrue(merged["complete"])
        self.assertEqual(merged["decided_count"], 4)
        self.assertEqual(merged["addition_count"], 2)
        self.assertEqual(len(state["template_decisions"]), 4)
        self.assertEqual(len({item["node_id"] for item in state["additions"]}), 2)
        for addition in state["additions"]:
            owner = state["addition_chapters"][addition["node_id"]]
            self.assertEqual(addition["parent_id"], chapter_parent_ids[owner])

    def test_bid_outline_sparse_chapter_unit_switches_to_contiguous_reading_mode(self) -> None:
        decision_workflow = load_outline_script("run_from_manifest").decision_workflow
        structure = {
            "schema_version": "template-structure.v1",
            "items": [
                {"number": "1", "title": "总体方案", "level": 1},
                {"number": "1.1", "title": "实施组织", "level": 2},
                {"number": "1.2", "title": "进度计划", "level": 2},
                {"number": "2", "title": "技术规范响应", "level": 1},
            ],
        }

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            normal = decision_workflow.next_decision_batch(root, structure)
            self.assertNotIn("authoring_mode", normal)
            self.assertFalse(any("连续通读" in step for step in normal["decision_steps"]))
            decision_workflow.submit_decision_batch(
                root,
                structure,
                {
                    "batch_token": normal["batch_token"],
                    "items": [
                        {
                            "target_id": item["target_id"],
                            "decision": "retain",
                            "reason": "历史模板专家经验仍适用。",
                        }
                        for item in normal["items"]
                    ],
                    "additions": [],
                },
            )
            sparse = decision_workflow.next_decision_batch(root, structure)
            replay = decision_workflow.next_decision_batch(root, structure)

        # 有二级节点的章走默认判定流程；空章切换为"圈定上级章节 + 大页连续通读"纪律
        self.assertEqual(len(sparse["items"]), 1)
        self.assertEqual(sparse["authoring_mode"], "sparse_chapter")
        self.assertTrue(any("连续通读" in step for step in sparse["decision_steps"]))
        self.assertTrue(any("最多 2 次" in step for step in sparse["decision_steps"]))
        self.assertTrue(any("--max-chars 30000" in step for step in sparse["decision_steps"]))
        # 未提交前重复 decision-next 返回同一活动批次，模式标记保持
        self.assertEqual(replay["batch_token"], sparse["batch_token"])
        self.assertEqual(replay["authoring_mode"], "sparse_chapter")

    def test_bid_outline_appendix_decision_progress_is_read_only(self) -> None:
        decision_workflow = load_outline_script("run_from_manifest").decision_workflow
        structure = {
            "schema_version": "template-structure.v1",
            "items": [{"number": "1", "title": "Technical proposal", "level": 1}],
        }
        inventory = [
            {
                "appendix_id": "APP-0001",
                "file_id": "TEN-1",
                "number": "Appendix B.1",
                "title": "Guaranteed data sheet",
                "raw_text": "Appendix B.1Guaranteed data sheet",
                "following_table_count": 1,
            },
        ]

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = decision_workflow.next_decision_batch(root, structure)
            decision_workflow.submit_decision_batch(
                root,
                structure,
                {
                    "batch_token": template["batch_token"],
                    "items": [
                        {
                            "target_id": template["items"][0]["target_id"],
                            "decision": "retain",
                            "reason": "历史模板专家经验仍适用。",
                        }
                    ],
                    "additions": [],
                },
            )
            pending = decision_workflow.appendix_decision_progress(
                root, structure, inventory
            )
            state_before_batch = json_load(root / "outline_decision_state.json")
            batch = decision_workflow.next_appendix_batch(root, structure, inventory)
            in_flight = decision_workflow.appendix_decision_progress(
                root, structure, inventory
            )
            decision_workflow.submit_appendix_batch(
                root,
                structure,
                {
                    "batch_token": batch["batch_token"],
                    "items": [
                        {
                            "appendix_id": "APP-0001",
                            "decision": "include",
                            "node_id": "ADD-B1",
                            "parent_id": "ADD-TECH-APPENDIX",
                            "reason": "Tender requires the completed form.",
                        }
                    ],
                    "root_addition": {
                        "node_id": "ADD-TECH-APPENDIX",
                        "reason": "Tender contains controlled appendices.",
                    },
                },
                inventory,
            )
            done = decision_workflow.appendix_decision_progress(
                root, structure, inventory
            )

        self.assertEqual(
            pending,
            {
                "decidedCount": 0,
                "remainingCount": 1,
                "activeBatch": False,
                "complete": False,
            },
        )
        # 进度查询不得发放批次令牌或改动状态
        self.assertEqual(
            state_before_batch["active_appendix_batch"],
            {"token": "", "appendix_ids": []},
        )
        self.assertFalse(in_flight["complete"])
        self.assertTrue(in_flight["activeBatch"])
        self.assertTrue(done["complete"])
        self.assertEqual(done["decidedCount"], 1)
        self.assertEqual(done["remainingCount"], 0)

    def test_bid_outline_appendix_batch_copies_inventory_metadata_for_include(self) -> None:
        decision_workflow = load_outline_script("run_from_manifest").decision_workflow
        structure = {
            "schema_version": "template-structure.v1",
            "items": [{"number": "1", "title": "Technical proposal", "level": 1}],
        }
        inventory = [
            {
                "appendix_id": "APP-0001",
                "file_id": "TEN-1",
                "number": "Appendix B.1",
                "title": "Guaranteed data sheet",
                "raw_text": "Appendix B.1Guaranteed data sheet",
                "following_table_count": 1,
            },
            {
                "appendix_id": "APP-0002",
                "file_id": "TEN-1",
                "number": "Appendix B.2",
                "title": "Reference only",
                "following_table_count": 0,
            },
        ]

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = decision_workflow.next_decision_batch(
                root, structure
            )
            initial_state = json_load(root / "outline_decision_state.json")
            with self.assertRaisesRegex(SystemExit, "node_id is missing or duplicate"):
                decision_workflow.submit_decision_batch(
                    root,
                    structure,
                    {
                        "batch_token": template["batch_token"],
                        "items": [
                            {
                                "target_id": template["items"][0]["target_id"],
                                "decision": "retain",
                            }
                        ],
                        "additions": [
                            {
                                "node_id": "TPL-0001",
                                "parent_id": None,
                                "number": "2",
                                "title": "技术附表",
                                "reason": "Conflicts with the template node ID.",
                            }
                        ],
                    },
                )
            self.assertEqual(json_load(root / "outline_decision_state.json"), initial_state)

            decision_workflow.submit_decision_batch(
                root,
                structure,
                {
                    "batch_token": template["batch_token"],
                    "items": [
                        {
                            "target_id": template["items"][0]["target_id"],
                            "decision": "retain",
                        }
                    ],
                    "additions": [
                        {
                            "node_id": "ADD-TECH-APPENDIX",
                            "parent_id": None,
                            "number": "2",
                            "title": "技术附表",
                            "reason": "Tender contains controlled appendices.",
                        }
                    ],
                },
            )
            with self.assertRaisesRegex(SystemExit, "duplicate"):
                decision_workflow.next_appendix_batch(
                    root, structure, [inventory[0], inventory[0]]
                )
            batch = decision_workflow.next_appendix_batch(root, structure, inventory)

            invalid_payloads = [
                [
                    {"appendix_id": "APP-0001", "decision": "include", "parent_id": "ROOT", "reason": "Required."},
                    {"appendix_id": "APP-0002", "decision": "exclude", "reason": "No."},
                ],
                [
                    {"appendix_id": "APP-0001", "decision": "include", "node_id": "ADD-B1", "reason": "Required."},
                    {"appendix_id": "APP-0002", "decision": "exclude", "reason": "No."},
                ],
                [
                    {"appendix_id": "APP-0001", "decision": "include", "node_id": "ADD-B1", "parent_id": "ROOT"},
                    {"appendix_id": "APP-0002", "decision": "exclude", "reason": "No."},
                ],
                [
                    {"appendix_id": "APP-0001", "decision": "include", "node_id": "ADD-B1", "parent_id": "ADD-TECH-APPENDIX", "reason": "Required."},
                    {"appendix_id": "APP-0002", "decision": "exclude"},
                ],
            ]
            for invalid_items in invalid_payloads:
                with self.subTest(invalid_items=invalid_items):
                    with self.assertRaisesRegex(SystemExit, "required"):
                        decision_workflow.submit_appendix_batch(
                            root,
                            structure,
                            {"batch_token": batch["batch_token"], "items": invalid_items},
                            inventory,
                        )

            decision_workflow.submit_appendix_batch(
                root,
                structure,
                {
                    "batch_token": batch["batch_token"],
                    "items": [
                        {
                            "appendix_id": "APP-0001",
                            "decision": "include",
                            "node_id": "ADD-B1",
                            "parent_id": "ADD-TECH-APPENDIX",
                            "reason": "Tender requires the completed form.",
                        },
                        {
                            "appendix_id": "APP-0002",
                            "decision": "exclude",
                            "reason": "Heading only; no independent form.",
                        },
                    ],
                },
                inventory,
            )
            state = json_load(root / "outline_decision_state.json")
            finalized = decision_workflow.finalize_decisions(
                root, structure, appendix_items=inventory
            )

        self.assertEqual(state["schema_version"], decision_workflow.STATE_SCHEMA)
        self.assertEqual(state["appendix_decisions"]["APP-0001"]["decision"], "include")
        self.assertEqual(state["appendix_decisions"]["APP-0001"]["reason"], "Tender requires the completed form.")
        self.assertEqual(state["appendix_decisions"]["APP-0002"]["decision"], "exclude")
        self.assertEqual(len(state["additions"]), 2)
        self.assertEqual(state["additions"][1]["number"], "Appendix B.1")
        self.assertEqual(state["additions"][1]["title"], "Guaranteed data sheet")
        self.assertEqual(
            state["additions"][1]["tender_basis"]["search_text"],
            "Appendix B.1Guaranteed data sheet",
        )
        self.assertTrue(finalized["decisionsFile"])

    def test_bid_outline_appendix_include_requires_unique_valid_root_parent_atomically(self) -> None:
        decision_workflow = load_outline_script("run_from_manifest").decision_workflow
        structure = {
            "schema_version": "template-structure.v1",
            "items": [{"number": "1", "title": "Technical proposal", "level": 1}],
        }
        inventory = [
            {
                "appendix_id": f"APP-{index:04d}",
                "file_id": "TEN-1",
                "number": f"Appendix D.{index}",
                "title": f"Controlled form {index}",
                "following_table_count": 1,
            }
            for index in range(1, 3)
        ]

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = decision_workflow.next_decision_batch(
                root, structure
            )
            decision_workflow.submit_decision_batch(
                root,
                structure,
                {
                    "batch_token": template["batch_token"],
                    "items": [
                        {
                            "target_id": template["items"][0]["target_id"],
                            "decision": "retain",
                        }
                    ],
                    "additions": [
                        {
                            "node_id": "ADD-TECH-APPENDIX",
                            "parent_id": None,
                            "number": "2",
                            "title": "技术附表",
                            "reason": "Valid root.",
                        },
                        {
                            "node_id": "ADD-NESTED-APPENDIX",
                            "parent_id": "TPL-0001",
                            "number": "1.1",
                            "title": "技术附表",
                            "reason": "Nested candidate.",
                        },
                        {
                            "node_id": "ADD-WRONG-TITLE",
                            "parent_id": None,
                            "number": "3",
                            "title": "技术附件",
                            "reason": "Wrong title candidate.",
                        },
                    ],
                },
            )
            batch = decision_workflow.next_appendix_batch(root, structure, inventory)
            original_state = json_load(root / "outline_decision_state.json")

            def include_payload(node_ids: list[str], parent_id: str) -> dict:
                return {
                    "batch_token": batch["batch_token"],
                    "items": [
                        {
                            "appendix_id": item["appendix_id"],
                            "decision": "include",
                            "node_id": node_ids[index],
                            "parent_id": parent_id,
                            "reason": "Required controlled form.",
                        }
                        for index, item in enumerate(batch["items"])
                    ],
                }

            invalid_root_states = []
            deleted_root_state = json.loads(json.dumps(original_state))
            deleted_root_state["additions"].append(
                {
                    "operation": "suggest_delete",
                    "target_id": "ADD-TECH-APPENDIX",
                }
            )
            invalid_root_states.append(("deleted root", deleted_root_state))
            duplicate_root_state = json.loads(json.dumps(original_state))
            duplicate_root_state["additions"].append(
                {
                    "operation": "add",
                    "node_id": "ADD-SECOND-TECH-APPENDIX",
                    "parent_id": None,
                    "number": "4",
                    "title": "技术附表",
                    "suggestion_action": "建议增加",
                    "suggestion_reason": "Ambiguous root.",
                }
            )
            invalid_root_states.append(("duplicate roots", duplicate_root_state))
            for label, invalid_state in invalid_root_states:
                with self.subTest(label=label):
                    (root / "outline_decision_state.json").write_text(
                        json.dumps(invalid_state, ensure_ascii=False), encoding="utf-8"
                    )
                    with self.assertRaisesRegex(SystemExit, "parent_id|root_addition"):
                        decision_workflow.submit_appendix_batch(
                            root,
                            structure,
                            include_payload(
                                ["ADD-VALID-1", "ADD-VALID-2"],
                                "ADD-TECH-APPENDIX",
                            ),
                            inventory,
                        )
                    self.assertEqual(
                        json_load(root / "outline_decision_state.json"), invalid_state
                    )
            (root / "outline_decision_state.json").write_text(
                json.dumps(original_state, ensure_ascii=False), encoding="utf-8"
            )

            invalid_cases = [
                ("template node", ["TPL-0001", "ADD-VALID-2"], "ADD-TECH-APPENDIX"),
                ("existing addition", ["ADD-TECH-APPENDIX", "ADD-VALID-2"], "ADD-TECH-APPENDIX"),
                ("same batch", ["ADD-DUPLICATE", "ADD-DUPLICATE"], "ADD-TECH-APPENDIX"),
                ("missing parent", ["ADD-VALID-1", "ADD-VALID-2"], "ADD-MISSING"),
                ("nested parent", ["ADD-VALID-1", "ADD-VALID-2"], "ADD-NESTED-APPENDIX"),
                ("wrong title", ["ADD-VALID-1", "ADD-VALID-2"], "ADD-WRONG-TITLE"),
            ]
            for label, node_ids, parent_id in invalid_cases:
                with self.subTest(label=label):
                    with self.assertRaisesRegex(SystemExit, "node_id|parent_id"):
                        decision_workflow.submit_appendix_batch(
                            root,
                            structure,
                            include_payload(node_ids, parent_id),
                            inventory,
                        )
                    self.assertEqual(
                        json_load(root / "outline_decision_state.json"), original_state
                    )

            decision_workflow.submit_appendix_batch(
                root,
                structure,
                {
                    "batch_token": batch["batch_token"],
                    "items": [
                        {
                            "appendix_id": item["appendix_id"],
                            "decision": "include",
                            "node_id": f"ADD-VALID-{index}",
                            "parent_id": "ADD-TECH-APPENDIX",
                            "reason": "Required controlled form.",
                        }
                        for index, item in enumerate(batch["items"], start=1)
                    ],
                },
                inventory,
            )
            final_state = json_load(root / "outline_decision_state.json")
            decision_workflow.finalize_decisions(
                root, structure, appendix_items=inventory
            )

        self.assertEqual(final_state["active_appendix_batch"]["appendix_ids"], [])
        self.assertEqual(set(final_state["appendix_decisions"]), {"APP-0001", "APP-0002"})

    def test_bid_outline_appendix_first_include_creates_root_and_later_batch_reuses_it(self) -> None:
        decision_workflow = load_outline_script("run_from_manifest").decision_workflow
        structure = {
            "schema_version": "template-structure.v1",
            "items": [{"number": "1", "title": "Technical proposal", "level": 1}],
        }
        inventory = [
            {
                "appendix_id": f"APP-{index:04d}",
                "file_id": "TEN-1",
                "number": f"Appendix R.{index}",
                "title": f"Controlled form {index}",
                "following_table_count": 1,
            }
            for index in range(1, 3)
        ]

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = decision_workflow.next_decision_batch(
                root, structure
            )
            decision_workflow.submit_decision_batch(
                root,
                structure,
                {
                    "batch_token": template["batch_token"],
                    "items": [
                        {
                            "target_id": template["items"][0]["target_id"],
                            "decision": "retain",
                        }
                    ],
                },
            )

            first = decision_workflow.next_appendix_batch(
                root, structure, inventory, max_items=1
            )
            decision_workflow.submit_appendix_batch(
                root,
                structure,
                {
                    "batch_token": first["batch_token"],
                    "root_addition": {
                        "node_id": "ADD-TECH-APPENDIX",
                        "parent_id": None,
                        "number": "第7章",
                        "title": "技术附表",
                        "reason": "Tender contains controlled appendices.",
                    },
                    "items": [
                        {
                            "appendix_id": "APP-0001",
                            "decision": "include",
                            "node_id": "ADD-APP-0001",
                            "parent_id": "ADD-TECH-APPENDIX",
                            "reason": "Include first form.",
                        }
                    ],
                },
                inventory,
            )
            first_state = json_load(root / "outline_decision_state.json")
            second = decision_workflow.next_appendix_batch(
                root, structure, inventory, max_items=1
            )
            second_state = json_load(root / "outline_decision_state.json")

            with self.assertRaisesRegex(SystemExit, "root_addition"):
                decision_workflow.submit_appendix_batch(
                    root,
                    structure,
                    {
                        "batch_token": second["batch_token"],
                        "root_addition": {
                            "node_id": "ADD-SECOND-ROOT",
                            "parent_id": None,
                            "number": "第8章",
                            "title": "技术附表",
                            "reason": "Duplicate root.",
                        },
                        "items": [
                            {
                                "appendix_id": "APP-0002",
                                "decision": "include",
                                "node_id": "ADD-APP-0002",
                                "parent_id": "ADD-SECOND-ROOT",
                                "reason": "Include second form.",
                            }
                        ],
                    },
                    inventory,
                )
            self.assertEqual(
                json_load(root / "outline_decision_state.json"), second_state
            )

            decision_workflow.submit_appendix_batch(
                root,
                structure,
                {
                    "batch_token": second["batch_token"],
                    "items": [
                        {
                            "appendix_id": "APP-0002",
                            "decision": "include",
                            "node_id": "ADD-APP-0002",
                            "parent_id": "ADD-TECH-APPENDIX",
                            "reason": "Include second form.",
                        }
                    ],
                },
                inventory,
            )
            finalized = decision_workflow.finalize_decisions(
                root, structure, appendix_items=inventory
            )
            decisions = json_load(Path(finalized["decisionsFile"]))
            outline, _ = decision_workflow.outline_composer.build_composition(
                structure, decisions
            )
            validation = decision_workflow.validate_finalized_decisions(
                root,
                structure,
                decisions,
                workflow_binding={},
                appendix_items=inventory,
                include_appendix_decisions=True,
            )

        self.assertEqual(
            first_state["additions"][0],
            {
                "operation": "add",
                "node_id": "ADD-TECH-APPENDIX",
                "parent_id": None,
                "number": "附录",
                "title": "技术附表",
                "suggestion_action": "建议增加",
                "suggestion_reason": "Tender contains controlled appendices.",
            },
        )
        self.assertEqual(
            outline["nodes"][-1],
            {
                "number": "附录",
                "title": "技术附表",
                "suggestion_action": "建议增加",
                "suggestion_reason": "Tender contains controlled appendices.",
                "children": [
                    {
                        "number": "Appendix R.1",
                        "title": "Controlled form 1",
                        "suggestion_action": "建议增加",
                        "suggestion_reason": "Include first form.",
                        "tender_basis": {
                            "file_id": "TEN-1",
                            "search_text": "Appendix R.1 Controlled form 1",
                        },
                        "children": [],
                    },
                    {
                        "number": "Appendix R.2",
                        "title": "Controlled form 2",
                        "suggestion_action": "建议增加",
                        "suggestion_reason": "Include second form.",
                        "tender_basis": {
                            "file_id": "TEN-1",
                            "search_text": "Appendix R.2 Controlled form 2",
                        },
                        "children": [],
                    },
                ],
            },
        )
        self.assertEqual(
            [item["decision"] for item in validation["appendixDecisions"]],
            ["include", "include"],
        )

    def test_bid_outline_appendix_root_addition_validation_is_atomic(self) -> None:
        decision_workflow = load_outline_script("run_from_manifest").decision_workflow
        structure = {
            "schema_version": "template-structure.v1",
            "items": [{"number": "1", "title": "Technical proposal", "level": 1}],
        }
        inventory = [
            {
                "appendix_id": "APP-0001",
                "file_id": "TEN-1",
                "number": "Appendix V.1",
                "title": "Validated form",
                "following_table_count": 1,
            }
        ]

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = decision_workflow.next_decision_batch(
                root, structure
            )
            decision_workflow.submit_decision_batch(
                root,
                structure,
                {
                    "batch_token": template["batch_token"],
                    "items": [
                        {
                            "target_id": template["items"][0]["target_id"],
                            "decision": "retain",
                        }
                    ],
                    "additions": [
                        {
                            "node_id": "ADD-EXISTING",
                            "parent_id": None,
                            "number": "2",
                            "title": "Other section",
                            "reason": "Existing non-appendix addition.",
                        }
                    ],
                },
            )
            batch = decision_workflow.next_appendix_batch(root, structure, inventory)
            original_state = json_load(root / "outline_decision_state.json")

            include_item = {
                "appendix_id": "APP-0001",
                "decision": "include",
                "node_id": "ADD-APP-0001",
                "parent_id": "ADD-TECH-APPENDIX",
                "reason": "Required form.",
            }
            valid_root = {
                "node_id": "ADD-TECH-APPENDIX",
                "parent_id": None,
                "number": "第7章",
                "title": "技术附表",
                "reason": "Required appendix root.",
            }
            invalid_payloads = [
                ("root_addition", {"batch_token": batch["batch_token"], "items": [include_item]}),
                (
                    "reason",
                    {
                        "batch_token": batch["batch_token"],
                        "root_addition": {key: value for key, value in valid_root.items() if key != "reason"},
                        "items": [include_item],
                    },
                ),
                (
                    "template node",
                    {
                        "batch_token": batch["batch_token"],
                        "root_addition": {**valid_root, "node_id": "TPL-0001"},
                        "items": [include_item],
                    },
                ),
                (
                    "existing addition",
                    {
                        "batch_token": batch["batch_token"],
                        "root_addition": {**valid_root, "node_id": "ADD-EXISTING"},
                        "items": [include_item],
                    },
                ),
                (
                    "parent_id",
                    {
                        "batch_token": batch["batch_token"],
                        "root_addition": {**valid_root, "parent_id": "TPL-0001"},
                        "items": [include_item],
                    },
                ),
                (
                    "number",
                    {
                        "batch_token": batch["batch_token"],
                        "root_addition": {**valid_root, "number": ""},
                        "items": [include_item],
                    },
                ),
                (
                    "title",
                    {
                        "batch_token": batch["batch_token"],
                        "root_addition": {**valid_root, "title": "技术附件"},
                        "items": [include_item],
                    },
                ),
                (
                    "parent_id",
                    {
                        "batch_token": batch["batch_token"],
                        "root_addition": valid_root,
                        "items": [{**include_item, "parent_id": "ADD-WRONG-ROOT"}],
                    },
                ),
                (
                    "node_id",
                    {
                        "batch_token": batch["batch_token"],
                        "root_addition": valid_root,
                        "items": [{**include_item, "node_id": "ADD-TECH-APPENDIX"}],
                    },
                ),
                (
                    "root_addition",
                    {
                        "batch_token": batch["batch_token"],
                        "root_addition": valid_root,
                        "items": [
                            {
                                "appendix_id": "APP-0001",
                                "decision": "exclude",
                                "reason": "Not required.",
                            }
                        ],
                    },
                ),
            ]
            for message, payload in invalid_payloads:
                with self.subTest(message=message):
                    with self.assertRaisesRegex(SystemExit, message):
                        decision_workflow.submit_appendix_batch(
                            root, structure, payload, inventory
                        )
                    self.assertEqual(
                        json_load(root / "outline_decision_state.json"), original_state
                    )

            decision_workflow.submit_appendix_batch(
                root,
                structure,
                {
                    "batch_token": batch["batch_token"],
                    "root_addition": {
                        "node_id": "ADD-TECH-APPENDIX",
                        "reason": "Required appendix root.",
                    },
                    "items": [include_item],
                },
                inventory,
            )
            final_state = json_load(root / "outline_decision_state.json")

        self.assertEqual(final_state["active_appendix_batch"]["appendix_ids"], [])
        self.assertEqual(len(final_state["additions"]), 3)
        appendix_root = next(
            item for item in final_state["additions"] if item.get("node_id") == "ADD-TECH-APPENDIX"
        )
        self.assertEqual(appendix_root["parent_id"], None)
        self.assertEqual(appendix_root["number"], "附录")
        self.assertEqual(appendix_root["title"], "技术附表")

    def test_bid_outline_appendix_cli_rejects_non_string_json_fields_atomically(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest, manifest_path = write_decision_context_fixture(
                root, heading_count=0
            )
            (root / "tender_appendix_inventory.json").write_text(
                json.dumps(
                    {
                        "schema_version": "tender-appendix-inventory.v1",
                        "items": [
                            {
                                "appendix_id": "7",
                                "file_id": "TEN-1",
                                "number": "Appendix T.1",
                                "title": "Typed form",
                                "following_table_count": 1,
                            }
                        ],
                    }
                ),
                encoding="utf-8",
            )
            template = outline_runner.dispatch_command(
                "decision-next", manifest, manifest_path, []
            )
            outline_runner.dispatch_command(
                "decision-batch",
                manifest,
                manifest_path,
                [
                    json.dumps(
                        {
                            "batch_token": template["batch_token"],
                            "items": [
                                {
                                    "target_id": template["items"][0]["target_id"],
                                    "decision": "retain",
                                }
                            ],
                            "additions": [],
                        }
                    )
                ],
            )
            batch = outline_runner.dispatch_command(
                "appendix-next", manifest, manifest_path, []
            )
            state_path = root / "outline_decision_state.json"
            original_state = state_path.read_text(encoding="utf-8")
            valid_root = {
                "node_id": "ADD-TECH-APPENDIX",
                "parent_id": None,
                "number": "第7章",
                "title": "技术附表",
                "reason": "Required root.",
            }
            valid_include = {
                "appendix_id": batch["items"][0]["appendix_id"],
                "decision": "include",
                "node_id": "ADD-APP-0001",
                "parent_id": "ADD-TECH-APPENDIX",
                "reason": "Required form.",
            }

            invalid_payloads: list[tuple[str, dict]] = []
            for field in ("node_id", "number", "title", "reason"):
                for value in ([], {"unexpected": True}, 7):
                    root_addition = {**valid_root, field: value}
                    item = dict(valid_include)
                    if field == "node_id":
                        item["parent_id"] = str(value)
                    invalid_payloads.append(
                        (
                            rf"root_addition\.{field} must be a string",
                            {
                                "batch_token": batch["batch_token"],
                                "root_addition": root_addition,
                                "items": [item],
                            },
                        )
                    )
            for value in ("TPL-0001", {"unexpected": True}):
                invalid_payloads.append(
                    (
                        r"root_addition\.parent_id must be null",
                        {
                            "batch_token": batch["batch_token"],
                            "root_addition": {**valid_root, "parent_id": value},
                            "items": [valid_include],
                        },
                    )
                )

            invalid_item_cases = [
                ("appendix_id", 7, r"items\[0\]\.appendix_id must be a string"),
                ("decision", ["include"], r"items\[0\]\.decision must be a string"),
                ("node_id", {"unexpected": True}, r"items\[0\]\.node_id must be a string"),
                ("reason", [], r"items\[0\]\.reason must be a string"),
            ]
            for field, value, message in invalid_item_cases:
                invalid_payloads.append(
                    (
                        message,
                        {
                            "batch_token": batch["batch_token"],
                            "root_addition": valid_root,
                            "items": [{**valid_include, field: value}],
                        },
                    )
                )
            invalid_payloads.append(
                (
                    r"items\[0\]\.parent_id must be a string",
                    {
                        "batch_token": batch["batch_token"],
                        "root_addition": {**valid_root, "node_id": "7"},
                        "items": [{**valid_include, "parent_id": 7}],
                    },
                )
            )
            for field, value, message in (
                ("appendix_id", 7, r"items\[0\]\.appendix_id must be a string"),
                ("decision", ["exclude"], r"items\[0\]\.decision must be a string"),
                ("reason", {"unexpected": True}, r"items\[0\]\.reason must be a string"),
            ):
                invalid_payloads.append(
                    (
                        message,
                        {
                            "batch_token": batch["batch_token"],
                            "items": [
                                {
                                    "appendix_id": batch["items"][0]["appendix_id"],
                                    "decision": "exclude",
                                    "reason": "Not required.",
                                    field: value,
                                }
                            ],
                        },
                    )
                )

            for message, payload in invalid_payloads:
                with self.subTest(message=message, payload=payload):
                    state_path.write_text(original_state, encoding="utf-8")
                    with self.assertRaisesRegex(SystemExit, message):
                        outline_runner.dispatch_command(
                            "appendix-decision-batch",
                            manifest,
                            manifest_path,
                            [json.dumps(payload, ensure_ascii=False)],
                        )
                    self.assertEqual(
                        state_path.read_text(encoding="utf-8"), original_state
                    )

    def test_bid_outline_appendix_later_first_include_creates_root(self) -> None:
        decision_workflow = load_outline_script("run_from_manifest").decision_workflow
        structure = {
            "schema_version": "template-structure.v1",
            "items": [{"number": "1", "title": "Technical proposal", "level": 1}],
        }
        inventory = [
            {
                "appendix_id": f"APP-{index:04d}",
                "file_id": "TEN-1",
                "number": f"Appendix L.{index}",
                "title": f"Later form {index}",
                "following_table_count": 1,
            }
            for index in range(1, 3)
        ]

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = decision_workflow.next_decision_batch(
                root, structure
            )
            decision_workflow.submit_decision_batch(
                root,
                structure,
                {
                    "batch_token": template["batch_token"],
                    "items": [
                        {
                            "target_id": template["items"][0]["target_id"],
                            "decision": "retain",
                        }
                    ],
                },
            )
            first = decision_workflow.next_appendix_batch(
                root, structure, inventory, max_items=1
            )
            decision_workflow.submit_appendix_batch(
                root,
                structure,
                {
                    "batch_token": first["batch_token"],
                    "items": [
                        {
                            "appendix_id": "APP-0001",
                            "decision": "exclude",
                            "reason": "Not required.",
                        }
                    ],
                },
                inventory,
            )
            self.assertEqual(
                json_load(root / "outline_decision_state.json")["additions"], []
            )

            second = decision_workflow.next_appendix_batch(
                root, structure, inventory, max_items=1
            )
            decision_workflow.submit_appendix_batch(
                root,
                structure,
                {
                    "batch_token": second["batch_token"],
                    "root_addition": {
                        "node_id": "ADD-TECH-APPENDIX",
                        "parent_id": None,
                        "number": "第7章",
                        "title": "技术附表",
                        "reason": "Later batch first include.",
                    },
                    "items": [
                        {
                            "appendix_id": "APP-0002",
                            "decision": "include",
                            "node_id": "ADD-APP-0002",
                            "parent_id": "ADD-TECH-APPENDIX",
                            "reason": "Required later form.",
                        }
                    ],
                },
                inventory,
            )
            finalized = decision_workflow.finalize_decisions(
                root, structure, appendix_items=inventory
            )
            decisions = json_load(Path(finalized["decisionsFile"]))
            outline, _ = decision_workflow.outline_composer.build_composition(
                structure, decisions
            )

        self.assertEqual(outline["nodes"][-1]["title"], "技术附表")
        self.assertEqual(
            [item["title"] for item in outline["nodes"][-1]["children"]],
            ["Later form 2"],
        )

    def test_bid_outline_appendix_all_exclude_finishes_without_root(self) -> None:
        decision_workflow = load_outline_script("run_from_manifest").decision_workflow
        structure = {
            "schema_version": "template-structure.v1",
            "items": [{"number": "1", "title": "Technical proposal", "level": 1}],
        }
        inventory = [
            {
                "appendix_id": f"APP-{index:04d}",
                "file_id": "TEN-1",
                "number": f"Appendix X.{index}",
                "title": f"Excluded form {index}",
                "following_table_count": 1,
            }
            for index in range(1, 65)
        ]

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = decision_workflow.next_decision_batch(
                root, structure
            )
            decision_workflow.submit_decision_batch(
                root,
                structure,
                {
                    "batch_token": template["batch_token"],
                    "items": [
                        {
                            "target_id": template["items"][0]["target_id"],
                            "decision": "retain",
                        }
                    ],
                },
            )

            while True:
                batch = decision_workflow.next_appendix_batch(
                    root, structure, inventory
                )
                if batch["complete"]:
                    break
                decision_workflow.submit_appendix_batch(
                    root,
                    structure,
                    {
                        "batch_token": batch["batch_token"],
                        "items": [
                            {
                                "appendix_id": item["appendix_id"],
                                "decision": "exclude",
                                "reason": "Not required by this tender.",
                            }
                            for item in batch["items"]
                        ],
                    },
                    inventory,
                )

            state = json_load(root / "outline_decision_state.json")
            finalized = decision_workflow.finalize_decisions(
                root, structure, appendix_items=inventory
            )
            decisions = json_load(Path(finalized["decisionsFile"]))
            outline, _ = decision_workflow.outline_composer.build_composition(
                structure, decisions
            )
            validation = decision_workflow.validate_finalized_decisions(
                root,
                structure,
                decisions,
                workflow_binding={},
                appendix_items=inventory,
                include_appendix_decisions=True,
            )

        self.assertEqual(state["additions"], [])
        self.assertEqual(decisions["changes"], [])
        self.assertNotIn("技术附表", json.dumps(outline, ensure_ascii=False))
        self.assertEqual(len(validation["appendixDecisions"]), 64)
        self.assertTrue(
            all(item["decision"] == "exclude" for item in validation["appendixDecisions"])
        )

    def test_bid_outline_decisions_reject_unjudged_appendix_candidates(self) -> None:
        decision_workflow = load_outline_script("decision_workflow")
        structure = {
            "schema_version": "template-structure.v1",
            "items": [{"number": "1", "title": "Technical proposal", "level": 1}],
        }
        inventory = [
            {
                "appendix_id": f"APP-{index:04d}",
                "file_id": "TEN-1",
                "number": f"Appendix C.{index}",
                "title": f"Form {index}",
                "following_table_count": 1,
            }
            for index in range(1, 22)
        ]

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = decision_workflow.next_decision_batch(
                root, structure
            )
            decision_workflow.submit_decision_batch(
                root,
                structure,
                {
                    "batch_token": template["batch_token"],
                    "items": [
                        {
                            "target_id": template["items"][0]["target_id"],
                            "decision": "retain",
                        }
                    ],
                },
            )
            with self.assertRaises(SystemExit) as raised:
                decision_workflow.finalize_decisions(
                    root, structure, appendix_items=inventory
                )

        message = str(raised.exception)
        self.assertIn("APP-0001", message)
        self.assertIn("APP-0020", message)
        self.assertNotIn("APP-0021", message)

    def test_bid_outline_appendix_submit_validates_complete_composition_atomically(self) -> None:
        decision_workflow = load_outline_script("run_from_manifest").decision_workflow
        structure = {
            "schema_version": "template-structure.v1",
            "items": [{"number": "1", "title": "Technical proposal", "level": 1}],
        }
        inventory = [
            {
                "appendix_id": "APP-0001",
                "file_id": "TEN-1",
                "number": "Appendix E.1",
                "title": "Validated form",
                "following_table_count": 1,
            }
        ]

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = decision_workflow.next_decision_batch(
                root, structure
            )
            decision_workflow.submit_decision_batch(
                root,
                structure,
                {
                    "batch_token": template["batch_token"],
                    "items": [
                        {
                            "target_id": template["items"][0]["target_id"],
                            "decision": "retain",
                        }
                    ],
                    "additions": [
                        {
                            "node_id": "ADD-TECH-APPENDIX",
                            "parent_id": None,
                            "number": "2",
                            "title": "技术附表",
                            "reason": "Controlled appendix root.",
                        },
                        {
                            "node_id": "ADD-INVALID-EXISTING",
                            "parent_id": None,
                            "number": "",
                            "title": "Invalid existing change",
                            "reason": "Empty number must be rejected by the builder.",
                        },
                    ],
                },
            )
            batch = decision_workflow.next_appendix_batch(root, structure, inventory)

            def include_payload(parent_id: str) -> dict:
                return {
                    "batch_token": batch["batch_token"],
                    "items": [
                        {
                            "appendix_id": "APP-0001",
                            "decision": "include",
                            "node_id": "ADD-APPENDIX-E1",
                            "parent_id": parent_id,
                            "reason": "Required controlled form.",
                        }
                    ],
                }

            corrected_parent_state = json_load(root / "outline_decision_state.json")
            with self.assertRaisesRegex(SystemExit, "changes\\[1\\]\\.number"):
                decision_workflow.submit_appendix_batch(
                    root,
                    structure,
                    include_payload("ADD-TECH-APPENDIX"),
                    inventory,
                )
            self.assertEqual(
                json_load(root / "outline_decision_state.json"), corrected_parent_state
            )

            corrected_parent_state["additions"][1]["number"] = "3"
            (root / "outline_decision_state.json").write_text(
                json.dumps(corrected_parent_state, ensure_ascii=False), encoding="utf-8"
            )
            decision_workflow.submit_appendix_batch(
                root,
                structure,
                include_payload("ADD-TECH-APPENDIX"),
                inventory,
            )
            decision_workflow.finalize_decisions(
                root, structure, appendix_items=inventory
            )
            final_state = json_load(root / "outline_decision_state.json")

        self.assertEqual(final_state["active_appendix_batch"]["appendix_ids"], [])
        self.assertEqual(final_state["appendix_decisions"]["APP-0001"]["decision"], "include")

    def test_bid_outline_decisions_reject_add_or_move_under_collapsed_parent(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest_path = root / "s2_input.json"
            manifest = {"workDir": str(root), "outputFile": str(root / "toc.json")}
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            structure = {
                "schema_version": "template-structure.v1",
                "items": [
                    {"number": "第1章", "title": "技术方案", "level": 1},
                    {"number": "1.1", "title": "待收敛节点", "level": 2},
                    {"number": "1.2", "title": "保留节点", "level": 2},
                    {"number": "1.2.1", "title": "待移动专题", "level": 3},
                ],
            }
            structure_path = root / "template_structure.json"
            structure_path.write_text(json.dumps(structure, ensure_ascii=False), encoding="utf-8")
            invalid_changes = [
                [
                    {
                        "operation": "collapse",
                        "target_id": "TPL-0002",
                        "reason": "由根节点统一承载。",
                    },
                    {
                        "operation": "add",
                        "node_id": "ADD-0001",
                        "parent_id": "TPL-0002",
                        "number": "1.1.1",
                        "title": "新增专题",
                        "suggestion_action": "建议增加",
                        "suggestion_reason": "招标要求独立提交。",
                    },
                ],
                [
                    {
                        "operation": "collapse",
                        "target_id": "TPL-0002",
                        "reason": "由根节点统一承载。",
                    },
                    {
                        "operation": "update",
                        "target_id": "TPL-0004",
                        "parent_id": "TPL-0002",
                        "reason": "调整专题归属。",
                    },
                ],
            ]

            for changes in invalid_changes:
                with self.subTest(operation=changes[-1]["operation"]), self.assertRaisesRegex(
                    SystemExit, "collapsed"
                ):
                    submit_outline_changes(outline_runner, manifest, manifest_path, changes)

    def test_bid_outline_finalize_rejects_output_not_reproducible_from_decisions(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            output = root / "toc.json"
            manifest_path = root / "s2_input.json"
            manifest = {
                "workDir": str(root),
                "outputFile": str(output),
                "requireComposedOutline": True,
            }
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            structure = {
                "schema_version": "template-structure.v1",
                "items": [{"number": "第1章", "title": "技术方案", "level": 1}],
            }
            (root / "template_structure.json").write_text(
                json.dumps(structure, ensure_ascii=False), encoding="utf-8"
            )
            fingerprint = outline_runner.outline_composer.annotate_template_structure(structure)[
                "input_fingerprint"
            ]
            outline_runner.submit_outline_decisions(
                manifest,
                manifest_path,
                {
                    "schema_version": "technical-outline-decisions.v1",
                    "input_fingerprint": fingerprint,
                    "template_decisions": [
                        {"target_id": "TPL-0001", "decision": "retain"}
                    ],
                    "changes": [],
                },
            )
            outline_runner.compose_manifest(manifest, manifest_path)
            payload = json_load(output)
            payload["nodes"][0]["title"] = "硬编码改写结果"
            output.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            report_path = root / "outline_compose_report.json"
            report = json_load(report_path)
            report["outputSha256"] = hashlib.sha256(output.read_bytes()).hexdigest()
            report_path.write_text(json.dumps(report, ensure_ascii=False), encoding="utf-8")

            with self.assertRaisesRegex(SystemExit, "与 decisions 合成结果不一致"):
                outline_runner.finalize_manifest(manifest, manifest_path)

    def test_bid_outline_template_structure_supplements_anchored_body_level_three(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "template.docx"
            doc = Document()
            if "TOC 1" not in [style.name for style in doc.styles]:
                doc.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
            if "TOC 2" not in [style.name for style in doc.styles]:
                doc.styles.add_style("TOC 2", WD_STYLE_TYPE.PARAGRAPH)
            doc.add_paragraph("第1章 总体技术方案 ........ 1", style="TOC 1")
            doc.add_paragraph("1.1 机组选型 ........ 2", style="TOC 2")
            doc.add_paragraph("1.2 供货范围 ........ 3", style="TOC 2")
            doc.add_paragraph("第1章 总体技术方案", style="Heading 1")
            doc.add_paragraph("1.1 机组选型", style="Heading 2")
            doc.add_paragraph("1.1.1 关键部件选型", style="Heading 3")
            doc.add_paragraph("1.1.1.1 叶片设计参数", style="Heading 4")
            doc.add_paragraph("1.2 供货范围", style="Heading 2")
            doc.add_paragraph("1.2.1 主机供货范围", style="Heading 3")
            doc.add_paragraph("第9章 正文噪声", style="Heading 1")
            doc.add_paragraph("9.1 非模板章节", style="Heading 2")
            doc.add_paragraph("9.1.1 不应补充", style="Heading 3")
            doc.save(template)

            result = outline_runner.extract_template_structure(template)

        self.assertEqual(result["source"], "automatic_toc")
        self.assertEqual(
            [(item["number"], item["title"], item["level"]) for item in result["items"]],
            [
                ("第1章", "总体技术方案", 1),
                ("1.1", "机组选型", 2),
                ("1.1.1", "关键部件选型", 3),
                ("1.2", "供货范围", 2),
                ("1.2.1", "主机供货范围", 3),
            ],
        )

    def test_bid_outline_template_structure_numbers_blank_body_level_three_in_document_order(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "template.docx"
            doc = Document()
            for style_name in ("TOC 1", "TOC 2"):
                if style_name not in [style.name for style in doc.styles]:
                    doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            doc.add_paragraph("\u7b2c1\u7ae0 Overview ........ 1", style="TOC 1")
            doc.add_paragraph("5.18 Digital Wind Farm ........ 2", style="TOC 2")
            doc.add_paragraph("\u7b2c1\u7ae0 Overview", style="Heading 1")
            doc.add_paragraph("5.18 Digital Wind Farm", style="Heading 2")
            doc.add_paragraph("SCADA System", style="Heading 3")
            doc.add_paragraph("Wind Farm Control", style="Heading 3")
            doc.save(template)

            result = outline_runner.extract_template_structure(template)

        self.assertEqual(
            [
                (item["number"], item["title"])
                for item in result["items"]
                if item["level"] == 3
            ],
            [
                ("5.18.1", "SCADA System"),
                ("5.18.2", "Wind Farm Control"),
            ],
        )

    def test_bid_outline_template_structure_deduplicates_existing_level_three(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "template.docx"
            doc = Document()
            for style_name in ("TOC 1", "TOC 2", "TOC 3"):
                if style_name not in [style.name for style in doc.styles]:
                    doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            doc.add_paragraph("第1章 总体技术方案 ........ 1", style="TOC 1")
            doc.add_paragraph("1.1 机组选型 ........ 2", style="TOC 2")
            doc.add_paragraph("1.1.1 关键部件选型 ........ 3", style="TOC 3")
            doc.add_paragraph("总体技术方案", style="Heading 1")
            doc.add_paragraph("机组选型", style="Heading 2")
            doc.add_paragraph("关键部件选型", style="Heading 3")
            doc.save(template)

            result = outline_runner.extract_template_structure(template)

        self.assertEqual(
            [(item["number"], item["title"], item["level"]) for item in result["items"]],
            [
                ("第1章", "总体技术方案", 1),
                ("1.1", "机组选型", 2),
                ("1.1.1", "关键部件选型", 3),
            ],
        )

    def test_bid_outline_template_structure_preserves_same_title_with_different_numbers(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "template.docx"
            doc = Document()
            for style_name in ("TOC 1", "TOC 2"):
                if style_name not in [style.name for style in doc.styles]:
                    doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            doc.add_paragraph("第1章 总体技术方案 ........ 1", style="TOC 1")
            doc.add_paragraph("1.1 认证情况 ........ 2", style="TOC 2")
            doc.add_paragraph("第1章 总体技术方案", style="Heading 1")
            doc.add_paragraph("1.1 认证情况", style="Heading 2")
            doc.add_paragraph("1.1.1 认证未完成或存在待解决项", style="Heading 3")
            doc.add_paragraph("1.1.2 认证未完成或存在待解决项", style="Heading 3")
            doc.save(template)

            result = outline_runner.extract_template_structure(template)

        self.assertEqual(
            [item["number"] for item in result["items"] if item["level"] == 3],
            ["1.1.1", "1.1.2"],
        )

    def test_bid_outline_template_structure_interleaves_supplemented_level_three(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "template.docx"
            doc = Document()
            for style_name in ("TOC 1", "TOC 2", "TOC 3", "TOC 4"):
                if style_name not in [style.name for style in doc.styles]:
                    doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            doc.add_paragraph("第1章 总体技术方案 ........ 1", style="TOC 1")
            doc.add_paragraph("1.1 机组选型 ........ 2", style="TOC 2")
            doc.add_paragraph("1.1.1 叶片专题 ........ 3", style="TOC 3")
            doc.add_paragraph("1.1.1.1 叶片参数 ........ 4", style="TOC 4")
            doc.add_paragraph("1.1.3 齿轮箱专题 ........ 5", style="TOC 3")
            doc.add_paragraph("第1章 总体技术方案", style="Heading 1")
            doc.add_paragraph("1.1 机组选型", style="Heading 2")
            doc.add_paragraph("1.1.1 叶片专题", style="Heading 3")
            doc.add_paragraph("1.1.1.1 叶片参数", style="Heading 4")
            doc.add_paragraph("1.1.2 变桨系统专题", style="Heading 3")
            doc.add_paragraph("1.1.3 齿轮箱专题", style="Heading 3")
            doc.save(template)

            result = outline_runner.extract_template_structure(template)

        self.assertEqual(
            [(item["number"], item["level"]) for item in result["items"]],
            [
                ("第1章", 1),
                ("1.1", 2),
                ("1.1.1", 3),
                ("1.1.1.1", 4),
                ("1.1.2", 3),
                ("1.1.3", 3),
            ],
        )

    def test_bid_outline_template_command_writes_tender_appendix_heading_inventory(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            manifest_path = root / "s2_input.json"

            template_doc = Document()
            template_doc.add_paragraph("第1章 技术方案", style="Heading 1")
            template_doc.save(template)

            tender_doc = Document()
            tender_doc.add_paragraph("附表H.1 工程进度表 177")
            tender_doc.add_paragraph("附表H.2 交货进度表 178")
            tender_doc.add_paragraph("技术附表I 技术条款偏差表 199")
            tender_doc.add_paragraph("技术附表H 进度表")
            tender_doc.add_paragraph("附表H.1 工程进度表")
            tender_doc.add_table(rows=1, cols=2)
            tender_doc.add_paragraph("附表H.2 交货进度表")
            tender_doc.add_table(rows=2, cols=2)
            tender_doc.add_paragraph("技术附表I 技术条款偏差表")
            tender_doc.add_table(rows=1, cols=3)
            tender_doc.add_paragraph("附表F.5-2 认证未完成或存在待解决项")
            tender_doc.add_table(rows=1, cols=2)
            tender_doc.add_paragraph("投标人应逐项填写，不得遗漏。")
            tender_doc.save(tender)

            manifest = {
                "projectId": "PRJ-TEST",
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": "tender.docx", "path": str(tender)}],
                "outputFile": str(root / "toc.json"),
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            result = outline_runner.write_template_structure(manifest, manifest_path)
            structure = json_load(root / "template_structure.json")
            inventory = json_load(root / "tender_appendix_inventory.json")

        self.assertEqual(result["inputFingerprint"], structure["input_fingerprint"])
        self.assertRegex(result["inputFingerprint"], r"^[0-9a-f]{64}$")
        self.assertEqual(result["tenderAppendixItemCount"], 5)
        self.assertEqual(result["tenderAppendixInventoryFile"], str(root / "tender_appendix_inventory.json"))
        self.assertEqual(inventory["schema_version"], "tender-appendix-inventory.v1")
        self.assertEqual(
            [(item["number"], item["title"], item["file_id"]) for item in inventory["items"]],
            [
                ("技术附表H", "进度表", "TEN-1"),
                ("附表H.1", "工程进度表", "TEN-1"),
                ("附表H.2", "交货进度表", "TEN-1"),
                ("技术附表I", "技术条款偏差表", "TEN-1"),
                ("附表F.5-2", "认证未完成或存在待解决项", "TEN-1"),
            ],
        )
        self.assertEqual([item["following_table_count"] for item in inventory["items"]], [0, 1, 1, 1, 1])

    def test_bid_outline_prepare_builds_ordered_tender_review_chunks(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            manifest_path = root / "s2_input.json"

            template_doc = Document()
            template_doc.add_paragraph("第1章 技术方案", style="Heading 1")
            template_doc.save(template)

            tender_doc = Document()
            tender_doc.add_paragraph("1. 总则", style="Heading 1")
            tender_doc.add_paragraph("投标人应提供总体技术方案。")
            table = tender_doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "项目"
            table.cell(0, 1).text = "要求"
            table.cell(0, 0).merge(table.cell(0, 1))
            table.cell(1, 0).text = "机型"
            table.cell(1, 1).text = "投标人填写"
            tender_doc.add_paragraph("2. 专题方案", style="Heading 1")
            tender_doc.add_paragraph("投标人须编制场址安全适应性报告。")
            tender_doc.save(tender)

            manifest = {
                "projectId": "PRJ-TEST",
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": "tender.docx", "path": str(tender)}],
                "outputFile": str(root / "toc.json"),
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            result = outline_runner.write_template_structure(manifest, manifest_path)
            chunks = json_load(root / "tender_review_chunks.json")
            state = json_load(root / "tender_review_state.json")

        self.assertEqual(result["tenderReviewChunkCount"], len(chunks["chunks"]))
        self.assertEqual(chunks["schema_version"], "tender-review-chunks.v1")
        self.assertEqual(chunks["source_block_count"], 5)
        self.assertEqual(
            [block["type"] for chunk in chunks["chunks"] for block in chunk["blocks"]],
            ["paragraph", "paragraph", "table", "paragraph", "paragraph"],
        )
        table_block = next(
            block for chunk in chunks["chunks"] for block in chunk["blocks"] if block["type"] == "table"
        )
        self.assertEqual(table_block["rows"][0]["cells"], ["项目 要求", "项目 要求"])
        self.assertEqual(state["schema_version"], "tender-review-state.v1")
        self.assertEqual(state["reviewed_chunk_count"], 0)
        self.assertEqual(state["pending_chunk_count"], len(chunks["chunks"]))

    def test_bid_outline_review_workspace_uses_lightweight_docx_xml_parser(self) -> None:
        source = (OUTLINE_SCRIPT_DIR / "review_workflow.py").read_text(encoding="utf-8")

        self.assertNotIn("from docx", source)
        self.assertNotIn("Document(str(path))", source)

    def test_bid_outline_headings_prefers_toc_and_skips_body_titles_without_marking_reviewed(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")
        review_workflow = load_outline_script("review_workflow")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            manifest_path = root / "s2_input.json"

            template_doc = Document()
            template_doc.add_paragraph("第1章 技术方案", style="Heading 1")
            template_doc.save(template)

            tender_doc = Document()
            if "TOC 1" not in [style.name for style in tender_doc.styles]:
                tender_doc.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
            tender_doc.add_paragraph("第1章 总体要求 ........ 1", style="TOC 1")
            tender_doc.add_paragraph("第1章 总体要求", style="Heading 1")
            tender_doc.add_paragraph("投标人应提供总体技术方案。")
            tender_doc.add_paragraph("2. 专题方案")
            tender_doc.add_paragraph("附表A.1 技术参数表")
            table = tender_doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "参数"
            table.cell(0, 1).text = "要求"
            table.cell(1, 0).text = "额定功率"
            table.cell(1, 1).text = "投标人填写"
            tender_doc.save(tender)

            manifest = {
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": "tender.docx", "path": str(tender)}],
                "outputFile": str(root / "toc.json"),
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
            outline_runner.write_template_structure(manifest, manifest_path)

            result = outline_runner.dispatch_command("headings", manifest, manifest_path, [])
            mapped_section = outline_runner.dispatch_command(
                "section",
                manifest,
                manifest_path,
                [result["files"][0]["items"][0]["section_id"]],
            )
            status = review_workflow.review_status(root)

        self.assertEqual(result["schema_version"], "tender-headings.v1")
        self.assertEqual(
            [(item["kind"], item["text"]) for item in result["files"][0]["items"]],
            [
                ("toc", "第1章 总体要求 ........ 1"),
            ],
        )
        self.assertEqual(result["files"][0]["source"], "toc")
        self.assertEqual(result["files"][0]["items"][0]["section_id"], "TEN-1:S0001")
        self.assertIn(
            "投标人应提供总体技术方案。",
            [item["text"] for item in mapped_section["records"]],
        )
        self.assertTrue(result["complete"])
        self.assertEqual(result["appendix_count"], 1)
        self.assertNotIn("appendices", result)
        self.assertEqual(status["reviewed_chunk_count"], 0)
        self.assertEqual(status["pending_chunk_count"], status["chunk_count"])

    def test_bid_outline_headings_pages_body_titles_when_toc_is_missing(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            manifest_path = root / "s2_input.json"

            template_doc = Document()
            template_doc.add_paragraph("Template", style="Heading 1")
            template_doc.save(template)

            tender_doc = Document()
            tender_doc.add_paragraph("Section one", style="Heading 1")
            tender_doc.add_paragraph("Section two", style="Heading 1")
            tender_doc.add_paragraph("Section three", style="Heading 1")
            tender_doc.save(tender)

            manifest = {
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": "tender.docx", "path": str(tender)}],
                "outputFile": str(root / "toc.json"),
            }
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            outline_runner.write_template_structure(manifest, manifest_path)

            first = outline_runner.dispatch_command(
                "headings", manifest, manifest_path, ["--page-size", "2"]
            )
            second = outline_runner.dispatch_command(
                "headings",
                manifest,
                manifest_path,
                ["--cursor", first["next_cursor"], "--page-size", "2"],
            )

        self.assertEqual(first["files"][0]["source"], "body_headings")
        self.assertEqual(
            [item["text"] for item in first["files"][0]["items"]],
            ["Section one", "Section two"],
        )
        self.assertFalse(first["complete"])
        self.assertEqual(first["next_cursor"], "2")
        self.assertEqual(
            [item["text"] for item in second["files"][0]["items"]],
            ["Section three"],
        )
        self.assertTrue(second["complete"])
        self.assertEqual(second["next_cursor"], "")

    def test_bid_outline_headings_exposes_stable_sections_and_section_reads_continuous_body(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            manifest_path = root / "s2_input.json"

            template_doc = Document()
            template_doc.add_paragraph("第1章 技术方案", style="Heading 1")
            template_doc.save(template)
            tender_doc = Document()
            tender_doc.add_paragraph("1 总体技术要求", style="Heading 1")
            tender_doc.add_paragraph("投标人应提交总体技术方案和项目组织方案。")
            tender_doc.add_paragraph("1.1 专题报告", style="Heading 2")
            tender_doc.add_paragraph("投标人应提交场址安全适应性专题报告。")
            tender_doc.add_paragraph("2 供货范围", style="Heading 1")
            tender_doc.add_paragraph("投标人应提供完整供货清单。")
            tender_doc.save(tender)
            manifest = {
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": tender.name, "path": str(tender)}],
                "outputFile": str(root / "toc.json"),
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
            outline_runner.write_template_structure(manifest, manifest_path)

            headings = outline_runner.dispatch_command("headings", manifest, manifest_path, [])
            section_id = headings["files"][0]["items"][0]["section_id"]
            first_page = outline_runner.dispatch_command(
                "section",
                manifest,
                manifest_path,
                [section_id, "--cursor", "0", "--max-chars", "20"],
            )
            second_page = outline_runner.dispatch_command(
                "section",
                manifest,
                manifest_path,
                [section_id, "--cursor", first_page["next_cursor"], "--max-chars", "200"],
            )
            review_headings = outline_runner.dispatch_command(
                "headings",
                manifest,
                manifest_path,
                ["--review", "--cursor", "0", "--page-size", "1"],
            )

        self.assertEqual(section_id, "TEN-1:S0001")
        self.assertEqual(first_page["section"]["title"], "1 总体技术要求")
        combined = first_page["records"] + second_page["records"]
        self.assertIn("投标人应提交场址安全适应性专题报告。", [item["text"] for item in combined])
        self.assertNotIn("2 供货范围", [item["text"] for item in combined])
        self.assertTrue(second_page["complete"])
        self.assertTrue(review_headings["review"])
        self.assertGreater(review_headings["returned_heading_count"], 0)

    def test_bid_outline_search_locates_full_text_without_marking_evidence_as_read(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")
        review_workflow = load_outline_script("review_workflow")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            manifest_path = root / "s2_input.json"
            template_doc = Document()
            template_doc.add_paragraph("第1章 技术方案", style="Heading 1")
            template_doc.save(template)
            tender_doc = Document()
            tender_doc.add_paragraph("1 安全要求", style="Heading 1")
            tender_doc.add_paragraph("投标人必须提交海上运输安全专项方案。")
            tender_doc.save(tender)
            manifest = {
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": tender.name, "path": str(tender)}],
                "outputFile": str(root / "toc.json"),
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
            outline_runner.write_template_structure(manifest, manifest_path)
            headings = outline_runner.dispatch_command("headings", manifest, manifest_path, [])

            result = outline_runner.dispatch_command(
                "search", manifest, manifest_path, ["运输安全", "--max-results", "10"]
            )
            evidence_id = result["results"][0]["evidence_id"]
            with self.assertRaisesRegex(SystemExit, "尚未通过受控阅读"):
                review_workflow.resolve_tender_basis(root, evidence_id)
            outline_runner.dispatch_command(
                "section",
                manifest,
                manifest_path,
                [headings["files"][0]["items"][0]["section_id"]],
            )
            basis = review_workflow.resolve_tender_basis(root, evidence_id)

        self.assertEqual(result["results"][0]["section_id"], "TEN-1:S0001")
        self.assertEqual(basis["evidence_id"], evidence_id)
        self.assertEqual(basis["file_id"], "TEN-1")
        self.assertEqual(basis["search_text"], "投标人必须提交海上运输安全专项方案。")

    def test_bid_outline_table_evidence_uses_meaningful_cell_text_for_location(self) -> None:
        review_workflow = load_outline_script("review_workflow")

        self.assertEqual(
            review_workflow.evidence_search_text(
                {"cells": ["2", "抗低温", "√"], "text": "2 | 抗低温 | √"}
            ),
            "抗低温",
        )
        self.assertEqual(
            review_workflow.evidence_search_text(
                {
                    "type": "table",
                    "rows": [
                        {"cells": ["序号", "货物名称", "品牌或制造商名称"]},
                        {"cells": ["1", "主控系统", "自主可控品牌"]},
                    ],
                    "text": "序号 | 货物名称 | 品牌或制造商名称 | 1 | 主控系统 | 自主可控品牌",
                }
            ),
            "自主可控品牌",
        )

    def test_bid_outline_tender_search_validation_uses_controlled_table_evidence(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")
        review_workflow = load_outline_script("review_workflow")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            manifest_path = root / "s2_input.json"
            template_doc = Document()
            template_doc.add_paragraph("第1章 技术方案", style="Heading 1")
            template_doc.save(template)
            tender_doc = Document()
            tender_doc.add_paragraph("1 环境适应性", style="Heading 1")
            table = tender_doc.add_table(rows=2, cols=3)
            for column, value in enumerate(["序号", "要求", "响应"]):
                table.cell(0, column).text = value
            for column, value in enumerate(["2", "抗低温", "√"]):
                table.cell(1, column).text = value
            tender_doc.save(tender)
            manifest = {
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": tender.name, "path": str(tender)}],
                "outputFile": str(root / "toc.json"),
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
            outline_runner.write_template_structure(manifest, manifest_path)
            table_result = review_workflow.read_table(root, "TEN-1:T0001", start=1, end=2)
            evidence_id = table_result["rows"][1]["evidence_id"]
            basis = review_workflow.resolve_tender_basis(root, evidence_id)
            nodes = [{"tender_basis": basis, "children": []}]

            access_path = root / "tender_evidence_access.json"
            access = json_load(access_path)
            access["evidence_ids"] = []
            access["events"] = []
            access_path.write_text(json.dumps(access, ensure_ascii=False), encoding="utf-8")

            outline_runner.validate_tender_search_texts(nodes, manifest, work_dir=root)
            nodes[0]["tender_basis"] = {**basis, "search_text": "被篡改的定位文本"}
            with self.assertRaisesRegex(SystemExit, "受控证据不一致"):
                outline_runner.validate_tender_search_texts(nodes, manifest, work_dir=root)

    def test_bid_outline_headings_pages_toc_items_with_same_cursor_contract(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            manifest_path = root / "s2_input.json"

            template_doc = Document()
            template_doc.add_paragraph("Template", style="Heading 1")
            template_doc.save(template)

            tender_doc = Document()
            if "TOC 1" not in [style.name for style in tender_doc.styles]:
                tender_doc.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
            for index in range(1, 6):
                tender_doc.add_paragraph(f"Section {index} ........ {index}", style="TOC 1")
            tender_doc.save(tender)

            manifest = {
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": "tender.docx", "path": str(tender)}],
                "outputFile": str(root / "toc.json"),
            }
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            outline_runner.write_template_structure(manifest, manifest_path)

            first = outline_runner.dispatch_command(
                "headings", manifest, manifest_path, ["--page-size", "2"]
            )
            self.assertEqual(first["returned_heading_count"], 2)
            self.assertEqual(first["next_cursor"], "2")
            self.assertFalse(first["complete"])

            second = outline_runner.dispatch_command(
                "headings",
                manifest,
                manifest_path,
                ["--cursor", first["next_cursor"], "--page-size", "2"],
            )
            self.assertEqual(second["returned_heading_count"], 2)
            self.assertEqual(second["next_cursor"], "4")
            self.assertFalse(second["complete"])

            third = outline_runner.dispatch_command(
                "headings",
                manifest,
                manifest_path,
                ["--cursor", second["next_cursor"], "--page-size", "2"],
            )
            self.assertEqual(third["returned_heading_count"], 1)
            self.assertEqual(third["next_cursor"], "")
            self.assertTrue(third["complete"])

    def test_bid_outline_headings_requires_full_review_when_no_structure_exists(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            manifest_path = root / "s2_input.json"
            template_doc = Document()
            template_doc.add_paragraph("第1章 技术方案", style="Heading 1")
            template_doc.save(template)
            tender_doc = Document()
            tender_doc.add_paragraph("本项目位于沿海区域。")
            tender_doc.add_paragraph("投标人应提交完整技术响应文件。")
            tender_doc.save(tender)
            manifest = {
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": tender.name, "path": str(tender)}],
                "outputFile": str(root / "toc.json"),
                "requireComposedOutline": True,
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
            outline_runner.write_template_structure(manifest, manifest_path)

            first = outline_runner.dispatch_command("headings", manifest, manifest_path, [])
            with self.assertRaisesRegex(SystemExit, "全文审阅"):
                outline_runner.dispatch_command("decision-next", manifest, manifest_path, [])
            batch = outline_runner.dispatch_command("next-batch", manifest, manifest_path, [])
            outline_runner.dispatch_command(
                "review-batch",
                manifest,
                manifest_path,
                [
                    json.dumps(
                        {
                            "chunk_ids": batch["chunk_ids"],
                            "review_summary": "已逐段审阅无结构招标文件。",
                            "requirements": [],
                        },
                        ensure_ascii=False,
                    )
                ],
            )
            second = outline_runner.dispatch_command("headings", manifest, manifest_path, [])

        self.assertEqual(first["heading_count"], 0)
        self.assertTrue(first["requires_full_review"])
        self.assertFalse(first["complete"])
        self.assertGreater(first["full_review_pending_chunk_count"], 0)
        self.assertTrue(second["complete"])
        self.assertEqual(second["full_review_pending_chunk_count"], 0)

    def test_bid_outline_strict_compose_requires_controlled_decisions_and_binds_state(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            output = root / "toc.json"
            manifest_path = root / "s2_input.json"
            template_doc = Document()
            template_doc.add_paragraph("第1章 技术方案", style="Heading 1")
            template_doc.save(template)
            tender_doc = Document()
            tender_doc.add_paragraph("第1章 招标技术要求", style="Heading 1")
            tender_doc.save(tender)
            manifest = {
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": tender.name, "path": str(tender)}],
                "outputFile": str(output),
                "requireComposedOutline": True,
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
            outline_runner.write_template_structure(manifest, manifest_path)
            with self.assertRaisesRegex(SystemExit, "完整读取招标目录"):
                outline_runner.compose_manifest(manifest, manifest_path)
            headings = outline_runner.dispatch_command("headings", manifest, manifest_path, [])
            section = outline_runner.dispatch_command(
                "section",
                manifest,
                manifest_path,
                [headings["files"][0]["items"][0]["section_id"]],
            )
            evidence_id = section["records"][0]["evidence_id"]
            structure = json_load(root / "template_structure.json")
            outline_runner.submit_outline_decisions(
                manifest,
                manifest_path,
                {
                    "schema_version": "technical-outline-decisions.v1",
                    "input_fingerprint": structure["input_fingerprint"],
                    "template_decisions": [
                        {"target_id": "TPL-0001", "decision": "retain"}
                    ],
                    "changes": [],
                },
            )
            with self.assertRaisesRegex(SystemExit, "受控决策"):
                outline_runner.compose_manifest(manifest, manifest_path)

            decision_batch = outline_runner.dispatch_command(
                "decision-next", manifest, manifest_path, []
            )
            outline_runner.dispatch_command(
                "read", manifest, manifest_path, [evidence_id]
            )
            outline_runner.dispatch_command(
                "decision-batch",
                manifest,
                manifest_path,
                [
                    json.dumps(
                        {
                            "batch_token": decision_batch["batch_token"],
                            "items": [
                                {
                                    "target_id": "TPL-0001",
                                    "decision": "retain",
                                    "evidence_id": evidence_id,
                                }
                            ],
                            "additions": [],
                        },
                        ensure_ascii=False,
                    )
                ],
            )
            outline_runner.dispatch_command(
                "read", manifest, manifest_path, [evidence_id]
            )
            complete_outline_review(outline_runner, manifest, manifest_path)
            outline_runner.dispatch_command("decisions", manifest, manifest_path, [])
            outline_runner.compose_manifest(manifest, manifest_path)
            report = json_load(root / "outline_compose_report.json")

            state_path = root / "outline_decision_state.json"
            state = json_load(state_path)
            state["tampered"] = True
            state_path.write_text(json.dumps(state, ensure_ascii=False), encoding="utf-8")
            with self.assertRaisesRegex(SystemExit, "decisionStateDigest"):
                outline_runner.finalize_manifest(manifest, manifest_path)

        self.assertTrue(report["tenderInputsDigest"])
        self.assertTrue(report["headingsStateDigest"])
        self.assertTrue(report["decisionStateDigest"])

    def test_bid_outline_review_navigation_resumes_from_first_pending_chunk(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")
        review_workflow = load_outline_script("review_workflow")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            manifest_path = root / "s2_input.json"
            template_doc = Document()
            template_doc.add_paragraph("第1章 技术方案", style="Heading 1")
            template_doc.save(template)
            tender_doc = Document()
            tender_doc.add_paragraph("1. 总则", style="Heading 1")
            tender_doc.add_paragraph("投标人应提供总体技术方案。")
            tender_doc.add_paragraph("2. 专题", style="Heading 1")
            tender_doc.add_paragraph("投标人应提供场址安全适应性报告。")
            tender_doc.save(tender)
            manifest = {
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": "tender.docx", "path": str(tender)}],
                "outputFile": str(root / "toc.json"),
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
            outline_runner.write_template_structure(manifest, manifest_path)

            first = review_workflow.next_review_chunk(root)
            review_workflow.submit_chunk_review(
                root,
                first["chunk"]["chunk_id"],
                {"review_summary": "已审阅总则要求。", "requirements": []},
            )
            second = review_workflow.next_review_chunk(root)
            state = review_workflow.review_status(root)

        self.assertNotEqual(first["chunk"]["chunk_id"], second["chunk"]["chunk_id"])
        self.assertEqual(first["remaining_chunk_count"], 2)
        self.assertEqual(second["remaining_chunk_count"], 1)
        self.assertEqual(state["reviewed_chunk_count"], 1)
        self.assertEqual(state["pending_chunk_count"], 1)

    def test_bid_outline_table_navigation_reports_continuation_and_truncation(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")
        review_workflow = load_outline_script("review_workflow")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            manifest_path = root / "s2_input.json"
            template_doc = Document()
            template_doc.add_paragraph("第1章 技术方案", style="Heading 1")
            template_doc.save(template)
            tender_doc = Document()
            tender_doc.add_paragraph("附表A.1 参数表")
            table = tender_doc.add_table(rows=30, cols=1)
            for row_index in range(30):
                table.cell(row_index, 0).text = "超长参数" * 80 if row_index == 0 else f"参数{row_index + 1}"
            tender_doc.save(tender)
            manifest = {
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": "tender.docx", "path": str(tender)}],
                "outputFile": str(root / "toc.json"),
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
            outline_runner.write_template_structure(manifest, manifest_path)

            first_page = review_workflow.read_table(root, "TEN-1:T0001", start=1, end=24, max_chars=2_400)
            truncated_row = review_workflow.read_table(root, "TEN-1:T0001", start=1, end=1, max_chars=80)
            state = json_load(root / "tender_review_state.json")

        self.assertEqual(first_page["table"]["row_count"], 30)
        self.assertEqual(first_page["returned_range"], {"start": 1, "end": 24})
        self.assertTrue(first_page["has_more"])
        self.assertEqual(first_page["next_range"], "25-30")
        self.assertEqual(truncated_row["truncated_rows"], [1])
        table_chunk = next(item for item in state["chunks"] if item["chunk_id"].endswith("C0002"))
        self.assertEqual(table_chunk["table_read_ranges"], [{"start": 1, "end": 24}])

    def test_bid_outline_chunk_review_validates_dynamic_requirement_dispositions(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")
        review_workflow = load_outline_script("review_workflow")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            manifest_path = root / "s2_input.json"
            template_doc = Document()
            template_doc.add_paragraph("第1章 技术方案", style="Heading 1")
            template_doc.save(template)
            tender_doc = Document()
            tender_doc.add_paragraph("5.8 项目风机各子系统专题", style="Heading 1")
            tender_doc.add_paragraph("投标人应提供叶片专题。")
            tender_doc.add_paragraph("投标人应提供场址安全适应性报告。")
            tender_doc.save(tender)
            manifest = {
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": "tender.docx", "path": str(tender)}],
                "outputFile": str(root / "toc.json"),
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
            outline_runner.write_template_structure(manifest, manifest_path)
            chunk = review_workflow.next_review_chunk(root)["chunk"]
            paragraph_ids = [block["evidence_id"] for block in chunk["blocks"] if block["type"] == "paragraph"]

            with self.assertRaisesRegex(SystemExit, "target_node"):
                review_workflow.submit_chunk_review(
                    root,
                    chunk["chunk_id"],
                    {
                        "review_summary": "已识别子系统专题。",
                        "requirements": [
                            {
                                "evidence_ids": [paragraph_ids[1]],
                                "obligation": "投标人应提供叶片专题",
                                "disposition": "map_existing",
                            }
                        ],
                    },
                )
            with self.assertRaisesRegex(SystemExit, "不属于当前分块"):
                review_workflow.submit_chunk_review(
                    root,
                    chunk["chunk_id"],
                    {
                        "review_summary": "证据校验。",
                        "requirements": [
                            {
                                "evidence_ids": ["TEN-1:B999999"],
                                "obligation": "不存在的义务",
                                "disposition": "reference_only",
                                "reason": "仅作参考",
                            }
                        ],
                    },
                )

            result = review_workflow.submit_chunk_review(
                root,
                chunk["chunk_id"],
                {
                    "review_summary": "已逐项判断两个独立响应要求。",
                    "requirements": [
                        {
                            "evidence_ids": [paragraph_ids[1]],
                            "obligation": "投标人应提供叶片专题",
                            "disposition": "map_existing",
                            "target_node": "5.8",
                        },
                        {
                            "evidence_ids": [paragraph_ids[2]],
                            "obligation": "投标人应提供场址安全适应性报告",
                            "disposition": "suggest_add",
                            "target_node": "5.7",
                            "proposed_title": "项目场址安全适应性报告",
                            "reason": "招标明确要求独立报告，模板没有语义等价节点。",
                        },
                    ],
                },
            )
            ledger = json_load(root / "requirement_ledger.json")

        self.assertEqual(result["added_requirement_count"], 2)
        self.assertEqual(ledger["requirement_count"], 2)
        self.assertEqual(
            [item["disposition"] for item in ledger["requirements"]],
            ["map_existing", "suggest_add"],
        )
        self.assertEqual(ledger["requirements"][1]["proposed_title"], "项目场址安全适应性报告")

    def test_bid_outline_finalize_allows_selective_review_and_realizes_dynamic_additions(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")
        review_workflow = load_outline_script("review_workflow")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            output = root / "toc.json"
            manifest_path = root / "s2_input.json"
            template_doc = Document()
            template_doc.add_paragraph("第1章 技术方案", style="Heading 1")
            template_doc.save(template)
            tender_doc = Document()
            tender_doc.add_paragraph("投标人应提供项目场址安全适应性报告。")
            tender_doc.save(tender)
            manifest = {
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": "tender.docx", "path": str(tender)}],
                "outputFile": str(output),
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
            outline_runner.write_template_structure(manifest, manifest_path)
            payload = {
                "schema_version": "technical-outline.v1",
                "nodes": [
                    {
                        "number": "第1章",
                        "title": "技术方案",
                        "suggestion_action": "必要",
                        "suggestion_reason": "",
                        "children": [],
                    }
                ],
            }
            output.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

            selective_result = outline_runner.finalize_manifest(manifest, manifest_path)
            self.assertEqual(selective_result["summary"]["reviewCoverage"], 0.0)

            chunk = review_workflow.next_review_chunk(root)["chunk"]
            evidence_id = chunk["blocks"][0]["evidence_id"]
            review_workflow.submit_chunk_review(
                root,
                chunk["chunk_id"],
                {
                    "review_summary": "发现一项独立专项报告要求。",
                    "requirements": [
                        {
                            "evidence_ids": [evidence_id],
                            "obligation": "投标人应提供项目场址安全适应性报告",
                            "disposition": "suggest_add",
                            "target_node": "第1章",
                            "proposed_title": "项目场址安全适应性报告",
                            "reason": "招标明确要求独立报告，模板没有等价节点。",
                        }
                    ],
                },
            )
            with self.assertRaisesRegex(SystemExit, "未落实到最终目录"):
                outline_runner.finalize_manifest(manifest, manifest_path)

            payload["nodes"][0]["children"].append(
                {
                    "number": "1.1",
                    "title": "项目场址安全适应性报告",
                    "suggestion_action": "建议增加",
                    "suggestion_reason": "招标明确要求独立报告，模板没有等价节点。",
                    "tender_basis": {
                        "file_id": "TEN-1",
                        "search_text": "投标人应提供项目场址安全适应性报告",
                    },
                    "children": [],
                }
            )
            output.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            result = outline_runner.finalize_manifest(manifest, manifest_path)

        self.assertEqual(result["summary"]["reviewCoverage"], 1.0)
        self.assertEqual(result["summary"]["requirementCount"], 1)
        self.assertEqual(result["summary"]["unfinishedTableCount"], 0)

    def test_bid_outline_finalize_does_not_require_reading_appendix_table_contents(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            output = root / "toc.json"
            manifest_path = root / "s2_input.json"

            template_doc = Document()
            template_doc.add_paragraph("第1章 技术方案", style="Heading 1")
            template_doc.save(template)
            tender_doc = Document()
            tender_doc.add_paragraph("附表A.1 技术参数表")
            table = tender_doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "参数"
            table.cell(0, 1).text = "要求"
            table.cell(1, 0).text = "额定功率"
            table.cell(1, 1).text = "投标人填写"
            tender_doc.save(tender)

            manifest = {
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": "tender.docx", "path": str(tender)}],
                "outputFile": str(output),
                "requireComposedOutline": True,
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
            outline_runner.write_template_structure(manifest, manifest_path)
            headings = outline_runner.dispatch_command("headings", manifest, manifest_path, [])
            self.assertTrue(headings["complete"])
            batch = outline_runner.dispatch_command("decision-next", manifest, manifest_path, [])
            appendix_heading = outline_runner.dispatch_command(
                "search", manifest, manifest_path, ["附表A.1"]
            )
            outline_runner.dispatch_command(
                "read",
                manifest,
                manifest_path,
                [appendix_heading["results"][0]["evidence_id"]],
            )
            outline_runner.dispatch_command(
                "decision-batch",
                manifest,
                manifest_path,
                [
                    json.dumps(
                        {
                            "batch_token": batch["batch_token"],
                            "items": [
                                {
                                    "target_id": item["target_id"],
                                    "decision": "retain",
                                    "reason": "历史模板专家经验保留。",
                                }
                                for item in batch["items"]
                            ],
                            "additions": [],
                        },
                        ensure_ascii=False,
                    )
                ],
            )
            appendix_batch = outline_runner.dispatch_command(
                "appendix-next", manifest, manifest_path, []
            )
            outline_runner.dispatch_command(
                "appendix-decision-batch",
                manifest,
                manifest_path,
                [
                    json.dumps(
                        {
                            "batch_token": appendix_batch["batch_token"],
                            "items": [
                                {
                                    "appendix_id": "APP-0001",
                                    "decision": "include",
                                    "node_id": "ADD-APPENDIX-A1",
                                    "parent_id": "ADD-APPENDIX",
                                    "reason": "招标文件包含独立填写表格。",
                                }
                            ],
                            "root_addition": {
                                "node_id": "ADD-APPENDIX",
                                "parent_id": None,
                                "number": "第2章",
                                "title": "技术附表",
                                "reason": "招标文件包含独立附表。",
                            },
                        },
                        ensure_ascii=False,
                    )
                ],
            )
            outline_runner.dispatch_command(
                "read",
                manifest,
                manifest_path,
                [appendix_batch["items"][0]["evidence_id"]],
            )
            complete_outline_review(outline_runner, manifest, manifest_path)
            outline_runner.dispatch_command("decisions", manifest, manifest_path, [])
            outline_runner.compose_manifest(manifest, manifest_path)

            result = outline_runner.finalize_manifest(manifest, manifest_path)

        self.assertEqual(result["summary"]["workflowStage"], "finalized")
        self.assertEqual(result["summary"]["reviewCoverage"], 0.0)
        self.assertEqual(result["summary"]["unfinishedTableCount"], 1)

    def test_bid_outline_finalize_rejects_nested_technical_appendix(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            output = root / "toc.json"
            manifest_path = root / "s2_input.json"
            manifest = {"workDir": str(root), "outputFile": str(output)}
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            output.write_text(
                json.dumps(
                    {
                        "schema_version": "technical-outline.v1",
                        "nodes": [
                            {
                                "number": "第6章",
                                "title": "技术方案",
                                "suggestion_action": "必要",
                                "suggestion_reason": "",
                                "children": [
                                    {
                                        "number": "6.6",
                                        "title": "技术附表",
                                        "suggestion_action": "建议增加",
                                        "suggestion_reason": "招标包含技术附表。",
                                        "children": [],
                                    }
                                ],
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            with self.assertRaisesRegex(SystemExit, "技术附表必须是唯一的最后一个根节点"):
                outline_runner.finalize_manifest(manifest, manifest_path)

    def test_bid_outline_finalize_rejects_missing_included_appendix(self) -> None:
        def remove_included_addition(state: dict, _items: list[dict]) -> None:
            state["additions"] = [
                addition
                for addition in state["additions"]
                if addition.get("node_id") != "ADD-APPENDIX-1"
            ]

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            outline_runner, manifest, manifest_path = self._prepare_finalized_appendix_outline(
                root,
                ["include", "exclude"],
                mutate_state=remove_included_addition,
            )

            with self.assertRaisesRegex(SystemExit, "APP-0001"):
                outline_runner.finalize_manifest(manifest, manifest_path)

    def test_bid_outline_finalize_rejects_excluded_appendix_in_output(self) -> None:
        def exclude_included_addition(state: dict, items: list[dict]) -> None:
            appendix_id = items[0]["appendix_id"]
            state["appendix_decisions"][appendix_id] = {
                "appendix_id": appendix_id,
                "decision": "exclude",
                "reason": "逐项核验后排除。",
            }

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            outline_runner, manifest, manifest_path = self._prepare_finalized_appendix_outline(
                root,
                ["include", "exclude"],
                mutate_state=exclude_included_addition,
            )

            with self.assertRaisesRegex(SystemExit, "APP-0001"):
                outline_runner.finalize_manifest(manifest, manifest_path)

    def test_bid_outline_finalize_accepts_exact_ai_appendix_decisions(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            outline_runner, manifest, manifest_path = self._prepare_finalized_appendix_outline(
                root,
                ["include", "exclude"],
            )

            result = outline_runner.finalize_manifest(manifest, manifest_path)

        self.assertEqual(result["summary"]["workflowStage"], "finalized")

    def test_bid_outline_finalize_allows_missing_appendices_and_free_number_title_split(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            output = root / "toc.json"
            manifest_path = root / "s2_input.json"
            manifest = {"workDir": str(root), "outputFile": str(output)}
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            (root / "template_structure.json").write_text(
                json.dumps(
                    {
                        "schema_version": "template-structure.v1",
                        "items": [{"number": "第1章", "title": "技术方案", "level": 1}],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            (root / "tender_appendix_inventory.json").write_text(
                json.dumps(
                    {
                        "schema_version": "tender-appendix-inventory.v1",
                        "items": [
                            {
                                "number": "附表A.1",
                                "title": "机型信息表",
                                "following_table_count": 1,
                            },
                            {
                                "number": "附表A.2",
                                "title": "特殊方案表",
                                "following_table_count": 1,
                            },
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            output.write_text(
                json.dumps(
                    {
                        "schema_version": "technical-outline.v1",
                        "nodes": [
                            {
                                "number": "第1章",
                                "title": "技术方案",
                                "suggestion_action": "必要",
                                "suggestion_reason": "",
                                "children": [],
                            },
                            {
                                "number": "第2章",
                                "title": "技术附表",
                                "suggestion_action": "建议增加",
                                "suggestion_reason": "招标文件新增独立附表章节。",
                                "children": [
                                    {
                                        "number": "A.1",
                                        "title": "附表A.1 机型信息表",
                                        "suggestion_action": "建议增加",
                                        "suggestion_reason": "招标文件新增独立表格。",
                                        "children": [],
                                    }
                                ],
                            },
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = outline_runner.finalize_manifest(manifest, manifest_path)
            payload = json_load(output)
            payload["nodes"] = payload["nodes"][:1]
            output.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            result_without_appendix = outline_runner.finalize_manifest(manifest, manifest_path)

        self.assertEqual(result["summary"]["workflowStage"], "finalized")
        self.assertEqual(result_without_appendix["summary"]["workflowStage"], "finalized")

    def test_bid_outline_agentic_commands_drive_review_without_raw_file_reads(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            manifest_path = root / "s2_input.json"
            template_doc = Document()
            template_doc.add_paragraph("第1章 技术方案", style="Heading 1")
            template_doc.save(template)
            tender_doc = Document()
            tender_doc.add_paragraph("1. 总则", style="Heading 1")
            tender_doc.add_paragraph("投标人应提供总体技术方案。")
            tender_doc.save(tender)
            manifest = {
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": "tender.docx", "path": str(tender)}],
                "outputFile": str(root / "toc.json"),
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")

            prepared = outline_runner.dispatch_command("prepare", manifest, manifest_path, [])
            next_payload = outline_runner.dispatch_command("next", manifest, manifest_path, [])
            chunk = next_payload["chunk"]
            evidence_id = chunk["blocks"][1]["evidence_id"]
            read_payload = outline_runner.dispatch_command("read", manifest, manifest_path, [evidence_id])
            window_payload = outline_runner.dispatch_command(
                "window",
                manifest,
                manifest_path,
                [evidence_id, "--before", "1", "--after", "1"],
            )
            submitted = outline_runner.dispatch_command(
                "review-chunk",
                manifest,
                manifest_path,
                [
                    chunk["chunk_id"],
                    json.dumps(
                        {
                            "review_summary": "总体技术方案由模板节点承接。",
                            "requirements": [
                                {
                                    "evidence_ids": [evidence_id],
                                    "obligation": "投标人应提供总体技术方案",
                                    "disposition": "map_existing",
                                    "target_node": "第1章",
                                }
                            ],
                        },
                        ensure_ascii=False,
                    ),
                ],
            )
            status = outline_runner.dispatch_command("status", manifest, manifest_path, [])

        self.assertEqual(prepared["tenderReviewChunkCount"], 1)
        self.assertEqual(read_payload["record"]["evidence_id"], evidence_id)
        self.assertEqual(len(window_payload["blocks"]), 2)
        self.assertEqual(submitted["status"], "reviewed")
        self.assertEqual(status["review_coverage"], 1.0)
        self.assertEqual(status["requirement_count"], 1)

    def test_bid_outline_batch_review_keeps_business_decisions_agentic(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")
        review_workflow = load_outline_script("review_workflow")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            manifest_path = root / "s2_input.json"
            template_doc = Document()
            template_doc.add_paragraph("第1章 技术方案", style="Heading 1")
            template_doc.save(template)
            tender_doc = Document()
            tender_doc.add_paragraph("1. 总则", style="Heading 1")
            tender_doc.add_paragraph("投标人应提供总体技术方案。")
            table = tender_doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "项目"
            table.cell(0, 1).text = "要求"
            table.cell(1, 0).text = "报告"
            table.cell(1, 1).text = "投标人填写"
            tender_doc.add_paragraph("2. 专题", style="Heading 1")
            tender_doc.add_paragraph("投标人应提供独立专题报告。")
            tender_doc.save(tender)
            manifest = {
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": "tender.docx", "path": str(tender)}],
                "outputFile": str(root / "toc.json"),
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
            outline_runner.write_template_structure(manifest, manifest_path)

            batch = review_workflow.next_review_batch(root, max_chunks=2, max_chars=30_000)
            chunk_ids = batch["chunk_ids"]
            repeated_batch = review_workflow.next_review_batch(root, max_chunks=8, max_chars=60_000)
            all_chunk_ids = [chunk["chunk_id"] for chunk in json_load(root / "tender_review_chunks.json")["chunks"]]
            next_chunk_id = next(chunk_id for chunk_id in all_chunk_ids if chunk_id not in chunk_ids)
            table_id = next(
                block["table_id"]
                for chunk in batch["chunks"]
                for block in chunk["blocks"]
                if block["type"] == "table"
            )
            first_paragraph_id = next(
                block["evidence_id"]
                for chunk in batch["chunks"]
                for block in chunk["blocks"]
                if block["type"] == "paragraph" and "投标人应提供" in block["text"]
            )

            with self.assertRaisesRegex(SystemExit, "当前受控批次完全一致"):
                review_workflow.submit_batch_review(
                    root,
                    [*chunk_ids, next_chunk_id],
                    {"review_summary": "试图扩展当前批次。", "requirements": []},
                )
            with self.assertRaisesRegex(SystemExit, "table chunk must be fully read"):
                review_workflow.submit_batch_review(
                    root,
                    chunk_ids,
                    {"review_summary": "表格尚未读完。", "requirements": []},
                )

            tables = review_workflow.read_tables(
                root,
                [table_id],
                start=1,
                end=24,
                max_chars=8_000,
            )
            progress_batch = review_workflow.next_review_batch(root, max_chunks=8, max_chars=60_000)
            table_progress = next(
                block
                for chunk in progress_batch["chunks"]
                for block in chunk["blocks"]
                if block["type"] == "table"
            )
            with self.assertRaisesRegex(SystemExit, "单一目录节点"):
                review_workflow.submit_batch_review(
                    root,
                    chunk_ids,
                    {
                        "review_summary": "错误地把多个承接节点写在一个 target_node。",
                        "requirements": [
                            {
                                "evidence_ids": [first_paragraph_id],
                                "obligation": "投标人应提供总体技术方案",
                                "disposition": "map_existing",
                                "target_node": "第1章/第2章",
                            }
                        ],
                    },
                )
            first_submitted = review_workflow.submit_batch_review(
                root,
                chunk_ids,
                {
                    "review_summary": "逐项审阅本批段落和表格。",
                    "requirements": [
                        {
                            "evidence_ids": [first_paragraph_id],
                            "obligation": "投标人应提供总体技术方案",
                            "disposition": "map_existing",
                            "target_node": "第1章",
                        }
                    ],
                },
            )
            second_batch = review_workflow.next_review_batch(root, max_chunks=8, max_chars=30_000)
            second_paragraph_id = next(
                block["evidence_id"]
                for chunk in second_batch["chunks"]
                for block in chunk["blocks"]
                if block["type"] == "paragraph" and "独立专题报告" in block["text"]
            )
            second_submitted = review_workflow.submit_batch_review(
                root,
                second_batch["chunk_ids"],
                {
                    "review_summary": "逐项审阅第二批段落。",
                    "requirements": [
                        {
                            "evidence_ids": [second_paragraph_id],
                            "obligation": "投标人应提供独立专题报告",
                            "disposition": "suggest_add",
                            "target_node": "第1章",
                            "proposed_title": "独立专题报告",
                            "reason": "招标明确要求独立报告，模板无语义等价节点。",
                        },
                    ],
                },
            )
            status = review_workflow.review_status(root)

        self.assertEqual(len(batch["chunks"]), 2)
        self.assertEqual(repeated_batch["chunk_ids"], chunk_ids)
        self.assertEqual(tables["table_count"], 1)
        self.assertEqual(tables["tables"][0]["table"]["table_id"], table_id)
        self.assertTrue(table_progress["fully_read"])
        self.assertEqual(table_progress["read_ranges"], [{"start": 1, "end": 2}])
        self.assertEqual(table_progress["truncated_rows"], [])
        self.assertEqual(first_submitted["reviewed_batch_chunk_count"], len(chunk_ids))
        self.assertEqual(first_submitted["added_requirement_count"], 1)
        self.assertEqual(second_submitted["added_requirement_count"], 1)
        self.assertEqual(status["pending_chunk_count"], 0)
        self.assertEqual(status["unfinished_table_count"], 0)

    def test_bid_outline_template_structure_falls_back_to_toc_page_then_body_headings(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_page = root / "toc-page.docx"
            toc_doc = Document()
            toc_doc.add_paragraph("目录")
            toc_doc.add_paragraph("第1章 总体技术方案 ........ 1")
            toc_doc.add_paragraph("1.1 机组选型 ........ 2")
            toc_doc.save(toc_page)

            body = root / "body.docx"
            body_doc = Document()
            body_doc.add_paragraph("第1章 总体技术方案", style="Heading 1")
            body_doc.add_paragraph("1.1 机组选型", style="Heading 2")
            body_doc.save(body)

            toc_result = outline_runner.extract_template_structure(toc_page)
            body_result = outline_runner.extract_template_structure(body)

        self.assertEqual(toc_result["source"], "toc_page")
        self.assertEqual(body_result["source"], "body_headings")
        self.assertEqual([item["level"] for item in body_result["items"]], [1, 2])

    def test_bid_outline_generator_runner_only_validates_inputs_and_writes_no_final_toc(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            output = root / "toc.json"
            manifest_path = root / "s2_input.json"

            template_doc = Document()
            template_doc.add_paragraph("第1章 技术方案", style="Heading 1")
            template_doc.save(template)

            tender_doc = Document()
            tender_doc.add_paragraph("投标人应提供实施方案。")
            tender_doc.save(tender)

            manifest = {
                "projectId": "PRJ-TEST",
                "projectName": "测试项目",
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": "tender.docx", "path": str(tender)}],
                "outputFile": str(output),
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
            result = outline_runner.run_manifest(manifest, manifest_path)

            self.assertEqual(result["summary"]["schema_version"], "technical-outline-inputs.v1")
            self.assertEqual(result["summary"]["templateFile"], str(template))
            self.assertEqual(result["summary"]["tenderFileCount"], 1)
            self.assertEqual(result["summary"]["outputFile"], str(output))
            self.assertFalse(output.exists())
            self.assertFalse((root / "agent_review_input.json").exists())

    def test_bid_outline_generator_finalize_validates_existing_final_outputs(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")
        review_workflow = load_outline_script("review_workflow")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            output = root / "toc.json"
            manifest_path = root / "s2_input.json"

            template_doc = Document()
            template_doc.add_paragraph("第1章 总体技术方案", style="Heading 1")
            template_doc.save(template)
            tender_doc = Document()
            tender_doc.add_paragraph("投标人应提供总体技术方案")
            tender_doc.save(tender)
            manifest = {
                "projectId": "PRJ-TEST",
                "projectName": "test",
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": "tender.docx", "path": str(tender)}],
                "outputFile": str(output),
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
            outline_runner.write_template_structure(manifest, manifest_path)
            chunk = review_workflow.next_review_chunk(root)["chunk"]
            evidence_id = chunk["blocks"][0]["evidence_id"]
            review_workflow.submit_chunk_review(
                root,
                chunk["chunk_id"],
                {
                    "review_summary": "总体技术方案要求由模板根节点承接。",
                    "requirements": [
                        {
                            "evidence_ids": [evidence_id],
                            "obligation": "投标人应提供总体技术方案",
                            "disposition": "map_existing",
                            "target_node": "第1章",
                        }
                    ],
                },
            )
            output.write_text(
                json.dumps(
                    {
                        "schema_version": "technical-outline.v1",
                        "nodes": [
                            {
                                "number": "第1章",
                                "title": "总体技术方案",
                                "suggestion_action": "必要",
                                "suggestion_reason": "",
                                "tender_basis": {
                                    "file_id": "TEN-1",
                                    "search_text": "投标人应提供总体技术方案",
                                },
                                "children": [
                                    {
                                        "number": "1.1",
                                        "title": "机组选型",
                                        "suggestion_action": "建议增加",
                                        "suggestion_reason": "招标文件要求单独编制机组选型方案。",
                                        "children": [],
                                    }
                                ],
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            result = outline_runner.finalize_manifest(manifest, manifest_path)

        self.assertEqual(result["schema_version"], "technical-outline.v1")
        self.assertEqual(result["summary"]["workflowStage"], "finalized")
        self.assertEqual(result["summary"]["total_nodes"], 2)
        self.assertEqual(result["summary"]["action_counts"], {"必要": 1, "建议增加": 1})
        self.assertNotIn("evidenceFile", result)

    def test_bid_outline_finalize_rejects_unlocatable_tender_basis(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")
        review_workflow = load_outline_script("review_workflow")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            output = root / "toc.json"
            manifest_path = root / "s2_input.json"
            template_doc = Document()
            template_doc.add_paragraph("第1章 总体技术方案", style="Heading 1")
            template_doc.save(template)
            tender_doc = Document()
            tender_doc.add_paragraph("投标人应提供真实、完整的总体技术方案。")
            tender_doc.save(tender)
            manifest = {
                "workDir": str(root),
                "templateFile": str(template),
                "outputFile": str(output),
                "tenderFiles": [{"id": "TEN-1", "name": "tender.docx", "path": str(tender)}],
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
            outline_runner.write_template_structure(manifest, manifest_path)
            chunk = review_workflow.next_review_chunk(root)["chunk"]
            evidence_id = chunk["blocks"][0]["evidence_id"]
            review_workflow.submit_chunk_review(
                root,
                chunk["chunk_id"],
                {
                    "review_summary": "总体技术方案要求由模板根节点承接。",
                    "requirements": [
                        {
                            "evidence_ids": [evidence_id],
                            "obligation": "投标人应提供真实、完整的总体技术方案",
                            "disposition": "map_existing",
                            "target_node": "第1章",
                        }
                    ],
                },
            )
            payload = {
                "schema_version": "technical-outline.v1",
                "nodes": [
                    {
                        "number": "第1章",
                        "title": "总体技术方案",
                        "suggestion_action": "必要",
                        "suggestion_reason": "",
                        "tender_basis": {"file_id": "TEN-1", "search_text": "原文不存在的编造依据"},
                        "children": [],
                    }
                ],
            }
            output.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

            with self.assertRaisesRegex(SystemExit, "search_text 无法在 tenderFile 定位"):
                outline_runner.finalize_manifest(manifest, manifest_path)

            payload["nodes"][0]["tender_basis"]["search_text"] = "投标人应提供真实、完整的总体技术方案"
            output.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            result = outline_runner.finalize_manifest(manifest, manifest_path)

        self.assertEqual(result["summary"]["total_nodes"], 1)

    def test_bid_outline_finalize_rejects_necessary_appendix_absent_from_template(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            output = root / "toc.json"
            manifest_path = root / "s2_input.json"
            manifest = {"workDir": str(root), "outputFile": str(output)}
            manifest_path.write_text(json.dumps(manifest), encoding="utf-8")
            (root / "template_structure.json").write_text(
                json.dumps(
                    {
                        "schema_version": "template-structure.v1",
                        "items": [{"number": "第1章", "title": "技术方案", "level": 1}],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            appendix_inventory = {
                "schema_version": "tender-appendix-inventory.v1",
                "items": [
                    {
                        "file_id": "TEN-1",
                        "file_name": "tender.docx",
                        "number": "附表A.1",
                        "title": "投标机型总方案信息表",
                        "raw_text": "附表A.1 投标机型总方案信息表",
                        "paragraph_index": 20,
                        "following_table_count": 0,
                        "following_text_count": 0,
                    },
                    {
                        "file_id": "TEN-1",
                        "file_name": "tender.docx",
                        "number": "附表B.1",
                        "title": "另一张实际表单",
                        "raw_text": "附表B.1 另一张实际表单",
                        "paragraph_index": 30,
                        "following_table_count": 1,
                        "following_text_count": 0,
                    },
                ],
            }
            (root / "tender_appendix_inventory.json").write_text(
                json.dumps(appendix_inventory, ensure_ascii=False),
                encoding="utf-8",
            )
            output.write_text(
                json.dumps(
                    {
                        "schema_version": "technical-outline.v1",
                        "nodes": [
                            {
                                "number": "第1章",
                                "title": "技术方案",
                                "suggestion_action": "必要",
                                "suggestion_reason": "",
                                "children": [],
                            },
                            {
                                "number": "第2章",
                                "title": "技术附表",
                                "suggestion_action": "必要",
                                "suggestion_reason": "",
                                "children": [
                                    {
                                        "number": "附表A.1",
                                        "title": "投标机型总方案信息表",
                                        "suggestion_action": "必要",
                                        "suggestion_reason": "",
                                        "children": [],
                                    }
                                ],
                            },
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            with self.assertRaisesRegex(SystemExit, "模板目录不存在.*建议增加"):
                outline_runner.finalize_manifest(manifest, manifest_path)

            payload = json_load(output)
            appendix = payload["nodes"][-1]
            appendix["suggestion_action"] = "建议增加"
            appendix["suggestion_reason"] = "招标文件新增技术附表。"
            appendix["children"][0]["suggestion_action"] = "建议增加"
            appendix["children"][0]["suggestion_reason"] = "招标文件新增独立表单。"
            output.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

            with self.assertRaisesRegex(SystemExit, "没有独立填写表格"):
                outline_runner.finalize_manifest(manifest, manifest_path)

            appendix_inventory["items"][0]["following_table_count"] = 1
            (root / "tender_appendix_inventory.json").write_text(
                json.dumps(appendix_inventory, ensure_ascii=False),
                encoding="utf-8",
            )
            appendix["children"].append(
                {
                    "number": "附表B.1",
                    "title": "另一张实际表单",
                    "suggestion_action": "建议增加",
                    "suggestion_reason": "招标文件新增独立表单。",
                    "children": [],
                }
            )
            output.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            result = outline_runner.finalize_manifest(manifest, manifest_path)

        self.assertEqual(result["summary"]["action_counts"], {"必要": 1, "建议增加": 3})

    def test_bid_outline_cli_accepts_navigation_options_after_manifest(self) -> None:
        outline_runner = load_outline_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            tender = root / "tender.docx"
            manifest_path = root / "s2_input.json"
            template_doc = Document()
            template_doc.add_paragraph("第1章 技术方案", style="Heading 1")
            template_doc.save(template)
            tender_doc = Document()
            tender_doc.add_paragraph("投标人应提供总体技术方案。")
            tender_doc.save(tender)
            manifest = {
                "workDir": str(root),
                "templateFile": str(template),
                "tenderFiles": [{"id": "TEN-1", "name": "tender.docx", "path": str(tender)}],
                "outputFile": str(root / "toc.json"),
            }
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
            outline_runner.write_template_structure(manifest, manifest_path)
            chunk = outline_runner.dispatch_command("next", manifest, manifest_path, [])["chunk"]
            evidence_id = chunk["blocks"][0]["evidence_id"]

            with patch.object(
                sys,
                "argv",
                [
                    "run_from_manifest.py",
                    "window",
                    str(manifest_path),
                    evidence_id,
                    "--before",
                    "0",
                    "--after",
                    "0",
                ],
            ), patch("builtins.print") as print_mock:
                result = outline_runner.main()

        self.assertEqual(result, 0)
        payload = json.loads(print_mock.call_args.args[0])
        self.assertEqual(payload["center"], evidence_id)
        self.assertEqual(len(payload["blocks"]), 1)

    def test_bid_assembler_parse_toc_accepts_current_s2_json(self) -> None:
        parse_toc = load_assembler_script("parse_toc")

        with tempfile.TemporaryDirectory() as tmp:
            toc_path = Path(tmp) / "投标文件-总目录.json"
            toc_path.write_text(
                """{
  "schema_version": "bid-toc-json-v1",
  "document_title": "测试项目投标文件总目录",
  "items": [
    {"order": 1, "level": 1, "number": "1", "title": "项目概况", "annotation": "保留"},
    {"order": 2, "level": 2, "number": "1.1", "title": "项目背景", "annotation": "适配"},
    {"order": 3, "level": 1, "number": "附表", "title": "", "annotation": "保留"}
  ]
}""",
                encoding="utf-8",
            )

            entries = parse_toc.parse_toc_json(toc_path)

        self.assertEqual(entries[0]["chapter_no_flat"], "1")
        self.assertEqual(entries[0]["title"], "项目概况")
        self.assertEqual(entries[1]["tag"], "适配")
        self.assertEqual(entries[2]["title"], "附表")

    def test_bid_table_filler_preserves_generic_appendix_and_highlights_manual_cells(self) -> None:
        table_filler = load_table_filler_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            blank = tmp_path / "附表A.1 投标机型总方案信息表.docx"
            doc = Document()
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            for index, text in enumerate(["序号", "项目", "投标响应", "备注"]):
                table.cell(0, index).text = text
            for row in (
                ["1", "投标机型", "", ""],
                ["2", "叶轮直径", "", "m"],
                ["3", "项目业主", "", ""],
                ["4", "场址空气密度", "", ""],
            ):
                cells = table.add_row().cells
                for index, text in enumerate(row):
                    cells[index].text = text
            doc.save(blank)

            manifest_path = tmp_path / "manifest.json"
            output = tmp_path / "filled.docx"
            manifest_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-tech-table-fill-v1",
                        "appendixTask": {
                            "id": "APPX-A1",
                            "title": "附表A.1 投标机型总方案信息表",
                            "docxPath": str(blank),
                        },
                        "projectTurbineModel": {
                            "model": "EW10.0-220上置",
                            "ratedPowerKw": 10000,
                            "rotorDiameterM": 220,
                        },
                        "parseFields": [{"id": "OWNER", "label": "项目业主", "value": "测试业主"}],
                        "outputFile": str(output),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = table_filler.run_from_manifest(manifest_path)

            self.assertEqual(result["schema_version"], "bid-tech-table-fill-v1")
            self.assertTrue(output.exists())
            self.assertEqual(result["fillReport"]["filledFieldCount"], 3)
            self.assertEqual(result["fillReport"]["unfilledFieldCount"], 1)
            filled_doc = Document(str(output))
            rows = filled_doc.tables[0].rows
            self.assertEqual(rows[1].cells[2].text, "EW10.0-220")
            self.assertEqual(rows[2].cells[2].text, "220")
            self.assertEqual(rows[3].cells[2].text, "测试业主")
            self.assertEqual(rows[4].cells[2].text, "[待人工补充：场址空气密度]")
            shd = rows[4].cells[2]._tc.tcPr.find(qn("w:shd"))
            self.assertIsNotNone(shd)
            self.assertEqual(shd.get(qn("w:fill")), "FFF2CC")

    def test_bid_word_placeholder_filler_replaces_parse_and_project_placeholders(self) -> None:
        word_filler = load_word_filler_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            template = tmp_path / "待填写-投标说明函.docx"
            doc = Document()
            doc.add_paragraph("投标机型：[投标机型，待填写]")
            doc.add_paragraph("项目要求：[招标要求，待填写]")
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            table.cell(0, 0).text = "项目名称"
            table.cell(0, 1).text = "[项目名称，待填写]"
            doc.save(template)

            manifest_path = tmp_path / "manifest.json"
            output = tmp_path / "filled.docx"
            manifest_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-tech-word-placeholder-fill-v1",
                        "projectName": "测试风电项目",
                        "projectTurbineModel": {"model": "EW10.0-220上置"},
                        "blankSource": {
                            "id": "RAW-TEMPLATE",
                            "title": "待填写-投标说明函.docx",
                            "docxPath": str(template),
                            "placeholderCount": 3,
                        },
                        "parseFields": [
                            {
                                "id": "REQ-001",
                                "label": "招标要求",
                                "value": "满足招标文件所有技术要求",
                                "sourceFile": "招标文件.docx",
                            }
                        ],
                        "outputFile": str(output),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = word_filler.run_from_manifest(manifest_path)

            self.assertEqual(result["schema_version"], "bid-tech-word-placeholder-fill-v1")
            self.assertEqual(result["fillReport"]["placeholderCount"], 3)
            self.assertEqual(result["fillReport"]["filledPlaceholderCount"], 3)
            self.assertEqual(result["unfilledFields"], [])
            filled_doc = Document(str(output))
            self.assertEqual(filled_doc.paragraphs[0].text, "投标机型：EW10.0-220上置")
            self.assertEqual(filled_doc.paragraphs[1].text, "项目要求：满足招标文件所有技术要求")
            self.assertEqual(filled_doc.tables[0].cell(0, 1).text, "测试风电项目")

    def test_bid_word_placeholder_filler_uses_project_identity_for_owner_placeholders(self) -> None:
        word_filler = load_word_filler_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            template = tmp_path / "固定-风电机组自主可控推广应用的承诺.docx"
            doc = Document()
            doc.add_paragraph("招标方：[招标方，待填写]")
            doc.add_paragraph("日期：[日期，待填写]")
            doc.save(template)

            manifest_path = tmp_path / "manifest.json"
            output = tmp_path / "filled.docx"
            manifest_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-tech-word-placeholder-fill-v1",
                        "projectName": "测试风电项目",
                        "projectIdentity": {
                            "owner": "华能集团",
                            "customerName": "华能集团",
                        },
                        "blankSource": {
                            "id": "RAW-0452",
                            "title": "固定-风电机组自主可控推广应用的承诺.docx",
                            "docxPath": str(template),
                            "placeholderCount": 2,
                        },
                        "outputFile": str(output),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = word_filler.run_from_manifest(manifest_path)

            self.assertEqual(result["fillReport"]["placeholderCount"], 2)
            self.assertEqual(result["fillReport"]["filledPlaceholderCount"], 2)
            self.assertEqual(result["unfilledFields"], [])
            filled_text = "\n".join(paragraph.text for paragraph in Document(str(output)).paragraphs)
            self.assertIn("招标方：华能集团", filled_text)
            self.assertNotIn("待填写", filled_text)

    def test_bid_word_placeholder_filler_uses_manufacturing_base_intro_from_materials(self) -> None:
        word_filler = load_word_filler_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            template = tmp_path / "待填写-供货保障能力.docx"
            doc = Document()
            doc.add_paragraph("本项目主机供货制造基地-[基地名称，待填写]")
            doc.add_paragraph("[基地介绍，待填写]生产基地情况详见5.11.1 本项目主机供货制造基地")
            doc.add_paragraph("本项目叶片供货制造基地-[基地名称，待填写]")
            doc.add_paragraph("[基地介绍，待填写]生产基地情况详见5.11.2 本项目叶片供货制造基地")
            doc.save(template)

            material = tmp_path / "固定-上海电气生产能力介绍.docx"
            source_doc = Document()
            source_doc.add_paragraph("上海电气生产能力介绍")
            source_doc.add_paragraph(
                "上海电气风电(张掖)叶片科技有限公司作为上海电气风电集团股份有限公司子公司成立于2021年12月24日，"
                "公司位于高台县南华镇工业园区。公司主要生产5XMW及以上风力发电机组和叶片，配套年产能500台机组、500套叶片。"
            )
            source_doc.add_paragraph(
                "上海电气高台生产基地于2023年6月投入生产运营，产品订单覆盖西北五省，目前已经累计生产5MW及上风力发电机组100台套。"
            )
            source_doc.save(material)

            manifest_path = tmp_path / "manifest.json"
            output = tmp_path / "filled.docx"
            manifest_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-tech-word-placeholder-fill-v1",
                        "blankSource": {
                            "id": "RAW-SUPPLY",
                            "title": template.name,
                            "docxPath": str(template),
                            "placeholderCount": 4,
                        },
                        "referenceMaterials": [
                            {"id": "RAW-PROD", "name": material.name, "path": str(material), "materialTier": "standard"}
                        ],
                        "outputFile": str(output),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = word_filler.run_from_manifest(manifest_path)

            self.assertEqual(result["unfilledFields"], [])
            filled_text = "\n".join(paragraph.text for paragraph in Document(str(output)).paragraphs)
            self.assertIn("本项目主机供货制造基地-上海电气高台生产基地", filled_text)
            self.assertIn("上海电气高台生产基地于2023年6月投入生产运营", filled_text)
            self.assertIn("本项目叶片供货制造基地-上海电气风电（张掖）叶片科技有限公司", filled_text)
            self.assertIn("配套年产能500台机组、500套叶片", filled_text)
            self.assertNotIn("待人工补充", filled_text)

    def test_bid_word_placeholder_filler_summarizes_wind_resource_report_from_project_facts(self) -> None:
        word_filler = load_word_filler_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            template = tmp_path / "风资源评估与机位排布方案.docx"
            doc = Document()
            doc.add_paragraph("风资源概况：[风资源报告，待填写]")
            doc.save(template)

            manifest_path = tmp_path / "manifest.json"
            output = tmp_path / "filled.docx"
            manifest_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-tech-word-placeholder-fill-v1",
                        "blankSource": {
                            "id": "RAW-WIND",
                            "title": "风资源评估与机位排布方案.docx",
                            "docxPath": str(template),
                            "placeholderCount": 1,
                        },
                        "projectFactTable": {
                            "status": "confirmed",
                            "fields": [
                                {"label": "轮毂高度", "value": "125", "unit": "m"},
                                {"label": "年平均风速", "value": "7.22", "unit": "m/s"},
                                {"label": "空气密度", "value": "1.16", "unit": "kg/m3"},
                                {"label": "湍流强度", "value": "0.1"},
                                {"label": "风剪切", "value": "0.1"},
                                {"label": "极端风速", "value": "53.51", "unit": "m/s"},
                                {"label": "安全等级", "value": "IEC S"},
                            ],
                        },
                        "outputFile": str(output),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = word_filler.run_from_manifest(manifest_path)

            self.assertEqual(result["fillReport"]["filledPlaceholderCount"], 1)
            self.assertEqual(result["unfilledFields"], [])
            filled_text = "\n".join(paragraph.text for paragraph in Document(str(output)).paragraphs)
            self.assertIn("7.22", filled_text)
            self.assertIn("1.16", filled_text)
            self.assertIn("IEC S", filled_text)
            self.assertNotIn("待填写", filled_text)

    def test_bid_word_placeholder_filler_uses_table_row_label_before_model_column_header(self) -> None:
        word_filler = load_word_filler_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            template = tmp_path / "待填写-投标关键数据一览表.docx"
            doc = Document()
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            for index, text in enumerate(["项目名称", "[项目名称，待填写]", "[项目名称，待填写]"]):
                table.cell(0, index).text = text
            for row in (
                ["承诺方式", "[投标方案，待填写]", "[投标方案，待填写]"],
                ["方案", "[投标方案，待填写]", "[投标方案，待填写]"],
                ["机型", "[投标方案，待填写]", "[投标方案，待填写]"],
                ["轮毂高度（m）", "[投标方案，待填写]", "[投标方案，待填写]"],
                ["台数", "[投标方案，待填写]", "[投标方案，待填写]"],
                ["容量（MW）", "[投标方案，待填写]", "[投标方案，待填写]"],
                ["净发电量（MWh/y）", "[投标方案，待填写]", "[投标方案，待填写]"],
                ["有效小时数（h）", "[投标方案，待填写]", "[投标方案，待填写]"],
            ):
                cells = table.add_row().cells
                for index, text in enumerate(row):
                    cells[index].text = text
            guarantee = doc.add_table(rows=1, cols=2)
            guarantee.style = "Table Grid"
            guarantee.cell(0, 0).text = "保证项"
            guarantee.cell(0, 1).text = "保证率"
            for row in (
                ["单台机组功率曲线保证率（%）", "[投标方案，待填写]"],
                ["单台机组时间可利用率保证值（%）", "[投标方案，待填写]"],
                ["全场机组时间可利用率保证值）（%）", "[投标方案，待填写]"],
            ):
                cells = guarantee.add_row().cells
                cells[0].text, cells[1].text = row
            doc.save(template)

            manifest_path = tmp_path / "manifest.json"
            output = tmp_path / "filled.docx"
            manifest_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-tech-word-placeholder-fill-v1",
                        "projectName": "测试项目",
                        "blankSource": {
                            "id": "RAW-KEY-DATA",
                            "title": "待填写-投标关键数据一览表.docx",
                            "docxPath": str(template),
                            "placeholderCount": 18,
                        },
                        "projectFactTable": {
                            "status": "confirmed",
                            "fields": [
                                {"label": "项目名称", "value": "测试项目"},
                                {"label": "投标方案", "value": "EW10.0-220-125"},
                                {"label": "投标机型", "value": "EW10.0-220上置"},
                                {"label": "轮毂高度", "value": "125", "unit": "m"},
                                {"label": "机组台数", "value": "60", "unit": "台"},
                                {"label": "总装机容量", "value": "600", "unit": "MW"},
                                {"label": "保证发电量", "value": "1701601", "unit": "MWh"},
                                {"label": "保证有效小时数", "value": "2836", "unit": "h"},
                                {"label": "功率曲线保证率", "value": "97%", "unit": "%"},
                                {"label": "单台可利用率", "value": "95%", "unit": "%"},
                                {"label": "全场可利用率", "value": "98%", "unit": "%"},
                            ],
                        },
                        "outputFile": str(output),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = word_filler.run_from_manifest(manifest_path)

            self.assertEqual(result["unfilledFields"], [])
            filled_doc = Document(str(output))
            rows = filled_doc.tables[0].rows
            self.assertEqual(rows[1].cells[1].text, "承诺保证值（75%折减）")
            self.assertEqual(rows[2].cells[1].text, "EW10.0-220-125")
            self.assertEqual(rows[3].cells[1].text, "EW10.0-220-125")
            self.assertEqual(rows[4].cells[1].text, "125")
            self.assertEqual(rows[5].cells[1].text, "60")
            self.assertEqual(rows[6].cells[1].text, "600")
            self.assertEqual(rows[7].cells[1].text, "1701601")
            self.assertEqual(rows[8].cells[1].text, "2836")
            guarantee_rows = filled_doc.tables[1].rows
            self.assertEqual(guarantee_rows[1].cells[1].text, "97%")
            self.assertEqual(guarantee_rows[2].cells[1].text, "95%")
            self.assertEqual(guarantee_rows[3].cells[1].text, "98%")
            self.assertEqual(result["fillReport"]["semanticCheckCount"], 21)
            self.assertEqual(result["fillReport"]["semanticFailedCount"], 0)
            self.assertEqual(result["fillReport"]["semanticValidationRate"], 1.0)

    def test_bid_table_filler_uses_first_empty_response_column_for_multi_model_tables(self) -> None:
        table_filler = load_table_filler_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            blank = tmp_path / "附表A.1 投标机型总方案信息表.docx"
            blank_doc = Document()
            table = blank_doc.add_table(rows=1, cols=6)
            table.style = "Table Grid"
            for index, text in enumerate(["编号", "项目", "项目", "投标机型1", "投标机型2", "备注"]):
                table.cell(0, index).text = text
            for row in (
                ["1", "投标机型", "投标机型", "", "", ""],
                ["2", "机组台数", "机组台数", "", "", ""],
            ):
                cells = table.add_row().cells
                for index, text in enumerate(row):
                    cells[index].text = text
            blank_doc.save(blank)

            manifest_path = tmp_path / "manifest.json"
            output = tmp_path / "filled.docx"
            manifest_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-tech-table-fill-v1",
                        "appendixTask": {
                            "id": "APPX-A1",
                            "title": "附表A.1 投标机型总方案信息表",
                            "docxPath": str(blank),
                        },
                        "projectTurbineModel": {"model": "EW10.0-220上置", "ratedPowerKw": 10000},
                        "parseFields": [{"id": "REQ-SCALE", "label": "标段规模", "value": "600MW"}],
                        "outputFile": str(output),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            table_filler.run_from_manifest(manifest_path)

            filled_doc = Document(str(output))
            rows = filled_doc.tables[0].rows
            self.assertEqual(rows[1].cells[3].text, "EW10.0-220")
            self.assertEqual(rows[1].cells[4].text, "")
            self.assertEqual(rows[2].cells[3].text, "60")

    def test_bid_table_filler_uses_bidder_response_column_and_requirement_fallback(self) -> None:
        table_filler = load_table_filler_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            blank = tmp_path / "附表E.2 风电场折减系数表.docx"
            blank_doc = Document()
            table = blank_doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            for index, text in enumerate(["风电场折减项目", "招标人要求值", "投标人响应值"]):
                table.cell(0, index).text = text
            for row in (
                ["尾流折减", "97%", ""],
                ["空气密度折减", "根据各厂家测算结果确定", ""],
            ):
                cells = table.add_row().cells
                for index, text in enumerate(row):
                    cells[index].text = text
            blank_doc.save(blank)

            manifest_path = tmp_path / "manifest.json"
            output = tmp_path / "filled.docx"
            manifest_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-tech-table-fill-v1",
                        "appendixTask": {
                            "id": "APPX-E2",
                            "title": "附表E.2 风电场折减系数表",
                            "docxPath": str(blank),
                        },
                        "outputFile": str(output),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = table_filler.run_from_manifest(manifest_path)

            self.assertEqual(result["fillReport"]["targetFieldCount"], 2)
            self.assertEqual(result["fillReport"]["filledFieldCount"], 1)
            self.assertEqual(result["fillReport"]["unfilledFieldCount"], 1)
            filled_doc = Document(str(output))
            rows = filled_doc.tables[0].rows
            self.assertEqual(rows[1].cells[2].text, "97%")
            self.assertEqual(rows[2].cells[2].text, "[待人工补充：空气密度折减]")
            shd = rows[2].cells[2]._tc.tcPr.find(qn("w:shd"))
            self.assertIsNotNone(shd)

    def test_bid_table_filler_counts_only_empty_or_placeholder_response_cells(self) -> None:
        table_filler = load_table_filler_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            blank = tmp_path / "附表A.1 投标机型总方案信息表.docx"
            blank_doc = Document()
            table = blank_doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            for index, text in enumerate(["编号", "项目", "投标响应"]):
                table.cell(0, index).text = text
            for row in (
                ["1", "投标机型", "已填写机型"],
                ["2", "轮毂高度", ""],
            ):
                cells = table.add_row().cells
                for index, text in enumerate(row):
                    cells[index].text = text
            blank_doc.save(blank)

            manifest_path = tmp_path / "manifest.json"
            output = tmp_path / "filled.docx"
            manifest_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-tech-table-fill-v1",
                        "appendixTask": {
                            "id": "APPX-A1",
                            "title": "附表A.1 投标机型总方案信息表",
                            "docxPath": str(blank),
                        },
                        "projectFactTable": {
                            "status": "confirmed",
                            "fields": [
                                {"label": "轮毂高度", "value": "125", "unit": "m"},
                                {"label": "投标机型", "value": "EW10.0-220上置"},
                            ],
                        },
                        "outputFile": str(output),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = table_filler.run_from_manifest(manifest_path)

            self.assertEqual(result["fillReport"]["targetFieldCount"], 1)
            self.assertEqual(result["fillReport"]["filledFieldCount"], 1)
            filled_doc = Document(str(output))
            rows = filled_doc.tables[0].rows
            self.assertEqual(rows[1].cells[2].text, "已填写机型")
            self.assertEqual(rows[2].cells[2].text, "125")

    def test_bid_table_filler_auto_selects_non_c_appendix_sources_from_material_index(self) -> None:
        table_filler = load_table_filler_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            blank = tmp_path / "附表B.2 供货范围响应表.docx"
            blank_doc = Document()
            table = blank_doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            for index, text in enumerate(["序号", "项目", "投标响应", "备注"]):
                table.cell(0, index).text = text
            for row in (
                ["1", "风力发电机组", "", ""],
                ["2", "塔筒", "", ""],
                ["3", "专用工具", "", ""],
            ):
                cells = table.add_row().cells
                for index, text in enumerate(row):
                    cells[index].text = text
            blank_doc.save(blank)

            supply = tmp_path / "固定-供货范围清单.docx"
            supply_doc = Document()
            source_table = supply_doc.add_table(rows=1, cols=2)
            source_table.style = "Table Grid"
            source_table.cell(0, 0).text = "项目"
            source_table.cell(0, 1).text = "投标响应"
            for row in (
                ["风力发电机组", "提供 EW10.0-220 风力发电机组"],
                ["塔筒", "包含配套钢塔筒"],
                ["专用工具", "提供安装维护专用工具一套"],
            ):
                cells = source_table.add_row().cells
                cells[0].text = row[0]
                cells[1].text = row[1]
            supply_doc.save(supply)

            wrong = tmp_path / "固定-风资源评估报告.docx"
            wrong_doc = Document()
            wrong_doc.add_paragraph("平均风速：7.2m/s")
            wrong_doc.save(wrong)

            manifest_path = tmp_path / "manifest.json"
            output = tmp_path / "filled.docx"
            manifest_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-tech-table-fill-v1",
                        "appendixTask": {
                            "id": "APPX-B2",
                            "title": "附表B.2 供货范围响应表",
                            "docxPath": str(blank),
                        },
                        "materialIndex": [
                            {
                                "id": "RAW-WIND",
                                "name": "固定-风资源评估报告.docx",
                                "folderPath": "技术标/项目素材/风资源评估",
                                "path": str(wrong),
                                "materialTier": "project",
                            },
                            {
                                "id": "RAW-SUPPLY",
                                "name": "固定-供货范围清单.docx",
                                "folderPath": "技术标/通用素材/供货范围",
                                "path": str(supply),
                                "materialTier": "standard",
                            },
                        ],
                        "recommendedMaterials": [
                            {
                                "id": "RAW-WIND",
                                "name": "固定-风资源评估报告.docx",
                                "path": str(wrong),
                            }
                        ],
                        "outputFile": str(output),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = table_filler.run_from_manifest(manifest_path)

            self.assertEqual(result["fillReport"]["filledFieldCount"], 3)
            selected_ids = [item["id"] for item in result["fillReport"]["sourceSelection"]["selected"]]
            self.assertIn("RAW-SUPPLY", selected_ids)
            self.assertEqual(result["fillReport"]["referenceSources"][0]["id"], "RAW-SUPPLY")
            self.assertEqual(result["fillReport"]["referenceSources"][0]["route"], "Wiki/materialIndex 自动选材")
            filled_doc = Document(str(output))
            rows = filled_doc.tables[0].rows
            self.assertEqual(rows[1].cells[2].text, "提供 EW10.0-220 风力发电机组")
            self.assertEqual(rows[2].cells[2].text, "包含配套钢塔筒")
            self.assertEqual(rows[3].cells[2].text, "提供安装维护专用工具一套")

    def test_bid_gap_planner_routes_appendix_sources_by_customer_matrix(self) -> None:
        gap_runner = load_gap_planner_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_path = root / "toc.json"
            parse_path = root / "parse.json"
            output_path = root / "gap_plan.json"
            toc_path.write_text(
                json.dumps(
                    {
                        "items": [
                            {"number": "附表C.1", "title": "附表C.1 总体技术参数与规格", "level": 2},
                            {"number": "附表D.3", "title": "附表D.3 功率曲线", "level": 2},
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            blank = root / "附表C.1 总体技术参数与规格.docx"
            doc = Document()
            table = doc.add_table(rows=1, cols=4)
            table.cell(0, 0).text = "序号"
            table.cell(0, 1).text = "参数名称"
            table.cell(0, 2).text = "投标人响应值"
            table.cell(0, 3).text = "单位"
            doc.save(blank)
            range_blank = root / "附表D.3 功率曲线.docx"
            doc.save(range_blank)
            parse_path.write_text(
                json.dumps(
                    {
                        "structured": {
                            "appendices": [
                                {
                                    "id": "APPX-C1",
                                    "title": "附表C.1 总体技术参数与规格",
                                    "docxPath": str(blank),
                                },
                                {
                                    "id": "APPX-D3",
                                    "title": "附表D.3 功率曲线",
                                    "docxPath": str(range_blank),
                                }
                            ]
                        }
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            manifest_path = root / "manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-HN",
                        "projectName": "华能项目",
                        "bidType": "技术标",
                        "customerName": "华能",
                        "tocJsonPath": str(toc_path),
                        "parseResultPath": str(parse_path),
                        "materialScope": {"paths": ["技术标/标准文件", "技术标/项目定制"]},
                        "appendixSourceMatrix": {
                            "rows": [
                                {
                                    "id": "Sheet1!R35",
                                    "customer": "华能",
                                    "tableTitle": "附表C.1 总体技术参数与规格",
                                    "projectSources": [],
                                    "standardSources": ["机型参数表"],
                                    "otherSources": [],
                                },
                                {
                                    "id": "Sheet1!R40",
                                    "customer": "华能",
                                    "tableTitle": "附表D.1-D.6",
                                    "projectSources": ["功率曲线"],
                                    "standardSources": [],
                                    "otherSources": [],
                                }
                            ]
                        },
                        "materialIndex": [
                            {
                                "id": "RAW-WIND",
                                "name": "风资源评估报告.docx",
                                "folderPath": "技术标/项目定制/风资源评估报告",
                                "materialTier": "project",
                            },
                            {
                                "id": "RAW-PARAM",
                                "name": "X2平台机型投标参数_20250106.xlsx",
                                "folderPath": "技术标/标准文件/机型参数表",
                                "materialTier": "standard",
                            },
                            {
                                "id": "RAW-POWER",
                                "name": "项目功率曲线.xlsx",
                                "folderPath": "技术标/项目定制/功率曲线",
                                "materialTier": "project",
                            },
                        ],
                        "outputFile": str(output_path),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = gap_runner.build_gap_plan(json_load(manifest_path))

            c_item = result["items"][0]
            c_task = c_item["appendixTasks"][0]
            self.assertEqual(c_item["usage"], "appendix_fill")
            self.assertEqual(c_task["sourceRouting"]["source"], "appendix_source_matrix")
            self.assertEqual(c_task["sourceRouting"]["standardSources"], ["机型参数表"])
            self.assertEqual(c_task["recommendedMaterials"][0]["id"], "RAW-PARAM")
            self.assertIn("standard 来源规定命中", c_task["recommendedMaterials"][0]["matchReason"])

            d_item = result["items"][1]
            d_task = d_item["appendixTasks"][0]
            self.assertEqual(d_task["sourceRouting"]["ruleId"], "Sheet1!R40")
            self.assertEqual(d_task["sourceRouting"]["projectSources"], ["功率曲线"])
            self.assertEqual(d_task["recommendedMaterials"][0]["id"], "RAW-POWER")
            self.assertIn("project 来源规定命中", d_task["recommendedMaterials"][0]["matchReason"])

    def test_bid_gap_planner_routes_tender_rule_to_all_project_tender_documents(self) -> None:
        gap_runner = load_gap_planner_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_path = root / "toc.json"
            parse_path = root / "parse.json"
            blank = root / "附表B.6 技术服务响应表.docx"
            Document().save(blank)
            toc_path.write_text(
                json.dumps({"items": [{"number": "附表B.6", "title": "附表B.6 技术服务响应表", "level": 2}]}, ensure_ascii=False),
                encoding="utf-8",
            )
            parse_path.write_text(
                json.dumps(
                    {
                        "documents": [
                            {"id": "TEN-1", "name": "技术规范.pdf", "status": "completed"},
                            {"id": "TEN-2", "name": "招标附图.docx", "status": "completed"},
                            {"id": "TEN-3", "name": "损坏文件.pdf", "status": "failed"},
                        ],
                        "structured": {
                            "appendices": [
                                {"id": "APPX-B6", "title": "附表B.6 技术服务响应表", "docxPath": str(blank)}
                            ]
                        },
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            plan = gap_runner.build_gap_plan(
                {
                    "projectId": "PRJ-TENDER",
                    "projectName": "招标文件填写项目",
                    "tocJsonPath": str(toc_path),
                    "parseResultPath": str(parse_path),
                    "materialIndex": [],
                    "appendixSourceMatrix": {
                        "rows": [
                            {
                                "id": "Sheet1!R8",
                                "tableTitle": "附表B.6 技术服务响应表",
                                "projectSources": [],
                                "standardSources": [],
                                "otherSources": ["响应招标文件填写"],
                            }
                        ]
                    },
                }
            )

        routing = plan["items"][0]["appendixTasks"][0]["sourceRouting"]
        self.assertEqual(routing["status"], "tender_parse_fields")
        self.assertTrue(routing["useTenderParseFields"])
        self.assertEqual(routing["tenderDocumentStatus"], "available")
        self.assertEqual(routing["tenderDocumentCount"], 2)
        self.assertEqual([item["name"] for item in routing["tenderDocuments"]], ["技术规范.pdf", "招标附图.docx"])

    def test_bid_gap_planner_parent_rule_covers_sub_numbered_appendix(self) -> None:
        """规则只写父级编号（附表F.2）时应覆盖子编号附表（F.2.1/F.2.2）。"""
        gap_runner = load_gap_planner_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_path = root / "toc.json"
            parse_path = root / "parse.json"
            output_path = root / "gap_plan.json"
            toc_path.write_text(
                json.dumps(
                    {
                        "items": [
                            {"number": "附表F.2.1", "title": "附表F.2.1 投标机组设计认证", "level": 2},
                            {"number": "附表G.2.1", "title": "附表G.2.1 场址载荷仿真关键计算方法", "level": 2},
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            blank = root / "附表F.2.1 投标机组设计认证.docx"
            Document().save(blank)
            other_blank = root / "附表G.2.1 场址载荷仿真关键计算方法.docx"
            Document().save(other_blank)
            parse_path.write_text(
                json.dumps(
                    {
                        "structured": {
                            "appendices": [
                                {"id": "APPX-F21", "title": "附表F.2.1 投标机组设计认证", "docxPath": str(blank)},
                                {
                                    "id": "APPX-G21",
                                    "title": "附表G.2.1 场址载荷仿真关键计算方法",
                                    "docxPath": str(other_blank),
                                },
                            ]
                        }
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            manifest_path = root / "manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-HN",
                        "projectName": "华能项目",
                        "bidType": "技术标",
                        "customerName": "华能",
                        "tocJsonPath": str(toc_path),
                        "parseResultPath": str(parse_path),
                        "materialScope": {"paths": ["技术标/标准文件", "技术标/项目定制"]},
                        "appendixSourceMatrix": {
                            "rows": [
                                {
                                    "id": "Sheet1!R33",
                                    "customer": "华能",
                                    "tableTitle": "附表F.2 投标机型整机认证",
                                    "projectSources": [],
                                    "standardSources": ["认证证书"],
                                    "otherSources": [],
                                }
                            ]
                        },
                        "materialIndex": [
                            {
                                "id": "RAW-CERT",
                                "name": "EW10.0-220上置型式认证证书.pdf",
                                "folderPath": "技术标/标准文件/EW10.0-220上置/认证证书",
                                "materialTier": "standard",
                            }
                        ],
                        "outputFile": str(output_path),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = gap_runner.build_gap_plan(json_load(manifest_path))

            f_item = result["items"][0]
            f_task = f_item["appendixTasks"][0]
            self.assertEqual(f_task["sourceRouting"]["source"], "appendix_source_matrix")
            self.assertEqual(f_task["sourceRouting"]["ruleId"], "Sheet1!R33")
            self.assertEqual(f_task["sourceRouting"]["standardSources"], ["认证证书"])
            self.assertEqual(f_task["recommendedMaterials"][0]["id"], "RAW-CERT")
            # 前缀不同的子编号（G.2.1）不应被 F.2 规则覆盖
            g_task = result["items"][1]["appendixTasks"][0]
            self.assertNotEqual(
                (g_task.get("sourceRouting") or {}).get("source"), "appendix_source_matrix"
            )

    def test_bid_table_filler_fills_same_shape_response_table_from_reference_docx(self) -> None:
        table_filler = load_table_filler_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            blank = tmp_path / "附表B.6 技术服务响应表.docx"
            blank_doc = Document()
            table = blank_doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            for index, text in enumerate(["序号", "服务项目", "投标人响应值", "备注"]):
                table.cell(0, index).text = text
            for row in (
                ["1", "现场培训", "", ""],
                ["2", "技术支持", "", ""],
                ["3", "资料交付", "", ""],
            ):
                cells = table.add_row().cells
                for index, text in enumerate(row):
                    cells[index].text = text
            blank_doc.save(blank)

            reference = tmp_path / "项目技术服务响应源表.docx"
            reference_doc = Document()
            source_table = reference_doc.add_table(rows=1, cols=4)
            source_table.style = "Table Grid"
            for index, text in enumerate(["编号", "服务项目", "响应内容", "说明"]):
                source_table.cell(0, index).text = text
            for row in (
                ["1", "现场培训", "提供不少于10人次现场培训", "项目服务"],
                ["2", "技术支持", "提供7x24小时远程技术支持", "项目服务"],
                ["3", "资料交付", "按招标要求提交全套技术资料", "项目服务"],
            ):
                cells = source_table.add_row().cells
                for index, text in enumerate(row):
                    cells[index].text = text
            reference_doc.save(reference)

            manifest_path = tmp_path / "manifest.json"
            output = tmp_path / "filled.docx"
            manifest_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-tech-table-fill-v1",
                        "appendixTask": {
                            "id": "APPX-B6",
                            "title": "附表B.6 技术服务响应表",
                            "docxPath": str(blank),
                        },
                        "referenceMaterials": [
                            {"id": "SERVICE", "name": reference.name, "path": str(reference), "materialTier": "project"}
                        ],
                        "outputFile": str(output),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = table_filler.run_from_manifest(manifest_path)

            self.assertEqual(result["fillReport"]["filledFieldCount"], 3)
            self.assertEqual(result["fillReport"]["unfilledFieldCount"], 0)
            self.assertTrue(any(item["action"] == "fill" and "同形行列结构" in item["reason"] for item in result["filledFieldDetails"]))
            filled_doc = Document(str(output))
            rows = filled_doc.tables[0].rows
            self.assertEqual(rows[1].cells[2].text, "提供不少于10人次现场培训")
            self.assertEqual(rows[2].cells[2].text, "提供7x24小时远程技术支持")
            self.assertEqual(rows[3].cells[2].text, "按招标要求提交全套技术资料")

    def test_bid_table_filler_reads_whole_tender_document_without_material_references(self) -> None:
        table_filler = load_table_filler_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            blank = tmp_path / "附表B.6 技术服务响应表.docx"
            blank_doc = Document()
            target = blank_doc.add_table(rows=1, cols=3)
            for index, text in enumerate(["序号", "服务项目", "投标人响应值"]):
                target.cell(0, index).text = text
            for row in (["1", "现场培训", ""], ["2", "资料交付", ""]):
                cells = target.add_row().cells
                for index, text in enumerate(row):
                    cells[index].text = text
            blank_doc.save(blank)

            tender = tmp_path / "完整招标文件.docx"
            tender_doc = Document()
            tender_doc.add_heading("1 项目概况", level=1)
            unrelated = tender_doc.add_table(rows=2, cols=2)
            unrelated.cell(0, 0).text = "付款方式"
            unrelated.cell(0, 1).text = "分期付款"
            unrelated.cell(1, 0).text = "交货地点"
            unrelated.cell(1, 1).text = "项目现场"
            tender_doc.add_heading("附表B.6 技术服务响应表", level=1)
            source = tender_doc.add_table(rows=1, cols=3)
            for index, text in enumerate(["编号", "服务项目", "响应内容"]):
                source.cell(0, index).text = text
            for row in (["1", "现场培训", "提供不少于10人次现场培训"], ["2", "资料交付", "提交全套技术资料"]):
                cells = source.add_row().cells
                for index, text in enumerate(row):
                    cells[index].text = text
            tender_doc.save(tender)

            output = tmp_path / "filled.docx"
            manifest_path = tmp_path / "manifest.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-tech-table-fill-v1",
                        "appendixTask": {"id": "APPX-B6", "title": "附表B.6 技术服务响应表", "docxPath": str(blank)},
                        "referenceMaterials": [],
                        "materialIndex": [],
                        "tenderDocuments": [
                            {
                                "id": "TEN-1",
                                "name": tender.name,
                                "sourcePath": str(tender),
                                "sourceType": "project_tender_document",
                            }
                        ],
                        "outputFile": str(output),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = table_filler.run_from_manifest(manifest_path)

            self.assertEqual(result["fillReport"]["referenceMaterialCount"], 1)
            self.assertEqual(result["fillReport"]["unfilledFieldCount"], 0)
            self.assertEqual(result["fillReport"]["sourceSelection"]["selected"][0]["route"], "项目招标文件全文")
            filled = Document(str(output)).tables[0]
            self.assertEqual(filled.cell(1, 2).text, "提供不少于10人次现场培训")
            self.assertEqual(filled.cell(2, 2).text, "提交全套技术资料")

    def test_bid_table_filler_uses_parse_fields_and_project_materials_for_project_specific_values(self) -> None:
        table_filler = load_table_filler_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            blank = tmp_path / "附表A.1 投标机型总方案信息表.docx"
            blank_doc = Document()
            table = blank_doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            for index, text in enumerate(["序号", "项目", "投标响应", "备注"]):
                table.cell(0, index).text = text
            for row in (
                ["1", "机组台数", "", ""],
                ["2", "总容量（MW）", "", ""],
                ["3", "场址空气密度", "", ""],
                ["4", "基础混凝土（m3）", "", ""],
                ["5", "基础钢筋（t）", "", ""],
                ["6", "箱变位置", "", ""],
            ):
                cells = table.add_row().cells
                for index, text in enumerate(row):
                    cells[index].text = text
            blank_doc.save(blank)

            foundation = tmp_path / "基础弯矩表.xlsx"
            wb = Workbook()
            ws = wb.active
            ws.title = "基础工程量"
            ws.append(["基础混凝土用量（m3）", "789.482"])
            ws.append(["垫层混凝土用量（m3）", "57.353"])
            ws.append(["基础钢筋用量（kg）", "85787.635"])
            wb.save(foundation)

            manifest_path = tmp_path / "manifest.json"
            output = tmp_path / "filled.docx"
            manifest_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-tech-table-fill-v1",
                        "appendixTask": {
                            "id": "APPX-A1",
                            "title": "附表A.1 投标机型总方案信息表",
                            "docxPath": str(blank),
                        },
                        "projectTurbineModel": {
                            "model": "EW10.0-220上置",
                            "layout": "变压器上置",
                        },
                        "parseFields": [
                            {"id": "REQ-NOISE-COUNT", "label": "项目基础信息", "value": "故障率≥10%，不足3台按3台计算"},
                            {"id": "REQ-NOISE-DENSITY", "label": "空气密度", "value": "湿度均有相应变化，相同风速蕴含的风能也会发生变化"},
                            {"id": "REQ-COUNT", "label": "项目概况", "value": "本工程计划安装60台单机容量10MW风力发电机组"},
                            {"id": "REQ-SCALE", "label": "标段规模", "value": "600MW"},
                            {"id": "REQ-DENSITY", "label": "空气密度", "value": "1.154 kg/m³"},
                        ],
                        "materialIndex": [
                            {
                                "id": "RAW-FOUNDATION",
                                "name": "基础弯矩表.xlsx",
                                "folderPath": "技术标/项目素材/基础弯矩表",
                                "path": str(foundation),
                                "materialTier": "project",
                            }
                        ],
                        "outputFile": str(output),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = table_filler.run_from_manifest(manifest_path)

            self.assertEqual(result["fillReport"]["unfilledFieldCount"], 0)
            filled_doc = Document(str(output))
            rows = filled_doc.tables[0].rows
            self.assertEqual(rows[1].cells[2].text, "60")
            self.assertEqual(rows[2].cells[2].text, "600")
            self.assertEqual(rows[3].cells[2].text, "1.154 kg/m³")
            self.assertEqual(rows[4].cells[2].text, "789.482")
            self.assertEqual(rows[5].cells[2].text, "85.788")
            self.assertEqual(rows[6].cells[2].text, "塔上上置机型")

    def test_bid_table_filler_copies_power_curve_matrix_from_project_excels(self) -> None:
        table_filler = load_table_filler_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            blank = tmp_path / "附表D.1 标准及风电场空气密度功率曲线.docx"
            doc = Document()
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            for index, text in enumerate(["风速区间（m/s）", "区间平均风速（m/s）", "标准空气密度下功率（kW）", "风电场空气密度下功率（kW）"]):
                table.cell(0, index).text = text
            for row in (
                ["0.00-0.25", "0", "", ""],
                ["2.75-3.25", "3.0", "", ""],
                ["3.25-3.75", "3.5", "", ""],
            ):
                cells = table.add_row().cells
                for index, text in enumerate(row):
                    cells[index].text = text
            doc.save(blank)

            standard = tmp_path / "W10.0-220_空气密度1.225_功率曲线.xlsx"
            wb = Workbook()
            ws = wb.active
            ws.title = "功率曲线与发电量"
            ws.append(["风速(m/s)", "修正功率(kW)(Ce_max=0.45)", "Ct"])
            ws.append([3.0, 210, 0.91])
            ws.append([3.5, 350, 0.87])
            wb.save(standard)

            site = tmp_path / "W10.0-220_空气密度1.16_功率曲线.xlsx"
            wb = Workbook()
            ws = wb.active
            ws.title = "功率曲线与发电量"
            ws.append(["风速(m/s)", "修正功率(kW)(Ce_max=0.45)", "Ct"])
            ws.append([3.0, 204, 0.9])
            ws.append([3.5, 345.4, 0.86])
            wb.save(site)

            manifest_path = tmp_path / "manifest.json"
            output = tmp_path / "filled.docx"
            manifest_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-tech-table-fill-v1",
                        "appendixTask": {
                            "id": "APPX-D1",
                            "title": "附表D.1 标准及风电场空气密度功率曲线",
                            "docxPath": str(blank),
                        },
                        "materialIndex": [
                            {"id": "STD", "name": standard.name, "path": str(standard), "materialTier": "project"},
                            {"id": "SITE", "name": site.name, "path": str(site), "materialTier": "project"},
                        ],
                        "outputFile": str(output),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = table_filler.run_from_manifest(manifest_path)

            self.assertEqual(result["fillReport"]["filledFieldCount"], 6)
            filled_doc = Document(str(output))
            rows = filled_doc.tables[0].rows
            self.assertEqual(rows[1].cells[2].text, "0")
            self.assertEqual(rows[1].cells[3].text, "0")
            self.assertEqual(rows[2].cells[2].text, "210")
            self.assertEqual(rows[2].cells[3].text, "204")
            self.assertEqual(rows[3].cells[2].text, "350")
            self.assertEqual(rows[3].cells[3].text, "345.4")

    def test_bid_table_filler_selects_curve_excels_for_blank_ct_curve_matrix(self) -> None:
        table_filler = load_table_filler_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            blank = tmp_path / "附表D.2 标准及风电场空气密度下推力系数曲线.docx"
            doc = Document()
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            for index, text in enumerate(["风速（m/s）", "标准空气密度下推力系数 Ct", "风电场空气密度下推力系数 Ct"]):
                table.cell(0, index).text = text
            for _ in range(3):
                table.add_row()
            doc.save(blank)

            params = tmp_path / "X2平台机型投标参数_20250106.xlsx"
            wb = Workbook()
            ws = wb.active
            ws.title = "主参数"
            ws.append(["对象", "项目", "参数名称", "单位", "W10.0-220"])
            ws.append(["", "", "叶轮直径", "m", 220])
            wb.save(params)

            standard = tmp_path / "RAW-0460-W10.0-220_空气密度1.225_功率曲线.xlsx"
            wb = Workbook()
            ws = wb.active
            ws.title = "功率曲线与发电量"
            ws.append(["风速(m/s)", "修正功率(kW)", "Ct"])
            ws.append([3.0, 210, 0.91])
            ws.append([3.5, 350, 0.87])
            ws.append([4.0, 510, 0.83])
            wb.save(standard)

            site = tmp_path / "RAW-0461-W10.0-220_空气密度1.16_功率曲线.xlsx"
            wb = Workbook()
            ws = wb.active
            ws.title = "功率曲线与发电量"
            ws.append(["风速(m/s)", "修正功率(kW)", "Ct"])
            ws.append([3.0, 204, 0.9])
            ws.append([3.5, 345.4, 0.86])
            ws.append([4.0, 500.2, 0.82])
            wb.save(site)

            manifest_path = tmp_path / "manifest.json"
            output = tmp_path / "filled.docx"
            manifest_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-tech-table-fill-v1",
                        "appendixTask": {
                            "id": "APPX-D2",
                            "title": "附表D.2 标准及风电场空气密度下推力系数曲线",
                            "docxPath": str(blank),
                        },
                        "materialIndex": [
                            {"id": "PARAM", "name": params.name, "path": str(params), "materialTier": "project"},
                            {"id": "STD", "name": standard.name, "path": str(standard), "materialTier": "project"},
                            {"id": "SITE", "name": site.name, "path": str(site), "materialTier": "project"},
                        ],
                        "outputFile": str(output),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = table_filler.run_from_manifest(manifest_path)

            selected_names = [item["name"] for item in result["fillReport"]["sourceSelection"]["selected"]]
            self.assertIn(standard.name, selected_names)
            self.assertIn(site.name, selected_names)
            self.assertEqual(result["fillReport"]["filledFieldCount"], 9)
            filled_doc = Document(str(output))
            rows = filled_doc.tables[0].rows
            self.assertEqual(rows[1].cells[0].text, "3")
            self.assertEqual(rows[1].cells[1].text, "0.91")
            self.assertEqual(rows[1].cells[2].text, "0.9")
            self.assertEqual(rows[3].cells[0].text, "4")
            self.assertEqual(rows[3].cells[1].text, "0.83")
            self.assertEqual(rows[3].cells[2].text, "0.82")

    def test_bid_table_filler_transplants_matching_structured_source_table(self) -> None:
        table_filler = load_table_filler_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            blank = tmp_path / "附表E.1 投标人风资源评估与机位排布方案.docx"
            doc = Document()
            table = doc.add_table(rows=1, cols=6)
            table.style = "Table Grid"
            for index, text in enumerate(["机位编号", "坐标", "海拔高度（m）", "韦布尔参数A", "平均风速（m/s）", "机型"]):
                table.cell(0, index).text = text
            for row in (["", "", "", "", "", ""], ["…", "", "", "", "", ""], ["平均值", "", "", "", "", ""]):
                cells = table.add_row().cells
                for index, text in enumerate(row):
                    cells[index].text = text
            doc.save(blank)

            source_docx = tmp_path / "风资源评估报告.docx"
            doc = Document()
            table = doc.add_table(rows=1, cols=6)
            table.style = "Table Grid"
            for index, text in enumerate(["风机/编号", "X/(m)", "Z/(m)", "韦布尔参数A", "平均风速(m/s)", "机型"]):
                table.cell(0, index).text = text
            for index in range(1, 12):
                cells = table.add_row().cells
                values = [f"A{index:03d}", f"40416{index}", str(700 + index), f"8.{index:02d}", f"7.{index:02d}", "EW10.0-220-125"]
                for col, value in enumerate(values):
                    cells[col].text = value
            doc.save(source_docx)

            manifest_path = tmp_path / "manifest.json"
            output = tmp_path / "filled.docx"
            manifest_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-tech-table-fill-v1",
                        "appendixTask": {
                            "id": "APPX-E1",
                            "title": "附表E.1 投标人风资源评估与机位排布方案",
                            "docxPath": str(blank),
                        },
                        "referenceMaterials": [
                            {"id": "WIND", "name": source_docx.name, "path": str(source_docx), "materialTier": "project"}
                        ],
                        "outputFile": str(output),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = table_filler.run_from_manifest(manifest_path)

            self.assertGreater(result["fillReport"]["filledFieldCount"], 60)
            filled_doc = Document(str(output))
            rows = filled_doc.tables[0].rows
            self.assertEqual(len(rows), 12)
            self.assertEqual(rows[0].cells[0].text, "风机/编号")
            self.assertEqual(rows[1].cells[0].text, "A001")
            self.assertEqual(rows[11].cells[5].text, "EW10.0-220-125")
            self.assertEqual(result["unfilledFields"], [])

    def test_bid_table_filler_generates_load_wind_parameter_group_from_position_table(self) -> None:
        table_filler = load_table_filler_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            blank = tmp_path / "附表G.2.2 投标人对招标项目场址载荷计算选取风参数结果.docx"
            doc = Document()
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            for col in range(5):
                table.cell(0, col).text = "载荷计算风参数分组1"
            for row in (
                ["序号", "机位编号", "机组坐标", "机组坐标", "投标机型"],
                ["序号", "机位编号", "X", "Y", "投标机型"],
                ["1", "1", "", "", ""],
                ["……", "", "", "", ""],
                ["备注：如采用全场包络，只需填写分组1。"] * 5,
            ):
                cells = table.add_row().cells
                for index, text in enumerate(row):
                    cells[index].text = text
            doc.save(blank)

            source_docx = tmp_path / "风资源评估报告.docx"
            doc = Document()
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            for index, text in enumerate(["机位编号", "X", "Y", "机型"]):
                table.cell(0, index).text = text
            for row in (
                ["北区", "北区", "北区", "北区"],
                ["A001", "40416757.857", "4786181.802", "EW10.0-220-125"],
                ["A002", "40416633.400", "4785238.500", "EW10.0-220-125"],
                ["A044", "40431040.541", "4780184.646", "EW10.0-220-125"],
                ["A34", "40425746.863", "4776165.318", "EW10.0-220-125"],
                ["南区", "南区", "南区", "南区"],
                ["BX01", "40436771.920", "4779768.358", "备选"],
            ):
                cells = table.add_row().cells
                for index, text in enumerate(row):
                    cells[index].text = text
            doc.save(source_docx)

            manifest_path = tmp_path / "manifest.json"
            output = tmp_path / "filled.docx"
            manifest_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-tech-table-fill-v1",
                        "appendixTask": {
                            "id": "APPX-G22",
                            "title": "附表G.2.2 投标人对招标项目场址载荷计算选取风参数结果",
                            "docxPath": str(blank),
                        },
                        "referenceMaterials": [
                            {"id": "WIND", "name": source_docx.name, "path": str(source_docx), "materialTier": "project"}
                        ],
                        "outputFile": str(output),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = table_filler.run_from_manifest(manifest_path)

            self.assertEqual(result["unfilledFields"], [])
            filled_doc = Document(str(output))
            rows = filled_doc.tables[0].rows
            self.assertEqual(len(rows), 8)
            self.assertEqual([cell.text for cell in rows[3].cells], ["1", "A001", "40416758", "4786182", "EW10.0-220-125"])
            self.assertEqual([cell.text for cell in rows[4].cells], ["2", "A002", "40416633", "4785239", "EW10.0-220-125"])
            self.assertEqual([cell.text for cell in rows[5].cells], ["3", "A34", "40425747", "4776165", "EW10.0-220-125"])
            self.assertEqual([cell.text for cell in rows[6].cells], ["4", "A044", "40431041", "4780185", "EW10.0-220-125"])
            self.assertEqual(rows[7].cells[0].text, "备注：如采用全场包络，只需填写分组1。")

    def test_bid_table_filler_generates_spare_parts_table_from_quote_xlsx(self) -> None:
        table_filler = load_table_filler_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            blank = tmp_path / "附表B.2 质量保证期备品备件、消耗品清单.docx"
            doc = Document()
            table = doc.add_table(rows=1, cols=8)
            table.style = "Table Grid"
            for index, text in enumerate(["序号", "名称", "型号和规格", "单位", "数量", "备注", "更换周期", "国内替代产品型号"]):
                table.cell(0, index).text = text
            table.add_row()
            doc.save(blank)

            quote = tmp_path / "项目报价文件.xlsx"
            wb = Workbook()
            ws = wb.active
            ws.title = "E 推荐备品备件（如果有）的分项报价"
            ws.append([""] * 10)
            ws.append(["表2 E"] + [""] * 9)
            ws.append(["单位：人民币万元"] + [""] * 9)
            ws.append(["序号", "名称", "规格型号", "单位", "数量", "价格", "总价", "产地", "生产厂家", "备注"])
            ws.append(["一、备品备件部分"] + [""] * 9)
            ws.append([1, "变桨限位开关_FD 2031_S1", "EW10.0-220-125", "EA", 12, "", "", "中国", "上海电气合供", ""])
            ws.append([2, "避雷器DXH06_FCS/3+1R40", "EW10.0-220-125", "EA", 12, "", "", "中国", "上海电气合供", ""])
            ws.append([None, "合计(单台机组)（万元）"] + [""] * 8)
            wb.save(quote)

            manifest_path = tmp_path / "manifest.json"
            output = tmp_path / "filled.docx"
            manifest_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-tech-table-fill-v1",
                        "appendixTask": {
                            "id": "APPX-B2",
                            "title": "附表B.2 质量保证期备品备件、消耗品清单",
                            "docxPath": str(blank),
                        },
                        "referenceMaterials": [
                            {"id": "QUOTE", "name": quote.name, "path": str(quote), "materialTier": "project"}
                        ],
                        "outputFile": str(output),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = table_filler.run_from_manifest(manifest_path)

            filled_doc = Document(str(output))
            rows = filled_doc.tables[0].rows
            self.assertEqual(len(rows), 4)
            self.assertEqual([cell.text for cell in rows[1].cells], ["一、备品备件部分"] * 8)
            self.assertEqual([cell.text for cell in rows[2].cells], ["1", "变桨限位开关_FD 2031_S1", "EW10.0-220-125", "EA", "12", "\\", "\\", "\\"])
            self.assertEqual(result["unfilledFields"], [])

    def test_bid_table_filler_replace_first_table_preserves_table_props(self) -> None:
        """整表替换必须保留目标表的样式引用与显式边框，否则渲染成无框线文本。"""
        table_filler = load_table_filler_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            blank = Path(tmp) / "附表D.7 性能及考核承诺保证表.docx"
            doc = Document()
            table = doc.add_table(rows=2, cols=4)
            tbl_pr = table._tbl.tblPr
            style_el = tbl_pr.find(qn("w:tblStyle"))
            if style_el is None:
                style_el = OxmlElement("w:tblStyle")
                tbl_pr.insert(0, style_el)
            style_el.set(qn("w:val"), "89")
            borders = OxmlElement("w:tblBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "single")
                el.set(qn("w:sz"), "4")
                borders.append(el)
            tbl_pr.append(borders)
            for index, text in enumerate(["项目", "保证值", "授权人签名", "日期"]):
                table.cell(0, index).text = text
            doc.save(blank)

            table_filler.replace_first_table(
                blank,
                [["项目", "保证值", "授权人签名", "日期"], ["功率曲线保证值", "97%", "", "2026-01-23"]],
            )

            filled = Document(str(blank))
            new_tbl_pr = filled.tables[0]._tbl.tblPr
            self.assertEqual(new_tbl_pr.find(qn("w:tblStyle")).get(qn("w:val")), "89")
            self.assertIsNotNone(new_tbl_pr.find(qn("w:tblBorders")))
            self.assertEqual(filled.tables[0].rows[1].cells[1].text, "97%")

    def test_bid_table_filler_batch_output_names_include_appendix_id_to_avoid_collisions(self) -> None:
        table_filler = load_table_filler_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            blanks = []
            for index, appendix_id in enumerate(("APPX-0014", "APPX-0080", "APPX-0014"), start=1):
                blank = tmp_path / f"{index}-{appendix_id}.docx"
                doc = Document()
                table = doc.add_table(rows=1, cols=3)
                table.style = "Table Grid"
                for index, text in enumerate(["序号", "项目", "投标响应"]):
                    table.cell(0, index).text = text
                cells = table.add_row().cells
                cells[0].text = "1"
                cells[1].text = "投标机型"
                cells[2].text = ""
                doc.save(blank)
                blanks.append((appendix_id, blank))

            manifest_path = tmp_path / "manifest.json"
            output_dir = tmp_path / "outputs"
            manifest_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-tech-table-fill-v1",
                        "projectTurbineModel": {"model": "EW10.0-220上置"},
                        "outputDir": str(output_dir),
                        "targets": [
                            {
                                "id": appendix_id,
                                "title": "附表B.8 出质保后备品备件服务 无",
                                "docxPath": str(blank),
                            }
                            for appendix_id, blank in blanks
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = table_filler.run_from_manifest(manifest_path)

            output_files = [Path(item) for item in result["outputFiles"]]
            self.assertEqual(len(output_files), 3)
            self.assertEqual(len({item.name for item in output_files}), 3)
            self.assertTrue(any("APPX-0014" in item.name for item in output_files))
            self.assertTrue(any("APPX-0080" in item.name for item in output_files))

    def test_bid_table_filler_copies_no_table_appendix_without_failing_batch(self) -> None:
        table_filler = load_table_filler_script("run_from_manifest")

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            no_table = tmp_path / "附表B.8 出质保后备品备件服务 无.docx"
            no_table_doc = Document()
            no_table_doc.add_paragraph("附表B.8 出质保后备品备件服务 无")
            no_table_doc.save(no_table)

            normal = tmp_path / "附表A.1 投标机型总方案信息表.docx"
            normal_doc = Document()
            table = normal_doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            for index, text in enumerate(["序号", "项目", "投标响应"]):
                table.cell(0, index).text = text
            cells = table.add_row().cells
            cells[0].text = "1"
            cells[1].text = "投标机型"
            cells[2].text = ""
            normal_doc.save(normal)

            manifest_path = tmp_path / "manifest.json"
            output_dir = tmp_path / "outputs"
            manifest_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-tech-table-fill-v1",
                        "projectTurbineModel": {"model": "EW10.0-220上置"},
                        "outputDir": str(output_dir),
                        "targets": [
                            {
                                "id": "APPX-B8",
                                "title": "附表B.8 出质保后备品备件服务 无",
                                "docxPath": str(no_table),
                            },
                            {
                                "id": "APPX-A1",
                                "title": "附表A.1 投标机型总方案信息表",
                                "docxPath": str(normal),
                            },
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            result = table_filler.run_from_manifest(manifest_path)

            self.assertEqual(result["fillReport"]["targetCount"], 2)
            self.assertEqual(result["fillReport"]["failedTargetCount"], 0)
            self.assertEqual(result["fillReport"]["successfulTargetCount"], 2)
            self.assertEqual(len(result["outputFiles"]), 2)
            no_table_result = result["targetResults"][0]
            self.assertEqual(no_table_result["fillReport"]["targetFieldCount"], 0)
            self.assertTrue(Path(no_table_result["outputFile"]).exists())
            copied_doc = Document(no_table_result["outputFile"])
            self.assertEqual(copied_doc.paragraphs[0].text, "附表B.8 出质保后备品备件服务 无")

    def test_bid_assembler_inject_prefix_collapses_heading_level_gaps(self) -> None:
        numbering_fixer = load_assembler_script("numbering_fixer")

        doc = Document()
        doc.add_paragraph("上海电气优势简介", style="Heading 2")
        doc.add_paragraph("基本情况", style="Heading 3")
        doc.add_paragraph("集团概况", style="Heading 4")
        doc.add_paragraph("载荷仿真分析能力", style="Heading 1")
        doc.add_paragraph("测试验证技术", style="Heading 6")

        stats = numbering_fixer.inject_prefix_to_headings(
            doc,
            "1.9",
            toc_title="上海电气优势简介",
            skip_first_if_match=True,
        )
        headings = [para.text.strip().replace("  ", " ") for para in doc.paragraphs if para.text.strip()]

        self.assertTrue(stats["skipped_first"])
        self.assertEqual(
            headings,
            [
                "1.9.1 基本情况",
                "1.9.1.1 集团概况",
                "1.9.2 载荷仿真分析能力",
                "1.9.2.1 测试验证技术",
            ],
        )
        self.assertFalse(any(".0." in item or item.endswith(".0") for item in headings))

    def test_bid_assembler_inject_prefix_preserves_style_and_updates_navigation(self) -> None:
        numbering_fixer = load_assembler_script("numbering_fixer")

        doc = Document()
        doc.add_paragraph("上海电气优势简介", style="Heading 2")
        stale = doc.add_paragraph("载荷仿真分析能力", style="Heading 1")
        p_pr = stale._p.get_or_add_pPr()
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), "0")
        p_pr.append(outline)

        numbering_fixer.inject_prefix_to_headings(
            doc,
            "1.9",
            toc_title="上海电气优势简介",
            skip_first_if_match=True,
        )

        remaining = [para for para in doc.paragraphs if "载荷仿真分析能力" in para.text][0]
        self.assertEqual(remaining.style.name, "Heading 1")
        p_pr = remaining._p.find(qn("w:pPr"))
        self.assertIsNotNone(p_pr)
        self.assertEqual(p_pr.find(qn("w:outlineLvl")).get(qn("w:val")), "2")

    def test_bid_assembler_demotes_material_headings_to_body(self) -> None:
        numbering_fixer = load_assembler_script("numbering_fixer")

        doc = Document()
        doc.add_paragraph("上海电气优势简介", style="Heading 2")
        doc.add_paragraph("基本情况", style="Heading 3")
        stale = doc.add_paragraph("载荷仿真分析能力", style="Heading 1")
        p_pr = stale._p.get_or_add_pPr()
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), "0")
        p_pr.append(outline)

        stats = numbering_fixer.demote_headings_to_body(
            doc,
            toc_title="上海电气优势简介",
            remove_first_if_match=True,
        )

        self.assertEqual(stats["removed"], 1)
        self.assertEqual(stats["demoted"], 2)
        self.assertEqual([para.text for para in doc.paragraphs], ["基本情况", "载荷仿真分析能力"])
        self.assertFalse(any((para.style.name or "").startswith("Heading") for para in doc.paragraphs))
        self.assertTrue(all(para._p.find(qn("w:pPr")).find(qn("w:outlineLvl")) is None for para in doc.paragraphs))

    def test_bid_assembler_demotes_direct_outline_without_body_style(self) -> None:
        numbering_fixer = load_assembler_script("numbering_fixer")

        doc = Document()
        styled = doc.add_paragraph("发电小时数承诺函", style="Heading 2")
        direct_only = doc.add_paragraph("发电量保证矩阵表")
        direct_p_pr = direct_only._p.get_or_add_pPr()
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), "0")
        direct_p_pr.append(outline)

        styles_el = doc.styles._element
        for style in list(styles_el.findall(qn("w:style"))):
            if style.get(qn("w:styleId")) == "Normal":
                styles_el.remove(style)

        stats = numbering_fixer.demote_headings_to_body(doc)

        self.assertEqual(stats["demoted"], 2)
        for para in (styled, direct_only):
            p_pr = para._p.find(qn("w:pPr"))
            self.assertIsNotNone(p_pr)
            self.assertIsNone(p_pr.find(qn("w:pStyle")))
            self.assertIsNone(p_pr.find(qn("w:outlineLvl")))

    def test_bid_assembler_keeps_s2_child_headings_from_material(self) -> None:
        numbering_fixer = load_assembler_script("numbering_fixer")

        doc = Document()
        doc.add_paragraph("项目技术承诺函", style="Heading 1")
        child = doc.add_paragraph("发电小时数承诺函", style="Heading 2")
        extra = doc.add_paragraph("发电量保证矩阵表")
        p_pr = extra._p.get_or_add_pPr()
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), "0")
        p_pr.append(outline)

        stats = numbering_fixer.demote_headings_to_body(
            doc,
            toc_title="项目技术承诺函",
            remove_first_if_match=True,
            keep_heading_map={
                "发电小时数承诺函": {
                    "chapter_no": "4.1",
                    "title": "发电小时数承诺函",
                    "level": 2,
                }
            },
        )

        self.assertEqual(stats["removed"], 1)
        self.assertEqual(stats["kept"], 1)
        self.assertEqual(stats["demoted"], 1)
        self.assertEqual(child.text.strip().replace("  ", " "), "4.1 发电小时数承诺函")
        self.assertEqual(child.style.name, "Heading 2")
        extra_p_pr = extra._p.find(qn("w:pPr"))
        self.assertIsNone(extra_p_pr.find(qn("w:pStyle")))
        self.assertIsNone(extra_p_pr.find(qn("w:outlineLvl")))

    def test_bid_assembler_remaps_material_headings_to_navigation(self) -> None:
        numbering_fixer = load_assembler_script("numbering_fixer")

        doc = Document()
        doc.add_paragraph("设备运行和维护专题", style="Heading 2")
        project_flow = doc.add_paragraph("项目流程", style="Heading 1")
        list_item = doc.add_paragraph("（3）机组调试")
        static = doc.add_paragraph()
        static.add_run("静态调试").bold = True
        table_heading = doc.add_paragraph("表C.1 总体技术参数与规格", style="Heading 2")

        stats = numbering_fixer.remap_material_headings_to_navigation(
            doc,
            toc_title="设备运行和维护专题",
            remove_first_if_match=True,
            parent_level=2,
        )

        self.assertEqual(stats["removed"], 1)
        self.assertEqual(stats["remapped"], 1)
        self.assertEqual(stats["bold_subheadings"], 0)
        self.assertEqual(stats["demoted"], 1)
        self.assertEqual(project_flow.style.name, "Heading 1")
        self.assertEqual(
            project_flow._p.find(qn("w:pPr")).find(qn("w:outlineLvl")).get(qn("w:val")),
            "2",
        )
        self.assertFalse((list_item.style.name or "").startswith("Heading"))
        self.assertFalse((static.style.name or "").startswith("Heading"))
        self.assertFalse((table_heading.style.name or "").startswith("Heading"))

    def test_bid_assembler_strips_heading_style_numbering(self) -> None:
        numbering_fixer = load_assembler_script("numbering_fixer")

        with tempfile.TemporaryDirectory() as tmp:
            docx_path = Path(tmp) / "numbered-heading.docx"
            doc = Document()
            heading_style = doc.styles["Heading 2"]
            p_pr = heading_style.element.get_or_add_pPr()
            num_pr = OxmlElement("w:numPr")
            ilvl = OxmlElement("w:ilvl")
            ilvl.set(qn("w:val"), "1")
            num_id = OxmlElement("w:numId")
            num_id.set(qn("w:val"), "1")
            num_pr.append(ilvl)
            num_pr.append(num_id)
            p_pr.append(num_pr)
            doc.add_paragraph("1.7 投标方案优势说明", style="Heading 2")

            self.assertEqual(numbering_fixer.strip_numPr_from_heading_styles(doc), 1)
            doc.save(docx_path)

            with zipfile.ZipFile(docx_path) as zf:
                styles_xml = zf.read("word/styles.xml").decode("utf-8")

        match = re.search(r'(<w:style[^>]+w:styleId="Heading2"[^>]*>.*?</w:style>)', styles_xml)
        self.assertIsNotNone(match)
        self.assertNotIn("<w:numPr>", match.group(1))

    def _make_hidden_numbered_custom_style(self, doc, style_name: str, base_name: str, num_id: str):
        """构造带隐藏自动编号的自定义标题样式（如"标题6-标书" basedOn Heading 6）。"""
        custom = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        custom.base_style = doc.styles[base_name]
        p_pr = custom.element.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        nid = OxmlElement("w:numId")
        nid.set(qn("w:val"), num_id)
        num_pr.append(ilvl)
        num_pr.append(nid)
        p_pr.append(num_pr)
        return custom

    def _add_paragraph_numpr(self, para, num_id: str) -> None:
        p_pr = para._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        nid = OxmlElement("w:numId")
        nid.set(qn("w:val"), num_id)
        num_pr.append(ilvl)
        num_pr.append(nid)
        p_pr.append(num_pr)

    def test_bid_assembler_remap_preserves_numid_zero_suppression(self) -> None:
        """1.7 回归：注入层级时不得删除抑制隐藏自动编号的 numId=0。"""
        numbering_fixer = load_assembler_script("numbering_fixer")

        doc = Document()
        self._make_hidden_numbered_custom_style(doc, "标题6-标书", "Heading 6", "7")
        para = doc.add_paragraph("总体技术路线", style="标题6-标书")
        self._add_paragraph_numpr(para, "0")  # 源文档用 numId=0 抑制样式隐藏编号

        stats = numbering_fixer.remap_material_headings_to_navigation(doc, parent_level=2)

        self.assertEqual(stats["remapped"], 1)
        num_pr = para._p.find(qn("w:pPr")).find(qn("w:numPr"))
        self.assertIsNotNone(num_pr)
        self.assertEqual(num_pr.find(qn("w:numId")).get(qn("w:val")), "0")

    def test_bid_assembler_remap_strips_active_paragraph_numbering(self) -> None:
        """注入层级时真正生效的段落自动编号（numId>0）仍要剥掉，避免双编号。"""
        numbering_fixer = load_assembler_script("numbering_fixer")

        doc = Document()
        para = doc.add_paragraph("总体技术路线", style="Heading 6")
        self._add_paragraph_numpr(para, "7")

        stats = numbering_fixer.remap_material_headings_to_navigation(doc, parent_level=2)

        self.assertEqual(stats["remapped"], 1)
        num_pr = para._p.find(qn("w:pPr")).find(qn("w:numPr"))
        self.assertIsNone(num_pr)

    def test_bid_assembler_strips_basedon_chain_heading_style_numbering(self) -> None:
        """1.7 回归：basedOn 链指向 Heading 的自定义样式也要剥样式级自动编号。"""
        numbering_fixer = load_assembler_script("numbering_fixer")

        doc = Document()
        custom = self._make_hidden_numbered_custom_style(doc, "标题6-标书", "Heading 6", "7")

        self.assertEqual(numbering_fixer.strip_numPr_from_heading_styles(doc), 1)
        self.assertIsNone(custom.element.find(qn("w:pPr")).find(qn("w:numPr")))

    def test_bid_assembler_enforce_invariant_clears_residual_numbering(self) -> None:
        """不变量：已写入文本编号的 Heading 不得再有有效 Word 自动编号。"""
        numbering_fixer = load_assembler_script("numbering_fixer")

        doc = Document()
        custom = self._make_hidden_numbered_custom_style(doc, "标题6-标书", "Heading 6", "7")
        # 场景1：段落 numId>0 → 改为段落级 numId=0
        active = doc.add_paragraph("1.7.3.1 总体技术路线", style="Heading 3")
        self._add_paragraph_numpr(active, "5")
        # 场景2：段落 numId=0 抑制仍在 → 不动
        suppressed = doc.add_paragraph("1.7.3.2 关键技术路线", style="标题6-标书")
        self._add_paragraph_numpr(suppressed, "0")
        # 场景3：样式链编号且无段落抑制 → 只给当前段落加 numId=0
        inherited = doc.add_paragraph("1.7.3.3 其他技术路线", style="标题6-标书")

        fixed = numbering_fixer.enforce_no_auto_numbering_on_numbered_headings(doc)

        self.assertEqual(fixed, 2)
        active_num_pr = active._p.find(qn("w:pPr")).find(qn("w:numPr"))
        self.assertEqual(active_num_pr.find(qn("w:numId")).get(qn("w:val")), "0")
        suppressed_num_pr = suppressed._p.find(qn("w:pPr")).find(qn("w:numPr"))
        self.assertIsNotNone(suppressed_num_pr)
        self.assertEqual(suppressed_num_pr.find(qn("w:numId")).get(qn("w:val")), "0")
        inherited_num_pr = inherited._p.find(qn("w:pPr")).find(qn("w:numPr"))
        self.assertEqual(inherited_num_pr.find(qn("w:numId")).get(qn("w:val")), "0")
        style_num_pr = custom.element.find(qn("w:pPr")).find(qn("w:numPr"))
        self.assertEqual(style_num_pr.find(qn("w:numId")).get(qn("w:val")), "7")

    def test_bid_assembler_invariant_keeps_shared_body_list_style_numbering(self) -> None:
        """标题局部抑制自动编号时，不得破坏同样式的正文列表。"""
        numbering_fixer = load_assembler_script("numbering_fixer")

        doc = Document()
        shared = self._make_hidden_numbered_custom_style(
            doc,
            "共享正文列表",
            "Normal",
            "7",
        )
        heading = doc.add_paragraph("1.7.4 列表样式标题", style=shared)
        heading_p_pr = heading._p.get_or_add_pPr()
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), "2")
        heading_p_pr.append(outline)
        body_item = doc.add_paragraph("正文列表项", style=shared)

        fixed = numbering_fixer.enforce_no_auto_numbering_on_numbered_headings(doc)

        self.assertEqual(fixed, 1)
        heading_num_pr = heading._p.find(qn("w:pPr")).find(qn("w:numPr"))
        self.assertEqual(heading_num_pr.find(qn("w:numId")).get(qn("w:val")), "0")
        style_num_pr = shared.element.find(qn("w:pPr")).find(qn("w:numPr"))
        self.assertEqual(style_num_pr.find(qn("w:numId")).get(qn("w:val")), "7")
        body_p_pr = body_item._p.find(qn("w:pPr"))
        self.assertIsNotNone(body_p_pr)
        self.assertIsNone(body_p_pr.find(qn("w:numPr")))

    def test_material_cleaner_does_not_fabricate_heading_levels(self) -> None:
        """1.9 回归：清洗不按"带编号、文字较短"把正文猜成标题。"""
        word_cleaner = load_material_cleaner_script("word_cleaner")

        doc = Document()
        body = doc.add_paragraph("1、载荷仿真分析能力")  # Normal 样式 + 手写编号
        heading = doc.add_paragraph("2、自主研发控制算法", style="Heading 1")

        word_cleaner._strip_numbered_heading_prefixes(doc)

        # 正文不升格、文本保持原样
        self.assertEqual(body.style.name, "Normal")
        self.assertEqual(body.text, "1、载荷仿真分析能力")
        self.assertIsNone(body._p.find(qn("w:pPr")))
        # 真 Heading 仍剥手写前缀
        self.assertEqual(heading.text, "自主研发控制算法")

    def test_material_cleaner_preserves_basedon_heading_level(self) -> None:
        """自定义样式 basedOn Heading 6 时按真实层级清理，不按文本编号推断。"""
        word_cleaner = load_material_cleaner_script("word_cleaner")

        doc = Document()
        custom = doc.styles.add_style("标题6-标书", WD_STYLE_TYPE.PARAGRAPH)
        custom.base_style = doc.styles["Heading 6"]
        heading = doc.add_paragraph("1.7 自定义技术路线", style=custom)

        normalized = word_cleaner._strip_numbered_heading_prefixes(doc)

        self.assertEqual(normalized, 1)
        self.assertEqual(heading.text, "自定义技术路线")
        self.assertEqual(heading.style.name, "Heading 6")
        outline = heading._p.find(qn("w:pPr")).find(qn("w:outlineLvl"))
        self.assertEqual(outline.get(qn("w:val")), "5")

    def test_bid_assembler_gap_plan_flags_unconfirmed_candidates(self) -> None:
        """1.1/4.7 回归：只有候选素材、未确认 matchedMaterials 时给出显式提示。"""
        build_assembly = load_assembler_script("build_assembly")

        with tempfile.TemporaryDirectory() as tmp:
            gap_plan_path = Path(tmp) / "gap_plan.json"
            gap_plan_path.write_text(
                json.dumps(
                    {
                        "items": [
                            {
                                "id": "GAP-0001",
                                "number": "1.1",
                                "title": "上海电气优势简介",
                                "matchedMaterials": [],
                                "candidateMaterials": [
                                    {"id": "RAW-0087", "matchScore": 0.87}
                                ],
                                "resolvedArtifacts": [],
                            }
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            plan = [
                {
                    "chapter_no_flat": "1.1",
                    "chapter_no": "1.1",
                    "title": "上海电气优势简介",
                    "paths": [],
                    "shifts": [],
                    "attach_modes": [],
                    "field_replace": False,
                    "status": "UNMATCHED",
                    "note": "",
                }
            ]

            updated = build_assembly.apply_gap_plan(plan, gap_plan_path)

        self.assertEqual(updated[0]["status"], "UNMATCHED")
        self.assertEqual(updated[0]["paths"], [])
        self.assertIn("候选素材未确认", updated[0]["note"])
        self.assertEqual(updated[0]["gap_plan_item_id"], "GAP-0001")

    def test_bid_assembler_gap_plan_flags_unfinished_ai_fill(self) -> None:
        """1.1/4.7 回归：AI 填写流程未产出 S7-ready 产物时给出显式提示。"""
        build_assembly = load_assembler_script("build_assembly")

        with tempfile.TemporaryDirectory() as tmp:
            gap_plan_path = Path(tmp) / "gap_plan.json"
            gap_plan_path.write_text(
                json.dumps(
                    {
                        "items": [
                            {
                                "id": "GAP-0002",
                                "number": "4.7",
                                "title": "项目技术承诺函",
                                "fillTasks": [{"id": "FILL-0002", "status": "pending"}],
                                "matchedMaterials": [],
                                "candidateMaterials": [
                                    {"id": "RAW-0149", "matchScore": 0.87},
                                    {"id": "RAW-0151", "matchScore": 0.85},
                                ],
                                "resolvedArtifacts": [],
                            }
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            plan = [
                {
                    "chapter_no_flat": "4.7",
                    "chapter_no": "4.7",
                    "title": "项目技术承诺函",
                    "paths": [],
                    "shifts": [],
                    "attach_modes": [],
                    "field_replace": False,
                    "status": "UNMATCHED",
                    "note": "",
                }
            ]

            updated = build_assembly.apply_gap_plan(plan, gap_plan_path)

        self.assertEqual(updated[0]["status"], "UNMATCHED")
        self.assertIn("AI 填写未完成", updated[0]["note"])
        self.assertEqual(updated[0]["gap_plan_item_id"], "GAP-0002")

    def test_bid_assembler_gap_plan_matches_appendix_number_plus_title(self) -> None:
        build_assembly = load_assembler_script("build_assembly")

        with tempfile.TemporaryDirectory() as tmp:
            gap_plan_path = Path(tmp) / "gap_plan.json"
            gap_plan_path.write_text(
                json.dumps(
                    {
                        "items": [
                            {
                                "id": "GAP-0058",
                                "number": "附表A.1",
                                "title": "投标机型总方案信息表",
                                "fillTasks": [{"id": "FILL-0058", "status": "completed"}],
                                "matchedMaterials": [
                                    {"path": "/tmp/投标机型总方案信息表_待填写.docx"}
                                ],
                                "resolvedArtifacts": [
                                    {
                                        "source": "ai_fill",
                                        "path": "/tmp/投标机型总方案信息表_AI填写.docx",
                                        "s7Ready": True,
                                        "qualityReport": {"status": "passed"},
                                    }
                                ],
                            }
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            plan = [
                {
                    "toc_idx": 57,
                    "level": 1,
                    "chapter_no_flat": "",
                    "chapter_no": "",
                    "title": "附表A.1 投标机型总方案信息表",
                    "paths": [],
                    "shifts": [],
                    "attach_modes": [],
                    "field_replace": False,
                    "status": "NEEDS_REVIEW",
                    "note": "招标/模板新增章节，需人工补素材",
                }
            ]

            updated = build_assembly.apply_gap_plan(plan, gap_plan_path)

        self.assertEqual(updated[0]["status"], "MATCHED")
        self.assertEqual(updated[0]["paths"], ["/tmp/投标机型总方案信息表_AI填写.docx"])
        self.assertEqual(updated[0]["gap_plan_item_id"], "GAP-0058")

    def test_bid_assembler_gap_plan_skips_unreviewed_ai_fill_artifact(self) -> None:
        build_assembly = load_assembler_script("build_assembly")

        with tempfile.TemporaryDirectory() as tmp:
            gap_plan_path = Path(tmp) / "gap_plan.json"
            gap_plan_path.write_text(
                json.dumps(
                    {
                        "items": [
                            {
                                "id": "GAP-0058",
                                "number": "附表A.1",
                                "title": "投标机型总方案信息表",
                                "fillTasks": [{"id": "FILL-0058", "status": "completed"}],
                                "matchedMaterials": [
                                    {"path": "/tmp/投标机型总方案信息表_待填写.docx"}
                                ],
                                "resolvedArtifacts": [
                                    {
                                        "source": "ai_fill",
                                        "path": "/tmp/未验收_AI填写.docx",
                                        "s7Ready": True,
                                        "qualityReport": {"status": "needs_review"},
                                    }
                                ],
                            }
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            plan = [
                {
                    "chapter_no_flat": "附表A.1",
                    "chapter_no": "",
                    "title": "投标机型总方案信息表",
                    "paths": [],
                    "shifts": [],
                    "attach_modes": [],
                    "field_replace": False,
                    "status": "NEEDS_REVIEW",
                    "note": "待人工复核",
                }
            ]

            updated = build_assembly.apply_gap_plan(plan, gap_plan_path)

        # 未通过质检的 AI 填写产物不得进入组装，且必须显式提示"AI 填写未完成"
        self.assertEqual(updated[0]["status"], "UNMATCHED")
        self.assertEqual(updated[0]["paths"], [])
        self.assertIn("AI 填写未完成", updated[0]["note"])
        self.assertEqual(updated[0]["gap_plan_item_id"], "GAP-0058")

    def test_bid_assembler_gap_plan_blocks_partially_reviewed_multi_fill_task_item(self) -> None:
        """R10-B07-02：多 fillTask 只复核一个，S7 不得因存在任一可用产物而放行整项。"""
        build_assembly = load_assembler_script("build_assembly")

        with tempfile.TemporaryDirectory() as tmp:
            gap_plan_path = Path(tmp) / "gap_plan.json"
            gap_plan_path.write_text(
                json.dumps(
                    {
                        "items": [
                            {
                                "id": "GAP-0058",
                                "number": "附表A.1",
                                "title": "投标机型总方案信息表",
                                "fillTasks": [
                                    {"id": "FILL-0058-A", "status": "completed"},
                                    {"id": "FILL-0058-B", "status": "pending"},
                                ],
                                "matchedMaterials": [],
                                "resolvedArtifacts": [
                                    {
                                        "source": "ai_fill",
                                        "fillTaskId": "FILL-0058-A",
                                        "path": "/tmp/机组参数表_AI填写.docx",
                                        "s7Ready": True,
                                        "qualityGate": "human_confirmed",
                                    }
                                ],
                            }
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            plan = [
                {
                    "chapter_no_flat": "附表A.1",
                    "chapter_no": "",
                    "title": "投标机型总方案信息表",
                    "paths": [],
                    "shifts": [],
                    "attach_modes": [],
                    "field_replace": False,
                    "status": "NEEDS_REVIEW",
                    "note": "待人工复核",
                }
            ]

            updated = build_assembly.apply_gap_plan(plan, gap_plan_path)

        # 塔筒参数表仍 pending：整项阻断并显式提示，不能只合并已复核的机组参数表。
        self.assertEqual(updated[0]["status"], "UNMATCHED")
        self.assertEqual(updated[0]["paths"], [])
        self.assertIn("AI 填写未完成", updated[0]["note"])
        self.assertEqual(updated[0]["gap_plan_item_id"], "GAP-0058")

    def test_bid_assembler_gap_plan_merges_multi_fill_task_item_after_all_reviewed(self) -> None:
        """R10-B07-02：所有 fillTask 完成且产物均放行后，S7 正常合并全部产物。"""
        build_assembly = load_assembler_script("build_assembly")

        with tempfile.TemporaryDirectory() as tmp:
            gap_plan_path = Path(tmp) / "gap_plan.json"
            gap_plan_path.write_text(
                json.dumps(
                    {
                        "items": [
                            {
                                "id": "GAP-0058",
                                "number": "附表A.1",
                                "title": "投标机型总方案信息表",
                                "fillTasks": [
                                    {"id": "FILL-0058-A", "status": "completed"},
                                    {"id": "FILL-0058-B", "status": "completed"},
                                ],
                                "matchedMaterials": [],
                                "resolvedArtifacts": [
                                    {
                                        "source": "ai_fill",
                                        "fillTaskId": "FILL-0058-A",
                                        "path": "/tmp/机组参数表_AI填写.docx",
                                        "s7Ready": True,
                                        "qualityGate": "human_confirmed",
                                    },
                                    {
                                        "source": "ai_fill",
                                        "fillTaskId": "FILL-0058-B",
                                        "path": "/tmp/塔筒参数表_AI填写.docx",
                                        "s7Ready": True,
                                        "qualityGate": "human_confirmed",
                                    },
                                ],
                            }
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            plan = [
                {
                    "chapter_no_flat": "附表A.1",
                    "chapter_no": "",
                    "title": "投标机型总方案信息表",
                    "paths": [],
                    "shifts": [],
                    "attach_modes": [],
                    "field_replace": False,
                    "status": "NEEDS_REVIEW",
                    "note": "待人工复核",
                }
            ]

            updated = build_assembly.apply_gap_plan(plan, gap_plan_path)

        self.assertEqual(updated[0]["status"], "MATCHED")
        self.assertEqual(
            updated[0]["paths"],
            ["/tmp/机组参数表_AI填写.docx", "/tmp/塔筒参数表_AI填写.docx"],
        )
        self.assertEqual(updated[0]["gap_plan_item_id"], "GAP-0058")

    def test_bid_assembler_gap_plan_manual_artifact_replaces_pending_fill_task(self) -> None:
        """R10-B07-02：人工上传/选材产物按决策终审可替代填写任务，不被 pending 任务误阻断。"""
        build_assembly = load_assembler_script("build_assembly")

        with tempfile.TemporaryDirectory() as tmp:
            gap_plan_path = Path(tmp) / "gap_plan.json"
            gap_plan_path.write_text(
                json.dumps(
                    {
                        "items": [
                            {
                                "id": "GAP-0058",
                                "number": "附表A.1",
                                "title": "投标机型总方案信息表",
                                "fillTasks": [{"id": "FILL-0058-A", "status": "pending"}],
                                "matchedMaterials": [],
                                "resolvedArtifacts": [
                                    {
                                        "source": "manual_upload",
                                        "path": "/tmp/人工上传_总方案信息表.docx",
                                        "s7Ready": True,
                                    }
                                ],
                            }
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            plan = [
                {
                    "chapter_no_flat": "附表A.1",
                    "chapter_no": "",
                    "title": "投标机型总方案信息表",
                    "paths": [],
                    "shifts": [],
                    "attach_modes": [],
                    "field_replace": False,
                    "status": "NEEDS_REVIEW",
                    "note": "待人工复核",
                }
            ]

            updated = build_assembly.apply_gap_plan(plan, gap_plan_path)

        self.assertEqual(updated[0]["status"], "MATCHED")
        self.assertEqual(updated[0]["paths"], ["/tmp/人工上传_总方案信息表.docx"])
        self.assertEqual(updated[0]["gap_plan_item_id"], "GAP-0058")

    def test_bid_assembler_gap_plan_preserves_structural_items(self) -> None:
        build_assembly = load_assembler_script("build_assembly")

        with tempfile.TemporaryDirectory() as tmp:
            gap_plan_path = Path(tmp) / "gap_plan.json"
            gap_plan_path.write_text(
                json.dumps(
                    {
                        "items": [
                            {
                                "id": "GAP-0137",
                                "number": "技术附表B",
                                "title": "供货范围、消耗品及安装调试人员计划",
                                "status": "structural",
                                "gapReason": "结构性目录项，不直接要求素材。",
                            }
                        ]
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            plan = [
                {
                    "toc_idx": 136,
                    "level": 1,
                    "chapter_no_flat": "",
                    "chapter_no": "",
                    "title": "技术附表B 供货范围、消耗品及安装调试人员计划",
                    "paths": [],
                    "shifts": [],
                    "attach_modes": [],
                    "field_replace": False,
                    "status": "NEEDS_REVIEW",
                    "note": "招标/模板新增章节，需人工补素材",
                }
            ]

            updated = build_assembly.apply_gap_plan(plan, gap_plan_path)

        self.assertEqual(updated[0]["status"], "STRUCTURAL")
        self.assertEqual(updated[0]["paths"], [])
        self.assertEqual(updated[0]["gap_plan_item_id"], "GAP-0137")

    def test_bid_assembler_merges_oversize_material_instead_of_placeholder(self) -> None:
        merger = load_assembler_script("merger")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            source = root / "lib" / "投标机型业绩情况.docx"
            out = root / "out.docx"
            prep = root / "prep"

            master = Document()
            master.add_paragraph("")
            master.save(template)

            source.parent.mkdir(parents=True)
            doc = Document()
            doc.add_paragraph("投标机型业绩情况", style="Heading 1")
            doc.add_paragraph("这里是业绩正文")
            doc.save(source)

            plan = [
                {
                    "status": "MATCHED",
                    "level": 2,
                    "title": "投标机型业绩情况",
                    "chapter_no": "1.8",
                    "chapter_no_flat": "1.8",
                    "paths": [source.name],
                }
            ]
            original_stat = Path.stat

            class FakeStat:
                def __init__(self, wrapped):
                    self._wrapped = wrapped
                    self.st_size = 234 * 1024 * 1024

                def __getattr__(self, name):
                    return getattr(self._wrapped, name)

            def fake_stat(path, *args, **kwargs):
                value = original_stat(path, *args, **kwargs)
                if Path(path).resolve() == source.resolve():
                    return FakeStat(value)
                return value

            with patch.object(Path, "stat", fake_stat):
                stats = merger.merge(template, plan, source.parent, {}, prep, out)

            result = Document(str(out))
            text = "\n".join(para.text for para in result.paragraphs)

        self.assertEqual(stats["merged_materials"], 1)
        self.assertNotIn("大素材跳过", text)
        self.assertIn("这里是业绩正文", text)

    def test_bid_assembler_merger_remaps_material_headings_in_navigation(self) -> None:
        merger = load_assembler_script("merger")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            source = root / "lib" / "上海电气优势简介.docx"
            out = root / "out.docx"
            prep = root / "prep"

            master = Document()
            master.add_paragraph("")
            master.save(template)

            source.parent.mkdir(parents=True)
            doc = Document()
            doc.add_paragraph("上海电气优势简介", style="Heading 2")
            doc.add_paragraph("基本情况", style="Heading 3")
            stale = doc.add_paragraph("载荷仿真分析能力", style="Heading 1")
            p_pr = stale._p.get_or_add_pPr()
            outline = OxmlElement("w:outlineLvl")
            outline.set(qn("w:val"), "0")
            p_pr.append(outline)
            doc.add_paragraph("这里是正文")
            doc.save(source)

            plan = [
                {
                    "status": "MATCHED",
                    "level": 2,
                    "title": "上海电气优势简介",
                    "chapter_no": "1.9",
                    "chapter_no_flat": "1.9",
                    "paths": [source.name],
                }
            ]
            merger.merge(template, plan, source.parent, {}, prep, out)

            result = Document(str(out))
            headings = [
                para.text.strip().replace("  ", " ")
                for para in result.paragraphs
                if (para.style.name or "").startswith("Heading") and para.text.strip()
            ]
            text = "\n".join(para.text for para in result.paragraphs)

        self.assertEqual(
            headings,
            ["1.9 上海电气优势简介", "1.9.1 基本情况", "1.9.2 载荷仿真分析能力"],
        )
        self.assertIn("基本情况", text)
        self.assertIn("载荷仿真分析能力", text)

    def test_bid_assembler_merger_keeps_matching_child_heading_in_place(self) -> None:
        merger = load_assembler_script("merger")

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            source = root / "lib" / "项目技术承诺函.docx"
            out = root / "out.docx"
            prep = root / "prep"

            master = Document()
            master.add_paragraph("")
            master.save(template)

            source.parent.mkdir(parents=True)
            doc = Document()
            doc.add_paragraph("项目技术承诺函", style="Heading 1")
            doc.add_paragraph("发电小时数承诺函", style="Heading 2")
            doc.add_paragraph("这里是 4.1 正文")
            doc.save(source)

            plan = [
                {
                    "status": "MATCHED",
                    "level": 1,
                    "title": "项目技术承诺函",
                    "chapter_no": "第四章",
                    "chapter_no_flat": "4",
                    "paths": [source.name],
                },
                {
                    "status": "UNMATCHED",
                    "level": 2,
                    "title": "发电小时数承诺函",
                    "chapter_no": "4.1",
                    "chapter_no_flat": "4.1",
                    "paths": [],
                },
            ]
            merger.merge(template, plan, source.parent, {}, prep, out)

            result = Document(str(out))
            headings = [
                para.text.strip().replace("  ", " ")
                for para in result.paragraphs
                if (para.style.name or "").startswith("Heading") and para.text.strip()
            ]
            text = "\n".join(para.text for para in result.paragraphs)

        self.assertEqual(headings, ["第四章 项目技术承诺函", "4.1 发电小时数承诺函"])
        self.assertIn("这里是 4.1 正文", text)
        self.assertNotIn("[缺失：发电小时数承诺函", text)

    def test_bid_assembler_verify_uses_direct_outline_level_priority(self) -> None:
        verify = load_assembler_script("verify")

        with tempfile.TemporaryDirectory() as tmp:
            docx_path = Path(tmp) / "stale-outline.docx"
            doc = Document()
            stale = doc.add_paragraph("1.9.4  载荷仿真分析能力", style="Heading 3")
            p_pr = stale._p.get_or_add_pPr()
            outline = OxmlElement("w:outlineLvl")
            outline.set(qn("w:val"), "0")
            p_pr.append(outline)
            doc.save(docx_path)

            scan = verify.scan_docx(docx_path)

        self.assertEqual(scan["heading_counts"], {"Heading 1": 1})
        self.assertIn("1.9.4  载荷仿真分析能力", scan["invalid_h1"])

if __name__ == "__main__":
    unittest.main()
