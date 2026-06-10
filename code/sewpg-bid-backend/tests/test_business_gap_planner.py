from __future__ import annotations

import json
import subprocess
import sys
import tempfile
import unittest
from base64 import b64encode
from pathlib import Path
from unittest.mock import patch
from docx import Document

from app.services.bid_outline_state import confirm_outline_state, save_generated_outline_state
from app.services.bid_parse_state import complete_parse_state, update_parse_result_state, update_template_files_state
from app.services.bid_project_state import update_template_fallback_state


def _update_parse_result_for_tests(store, project_id: str, parse_result: dict, *, parse_storage: dict | None = None) -> dict:
    project = store.require_project_for_update(project_id)
    payload = update_parse_result_state(project, parse_result, parse_storage=parse_storage)
    store.persist_project_state(project)
    return payload


def _complete_parse_for_tests(
    store,
    project_id: str,
    tender_files: list[dict],
    template_files: list[dict],
    *,
    summary: dict | None = None,
    parse_storage: dict | None = None,
) -> dict:
    project = store.require_project_for_update(project_id)
    payload = complete_parse_state(
        project,
        tender_files,
        template_files,
        summary=summary,
        parse_storage=parse_storage,
    )
    store.persist_project_state(project)
    return payload


def _update_template_files_for_tests(store, project_id: str, template_files: list[dict]) -> dict:
    project = store.require_project_for_update(project_id)
    payload = update_template_files_state(project, template_files)
    store.persist_project_state(project)
    return payload


def _update_template_fallback_for_tests(store, project_id: str, data: dict) -> dict:
    project = store.require_project_for_update(project_id)
    payload = update_template_fallback_state(project, data)
    store.persist_project_state(project)
    return payload


def _save_generated_outline_for_tests(
    store,
    project_id: str,
    *,
    nodes: list[dict],
    generated_at: str,
    summary: str,
) -> dict:
    project = store.require_project_for_update(project_id)
    payload = save_generated_outline_state(
        project,
        nodes=nodes,
        generated_at=generated_at,
        summary=summary,
    )
    store.persist_project_state(project)
    return payload


def _confirm_outline_for_tests(store, project_id: str) -> dict:
    project = store.require_project_for_update(project_id)
    payload = confirm_outline_state(project)
    store.persist_project_state(project)
    return payload


class BusinessGapPlannerTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        self.client = None

    def _setup_app_test(self) -> None:
        from fastapi.testclient import TestClient
        from app.core.config import settings
        from app.main import app
        from app.services.store import store

        base = Path(self.temp_dir.name)
        settings.uploads_dir = base / "uploads"
        settings.documents_dir = base / "documents"
        settings.parsed_dir = base / "parsed"
        settings.ensure_dirs()
        store.reset_for_tests()
        self.client = TestClient(app, base_url="http://127.0.0.1:8000")

    def tearDown(self) -> None:
        if self.client is not None:
            self.client.close()
        self.temp_dir.cleanup()

    def test_businessgap_runner_generates_toc_based_plan_and_parser_artifacts(self) -> None:
        backend_root = Path(__file__).resolve().parents[1]
        script_path = backend_root / "opencode" / "skills" / "bid-business-gap-planner" / "scripts" / "run_from_manifest.py"
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_path = root / "toc.json"
            parse_path = root / "parse_result.json"
            output_path = root / "business_gap_plan.json"
            letter_path = root / "保密承诺书.docx"
            letter_path.write_bytes(b"fake-docx")
            toc_path.write_text(
                json.dumps(
                    {
                        "schema_version": "bid-toc-json-v1",
                        "items": [
                            {"order": 1, "number": "一", "title": "投标函", "level": 1},
                            {"order": 2, "number": "九", "title": "投标人需要说明的其他内容", "level": 1},
                            {"order": 3, "number": "9.1", "title": "保密承诺书", "level": 2},
                            {"order": 4, "number": "七", "title": "机型认证证书", "level": 1},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            parse_path.write_text(
                json.dumps(
                    {
                        "status": "completed",
                        "structured": {
                            "schemaVersion": "bid-business-tender-structured-v1",
                            "commitmentLetters": [
                                {
                                    "id": "COMMIT-0001",
                                    "title": "保密承诺书",
                                    "docxPath": str(letter_path),
                                    "assetReviewStatus": "pending_review",
                                    "triggerText": "投标人须提供保密承诺书。",
                                }
                            ],
                            "appendices": [],
                        },
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            manifest_path = root / "business_gap_input.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-BG-001",
                        "projectName": "商务缺口测试",
                        "bidType": "商务标",
                        "workDir": str(root),
                        "tocJsonPath": str(toc_path),
                        "parseResultPath": str(parse_path),
                        "businessWikiDir": "",
                        "materialScope": {"bidType": "商务标", "readableScopes": []},
                        "materialIndex": [
                            {
                                "id": "RAW-0001",
                                "name": "EW6.25机型认证证书.pdf",
                                "folderPath": "商务标/通用素材/专题证书库/机型认证证书",
                                "materialTier": "standard",
                                "turbineModelLabel": "EW6.25",
                            }
                        ],
                        "selectedBusinessTurbineModel": {"model": "EW6.25"},
                        "outputFile": str(output_path),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            completed = subprocess.run(
                [sys.executable, str(script_path), "--manifest", str(manifest_path), "--response", "summary"],
                check=True,
                capture_output=True,
                text=True,
            )
            summary = json.loads(completed.stdout)
            plan = json.loads(output_path.read_text(encoding="utf-8"))

        self.assertEqual(summary["schemaVersion"], "bid-business-gap-plan-v1")
        self.assertEqual(summary["coverageStatus"], "complete")
        self.assertEqual(len(plan["tocRefs"]), 4)
        titles = [task["title"] for task in plan["tasks"]]
        self.assertIn("投标函", titles)
        commitment = next(task for task in plan["tasks"] if task["title"] == "保密承诺书")
        self.assertEqual(commitment["status"], "review_required")
        self.assertEqual(commitment["resolvedArtifacts"][0]["artifactId"], "COMMIT-0001")
        self.assertEqual(commitment["assemblyMode"], "template_fill_docx")
        self.assertEqual(commitment["fillPlan"]["mode"], "template_fill_docx")
        certificate = next(task for task in plan["tasks"] if "机型认证" in task["title"])
        self.assertTrue(certificate["candidateMaterials"])
        self.assertEqual(certificate["assemblyMode"], "embed_scan_or_image")
        self.assertEqual(certificate["materialUsage"], "embed_scan")

    def test_businessgap_runner_uses_business_wiki_index_for_candidates_and_risks(self) -> None:
        backend_root = Path(__file__).resolve().parents[1]
        script_path = backend_root / "opencode" / "skills" / "bid-business-gap-planner" / "scripts" / "run_from_manifest.py"
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_path = root / "toc.json"
            parse_path = root / "parse_result.json"
            output_path = root / "business_gap_plan.json"
            toc_path.write_text(
                json.dumps(
                    {
                        "schema_version": "bid-toc-json-v1",
                        "items": [
                            {"order": 1, "number": "七", "title": "机型认证证书", "level": 1},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            parse_path.write_text(json.dumps({"status": "completed", "structured": {}}, ensure_ascii=False), encoding="utf-8")
            manifest_path = root / "business_gap_input.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-BG-WIKI",
                        "projectName": "商务 Wiki 候选测试",
                        "bidType": "商务标",
                        "workDir": str(root),
                        "tocJsonPath": str(toc_path),
                        "parseResultPath": str(parse_path),
                        "businessWikiDir": "",
                        "businessWikiIndex": {
                            "schemaVersion": "bid-business-wiki-index-v1",
                            "mappingRows": [
                                {
                                    "module_name": "08-资格证明文件模块（附件7）",
                                    "module_code": "BM-08",
                                    "candidate_card_ids": "biz-card-RAW-0007",
                                    "usage_mode": "attach_whole",
                                }
                            ],
                            "evidenceCards": [
                                {
                                    "card_id": "biz-card-RAW-0007",
                                    "material_id": "RAW-0007",
                                    "title": "EW6.25机型认证证书",
                                    "path": "商务标/通用素材/专题证书库/机型认证证书/EW6.25机型认证证书.pdf",
                                    "material_tier": "通用素材",
                                    "business_category": "专题证书",
                                    "usage_mode": "attach_whole",
                                    "validity_status": "pending_verify",
                                    "expiry_date": "待核验",
                                    "turbine_models": ["EW6.25"],
                                    "needs_human_confirm": "yes",
                                    "ocr_confidence": "0.60",
                                }
                            ],
                            "evidenceSegments": [
                                {
                                    "segment_id": "biz-seg-RAW-0007-cert",
                                    "card_id": "biz-card-RAW-0007",
                                    "material_id": "RAW-0007",
                                    "title": "EW6.25机型认证证书扫描件",
                                    "segment_type": "pdf_attachment",
                                    "segment_scope": "card_primary",
                                    "business_category": "专题证书",
                                    "document_type": "证书/资质文件",
                                    "usage_mode": "attach_whole",
                                    "path": "商务标/通用素材/专题证书库/机型认证证书/EW6.25机型认证证书.pdf",
                                    "source_pages": "第 1-2 页",
                                    "summary": "EW6.25 机型认证证书，需核验证书有效期。",
                                    "keywords": ["EW6.25", "机型认证", "证书"],
                                    "validity_status": "pending_verify",
                                    "expiry_date": "待核验",
                                    "turbine_models": ["EW6.25"],
                                    "needs_human_confirm": "yes",
                                    "ocr_confidence": "0.60",
                                }
                            ],
                        },
                        "materialScope": {"bidType": "商务标", "readableScopes": []},
                        "materialIndex": [
                            {
                                "id": "RAW-0007",
                                "name": "证书扫描件.pdf",
                                "folderPath": "商务标/通用素材/专题证书库/机型认证证书",
                                "materialTier": "standard",
                                "turbineModelLabel": "",
                            }
                        ],
                        "selectedBusinessTurbineModel": {"model": "EW6.25"},
                        "outputFile": str(output_path),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            subprocess.run(
                [sys.executable, str(script_path), "--manifest", str(manifest_path), "--response", "summary"],
                check=True,
                capture_output=True,
                text=True,
            )
            plan = json.loads(output_path.read_text(encoding="utf-8"))

        certificate = next(task for task in plan["tasks"] if "机型认证" in task["title"])
        candidate = certificate["candidateMaterials"][0]
        self.assertEqual(candidate["materialId"], "RAW-0007")
        self.assertEqual(candidate["wikiCardId"], "biz-card-RAW-0007")
        self.assertEqual(candidate["evidenceSegmentId"], "biz-seg-RAW-0007-cert")
        self.assertEqual(candidate["evidenceSegmentTitle"], "EW6.25机型认证证书扫描件")
        self.assertEqual(candidate["evidenceSourcePages"], "第 1-2 页")
        self.assertIn("Wiki", candidate["reason"])
        self.assertEqual(certificate["assemblyMode"], "attach_whole_file")
        self.assertEqual(certificate["materialUsage"], "attach_whole")
        self.assertEqual(certificate["selectedEvidenceSegments"][0]["segmentId"], "biz-seg-RAW-0007-cert")
        self.assertEqual(certificate["fillPlan"]["mode"], "attach_whole_file")
        self.assertIn("wiki_validity_pending_verify", certificate["riskFlags"])
        self.assertIn("wiki_ocr_low_confidence", certificate["riskFlags"])

    def test_businessgap_runner_prioritizes_manual_material_feedback(self) -> None:
        backend_root = Path(__file__).resolve().parents[1]
        script_path = backend_root / "opencode" / "skills" / "bid-business-gap-planner" / "scripts" / "run_from_manifest.py"
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_path = root / "toc.json"
            parse_path = root / "parse_result.json"
            output_path = root / "business_gap_plan.json"
            toc_path.write_text(
                json.dumps(
                    {
                        "schema_version": "bid-toc-json-v1",
                        "items": [
                            {"order": 1, "number": "7.1", "title": "机型认证证书", "level": 1},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            parse_path.write_text(json.dumps({"status": "completed", "structured": {}}, ensure_ascii=False), encoding="utf-8")
            manifest_path = root / "business_gap_input.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-BG-FEEDBACK",
                        "projectName": "商务反馈候选测试",
                        "bidType": "商务标",
                        "workDir": str(root),
                        "tocJsonPath": str(toc_path),
                        "parseResultPath": str(parse_path),
                        "businessWikiDir": "",
                        "businessWikiIndex": {"schemaVersion": "bid-business-wiki-index-v1"},
                        "materialScope": {"bidType": "商务标", "readableScopes": []},
                        "materialIndex": [
                            {
                                "id": "RAW-FB-001",
                                "name": "人工纠偏过的EW6.25机型认证证书.pdf",
                                "folderPath": "商务标/通用素材/专题证书库/机型认证证书",
                                "materialTier": "standard",
                            }
                        ],
                        "materialFeedback": {
                            "schemaVersion": "bid-business-material-feedback-v1",
                            "items": [
                                {
                                    "feedbackKey": "biz-feedback-test",
                                    "taskTitle": "机型认证证书",
                                    "moduleKey": "qualification_compliance_certificates",
                                    "taskType": "certificate",
                                    "assemblyMode": "embed_scan_or_image",
                                    "materialUsage": "embed_scan",
                                    "materialId": "RAW-FB-001",
                                    "materialName": "人工纠偏过的EW6.25机型认证证书.pdf",
                                    "folderPath": "商务标/通用素材/专题证书库/机型认证证书",
                                    "evidenceSegmentId": "SEG-FB-001",
                                    "evidenceSegmentTitle": "EW6.25机型认证证书页",
                                    "evidenceSummary": "人工在商务 S3 中确认该材料支撑机型认证证书任务。",
                                }
                            ],
                        },
                        "selectedBusinessTurbineModel": {"model": "EW6.25"},
                        "outputFile": str(output_path),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            subprocess.run(
                [sys.executable, str(script_path), "--manifest", str(manifest_path), "--response", "summary"],
                check=True,
                capture_output=True,
                text=True,
            )
            plan = json.loads(output_path.read_text(encoding="utf-8"))

        certificate = next(task for task in plan["tasks"] if "机型认证" in task["title"])
        candidate = certificate["candidateMaterials"][0]
        self.assertEqual(candidate["materialId"], "RAW-FB-001")
        self.assertEqual(candidate["feedbackKey"], "biz-feedback-test")
        self.assertIn("人工纠偏反馈命中", candidate["reason"])
        self.assertIn("manual_feedback_candidate", certificate["riskFlags"])
        self.assertEqual(certificate["selectedEvidenceSegments"][0]["segmentId"], "SEG-FB-001")

    def test_businessgap_runner_attaches_scoring_asset_to_scoring_index_task(self) -> None:
        backend_root = Path(__file__).resolve().parents[1]
        script_path = backend_root / "opencode" / "skills" / "bid-business-gap-planner" / "scripts" / "run_from_manifest.py"
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_path = root / "toc.json"
            parse_path = root / "parse_result.json"
            output_path = root / "business_gap_plan.json"
            scoring_path = root / "商务评分标准表.docx"
            scoring_path.write_bytes(b"fake-scoring-docx")
            toc_path.write_text(
                json.dumps(
                    {
                        "schema_version": "bid-toc-json-v1",
                        "items": [
                            {"order": 1, "number": "一", "title": "商务评分索引表", "level": 1},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            parse_path.write_text(
                json.dumps(
                    {
                        "status": "completed",
                        "structured": {
                            "businessScoringAsset": {
                                "id": "BIZ-SCORING-001",
                                "fileName": "商务评分标准表.docx",
                                "docxPath": str(scoring_path),
                                "reviewStatus": "approved",
                            }
                        },
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            manifest_path = root / "business_gap_input.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-BG-SCORING",
                        "projectName": "商务评分索引测试",
                        "bidType": "商务标",
                        "workDir": str(root),
                        "tocJsonPath": str(toc_path),
                        "parseResultPath": str(parse_path),
                        "businessWikiDir": "",
                        "materialScope": {"bidType": "商务标", "readableScopes": []},
                        "materialIndex": [],
                        "selectedBusinessTurbineModel": {},
                        "outputFile": str(output_path),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            subprocess.run(
                [sys.executable, str(script_path), "--manifest", str(manifest_path), "--response", "summary"],
                check=True,
                capture_output=True,
                text=True,
            )
            plan = json.loads(output_path.read_text(encoding="utf-8"))

        scoring_task = next(task for task in plan["tasks"] if task["title"] == "商务评分索引表")
        self.assertEqual(scoring_task["status"], "ready")
        self.assertEqual(scoring_task["resolvedArtifacts"][0]["artifactId"], "BIZ-SCORING-001")
        self.assertEqual(scoring_task["resolvedArtifacts"][0]["sourceMode"], "parsed_business_scoring")

    def test_businessgap_runner_negative_rules_avoid_performance_material_for_bid_letter(self) -> None:
        backend_root = Path(__file__).resolve().parents[1]
        script_path = backend_root / "opencode" / "skills" / "bid-business-gap-planner" / "scripts" / "run_from_manifest.py"
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_path = root / "toc.json"
            parse_path = root / "parse_result.json"
            output_path = root / "business_gap_plan.json"
            toc_path.write_text(
                json.dumps(
                    {
                        "schema_version": "bid-toc-json-v1",
                        "items": [
                            {"order": 1, "number": "一", "title": "投标函", "level": 1},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            parse_path.write_text(json.dumps({"status": "completed", "structured": {}}, ensure_ascii=False), encoding="utf-8")
            manifest_path = root / "business_gap_input.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-BG-NEGATIVE",
                        "projectName": "商务负面规则测试",
                        "bidType": "商务标",
                        "workDir": str(root),
                        "tocJsonPath": str(toc_path),
                        "parseResultPath": str(parse_path),
                        "businessWikiDir": "",
                        "materialScope": {"bidType": "商务标", "readableScopes": []},
                        "materialIndex": [
                            {
                                "id": "RAW-PERF-001",
                                "name": "投标项目业绩合同扫描件.pdf",
                                "folderPath": "商务标/共用业绩库/合同扫描件",
                                "materialTier": "standard",
                                "sourceType": "performance_library",
                                "businessMaterialKind": "performance",
                                "businessMaterialKindLabel": "共用业绩",
                            }
                        ],
                        "selectedBusinessTurbineModel": {},
                        "outputFile": str(output_path),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            subprocess.run(
                [sys.executable, str(script_path), "--manifest", str(manifest_path), "--response", "summary"],
                check=True,
                capture_output=True,
                text=True,
            )
            plan = json.loads(output_path.read_text(encoding="utf-8"))

        bid_letter = next(task for task in plan["tasks"] if task["title"] == "投标函")
        self.assertEqual(bid_letter["candidateMaterials"], [])
        self.assertEqual(bid_letter["status"], "needs_input")

    def test_businessgap_runner_uses_shared_performance_library_for_performance_task(self) -> None:
        backend_root = Path(__file__).resolve().parents[1]
        script_path = backend_root / "opencode" / "skills" / "bid-business-gap-planner" / "scripts" / "run_from_manifest.py"
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_path = root / "toc.json"
            parse_path = root / "parse_result.json"
            output_path = root / "business_gap_plan.json"
            toc_path.write_text(
                json.dumps(
                    {
                        "schema_version": "bid-toc-json-v1",
                        "items": [
                            {"order": 1, "number": "一", "title": "投标函", "level": 1},
                            {"order": 2, "number": "九", "title": "近年类似项目业绩表", "level": 1},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            parse_path.write_text(json.dumps({"status": "completed", "structured": {}}, ensure_ascii=False), encoding="utf-8")
            manifest_path = root / "business_gap_input.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-BG-PERFORMANCE",
                        "projectName": "商务业绩候选测试",
                        "bidType": "商务标",
                        "workDir": str(root),
                        "tocJsonPath": str(toc_path),
                        "parseResultPath": str(parse_path),
                        "businessWikiDir": "",
                        "materialScope": {"bidType": "商务标", "readableScopes": []},
                        "materialIndex": [
                            {
                                "id": "PERF-0009",
                                "materialId": "PERF-0009",
                                "name": "华能风电机组供货业绩",
                                "fileName": "华能风电机组供货业绩.docx",
                                "folderPath": "商务标/共用业绩库/华能集团",
                                "path": "商务标/共用业绩库/华能集团/华能风电机组供货业绩",
                                "materialTier": "customer",
                                "libraryScope": "customer",
                                "sourceType": "performance_library",
                                "candidateType": "performance_record",
                                "businessMaterialKind": "performance",
                                "businessMaterialKindLabel": "共用业绩",
                                "cleanStatus": "original_only",
                                "tags": ["业绩", "合同", "中标"],
                                "keywords": ["业绩", "业绩证明", "合同", "中标通知书", "240h"],
                                "summary": "华能集团；风电机组供货；中标通知书；合同；240h试运行",
                                "businessCategory": "业绩证明",
                                "documentType": "业绩记录",
                            }
                        ],
                        "selectedBusinessTurbineModel": {},
                        "outputFile": str(output_path),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            subprocess.run(
                [sys.executable, str(script_path), "--manifest", str(manifest_path), "--response", "summary"],
                check=True,
                capture_output=True,
                text=True,
            )
            plan = json.loads(output_path.read_text(encoding="utf-8"))

        performance_task = next(task for task in plan["tasks"] if "业绩" in task["title"])
        performance_candidate = performance_task["candidateMaterials"][0]
        self.assertEqual(performance_candidate["materialId"], "PERF-0009")
        self.assertEqual(performance_candidate["sourceType"], "performance_library")
        self.assertEqual(performance_candidate["candidateType"], "performance_record")
        self.assertIn("共用业绩库候选", performance_candidate["reason"])
        bid_letter = next(task for task in plan["tasks"] if task["title"] == "投标函")
        self.assertEqual(bid_letter["candidateMaterials"], [])

    def test_businessgap_runner_uses_performance_package_candidates_for_performance_task(self) -> None:
        backend_root = Path(__file__).resolve().parents[1]
        script_path = backend_root / "opencode" / "skills" / "bid-business-gap-planner" / "scripts" / "run_from_manifest.py"
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_path = root / "toc.json"
            parse_path = root / "parse_result.json"
            output_path = root / "business_gap_plan.json"
            toc_path.write_text(
                json.dumps(
                    {
                        "schema_version": "bid-toc-json-v1",
                        "items": [
                            {"order": 1, "number": "一", "title": "投标函", "level": 1},
                            {"order": 2, "number": "九", "title": "近年类似项目业绩表", "level": 1},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            parse_path.write_text(json.dumps({"status": "completed", "structured": {}}, ensure_ascii=False), encoding="utf-8")
            manifest_path = root / "business_gap_input.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-BG-PERFORMANCE-PKG",
                        "projectName": "商务业绩包候选测试",
                        "bidType": "商务标",
                        "workDir": str(root),
                        "tocJsonPath": str(toc_path),
                        "parseResultPath": str(parse_path),
                        "businessWikiDir": "",
                        "materialScope": {"bidType": "商务标", "readableScopes": []},
                        "materialIndex": [
                            {
                                "id": "PERITEM-0268",
                                "materialId": "PERITEM-0268",
                                "categoryId": "PERCAT-0011",
                                "name": "华电新疆喀什 2x66 万千瓦",
                                "fileName": "001-华电新疆喀什 2x66 万千瓦_合同.docx",
                                "folderPath": "商务标/共用业绩库/陆上6MW业绩",
                                "path": "商务标/共用业绩库/陆上6MW业绩/华电新疆喀什 2x66 万千瓦",
                                "materialTier": "standard",
                                "libraryScope": "standard",
                                "sourceType": "performance_package",
                                "candidateType": "performance_item",
                                "businessMaterialKind": "performance",
                                "businessMaterialKindLabel": "共用业绩",
                                "cleanStatus": "original_only",
                                "tags": ["陆上"],
                                "keywords": ["陆上6MW业绩", "业绩", "华电", "11-230", "2024"],
                                "summary": "华电；11-230；2 台；2024",
                                "businessCategory": "业绩证明",
                                "documentType": "业绩明细",
                                "attachments": [
                                    {
                                        "id": "PERITEMATT-0118",
                                        "categoryId": "PERCAT-0011",
                                        "itemId": "PERITEM-0268",
                                        "fileName": "001-华电新疆喀什 2x66 万千瓦_合同.docx",
                                        "matchConfidence": 95,
                                        "matchMethod": "project_name",
                                    }
                                ],
                            }
                        ],
                        "selectedBusinessTurbineModel": {},
                        "outputFile": str(output_path),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            subprocess.run(
                [sys.executable, str(script_path), "--manifest", str(manifest_path), "--response", "summary"],
                check=True,
                capture_output=True,
                text=True,
            )
            plan = json.loads(output_path.read_text(encoding="utf-8"))

        performance_task = next(task for task in plan["tasks"] if "业绩" in task["title"])
        performance_candidate = performance_task["candidateMaterials"][0]
        self.assertEqual(performance_candidate["materialId"], "PERITEM-0268")
        self.assertEqual(performance_candidate["sourceType"], "performance_package")
        self.assertEqual(performance_candidate["candidateType"], "performance_item")
        self.assertIn("共用业绩库候选", performance_candidate["reason"])
        bid_letter = next(task for task in plan["tasks"] if task["title"] == "投标函")
        self.assertEqual(bid_letter["candidateMaterials"], [])

    def test_performance_package_candidate_builders_map_items_and_categories(self) -> None:
        from app.services.business_gap_planning import (
            _performance_package_candidate_from_category,
            _performance_package_candidate_from_item,
        )

        category = {
            "id": "PERCAT-0011",
            "name": "陆上6MW业绩",
            "scene": "陆上",
            "powerRating": "6MW及以上",
            "scope": "standard",
            "tags": ["陆上"],
            "turbineModels": ["6.25-182", "11-230"],
            "summaryFileName": "陆上6MW业绩_汇总表.docx",
            "summary": "陆上 6MW 及以上业绩汇总",
            "reviewStatus": "draft",
            "itemCount": 104,
        }
        item = {
            "id": "PERITEM-0268",
            "categoryId": "PERCAT-0011",
            "rowIndex": 1,
            "projectName": "华电新疆喀什 2x66 万千瓦",
            "customerName": "华电",
            "turbineModels": ["11-230"],
            "contractQuantity": "2",
            "commissionedCapacityMw": "132",
            "deliveryOrOperationTime": "2024",
            "operationYear": 2024,
            "values": {"项目名称": "华电新疆喀什 2x66 万千瓦", "业主单位": "华电"},
            "attachments": [
                {
                    "id": "PERITEMATT-0118",
                    "categoryId": "PERCAT-0011",
                    "itemId": "PERITEM-0268",
                    "attachmentType": "contract_item",
                    "fileName": "001-华电新疆喀什 2x66 万千瓦_合同.docx",
                    "sizeBytes": 2048,
                    "matchConfidence": 95,
                    "matchMethod": "project_name",
                }
            ],
        }

        category_candidate = _performance_package_candidate_from_category(category)
        self.assertEqual(category_candidate["sourceType"], "performance_package")
        self.assertEqual(category_candidate["candidateType"], "performance_category")
        self.assertEqual(category_candidate["path"], "商务标/共用业绩库/陆上6MW业绩")
        self.assertEqual(category_candidate["businessMaterialKind"], "performance")
        self.assertIn("陆上6MW业绩", category_candidate["keywords"])

        item_candidate = _performance_package_candidate_from_item(category, item)
        self.assertEqual(item_candidate["materialId"], "PERITEM-0268")
        self.assertEqual(item_candidate["candidateType"], "performance_item")
        self.assertEqual(item_candidate["cleanStatus"], "original_only")
        self.assertEqual(item_candidate["fileName"], "001-华电新疆喀什 2x66 万千瓦_合同.docx")
        self.assertEqual(item_candidate["attachments"][0]["matchMethod"], "project_name")
        self.assertEqual(item_candidate["attachments"][0]["matchConfidence"], 95)
        self.assertIn("11-230", item_candidate["keywords"])
        self.assertIn("华电", item_candidate["summary"])

    def test_businessgap_runner_emits_template_candidates_from_template_index(self) -> None:
        backend_root = Path(__file__).resolve().parents[1]
        script_path = backend_root / "opencode" / "skills" / "bid-business-gap-planner" / "scripts" / "run_from_manifest.py"
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            toc_path = root / "toc.json"
            parse_path = root / "parse_result.json"
            output_path = root / "business_gap_plan.json"
            template_path = root / "投标函格式模板.docx"
            doc = Document()
            doc.add_paragraph("投标函")
            doc.add_paragraph("项目名称：")
            doc.save(template_path)
            toc_path.write_text(
                json.dumps(
                    {
                        "schema_version": "bid-toc-json-v1",
                        "items": [
                            {"order": 1, "number": "一", "title": "投标函", "level": 1},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            parse_path.write_text(json.dumps({"status": "completed", "structured": {}}, ensure_ascii=False), encoding="utf-8")
            manifest_path = root / "business_gap_input.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-BG-TEMPLATE",
                        "projectName": "商务模板候选测试",
                        "bidType": "商务标",
                        "workDir": str(root),
                        "tocJsonPath": str(toc_path),
                        "parseResultPath": str(parse_path),
                        "businessWikiDir": "",
                        "materialScope": {"bidType": "商务标", "readableScopes": []},
                        "materialIndex": [],
                        "templateIndex": [
                            {
                                "templateId": "TPL-PROJECT-001",
                                "templateName": "投标函格式模板.docx",
                                "fileName": template_path.name,
                                "filePath": str(template_path),
                                "sourceMode": "project_uploaded_bid_template",
                                "sourceLabel": "项目上传模板",
                                "templateScope": "project",
                                "score": 0.92,
                                "reason": "S1/S2 项目上传投标模板",
                            }
                        ],
                        "selectedBusinessTurbineModel": {},
                        "outputFile": str(output_path),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            subprocess.run(
                [sys.executable, str(script_path), "--manifest", str(manifest_path), "--response", "summary"],
                check=True,
                capture_output=True,
                text=True,
            )
            plan = json.loads(output_path.read_text(encoding="utf-8"))

        bid_letter = next(task for task in plan["tasks"] if task["title"] == "投标函")
        self.assertEqual(bid_letter["status"], "review_required")
        self.assertEqual(bid_letter["assemblyMode"], "template_fill_docx")
        self.assertEqual(bid_letter["materialUsage"], "fill_template")
        self.assertEqual(bid_letter["templateCandidates"][0]["templateId"], "TPL-PROJECT-001")
        self.assertEqual(bid_letter["templateCandidates"][0]["sourceMode"], "project_uploaded_bid_template")
        self.assertEqual(bid_letter["candidateMaterials"], [])
        self.assertIn("candidate_template_unconfirmed", bid_letter["riskFlags"])

    def test_business_gap_api_uses_business_workspace_and_keeps_technical_gap_state_empty(self) -> None:
        self._setup_app_test()
        from app.core.config import settings
        from app.services.bid_runtime_state import now_iso
        from app.services.store import store
        from app.services.workspace_artifacts import business_workspace_dir, technical_workspace_dir

        project = store.create_project({"name": "商务S3项目", "customerName": "华能集团", "bidType": "商务标"})
        project_id = project["id"]
        _update_template_fallback_for_tests(store, project_id, {"enabled": False})
        business_workspace = business_workspace_dir(project_id)
        business_workspace.mkdir(parents=True, exist_ok=True)
        letter_path = business_workspace / "commitment-letters" / "保密承诺书.docx"
        letter_path.parent.mkdir(parents=True, exist_ok=True)
        letter_path.write_bytes(b"fake-docx")
        _complete_parse_for_tests(
            store,
            project_id,
            tender_files=[{"id": "TEN-1", "name": "商务招标文件.md", "path": str(settings.uploads_dir / "tender.md"), "size_label": "1 KB"}],
            template_files=[],
            summary={"fileCount": 1, "extractedCount": 1, "textLength": 20, "textPreview": "", "warnings": []},
            parse_storage={
                "projectDir": str(business_workspace),
                "combinedTextPath": str(business_workspace / "combined.txt"),
                "manifestPath": "",
            },
        )
        _update_parse_result_for_tests(
            store,
            project_id,
            {
                "status": "completed",
                "structured": {
                    "schemaVersion": "bid-business-tender-structured-v1",
                    "commitmentLetters": [
                        {
                            "id": "COMMIT-0001",
                            "title": "保密承诺书",
                            "docxPath": str(letter_path),
                            "assetReviewStatus": "pending_review",
                        }
                    ],
                    "appendices": [],
                    "projectFactFields": [
                        {
                            "fieldKey": "projectName",
                            "label": "项目名称",
                            "value": "商务S3项目",
                            "category": "项目基础信息",
                            "required": True,
                            "confidence": 0.95,
                        },
                        {
                            "fieldKey": "tenderNo",
                            "label": "招标编号",
                            "value": "BIZ-2026-001",
                            "category": "项目基础信息",
                            "required": True,
                            "confidence": 0.94,
                        },
                        {
                            "fieldKey": "tenderer",
                            "label": "招标人",
                            "value": "华能集团",
                            "category": "项目基础信息",
                            "required": True,
                            "confidence": 0.9,
                        },
                    ],
                },
            },
        )
        _save_generated_outline_for_tests(
            store,
            project_id=project_id,
            nodes=[
                {"id": "OL-1", "title": "投标函", "children": []},
                {"id": "OL-2", "title": "投标人需要说明的其他内容", "children": [
                    {"id": "OL-2-1", "title": "保密承诺书", "children": []},
                ]},
            ],
            generated_at=now_iso(),
            summary="商务目录已生成。",
        )
        _confirm_outline_for_tests(store, project_id)
        store.update_stage(project_id, 2, {"status": "completed"})

        async def empty_raw_files(**kwargs):
            self.assertNotIn("bid_type", kwargs)
            self.assertTrue(str(kwargs.get("folder_path") or "").startswith("商务标/"))
            return {"items": [], "total": 0}

        with patch(
            "app.services.business_gap_planning.OpencodeClient.run_bid_business_gap_planner_with_trace",
            side_effect=RuntimeError("offline test fallback"),
        ), patch(
            "app.services.business_gap_planning.business_material_store.raw_files",
            side_effect=empty_raw_files,
        ):
            response = self.client.post(f"/api/business/projects/{project_id}/business-gaps/run")
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["status"], "completed")
        self.assertTrue(str(payload["plan"]["planFile"]).endswith("business_gap_plan.json"))
        self.assertIn("business-workspace/gaps", str(payload["plan"]["planFile"]))
        self.assertEqual(store._require(project_id)["gap_state"]["recognitionStatus"], "idle")
        self.assertTrue((business_workspace / "gaps" / "business_gap_input.json").exists())
        self.assertFalse((technical_workspace_dir(project_id) / "s4_gap_workdir" / "gap_plan.json").exists())

        get_response = self.client.get(f"/api/business/projects/{project_id}/business-gaps")
        self.assertEqual(get_response.status_code, 200)
        self.assertEqual(len(get_response.json()["tocRefs"]), 3)
        empty_ref = next(item for item in get_response.json()["tocRefs"] if item["title"] == "投标人需要说明的其他内容")
        manual_response = self.client.post(
            f"/api/business/projects/{project_id}/business-gaps/toc/{empty_ref['nodeId']}/manual-task",
            json={"title": "本章节补充说明材料"},
        )
        self.assertEqual(manual_response.status_code, 200)
        manual_task = manual_response.json()["task"]
        self.assertEqual(manual_task["sourceType"], "manual_user")
        self.assertEqual(manual_task["tocTarget"]["nodeId"], empty_ref["nodeId"])
        self.assertEqual(manual_task["status"], "needs_input")
        manual_ref = next(item for item in manual_response.json()["plan"]["tocRefs"] if item["nodeId"] == empty_ref["nodeId"])
        self.assertIn(manual_task["id"], manual_ref["taskIds"])
        manual_upload_response = self.client.post(
            f"/api/business/projects/{project_id}/business-gaps/tasks/{manual_task['id']}/upload",
            json={
                "files": [
                    {
                        "name": "本章节补充说明.png",
                        "mimeType": "image/png",
                        "data": "data:image/png;base64," + b64encode(b"fake-image").decode("ascii"),
                    }
                ]
            },
        )
        self.assertEqual(manual_upload_response.status_code, 200)
        self.assertEqual(manual_upload_response.json()["task"]["status"], "ready")
        self.assertIn("business-workspace/gaps/uploads", manual_upload_response.json()["artifact"]["filePath"])
        manual_ref_after_upload = next(
            item
            for item in manual_upload_response.json()["plan"]["tocRefs"]
            if item["nodeId"] == empty_ref["nodeId"]
        )
        self.assertEqual(manual_ref_after_upload["status"], "ready")
        manual_artifact = manual_upload_response.json()["artifact"]
        remove_manual_response = self.client.delete(
            f"/api/business/projects/{project_id}/business-gaps/tasks/{manual_task['id']}/artifacts/{manual_artifact['artifactId']}"
        )
        self.assertEqual(remove_manual_response.status_code, 200)
        self.assertEqual(remove_manual_response.json()["task"]["status"], "needs_input")
        manual_ref_after_remove = next(
            item
            for item in remove_manual_response.json()["plan"]["tocRefs"]
            if item["nodeId"] == empty_ref["nodeId"]
        )
        self.assertEqual(manual_ref_after_remove["status"], "partial")
        self.assertFalse(Path(manual_artifact["filePath"]).exists())
        task = next(item for item in get_response.json()["tasks"] if item["title"] == "保密承诺书")
        confirm_response = self.client.post(
            f"/api/business/projects/{project_id}/business-gaps/tasks/{task['id']}/confirm-artifact",
            json={"artifactId": "COMMIT-0001", "confirmed": True},
        )
        self.assertEqual(confirm_response.status_code, 200)
        self.assertEqual(confirm_response.json()["task"]["status"], "ready")

        bid_letter_task = next(item for item in get_response.json()["tasks"] if item["title"] == "投标函")
        mode_response = self.client.patch(
            f"/api/business/projects/{project_id}/business-gaps/tasks/{bid_letter_task['id']}",
            json={"assemblyMode": "template_fill_docx"},
        )
        self.assertEqual(mode_response.status_code, 200)
        mode_task = mode_response.json()["task"]
        self.assertEqual(mode_task["status"], "needs_input")
        self.assertEqual(mode_task["decision"], "fill_required")
        self.assertEqual(mode_task["materialUsage"], "fill_template")
        self.assertIn("missing_material", mode_task["riskFlags"])
        self.assertIn("template_missing_for_fill", mode_task["riskFlags"])
        self.assertFalse(mode_task.get("templateCandidates"))

        stored_project = store._require(project_id)
        stored_plan = stored_project["business_gap_state"]["plan"]
        bid_letter_stored = next(item for item in stored_plan["tasks"] if item["id"] == bid_letter_task["id"])
        bid_letter_stored["candidateMaterials"] = [
            {
                "materialId": "RAW-TPL-001",
                "materialName": "投标函格式模板.docx",
                "folderPath": "商务标/通用素材/通用模板底稿库/投标函空白模板",
                "wikiUsageMode": "fill_template",
                "score": 0.86,
            }
        ]
        store._persist_project(stored_project)
        candidate_mode_response = self.client.patch(
            f"/api/business/projects/{project_id}/business-gaps/tasks/{bid_letter_task['id']}",
            json={"assemblyMode": "template_fill_docx"},
        )
        self.assertEqual(candidate_mode_response.status_code, 200)
        candidate_mode_task = candidate_mode_response.json()["task"]
        self.assertEqual(candidate_mode_task["status"], "review_required")
        self.assertEqual(candidate_mode_task["decision"], "review_required")
        self.assertIn("candidate_template_unconfirmed", candidate_mode_task["riskFlags"])
        template_source = business_workspace / "项目上传投标函模板.docx"
        doc = Document()
        doc.add_paragraph("投标函")
        doc.add_paragraph("项目名称：")
        doc.save(template_source)
        _update_template_files_for_tests(
            store,
            project_id,
            [
                {
                    "id": "TPL-PROJECT-001",
                    "name": "项目上传投标函模板.docx",
                    "stored_name": template_source.name,
                    "size_bytes": template_source.stat().st_size,
                    "size_label": "1 KB",
                    "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "path": str(template_source),
                }
            ],
        )
        stored_project = store._require(project_id)
        stored_plan = stored_project["business_gap_state"]["plan"]
        bid_letter_stored = next(item for item in stored_plan["tasks"] if item["id"] == bid_letter_task["id"])
        bid_letter_stored["candidateMaterials"] = []
        bid_letter_stored["templateCandidates"] = []
        store._persist_project(stored_project)
        backfill_response = self.client.get(f"/api/business/projects/{project_id}/business-gaps")
        self.assertEqual(backfill_response.status_code, 200)
        backfilled_task = next(item for item in backfill_response.json()["tasks"] if item["id"] == bid_letter_task["id"])
        self.assertEqual(backfilled_task["status"], "review_required")
        self.assertEqual(backfilled_task["templateCandidates"][0]["templateId"], "TPL-PROJECT-001")
        self.assertEqual(backfilled_task["templateCandidates"][0]["sourceMode"], "project_uploaded_bid_template")
        select_template_response = self.client.post(
            f"/api/business/projects/{project_id}/business-gaps/tasks/{bid_letter_task['id']}/select-template",
            json={"template": {"templateId": "TPL-PROJECT-001"}},
        )
        self.assertEqual(select_template_response.status_code, 200)
        selected_template_task = select_template_response.json()["task"]
        self.assertEqual(selected_template_task["status"], "ready")
        self.assertEqual(selected_template_task["assemblyMode"], "template_fill_docx")
        self.assertEqual(select_template_response.json()["artifact"]["sourceMode"], "project_uploaded_bid_template")
        self.assertIn("business-workspace/gaps/selected-templates", select_template_response.json()["artifact"]["filePath"])

        upload_response = self.client.post(
            f"/api/business/projects/{project_id}/business-gaps/tasks/{bid_letter_task['id']}/upload",
            json={
                "files": [
                    {
                        "name": "补充授权材料.pdf",
                        "mimeType": "application/pdf",
                        "data": "data:application/pdf;base64," + b64encode(b"%PDF-business-gap").decode("ascii"),
                    }
                ]
            },
        )
        self.assertEqual(upload_response.status_code, 200)
        uploaded = upload_response.json()["artifact"]
        self.assertEqual(upload_response.json()["task"]["status"], "ready")
        self.assertEqual(upload_response.json()["task"]["handlingMode"], "manual_upload")
        self.assertEqual(uploaded["sourceMode"], "uploaded_in_business_s3")
        self.assertIn("business-workspace/gaps/uploads", uploaded["filePath"])
        self.assertTrue(Path(uploaded["filePath"]).exists())
        self.assertFalse((technical_workspace_dir(project_id) / "s4_gap_workdir" / "manual_upload").exists())
        self.assertEqual(uploaded["materialSyncStatus"], "not_synced")
        self.assertEqual(uploaded["materialSyncPolicy"], "manual_project_only")
        self.assertIn(f"商务标/项目素材/{project_id}/资格审查与商务响应成册", uploaded["materialTargetPath"])

        multipart_upload_response = self.client.post(
            f"/api/business/projects/{project_id}/business-gaps/tasks/{bid_letter_task['id']}/upload-files",
            data={"operator": "测试用户"},
            files=[
                (
                    "files",
                    (
                        "S3补充模板.docx",
                        b"business-gap-docx-bytes",
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
            ],
        )
        self.assertEqual(multipart_upload_response.status_code, 200)
        multipart_uploaded = multipart_upload_response.json()["artifact"]
        self.assertEqual(multipart_upload_response.json()["task"]["status"], "ready")
        self.assertEqual(multipart_uploaded["sourceMode"], "uploaded_in_business_s3")
        self.assertEqual(multipart_uploaded["operator"], "测试用户")
        self.assertEqual(Path(multipart_uploaded["filePath"]).read_bytes(), b"business-gap-docx-bytes")

        ai_mode_response = self.client.patch(
            f"/api/business/projects/{project_id}/business-gaps/tasks/{bid_letter_task['id']}",
            json={"assemblyMode": "ai_draft"},
        )
        self.assertEqual(ai_mode_response.status_code, 200)
        ai_mode_task = ai_mode_response.json()["task"]
        self.assertEqual(ai_mode_task["status"], "needs_input")
        self.assertEqual(ai_mode_task["decision"], "ai_draft_required")
        self.assertIn("ai_draft_required", ai_mode_task["riskFlags"])
        ai_draft_response = self.client.post(
            f"/api/business/projects/{project_id}/business-gaps/tasks/{bid_letter_task['id']}/ai-draft",
            json={"operator": "测试用户"},
        )
        self.assertEqual(ai_draft_response.status_code, 200)
        ai_draft_task = ai_draft_response.json()["task"]
        self.assertEqual(ai_draft_task["status"], "ready")
        self.assertEqual(ai_draft_task["decision"], "ready")
        self.assertNotIn("ai_draft_required", ai_draft_task["riskFlags"])

        async def fake_raw_upload(**kwargs):
            self.assertNotIn("bid_type", kwargs)
            self.assertEqual(kwargs["material_tier"], "project")
            self.assertIn(f"商务标/项目素材/{project_id}/资格审查与商务响应成册", kwargs["target_path"])
            return {
                "message": "mock upload",
                "items": [
                    {
                        "id": "RAW-SYNC-001",
                        "name": "补充授权材料.pdf",
                        "folderPath": kwargs["target_path"],
                        "bidType": "商务标",
                        "materialTier": "project",
                        "projectId": project_id,
                        "cleanStatus": "",
                        "cleanMessage": "",
                    }
                ],
                "cleaning": {"queued": 0, "jobs": []},
            }

        with patch("app.services.business_gap_service.business_material_store.raw_upload", side_effect=fake_raw_upload):
            sync_response = self.client.post(
                f"/api/business/projects/{project_id}/business-gaps/tasks/{bid_letter_task['id']}/sync-artifact-material",
                json={"artifactId": uploaded["artifactId"]},
            )
        self.assertEqual(sync_response.status_code, 200)
        synced = sync_response.json()["artifact"]
        self.assertEqual(synced["materialSyncStatus"], "synced_to_project_material")
        self.assertEqual(synced["wikiSyncStatus"], "wiki_rebuild_required")
        self.assertIn(f"商务标/项目素材/{project_id}/资格审查与商务响应成册", synced["materialTargetPath"])
        self.assertTrue(sync_response.json()["wikiRebuildRequired"])
        synced_material = sync_response.json()["material"]
        self.assertEqual(synced_material["bidType"], "商务标")
        self.assertEqual(synced_material["materialTier"], "project")
        self.assertEqual(synced_material["projectId"], project_id)
        self.assertIn("商务标/项目素材", synced_material["folderPath"])
        self.assertFalse(synced_material["folderPath"].startswith("技术标/"))

        async def fake_download_content(material_id: str) -> dict[str, str]:
            return {
                "fileId": material_id,
                "fileName": "商务素材证书.pdf",
                "bucket": "mock-bucket",
                "key": "mock-key",
                "mimeType": "application/pdf",
            }

        def fake_download_file(bucket: str, key: str, target_path: Path) -> Path:
            target_path.parent.mkdir(parents=True, exist_ok=True)
            target_path.write_bytes(b"%PDF-selected-material")
            return target_path

        async def fake_raw_files(**kwargs):
            self.assertNotIn("bid_type", kwargs)
            self.assertTrue(str(kwargs.get("folder_path") or "").startswith("商务标/"))
            return {
                "items": [
                    {
                        "id": "RAW-9001",
                        "name": "商务素材证书.pdf",
                        "folderPath": "商务标/通用素材/专题证书库",
                        "materialTier": "standard",
                        "cleanStatus": "",
                        "hasCleanedWord": False,
                        "size": 128,
                    }
                ],
                "total": 1,
            }

        with patch(
            "app.services.business_gap_service.business_material_store.raw_download_cleaned_content",
            side_effect=RuntimeError("no cleaned content in this test"),
        ), patch(
            "app.services.business_gap_service.business_material_store.raw_download_content",
            side_effect=fake_download_content,
        ), patch(
            "app.services.business_gap_service.minio_client.download_file",
            side_effect=fake_download_file,
        ):
            with patch(
                "app.services.business_gap_planning.business_material_store.raw_files",
                side_effect=fake_raw_files,
            ):
                selectable_response = self.client.get(
                    f"/api/business/projects/{project_id}/business-gaps/selectable-materials?keyword=证书"
                )
            self.assertEqual(selectable_response.status_code, 200)
            selectable_payload = selectable_response.json()
            self.assertEqual(selectable_payload["bidType"], "商务标")
            self.assertEqual(selectable_payload["items"][0]["materialId"], "RAW-9001")
            self.assertEqual(selectable_payload["segments"][0]["materialId"], "RAW-9001")
            self.assertTrue(selectable_payload["segments"][0]["evidenceSegmentId"])

            select_response = self.client.post(
                f"/api/business/projects/{project_id}/business-gaps/tasks/{bid_letter_task['id']}/select-material",
                json={
                    "materials": [
                        {
                            "materialId": "RAW-9001",
                            "materialName": "商务素材证书.pdf",
                            "folderPath": "商务标/通用素材/专题证书库",
                            "materialTier": "standard",
                            "businessMaterialKind": "fixed",
                            "businessMaterialKindLabel": "固定素材",
                            "evidenceSegmentId": selectable_payload["segments"][0]["evidenceSegmentId"],
                            "evidenceSegmentTitle": selectable_payload["segments"][0]["evidenceSegmentTitle"],
                            "evidenceSourcePages": selectable_payload["segments"][0]["evidenceSourcePages"],
                            "evidenceSummary": selectable_payload["segments"][0]["evidenceSummary"],
                        }
                    ]
                },
            )
        self.assertEqual(select_response.status_code, 200)
        selected_task = select_response.json()["task"]
        self.assertEqual(selected_task["status"], "ready")
        self.assertEqual(selected_task["handlingMode"], "fixed_material")
        self.assertEqual(selected_task["selectedMaterialRefs"][0]["materialId"], "RAW-9001")
        self.assertEqual(selected_task["selectedMaterialRefs"][0]["businessMaterialKind"], "fixed")
        self.assertEqual(selected_task["selectedEvidenceSegments"][0]["segmentId"], selectable_payload["segments"][0]["evidenceSegmentId"])
        self.assertEqual(select_response.json()["artifact"]["sourceMode"], "selected_from_business_material_library")
        self.assertEqual(select_response.json()["artifact"]["businessMaterialKind"], "fixed")
        self.assertEqual(select_response.json()["artifact"]["evidenceSegmentId"], selectable_payload["segments"][0]["evidenceSegmentId"])
        self.assertIn("business-workspace/gaps/selected-materials", select_response.json()["artifact"]["filePath"])
        self.assertFalse((technical_workspace_dir(project_id) / "s4_gap_workdir" / "selected_material").exists())

        facts_response = self.client.post(f"/api/business/projects/{project_id}/business-gaps/facts/build")
        self.assertEqual(facts_response.status_code, 200)
        facts = facts_response.json()
        labels = {field["label"]: field for field in facts["fields"]}
        self.assertEqual(len(facts["fields"]), 15)
        self.assertEqual(
            list(labels),
            [
                "招标项目名称",
                "招标编号",
                "招标人",
                "招标项目单位",
                "招标代理机构",
                "风机型号",
                "投标项目标段名称",
                "投标人",
                "投标人地址",
                "投标人电话",
                "法定代表人姓名/性别/年龄/职务",
                "委托人姓名/身份证",
                "营业执照信息注册资本/信用代码/类型（可选）",
                "存款账户号码/银行/编号（不确定）",
                "日期",
            ],
        )
        self.assertEqual(labels["招标项目名称"]["value"], "商务S3项目")
        self.assertEqual(labels["招标编号"]["value"], "BIZ-2026-001")
        self.assertIn("投标人", labels)
        self.assertNotIn("投标报价", labels)
        self.assertNotIn("币种", labels)
        self.assertNotIn("商务偏差说明", labels)
        confirm_facts_response = self.client.put(
            f"/api/business/projects/{project_id}/business-gaps/facts",
            json={"fields": facts["fields"], "confirm": True, "operator": "测试用户"},
        )
        self.assertEqual(confirm_facts_response.status_code, 200)
        self.assertEqual(confirm_facts_response.json()["status"], "confirmed")
        stored_project = store._require(project_id)
        self.assertEqual(stored_project["business_gap_state"]["projectFactTable"]["status"], "confirmed")
        self.assertEqual(stored_project["gap_state"]["projectFactTable"], {})

    def test_business_gap_api_consumes_published_s1_handoff_fields(self) -> None:
        self._setup_app_test()
        from app.services.bid_runtime_state import now_iso
        from app.services.store import store
        from app.services.workspace_artifacts import business_workspace_dir

        project = store.create_project({"name": "旧项目名不应使用", "customerName": "旧招标人", "bidType": "商务标"})
        project_id = project["id"]
        business_workspace = business_workspace_dir(project_id)
        parse_dir = business_workspace / "parse"
        parse_dir.mkdir(parents=True, exist_ok=True)
        structured_path = parse_dir / "s1_structured_result.json"
        structured_path.write_text(
            json.dumps(
                {
                    "schemaVersion": "bid-business-tender-structured-v1",
                    "items": [],
                    "structured": {
                        "schemaVersion": "bid-business-tender-structured-v1",
                        "fieldGroups": {
                            "projectBasics": [
                                {"fieldKey": "projectName", "title": "项目名称", "value": "PWF交接项目", "confidence": 0.96},
                                {"fieldKey": "tenderNo", "title": "招标编号", "value": "PWF-2026-001", "confidence": 0.95},
                                {
                                    "fieldKey": "tenderer",
                                    "title": "招标人",
                                    "value": "PWF能源集团有限公司",
                                    "confidence": 0.94,
                                },
                            ]
                        },
                        "appendices": [],
                        "commitmentLetters": [],
                        "projectFactFields": [
                            {
                                "fieldKey": "tenderAgency",
                                "label": "招标代理机构",
                                "value": "PWF招标代理有限公司",
                                "confidence": 0.9,
                            }
                        ],
                    },
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        _update_parse_result_for_tests(
            store,
            project_id,
            {
                "status": "completed",
                "structured": {
                    "projectFactFields": [
                        {"fieldKey": "projectName", "label": "项目名称", "value": "旧解析项目"},
                        {"fieldKey": "tenderNo", "label": "招标编号", "value": "OLD-001"},
                    ]
                },
            },
            parse_storage={
                "projectDir": str(business_workspace),
                "parseDir": str(parse_dir),
                "structuredResultPath": str(parse_dir / "legacy_should_not_be_used.json"),
            },
        )
        record = store._require(project_id)
        record["stageArtifacts"] = {
            "s1": {
                "schemaVersion": "business-s1-handoff-v1",
                "status": "published",
                "version": 3,
                "projectId": project_id,
                "bidType": "商务标",
                "parseProfile": "business",
                "publishedAt": "2026-06-05T00:00:00+08:00",
                "paths": {
                    "workspaceRoot": str(business_workspace),
                    "parseDir": str(parse_dir),
                    "structuredResultPath": str(structured_path),
                    "combinedTextPath": str(parse_dir / "combined.txt"),
                    "businessSectionTreePath": str(parse_dir / "business_section_tree.json"),
                    "skillManifestPath": str(parse_dir / "s1_parse_manifest.json"),
                    "manifestPath": str(parse_dir / "manifest.json"),
                    "appendicesDir": str(business_workspace / "appendices"),
                    "commitmentLettersDir": str(business_workspace / "commitment-letters"),
                },
            }
        }
        store._persist_project(record)
        _save_generated_outline_for_tests(
            store,
            project_id=project_id,
            nodes=[{"id": "OL-1", "title": "投标函", "children": []}],
            generated_at=now_iso(),
            summary="商务目录已生成。",
        )
        _confirm_outline_for_tests(store, project_id)

        async def empty_raw_files(**kwargs):
            return {"items": [], "total": 0}

        with patch(
            "app.services.business_gap_planning.OpencodeClient.run_bid_business_gap_planner_with_trace",
            side_effect=RuntimeError("offline test fallback"),
        ), patch(
            "app.services.business_gap_planning.business_material_store.raw_files",
            side_effect=empty_raw_files,
        ):
            response = self.client.post(f"/api/business/projects/{project_id}/business-gaps/run")
        self.assertEqual(response.status_code, 200)
        plan = response.json()["plan"]
        self.assertEqual(plan["s1Consumption"]["source"], "stageArtifacts.s1")
        self.assertEqual(plan["s1Consumption"]["handoff"]["version"], 3)
        manifest_path = business_workspace / "gaps" / "business_gap_input.json"
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        self.assertEqual(manifest["s1Consumption"]["source"], "stageArtifacts.s1")
        self.assertEqual(manifest["s1Handoff"]["version"], 3)
        parse_snapshot = json.loads((business_workspace / "gaps" / "parse_result.json").read_text(encoding="utf-8"))
        self.assertEqual(parse_snapshot["structured"]["fieldGroups"]["projectBasics"][0]["value"], "PWF交接项目")

        facts_response = self.client.post(f"/api/business/projects/{project_id}/business-gaps/facts/build")
        self.assertEqual(facts_response.status_code, 200)
        labels = {field["label"]: field for field in facts_response.json()["fields"]}
        self.assertEqual(labels["招标项目名称"]["value"], "PWF交接项目")
        self.assertEqual(labels["招标编号"]["value"], "PWF-2026-001")
        self.assertEqual(labels["招标人"]["value"], "PWF能源集团有限公司")
        self.assertEqual(labels["招标代理机构"]["value"], "PWF招标代理有限公司")
        self.assertNotEqual(labels["招标项目名称"]["value"], "旧解析项目")

    def test_business_gap_rejects_unpublished_s1_handoff(self) -> None:
        self._setup_app_test()
        from app.services.bid_runtime_state import now_iso
        from app.services.store import store
        from app.services.workspace_artifacts import business_workspace_dir

        project = store.create_project({"name": "未发布交接测试", "customerName": "测试业主", "bidType": "商务标"})
        project_id = project["id"]
        business_workspace = business_workspace_dir(project_id)
        parse_dir = business_workspace / "parse"
        parse_dir.mkdir(parents=True, exist_ok=True)
        structured_path = parse_dir / "s1_structured_result.json"
        structured_path.write_text(json.dumps({"structured": {"projectFactFields": []}}, ensure_ascii=False), encoding="utf-8")
        record = store._require(project_id)
        record["stageArtifacts"] = {
            "s1": {
                "schemaVersion": "business-s1-handoff-v1",
                "status": "readyForReview",
                "version": 1,
                "projectId": project_id,
                "bidType": "商务标",
                "parseProfile": "business",
                "paths": {"structuredResultPath": str(structured_path)},
            }
        }
        store._persist_project(record)
        _save_generated_outline_for_tests(
            store,
            project_id=project_id,
            nodes=[{"id": "OL-1", "title": "投标函", "children": []}],
            generated_at=now_iso(),
            summary="商务目录已生成。",
        )
        _confirm_outline_for_tests(store, project_id)

        response = self.client.post(f"/api/business/projects/{project_id}/business-gaps/run")
        self.assertEqual(response.status_code, 400)
        self.assertIn("尚未发布", response.json()["detail"])

    def test_table_fill_lookup_does_not_pollute_specific_labels(self) -> None:
        import importlib.util

        backend_root = Path(__file__).resolve().parents[1]
        script_path = backend_root / "opencode" / "skills" / "bid-business-table-fill" / "scripts" / "run_from_manifest.py"
        spec = importlib.util.spec_from_file_location("business_table_fill_runner", script_path)
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)
        facts = {"投标人": "上海电气风电集团股份有限公司", "招标项目名称": "某风电项目"}
        self.assertEqual(module.lookup("投标人", facts), "上海电气风电集团股份有限公司")
        self.assertEqual(module.lookup("投标人（盖章）", facts), "上海电气风电集团股份有限公司")
        self.assertEqual(module.lookup("投标人名称", facts), "上海电气风电集团股份有限公司")
        self.assertEqual(module.lookup("投标人地址", facts), "")
        self.assertEqual(module.lookup("投标人电话", facts), "")
        self.assertEqual(module.lookup("项目名称", facts), "某风电项目")

    def test_business_gap_table_fill_creates_artifact_from_target_and_sources(self) -> None:
        self._setup_app_test()
        from app.services.business_gap_fact_table import PROJECT_FACT_TABLE_SCHEMA_VERSION
        from app.services.bid_runtime_state import now_iso
        from app.services.store import store
        from app.services.workspace_artifacts import business_workspace_dir, technical_workspace_dir

        project = store.create_project({"name": "商务AI填表项目", "customerName": "华能集团", "bidType": "商务标"})
        project_id = project["id"]
        business_workspace = business_workspace_dir(project_id)
        target_path = business_workspace / "gaps" / "templates" / "投标函模板.docx"
        target_path.parent.mkdir(parents=True, exist_ok=True)
        target_doc = Document()
        table = target_doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "字段"
        table.rows[0].cells[1].text = "投标人响应"
        table.rows[1].cells[0].text = "项目名称"
        table.rows[1].cells[1].text = ""
        target_doc.save(target_path)
        source_path = business_workspace / "source.xlsx"
        source_path.write_bytes(b"fake-xlsx")

        stored_project = store._require(project_id)
        stored_project["business_gap_state"].update(
            {
                "recognitionStatus": "completed",
                "recognizedAt": now_iso(),
                "plan": {
                    "schemaVersion": "bid-business-gap-plan-v1",
                    "tocRefs": [{"nodeId": "TOC-1", "title": "投标函", "taskIds": ["BT-001"], "status": "partial"}],
                    "tasks": [
                        {
                            "id": "BT-001",
                            "title": "投标函",
                            "taskType": "table",
                            "decision": "fill_required",
                            "status": "needs_input",
                            "moduleKey": "structured_response_tables",
                            "assemblyMode": "template_fill_docx",
                            "materialUsage": "fill_table",
                            "templateCandidates": [
                                {
                                    "templateId": "TPL-001",
                                    "templateName": "投标函模板.docx",
                                    "fileName": target_path.name,
                                    "filePath": str(target_path),
                                    "assemblyMode": "template_fill_docx",
                                    "materialUsage": "fill_table",
                                    "sourceMode": "project_uploaded_bid_template",
                                }
                            ],
                            "resolvedArtifacts": [],
                            "riskFlags": ["missing_material"],
                        }
                    ],
                    "summary": {},
                },
                "integrity": {},
                "projectFactTable": {
                    "schemaVersion": PROJECT_FACT_TABLE_SCHEMA_VERSION,
                    "status": "confirmed",
                    "fields": [{"label": "项目名称", "value": "商务AI填表项目", "status": "confirmed"}],
                },
            }
        )
        store._persist_project(stored_project)

        async def fake_download_payload(material_id: str) -> tuple[dict[str, str], str]:
            self.assertEqual(material_id, "RAW-TABLE-001")
            return {
                "fileId": material_id,
                "fileName": "报价数据.xlsx",
                "bucket": "mock-bucket",
                "key": "mock-key",
                "mimeType": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            }, "raw"

        def fake_download_file(bucket: str, key: str, output_path: Path) -> Path:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_bytes(source_path.read_bytes())
            return output_path

        def fake_runner(manifest_path: Path) -> dict[str, object]:
            manifest = json.loads(Path(manifest_path).read_text(encoding="utf-8"))
            self.assertEqual(manifest["target"]["templateId"], "TPL-001")
            self.assertEqual(manifest["sourceMaterials"][0]["materialId"], "RAW-TABLE-001")
            self.assertEqual(manifest["s1Consumption"]["source"], "legacy_parse_result")
            output_path = Path(manifest["outputFile"])
            output_path.parent.mkdir(parents=True, exist_ok=True)
            Document(str(target_path)).save(output_path)
            return {
                "schemaVersion": "bid-business-table-fill-v1",
                "outputFile": str(output_path),
                "fillReport": {"filledFieldCount": 1},
                "unfilledFields": [],
                "evidenceRefs": [{"materialId": "RAW-TABLE-001", "factCount": 1}],
            }

        with patch("app.services.business_gap_table_fill.downloadable_business_fill_source_payload", side_effect=fake_download_payload), patch(
            "app.services.business_gap_table_fill.minio_client.download_file",
            side_effect=fake_download_file,
        ), patch(
            "app.services.business_gap_service.run_business_table_fill_skill",
            side_effect=fake_runner,
        ):
            response = self.client.post(
                f"/api/business/projects/{project_id}/business-gaps/tasks/BT-001/table-fill",
                json={
                    "target": {"templateId": "TPL-001"},
                    "sourceMaterials": [{"materialId": "RAW-TABLE-001", "materialName": "报价数据.xlsx"}],
                    "operator": "测试用户",
                },
            )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["artifact"]["sourceMode"], "generated_by_business_table_fill")
        self.assertEqual(payload["artifact"]["assemblyMode"], "table_fill_from_material")
        self.assertEqual(payload["task"]["status"], "ready")
        self.assertEqual(payload["task"]["decision"], "ready")
        self.assertEqual(payload["task"]["handlingMode"], "ai_table_fill")
        self.assertEqual(payload["task"]["resolvedArtifacts"][0]["operator"], "测试用户")
        self.assertTrue(Path(payload["artifact"]["filePath"]).exists())
        self.assertFalse((technical_workspace_dir(project_id) / "s4_gap_workdir" / "table-fill").exists())

    def test_business_gap_table_fill_allows_project_fact_table_only(self) -> None:
        self._setup_app_test()
        from app.services.business_gap_fact_table import PROJECT_FACT_TABLE_SCHEMA_VERSION
        from app.services.bid_runtime_state import now_iso
        from app.services.store import store
        from app.services.workspace_artifacts import business_workspace_dir

        project = store.create_project({"name": "商务事实表填表项目", "customerName": "华能集团", "bidType": "商务标"})
        project_id = project["id"]
        business_workspace = business_workspace_dir(project_id)
        target_path = business_workspace / "gaps" / "templates" / "投标函模板.docx"
        target_path.parent.mkdir(parents=True, exist_ok=True)
        target_doc = Document()
        table = target_doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "字段"
        table.rows[0].cells[1].text = "投标人响应"
        table.rows[1].cells[0].text = "项目名称"
        table.rows[1].cells[1].text = ""
        target_doc.save(target_path)

        stored_project = store._require(project_id)
        stored_project["business_gap_state"].update(
            {
                "recognitionStatus": "completed",
                "recognizedAt": now_iso(),
                "plan": {
                    "schemaVersion": "bid-business-gap-plan-v1",
                    "tocRefs": [{"nodeId": "TOC-1", "title": "投标函", "taskIds": ["BT-001"], "status": "partial"}],
                    "tasks": [
                        {
                            "id": "BT-001",
                            "title": "投标函",
                            "taskType": "table",
                            "decision": "fill_required",
                            "status": "needs_input",
                            "moduleKey": "structured_response_tables",
                            "assemblyMode": "template_fill_docx",
                            "materialUsage": "fill_table",
                            "templateCandidates": [
                                {
                                    "templateId": "TPL-001",
                                    "templateName": "投标函模板.docx",
                                    "fileName": target_path.name,
                                    "filePath": str(target_path),
                                    "assemblyMode": "template_fill_docx",
                                    "materialUsage": "fill_table",
                                    "sourceMode": "parsed_from_tender_attachment_template",
                                }
                            ],
                            "resolvedArtifacts": [],
                            "riskFlags": ["missing_material"],
                        }
                    ],
                    "summary": {},
                },
                "integrity": {},
                "projectFactTable": {
                    "schemaVersion": PROJECT_FACT_TABLE_SCHEMA_VERSION,
                    "status": "confirmed",
                    "fields": [{"label": "项目名称", "value": "商务事实表填表项目", "status": "confirmed"}],
                },
            }
        )
        store._persist_project(stored_project)

        def fake_runner(manifest_path: Path) -> dict[str, object]:
            manifest = json.loads(Path(manifest_path).read_text(encoding="utf-8"))
            self.assertEqual(manifest["sourceMaterials"], [])
            self.assertEqual(manifest["projectFactTable"]["status"], "confirmed")
            output_path = Path(manifest["outputFile"])
            output_path.parent.mkdir(parents=True, exist_ok=True)
            Document(str(target_path)).save(output_path)
            return {
                "schemaVersion": "bid-business-table-fill-v1",
                "outputFile": str(output_path),
                "fillReport": {"filledFieldCount": 1},
                "unfilledFields": [],
                "evidenceRefs": [],
            }

        with patch("app.services.business_gap_service.run_business_table_fill_skill", side_effect=fake_runner):
            response = self.client.post(
                f"/api/business/projects/{project_id}/business-gaps/tasks/BT-001/table-fill",
                json={"target": {"templateId": "TPL-001"}, "operator": "测试用户"},
            )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["artifact"]["sourceMaterials"], [])
        self.assertEqual(payload["projectFactTable"]["status"], "confirmed")
        self.assertEqual(payload["task"]["handlingMode"], "ai_table_fill")

    def test_business_table_fill_runner_fallback_docx_has_no_internal_explanation(self) -> None:
        backend_root = Path(__file__).resolve().parents[1]
        script_path = backend_root / "opencode" / "skills" / "bid-business-table-fill" / "scripts" / "run_from_manifest.py"
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            target_path = root / "纯文本目标.txt"
            target_path.write_text("fallback", encoding="utf-8")
            output_path = root / "AI填表.docx"
            manifest_path = root / "business_table_fill_input.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-table-fill-v1",
                        "projectId": "PRJ-TABLE-FILL",
                        "projectName": "商务填表项目",
                        "task": {"id": "BT-001", "title": "商务评分索引表", "requirement": ""},
                        "target": {"fileName": target_path.name, "filePath": str(target_path)},
                        "sourceMaterials": [],
                        "projectFactTable": {
                            "schemaVersion": "bid-project-fact-table-v1",
                            "status": "confirmed",
                            "fields": [{"label": "项目名称", "value": "商务填表项目"}],
                        },
                        "facts": {"项目名称": "商务填表项目"},
                        "outputFile": str(output_path),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            subprocess.run(
                [sys.executable, str(script_path), "--manifest", str(manifest_path), "--response", "summary"],
                check=True,
                capture_output=True,
                text=True,
            )
            output_doc = Document(str(output_path))
            text_parts = [paragraph.text for paragraph in output_doc.paragraphs]
            for table in output_doc.tables:
                for row in table.rows:
                    text_parts.extend(cell.text for cell in row.cells)
            text = "\n".join(text_parts)

        self.assertIn("商务评分索引表", text)
        self.assertNotIn("当前目标文件暂不支持原格式写入", text)
        self.assertNotIn("已根据项目事实表和数据来源生成", text)
        self.assertNotIn("AI填表结果", text)
        self.assertNotIn("AI填表内容", text)
        self.assertNotIn("未从项目事实表或素材库提取到可填写内容", text)

    def test_business_gap_task_ignore_sets_handling_mode(self) -> None:
        self._setup_app_test()
        from app.services.bid_runtime_state import now_iso
        from app.services.store import store

        project = store.create_project({"name": "商务忽略任务项目", "customerName": "华能集团", "bidType": "商务标"})
        project_id = project["id"]
        stored_project = store._require(project_id)
        stored_project["business_gap_state"].update(
            {
                "recognitionStatus": "completed",
                "recognizedAt": now_iso(),
                "plan": {
                    "schemaVersion": "bid-business-gap-plan-v1",
                    "tocRefs": [{"nodeId": "TOC-1", "title": "承诺函", "taskIds": ["BT-IGNORE"], "status": "partial"}],
                    "tasks": [
                        {
                            "id": "BT-IGNORE",
                            "title": "可忽略事项",
                            "taskType": "attachment",
                            "decision": "review_required",
                            "status": "review_required",
                            "moduleKey": "commitments_and_notes",
                            "candidateMaterials": [],
                            "resolvedArtifacts": [],
                        }
                    ],
                    "summary": {},
                },
                "integrity": {},
            }
        )
        store._persist_project(stored_project)

        response = self.client.patch(
            f"/api/business/projects/{project_id}/business-gaps/tasks/BT-IGNORE",
            json={"status": "ignored", "notes": "无需响应"},
        )
        self.assertEqual(response.status_code, 200)
        task = response.json()["task"]
        self.assertEqual(task["status"], "ignored")
        self.assertEqual(task["handlingMode"], "ignored")
        self.assertEqual(response.json()["plan"]["summary"]["handlingModeCounts"]["ignored"], 1)


if __name__ == "__main__":
    unittest.main()
