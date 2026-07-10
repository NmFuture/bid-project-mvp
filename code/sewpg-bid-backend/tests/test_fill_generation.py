from __future__ import annotations

import asyncio
import copy
import json
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from fastapi.testclient import TestClient
from starlette.datastructures import URL

from app.main import app
from app.core.config import settings
from app.services.bid_fill_generation_state import save_fill_generation_result_state, start_fill_generation_state
from app.services.bid_outline_state import confirm_outline_state, save_generated_outline_state
from app.services.bid_parse_state import complete_parse_state
from app.services.bid_type import TECHNICAL_BID_TYPE
from app.services.bid_runtime_state import now_iso
from app.services.store import store
from app.services.technical_gap_repository import persist_technical_gap_project, require_technical_gap_project_for_update
from app.services.technical_gap_review import confirm_technical_review, prepare_technical_review_document
from app.services.technical_gap_service import technical_gap_service
from app.services.technical_gap_state import ensure_technical_gap_state
from app.services.workspace_artifacts import technical_workspace_dir, technical_workspace_stage_dir


class _DummyRequest:
    base_url = URL("http://testserver/")
    url = URL("http://testserver/")


def _confirm_technical_review_for_tests(project_id: str) -> None:
    project = require_technical_gap_project_for_update(project_id)
    gap_state = ensure_technical_gap_state(project)
    prepare_technical_review_document(project, gap_state)
    confirm_technical_review(project, gap_state)
    persist_technical_gap_project(project)


def _fill_state_for_tests(project_id: str) -> dict:
    return copy.deepcopy(store.get_project_runtime_state(project_id)["fill_state"])


def _start_fill_generation_for_tests(project_id: str) -> dict:
    project = store.require_project_for_update(project_id)
    payload = start_fill_generation_state(project)
    store.persist_project_state(project)
    return payload


def _save_fill_generation_result_for_tests(
    project_id: str,
    *,
    summary: str,
    sections: list[dict],
    content: str,
    filled_at: str,
    run_duration_sec: int,
    file_size_bytes: int,
    opencode_output: dict | None = None,
    file_name: str | None = None,
    coverage: dict | None = None,
    assembly: dict | None = None,
) -> dict:
    project = store.require_project_for_update(project_id)
    payload = save_fill_generation_result_state(
        project,
        project_id=project_id,
        summary=summary,
        sections=sections,
        content=content,
        filled_at=filled_at,
        run_duration_sec=run_duration_sec,
        file_size_bytes=file_size_bytes,
        opencode_output=opencode_output,
        file_name=file_name,
        coverage=coverage,
        assembly=assembly,
    )
    store.persist_project_state(project)
    return payload


def _complete_parse_for_tests(
    project_id: str,
    tender_files: list[dict],
    template_files: list[dict],
    *,
    summary: dict | None = None,
    parse_storage: dict | None = None,
) -> dict:
    project = store.require_project_for_update(project_id)
    payload = complete_parse_state(
        project,
        tender_files,
        template_files,
        summary=summary,
        parse_storage=parse_storage,
    )
    store.persist_project_state(project)
    return payload


def _save_generated_outline_for_tests(
    project_id: str,
    *,
    nodes: list[dict],
    generated_at: str,
    summary: str,
) -> dict:
    project = store.require_project_for_update(project_id)
    payload = save_generated_outline_state(
        project,
        nodes=nodes,
        generated_at=generated_at,
        summary=summary,
    )
    store.persist_project_state(project)
    return payload


def _confirm_outline_for_tests(project_id: str) -> dict:
    project = store.require_project_for_update(project_id)
    payload = confirm_outline_state(project)
    store.persist_project_state(project)
    return payload


class FillGenerationTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        base = Path(self.temp_dir.name)
        settings.uploads_dir = base / "uploads"
        settings.documents_dir = base / "documents"
        settings.parsed_dir = base / "parsed"
        settings.ensure_dirs()

        store.reset_for_tests()
        self.client = TestClient(app, base_url="http://127.0.0.1:8000")
        self.gap_planner_patcher = patch(
            "app.services.technical_gap_planner.OpencodeClient.run_bid_tech_gap_planner_with_trace",
            side_effect=RuntimeError("offline test fallback"),
        )
        self.gap_planner_patcher.start()
        login = self.client.post("/api/auth/login", json={"email": "admin@sewpg.com", "password": "123456"})
        self.assertEqual(login.status_code, 200)
        self.headers = {"Authorization": f"Bearer {login.json()['token']}"}

    def tearDown(self) -> None:
        self.gap_planner_patcher.stop()
        self.client.close()
        self.temp_dir.cleanup()

    def _prepare_project_after_outline(self) -> str:
        response = self.client.post(
            "/api/technical/projects",
            json={
                "name": "S4生成标书项目",
                "customerName": "测试业主",
            },
        )
        response.raise_for_status()
        project_id = response.json()["id"]

        project_dir = technical_workspace_dir(project_id) / "parse"
        project_dir.mkdir(parents=True, exist_ok=True)
        combined_text_path = project_dir / "combined.txt"
        combined_text_path.write_text(
            "\n".join(
                [
                    "# 文件：招标文件.docx",
                    "",
                    "第一章 项目概况",
                    "第二章 技术方案",
                    "第三章 实施与保障",
                    "项目要求投标文件包含总体方案、关键参数、实施组织与风险控制。",
                ]
            ),
            encoding="utf-8",
        )

        tender_path = settings.uploads_dir / project_id / "tender" / "招标文件.docx"
        tender_path.parent.mkdir(parents=True, exist_ok=True)
        tender_path.write_text("dummy", encoding="utf-8")

        _complete_parse_for_tests(
            project_id,
            tender_files=[
                {
                    "id": "TEN-1",
                    "name": "招标文件.docx",
                    "path": str(tender_path),
                    "size_label": "1.0 MB",
                }
            ],
            template_files=[],
            summary={
                "fileCount": 1,
                "extractedCount": 0,
                "textLength": 128,
                "textPreview": "",
                "warnings": [],
            },
            parse_storage={
                "projectDir": str(project_dir),
                "combinedTextPath": str(combined_text_path),
                "manifestPath": "",
                "documents": [],
            },
        )

        _save_generated_outline_for_tests(
            project_id=project_id,
            nodes=[
                {
                    "id": "OL-1",
                    "title": "项目概况",
                    "children": [
                        {"id": "OL-1-1", "title": "项目背景", "children": []},
                    ],
                },
                {
                    "id": "OL-2",
                    "title": "技术方案",
                    "children": [
                        {"id": "OL-2-1", "title": "总体方案", "children": []},
                        {"id": "OL-2-2", "title": "关键参数响应", "children": []},
                    ],
                },
                {
                    "id": "OL-3",
                    "title": "实施与保障",
                    "children": [
                        {"id": "OL-3-1", "title": "实施组织", "children": []},
                        {"id": "OL-3-2", "title": "风险控制", "children": []},
                    ],
                },
            ],
            generated_at=now_iso(),
            summary="目录已生成。",
        )
        _confirm_outline_for_tests(project_id)
        return project_id

    def _prepare_project_for_s7(self) -> str:
        project_id = self._prepare_project_after_outline()
        technical_gap_service.run_detection(project_id)
        for item in asyncio.run(technical_gap_service.gaps(project_id, _DummyRequest()))["items"]:
            asyncio.run(
                technical_gap_service.update_gap(
                    project_id,
                    item["id"],
                    {"status": "skipped", "reason": "测试中人工确认忽略"},
                )
            )
        asyncio.run(technical_gap_service.recheck(project_id))
        asyncio.run(technical_gap_service.submit_review(project_id))
        _confirm_technical_review_for_tests(project_id)
        return project_id

    def test_run_fill_generation_returns_running_state_immediately(self) -> None:
        project_id = self._prepare_project_for_s7()

        with patch("app.services.bid_generation_flow._schedule_fill_generation_job"):
            response = self.client.post(f"/api/technical/projects/{project_id}/fill-generation/run", headers=self.headers)

        self.assertEqual(response.status_code, 202)
        payload = response.json()
        self.assertEqual(payload["status"], "running")
        self.assertEqual(payload["percentage"], 5)
        self.assertEqual(payload["tasks"][0]["status"], "running")
        self.assertEqual(payload["tasks"][1]["status"], "pending")
        self.assertEqual(payload["events"][0]["step"], "bootstrap")
        self.assertEqual(payload["opencodeOutput"]["status"], "idle")
        audit = self.client.get("/api/technical/audit", headers=self.headers)
        self.assertEqual(audit.status_code, 200)
        generation_logs = [item for item in audit.json()["items"] if item["action"] == "开始生成标书"]
        self.assertGreaterEqual(len(generation_logs), 1)
        self.assertEqual(generation_logs[0]["actionType"], "generate")
        self.assertEqual(generation_logs[0]["module"], "generation")

    def test_background_job_updates_running_state_then_writes_real_docx(self) -> None:
        from app.services.bid_generation_flow import _handle_fill_progress, _run_fill_generation_job

        project_id = self._prepare_project_for_s7()
        _start_fill_generation_for_tests(project_id)

        _handle_fill_progress(
            project_id,
            "inputs_ready",
            {"wikiCardCount": 3, "exportedMaterialCount": 2},
            bid_type=TECHNICAL_BID_TYPE,
        )
        running_state = _fill_state_for_tests(project_id)
        self.assertEqual(running_state["status"], "running")
        self.assertEqual(running_state["percentage"], 30)
        self.assertEqual(running_state["tasks"][1]["status"], "running")
        self.assertEqual(running_state["events"][-1]["step"], "inputs_ready")

        def fake_assemble(fake_project_id, data, progress_callback=None):
            from app.services.onlyoffice_documents import document_path, write_document

            if progress_callback:
                progress_callback(
                    "calling_assembler",
                    {
                        "manifestPath": "/tmp/s7_assembly_input.json",
                        "workDir": "/tmp/s7_assembly_workdir",
                    },
                )
                progress_callback(
                    "assembling_result",
                    {
                        "sectionCount": 3,
                        "usedMaterialCount": 2,
                        "unassembledMaterialCount": 1,
                    },
                )
            sections = [
                    {
                        "nodeId": "OL-1",
                        "title": "项目概况",
                        "generationMode": "generated",
                        "content": "已拼装：项目背景",
                        "riskFlags": [],
                    },
                    {
                        "nodeId": "OL-2",
                        "title": "技术方案",
                        "generationMode": "generated_with_placeholder",
                        "content": "已拼装：总体方案\n【待填写：关键参数实测值】",
                        "riskFlags": ["FACT_REQUIRED"],
                    },
                    {
                        "nodeId": "OL-3",
                        "title": "实施与保障",
                        "generationMode": "generated",
                        "content": "已拼装：实施组织",
                        "riskFlags": [],
                    },
                ]
            content = "# S4生成标书项目 正文\n\n## 项目概况\n项目背景\n\n## 关键参数响应\n【待填写：关键参数实测值】"
            target = document_path(fake_project_id)
            write_document(target, "S4生成标书项目_正文.docx", content)
            return _save_fill_generation_result_for_tests(
                project_id=fake_project_id,
                summary="技术标正文拼装完成。",
                sections=sections,
                content=content,
                filled_at=now_iso(),
                run_duration_sec=3,
                file_size_bytes=target.stat().st_size,
                file_name="S4生成标书项目_正文.docx",
                opencode_output={
                    "status": "received",
                    "sessionId": "/tmp/s7_assembly_input.json",
                    "providerId": "local-skill",
                    "modelId": "bid-tech-assembler",
                    "receivedAt": "2026-04-20T00:00:00Z",
                    "parts": [
                        {"type": "text", "text": "{\"summary\":\"技术标正文拼装完成。\"}"},
                    ],
                },
                coverage={
                    "percentage": 67,
                    "fullCover": 2,
                    "partialCover": 0,
                    "noCover": 1,
                    "tree": [],
                    "partialItems": [],
                    "noCoverItems": [{"id": "RAW-0003", "title": "未拼素材", "nodeTitle": "素材未出现在 S2 目录 JSON 或拼装计划中"}],
                },
            )

        with patch(
            "app.services.technical_draft_generation.assemble_tech_bid_for_project_with_progress",
            side_effect=fake_assemble,
        ):
            _run_fill_generation_job(project_id, {}, bid_type=TECHNICAL_BID_TYPE)

        payload = _fill_state_for_tests(project_id)
        self.assertEqual(payload["status"], "completed")
        self.assertEqual(payload["percentage"], 100)
        self.assertEqual(payload["output"]["fileType"], "docx")
        self.assertEqual(len(payload["sections"]), 3)
        self.assertEqual(payload["opencodeOutput"]["status"], "received")
        self.assertEqual(payload["opencodeOutput"]["parts"][0]["type"], "text")
        audit = self.client.get("/api/technical/audit", headers=self.headers)
        self.assertEqual(audit.status_code, 200)
        actions = [item["action"] for item in audit.json()["items"]]
        self.assertIn("生成标书完成", actions)

        document_response = self.client.get(f"/api/technical/projects/{project_id}/document/file")
        self.assertEqual(document_response.status_code, 200)
        self.assertEqual(document_response.content[:2], b"PK")

        local_path = settings.documents_dir / f"{project_id}.docx"
        doc = Document(local_path)
        full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        self.assertIn("项目概况", full_text)
        self.assertIn("关键参数响应", full_text)
        self.assertIn("【待填写：关键参数实测值】", full_text)

    def test_generation_failure_before_inputs_marks_prepare_task_failed(self) -> None:
        from app.services.bid_generation_flow import _run_fill_generation_job

        project_id = self._prepare_project_after_outline()
        _start_fill_generation_for_tests(project_id)

        with patch(
            "app.services.bid_generation_flow.generate_technical_draft_for_project_with_progress",
            side_effect=RuntimeError("Remote end closed connection without response"),
        ):
            _run_fill_generation_job(project_id, {}, bid_type=TECHNICAL_BID_TYPE)

        payload = _fill_state_for_tests(project_id)
        self.assertEqual(payload["status"], "failed")
        self.assertEqual(payload["tasks"][0]["status"], "failed")
        self.assertEqual(payload["tasks"][1]["status"], "pending")
        self.assertEqual(payload["tasks"][2]["status"], "pending")
        self.assertIn("Remote end closed", payload["summary"])

    def test_get_coverage_returns_tree_after_fill_generation(self) -> None:
        from app.services.bid_generation_flow import _run_fill_generation_job

        project_id = self._prepare_project_for_s7()
        _start_fill_generation_for_tests(project_id)

        def fake_assemble(fake_project_id, data, progress_callback=None):
            from app.services.onlyoffice_documents import document_path, write_document

            target = document_path(fake_project_id)
            write_document(target, "S4生成标书项目_正文.docx", "# 正文\n\n已拼装。")
            return _save_fill_generation_result_for_tests(
                project_id=fake_project_id,
                summary="技术标正文拼装完成。",
                sections=[],
                content="# 正文\n\n已拼装。",
                filled_at=now_iso(),
                run_duration_sec=2,
                file_size_bytes=target.stat().st_size,
                file_name="S4生成标书项目_正文.docx",
                coverage={
                    "percentage": 50,
                    "fullCover": 1,
                    "partialCover": 1,
                    "noCover": 1,
                    "tree": [
                        {
                            "id": "通用/风资源",
                            "title": "通用/风资源",
                            "coverage": 50,
                            "status": "partial",
                            "children": [
                                {"id": "RAW-0001", "title": "已拼素材", "coverage": 100, "status": "full", "children": []},
                                {"id": "RAW-0002", "title": "未拼素材", "coverage": 0, "status": "none", "children": []},
                            ],
                        }
                    ],
                    "partialItems": [{"id": "TOC-2", "title": "目录未匹配", "nodeTitle": "目录项未匹配素材"}],
                    "noCoverItems": [{"id": "RAW-0002", "title": "未拼素材", "nodeTitle": "素材未出现在 S2 目录 JSON 或拼装计划中"}],
                },
            )

        with patch("app.services.technical_draft_generation.assemble_tech_bid_for_project_with_progress", side_effect=fake_assemble):
            _run_fill_generation_job(project_id, {}, bid_type=TECHNICAL_BID_TYPE)

        coverage_response = self.client.get(f"/api/technical/projects/{project_id}/coverage")
        self.assertEqual(coverage_response.status_code, 200)
        coverage = coverage_response.json()
        self.assertIn("tree", coverage)
        self.assertIn("partialItems", coverage)
        self.assertIn("noCoverItems", coverage)
        self.assertGreater(len(coverage["tree"]), 0)

    def test_material_coverage_matches_recovered_original_paths(self) -> None:
        from app.services.tech_assembly import _build_material_coverage

        original_path = "/data/documents/PRJ-0005/technical-workspace/s4_gap_workdir/ai_fill/GAP-0058/投标机型总方案信息表_AI填写.docx"
        coverage = _build_material_coverage(
            [
                {
                    "toc_idx": 58,
                    "title": "投标机型总方案信息表",
                    "status": "MATCHED",
                    "paths": [original_path],
                }
            ],
            [
                {
                    "id": original_path,
                    "title": "投标机型总方案信息表",
                    "path": "投标资料库-定制/缺口处理/投标机型总方案信息表_AI填写.docx",
                    "originalPath": original_path,
                    "scope": "定制",
                    "category": "缺口处理",
                    "available": True,
                }
            ],
        )

        self.assertEqual(coverage["fullCover"], 1)
        self.assertEqual(coverage["noCover"], 0)
        self.assertEqual(coverage["percentage"], 100)
        self.assertEqual(coverage["tree"][0]["status"], "full")
        self.assertEqual(coverage["tree"][0]["children"][0]["status"], "full")

    def test_technical_stage_skips_s3_after_outline_confirmation(self) -> None:
        project_id = self._prepare_project_after_outline()

        project = store.get_project(project_id)
        self.assertEqual(project["currentStage"], 4)

        stages = store.get_stages(project_id)
        active = next(stage for stage in stages if stage["status"] == "active")
        skipped = next(stage for stage in stages if stage["id"] == 3)
        self.assertEqual(active["id"], 4)
        self.assertEqual(skipped["status"], "completed")
        self.assertTrue(skipped["isSkipped"])

    def test_s7_manifest_allows_missing_gap_plan_for_technical_bid(self) -> None:
        from app.services import tech_assembly

        project_id = self._prepare_project_after_outline()
        manifest_payloads = []

        def fake_prepare_wiki_dir(project, parse_storage, work_dir):
            wiki_dir = work_dir / "wiki"
            cards_dir = wiki_dir / "卡片"
            cards_dir.mkdir(parents=True, exist_ok=True)
            return wiki_dir

        def fake_export_material_library(wiki_dir, library_dir):
            library_dir.mkdir(parents=True, exist_ok=True)
            return library_dir, []

        def fake_run_assembler_manifest(manifest_path, progress_callback=None):
            manifest_payloads.append(json.loads(Path(manifest_path).read_text(encoding="utf-8")))
            output_file = Path(manifest_payloads[-1]["outputFile"])
            output_file.parent.mkdir(parents=True, exist_ok=True)
            doc = Document()
            doc.add_paragraph("技术方案")
            doc.save(output_file)
            plan_file = output_file.parent / "assembly_plan.json"
            plan_file.write_text(
                json.dumps(
                    [
                        {
                            "toc_idx": 1,
                            "level": 1,
                            "title": "技术方案",
                            "status": "STRUCTURAL",
                            "paths": [],
                        }
                    ],
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            report = output_file.parent / "assembly_report.md"
            review = output_file.parent / "needs_review.md"
            report.write_text("ok", encoding="utf-8")
            review.write_text("ok", encoding="utf-8")
            return {
                "schema_version": "bid-tech-assembly-v1",
                "outputFile": str(output_file),
                "planFile": str(plan_file),
                "assemblyReport": str(report),
                "needsReview": str(review),
                "summary": {"total": 1, "byStatus": {"STRUCTURAL": 1}, "usedPathCount": 0},
            }

        with patch.object(tech_assembly, "_prepare_wiki_dir", side_effect=fake_prepare_wiki_dir), \
            patch.object(tech_assembly, "_augment_wiki_with_material_cards", return_value=0), \
            patch.object(tech_assembly, "_export_material_library", side_effect=fake_export_material_library), \
            patch.object(tech_assembly, "_run_assembler_manifest", side_effect=fake_run_assembler_manifest):
            result = tech_assembly.assemble_tech_bid_for_project_with_progress(project_id)

        self.assertEqual(result["status"], "completed")
        self.assertEqual(len(manifest_payloads), 1)
        self.assertEqual(manifest_payloads[0]["gapPlanPath"], "")
        self.assertEqual(result["assembly"]["gapPlanPath"], "")
        self.assertEqual(result["assembly"]["formatClean"]["status"], "completed")
        self.assertTrue(Path(result["assembly"]["formatClean"]["outputFile"]).exists())
        self.assertTrue(Path(result["assembly"]["formatClean"]["reportFile"]).exists())

    def test_s7_gap_plan_recovers_s3_ai_fill_outputs_without_review_confirmation(self) -> None:
        from app.services import tech_assembly

        project_id = self._prepare_project_after_outline()
        project = store._require(project_id)
        project["gap_state"] = {
            "plan": {
                "schemaVersion": "bid-tech-gap-plan-v1",
                "status": "ready",
                "items": [
                    {
                        "id": "GAP-0058",
                        "number": "附表A.1",
                        "title": "投标机型总方案信息表",
                        "status": "needs_input",
                        "matchedMaterials": [
                            {
                                "id": "RAW-OLD",
                                "path": "技术标/原始待填模板/投标机型总方案信息表.docx",
                                "title": "原始待填模板",
                            }
                        ],
                        "resolvedArtifacts": [],
                    }
                ],
            },
            "integrity": {"status": "blocked", "blockingCount": 1},
        }
        store._persist_project(project)

        ai_fill_dir = technical_workspace_stage_dir(project_id, "s4_gap_workdir") / "ai_fill" / "GAP-0058"
        ai_fill_dir.mkdir(parents=True, exist_ok=True)
        filled_docx = ai_fill_dir / "投标机型总方案信息表_AI填写.docx"
        doc = Document()
        doc.add_paragraph("已填写表格")
        doc.save(filled_docx)
        (ai_fill_dir / "投标机型总方案信息表_AI填写.fill_report.json").write_text(
            json.dumps(
                {
                    "fillReport": {"title": "投标机型总方案信息表"},
                    "qualityReport": {"status": "passed"},
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        work_dir = Path(self.temp_dir.name) / "s7"
        gap_plan_path = tech_assembly._prepare_gap_plan(project_id, work_dir)

        self.assertIsNotNone(gap_plan_path)
        recovered = json.loads(Path(gap_plan_path).read_text(encoding="utf-8"))
        item = recovered["items"][0]
        self.assertEqual(item["status"], "resolved")
        self.assertEqual(item["matchedMaterials"], [])
        self.assertEqual(len(item["resolvedArtifacts"]), 1)
        self.assertEqual(item["resolvedArtifacts"][0]["source"], "ai_fill")
        self.assertEqual(item["resolvedArtifacts"][0]["path"], str(filled_docx))
        self.assertTrue(item["resolvedArtifacts"][0]["s7Ready"])
        self.assertEqual(recovered["s7RecoveredAiFillArtifactCount"], 1)

    def test_s7_recovered_ai_fill_without_quality_pass_stays_blocked(self) -> None:
        from app.services import tech_assembly

        project_id = self._prepare_project_after_outline()
        project = store._require(project_id)
        project["gap_state"] = {
            "plan": {
                "schemaVersion": "bid-tech-gap-plan-v1",
                "status": "ready",
                "items": [
                    {
                        "id": "GAP-0058",
                        "number": "附表A.1",
                        "title": "投标机型总方案信息表",
                        "status": "needs_input",
                        "matchedMaterials": [],
                        "resolvedArtifacts": [],
                    }
                ],
            },
            "integrity": {"status": "blocked", "blockingCount": 1},
        }
        store._persist_project(project)

        ai_fill_dir = technical_workspace_stage_dir(project_id, "s4_gap_workdir") / "ai_fill" / "GAP-0058"
        ai_fill_dir.mkdir(parents=True, exist_ok=True)
        filled_docx = ai_fill_dir / "投标机型总方案信息表_AI填写.docx"
        doc = Document()
        doc.add_paragraph("未验收填写结果")
        doc.save(filled_docx)

        recovered = tech_assembly._with_recovered_ai_fill_artifacts(
            project_id,
            json.loads(json.dumps(project["gap_state"]["plan"], ensure_ascii=False)),
        )

        artifact = recovered["items"][0]["resolvedArtifacts"][0]
        self.assertFalse(artifact["s7Ready"])
        self.assertEqual(artifact["qualityGate"], "needs_review")
        self.assertFalse(tech_assembly._gap_plan_has_resolved_artifacts(recovered))

    def test_wiki_export_failure_keeps_runtime_wiki_available(self) -> None:
        from app.services import tech_assembly

        project_id = self._prepare_project_after_outline()
        project = store.get_project(project_id)
        work_dir = Path(self.temp_dir.name) / "wiki-fallback-workdir"

        with patch.object(tech_assembly, "export_wiki", side_effect=RuntimeError("remote closed")):
            wiki_dir = tech_assembly._prepare_wiki_dir(project, {}, work_dir)

        self.assertTrue(wiki_dir.exists())
        self.assertTrue((wiki_dir / "卡片").exists())

    def test_runtime_material_card_matching_prefers_specific_child_section(self) -> None:
        from app.services.tech_assembly import _best_toc_section_for_material

        section, score = _best_toc_section_for_material(
            {
                "name": "技术标-投标项目塔筒专题.docx",
                "cleanedFileName": "技术标-投标项目塔筒专题.docx",
                "folderPath": "技术标/通用素材/技术标-专题方案要求",
            },
            [
                {"section": "5", "title": "专题方案要求"},
                {"section": "5.10", "title": "投标项目塔筒专题"},
            ],
        )

        self.assertGreaterEqual(score, 8)
        self.assertEqual(section, "5.10")

    def test_tech_format_outline_restores_appendix_depth_from_numbering(self) -> None:
        from app.services.tech_assembly import _tech_format_sections_from_toc_items

        sections = _tech_format_sections_from_toc_items(
            [
                {"number": "第1章", "title": "正文", "level": 1},
                {"number": "1.1", "title": "二级标题", "level": 2},
                {"number": "1.1.1", "title": "三级标题", "level": 1},
                {"number": "1.1.1.1", "title": "四级标题", "level": 1},
                {"number": "附表B.1", "title": "供货范围", "level": 1},
                {"number": "附表B.1.1", "title": "风机供货范围", "level": 1},
                {"number": "附表B.1.1.1", "title": "叶片供货范围", "level": 1},
                {"number": "附表F.5-2", "title": "认证未完成项", "level": 1},
            ]
        )

        flattened: list[dict] = []

        def visit(nodes):
            for node in nodes:
                flattened.append(node)
                visit(node.get("children") or [])

        visit(sections)
        by_number = {item["number"]: item["level"] for item in flattened}
        self.assertEqual(by_number["1.1.1"], 3)
        self.assertEqual(by_number["1.1.1.1"], 4)
        self.assertEqual(by_number["附表B.1"], 2)
        self.assertEqual(by_number["附表B.1.1"], 3)
        self.assertEqual(by_number["附表B.1.1.1"], 4)
        self.assertEqual(by_number["附表F.5-2"], 3)

    def test_runtime_material_card_writes_clean_name_for_custom_override_rules(self) -> None:
        from app.services.tech_assembly import _render_runtime_material_card

        card_text = _render_runtime_material_card(
            {
                "id": "RAW-0081",
                "name": "项目技术承诺函.docx",
                "cleanedFileName": "项目技术承诺函.docx",
                "materialTier": "customer",
                "folderPath": "定制素材/技术标",
                "bidType": "技术标",
            },
            "4",
        )

        self.assertIn('name: "项目技术承诺函"', card_text)
        self.assertIn('skeleton_section: "4"', card_text)

        with self.assertRaises(ValueError):
            _render_runtime_material_card({"id": "RAW-0082", "name": "缺标类素材.docx"}, "4")

    def test_s7_manifest_includes_confirmed_gap_plan_path(self) -> None:
        from app.services import tech_assembly

        project_id = self._prepare_project_for_s7()
        project = store._require(project_id)
        gap_state = project["gap_state"]
        gap_state["plan"] = {
            "schemaVersion": "bid-tech-gap-plan-v1",
            "status": "ready",
            "items": [
                {
                    "id": "GAP-1",
                    "number": "1.1",
                    "title": "项目背景",
                    "status": "matched",
                    "matchedMaterials": [{"id": "RAW-0001", "path": "技术标/通用素材/项目背景.docx"}],
                    "resolvedArtifacts": [],
                }
            ],
        }
        gap_state["planFile"] = ""
        gap_state["integrity"] = {"status": "passed", "blockingCount": 0}
        gap_state["reviewConfirmed"] = True
        store._persist_project(project)

        manifest_payloads = []

        def fake_prepare_wiki_dir(project, parse_storage, work_dir):
            wiki_dir = work_dir / "wiki"
            cards_dir = wiki_dir / "卡片"
            cards_dir.mkdir(parents=True, exist_ok=True)
            (cards_dir / "项目背景.md").write_text(
                "---\n"
                "name: 项目背景\n"
                "path: 技术标/通用素材/项目背景.docx\n"
                "scope: 通用\n"
                "category: 技术标\n"
                "material_id: RAW-0001\n"
                "skeleton_section: \"1.1\"\n"
                "deprecated: false\n"
                "---\n",
                encoding="utf-8",
            )
            return wiki_dir

        def fake_export_material_library(wiki_dir, library_dir):
            library_dir.mkdir(parents=True, exist_ok=True)
            return library_dir, [{"id": "RAW-0001", "path": "技术标/通用素材/项目背景.docx", "available": True}]

        def fake_run_assembler_manifest(manifest_path, progress_callback=None):
            manifest_payloads.append(json.loads(Path(manifest_path).read_text(encoding="utf-8")))
            output_file = Path(manifest_payloads[-1]["outputFile"])
            output_file.parent.mkdir(parents=True, exist_ok=True)
            doc = Document()
            doc.add_paragraph("项目背景")
            doc.save(output_file)
            plan_file = output_file.parent / "assembly_plan.json"
            plan_file.write_text(
                json.dumps(
                    [
                        {
                            "toc_idx": 1,
                            "level": 1,
                            "title": "项目背景",
                            "status": "MATCHED",
                            "paths": ["技术标/通用素材/项目背景.docx"],
                        }
                    ],
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            report = output_file.parent / "assembly_report.md"
            review = output_file.parent / "needs_review.md"
            report.write_text("ok", encoding="utf-8")
            review.write_text("ok", encoding="utf-8")
            return {
                "schema_version": "bid-tech-assembly-v1",
                "outputFile": str(output_file),
                "planFile": str(plan_file),
                "assemblyReport": str(report),
                "needsReview": str(review),
                "summary": {"total": 1, "byStatus": {"MATCHED": 1}, "usedPathCount": 1},
            }

        with patch.object(tech_assembly, "_prepare_wiki_dir", side_effect=fake_prepare_wiki_dir), \
            patch.object(tech_assembly, "_export_material_library", side_effect=fake_export_material_library), \
            patch.object(tech_assembly, "_run_assembler_manifest", side_effect=fake_run_assembler_manifest):
            result = tech_assembly.assemble_tech_bid_for_project_with_progress(project_id)

        self.assertEqual(len(manifest_payloads), 1)
        self.assertIn("gapPlanPath", manifest_payloads[0])
        gap_plan_path = Path(manifest_payloads[0]["gapPlanPath"])
        self.assertTrue(gap_plan_path.exists())
        gap_plan = json.loads(gap_plan_path.read_text(encoding="utf-8"))
        self.assertEqual(gap_plan["schemaVersion"], "bid-tech-gap-plan-v1")
        self.assertEqual(result["assembly"]["formatClean"]["status"], "completed")


if __name__ == "__main__":
    unittest.main()
