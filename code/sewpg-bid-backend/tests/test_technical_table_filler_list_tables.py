"""清单型附表列检测与无需填写终态的单测（B 系列 0 填写攻坚）。

金标反评定位：
1. 清单型表（左侧行标签 + 右侧多列待填）与参数表（一列字段名 + 一列响应值）
   共用列检测，表头里的「规格」「内容」命中响应列词表、「名称」让待填列进了
   字段列候选，最终字段列指向空列，整表 0 字段（B.1.1 43 行全丢）。
2. fill_doc 按 (表, 行) 去重，清单表同一行多列决策互相覆盖只落最后一格。
3. 招标原文已写满的附表填 0 格被判 needs_review，成为假缺口（B.1.2）。
"""
from __future__ import annotations

import importlib.util
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

_SRC = (
    Path(__file__).resolve().parents[1]
    / "opencode"
    / "skills"
    / "bid-tech-table-filler"
    / "scripts"
    / "run_from_manifest.py"
)
_SPEC = importlib.util.spec_from_file_location("tech_table_filler_list_under_test", _SRC)
filler = importlib.util.module_from_spec(_SPEC)
assert _SPEC.loader is not None
sys.modules["tech_table_filler_list_under_test"] = filler
_SPEC.loader.exec_module(filler)


def _build_docx(path: Path, heading: str, rows: list[list[str]]) -> Path:
    doc = Document()
    doc.add_paragraph(heading)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for ri, row in enumerate(rows):
        for ci, value in enumerate(row):
            table.rows[ri].cells[ci].text = value
    doc.save(str(path))
    return path


# 供货范围清单：序号空、货物名称有值、品牌/型号/原产地/数量待填、备注部分有值
_SUPPLY_ROWS = [
    ["序号", "货物名称", "品牌或制造商名称", "型号和规格", "原产地", "数量", "备注"],
    ["", "主控系统", "", "", "", "", ""],
    ["", "发电机", "", "", "", "", "含发电机轴承"],
    ["", "变流器", "", "", "", "", ""],
    ["", "齿轮箱", "", "", "", "", ""],
]

# 参数表：编号 + 项目 + 主要项目 + 技术参数与规格（单列响应）+ 计量单位 + 备注
_PARAM_ROWS = [
    ["编号", "项目", "主要项目", "技术参数与规格", "计量单位", "备注"],
    ["1", "机型总体参数", "投标机型", "", "", ""],
    ["2", "机型总体参数", "单机容量", "", "kW", ""],
    ["3", "机型总体参数", "叶轮直径", "", "m", ""],
]


class DetectListTableLayoutTests(unittest.TestCase):
    def setUp(self) -> None:
        self._tmp = tempfile.TemporaryDirectory()
        self.base = Path(self._tmp.name)

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def test_supply_list_detected_with_multiple_fill_cols(self) -> None:
        path = _build_docx(self.base / "supply.docx", "附表B.1.1 供货范围清单", _SUPPLY_ROWS)
        table = Document(str(path)).tables[0]
        layout = filler.detect_list_table_layout(table)
        self.assertIsNotNone(layout)
        header_row, label_col, fill_cols = layout
        self.assertEqual(header_row, 0)
        self.assertEqual(label_col, 1)  # 序号列为空，行标签是「货物名称」
        # 备注列有值且在 NON_VALUE_HEADERS 里，不算待填列
        self.assertEqual(fill_cols, (2, 3, 4, 5))

    def test_param_table_not_treated_as_list(self) -> None:
        path = _build_docx(self.base / "param.docx", "附表C.1 总体技术参数与规格", _PARAM_ROWS)
        table = Document(str(path)).tables[0]
        # 参数表只有一列响应值，不该被判成清单型（否则「计量单位」也会当待填列）
        self.assertIsNone(filler.detect_list_table_layout(table))

    def test_data_row_not_mistaken_for_header(self) -> None:
        # 参数表首个数据行形如「1 | 机型总体参数 | 投标机型 | | |」，右侧多列空；
        # 若被当表头会误判为清单型
        self.assertFalse(filler.row_is_header_like(["1", "机型总体参数", "投标机型", "", "", ""]))
        self.assertTrue(filler.row_is_header_like(["序号", "货物名称", "品牌或制造商名称", "数量"]))


class ListColumnAxisTests(unittest.TestCase):
    def test_axis_recognition(self) -> None:
        self.assertEqual(filler.list_column_axis("品牌或制造商名称"), "brand")
        self.assertEqual(filler.list_column_axis("生产厂家"), "brand")
        self.assertEqual(filler.list_column_axis("型号和规格"), "model")
        self.assertEqual(filler.list_column_axis("原产地"), "origin")
        self.assertEqual(filler.list_column_axis("数量"), "quantity")
        self.assertEqual(filler.list_column_axis("货物名称"), "")

    def test_only_brand_axis_accepts_unaxised_fact(self) -> None:
        fact = {"value": "上海电气"}
        # 短名单给的事实是「部件→品牌」，标签里没有列维度；若各列都接收，
        # 同一个品牌值会被灌进型号/原产地/数量（金标反评「齿轮箱 数量 = 上海电气」）
        self.assertTrue(filler.list_value_fits_axis("brand", fact))
        for axis in ("model", "origin", "quantity", "unit"):
            self.assertFalse(filler.list_value_fits_axis(axis, fact))

    def test_numeric_value_not_accepted_as_brand(self) -> None:
        self.assertFalse(filler.list_value_fits_axis("brand", {"value": "60"}))

    def test_score_zero_when_axis_mismatch(self) -> None:
        field = {
            "field": "齿轮箱 数量",
            "concepts": [],
            "listColumnLabel": "数量",
            "listColumnAxis": "quantity",
        }
        fact = {
            "label": "齿轮箱",
            "value": "上海电气",
            "concepts": [],
            "sourceKind": "xlsx",
            "baseConfidence": 0.8,
            "unit": "",
        }
        self.assertEqual(filler.score(field, fact, "auto_or_manual"), 0.0)


class FillDocMultiCellPerRowTests(unittest.TestCase):
    """清单型附表一行要落多格；去重键必须含列，否则同行决策互相覆盖。"""

    def setUp(self) -> None:
        self._tmp = tempfile.TemporaryDirectory()
        self.base = Path(self._tmp.name)

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def test_same_row_multiple_columns_all_written(self) -> None:
        path = _build_docx(self.base / "supply.docx", "附表B.1.1 供货范围清单", _SUPPLY_ROWS)
        spec = filler.AppendixSpec(
            appendix_id="APPX-TEST",
            prefix="B1",
            title="附表B.1.1 供货范围清单",
            source=path,
            table_index=0,
            header_row=0,
            field_col=1,
            value_col=2,
            unit_col=None,
            remark_col=6,
            own_tables=1,
        )
        mapping = {
            "decisions": [
                {
                    "targetFieldId": f"B1-T0R01C{col:02d}",
                    "rowIndex": 1,
                    "tableIndex": 0,
                    "valueCol": col,
                    "unitCol": None,
                    "field": f"主控系统 col{col}",
                    "action": "fill",
                    "value": value,
                    "unit": "",
                    "confidence": 0.8,
                }
                for col, value in ((2, "上海电气"), (3, "EW10.0-220"), (4, "上海"), (5, "60"))
            ]
        }
        output = self.base / "out.docx"
        filler.fill_doc(spec, mapping, output)
        row = Document(str(output)).tables[0].rows[1]
        self.assertEqual(
            [row.cells[col].text.strip() for col in (2, 3, 4, 5)],
            ["上海电气", "EW10.0-220", "上海", "60"],
        )


class NoFillRequiredTests(unittest.TestCase):
    def setUp(self) -> None:
        self._tmp = tempfile.TemporaryDirectory()
        self.base = Path(self._tmp.name)

    def tearDown(self) -> None:
        self._tmp.cleanup()

    def _spec(self, path: Path) -> object:
        return filler.AppendixSpec(
            appendix_id="APPX-TEST",
            prefix="B1",
            title=path.stem,
            source=path,
            table_index=0,
            header_row=0,
            field_col=1,
            value_col=2,
            unit_col=None,
            remark_col=3,
            own_tables=1,
        )

    def test_fully_filled_template_needs_no_fill(self) -> None:
        # 招标原文已写满，只有备注列空（备注在 NON_VALUE_HEADERS，不算待填）
        rows = [
            ["序号", "货物名称", "品牌或制造商名称", "备注"],
            ["1", "主控系统", "国产自主可控", ""],
            ["2", "发电机", "湘潭电机、中车电机", ""],
        ]
        path = _build_docx(self.base / "brand.docx", "附表B.1.2 机型配置品牌表1", rows)
        spec = self._spec(path)
        self.assertEqual(filler.appendix_pending_cell_count(spec), 0)
        self.assertTrue(filler.appendix_has_no_fill_target(spec))

    def test_template_with_pending_cell_is_not_no_fill(self) -> None:
        rows = [
            ["序号", "货物名称", "品牌或制造商名称", "备注"],
            ["1", "主控系统", "国产自主可控", ""],
            ["2", "机舱起重机", "", ""],
        ]
        path = _build_docx(self.base / "brand2.docx", "附表B.1.3 机型配置品牌表2", rows)
        spec = self._spec(path)
        self.assertEqual(filler.appendix_pending_cell_count(spec), 1)
        self.assertFalse(filler.appendix_has_no_fill_target(spec))


if __name__ == "__main__":
    unittest.main()
