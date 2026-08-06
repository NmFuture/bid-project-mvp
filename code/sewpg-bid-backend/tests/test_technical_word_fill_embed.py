"""正文「待插入」类型单测：占位符文字命中素材名 → 整份 Word 嵌入。

待填写与待插入同构，共用一套精确定位，差别只在查表对象与执行动作：待填写拿占位符
文字查事实表字段取一个值，待插入拿占位符文字查 manifest.embedSources 取一份素材整个
嵌进去。后缀是类型的权威标记，因此混合文件天然支持。
fixture 为脱敏合成数据（通用领域词面，无真实项目数据）。
"""
from __future__ import annotations

import importlib.util
import json
import sys
import tempfile
import unittest
from pathlib import Path

_SRC = (
    Path(__file__).resolve().parents[1]
    / "opencode"
    / "skills"
    / "bid-tech-word-placeholder-filler"
    / "scripts"
    / "run_from_manifest.py"
)
_SPEC = importlib.util.spec_from_file_location("tech_word_filler_embed_under_test", _SRC)
filler = importlib.util.module_from_spec(_SPEC)
assert _SPEC.loader is not None
sys.modules["tech_word_filler_embed_under_test"] = filler
_SPEC.loader.exec_module(filler)


class PlaceholderKindTests(unittest.TestCase):
    def test_embed_suffix_is_recognized_as_embed_kind(self) -> None:
        found = filler.find_placeholders("[设备清单，待插入]")

        self.assertEqual(len(found), 1)
        self.assertEqual(found[0]["kind"], "embed")
        self.assertEqual(found[0]["label"], "设备清单")

    def test_fill_suffixes_stay_fill_kind(self) -> None:
        for text in ("[安全等级，待填写]", "【安全等级, 待补充】", "[安全等级，待确认]"):
            found = filler.find_placeholders(text)
            self.assertEqual(len(found), 1, text)
            self.assertEqual(found[0]["kind"], "fill", text)
            self.assertEqual(found[0]["label"], "安全等级", text)

    def test_prefix_form_without_captured_suffix_stays_fill_kind(self) -> None:
        # `[待填写：字段名]` 捕不到后缀组，仍须按 fill 识别，不能被新分型漏掉
        found = filler.find_placeholders("[待填写：安全等级]")

        self.assertEqual(len(found), 1)
        self.assertEqual(found[0]["kind"], "fill")
        self.assertEqual(found[0]["label"], "安全等级")

    def test_plain_bracket_text_is_not_a_placeholder(self) -> None:
        self.assertEqual(filler.find_placeholders("参见[附录三]的说明"), [])

    def test_embed_placeholder_key_matches_spec_style_key(self) -> None:
        # 归一化后与素材名可比：全半角括号、分隔符差异不影响命中
        self.assertEqual(filler.placeholder_key("[设备清单，待插入]"), filler.norm("设备清单"))

    def test_standalone_embed_requires_placeholder_to_own_the_paragraph(self) -> None:
        text = "[设备清单，待插入]"
        self.assertTrue(filler.standalone_embed(text, filler.find_placeholders(text)))

        mixed = "详见[设备清单，待插入]后附材料。"
        self.assertFalse(filler.standalone_embed(mixed, filler.find_placeholders(mixed)))

        fill_only = "[安全等级，待填写]"
        self.assertFalse(filler.standalone_embed(fill_only, filler.find_placeholders(fill_only)))


def _write_embed_source(path: Path, heading: str) -> None:
    from docx import Document

    document = Document()
    document.add_paragraph(heading)
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "部件"
    table.rows[0].cells[1].text = "数量"
    table.rows[1].cells[0].text = "机架"
    table.rows[1].cells[1].text = "12"
    document.save(str(path))


def _write_blank_source(path: Path) -> None:
    from docx import Document

    document = Document()
    document.add_paragraph("一、总体方案")
    document.add_paragraph("本项目安全等级为[安全等级，待填写]。")
    document.add_paragraph("[设备清单，待插入]")
    document.add_paragraph("[价格表，待插入]")
    document.add_paragraph("详见[附件材料，待插入]后附。")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "清单"
    table.rows[0].cells[1].text = "[表内素材，待插入]"
    document.save(str(path))


class EmbedRunTests(unittest.TestCase):
    def _run(self, tmp: Path) -> dict:
        blank = tmp / "待填写-总体方案.docx"
        embed_source = tmp / "设备清单.docx"
        _write_blank_source(blank)
        _write_embed_source(embed_source, "设备清单明细")
        manifest = {
            "schemaVersion": "bid-tech-word-placeholder-fill-v1",
            "title": "总体方案",
            "blankSource": {"docxPath": str(blank), "title": "待填写-总体方案.docx"},
            "outputFile": str(tmp / "out.docx"),
            "projectFactTable": {
                "status": "confirmed",
                "fields": [
                    {
                        "label": "安全等级",
                        "value": "一级",
                        "placeholder": "[安全等级，待填写]",
                        "targetFile": "待填写-总体方案.docx",
                    }
                ],
            },
            "embedSources": [
                {
                    "placeholder": "设备清单",
                    "materialId": "RAW-0001",
                    "name": "设备清单.docx",
                    "materialTier": "project",
                    "status": "ready",
                    "docxPath": str(embed_source),
                },
                {
                    "placeholder": "价格表",
                    "materialId": "RAW-0002",
                    "name": "价格表.xlsx",
                    "materialTier": "project",
                    "status": "unsupported_format",
                    "statusMessage": "「价格表.xlsx」是 Excel 素材，请人工另存为 Word 后重新上传。",
                },
            ],
        }
        manifest_path = tmp / "word_fill_input.json"
        manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
        return filler.run_from_manifest(manifest_path)

    def test_embed_counts_are_separate_from_filled_values(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            report = self._run(Path(raw))["fillReport"]

        # 整份嵌入不混进「填了几个值」，否则审核界面看不出哪些是整份文档
        self.assertEqual(report["filledPlaceholderCount"], 1)
        self.assertEqual(report["embeddedCount"], 1)
        self.assertEqual(report["manualEmbedCount"], 3)
        self.assertEqual(report["unfilledPlaceholderCount"], 3)
        self.assertEqual(report["manualEmbedMarker"], "[待人工插入：素材名]")

    def test_ready_source_is_inserted_and_placeholder_paragraph_removed(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as raw:
            result = self._run(Path(raw))
            document = Document(result["outputFile"])
            texts = [paragraph.text for paragraph in document.paragraphs]
            tables = len(document.tables)

        self.assertIn("设备清单明细", texts)
        self.assertNotIn("[设备清单，待插入]", texts)
        self.assertIn("本项目安全等级为一级。", texts)
        # 空白模板自带 1 张表 + 素材带进来的 1 张
        self.assertEqual(tables, 2)

    def test_unsupported_ambiguous_and_mixed_placeholders_fall_back_to_manual(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as raw:
            result = self._run(Path(raw))
            document = Document(result["outputFile"])
            texts = [paragraph.text for paragraph in document.paragraphs]
            cell_text = document.tables[-1].rows[0].cells[1].text

        self.assertIn("[待人工插入：价格表]", texts)
        self.assertIn("详见[待人工插入：附件材料]后附。", texts)
        self.assertEqual(cell_text, "[待人工插入：表内素材]")

    def test_manual_embed_reason_is_reported_not_swallowed(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            report = self._run(Path(raw))["fillReport"]

        by_label = {item["label"]: item for item in report["embedDetails"]}
        self.assertEqual(by_label["设备清单"]["action"], "embed")
        self.assertIn("Excel 素材", by_label["价格表"]["message"])
        self.assertEqual(by_label["附件材料"]["status"], "not_standalone")
        self.assertEqual(by_label["表内素材"]["status"], "not_standalone")

    def test_missing_embed_source_is_marked_manual_with_reason(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            tmp = Path(raw)
            blank = tmp / "待填写-总体方案.docx"
            _write_blank_source(blank)
            manifest_path = tmp / "word_fill_input.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-tech-word-placeholder-fill-v1",
                        "blankSource": {"docxPath": str(blank), "title": "待填写-总体方案.docx"},
                        "outputFile": str(tmp / "out.docx"),
                        "projectFactTable": {
                            "status": "confirmed",
                            "fields": [
                                {
                                    "label": "安全等级",
                                    "value": "一级",
                                    "placeholder": "[安全等级，待填写]",
                                    "targetFile": "待填写-总体方案.docx",
                                }
                            ],
                        },
                        "embedSources": [],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            report = filler.run_from_manifest(manifest_path)["fillReport"]

        self.assertEqual(report["embeddedCount"], 0)
        self.assertEqual(report["manualEmbedCount"], 4)
        by_label = {item["label"]: item for item in report["embedDetails"]}
        self.assertEqual(by_label["设备清单"]["status"], "not_found")


if __name__ == "__main__":
    unittest.main()
