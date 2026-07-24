from __future__ import annotations

import asyncio
import base64
import io
import json
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from fastapi.testclient import TestClient
from openpyxl import Workbook

from app.main import app
from app.core.config import settings
from app.services.bid_outline_state import confirm_outline_state, save_generated_outline_state
from app.services.bid_runtime_state import now_iso
from app.services.store import store
from app.services.technical_gap_repository import persist_technical_gap_project, require_technical_gap_project_for_update
from app.services.technical_gap_review import confirm_technical_review, prepare_technical_review_document
from app.services.technical_gap_service import technical_gap_service
from app.services.technical_gap_state import ensure_technical_gap_state
from app.services.turbine_models import extract_turbine_model_options_from_xlsx_bytes
from app.services.workspace_artifacts import technical_workspace_dir


def _xlsx_payload() -> str:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "01-X2平台机组主参数20260120"
    sheet.cell(row=1, column=5, value="在役")
    sheet.cell(row=2, column=5, value="变压器上置")
    sheet.cell(row=4, column=5, value="X2A")
    sheet.cell(row=5, column=5, value="EW10.0-230上置")
    sheet.cell(row=6, column=5, value=10000)
    sheet.cell(row=8, column=5, value=226)
    sheet.cell(row=1, column=6, value="研发中")
    sheet.cell(row=2, column=6, value="变压器下置")
    sheet.cell(row=4, column=6, value="X2E-2")
    sheet.cell(row=5, column=6, value="EW10.0-220下置")
    sheet.cell(row=6, column=6, value=10000)
    sheet.cell(row=8, column=6, value=220)
    sheet.cell(row=5, column=7, value="CGC202546131009230")
    sheet.cell(row=5, column=8, value="EW-DFIG60")
    sheet.cell(row=5, column=9, value="20250106")
    buffer = io.BytesIO()
    workbook.save(buffer)
    return base64.b64encode(buffer.getvalue()).decode("ascii")


def _confirm_technical_review_for_tests(project_id: str) -> None:
    project = require_technical_gap_project_for_update(project_id)
    gap_state = ensure_technical_gap_state(project)
    prepare_technical_review_document(project, gap_state)
    confirm_technical_review(project, gap_state)
    persist_technical_gap_project(project)


def _save_generated_outline_for_tests(
    project_id: str,
    *,
    nodes: list[dict],
    generated_at: str,
    summary: str,
) -> dict:
    project = store.require_project_for_update(project_id)
    payload = save_generated_outline_state(
        project,
        nodes=nodes,
        generated_at=generated_at,
        summary=summary,
    )
    store.persist_project_state(project)
    return payload


def _confirm_outline_for_tests(project_id: str) -> dict:
    project = store.require_project_for_update(project_id)
    payload = confirm_outline_state(project)
    store.persist_project_state(project)
    return payload


class TurbineModelSelectionTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        base = Path(self.temp_dir.name)
        settings.uploads_dir = base / "uploads"
        settings.documents_dir = base / "documents"
        settings.parsed_dir = base / "parsed"
        settings.ensure_dirs()
        store.reset_for_tests(clear_persistent=True)
        self.client = TestClient(app, base_url="http://127.0.0.1:8000")

    def tearDown(self) -> None:
        self.client.close()
        store.reset_for_tests(clear_persistent=True)
        self.temp_dir.cleanup()

    def _create_project(self) -> str:
        response = self.client.post(
            "/api/technical/projects",
            json={
                "name": "机型字段验证项目",
                "customerName": "华能集团",
                "bidType": "技术标",
                "turbineModel": {
                    "model": "EW10.0-220下置",
                    "platform": "X2E-2",
                    "ratedPowerKw": 10000,
                    "rotorDiameterM": 220,
                    "layout": "变压器下置",
                    "status": "research",
                    "source": "manual",
                    "aliases": ["EW10.0-220", "W10.0-220"],
                },
            },
        )
        self.assertEqual(response.status_code, 200, response.text)
        return str(response.json()["id"])

    def test_turbine_model_options_come_from_real_parameter_sheet_without_noise(self) -> None:
        options = extract_turbine_model_options_from_xlsx_bytes(
            base64.b64decode(_xlsx_payload()),
            source_file_id="RAW-0001",
            source_file_name="X2平台机型投标参数_20250106.xlsx",
            folder_path="技术标/通用素材/投标机型参数表（产品管理部：在线文档）",
        )

        labels = [item["model"] for item in options]
        self.assertIn("EW10.0-230上置", labels)
        self.assertIn("EW10.0-220下置", labels)
        self.assertNotIn("CGC202546131009230", labels)
        self.assertNotIn("EW-DFIG60", labels)
        self.assertNotIn("20250106", labels)
        research = next(item for item in options if item["model"] == "EW10.0-220下置")
        self.assertEqual(research["platform"], "X2E-2")
        self.assertEqual(research["ratedPowerKw"], 10000)
        self.assertEqual(research["rotorDiameterM"], 220)
        self.assertEqual(research["status"], "research")
        self.assertTrue(research["evidence"])

    def test_project_persists_selected_turbine_model(self) -> None:
        project_id = self._create_project()

        payload = self.client.get(f"/api/technical/projects/{project_id}").json()

        self.assertEqual(payload["turbineModel"]["model"], "EW10.0-220下置")
        self.assertEqual(payload["turbineModel"]["platform"], "X2E-2")
        self.assertEqual(payload["turbineModelLabel"], "EW10.0-220下置")
        self.assertIn("W10.0-220", payload["turbineModel"]["aliases"])

    def test_gap_manifest_ai_fill_and_assembly_carry_selected_turbine_model(self) -> None:
        from app.services import tech_assembly

        project_id = self._create_project()
        project_dir = technical_workspace_dir(project_id)
        toc_dir = project_dir / "s2_toc_workdir"
        toc_dir.mkdir(parents=True, exist_ok=True)
        toc_path = toc_dir / "投标文件-总目录.json"
        toc_path.write_text(
            json.dumps(
                {
                    "schema_version": "bid-toc-json-v1",
                    "items": [
                        {
                            "order": 1,
                            "number": "1.1",
                            "title": "性能保证",
                            "level": 2,
                            "annotation": "新增-招标要求",
                            "source": "tender",
                            "material_refs": [],
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        _save_generated_outline_for_tests(
            project_id=project_id,
            nodes=[{"id": "OL-1", "title": "性能保证", "children": []}],
            generated_at=now_iso(),
            summary="目录已生成。",
        )
        _confirm_outline_for_tests(project_id)
        project = store._require(project_id)
        project["parse_result"] = {
            "status": "completed",
            "structured": {
                "appendices": [
                    {"id": "APP-1", "title": "性能保证附表", "sourceFile": "招标文件.docx"}
                ]
            },
        }
        store._persist_project(project)

        gap_manifests: list[dict] = []

        def fake_gap_planner(manifest_path):
            manifest = json.loads(Path(manifest_path).read_text(encoding="utf-8"))
            gap_manifests.append(manifest)
            output_file = Path(manifest["outputFile"])
            plan = {
                "schemaVersion": "bid-tech-gap-plan-v1",
                "projectId": project_id,
                "projectTurbineModel": manifest["projectTurbineModel"],
                "status": "ready",
                "summary": {
                    "totalTocItems": 1,
                    "matchedCount": 0,
                    "missingCount": 1,
                    "resolvedCount": 0,
                    "ignoredCount": 0,
                    "structuralCount": 0,
                    "fillableTaskCount": 1,
                    "blockingCount": 1,
                },
                "items": [
                    {
                        "id": "GAP-0001",
                        "number": "1.1",
                        "title": "性能保证",
                        "status": "needs_input",
                        "priority": "high",
                        "matchedMaterials": [],
                        "requiredInputs": [{"type": "ai_fill", "label": "选择参考素材并填写空表"}],
                        "fillTasks": [
                            {
                                "id": "FILL-GAP-0001",
                                "skill": "bid-tech-table-filler",
                                "status": "pending",
                                "title": "填写性能保证",
                                "blankSource": {"id": "APP-1", "title": "性能保证附表"},
                                "requiredReferences": ["素材库文件", "招标解析字段"],
                            }
                        ],
                        "resolvedArtifacts": [],
                        "reviewNotes": [],
                        "gapReason": "待按投标机型填写。",
                    }
                ],
                "integrity": {"status": "blocked", "blockingCount": 1},
            }
            output_file.write_text(json.dumps(plan, ensure_ascii=False, indent=2), encoding="utf-8")
            return {"schema_version": "bid-tech-gap-plan-v1", "outputFile": str(output_file)}

        with patch("app.services.technical_gap_planner.run_technical_gap_planner_skill", side_effect=fake_gap_planner):
            detection = self.client.post(f"/api/technical/projects/{project_id}/gaps-detection/run")
        self.assertEqual(detection.status_code, 200, detection.text)
        self.assertEqual(gap_manifests[0]["projectTurbineModel"]["model"], "EW10.0-220下置")
        self.assertEqual(detection.json()["gapPlan"]["projectTurbineModel"]["platform"], "X2E-2")
        facts = self.client.post(f"/api/technical/projects/{project_id}/gaps/facts/build")
        self.assertEqual(facts.status_code, 200, facts.text)
        confirm_facts = self.client.put(
            f"/api/technical/projects/{project_id}/gaps/facts",
            json={"fields": facts.json()["fields"], "confirm": True, "operator": "测试用户"},
        )
        self.assertEqual(confirm_facts.status_code, 200, confirm_facts.text)

        fill_manifests: list[dict] = []

        def fake_table_filler(manifest_path, progress_callback=None):
            manifest = json.loads(Path(manifest_path).read_text(encoding="utf-8"))
            fill_manifests.append(manifest)
            output_file = Path(manifest["outputFile"])
            doc = Document()
            doc.add_paragraph("按 EW10.0-220下置 填写性能保证。")
            doc.save(output_file)
            return {
                "schema_version": "bid-tech-table-fill-v1",
                "outputFile": str(output_file),
                "unfilledFields": [],
                "evidenceRefs": [{"field": "性能保证", "source": "projectTurbineModel"}],
                "fillReport": {
                    "targetFieldCount": 1,
                    "filledFieldCount": 1,
                    "unfilledFieldCount": 0,
                    "semanticCheckCount": 1,
                    "semanticFailedCount": 0,
                    "semanticValidationRate": 1,
                },
            }

        with patch("app.services.technical_gap_ai_fill.run_technical_table_filler_skill", side_effect=fake_table_filler):
            fill = self.client.post(
                f"/api/technical/projects/{project_id}/gaps/GAP-0001/ai-fill",
                json={"fillTaskId": "FILL-GAP-0001", "referenceMaterialIds": ["RAW-0001"], "parseFieldIds": ["APP-1"]},
            )
        self.assertEqual(fill.status_code, 200, fill.text)
        self.assertEqual(fill_manifests[0]["projectTurbineModel"]["model"], "EW10.0-220下置")

        asyncio.run(technical_gap_service.recheck(project_id))
        asyncio.run(technical_gap_service.submit_review(project_id))
        _confirm_technical_review_for_tests(project_id)
        assembly_manifests: list[dict] = []

        def fake_prepare_wiki_dir(project, parse_storage, work_dir):
            wiki_dir = work_dir / "wiki"
            (wiki_dir / "卡片").mkdir(parents=True, exist_ok=True)
            return wiki_dir

        def fake_export_material_library(wiki_dir, library_dir):
            library_dir.mkdir(parents=True, exist_ok=True)
            return library_dir, []

        def fake_run_assembler(manifest_path, progress_callback=None):
            manifest = json.loads(Path(manifest_path).read_text(encoding="utf-8"))
            assembly_manifests.append(manifest)
            output_file = Path(manifest["outputFile"])
            doc = Document()
            doc.add_paragraph("正文")
            doc.save(output_file)
            plan_file = output_file.parent / "assembly_plan.json"
            report = output_file.parent / "assembly_report.md"
            review = output_file.parent / "needs_review.md"
            plan_file.write_text("[]", encoding="utf-8")
            report.write_text("ok", encoding="utf-8")
            review.write_text("ok", encoding="utf-8")
            return {
                "schema_version": "bid-tech-assembly-v1",
                "outputFile": str(output_file),
                "planFile": str(plan_file),
                "assemblyReport": str(report),
                "needsReview": str(review),
                "summary": {"total": 0, "byStatus": {}, "usedPathCount": 0},
            }

        with patch.object(tech_assembly, "_prepare_wiki_dir", side_effect=fake_prepare_wiki_dir), \
            patch.object(tech_assembly, "_export_material_library", side_effect=fake_export_material_library), \
            patch.object(tech_assembly, "_run_assembler_manifest", side_effect=fake_run_assembler):
            tech_assembly.assemble_tech_bid_for_project_with_progress(project_id)

        self.assertEqual(assembly_manifests[0]["projectTurbineModel"]["model"], "EW10.0-220下置")
        self.assertEqual(assembly_manifests[0]["projectParams"]["turbine_model"], "EW10.0-220下置")


if __name__ == "__main__":
    unittest.main()
