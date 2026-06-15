from __future__ import annotations

import copy
import json
import base64
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from app.services.bid_fill_generation_state import start_fill_generation_state
from app.services.bid_outline_state import confirm_outline_state, save_generated_outline_state
from app.services.bid_parse_state import complete_parse_state, update_parse_result_state
from app.services.bid_type import BUSINESS_BID_TYPE


def _update_parse_result_for_tests(store, project_id: str, parse_result: dict, *, parse_storage: dict | None = None) -> dict:
    project = store.require_project_for_update(project_id)
    payload = update_parse_result_state(project, parse_result, parse_storage=parse_storage)
    store.persist_project_state(project)
    return payload


def _complete_parse_for_tests(
    store,
    project_id: str,
    tender_files: list[dict],
    template_files: list[dict],
    *,
    summary: dict | None = None,
    parse_storage: dict | None = None,
) -> dict:
    project = store.require_project_for_update(project_id)
    payload = complete_parse_state(
        project,
        tender_files,
        template_files,
        summary=summary,
        parse_storage=parse_storage,
    )
    store.persist_project_state(project)
    return payload


def _fill_state_for_tests(store, project_id: str) -> dict:
    return copy.deepcopy(store.get_project_runtime_state(project_id)["fill_state"])


def _start_fill_generation_for_tests(store, project_id: str) -> dict:
    project = store.require_project_for_update(project_id)
    payload = start_fill_generation_state(project)
    store.persist_project_state(project)
    return payload


def _save_generated_outline_for_tests(
    store,
    project_id: str,
    *,
    nodes: list[dict],
    generated_at: str,
    summary: str,
) -> dict:
    project = store.require_project_for_update(project_id)
    payload = save_generated_outline_state(
        project,
        nodes=nodes,
        generated_at=generated_at,
        summary=summary,
    )
    store.persist_project_state(project)
    return payload


def _confirm_outline_for_tests(store, project_id: str) -> dict:
    project = store.require_project_for_update(project_id)
    payload = confirm_outline_state(project)
    store.persist_project_state(project)
    return payload


class BusinessAssemblyRunnerTests(unittest.TestCase):
    def test_business_assembler_runner_generates_docx_and_review_outputs(self) -> None:
        backend_root = Path(__file__).resolve().parents[1]
        script_path = backend_root / "opencode" / "skills" / "bid-business-assembler" / "scripts" / "run_from_manifest.py"
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            materials = root / "materials" / "商务标" / "项目素材" / "项目模板底稿与过程文件"
            materials.mkdir(parents=True)
            template_path = materials / "投标函空白模板.docx"
            doc = Document()
            doc.add_paragraph("项目名称：{{项目名称}}")
            doc.add_paragraph("招标编号：{{招标编号}}")
            doc.save(template_path)

            scoring_path = root / "商务评分标准.docx"
            doc = Document()
            doc.add_paragraph("商务评分标准正文")
            doc.save(scoring_path)

            toc_path = root / "toc.json"
            gap_path = root / "business_gap_plan.json"
            facts_path = root / "project_fact_table.json"
            parse_path = root / "parse_result.json"
            output_path = root / "商务投标文件.docx"
            toc_path.write_text(
                json.dumps(
                    {
                        "schema_version": "bid-toc-json-v1",
                        "items": [
                            {"number": "一", "title": "投标函", "level": 1},
                            {"number": "二", "title": "商务评分标准", "level": 1},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            gap_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-gap-plan-v1",
                        "tasks": [
                            {
                                "id": "BTASK-0010",
                                "title": "投标函",
                                "status": "ready",
                                "decision": "ready",
                                "tocTarget": {"title": "投标函"},
                                "resolvedArtifacts": [],
                            }
                        ],
                        "tocRefs": [],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            facts_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-project-fact-table-v1",
                        "status": "draft",
                        "fields": [
                            {"label": "项目名称", "value": "商务 S4 项目"},
                            {"label": "招标编号", "value": "BIZ-S4-001"},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            parse_path.write_text(
                json.dumps(
                    {
                        "status": "completed",
                        "structured": {
                            "businessScoringAsset": {
                                "id": "BIZ-SCORING-001",
                                "fileName": "商务评分标准.docx",
                                "docxPath": str(scoring_path),
                                "reviewStatus": "approved",
                            }
                        },
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            manifest_path = root / "business_assembly_input.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-BIZ-S4",
                        "projectName": "商务 S4 项目",
                        "bidType": "商务标",
                        "workDir": str(root),
                        "tocJsonPath": str(toc_path),
                        "businessGapPlanPath": str(gap_path),
                        "projectFactTablePath": str(facts_path),
                        "parseResultPath": str(parse_path),
                        "materialLibraryDir": str(root / "materials"),
                        "outputFile": str(output_path),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            completed = subprocess.run(
                [sys.executable, str(script_path), "--manifest", str(manifest_path), "--response", "summary"],
                check=True,
                capture_output=True,
                text=True,
            )
            summary = json.loads(completed.stdout)
            output_doc = Document(str(output_path))
            text = "\n".join(paragraph.text for paragraph in output_doc.paragraphs)

            self.assertEqual(summary["schema_version"], "bid-business-assembly-v1")
            self.assertTrue(output_path.exists())
            self.assertTrue((root / "business_assembly_plan.json").exists())
            self.assertTrue((root / "business_needs_review.md").exists())
            self.assertIn("商务 S4 项目", text)
            self.assertIn("BIZ-S4-001", text)
            self.assertIn("商务评分标准正文", text)
            self.assertEqual(text.count("商务评分标准正文"), 1)

    def test_business_assembler_runner_writes_scoring_table_from_structured_json(self) -> None:
        backend_root = Path(__file__).resolve().parents[1]
        script_path = backend_root / "opencode" / "skills" / "bid-business-assembler" / "scripts" / "run_from_manifest.py"
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            materials = root / "materials"
            materials.mkdir(parents=True)
            toc_path = root / "toc.json"
            gap_path = root / "business_gap_plan.json"
            facts_path = root / "project_fact_table.json"
            parse_path = root / "parse_result.json"
            output_path = root / "商务投标文件.docx"
            toc_path.write_text(
                json.dumps(
                    {
                        "schema_version": "bid-toc-json-v1",
                        "items": [{"number": "一", "title": "商务评分标准", "level": 1}],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            gap_path.write_text(
                json.dumps({"schemaVersion": "bid-business-gap-plan-v1", "tasks": [], "tocRefs": []}, ensure_ascii=False),
                encoding="utf-8",
            )
            facts_path.write_text(
                json.dumps({"schemaVersion": "bid-project-fact-table-v1", "status": "confirmed", "fields": []}, ensure_ascii=False),
                encoding="utf-8",
            )
            parse_path.write_text(
                json.dumps(
                    {
                        "status": "completed",
                        "structured": {
                            "schemaVersion": "bid-business-tender-structured-v1",
                            "scoringCriteria": {
                                "business": [
                                    {
                                        "order": "1",
                                        "scoringItem": "企业业绩",
                                        "score": "10分",
                                        "scorePoint": "近年同类项目业绩",
                                        "proofRequirement": "合同扫描件",
                                    }
                                ]
                            },
                        },
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            manifest_path = root / "business_assembly_input.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-BIZ-S4",
                        "projectName": "商务 S4 项目",
                        "bidType": "商务标",
                        "workDir": str(root),
                        "tocJsonPath": str(toc_path),
                        "businessGapPlanPath": str(gap_path),
                        "projectFactTablePath": str(facts_path),
                        "parseResultPath": str(parse_path),
                        "materialLibraryDir": str(materials),
                        "outputFile": str(output_path),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            subprocess.run(
                [sys.executable, str(script_path), "--manifest", str(manifest_path), "--response", "summary"],
                check=True,
                capture_output=True,
                text=True,
            )

            output_doc = Document(str(output_path))
            text = "\n".join(paragraph.text for paragraph in output_doc.paragraphs)
            table_text = "\n".join(cell.text for table in output_doc.tables for row in table.rows for cell in row.cells)
            self.assertIn("商务评分标准", text)
            self.assertIn("企业业绩", table_text)
            self.assertIn("合同扫描件", table_text)

    def test_business_assembler_dedupes_scoring_asset_from_task_and_parse_result(self) -> None:
        backend_root = Path(__file__).resolve().parents[1]
        script_path = backend_root / "opencode" / "skills" / "bid-business-assembler" / "scripts" / "run_from_manifest.py"
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            scoring_path = root / "商务评分标准.docx"
            doc = Document()
            doc.add_paragraph("商务评分标准正文唯一")
            doc.save(scoring_path)
            toc_path = root / "toc.json"
            gap_path = root / "business_gap_plan.json"
            facts_path = root / "project_fact_table.json"
            parse_path = root / "parse_result.json"
            output_path = root / "商务投标文件.docx"
            toc_path.write_text(
                json.dumps({"schema_version": "bid-toc-json-v1", "items": [{"number": "一", "title": "商务评分标准", "level": 1}]}, ensure_ascii=False),
                encoding="utf-8",
            )
            gap_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-gap-plan-v1",
                        "tasks": [
                            {
                                "id": "BTASK-SCORING",
                                "title": "商务评分标准",
                                "status": "ready",
                                "decision": "ready",
                                "tocTarget": {"title": "商务评分标准"},
                                "resolvedArtifacts": [
                                    {
                                        "artifactId": "BIZ-SCORING-001",
                                        "artifactType": "parse_business_scoring",
                                        "fileName": scoring_path.name,
                                        "filePath": str(scoring_path),
                                        "sourceMode": "parsed_business_scoring",
                                        "confirmed": True,
                                    }
                                ],
                            }
                        ],
                        "tocRefs": [],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            facts_path.write_text(
                json.dumps({"schemaVersion": "bid-project-fact-table-v1", "status": "confirmed", "fields": []}, ensure_ascii=False),
                encoding="utf-8",
            )
            parse_path.write_text(
                json.dumps(
                    {
                        "status": "completed",
                        "structured": {
                            "businessScoringAsset": {
                                "id": "BIZ-SCORING-001",
                                "fileName": scoring_path.name,
                                "docxPath": str(scoring_path),
                                "reviewStatus": "approved",
                            }
                        },
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            manifest_path = root / "business_assembly_input.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-BIZ-S4-DEDUP",
                        "projectName": "商务 S4 去重项目",
                        "bidType": "商务标",
                        "workDir": str(root),
                        "tocJsonPath": str(toc_path),
                        "businessGapPlanPath": str(gap_path),
                        "projectFactTablePath": str(facts_path),
                        "parseResultPath": str(parse_path),
                        "materialLibraryDir": str(root / "materials"),
                        "outputFile": str(output_path),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            subprocess.run([sys.executable, str(script_path), "--manifest", str(manifest_path), "--response", "summary"], check=True, capture_output=True, text=True)
            output_doc = Document(str(output_path))
            text = "\n".join(paragraph.text for paragraph in output_doc.paragraphs)
            plan = json.loads((root / "business_assembly_plan.json").read_text(encoding="utf-8"))

        self.assertEqual(text.count("商务评分标准正文唯一"), 1)
        scoring_sections = [section for section in plan["sections"] if section["title"] == "商务评分标准"]
        self.assertEqual(len(scoring_sections), 1)

    def test_business_assembler_fills_adjacent_blank_cells_in_template_tables(self) -> None:
        backend_root = Path(__file__).resolve().parents[1]
        script_path = backend_root / "opencode" / "skills" / "bid-business-assembler" / "scripts" / "run_from_manifest.py"
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            materials = root / "materials" / "商务标" / "项目素材" / "项目模板底稿与过程文件"
            materials.mkdir(parents=True)
            template_path = materials / "投标函格式模板.docx"
            doc = Document()
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "项目名称"
            table.cell(0, 1).text = ""
            table.cell(1, 0).text = "招标编号"
            table.cell(1, 1).text = "待填写"
            doc.save(template_path)
            toc_path = root / "toc.json"
            gap_path = root / "business_gap_plan.json"
            facts_path = root / "project_fact_table.json"
            parse_path = root / "parse_result.json"
            output_path = root / "商务投标文件.docx"
            toc_path.write_text(
                json.dumps({"schema_version": "bid-toc-json-v1", "items": [{"number": "一", "title": "投标函", "level": 1}]}, ensure_ascii=False),
                encoding="utf-8",
            )
            gap_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-gap-plan-v1",
                        "tasks": [
                            {
                                "id": "BTASK-LETTER",
                                "title": "投标函",
                                "status": "ready",
                                "decision": "ready",
                                "assemblyMode": "template_fill_docx",
                                "tocTarget": {"title": "投标函"},
                                "resolvedArtifacts": [],
                            }
                        ],
                        "tocRefs": [],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            facts_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-project-fact-table-v1",
                        "status": "confirmed",
                        "fields": [
                            {"label": "项目名称", "value": "模板填充项目"},
                            {"label": "招标编号", "value": "TPL-2026-001"},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            parse_path.write_text(json.dumps({"status": "completed", "structured": {}}, ensure_ascii=False), encoding="utf-8")
            manifest_path = root / "business_assembly_input.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-BIZ-S4-TPL",
                        "projectName": "模板填充项目",
                        "bidType": "商务标",
                        "workDir": str(root),
                        "tocJsonPath": str(toc_path),
                        "businessGapPlanPath": str(gap_path),
                        "projectFactTablePath": str(facts_path),
                        "parseResultPath": str(parse_path),
                        "materialLibraryDir": str(root / "materials"),
                        "outputFile": str(output_path),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            subprocess.run([sys.executable, str(script_path), "--manifest", str(manifest_path), "--response", "summary"], check=True, capture_output=True, text=True)
            output_doc = Document(str(output_path))
            table_text = "\n".join(cell.text for table in output_doc.tables for row in table.rows for cell in row.cells)
            field_report = json.loads((root / "field_fill_report.json").read_text(encoding="utf-8"))

        self.assertIn("模板填充项目", table_text)
        self.assertIn("TPL-2026-001", table_text)
        self.assertTrue(any(item["status"] == "filled_table_cell" for item in field_report["items"]))

    def test_business_assembler_extracts_only_matching_fragment_from_large_template(self) -> None:
        backend_root = Path(__file__).resolve().parents[1]
        script_path = backend_root / "opencode" / "skills" / "bid-business-assembler" / "scripts" / "run_from_manifest.py"
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            materials = root / "materials"
            materials.mkdir(parents=True)
            template_path = root / "商务投标文件大模板.docx"
            doc = Document()
            doc.add_paragraph("附件1 投标函")
            doc.add_paragraph("项目名称：{{项目名称}}")
            doc.add_paragraph("招标编号：{{招标编号}}")
            doc.add_paragraph("附件2 法定代表人授权书")
            doc.add_paragraph("授权书正文不应被投标函任务带入")
            doc.save(template_path)

            toc_path = root / "toc.json"
            gap_path = root / "business_gap_plan.json"
            facts_path = root / "project_fact_table.json"
            parse_path = root / "parse_result.json"
            output_path = root / "商务投标文件.docx"
            toc_path.write_text(
                json.dumps({"schema_version": "bid-toc-json-v1", "items": [{"number": "一", "title": "投标函", "level": 1}]}, ensure_ascii=False),
                encoding="utf-8",
            )
            gap_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-gap-plan-v1",
                        "tasks": [
                            {
                                "id": "BTASK-TPL-FRAGMENT",
                                "title": "投标函",
                                "status": "ready",
                                "decision": "ready",
                                "assemblyMode": "template_fill_docx",
                                "tocTarget": {"title": "投标函"},
                                "resolvedArtifacts": [
                                    {
                                        "artifactId": "BART-TPL-FRAGMENT",
                                        "artifactType": "selected_bid_template",
                                        "fileName": template_path.name,
                                        "filePath": str(template_path),
                                        "sourceMode": "project_uploaded_bid_template",
                                        "assemblyMode": "template_fill_docx",
                                        "materialUsage": "fill_template",
                                    }
                                ],
                            }
                        ],
                        "tocRefs": [],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            facts_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-project-fact-table-v1",
                        "status": "confirmed",
                        "fields": [
                            {"label": "项目名称", "value": "片段提取项目"},
                            {"label": "招标编号", "value": "FRAG-2026-001"},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            parse_path.write_text(json.dumps({"status": "completed", "structured": {}}, ensure_ascii=False), encoding="utf-8")
            manifest_path = root / "business_assembly_input.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-BIZ-S4-FRAG",
                        "projectName": "片段提取项目",
                        "bidType": "商务标",
                        "workDir": str(root),
                        "tocJsonPath": str(toc_path),
                        "businessGapPlanPath": str(gap_path),
                        "projectFactTablePath": str(facts_path),
                        "parseResultPath": str(parse_path),
                        "materialLibraryDir": str(materials),
                        "outputFile": str(output_path),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            subprocess.run([sys.executable, str(script_path), "--manifest", str(manifest_path), "--response", "summary"], check=True, capture_output=True, text=True)

            output_doc = Document(str(output_path))
            text = "\n".join(paragraph.text for paragraph in output_doc.paragraphs)
            attachment_manifest = json.loads((root / "attachment_manifest.json").read_text(encoding="utf-8"))

        self.assertIn("片段提取项目", text)
        self.assertIn("FRAG-2026-001", text)
        self.assertIn("附件1 投标函", text)
        self.assertNotIn("授权书正文不应被投标函任务带入", text)
        self.assertEqual(attachment_manifest["items"][0]["mode"], "template_fragment_fill")

    def test_business_assembler_skips_ignored_and_unconfirmed_tasks_without_template_fallback(self) -> None:
        backend_root = Path(__file__).resolve().parents[1]
        script_path = backend_root / "opencode" / "skills" / "bid-business-assembler" / "scripts" / "run_from_manifest.py"
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            materials = root / "materials" / "商务标" / "项目素材" / "项目模板底稿与过程文件"
            materials.mkdir(parents=True)
            template_path = materials / "商务响应文件模板.docx"
            doc = Document()
            doc.add_paragraph("变流器型式认证")
            doc.add_paragraph("{{安全生产许可证}}")
            doc.add_paragraph("{{质量管理体系认证证书（DNV-Business Assurance）}}")
            doc.add_paragraph("{{2024年财务情况}}")
            doc.add_paragraph("{{新疆新能博乐三台50MW项目(WH6.25N-182*8)}}")
            doc.save(template_path)

            toc_path = root / "toc.json"
            gap_path = root / "business_gap_plan.json"
            facts_path = root / "project_fact_table.json"
            parse_path = root / "parse_result.json"
            output_path = root / "商务投标文件.docx"
            toc_path.write_text(
                json.dumps(
                    {
                        "schema_version": "bid-toc-json-v1",
                        "items": [
                            {"number": "4", "title": "变流器型式认证", "level": 1},
                            {"number": "5", "title": "安全生产许可证", "level": 1},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            gap_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-gap-plan-v1",
                        "tasks": [
                            {
                                "id": "BTASK-IGNORED-CERT",
                                "title": "变流器型式认证",
                                "status": "ignored",
                                "decision": "ready",
                                "handlingMode": "ignored",
                                "assemblyMode": "template_fill_docx",
                                "tocTarget": {"title": "变流器型式认证"},
                                "resolvedArtifacts": [],
                            },
                            {
                                "id": "BTASK-UNCONFIRMED-LICENSE",
                                "title": "安全生产许可证",
                                "status": "review_required",
                                "decision": "review_required",
                                "assemblyMode": "template_fill_docx",
                                "tocTarget": {"title": "安全生产许可证"},
                                "resolvedArtifacts": [],
                            },
                        ],
                        "tocRefs": [],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            facts_path.write_text(
                json.dumps({"schemaVersion": "bid-project-fact-table-v1", "status": "confirmed", "fields": []}, ensure_ascii=False),
                encoding="utf-8",
            )
            parse_path.write_text(json.dumps({"status": "completed", "structured": {}}, ensure_ascii=False), encoding="utf-8")
            manifest_path = root / "business_assembly_input.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-BIZ-S4-SKIP",
                        "projectName": "跳过未处理任务项目",
                        "bidType": "商务标",
                        "workDir": str(root),
                        "tocJsonPath": str(toc_path),
                        "businessGapPlanPath": str(gap_path),
                        "projectFactTablePath": str(facts_path),
                        "parseResultPath": str(parse_path),
                        "materialLibraryDir": str(root / "materials"),
                        "outputFile": str(output_path),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            subprocess.run([sys.executable, str(script_path), "--manifest", str(manifest_path), "--response", "summary"], check=True, capture_output=True, text=True)

            output_doc = Document(str(output_path))
            text = "\n".join(paragraph.text for paragraph in output_doc.paragraphs)
            plan = json.loads((root / "business_assembly_plan.json").read_text(encoding="utf-8"))
            review_text = (root / "business_needs_review.md").read_text(encoding="utf-8")

        self.assertIn("4 变流器型式认证", text)
        self.assertIn("5 安全生产许可证", text)
        self.assertNotIn("[待填写：安全生产许可证]", text)
        self.assertNotIn("[待填写：质量管理体系认证证书", text)
        self.assertNotIn("[待填写：2024年财务情况]", text)
        self.assertNotIn("[待填写：新疆新能博乐三台50MW项目", text)
        self.assertEqual(plan["summary"]["assembledCount"], 0)
        self.assertEqual(plan["sections"][0]["status"], "ignored")
        self.assertEqual(plan["sections"][1]["status"], "review_required")
        self.assertIn("正文不写入占位或兜底模板", review_text)
        self.assertIn("任务尚未确认素材", review_text)

    def test_business_assembler_uses_extract_segment_intent_without_merging_whole_docx(self) -> None:
        backend_root = Path(__file__).resolve().parents[1]
        script_path = backend_root / "opencode" / "skills" / "bid-business-assembler" / "scripts" / "run_from_manifest.py"
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            materials = root / "materials"
            materials.mkdir(parents=True)
            source_docx = root / "业绩合同大文件.docx"
            doc = Document()
            doc.add_paragraph("这是不应整份拼入正文的大文件正文")
            doc.save(source_docx)

            toc_path = root / "toc.json"
            gap_path = root / "business_gap_plan.json"
            facts_path = root / "project_fact_table.json"
            parse_path = root / "parse_result.json"
            output_path = root / "商务投标文件.docx"
            toc_path.write_text(
                json.dumps(
                    {
                        "schema_version": "bid-toc-json-v1",
                        "items": [{"number": "五", "title": "同类项目业绩证明", "level": 1}],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            gap_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-gap-plan-v1",
                        "tasks": [
                            {
                                "id": "BTASK-0050",
                                "title": "同类项目业绩证明",
                                "status": "ready",
                                "decision": "ready",
                                "taskType": "bundle",
                                "assemblyMode": "extract_segment",
                                "materialUsage": "extract_segment",
                                "tocTarget": {"title": "同类项目业绩证明"},
                                "selectedEvidenceSegments": [
                                    {
                                        "segmentId": "SEG-PERF-001",
                                        "title": "50MW 风电项目合同业绩片段",
                                        "sourcePages": "第 3 页",
                                        "summary": "证明投标人具备同类风电项目合同业绩。",
                                    }
                                ],
                                "resolvedArtifacts": [
                                    {
                                        "artifactId": "BART-PERF-001",
                                        "artifactType": "selected_material",
                                        "fileName": source_docx.name,
                                        "filePath": str(source_docx),
                                        "sourceMode": "selected_from_business_material_library",
                                        "assemblyMode": "extract_segment",
                                        "materialUsage": "extract_segment",
                                        "evidenceSegmentId": "SEG-PERF-001",
                                        "evidenceSegmentTitle": "50MW 风电项目合同业绩片段",
                                        "evidenceSourcePages": "第 3 页",
                                        "evidenceSummary": "证明投标人具备同类风电项目合同业绩。",
                                    }
                                ],
                            }
                        ],
                        "tocRefs": [],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            facts_path.write_text(
                json.dumps({"schemaVersion": "bid-project-fact-table-v1", "status": "confirmed", "fields": []}, ensure_ascii=False),
                encoding="utf-8",
            )
            parse_path.write_text(json.dumps({"status": "completed", "structured": {}}, ensure_ascii=False), encoding="utf-8")
            manifest_path = root / "business_assembly_input.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-BIZ-S4-SEG",
                        "projectName": "商务 S4 片段项目",
                        "bidType": "商务标",
                        "workDir": str(root),
                        "tocJsonPath": str(toc_path),
                        "businessGapPlanPath": str(gap_path),
                        "projectFactTablePath": str(facts_path),
                        "parseResultPath": str(parse_path),
                        "materialLibraryDir": str(materials),
                        "outputFile": str(output_path),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            subprocess.run(
                [sys.executable, str(script_path), "--manifest", str(manifest_path), "--response", "summary"],
                check=True,
                capture_output=True,
                text=True,
            )

            output_doc = Document(str(output_path))
            text = "\n".join(paragraph.text for paragraph in output_doc.paragraphs)
            attachment_manifest = json.loads((root / "attachment_manifest.json").read_text(encoding="utf-8"))

        self.assertIn("已确认引用以下素材证据片段", text)
        self.assertIn("50MW 风电项目合同业绩片段", text)
        self.assertIn("证明投标人具备同类风电项目合同业绩", text)
        self.assertNotIn("这是不应整份拼入正文的大文件正文", text)
        self.assertEqual(attachment_manifest["items"][0]["mode"], "extract_segment_reference")
        self.assertEqual(attachment_manifest["items"][0]["assemblyMode"], "extract_segment")

    def test_business_assembler_embeds_certificate_images_without_wiki_locator_text(self) -> None:
        backend_root = Path(__file__).resolve().parents[1]
        script_path = backend_root / "opencode" / "skills" / "bid-business-assembler" / "scripts" / "run_from_manifest.py"
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            materials = root / "materials"
            materials.mkdir(parents=True)
            image_path = root / "certificate.png"
            image_path.write_bytes(
                base64.b64decode(
                    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAFgwJ/luzD4gAAAABJRU5ErkJggg=="
                )
            )
            raw_docx = root / "机型认证证书.docx"
            doc = Document()
            doc.add_paragraph("机型认证证书")
            doc.add_picture(str(image_path))
            doc.save(raw_docx)
            cleaned_docx = root / "机型认证证书-清洗稿.docx"
            doc = Document()
            doc.add_paragraph("[embedded_image_1 image1.png]")
            doc.add_paragraph("EW6.25-220型式认证")
            doc.add_paragraph("清洗稿标题片段：EW6.25-220型式认证")
            doc.save(cleaned_docx)

            toc_path = root / "toc.json"
            gap_path = root / "business_gap_plan.json"
            facts_path = root / "project_fact_table.json"
            parse_path = root / "parse_result.json"
            output_path = root / "商务投标文件.docx"
            toc_path.write_text(
                json.dumps({"schema_version": "bid-toc-json-v1", "items": [{"number": "七", "title": "机型认证证书", "level": 1}]}, ensure_ascii=False),
                encoding="utf-8",
            )
            gap_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-gap-plan-v1",
                        "tasks": [
                            {
                                "id": "BTASK-CERT-IMAGE",
                                "title": "机型认证证书",
                                "status": "ready",
                                "decision": "ready",
                                "taskType": "certificate",
                                "assemblyMode": "extract_and_summarize",
                                "materialUsage": "extract_and_summarize",
                                "tocTarget": {"title": "机型认证证书"},
                                "selectedEvidenceSegments": [
                                    {
                                        "segmentId": "SEG-CERT-001",
                                        "title": "EW6.25-220型式认证",
                                        "sourcePages": "清洗稿标题/待页码定位",
                                        "summary": "清洗稿标题片段：EW6.25-220型式认证",
                                    }
                                ],
                                "resolvedArtifacts": [
                                    {
                                        "artifactId": "BART-CERT-001",
                                        "artifactType": "selected_material",
                                        "fileName": cleaned_docx.name,
                                        "filePath": str(cleaned_docx),
                                        "rawFileName": raw_docx.name,
                                        "rawFilePath": str(raw_docx),
                                        "materialName": "机型认证证书.docx",
                                        "folderPath": "商务标/通用素材/专题证书库/机型认证证书",
                                        "sourceMode": "selected_from_business_material_library",
                                        "assemblyMode": "extract_and_summarize",
                                        "materialUsage": "extract_and_summarize",
                                        "evidenceSegmentId": "SEG-CERT-001",
                                        "evidenceSegmentTitle": "EW6.25-220型式认证",
                                        "evidenceSourcePages": "清洗稿标题/待页码定位",
                                        "evidenceSummary": "清洗稿标题片段：EW6.25-220型式认证",
                                        "wikiEvidence": {
                                            "summary": "[embedded_image_1 image1.png]\n清洗稿标题片段：EW6.25-220型式认证",
                                            "sourcePages": "原始文档未分页索引",
                                        },
                                    }
                                ],
                            }
                        ],
                        "tocRefs": [],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            facts_path.write_text(json.dumps({"schemaVersion": "bid-project-fact-table-v1", "status": "confirmed", "fields": []}, ensure_ascii=False), encoding="utf-8")
            parse_path.write_text(json.dumps({"status": "completed", "structured": {}}, ensure_ascii=False), encoding="utf-8")
            manifest_path = root / "business_assembly_input.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-BIZ-S4-CERT-IMAGE",
                        "projectName": "证书图片项目",
                        "bidType": "商务标",
                        "workDir": str(root),
                        "tocJsonPath": str(toc_path),
                        "businessGapPlanPath": str(gap_path),
                        "projectFactTablePath": str(facts_path),
                        "parseResultPath": str(parse_path),
                        "materialLibraryDir": str(materials),
                        "outputFile": str(output_path),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            subprocess.run(
                [sys.executable, str(script_path), "--manifest", str(manifest_path), "--response", "summary"],
                check=True,
                capture_output=True,
                text=True,
            )

            output_doc = Document(str(output_path))
            text = "\n".join(paragraph.text for paragraph in output_doc.paragraphs)
            image_rel_count = len(output_doc.part._package.image_parts)
            attachment_manifest = json.loads((root / "attachment_manifest.json").read_text(encoding="utf-8"))

        self.assertGreaterEqual(image_rel_count, 1)
        self.assertNotIn("支撑材料来源", text)
        self.assertNotIn("根据已确认素材", text)
        self.assertNotIn("清洗稿标题片段", text)
        self.assertNotIn("原始文档未分页索引", text)
        self.assertEqual(attachment_manifest["items"][0]["mode"], "extract_and_summarize")

    def test_business_assembler_extract_summary_prefers_raw_docx_over_cleaned_locator_segments(self) -> None:
        backend_root = Path(__file__).resolve().parents[1]
        script_path = backend_root / "opencode" / "skills" / "bid-business-assembler" / "scripts" / "run_from_manifest.py"
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            materials = root / "materials"
            materials.mkdir(parents=True)
            raw_docx = root / "18.公司生产能力介绍.docx"
            doc = Document()
            doc.add_paragraph("公司生产能力介绍")
            doc.add_paragraph("公司拥有完善的风电设备生产体系，具备叶片、塔筒、机舱等关键环节协同制造能力。")
            doc.add_paragraph("公司配置多条自动化装配产线，能够满足本项目供货周期及质量控制要求。")
            doc.save(raw_docx)
            cleaned_docx = root / "18.公司生产能力介绍-清洗稿.docx"
            doc = Document()
            doc.add_paragraph("18.公司生产能力介绍（清洗稿段落3）")
            doc.add_paragraph("18.公司生产能力介绍（清洗稿段落4）")
            doc.save(cleaned_docx)

            toc_path = root / "toc.json"
            gap_path = root / "business_gap_plan.json"
            facts_path = root / "project_fact_table.json"
            parse_path = root / "parse_result.json"
            output_path = root / "商务投标文件.docx"
            toc_path.write_text(
                json.dumps({"schema_version": "bid-toc-json-v1", "items": [{"number": "四", "title": "生产能力介绍", "level": 1}]}, ensure_ascii=False),
                encoding="utf-8",
            )
            gap_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-gap-plan-v1",
                        "tasks": [
                            {
                                "id": "BTASK-CAPACITY",
                                "title": "生产能力介绍",
                                "status": "ready",
                                "decision": "ready",
                                "taskType": "attachment",
                                "assemblyMode": "extract_and_summarize",
                                "materialUsage": "extract_and_summarize",
                                "tocTarget": {"title": "生产能力介绍"},
                                "selectedEvidenceSegments": [
                                    {
                                        "segmentId": "SEG-CAP-003",
                                        "title": "18.公司生产能力介绍（清洗稿段落3）",
                                        "summary": "18.公司生产能力介绍（清洗稿段落4）",
                                    }
                                ],
                                "resolvedArtifacts": [
                                    {
                                        "artifactId": "BART-CAPACITY",
                                        "artifactType": "selected_material",
                                        "fileName": cleaned_docx.name,
                                        "filePath": str(cleaned_docx),
                                        "rawFileName": raw_docx.name,
                                        "rawFilePath": str(raw_docx),
                                        "cleanedFileName": cleaned_docx.name,
                                        "cleanedFilePath": str(cleaned_docx),
                                        "materialName": raw_docx.name,
                                        "sourceMode": "selected_from_business_material_library",
                                        "assemblyMode": "extract_and_summarize",
                                        "materialUsage": "extract_and_summarize",
                                        "selectedEvidenceSegments": [
                                            {
                                                "segmentId": "SEG-CAP-003",
                                                "title": "18.公司生产能力介绍（清洗稿段落3）",
                                                "summary": "18.公司生产能力介绍（清洗稿段落4）",
                                            }
                                        ],
                                    }
                                ],
                            }
                        ],
                        "tocRefs": [],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            facts_path.write_text(json.dumps({"schemaVersion": "bid-project-fact-table-v1", "status": "confirmed", "fields": []}, ensure_ascii=False), encoding="utf-8")
            parse_path.write_text(json.dumps({"status": "completed", "structured": {}}, ensure_ascii=False), encoding="utf-8")
            manifest_path = root / "business_assembly_input.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-BIZ-S4-CAPACITY",
                        "projectName": "生产能力项目",
                        "bidType": "商务标",
                        "workDir": str(root),
                        "tocJsonPath": str(toc_path),
                        "businessGapPlanPath": str(gap_path),
                        "projectFactTablePath": str(facts_path),
                        "parseResultPath": str(parse_path),
                        "materialLibraryDir": str(materials),
                        "outputFile": str(output_path),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            subprocess.run([sys.executable, str(script_path), "--manifest", str(manifest_path), "--response", "summary"], check=True, capture_output=True, text=True)
            output_doc = Document(str(output_path))
            text = "\n".join(paragraph.text for paragraph in output_doc.paragraphs)

        self.assertIn("风电设备生产体系", text)
        self.assertIn("自动化装配产线", text)
        self.assertNotIn("根据已确认素材", text)
        self.assertNotIn("支撑材料来源", text)
        self.assertNotIn("清洗稿段落3", text)
        self.assertNotIn("清洗稿段落4", text)

    def test_business_assembler_does_not_emit_table_summary_prefix(self) -> None:
        backend_root = Path(__file__).resolve().parents[1]
        script_path = backend_root / "opencode" / "skills" / "bid-business-assembler" / "scripts" / "run_from_manifest.py"
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source_docx = root / "企业获奖情况.docx"
            doc = Document()
            table = doc.add_table(rows=2, cols=3)
            table.rows[0].cells[0].text = "序号"
            table.rows[0].cells[1].text = "奖项名称"
            table.rows[0].cells[2].text = "获奖年度"
            table.rows[1].cells[0].text = "1"
            table.rows[1].cells[1].text = "国家优质工程奖"
            table.rows[1].cells[2].text = "2024"
            doc.save(source_docx)
            toc_path = root / "toc.json"
            gap_path = root / "business_gap_plan.json"
            facts_path = root / "project_fact_table.json"
            parse_path = root / "parse_result.json"
            output_path = root / "商务投标文件.docx"
            toc_path.write_text(
                json.dumps({"schema_version": "bid-toc-json-v1", "items": [{"number": "七", "title": "企业获奖情况", "level": 1}]}, ensure_ascii=False),
                encoding="utf-8",
            )
            gap_path.write_text(
                json.dumps(
                    {
                        "schemaVersion": "bid-business-gap-plan-v1",
                        "tasks": [
                            {
                                "id": "BTASK-AWARD",
                                "title": "企业获奖情况",
                                "status": "ready",
                                "decision": "ready",
                                "taskType": "attachment",
                                "assemblyMode": "extract_and_summarize",
                                "materialUsage": "extract_and_summarize",
                                "tocTarget": {"title": "企业获奖情况"},
                                "resolvedArtifacts": [
                                    {
                                        "artifactId": "BART-AWARD",
                                        "artifactType": "selected_material",
                                        "fileName": source_docx.name,
                                        "filePath": str(source_docx),
                                        "sourceMode": "selected_from_business_material_library",
                                        "assemblyMode": "extract_and_summarize",
                                        "materialUsage": "extract_and_summarize",
                                    }
                                ],
                            }
                        ],
                        "tocRefs": [],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            facts_path.write_text(json.dumps({"schemaVersion": "bid-project-fact-table-v1", "status": "confirmed", "fields": []}, ensure_ascii=False), encoding="utf-8")
            parse_path.write_text(json.dumps({"status": "completed", "structured": {}}, ensure_ascii=False), encoding="utf-8")
            manifest_path = root / "business_assembly_input.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "projectId": "PRJ-BIZ-S4-AWARD",
                        "projectName": "获奖情况项目",
                        "bidType": "商务标",
                        "workDir": str(root),
                        "tocJsonPath": str(toc_path),
                        "businessGapPlanPath": str(gap_path),
                        "projectFactTablePath": str(facts_path),
                        "parseResultPath": str(parse_path),
                        "materialLibraryDir": str(root / "materials"),
                        "outputFile": str(output_path),
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            subprocess.run([sys.executable, str(script_path), "--manifest", str(manifest_path), "--response", "summary"], check=True, capture_output=True, text=True)
            output_doc = Document(str(output_path))
            text = "\n".join(paragraph.text for paragraph in output_doc.paragraphs)

        self.assertIn("国家优质工程奖", text)
        self.assertNotIn("表格信息：", text)


class BusinessAssemblyServiceTests(unittest.TestCase):
    def setUp(self) -> None:
        from fastapi.testclient import TestClient
        from app.core.config import settings
        from app.main import app
        from app.services.store import store

        self.temp_dir = tempfile.TemporaryDirectory()
        base = Path(self.temp_dir.name)
        settings.uploads_dir = base / "uploads"
        settings.documents_dir = base / "documents"
        settings.parsed_dir = base / "parsed"
        settings.ensure_dirs()
        store.reset_for_tests()
        self.client = TestClient(app, base_url="http://127.0.0.1:8000")
        login = self.client.post("/api/auth/login", json={"email": "admin@sewpg.com", "password": "123456"})
        self.assertEqual(login.status_code, 200)
        self.headers = {"Authorization": f"Bearer {login.json()['token']}"}

    def tearDown(self) -> None:
        self.client.close()
        self.temp_dir.cleanup()

    def test_business_material_export_uses_business_material_store(self) -> None:
        from app.services import business_assembly

        async def fake_raw_files(**kwargs):
            self.assertNotIn("bid_type", kwargs)
            self.assertTrue(str(kwargs.get("folder_path") or "").startswith("商务标/"))
            return {"items": [], "total": 0}

        project = {
            "id": "PRJ-BIZ-MATERIAL-SCOPE",
            "name": "商务素材范围项目",
            "bidType": "商务标",
            "customerName": "华能集团",
        }
        with patch("app.services.business_assembly.business_material_store.raw_files", side_effect=fake_raw_files) as raw_files:
            _, exported = business_assembly._export_business_material_library(
                project,
                Path(self.temp_dir.name) / "business-material-export",
            )

        self.assertEqual(exported, 0)
        self.assertGreaterEqual(raw_files.call_count, 1)

    def test_business_fill_generation_uses_business_assembler_without_technical_gap_state(self) -> None:
        from app.core.config import settings
        from app.services.bid_runtime_state import now_iso
        from app.services.store import store
        from app.services.workspace_artifacts import business_workspace_dir, technical_workspace_dir

        project = store.create_project({"name": "商务S4集成项目", "customerName": "华能集团", "bidType": "商务标"})
        project_id = project["id"]
        business_workspace = business_workspace_dir(project_id)
        parse_dir = business_workspace / "parse"
        parse_dir.mkdir(parents=True, exist_ok=True)
        tender_path = settings.uploads_dir / project_id / "tender" / "商务招标文件.docx"
        tender_path.parent.mkdir(parents=True, exist_ok=True)
        tender_path.write_text("dummy", encoding="utf-8")
        scoring_path = business_workspace / "parse" / "商务评分标准.docx"
        doc = Document()
        doc.add_paragraph("商务评分标准正文")
        doc.save(scoring_path)

        _complete_parse_for_tests(
            store,
            project_id,
            tender_files=[{"id": "TEN-1", "name": "商务招标文件.docx", "path": str(tender_path), "size_label": "1 KB"}],
            template_files=[],
            summary={"fileCount": 1, "extractedCount": 1, "textLength": 10, "textPreview": "", "warnings": []},
            parse_storage={"projectDir": str(business_workspace), "combinedTextPath": str(parse_dir / "combined.txt"), "manifestPath": ""},
        )
        _update_parse_result_for_tests(
            store,
            project_id,
            {
                "status": "completed",
                "structured": {
                    "schemaVersion": "bid-business-tender-structured-v1",
                    "businessScoringAsset": {
                        "id": "BIZ-SCORING-001",
                        "fileName": "商务评分标准.docx",
                        "docxPath": str(scoring_path),
                        "reviewStatus": "approved",
                    },
                    "projectFactFields": [
                        {"label": "项目名称", "value": "商务S4集成项目", "category": "项目基础信息", "required": True},
                        {"label": "招标编号", "value": "BIZ-S4-INT", "category": "项目基础信息", "required": True},
                    ],
                },
            },
        )
        _save_generated_outline_for_tests(
            store,
            project_id=project_id,
            nodes=[
                {"id": "OL-1", "title": "投标函", "children": []},
                {"id": "OL-2", "title": "商务评分标准", "children": []},
            ],
            generated_at=now_iso(),
            summary="商务目录已生成。",
        )
        _confirm_outline_for_tests(store, project_id)
        project_record = store._require(project_id)
        business_gap_state = project_record["business_gap_state"]
        business_gap_state.update(
            {
                "recognitionStatus": "completed",
                "recognizedAt": now_iso(),
                "reviewConfirmed": False,
                "plan": {
                    "schemaVersion": "bid-business-gap-plan-v1",
                    "tasks": [
                        {
                            "id": "BTASK-0010",
                            "title": "投标函",
                            "status": "needs_input",
                            "decision": "fill_required",
                            "tocTarget": {"title": "投标函"},
                            "resolvedArtifacts": [],
                        }
                    ],
                    "tocRefs": [],
                },
                "projectFactTable": {},
            }
        )
        store._persist_project(project_record)

        with (
            patch("app.services.business_assembly.OpencodeClient.run_bid_business_assembler_with_trace", side_effect=RuntimeError("offline")),
            patch("app.services.business_assembly.OpencodeClient.run_bid_business_format_cleaner_with_trace", side_effect=RuntimeError("offline")),
        ):
            _start_fill_generation_for_tests(store, project_id)
            from app.services.bid_generation_flow import _run_fill_generation_job

            _run_fill_generation_job(project_id, {}, {"id": "u1", "name": "测试"}, bid_type=BUSINESS_BID_TYPE)

        state = _fill_state_for_tests(store, project_id)
        self.assertEqual(state["status"], "completed")
        self.assertEqual((state.get("assembly") or {}).get("skill"), "bid-business-assembler")
        self.assertIn("business-workspace/s4_assembly_workdir", (state.get("assembly") or {}).get("workDir", ""))
        self.assertFalse((technical_workspace_dir(project_id) / "s7_assembly_workdir").exists())
        self.assertTrue((business_workspace / "s4_assembly_workdir" / "business_assembly_input.json").exists())
        self.assertTrue((business_workspace / "s4_assembly_workdir" / "business_format_clean_input.json").exists())
        self.assertTrue((business_workspace / "s4_assembly_workdir" / "business_format_outline.json").exists())
        format_clean = (state.get("assembly") or {}).get("formatClean") or {}
        self.assertEqual(format_clean.get("status"), "completed")
        self.assertEqual(format_clean.get("skill"), "bid-business-format-cleaner")
        self.assertTrue(Path(str(format_clean.get("outputFile") or "")).exists())
        self.assertTrue(Path(str(format_clean.get("reportFile") or "")).exists())


if __name__ == "__main__":
    unittest.main()
