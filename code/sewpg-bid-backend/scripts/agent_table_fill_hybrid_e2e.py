#!/usr/bin/env python3
"""Hybrid per-table Agent planning + deterministic DOCX execution experiment.

This is an experiment harness, not production S3/S4 code.

Compared with agent_table_fill_e2e.py, the Agent no longer emits final DOCX
coordinates as the main artifact. It emits a semantic fill plan:

- fields: scalar values identified by row/field label
- tableRows: row-keyed matrix rows
- columnFills: repeat a value down rows matching a condition
- copyInstructions: ask the deterministic executor to copy matching source
  table data by row key and column header

The executor handles DOCX geometry, simple merged-cell-safe writes, table row
matching, column matching, unit conversion, and final scoring.
"""

from __future__ import annotations

import argparse
import concurrent.futures
import json
import re
import shutil
import sys
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document

import agent_table_fill_e2e as base


DEFAULT_OUTPUT_DIR = Path(
    "/data/documents/PRJ-0003/technical-workspace/s4_gap_workdir/"
    "agent_table_hybrid_e2e_v1"
)
PLAN_SCHEMA = "agent-table-fill-plan-v1"


@dataclass(frozen=True)
class SourceTable:
    source_id: str
    name: str
    path: str
    table_name: str
    rows: list[list[str]]


def tokenize(value: Any) -> list[str]:
    text = base.norm_text(value)
    tokens = re.findall(r"[A-Za-z][A-Za-z0-9_.-]*|\d+(?:\.\d+)?|[\u4e00-\u9fff]{2,}", text)
    return [token for token in tokens if len(token) >= 1]


def token_score(left: Any, right: Any) -> float:
    left_tokens = set(tokenize(left))
    right_tokens = set(tokenize(right))
    if not left_tokens or not right_tokens:
        return 0.0
    overlap = len(left_tokens & right_tokens)
    return overlap / max(1, min(len(left_tokens), len(right_tokens)))


def compact_cell(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def header_text(table: Any, row_idx: int, col_idx: int) -> str:
    parts: list[str] = []
    for r in range(0, min(row_idx + 1, 4)):
        if r < len(table.rows) and col_idx < len(table.rows[r].cells):
            text = compact_cell(table.rows[r].cells[col_idx].text)
            if text:
                parts.append(text)
    return " / ".join(parts)


def row_text(row: Any, upto_col: int | None = None) -> str:
    cells = row.cells if upto_col is None else row.cells[: max(0, upto_col)]
    return " / ".join(compact_cell(cell.text) for cell in cells if compact_cell(cell.text))


def is_empty_cell(cell: Any) -> bool:
    return not base.norm_text(cell.text)


def writable_cols(row: Any, header_row: Any | None = None) -> list[int]:
    result: list[int] = []
    for col_idx, cell in enumerate(row.cells):
        if not is_empty_cell(cell):
            continue
        head = compact_cell(header_row.cells[col_idx].text) if header_row is not None and col_idx < len(header_row.cells) else ""
        left = row_text(row, col_idx)
        if col_idx >= 1 and (head or left):
            result.append(col_idx)
    return result


def best_row_for_label(doc: Document, label: str) -> tuple[int, int, float]:
    best = (-1, -1, 0.0)
    label_norm = base.norm_text(label)
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            text = row_text(row)
            score = token_score(label, text)
            if label_norm and label_norm in base.norm_text(text):
                score = max(score, 0.95)
            if score > best[2]:
                best = (table_idx, row_idx, score)
    return best


def best_col_for_label(table: Any, row_idx: int, label: str, role: str = "") -> tuple[int, float]:
    row = table.rows[row_idx]
    header = table.rows[0] if table.rows else None
    role_norm = base.norm_text(role or label)
    blanks = writable_cols(row, header)
    if not blanks:
        return (-1, 0.0)
    if role_norm in {"all", "allresponse", "all_response", "全部响应列", "通用承诺值"}:
        return (blanks[0], 1.0)
    best = (blanks[0], 0.01)
    for col_idx in blanks:
        head = header_text(table, row_idx, col_idx)
        score = token_score(label + " " + role, head)
        if "投标机型1" in head or "投标" in head or "响应" in head:
            score += 0.25
        if score > best[1]:
            best = (col_idx, score)
    return best


def normalize_number_text(value: Any) -> str:
    text = compact_cell(value)
    if re.fullmatch(r"[-+]?\d+\.0", text):
        return text[:-2]
    return text


def valid_fill_value(value: Any) -> bool:
    text = base.norm_text(value)
    if not text:
        return False
    return text not in {"待填写", "无依据", "暂无", "none", "null"}


def convert_value_for_target(value: Any, source_context: str, target_context: str) -> str:
    text = normalize_number_text(value)
    number = base.parse_simple_number(text)
    if number is None:
        return text
    source_norm = base.norm_text(source_context)
    target_norm = base.norm_text(target_context)
    converted = number
    if "kg" in source_norm.lower() and ("(t)" in target_context.lower() or "（t" in target_context or "吨" in target_context):
        converted = number / 1000.0
    elif "mwh" in source_norm.lower() and ("万kwh" in target_norm.lower() or "万千瓦时" in target_norm):
        converted = number / 10.0
    elif "kwh" in source_norm.lower() and "万kwh" in target_norm.lower():
        converted = number / 10000.0
    if abs(converted - round(converted)) < 1e-9:
        return str(int(round(converted)))
    return f"{converted:.6f}".rstrip("0").rstrip(".")


def apply_scalar_field(doc: Document, field: dict[str, Any], edits: list[dict[str, Any]]) -> int:
    label = str(field.get("target") or field.get("label") or field.get("rowLabel") or "").strip()
    value = str(field.get("value") or "").strip()
    if not label or not valid_fill_value(value):
        return 0
    table_idx, row_idx, row_score = best_row_for_label(doc, label)
    if table_idx < 0 or row_score < 0.28:
        return 0
    table = doc.tables[table_idx]
    column_role = str(field.get("targetColumn") or field.get("column") or field.get("columnRole") or "")
    col_idx, col_score = best_col_for_label(table, row_idx, label, column_role)
    if col_idx < 0:
        return 0
    cols = [col_idx]
    if base.norm_text(column_role) in {"all", "allresponse", "all_response", "全部响应列", "通用承诺值"}:
        cols = writable_cols(table.rows[row_idx], table.rows[0] if table.rows else None)
    applied = 0
    for col in cols:
        cell = table.rows[row_idx].cells[col]
        if not is_empty_cell(cell):
            continue
        base.set_cell_text(cell, value)
        edits.append(
            {
                "kind": "field",
                "table": table_idx,
                "row": row_idx,
                "col": col,
                "text": value,
                "label": label,
                "rowScore": round(row_score, 3),
                "colScore": round(col_score, 3),
                "evidence": field.get("evidence"),
            }
        )
        applied += 1
    return applied


def source_tables_for_sources(sources: list[dict[str, str]]) -> list[SourceTable]:
    tables: list[SourceTable] = []
    for source in sources:
        path = Path(source.get("path") or "")
        if not path.exists():
            continue
        source_id = source.get("id") or ""
        name = source.get("name") or path.name
        suffix = path.suffix.lower()
        try:
            if suffix == ".docx":
                doc = Document(str(path))
                for table_idx, table in enumerate(doc.tables[:12]):
                    rows = [[compact_cell(cell.text) for cell in row.cells[:40]] for row in table.rows[:700]]
                    if rows:
                        tables.append(SourceTable(source_id, name, str(path), f"docx_table_{table_idx}", rows))
            elif suffix in {".xlsx", ".xlsm"}:
                import openpyxl

                workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
                for sheet_name in workbook.sheetnames[:12]:
                    sheet = workbook[sheet_name]
                    rows: list[list[str]] = []
                    for row_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                        values = [compact_cell(value) for value in row[:40]]
                        if any(base.norm_text(value) for value in values):
                            rows.append(values)
                        if row_idx >= 1200:
                            break
                    if rows:
                        tables.append(SourceTable(source_id, name, str(path), sheet_name, rows))
                workbook.close()
        except Exception:
            continue
    return tables


def match_source_by_hint(source_tables: list[SourceTable], hint: str) -> list[SourceTable]:
    if not hint:
        return source_tables
    hint_norm = base.norm_text(hint)
    matched = [
        table
        for table in source_tables
        if hint_norm in base.norm_text(table.source_id + table.name + table.table_name + table.path)
        or token_score(hint, table.source_id + " " + table.name + " " + table.table_name) >= 0.35
    ]
    return matched or source_tables


def source_header(rows: list[list[str]], col_idx: int, row_idx: int) -> str:
    parts: list[str] = []
    for r in range(0, min(row_idx + 1, 4)):
        if r < len(rows) and col_idx < len(rows[r]):
            text = rows[r][col_idx]
            if text:
                parts.append(text)
    return " / ".join(parts)


def find_source_row(rows: list[list[str]], key: str) -> int:
    key_norm = base.norm_text(key)
    if not key_norm:
        return -1
    for row_idx, row in enumerate(rows):
        first_cells = " ".join(row[:5])
        if key_norm in base.norm_text(first_cells):
            return row_idx
    best = (-1, 0.0)
    for row_idx, row in enumerate(rows):
        score = token_score(key, " ".join(row[:8]))
        if score > best[1]:
            best = (row_idx, score)
    return best[0] if best[1] >= 0.7 else -1


def find_source_col(rows: list[list[str]], row_idx: int, target_header: str, used_cols: set[int]) -> int:
    if row_idx < 0 or row_idx >= len(rows):
        return -1
    best = (-1, 0.0)
    for col_idx, value in enumerate(rows[row_idx]):
        if col_idx in used_cols or not base.norm_text(value):
            continue
        context = source_header(rows, col_idx, row_idx)
        score = token_score(target_header, context)
        if "万kwh" in base.norm_text(target_header).lower() and "mwh" in base.norm_text(context).lower():
            score += 0.35
        if score > best[1]:
            best = (col_idx, score)
    if best[1] >= 0.15:
        return best[0]
    # Fallback to the next non-empty numeric/text value after the likely key columns.
    for col_idx, value in enumerate(rows[row_idx]):
        if col_idx not in used_cols and col_idx >= 1 and base.norm_text(value):
            return col_idx
    return -1


def likely_row_key(row: Any, col_idx: int | None = None) -> str:
    cells = row.cells if col_idx is None else row.cells[: max(1, col_idx)]
    texts = [compact_cell(cell.text) for cell in cells if compact_cell(cell.text)]
    for text in texts:
        if re.fullmatch(r"[A-Za-z]?\d+(?:\.\d+)?(?:[-~至]\d+(?:\.\d+)?)?", text):
            return text
    return texts[-1] if texts else ""


def apply_table_rows(doc: Document, table_rows: list[dict[str, Any]], edits: list[dict[str, Any]]) -> int:
    applied = 0
    for item in table_rows:
        if not isinstance(item, dict):
            continue
        row_key = str(item.get("rowKey") or item.get("key") or "").strip()
        values = item.get("values") if isinstance(item.get("values"), dict) else {}
        if not row_key or not values:
            continue
        best = (-1, -1, 0.0)
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                score = token_score(row_key, row_text(row))
                if base.norm_text(row_key) in base.norm_text(row_text(row)):
                    score = max(score, 0.95)
                if score > best[2]:
                    best = (table_idx, row_idx, score)
        if best[0] < 0 or best[2] < 0.45:
            continue
        table = doc.tables[best[0]]
        used_cols: set[int] = set()
        for col_label, raw_value in values.items():
            value = str(raw_value or "").strip()
            if not valid_fill_value(value):
                continue
            col_idx, _ = best_col_for_label(table, best[1], str(col_label), str(col_label))
            if col_idx in used_cols:
                continue
            if col_idx < 0:
                continue
            cell = table.rows[best[1]].cells[col_idx]
            if not is_empty_cell(cell):
                continue
            value = convert_value_for_target(value, str(col_label), row_text(table.rows[best[1]]) + " " + header_text(table, best[1], col_idx))
            base.set_cell_text(cell, value)
            used_cols.add(col_idx)
            edits.append({"kind": "tableRow", "table": best[0], "row": best[1], "col": col_idx, "text": value, "rowKey": row_key, "column": col_label})
            applied += 1
    return applied


def apply_column_fills(doc: Document, column_fills: list[dict[str, Any]], edits: list[dict[str, Any]]) -> int:
    applied = 0
    for item in column_fills:
        if not isinstance(item, dict):
            continue
        value = str(item.get("value") or "").strip()
        column = str(item.get("column") or item.get("targetColumn") or "")
        row_filter = str(item.get("rowFilter") or item.get("targetRows") or "")
        if not valid_fill_value(value):
            continue
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows[1:], start=1):
                text = row_text(row)
                if row_filter and token_score(row_filter, text) < 0.2 and base.norm_text(row_filter) not in base.norm_text(text):
                    continue
                col_idx, _ = best_col_for_label(table, row_idx, column or row_filter, column)
                if col_idx < 0:
                    continue
                cell = row.cells[col_idx]
                if not is_empty_cell(cell):
                    continue
                base.set_cell_text(cell, value)
                edits.append({"kind": "columnFill", "table": table_idx, "row": row_idx, "col": col_idx, "text": value, "column": column, "rowFilter": row_filter})
                applied += 1
    return applied


def apply_copy_instruction(
    doc: Document,
    instruction: dict[str, Any],
    source_tables: list[SourceTable],
    edits: list[dict[str, Any]],
) -> int:
    hints = " ".join(str(x) for x in instruction.get("targetHints") or [])
    source_hint = str(instruction.get("sourceHint") or instruction.get("source") or "")
    matched_sources = match_source_by_hint(source_tables, source_hint)
    applied = 0
    for table_idx, dest_table in enumerate(doc.tables):
        table_hint_score = token_score(hints, " ".join(row_text(row) for row in dest_table.rows[:4]))
        if hints and table_hint_score < 0.1:
            continue
        for row_idx, row in enumerate(dest_table.rows):
            blanks = writable_cols(row, dest_table.rows[0] if dest_table.rows else None)
            if not blanks:
                continue
            row_key = likely_row_key(row)
            if not row_key:
                continue
            best_source: tuple[SourceTable | None, int, float] = (None, -1, 0.0)
            for src in matched_sources:
                src_row_idx = find_source_row(src.rows, row_key)
                if src_row_idx < 0:
                    continue
                score = token_score(row_key, " ".join(src.rows[src_row_idx][:8]))
                if base.norm_text(row_key) in base.norm_text(" ".join(src.rows[src_row_idx][:8])):
                    score = max(score, 0.95)
                if score > best_source[2]:
                    best_source = (src, src_row_idx, score)
            src, src_row_idx, src_score = best_source
            if src is None or src_score < 0.45:
                continue
            used_source_cols: set[int] = set()
            for dest_col in blanks:
                target_context = row_text(row, dest_col) + " " + header_text(dest_table, row_idx, dest_col)
                source_col = find_source_col(src.rows, src_row_idx, target_context, used_source_cols)
                if source_col < 0:
                    continue
                raw_value = src.rows[src_row_idx][source_col]
                if not valid_fill_value(raw_value):
                    continue
                source_context = source_header(src.rows, source_col, src_row_idx)
                value = convert_value_for_target(raw_value, source_context, target_context)
                base.set_cell_text(row.cells[dest_col], value)
                used_source_cols.add(source_col)
                edits.append(
                    {
                        "kind": "copyInstruction",
                        "table": table_idx,
                        "row": row_idx,
                        "col": dest_col,
                        "text": value,
                        "rowKey": row_key,
                        "source": f"{src.source_id} {src.table_name} R{src_row_idx}C{source_col}",
                    }
                )
                applied += 1
    return applied


def apply_hybrid_plan(
    blank_docx: Path,
    output_path: Path,
    plan: dict[str, Any],
    sources: list[dict[str, str]],
) -> tuple[int, list[dict[str, Any]]]:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(blank_docx, output_path)
    doc = Document(str(output_path))
    edits: list[dict[str, Any]] = []

    for field in plan.get("fields") or []:
        if isinstance(field, dict):
            apply_scalar_field(doc, field, edits)
    table_rows = plan.get("tableRows") if isinstance(plan.get("tableRows"), list) else []
    apply_table_rows(doc, table_rows, edits)
    column_fills = plan.get("columnFills") if isinstance(plan.get("columnFills"), list) else []
    apply_column_fills(doc, column_fills, edits)

    source_tables = source_tables_for_sources(sources)
    for instruction in plan.get("copyInstructions") or []:
        if isinstance(instruction, dict):
            apply_copy_instruction(doc, instruction, source_tables, edits)

    # Backward-compatible escape hatch for the model if it can identify exact cells.
    for edit in plan.get("edits") or []:
        if not isinstance(edit, dict):
            continue
        try:
            table_idx = int(edit.get("table", 0))
            row_idx = int(edit.get("row"))
            col_idx = int(edit.get("col"))
        except Exception:
            continue
        value = str(edit.get("text") or "").strip()
        if not valid_fill_value(value) or table_idx >= len(doc.tables) or row_idx >= len(doc.tables[table_idx].rows):
            continue
        row = doc.tables[table_idx].rows[row_idx]
        if col_idx < 0 or col_idx >= len(row.cells) or not is_empty_cell(row.cells[col_idx]):
            continue
        base.set_cell_text(row.cells[col_idx], value)
        edits.append({"kind": "exactFallback", "table": table_idx, "row": row_idx, "col": col_idx, "text": value, "evidence": edit.get("evidence")})

    doc.save(str(output_path))
    return len(edits), edits


def build_hybrid_prompt(
    manifest: dict[str, Any],
    parse_items: list[dict[str, str]],
    target: base.Target,
    output_path: Path,
    report_path: Path,
    workdir: Path,
    sources: list[dict[str, str]],
    table_preview: str,
    source_context: str,
    material_index_dir: Path,
) -> str:
    project_context = {
        "projectId": manifest.get("projectId"),
        "projectName": manifest.get("projectName"),
        "projectTurbineModel": manifest.get("projectTurbineModel"),
        "selectedTenderParseItems": parse_items,
    }
    return f"""
你是“单表独立语义规划 Agent”。每张表都是全新会话。不要调用工具，不要调用 Bash，不要读写文件，不要使用人工填写文件。

你的任务不是给最终 Word 坐标，而是输出一个语义填表计划；确定性执行器会负责 DOCX 表格坐标、合并单元格、矩阵搬运、单位换算和写入。

本表：
- appendixId: {target.appendix_id}
- title: {target.title}
- blankDocx: {target.docx_path}
- outputDocx: {output_path}
- reportJson: {report_path}
- workdir: {workdir}
- materialIndexDir: {material_index_dir}

项目上下文：
```json
{json.dumps(project_context, ensure_ascii=False, indent=2)}
```

候选来源：
```json
{json.dumps(sources, ensure_ascii=False, indent=2)}
```

空表预览：
```
{table_preview}
```

候选来源摘录：
```
{source_context}
```

请输出严格 JSON，schema 如下：
{{
  "schema_version": "{PLAN_SCHEMA}",
  "appendixId": "{target.appendix_id}",
  "title": "{target.title}",
  "fields": [
    {{"target": "表内行名或字段名，如 投标机型", "value": "要填写的值", "targetColumn": "投标机型1/投标响应/全部响应列/auto", "evidence": "来源和推导说明", "confidence": 0.9}}
  ],
  "tableRows": [
    {{"rowKey": "表内行键，如 5.2 或 A001", "values": {{"目标列语义，如 风电场保证年上网电量（万kWh）": "值"}}, "evidence": "来源表行"}}
  ],
  "columnFills": [
    {{"rowFilter": "匹配哪些行，可为空", "column": "目标列语义", "value": "重复填入值", "evidence": "来源"}}
  ],
  "copyInstructions": [
    {{"sourceHint": "来源材料ID/文件名/表名", "targetHints": ["目标表关键词"], "mode": "match_row_key_and_column_header", "unitRules": ["MWh->万kWh /10", "kg->t /1000"]}}
  ],
  "edits": [],
  "notes": [],
  "unfilledHints": []
}}

规划原则：
1. 只输出 JSON，不要 Markdown，不要解释文字，不要工具调用。
2. 优先使用 fields 表达标量；大表/曲线/矩阵优先用 copyInstructions 或 tableRows，不要手写几百个坐标。
3. 不要改编号、项目名、备注、表头；只规划投标人响应值。
4. 数字保留来源精度；kg 到 t 除以 1000，MWh 到 万kWh 除以 10，kWh 到 万kWh 除以 10000。
5. 按风速索引的承诺矩阵要用项目指定测风塔年平均风速选行；不要把机位坐标表当成电量表。
6. “投标机型”优先用正式机型名，不带“上置”；布局信息用于箱变位置。
7. 如果某附表标题或正文明确为“无”，可以输出空计划，并在 notes 说明。
""".strip()


def parse_plan_or_repair(base_url: str, session_id: str, raw_text: str, timeout: int) -> dict[str, Any]:
    try:
        return base.parse_json_from_text(raw_text)
    except Exception:
        repair_prompt = f"""
把下面内容修复为严格 JSON。只输出 JSON，不要 Markdown，不要解释。必须保留 schema_version、appendixId、fields/tableRows/columnFills/copyInstructions/edits/notes/unfilledHints 等语义，不要新增事实。

原始内容：
{raw_text[:12000]}
""".strip()
        response = base.send_prompt(base_url, session_id, repair_prompt, timeout=min(timeout, 180))
        return base.parse_json_from_text(base.assistant_text(response))


def run_one(
    base_url: str,
    target: base.Target,
    manifest: dict[str, Any],
    output_dir: Path,
    selected_by_appendix: dict[str, list[dict[str, Any]]],
    catalog: list[dict[str, Any]],
    material_index_dir: Path,
    timeout: int,
) -> dict[str, Any]:
    started = time.time()
    target_dir = output_dir / f"{target.index:03d}-{target.appendix_id}"
    target_dir.mkdir(parents=True, exist_ok=True)
    output_path = target_dir / f"{target.index:03d}-{target.appendix_id}-{base.slug(target.title)}_Hybrid填写.docx"
    report_path = target_dir / "hybrid_fill_report.json"
    trace_path = target_dir / "opencode_response.json"
    prompt_path = target_dir / "prompt.txt"

    sources = base.compact_sources(target, selected_by_appendix, catalog)
    table_preview = base.compact_table_preview(target.docx_path)
    tokens = base.query_tokens(target, table_preview, manifest)
    parse_result_path = target.docx_path.parents[1] / "s4_gap_workdir" / "parse_result.json"
    parse_items = base.select_parse_items(parse_result_path, tokens)
    source_context = base.source_context_for_agent(sources, tokens)
    prompt = build_hybrid_prompt(
        manifest,
        parse_items,
        target,
        output_path,
        report_path,
        target_dir,
        sources,
        table_preview,
        source_context,
        material_index_dir,
    )
    prompt_path.write_text(prompt, encoding="utf-8")

    session_id = ""
    try:
        session_id = base.create_session(base_url, f"Hybrid填表 {target.appendix_id} {target.title}", timeout=30)
        response = base.send_prompt(base_url, session_id, prompt, timeout=timeout)
        trace_path.write_text(json.dumps(response, ensure_ascii=False, indent=2), encoding="utf-8")
        plan = parse_plan_or_repair(base_url, session_id, base.assistant_text(response), timeout)
        applied_count, applied_edits = apply_hybrid_plan(target.docx_path, output_path, plan, sources)
        report = {
            "schema_version": "agent-table-hybrid-fill-report-v1",
            "appendixId": target.appendix_id,
            "title": target.title,
            "blankDocx": str(target.docx_path),
            "outputFile": str(output_path),
            "sources": sources,
            "plan": plan,
            "appliedEdits": applied_edits,
            "appliedEditCount": applied_count,
        }
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        ok = output_path.exists() and output_path.stat().st_size > 0
        return {
            "appendixId": target.appendix_id,
            "title": target.title,
            "index": target.index,
            "ok": ok,
            "sessionId": session_id,
            "outputFile": str(output_path),
            "reportFile": str(report_path),
            "appliedEditCount": applied_count,
            "durationSec": round(time.time() - started, 2),
        }
    except Exception as exc:
        fallback = target_dir / f"{target.index:03d}-{target.appendix_id}-{base.slug(target.title)}_Hybrid失败原表.docx"
        if target.docx_path.exists() and not fallback.exists():
            shutil.copyfile(target.docx_path, fallback)
        return {
            "appendixId": target.appendix_id,
            "title": target.title,
            "index": target.index,
            "ok": False,
            "sessionId": session_id,
            "outputFile": str(fallback),
            "reportFile": str(report_path),
            "error": str(exc),
            "durationSec": round(time.time() - started, 2),
        }


def write_markdown_report(path: Path, run_payload: dict[str, Any], score: dict[str, Any]) -> None:
    base.write_markdown_report(path, run_payload, score)


def main(argv: list[str]) -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--manifest", type=Path, default=base.DEFAULT_MANIFEST)
    parser.add_argument("--source-selection-report", type=Path, default=base.DEFAULT_SOURCE_SELECTION_REPORT)
    parser.add_argument("--human-docx", type=Path, default=base.DEFAULT_HUMAN_DOCX)
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR)
    parser.add_argument("--opencode-url", default=base.DEFAULT_OPENCODE_URL)
    parser.add_argument("--limit", type=int, default=0)
    parser.add_argument("--ids", default="", help="comma-separated appendix ids")
    parser.add_argument("--concurrency", type=int, default=3)
    parser.add_argument("--timeout", type=int, default=420)
    args = parser.parse_args(argv)

    manifest, targets = base.load_targets(args.manifest)
    if args.ids.strip():
        wanted = {item.strip() for item in args.ids.split(",") if item.strip()}
        targets = [target for target in targets if target.appendix_id in wanted]
    if args.limit > 0:
        targets = targets[: args.limit]
    if not targets:
        raise SystemExit("no targets selected")

    material_index_dir = args.manifest.parent.parent / "material_index"
    selected_by_appendix = base.load_source_selection(args.source_selection_report)
    catalog = base.material_catalog(manifest)
    args.output_dir.mkdir(parents=True, exist_ok=True)

    results: list[dict[str, Any]] = []
    print(
        json.dumps(
            {
                "event": "start",
                "mode": "hybrid",
                "targets": len(targets),
                "outputDir": str(args.output_dir),
                "concurrency": args.concurrency,
            },
            ensure_ascii=False,
        ),
        flush=True,
    )
    with concurrent.futures.ThreadPoolExecutor(max_workers=max(1, args.concurrency)) as pool:
        future_map = {
            pool.submit(
                run_one,
                args.opencode_url,
                target,
                manifest,
                args.output_dir,
                selected_by_appendix,
                catalog,
                material_index_dir,
                args.timeout,
            ): target
            for target in targets
        }
        for future in concurrent.futures.as_completed(future_map):
            target = future_map[future]
            result = future.result()
            result["blankDocx"] = str(target.docx_path)
            results.append(result)
            print(
                json.dumps(
                    {
                        "event": "target_done",
                        "appendixId": result.get("appendixId"),
                        "ok": result.get("ok"),
                        "appliedEditCount": result.get("appliedEditCount"),
                        "durationSec": result.get("durationSec"),
                        "error": result.get("error"),
                    },
                    ensure_ascii=False,
                ),
                flush=True,
            )

    results.sort(key=lambda item: int(item.get("index") or 0))
    run_payload = {
        "schema_version": "agent-table-hybrid-e2e-run-v1",
        "manifest": str(args.manifest),
        "outputDir": str(args.output_dir),
        "summary": {
            "targetCount": len(results),
            "successCount": sum(1 for item in results if item.get("ok")),
            "failedCount": sum(1 for item in results if not item.get("ok")),
            "appliedEditCount": sum(int(item.get("appliedEditCount") or 0) for item in results),
        },
        "results": results,
    }
    run_json = args.output_dir / "hybrid_e2e_run.json"
    run_json.write_text(json.dumps(run_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    score = base.score_outputs(results, args.human_docx)
    score_json = args.output_dir / "human_comparison_score.json"
    score_json.write_text(json.dumps(score, ensure_ascii=False, indent=2), encoding="utf-8")
    write_markdown_report(args.output_dir / "human_comparison_report.md", run_payload, score)

    print(
        json.dumps(
            {
                "schema_version": "agent-table-hybrid-e2e-summary-v1",
                "outputDir": str(args.output_dir),
                "runFile": str(run_json),
                "scoreFile": str(score_json),
                "reportFile": str(args.output_dir / "human_comparison_report.md"),
                "summary": run_payload["summary"],
                "scoreSummary": score.get("summary") or {"error": score.get("error")},
            },
            ensure_ascii=False,
            indent=2,
        ),
        flush=True,
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
