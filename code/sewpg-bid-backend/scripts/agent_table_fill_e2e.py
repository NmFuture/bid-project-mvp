#!/usr/bin/env python3
"""Run a per-appendix Agent table filling experiment.

This script is intentionally outside the production S3/S4 skill flow. It
creates one OpenCode session per appendix, provides a generic extraction of the
blank table and candidate source snippets, asks the session to produce a JSON
cell patch, and applies that patch to the DOCX. The harness does not encode
appendix-specific filling rules; it only dispatches, patches, collects, and
scores the results.

Expected runtime location: the opencode container, where /data is mounted and
the OpenCode server is reachable at http://127.0.0.1:4096.
"""

from __future__ import annotations

import argparse
import concurrent.futures
import json
import re
import shutil
import sys
import time
import urllib.error
import urllib.request
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document


DEFAULT_MANIFEST = Path(
    "/data/documents/PRJ-0003/technical-workspace/s4_gap_workdir/"
    "table_filler_e2e_all/appendix_only_e2e/appendix_only_manifest.json"
)
DEFAULT_SOURCE_SELECTION_REPORT = Path(
    "/data/documents/PRJ-0003/technical-workspace/s4_gap_workdir/"
    "table_filler_e2e_all/fixed_probe_outputs_v4/batch_fill_report.json"
)
DEFAULT_OUTPUT_DIR = Path(
    "/data/documents/PRJ-0003/technical-workspace/s4_gap_workdir/"
    "agent_table_e2e_v1"
)
DEFAULT_HUMAN_DOCX = Path("/tmp/human_appendices.docx")
DEFAULT_OPENCODE_URL = "http://127.0.0.1:4096"
MODEL = {"providerID": "opencode", "modelID": "big-pickle"}
MAX_SOURCE_CHARS = 44_000


@dataclass(frozen=True)
class Target:
    index: int
    appendix_id: str
    title: str
    docx_path: Path
    gap_id: str
    number: str
    recommended_materials: list[dict[str, Any]]


def norm_text(value: Any) -> str:
    text = str(value or "")
    text = text.replace("\u3000", " ").replace("\xa0", " ")
    text = re.sub(r"\s+", "", text)
    text = text.replace("－", "-").replace("—", "-").replace("–", "-")
    text = text.replace("：", ":").replace("（", "(").replace("）", ")")
    text = text.replace("，", ",").replace("。", ".")
    return text.strip()


def display_text(value: Any, limit: int = 120) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    if len(text) <= limit:
        return text
    return text[: limit - 3] + "..."


def slug(value: str) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|\r\n\t]+", "_", value).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned[:90] or "appendix"


def load_json(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as handle:
        payload = json.load(handle)
    if not isinstance(payload, dict):
        raise RuntimeError(f"JSON root must be an object: {path}")
    return payload


def post_json(url: str, payload: dict[str, Any], timeout: int) -> dict[str, Any]:
    data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    request = urllib.request.Request(
        url,
        data=data,
        headers={"content-type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            raw = response.read().decode("utf-8", errors="replace")
    except urllib.error.HTTPError as exc:
        raw = exc.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"HTTP {exc.code} from {url}: {raw[:800]}") from exc
    if not raw.strip():
        raise RuntimeError(f"empty response from {url}")
    try:
        return json.loads(raw)
    except json.JSONDecodeError as exc:
        raise RuntimeError(f"non-JSON response from {url}: {raw[:800]}") from exc


def create_session(base_url: str, title: str, timeout: int) -> str:
    payload = post_json(f"{base_url.rstrip('/')}/session", {"title": title}, timeout)
    session_id = str(payload.get("id") or "").strip()
    if not session_id:
        raise RuntimeError(f"OpenCode did not return a session id: {payload}")
    return session_id


def send_prompt(base_url: str, session_id: str, prompt: str, timeout: int) -> dict[str, Any]:
    return post_json(
        f"{base_url.rstrip('/')}/session/{session_id}/message",
        {
            "model": MODEL,
            "parts": [{"type": "text", "text": prompt}],
        },
        timeout,
    )


def assistant_text(response: dict[str, Any]) -> str:
    parts = response.get("parts") if isinstance(response, dict) else []
    text_parts = [
        str(part.get("text") or "")
        for part in parts or []
        if isinstance(part, dict) and part.get("type") == "text"
    ]
    return "\n".join(part for part in text_parts if part).strip()


def parse_json_from_text(text: str) -> dict[str, Any]:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```[a-zA-Z0-9_-]*\n?", "", cleaned)
        cleaned = re.sub(r"\n?```$", "", cleaned).strip()
    try:
        payload = json.loads(cleaned)
    except json.JSONDecodeError:
        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start == -1 or end <= start:
            raise
        payload = json.loads(cleaned[start : end + 1])
    if not isinstance(payload, dict):
        raise RuntimeError("assistant JSON root is not an object")
    return payload


def load_targets(manifest_path: Path) -> tuple[dict[str, Any], list[Target]]:
    manifest = load_json(manifest_path)
    targets: list[Target] = []
    for idx, raw in enumerate(manifest.get("targets") or [], start=1):
        if not isinstance(raw, dict):
            continue
        docx_path = Path(str(raw.get("docxPath") or ""))
        targets.append(
            Target(
                index=idx,
                appendix_id=str(raw.get("id") or f"APPX-{idx:04d}"),
                title=str(raw.get("title") or docx_path.stem),
                docx_path=docx_path,
                gap_id=str(raw.get("gapId") or ""),
                number=str(raw.get("number") or ""),
                recommended_materials=[
                    item for item in raw.get("recommendedMaterials") or [] if isinstance(item, dict)
                ],
            )
        )
    return manifest, targets


def load_source_selection(path: Path) -> dict[str, list[dict[str, Any]]]:
    if not path.exists():
        return {}
    payload = load_json(path)
    result: dict[str, list[dict[str, Any]]] = {}
    report = payload.get("fillReport") or {}
    for item in report.get("sourceSelections") or []:
        if not isinstance(item, dict):
            continue
        appendix_id = str(item.get("appendixId") or "")
        selected = [entry for entry in item.get("selected") or [] if isinstance(entry, dict)]
        if appendix_id and selected:
            result[appendix_id] = selected
    return result


def material_catalog(manifest: dict[str, Any]) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    for raw in manifest.get("materialIndex") or []:
        if not isinstance(raw, dict):
            continue
        items.append(raw)
    return items


def compact_table_preview(path: Path, max_rows: int = 80, max_cols: int = 12) -> str:
    if not path.exists():
        return f"[missing] {path}"
    doc = Document(str(path))
    if not doc.tables:
        paragraphs = [display_text(p.text, 180) for p in doc.paragraphs if norm_text(p.text)]
        return "无表格。正文摘录：\n" + "\n".join(paragraphs[:20])
    chunks: list[str] = []
    for table_idx, table in enumerate(doc.tables[:3]):
        rows = table.rows[:max_rows]
        chunks.append(f"table[{table_idx}] rows={len(table.rows)} cols={len(table.columns)}")
        for row_idx, row in enumerate(rows):
            values = [display_text(cell.text, 70) for cell in row.cells[:max_cols]]
            chunks.append(f"R{row_idx}: " + " | ".join(values))
        if len(table.rows) > max_rows:
            chunks.append(f"... ({len(table.rows) - max_rows} more rows)")
    return "\n".join(chunks)


def query_tokens(target: Target, table_preview: str, manifest: dict[str, Any]) -> list[str]:
    raw_parts = [
        target.title,
        target.number,
        table_preview,
        json.dumps(manifest.get("projectTurbineModel") or {}, ensure_ascii=False),
        str(manifest.get("projectName") or ""),
    ]
    text = "\n".join(raw_parts)
    tokens: list[str] = []
    for token in re.findall(r"[A-Za-z][A-Za-z0-9_.-]{1,}|[0-9]+(?:\.[0-9]+)?|[\u4e00-\u9fff]{2,}", text):
        normalized = norm_text(token)
        if len(normalized) >= 2 and normalized not in tokens:
            tokens.append(normalized)
    priority = ["W10.0-220", "10000", "220", "投标机型", "总容量", "机组台数"]
    for token in reversed(priority):
        normalized = norm_text(token)
        if normalized in tokens:
            tokens.remove(normalized)
        tokens.insert(0, normalized)
    return tokens[:180]


def row_score(text: str, tokens: list[str]) -> int:
    hay = norm_text(text)
    if not hay:
        return 0
    score = 0
    for token in tokens:
        if token and token in hay:
            score += min(12, max(2, len(token)))
    if re.search(r"\d", hay):
        score += 1
    return score


def extract_docx_snippets(path: Path, tokens: list[str], limit: int = 120) -> list[str]:
    doc = Document(str(path))
    candidates: list[tuple[int, int, str]] = []
    must_keep: list[str] = []
    order = 0
    for para in doc.paragraphs:
        text = display_text(para.text, 260)
        if norm_text(text):
            candidates.append((row_score(text, tokens), order, f"P{order}: {text}"))
            order += 1
    for table_idx, table in enumerate(doc.tables[:8]):
        small_lookup_table = len(table.rows) <= 90 and len(table.columns) <= 8
        for row_idx, row in enumerate(table.rows[:260]):
            cells = [display_text(cell.text, 120) for cell in row.cells[:14]]
            text = " | ".join(cells)
            if norm_text(text):
                line = f"T{table_idx}R{row_idx}: {text}"
                if row_idx < 3 or small_lookup_table:
                    must_keep.append(line)
                candidates.append((row_score(text, tokens), order, line))
                order += 1
    head = [item for item in candidates[:12]]
    scored = sorted((item for item in candidates if item[0] > 0), key=lambda item: (item[0], -item[1]), reverse=True)
    chosen: list[str] = []
    seen: set[str] = set()
    for text in must_keep:
        if text not in seen:
            seen.add(text)
            chosen.append(text)
        if len(chosen) >= limit:
            break
    for _, _, text in head + scored:
        if text not in seen:
            seen.add(text)
            chosen.append(text)
        if len(chosen) >= limit:
            break
    return chosen


def extract_xlsx_snippets(path: Path, tokens: list[str], limit: int = 120) -> list[str]:
    try:
        import openpyxl
    except Exception as exc:  # pragma: no cover - dependency exists in container
        return [f"[xlsx extraction unavailable: {exc}]"]

    workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    snippets: list[str] = []
    for sheet_name in workbook.sheetnames[:8]:
        sheet = workbook[sheet_name]
        rows: list[tuple[int, str]] = []
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            values = [display_text(value, 80) for value in row[:24]]
            text = " | ".join(value for value in values if value)
            if not norm_text(text):
                continue
            rows.append((row_idx, text))
            if row_idx >= 600:
                break
        selected_indices: set[int] = set(range(min(8, len(rows))))
        scored = sorted(
            ((row_score(text, tokens), pos, row_idx, text) for pos, (row_idx, text) in enumerate(rows)),
            key=lambda item: (item[0], -item[1]),
            reverse=True,
        )
        for score, pos, _, _ in scored:
            if score <= 0:
                break
            for nearby in range(max(0, pos - 3), min(len(rows), pos + 4)):
                selected_indices.add(nearby)
            if len(selected_indices) >= 40:
                break
        snippets.append(f"[sheet] {sheet_name}")
        for pos in sorted(selected_indices)[:45]:
            row_idx, text = rows[pos]
            snippets.append(f"{sheet_name}!R{row_idx}: {text}")
        if len(snippets) >= limit:
            break
    workbook.close()
    return snippets[:limit]


def extract_pdf_snippets(path: Path, tokens: list[str], limit: int = 40) -> list[str]:
    try:
        import fitz  # type: ignore
    except Exception as exc:  # pragma: no cover - dependency exists in container
        return [f"[pdf extraction unavailable: {exc}]"]

    doc = fitz.open(str(path))
    candidates: list[tuple[int, int, str]] = []
    for page_idx in range(min(len(doc), 40)):
        text = doc[page_idx].get_text("text")
        parts = [display_text(line, 260) for line in text.splitlines() if norm_text(line)]
        for line_idx, line in enumerate(parts[:120]):
            candidates.append((row_score(line, tokens), page_idx * 1000 + line_idx, f"P{page_idx + 1}: {line}"))
    doc.close()
    head = candidates[:12]
    scored = sorted((item for item in candidates if item[0] > 0), key=lambda item: (item[0], -item[1]), reverse=True)
    chosen: list[str] = []
    seen: set[str] = set()
    for _, _, text in head + scored:
        if text not in seen:
            seen.add(text)
            chosen.append(text)
        if len(chosen) >= limit:
            break
    return chosen


def source_context_for_agent(sources: list[dict[str, str]], tokens: list[str]) -> str:
    chunks: list[str] = []
    total_chars = 0
    for source in sources:
        path_text = source.get("path") or ""
        path = Path(path_text)
        header = f"## {source.get('id') or ''} {source.get('name') or path.name}\npath: {path_text}\nroute: {source.get('route') or ''}"
        if not path.exists():
            body = ["[missing source file]"]
        else:
            suffix = path.suffix.lower()
            try:
                if suffix == ".docx":
                    body = extract_docx_snippets(path, tokens)
                elif suffix in {".xlsx", ".xlsm"}:
                    body = extract_xlsx_snippets(path, tokens)
                elif suffix == ".pdf":
                    body = extract_pdf_snippets(path, tokens)
                else:
                    body = [f"[unsupported source type: {suffix}]"]
            except Exception as exc:
                body = [f"[source extraction error: {exc}]"]
        chunk = header + "\n" + "\n".join(body[:140])
        if total_chars + len(chunk) > MAX_SOURCE_CHARS and chunks:
            break
        chunks.append(chunk)
        total_chars += len(chunk)
    return "\n\n".join(chunks)


def select_parse_items(parse_result_path: Path, tokens: list[str], limit: int = 80) -> list[dict[str, str]]:
    if not parse_result_path.exists():
        return []
    try:
        payload = load_json(parse_result_path)
    except Exception:
        return []
    selected: list[tuple[int, int, dict[str, str]]] = []
    for idx, item in enumerate(payload.get("items") or []):
        if not isinstance(item, dict):
            continue
        text = "\n".join(
            str(item.get(key) or "")
            for key in ("title", "value", "keyValue", "keyEntity", "section", "evidence", "category")
        )
        score = row_score(text, tokens)
        category = str(item.get("category") or "")
        if category in {"project_basics", "performance_guarantees", "turbine_parameters"}:
            score += 2
        if score <= 0:
            continue
        selected.append(
            (
                score,
                idx,
                {
                    "id": str(item.get("id") or ""),
                    "title": display_text(item.get("title"), 120),
                    "value": display_text(item.get("value") or item.get("keyValue"), 220),
                    "section": display_text(item.get("section"), 160),
                    "evidence": display_text(item.get("evidence"), 300),
                    "category": str(item.get("category") or ""),
                    "location": str(item.get("evidenceLocation") or ""),
                },
            )
        )
    selected.sort(key=lambda item: (item[0], -item[1]), reverse=True)
    return [item for _, _, item in selected[:limit]]


def compact_sources(
    target: Target,
    selected_by_appendix: dict[str, list[dict[str, Any]]],
    catalog: list[dict[str, Any]],
    max_selected: int = 12,
) -> list[dict[str, str]]:
    seen: set[str] = set()
    compact: list[dict[str, str]] = []

    def add(raw: dict[str, Any], route: str = "") -> None:
        path = str(raw.get("path") or "")
        material_id = str(raw.get("id") or "")
        name = str(raw.get("name") or raw.get("title") or raw.get("fileName") or Path(path).name)
        key = path or material_id or name
        if not key or key in seen:
            return
        seen.add(key)
        compact.append(
            {
                "id": material_id,
                "name": name,
                "path": path,
                "route": route or str(raw.get("route") or raw.get("folderPath") or ""),
            }
        )

    for raw in target.recommended_materials:
        add(raw, "target.recommendedMaterials")
    for raw in selected_by_appendix.get(target.appendix_id, [])[:max_selected]:
        add(raw, str(raw.get("route") or "sourceSelection"))

    title_query = norm_text(target.title)
    scored: list[tuple[int, dict[str, Any]]] = []
    for raw in catalog:
        hay = norm_text(" ".join(str(raw.get(k) or "") for k in ("name", "folderPath", "fileName")))
        score = 0
        for token in re.findall(r"[A-Za-z0-9]+|[\u4e00-\u9fff]{2,}", target.title):
            if token and norm_text(token) in hay:
                score += len(token)
        if "参数" in title_query and raw.get("path", "").endswith(".xlsx"):
            score += 8
        if score:
            scored.append((score, raw))
    for _, raw in sorted(scored, key=lambda item: item[0], reverse=True)[:4]:
        add(raw, "catalog-title-overlap")

    return compact[:max_selected]


def build_prompt(
    manifest: dict[str, Any],
    parse_items: list[dict[str, str]],
    target: Target,
    output_path: Path,
    report_path: Path,
    workdir: Path,
    sources: list[dict[str, str]],
    table_preview: str,
    source_context: str,
    material_index_dir: Path,
) -> str:
    sources_json = json.dumps(sources, ensure_ascii=False, indent=2)
    project_context = {
        "projectId": manifest.get("projectId"),
        "projectName": manifest.get("projectName"),
        "projectTurbineModel": manifest.get("projectTurbineModel"),
        "parseFields": manifest.get("parseFields") or [],
        "selectedTenderParseItems": parse_items,
    }
    project_json = json.dumps(project_context, ensure_ascii=False, indent=2)
    return f"""
你是一个“单表独立填表 Agent”。这是一次泛化性实验：每张表都是全新会话，禁止调用生产规则脚本 `s4fill`，也不要使用人工填写文件。

当前执行环境的工具调用通道不可用，所以请不要调用 Bash/工具。你要基于下面提供的空表坐标、项目上下文和候选来源摘录，自主判断每个可填写单元格，并输出通用 JSON 补丁。外层 harness 只负责把你的补丁写入 Word，不包含任何表类型规则。

目标：
- 尽量把空白副表填到接近人工填写效果；不要只填显而易见的几个字段。
- 对每一个能从来源材料、招标解析信息或表格上下文合理推出的格子，都应尝试填写。
- 找不到依据时可以留空，但不要把“待填写”“/”“-”“无依据”当成有效填充。
- 保持原 Word 表格格式、行列结构和标题不变，只改应填写的单元格。

本表信息：
- appendixId: {target.appendix_id}
- title: {target.title}
- blankDocx: {target.docx_path}
- outputDocx: {output_path}
- reportJson: {report_path}
- workdir: {workdir}
- materialIndexDir: {material_index_dir}

项目上下文：
```json
{project_json}
```

候选来源材料：
{sources_json}

空表预览（仅为导航，最终请以 DOCX 原文件为准）：
```
{table_preview}
```

候选来源摘录：
```
{source_context}
```

输出要求：
1. 只输出 JSON，不要 Markdown，不要解释文字，不要调用工具。
2. `edits` 是你要填入 Word 的通用补丁；坐标使用空表预览中的 0-based table,row,col。
3. 只填写“投标机型1”或明显需要投标人填写的空白响应格；不要改编号、项目名、备注、表头。
4. 能从来源或项目上下文合理推导的格子都要填；同一行如果有“投标机型1/投标机型2”，通常只填投标机型1列。
5. 每个 edit 必须带 evidence，说明来自哪份来源或如何推导。
6. 口径原则：
- 数字优先保留来源精度，不要随意四舍五入；kg→t 要除以1000，MWh→万kWh 要除以10，kWh→万kWh 要除以10000。
- 如果表格单位是“万kWh”，而来源给的是“等效满负荷小时数(h)”和单机容量(kW)，单机电量=小时数×单机容量/10000；全场电量=来源给出的全场MWh/10，或单机电量×台数。
- 如果来源是按风速索引的矩阵表，应结合项目/招标解析中的指定测风塔年平均风速，选择相同或最近的风速行；不要把机位坐标表当成电量矩阵。
- “投标机型”字段优先用不含布局后缀的正式机型名；“上置/下置/内置”等布局信息用于“箱变位置”等字段。
- 遇到同一行多个投标响应空列时：如果这些列代表同一项目的通用承诺值，可以填相同值；如果明确代表不同机型，则只填已有机型对应列。
7. 最终 JSON schema：
{{
  "schema_version": "agent-table-fill-e2e-v1",
  "appendixId": "{target.appendix_id}",
  "title": "{target.title}",
  "outputFile": "{output_path}",
  "reportFile": "{report_path}",
  "edits": [
    {{"table": 0, "row": 1, "col": 3, "text": "示例值", "evidence": "来源文件/行/推导说明"}}
  ],
  "unfilledHints": [],
  "notes": []
}}
""".strip()


def run_one(
    base_url: str,
    target: Target,
    manifest: dict[str, Any],
    output_dir: Path,
    selected_by_appendix: dict[str, list[dict[str, Any]]],
    catalog: list[dict[str, Any]],
    material_index_dir: Path,
    timeout: int,
) -> dict[str, Any]:
    started = time.time()
    target_dir = output_dir / f"{target.index:03d}-{target.appendix_id}"
    target_dir.mkdir(parents=True, exist_ok=True)
    output_path = target_dir / f"{target.index:03d}-{target.appendix_id}-{slug(target.title)}_Agent填写.docx"
    report_path = target_dir / "agent_fill_report.json"
    trace_path = target_dir / "opencode_response.json"
    prompt_path = target_dir / "prompt.txt"

    sources = compact_sources(target, selected_by_appendix, catalog)
    table_preview = compact_table_preview(target.docx_path)
    tokens = query_tokens(target, table_preview, manifest)
    parse_result_path = target.docx_path.parents[1] / "s4_gap_workdir" / "parse_result.json"
    parse_items = select_parse_items(parse_result_path, tokens)
    source_context = source_context_for_agent(sources, tokens)
    prompt = build_prompt(
        manifest,
        parse_items,
        target,
        output_path,
        report_path,
        target_dir,
        sources,
        table_preview,
        source_context,
        material_index_dir,
    )
    prompt_path.write_text(prompt, encoding="utf-8")

    session_id = ""
    try:
        session_id = create_session(base_url, f"Agent填表 {target.appendix_id} {target.title}", timeout=30)
        response = send_prompt(base_url, session_id, prompt, timeout=timeout)
        trace_path.write_text(json.dumps(response, ensure_ascii=False, indent=2), encoding="utf-8")
        text = assistant_text(response)
        parsed = parse_json_from_text(text)
        edit_count = apply_agent_patch(target.docx_path, output_path, parsed)
        report = {
            "schema_version": "agent-table-fill-e2e-report-v1",
            "appendixId": target.appendix_id,
            "title": target.title,
            "blankDocx": str(target.docx_path),
            "outputFile": str(output_path),
            "sources": sources,
            "edits": parsed.get("edits") if isinstance(parsed.get("edits"), list) else [],
            "unfilledHints": parsed.get("unfilledHints") if isinstance(parsed.get("unfilledHints"), list) else [],
            "notes": parsed.get("notes") if isinstance(parsed.get("notes"), list) else [],
            "appliedEditCount": edit_count,
        }
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        ok = output_path.exists() and output_path.stat().st_size > 0
        return {
            "appendixId": target.appendix_id,
            "title": target.title,
            "index": target.index,
            "ok": ok,
            "sessionId": session_id,
            "outputFile": str(output_path),
            "reportFile": str(report_path),
            "assistantJson": parsed,
            "appliedEditCount": edit_count,
            "durationSec": round(time.time() - started, 2),
        }
    except Exception as exc:
        fallback = target_dir / f"{target.index:03d}-{target.appendix_id}-{slug(target.title)}_Agent失败原表.docx"
        if target.docx_path.exists() and not fallback.exists():
            shutil.copyfile(target.docx_path, fallback)
        return {
            "appendixId": target.appendix_id,
            "title": target.title,
            "index": target.index,
            "ok": False,
            "sessionId": session_id,
            "outputFile": str(fallback),
            "reportFile": str(report_path),
            "error": str(exc),
            "durationSec": round(time.time() - started, 2),
        }


def set_cell_text(cell: Any, text: str) -> None:
    if cell.paragraphs:
        paragraph = cell.paragraphs[0]
        for extra in cell.paragraphs[1:]:
            element = extra._element
            element.getparent().remove(element)
        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(text)
    else:
        cell.text = text


def apply_agent_patch(blank_docx: Path, output_path: Path, payload: dict[str, Any]) -> int:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(blank_docx, output_path)
    doc = Document(str(output_path))
    edits = payload.get("edits")
    if not isinstance(edits, list):
        doc.save(str(output_path))
        return 0
    applied = 0
    seen: set[tuple[int, int, int]] = set()
    for edit in edits:
        if not isinstance(edit, dict):
            continue
        try:
            table_idx = int(edit.get("table", 0))
            row_idx = int(edit.get("row"))
            col_idx = int(edit.get("col"))
        except Exception:
            continue
        text = str(edit.get("text") or "").strip()
        if not text or text in {"/", "-", "待填写", "无依据"}:
            continue
        key = (table_idx, row_idx, col_idx)
        if key in seen:
            continue
        seen.add(key)
        if table_idx < 0 or table_idx >= len(doc.tables):
            continue
        table = doc.tables[table_idx]
        if row_idx < 0 or row_idx >= len(table.rows):
            continue
        row = table.rows[row_idx]
        if col_idx < 0 or col_idx >= len(row.cells):
            continue
        set_cell_text(row.cells[col_idx], text)
        applied += 1
    doc.save(str(output_path))
    return applied


def table_cells(path: Path) -> list[list[list[str]]]:
    if not path.exists():
        return []
    doc = Document(str(path))
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        tables.append(rows)
    return tables


def table_static_tokens(table: list[list[str]]) -> set[str]:
    tokens: set[str] = set()
    for row in table:
        for cell in row:
            text = norm_text(cell)
            if len(text) >= 2:
                tokens.add(text)
    return tokens


def table_similarity(a: list[list[str]], b: list[list[str]]) -> float:
    a_tokens = table_static_tokens(a)
    b_tokens = table_static_tokens(b)
    if not a_tokens or not b_tokens:
        return 0.0
    overlap = len(a_tokens & b_tokens)
    return overlap / max(1, min(len(a_tokens), len(b_tokens)))


def best_human_table(blank_table: list[list[str]], human_tables: list[list[list[str]]]) -> tuple[int, float]:
    best_idx = -1
    best_score = 0.0
    for idx, human_table in enumerate(human_tables):
        score = table_similarity(blank_table, human_table)
        if score > best_score:
            best_idx = idx
            best_score = score
    return best_idx, best_score


def cell_at(table: list[list[str]], row: int, col: int) -> str:
    if row < 0 or row >= len(table):
        return ""
    if col < 0 or col >= len(table[row]):
        return ""
    return table[row][col]


def parse_simple_number(value: str) -> float | None:
    text = str(value or "").strip()
    if not text:
        return None
    normalized = text.replace(",", "").replace("，", "")
    if not re.fullmatch(r"~?\s*[-+]?\d+(?:\.\d+)?\s*", normalized):
        return None
    try:
        return float(normalized.replace("~", "").strip())
    except ValueError:
        return None


def tolerant_equal(left: str, right: str) -> bool:
    left_norm = norm_text(left)
    right_norm = norm_text(right)
    if left_norm == right_norm:
        return True
    left_num = parse_simple_number(str(left))
    right_num = parse_simple_number(str(right))
    if left_num is None or right_num is None:
        return False
    tolerance = max(0.05, abs(right_num) * 0.0005)
    return abs(left_num - right_num) <= tolerance


def score_outputs(results: list[dict[str, Any]], human_docx: Path, min_align: float = 0.35) -> dict[str, Any]:
    if not human_docx.exists():
        return {"error": f"human docx not found: {human_docx}"}
    human_tables = table_cells(human_docx)
    rows: list[dict[str, Any]] = []
    totals = {
        "targetTables": 0,
        "alignedTables": 0,
        "humanChangedCells": 0,
        "agentCorrectCells": 0,
        "agentTolerantCorrectCells": 0,
        "agentTouchedCells": 0,
        "sameRowAttemptCells": 0,
    }
    for result in results:
        blank = table_cells(Path(str(result.get("blankDocx") or "")))
        agent = table_cells(Path(str(result.get("outputFile") or "")))
        for table_idx, blank_table in enumerate(blank):
            totals["targetTables"] += 1
            agent_table = agent[table_idx] if table_idx < len(agent) else []
            human_idx, align_score = best_human_table(blank_table, human_tables)
            blank_human_idx = human_idx
            blank_align_score = align_score
            agent_human_idx, agent_align_score = (-1, 0.0)
            alignment_method = "blank"
            if agent_table:
                agent_human_idx, agent_align_score = best_human_table(agent_table, human_tables)
                output_expanded = len(agent_table) >= len(blank_table) + 3
                output_clearly_better = agent_align_score >= align_score + 0.08
                output_high_confidence = agent_align_score >= (0.85 if blank_align_score < 0.8 else 0.995)
                if agent_human_idx >= 0 and output_expanded and output_clearly_better and output_high_confidence:
                    human_idx = agent_human_idx
                    align_score = agent_align_score
                    alignment_method = "agent-expanded"
            if human_idx < 0 or align_score < min_align:
                rows.append(
                    {
                        "appendixId": result.get("appendixId"),
                        "title": result.get("title"),
                        "table": table_idx,
                        "aligned": False,
                        "alignScore": round(align_score, 3),
                        "blankHumanTable": blank_human_idx,
                        "blankAlignScore": round(blank_align_score, 3),
                        "agentHumanTable": agent_human_idx,
                        "agentAlignScore": round(agent_align_score, 3),
                        "alignmentMethod": alignment_method,
                    }
                )
                continue
            totals["alignedTables"] += 1
            human_table = human_tables[human_idx]
            changed = correct = touched = same_row_attempt = 0
            tolerant_correct = 0
            max_rows = max(len(blank_table), len(human_table), len(agent_table))
            for row_idx in range(max_rows):
                max_cols = max(
                    len(blank_table[row_idx]) if row_idx < len(blank_table) else 0,
                    len(human_table[row_idx]) if row_idx < len(human_table) else 0,
                    len(agent_table[row_idx]) if row_idx < len(agent_table) else 0,
                )
                human_row_values = {norm_text(cell) for cell in (human_table[row_idx] if row_idx < len(human_table) else [])}
                agent_row_values = {norm_text(cell) for cell in (agent_table[row_idx] if row_idx < len(agent_table) else [])}
                for col_idx in range(max_cols):
                    blank_text = norm_text(cell_at(blank_table, row_idx, col_idx))
                    human_text = norm_text(cell_at(human_table, row_idx, col_idx))
                    agent_text = norm_text(cell_at(agent_table, row_idx, col_idx))
                    if human_text and human_text != blank_text:
                        changed += 1
                        if agent_text == human_text:
                            correct += 1
                        if tolerant_equal(cell_at(agent_table, row_idx, col_idx), cell_at(human_table, row_idx, col_idx)):
                            tolerant_correct += 1
                        elif human_text in agent_row_values:
                            same_row_attempt += 1
                    if agent_text and agent_text != blank_text:
                        touched += 1
            totals["humanChangedCells"] += changed
            totals["agentCorrectCells"] += correct
            totals["agentTolerantCorrectCells"] += tolerant_correct
            totals["agentTouchedCells"] += touched
            totals["sameRowAttemptCells"] += same_row_attempt
            rows.append(
                {
                    "appendixId": result.get("appendixId"),
                    "title": result.get("title"),
                    "table": table_idx,
                    "aligned": True,
                    "humanTable": human_idx,
                    "alignScore": round(align_score, 3),
                    "blankHumanTable": blank_human_idx,
                    "blankAlignScore": round(blank_align_score, 3),
                    "agentHumanTable": agent_human_idx,
                    "agentAlignScore": round(agent_align_score, 3),
                    "alignmentMethod": alignment_method,
                    "humanChangedCells": changed,
                    "agentCorrectCells": correct,
                    "agentTolerantCorrectCells": tolerant_correct,
                    "agentTouchedCells": touched,
                    "sameRowAttemptCells": same_row_attempt,
                    "strictCoverage": round(correct / changed, 4) if changed else None,
                    "tolerantCoverage": round(tolerant_correct / changed, 4) if changed else None,
                }
            )
    totals["strictCoverage"] = (
        round(totals["agentCorrectCells"] / totals["humanChangedCells"], 4)
        if totals["humanChangedCells"]
        else None
    )
    totals["rowAttemptCoverage"] = (
        round((totals["agentCorrectCells"] + totals["sameRowAttemptCells"]) / totals["humanChangedCells"], 4)
        if totals["humanChangedCells"]
        else None
    )
    totals["tolerantCoverage"] = (
        round(totals["agentTolerantCorrectCells"] / totals["humanChangedCells"], 4)
        if totals["humanChangedCells"]
        else None
    )
    totals["agentTouchToHumanRatio"] = (
        round(totals["agentTouchedCells"] / totals["humanChangedCells"], 4)
        if totals["humanChangedCells"]
        else None
    )
    return {"summary": totals, "tables": rows}


def write_markdown_report(path: Path, run_payload: dict[str, Any], score: dict[str, Any]) -> None:
    summary = run_payload.get("summary") or {}
    score_summary = score.get("summary") or {}
    lines = [
        "# Agent Table Fill E2E",
        "",
        "## Run",
        "",
        f"- targets: {summary.get('targetCount')}",
        f"- success: {summary.get('successCount')}",
        f"- failed: {summary.get('failedCount')}",
        f"- outputDir: {run_payload.get('outputDir')}",
        "",
        "## Human Comparison",
        "",
    ]
    if score_summary:
        lines.extend(
            [
                f"- alignedTables: {score_summary.get('alignedTables')}/{score_summary.get('targetTables')}",
                f"- humanChangedCells: {score_summary.get('humanChangedCells')}",
                f"- agentCorrectCells: {score_summary.get('agentCorrectCells')}",
                f"- agentTolerantCorrectCells: {score_summary.get('agentTolerantCorrectCells')}",
                f"- agentTouchedCells: {score_summary.get('agentTouchedCells')}",
                f"- strictCoverage: {score_summary.get('strictCoverage')}",
                f"- tolerantCoverage: {score_summary.get('tolerantCoverage')}",
                f"- rowAttemptCoverage: {score_summary.get('rowAttemptCoverage')}",
                "",
            ]
        )
    else:
        lines.append(f"- scoreError: {score.get('error')}")
        lines.append("")

    rows = sorted(
        (score.get("tables") or []),
        key=lambda row: (
            row.get("strictCoverage") is None,
            row.get("strictCoverage") if row.get("strictCoverage") is not None else -1,
        ),
    )
    lines.extend(["## Lowest Tables", ""])
    for row in rows[:30]:
        lines.append(
            "- {appendixId} {title}: changed={humanChangedCells}, correct={agentCorrectCells}, "
            "tolerant={agentTolerantCorrectCells}, touched={agentTouchedCells}, "
            "coverage={strictCoverage}, tolerantCoverage={tolerantCoverage}, align={alignScore}".format(
                appendixId=row.get("appendixId"),
                title=row.get("title"),
                humanChangedCells=row.get("humanChangedCells"),
                agentCorrectCells=row.get("agentCorrectCells"),
                agentTolerantCorrectCells=row.get("agentTolerantCorrectCells"),
                agentTouchedCells=row.get("agentTouchedCells"),
                strictCoverage=row.get("strictCoverage"),
                tolerantCoverage=row.get("tolerantCoverage"),
                alignScore=row.get("alignScore"),
            )
        )
    lines.append("")
    path.write_text("\n".join(lines), encoding="utf-8")


def main(argv: list[str]) -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--manifest", type=Path, default=DEFAULT_MANIFEST)
    parser.add_argument("--source-selection-report", type=Path, default=DEFAULT_SOURCE_SELECTION_REPORT)
    parser.add_argument("--human-docx", type=Path, default=DEFAULT_HUMAN_DOCX)
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR)
    parser.add_argument("--opencode-url", default=DEFAULT_OPENCODE_URL)
    parser.add_argument("--limit", type=int, default=0)
    parser.add_argument("--ids", default="", help="comma-separated appendix ids")
    parser.add_argument("--concurrency", type=int, default=2)
    parser.add_argument("--timeout", type=int, default=1800)
    args = parser.parse_args(argv)

    manifest, targets = load_targets(args.manifest)
    if args.ids.strip():
        wanted = {item.strip() for item in args.ids.split(",") if item.strip()}
        targets = [target for target in targets if target.appendix_id in wanted]
    if args.limit > 0:
        targets = targets[: args.limit]
    if not targets:
        raise SystemExit("no targets selected")

    material_index_dir = args.manifest.parent.parent / "material_index"
    selected_by_appendix = load_source_selection(args.source_selection_report)
    catalog = material_catalog(manifest)

    args.output_dir.mkdir(parents=True, exist_ok=True)
    results: list[dict[str, Any]] = []
    print(
        json.dumps(
            {
                "event": "start",
                "targets": len(targets),
                "outputDir": str(args.output_dir),
                "concurrency": args.concurrency,
            },
            ensure_ascii=False,
        ),
        flush=True,
    )

    with concurrent.futures.ThreadPoolExecutor(max_workers=max(1, args.concurrency)) as pool:
        future_map = {}
        for target in targets:
            future = pool.submit(
                run_one,
                args.opencode_url,
                target,
                manifest,
                args.output_dir,
                selected_by_appendix,
                catalog,
                material_index_dir,
                args.timeout,
            )
            future_map[future] = target
        for future in concurrent.futures.as_completed(future_map):
            target = future_map[future]
            result = future.result()
            result["blankDocx"] = str(target.docx_path)
            results.append(result)
            print(
                json.dumps(
                    {
                        "event": "target_done",
                        "appendixId": result.get("appendixId"),
                        "ok": result.get("ok"),
                        "durationSec": result.get("durationSec"),
                        "error": result.get("error"),
                    },
                    ensure_ascii=False,
                ),
                flush=True,
            )

    results.sort(key=lambda item: int(item.get("index") or 0))
    run_payload = {
        "schema_version": "agent-table-fill-e2e-run-v1",
        "manifest": str(args.manifest),
        "outputDir": str(args.output_dir),
        "summary": {
            "targetCount": len(results),
            "successCount": sum(1 for item in results if item.get("ok")),
            "failedCount": sum(1 for item in results if not item.get("ok")),
        },
        "results": results,
    }
    run_json = args.output_dir / "agent_table_e2e_run.json"
    run_json.write_text(json.dumps(run_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    score = score_outputs(results, args.human_docx)
    score_json = args.output_dir / "human_comparison_score.json"
    score_json.write_text(json.dumps(score, ensure_ascii=False, indent=2), encoding="utf-8")
    write_markdown_report(args.output_dir / "human_comparison_report.md", run_payload, score)

    print(
        json.dumps(
            {
                "schema_version": "agent-table-fill-e2e-summary-v1",
                "outputDir": str(args.output_dir),
                "runFile": str(run_json),
                "scoreFile": str(score_json),
                "reportFile": str(args.output_dir / "human_comparison_report.md"),
                "summary": run_payload["summary"],
                "scoreSummary": score.get("summary") or {"error": score.get("error")},
            },
            ensure_ascii=False,
            indent=2,
        ),
        flush=True,
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
