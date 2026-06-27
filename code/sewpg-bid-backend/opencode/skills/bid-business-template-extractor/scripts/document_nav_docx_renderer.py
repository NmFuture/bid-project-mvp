from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document


def safe_name(value: str, fallback: str) -> str:
    name = re.sub(r'[<>:"/\\|?*\r\n\t]+', "_", str(value or "").strip())
    name = re.sub(r"\s+", " ", name).strip(" .")
    return (name or fallback)[:80]


def _block_map(blocks: list[dict[str, Any]]) -> dict[int, dict[str, Any]]:
    items: dict[int, dict[str, Any]] = {}
    for block in blocks:
        if not isinstance(block, dict):
            continue
        try:
            block_id = int(block.get("blockId"))
        except (TypeError, ValueError):
            continue
        items[block_id] = block
    return items


def _add_table(document: Document, rows: list[Any]) -> int:
    normalized_rows = [row for row in rows if isinstance(row, list)]
    column_count = max((len(row) for row in normalized_rows), default=0)
    if not normalized_rows or column_count <= 0:
        return 0
    table = document.add_table(rows=len(normalized_rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, row in enumerate(normalized_rows):
        for column_index in range(column_count):
            table.cell(row_index, column_index).text = str(row[column_index] if column_index < len(row) else "")
    return len(normalized_rows)


def _write_template_docx(
    *,
    target: Path,
    title: str,
    selected_blocks: list[dict[str, Any]],
) -> int:
    document = Document()
    document.add_heading(title, level=1)
    row_count = 0
    for block in selected_blocks:
        rows = block.get("rows") if isinstance(block.get("rows"), list) else []
        if rows:
            row_count += _add_table(document, rows)
            continue
        text = str(block.get("text") or "").strip()
        if not text:
            continue
        if str(block.get("type") or "").lower() == "heading":
            document.add_heading(text, level=2)
        else:
            document.add_paragraph(text)
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(target))
    return row_count


def render_document_nav_templates(
    *,
    source_document: dict[str, Any],
    blocks: list[dict[str, Any]],
    boundaries: dict[str, Any],
    output_dir: Path,
) -> dict[str, list[dict[str, Any]]]:
    templates_dir = output_dir / "templates"
    templates_dir.mkdir(parents=True, exist_ok=True)
    blocks_by_id = _block_map(blocks)
    rendered: list[dict[str, Any]] = []
    source_engine = str(source_document.get("documentParseEngine") or source_document.get("sourceEngine") or "mineru")
    for template in boundaries.get("templates") or []:
        if not isinstance(template, dict):
            continue
        start = int(template["startBlockId"])
        end = int(template["endBlockId"])
        title = str(template.get("title") or template.get("templateTitle") or "商务模板").strip()
        selected = [blocks_by_id[block_id] for block_id in range(start, end + 1) if block_id in blocks_by_id]
        filename = f"{template['id']}-{safe_name(title, 'business-template')}.docx"
        target = templates_dir / filename
        row_count = _write_template_docx(target=target, title=title, selected_blocks=selected)
        item = dict(template)
        item["outputPath"] = str(Path("templates") / filename).replace("\\", "/")
        item["sourceEngine"] = source_engine
        item["sourceDocumentId"] = str(source_document.get("id") or "")
        item["rowCount"] = row_count
        quality = item.get("quality") if isinstance(item.get("quality"), dict) else {}
        item["quality"] = {**quality, "sourceEngine": source_engine}
        rendered.append(item)
    return {"templates": rendered}
