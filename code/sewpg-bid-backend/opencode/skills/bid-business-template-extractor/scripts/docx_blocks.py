from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document


WORD_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


@dataclass(frozen=True)
class Block:
    block_id: int
    body_index: int
    type: str
    text: str
    rows: list[list[str]]
    style_name: str
    is_centered: bool
    has_page_break_before: bool
    has_page_break_after: bool
    page_segment: int
    position_in_page_segment: int | None
    is_page_first_non_empty: bool

    def to_dict(self) -> dict[str, Any]:
        return {
            "blockId": self.block_id,
            "bodyIndex": self.body_index,
            "type": self.type,
            "text": self.text,
            "rows": self.rows,
            "styleName": self.style_name,
            "isCentered": self.is_centered,
            "hasPageBreakBefore": self.has_page_break_before,
            "hasPageBreakAfter": self.has_page_break_after,
            "pageSegment": self.page_segment,
            "positionInPageSegment": self.position_in_page_segment,
            "isPageFirstNonEmpty": self.is_page_first_non_empty,
        }


def _paragraph_text(element: Any) -> str:
    return "".join(node.text or "" for node in element.iter(f"{WORD_NS}t")).strip()


def _xml_attr(element: Any, name: str) -> str:
    if element is None:
        return ""
    return str(element.get(f"{WORD_NS}{name}") or element.get(name) or "").strip()


def _paragraph_has_page_break(element: Any) -> bool:
    for br in element.iter(f"{WORD_NS}br"):
        if _xml_attr(br, "type").lower() == "page":
            return True
    return False


def _paragraph_page_break_before(element: Any) -> bool:
    p_pr = element.find(f"{WORD_NS}pPr")
    return p_pr is not None and p_pr.find(f"{WORD_NS}pageBreakBefore") is not None


def _paragraph_alignment(element: Any, paragraph: Any) -> str:
    alignment = getattr(paragraph, "alignment", None) if paragraph is not None else None
    if alignment is not None:
        value = getattr(alignment, "value", alignment)
        if str(value) in {"1", "CENTER", "WD_ALIGN_PARAGRAPH.CENTER"}:
            return "center"
    p_pr = element.find(f"{WORD_NS}pPr")
    jc = p_pr.find(f"{WORD_NS}jc") if p_pr is not None else None
    return _xml_attr(jc, "val").lower()


def _paragraph_style_name(paragraph: Any) -> str:
    return str(getattr(getattr(paragraph, "style", None), "name", "") or "")


def _table_rows(table: Any) -> list[list[str]]:
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _contains_section_properties(element: Any) -> bool:
    return element.find(f".//{WORD_NS}sectPr") is not None


def extract_blocks(path: Path) -> list[dict[str, Any]]:
    doc = Document(str(path))
    tables = iter(doc.tables)
    paragraphs = iter(doc.paragraphs)
    blocks: list[Block] = []
    page_segment = 1
    non_empty_position = 0
    next_block_has_page_break_before = False

    for body_index, child in enumerate(doc.element.body.iterchildren()):
        has_break_before = next_block_has_page_break_before
        next_block_has_page_break_before = False
        if child.tag == f"{WORD_NS}p":
            paragraph = next(paragraphs, None)
            text = _paragraph_text(child)
            explicit_break_before = _paragraph_page_break_before(child)
            has_break_before = has_break_before or explicit_break_before
            if has_break_before:
                page_segment += 1
                non_empty_position = 0
            if text:
                non_empty_position += 1
                position_in_page_segment: int | None = non_empty_position
            else:
                position_in_page_segment = None
            has_break_after = _paragraph_has_page_break(child) or _contains_section_properties(child)
            block = Block(
                block_id=len(blocks) + 1,
                body_index=body_index,
                type="paragraph",
                text=text,
                rows=[],
                style_name=_paragraph_style_name(paragraph),
                is_centered=_paragraph_alignment(child, paragraph) == "center",
                has_page_break_before=has_break_before,
                has_page_break_after=has_break_after,
                page_segment=page_segment,
                position_in_page_segment=position_in_page_segment,
                is_page_first_non_empty=position_in_page_segment == 1,
            )
            blocks.append(block)
            if has_break_after:
                next_block_has_page_break_before = True
        elif child.tag == f"{WORD_NS}tbl":
            table = next(tables, None)
            rows = _table_rows(table) if table is not None else []
            if has_break_before:
                page_segment += 1
                non_empty_position = 0
            non_empty_position += 1
            has_break_after = _contains_section_properties(child)
            blocks.append(
                Block(
                    block_id=len(blocks) + 1,
                    body_index=body_index,
                    type="table",
                    text="\n".join(" | ".join(row) for row in rows),
                    rows=rows,
                    style_name="",
                    is_centered=False,
                    has_page_break_before=has_break_before,
                    has_page_break_after=has_break_after,
                    page_segment=page_segment,
                    position_in_page_segment=non_empty_position,
                    is_page_first_non_empty=non_empty_position == 1,
                )
            )
            if has_break_after:
                next_block_has_page_break_before = True

    return [block.to_dict() for block in blocks]
