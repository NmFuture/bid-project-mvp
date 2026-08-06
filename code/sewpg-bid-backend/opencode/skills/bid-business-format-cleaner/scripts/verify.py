#!/usr/bin/env python3
"""生成商务标格式清洗报告。"""

from __future__ import annotations

import argparse
import json
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn

try:
    from outline_matcher import flatten_outline, load_outline
except ImportError:  # pragma: no cover
    from .outline_matcher import flatten_outline, load_outline


def verify_cleaned_docx(
    *,
    output_file: str | Path,
    outline_file: str | Path,
    report_file: str | Path,
    clean_result: dict[str, Any] | None = None,
    style_spec_path: str | Path | None = None,
) -> dict[str, Any]:
    output_path = Path(output_file)
    outline_path = Path(outline_file)
    report_path = Path(report_file)
    style_path = Path(style_spec_path) if style_spec_path else None

    outline_items = flatten_outline(load_outline(outline_path))
    doc = Document(str(output_path))
    toc_present = _doc_has_toc(doc)
    header_text = "\n".join(
        paragraph.text
        for section in doc.sections
        for header in (section.header, section.first_page_header, section.even_page_header)
        for paragraph in header.paragraphs
    )
    residual_terms = _load_residual_terms(style_path)
    header_cleaned = not any(term in header_text for term in residual_terms)

    matched = list((clean_result or {}).get("matchedHeadings") or [])
    unmatched = list((clean_result or {}).get("unmatchedHeadings") or [])
    if not matched and not unmatched:
        heading_texts = {paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.style.name.startswith("Heading")}
        matched, unmatched = _infer_matches_from_doc(outline_items, heading_texts)

    risks = list((clean_result or {}).get("formatRisks") or [])
    risks.extend(_structural_risks(doc, toc_present, header_cleaned, len(outline_items), len(unmatched)))
    risks = _dedupe(risks)

    report = {
        "outlineCount": len(outline_items),
        "matchedHeadingCount": len(matched),
        "unmatchedHeadings": unmatched,
        "tocPresent": toc_present,
        "headerCleaned": header_cleaned,
        "pagination": dict((clean_result or {}).get("pagination") or {}),
        "fontFamilies": list((clean_result or {}).get("fontFamilies") or []),
        "formatRisks": risks,
    }
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text(_render_report(report, output_path, outline_path), encoding="utf-8")
    return report


def _doc_has_toc(doc: Document) -> bool:
    return any((node.text or "").upper().find("TOC") >= 0 for node in doc.element.iter(qn("w:instrText")))


def _load_residual_terms(style_path: Path | None) -> list[str]:
    if style_path is None:
        style_path = Path(__file__).resolve().parents[1] / "references" / "business_heading_style.json"
    if style_path and style_path.exists():
        data = json.loads(style_path.read_text(encoding="utf-8"))
        return [str(term) for term in data.get("header", {}).get("residual_terms", [])]
    return []


def _infer_matches_from_doc(outline_items: list[dict[str, Any]], heading_texts: set[str]):
    matched: list[dict[str, Any]] = []
    unmatched: list[dict[str, Any]] = []
    for item in outline_items:
        expected = _heading_text(item)
        record = {
            "id": item.get("id"),
            "number": item.get("number", ""),
            "title": item.get("title", ""),
            "level": item.get("level"),
            "order": item.get("order"),
        }
        if expected in heading_texts:
            matched.append(record)
        else:
            unmatched.append(record)
    return matched, unmatched


def _heading_text(item: dict[str, Any]) -> str:
    number = str(item.get("number") or "").strip()
    title = str(item.get("title") or "").strip()
    return f"{number} {title}".strip() if number else title


def _structural_risks(
    doc: Document,
    toc_present: bool,
    header_cleaned: bool,
    outline_count: int,
    unmatched_count: int,
) -> list[str]:
    risks: list[str] = []
    if outline_count == 0:
        risks.append("outline 为空，未能定位任何商务标标题。")
    if unmatched_count:
        risks.append("存在未匹配标题，脚本不会插入或删除正文，请人工核对缺失章节。")
    if not toc_present:
        risks.append("未检测到 TOC 域，目录页需要人工补充。")
    if not header_cleaned:
        risks.append("页眉仍疑似存在技术标残留语义，请人工检查页眉。")
    if not doc.paragraphs:
        risks.append("输出文档未检测到正文段落。")
    if not risks:
        risks.append("未发现脚本可识别的格式风险；仍建议在 Word/WPS 中打开后刷新目录并人工抽检。")
    return risks


def _dedupe(items: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for item in items:
        if item in seen:
            continue
        seen.add(item)
        result.append(item)
    return result


def _render_report(report: dict[str, Any], output_path: Path, outline_path: Path) -> str:
    lines = [
        "# 商务标格式清洗报告",
        "",
        f"- 输出文件：`{output_path}`",
        f"- outline 文件：`{outline_path}`",
        f"- outline 总数：{report['outlineCount']}",
        f"- 成功匹配标题数：{report['matchedHeadingCount']}",
        f"- 未匹配标题数：{len(report['unmatchedHeadings'])}",
        f"- TOC 是否插入：{'是' if report['tocPresent'] else '否'}",
        f"- 页眉是否清理：{'是' if report['headerCleaned'] else '否'}",
        f"- 实际写入字体族：{', '.join(report.get('fontFamilies') or []) or '未记录'}",
        f"- 新增小节分页符数：{report.get('pagination', {}).get('insertedPageBreaks', 0)}",
        f"- 清理多余分页符数：{report.get('pagination', {}).get('removedBlankPageBreaks', 0)}",
        f"- 清理分页间空段数：{report.get('pagination', {}).get('removedBlankParagraphs', 0)}",
        "",
        "## 未匹配标题清单",
        "",
    ]
    if report["unmatchedHeadings"]:
        for item in report["unmatchedHeadings"]:
            text = _heading_text(item)
            lines.append(f"- {text}（id={item.get('id')}, level={item.get('level')}）")
    else:
        lines.append("- 无")
    lines.extend(["", "## 格式风险提示", ""])
    for risk in report["formatRisks"]:
        lines.append(f"- {risk}")
    lines.append("")
    return "\n".join(lines)


def _settings_update_fields(docx_path: Path) -> bool:
    with zipfile.ZipFile(docx_path) as archive:
        try:
            settings = archive.read("word/settings.xml").decode("utf-8", errors="replace")
        except KeyError:
            return False
    return "updateFields" in settings


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="验证商务标格式清洗结果")
    parser.add_argument("--output", required=True)
    parser.add_argument("--outline", required=True)
    parser.add_argument("--report", required=True)
    parser.add_argument("--result-json")
    parser.add_argument("--style")
    args = parser.parse_args(argv)

    clean_result = None
    if args.result_json:
        clean_result = json.loads(Path(args.result_json).read_text(encoding="utf-8"))
    report = verify_cleaned_docx(
        output_file=args.output,
        outline_file=args.outline,
        report_file=args.report,
        clean_result=clean_result,
        style_spec_path=args.style,
    )
    report["updateFields"] = _settings_update_fields(Path(args.output))
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
