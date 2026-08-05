from __future__ import annotations

import importlib.util
import json
import sys
import unittest
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_BREAK
from xml.etree import ElementTree as ET


ROOT = Path(__file__).resolve().parents[1]
RUNNER_PATH = ROOT / "scripts" / "run_from_manifest.py"
SCRIPT_DIR = str(ROOT / "scripts")
if SCRIPT_DIR not in sys.path:
    sys.path.insert(0, SCRIPT_DIR)


def _run_manifest(manifest_path: Path, *, response: str = "summary"):
    module_name = "bid_business_format_cleaner_run_from_manifest"
    inserted = False
    if SCRIPT_DIR not in sys.path:
        sys.path.insert(0, SCRIPT_DIR)
        inserted = True
    try:
        spec = importlib.util.spec_from_file_location(module_name, RUNNER_PATH)
        if spec is None or spec.loader is None:
            raise ImportError(f"cannot load runner: {RUNNER_PATH}")
        module = importlib.util.module_from_spec(spec)
        sys.modules[module_name] = module
        spec.loader.exec_module(module)
        return module.run_manifest(manifest_path, response=response)
    finally:
        sys.modules.pop(module_name, None)
        if inserted:
            sys.path.remove(SCRIPT_DIR)


W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _read_docx_part(path: Path, name: str) -> str:
    with zipfile.ZipFile(path) as archive:
        return archive.read(name).decode("utf-8")


def _style_by_name(styles_xml: str, style_name: str) -> ET.Element:
    root = ET.fromstring(styles_xml)
    for style in root.findall("w:style", W_NS):
        name = style.find("w:name", W_NS)
        if name is not None and name.attrib.get(f"{{{W_NS['w']}}}val") == style_name:
            return style
    raise AssertionError(f"style not found: {style_name}")


def _style_spacing(style: ET.Element) -> ET.Element:
    spacing = style.find("w:pPr/w:spacing", W_NS)
    if spacing is None:
        raise AssertionError("style spacing not found")
    return spacing


def _style_indent(style: ET.Element) -> ET.Element:
    indent = style.find("w:pPr/w:ind", W_NS)
    if indent is None:
        raise AssertionError("style indent not found")
    return indent


def _style_tab(style: ET.Element) -> ET.Element:
    tab = style.find("w:pPr/w:tabs/w:tab", W_NS)
    if tab is None:
        raise AssertionError("style tab stop not found")
    return tab


def _style_rfonts(style: ET.Element) -> ET.Element:
    rfonts = style.find("w:rPr/w:rFonts", W_NS)
    if rfonts is None:
        raise AssertionError("style rFonts not found")
    return rfonts


def _style_size(style: ET.Element) -> ET.Element:
    size = style.find("w:rPr/w:sz", W_NS)
    if size is None:
        raise AssertionError("style size not found")
    return size


def _style_color(style: ET.Element) -> ET.Element:
    color = style.find("w:rPr/w:color", W_NS)
    if color is None:
        raise AssertionError("style color not found")
    return color


def _wval(element: ET.Element, attr: str = "val") -> str | None:
    return element.attrib.get(f"{{{W_NS['w']}}}{attr}")


def _toc_instruction(document_xml: str) -> str:
    root = ET.fromstring(document_xml)
    return "".join(node.text or "" for node in root.findall(".//w:instrText", W_NS))


def _document_body_children(document_xml: str) -> list[ET.Element]:
    root = ET.fromstring(document_xml)
    body = root.find("w:body", W_NS)
    if body is None:
        raise AssertionError("document body not found")
    return [child for child in list(body) if child.tag != f"{{{W_NS['w']}}}sectPr"]


def _element_text(element: ET.Element) -> str:
    return "".join(node.text or "" for node in element.findall(".//w:t", W_NS))


def _has_page_break(element: ET.Element) -> bool:
    return any(br.attrib.get(f"{{{W_NS['w']}}}type") == "page" for br in element.findall(".//w:br", W_NS))


def _sect_pr(element: ET.Element) -> ET.Element | None:
    return element.find("w:pPr/w:sectPr", W_NS)


def _sect_orientation(sect_pr: ET.Element | None) -> str | None:
    if sect_pr is None:
        return None
    pg_sz = sect_pr.find("w:pgSz", W_NS)
    if pg_sz is None:
        return None
    return _wval(pg_sz, "orient")


def _body_final_sect_orientation(document_xml: str) -> str | None:
    root = ET.fromstring(document_xml)
    sect_pr = root.find("w:body/w:sectPr", W_NS)
    return _sect_orientation(sect_pr)


def _is_blank_paragraph(element: ET.Element) -> bool:
    return element.tag == f"{{{W_NS['w']}}}p" and not _element_text(element).strip() and not _has_page_break(element)


def _heading_style_value(element: ET.Element) -> str | None:
    p_style = element.find("w:pPr/w:pStyle", W_NS)
    if p_style is None:
        return None
    return _wval(p_style)


def _paragraph_by_text(document_xml: str, text: str) -> ET.Element:
    for child in _document_body_children(document_xml):
        if child.tag == f"{{{W_NS['w']}}}p" and text in _element_text(child):
            return child
    raise AssertionError(f"paragraph not found: {text}")


def _first_run_rpr(element: ET.Element) -> ET.Element:
    r_pr = element.find(".//w:rPr", W_NS)
    if r_pr is None:
        raise AssertionError("run properties not found")
    return r_pr


def _index_by_text(children: list[ET.Element], text: str) -> int:
    for index, child in enumerate(children):
        if text in _element_text(child):
            return index
    raise AssertionError(f"text not found: {text}")


def _index_by_toc_field(children: list[ET.Element]) -> int:
    for index, child in enumerate(children):
        instruction = "".join(node.text or "" for node in child.findall(".//w:instrText", W_NS))
        if "TOC" in instruction.upper():
            return index
    raise AssertionError("TOC field not found")


def _has_page_break_between(children: list[ET.Element], before_text: str, after_text: str) -> bool:
    before = _index_by_text(children, before_text)
    after = _index_by_text(children, after_text)
    return any(_has_page_break(child) for child in children[before + 1 : after])


def _page_breaks_between(children: list[ET.Element], before_text: str, after_text: str) -> list[int]:
    before = _index_by_text(children, before_text)
    after = _index_by_text(children, after_text)
    return [index for index in range(before + 1, after) if _has_page_break(children[index])]


def _paragraph_child_tags(element: ET.Element) -> list[str]:
    return [child.tag.rsplit("}", 1)[-1] for child in list(element)]


def _has_blank_paragraph_between_page_breaks(children: list[ET.Element]) -> bool:
    for index, child in enumerate(children[1:-1], start=1):
        if not _is_blank_paragraph(child):
            continue
        if _has_page_break(children[index - 1]) and _has_page_break(children[index + 1]):
            return True
    return False


def _page_break_count(document_xml: str) -> int:
    root = ET.fromstring(document_xml)
    return sum(1 for br in root.findall(".//w:br", W_NS) if br.attrib.get(f"{{{W_NS['w']}}}type") == "page")


def _write_style(path: Path) -> None:
    toc_style_path = path.with_name("toc_style.json")
    _write_toc_style(toc_style_path)
    style = {
        "schema_version": "business_heading_style.v1",
        "page": {
            "top_cm": 2.0,
            "bottom_cm": 2.0,
            "left_cm": 2.5,
            "right_cm": 2.5,
            "header_top_cm": 1.0,
            "footer_bottom_cm": 0.6,
            "preserve_section_orientation": True,
        },
        "heading": {
            "1": {
                "zh_font": "黑体",
                "en_font": "Arial",
                "size_pt": 16,
                "font_color": "000000",
                "bold": True,
                "align": "center",
                "space_before_pt": 6,
                "space_after_pt": 6,
                "line_spacing": 1.5,
                "first_line_indent_chars": 0,
                "left_indent_cm": 0,
            },
            "2": {
                "zh_font": "楷体",
                "en_font": "Arial",
                "size_pt": 14,
                "font_color": "000000",
                "bold": True,
                "align": "left",
                "space_before_pt": 6,
                "space_after_pt": 6,
                "line_spacing": 1.5,
                "first_line_indent_chars": 0,
                "left_indent_cm": 0,
            },
        },
        "body": {
            "zh_font": "仿宋",
            "en_font": "Arial",
            "size_pt": 12,
            "font_color": "000000",
            "bold": False,
            "align": "both",
            "space_before_pt": 0,
            "space_after_pt": 0,
            "line_spacing": 1.5,
            "first_line_indent_chars": 2,
        },
        "table_cell": {
            "zh_font": "宋体",
            "en_font": "宋体",
            "size_pt": 10.5,
            "font_color": "000000",
            "bold": False,
            "align": "center",
            "line_spacing": 1.0,
            "table_align": "center",
        },
        "toc": {
            "title": "目 录",
            "title_style": "TOC Heading",
            "field_instruction": " TOC \\o \"1-4\" \\h \\z \\u ",
            "placeholder": "目录将在 Word 打开时自动更新",
            "insert_when_missing": True,
            "style_spec_path": str(toc_style_path),
        },
        "header": {
            "zh_font": "宋体",
            "en_font": "宋体",
            "size_pt": 7.5,
            "font_color": "000000",
            "bold": False,
            "align": "right",
            "line_spacing": 1.5,
            "text_template": "{projectName}投标文件-商务部分",
            "residual_terms": ["技术部分", "技术卷", "技术标"],
        },
    }
    path.write_text(json.dumps(style, ensure_ascii=False, indent=2), encoding="utf-8")


def _write_toc_style(path: Path) -> None:
    toc_style = {
        "schema_version": "business_toc_style.v1",
        "source": "test fixture",
        "title": {
            "text": "目 录",
            "style_name": "TOC Heading",
            "zh_font": "宋体",
            "en_font": "Times New Roman",
            "size_pt": 16,
            "font_color": "000000",
            "bold": True,
            "align": "center",
            "first_line_indent_chars": 0,
            "left_indent_cm": 0,
        },
        "field": {
            "instruction": " TOC \\o \"1-4\" \\h \\z \\u ",
            "placeholder": "目录将在 Word 打开时自动更新",
            "style_name": "TOC 1",
            "level": "1",
            "page_break_after": True,
        },
        "entry_styles": {
            "1": {
                "style_name": "TOC 1",
                "zh_font": "宋体",
                "en_font": "Times New Roman",
                "size_pt": 12,
                "font_color": "000000",
                "bold": False,
                "align": "left",
                "space_before_pt": 0,
                "space_after_pt": 0,
                "line_spacing": 1.5,
                "line_spacing_rule": "auto",
                "first_line_indent_twips": 0,
                "left_indent_twips": 0,
                "tab_stops": [
                    {"alignment": "right", "leader": "dot", "position_twips": 9060}
                ],
            },
            "2": {
                "style_name": "TOC 2",
                "zh_font": "宋体",
                "en_font": "Times New Roman",
                "size_pt": 12,
                "font_color": "000000",
                "bold": False,
                "align": "left",
                "space_before_pt": 0,
                "space_after_pt": 0,
                "line_spacing": 1.5,
                "line_spacing_rule": "auto",
                "first_line_indent_twips": 0,
                "left_indent_twips": 420,
                "tab_stops": [
                    {"alignment": "right", "leader": "dot", "position_twips": 9060}
                ],
            },
            "3": {
                "style_name": "TOC 3",
                "zh_font": "宋体",
                "en_font": "Times New Roman",
                "size_pt": 12,
                "font_color": "000000",
                "bold": False,
                "align": "left",
                "space_before_pt": 0,
                "space_after_pt": 0,
                "line_spacing": 1.5,
                "line_spacing_rule": "auto",
                "first_line_indent_twips": 0,
                "left_indent_twips": 840,
                "tab_stops": [
                    {"alignment": "right", "leader": "dot", "position_twips": 9060}
                ],
            },
            "4": {
                "style_name": "TOC 4",
                "zh_font": "宋体",
                "en_font": "Times New Roman",
                "size_pt": 12,
                "font_color": "000000",
                "bold": False,
                "align": "left",
                "space_before_pt": 0,
                "space_after_pt": 0,
                "line_spacing": 1.5,
                "line_spacing_rule": "auto",
                "first_line_indent_twips": 0,
                "left_indent_twips": 1260,
                "tab_stops": [
                    {"alignment": "right", "leader": "dot", "position_twips": 9060}
                ],
            },
        },
    }
    path.write_text(json.dumps(toc_style, ensure_ascii=False, indent=2), encoding="utf-8")


def _write_outline(path: Path) -> None:
    outline = {
        "schema_version": "business_bid_outline.v1",
        "sections": [
            {
                "id": "sec-001",
                "title": "投标函",
                "number": "一、",
                "level": 1,
                "children": [
                    {
                        "id": "sec-001-001",
                        "title": "授权委托书",
                        "number": "1.1",
                        "level": 2,
                        "children": [],
                    },
                    {
                        "id": "sec-001-002",
                        "title": "未出现标题",
                        "number": "1.2",
                        "level": 2,
                        "children": [],
                    },
                ],
            }
        ],
    }
    path.write_text(json.dumps(outline, ensure_ascii=False, indent=2), encoding="utf-8")


def _write_docx(path: Path) -> None:
    doc = Document()
    doc.sections[0].header.paragraphs[0].text = "旧项目投标文件-技术部分 技术卷"
    doc.add_paragraph("一、 投标函")
    doc.add_paragraph("这是正文段落。")
    doc.add_paragraph("授权委托书")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "表格内容"
    doc.save(path)


def _write_docx_with_cover_and_index(path: Path) -> None:
    doc = Document()
    cover = doc.add_paragraph("封面")
    cover.add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("商务评分索引表")
    doc.add_paragraph("一、 投标函")
    doc.add_paragraph("授权委托书")
    doc.save(path)


def _write_docx_with_cover_and_landscape_index(path: Path) -> None:
    doc = Document()
    cover = doc.add_paragraph("COVER")
    cover.add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("BUSINESS_SCORE_INDEX")
    doc.add_paragraph("BODY")
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    doc.save(path)


def _write_outline_for_transparent_page_breaks(path: Path) -> None:
    outline = {
        "schema_version": "business_bid_outline.v1",
        "sections": [
            {
                "id": "sec-001",
                "title": "Existing Break Section",
                "number": "1.1",
                "level": 2,
                "children": [],
            },
            {
                "id": "sec-002",
                "title": "Next Section",
                "number": "1.2",
                "level": 2,
                "children": [],
            },
        ],
    }
    path.write_text(json.dumps(outline, ensure_ascii=False, indent=2), encoding="utf-8")


def _write_outline_for_internal_page_breaks(path: Path) -> None:
    outline = {
        "schema_version": "business_bid_outline.v1",
        "sections": [
            {
                "id": "sec-001",
                "title": "Internal Break Section",
                "number": "1.1",
                "level": 2,
                "children": [],
            },
            {
                "id": "sec-002",
                "title": "Next Section",
                "number": "1.2",
                "level": 2,
                "children": [],
            },
        ],
    }
    path.write_text(json.dumps(outline, ensure_ascii=False, indent=2), encoding="utf-8")


def _write_docx_for_transparent_page_breaks(path: Path, duplicate_page_break: bool = False) -> None:
    doc = Document()
    first = doc.add_paragraph("1.1 Existing Break Section")
    first.style = doc.styles["Heading 2"]
    doc.add_paragraph("Section body")
    doc.add_paragraph("").add_run().add_break(WD_BREAK.PAGE)
    if duplicate_page_break:
        doc.add_paragraph("").add_run().add_break(WD_BREAK.PAGE)
    second = doc.add_paragraph("1.2 Next Section")
    second.style = doc.styles["Heading 2"]
    doc.add_paragraph("Next body")
    doc.save(path)

    with zipfile.ZipFile(path, "r") as zin:
        names = zin.namelist()
        data = {name: zin.read(name) for name in names}
    xml = data["word/document.xml"].decode("utf-8")
    if duplicate_page_break:
        first_break = xml.index('<w:br w:type="page"/>')
        second_break = xml.index('<w:br w:type="page"/>', first_break + 1)
        insert_at = xml.rfind("<w:p", 0, second_break)
    else:
        marker = '<w:pStyle w:val="Heading2"/>'
        second_heading = xml.rindex(marker)
        insert_at = xml.rfind("<w:p", 0, second_heading)
    xml = xml[:insert_at] + '<w:bookmarkEnd w:id="42"/>' + xml[insert_at:]
    data["word/document.xml"] = xml.encode("utf-8")
    tmp = path.with_suffix(".bookmark.tmp.docx")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in names:
            zout.writestr(name, data[name])
    tmp.replace(path)


def _write_docx_for_internal_page_breaks(path: Path) -> None:
    doc = Document()
    first = doc.add_paragraph("1.1 Internal Break Section")
    first.style = doc.styles["Heading 2"]
    doc.add_paragraph("Body before internal break")
    doc.add_paragraph("").add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("Body after internal break")
    doc.add_paragraph("").add_run().add_break(WD_BREAK.PAGE)
    second = doc.add_paragraph("1.2 Next Section")
    second.style = doc.styles["Heading 2"]
    doc.add_paragraph("Next body")
    doc.save(path)


def _write_docx_for_internal_page_breaks_with_boundary_break_inside_content(path: Path) -> None:
    doc = Document()
    first = doc.add_paragraph("1.1 Internal Break Section")
    first.style = doc.styles["Heading 2"]
    doc.add_paragraph("Body before internal break")
    doc.add_paragraph("").add_run().add_break(WD_BREAK.PAGE)
    boundary = doc.add_paragraph("Body after internal break")
    boundary.add_run().add_break(WD_BREAK.PAGE)
    second = doc.add_paragraph("1.2 Next Section")
    second.style = doc.styles["Heading 2"]
    doc.add_paragraph("Next body")
    doc.save(path)


def _write_docx_for_internal_page_breaks_with_non_heading_style_ids(path: Path) -> None:
    _write_docx_for_internal_page_breaks(path)

    with zipfile.ZipFile(path, "r") as zin:
        names = zin.namelist()
        data = {name: zin.read(name) for name in names}
    xml = data["word/document.xml"].decode("utf-8")
    xml = xml.replace('<w:pStyle w:val="Heading2"/>', '<w:pStyle w:val="5"/>')
    data["word/document.xml"] = xml.encode("utf-8")
    tmp = path.with_suffix(".style-id.tmp.docx")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in names:
            zout.writestr(name, data[name])
    tmp.replace(path)


def _write_docx_for_next_heading_with_leading_page_break(path: Path) -> None:
    doc = Document()
    first = doc.add_paragraph("1.1 Internal Break Section")
    first.style = doc.styles["Heading 2"]
    doc.add_paragraph("Section body")
    doc.add_paragraph("").add_run().add_break(WD_BREAK.PAGE)
    second = doc.add_paragraph()
    second.style = doc.styles["Heading 2"]
    second.add_run().add_break(WD_BREAK.PAGE)
    second.add_run("1.2 Next Section")
    doc.add_paragraph("Next body")
    doc.save(path)


def _write_docx_for_toc_trailing_blank_page(path: Path) -> None:
    doc = Document()
    cover = doc.add_paragraph("COVER")
    cover.add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("目 录")
    doc.add_paragraph("目录将在 Word 打开时自动更新")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("").add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("BUSINESS_SCORE_INDEX")
    doc.save(path)

    with zipfile.ZipFile(path, "r") as zin:
        names = zin.namelist()
        data = {name: zin.read(name) for name in names}
    xml = data["word/document.xml"].decode("utf-8")
    marker = "<w:t>目录将在 Word 打开时自动更新</w:t>"
    insert_at = xml.index(marker)
    run_start = xml.rfind("<w:r", 0, insert_at)
    run_end = xml.index("</w:r>", insert_at) + len("</w:r>")
    toc_run = (
        '<w:r><w:fldChar w:fldCharType="begin"/></w:r>'
        '<w:r><w:instrText xml:space="preserve"> TOC \\\\o "1-4" \\\\h \\\\z \\\\u </w:instrText></w:r>'
        '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
        + xml[run_start:run_end]
        + '<w:r><w:fldChar w:fldCharType="end"/></w:r>'
    )
    xml = xml[:run_start] + toc_run + xml[run_end:]
    data["word/document.xml"] = xml.encode("utf-8")
    tmp = path.with_suffix(".toc.tmp.docx")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in names:
            zout.writestr(name, data[name])
    tmp.replace(path)


def _write_outline_for_page_breaks(path: Path) -> None:
    outline = {
        "schema_version": "business_bid_outline.v1",
        "sections": [
            {
                "id": "sec-001",
                "title": "一级有正文",
                "number": "一、",
                "level": 1,
                "children": [],
            },
            {
                "id": "sec-002",
                "title": "一级无正文",
                "number": "二、",
                "level": 1,
                "children": [
                    {
                        "id": "sec-002-001",
                        "title": "二级有表格正文",
                        "number": "2.1",
                        "level": 2,
                        "children": [],
                    }
                ],
            },
            {
                "id": "sec-003",
                "title": "末尾章节",
                "number": "三、",
                "level": 1,
                "children": [],
            },
        ],
    }
    path.write_text(json.dumps(outline, ensure_ascii=False, indent=2), encoding="utf-8")


def _write_docx_for_page_breaks(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("一、 一级有正文")
    doc.add_paragraph("第一节正文")
    doc.add_paragraph("")
    doc.add_paragraph("二、 一级无正文")
    doc.add_paragraph("二级有表格正文")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "表格正文"
    doc.add_paragraph("")
    doc.add_paragraph("").add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("")
    doc.add_paragraph("").add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("三、 末尾章节")
    doc.add_paragraph("末尾正文")
    doc.save(path)


class BusinessFormatCleanerTest(unittest.TestCase):
    def setUp(self) -> None:
        import tempfile

        self._tmpdir = tempfile.TemporaryDirectory()
        self.tmp_path = Path(self._tmpdir.name)

    def tearDown(self) -> None:
        self._tmpdir.cleanup()

    def test_flatten_business_outline_recurses_sections_and_children(self):
        from outline_matcher import flatten_outline, load_outline

        outline_path = self.tmp_path / "outline.json"
        _write_outline(outline_path)

        items = flatten_outline(load_outline(outline_path))

        self.assertEqual([item["id"] for item in items], ["sec-001", "sec-001-001", "sec-001-002"])
        self.assertEqual([item["title"] for item in items], ["投标函", "授权委托书", "未出现标题"])
        self.assertEqual([item["number"] for item in items], ["一、", "1.1", "1.2"])
        self.assertEqual([item["level"] for item in items], [1, 2, 2])

    def test_builtin_style_writes_official_font_names_to_docx_xml(self):
        input_docx = self.tmp_path / "input_builtin_fonts.docx"
        outline_path = self.tmp_path / "outline_builtin_fonts.json"
        output_docx = self.tmp_path / "output_builtin_fonts.docx"
        manifest_path = self.tmp_path / "manifest_builtin_fonts.json"

        _write_docx(input_docx)
        _write_outline(outline_path)
        manifest_path.write_text(
            json.dumps(
                {
                    "inputFile": str(input_docx),
                    "outlineFile": str(outline_path),
                    "outputFile": str(output_docx),
                    "projectName": "测试项目",
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )

        _run_manifest(manifest_path, response="summary")

        with zipfile.ZipFile(output_docx) as archive:
            document_xml = "\n".join(
                archive.read(name).decode("utf-8", errors="ignore")
                for name in archive.namelist()
                if name.startswith("word/") and name.endswith(".xml")
            )
        for family in ("等线", "等线 Light", "宋体", "Times New Roman"):
            self.assertIn(family, document_xml)
        for family in ("Noto Sans CJK SC", "Noto Serif CJK SC", "Liberation Serif"):
            self.assertNotIn(family, document_xml)

    def test_run_manifest_promotes_plain_body_headings_and_preserves_input(self):

        input_docx = self.tmp_path / "input.docx"
        outline_path = self.tmp_path / "outline.json"
        style_path = self.tmp_path / "style.json"
        output_docx = self.tmp_path / "output.docx"
        manifest_path = self.tmp_path / "manifest.json"

        _write_docx(input_docx)
        _write_outline(outline_path)
        _write_style(style_path)
        original_bytes = input_docx.read_bytes()

        manifest = {
            "inputFile": str(input_docx),
            "outlineFile": str(outline_path),
            "outputFile": str(output_docx),
            "projectName": "测试项目",
            "styleSpecPath": str(style_path),
        }
        manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

        response = _run_manifest(manifest_path, response="summary")

        self.assertEqual(input_docx.read_bytes(), original_bytes)
        self.assertEqual(response["schema_version"], "bid-business-format-clean-v1")
        self.assertEqual(response["inputFile"], str(input_docx))
        self.assertEqual(response["outlineFile"], str(outline_path))
        self.assertEqual(response["outputFile"], str(output_docx))
        report_path = Path(response["reportFile"])
        self.assertTrue(output_docx.exists())
        self.assertTrue(report_path.exists())

        cleaned = Document(output_docx)
        headings = [(p.text, p.style.name) for p in cleaned.paragraphs if p.style.name.startswith("Heading")]
        self.assertIn(("一、 投标函", "Heading 1"), headings)
        self.assertIn(("1.1 授权委托书", "Heading 2"), headings)

        header_text = "\n".join(paragraph.text for paragraph in cleaned.sections[0].header.paragraphs)
        self.assertIn("测试项目投标文件-商务部分", header_text)
        self.assertNotIn("技术部分", header_text)
        self.assertNotIn("技术卷", header_text)

        document_xml = _read_docx_part(output_docx, "word/document.xml")
        settings_xml = _read_docx_part(output_docx, "word/settings.xml")
        self.assertIn("TOC", document_xml)
        self.assertIn("updateFields", settings_xml)
        self.assertNotIn("<w:numPr>", document_xml)
        self.assertIn("仿宋", document_xml)
        self.assertIn("宋体", document_xml)
        styles_xml = _read_docx_part(output_docx, "word/styles.xml")
        heading1_style = _style_by_name(styles_xml, "heading 1")
        self.assertEqual(_wval(_style_color(heading1_style)), "000000")
        self.assertIsNone(_wval(_style_color(heading1_style), "themeColor"))

        report = report_path.read_text(encoding="utf-8")
        self.assertIn("outline 总数：3", report)
        self.assertIn("成功匹配标题数：2", report)
        self.assertIn("1.2 未出现标题", report)
        self.assertIn("TOC 是否插入：是", report)
        self.assertIn("页眉是否清理：是", report)

    def test_run_manifest_fuzzy_matches_numbered_headings_and_cleans_residual_heading_styles(self):
        input_docx = self.tmp_path / "input_fuzzy.docx"
        outline_path = self.tmp_path / "outline.json"
        style_path = self.tmp_path / "style.json"
        output_docx = self.tmp_path / "output_fuzzy.docx"
        manifest_path = self.tmp_path / "manifest_fuzzy.json"

        doc = Document()
        first = doc.add_paragraph()
        first.add_run("1. 投标函：").bold = False
        first.add_run(" 附加说明").italic = True
        stale_heading = doc.add_paragraph("旧模板残留标题")
        stale_heading.style = doc.styles["Heading 2"]
        stale_heading.add_run("，应转为正文").bold = True
        body = doc.add_paragraph()
        run = body.add_run("正文应统一格式")
        run.bold = True
        run.italic = True
        doc.save(input_docx)

        outline = {
            "schema_version": "business_bid_outline.v1",
            "sections": [{"id": "sec-001", "title": "投标函", "number": "一、", "level": 1, "children": []}],
        }
        outline_path.write_text(json.dumps(outline, ensure_ascii=False, indent=2), encoding="utf-8")
        _write_style(style_path)
        manifest_path.write_text(
            json.dumps(
                {
                    "inputFile": str(input_docx),
                    "outlineFile": str(outline_path),
                    "outputFile": str(output_docx),
                    "projectName": "测试项目",
                    "styleSpecPath": str(style_path),
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )

        response = _run_manifest(manifest_path, response="summary")

        self.assertEqual(response["summary"]["matchedHeadingCount"], 1)
        document_xml = _read_docx_part(output_docx, "word/document.xml")
        promoted = _paragraph_by_text(document_xml, "一、 投标函")
        stale = _paragraph_by_text(document_xml, "旧模板残留标题")
        body_para = _paragraph_by_text(document_xml, "正文应统一格式")
        self.assertEqual(_heading_style_value(promoted), "Heading1")
        promoted_color = _first_run_rpr(promoted).find("w:color", W_NS)
        self.assertIsNotNone(promoted_color)
        self.assertEqual(_wval(promoted_color), "000000")
        self.assertIsNone(_wval(promoted_color, "themeColor"))
        self.assertNotEqual(_heading_style_value(stale), "Heading2")
        body_rpr = _first_run_rpr(body_para)
        self.assertEqual(_wval(body_rpr.find("w:b", W_NS)), "0")
        self.assertEqual(_wval(body_rpr.find("w:i", W_NS)), "0")
        self.assertEqual(_wval(body_rpr.find("w:rFonts", W_NS), "eastAsia"), "仿宋")

    def test_toc_title_and_entry_styles_are_loaded_from_reference(self):

        input_docx = self.tmp_path / "input.docx"
        outline_path = self.tmp_path / "outline.json"
        style_path = self.tmp_path / "style.json"
        output_docx = self.tmp_path / "output.docx"
        manifest_path = self.tmp_path / "manifest.json"

        _write_docx(input_docx)
        _write_outline(outline_path)
        _write_style(style_path)

        manifest = {
            "inputFile": str(input_docx),
            "outlineFile": str(outline_path),
            "outputFile": str(output_docx),
            "projectName": "测试项目",
            "styleSpecPath": str(style_path),
        }
        manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

        _run_manifest(manifest_path, response="summary")

        document_xml = _read_docx_part(output_docx, "word/document.xml")
        styles_xml = _read_docx_part(output_docx, "word/styles.xml")

        self.assertIn("目 录", document_xml)
        self.assertIn("目录将在 Word 打开时自动更新", document_xml)
        self.assertIn('w:val="TOC1"', document_xml)
        self.assertEqual(_toc_instruction(document_xml), " TOC \\o \"1-4\" \\h \\z \\u ")
        self.assertLess(document_xml.index("目 录"), document_xml.index("目录将在 Word 打开时自动更新"))

        toc1 = _style_by_name(styles_xml, "TOC 1")
        toc2 = _style_by_name(styles_xml, "TOC 2")
        toc3 = _style_by_name(styles_xml, "TOC 3")
        for style in (toc1, toc2, toc3):
            self.assertEqual(_wval(_style_spacing(style), "line"), "360")
            self.assertEqual(_wval(_style_spacing(style), "lineRule"), "auto")
            self.assertEqual(_wval(_style_tab(style), "val"), "right")
            self.assertEqual(_wval(_style_tab(style), "leader"), "dot")
            self.assertEqual(_wval(_style_tab(style), "pos"), "9060")
            self.assertEqual(_wval(_style_rfonts(style), "eastAsia"), "宋体")
            self.assertEqual(_wval(_style_rfonts(style), "ascii"), "Times New Roman")
            self.assertEqual(_wval(_style_size(style)), "24")
        self.assertEqual(_wval(_style_indent(toc1), "left"), "0")
        self.assertEqual(_wval(_style_indent(toc1), "firstLine"), "0")
        self.assertEqual(_wval(_style_indent(toc2), "left"), "420")
        self.assertEqual(_wval(_style_indent(toc3), "left"), "840")
        toc4 = _style_by_name(styles_xml, "TOC 4")
        if toc4 is not None:
            self.assertEqual(_wval(_style_indent(toc4), "left"), "1260")

    def test_builtin_toc_reference_matches_reference_document_format(self):
        style = json.loads((ROOT / "references" / "business_toc_style.json").read_text(encoding="utf-8"))

        self.assertEqual(style["title"]["text"], "目 录")
        self.assertEqual(style["title"]["zh_font"], "宋体")
        self.assertEqual(style["title"]["en_font"], "Times New Roman")
        self.assertEqual(style["title"]["size_pt"], 16)
        self.assertEqual(style["field"]["instruction"], " TOC \\o \"1-4\" \\h \\z \\u ")
        self.assertEqual(set(style["entry_styles"].keys()), {"1", "2", "3", "4"})
        self.assertEqual(style["entry_styles"]["1"]["left_indent_twips"], 0)
        self.assertEqual(style["entry_styles"]["2"]["left_indent_twips"], 420)
        self.assertEqual(style["entry_styles"]["3"]["left_indent_twips"], 840)
        self.assertEqual(style["entry_styles"]["4"]["left_indent_twips"], 1260)
        for entry in style["entry_styles"].values():
            self.assertEqual(entry["zh_font"], "宋体")
            self.assertEqual(entry["en_font"], "Times New Roman")
            self.assertEqual(entry["size_pt"], 12)
            self.assertEqual(entry["line_spacing"], 1.5)
            self.assertEqual(entry["line_spacing_rule"], "auto")
            self.assertEqual(entry["tab_stops"][0]["position_twips"], 9060)
            self.assertEqual(entry["tab_stops"][0]["leader"], "dot")

    def test_toc_is_inserted_after_cover_page_break_before_business_score_index(self):

        input_docx = self.tmp_path / "input.docx"
        outline_path = self.tmp_path / "outline.json"
        style_path = self.tmp_path / "style.json"
        output_docx = self.tmp_path / "output.docx"
        manifest_path = self.tmp_path / "manifest.json"

        _write_docx_with_cover_and_index(input_docx)
        _write_outline(outline_path)
        _write_style(style_path)

        manifest = {
            "inputFile": str(input_docx),
            "outlineFile": str(outline_path),
            "outputFile": str(output_docx),
            "projectName": "测试项目",
            "styleSpecPath": str(style_path),
        }
        manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

        _run_manifest(manifest_path, response="summary")

        document_xml = _read_docx_part(output_docx, "word/document.xml")
        self.assertLess(document_xml.index("封面"), document_xml.index("目 录"))
        self.assertLess(document_xml.index("目 录"), document_xml.index("商务评分索引表"))

    def test_inserted_toc_gets_portrait_section_before_landscape_score_index(self):

        input_docx = self.tmp_path / "input.docx"
        outline_path = self.tmp_path / "outline.json"
        style_path = self.tmp_path / "style.json"
        output_docx = self.tmp_path / "output.docx"
        manifest_path = self.tmp_path / "manifest.json"

        _write_docx_with_cover_and_landscape_index(input_docx)
        outline_path.write_text(
            json.dumps(
                {
                    "schema_version": "business_bid_outline.v1",
                    "sections": [],
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        _write_style(style_path)

        manifest = {
            "inputFile": str(input_docx),
            "outlineFile": str(outline_path),
            "outputFile": str(output_docx),
            "projectName": "Project",
            "styleSpecPath": str(style_path),
        }
        manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

        _run_manifest(manifest_path, response="summary")

        document_xml = _read_docx_part(output_docx, "word/document.xml")
        children = _document_body_children(document_xml)
        toc_index = _index_by_toc_field(children)
        score_index = _index_by_text(children, "BUSINESS_SCORE_INDEX")
        toc_section_breaks = [
            _sect_orientation(sect_pr)
            for child in children[toc_index:score_index]
            for sect_pr in [_sect_pr(child)]
            if sect_pr is not None
        ]

        self.assertEqual(toc_section_breaks, [None])
        self.assertEqual(_body_final_sect_orientation(document_xml), "landscape")

    def test_page_break_cleanup_adds_missing_breaks_after_section_body_and_collapses_blank_pages(self):

        input_docx = self.tmp_path / "input.docx"
        outline_path = self.tmp_path / "outline.json"
        style_path = self.tmp_path / "style.json"
        output_docx = self.tmp_path / "output.docx"
        manifest_path = self.tmp_path / "manifest.json"

        _write_docx_for_page_breaks(input_docx)
        _write_outline_for_page_breaks(outline_path)
        _write_style(style_path)

        manifest = {
            "inputFile": str(input_docx),
            "outlineFile": str(outline_path),
            "outputFile": str(output_docx),
            "projectName": "测试项目",
            "styleSpecPath": str(style_path),
        }
        manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

        response = _run_manifest(manifest_path, response="details")

        document_xml = _read_docx_part(output_docx, "word/document.xml")
        children = _document_body_children(document_xml)
        self.assertTrue(_has_page_break_between(children, "第一节正文", "二、 一级无正文"))
        self.assertFalse(_has_page_break_between(children, "二、 一级无正文", "2.1 二级有表格正文"))
        self.assertTrue(_has_page_break_between(children, "表格正文", "三、 末尾章节"))
        self.assertFalse(_has_blank_paragraph_between_page_breaks(children))
        self.assertEqual(response["details"]["cleanResult"]["pagination"]["insertedPageBreaks"], 1)
        self.assertGreaterEqual(response["details"]["cleanResult"]["pagination"]["removedBlankPageBreaks"], 1)
        self.assertEqual(response["summary"]["insertedPageBreaks"], 1)
        self.assertGreaterEqual(response["summary"]["removedBlankPageBreaks"], 1)

    def test_page_break_cleanup_is_idempotent(self):
        from pagination_cleaner import clean_section_page_breaks

        input_docx = self.tmp_path / "input.docx"
        outline_path = self.tmp_path / "outline.json"
        output_docx = self.tmp_path / "output.docx"

        _write_docx_for_page_breaks(input_docx)
        _write_outline_for_page_breaks(outline_path)
        input_docx.replace(output_docx)

        first = clean_section_page_breaks(output_docx, outline_path)
        after_first_xml = _read_docx_part(output_docx, "word/document.xml")
        second = clean_section_page_breaks(output_docx, outline_path)
        after_second_xml = _read_docx_part(output_docx, "word/document.xml")

        self.assertEqual(first["insertedPageBreaks"], 1)
        self.assertEqual(second["insertedPageBreaks"], 0)
        self.assertEqual(_page_break_count(after_first_xml), _page_break_count(after_second_xml))

    def test_existing_break_before_transparent_bookmark_does_not_get_duplicated(self):
        from pagination_cleaner import clean_section_page_breaks

        input_docx = self.tmp_path / "input.docx"
        outline_path = self.tmp_path / "outline.json"

        _write_docx_for_transparent_page_breaks(input_docx)
        _write_outline_for_transparent_page_breaks(outline_path)
        before_xml = _read_docx_part(input_docx, "word/document.xml")

        result = clean_section_page_breaks(input_docx, outline_path)

        after_xml = _read_docx_part(input_docx, "word/document.xml")
        self.assertEqual(result["insertedPageBreaks"], 0)
        self.assertEqual(_page_break_count(after_xml), _page_break_count(before_xml))

    def test_duplicate_breaks_separated_by_transparent_bookmark_are_collapsed(self):
        from pagination_cleaner import clean_section_page_breaks

        input_docx = self.tmp_path / "input.docx"
        outline_path = self.tmp_path / "outline.json"

        _write_docx_for_transparent_page_breaks(input_docx, duplicate_page_break=True)
        _write_outline_for_transparent_page_breaks(outline_path)

        result = clean_section_page_breaks(input_docx, outline_path)

        after_xml = _read_docx_part(input_docx, "word/document.xml")
        self.assertEqual(_page_break_count(after_xml), 1)
        self.assertEqual(result["removedBlankPageBreaks"], 1)

    def test_internal_section_breaks_are_removed_when_boundary_break_exists(self):
        from pagination_cleaner import clean_section_page_breaks

        input_docx = self.tmp_path / "input.docx"
        outline_path = self.tmp_path / "outline.json"

        _write_docx_for_internal_page_breaks(input_docx)
        _write_outline_for_internal_page_breaks(outline_path)

        result = clean_section_page_breaks(input_docx, outline_path)

        after_xml = _read_docx_part(input_docx, "word/document.xml")
        children = _document_body_children(after_xml)
        self.assertEqual(_page_break_count(after_xml), 1)
        self.assertFalse(_has_page_break_between(children, "Body before internal break", "Body after internal break"))
        self.assertTrue(_has_page_break_between(children, "Body after internal break", "1.2 Next Section"))
        self.assertEqual(result["removedBlankPageBreaks"], 1)

    def test_outline_matching_drives_pagination_when_heading_style_ids_are_localized(self):
        from pagination_cleaner import clean_section_page_breaks

        input_docx = self.tmp_path / "input.docx"
        outline_path = self.tmp_path / "outline.json"

        _write_docx_for_internal_page_breaks_with_non_heading_style_ids(input_docx)
        _write_outline_for_internal_page_breaks(outline_path)

        result = clean_section_page_breaks(input_docx, outline_path)

        after_xml = _read_docx_part(input_docx, "word/document.xml")
        children = _document_body_children(after_xml)
        self.assertEqual(_page_break_count(after_xml), 1)
        self.assertFalse(_has_page_break_between(children, "Body before internal break", "Body after internal break"))
        self.assertTrue(_has_page_break_between(children, "Body after internal break", "1.2 Next Section"))
        self.assertEqual(result["removedBlankPageBreaks"], 1)

    def test_internal_break_is_removed_when_boundary_break_is_inside_last_content_paragraph(self):
        from pagination_cleaner import clean_section_page_breaks

        input_docx = self.tmp_path / "input.docx"
        outline_path = self.tmp_path / "outline.json"

        _write_docx_for_internal_page_breaks_with_boundary_break_inside_content(input_docx)
        _write_outline_for_internal_page_breaks(outline_path)

        result = clean_section_page_breaks(input_docx, outline_path)

        after_xml = _read_docx_part(input_docx, "word/document.xml")
        children = _document_body_children(after_xml)
        boundary_index = _index_by_text(children, "Body after internal break")
        self.assertEqual(_page_break_count(after_xml), 1)
        self.assertFalse(_has_page_break_between(children, "Body before internal break", "Body after internal break"))
        self.assertTrue(_has_page_break(children[boundary_index]))
        self.assertEqual(result["removedBlankPageBreaks"], 1)

    def test_leading_page_break_is_removed_from_next_heading_when_boundary_break_exists(self):
        from pagination_cleaner import clean_section_page_breaks

        input_docx = self.tmp_path / "input.docx"
        outline_path = self.tmp_path / "outline.json"

        _write_docx_for_next_heading_with_leading_page_break(input_docx)
        _write_outline_for_internal_page_breaks(outline_path)

        result = clean_section_page_breaks(input_docx, outline_path)

        after_xml = _read_docx_part(input_docx, "word/document.xml")
        children = _document_body_children(after_xml)
        next_heading_index = _index_by_text(children, "1.2 Next Section")
        self.assertEqual(_page_break_count(after_xml), 1)
        self.assertEqual(_page_breaks_between(children, "Section body", "1.2 Next Section"), [next_heading_index - 1])
        self.assertFalse(_has_page_break(children[next_heading_index]))
        self.assertEqual(result["removedBlankPageBreaks"], 1)

    def test_toc_trailing_blank_paragraphs_and_page_break_are_removed_before_score_index(self):
        from pagination_cleaner import clean_section_page_breaks

        input_docx = self.tmp_path / "input.docx"

        _write_docx_for_toc_trailing_blank_page(input_docx)

        result = clean_section_page_breaks(input_docx)

        after_xml = _read_docx_part(input_docx, "word/document.xml")
        children = _document_body_children(after_xml)
        toc_index = _index_by_toc_field(children)
        score_index = _index_by_text(children, "BUSINESS_SCORE_INDEX")
        between = children[toc_index + 1 : score_index]
        self.assertFalse(any(_has_page_break(child) for child in between))
        self.assertFalse(any(_is_blank_paragraph(child) for child in between))
        self.assertGreaterEqual(result["removedBlankParagraphs"], 2)
        self.assertEqual(result["removedBlankPageBreaks"], 1)


if __name__ == "__main__":
    unittest.main()
