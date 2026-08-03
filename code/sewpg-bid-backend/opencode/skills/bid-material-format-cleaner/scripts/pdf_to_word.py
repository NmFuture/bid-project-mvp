#!/usr/bin/env python3
"""
pdf_to_word.py — 将 PDF 按页拆分成图片并合并成 Word 文档。

用法:
    python3 pdf_to_word.py <pdf_or_dir> [<pdf_or_dir> ...]

每个 PDF 的处理流程:
  1. 在 PDF 所在目录创建与 PDF 同名的子文件夹
  2. 用 PyMuPDF 将每页导出为 page_N.png (2x 缩放)
  3. 将图片按页码合并到一个 docx 文件中, 放在 PDF 所在目录

Word 文档格式严格按投标技术标模板:
  - A4 竖排, 页边距 上下1.5cm / 左右2cm
  - 行距 1.15
  - P0: 标题(默认 Heading 1, 加粗, Noto Serif CJK SC, 首行缩进482 twips, 左对齐)
  - P1: 所有图片(首行缩进0, 图片尺寸 414.85pt x 586.75pt)
"""

from __future__ import annotations

import argparse
import glob
import os
import sys
from pathlib import Path
from typing import Iterable, List


def _check_deps() -> None:
    """检查依赖是否可用, 给出明确的安装提示。"""
    missing = []
    try:
        import fitz  # noqa: F401  (PyMuPDF)
    except ImportError:
        missing.append("pymupdf")
    try:
        import docx  # noqa: F401
    except ImportError:
        missing.append("python-docx")
    if missing:
        sys.stderr.write(
            f"缺少依赖: {', '.join(missing)}\n"
            f"请运行: pip install {' '.join(missing)}\n"
        )
        sys.exit(1)


# ---------- 1. PDF -> images ----------


def split_pdf_to_images(pdf_path: Path, out_dir: Path, zoom: float = 2.0) -> List[Path]:
    """将 PDF 按页拆分为 PNG, 返回按页码排序的图片路径列表。

    zoom=2.0 对应约 144 dpi, 清晰度较好且文件大小适中。
    """
    import fitz

    out_dir.mkdir(parents=True, exist_ok=True)
    doc = fitz.open(str(pdf_path))
    mat = fitz.Matrix(zoom, zoom)
    paths: List[Path] = []
    try:
        for i, page in enumerate(doc):
            pix = page.get_pixmap(matrix=mat)
            img_path = out_dir / f"page_{i + 1}.png"
            pix.save(str(img_path))
            paths.append(img_path)
    finally:
        doc.close()
    return paths


def collect_existing_images(out_dir: Path) -> List[Path]:
    """若同名子文件夹已有 page_N.png, 直接复用, 避免重复工作。"""
    files = glob.glob(str(out_dir / "page_*.png"))
    if not files:
        return []
    return sorted(
        (Path(f) for f in files),
        key=lambda p: int(p.stem.split("_")[1]),
    )


# ---------- 2. images -> docx ----------


def build_docx(title: str, image_paths: List[Path], out_path: Path) -> None:
    """生成符合模板的 Word 文档。

    模板关键点:
      - 页面 EMU 值与参考模板精确匹配
      - P0 标题: 加粗, 首行缩进 482 twips, Noto Serif CJK SC, 左对齐
      - P1 图片: 首行缩进 0, 所有图片放同一段落, 每张 414.85pt x 586.75pt
    """
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    IMG_WIDTH = Pt(414.85)
    IMG_HEIGHT = Pt(586.75)

    doc = Document()

    # 页面设置 (精确 EMU)
    section = doc.sections[0]
    section.page_width = 7560310
    section.page_height = 10692130
    section.top_margin = 540385
    section.bottom_margin = 540385
    section.left_margin = 720090
    section.right_margin = 720090
    section.header_distance = 360045
    section.footer_distance = 215900

    # Normal 样式: 行距 1.15, 段前段后 0
    style = doc.styles["Normal"]
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # 删除默认空段落, 避免多一行间距
    body = doc.element.body
    for p in body.findall(qn("w:p")):
        body.remove(p)

    def _set_serif_font(run) -> None:
        """给 run 设置开源中文衬线字体。"""
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), "Noto Serif CJK SC")

    def _set_first_line_indent(paragraph, twips: int) -> None:
        pPr = paragraph._element.get_or_add_pPr()
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        ind.set(qn("w:firstLine"), str(twips))

    # ---- P0: 标题 + 图片（同一段落，标题默认带 Heading 1） ----
    # 将标题和图片放在同一段落中，Word 会让它们紧挨着排在第一页
    # 不会在标题和第一张图片之间插入分页符
    p0 = doc.add_paragraph()
    for style_name in ("Heading 1", "标题 1"):
        try:
            p0.style = doc.styles[style_name]
            break
        except KeyError:
            continue
    p0.paragraph_format.line_spacing = 1.15
    _set_first_line_indent(p0, 482)
    title_run = p0.add_run(title)
    title_run.bold = True
    _set_serif_font(title_run)

    # 图片紧接着放在同一段落
    for img in image_paths:
        r = p0.add_run()
        r.add_picture(str(img), width=IMG_WIDTH, height=IMG_HEIGHT)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


# ---------- 3. 主流程 ----------


def process_pdf(pdf_path: Path, output_dir: Path | None = None) -> Path:
    """处理单个 PDF, 返回生成的 docx 路径。

    Args:
        pdf_path: 源 PDF 文件路径
        output_dir: 指定输出目录（可选）。若提供，图片和 docx 均写入此目录，
                    不在原始 PDF 同目录下留下中间产物。若为 None，
                    则行为与原来一致（输出到 PDF 同目录）。
    """
    pdf_path = pdf_path.resolve()
    if not pdf_path.exists():
        raise FileNotFoundError(pdf_path)

    stem = pdf_path.stem  # 不含扩展名的文件名

    if output_dir is not None:
        # 所有产物写入指定目录，与原始 PDF 所在目录完全隔离
        img_dir = output_dir / stem
        docx_path = output_dir / f"{stem}.docx"
        img_dir.mkdir(parents=True, exist_ok=True)

        print(f"  拆分 PDF -> {img_dir}")
        images = split_pdf_to_images(pdf_path, img_dir)
        print(f"  已导出 {len(images)} 张图片")

        print(f"  生成 Word: {docx_path.name}")
        build_docx(title=stem, image_paths=images, out_path=docx_path)
        return docx_path

    # 兼容模式：不指定 output_dir 时，行为与原版一致
    base = pdf_path.parent
    img_dir = base / stem
    docx_path = base / f"{stem}.docx"

    # 尝试复用已有的图片
    images = collect_existing_images(img_dir)
    if images:
        print(f"  复用已有图片 ({len(images)} 张): {img_dir}")
    else:
        print(f"  拆分 PDF -> {img_dir}")
        images = split_pdf_to_images(pdf_path, img_dir)
        print(f"  已导出 {len(images)} 张图片")

    print(f"  生成 Word: {docx_path.name}")
    build_docx(title=stem, image_paths=images, out_path=docx_path)
    return docx_path


def expand_inputs(inputs: Iterable[str]) -> List[Path]:
    """将输入(文件或目录)展开为 PDF 路径列表。"""
    pdfs: List[Path] = []
    for item in inputs:
        p = Path(item)
        if p.is_dir():
            pdfs.extend(sorted(p.glob("*.pdf")))
        elif p.suffix.lower() == ".pdf":
            pdfs.append(p)
        else:
            sys.stderr.write(f"跳过非 PDF: {item}\n")
    return pdfs


def main(argv: List[str] | None = None) -> int:
    _check_deps()

    parser = argparse.ArgumentParser(
        description="将 PDF 按页拆成图片并合并成 Word 文档"
    )
    parser.add_argument(
        "inputs",
        nargs="+",
        help="一个或多个 PDF 文件或包含 PDF 的目录",
    )
    parser.add_argument(
        "--out-dir",
        dest="output_dir",
        metavar="DIR",
        help="所有输出（图片子文件夹 + docx）写入指定目录，不再污染原始 PDF 所在目录",
    )
    args = parser.parse_args(argv)

    pdfs = expand_inputs(args.inputs)
    if not pdfs:
        sys.stderr.write("未找到 PDF 文件\n")
        return 1

    out_dir = Path(args.output_dir) if args.output_dir else None
    if out_dir is not None:
        out_dir.mkdir(parents=True, exist_ok=True)

    print(f"共 {len(pdfs)} 个 PDF:")
    for p in pdfs:
        print(f"  - {p}")
    print()

    for pdf in pdfs:
        print(f"[{pdf.name}]")
        try:
            process_pdf(pdf, output_dir=out_dir)
        except Exception as e:
            sys.stderr.write(f"  失败: {e}\n")
            continue
        print()

    print("全部完成")
    return 0


if __name__ == "__main__":
    sys.exit(main())
