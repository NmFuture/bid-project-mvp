#!/usr/bin/env python3
"""
word_cleaner.py — officecli 的纯 Python 后备方案，用于 Word 文档清洗。

当 officecli 因 Git Bash 路径转换、中文编码等问题失败时，
使用此脚本替代 officecli 完成探针和切割操作。

子命令:
    peek <docx>                                  提取前100段的纯文本+样式名
    list-body <docx>                             列出 body 所有一级子元素
    trim <docx> --anchor-para-id ID [--anchor-text TEXT]
                                                优先按 paraId（可选文本校验）定位锚点
    trim <docx> --anchor-index N [--anchor-text TEXT]
                                                仅在 paraId 不可用时按 body index 定位
    normalize <docx>                             格式规范化（标题去编号、清除空标题元数据、合并连续空行、清理文末空白）

中文路径兼容: 自动复制到临时 ASCII 路径处理，再拷回。
"""

from __future__ import annotations

import argparse
import io
import json
import os
import re
import secrets
import shutil
import sys
import tempfile
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path


_BLANK_HEADING_BOOKMARK_PREFIX = "_MaterialCleanerBlankHeading_"

# Windows 终端默认使用 GBK 编码，无法稳定输出中文 Unicode 字符。
# 强制 stdout/stderr 使用 UTF-8，确保 JSON、日志和报错正常输出。
if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
if sys.stderr.encoding and sys.stderr.encoding.lower() != "utf-8":
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")


def _check_deps() -> None:
    missing = []
    try:
        import docx  # noqa: F401
    except ImportError:
        missing.append("python-docx")
    try:
        import lxml  # noqa: F401
    except ImportError:
        missing.append("lxml")
    if missing:
        sys.stderr.write(
            f"缺少依赖: {', '.join(missing)}\n"
            f"请运行: pip install {' '.join(missing)}\n"
        )
        sys.exit(1)


def _trust_path_mode() -> bool:
    """是否启用受信任路径模式（由 driver 的单临时副本事务使用）。"""
    return os.environ.get("FORMAT_CLEANER_TRUST_PATH") == "1"


def _replace_with_retry(src: str, dst: str, retries: int = 5, delay: float = 0.2) -> None:
    """原子替换文件，遇到短暂锁定时重试。"""
    import time
    for attempt in range(retries):
        try:
            os.replace(src, dst)
            return
        except PermissionError:
            if attempt < retries - 1:
                time.sleep(delay)
            else:
                raise


def _safe_open(docx_path: Path):
    """打开 docx 文件，中文路径时自动复制到临时目录再打开。"""
    from docx import Document

    if _trust_path_mode():
        return Document(str(docx_path)), None

    try:
        return Document(str(docx_path)), None
    except Exception:
        # 中文路径可能导致 python-docx 失败，复制到临时 ASCII 路径
        tmp_dir = tempfile.mkdtemp(prefix="fc_")
        tmp_path = Path(tmp_dir) / "doc.docx"
        shutil.copy2(str(docx_path), str(tmp_path))
        return Document(str(tmp_path)), tmp_dir


def _copy_with_retry(src: str, dst: str, retries: int = 5, delay: float = 1.0) -> None:
    """拷贝文件，遇到 PermissionError 时重试。

    Windows 上 officecli、杀毒软件、搜索索引等可能短暂锁定文件，
    导致紧接着的写入失败。等待片刻后重试通常能解决。
    """
    import time
    for attempt in range(retries):
        try:
            shutil.copy2(src, dst)
            return
        except PermissionError:
            if attempt < retries - 1:
                time.sleep(delay)
            else:
                raise


def _safe_save(doc, docx_path: Path, tmp_dir: str | None) -> None:
    """保存 docx，始终先写临时文件再拷回原路径（避免文件锁冲突）。"""
    if _trust_path_mode() and tmp_dir is None:
        tmp_path = docx_path.with_name(f".{docx_path.stem}.tmp{docx_path.suffix}")
        try:
            if tmp_path.exists():
                tmp_path.unlink()
            doc.save(str(tmp_path))
            _replace_with_retry(str(tmp_path), str(docx_path))
        finally:
            if tmp_path.exists():
                tmp_path.unlink(missing_ok=True)
        return

    if tmp_dir is not None:
        tmp_path = Path(tmp_dir) / "doc.docx"
        doc.save(str(tmp_path))
        _copy_with_retry(str(tmp_path), str(docx_path))
        shutil.rmtree(tmp_dir, ignore_errors=True)
    else:
        # 即使无中文路径问题，也先存到临时文件再拷回
        # 避免 officecli/杀毒/索引短暂锁定原文件导致 PermissionError
        save_tmp = tempfile.mkdtemp(prefix="fc_save_")
        save_path = Path(save_tmp) / "doc.docx"
        try:
            doc.save(str(save_path))
            _copy_with_retry(str(save_path), str(docx_path))
        finally:
            shutil.rmtree(save_tmp, ignore_errors=True)


def _get_element_tag(el) -> str:
    """获取元素的简短标签名（去掉命名空间）。"""
    tag = el.tag
    if "}" in tag:
        tag = tag.split("}")[1]
    return tag


def _get_element_text(el) -> str:
    """递归提取元素内所有文本。"""
    from docx.oxml.ns import qn
    texts = []
    for t in el.iter(qn("w:t")):
        if t.text:
            texts.append(t.text)
    return "".join(texts)


def _get_para_style(para) -> str:
    """获取段落样式名。"""
    try:
        return para.style.name or "Normal"
    except Exception:
        return "Unknown"


CHINESE_NUMERAL_CHARS = "零〇一二三四五六七八九十百千万两"
HEADING_LEVEL_RE = re.compile(r"^(?:heading|标题)\s*([1-9]\d*)$", re.I)
NUMBERED_HEADING_RE = re.compile(
    r"^\s*(?P<prefix>\d+\s*(?:[.．]\s*\d+\s*)*(?:[.．、)）]\s*|\s+)?)"
    r"(?P<title>\S.*)$"
)
CHINESE_NUMBERED_HEADING_RE = re.compile(
    rf"^\s*(?P<prefix>(?:（\s*)?[{CHINESE_NUMERAL_CHARS}]+\s*(?:[、.．)）]))\s*(?P<title>\S.*)$"
)


def _extract_heading_level_from_style(style_name: str | None) -> int | None:
    """从段落样式名中提取 Heading 层级。"""
    match = HEADING_LEVEL_RE.match((style_name or "").strip())
    if not match:
        return None
    return max(1, min(int(match.group(1)), 9))


def _style_outline_level(style) -> int | None:
    """读取单个样式直接声明的 outlineLvl。"""
    from docx.oxml.ns import qn

    style_element = getattr(style, "element", None)
    if style_element is None:
        return None
    p_pr = style_element.find(qn("w:pPr"))
    outline = p_pr.find(qn("w:outlineLvl")) if p_pr is not None else None
    if outline is None:
        return None
    try:
        value = int(outline.get(qn("w:val")))
    except (ValueError, TypeError):
        return None
    return value if 0 <= value <= 8 else None


def _style_chain_heading_level(para) -> int | None:
    """沿 basedOn 链读取真实 Heading 名称或祖先 outlineLvl。"""
    style = getattr(para, "style", None)
    seen: set[int] = set()
    while style is not None:
        element = getattr(style, "element", None)
        marker = id(element) if element is not None else id(style)
        if marker in seen:
            break
        seen.add(marker)

        level = _extract_heading_level_from_style(getattr(style, "name", "") or "")
        if level is not None:
            return level
        outline_level = _style_outline_level(style)
        if outline_level is not None:
            return outline_level + 1
        style = getattr(style, "base_style", None)
    return None


def _parse_numbered_heading_prefix(text: str) -> tuple[str, str, int] | None:
    """解析标题前缀编号，返回 (prefix, title, depth)。"""
    raw_text = text or ""

    match = NUMBERED_HEADING_RE.match(raw_text)
    if match:
        prefix = match.group("prefix")
        title = match.group("title").strip()
        if not title:
            return None

        numeric_parts = re.findall(r"\d+", prefix)
        if not numeric_parts:
            return None

        has_explicit_separator = prefix != prefix.rstrip() or bool(re.search(r"[.．、)）]\s*$", prefix))
        if len(numeric_parts) == 1 and not has_explicit_separator:
            return None

        return prefix, title, min(len(numeric_parts), 9)

    match = CHINESE_NUMBERED_HEADING_RE.match(raw_text)
    if not match:
        return None

    prefix = match.group("prefix")
    title = match.group("title").strip()
    if not title:
        return None

    numeral_text = re.sub(r"[\s（）()、.．)）]", "", prefix)
    if not numeral_text:
        return None

    return prefix, title, 1


def _remove_paragraph_numpr(para) -> bool:
    """移除段落级自动编号设置。"""
    from docx.oxml.ns import qn

    pPr = para._element.find(qn("w:pPr"))
    if pPr is None:
        return False

    num_pr = pPr.find(qn("w:numPr"))
    if num_pr is None:
        return False

    pPr.remove(num_pr)
    return True


def _remove_style_numpr(style) -> bool:
    """移除样式上的自动编号设置。"""
    from docx.oxml.ns import qn

    style_element = getattr(style, "element", None)
    if style_element is None:
        return False

    pPr = style_element.find(qn("w:pPr"))
    if pPr is None:
        return False

    num_pr = pPr.find(qn("w:numPr"))
    if num_pr is None:
        return False

    pPr.remove(num_pr)
    return True


def _paragraph_has_numpr(para) -> bool:
    """段落级是否带自动编号。"""
    from docx.oxml.ns import qn

    pPr = para._element.find(qn("w:pPr"))
    return pPr is not None and pPr.find(qn("w:numPr")) is not None


def _resolve_heading_level(para, inferred_depth: int) -> int:
    """优先保留段落及 basedOn 样式链的 Heading/outline 层级。"""
    style_level = _extract_heading_level_from_style(_get_para_style(para))
    if style_level is not None:
        return style_level

    outline_level = _get_outline_level(para._element)
    if outline_level is not None:
        return max(1, min(outline_level + 1, 9))

    style_chain_level = _style_chain_heading_level(para)
    if style_chain_level is not None:
        return style_chain_level

    numpr_level = _get_numpr_level(para)
    if numpr_level is not None:
        return max(1, min(numpr_level + 1, 9))

    return max(1, min(inferred_depth, 9))


def _set_outline_level(para_element, level: int) -> None:
    """为段落设置 outlineLvl。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = para_element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        para_element.insert(0, pPr)

    outline = pPr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        pPr.append(outline)
    outline.set(qn("w:val"), str(max(0, min(level - 1, 8))))


def _apply_heading_level(para, level: int) -> None:
    """为段落保留或补齐对应 Heading。"""
    level = max(1, min(level, 9))
    for style_name in (f"Heading {level}", f"标题 {level}", f"标题{level}"):
        try:
            para.style = style_name
            break
        except (KeyError, ValueError):
            continue
    _remove_style_numpr(getattr(para, "style", None))
    _set_outline_level(para._element, level)


def _looks_like_heading_para(para) -> bool:
    """段落是否应按标题处理。

    只认 docx 里真实存在的标题证据（Heading 样式名、段落 outlineLvl、
    basedOn 样式链上的 Heading/outlineLvl），不按"带编号、文字较短"猜标题——
    猜测会把"1、载荷仿真分析能力"这类正文误升为 Heading，凭空造出原文档
    没有的层级。
    """
    style_level = _extract_heading_level_from_style(_get_para_style(para))
    outline_level = _get_outline_level(para._element)
    style_chain_level = _style_chain_heading_level(para)
    return (
        style_level is not None
        or outline_level is not None
        or style_chain_level is not None
    )


def _replace_para_text_preserve_runs(para, new_text: str) -> None:
    """尽量保留原 run 结构，仅替换可见文本。"""
    from docx.oxml.ns import qn

    text_nodes = [node for node in para._element.iter(qn("w:t"))]
    if not text_nodes:
        para.text = new_text
        return

    text_nodes[0].text = new_text
    for node in text_nodes[1:]:
        node.text = ""


def _strip_numbered_heading_prefixes(doc) -> int:
    """去掉标题前的数字编号，并保留或补齐对应 Heading。"""
    normalized = 0
    for para in doc.paragraphs:
        raw_text = para.text or ""
        stripped_text = raw_text.strip()
        if not stripped_text:
            continue

        parsed = _parse_numbered_heading_prefix(raw_text)
        cleaned_text = stripped_text
        inferred_depth = 1
        had_visible_prefix = False
        if parsed is not None:
            _, cleaned_text, inferred_depth = parsed
            had_visible_prefix = cleaned_text != stripped_text

        had_numpr = _paragraph_has_numpr(para)
        if not _looks_like_heading_para(para):
            continue

        changed = False
        if had_visible_prefix:
            _replace_para_text_preserve_runs(para, cleaned_text)
            changed = True

        if had_numpr:
            changed = _remove_paragraph_numpr(para) or changed

        style_level = _extract_heading_level_from_style(_get_para_style(para))
        outline_level = _get_outline_level(para._element)
        style_chain_level = _style_chain_heading_level(para)
        if (
            had_visible_prefix
            or had_numpr
            or style_level is not None
            or outline_level is not None
            or style_chain_level is not None
        ):
            _apply_heading_level(para, _resolve_heading_level(para, inferred_depth))

        if changed:
            normalized += 1

    return normalized


def _collect_body_elements(body, *, include_full_text: bool = False) -> list[dict]:
    """收集 body 一级子元素信息。

    index 始终表示 /body 的一级子元素索引，不是段落序号或 outline 序号。
    """
    from docx.oxml.ns import qn

    elements = []
    for i, child in enumerate(body):
        tag = _get_element_tag(child)
        text = _get_element_text(child).strip()
        text_preview = text[:60] + ("..." if len(text) > 60 else "") if text else ""

        entry = {
            "index": i,
            "tag": tag,
            "text_preview": text_preview,
        }
        if include_full_text:
            entry["_full_text"] = text

        para_id = child.get(qn("w14:paraId")) or child.get(qn("w:paraId"))
        if para_id:
            entry["paraId"] = para_id

        if tag == "p":
            pPr = child.find(qn("w:pPr"))
            if pPr is not None:
                pStyle = pPr.find(qn("w:pStyle"))
                if pStyle is not None:
                    entry["style"] = pStyle.get(qn("w:val"), "")

            for br in child.iter(qn("w:br")):
                if br.get(qn("w:type"), "") == "page":
                    entry["has_page_break"] = True
                    break

        elements.append(entry)

    return elements


def _normalize_anchor_text(text: str | None) -> str:
    """归一化锚点文本，便于稳健比对。"""
    return " ".join((text or "").split())


def _describe_body_element(entry: dict) -> str:
    """生成用于日志/报错的 body 元素摘要。"""
    parts = [f"body index={entry['index']}", f"tag={entry['tag']}"]
    para_id = entry.get("paraId")
    if para_id:
        parts.append(f"paraId={para_id}")

    text = entry.get("_full_text") or entry.get("text_preview") or ""
    if text:
        preview = text[:60] + ("..." if len(text) > 60 else "")
        parts.append(f'text="{preview}"')

    return ", ".join(parts)


def _resolve_trim_anchor(
    elements: list[dict],
    *,
    anchor_index: int | None,
    anchor_para_id: str | None,
    anchor_text: str | None,
) -> dict:
    """解析 trim 命中的锚点，并阻止混用不同坐标系。"""
    entry_by_index = None
    if anchor_index is not None:
        if anchor_index < 0 or anchor_index >= len(elements):
            raise ValueError(f"anchor_index {anchor_index} 超出范围 [0, {len(elements)-1}]")
        entry_by_index = elements[anchor_index]

    entry_by_para_id = None
    normalized_para_id = (anchor_para_id or "").strip().upper()
    if normalized_para_id:
        matches = [
            entry for entry in elements
            if (entry.get("paraId") or "").upper() == normalized_para_id
        ]
        if not matches:
            raise ValueError(f"未找到 paraId={normalized_para_id} 对应的 body 一级元素")
        if len(matches) > 1:
            raise ValueError(f"paraId={normalized_para_id} 命中了多个 body 一级元素，请改用更明确的锚点")
        entry_by_para_id = matches[0]

    if entry_by_index and entry_by_para_id and entry_by_index["index"] != entry_by_para_id["index"]:
        raise ValueError(
            "--anchor-index 与 --anchor-para-id 指向了不同元素；"
            "这通常说明把 outline/paragraph 序号误当成了 body index。\n"
            f"  index 命中: {_describe_body_element(entry_by_index)}\n"
            f"  paraId 命中: {_describe_body_element(entry_by_para_id)}"
        )

    entry = entry_by_para_id or entry_by_index
    if entry is None:
        raise ValueError("trim 至少需要提供 --anchor-index 或 --anchor-para-id 之一")

    normalized_expected = _normalize_anchor_text(anchor_text)
    if normalized_expected:
        actual_text = _normalize_anchor_text(entry.get("_full_text"))
        if normalized_expected.casefold() not in actual_text.casefold():
            raise ValueError(
                "--anchor-text 与命中的 body 元素不一致；"
                "这通常说明锚点 index/paraId 取错了。\n"
                f"  期望文本: {anchor_text}\n"
                f"  实际命中: {_describe_body_element(entry)}"
            )

    return entry


# ─── 子命令: peek ───


def cmd_peek(docx_path: Path, max_paras: int = 100) -> None:
    """提取前 N 段的纯文本 + 样式名，输出到 stdout。

    格式模仿 officecli view text/outline 的混合输出，方便 Claude 阅读。
    """
    doc, tmp_dir = _safe_open(docx_path)

    print(f"=== 文档探针: {docx_path.name} ===")
    print(f"总段落数: {len(doc.paragraphs)}")
    print()

    # 先输出 outline（仅 heading 段落）
    print("--- OUTLINE ---")
    for i, para in enumerate(doc.paragraphs):
        style = _get_para_style(para)
        if "heading" in style.lower() or "标题" in style:
            level = ""
            try:
                if hasattr(para.style, "name"):
                    # 从样式名提取层级: "Heading 1" -> 1
                    parts = para.style.name.split()
                    if len(parts) >= 2 and parts[-1].isdigit():
                        level = f" (heading {parts[-1]})"
            except Exception:
                pass
            text = para.text.strip()
            if text:
                print(f"  [{i:3d}]{level} {text}")

    print()
    print("--- TEXT (前{}段) ---".format(min(max_paras, len(doc.paragraphs))))

    for i, para in enumerate(doc.paragraphs[:max_paras]):
        style = _get_para_style(para)
        text = para.text.strip()
        marker = f"[{i:3d}][{style}]"

        if not text:
            # 检查是否含分页符
            from docx.oxml.ns import qn
            has_break = False
            for br in para._element.iter(qn("w:br")):
                br_type = br.get(qn("w:type"), "")
                if br_type == "page":
                    has_break = True
                    break
            if has_break:
                print(f"  {marker} <PAGE BREAK>")
            else:
                print(f"  {marker} <EMPTY>")
        else:
            # 截断过长文本
            display = text[:80] + ("..." if len(text) > 80 else "")
            print(f"  {marker} {display}")

    if tmp_dir:
        shutil.rmtree(tmp_dir, ignore_errors=True)


# ─── 子命令: list-body ───


def cmd_list_body(docx_path: Path) -> None:
    """列出 body 所有一级子元素的类型、文本摘要和索引，输出 JSON。

    替代 officecli get '<file>' '/body' --depth 1 --json
    index 为 /body 一级子元素索引，不是段落序号或 outline 序号。
    """
    doc, tmp_dir = _safe_open(docx_path)

    body = doc.element.body
    elements = _collect_body_elements(body)

    result = {
        "file": str(docx_path.name),
        "total_elements": len(elements),
        "index_semantics": "body-child-index",
        "elements": elements,
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))

    if tmp_dir:
        shutil.rmtree(tmp_dir, ignore_errors=True)


# ─── 子命令: trim ───


def cmd_trim(
    docx_path: Path,
    anchor_index: int | None,
    anchor_para_id: str | None = None,
    anchor_text: str | None = None,
) -> None:
    """删除 body 中锚点之前的所有元素。

    推荐使用 paraId（可选配合 anchor_text）定位锚点。
    anchor_index 仅表示 list-body 输出中的 body index，不是段落序号或 outline 序号。
    索引 anchor 处的元素及其之后的所有内容被完整保留。

    操作直接在文件上进行（已通过 _safe_open/_safe_save 保护中文路径）。
    """
    doc, tmp_dir = _safe_open(docx_path)

    body = doc.element.body
    children = list(body)
    elements = _collect_body_elements(body, include_full_text=True)

    try:
        anchor = _resolve_trim_anchor(
            elements,
            anchor_index=anchor_index,
            anchor_para_id=anchor_para_id,
            anchor_text=anchor_text,
        )
    except ValueError as exc:
        sys.stderr.write(f"{exc}\n")
        if tmp_dir:
            shutil.rmtree(tmp_dir, ignore_errors=True)
        sys.exit(1)

    resolved_index = anchor["index"]
    print(f"trim 锚点确认: {_describe_body_element(anchor)}")

    # 从前往后删除，每次删第一个（因为删除后后续元素前移）
    removed = 0
    for _ in range(resolved_index):
        first_child = body[0]
        # 不删 sectPr（文档最后的节属性必须保留）
        if _get_element_tag(first_child) == "sectPr":
            break
        body.remove(first_child)
        removed += 1

    _safe_save(doc, docx_path, tmp_dir)
    print(f"已删除 {removed} 个元素，文档现在从索引 0 开始（原索引 {resolved_index}）")

    # 验证：输出新的第一个元素
    doc2, tmp2 = _safe_open(docx_path)
    body2 = doc2.element.body
    if len(body2) > 0:
        first = body2[0]
        tag = _get_element_tag(first)
        text = _get_element_text(first).strip()[:60]
        print(f"新首元素: <{tag}> {text}")
    if tmp2:
        shutil.rmtree(tmp2, ignore_errors=True)


# ─── 子命令: normalize ───


def _has_image(para_element) -> bool:
    """段落是否含图片（drawing 或 pict）。"""
    from docx.oxml.ns import qn
    if para_element.find(qn("w:r")) is not None:
        for r in para_element.iter(qn("w:r")):
            if r.find(qn("w:drawing")) is not None or r.find(qn("w:pict")) is not None:
                return True
    return False


def _has_page_break(para_element) -> bool:
    """段落是否含分页符。"""
    from docx.oxml.ns import qn
    for br in para_element.iter(qn("w:br")):
        if br.get(qn("w:type"), "") == "page":
            return True
    return False


def _has_sect_pr(para_element) -> bool:
    """段落 pPr 中是否含节属性（sectPr）。"""
    from docx.oxml.ns import qn
    pPr = para_element.find(qn("w:pPr"))
    if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
        return True
    return False


def _get_outline_level(para_element) -> int | None:
    """获取段落的 outlineLvl 值（段落级别覆盖），None 表示无。"""
    from docx.oxml.ns import qn
    pPr = para_element.find(qn("w:pPr"))
    if pPr is not None:
        ol = pPr.find(qn("w:outlineLvl"))
        if ol is not None:
            try:
                return int(ol.get(qn("w:val")))
            except (ValueError, TypeError):
                pass
    return None


def _get_numpr_level(para) -> int | None:
    """获取段落级或样式级自动编号层级。"""
    from docx.oxml.ns import qn

    pPr = para._element.find(qn("w:pPr"))
    if pPr is not None:
        num_pr = pPr.find(qn("w:numPr"))
        if num_pr is not None:
            ilvl = num_pr.find(qn("w:ilvl"))
            if ilvl is not None:
                try:
                    return int(ilvl.get(qn("w:val")))
                except (ValueError, TypeError):
                    pass

    style = getattr(para, "style", None)
    style_element = getattr(style, "element", None)
    if style_element is None:
        return None

    style_pPr = style_element.find(qn("w:pPr"))
    if style_pPr is None:
        return None

    num_pr = style_pPr.find(qn("w:numPr"))
    if num_pr is None:
        return None

    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        return None

    try:
        return int(ilvl.get(qn("w:val")))
    except (ValueError, TypeError):
        return None


def _is_empty_para(para) -> bool:
    """段落是否无可见文字且不含图片。"""
    if para.text.strip():
        return False
    if _has_image(para._element):
        return False
    return True


def _is_page_break_only(para_element) -> bool:
    """段落是否仅含分页符（无可见文字、无图片）。"""
    if _get_element_text(para_element).strip():
        return False
    if _has_image(para_element):
        return False
    return _has_page_break(para_element)


def _is_empty_or_page_break_only_body_para(body_child) -> bool:
    """body 一级子元素是否为空段落或纯分页符段落。"""
    if _get_element_tag(body_child) != "p":
        return False
    if _has_sect_pr(body_child):
        return False
    if _get_element_text(body_child).strip():
        return False
    if _has_image(body_child):
        return False
    return True


def _is_empty_body_para(body_child) -> bool:
    """body 一级子元素是否为可参与空行群统计的空段落。"""
    if _is_preserved_blank_heading(body_child):
        return False
    if _get_element_tag(body_child) != "p":
        return False
    if _has_sect_pr(body_child):
        return False
    if _get_element_text(body_child).strip():
        return False
    if _has_image(body_child):
        return False
    return True


def _is_substantive_body_child(body_child) -> bool:
    """body 一级子元素是否承载实际内容。"""
    if _is_preserved_blank_heading(body_child):
        return True
    tag = _get_element_tag(body_child)
    if tag == "sectPr":
        return False
    if tag == "p":
        return bool(_get_element_text(body_child).strip()) or _has_image(body_child)
    if tag in {"tbl", "sdt"}:
        return True
    if _get_element_text(body_child).strip():
        return True
    if _has_image(body_child):
        return True
    return False


def _iter_page_breaks(body_child):
    """按文档顺序迭代 body 子元素中的显式分页符。"""
    from docx.oxml.ns import qn
    for br in body_child.iter(qn("w:br")):
        if br.get(qn("w:type"), "") == "page":
            yield br


@dataclass
class _PageBreakEvent:
    body_idx: int
    body_child: object
    break_el: object


def _collect_page_break_events(body_children) -> list[_PageBreakEvent]:
    """收集 body 一级子元素中所有显式分页符事件。"""
    events: list[_PageBreakEvent] = []
    for idx, child in enumerate(body_children):
        for br in _iter_page_breaks(child):
            events.append(_PageBreakEvent(body_idx=idx, body_child=child, break_el=br))
    return events


def _node_has_substantive_content(node) -> bool:
    """节点本身是否承载可见内容。"""
    from docx.oxml.ns import qn
    if node.tag == qn("w:t"):
        return bool((node.text or "").strip())
    return node.tag in {qn("w:drawing"), qn("w:pict")}


def _is_blank_sectpr_para(body_child) -> bool:
    """body 一级子元素是否为仅承载 sectPr 的空段落包装。"""
    return (
        _get_element_tag(body_child) == "p"
        and _has_sect_pr(body_child)
        and not _get_element_text(body_child).strip()
        and not _has_image(body_child)
        and not _has_page_break(body_child)
    )


def _is_removable_blank_residue(body_child) -> bool:
    """可安全删除的空白页残留：普通空段落、纯分页符段落。"""
    if _is_preserved_blank_heading(body_child):
        return False
    return _is_empty_or_page_break_only_body_para(body_child)


def _mark_preserved_blank_heading(para_element) -> None:
    """用不可见书签标记空白标题，确保重复 normalize 时仍保留段落。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for bookmark in para_element.iter(qn("w:bookmarkStart")):
        if str(bookmark.get(qn("w:name")) or "").startswith(_BLANK_HEADING_BOOKMARK_PREFIX):
            return

    bookmark_id = str(secrets.randbelow(2_000_000_000) + 1)
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), bookmark_id)
    bookmark_start.set(qn("w:name"), f"{_BLANK_HEADING_BOOKMARK_PREFIX}{bookmark_id}")
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), bookmark_id)
    para_element.append(bookmark_start)
    para_element.append(bookmark_end)


def _is_preserved_blank_heading(body_child) -> bool:
    """段落是否为本次 normalize 必须保留的空白标题段落。"""
    if _get_element_tag(body_child) != "p":
        return False
    from docx.oxml.ns import qn

    return any(
        str(bookmark.get(qn("w:name")) or "").startswith(_BLANK_HEADING_BOOKMARK_PREFIX)
        for bookmark in body_child.iter(qn("w:bookmarkStart"))
    )


def _is_blank_page_residue(body_child) -> bool:
    """body 一级子元素是否属于空白页中的布局残留。"""
    tag = _get_element_tag(body_child)
    if tag == "sectPr":
        return True
    if tag != "p":
        return False
    if _has_sect_pr(body_child):
        return True
    return _is_removable_blank_residue(body_child)


def _get_para_sectpr(para_element):
    """获取段落 pPr 中的 sectPr 节点。"""
    from docx.oxml.ns import qn

    pPr = para_element.find(qn("w:pPr"))
    if pPr is None:
        return None
    return pPr.find(qn("w:sectPr"))


def _ensure_para_ppr(para_element):
    """确保段落存在 pPr，并返回该节点。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = para_element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        para_element.insert(0, pPr)
    return pPr


def _find_sectpr_rehome_target(body, body_children, carrier_idx: int):
    """为待迁移的 sectPr 找到最接近的保留位置。"""
    for idx in range(carrier_idx - 1, -1, -1):
        candidate = body_children[idx]
        if candidate.getparent() is not body:
            continue
        tag = _get_element_tag(candidate)
        if tag == "p":
            if _has_sect_pr(candidate):
                break
            if _is_removable_blank_residue(candidate) and not _has_page_break(candidate):
                continue
            return candidate
        if tag == "sectPr":
            break
        if _is_substantive_body_child(candidate):
            break

    from docx.oxml import OxmlElement

    carrier = body_children[carrier_idx]
    new_para = OxmlElement("w:p")
    carrier.addprevious(new_para)
    return new_para


def _relocate_blank_sectpr_para(body, body_children, carrier_idx: int) -> bool:
    """迁移空白段落上的 sectPr，避免删空白页时丢失分节信息。"""
    carrier = body_children[carrier_idx]
    if carrier.getparent() is not body or not _is_blank_sectpr_para(carrier):
        return False

    sectPr = _get_para_sectpr(carrier)
    if sectPr is None:
        return False

    target = _find_sectpr_rehome_target(body, body_children, carrier_idx)
    target_pPr = _ensure_para_ppr(target)
    target_pPr.append(deepcopy(sectPr))
    body.remove(carrier)
    return True


def _next_non_residue_index(body_children, start_idx: int) -> tuple[int, bool]:
    """返回下一个非空白残留元素索引，以及是否跳过了残留元素。"""
    idx = start_idx
    saw_residue = False
    while idx < len(body_children) and _is_blank_page_residue(body_children[idx]):
        saw_residue = True
        idx += 1
    return idx, saw_residue


def _prev_non_residue_index(body_children, start_idx: int) -> tuple[int, bool]:
    """返回上一个非空白残留元素索引，以及是否跳过了残留元素。"""
    idx = start_idx
    saw_residue = False
    while idx >= 0 and _is_blank_page_residue(body_children[idx]):
        saw_residue = True
        idx -= 1
    return idx, saw_residue


def _has_substantive_before_break(body_child, break_el) -> bool:
    """指定分页符之前是否已有可见内容。"""
    for node in body_child.iter():
        if node is break_el:
            return False
        if _node_has_substantive_content(node):
            return True
    return False


def _has_substantive_after_break(body_child, break_el) -> bool:
    """指定分页符之后是否仍有可见内容。"""
    seen_break = False
    for node in body_child.iter():
        if seen_break and _node_has_substantive_content(node):
            return True
        if node is break_el:
            seen_break = True
    return False


def _has_substantive_before_event(body_children, event: _PageBreakEvent) -> bool:
    """指定分页事件之前是否已有实质内容。"""
    if _is_preserved_blank_heading(event.body_child):
        return True
    for child in body_children[:event.body_idx]:
        if _is_substantive_body_child(child):
            return True
    return _has_substantive_before_break(event.body_child, event.break_el)


def _has_substantive_after_event(body_children, event: _PageBreakEvent) -> bool:
    """指定分页事件之后是否仍有实质内容。"""
    if _is_preserved_blank_heading(event.body_child):
        return True
    if _has_substantive_after_break(event.body_child, event.break_el):
        return True
    for child in body_children[event.body_idx + 1:]:
        if _is_substantive_body_child(child):
            return True
    return False


def _remove_page_break(break_el) -> None:
    """删除单个显式分页符节点。"""
    parent = break_el.getparent()
    if parent is not None:
        parent.remove(break_el)


def _remove_body_child_if_orphaned(body, body_child) -> None:
    """若 body 子元素已变成无用空壳且不再含分页符，则删除。"""
    if body_child.getparent() is not body:
        return
    if _get_element_tag(body_child) == "p" and _has_page_break(body_child):
        return
    if _is_removable_blank_residue(body_child):
        body.remove(body_child)


def _range_has_substantive_between_breaks(
    body_children,
    left_event: _PageBreakEvent,
    right_event: _PageBreakEvent,
) -> bool:
    """两个相邻分页符之间是否存在实质内容。"""
    if _is_preserved_blank_heading(left_event.body_child) or _is_preserved_blank_heading(right_event.body_child):
        return True
    if left_event.body_idx == right_event.body_idx:
        seen_left = False
        for node in left_event.body_child.iter():
            if node is left_event.break_el:
                seen_left = True
                continue
            if node is right_event.break_el:
                return False
            if seen_left and _node_has_substantive_content(node):
                return True
        return False

    if _has_substantive_after_break(left_event.body_child, left_event.break_el):
        return True

    for child in body_children[left_event.body_idx + 1:right_event.body_idx]:
        if _is_substantive_body_child(child):
            return True

    if _has_substantive_before_break(right_event.body_child, right_event.break_el):
        return True

    return False


def _cleanup_leading_blank_page(
    body,
    body_children,
    first_event: _PageBreakEvent,
) -> None:
    """清理文档开头由分页符和空白段落组成的空白页。"""
    _remove_page_break(first_event.break_el)

    for idx in range(first_event.body_idx - 1, -1, -1):
        child = body_children[idx]
        if child.getparent() is not body:
            continue
        if _is_blank_sectpr_para(child):
            _relocate_blank_sectpr_para(body, body_children, idx)
            continue
        if _is_removable_blank_residue(child):
            body.remove(child)

    _remove_body_child_if_orphaned(body, first_event.body_child)


def _cleanup_blank_page_gap(
    body,
    body_children,
    left_event: _PageBreakEvent,
    right_event: _PageBreakEvent,
) -> None:
    """清理空白页区间中的多余分页符与空白段落。"""
    _remove_page_break(right_event.break_el)

    for idx in range(right_event.body_idx - 1, left_event.body_idx, -1):
        child = body_children[idx]
        if child.getparent() is not body:
            continue
        if _is_blank_sectpr_para(child):
            _relocate_blank_sectpr_para(body, body_children, idx)
            continue
        if _is_removable_blank_residue(child):
            body.remove(child)

    _remove_body_child_if_orphaned(body, right_event.body_child)


def _cleanup_trailing_blank_page(
    body,
    body_children,
    last_event: _PageBreakEvent,
) -> None:
    """清理文档末尾由分页符和空白段落组成的空白页。"""
    _remove_page_break(last_event.break_el)

    for idx in range(len(body_children) - 1, last_event.body_idx, -1):
        child = body_children[idx]
        if child.getparent() is not body:
            continue
        if _is_blank_sectpr_para(child):
            _relocate_blank_sectpr_para(body, body_children, idx)
            continue
        if _is_removable_blank_residue(child):
            body.remove(child)

    _remove_body_child_if_orphaned(body, last_event.body_child)


def _has_body_sectpr_after(body, body_children, start_idx: int) -> bool:
    """指定位置之后是否仍保留 body 级 sectPr。"""
    for idx in range(start_idx + 1, len(body_children)):
        child = body_children[idx]
        if child.getparent() is body and _get_element_tag(child) == "sectPr":
            return True
    return False


def _cleanup_trailing_blank_tail(body) -> int:
    """清理文末尾巴上的空段落、纯分页符段落和空 sectPr 包装段落。"""
    body_children = list(body)
    if not body_children:
        return 0

    idx = len(body_children) - 1
    while idx >= 0 and _get_element_tag(body_children[idx]) == "sectPr":
        idx -= 1

    removable_indices = []
    while idx >= 0:
        child = body_children[idx]
        if _is_blank_sectpr_para(child) or _is_removable_blank_residue(child):
            removable_indices.append(idx)
            idx -= 1
            continue
        break

    if not removable_indices:
        return 0

    removed = 0
    for remove_idx in reversed(removable_indices):
        child = body_children[remove_idx]
        if child.getparent() is not body:
            continue
        if _is_blank_sectpr_para(child):
            if _has_body_sectpr_after(body, body_children, remove_idx):
                body.remove(child)
                removed += 1
            elif _relocate_blank_sectpr_para(body, body_children, remove_idx):
                removed += 1
            continue
        if _is_removable_blank_residue(child):
            body.remove(child)
            removed += 1

    return removed


def _collapse_blank_break_runs(body) -> int:
    """折叠两个实质内容块之间仅由空白残留组成的区间。"""
    body_children = list(body)
    runs: list[tuple[int, int]] = []
    run_start = None

    for idx, child in enumerate(body_children):
        if _is_removable_blank_residue(child):
            if run_start is None:
                run_start = idx
        else:
            if run_start is not None:
                runs.append((run_start, idx - 1))
                run_start = None
    if run_start is not None:
        runs.append((run_start, len(body_children) - 1))

    removed = 0
    for start_idx, end_idx in reversed(runs):
        body_children = list(body)
        run_children = body_children[start_idx:end_idx + 1]
        if not run_children:
            continue

        prev_idx = start_idx - 1
        while prev_idx >= 0 and not _is_substantive_body_child(body_children[prev_idx]):
            prev_idx -= 1

        next_idx = end_idx + 1
        while next_idx < len(body_children) and not _is_substantive_body_child(body_children[next_idx]):
            next_idx += 1

        if prev_idx < 0 or next_idx >= len(body_children):
            continue

        break_positions = [
            rel_idx for rel_idx, child in enumerate(run_children)
            if _get_element_tag(child) == "p" and _has_page_break(child)
        ]
        if not break_positions:
            continue

        prev_has_break = _has_page_break(body_children[prev_idx])
        keep_rel_idx = None if prev_has_break else break_positions[-1]

        for rel_idx in range(len(run_children) - 1, -1, -1):
            if keep_rel_idx is not None and rel_idx == keep_rel_idx:
                continue
            child = run_children[rel_idx]
            if child.getparent() is not body and not _is_blank_sectpr_para(child):
                continue
            if _is_blank_sectpr_para(child):
                if _relocate_blank_sectpr_para(body, body_children, start_idx + rel_idx):
                    removed += 1
                continue
            if child.getparent() is body and _is_removable_blank_residue(child):
                body.remove(child)
                removed += 1

    return removed


def _count_blank_pages_from_events(body_children) -> int:
    """按分页事件分布估算当前文档中的物理空白页数量。"""
    events = _collect_page_break_events(body_children)
    if not events:
        return 0

    blank_pages = 0
    first_event = events[0]
    if not _has_substantive_before_event(body_children, first_event):
        blank_pages += 1

    for left_event, right_event in zip(events, events[1:]):
        if not _range_has_substantive_between_breaks(body_children, left_event, right_event):
            blank_pages += 1

    last_event = events[-1]
    if not _has_substantive_after_event(body_children, last_event):
        blank_pages += 1

    return blank_pages


def _remove_blank_pages_via_xml(docx_path: Path) -> int:
    """按前导/相邻/尾随分页事件之间是否仅有空白内容来清理空白页。"""
    tmp_dir = None

    try:
        doc, tmp_dir = _safe_open(docx_path)
        body = doc.element.body
        blank_pages = _count_blank_pages_from_events(list(body))
        changed = False

        while True:
            body_children = list(body)
            events = _collect_page_break_events(body_children)
            if not events:
                break

            removed_in_pass = False
            first_event = events[0]
            if not _has_substantive_before_event(body_children, first_event):
                _cleanup_leading_blank_page(body, body_children, first_event)
                changed = True
                removed_in_pass = True
            else:
                for left_event, right_event in zip(events, events[1:]):
                    if _range_has_substantive_between_breaks(body_children, left_event, right_event):
                        continue

                    _cleanup_blank_page_gap(body, body_children, left_event, right_event)
                    changed = True
                    removed_in_pass = True
                    break

                if not removed_in_pass:
                    last_event = events[-1]
                    if not _has_substantive_after_event(body_children, last_event):
                        _cleanup_trailing_blank_page(body, body_children, last_event)
                        changed = True
                        removed_in_pass = True

            if not removed_in_pass:
                break

        if changed:
            _safe_save(doc, docx_path, tmp_dir)
            tmp_dir = None
        elif tmp_dir is not None:
            shutil.rmtree(tmp_dir, ignore_errors=True)
            tmp_dir = None

        return blank_pages
    except Exception as exc:
        if tmp_dir is not None:
            shutil.rmtree(tmp_dir, ignore_errors=True)
        raise RuntimeError(str(exc)) from exc


def cmd_normalize(docx_path: Path) -> None:
    """格式规范化：标题去编号并保留/补齐 Heading，清除空段落的标题元数据、合并连续空行、按显式分页符清理空白页、清理文末空白。"""
    blank_pages_removed = 0
    blank_page_error = None

    doc, tmp_dir = _safe_open(docx_path)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    print(f"格式规范化: {docx_path.name}")

    # ── 操作 A：标题去编号，并保留/补齐 Heading ──
    heading_number_stripped = _strip_numbered_heading_prefixes(doc)

    # ── 操作 B：清除空段落的 Heading 样式和目录级别 ──
    heading_cleared = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            continue  # 有文字的段落不处理

        para_outline_lvl = _get_outline_level(para._element)
        if not _looks_like_heading_para(para):
            continue

        # 仅清除标题元数据；段落、换行符和分页符保持不变。
        _mark_preserved_blank_heading(para._element)
        para.style = doc.styles["Normal"]

        if para_outline_lvl is not None:
            pPr = para._element.find(qn("w:pPr"))
            if pPr is not None:
                ol = pPr.find(qn("w:outlineLvl"))
                if ol is not None:
                    pPr.remove(ol)

        heading_cleared += 1

    # ── 操作 C：连续空段落合并为分页符 ──
    # 在 body XML 层面操作，避免跨越实际内容元素误合并空行群
    body = doc.element.body
    body_children = list(body)

    # 第一遍：按 body 一级子元素标记连续空段落群
    groups = []  # [(start_body_idx, end_body_idx), ...]
    run_start = None
    for i, child in enumerate(body_children):
        if _is_empty_body_para(child):
            if run_start is None:
                run_start = i
        else:
            if run_start is not None and (i - run_start) >= 3:
                groups.append((run_start, i - 1))
            run_start = None
    # 处理末尾
    if run_start is not None and (len(body_children) - run_start) >= 3:
        groups.append((run_start, len(body_children) - 1))

    # 第二遍：合并（从后往前操作，避免索引偏移）
    empty_merged_groups = 0
    empty_merged_paras = 0
    for start_idx, end_idx in reversed(groups):
        body_children = list(body)
        count = end_idx - start_idx + 1

        # 判断该空行群之后是否还有实际内容元素
        has_content_after = False
        for child in body_children[end_idx + 1:]:
            if _is_substantive_body_child(child):
                has_content_after = True
                break

        # 找到群中是否已有分页符段落
        existing_break_idx = None
        for i in range(start_idx, end_idx + 1):
            if _has_page_break(body_children[i]):
                existing_break_idx = i
                break

        if not has_content_after:
            # 文末：直接删除所有空段落，不插入分页符
            for i in range(end_idx, start_idx - 1, -1):
                el = body_children[i]
                if _has_sect_pr(el):
                    continue
                body.remove(el)
            empty_merged_paras += count
        elif existing_break_idx is not None:
            # 已有分页符段落：保留它，删除其余
            for i in range(end_idx, start_idx - 1, -1):
                if i == existing_break_idx:
                    continue
                body.remove(body_children[i])
            empty_merged_paras += count - 1
        else:
            # 无分页符：在第一个空段落位置插入分页符，删除其余
            first_el = body_children[start_idx]
            r = first_el.find(qn("w:r"))
            if r is None:
                r = OxmlElement("w:r")
                first_el.append(r)
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "page")
            r.append(br)

            for i in range(end_idx, start_idx, -1):
                body.remove(body_children[i])
            empty_merged_paras += count - 1

        empty_merged_groups += 1

    _safe_save(doc, docx_path, tmp_dir)

    try:
        blank_pages_removed = _remove_blank_pages_via_xml(docx_path)
    except RuntimeError as exc:
        blank_page_error = str(exc)

    doc, tmp_dir = _safe_open(docx_path)
    body = doc.element.body
    collapsed_break_runs = _collapse_blank_break_runs(body)
    trailing_tail_removed = _cleanup_trailing_blank_tail(body)
    _safe_save(doc, docx_path, tmp_dir)

    # 统计输出
    print(f"  标题去编号: {heading_number_stripped} 个标题段落")
    print(f"  空标题清理: {heading_cleared} 个段落已清除 Heading/outlineLvl")
    print(f"  空行合并: {empty_merged_groups} 处连续空行群（共删除 {empty_merged_paras} 个空段落）")
    if blank_pages_removed > 0:
        print(f"  空白页清除: 删除 {blank_pages_removed} 张物理空白页")
    elif blank_page_error is not None:
        print(f"  空白页清除: 已跳过（{blank_page_error}）")
    if collapsed_break_runs > 0:
        print(f"  空白残留折叠: 删除 {collapsed_break_runs} 个残留段落")
    if trailing_tail_removed > 0:
        print(f"  文末清理: 删除 {trailing_tail_removed} 个尾部残留段落")


def main(argv: list[str] | None = None) -> int:
    _check_deps()

    parser = argparse.ArgumentParser(
        description="Word 文档清洗工具（officecli 后备方案）"
    )
    sub = parser.add_subparsers(dest="command", required=True)

    # peek
    p_peek = sub.add_parser("peek", help="提取前100段文本+样式")
    p_peek.add_argument("docx", help="Word 文档路径")
    p_peek.add_argument("--max-paras", type=int, default=100, help="最大段落数")

    # list-body
    p_list = sub.add_parser("list-body", help="列出 body 一级子元素")
    p_list.add_argument("docx", help="Word 文档路径")

    # trim
    p_trim = sub.add_parser("trim", help="删除锚点之前的所有元素")
    p_trim.add_argument("docx", help="Word 文档路径")
    p_trim.add_argument(
        "--anchor-index", type=int,
        help="正文起始元素的 body index（仅 list-body 输出中的 index，不是段落序号或 outline 序号）",
    )
    p_trim.add_argument(
        "--anchor-para-id",
        help="正文起始段落的 paraId；推荐优先使用，避免手工换算 body index",
    )
    p_trim.add_argument(
        "--anchor-text",
        help="用于校验锚点的期望文本；若与命中元素不一致则直接报错",
    )

    # normalize
    p_norm = sub.add_parser("normalize", help="格式规范化（标题去编号+显式分页符空白页清理+空标题元数据清理+空行合并+文末清理）")
    p_norm.add_argument("docx", help="Word 文档路径")

    args = parser.parse_args(argv)
    docx_path = Path(args.docx)

    if not docx_path.exists():
        sys.stderr.write(f"文件不存在: {docx_path}\n")
        return 1

    if args.command == "peek":
        cmd_peek(docx_path, max_paras=args.max_paras)
    elif args.command == "list-body":
        cmd_list_body(docx_path)
    elif args.command == "trim":
        cmd_trim(
            docx_path,
            anchor_index=args.anchor_index,
            anchor_para_id=args.anchor_para_id,
            anchor_text=args.anchor_text,
        )
    elif args.command == "normalize":
        cmd_normalize(docx_path)

    return 0


if __name__ == "__main__":
    sys.exit(main())
