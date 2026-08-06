#!/usr/bin/env python3
"""
excel_to_word.py — 将 Excel 文件转换为 Word 文档（表格形式）。

用法:
    python excel_to_word.py <excel_path> <output_path>

处理流程:
  1. 用 openpyxl 逐格读取所有 sheet，按单元格数字格式渲染成显示值
  2. 裁剪前导/尾部/中间的全空行和全空列
  3. 第一行作为表头（加粗），后续行为数据
  4. Sheet 未命名（Sheet1/Sheet2...）时以文件名替代
  5. 输出为 .docx 文件
"""

from __future__ import annotations

import argparse
import re
import sys
from datetime import date, datetime, time
from pathlib import Path


def _check_deps() -> None:
    """检查依赖是否可用。"""
    missing = []
    try:
        import pandas  # noqa: F401
    except ImportError:
        missing.append("pandas")
    try:
        import openpyxl  # noqa: F401
    except ImportError:
        missing.append("openpyxl")
    try:
        import docx  # noqa: F401
    except ImportError:
        missing.append("python-docx")
    if missing:
        sys.stderr.write(
            f"缺少依赖: {', '.join(missing)}\n"
            f"请运行: pip install {' '.join(missing)}\n"
        )
        sys.exit(1)


def _is_default_sheet_name(name: str) -> bool:
    """判断是否为 Excel 默认 sheet 名（Sheet1, Sheet2, 工作表1...）。"""
    return bool(re.match(r"^(Sheet|工作表)\d+$", name, re.IGNORECASE))


# 格式码里的字面量（"吨"）、颜色与条件段（[红色]、[>100]）、转义字符都不参与小数位判断
_FORMAT_LITERAL_RE = re.compile(r'"[^"]*"|\[[^\]]*\]|\\.')


def _format_spec(number_format: str) -> str:
    """取格式码的正数段并剥掉字面量、颜色与条件段。"""
    return _FORMAT_LITERAL_RE.sub("", str(number_format or "").split(";")[0])


def _general_number(value) -> str:
    """General 格式：整数值的浮点不带小数尾巴，其余原样。"""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value)


def _format_number(value, number_format: str) -> str:
    """按 Excel 数字格式渲染显示值。

    Excel 单元格分「底层值」与「显示格式」两层，读到的底层值可能远比表格里显示的精细：
    基础弯矩表把荷载列设成 `0`（整数显示），底层却存着 -149696.921875，直接 str() 会把
    12 位小数写进投标材料，凭空改变原表的有效数字约定。这里按格式码的小数位取整，
    只覆盖投标表格实际出现的几类（整数/定点小数/百分比/千分位），其余回落 General。
    """
    spec = _format_spec(number_format)
    if not spec.strip() or "general" in spec.lower():
        return _general_number(value)
    percent = "%" in spec
    if percent:
        value = value * 100
    decimal_match = re.search(r"\.([0#?]+)", spec)
    decimals = len(decimal_match.group(1)) if decimal_match else 0
    grouping = "," if re.search(r"[0#?],[0#?]", spec) else ""
    text = f"{value:{grouping}.{decimals}f}"
    return f"{text}%" if percent else text


def _format_datetime(value) -> str:
    if isinstance(value, datetime):
        if (value.hour, value.minute, value.second) == (0, 0, 0):
            return value.strftime("%Y-%m-%d")
        return value.strftime("%Y-%m-%d %H:%M:%S")
    if isinstance(value, date):
        return value.strftime("%Y-%m-%d")
    return value.strftime("%H:%M:%S")


def _cell_display(cell) -> str | None:
    """单元格的显示文本；空单元格返回 None 供后续裁剪识别。"""
    value = cell.value
    if value is None:
        return None
    if isinstance(value, str):
        return value.strip() or None
    if isinstance(value, bool):
        return "TRUE" if value else "FALSE"
    if isinstance(value, (datetime, date, time)):
        return _format_datetime(value)
    if isinstance(value, (int, float)):
        return _format_number(value, cell.number_format)
    return str(value)


def _sheet_frame(worksheet):
    import pandas as pd

    return pd.DataFrame([[_cell_display(cell) for cell in row] for row in worksheet.iter_rows()])


def _clean_dataframe(df):
    """清洗 DataFrame：去除全空行和全空列，裁剪到有效数据区域。"""
    import pandas as pd

    if df.empty:
        return df

    # 1. 删除全空列（包括前导、中间、尾部）
    non_empty_cols = [col for col in df.columns if df[col].notna().any()]
    df = df[non_empty_cols]

    if df.empty:
        return df

    # 2. 删除全空行（包括前导、中间、尾部）
    non_empty_mask = df.notna().any(axis=1)
    df = df[non_empty_mask]

    # 3. 重置索引
    df = df.reset_index(drop=True)

    return df


def convert_excel_to_word(excel_path: Path, output_path: Path) -> None:
    """读取 Excel 所有 sheet，清洗脏数据后转为 Word 表格。"""
    import openpyxl
    import pandas as pd
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()

    # A4 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Normal 样式
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    file_stem = excel_path.stem
    workbook = openpyxl.load_workbook(excel_path, data_only=True)
    sheet_names = list(workbook.sheetnames)
    sheets_written = 0

    for idx, sheet_name in enumerate(sheet_names):
        # 逐格读取（无表头模式），让我们自己判断哪行是表头
        df = _sheet_frame(workbook[sheet_name])

        # 清洗：删除全空行和全空列
        df = _clean_dataframe(df)

        if df.empty or len(df.columns) == 0:
            print(f"  跳过空 sheet: {sheet_name}")
            continue

        # 第一行作为表头
        header_row_data = [str(v) if pd.notna(v) else "" for v in df.iloc[0]]
        data_df = df.iloc[1:].reset_index(drop=True)

        # 多 sheet 时加分页符（非首个有效 sheet）
        if sheets_written > 0:
            doc.add_page_break()

        # Sheet 标题：未命名时用文件名
        title = file_stem if _is_default_sheet_name(sheet_name) else sheet_name
        # 多 sheet 且全部未命名时加序号区分
        if _is_default_sheet_name(sheet_name) and len(sheet_names) > 1:
            title = f"{file_stem} ({idx + 1})"
        doc.add_heading(title, level=2)

        rows_count = len(data_df) + 1  # 数据行 + 表头行
        cols_count = len(df.columns)

        table = doc.add_table(rows=rows_count, cols=cols_count, style="Table Grid")
        table.autofit = True

        # 填充表头
        for j, col_text in enumerate(header_row_data):
            cell = table.rows[0].cells[j]
            cell.text = col_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10.5)
            # 表头浅蓝底色
            shading = OxmlElement("w:shd")
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:fill"), "D9E2F3")
            cell.paragraphs[0]._element.get_or_add_pPr().append(shading)

        # 填充数据行
        for i in range(len(data_df)):
            for j in range(cols_count):
                cell = table.rows[i + 1].cells[j]
                val = data_df.iloc[i, j]
                cell.text = "" if pd.isna(val) else str(val)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10.5)

        sheets_written += 1

    if sheets_written == 0:
        print(f"  警告: {excel_path.name} 所有 sheet 均为空，生成空文档")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main(argv: list[str] | None = None) -> int:
    _check_deps()

    parser = argparse.ArgumentParser(
        description="将 Excel 文件转换为 Word 文档（表格形式）"
    )
    parser.add_argument("excel_path", help="输入的 Excel 文件路径")
    parser.add_argument("output_path", help="输出的 Word 文件路径")
    args = parser.parse_args(argv)

    excel_path = Path(args.excel_path)
    output_path = Path(args.output_path)

    if not excel_path.exists():
        sys.stderr.write(f"文件不存在: {excel_path}\n")
        return 1

    if excel_path.suffix.lower() not in (".xlsx", ".xls", ".xlsm"):
        sys.stderr.write(f"不支持的文件格式: {excel_path.suffix}\n")
        return 1

    try:
        print(f"  读取 Excel: {excel_path.name}")
        convert_excel_to_word(excel_path, output_path)
        print(f"  生成 Word: {output_path.name}")
    except Exception as e:
        sys.stderr.write(f"  转换失败: {e}\n")
        return 1

    return 0


if __name__ == "__main__":
    sys.exit(main())
