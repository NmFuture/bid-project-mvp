#!/usr/bin/env python3
"""bid-tech-fact-curator 的证据简报生成器。

OpenCode 先调用一次 `factcurate <manifest>`，本脚本做确定性机械工作：
- 汇总 manifest 给定的招标文件解析产物与素材为可检索文本；
- 按字段 label/reviewLabel 检索候选原文片段；
- 对 extracted 字段打机械脏数据标记（serial-text / range / unit-missing）。

本脚本不做最终事实判断：取值、修正、口径结论由 agent 读简报与原文后写入
manifest 的 outputFile。
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

SCHEMA_VERSION = "bid-tech-fact-curate-v1"

# 单来源与单字段的体量上限，避免简报失控
MAX_SOURCE_CHARS = 300_000
MAX_SNIPPETS_PER_FIELD = 4
SNIPPET_WINDOW = 100

# 常见合理区间（与后端 technical_gap_fact_table.clean_fact_value 同源）
NUMERIC_RANGES = {
    "年平均风速": (2, 15),
    "极端风速": (20, 100),
    "轮毂高度": (40, 250),
    "叶轮直径": (50, 350),
    "空气密度": (0.7, 1.5),
    "湍流强度": (0, 1),
    "风剪切": (0, 1),
    "机组台数": (1, 1000),
    "保证有效小时数": (1, 8760),
}


def _load_json(path_text: str) -> Any:
    path = Path(path_text)
    if not path.is_file():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None


def _flatten_json_text(value: Any, limit: int = MAX_SOURCE_CHARS) -> str:
    """把结构化 JSON 拍平为可检索文本（键与值拼行）。"""
    lines: list[str] = []

    def visit(node: Any) -> None:
        if sum(len(line) for line in lines) >= limit:
            return
        if isinstance(node, dict):
            for key, child in node.items():
                if isinstance(child, (dict, list)):
                    visit(child)
                elif child not in (None, ""):
                    lines.append(f"{key}: {child}")
        elif isinstance(node, list):
            for child in node:
                visit(child)

    visit(value)
    return "\n".join(lines)[:limit]


def _read_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        return ""
    try:
        document = Document(str(path))
    except Exception:
        return ""
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_xlsx_text(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return ""
    try:
        workbook = load_workbook(path, data_only=True, read_only=True)
    except Exception:
        return ""
    parts: list[str] = []
    for worksheet in workbook.worksheets:
        for row in worksheet.iter_rows(values_only=True):
            cells = [str(cell).strip() for cell in row if cell not in (None, "")]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_source_text(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix in {".txt", ".md", ".csv"}:
            return path.read_text(encoding="utf-8", errors="ignore")[:MAX_SOURCE_CHARS]
        if suffix == ".json":
            return _flatten_json_text(_load_json(str(path)))
        if suffix in {".docx", ".doc"}:
            return _read_docx_text(path)[:MAX_SOURCE_CHARS]
        if suffix in {".xlsx", ".xlsm"}:
            return _read_xlsx_text(path)[:MAX_SOURCE_CHARS]
    except OSError:
        return ""
    return ""


def _corpus(manifest: dict[str, Any]) -> list[dict[str, str]]:
    """manifest 给定的招标文件解析产物 + 素材 → [(名称, 文本)]。"""
    corpus: list[dict[str, str]] = []
    for source in manifest.get("tenderSources") or []:
        if not isinstance(source, dict):
            continue
        path = Path(str(source.get("path") or ""))
        text = _read_source_text(path)
        if text:
            corpus.append({"name": f"招标解析/{source.get('kind') or path.name}", "text": text})
    for material in manifest.get("materials") or []:
        if not isinstance(material, dict):
            continue
        path = Path(str(material.get("path") or ""))
        text = _read_source_text(path)
        if text:
            corpus.append({"name": str(material.get("name") or path.name), "text": text})
    return corpus


def _search_terms(field: dict[str, Any]) -> list[str]:
    terms: list[str] = []
    for raw in (field.get("label"), field.get("reviewLabel")):
        text = str(raw or "").strip()
        if not text:
            continue
        # 去掉括号内单位标注（如 招标单机容量（出口端，MW））
        stripped = re.sub(r"[（(][^）)]*[）)]", "", text).strip()
        for term in (text, stripped):
            if term and term not in terms:
                terms.append(term)
    return terms


def _snippets(field: dict[str, Any], corpus: list[dict[str, str]]) -> list[dict[str, str]]:
    found: list[dict[str, str]] = []
    for term in _search_terms(field):
        for source in corpus:
            for match in re.finditer(re.escape(term), source["text"]):
                start = max(0, match.start() - SNIPPET_WINDOW)
                end = min(len(source["text"]), match.end() + SNIPPET_WINDOW)
                fragment = re.sub(r"\s+", " ", source["text"][start:end]).strip()
                found.append({"source": source["name"], "text": fragment})
                if len(found) >= MAX_SNIPPETS_PER_FIELD:
                    return found
    return found


def _first_number(value: str) -> float | None:
    match = re.search(r"([0-9]+(?:\.[0-9]+)?)", value)
    return float(match.group(1)) if match else None


def _dirty_flags(field: dict[str, Any]) -> list[str]:
    """extracted 字段的机械脏数据标记，只报线索不下结论。"""
    value = str(field.get("value") or "").strip()
    if not value:
        return []
    flags: list[str] = []
    # 跨列串行：多数值用 / 连接，或数值后面拖着另一字段的中文名
    if re.search(r"\d+(?:\.\d+)?\s*/\s*\d", value):
        flags.append("serial-text")
    elif re.match(r"^[0-9.,/\s]+[一-鿿（(]", value):
        flags.append("serial-text")
    label = re.sub(r"[（(][^）)]*[）)]", "", str(field.get("label") or "")).strip()
    number = _first_number(value)
    if label in NUMERIC_RANGES and number is not None:
        low, high = NUMERIC_RANGES[label]
        if not (low <= number <= high):
            flags.append("range")
    if re.search(r"[（(][^）)]*(?:m/s|MW|kW|m|MWh|h|%|台)[^）)]*[）)]", str(field.get("label") or "")) and not str(
        field.get("unit") or ""
    ).strip():
        flags.append("unit-missing")
    return flags


def build_brief(manifest: dict[str, Any]) -> dict[str, Any]:
    table = manifest.get("projectFactTable") if isinstance(manifest.get("projectFactTable"), dict) else {}
    fields = [field for field in (table.get("fields") or []) if isinstance(field, dict)]
    targets = manifest.get("targets") if isinstance(manifest.get("targets"), dict) else {}
    bucket_by_key: dict[str, str] = {}
    for bucket, action in (("fill", "fill"), ("fix", "fix"), ("confirmAdvice", "confirm-advice")):
        for key in targets.get(bucket) or []:
            bucket_by_key[str(key)] = action

    corpus = _corpus(manifest)
    brief_fields: list[dict[str, Any]] = []
    for field in fields:
        field_key = str(field.get("key") or "").strip()
        action = bucket_by_key.get(field_key)
        if not field_key or not action:
            continue
        brief_fields.append(
            {
                "fieldKey": field_key,
                "action": action,
                "label": str(field.get("label") or ""),
                "reviewLabel": str(field.get("reviewLabel") or ""),
                "currentValue": str(field.get("value") or ""),
                "unit": str(field.get("unit") or ""),
                "specKey": str(field.get("specKey") or ""),
                "specSeq": field.get("specSeq"),
                "sourceKind": str(field.get("sourceKind") or ""),
                "needsConfirmation": bool(field.get("needsConfirmation")),
                "flags": _dirty_flags(field) if action == "fix" else [],
                "snippets": _snippets(field, corpus),
            }
        )
    return {
        "schema": SCHEMA_VERSION,
        "phase": "brief",
        "projectId": str(manifest.get("projectId") or ""),
        "sourceCount": len(corpus),
        "fields": brief_fields,
    }


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="bid-tech-fact-curator 证据简报生成器")
    parser.add_argument("--manifest", required=True)
    parser.add_argument("--response", default="summary")
    args = parser.parse_args(argv)

    manifest = _load_json(args.manifest)
    if not isinstance(manifest, dict):
        print(json.dumps({"schema": SCHEMA_VERSION, "error": f"manifest 不可读：{args.manifest}"}, ensure_ascii=False))
        return 1

    brief = build_brief(manifest)
    brief_file = Path(str(manifest.get("briefFile") or ""))
    if not brief_file.parent.exists():
        brief_file.parent.mkdir(parents=True, exist_ok=True)
    brief_file.write_text(json.dumps(brief, ensure_ascii=False, indent=2), encoding="utf-8")

    counts = {"fill": 0, "fix": 0, "confirmAdvice": 0}
    for field in brief["fields"]:
        bucket = {"fill": "fill", "fix": "fix", "confirm-advice": "confirmAdvice"}[field["action"]]
        counts[bucket] += 1
    summary = {
        "schema": SCHEMA_VERSION,
        "phase": "brief",
        "briefFile": str(brief_file),
        "outputFile": str(manifest.get("outputFile") or ""),
        "sourceCount": brief["sourceCount"],
        "counts": counts,
        "flaggedFields": [field["fieldKey"] for field in brief["fields"] if field["flags"]],
    }
    print(json.dumps(summary, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    sys.exit(main())
