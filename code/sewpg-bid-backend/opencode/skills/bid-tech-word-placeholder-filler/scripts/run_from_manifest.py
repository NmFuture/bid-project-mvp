#!/usr/bin/env python3
"""Fill placeholders in technical-bid material Word files from manifest data."""

from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass
from datetime import UTC, datetime
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import load_workbook


SCHEMA_VERSION = "bid-tech-word-placeholder-fill-v1"
PLACEHOLDER_RE = re.compile(
    r"([\[【])\s*(?:待填写[:：]\s*)?([^\]】\r\n]{1,80}?)(?:[,，、:：\s]*(?:待填写|待补充|待确认))?\s*([\]】])"
)
MANUAL_FILL = "FFF2CC"
GENERIC_SOURCE_LIMIT = 500


@dataclass
class Source:
    name: str
    path: Path
    kind: str
    priority: int
    route: str
    material_id: str = ""


def now_iso() -> str:
    return datetime.now(UTC).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def clean(value: Any) -> str:
    text = str(value or "").replace("\u3000", " ").replace("\xa0", " ")
    return re.sub(r"\s+", " ", text.replace("\n", " / ")).strip()


# 归一化剥离的标点：字符类里的 `]` 必须转义，否则字符类在 `\\]` 处提前闭合，
# 整个表达式退化成「一个标点后面必须紧跟字面 `【】]`」而几乎永不匹配——
# 修复前 norm() 实际不剥离任何标点，全半角括号差异会让占位符与字段名对不上。
NORM_PUNCT_RE = re.compile(r"[\s（）()、/\\:：；;，,。\-_—×*\[\]【】]+")


def norm(value: Any) -> str:
    return NORM_PUNCT_RE.sub("", clean(value).lower())


def placeholderish(value: Any) -> bool:
    text = clean(value)
    return any(marker in text for marker in ("待填写", "待补充", "待确认", "待人工补充", "[", "【"))


def placeholder_key(raw: Any) -> str:
    """占位符归一化键：清单列与 Word 正文共用同一套解析，消除括号/全半角/分隔符差异。

    清单里写 `[安全等级，待填写]`、Word 里写 `【安全等级, 待填写】`，归一化后都是
    `安全等级`。
    """
    text = clean(raw)
    if not text:
        return ""
    match = PLACEHOLDER_RE.search(text)
    body = match.group(2) if match else text
    return norm(placeholder_label(body))


def file_key(value: Any) -> str:
    """待填写文件归一化键：只取文件名去扩展名。

    清单第 2 列带素材库路径（`华能/待填写-塔筒设计方案专题报告.docx`），实际 blankSource
    可能来自别处；18 个待填写文件的文件名无重名，按文件名匹配即可跨客户/项目复用清单。
    """
    text = clean(value).replace("\\", "/").split("/")[-1]
    return norm(re.sub(r"\.(docx|doc)$", "", text, flags=re.IGNORECASE))


def split_field_enumeration(value: Any) -> list[str]:
    """按顿号/逗号拆字段枚举，但括号内的分隔符不算。

    字段名自带括号逗号很常见（`折减系数（考核值，%）`），无条件按逗号拆会把单个字段名
    切碎；`[机型认证湍流强度，机型认证10分钟平均极限风速（m/s），待填写]` 这种真正的
    多字段枚举又确实用逗号分隔，只认顿号会漏。屏蔽括号内容后再定位分隔符位置。
    """
    text = clean(value)
    if not text:
        return []
    masked = re.sub(r"[（(][^）)]*[）)]", lambda match: "\x01" * len(match.group(0)), text)
    parts: list[str] = []
    start = 0
    for position, char in enumerate(masked):
        if char in "、,，":
            parts.append(text[start:position])
            start = position + 1
    parts.append(text[start:])
    return [part.strip() for part in parts if part.strip()]


def split_spec_cell(value: Any) -> list[str]:
    """清单单元格拆多值：分号（全/半角）与换行都算分隔符。

    不走 clean()——它会把换行压成 " / "，而换行正是清单里一格多文件的分隔符。
    """
    text = str(value or "").replace("　", " ").replace("\xa0", " ")
    return [part.strip() for part in re.split(r"[；;\n]+", text) if part.strip()]


def compact_number(value: Any) -> str:
    text = clean(value)
    if re.fullmatch(r"\d+\.0+", text):
        return text.split(".", 1)[0]
    return text


def object_items(value: Any) -> list[dict[str, Any]]:
    return [item for item in value if isinstance(item, dict)] if isinstance(value, list) else []


def material_label(material: dict[str, Any]) -> str:
    return clean(
        material.get("name")
        or material.get("cleanedFileName")
        or material.get("fileName")
        or material.get("title")
        or material.get("id")
        or material.get("path")
    )


def first_path(values: Iterable[Any], base: Path) -> Path | None:
    for value in values:
        text = clean(value)
        if not text:
            continue
        path = Path(text).expanduser()
        if not path.is_absolute():
            path = base / path
        if path.exists():
            return path
    return None


def load_manifest(path: Path) -> dict[str, Any]:
    payload = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(payload, dict):
        raise RuntimeError("manifest root must be an object")
    return payload


def template_path(manifest: dict[str, Any], manifest_path: Path) -> Path:
    blank = manifest.get("blankSource") if isinstance(manifest.get("blankSource"), dict) else {}
    appendix = manifest.get("appendixTask") if isinstance(manifest.get("appendixTask"), dict) else {}
    path = first_path(
        (
            blank.get("docxPath"),
            blank.get("path"),
            blank.get("workspacePath"),
            appendix.get("docxPath"),
            manifest.get("templatePath"),
        ),
        manifest_path.parent,
    )
    if path is None or path.suffix.lower() != ".docx":
        raise RuntimeError("manifest does not contain a readable Word template path")
    return path


def output_path(manifest: dict[str, Any], manifest_path: Path, source: Path) -> Path:
    explicit = clean(manifest.get("outputFile"))
    if explicit:
        path = Path(explicit).expanduser()
        if not path.is_absolute():
            path = manifest_path.parent / path
        return path
    return manifest_path.with_name(f"{source.stem}_AI填写.docx")


def material_path(material: dict[str, Any], base: Path) -> Path | None:
    return first_path(
        (
            material.get("path"),
            material.get("docxPath"),
            material.get("workspacePath"),
            material.get("filePath"),
        ),
        base,
    )


def same_as_blank(material: dict[str, Any], blank: dict[str, Any]) -> bool:
    ids = {clean(material.get("id")), clean(material.get("materialId"))} - {""}
    blank_ids = {clean(blank.get("id")), clean(blank.get("materialId"))} - {""}
    if ids & blank_ids:
        return True
    material_name = material_label(material)
    return bool(material_name and material_name == clean(blank.get("title")))


def add_source(
    sources: list[Source],
    material: dict[str, Any],
    base: Path,
    *,
    priority: int,
    route: str,
) -> None:
    path = material_path(material, base)
    if path is None or path.suffix.lower() not in {".docx", ".xlsx", ".xlsm"}:
        return
    sources.append(
        Source(
            name=material_label(material) or path.name,
            path=path,
            kind="xlsx" if path.suffix.lower() in {".xlsx", ".xlsm"} else "docx",
            priority=priority,
            route=route,
            material_id=clean(material.get("id") or material.get("materialId")),
        )
    )


def select_sources(manifest: dict[str, Any], manifest_path: Path) -> list[Source]:
    blank = manifest.get("blankSource") if isinstance(manifest.get("blankSource"), dict) else {}
    contexts = (
        object_items(manifest.get("referenceMaterials"))
        + object_items(manifest.get("selectedReferenceMaterials"))
        + object_items(manifest.get("materialIndex"))
        + object_items(manifest.get("recommendedMaterials"))
    )
    sources: list[Source] = []
    for material in contexts:
        if same_as_blank(material, blank):
            continue
        priority = 100 if material in object_items(manifest.get("referenceMaterials")) else 60
        route = "人工指定/manifest referenceMaterials" if priority == 100 else "materialIndex 自动参考"
        add_source(sources, material, manifest_path.parent, priority=priority, route=route)

    deduped: dict[str, Source] = {}
    for source in sources:
        key = str(source.path.resolve())
        previous = deduped.get(key)
        if previous is None or source.priority > previous.priority:
            deduped[key] = source
    return sorted(deduped.values(), key=lambda item: (-item.priority, item.kind != "xlsx", item.name))[:80]


def add_fact(
    facts: list[dict[str, Any]],
    *,
    label: Any,
    value: Any,
    source: Any,
    confidence: float = 0.75,
    location: str = "",
    fact_type: str = "generic",
) -> None:
    label_text = clean(label)
    value_text = clean(value)
    if not label_text or not value_text or value_text in {"/", "-", "待填写", "待补充"} or placeholderish(value_text):
        return
    facts.append(
        {
            "id": f"F{len(facts) + 1:04d}",
            "label": label_text,
            "value": value_text,
            "source": source.name if isinstance(source, Source) else str(source),
            "sourcePath": str(source.path) if isinstance(source, Source) else "",
            "sourceKind": source.kind if isinstance(source, Source) else "manifest",
            "sourcePriority": source.priority if isinstance(source, Source) else 90,
            "location": location,
            "factType": fact_type,
            "confidence": confidence,
        }
    )


def add_derived_fact(
    facts: list[dict[str, Any]],
    label: str,
    value: Any,
    *,
    source: str,
    confidence: float = 0.9,
    location: str = "",
    fact_type: str = "derived",
) -> None:
    add_fact(facts, label=label, value=value, source=source, confidence=confidence, location=location, fact_type=fact_type)


def readable_model(model: Any, aliases: Any = None) -> str:
    candidates = [clean(item) for item in aliases or [] if clean(item)]
    for item in candidates:
        if item and "上置" not in item and "下置" not in item:
            return item
    return clean(model)


def project_fact_table_facts(manifest: dict[str, Any]) -> list[dict[str, Any]]:
    facts: list[dict[str, Any]] = []
    table = manifest.get("projectFactTable") if isinstance(manifest.get("projectFactTable"), dict) else {}
    table_status = clean(table.get("status"))
    for field in object_items(table.get("fields")):
        status = clean(field.get("status")) or ("confirmed" if table_status == "confirmed" else "")
        if status not in {"confirmed", "candidate"}:
            continue
        confidence = 0.97 if status == "confirmed" else 0.84
        before = len(facts)
        add_fact(
            facts,
            label=field.get("label"),
            value=field.get("value"),
            source="projectFactTable",
            confidence=confidence,
            location=status,
            fact_type="confirmed_project_fact" if status == "confirmed" else "candidate_project_fact",
        )
        if len(facts) > before:
            # 清单第 2/3 列随字段下发：正文按「待填写文件 + 占位符原文」定位，不靠字面相似度猜
            facts[-1]["specPlaceholders"] = [key for key in map(placeholder_key, split_spec_cell(field.get("placeholder"))) if key]
            facts[-1]["specTargets"] = [key for key in map(file_key, split_spec_cell(field.get("targetFile"))) if key]
            facts[-1]["reviewLabel"] = clean(field.get("reviewLabel"))
    return facts


def collect_project_facts(manifest: dict[str, Any]) -> list[dict[str, Any]]:
    facts: list[dict[str, Any]] = project_fact_table_facts(manifest)
    identity = manifest.get("projectIdentity") if isinstance(manifest.get("projectIdentity"), dict) else {}
    project = manifest.get("projectTurbineModel") if isinstance(manifest.get("projectTurbineModel"), dict) else {}
    model = project.get("model") or project.get("turbineModel")
    model_display = readable_model(model, project.get("aliases"))
    add_fact(facts, label="项目名称", value=manifest.get("projectName"), source="manifest", confidence=0.75, fact_type="project")
    owner = (
        identity.get("owner")
        or identity.get("customerCanonicalName")
        or identity.get("customerName")
        or manifest.get("customerName")
    )
    add_fact(facts, label="招标方", value=owner, source="projectIdentity", confidence=0.92, fact_type="project_identity")
    add_fact(facts, label="招标人", value=owner, source="projectIdentity", confidence=0.92, fact_type="project_identity")
    add_fact(facts, label="客户名称", value=identity.get("customerName") or manifest.get("customerName"), source="projectIdentity", confidence=0.9, fact_type="project_identity")
    add_fact(facts, label="客户标准名称", value=identity.get("customerCanonicalName"), source="projectIdentity", confidence=0.9, fact_type="project_identity")
    add_fact(facts, label="投标机型", value=model_display or model, source="projectTurbineModel", confidence=0.98, fact_type="project")
    add_fact(facts, label="投标机型完整", value=model, source="projectTurbineModel", confidence=0.9, fact_type="project")
    rated_kw = project.get("ratedPowerKw")
    add_fact(facts, label="单机容量", value=compact_number(float(rated_kw) / 1000) if isinstance(rated_kw, (int, float)) else rated_kw, source="projectTurbineModel", confidence=0.9, fact_type="project")
    add_fact(facts, label="叶轮直径", value=project.get("rotorDiameterM"), source="projectTurbineModel", confidence=0.9, fact_type="project")
    add_fact(facts, label="轮毂高度", value=project.get("hubHeightM"), source="projectTurbineModel", confidence=0.86, fact_type="project")
    for field in object_items(manifest.get("parseFields")):
        label = field.get("label") or field.get("title") or field.get("keyEntity") or field.get("id")
        value = field.get("value") or field.get("keyValue")
        location = clean(field.get("sourceFile") or field.get("evidenceLocation"))
        add_fact(facts, label=label, value=value, source="parseFields", confidence=0.86, location=location, fact_type="parse_field")
        evidence = clean(field.get("evidence"))
        if evidence and evidence != clean(value):
            add_fact(facts, label=label, value=evidence, source="parseFields", confidence=0.72, location=location, fact_type="evidence_text")
    return facts


def doc_facts(source: Source) -> list[dict[str, Any]]:
    facts: list[dict[str, Any]] = []
    try:
        doc = Document(str(source.path))
    except Exception:
        return facts
    for idx, paragraph in enumerate(doc.paragraphs, start=1):
        text = clean(paragraph.text)
        if not text:
            continue
        match = re.match(r"^([^:：]{2,50})[:：]\s*(.{1,500})$", text)
        if match:
            add_fact(facts, label=match.group(1), value=match.group(2), source=source, confidence=0.7, location=f"P{idx}", fact_type="paragraph_kv")
        elif len(text) <= 500:
            add_fact(facts, label=source.name, value=text, source=source, confidence=0.52, location=f"P{idx}", fact_type="paragraph")
        if len(facts) >= GENERIC_SOURCE_LIMIT:
            return facts
    for table_idx, table in enumerate(doc.tables, start=1):
        for row_idx, row in enumerate(table.rows, start=1):
            values = [clean(cell.text) for cell in row.cells]
            nonempty = [value for value in values if value]
            if len(nonempty) >= 2:
                add_fact(
                    facts,
                    label=nonempty[0],
                    value=nonempty[1] if len(nonempty) == 2 else " / ".join(nonempty[1:]),
                    source=source,
                    confidence=0.74,
                    location=f"T{table_idx}/R{row_idx}",
                    fact_type="table_row",
                )
            if len(facts) >= GENERIC_SOURCE_LIMIT:
                return facts
    return facts


def xlsx_facts(source: Source) -> list[dict[str, Any]]:
    facts: list[dict[str, Any]] = []
    try:
        wb = load_workbook(source.path, data_only=True, read_only=True)
    except Exception:
        return facts
    for ws in wb.worksheets:
        for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
            values = [clean(value) for value in row]
            nonempty = [value for value in values if value]
            if len(nonempty) >= 2:
                add_fact(
                    facts,
                    label=nonempty[0],
                    value=nonempty[1] if len(nonempty) == 2 else " / ".join(nonempty[1:]),
                    source=source,
                    confidence=0.7,
                    location=f"{ws.title}/R{row_idx}",
                    fact_type="excel_row",
                )
            if len(facts) >= GENERIC_SOURCE_LIMIT:
                wb.close()
                return facts
    wb.close()
    return facts


def doc_text_and_tables(source: Source) -> tuple[list[str], list[list[list[str]]]]:
    try:
        doc = Document(str(source.path))
    except Exception:
        return [], []
    paragraphs = [clean(paragraph.text) for paragraph in doc.paragraphs if clean(paragraph.text)]
    tables = [[[clean(cell.text) for cell in row.cells] for row in table.rows] for table in doc.tables]
    return paragraphs, tables


def first_fact_value(facts: list[dict[str, Any]], *labels: str) -> str:
    wanted = [norm(label) for label in labels]
    for label in wanted:
        for fact in facts:
            if norm(fact.get("label")) == label and clean(fact.get("value")):
                return clean(fact.get("value"))
    for label in wanted:
        for fact in facts:
            fact_label = norm(fact.get("label"))
            if label and (label in fact_label or fact_label in label) and clean(fact.get("value")):
                return clean(fact.get("value"))
    return ""


def derive_textual_facts(manifest: dict[str, Any], sources: list[Source], facts: list[dict[str, Any]]) -> None:
    model = first_fact_value(facts, "投标机型")
    hub_height = first_fact_value(facts, "轮毂高度")
    turbine_count = ""
    capacity_mw = ""
    for source in sources:
        if source.kind != "docx":
            continue
        paragraphs, tables = doc_text_and_tables(source)
        full_text = "\n".join(paragraphs)
        if "招标文件" in source.name:
            for paragraph in paragraphs[:40]:
                if "项目" in paragraph and "风电项目" in paragraph and "目录" not in paragraph and len(paragraph) >= 20:
                    add_derived_fact(facts, "项目名称", paragraph.lstrip("*"), source=source.name, confidence=0.94, location="frontmatter", fact_type="tender_project")
                    break
            for table_idx, table in enumerate(tables, start=1):
                for row_idx, row in enumerate(table, start=1):
                    cells = [cell for cell in row if cell]
                    if len(cells) >= 2 and cells[0] in {"招标人", "管理单位", "招标代理机构"}:
                        add_derived_fact(facts, cells[0], cells[-1], source=source.name, confidence=0.92, location=f"T{table_idx}/R{row_idx}", fact_type="tender_party")
                        if cells[0] == "招标人":
                            add_derived_fact(facts, "招标方", cells[-1], source=source.name, confidence=0.92, location=f"T{table_idx}/R{row_idx}", fact_type="tender_party")
                    if len(cells) >= 3 and "功率曲线保证率折减" in cells[0]:
                        add_derived_fact(facts, "功率曲线保证率", cells[1], source=source.name, confidence=0.86, location=f"T{table_idx}/R{row_idx}", fact_type="tender_requirement")
            match = re.search(r"风电场机组年平均可利用率≥\s*(\d+%)；单台机组年平均可利用率≥\s*(\d+%)", full_text)
            if match:
                add_derived_fact(facts, "全场可利用率", match.group(1), source=source.name, confidence=0.9, location="性能保证", fact_type="tender_requirement")
                add_derived_fact(facts, "单台可利用率", match.group(2), source=source.name, confidence=0.9, location="性能保证", fact_type="tender_requirement")
            match = re.search(r"全部机组平均可利用率≥\s*(\d+%)", full_text)
            if match:
                add_derived_fact(facts, "全场可利用率", match.group(1), source=source.name, confidence=0.88, location="性能考核", fact_type="tender_requirement")
            match = re.search(r"主要部件更换率低于\s*(\d+%)", full_text)
            if match:
                add_derived_fact(facts, "主要部件更换率", match.group(1), source=source.name, confidence=0.9, location="性能考核", fact_type="tender_requirement")
            match = re.search(r"免费提供(\d+)套[^。\n]{0,80}单叶片吊装", full_text)
            if match:
                add_derived_fact(facts, "单叶片吊装及辅助工装工具套数", match.group(1), source=source.name, confidence=0.9, location="安装与吊装工具", fact_type="tender_requirement")
            match = re.search(r"运行和维护所需的(\d+)套完整工具", full_text)
            if match:
                add_derived_fact(facts, "运行维护工具套数", match.group(1), source=source.name, confidence=0.9, location="运行维护工具", fact_type="tender_requirement")
            match = re.search(r"升降机采用([^。；]+导向方式)", full_text)
            if match:
                add_derived_fact(facts, "升降机导向方式", match.group(1), source=source.name, confidence=0.86, location="升降机", fact_type="tender_requirement")
        if "物流解决方案" in source.name:
            for table in tables[:1]:
                for row in table[1:]:
                    if len(row) >= 4:
                        name, count, _, origin = row[:4]
                        if "机舱" in name and count:
                            turbine_count = count
                            add_derived_fact(facts, "台数", count, source=source.name, confidence=0.92, location="T1", fact_type="logistics")
                            add_derived_fact(facts, "主机供货制造基地", f"上海电气{origin}生产基地", source=source.name, confidence=0.86, location="T1", fact_type="logistics")
                        if "叶片" in name and origin:
                            add_derived_fact(facts, "叶片供货制造基地", f"中材叶片{origin}生产基地", source=source.name, confidence=0.82, location="T1", fact_type="logistics")
        if "上海电气生产能力介绍" in source.name:
            high_tai_intro = next((paragraph for paragraph in paragraphs if "高台生产基地于" in paragraph), "")
            high_tai_capacity = next((paragraph for paragraph in paragraphs if "高台县南华镇工业园区" in paragraph and "年产能" in paragraph), "")
            if high_tai_intro or high_tai_capacity:
                add_derived_fact(facts, "主机供货制造基地", "上海电气高台生产基地", source=source.name, confidence=0.86, location="高台工厂", fact_type="manufacturing_base")
                add_derived_fact(facts, "叶片供货制造基地", "上海电气风电（张掖）叶片科技有限公司", source=source.name, confidence=0.84, location="高台工厂", fact_type="manufacturing_base")
            if high_tai_intro:
                add_derived_fact(facts, "主机供货制造基地介绍", high_tai_intro, source=source.name, confidence=0.84, location="高台工厂", fact_type="manufacturing_base")
            if high_tai_capacity:
                add_derived_fact(facts, "叶片供货制造基地介绍", high_tai_capacity, source=source.name, confidence=0.84, location="高台工厂", fact_type="manufacturing_base")
        if "塔架与基础工程量" in source.name:
            for table in tables[:1]:
                for row in table:
                    if len(row) >= 3 and "轮毂高度" in row[0]:
                        hub_height = compact_number(row[2])
                        add_derived_fact(facts, "轮毂高度", hub_height, source=source.name, confidence=0.9, location="T1", fact_type="tower")
        if "风资源评估报告" in source.name:
            for table_idx, table in enumerate(tables, start=1):
                for row_idx, row in enumerate(table, start=1):
                    if len(row) >= 5 and row[1] == "项目名称":
                        add_derived_fact(facts, "项目名称", row[3], source=source.name, confidence=0.96, location=f"T{table_idx}/R{row_idx}", fact_type="wind_resource")
                    if len(row) >= 5 and row[1] == "建设容量":
                        capacity_mw = compact_number(row[3])
                        add_derived_fact(facts, "容量", capacity_mw, source=source.name, confidence=0.94, location=f"T{table_idx}/R{row_idx}", fact_type="wind_resource")
                    if len(row) >= 5 and row[1] == "方案":
                        add_derived_fact(facts, "投标方案", row[3], source=source.name, confidence=0.92, location=f"T{table_idx}/R{row_idx}", fact_type="wind_resource")
                    if len(row) >= 5 and row[0] == "轮毂高度处代表年平均风速":
                        add_derived_fact(facts, "测风塔210315年平均风速", row[-1], source=source.name, confidence=0.9, location=f"T{table_idx}/R{row_idx}", fact_type="wind_resource")
                    if len(row) >= 5 and row[0] == "位置":
                        add_derived_fact(facts, "测风塔210315坐标", row[-1], source=source.name, confidence=0.9, location=f"T{table_idx}/R{row_idx}", fact_type="wind_resource")
        if "发电量担保" in source.name:
            is_assessment = "考核值" in source.name
            for table_idx, table in enumerate(tables[:1], start=1):
                for row_idx, row in enumerate(table, start=1):
                    if len(row) >= 3 and compact_number(row[0]) == "7.2":
                        prefix = "考核" if is_assessment else "保证"
                        add_derived_fact(facts, f"{prefix}净发电量", row[1], source=source.name, confidence=0.92, location=f"T{table_idx}/R{row_idx}", fact_type="energy_commitment")
                        add_derived_fact(facts, f"{prefix}有效小时数", row[2], source=source.name, confidence=0.92, location=f"T{table_idx}/R{row_idx}", fact_type="energy_commitment")
        if "业绩情况" in source.name and "待填写" not in source.name:
            for paragraph in paragraphs:
                for label, pattern in (
                    ("试运行业绩", r"试运行业绩共计(\d+)台，共计([\d.]+)MW"),
                    ("合同业绩", r"合同业绩共计(\d+)台，共计([\d.]+)MW.*?提供(\d+)台"),
                    ("正在供货和新承接项目", r"正在供货和新承接[^。]*?共计(\d+)台[^。]*?共计([\d.]+)MW"),
                ):
                    match = re.search(pattern, paragraph)
                    if not match:
                        continue
                    add_derived_fact(facts, f"{label}台数", match.group(1), source=source.name, confidence=0.93, location="paragraph", fact_type="performance")
                    add_derived_fact(facts, f"{label}容量", match.group(2), source=source.name, confidence=0.93, location="paragraph", fact_type="performance")
                    if label == "合同业绩" and len(match.groups()) >= 3:
                        add_derived_fact(facts, "合同发票台数", match.group(3), source=source.name, confidence=0.9, location="paragraph", fact_type="performance")
    single_mw = first_fact_value(facts, "单机容量")
    turbine_count = turbine_count or first_fact_value(facts, "台数")
    capacity_mw = capacity_mw or first_fact_value(facts, "容量")
    if not capacity_mw and turbine_count and single_mw:
        try:
            capacity_mw = compact_number(float(turbine_count) * float(single_mw))
            add_derived_fact(facts, "容量", capacity_mw, source="derived", confidence=0.84, fact_type="computed_project")
        except ValueError:
            pass
    if model and hub_height:
        add_derived_fact(facts, "方案", f"{turbine_count}*{model}-{hub_height}" if turbine_count else f"{model}-{hub_height}", source="derived", confidence=0.86, fact_type="computed_project")
        add_derived_fact(facts, "投标方案", f"{turbine_count}台{model}-{hub_height}" if turbine_count else f"{model}-{hub_height}", source="derived", confidence=0.86, fact_type="computed_project")


def collect_facts(manifest: dict[str, Any], sources: list[Source]) -> list[dict[str, Any]]:
    facts = collect_project_facts(manifest)
    for source in sources:
        facts.extend(xlsx_facts(source) if source.kind == "xlsx" else doc_facts(source))
    derive_textual_facts(manifest, sources, facts)
    for index, fact in enumerate(facts, start=1):
        fact["id"] = f"F{index:04d}"
    return facts


def placeholder_label(raw: str) -> str:
    label = clean(raw)
    label = re.sub(r"^(待填写|缺失)[:：]", "", label).strip()
    label = re.sub(r"[,，、:：\s]*(待填写|待补充|待确认)$", "", label).strip()
    return label or "待填写内容"


def find_placeholders(text: str) -> list[dict[str, Any]]:
    result: list[dict[str, Any]] = []
    for match in PLACEHOLDER_RE.finditer(text or ""):
        body = placeholder_label(match.group(2))
        full = match.group(0)
        if "待填写" not in full and "待补充" not in full and "待确认" not in full:
            continue
        result.append({"full": full, "label": body, "start": match.start(), "end": match.end()})
    return result


def label_score(label: str, context: str, fact: dict[str, Any]) -> float:
    label_norm = norm(label)
    fact_label = norm(fact.get("label"))
    if not label_norm or not fact_label:
        return 0.0
    material_fact = str(fact.get("sourceKind") or "") in {"docx", "xlsx"}
    score = 0.0
    if label_norm == fact_label:
        score = 0.94
    elif label_norm in fact_label:
        # 方向性约束（金标反评/商务标 519f8a4 同款）：占位符标签 ⊂ 事实键才算强包含
        score = 0.82
    elif fact_label in label_norm:
        # 反向包含只对足够长的事实键放行——短键蹭长标签是污染源
        if len(fact_label) < 5:
            return 0.0
        score = 0.78
    else:
        if material_fact:
            # 素材文档抽的事实禁走模糊路径：金标反评发现无关素材参数
            # （如低温冲击功）靠 SequenceMatcher 蹭分批量污染承诺值占位符。
            return 0.0
        score = SequenceMatcher(None, label_norm, fact_label).ratio() * 0.78
    context_norm = norm(context)
    value_norm = norm(fact.get("value"))
    if context_norm and fact_label in context_norm:
        score += 0.08
    if context_norm and value_norm and len(value_norm) >= 4 and value_norm in context_norm:
        score += 0.04
    if fact.get("sourceKind") == "manifest":
        score += 0.04
    score *= float(fact.get("confidence") or 0.5)
    return round(min(score, 0.99), 3)


def choose_fact(
    label: str,
    context: str,
    facts: list[dict[str, Any]],
    used_material_facts: dict[tuple[str, str], str] | None = None,
) -> tuple[dict[str, Any] | None, list[dict[str, Any]]]:
    candidates = []
    for fact in facts:
        score = label_score(label, context, fact)
        if score > 0:
            item = dict(fact)
            item["score"] = score
            # 事实表/派生事实优先（金标反评：错值的正确答案大多躺在已确认事实表里，
            # 却被素材证据以相近分捞走）——同分段位内 manifest 来源先选。
            item["_trustBoost"] = 0.1 if str(item.get("sourceKind") or "") == "manifest" else 0.0
            candidates.append(item)
    candidates.sort(key=lambda item: (item["score"] + item["_trustBoost"], item.get("sourcePriority", 0)), reverse=True)

    def usable(item: dict[str, Any]) -> bool:
        if item["score"] < 0.62:
            return False
        if used_material_facts is not None and str(item.get("sourceKind") or "") in {"docx", "xlsx"}:
            # 素材证据单次使用：同一条素材事实不得绑定到语义不同的占位符
            #（金标反评：一条证据喷洒 23 处占位符）。同标签复用（表格重复行）放行。
            key = (norm(item.get("label")), norm(item.get("value")))
            bound_to = used_material_facts.get(key)
            if bound_to is not None and bound_to != norm(label):
                return False
        return True

    selected = next((item for item in candidates if usable(item)), None)
    return selected, candidates[:5]


# ---------- 清单驱动定位：待填写文件 → 占位符 → 候选内上下文消歧 ----------
#
# 事实表清单的第 2/3 列本身就是一张「字段 → 填进哪个 Word → 占位符长什么样」的路由表，
# 不需要靠字面相似度反猜字段。清单当前一个占位符可对多个字段（`[技术方案，待填写]`
# 在塔筒专题里对应 58 个），业务侧后续会把占位符逐步拆细；拆得越细，走确定性路径的
# 比例越高，下面的上下文消歧自然越少触发，无需改代码。

CONTEXT_MATCH_MIN = 0.62  # 候选内消歧：字段名被上下文覆盖的最低比例
CONTEXT_MATCH_MARGIN = 0.12  # 冠亚军差距不足说明分不开，宁空勿错


class SpecIndex:
    """清单定位索引。

    两张表：
    - `by_label`：字段名（含复核列别名）→ 字段。素材库里的占位符大多已经拆细成字段名
      本身（`[单台机组功率曲线保证率（%），待填写]`），这条是零歧义的确定性路径。
    - `by_placeholder`：清单第 3 列 → 候选字段。清单粒度粗于素材时靠它收敛候选，
      再交给上下文消歧。
    """

    def __init__(
        self,
        facts: list[dict[str, Any]],
        blank_keys: Iterable[str],
        known_placeholders: Iterable[str],
    ) -> None:
        self.blank_keys = {key for key in blank_keys if key}
        self.known_placeholders = {key for key in known_placeholders if key}
        self.by_placeholder: dict[str, list[dict[str, Any]]] = {}
        for fact in facts:
            for key in fact.get("specPlaceholders") or []:
                self.by_placeholder.setdefault(key, []).append(fact)
        # 只认带清单元数据的事实表字段。派生事实（招标方/项目名称等）与无清单信息的
        # 历史 manifest 字段不得进这张表：它们的 label 可能恰好等于泛占位符文字
        #（如「投标方案」），会劫持整份文档里该占位符的所有位置。
        spec_facts = [fact for fact in facts if fact.get("specPlaceholders") or fact.get("specTargets")]
        # 字段名先占位，复核列别名只填补空位——别名与其他字段的正式名冲突时不得覆盖
        self.by_label: dict[str, dict[str, Any]] = {}
        for name_key in ("label", "reviewLabel"):
            for fact in spec_facts:
                key = norm(fact.get(name_key))
                if key and key not in self.by_label:
                    self.by_label[key] = fact
        self.enabled = bool(self.known_placeholders or self.by_label)

    def knows(self, placeholder_label_text: str) -> bool:
        return norm(placeholder_label_text) in self.known_placeholders

    def by_field_name(self, placeholder_label_text: str) -> dict[str, Any] | None:
        return self.by_label.get(norm(placeholder_label_text))

    def candidates(self, placeholder_label_text: str) -> tuple[list[dict[str, Any]], bool]:
        entries = self.by_placeholder.get(norm(placeholder_label_text)) or []
        if not entries:
            return [], False
        # 软过滤：清单指定了当前文件的候选优先；一条都没指定时不做排除——清单文件名与
        # 实际待填写文件对不上（改名/换客户）时硬过滤会让整份文件静默一个字段都填不进。
        scoped = [fact for fact in entries if self.blank_keys & set(fact.get("specTargets") or [])]
        return (scoped or entries), bool(scoped)


def build_spec_index(manifest: dict[str, Any], facts: list[dict[str, Any]], blank_keys: Iterable[str]) -> SpecIndex:
    table = manifest.get("projectFactTable") if isinstance(manifest.get("projectFactTable"), dict) else {}
    known: list[str] = []
    for field in object_items(table.get("fields")):
        known.extend(key for key in map(placeholder_key, split_spec_cell(field.get("placeholder"))) if key)
    return SpecIndex(facts, blank_keys, known)


def context_fragments(context: str) -> list[str]:
    """上下文切片：表格上下文由 " / " 连接（表格坐标 / 行标签 / 列头 / 单元格文本），
    段落上下文是整句，都按同一套切分后归一化。"""
    return [fragment for fragment in (norm(part) for part in str(context or "").split(" / ")) if fragment]


def context_coverage(label: Any, fragments: list[str]) -> float:
    """字段名与上下文片段的双向包含度。

    两种形态都要认：
    - 表格：字段名「第1段（底）塔节底部直径（m）」被拆成行标签「第1段（底）」+
      列头「底部直径（m）」，片段是字段名的组成部分，按覆盖字符数累计；
    - 段落：字段名整体出现在句子里，反向包含，直接给满分。

    刻意不用 SequenceMatcher 模糊比：「底部直径」和「顶部直径」模糊分只差 0.09，
    分不开同一张表里的相邻列；完整包含是二值判据，同场景差距拉到 0.45。
    """
    target = norm(label)
    if not target or not fragments:
        return 0.0
    # 按字符位置累计而不是按片段长度求和：表头列名会同时出现在 header 段和近邻列段，
    # 重复计数会把所有含该列名的候选一起顶到 1.0，同列相邻字段就再也分不开。
    covered: set[int] = set()
    for fragment in fragments:
        if len(fragment) < 2:
            continue
        if target in fragment:
            return 1.0
        start = target.find(fragment)
        while start >= 0:
            covered.update(range(start, start + len(fragment)))
            start = target.find(fragment, start + 1)
    return len(covered) / len(target)


def disambiguate_by_context(
    candidates: list[dict[str, Any]],
    context: str,
) -> tuple[dict[str, Any] | None, list[dict[str, Any]]]:
    """候选内消歧：同一占位符下候选的占位符文字完全相同，只能靠文档上下文区分。"""
    fragments = context_fragments(context)
    scored: list[tuple[float, dict[str, Any]]] = []
    for fact in candidates:
        score = max(
            context_coverage(fact.get("label"), fragments),
            context_coverage(fact.get("reviewLabel"), fragments),
        )
        scored.append((score, fact))
    scored.sort(key=lambda item: -item[0])
    alternatives = [fact for _, fact in scored[:4]]
    if not scored or scored[0][0] < CONTEXT_MATCH_MIN:
        return None, alternatives
    if len(scored) > 1 and scored[0][0] - scored[1][0] < CONTEXT_MATCH_MARGIN:
        return None, alternatives
    picked = dict(scored[0][1])
    picked["score"] = round(scored[0][0], 3)
    return picked, [fact for _, fact in scored[1:4]]


def spec_locate(
    spec_index: SpecIndex | None,
    placeholder: dict[str, Any],
    context: str,
) -> tuple[dict[str, Any] | None, list[dict[str, Any]], str]:
    """按清单定位字段，返回 (选中事实, 备选, 状态)。

    状态：field_name=占位符文字就是字段名，直接命中；unique=清单第 3 列唯一对应一个
    字段；disambiguated=候选内经上下文选定；ambiguous=候选内分不开；
    no_value=清单有该占位符但字段没值；not_in_spec=清单里根本没有这个占位符；
    skipped=清单未下发，走旧模糊链路。
    """
    if spec_index is None or not spec_index.enabled:
        return None, [], "skipped"
    label = placeholder["label"]
    named = spec_index.by_field_name(label)
    if named is not None:
        picked = dict(named)
        picked["score"] = 0.99
        return picked, [], "field_name"
    if not spec_index.knows(label):
        return None, [], "not_in_spec"
    candidates, _ = spec_index.candidates(label)
    if not candidates:
        return None, [], "no_value"
    if len(candidates) == 1:
        picked = dict(candidates[0])
        picked["score"] = 0.99
        return picked, [], "unique"
    picked, alternatives = disambiguate_by_context(candidates, context)
    return picked, alternatives, "disambiguated" if picked else "ambiguous"


# 宁空勿错哨兵：上下文规则明确该槽位属于某个事实、但事实缺失时返回它，
# 强制走人工占位，禁止级联到后续规则或模糊匹配错填（金标反评：单机容量缺失
# 级联到总装机容量规则，10MW 被填成 600）。
MANUAL_SENTINEL: dict[str, Any] = {"__manual__": True}

NUMERIC_HINT_LABELS = (
    "数量", "台数", "小时", "容量", "坐标", "高度", "风速", "直径",
    "功率", "保证率", "可利用率", "利用率", "更换率", "电量", "年限", "期限",
)


def numeric_slot(text: str, placeholder: dict[str, Any], label: str) -> bool:
    """占位符槽位是否期望数字值（比较符前缀 / 单位后缀 / 数字型标签）。"""
    start = int(placeholder.get("start") or 0)
    end = int(placeholder.get("end") or 0)
    before = clean(text[max(0, start - 8):start])
    after = clean(text[end:end + 4])
    if any(ch in before[-3:] for ch in ("≥", "≤", ">", "<", "＞", "＜")):
        return True
    if re.match(r"^(台|套|米|m\b|mw|kw|h\b|小时|年|%|℃|mm)", after, re.IGNORECASE):
        return True
    return any(key in label for key in NUMERIC_HINT_LABELS)


def value_fits_numeric_slot(value: Any) -> bool:
    """数字槽位守卫：值必须含数字且是短值——长混合串（型号+参数拼接）判不合格。"""
    text = clean(value)
    return bool(re.search(r"\d", text)) and len(norm(text)) <= 24


def synthetic_fact(label: str, value: str, source: str, confidence: float = 0.9) -> dict[str, Any] | None:
    value_text = clean(value)
    if not value_text or placeholderish(value_text):
        return None
    return {
        "id": "SYN",
        "label": label,
        "value": value_text,
        "source": source,
        "sourcePath": "",
        "sourceKind": "derived",
        "sourcePriority": 95,
        "location": "context_rule",
        "factType": "context_rule",
        "confidence": confidence,
        "score": confidence,
    }


def value_fact(facts: list[dict[str, Any]], label: str, *aliases: str, confidence: float = 0.9) -> dict[str, Any] | None:
    value = first_fact_value(facts, label, *aliases)
    return synthetic_fact(label, value, "derived facts", confidence=confidence)


def context_without_placeholders(context: str) -> str:
    return clean(PLACEHOLDER_RE.sub(" ", context or ""))


def value_with_unit(value: str, unit: str) -> str:
    text = clean(value)
    if not text:
        return ""
    if re.search(r"[a-zA-Z%/³3]|米|台|小时|千瓦", text):
        return text
    return f"{text}{unit}"


def wind_resource_report_fact(facts: list[dict[str, Any]]) -> dict[str, Any] | None:
    hub = first_fact_value(facts, "轮毂高度", "轮毂中心高度")
    speed = first_fact_value(facts, "年平均风速", "测风塔210315年平均风速")
    density = first_fact_value(facts, "空气密度", "场址空气密度")
    turbulence = first_fact_value(facts, "湍流强度", "特征湍流强度")
    shear = first_fact_value(facts, "风剪切")
    extreme = first_fact_value(facts, "极端风速", "50年一遇极端风速")
    safety = first_fact_value(facts, "安全等级", "机组安全设计等级")
    parts = []
    if hub:
        parts.append(f"轮毂高度{value_with_unit(hub, 'm')}")
    if speed:
        parts.append(f"年平均风速{value_with_unit(speed, 'm/s')}")
    if density:
        parts.append(f"空气密度{value_with_unit(density, 'kg/m3')}")
    if turbulence:
        parts.append(f"湍流强度{turbulence}")
    if shear:
        parts.append(f"风剪切{shear}")
    if extreme:
        parts.append(f"50年一遇极端风速{value_with_unit(extreme, 'm/s')}")
    if safety:
        parts.append(f"安全等级{safety}")
    if len(parts) < 2:
        return None
    return synthetic_fact(
        "风资源报告",
        f"本项目风资源参数采用项目风资源评估及确认事实表：{'，'.join(parts)}。",
        "projectFactTable wind resource facts",
        confidence=0.9,
    )


def context_fact(label: str, context: str, placeholder: dict[str, Any], occurrence: int, facts: list[dict[str, Any]]) -> dict[str, Any] | None:
    label_norm = norm(label)
    ctx = norm(context_without_placeholders(context))
    raw_text = str(placeholder.get("rawText") or context)
    suffix = norm(raw_text[placeholder.get("end", 0) : placeholder.get("end", 0) + 20]) if isinstance(placeholder.get("end"), int) else ""
    if "风资源报告" in label or ("风资源" in label and "报告" in label):
        return wind_resource_report_fact(facts)
    if "投标机型" in label:
        return value_fact(facts, "投标机型", confidence=0.98)
    if "项目名称" in label:
        return value_fact(facts, "项目名称", confidence=0.95)
    if "招标方" in label or "招标人" in label:
        return value_fact(facts, "招标方", "招标人", confidence=0.92)
    if "日期" in label:
        return synthetic_fact("日期", datetime.now(UTC).strftime("%Y年%m月%d日"), "system date", confidence=0.62)
    if "基地名称" in label:
        if "叶片" in ctx:
            return value_fact(facts, "叶片供货制造基地", confidence=0.86)
        return value_fact(facts, "主机供货制造基地", confidence=0.86)
    if "基地介绍" in label:
        if "叶片" in ctx:
            return value_fact(facts, "叶片供货制造基地介绍", "叶片供货制造基地", confidence=0.78)
        return value_fact(facts, "主机供货制造基地介绍", "主机供货制造基地", confidence=0.78)
    if label_norm in {"投标方案", "项目方案", "技术方案", "招标要求", "投标响应"} or any(key in label for key in ("性能指标承诺", "风资源数据", "发电小时数承诺", "章节索引")):
        # 章节索引列最先短路：它是目录引用不是数据值，绝不能被下面的
        # 机型/容量等上下文规则错填（金标反评：18 行索引列被机型名覆盖）。
        if "章节索引" in label:
            return synthetic_fact("章节索引", "详见对应技术方案章节及技术附表", "directory context", confidence=0.65)
        if "单叶片吊装" in ctx:
            return value_fact(facts, "单叶片吊装及辅助工装工具套数", confidence=0.9) or MANUAL_SENTINEL
        if "运行维护工具" in ctx:
            return value_fact(facts, "运行维护工具套数", confidence=0.9) or MANUAL_SENTINEL
        if "超细干粉灭火装置" in ctx and ("主机舱" in ctx or "副机舱" in ctx):
            return synthetic_fact("自动消防配置数量", "3", "standard fire protection configuration", confidence=0.78)
        if "机舱烟感温感组合" in ctx:
            return synthetic_fact("自动消防配置数量", "2", "standard fire protection configuration", confidence=0.78)
        if any(item in ctx for item in ("声光报警器", "消防控制器", "气溶胶灭火装置", "温感探测器", "超细干粉灭火装置")):
            return synthetic_fact("自动消防配置数量", "1", "standard fire protection configuration", confidence=0.76)
        if "升降机" in ctx and "导向" in ctx:
            return value_fact(facts, "升降机导向方式", confidence=0.86) or MANUAL_SENTINEL
        if "主机" in ctx and "基地" in ctx:
            return value_fact(facts, "主机供货制造基地", confidence=0.86) or MANUAL_SENTINEL
        if "叶片" in ctx and "基地" in ctx:
            return value_fact(facts, "叶片供货制造基地", confidence=0.86) or MANUAL_SENTINEL
        if "项目名称" in ctx:
            return value_fact(facts, "项目名称", confidence=0.95) or MANUAL_SENTINEL
        if "承诺方式" in ctx:
            if "考核" in ctx:
                return synthetic_fact("承诺方式", "承诺考核值（77%折减）", "tender/performance facts", confidence=0.82)
            return synthetic_fact("承诺方式", "承诺保证值（75%折减）", "tender/performance facts", confidence=0.82)
        if "轮毂高度" in ctx or "轮毂中心高度" in ctx:
            return value_fact(facts, "轮毂高度", confidence=0.9) or MANUAL_SENTINEL
        if "净发电量" in ctx and "考核" in ctx:
            return value_fact(facts, "考核发电量", "考核净发电量", confidence=0.9) or MANUAL_SENTINEL
        if "净发电量" in ctx:
            return value_fact(facts, "保证发电量", "保证净发电量", confidence=0.9) or MANUAL_SENTINEL
        if "有效小时" in ctx and "考核" in ctx:
            return value_fact(facts, "考核有效小时数", confidence=0.9) or MANUAL_SENTINEL
        if "有效小时" in ctx:
            return value_fact(facts, "保证有效小时数", confidence=0.9) or MANUAL_SENTINEL
        if "功率曲线" in ctx or "功率曲线保证" in label:
            return value_fact(facts, "功率曲线保证率", confidence=0.88) or MANUAL_SENTINEL
        if ("单台" in ctx or "每台" in ctx) and "可利用率" in ctx:
            return value_fact(facts, "单台可利用率", confidence=0.88) or MANUAL_SENTINEL
        if "设备利用率" in ctx:
            return value_fact(facts, "单台可利用率", confidence=0.86) or MANUAL_SENTINEL
        if ("全场" in ctx or "全部" in ctx) and "可利用率" in ctx:
            return value_fact(facts, "全场可利用率", confidence=0.88) or MANUAL_SENTINEL
        if "主要部件更换率" in ctx:
            return value_fact(facts, "主要部件更换率", confidence=0.88) or MANUAL_SENTINEL
        if "风机数量" in ctx or "台数" in ctx or "全场配置" in ctx or suffix.startswith("台"):
            return value_fact(facts, "机组台数", "台数", confidence=0.9) or MANUAL_SENTINEL
        if "单机容量" in ctx:
            return value_fact(facts, "单机容量", confidence=0.9) or MANUAL_SENTINEL
        if "总装机容量" in ctx or "容量mw" in ctx or "项目总装机容量" in ctx:
            return value_fact(facts, "总装机容量", "容量", confidence=0.9) or MANUAL_SENTINEL
        if "容量" in ctx and "mw" in ctx:
            return value_fact(facts, "总装机容量", "容量", confidence=0.9) or MANUAL_SENTINEL
        if "方案" in ctx and "考核" in ctx:
            return value_fact(facts, "方案", confidence=0.86) or MANUAL_SENTINEL
        if "方案" in ctx:
            return value_fact(facts, "投标方案", "方案", confidence=0.86) or MANUAL_SENTINEL
        if "机型" in ctx:
            return value_fact(facts, "投标方案", "投标机型", confidence=0.95) or MANUAL_SENTINEL
        if "叶轮直径" in ctx:
            return value_fact(facts, "叶轮直径", confidence=0.9) or MANUAL_SENTINEL
        if "年平均风速" in ctx:
            return value_fact(facts, "测风塔210315年平均风速", confidence=0.9) or MANUAL_SENTINEL
        if "测风塔" in ctx or "风资源数据" in label:
            tower = first_fact_value(facts, "测风塔210315坐标")
            speed = first_fact_value(facts, "测风塔210315年平均风速")
            if tower and speed:
                return synthetic_fact("风资源数据", f"210315#（{tower}）125m高度处年平均风速{speed}m/s", "wind resource facts", confidence=0.86)
    if label_norm == "待填写内容":
        if "t1r" in ctx and "共计" in ctx:
            return value_fact(facts, "试运行业绩台数" if suffix.startswith("台") else "试运行业绩容量", confidence=0.92)
        if "t2r" in ctx and "共计" in ctx:
            if "发票" in ctx and (suffix.startswith("台") or "m台" in ctx):
                return value_fact(facts, "合同发票台数", confidence=0.9)
            return value_fact(facts, "合同业绩台数" if suffix.startswith("台") else "合同业绩容量", confidence=0.92)
        if "t3r" in ctx and "共计" in ctx:
            return value_fact(facts, "正在供货和新承接项目台数" if suffix.startswith("台") else "正在供货和新承接项目容量", confidence=0.92)
        if "试运行业绩" in ctx and (suffix.startswith("台") or occurrence % 2 == 0):
            return value_fact(facts, "试运行业绩台数", confidence=0.92)
        if "试运行业绩" in ctx:
            return value_fact(facts, "试运行业绩容量", confidence=0.92)
        if "合同业绩" in ctx and "发票" in ctx:
            if suffix.startswith("台") or "m台" in ctx:
                return value_fact(facts, "合同发票台数", confidence=0.9)
        if "合同业绩" in ctx and (suffix.startswith("台") or occurrence % 2 == 0):
            return value_fact(facts, "合同业绩台数", confidence=0.92)
        if "合同业绩" in ctx:
            return value_fact(facts, "合同业绩容量", confidence=0.92)
        if "正在供货" in ctx and (suffix.startswith("台") or occurrence % 2 == 0):
            return value_fact(facts, "正在供货和新承接项目台数", confidence=0.92)
        if "正在供货" in ctx:
            return value_fact(facts, "正在供货和新承接项目容量", confidence=0.92)
    return None


def set_paragraph_text(paragraph: Any, text: str, *, highlight: bool = False) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def set_cell_shading(cell: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), MANUAL_FILL)


def expected_value_for_key_row(row_label: str, facts: list[dict[str, Any]]) -> str:
    row_norm = norm(row_label)
    if not row_norm:
        return ""
    if "项目名称" in row_label:
        return first_fact_value(facts, "项目名称")
    if "承诺方式" in row_label:
        return "承诺考核值（77%折减）" if "考核" in row_label else "承诺保证值（75%折减）"
    if row_norm == "方案":
        return first_fact_value(facts, "投标方案", "方案")
    if row_norm == "机型":
        return first_fact_value(facts, "投标方案", "投标机型")
    if "轮毂高度" in row_label or "轮毂中心高度" in row_label:
        return first_fact_value(facts, "轮毂高度")
    if "净发电量" in row_label:
        if "考核" in row_label:
            return first_fact_value(facts, "考核发电量", "考核净发电量")
        return first_fact_value(facts, "保证发电量", "保证净发电量")
    if "有效小时" in row_label:
        if "考核" in row_label:
            return first_fact_value(facts, "考核有效小时数")
        return first_fact_value(facts, "保证有效小时数")
    if "功率曲线" in row_label:
        return first_fact_value(facts, "功率曲线保证率")
    if ("单台" in row_label or "每台" in row_label) and "可利用率" in row_label:
        return first_fact_value(facts, "单台可利用率")
    if ("全场" in row_label or "全部" in row_label) and "可利用率" in row_label:
        return first_fact_value(facts, "全场可利用率")
    if "风机数量" in row_label or row_norm == "台数" or "机组台数" in row_label:
        return first_fact_value(facts, "机组台数", "台数")
    if "单机容量" in row_label:
        return first_fact_value(facts, "单机容量")
    if "容量" in row_label and ("mw" in row_norm or "总装机" in row_label):
        return first_fact_value(facts, "总装机容量", "容量")
    return ""


def numeric_tokens(text: str) -> list[str]:
    tokens: list[str] = []
    for token in re.findall(r"\d+(?:\.\d+)?%?", clean(text)):
        tokens.append(token.rstrip("%"))
    return tokens


def validated_value_matches(actual: str, expected: str) -> bool:
    actual_text = clean(actual)
    expected_text = clean(expected)
    if not actual_text or placeholderish(actual_text) or not expected_text:
        return False
    expected_numbers = numeric_tokens(expected_text)
    if expected_numbers:
        actual_numbers = numeric_tokens(actual_text)
        return all(token in actual_numbers for token in expected_numbers)
    actual_norm = norm(actual_text)
    expected_norm = norm(expected_text)
    return bool(actual_norm and expected_norm and (actual_norm == expected_norm or expected_norm in actual_norm or actual_norm in expected_norm))


def validate_key_data_tables(output_file: Path, facts: list[dict[str, Any]], filled_locations: set[str]) -> dict[str, Any]:
    doc = Document(str(output_file))
    checks: list[dict[str, Any]] = []
    for table_idx, table in enumerate(doc.tables, start=1):
        for row_idx, row in enumerate(table.rows, start=1):
            cells = [clean(cell.text) for cell in row.cells]
            if len(cells) < 2:
                continue
            expected = expected_value_for_key_row(cells[0], facts)
            if not expected:
                continue
            for col_idx, value in enumerate(cells[1:], start=2):
                location = f"T{table_idx}/R{row_idx}/C{col_idx}"
                if location not in filled_locations:
                    continue
                if not value:
                    continue
                passed = validated_value_matches(value, expected)
                checks.append(
                    {
                        "location": location,
                        "rowLabel": cells[0],
                        "actual": value,
                        "expected": expected,
                        "passed": passed,
                    }
                )
    failed = [item for item in checks if not item["passed"]]
    passed_count = len(checks) - len(failed)
    rate = passed_count / len(checks) if checks else None
    return {
        "semanticCheckCount": len(checks),
        "semanticPassedCount": passed_count,
        "semanticFailedCount": len(failed),
        "semanticValidationRate": round(rate, 4) if rate is not None else None,
        "semanticFailures": failed[:20],
    }


def replace_text(
    text: str,
    facts: list[dict[str, Any]],
    context: str,
    used_material_facts: dict[tuple[str, str], str] | None = None,
    spec_index: SpecIndex | None = None,
) -> tuple[str, list[dict[str, Any]], list[str], bool]:
    decisions: list[dict[str, Any]] = []
    unfilled: list[str] = []
    highlighted = False
    placeholders = find_placeholders(text)
    if not placeholders:
        return text, decisions, unfilled, highlighted
    parts: list[str] = []
    last = 0
    for occurrence, placeholder in enumerate(placeholders):
        parts.append(text[last : placeholder["start"]])
        placeholder["rawText"] = text
        selected, alternatives, spec_status = spec_locate(spec_index, placeholder, context)
        if spec_status == "skipped":
            # 清单未下发（历史 manifest）：退回上下文规则 + 全库模糊匹配的旧链路
            selected = context_fact(placeholder["label"], context, placeholder, occurrence, facts)
            alternatives = []
            if selected is MANUAL_SENTINEL:
                selected = None  # 宁空勿错：规则已认领该槽位但事实缺失，禁止模糊兜底
            elif selected is None:
                selected, alternatives = choose_fact(placeholder["label"], context, facts, used_material_facts)
            if selected and numeric_slot(text, placeholder, placeholder["label"]) and not value_fits_numeric_slot(selected.get("value")):
                # 类型守卫：数字槽位拒绝非数字/超长混合值（金标反评：QT350-22AL 参数串
                # 覆盖 97%/2836h 等承诺数字）。类型守卫只保护模糊链路；清单定位是人工
                # 指定的字段，取值不合预期属于清单或事实表的问题，应如实填入并暴露。
                alternatives = ([selected] + alternatives)[:4]
                selected = None
        if selected:
            value = clean(selected["value"])
            if used_material_facts is not None and str(selected.get("sourceKind") or "") in {"docx", "xlsx"}:
                used_material_facts.setdefault((norm(selected.get("label")), norm(value)), norm(placeholder["label"]))
            decisions.append(
                {
                    "placeholder": placeholder["full"],
                    "label": placeholder["label"],
                    "action": "fill",
                    "value": value,
                    "confidence": selected["score"],
                    "evidence": selected,
                    "alternatives": alternatives[1:4] if spec_status == "skipped" else alternatives[:3],
                    "specStatus": spec_status,
                }
            )
        else:
            value = f"[待人工补充：{placeholder['label']}]"
            highlighted = True
            unfilled.append(placeholder["label"])
            decisions.append(
                {
                    "placeholder": placeholder["full"],
                    "label": placeholder["label"],
                    "action": "manual",
                    "value": value,
                    "confidence": 0,
                    "evidence": None,
                    "alternatives": alternatives[:4],
                    "specStatus": spec_status,
                    "candidateLabels": [clean(item.get("label")) for item in alternatives[:6]],
                }
            )
        parts.append(value)
        last = placeholder["end"]
    parts.append(text[last:])
    result = "".join(parts)
    return result, decisions, unfilled, highlighted


def fill_docx(
    source_path: Path,
    output_file: Path,
    facts: list[dict[str, Any]],
    spec_index: SpecIndex | None = None,
) -> dict[str, Any]:
    doc = Document(str(source_path))
    decisions: list[dict[str, Any]] = []
    unfilled: list[str] = []
    # 素材证据单次使用登记（文档级）：(事实标签, 值) -> 首次绑定的占位符标签
    used_material_facts: dict[tuple[str, str], str] = {}

    for idx, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text
        if not find_placeholders(text):
            continue
        context = clean(text)
        replaced, local_decisions, local_unfilled, highlight = replace_text(
            text, facts, context, used_material_facts, spec_index
        )
        set_paragraph_text(paragraph, replaced, highlight=highlight)
        for decision in local_decisions:
            decision["location"] = f"P{idx}"
        decisions.extend(local_decisions)
        unfilled.extend(local_unfilled)

    for table_idx, table in enumerate(doc.tables, start=1):
        seen_cells: set[Any] = set()
        for row_idx, row in enumerate(table.rows, start=1):
            row_context = " / ".join(clean(cell.text) for cell in row.cells if clean(cell.text))
            for col_idx, cell in enumerate(row.cells, start=1):
                cell_key = cell._tc
                if cell_key in seen_cells:
                    continue
                seen_cells.add(cell_key)
                text = cell.text
                if not find_placeholders(text):
                    continue
                column_context = " / ".join(
                    clean(table.rows[prev_idx].cells[col_idx - 1].text)
                    for prev_idx in range(max(0, row_idx - 4), row_idx - 1)
                    if col_idx - 1 < len(table.rows[prev_idx].cells) and clean(table.rows[prev_idx].cells[col_idx - 1].text)
                )
                # 表头始终单列一段：列头只在首行，近邻窗口越过 4 行后就丢了，
                # 分段参数表从第 5 行起会拿不到「底部直径」这类列名而无法消歧。
                header_cells = table.rows[0].cells if table.rows else []
                header_context = clean(header_cells[col_idx - 1].text) if col_idx - 1 < len(header_cells) else ""
                table_marker = f"T{table_idx}R{row_idx}C{col_idx}"
                context = " / ".join(
                    part for part in (table_marker, row_context, header_context, column_context, clean(text)) if part
                )
                replaced, local_decisions, local_unfilled, highlight = replace_text(
                    text, facts, context, used_material_facts, spec_index
                )
                cell.text = ""
                set_paragraph_text(cell.paragraphs[0], replaced, highlight=highlight)
                if highlight:
                    set_cell_shading(cell)
                for decision in local_decisions:
                    decision["location"] = f"T{table_idx}/R{row_idx}/C{col_idx}"
                decisions.extend(local_decisions)
                unfilled.extend(local_unfilled)

    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_file))
    return {
        "decisions": decisions,
        "unfilled": unfilled,
    }


def write_reports(output_file: Path, result: dict[str, Any]) -> tuple[Path, Path]:
    json_path = output_file.with_suffix(".fill_report.json")
    md_path = output_file.with_suffix(".fill_report.md")
    json_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    report = result["fillReport"]
    lines = [
        f"# {report.get('title') or '待填写 Word 填写报告'}",
        "",
        f"- 输出文件：`{result['outputFile']}`",
        f"- 占位符：{report['placeholderCount']}",
        f"- 已填写：{report['filledPlaceholderCount']}",
        f"- 待人工：{report['unfilledPlaceholderCount']}",
        "",
        "## 参考来源",
    ]
    for source in report.get("referenceSources") or []:
        lines.append(f"- {source['name']}（{source['route']}）")
    if report.get("specDriven"):
        lines.extend(
            [
                "",
                "## 清单定位",
                "",
                f"- 占位符即字段名，直接命中：{report.get('specFieldNameHitCount', 0)}",
                f"- 清单占位符列唯一命中：{report.get('specLocatedUniqueCount', 0)}",
                f"- 候选内上下文消歧：{report.get('specDisambiguatedCount', 0)}",
                f"- 候选内分不开（待拆细占位符）：{report.get('specAmbiguousCount', 0)}",
                f"- 清单外占位符（清单漏字段）：{report.get('specNotInSpecCount', 0)}",
                f"- 复合占位符（一格多字段，需产品决策）：{report.get('specCompositeCount', 0)}",
                f"- 清单有占位符但事实表没值：{report.get('specNoValueCount', 0)}",
            ]
        )
        diagnostics = (
            ("待拆细的歧义占位符", "ambiguousPlaceholders", ("location", "placeholder", "candidateCount")),
            ("复合占位符（一格多字段）", "compositePlaceholders", ("location", "placeholder", "fields")),
            ("清单外占位符", "placeholdersNotInSpec", ("location", "placeholder", "label")),
            ("事实表缺值字段", "fieldsWithoutValue", ("location", "placeholder", "label")),
            ("清单指定本文件但文档未出现的字段", "fieldsNotFoundInDoc", ("label", "placeholder")),
        )
        for heading, key, columns in diagnostics:
            rows = report.get(key) or []
            if not rows:
                continue
            lines.extend(
                [
                    "",
                    f"### {heading}（{len(rows)}）",
                    "",
                    "| " + " | ".join(columns) + " |",
                    "|" + "---|" * len(columns),
                ]
            )
            for row in rows:
                lines.append("| " + " | ".join(str(row.get(column, "")) for column in columns) + " |")
    lines.extend(["", "## 占位符明细", "", "| 位置 | 占位符 | 动作 | 值 | 置信度 |", "|---|---|---|---|---:|"])
    for item in result.get("filledFieldDetails") or []:
        action = "填写" if item["action"] == "fill" else "待人工"
        lines.append(f"| {item['location']} | {item['placeholder']} | {action} | {item['value']} | {item['confidence']} |")
    md_path.write_text("\n".join(lines), encoding="utf-8")
    return json_path, md_path


def composite_field_names(spec_index: SpecIndex | None, placeholder_label_text: str) -> list[str]:
    """复合占位符识别：`[投标机型、台数，待填写]` 这类一格要填多个字段值的写法。

    这类占位符不自动填：一格填几个值、用什么分隔是产品决策，先如实标黄并列入诊断。
    """
    if spec_index is None:
        return []
    parts = split_field_enumeration(placeholder_label_text)
    if len(parts) < 2:
        return []
    hits = [clean(spec_index.by_field_name(part).get("label")) for part in parts if spec_index.by_field_name(part)]
    return hits


def build_spec_diagnostics(
    manifest: dict[str, Any],
    blank_keys: set[str],
    decisions: list[dict[str, Any]],
    spec_index: SpecIndex | None = None,
) -> dict[str, Any]:
    """清单定位诊断：把没填上的原因拆成可执行的三类，供业务侧按量排优先级拉齐清单。

    - ambiguousPlaceholders：占位符在本文件对应多个字段且上下文分不开 → 该拆细占位符
    - placeholdersNotInSpec：文档里有、清单第 3 列没有 → 清单漏字段
    - fieldsWithoutValue：清单有该占位符但事实表没值 → 该补事实表
    - fieldsNotFoundInDoc：清单说要填进本文件、文档里却没有该占位符 → 文件填错或占位符改过
    """
    seen_keys = {norm(item.get("label")) for item in decisions}
    ambiguous: list[dict[str, Any]] = []
    not_in_spec: list[dict[str, Any]] = []
    composite: list[dict[str, Any]] = []
    no_value: list[dict[str, Any]] = []
    counts: dict[str, int] = {}
    for item in decisions:
        status = str(item.get("specStatus") or "")
        counts[status] = counts.get(status, 0) + 1
        entry = {
            "location": str(item.get("location") or ""),
            "placeholder": str(item.get("placeholder") or ""),
            "label": str(item.get("label") or ""),
        }
        if status == "ambiguous":
            candidates = item.get("candidateLabels") or []
            ambiguous.append({**entry, "candidateCount": len(candidates), "candidates": candidates})
        elif status == "not_in_spec":
            fields = composite_field_names(spec_index, entry["label"])
            if fields:
                composite.append({**entry, "fields": fields})
            else:
                not_in_spec.append(entry)
        elif status == "no_value":
            no_value.append(entry)

    table = manifest.get("projectFactTable") if isinstance(manifest.get("projectFactTable"), dict) else {}
    not_found: list[dict[str, Any]] = []
    for field in object_items(table.get("fields")):
        targets = {key for key in map(file_key, split_spec_cell(field.get("targetFile"))) if key}
        if not targets or not (blank_keys & targets):
            continue
        keys = {key for key in map(placeholder_key, split_spec_cell(field.get("placeholder"))) if key}
        if keys and not (keys & seen_keys):
            not_found.append({"label": clean(field.get("label")), "placeholder": clean(field.get("placeholder"))})

    return {
        "specFieldNameHitCount": counts.get("field_name", 0),
        "specLocatedUniqueCount": counts.get("unique", 0),
        "specDisambiguatedCount": counts.get("disambiguated", 0),
        "specAmbiguousCount": counts.get("ambiguous", 0),
        "specNotInSpecCount": counts.get("not_in_spec", 0),
        "specCompositeCount": len(composite),
        "specNoValueCount": counts.get("no_value", 0),
        "specSkippedCount": counts.get("skipped", 0),
        "ambiguousPlaceholders": ambiguous[:50],
        "compositePlaceholders": composite[:50],
        "placeholdersNotInSpec": not_in_spec[:50],
        "fieldsWithoutValue": no_value[:50],
        "fieldsNotFoundInDoc": not_found[:50],
    }


def run_from_manifest(manifest_path: Path) -> dict[str, Any]:
    manifest = load_manifest(manifest_path)
    source_path = template_path(manifest, manifest_path)
    blank_source = manifest.get("blankSource") if isinstance(manifest.get("blankSource"), dict) else {}
    title = clean(manifest.get("title") or blank_source.get("title") or source_path.stem)
    output_file = output_path(manifest, manifest_path, source_path)
    sources = select_sources(manifest, manifest_path)
    facts = collect_facts(manifest, sources)
    blank_keys = {file_key(source_path.name), file_key(blank_source.get("title")), file_key(blank_source.get("fileName"))}
    spec_index = build_spec_index(manifest, facts, blank_keys)
    filled = fill_docx(source_path, output_file, facts, spec_index)
    decisions = filled["decisions"]
    filled_locations = {str(item.get("location")) for item in decisions if item.get("action") == "fill" and item.get("location")}
    semantic_validation = validate_key_data_tables(output_file, facts, filled_locations)
    reference_sources = [
        {
            "id": source.material_id,
            "name": source.name,
            "path": str(source.path),
            "kind": source.kind,
            "route": source.route,
            "priority": source.priority,
        }
        for source in sources
    ]
    evidence_refs = []
    for decision in decisions:
        evidence = decision.get("evidence")
        if not evidence:
            continue
        evidence_refs.append(
            {
                "type": "selected_fact",
                "label": decision["label"],
                "source": evidence.get("source"),
                "sourcePath": evidence.get("sourcePath"),
                "location": evidence.get("location"),
                "factType": evidence.get("factType"),
                "confidence": evidence.get("confidence"),
            }
        )
    evidence_refs.append({"type": "blank_source", "path": str(source_path), "title": source_path.name})
    result: dict[str, Any] = {
        "schema_version": SCHEMA_VERSION,
        "outputFile": str(output_file),
        "unfilledFields": list(dict.fromkeys(filled["unfilled"])),
        "evidenceRefs": evidence_refs,
        "filledFieldDetails": decisions,
        "unfilledFieldDetails": [item for item in decisions if item["action"] == "manual"],
        "fillReport": {
            "title": title,
            "placeholderCount": len(decisions),
            "filledPlaceholderCount": sum(1 for item in decisions if item["action"] == "fill"),
            "unfilledPlaceholderCount": sum(1 for item in decisions if item["action"] == "manual"),
            "referenceMaterialCount": len(sources),
            "referenceSources": reference_sources,
            "blankDocxPath": str(source_path),
            "preservedOriginalStructure": True,
            "manualMarker": "[待人工补充：字段名]",
            "manualHighlight": MANUAL_FILL,
            "specDriven": spec_index.enabled,
            **semantic_validation,
            **build_spec_diagnostics(manifest, {key for key in blank_keys if key}, decisions, spec_index),
        },
        "filledAt": now_iso(),
    }
    json_path, md_path = write_reports(output_file, result)
    result["fillReport"]["reportJsonPath"] = str(json_path)
    result["fillReport"]["reportMarkdownPath"] = str(md_path)
    json_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    return result


def compact_summary(result: dict[str, Any]) -> dict[str, Any]:
    return {
        "schema_version": result["schema_version"],
        "outputFile": result["outputFile"],
        "unfilledFields": result["unfilledFields"],
        "evidenceRefs": result["evidenceRefs"],
        "fillReport": result["fillReport"],
        "filledAt": result["filledAt"],
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--manifest", required=True)
    parser.add_argument("--response", choices=("summary", "full"), default="summary")
    args = parser.parse_args()
    result = run_from_manifest(Path(args.manifest).expanduser())
    print(json.dumps(compact_summary(result) if args.response == "summary" else result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
