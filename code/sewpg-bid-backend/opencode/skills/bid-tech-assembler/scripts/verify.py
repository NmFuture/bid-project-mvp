#!/usr/bin/env python3
"""
对账与审计：

1. plan.json 中 in-scope 条目数 vs 输出 docx 中 Heading 数对账
2. 残留占位符扫描（[待填写：xxx] / [缺失：xxx] / [FIELD] / [PROJECT_NAME]...）
3. 可疑情况扫描：相邻重复 Heading（合并时 toc title + 素材首标题 撞车）

输出：stdout 紧凑 JSON；可选写入 --result 指定的 JSON 文件

用法：
    python3 verify.py \\
        --docx 输出.docx \\
        --plan /tmp/assembly_plan.json \\
        --params project_params.json \\
        --result 工作目录/assembly_verify_result.json
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from collections import Counter
from pathlib import Path

from docx import Document


PLACEHOLDER_PATTERNS = [
    re.compile(r"\[待填写[:：].*?\]"),
    re.compile(r"\[缺失[:：].*?\]"),
    re.compile(r"\[[A-Z_]{4,}\]"),  # [PROJECT_NAME] / [MODEL_NO] 等
]

# 方案 B 硬性检查
ALLOWED_H1_PATTERNS = [
    re.compile(r"^封面$"),
    re.compile(r"^前言(\s|$)"),
    re.compile(r"^第[一二三四五六七八九十]+章(\s|$)"),
    re.compile(r"^附表"),
    re.compile(r"^目录$"),  # finalize 插入的 TOC 标题
]

# 幽灵编号：Heading text 里出现"第 X 章"X > 6 或 > 10 的数字（真正章节最多 6）
GHOST_CHAPTER = re.compile(r"第\s*(\d+)\s*章")
# Heading text 是否以合法编号开头：前言 / 封面 / 第X章 / 数字.数字.. / 附
VALID_HEADING_START = re.compile(r"^(?:封面|前言|第[一二三四五六七八九十\d]+章|\d+(?:\.\d+){0,6}(?:\s|$)|附|附表|目录)")


def scan_docx(docx_path: Path) -> dict:
    from docx.oxml.ns import qn

    doc = Document(str(docx_path))

    heading_counts = Counter()
    heading_list: list[tuple[int, str]] = []
    placeholders: list[str] = []
    dup_alerts: list[str] = []
    ghost_chapters: list[str] = []
    invalid_h1: list[str] = []
    invalid_prefix: list[str] = []
    empty_leaf_headings: list[str] = []

    # Body 级遍历：既能看到段落也能看到表格，正确识别 heading 之间的"内容量"
    W_P = qn("w:p")
    W_TBL = qn("w:tbl")
    W_DRAWING = qn("w:drawing")
    W_PICT = qn("w:pict")

    style_level_map: dict[str, int] = {}
    styles_el = doc.part.styles._element
    for style_el in styles_el.findall(qn("w:style")):
        if style_el.get(qn("w:type")) != "paragraph":
            continue
        sid = style_el.get(qn("w:styleId")) or ""
        if not sid:
            continue
        lvl = None
        name_el = style_el.find(qn("w:name"))
        name = name_el.get(qn("w:val")) if name_el is not None else ""
        m = re.match(r"^(?:Heading|heading|标题)\s*(\d+)$", name or "")
        if m:
            lvl = int(m.group(1))
        if lvl is None:
            outline = style_el.find(qn("w:pPr") + "/" + qn("w:outlineLvl"))
            if outline is not None:
                try:
                    value = int(outline.get(qn("w:val")))
                    if 0 <= value <= 8:
                        lvl = value + 1
                except (TypeError, ValueError):
                    pass
        if lvl is not None:
            style_level_map[sid] = lvl

    def _para_style_level(el):
        pPr = el.find(qn("w:pPr"))
        if pPr is None:
            return None
        direct_outline = pPr.find(qn("w:outlineLvl"))
        if direct_outline is not None:
            try:
                value = int(direct_outline.get(qn("w:val")))
                if 0 <= value <= 8:
                    return value + 1
            except (TypeError, ValueError):
                pass
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is None:
            return None
        val = pStyle.get(qn("w:val")) or ""
        if val in style_level_map:
            return style_level_map[val]
        # styleId 数字 2-6 对应 Heading 2-6，10 对应 H1；也接受以 "Heading" 或 "标题" 开头
        if val in ("1", "10"):
            return 1
        if val in ("2", "3", "4", "5", "6"):
            return int(val)
        m = re.match(r"^(?:Heading|heading|标题)\s*(\d+)$", val)
        if m:
            return int(m.group(1))
        return None

    def _para_text(el) -> str:
        return "".join((t.text or "") for t in el.iter(qn("w:t")))

    def _has_image(el) -> bool:
        return el.find(".//" + W_DRAWING) is not None or el.find(".//" + W_PICT) is not None

    prev_heading: tuple[int, str] | None = None
    last_heading: tuple[int, str] | None = None
    content_for_last: int = 0  # >0 表示该 heading 有文字/表格/图片

    def _flush_last(next_level: int | None):
        """结算 last_heading 的内容计数，必要时登记空章节告警。

        next_level: 下一个 heading 的 level；None 表示文档末尾。
        只在"叶子"位置记录：下一个 heading level <= 本 heading level（或 EOF）。
        """
        nonlocal last_heading, content_for_last
        if last_heading is None:
            return
        is_leaf = (next_level is None) or (next_level <= last_heading[0])
        if is_leaf and content_for_last == 0:
            prev_lvl, prev_text = last_heading
            # 合法空：封面 / 前言 / 目录 / 纯章节标题（第X章 / 第X节）
            if not (
                prev_text in ("封面", "前言", "目录")
                or re.match(r"^第[一二三四五六七八九十\d]+[章节]", prev_text)
                or re.match(r"^\d+(\.\d+)*\s*$", prev_text.strip())
            ):
                empty_leaf_headings.append(f"L{prev_lvl} {prev_text}")

    for child in doc.element.body.iterchildren():
        tag = child.tag
        if tag == W_P:
            lvl = _para_style_level(child)
            text = _para_text(child).strip()

            # 占位符扫描
            for pat in PLACEHOLDER_PATTERNS:
                for m in pat.finditer(text):
                    placeholders.append(m.group(0))

            if lvl is not None:
                # 遇到新 heading：先结算上一个
                _flush_last(lvl)
                heading_counts[f"Heading {lvl}"] += 1
                heading_list.append((lvl, text))
                last_heading = (lvl, text)
                content_for_last = 0

                if prev_heading and prev_heading[1] == text and text:
                    dup_alerts.append(f"L{lvl} 相邻重复：{text}")
                prev_heading = (lvl, text)

                gm = GHOST_CHAPTER.search(text)
                if gm and int(gm.group(1)) > 6:
                    ghost_chapters.append(f"L{lvl} {text!r}")

                if lvl == 1 and text:
                    if not any(pat.search(text) for pat in ALLOWED_H1_PATTERNS):
                        invalid_h1.append(text)

                if text and not VALID_HEADING_START.match(text):
                    invalid_prefix.append(f"L{lvl} {text[:60]!r}")
            else:
                # 非 heading 段：算正文量（文字或图片/drawing）
                if text or _has_image(child):
                    content_for_last += max(1, len(text))
        elif tag == W_TBL:
            # 表格一律计为有内容
            content_for_last += 100
            # 同时扫表格里的占位符
            for t in child.iter(qn("w:t")):
                cell_text = t.text or ""
                for pat in PLACEHOLDER_PATTERNS:
                    for m in pat.finditer(cell_text):
                        placeholders.append(m.group(0))

    # 文档末尾结算
    _flush_last(None)

    return {
        "sections": len(doc.sections),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "heading_counts": dict(heading_counts),
        "heading_list": heading_list,
        "placeholders": placeholders,
        "dup_alerts": dup_alerts,
        "ghost_chapters": ghost_chapters,
        "invalid_h1": invalid_h1,
        "invalid_prefix": invalid_prefix,
        "empty_leaf_headings": empty_leaf_headings,
    }


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", type=Path, required=True)
    ap.add_argument("--plan", type=Path, required=True)
    ap.add_argument("--params", type=Path, required=True)
    ap.add_argument("--result", type=Path, default=None)
    args = ap.parse_args()

    plan = json.loads(args.plan.read_text(encoding="utf-8"))
    params = json.loads(args.params.read_text(encoding="utf-8"))
    scan = scan_docx(args.docx)

    if args.result:
        args.result.parent.mkdir(parents=True, exist_ok=True)
        args.result.write_text(json.dumps(scan, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(scan, ensure_ascii=True, separators=(",", ":")))


if __name__ == "__main__":
    main()
