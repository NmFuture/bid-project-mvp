#!/usr/bin/env python3
"""
post-fix：扫描 merger 输出的 docx，把 Heading 级别与 text 编号深度不一致的段落修正。

典型问题（v2 遗留）：
- "5.7.1  偏航系统" style=Heading 1（应 Heading 3）
- "3.8.1  风资源评估..." style=Heading 1（应 Heading 3）
- "1  偏航系统" style=Heading 1（文本前缀不完整，也可能降级处理）

修正规则：
- Heading text 以 "X.Y[.Z...]" 开头 → 降级到 Heading {depth}
- 只"降级"（N→M 且 M≥N），不升级（避免误伤）
- Heading 1 text 只应为: "前言..."/"第X章..."/"附表..."/"目录"；其他降级

用法：
    python3 fix_invalid_headings.py <in.docx> <out.docx>
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document


_NUM_PREFIX = re.compile(r"^\s*(\d+(?:\.\d+)*)\s")


def _heading_level(style_name: str):
    if not style_name:
        return None
    m = re.match(r"^(?:Heading|heading)\s+(\d+)$", style_name.strip())
    if m:
        return int(m.group(1))
    m = re.match(r"^标题\s*(\d+)$", style_name.strip())
    if m:
        return int(m.group(1))
    return None


def _build_style_outline_map(doc) -> dict:
    """扫 doc.styles XML，返回 {styleId: outlineLvl(0-based)}。"""
    from docx.oxml.ns import qn
    result = {}
    styles_el = doc.part.styles._element
    for s in styles_el.findall(qn("w:style")):
        sid = s.get(qn("w:styleId"))
        if not sid:
            continue
        pPr = s.find(qn("w:pPr"))
        if pPr is None:
            continue
        ol = pPr.find(qn("w:outlineLvl"))
        if ol is None:
            continue
        try:
            result[sid] = int(ol.get(qn("w:val")))
        except (ValueError, TypeError):
            pass
    return result


def normalize_headings_by_outline_level(doc) -> int:
    """共性修法：识别 heading 不靠 style 名，靠 w:outlineLvl。
    把所有 effective outlineLvl ∈ [0,5] 的段落强制 style = Heading {lvl+1}。

    覆盖所有 WPS 自定义 heading style（章标题、附件标题1、专题标题、二级标题等）。
    """
    from docx.oxml.ns import qn
    style_outline_map = _build_style_outline_map(doc)

    count = 0
    for para in doc.paragraphs:
        pPr = para._p.find(qn("w:pPr"))
        if pPr is None:
            continue
        # 1) 段落直接 outlineLvl 优先
        direct_ol = pPr.find(qn("w:outlineLvl"))
        effective_lvl = None
        if direct_ol is not None:
            try:
                effective_lvl = int(direct_ol.get(qn("w:val")))
            except (ValueError, TypeError):
                pass
        # 2) 否则查段落 style 的 outlineLvl
        if effective_lvl is None:
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None:
                val = pStyle.get(qn("w:val")) or ""
                if val in style_outline_map:
                    effective_lvl = style_outline_map[val]
        if effective_lvl is None:
            continue
        if effective_lvl < 0 or effective_lvl > 5:
            continue

        target_style_name = f"Heading {effective_lvl + 1}"
        cur_name = para.style.name if para.style else ""
        if cur_name == target_style_name:
            continue
        try:
            para.style = doc.styles[target_style_name]
            count += 1
        except KeyError:
            pass
    return count


def _pure_title(text: str) -> str:
    """语义归一化：同义变体（"附表X"/"表X"、"xxx（如有）"/"xxx"等）收敛到同一 key。"""
    import unicodedata
    t = _NUM_PREFIX.sub("", text or "").strip()
    t = unicodedata.normalize("NFKC", t)
    # 前缀星号/修饰
    t = re.sub(r"^[*＊★☆※◆◎○●]+\s*", "", t)
    # 前缀"附"字（可选跟"表"）
    t = re.sub(r"^附(表)?(?=\S)", r"\1", t)
    # 末尾括号修饰（如"（如有）"/"（可选）"/"（北区）"）
    t = re.sub(r"[（(][^）)]{1,6}[）)]\s*$", "", t).strip()
    t = re.sub(r"[\s:：，,。.！!？?\-_/'\"“”‘’·（）()【】\[\]]+", "", t)
    return t


def _parent_prefix(text: str) -> str:
    """'4.1 xxx' → '4'；'5.8.1 yyy' → '5.8'；'第四章 xx' → 'ch4'；'前言 xx' → 'preface'。"""
    if not text:
        return ""
    text = text.strip()
    m = _NUM_PREFIX.match(text)
    if m:
        parts = m.group(1).split(".")
        if len(parts) <= 1:
            return ""
        return ".".join(parts[:-1])
    if text.startswith("第"):
        # '第四章' → ch4
        cm = re.match(r"^第([一二三四五六七八九十\d]+)章", text)
        if cm:
            return f"ch_{cm.group(1)}"
    if text.startswith("前言"):
        return "preface"
    return ""


def fix_invalid_h1(doc) -> int:
    """H1 不是合法前缀（前言/第X章/附表/目录/封面）→ 按数字深度降级。"""
    demoted = 0
    for para in doc.paragraphs:
        st = para.style.name if para.style else ""
        cur_lvl = _heading_level(st)
        if cur_lvl != 1:
            continue
        text = (para.text or "").strip()
        if not text:
            continue
        if text.startswith(("前言", "第", "附表", "目录", "封面")):
            continue
        m = _NUM_PREFIX.match(text)
        if m:
            depth = m.group(1).count(".") + 1
            target = max(2, min(6, depth))
            if target > cur_lvl:
                try:
                    para.style = doc.styles[f"Heading {target}"]
                    demoted += 1
                except KeyError:
                    pass
    return demoted


def dedupe_placeholder_headings(doc) -> int:
    """删除"占位 heading"（下一段是 [待填写]）若同父章节下已有相同 pure_title 的实际 heading。"""
    from docx.oxml.ns import qn

    # 建索引：{parent_prefix: {pure_title: first_seen_idx}}
    heading_paras = []
    for i, para in enumerate(doc.paragraphs):
        st = para.style.name if para.style else ""
        lvl = _heading_level(st)
        if lvl is None:
            continue
        heading_paras.append((i, para, lvl, (para.text or "").strip()))

    # 找每个 Heading 的"下一段"是否是 [待填写] 占位
    all_paras = list(doc.paragraphs)

    def _next_non_empty(idx):
        for j in range(idx + 1, min(idx + 3, len(all_paras))):
            t = (all_paras[j].text or "").strip()
            if t:
                return all_paras[j], j
        return None, None

    # 第一遍：收集非占位 Heading → 建 {parent: set(pure)}
    non_placeholder_titles_by_parent: dict[str, set[str]] = {}
    placeholder_targets: list[tuple[int, int, object, object]] = []  # (heading_idx, placeholder_idx, heading_para, placeholder_para)

    for i, para, lvl, text in heading_paras:
        next_para, next_idx = _next_non_empty(i)
        is_placeholder = False
        if next_para is not None:
            next_text = (next_para.text or "").strip()
            if next_text.startswith("[待填写") or next_text.startswith("[缺失"):
                is_placeholder = True

        pure = _pure_title(text)
        parent = _parent_prefix(text)
        if is_placeholder:
            placeholder_targets.append((i, next_idx, para, next_para))
        else:
            non_placeholder_titles_by_parent.setdefault(parent, set()).add(pure)

    # 第二遍：对每个 placeholder，若其 pure_title 已在同 parent 下有非占位 heading → 删除
    removed = 0
    for h_idx, ph_idx, h_para, ph_para in placeholder_targets:
        text = (h_para.text or "").strip()
        pure = _pure_title(text)
        parent = _parent_prefix(text)
        existing = non_placeholder_titles_by_parent.get(parent, set())
        if pure in existing:
            # 删除 heading + 占位段
            h_para._p.getparent().remove(h_para._p)
            if ph_para is not None:
                ph_para._p.getparent().remove(ph_para._p)
            removed += 1

    return removed


def renumber_headings_by_parent(doc) -> int:
    """同父章节下 Heading 重编号：按 doc 顺序，按 level 维护 counter。

    解决：
    1. 素材 A inject 发 4.1-4.N，素材 B inject 又从 4.1 开始的撞号
    2. appendix 素材 inject 时 parent="" 导致 text="1 xxx" 无效前缀
       → 用"最近的父上下文"作为 parent 续号
    """
    changed = 0

    # parent_counters：每个 parent_prefix 下最深一级已用的编号
    parent_counters: dict[str, int] = {}
    # recent_context_parent：最近见过的 depth>=2 heading 的 parent，作为 appendix fallback
    recent_context_parent = ""
    # recent_context_depth：最近 depth>=2 heading 的深度（用于推断 level）
    recent_context_depth = 0

    for para in doc.paragraphs:
        st = para.style.name if para.style else ""
        lvl = _heading_level(st)
        if lvl is None:
            continue
        text = (para.text or "").strip()
        if not text:
            continue
        # 章级/前言跳过
        if text.startswith(("前言", "第", "附表", "目录", "封面")):
            continue

        m = _NUM_PREFIX.match(text)
        if not m:
            continue
        prefix = m.group(1)
        depth = prefix.count(".") + 1
        parts = prefix.split(".")

        if depth >= 2:
            # 正常多级：按 parent 续号（仅对与 parent 紧挨着的层）
            parent = ".".join(parts[:-1])
            # 记录上下文供后续单级 heading 用
            recent_context_parent = parent
            recent_context_depth = depth - 1  # parent 深度

            # 仅对"恰好是 parent 下一层"的做重编号
            cnt = parent_counters.get(parent, 0) + 1
            parent_counters[parent] = cnt
            new_prefix = f"{parent}.{cnt}"
            new_depth = new_prefix.count(".") + 1
            new_text = f"{new_prefix}  {_NUM_PREFIX.sub('', text).strip()}"
            # 强制 style level 与编号深度一致
            target_lvl = max(1, min(6, new_depth))
            if target_lvl != lvl:
                try:
                    para.style = doc.styles[f"Heading {target_lvl}"]
                except KeyError:
                    pass
            if new_text != text:
                _rewrite_para(para, new_text)
                changed += 1
        else:
            # 单级编号 "N xxx"：多半是 appendix 素材 inject 时 parent="" 产生的
            # 用最近的 recent_context_parent 续号
            if recent_context_parent:
                parent = recent_context_parent
                cnt = parent_counters.get(parent, 0) + 1
                parent_counters[parent] = cnt
                new_text = f"{parent}.{cnt}  {_NUM_PREFIX.sub('', text).strip()}"
                # 同时降级 style
                target_lvl = max(2, min(6, len(parent.split(".")) + 1))
                if target_lvl != lvl:
                    try:
                        para.style = doc.styles[f"Heading {target_lvl}"]
                    except KeyError:
                        pass
                _rewrite_para(para, new_text)
                changed += 1
    return changed


def _rewrite_para(para, new_text: str) -> None:
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.text = new_text


def remove_same_level_duplicate_subtrees(doc) -> int:
    """删除同 level 下同 pure_title 的"后来者"heading + 其整个子树。

    子树定义：从该 heading 起，到下一个 level <= 当前 level 的 heading 之前为止
    （包含中间的正文段落、表格）。

    用于解决：素材内部冗余结构（如既有"表1 风资源..."又有"附表1 风资源..."L2，
    归一化后同义 → 后者被视为重复子树删除）。
    """
    from docx.oxml.ns import qn

    body = doc.element.body
    children = list(body)

    def _get_lvl(el):
        if el.tag != qn("w:p"):
            return None
        pPr = el.find(qn("w:pPr"))
        if pPr is None:
            return None
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is None:
            return None
        val = pStyle.get(qn("w:val")) or ""
        # 母版 styleId 映射：10→1, 2→2, ..., 9→9
        _MAP = {"10": 1, "2": 2, "3": 3, "4": 4, "5": 5, "6": 6, "7": 7, "8": 8, "9": 9}
        if val in _MAP:
            return _MAP[val]
        # 兜底：text 名
        m = re.match(r"^(?:Heading|heading)\s+(\d+)$", val)
        if m:
            return int(m.group(1))
        return None

    def _get_text(el):
        return "".join(t.text or "" for t in el.iter(qn("w:t")))

    seen_by_level: dict[int, set[str]] = {}
    to_remove: list = []

    i = 0
    while i < len(children):
        el = children[i]
        lvl = _get_lvl(el)
        if lvl is None or lvl < 2:  # 不去重 L1（章级）
            i += 1
            continue

        text = _get_text(el).strip()
        if not text:
            i += 1
            continue
        pure = _pure_title(text)
        if not pure:
            i += 1
            continue

        seen = seen_by_level.setdefault(lvl, set())
        if pure in seen:
            # 找 subtree 边界
            j = i + 1
            while j < len(children):
                lvl_next = _get_lvl(children[j])
                if lvl_next is not None and lvl_next <= lvl:
                    break
                j += 1
            # 标记 [i, j) 为删除
            to_remove.extend(children[i:j])
            i = j
        else:
            seen.add(pure)
            i += 1

    for el in to_remove:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
    return len(to_remove)


def realign_prefix_by_parent_stack(doc) -> int:
    """按文档顺序维护父 heading 栈，把每个 heading 的编号前缀重新对齐到栈顶。

    解决：inject 用 toc.chapter_no_flat 作 parent，但实际 doc 中编号被 renumber 改动后
    深层 heading 的 prefix 仍是旧的 toc 编号，导致挂错位置。

    算法：
        stack[l] = 当前 level l 最近的编号前缀
        parent_children_count[parent_prefix] = parent 下已出现的直接子节数
        for each heading p (按 doc 顺序):
            cur_lvl = heading level
            if cur_lvl == 1:
                new_prefix = 保留原 prefix（章级不改）
            else:
                parent = stack[cur_lvl - 1]
                if not parent: fallback 保留原 prefix
                else:
                    cnt = parent_children_count[parent] + 1
                    new_prefix = f"{parent}.{cnt}"
            stack[cur_lvl] = new_prefix
            for d in range(cur_lvl+1, 8): stack[d] = ""
            rewrite heading text with new_prefix
    """
    stack = ["" for _ in range(8)]  # 1..7
    parent_cnt: dict[str, int] = {}
    changed = 0

    for para in doc.paragraphs:
        st = para.style.name if para.style else ""
        lvl = _heading_level(st)
        if lvl is None:
            continue
        text = (para.text or "").strip()
        if not text:
            continue
        # 章级/前言跳过（L1 特殊）
        if text.startswith(("前言", "第", "附表", "目录", "封面")):
            # 把 stack[1] 置空，之后深层不依赖 L1 prefix
            stack = ["" for _ in range(8)]
            continue

        m = _NUM_PREFIX.match(text)
        if not m:
            continue
        old_prefix = m.group(1)

        if lvl < 1:
            continue
        if lvl == 1:
            new_prefix = old_prefix
        else:
            parent = stack[lvl - 1]
            if not parent:
                # 无父上下文；保留原 prefix（或降级）
                new_prefix = old_prefix
            else:
                cnt = parent_cnt.get(parent, 0) + 1
                parent_cnt[parent] = cnt
                new_prefix = f"{parent}.{cnt}"

        # 更新栈
        stack[lvl] = new_prefix
        for d in range(lvl + 1, 8):
            stack[d] = ""

        if new_prefix != old_prefix:
            new_text = f"{new_prefix}  {_NUM_PREFIX.sub('', text).strip()}"
            _rewrite_para(para, new_text)
            # 同时校准 style level 与 new_prefix 深度一致
            new_depth = new_prefix.count(".") + 1
            target_lvl = max(1, min(6, new_depth))
            if target_lvl != lvl:
                try:
                    para.style = doc.styles[f"Heading {target_lvl}"]
                except KeyError:
                    pass
            changed += 1
    return changed


def inject_prefix_to_unprefixed_headings(doc) -> int:
    """对 text 不带编号前缀的 Heading 段按父栈补全编号。

    场景：素材 heading 的 style 是 WPS 自定义（如"章标题"），preprocess 时
    没被 inject_prefix_to_headings 处理；经 normalize_headings_by_outline_level
    归并到 Heading N 后，text 仍是纯标题（如"范围"/"结论"）。
    本函数给它们加编号前缀。
    """
    stack = ["" for _ in range(8)]
    parent_cnt: dict[str, int] = {}
    count = 0

    for para in doc.paragraphs:
        st = para.style.name if para.style else ""
        lvl = _heading_level(st)
        if lvl is None:
            continue
        text = (para.text or "").strip()
        if not text:
            continue
        if text.startswith(("前言", "第", "附表", "目录", "封面")):
            stack = ["" for _ in range(8)]
            continue

        m = _NUM_PREFIX.match(text)
        if m:
            # 已有编号前缀 — 按现有算法更新栈
            stack[lvl] = m.group(1)
            for d in range(lvl + 1, 8):
                stack[d] = ""
            continue

        # 无前缀：按栈补
        if lvl == 1:
            continue  # L1 不补
        parent = stack[lvl - 1]
        if not parent:
            continue
        cnt = parent_cnt.get(parent, 0) + 1
        parent_cnt[parent] = cnt
        new_prefix = f"{parent}.{cnt}"
        new_text = f"{new_prefix}  {text}"
        _rewrite_para(para, new_text)
        stack[lvl] = new_prefix
        for d in range(lvl + 1, 8):
            stack[d] = ""
        count += 1
    return count


def fix(in_path: Path, out_path: Path) -> dict:
    doc = Document(str(in_path))
    stats = {
        "heading_by_outline": normalize_headings_by_outline_level(doc),
        "prefix_inj_for_unprefixed": inject_prefix_to_unprefixed_headings(doc),
        "demoted": fix_invalid_h1(doc),
        "dup_subtree_removed": remove_same_level_duplicate_subtrees(doc),
        "dedup_removed": dedupe_placeholder_headings(doc),
        "renumbered": renumber_headings_by_parent(doc),
        "realigned": realign_prefix_by_parent_stack(doc),
    }
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return stats


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input", type=Path)
    ap.add_argument("output", type=Path)
    args = ap.parse_args()

    if not args.input.exists():
        print(f"[ERR] not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    stats = fix(args.input, args.output)
    print(f"[OK] {args.input} → {args.output}")
    print(f"     {stats}")


if __name__ == "__main__":
    main()
