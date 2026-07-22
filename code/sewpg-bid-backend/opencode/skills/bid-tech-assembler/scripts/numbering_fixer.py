"""
Heading 编号处理器（方案 B）

两套能力：
1. strip_prefix / strip_heading_prefixes_in_doc — 剥离 Heading 文本里的章节号前缀
2. inject_prefix_to_headings — 按父章节号注入编号前缀（方案 B 核心）
3. strip_handwritten_numbering_in_body — 正文手写编号擦除（"7.10 xxx" 之类）

方案 B 下 Heading text 直接带章节号字符串，禁用 Word 多级列表。
"""

from __future__ import annotations

import re
from typing import Optional

# Heading 段落开头的编号前缀正则。按优先级从长到短匹配。
_PREFIX_PATTERNS = [
    # "第X章/节/篇/部分/卷"（中文数字或阿拉伯）
    re.compile(r"^\s*第[一二三四五六七八九十百千万零〇0-9]+[章节篇部分卷]\s*[::]?\s*"),
    # 阿拉伯数字点分前缀 "1.2.3 " 或 "1.2.3. "
    re.compile(r"^\s*\d+(?:\.\d+){0,6}[.．、\s]+"),
    # 括号中文序号 "（一）/（1）/(一)"
    re.compile(r"^\s*[（(][一二三四五六七八九十百千万\d]+[）)]\s*[::]?\s*"),
    # 中文数字序号 "一、/二．/三."
    re.compile(r"^\s*[一二三四五六七八九十百千万零〇]+[、.．:：]\s*"),
    # 罗马数字 + 点或括号
    re.compile(r"^\s*(?:[IVX]+|[ivx]+)[.．、)]\s*"),
    # 英文字母序号 "A./a)/A、"
    re.compile(r"^\s*[A-Za-z][.．、)]\s*"),
    # 前导"附"字（附录/附表类）
    re.compile(r"^\s*附\s*[:：]?\s*"),
]


def strip_prefix(text: str) -> str:
    """从单段 Heading 文本剥除所有可识别的章节号前缀。

    会连续尝试所有 pattern 直到没有匹配为止（处理 "第一章 1.1 xxx" 之类的复合前缀）。
    """
    if not text:
        return text
    prev = None
    current = text
    while prev != current:
        prev = current
        for pat in _PREFIX_PATTERNS:
            current = pat.sub("", current, count=1)
    return current.strip()


def strip_heading_prefixes_in_doc(doc, *, only_heading_styles: bool = True) -> int:
    """对 python-docx Document 中所有 Heading 段落剥除编号前缀。"""
    count = 0
    for para in doc.paragraphs:
        style_name = (para.style.name or "") if para.style else ""
        if only_heading_styles and not _is_heading_style(style_name):
            continue
        old = para.text
        new = strip_prefix(old)
        if new != old and new:
            _replace_paragraph_text_preserve_format(para, new)
            count += 1
    return count


def _is_heading_style(style_name: str) -> bool:
    """兼容英文 Heading X 和中文 标题 X。"""
    if not style_name:
        return False
    s = style_name.strip()
    if s.startswith("Heading ") or s.lower().startswith("heading "):
        return True
    if s.startswith("标题 ") or s == "标题":
        return True
    return False


def _heading_level(style_name: str) -> Optional[int]:
    """从样式名提取 Heading level（1-9）。返回 None 表示不是 Heading。"""
    if not style_name:
        return None
    s = style_name.strip()
    m = re.match(r"^(?:Heading|heading)\s+(\d+)$", s)
    if m:
        return int(m.group(1))
    m = re.match(r"^标题\s*(\d+)$", s)
    if m:
        return int(m.group(1))
    return None


def _direct_outline_level(para) -> Optional[int]:
    from docx.oxml.ns import qn

    p_pr = para._p.find(qn("w:pPr"))
    if p_pr is None:
        return None
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        return None
    try:
        value = int(outline.get(qn("w:val")))
    except (TypeError, ValueError):
        return None
    if 0 <= value <= 8:
        return value + 1
    return None


def _style_heading_level(style) -> Optional[int]:
    visited: set[str] = set()
    while style is not None:
        style_id = str(getattr(style, "style_id", "") or "")
        if style_id in visited:
            break
        visited.add(style_id)

        named_level = _heading_level(str(getattr(style, "name", "") or ""))
        if named_level is not None:
            return named_level

        p_pr = style.element.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
        if p_pr is not None:
            outline = p_pr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl")
            if outline is not None:
                from docx.oxml.ns import qn

                try:
                    value = int(outline.get(qn("w:val")))
                except (TypeError, ValueError):
                    value = -1
                if 0 <= value <= 8:
                    return value + 1
        style = getattr(style, "base_style", None)
    return None


def _paragraph_heading_level(para) -> Optional[int]:
    direct = _direct_outline_level(para)
    if direct is not None:
        return direct
    return _style_heading_level(para.style)


def _num_pr_suppresses_numbering(num_pr) -> bool:
    """Word uses direct numId=0 to suppress numbering inherited from a style."""
    from docx.oxml.ns import qn

    num_id = num_pr.find(qn("w:numId"))
    return num_id is not None and num_id.get(qn("w:val")) == "0"


def _replace_paragraph_text_preserve_format(para, new_text: str) -> None:
    """替换段落文本，尽量保留第一个 run 的格式（粗体/字体/字号）。"""
    if not para.runs:
        para.text = new_text
        return
    first_run = para.runs[0]
    first_run.text = new_text
    for extra_run in para.runs[1:]:
        extra_run.text = ""


def _replace_heading_number_preserve_format(para, pure_text: str, chapter_no: str) -> None:
    """只替换标题编号，保留标题正文所在 Run 及其格式。"""
    prefix = f"{chapter_no}  " if chapter_no else ""
    old_text = para.text or ""
    start = old_text.find(pure_text)
    if start < 0 or not para.runs:
        _replace_paragraph_text_preserve_format(para, prefix + pure_text)
        return

    end = start + len(pure_text)
    cursor = 0
    first_content_run = None
    for run in para.runs:
        run_text = run.text or ""
        run_start = cursor
        run_end = cursor + len(run_text)
        keep_start = max(start, run_start)
        keep_end = min(end, run_end)
        if keep_start < keep_end:
            run.text = run_text[keep_start - run_start : keep_end - run_start]
            if first_content_run is None:
                first_content_run = run
        else:
            run.text = ""
        cursor = run_end

    if first_content_run is None:
        _replace_paragraph_text_preserve_format(para, prefix + pure_text)
        return
    first_content_run.text = prefix + first_content_run.text


def _set_direct_heading_level(para, level: int) -> None:
    """写入导航层级但不替换素材段落样式。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = para._p.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is not None:
        p_pr.remove(outline)
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None and not _num_pr_suppresses_numbering(num_pr):
        p_pr.remove(num_pr)
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(max(1, min(int(level), 9)) - 1))
    p_pr.append(outline)


def _clear_direct_outline_and_numbering(para) -> None:
    """Remove paragraph-level outline/numbering overrides.

    Word/OnlyOffice navigation gives direct ``w:outlineLvl`` priority over the
    paragraph style. If a source H1 paragraph is later restyled to Heading 3 but
    keeps ``outlineLvl=0``, it still appears as a top-level item in the left
    outline pane.
    """
    from docx.oxml.ns import qn

    p_pr = para._p.find(qn("w:pPr"))
    if p_pr is None:
        return
    for tag in ("w:outlineLvl", "w:numPr"):
        el = p_pr.find(qn(tag))
        if el is not None:
            p_pr.remove(el)


def _set_body_style_or_clear(para, doc) -> None:
    """Switch to a body style; if the source doc lacks one, clear pStyle."""
    from docx.oxml.ns import qn

    for name in ("Normal", "正文"):
        try:
            para.style = doc.styles[name]
            return
        except KeyError:
            continue

    p_pr = para._p.find(qn("w:pPr"))
    if p_pr is None:
        return
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is not None:
        p_pr.remove(p_style)


# ---------- 方案 B 核心：前缀注入 ----------

def inject_prefix_to_headings(
    doc,
    parent_chapter_no: str,
    *,
    toc_title: Optional[str] = None,
    skip_first_if_match: bool = True,
    max_level: int = 6,
    l1_offset: int = 0,
) -> dict:
    """按父章节号给素材内部 Heading 注入编号前缀（方案 B 核心）。

    关键：用"**相对深度**"而非绝对 Heading level。素材起始 Heading 是 H3 还是 H1
    不影响结果——以素材内部**最小 level** 为 rel=1，按相对深度展开到父章节号下。

    Args:
        doc: python-docx Document
        parent_chapter_no: 父章节号（如 "5.8.7"）
        toc_title: toc 条目的纯标题（用于首 Heading 去重）
        skip_first_if_match: 首 Heading 文本匹配 toc_title 时是否剥除
        max_level: 只处理到这一级（绝对 Heading level，过深的忽略）

    Returns:
        {"injected": int, "skipped_first": bool, "removed": int, "min_level": int}
    """
    stats = {"injected": 0, "skipped_first": False, "removed": 0, "min_level": 0}

    # 第一遍：收集所有 Heading 段落（带 level）、识别首个需要 skip 的
    heading_entries: list[tuple] = []  # (para, lvl, pure)
    skip_idx: Optional[int] = None

    for para in doc.paragraphs:
        lvl = _paragraph_heading_level(para)
        if lvl is None or lvl > max_level:
            continue
        pure = strip_prefix(para.text)
        if not pure:
            continue
        heading_entries.append((para, lvl, pure))

    if not heading_entries:
        return stats

    if skip_first_if_match and toc_title:
        first_para, first_lvl, first_pure = heading_entries[0]
        if _normalize_for_match(first_pure) == _normalize_for_match(toc_title):
            skip_idx = 0
            stats["skipped_first"] = True

    if not any(i != skip_idx for i in range(len(heading_entries))):
        # 整份素材只有一个首 Heading 且被 skip
        if skip_idx is not None:
            p = heading_entries[0][0]
            p._p.getparent().remove(p._p)
            stats["removed"] = 1
        return stats
    stats["min_level"] = min(lvl for i, (_, lvl, _) in enumerate(heading_entries) if i != skip_idx)

    parent_parts = [p for p in (parent_chapter_no or "").split(".") if p]
    parent_depth = len(parent_parts)
    # 目标起始 level：父章节深度 + 1（封顶 6）
    target_first_level = min(parent_depth + 1, 6) if parent_depth > 0 else 1
    stats["target_first_level"] = target_first_level

    counters = [0] * (max_level + 2)  # 1-based
    counters[1] = l1_offset  # 用于避免 appendix 与前序 MATCHED 撞号
    paras_to_remove = []
    level_stack: list[tuple[int, int]] = []

    for i, (para, lvl, pure) in enumerate(heading_entries):
        if i == skip_idx:
            paras_to_remove.append(para)
            continue

        # 用行文顺序维护相对层级，而不是用全局最小 Heading level。
        # 很多素材内部会混用 H1/H2/H5，若直接用绝对 level 差值会生成
        # 1.9.0.0.1 这类空洞编号。这里把当前 heading 挂到最近的上级
        # heading 下面，缺失的中间层级会自动压缩。
        while level_stack and level_stack[-1][0] >= lvl:
            level_stack.pop()
        rel = level_stack[-1][1] + 1 if level_stack else 1
        rel = max(1, min(rel, max_level))
        level_stack.append((lvl, rel))

        counters[rel] += 1
        for d in range(rel + 1, len(counters)):
            counters[d] = 0

        # 只更新导航层级，不替换素材自己的段落样式。
        new_level = max(1, min(6, target_first_level + rel - 1))
        _set_direct_heading_level(para, new_level)

        my_parts = parent_parts + [str(counters[k]) for k in range(1, rel + 1)]
        _replace_heading_number_preserve_format(para, pure, ".".join(my_parts))
        stats["injected"] += 1

    for p in paras_to_remove:
        p._p.getparent().remove(p._p)
        stats["removed"] += 1

    # 返回本次注入最终 L1（相对）计数器值，供 merger 追踪 parent 下已用编号
    stats["top_count"] = counters[1] - l1_offset
    stats["final_l1"] = counters[1]
    return stats


def _set_heading_style(para, level: int, doc) -> None:
    """把段落样式切换到 Heading {level}，兼容中文 "标题 N"。"""
    target = f"Heading {level}"
    try:
        para.style = doc.styles[target]
        _clear_direct_outline_and_numbering(para)
        return
    except KeyError:
        pass
    cn = f"标题 {level}"
    try:
        para.style = doc.styles[cn]
        _clear_direct_outline_and_numbering(para)
    except KeyError:
        pass


def _normalize_for_match(s: str) -> str:
    """归一化：去空白、去全半角差异、统一括号，便于比较首 heading 与 toc title。"""
    if not s:
        return ""
    s = strip_prefix(s)
    s = s.strip()
    s = re.sub(r"\s+", "", s)
    # 全角 → 半角常用
    s = s.replace("（", "(").replace("）", ")")
    s = s.replace("：", ":").replace("，", ",")
    return s


def _normalize_title_key(s: str) -> str:
    """Stronger title key for matching material subheadings to S2 TOC items."""
    s = _normalize_for_match(s)
    return re.sub(r"[^\w\u4e00-\u9fff]+", "", s)


def demote_headings_to_body(
    doc,
    *,
    toc_title: Optional[str] = None,
    remove_first_if_match: bool = True,
    keep_heading_map: Optional[dict] = None,
) -> dict:
    """Convert material-internal headings to body paragraphs.

    S7 inserts navigational headings from the S2 TOC. Source-material headings
    are useful as in-body subheadings. If a source heading matches a known S2
    child item, keep it as the official numbered TOC heading; otherwise it must
    not enter Word/OnlyOffice navigation. This keeps text, removes existing
    numbering prefixes, and clears both Heading styles and direct outline levels
    for non-TOC material headings.
    """
    normalized_keep: dict[str, dict] = {}
    for key, value in (keep_heading_map or {}).items():
        if not isinstance(value, dict):
            continue
        normalized_keep[_normalize_title_key(str(key))] = value

    stats = {"demoted": 0, "kept": 0, "removed": 0, "skipped_first": False}
    heading_entries: list[tuple] = []
    for para in doc.paragraphs:
        lvl = _paragraph_heading_level(para)
        if lvl is None:
            continue
        pure = strip_prefix(para.text)
        if not pure:
            continue
        heading_entries.append((para, pure))

    if not heading_entries:
        return stats

    first_to_remove = None
    if remove_first_if_match and toc_title:
        first_para, first_pure = heading_entries[0]
        if _normalize_for_match(first_pure) == _normalize_for_match(toc_title):
            first_to_remove = first_para
            stats["skipped_first"] = True

    for para, pure in heading_entries:
        if para is first_to_remove:
            continue
        keep = normalized_keep.get(_normalize_title_key(pure))
        if keep:
            title = str(keep.get("title") or pure).strip()
            chapter_no = str(keep.get("chapter_no") or "").strip()
            try:
                level = int(keep.get("level") or 1)
            except (TypeError, ValueError):
                level = 1
            new_text = f"{chapter_no}  {title}" if chapter_no else title
            _replace_paragraph_text_preserve_format(para, new_text)
            _set_heading_style(para, level, doc)
            stats["kept"] += 1
            continue
        if para.text != pure:
            _replace_paragraph_text_preserve_format(para, pure)
        _set_body_style_or_clear(para, doc)
        _clear_direct_outline_and_numbering(para)
        for run in para.runs:
            run.bold = True
        stats["demoted"] += 1

    if first_to_remove is not None:
        first_to_remove._p.getparent().remove(first_to_remove._p)
        stats["removed"] = 1

    return stats


def remap_material_headings_to_navigation(
    doc,
    *,
    toc_title: Optional[str] = None,
    remove_first_if_match: bool = True,
    keep_heading_map: Optional[dict] = None,
    parent_chapter_no: str = "",
    parent_level: int = 2,
    max_target_level: int = 6,
    l1_offset: int = 0,
) -> dict:
    """按素材真实大纲树连续编号，并保留素材原段落样式。"""
    # 参数保留用于兼容现有调用；素材标题编号不再依赖 S2 文本匹配。
    _ = keep_heading_map
    max_target_level = max(1, min(int(max_target_level or 6), 9))
    parent_level = max(1, min(int(parent_level or 1), max_target_level))
    base_child_level = min(parent_level + 1, max_target_level)

    stats = {
        "remapped": 0,
        "kept": 0,
        "removed": 0,
        "demoted": 0,
        "bold_subheadings": 0,
        "skipped_first": False,
    }

    heading_entries: list[tuple] = []
    for para in doc.paragraphs:
        lvl = _paragraph_heading_level(para)
        if lvl is None:
            continue
        pure = strip_prefix(para.text)
        if not pure:
            continue
        heading_entries.append((para, lvl, pure))

    first_to_remove = None
    first_to_remove_key = None
    if heading_entries and remove_first_if_match and toc_title:
        first_para, _first_level, first_pure = heading_entries[0]
        if _normalize_for_match(first_pure) == _normalize_for_match(toc_title):
            first_to_remove = first_para
            first_to_remove_key = id(first_para._p)
            stats["skipped_first"] = True

    heading_stack: list[tuple[int, int]] = []
    counters = [0] * 10
    counters[1] = max(0, int(l1_offset or 0))
    parent_parts = [part for part in parent_chapter_no.split(".") if part]

    for para, source_level, pure in heading_entries:
        if id(para._p) == first_to_remove_key:
            continue

        while heading_stack and heading_stack[-1][0] >= source_level:
            heading_stack.pop()
        relative_level = heading_stack[-1][1] + 1 if heading_stack else 1
        relative_level = max(1, min(relative_level, 9))
        heading_stack.append((source_level, relative_level))

        counters[relative_level] += 1
        for index in range(relative_level + 1, len(counters)):
            counters[index] = 0

        target_level = min(base_child_level + relative_level - 1, max_target_level)
        number_parts = parent_parts + [
            str(counters[index]) for index in range(1, relative_level + 1)
        ]
        _replace_heading_number_preserve_format(para, pure, ".".join(number_parts))
        _set_direct_heading_level(para, target_level)
        stats["remapped"] += 1

    if first_to_remove is not None:
        first_to_remove._p.getparent().remove(first_to_remove._p)
        stats["removed"] = 1

    stats["top_count"] = counters[1] - max(0, int(l1_offset or 0))
    stats["final_l1"] = counters[1]
    return stats


def _clean_text_for_heading_shape(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "").replace("\u3000", " ")).strip()


def _looks_like_caption_or_table_title(text: str) -> bool:
    clean = _clean_text_for_heading_shape(text)
    compact = re.sub(r"\s+", "", clean)
    if not compact:
        return False
    if re.match(r"^(?:图|表)\s+\S", clean):
        return True
    if re.match(r"^(?:图|表)[：:]\S", clean):
        return True
    if re.match(r"^(?:图|表)[A-Za-z]?[.-]?\d", compact):
        return True
    if re.match(r"^(?:图|表)[一二三四五六七八九十]+", compact):
        return True
    if re.match(r"^图[\u4e00-\u9fffA-Za-z0-9]{2,40}(?:示意图|结构图|流程图|布置图|接线图|曲线图|照片)$", compact):
        return True
    if "一览表" in compact or "统计表" in compact or "参数表" in compact:
        return True
    if compact.endswith(("表", "清单")) and len(compact) <= 28:
        return True
    return False


def _looks_like_bold_body_subheading(para, text: str) -> bool:
    if not text:
        return False
    if len(text) > 48:
        return False
    if text.startswith(("[", "【", "*")):
        return False
    if re.match(r"^(?:备注|注|说明)[:：]", text):
        return False
    if re.search(r"[。；;！!？?：:]$", text):
        return False
    if _looks_like_caption_or_table_title(text):
        return False
    if re.match(r"^[（(][一二三四五六七八九十\d]+[）)]\s*\S", text):
        return False
    if re.match(r"^\d+[）)、.]\s*\S", text):
        return False
    if re.match(r"^[一二三四五六七八九十]+[、.]\s*\S", text):
        return False
    if _bold_text_ratio(para) < 0.65:
        return False
    if re.search(r"[，,。；;！!？?：:]", text):
        return False
    return len(re.findall(r"[\u4e00-\u9fff]", text)) >= 2


def _bold_text_ratio(para) -> float:
    total = 0
    bold = 0
    for run in para.runs:
        text = str(run.text or "")
        visible = len(re.sub(r"\s+", "", text))
        if not visible:
            continue
        total += visible
        if run.bold is True or getattr(run.font, "bold", None) is True:
            bold += visible
    return bold / total if total else 0.0


# ---------- 正文手写编号擦除 ----------

# 常见单位字（开头一个字就能判定是单位的）；后接其它字也不应被当章节号
_UNIT_CHARS = (
    "秒分时日天年月周"        # 时间
    "米厘毫千公英里"           # 长度
    "倍次个台件份块条根本页条次"  # 量词
    "度摄"                   # 温度
    "斤吨克"                 # 质量
    "瓦伏安"                 # 电功率/电压/电流
    "帕兆"                   # 压力/兆
    "微纳"                   # 微/纳
    "元万亿"                 # 货币
    "小"                     # "小时"
    "立方平方"                # 面积/体积（用"立"/"平"拦截）
    "赫摩尔焦"               # 物理单位
    "%％℃°"                 # 符号（顺便加进负 lookahead）
)

# 仅匹配段首多级数字前缀 + 空白 + 中文字符。要求：
# - 至少 X.Y（两级）
# - 紧邻后一个汉字非单位字
# - 再后还有 2+ 汉字（避免 "X.Y 秒" / "X.Y 倍" 误伤）
_HAND_NUMBER_PREFIX = re.compile(
    r"^(\s*)(\d+(?:\.\d+){1,6})([\.\s　]+)"
    rf"(?=[\u4e00-\u9fff])(?![{re.escape(_UNIT_CHARS)}])"
    r"(?=[\u4e00-\u9fff]{2,})"
)


def strip_handwritten_numbering_in_body(doc, *, only_normal_style: bool = True) -> int:
    """擦除正文段落首的手写多级编号（"5.1 xxx" → "xxx"）。

    只处理 Normal / 正文 / 空样式段落，不碰 Heading 样式。

    Returns:
        修改段落数
    """
    count = 0
    for para in doc.paragraphs:
        if only_normal_style:
            if _paragraph_heading_level(para) is not None:
                continue
        old = para.text
        if not old:
            continue
        new = _HAND_NUMBER_PREFIX.sub(r"\1", old, count=1)
        if new != old:
            _replace_paragraph_text_preserve_format(para, new)
            count += 1
    return count


# ---------- 样式归一化 ----------

def normalize_heading_style_names(doc) -> int:
    """把段落样式名 '标题 N' 统一改为 'Heading N'。"""
    import re as _re
    pat = _re.compile(r"^标题\s*(\d+)$")
    count = 0
    for para in doc.paragraphs:
        st = para.style.name if para.style else ""
        m = pat.match(st or "")
        if not m:
            continue
        n = m.group(1)
        en_name = f"Heading {n}"
        try:
            para.style = doc.styles[en_name]
            count += 1
        except KeyError:
            pass
    return count


# ---------- numPr 剥离 ----------

def strip_numPr_from_body(doc, *, only_heading_styles: bool = True) -> int:
    """剥掉 Heading 样式段落的 w:numPr（破多级列表自动编号绑定）。

    默认 only_heading_styles=True：只对 Heading 样式段剥，**保留正文列表**
    （"1)xxx; 2)xxx; ..." 这种 Word 自动编号列表，用户期望保留）。
    """
    from docx.oxml.ns import qn
    count = 0

    def _should_strip(para) -> bool:
        if not only_heading_styles:
            return True
        return _paragraph_heading_level(para) is not None

    for para in doc.paragraphs:
        if not _should_strip(para):
            continue
        pPr = para._p.find(qn("w:pPr"))
        if pPr is None:
            continue
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None and not _num_pr_suppresses_numbering(numPr):
            pPr.remove(numPr)
            count += 1
    # 表格里的段落也要处理
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if not _should_strip(para):
                        continue
                    pPr = para._p.find(qn("w:pPr"))
                    if pPr is None:
                        continue
                    numPr = pPr.find(qn("w:numPr"))
                    if numPr is not None and not _num_pr_suppresses_numbering(numPr):
                        pPr.remove(numPr)
                        count += 1
    return count


def strip_numPr_from_heading_styles(doc) -> int:
    """剥掉 Heading 样式自身绑定的 Word 自动编号。

    技术标正文使用"文本编号 + Heading 样式"的方案。如果母版 Heading
    style 仍绑定了多级列表，Word/OnlyOffice 会在显示层再自动加一次编号，
    形成 "1.7 1.7 标题"。此函数清掉样式级 numPr，保留字体、字号等样式。
    """
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn

    count = 0
    for style in doc.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        pPr = style.element.find(qn("w:pPr"))
        if pPr is None:
            continue
        if _style_heading_level(style) is None:
            continue
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            pPr.remove(numPr)
            count += 1
    return count


if __name__ == "__main__":
    # 前缀剥离 smoke test
    print("== strip_prefix ==")
    samples = [
        "1.1 发电机设计原理",
        "第一章 标前概述",
        "1.1.2 外壳结构",
        "一、总体方案",
        "（二）技术参数",
        "第一章  1.1  双编号",
        "纯标题",
        "附  校核报告",
    ]
    for s in samples:
        print(f"  {s!r:35s} → {strip_prefix(s)!r}")

    # 手写编号擦除 smoke test
    print("\n== strip_handwritten (regex only) ==")
    body_samples = [
        "7.10 符合招标公告及招标文件要求的业绩情况",  # 应擦
        "5.1 项目概况",                                # 应擦
        "1. 第一步",                                   # 不擦（单级列表）
        "7.10 m/s 额定风速",                           # 不擦（单位）
        "2.35 秒启动时间",                             # 不擦（数字+单位）
        "一、总体方案",                                # 不擦（中文序号）
        "普通段落",                                    # 不擦
        "1.5 倍超速保护",                              # 不擦（后接"倍"）
    ]
    for s in body_samples:
        new = _HAND_NUMBER_PREFIX.sub("", s, count=1)
        marker = "✗" if new == s else "✓"
        print(f"  {marker}  {s!r:40s} → {new!r}")
