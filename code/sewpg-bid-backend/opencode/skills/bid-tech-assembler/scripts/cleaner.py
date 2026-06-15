"""
招标文件数据清洗模块

从 .docx 招标文件中提取干净、结构化的文本，为后续 LLM 解析做准备。

清洗规则:
1. 去噪: 删除页眉/页脚重复文本、连续空行、目录区域
2. 标题归一化: 统一中文数字前缀格式
3. 表格结构化: 提取为 List[List[str]] 便于 LLM 读取
4. 区域分割: 识别"正文区"和"附件模板区"的边界
"""

import logging
import os
import re
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Tuple

from docx import Document

logger = logging.getLogger(__name__)


@dataclass
class TableData:
    """结构化表格"""
    headers: List[str]
    rows: List[List[str]]
    context: str = ""  # 表格前方最近的标题/段落文本, 用于推断语义

    def to_markdown(self) -> str:
        """转为 markdown 表格字符串"""
        if not self.headers:
            return ""
        lines = ["| " + " | ".join(self.headers) + " |"]
        lines.append("| " + " | ".join(["---"] * len(self.headers)) + " |")
        for row in self.rows:
            # 补齐列数
            padded = row + [""] * (len(self.headers) - len(row))
            lines.append("| " + " | ".join(padded[:len(self.headers)]) + " |")
        return "\n".join(lines)


@dataclass
class Section:
    """文档章节"""
    title: str
    level: int  # 1=顶级, 2=子章节, ...
    content: str  # 该章节下的文本内容(不含子章节)
    tables: List[TableData] = field(default_factory=list)
    children: List["Section"] = field(default_factory=list)

    def full_text(self, include_tables: bool = True) -> str:
        """获取完整文本(含子章节)"""
        parts = [self.content]
        if include_tables:
            for t in self.tables:
                parts.append(t.to_markdown())
        for child in self.children:
            parts.append(f"{'#' * (child.level + 1)} {child.title}")
            parts.append(child.full_text(include_tables))
        return "\n\n".join(p for p in parts if p.strip())


@dataclass
class CleanedDoc:
    """清洗后的文档"""
    full_text: str                     # 清洗后全文(纯文本)
    sections: List[Section]            # 按标题层级拆分的章节树
    tables: List[TableData]            # 所有表格(结构化)
    appendix_text: str = ""            # 附件/模板区域文本
    main_text: str = ""                # 正文区域文本(不含附件区)
    raw_paragraphs_count: int = 0      # 原始段落数
    cleaned_paragraphs_count: int = 0  # 清洗后段落数

    def summary(self) -> str:
        return (
            f"清洗结果: {self.raw_paragraphs_count} → {self.cleaned_paragraphs_count} 段落, "
            f"{len(self.tables)} 个表格, "
            f"{len(self.sections)} 个顶级章节, "
            f"正文 {len(self.main_text)} 字, "
            f"附件区 {len(self.appendix_text)} 字"
        )


class DocxCleaner:
    """招标文件原始数据清洗器"""

    # 目录页特征行(匹配其中任一则认为是目录区域)
    TOC_PATTERNS = [
        re.compile(r'^\s*目\s*录\s*$'),
        re.compile(r'^\s*目\s*次\s*$'),
    ]

    # 附件模板区起始锚点(支持多种编号格式, 按优先级排列)
    APPENDIX_ANCHORS = [
        # 华能格式
        re.compile(r'投标函的格式\s*[\(（]\s*1\s*[Aａ]\s*[\)）]', re.IGNORECASE),
        # 通用格式: "附件1" / "附录一" / "Appendix 1"
        re.compile(r'附件\s*[1１一]\s*[：:.\s]'),
        re.compile(r'附录\s*[1１一]\s*[：:.\s]'),
        # 投标文件格式要求
        re.compile(r'投标文件格式[及与和]要求'),
        re.compile(r'投标文件组成及格式'),
    ]

    # 页眉/页脚重复文本特征
    HEADER_FOOTER_PATTERNS = [
        re.compile(r'第\s*\d+\s*页\s*[共/]?\s*\d*\s*页?'),
        re.compile(r'Page\s+\d+\s+of\s+\d+', re.IGNORECASE),
    ]

    # 中文数字到阿拉伯数字映射
    CN_NUMS = {
        '一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
        '六': 6, '七': 7, '八': 8, '九': 9, '十': 10,
        '十一': 11, '十二': 12, '十三': 13, '十四': 14, '十五': 15,
    }

    def clean(self, docx_path: str) -> CleanedDoc:
        """
        清洗 docx 招标文件。

        Args:
            docx_path: 招标文件 .docx 路径

        Returns:
            CleanedDoc 实例
        """
        logger.info(f"开始清洗: {docx_path}")
        doc = Document(docx_path)

        # Step 1: 提取原始段落和表格
        raw_paragraphs = self._extract_paragraphs(doc)
        raw_tables = self._extract_tables(doc)
        logger.info(f"原始数据: {len(raw_paragraphs)} 段落, {len(raw_tables)} 表格")

        # Step 2: 去噪
        cleaned_paragraphs = self._remove_noise(raw_paragraphs)
        logger.info(f"去噪后: {len(cleaned_paragraphs)} 段落")

        # Step 3: 识别目录区并移除
        cleaned_paragraphs = self._remove_toc(cleaned_paragraphs)

        # Step 4: 分割正文区和附件模板区
        main_paras, appendix_paras = self._split_appendix(cleaned_paragraphs)

        # Step 5: 构建清洗后全文
        full_text = "\n".join(cleaned_paragraphs)
        main_text = "\n".join(main_paras)
        appendix_text = "\n".join(appendix_paras)

        # Step 6: 章节树构建
        sections = self._build_section_tree(main_paras, raw_tables)

        result = CleanedDoc(
            full_text=full_text,
            sections=sections,
            tables=raw_tables,
            appendix_text=appendix_text,
            main_text=main_text,
            raw_paragraphs_count=len(raw_paragraphs),
            cleaned_paragraphs_count=len(cleaned_paragraphs),
        )
        logger.info(result.summary())
        return result

    # ================================================================
    # 提取
    # ================================================================

    def _extract_paragraphs(self, doc: Document) -> List[str]:
        """提取所有段落文本(含样式信息用于后续标题识别)"""
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                paragraphs.append("")
                continue

            # 保留标题层级信息: 在文本前添加标记
            level = self._get_heading_level(para)
            if level and level <= 4:
                text = f"{'#' * level} {text}"

            paragraphs.append(text)
        return paragraphs

    def _get_heading_level(self, para) -> Optional[int]:
        """从段落样式中提取标题层级"""
        if not para.style or not para.style.name:
            return None
        style_name = para.style.name
        if 'Heading' in style_name:
            try:
                return int(style_name.replace('Heading', '').strip())
            except ValueError:
                return None
        # 一些文档使用自定义样式名
        if style_name.startswith('标题'):
            try:
                return int(style_name.replace('标题', '').strip())
            except ValueError:
                return None
        return None

    def _extract_tables(self, doc: Document) -> List[TableData]:
        """提取所有表格为结构化数据"""
        tables = []
        body = doc.element.body

        for table in doc.tables:
            if len(table.rows) < 1:
                continue

            # 获取表格前方段落作为上下文
            context = self._get_table_context(doc, table)

            # 提取表头(第一行)
            headers = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells]

            # 去重相邻的合并单元格
            headers = self._dedup_merged_cells(headers)

            # 提取数据行
            rows = []
            for row_idx in range(1, len(table.rows)):
                cells = [cell.text.strip().replace('\n', ' ')
                         for cell in table.rows[row_idx].cells]
                cells = self._dedup_merged_cells(cells)
                # 跳过完全空的行
                if any(c.strip() for c in cells):
                    rows.append(cells)

            tables.append(TableData(
                headers=headers,
                rows=rows,
                context=context,
            ))

        return tables

    def _get_table_context(self, doc: Document, table) -> str:
        """获取表格前方最近的非空段落文本"""
        table_element = table._tbl
        body = doc.element.body
        prev_texts = []
        for child in body:
            if child is table_element:
                break
            if child.tag.endswith('}p'):
                text_parts = []
                for run in child.iter():
                    if run.text:
                        text_parts.append(run.text)
                text = ''.join(text_parts).strip()
                if text:
                    prev_texts.append(text)
        # 返回最近的 2 个非空段落
        return " | ".join(prev_texts[-2:]) if prev_texts else ""

    @staticmethod
    def _dedup_merged_cells(cells: List[str]) -> List[str]:
        """去除 python-docx 因合并单元格产生的相邻重复值"""
        if not cells:
            return cells
        result = [cells[0]]
        for c in cells[1:]:
            if c != result[-1]:
                result.append(c)
            # 如果和前一个完全相同, 说明是合并单元格, 跳过
        return result

    # ================================================================
    # 去噪
    # ================================================================

    def _remove_noise(self, paragraphs: List[str]) -> List[str]:
        """去除噪音段落"""
        result = []
        consecutive_empty = 0

        for text in paragraphs:
            # 空行: 最多保留 1 个
            if not text.strip():
                consecutive_empty += 1
                if consecutive_empty <= 1:
                    result.append("")
                continue
            consecutive_empty = 0

            # 跳过页码行
            if any(p.search(text) for p in self.HEADER_FOOTER_PATTERNS):
                continue

            # 跳过纯装饰线
            stripped = text.strip('#').strip()
            if stripped and all(c in '=—-_─━' for c in stripped):
                continue

            result.append(text)

        return result

    def _remove_toc(self, paragraphs: List[str]) -> List[str]:
        """识别并移除目录区域"""
        toc_start = -1
        toc_end = -1

        for i, text in enumerate(paragraphs):
            clean = text.lstrip('#').strip()
            if any(p.match(clean) for p in self.TOC_PATTERNS):
                toc_start = i
                break

        if toc_start < 0:
            return paragraphs

        # 目录区通常到下一个正文标题结束
        # 目录条目特征(多种格式):
        #   1) "第X章 xxx ....... 页码"  (点线格式)
        #   2) "1. 招标条件 2"            (标题+tab/空格+页码)
        #   3) "1.1 xxx     16"          (多级编号+页码)
        toc_entry_patterns = [
            # 点线格式: "xxx......N"
            re.compile(r'.+[\.…·]{3,}\s*\d+\s*$'),
            # 编号+标题+尾部页码: "1.1 xxx 16" (末尾是 1-4 位数字)
            re.compile(r'^(?:第[一二三四五六七八九十\d]+[章卷篇节]|'
                       r'\d+(?:\.\d+)*[.\s])\s*.+\s+\d{1,4}\s*$'),
            # 纯标题行(无内容, 只有编号和短标题, 跟页码): "投标人须知前附表    11"
            re.compile(r'^[\u4e00-\u9fa5]{2,20}\s+\d{1,4}\s*$'),
        ]

        def _is_toc_entry(line: str) -> bool:
            return any(p.match(line) for p in toc_entry_patterns)

        for i in range(toc_start + 1, min(toc_start + 300, len(paragraphs))):
            text = paragraphs[i].lstrip('#').strip()
            if not text:
                continue
            # 如果遇到不像目录条目的行, 检查是否目录已结束
            if not _is_toc_entry(text):
                # 往后看 5 行, 如果仍有目录条目则继续
                looks_like_toc = False
                for j in range(i + 1, min(i + 5, len(paragraphs))):
                    check = paragraphs[j].lstrip('#').strip()
                    if check and _is_toc_entry(check):
                        looks_like_toc = True
                        break
                if not looks_like_toc:
                    toc_end = i
                    break

        if toc_end <= toc_start:
            toc_end = toc_start + 1  # 只有目录标题, 删掉它

        removed = toc_end - toc_start
        logger.info(f"移除目录区: 第 {toc_start}-{toc_end} 行 ({removed} 行)")
        return paragraphs[:toc_start] + paragraphs[toc_end:]

    # ================================================================
    # 区域分割
    # ================================================================

    def _split_appendix(self, paragraphs: List[str]) -> Tuple[List[str], List[str]]:
        """分割正文区和附件模板区"""
        split_idx = len(paragraphs)

        for i, text in enumerate(paragraphs):
            clean = text.lstrip('#').strip()
            for pattern in self.APPENDIX_ANCHORS:
                if pattern.search(clean):
                    # 确保不是目录中的同名引用(通常在文档后半部分)
                    if i > len(paragraphs) * 0.3:  # 至少在文档 30% 之后
                        split_idx = i
                        logger.info(f"附件区起始: 第 {i} 行 — {clean[:50]}")
                        break
            if split_idx < len(paragraphs):
                break

        main = paragraphs[:split_idx]
        appendix = paragraphs[split_idx:]
        return main, appendix

    # ================================================================
    # 章节树构建
    # ================================================================

    # 标题模式匹配(按优先级排列)
    HEADING_PATTERNS = [
        # markdown 标题 (来自 _extract_paragraphs 的标记)
        (re.compile(r'^(#{1,4})\s+(.+)$'), lambda m: (len(m.group(1)), m.group(2).strip())),
        # "第X章 xxx" / "第X卷 xxx"
        (re.compile(r'^第[一二三四五六七八九十\d]+[章卷篇]\s+(.+)$'),
         lambda m: (1, m.group(0))),
        # "一、xxx" / "（一）xxx"
        (re.compile(r'^[一二三四五六七八九十]+[、．.]\s*(.+)$'),
         lambda m: (1, m.group(0))),
        (re.compile(r'^[（(][一二三四五六七八九十\d]+[）)]\s*(.+)$'),
         lambda m: (2, m.group(0))),
        # "1. xxx" / "1.1 xxx" / "1.1.1 xxx"
        (re.compile(r'^(\d+(?:\.\d+)+)\s+(.+)$'),
         lambda m: (len(m.group(1).split('.')), m.group(0))),
        (re.compile(r'^(\d+)[.、]\s+(.+)$'),
         lambda m: (2, m.group(0))),
    ]

    def _build_section_tree(self, paragraphs: List[str],
                            tables: List[TableData]) -> List[Section]:
        """从段落构建章节树"""
        # 先识别所有标题及其位置
        headings = []  # [(line_idx, level, title)]
        for i, text in enumerate(paragraphs):
            level, title = self._detect_heading(text)
            if level:
                headings.append((i, level, title))

        if not headings:
            # 没有可识别的标题, 整个文档作为一个章节
            return [Section(
                title="全文",
                level=0,
                content="\n".join(paragraphs),
                tables=tables,
            )]

        # 构建章节列表(平铺)
        sections_flat = []
        for idx, (line_idx, level, title) in enumerate(headings):
            # 内容范围: 当前标题到下一个标题之间
            next_line = headings[idx + 1][0] if idx + 1 < len(headings) else len(paragraphs)
            content_lines = paragraphs[line_idx + 1:next_line]
            content = "\n".join(line for line in content_lines if line.strip())

            sections_flat.append(Section(
                title=title,
                level=level,
                content=content,
            ))

        # 构建树结构
        return self._nest_sections(sections_flat)

    def _detect_heading(self, text: str) -> Tuple[Optional[int], str]:
        """检测一行文本是否是标题, 返回 (level, title) 或 (None, "")"""
        if not text.strip():
            return None, ""

        for pattern, extractor in self.HEADING_PATTERNS:
            m = pattern.match(text.strip())
            if m:
                level, title = extractor(m)
                # 过滤掉太短或太长的"标题"
                clean_title = title.lstrip('#').strip()
                if len(clean_title) < 2 or len(clean_title) > 100:
                    continue
                return level, clean_title

        return None, ""

    @staticmethod
    def _nest_sections(flat_sections: List[Section]) -> List[Section]:
        """将平铺的章节列表构建为树结构"""
        if not flat_sections:
            return []

        root = []
        stack = []  # [(section, level)]

        for section in flat_sections:
            # 弹出栈中层级 >= 当前的所有项
            while stack and stack[-1][1] >= section.level:
                stack.pop()

            if stack:
                # 作为栈顶的子节点
                stack[-1][0].children.append(section)
            else:
                # 顶级节点
                root.append(section)

            stack.append((section, section.level))

        return root


# ================================================================
# 便捷函数
# ================================================================

def clean_bid_document(docx_path: str) -> CleanedDoc:
    """便捷函数: 清洗招标文件"""
    cleaner = DocxCleaner()
    return cleaner.clean(docx_path)


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("用法: python cleaner.py <招标文件.docx> [<输出.txt>]")
        print("注意: 本脚本是招标文件分析器，不属于 bid-tech-assembler 正式流程；")
        print("      必须显式给出输出路径，禁止在素材库里产生 *_cleaned.txt。")
        sys.exit(1)

    if len(sys.argv) < 3:
        print("错误：需要显式指定输出路径作为第 2 个参数，避免污染素材库。")
        sys.exit(2)

    out_path = sys.argv[2]
    # 硬禁止写入素材库
    if "素材库" in out_path or "投标资料库" in out_path:
        print(f"错误：输出路径 {out_path} 在素材库内，禁止。请放到 /tmp 或工作目录外。")
        sys.exit(3)

    logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
    result = clean_bid_document(sys.argv[1])
    print(result.summary())

    with open(out_path, 'w', encoding='utf-8') as f:
        f.write("=== 正文区 ===\n\n")
        f.write(result.main_text)
        f.write("\n\n=== 附件模板区 ===\n\n")
        f.write(result.appendix_text)
    print(f"清洗结果已保存: {out_path}")

    # 输出章节树
    print(f"\n=== 章节树 ({len(result.sections)} 个顶级章节) ===")
    for s in result.sections:
        print(f"  {'  ' * s.level}[L{s.level}] {s.title} ({len(s.content)}字, {len(s.children)}个子节)")

    # 输出表格摘要
    print(f"\n=== 表格 ({len(result.tables)} 个) ===")
    for i, t in enumerate(result.tables):
        print(f"  表{i+1}: {len(t.headers)}列 × {len(t.rows)}行 | 上下文: {t.context[:60]}")
