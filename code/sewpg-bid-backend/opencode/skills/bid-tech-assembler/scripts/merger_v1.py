#!/usr/bin/env python3
"""
按 assembly_plan.json 逐条合并到技术标母版。

对每条 plan 项：
  - OUT_OF_SCOPE：跳过
  - STRUCTURAL：只插 Heading（伞形章节，标题占位）
  - NEEDS_REVIEW：插 Heading + '[待填写：title]' 占位段
  - MATCHED/ADAPTED：
      - preprocess 每个 path → /tmp/bid_prep/<hash>.docx
      - 在 master 追加一个 Heading（toc 的 level 和 title）
      - 用 docxcompose 把预处理后的素材追加进来

依赖：docxcompose、python-docx。

用法：
    python3 merger.py \\
        --template templates/技术投标母版模板.docx \\
        --plan /tmp/assembly_plan.json \\
        --lib  /Users/wlb/Downloads/技术标/素材库 \\
        --params /tmp/project_params.json \\
        --prep-dir /tmp/bid_prep \\
        --out /tmp/bid_merged.docx
"""

from __future__ import annotations

import argparse
import hashlib
import json
import logging
import os
import shutil
import sys
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt
from docxcompose.composer import Composer

sys.path.insert(0, str(Path(__file__).resolve().parent))
from preprocess import preprocess

logging.basicConfig(level=logging.INFO, format="[merger] %(message)s")
log = logging.getLogger("merger")


# ---------- 工具 ----------

def _hash_path(p: Path) -> str:
    return hashlib.md5(str(p).encode("utf-8")).hexdigest()[:10]


def add_heading(doc: Document, text: str, level: int) -> None:
    """在 doc 末尾追加 Heading N 段。level clamp 到 1-6。"""
    level = max(1, min(6, level))
    style_name = f"Heading {level}"
    try:
        p = doc.add_paragraph(text, style=style_name)
    except KeyError:
        # fallback
        p = doc.add_paragraph(text)
    return p


def add_placeholder_paragraph(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_preface_heading(doc: Document, text: str) -> None:
    """前言段：无编号 Heading 1。技术母版的 Heading 1 带多级列表编号；
    暂用 Normal 样式加粗大字代替，避免被多级列表吃掉。
    """
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    # 从 heading_style.json 看 H1 是等线 Light 小三
    run.font.name = "等线 Light"
    # 中文字体 rFonts 注入
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), "等线 Light")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")


# ---------- 主 ----------

def merge(
    template_path: Path,
    plan: list[dict],
    lib_root: Path,
    params: dict,
    prep_dir: Path,
    out_path: Path,
) -> dict:
    prep_dir.mkdir(parents=True, exist_ok=True)

    # 打开母版；母版的 body 只有 1 个空段，删掉它以便追加从头开始
    doc = Document(str(template_path))
    body = doc.element.body
    from docx.oxml.ns import qn
    # 保留 sectPr，删其它
    sect_pr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is sect_pr:
            continue
        body.remove(child)

    composer = Composer(doc)

    stats = {
        "inserted_headings": 0,
        "inserted_placeholders": 0,
        "merged_materials": 0,
        "skipped_oos": 0,
        "errors": 0,
    }

    in_scope_count = sum(1 for p in plan if p["status"] != "OUT_OF_SCOPE")
    log.info(f"in-scope entries: {in_scope_count}/{len(plan)}")

    for i, entry in enumerate(plan):
        status = entry["status"]
        level = entry["level"]
        title = entry["title"]

        if status == "OUT_OF_SCOPE":
            stats["skipped_oos"] += 1
            continue

        if status == "STRUCTURAL":
            # 伞形章节只插 Heading（由多级列表自动编号）
            add_heading(doc, title, level)
            stats["inserted_headings"] += 1
            continue

        if status == "NEEDS_REVIEW":
            add_heading(doc, title, level)
            add_placeholder_paragraph(doc, f"[待填写：{title}——本节由招标/模板新增，请补素材]")
            stats["inserted_headings"] += 1
            stats["inserted_placeholders"] += 1
            continue

        if status in ("MATCHED", "ADAPTED"):
            # 前言段：特殊 Heading
            if entry.get("is_preface"):
                add_preface_heading(doc, title)
            elif entry.get("is_appendix"):
                # 附字头不另插 Heading；素材本身有 heading
                pass
            else:
                add_heading(doc, title, level)
            stats["inserted_headings"] += 1

            # 逐份素材 preprocess + compose.append
            for path_rel, shift in zip(entry.get("paths", []), entry.get("shifts", [])):
                src = lib_root / path_rel
                if not src.exists():
                    log.warning(f"  [{i}] 素材不存在: {src}")
                    stats["errors"] += 1
                    continue
                prep_path = prep_dir / f"{_hash_path(src)}_{src.name}"
                try:
                    preprocess(src, prep_path, shift=shift, params=params)
                except Exception as e:
                    log.exception(f"  [{i}] preprocess 失败 {src.name}: {e}")
                    stats["errors"] += 1
                    continue
                try:
                    sub_doc = Document(str(prep_path))
                    composer.append(sub_doc)
                    stats["merged_materials"] += 1
                except Exception as e:
                    log.exception(f"  [{i}] compose 失败 {src.name}: {e}")
                    stats["errors"] += 1
            continue

        if status == "UNMATCHED":
            add_heading(doc, title, level)
            add_placeholder_paragraph(doc, f"[缺失：{title}——wiki 无匹配卡片，请人工处理]")
            stats["inserted_headings"] += 1
            stats["inserted_placeholders"] += 1
            continue

    out_path.parent.mkdir(parents=True, exist_ok=True)
    composer.save(str(out_path))

    log.info(f"merged → {out_path} ({out_path.stat().st_size / 1024 / 1024:.1f} MB)")
    log.info(f"stats: {stats}")
    return stats


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--template", type=Path, required=True)
    ap.add_argument("--plan", type=Path, required=True)
    ap.add_argument("--lib", type=Path, required=True, help="素材库根")
    ap.add_argument("--params", type=Path, default=None)
    ap.add_argument("--prep-dir", type=Path, default=Path("/tmp/bid_prep"))
    ap.add_argument("--out", type=Path, required=True)
    args = ap.parse_args()

    plan = json.loads(args.plan.read_text(encoding="utf-8"))
    params = {}
    if args.params and args.params.exists():
        params = json.loads(args.params.read_text(encoding="utf-8"))

    merge(args.template, plan, args.lib, params, args.prep_dir, args.out)


if __name__ == "__main__":
    main()
