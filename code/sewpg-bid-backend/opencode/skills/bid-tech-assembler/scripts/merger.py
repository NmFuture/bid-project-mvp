#!/usr/bin/env python3
"""
merger v2（方案 B）：按 assembly_plan.json 合并素材到技术标母版。

关键特性：
- 使用 docxcompose 做 section 级合并（媒体/样式自动处理）
- 手插 toc heading：text="{chapter_no}  {title}"，样式=Heading N
- 素材内部 Heading 按父章节相对映射为 3/4 级标题，保留正文结构
- 素材首 Heading 若匹配 toc_title 则去重（物理移除）
- 前言段特殊：style=Heading 1，text="前言  投标说明函"
- 封面段：attach_mode=cover 的卡片优先 compose，放在文档最前
- STRUCTURAL / NEEDS_REVIEW 只插 Heading（+ 占位）

依赖：docxcompose、python-docx

用法：
    python3 merger.py \\
        --template templates/技术投标母版模板.docx \\
        --plan /tmp/assembly_plan.json \\
        --lib /Users/wlb/Downloads/技术标/素材库 \\
        --params /tmp/project_params.json \\
        --prep-dir /tmp/bid_prep \\
        --out /tmp/bid_merged.docx
"""

from __future__ import annotations

import argparse
import hashlib
import json
import logging
import os
import re
import sys
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docxcompose.composer import Composer
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy

sys.path.insert(0, str(Path(__file__).resolve().parent))
from preprocess import preprocess
from numbering_fixer import inject_prefix_to_headings, remap_material_headings_to_navigation, strip_numPr_from_body, strip_numPr_from_heading_styles
from docx_style_pruner import prune_unused_styles


def _path_exists(path: Path) -> bool:
    return os.path.exists(os.fspath(path))


def _isolate_section(doc) -> bool:
    """在 doc body **尾部 sectPr 之前**插入一个 nextPage section break，其 sectPr =
    自身尾部 sectPr 的克隆。

    核心原则：素材横版 → 合入后横版；素材竖版 → 合入后竖版；横竖拼接 → 分页。

    docxcompose 合并多份单 sectPr 素材时，每份 sub 的尾部 sectPr 只是替换 master 尾部，
    **不会自动封住** sub 自己的段落 section。后来的素材带 landscape 时就把前面单 sectPr
    素材的段落吞进同一节变横版。本函数在每份素材 body 末尾（tail sectPr 之前）显式插入
    一个带"自身 sectPr 克隆 + type=nextPage"的段落，让 sub 的段落所在 section 被自己
    的 sectPr 明确封住，后面素材无论什么 orientation 都不污染。
    """
    body = doc.element.body
    tail_sect = body.find(qn("w:sectPr"))
    if tail_sect is None:
        return False
    clone = deepcopy(tail_sect)
    # type=nextPage：orientation 变化时会强制换页，符合"横竖拼接分页"原则
    t = clone.find(qn("w:type"))
    if t is None:
        t = OxmlElement("w:type")
        clone.insert(0, t)
    t.set(qn("w:val"), "nextPage")
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pPr.append(clone)
    p.append(pPr)
    # 插到 body 最后一个 sectPr 之前（tail_sect 之前）
    tail_idx = list(body).index(tail_sect)
    body.insert(tail_idx, p)
    return True

logging.basicConfig(level=logging.INFO, format="[merger] %(message)s")
log = logging.getLogger("merger")


class BatchComposer(Composer):
    """将全量 ID 扫描延迟到所有素材追加完成之后。"""

    def __init__(self, doc):
        self._defer_global_ids = True
        super().__init__(doc)

    def renumber_bookmarks(self):
        if not self._defer_global_ids:
            super().renumber_bookmarks()

    def renumber_docpr_ids(self):
        if not self._defer_global_ids:
            super().renumber_docpr_ids()

    def renumber_nvpicpr_ids(self):
        if not self._defer_global_ids:
            super().renumber_nvpicpr_ids()

    def finalize_global_ids(self):
        if not self._defer_global_ids:
            return
        self._defer_global_ids = False
        super().renumber_bookmarks()
        super().renumber_docpr_ids()
        super().renumber_nvpicpr_ids()


# ---------- 工具 ----------

def _hash_path(p: Path) -> str:
    return hashlib.md5(str(p).encode("utf-8")).hexdigest()[:10]


def _master_doc_clear_body(doc) -> None:
    """清空母版 body 所有子元素（保留最后 sectPr）。"""
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is sect_pr:
            continue
        body.remove(child)


def _add_heading(doc, text: str, level: int):
    level = max(1, min(6, level))
    try:
        p = doc.add_paragraph(text, style=f"Heading {level}")
    except KeyError:
        p = doc.add_paragraph(text)
    p_pr = p._p.find(qn("w:pPr"))
    if p_pr is not None:
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is not None:
            p_pr.remove(num_pr)
    return p


def _add_body_paragraph(doc, text: str):
    return doc.add_paragraph(text)


# ---------- 父章节号推导 ----------

def _parent_of(chapter_no_flat: str) -> str:
    """'4.1' → '4'；'5.8.1' → '5.8'；'1' → ''。"""
    if not chapter_no_flat:
        return ""
    parts = chapter_no_flat.split(".")
    if len(parts) <= 1:
        return ""
    return ".".join(parts[:-1])


def _normalize_title_for_dedup(s: str) -> str:
    """语义归一化：把"附表1"/"表1"/"附 xxx"/"xxx（如有）" 等同义变体收敛到同一 key。

    策略（按顺序应用）：
    1. Unicode NFKC 归一（全半角统一）
    2. 去标准可变修饰：
       - 前缀"附"字（"附表X"→"表X"，"附 xxx"→"xxx"）
       - 末尾括号修饰（"（如有）"/"（可选）"/"（待定）"/"（北区）"等）
       - 前缀"★"/"*" 星号
    3. 去所有空白、标点
    """
    if not s:
        return ""
    import unicodedata
    s = unicodedata.normalize("NFKC", s).strip()

    # 前缀"附"字（可选跟"表"）：剥一次
    s = re.sub(r"^附(表)?(?=\S)", r"\1", s)
    # 前缀星号/修饰符
    s = re.sub(r"^[*★☆※◆◎○●]+\s*", "", s)
    # 末尾括号修饰（中文/英文括号）
    s = re.sub(r"[（(][^）)]{1,6}[）)]\s*$", "", s).strip()

    # 去所有空白、标点（注意"附"可能再出现在中间，不再动）
    s = re.sub(r"[\s:：，,。.！!？?\\\-_/'\"“”‘’·（）()【】\[\]]+", "", s)
    return s


def _collect_heading_titles(doc, target_set: set) -> None:
    """扫描 doc 内所有 Heading 段落，把 pure title（去掉 inject 前缀后）加入 target_set。"""
    import re as _re
    # inject 后 text 形如 "5.8.7.1  xxx" 或 "4.1  xxx"；也可能是未被 inject 的原文
    prefix_strip = _re.compile(r"^\s*\d+(?:\.\d+){0,6}[\.\s　]+")
    for para in doc.paragraphs:
        st = para.style.name if para.style else ""
        if not (st.startswith("Heading") or st.startswith("标题")):
            continue
        t = (para.text or "").strip()
        if not t:
            continue
        pure = prefix_strip.sub("", t).strip()
        if pure:
            target_set.add(_normalize_title_for_dedup(pure))


def _derive_parent_chapter(entry: dict) -> str:
    """根据 toc entry 推导父章节号（用于 inject_prefix_to_headings）。

    - chapter_no_flat 非空 → 直接用（"1.1" / "5.8.7"）
    - "第N章" → chapter_no_flat 应已填阿拉伯数字（"1" / "2" ...）
    - 前言 / 附 → 空（退化为 chapter 根）
    """
    flat = entry.get("chapter_no_flat") or ""
    return flat


# ---------- 主合并 ----------

def merge(
    template_path: Path,
    plan: list[dict],
    lib_root: Path,
    params: dict,
    prep_dir: Path,
    out_path: Path,
) -> dict:
    os.makedirs(os.fspath(prep_dir), exist_ok=True)

    # 打开母版并清空 body（只保留 sectPr）
    master_doc = Document(str(template_path))
    _master_doc_clear_body(master_doc)
    prune_unused_styles(master_doc)
    composer = BatchComposer(master_doc)

    stats = {
        "cover_merged": 0,
        "inserted_headings": 0,
        "inserted_placeholders": 0,
        "merged_materials": 0,
        "skipped_oos": 0,
        "inject_skipped_first": 0,
        "errors": 0,
    }
    warning_counts = {
        "MATERIAL_MISSING": 0,
        "MATERIAL_MERGE_FAILED": 0,
        "DIRECTORY_WITHOUT_MATERIAL": 0,
    }

    # Step 0: 如果 plan 首条是 COVER，先处理它
    cover_entries = [e for e in plan if e.get("status") == "COVER"]
    non_cover = [e for e in plan if e.get("status") != "COVER"]

    for cover in cover_entries:
        for path_rel in cover.get("paths", []):
            src = lib_root / path_rel
            if not _path_exists(src):
                log.warning(f"封面素材不存在: {src}")
                stats["errors"] += 1
                warning_counts["MATERIAL_MISSING"] += 1
                continue
            prep = prep_dir / f"cover_{_hash_path(src)}_{src.name}"
            try:
                preprocess(src, prep, params)
                sub = Document(str(prep))
                composer.append(sub)
                stats["cover_merged"] += 1
                log.info(f"合并封面: {src.name}")
            except Exception as e:
                log.exception(f"封面合并失败 {src.name}: {e}")
                stats["errors"] += 1
                warning_counts["MATERIAL_MERGE_FAILED"] += 1

    in_scope = [e for e in non_cover if e.get("status") != "OUT_OF_SCOPE"]
    log.info(f"in-scope entries: {len(in_scope)}/{len(non_cover)}")

    chapter_master_prefixes = {
        str(entry.get("chapter_no_flat") or "").strip()
        for entry in non_cover
        if str(entry.get("coverage_role") or entry.get("coverageRole") or "").strip() == "chapter_master"
        and entry.get("status") in ("MATCHED", "ADAPTED")
        and entry.get("paths")
        and str(entry.get("chapter_no_flat") or "").strip()
    }

    # 状态：记录"当前父章节号"。appendix (附 xxx) 素材用这个作 parent
    current_parent_chapter = ""
    # 每个父章节下已经由素材 inject 生成的子 heading pure titles（去前缀去空白）
    # 用于让后续同父下的 NEEDS_REVIEW 子条目判"已被父素材覆盖"，避免重复占位
    seen_titles_by_parent: dict[str, set[str]] = {}
    # Parent chapter -> raw child title -> official S2 heading metadata. Whole
    # chapter materials often contain headings for their children (e.g. 4.1-4.7).
    # Keep only those matched child titles as navigation headings; demote all
    # other material-internal headings.
    toc_children_by_parent: dict[str, dict[str, dict]] = {}
    for child in non_cover:
        if child.get("status") == "OUT_OF_SCOPE" or child.get("is_appendix"):
            continue
        child_flat = child.get("chapter_no_flat") or ""
        parent = _parent_of(child_flat)
        if not parent:
            continue
        title_key = str(child.get("title") or "").strip()
        if not title_key:
            continue
        toc_children_by_parent.setdefault(parent, {})[title_key] = {
            "chapter_no": child.get("chapter_no") or "",
            "title": child.get("title") or "",
            "level": child.get("level") or 1,
        }

    # Step 1: 遍历正文 plan
    for i, entry in enumerate(non_cover):
        status = entry["status"]
        level = entry["level"]
        title = entry["title"]
        chapter_no = entry.get("chapter_no") or ""
        chapter_no_flat = str(entry.get("chapter_no_flat") or "").strip()

        if any(chapter_no_flat.startswith(f"{prefix}.") for prefix in chapter_master_prefixes):
            stats.setdefault("superseded", 0)
            stats["superseded"] += 1
            continue

        if status == "OUT_OF_SCOPE":
            stats["skipped_oos"] += 1
            continue

        # toc heading text
        if entry.get("is_preface"):
            heading_text = f"前言  {title}" if title != "前言" else "前言"
            heading_level = 1
        elif entry.get("is_appendix"):
            # 附 开头：不手插 heading（素材内部的 heading 就是"附 xxx"），但不去重
            heading_text = None
            heading_level = level
        else:
            heading_text = f"{chapter_no}  {title}" if chapter_no else title
            heading_level = level

        if status == "STRUCTURAL":
            if heading_text:
                _add_heading(master_doc, heading_text, heading_level)
                stats["inserted_headings"] += 1
            continue

        if status == "NEEDS_REVIEW":
            # 若父章节的素材内部 heading 树已包含本条 title（去前缀归一化后），
            # 视为已被父素材覆盖，不再手插空 heading + 占位（避免与素材子节重复）
            parent_prefix = _parent_of(entry.get("chapter_no_flat"))
            seen = seen_titles_by_parent.get(parent_prefix, set())
            if _normalize_title_for_dedup(title) in seen:
                stats.setdefault("superseded", 0)
                stats["superseded"] += 1
                continue
            if heading_text:
                _add_heading(master_doc, heading_text, heading_level)
                stats["inserted_headings"] += 1
            _add_body_paragraph(master_doc, f"[待填写：{title}——招标/模板新增，请补素材]")
            stats["inserted_placeholders"] += 1
            continue

        if status == "UNMATCHED":
            # 若父素材 inject 后已包含本条 title（归一化后），视为已被父素材覆盖
            parent_prefix = _parent_of(entry.get("chapter_no_flat"))
            seen = seen_titles_by_parent.get(parent_prefix, set())
            if _normalize_title_for_dedup(title) in seen:
                stats.setdefault("superseded", 0)
                stats["superseded"] += 1
                continue
            if heading_text:
                _add_heading(master_doc, heading_text, heading_level)
                stats["inserted_headings"] += 1
            _add_body_paragraph(master_doc, f"[缺失：{title}——wiki 无匹配卡片，请人工处理]")
            stats["inserted_placeholders"] += 1
            continue

        if status in ("MATCHED", "ADAPTED"):
            # 手插 toc heading（method: 素材首 Heading 会由 inject 去重）
            if heading_text:
                _add_heading(master_doc, heading_text, heading_level)
                stats["inserted_headings"] += 1

            # 更新"当前父章节号"状态（非 appendix 条目更新；appendix 沿用）
            if entry.get("is_appendix"):
                # appendix 素材用当前父章节号作 parent，inject 编号接续
                parent_chapter = current_parent_chapter
            else:
                parent_chapter = _derive_parent_chapter(entry)
                # 对非 appendix 条目：若它自身是 toc 的章节（如 5.7），更新 current_parent
                if parent_chapter:
                    current_parent_chapter = parent_chapter

            merged_for_entry = 0
            for path_idx, path_rel in enumerate(entry.get("paths", [])):
                src = lib_root / path_rel
                if not _path_exists(src):
                    log.warning(f"  [{i}] 素材不存在: {src}")
                    stats["errors"] += 1
                    warning_counts["MATERIAL_MISSING"] += 1
                    continue

                size_mb = os.path.getsize(os.fspath(src)) / 1024 / 1024
                if size_mb > 100:
                    log.warning(f"  [{i}] 合并超大素材 ({size_mb:.0f} MB): {src.name}")

                prep = prep_dir / f"{_hash_path(src)}_{src.name}"
                try:
                    preprocess(src, prep, params)
                except Exception as e:
                    log.exception(f"  [{i}] preprocess 失败 {src.name}: {e}")
                    stats["errors"] += 1
                    warning_counts["MATERIAL_MERGE_FAILED"] += 1
                    continue

                # 关键：S2 TOC 条目由 merger 手插为真正的导航 Heading。
                # 素材内部 Heading 不能再整体降级；这里按当前 TOC 父级
                # 相对映射到最终 3/4 级，避免最后 cleaner 从正文里猜层级。
                # remove_first_if_match 只对本条 entry 首份成功素材启用（避免前序失败
                # 漏去重，也避免叠加场景下多份素材首 heading 都被删）。
                try:
                    sub_doc = Document(str(prep))
                    material_heading_titles: set[str] = set()
                    _collect_heading_titles(sub_doc, material_heading_titles)
                    is_chapter_master = str(
                        entry.get("coverage_role") or entry.get("coverageRole") or ""
                    ).strip() == "chapter_master"
                    if is_chapter_master:
                        inject_stats = inject_prefix_to_headings(
                            sub_doc,
                            parent_chapter,
                            toc_title=title,
                            skip_first_if_match=(merged_for_entry == 0),
                        )
                        if inject_stats["skipped_first"]:
                            stats["inject_skipped_first"] += 1
                        stats.setdefault("material_headings_injected", 0)
                        stats["material_headings_injected"] += inject_stats["injected"]
                    else:
                        remap_stats = remap_material_headings_to_navigation(
                            sub_doc,
                            toc_title=title,
                            remove_first_if_match=(merged_for_entry == 0),
                            keep_heading_map=toc_children_by_parent.get(parent_chapter, {}),
                            parent_level=heading_level,
                            max_target_level=4,
                        )
                        if remap_stats["skipped_first"]:
                            stats["inject_skipped_first"] += 1
                        stats.setdefault("material_headings_remapped", 0)
                        stats["material_headings_remapped"] += remap_stats["remapped"]
                        stats.setdefault("material_bold_subheadings_promoted", 0)
                        stats["material_bold_subheadings_promoted"] += remap_stats["bold_subheadings"]
                        stats.setdefault("material_headings_kept", 0)
                        stats["material_headings_kept"] += remap_stats.get("kept", 0)
                        stats.setdefault("material_headings_demoted", 0)
                        stats["material_headings_demoted"] += remap_stats.get("demoted", 0)

                    # 保存处理后的副本
                    inj_path = prep_dir / f"inj_{_hash_path(src)}_{src.name}"
                    sub_doc.save(str(inj_path))
                    # 重新打开用于 compose；先隔离 section，避免 landscape 串扰
                    sub_doc2 = Document(str(inj_path))
                    _isolate_section(sub_doc2)
                    composer.append(sub_doc2)
                    seen_titles_by_parent.setdefault(parent_chapter, set()).update(
                        material_heading_titles
                    )
                    stats["merged_materials"] += 1
                    merged_for_entry += 1
                except Exception as e:
                    log.exception(f"  [{i}] compose 失败 {src.name}: {e}")
                    stats["errors"] += 1
                    warning_counts["MATERIAL_MERGE_FAILED"] += 1

            if merged_for_entry == 0:
                _add_body_paragraph(master_doc, f"[缺失：{title}——没有可用素材，请补充后重试]")
                stats["inserted_placeholders"] += 1
                warning_counts["DIRECTORY_WITHOUT_MATERIAL"] += 1

    # Save
    strip_numPr_from_heading_styles(master_doc)
    strip_numPr_from_body(master_doc)
    os.makedirs(os.fspath(out_path.parent), exist_ok=True)
    composer.finalize_global_ids()
    composer.save(str(out_path))

    warning_messages = {
        "MATERIAL_MISSING": "{count} 份素材不存在，已跳过",
        "MATERIAL_MERGE_FAILED": "{count} 份素材处理或合并失败，已跳过",
        "DIRECTORY_WITHOUT_MATERIAL": "{count} 个目录节点没有可用素材，已保留标题并插入占位提示",
    }
    stats["warnings"] = [
        {
            "code": code,
            "message": warning_messages[code].format(count=count),
            "count": count,
        }
        for code, count in warning_counts.items()
        if count
    ]

    log.info(f"merged → {out_path} ({os.path.getsize(os.fspath(out_path)) / 1024 / 1024:.1f} MB)")
    log.info(f"stats: {stats}")
    return stats


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--template", type=Path, required=True)
    ap.add_argument("--plan", type=Path, required=True)
    ap.add_argument("--lib", type=Path, required=True)
    ap.add_argument("--params", type=Path, default=None)
    ap.add_argument("--prep-dir", type=Path, default=Path("/tmp/bid_prep"))
    ap.add_argument("--out", type=Path, required=True)
    ap.add_argument("--result", type=Path, default=None)
    args = ap.parse_args()

    plan = json.loads(args.plan.read_text(encoding="utf-8"))
    params = {}
    if args.params and args.params.exists():
        params = json.loads(args.params.read_text(encoding="utf-8"))

    result = merge(args.template, plan, args.lib, params, args.prep_dir, args.out)
    if args.result:
        args.result.parent.mkdir(parents=True, exist_ok=True)
        args.result.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
