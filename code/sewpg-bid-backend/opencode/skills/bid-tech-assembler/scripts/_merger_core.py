"""
大附件合并器

处理投标文件中的大型 docx 附件（资质证书、业绩表、专利列表等），
将它们合并到主文档中，保留原始格式和嵌入图片。

支持两种合并模式：
1. append  - 将整个附件文档追加到主文档末尾
2. section - 在主文档指定位置（通过标记段落）插入附件内容
"""

import copy
import io
import logging
import os
import re
import shutil
import warnings
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple
from lxml import etree
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.image.exceptions import UnrecognizedImageError
from docx.parts.image import ImagePart
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

# 商务标专属 CATEGORY_TITLE_ALIASES 桩；技术标不依赖分类别名
CATEGORY_TITLE_ALIASES: Dict[str, List[str]] = {}

logger = logging.getLogger(__name__)
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"


class SmartNumberingFixer:
    """
    智能标题编号修正器（移植自旧版 bid-assembler v2.0）

    自动将模板文件内部的编号（如 3.1, 3.2）转换为
    正确的章节前缀编号（如 7.3.5.1, 7.3.5.2）
    """

    # 章节前缀映射：附件类别 → 目标编号前缀
    # 注意：多级编号由 Word 自动管理，此处前缀仅用于内部小节编号修正
    CHAPTER_PREFIXES = {
        '服务能力': '',
        '质量专题': '',
        '制造基地': '',
    }

    # 章节标题修正映射（不带数字前缀，编号由 Word 多级列表自动生成）
    CORRECT_TITLES = {
        '供货保障': '供货保障专题',
        '资质': '商务部分摘要表',
        '公司介绍': '公司基本情况介绍',
        '制造基地': '生产能力介绍',
        '服务能力': '服务能力介绍',
        '质量专题': '质量管理介绍',
        '获奖情况': '获奖情况',
        '专利情况': '专利情况',
        '业绩_陆上': '近年完成的6.25MW及以上容量等级风电机组业绩表（陆上）',
        '业绩_海上': '近年完成的6.25MW及以上容量等级风电机组业绩表（海上）',
        '感谢信': '商业信誉（感谢信）',
    }

    # 业绩编号基数（多级编号由 Word 自动管理，不再硬编码章节前缀）
    PERF_BASE_NUMBERS = {}

    FRAMEWORK_HEADING_CATEGORIES = {
        '公司介绍',
        '制造基地',
        '服务能力',
        '质量专题',
    }

    @staticmethod
    def _normalize_heading_text(text: str) -> str:
        text = text.strip().lstrip("#").strip()
        text = re.sub(
            r"^\s*(?:(?:\d+\.){1,5}\d*\s+|"
            r"第[一二三四五六七八九十百千万\d]+[章节篇部分卷]\s*|"
            r"[一二三四五六七八九十百千万\d]+[、.．]\s*)",
            "",
            text,
        ).strip()
        text = re.sub(r"[（(【\[].*?[）)】\]]", "", text)
        text = re.sub(r"[\s:：·\\-_/\"“”‘’]+", "", text)
        return text.lower()

    @classmethod
    def _looks_like_heading(cls, text: str, category: str, correct_title: str) -> bool:
        normalized = cls._normalize_heading_text(text)
        correct_norm = cls._normalize_heading_text(correct_title)
        category_norm = cls._normalize_heading_text(category)

        if normalized and (normalized == correct_norm or normalized == category_norm):
            return True
        if normalized and correct_norm and (normalized in correct_norm or correct_norm in normalized):
            return True

        return (
            len(text) <= 40
            and not re.search(r"[。；;！？!?]", text)
            and (
                bool(re.match(r"^(?:\d+(?:\.\d+)*\s+|第[一二三四五六七八九十百千万\d]+[章节篇部分卷]|[一二三四五六七八九十百千万\d]+[、.．])", text))
                or text.endswith(("介绍", "说明", "承诺书", "概述", "简介"))
            )
        )

    @classmethod
    def fix_heading(cls, doc: Document, category: str) -> int:
        """修正文档首个标题为标准章节标题，返回修正数"""
        correct_title = cls.CORRECT_TITLES.get(category)
        if not correct_title:
            return 0
        for para in doc.paragraphs[:5]:
            text = para.text.strip()
            if not text:
                continue
            if cls._looks_like_heading(text, category, correct_title):
                para.clear()
                run = para.add_run(correct_title)
                run.bold = True
                return 1
            if len(text) > 40 or re.search(r"[。；;！？!?]", text):
                break
        return 0

    @classmethod
    def fix_internal_numbering(cls, doc: Document, category: str) -> int:
        """修正文档内部编号（3.1 → 7.3.5.1），返回修正数"""
        prefix = cls.CHAPTER_PREFIXES.get(category)
        if not prefix:
            return 0
        count = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            m = re.match(r'^(\d+(?:\.\d+)*)\s+(.+)$', text)
            if m:
                old_number, title = m.groups()
                if old_number == prefix or old_number.startswith(f"{prefix}."):
                    continue
                parts = old_number.split('.')
                if len(parts) == 1:
                    new_number = f"{prefix}.{old_number}"
                else:
                    remaining = '.'.join(parts[1:])
                    new_number = f"{prefix}.{remaining}"
                if old_number != new_number:
                    has_bold = any(r.bold for r in para.runs) if para.runs else False
                    para.clear()
                    run = para.add_run(f"{new_number} {title}")
                    if has_bold:
                        run.bold = True
                    count += 1
        return count

    @classmethod
    def add_performance_numbering(cls, doc: Document, category: str) -> int:
        """为业绩文件中的项目添加四级编号，返回编号数"""
        base = cls.PERF_BASE_NUMBERS.get(category)
        if not base:
            return 0
        counter = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            m = re.match(r'^(\d+)\s+(.+项目.*)$', text)
            if m and para.runs and any(r.bold for r in para.runs):
                counter += 1
                _, project_name = m.groups()
                para.clear()
                run = para.add_run(f"{base}.{counter} {project_name}")
                run.bold = True
        return counter

    @classmethod
    def process_document(cls, doc: Document, category: str) -> Dict[str, int]:
        """对文档执行全部智能修正，返回各项修正计数"""
        stats = {
            'heading_fixes': cls.fix_heading(doc, category),
            'numbering_fixes': cls.fix_internal_numbering(doc, category),
            'perf_numbering': cls.add_performance_numbering(doc, category),
        }
        total = sum(stats.values())
        if total > 0:
            logger.info(f"  智能修正: 标题{stats['heading_fixes']}处, "
                        f"编号{stats['numbering_fixes']}处, "
                        f"业绩{stats['perf_numbering']}项")
        return stats


@dataclass
class MergeRecord:
    """单次合并记录"""
    category: str
    filename: str
    source_paragraphs: int = 0
    source_tables: int = 0
    added_paragraphs: int = 0
    added_tables: int = 0
    complete: bool = True
    numbering_stats: Dict[str, int] = field(default_factory=dict)


class DocxMerger:
    """
    高性能 docx 文件合并器

    通过直接操作 docx 的 XML 和 ZIP 结构来合并文档，
    避免 python-docx 高层 API 丢失格式和图片的问题。
    """

    def __init__(self, master_path: str):
        self.master_path = master_path
        self.master_doc = Document(master_path)
        self._media_counter = self._count_existing_media()
        self._merged_count = 0
        self.merge_records: List[MergeRecord] = []
        self._next_drawing_id = 1_000_000
        self._next_bookmark_id = 1_000_000

    def _count_existing_media(self) -> int:
        """统计主文档已有的 media 文件数量"""
        try:
            with zipfile.ZipFile(self.master_path, 'r') as z:
                media_files = [n for n in z.namelist() if n.startswith('word/media/')]
                return len(media_files)
        except Exception:
            return 0

    @staticmethod
    def _remove_element(element) -> None:
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    @staticmethod
    def _iter_insertable_children(body) -> List[etree._Element]:
        """收集 body 中可插入的内容元素，并保留 body 级 sectPr 的页面方向。

        OOXML 中页面方向存储在两个位置：
        1. 段落级 sectPr（w:pPr/w:sectPr）—— 定义该段落所结束 section 的页面设置
        2. Body 级 sectPr（w:body/w:sectPr）—— 定义最后一个 section 的页面设置

        之前只收集 p/tbl/sdt，丢弃了 body 级 sectPr。如果子文档末尾
        需要从横版切回竖版（常见于专利表格→证书页、业绩表→详情页），
        这个切换就丢失了。

        修复策略：检测 body 级 sectPr 与最后一个段落级 sectPr 的页面方向
        是否不同；若不同，在末尾追加一个空段落并注入 body 级 sectPr，
        确保方向切换被保留。
        """
        elements = []
        body_sect_pr = None
        for child in body:
            tag = etree.QName(child.tag).localname
            if tag in ('p', 'tbl', 'sdt'):
                elements.append(child)
            elif tag == 'sectPr':
                body_sect_pr = child

        if body_sect_pr is not None and elements:
            # 找最后一个段落级 sectPr 的页面方向
            last_para_orient = None
            for elem in reversed(elements):
                if etree.QName(elem.tag).localname != 'p':
                    continue
                ppr = elem.find(f"{{{W_NS}}}pPr")
                if ppr is not None:
                    inner_sect = ppr.find(f"{{{W_NS}}}sectPr")
                    if inner_sect is not None:
                        pg = inner_sect.find(f"{{{W_NS}}}pgSz")
                        if pg is not None:
                            iw = int(pg.get(f"{{{W_NS}}}w", "0"))
                            ih = int(pg.get(f"{{{W_NS}}}h", "0"))
                            last_para_orient = "landscape" if iw > ih else "portrait"
                        break

            # body 级 sectPr 的页面方向
            body_orient = None
            bpg = body_sect_pr.find(f"{{{W_NS}}}pgSz")
            if bpg is not None:
                bw = int(bpg.get(f"{{{W_NS}}}w", "0"))
                bh = int(bpg.get(f"{{{W_NS}}}h", "0"))
                body_orient = "landscape" if bw > bh else "portrait"

            # 方向不同时，注入 body 级 sectPr 到一个空段落，保留方向切换
            if (last_para_orient and body_orient
                    and last_para_orient != body_orient):
                from copy import deepcopy
                sect_copy = deepcopy(body_sect_pr)
                # 移除 header/footer 引用，防止污染主文档
                for ch in list(sect_copy):
                    if etree.QName(ch.tag).localname in (
                            "headerReference", "footerReference"):
                        sect_copy.remove(ch)
                # 设为连续节分隔（不额外分页）
                sect_type = sect_copy.find(f"{{{W_NS}}}type")
                if sect_type is None:
                    sect_type = etree.SubElement(sect_copy, f"{{{W_NS}}}type")
                sect_type.set(f"{{{W_NS}}}val", "continuous")
                # 构造空段落并嵌入 sectPr
                empty_p = etree.Element(f"{{{W_NS}}}p")
                p_pr = etree.SubElement(empty_p, f"{{{W_NS}}}pPr")
                p_pr.append(sect_copy)
                elements.append(empty_p)

        return elements

    @staticmethod
    def _collect_relationship_rids(elements: List[etree._Element]) -> Set[str]:
        refs: Set[str] = set()
        ignored_elements = {"headerReference", "footerReference"}
        for element in elements:
            for node in element.iter():
                if etree.QName(node.tag).localname in ignored_elements:
                    continue
                for attr_name, attr_value in node.attrib.items():
                    qname = etree.QName(attr_name)
                    if qname.namespace != R_NS:
                        continue
                    if isinstance(attr_value, str) and attr_value.startswith("rId"):
                        refs.add(attr_value)
        return refs

    @staticmethod
    def _sanitize_section_properties(sect_pr: etree._Element) -> None:
        """保留页面设置，但移除会污染主文档 header/footer 的关系引用。"""
        removable = {"headerReference", "footerReference"}
        for child in list(sect_pr):
            if etree.QName(child.tag).localname in removable:
                sect_pr.remove(child)

    def _find_marker_element(self, marker: str):
        body = self.master_doc.element.body
        for para in body.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
            text = ''.join(
                node.text or ''
                for node in para.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            )
            if marker in text:
                return para
        return None

    def has_marker(self, marker: str) -> bool:
        return self._find_marker_element(marker) is not None

    @staticmethod
    def _paragraph_has_page_break(paragraph) -> bool:
        has_break = any(
            br.get(f"{{{W_NS}}}type") == "page"
            for br in paragraph._element.iter(f"{{{W_NS}}}br")
        )
        has_rendered_break = any(
            True for _ in paragraph._element.iter(f"{{{W_NS}}}lastRenderedPageBreak")
        )
        return has_break or has_rendered_break

    @classmethod
    def _is_page_break_only_paragraph(cls, paragraph) -> bool:
        return not paragraph.text.strip() and cls._paragraph_has_page_break(paragraph)

    @staticmethod
    def _paragraph_has_section_break(paragraph) -> bool:
        ppr = paragraph._element.find(qn('w:pPr'))
        if ppr is None:
            return False
        return ppr.find(qn('w:sectPr')) is not None

    @classmethod
    def _has_sensitive_page_layout(cls, sub_doc: Document) -> bool:
        try:
            if len(sub_doc.sections) > 1:
                return True
            if any(section.page_width > section.page_height for section in sub_doc.sections):
                return True
        except Exception:
            pass

        return any(cls._paragraph_has_section_break(paragraph) for paragraph in sub_doc.paragraphs)

    @staticmethod
    def _clean_material_basename(doc_path: str) -> str:
        basename = os.path.splitext(os.path.basename(doc_path))[0]
        basename = re.sub(r'[\-_]\d{6,8}$', '', basename)
        basename = re.sub(r'\s*[\(（]\d+[\)）]$', '', basename)
        return basename

    @staticmethod
    def _normalize_title(text: str) -> str:
        text = text.strip().lstrip("#").strip()
        prev = None
        while prev != text:
            prev = text
            text = re.sub(
                r"^\s*(?:(?:\d+\.){1,5}\d*\s+|"
                r"第[一二三四五六七八九十百千万\d]+[章节篇部分卷]\s*|"
                r"[一二三四五六七八九十百千万\d]+[、.．]\s*)",
                "",
                text,
            ).strip()
        text = re.sub(r"[（(【\[].*?[）)】\]]", "", text)
        text = re.sub(r"[\s:：·\\-_/\"“”‘’]+", "", text)
        text = text.replace("及以上", "")
        text = text.replace("最新版", "")
        text = text.replace("最新", "")
        text = text.replace("更新", "")
        return text.lower()

    def _build_material_title_candidates(self, doc_path: str, category: str) -> List[str]:
        candidates = [
            self._clean_material_basename(doc_path),
            category,
            SmartNumberingFixer.CORRECT_TITLES.get(category, ""),
        ]
        candidates.extend(CATEGORY_TITLE_ALIASES.get(category, []))
        return [candidate for candidate in candidates if candidate]

    def _is_redundant_material_title(self, text: str, candidates: List[str]) -> bool:
        text_norm = self._normalize_title(text)
        if not text_norm or len(text_norm) > 30:
            return False

        for candidate in candidates:
            cand_norm = self._normalize_title(candidate)
            if not cand_norm:
                continue
            if text_norm == cand_norm:
                return True
            if text_norm in cand_norm or cand_norm in text_norm:
                shorter = min(len(text_norm), len(cand_norm))
                longer = max(len(text_norm), len(cand_norm))
                if shorter >= 2 and shorter / longer >= 0.6:
                    return True
        return False

    def _is_minor_leading_heading(self, text: str) -> bool:
        stripped = text.strip()
        if not stripped:
            return False

        compact = re.sub(r"^\d+(?:\.\d+)*\s*", "", stripped).strip("、.．:： ")
        compact = compact.strip("“”\"'")
        return (
            len(compact) <= 20
            and not re.search(r"[。；;！？!?]", compact)
            and (
                bool(re.match(r"^\d+(?:\.\d+)*\s*", stripped))
                or stripped.endswith(("介绍", "说明", "承诺书", "概述", "简介"))
            )
        )

    def _is_inline_noise_heading(self, text: str) -> bool:
        stripped = text.strip()
        if not stripped or re.search(r"[。；;！？!?：:]", stripped):
            return False

        compact = re.sub(r"^\d+(?:\.\d+){2,}\s*", "", stripped).strip("、.．:： ")
        compact = compact.strip("“”\"'")

        if re.match(r"^\d+(?:\.\d+){2,}\s*", stripped) and len(compact) <= 30:
            return True

        return (
            len(compact) <= 12
            and (
                compact.endswith(("外景", "内景", "图片", "照片", "示意图", "现场图"))
                or compact.startswith(("图", "表"))
            )
        )

    @staticmethod
    def _is_inline_section_heading(text: str) -> bool:
        stripped = text.strip()
        if (
            not stripped
            or len(stripped) > 40
            or re.search(r"[。；;！？!?：:]", stripped)
        ):
            return False

        prefix_re = (
            r"^(?:"
            r"[（(]?\d+[)）]\s*|"
            r"\d+[)）]\s*|"
            r"[一二三四五六七八九十百千万]+[、.．]\s*|"
            r"(?:\d+\.){1,5}\d*\s+"
            r")"
        )
        if not re.match(prefix_re, stripped):
            return False

        compact = re.sub(prefix_re, "", stripped).strip("、.．:： ")
        compact = compact.strip("“”\"'")
        if not compact or len(compact) > 30:
            return False

        # 普通章节中若是明显的项目/合同名称则保留，短小节标题继续清洗。
        if len(compact) > 18 and re.search(r"(项目|合同|采购|风电|场址|基地|MW|兆瓦|EPC)", compact):
            return False

        return True

    @staticmethod
    def _is_performance_subheading(text: str) -> bool:
        stripped = text.strip()
        if not stripped or len(stripped) > 40 or re.search(r"[。；;！？!?]", stripped):
            return False

        return bool(
            re.match(r"^(?:陆上|海上)\s*\d+(?:\.\d+)?MW业绩$", stripped)
            or re.match(r"^\d+(?:\.\d+)?MW及以上机型合同业绩表$", stripped)
            or stripped in {
                "海上业绩情况",
                "陆上业绩情况",
                "合同业绩情况",
                "合同业绩表",
            }
        )

    @staticmethod
    def _is_performance_project_label(text: str) -> bool:
        stripped = text.strip()
        if not stripped or len(stripped) > 140 or re.search(r"[。；;！？!?]", stripped):
            return False

        project_tail = r".*(项目|合同|采购|风电场|场址|基地|示范工程|供货合同).*$"
        return bool(
            re.match(rf"^\d+\s+{project_tail}", stripped)
            or re.match(rf"^(?:\d+\.){{2,}}\d+\s+{project_tail}", stripped)
            or re.match(rf"^[（(]\d+[)）]\s*{project_tail}", stripped)
        )

    def _cleanup_material_doc(self, sub_doc: Document, doc_path: str, category: str) -> Dict[str, int]:
        """
        素材清洗入口。

        当前优先级是“全程保留原始 Word 版式”，因此对外部 docx
        素材不再做结构性删段，只保留 SmartNumberingFixer 的原位编号修正。
        """
        stats = {
            "page_break_paragraphs": 0,
            "leading_blank_paragraphs": 0,
            "leading_title_paragraphs": 0,
            "leading_minor_headings": 0,
            "inline_noise_headings": 0,
            "inline_section_headings": 0,
            "performance_project_labels": 0,
        }

        logger.info("  素材清洗: 保留原始 Word 结构，跳过结构性删段")
        return stats

    def _prepare_sub_doc(self, doc_path: str, sub_doc: Document, category: str = "") -> Dict[str, int]:
        """按合并场景对素材进行预处理，不影响招标原模板。"""
        stats: Dict[str, int] = {}
        if not category:
            return stats

        stats = SmartNumberingFixer.process_document(sub_doc, category)
        self._cleanup_material_doc(sub_doc, doc_path, category)
        return stats

    def _sanitize_insert_element(self, element: etree._Element) -> None:
        """
        保留页面方向等节属性，但移除会污染主文档页眉页脚的引用。
        同时对 bookmark / drawing ID 去重，防止 OOXML schema 验证报 w:id 冲突。
        修正 pPr 子元素顺序，确保 pStyle 始终为第一个子元素。
        """
        for sect_pr in list(element.iter(f"{{{W_NS}}}sectPr")):
            DocxMerger._sanitize_section_properties(sect_pr)
        self._renumber_drawing_ids(element)
        self._renumber_bookmark_ids(element)
        self._fix_pstyle_order(element)

    @staticmethod
    def _detect_material_headings(sub_doc: Document) -> Dict[int, int]:
        """
        扫描素材文档，返回 {body子元素索引: 目标Heading级别} 映射。

        素材内容合并在框架 H2 下面，因此：
          素材 H1/H2/macro/自定义标题 → Heading 3（ilvl=2，编号如 9.3.1）
          素材 H3+                   → Heading 4（ilvl=3，编号如 9.3.1.1）

        索引对应 _iter_insertable_children 的顺序。
        """
        # 先构建 style_id → style_name 映射（解决 id="2" → "Heading 2" 等情况）
        id_to_name: Dict[str, str] = {}
        for s in sub_doc.styles:
            if s.style_id:
                id_to_name[s.style_id] = s.name or ''

        remap: Dict[int, int] = {}
        body = sub_doc.element.body
        insertable_idx = 0

        for child in body:
            tag = etree.QName(child.tag).localname
            if tag not in ('p', 'tbl', 'sdt'):
                continue

            if tag == 'p':
                pPr = child.find(f"{{{W_NS}}}pPr")
                pStyle = pPr.find(f"{{{W_NS}}}pStyle") if pPr is not None else None
                style_id = pStyle.get(f"{{{W_NS}}}val", '') if pStyle is not None else ''
                style_name = id_to_name.get(style_id, style_id)
                style_lower = style_name.lower()

                source_level = 0
                is_heading = False

                if style_lower.startswith('heading ') or style_lower.startswith('heading'):
                    digits = ''.join(c for c in style_lower if c.isdigit())
                    if digits:
                        source_level = int(digits)
                        is_heading = True
                elif style_lower in ('macro', 'toc heading', 'title') or '标题' in style_name:
                    is_heading = True
                    source_level = 1

                if is_heading:
                    target_level = 3 if source_level <= 2 else 4
                    remap[insertable_idx] = target_level

            insertable_idx += 1

        return remap

    def _set_element_heading_style(self, element: etree._Element, heading_level: int) -> None:
        """在 XML 级别将段落的 pStyle 设为 Heading N（使用主文档的样式 ID）。"""
        # 从主文档查找实际 style ID（可能是 "5"、"Heading3" 等）
        target_name = f"Heading {heading_level}"
        style_id = None
        for s in self.master_doc.styles:
            if s.name == target_name:
                style_id = s.style_id
                break
        if not style_id:
            style_id = f"Heading{heading_level}"  # fallback
        pPr = element.find(f"{{{W_NS}}}pPr")
        if pPr is None:
            pPr = etree.SubElement(element, f"{{{W_NS}}}pPr")
        pStyle = pPr.find(f"{{{W_NS}}}pStyle")
        if pStyle is None:
            pStyle = etree.SubElement(pPr, f"{{{W_NS}}}pStyle")
            # 确保 pStyle 是 pPr 的第一个子元素
            pPr.remove(pStyle)
            pPr.insert(0, pStyle)
        pStyle.set(f"{{{W_NS}}}val", style_id)

    @staticmethod
    def _strip_outline_level(element: etree._Element) -> None:
        """移除段落的 outlineLvl，防止非 Heading 段落出现在 Word 导航窗格。"""
        tag = etree.QName(element.tag).localname
        if tag != 'p':
            return
        pPr = element.find(f"{{{W_NS}}}pPr")
        if pPr is None:
            return
        outline = pPr.find(f"{{{W_NS}}}outlineLvl")
        if outline is not None:
            pPr.remove(outline)

    @staticmethod
    def _fix_pstyle_order(element: etree._Element) -> None:
        """确保所有 pPr 中 pStyle 是第一个子元素（OOXML schema 要求）。"""
        for pPr in element.iter(f"{{{W_NS}}}pPr"):
            pstyle = pPr.find(f"{{{W_NS}}}pStyle")
            if pstyle is not None and list(pPr)[0] is not pstyle:
                pPr.remove(pstyle)
                pPr.insert(0, pstyle)

    def _renumber_drawing_ids(self, element: etree._Element) -> None:
        for xpath in (
            f".//{{{WP_NS}}}docPr",
            f".//{{{PIC_NS}}}cNvPr",
            f".//{{{A_NS}}}cNvPr",
        ):
            for node in element.iterfind(xpath):
                node.set("id", str(self._next_drawing_id))
                self._next_drawing_id += 1

    def _renumber_bookmark_ids(self, element: etree._Element) -> None:
        """重新分配 bookmark ID，防止多文档合并后 w:id 冲突。"""
        # bookmarkStart 和 bookmarkEnd 共享同一 w:id，需成对重映射
        id_map: Dict[str, str] = {}
        w_id = f"{{{W_NS}}}id"
        for tag_local in ("bookmarkStart", "bookmarkEnd"):
            for node in element.iter(f"{{{W_NS}}}{tag_local}"):
                old_id = node.get(w_id)
                if old_id is None:
                    continue
                if old_id not in id_map:
                    id_map[old_id] = str(self._next_bookmark_id)
                    self._next_bookmark_id += 1
                node.set(w_id, id_map[old_id])

    def append_document(self, doc_path: str, section_title: Optional[str] = None,
                        add_page_break: bool = True,
                        category: str = "") -> bool:
        """
        将一个 docx 文件追加到主文档末尾

        Args:
            doc_path: 要追加的文档路径
            section_title: 可选的章节标题（追加在内容之前）
            add_page_break: 是否在追加前插入分页符
            category: 附件类别（用于智能编号修正）
        """
        if not os.path.exists(doc_path):
            logger.error(f"文件不存在: {doc_path}")
            return False

        file_size = os.path.getsize(doc_path) / (1024 * 1024)
        logger.info(f"正在合并: {os.path.basename(doc_path)} ({file_size:.1f}MB)")

        try:
            sub_doc = Document(doc_path)
        except Exception as e:
            logger.error(f"无法打开文档 {doc_path}: {e}")
            return False

        num_stats = self._prepare_sub_doc(doc_path, sub_doc, category)

        # 记录素材中需要重映射的标题段落索引及目标级别
        heading_remap = {}
        if category:
            heading_remap = self._detect_material_headings(sub_doc)

        # 合并前统计
        before_paras = len(self.master_doc.paragraphs)
        before_tables = len(self.master_doc.tables)
        source_paras = len(sub_doc.paragraphs)
        source_tables = len(sub_doc.tables)

        if add_page_break:
            self.master_doc.add_page_break()

        if section_title:
            heading = self.master_doc.add_heading(section_title, level=2)
            for run in heading.runs:
                run.font.name = '等线 Light'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                from docx.oxml.ns import qn as _qn
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.find(_qn('w:rFonts'))
                if rFonts is None:
                    rFonts = run._element.makeelement(_qn('w:rFonts'), {})
                    rPr.insert(0, rFonts)
                rFonts.set(_qn('w:eastAsia'), '等线 Light')

        source_elements = self._iter_insertable_children(sub_doc.element.body)
        rId_map = self._copy_relationships(
            doc_path,
            sub_doc,
            required_rids=self._collect_relationship_rids(source_elements),
        )

        for idx, child in enumerate(source_elements):
            new_elem = copy.deepcopy(child)
            self._remap_relationship_rids(new_elem, rId_map)
            self._sanitize_insert_element(new_elem)
            # 素材标题样式重映射（XML 级别直接改 pStyle）
            if idx in heading_remap:
                self._set_element_heading_style(new_elem, heading_remap[idx])
            elif category:
                self._strip_outline_level(new_elem)
            self.master_doc.element.body.append(new_elem)

        # 合并后验证
        after_paras = len(self.master_doc.paragraphs)
        after_tables = len(self.master_doc.tables)
        added_paras = after_paras - before_paras
        added_tables = after_tables - before_tables
        complete = (added_paras >= source_paras * 0.95 and added_tables >= source_tables)

        record = MergeRecord(
            category=category or os.path.basename(doc_path),
            filename=os.path.basename(doc_path),
            source_paragraphs=source_paras,
            source_tables=source_tables,
            added_paragraphs=added_paras,
            added_tables=added_tables,
            complete=complete,
            numbering_stats=num_stats,
        )
        self.merge_records.append(record)

        if not complete:
            logger.warning(f"  ⚠ 合并可能不完整: 段落 {added_paras}/{source_paras}, "
                           f"表格 {added_tables}/{source_tables}")
        else:
            logger.info(f"  ✓ {os.path.basename(doc_path)}: "
                        f"段落 {added_paras}, 表格 {added_tables}")

        self._merged_count += 1
        return True

    def _copy_relationships(
        self,
        doc_path: str,
        sub_doc: Document,
        required_rids: Optional[Set[str]] = None,
    ) -> Dict[str, str]:
        """
        从子文档复制被插入 XML 实际引用到的关系，返回 rId 映射关系。

        图片关系使用新的 image part；其余内部关系（如 SmartArt/diagram）
        直接挂接到主文档 package，让 python-docx 在保存时一并写出依赖部件。
        """
        rId_map = {}

        try:
            sub_rels = sub_doc.part.rels
        except Exception:
            return rId_map

        selected_rids = set(sub_rels.keys()) if required_rids is None else set(required_rids)
        selected_reltypes = {
            sub_rels[rid].reltype
            for rid in selected_rids
            if rid in sub_rels
        }
        for family_keyword in ("diagram", "chart"):
            if any(family_keyword in reltype for reltype in selected_reltypes):
                for rid, rel in sub_rels.items():
                    if family_keyword in rel.reltype:
                        selected_rids.add(rid)

        def sort_key(rel_id: str):
            suffix = rel_id[3:] if rel_id.startswith("rId") else rel_id
            return (0, int(suffix)) if str(suffix).isdigit() else (1, rel_id)

        for rId in sorted(selected_rids, key=sort_key):
            rel = sub_rels.get(rId)
            if rel is None:
                logger.warning("关系不存在，已跳过: %s (%s)", rId, os.path.basename(doc_path))
                continue

            if rel.reltype == RT.IMAGE:
                try:
                    if rel.is_external:
                        logger.debug(
                            "跳过外链图片关系: %s (%s)",
                            rId,
                            os.path.basename(doc_path),
                        )
                        continue

                    image_part = rel.target_part
                    try:
                        image_stream = io.BytesIO(image_part.blob)
                        new_rId, _ = self.master_doc.part.get_or_add_image(image_stream)
                    except UnrecognizedImageError:
                        fallback_part = self._get_or_add_raw_image_part(
                            image_part.blob,
                            image_part.content_type,
                            image_part.sha1,
                        )
                        new_rId = self.master_doc.part.relate_to(fallback_part, RT.IMAGE)
                    rId_map[rId] = new_rId
                except Exception as e:
                    logger.warning(
                        "复制图片失败 (rId=%s, type=%s): %s",
                        rId,
                        type(e).__name__,
                        e,
                    )
                continue

            try:
                if rel.is_external:
                    new_rId = self.master_doc.part.relate_to(
                        rel.target_ref,
                        rel.reltype,
                        is_external=True,
                    )
                else:
                    new_rId = self.master_doc.part.relate_to(rel.target_part, rel.reltype)
                rId_map[rId] = new_rId
            except Exception as e:
                logger.warning(
                    "复制关系失败 (rId=%s, reltype=%s): %s",
                    rId,
                    rel.reltype,
                    e,
                )

        logger.debug(f"复制了 {len(rId_map)} 个关系")
        return rId_map

    def _get_or_add_raw_image_part(
        self,
        blob: bytes,
        content_type: str,
        sha1: str,
    ) -> ImagePart:
        """
        为 python-docx 不能解析的图片格式（如 EMF/WMF）创建/复用 image part。
        """
        image_parts = self.master_doc.part.package.image_parts
        for existing_part in image_parts:
            if existing_part.sha1 == sha1:
                return existing_part

        ext = self._content_type_to_ext(content_type).lstrip('.')
        partname = image_parts._next_image_partname(ext)
        new_part = ImagePart(partname, content_type, blob)
        image_parts.append(new_part)
        return new_part

    def _remap_relationship_rids(self, element: etree._Element, rId_map: Dict[str, str]):
        """递归替换 XML 元素中的关系 rId 引用。"""
        if not rId_map:
            return

        for node in element.iter():
            for attr_name, attr_value in list(node.attrib.items()):
                qname = etree.QName(attr_name)
                if qname.namespace != R_NS:
                    continue
                if not isinstance(attr_value, str) or not attr_value.startswith("rId"):
                    continue

                if attr_value in rId_map:
                    node.set(attr_name, rId_map[attr_value])
                else:
                    node.attrib.pop(attr_name, None)

    @staticmethod
    def _content_type_to_ext(content_type: str) -> str:
        mapping = {
            'image/png': '.png',
            'image/jpeg': '.jpg',
            'image/gif': '.gif',
            'image/bmp': '.bmp',
            'image/tiff': '.tiff',
            'image/x-emf': '.emf',
            'image/x-wmf': '.wmf',
        }
        return mapping.get(content_type, '.png')

    def insert_at_marker(self, doc_path: str, marker: str,
                         remove_marker: bool = True,
                         category: str = "") -> bool:
        """
        在主文档中查找标记文本，在该位置插入附件内容

        Args:
            doc_path: 要插入的文档路径
            marker: 标记文本，例如 "{{业绩表}}"
            remove_marker: 是否移除标记段落
            category: 素材类别；仅素材库文档会做标题/空白清洗
        """
        if not os.path.exists(doc_path):
            logger.error(f"文件不存在: {doc_path}")
            return False

        marker_elem = self._find_marker_element(marker)

        if marker_elem is None:
            logger.warning(f"未找到标记: {marker}")
            return False

        logger.info(f"在标记 '{marker}' 处插入: {os.path.basename(doc_path)}")

        try:
            sub_doc = Document(doc_path)
        except Exception as e:
            logger.error(f"无法打开文档 {doc_path}: {e}")
            return False

        num_stats = self._prepare_sub_doc(doc_path, sub_doc, category)

        # 素材内部标题样式映射
        heading_remap = {}
        if category:
            heading_remap = self._detect_material_headings(sub_doc)

        source_elements = self._iter_insertable_children(sub_doc.element.body)
        source_paras = len(sub_doc.paragraphs)
        source_tables = len(sub_doc.tables)
        before_paras = len(self.master_doc.paragraphs)
        before_tables = len(self.master_doc.tables)

        rId_map = self._copy_relationships(
            doc_path,
            sub_doc,
            required_rids=self._collect_relationship_rids(source_elements),
        )

        insert_elements = []
        for idx, child in enumerate(source_elements):
            new_elem = copy.deepcopy(child)
            self._remap_relationship_rids(new_elem, rId_map)
            self._sanitize_insert_element(new_elem)
            if idx in heading_remap:
                self._set_element_heading_style(new_elem, heading_remap[idx])
            elif category:
                # 非标题段落：清除 outlineLvl，防止在导航窗格中显示为无编号项
                self._strip_outline_level(new_elem)
            insert_elements.append(new_elem)

        parent = marker_elem.getparent()
        idx = list(parent).index(marker_elem)
        for i, elem in enumerate(insert_elements):
            parent.insert(idx + 1 + i, elem)

        if remove_marker:
            parent.remove(marker_elem)

        after_paras = len(self.master_doc.paragraphs)
        after_tables = len(self.master_doc.tables)
        removed_marker_paras = 1 if remove_marker else 0
        added_paras = (after_paras - before_paras) + removed_marker_paras
        added_tables = after_tables - before_tables
        complete = (added_paras >= source_paras * 0.95 and added_tables >= source_tables)

        self.merge_records.append(
            MergeRecord(
                category=category or os.path.basename(doc_path),
                filename=os.path.basename(doc_path),
                source_paragraphs=source_paras,
                source_tables=source_tables,
                added_paragraphs=added_paras,
                added_tables=added_tables,
                complete=complete,
                numbering_stats=num_stats,
            )
        )

        self._merged_count += 1
        return True

    def cleanup_placeholder_markers(self) -> int:
        patterns = ("{{附件:", "{{招标模板:")
        removed = 0
        body = self.master_doc.element.body
        for para in list(body.iter(f"{{{W_NS}}}p")):
            text = "".join(
                node.text or ""
                for node in para.iter(f"{{{W_NS}}}t")
            ).strip()
            if text and any(pattern in text for pattern in patterns):
                parent = para.getparent()
                if parent is not None:
                    parent.remove(para)
                    removed += 1
        return removed

    def get_validation_report(self) -> Dict:
        """生成完整性验证报告"""
        total = len(self.merge_records)
        successful = sum(1 for r in self.merge_records if r.complete)
        total_paras = len(self.master_doc.paragraphs)
        total_tables = len(self.master_doc.tables)
        total_numbering = sum(
            sum(r.numbering_stats.values()) for r in self.merge_records
        )
        total_perf = sum(
            r.numbering_stats.get('perf_numbering', 0) for r in self.merge_records
        )

        return {
            'total_merges': total,
            'successful': successful,
            'total_paragraphs': total_paras,
            'total_tables': total_tables,
            'numbering_fixes': total_numbering,
            'perf_numbered': total_perf,
            'records': self.merge_records,
        }

    def print_validation_report(self):
        """打印完整性验证报告"""
        rpt = self.get_validation_report()
        logger.info('=' * 60)
        logger.info('【完整性统计】')
        logger.info(f"  合并文件数: {rpt['successful']}/{rpt['total_merges']} ✓")
        logger.info(f"  总段落数: {rpt['total_paragraphs']:,}")
        logger.info(f"  总表格数: {rpt['total_tables']:,}")
        if rpt['numbering_fixes'] > 0:
            logger.info(f"  智能编号修正: {rpt['numbering_fixes']} 处")
        if rpt['perf_numbered'] > 0:
            logger.info(f"  业绩项目编号: {rpt['perf_numbered']} 个项目")

        warnings = [r for r in rpt['records'] if not r.complete]
        if warnings:
            logger.info('【警告】')
            for r in warnings:
                logger.info(f"  ⚠ {r.filename}: 段落 {r.added_paragraphs}/{r.source_paragraphs}")
        logger.info('=' * 60)

    def save(self, output_path: Optional[str] = None):
        """保存合并后的文档"""
        path = output_path or self.master_path
        os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
        with warnings.catch_warnings():
            warnings.filterwarnings("ignore", message=r"Duplicate name: .*", category=UserWarning)
            self.master_doc.save(path)
        deduped = self._dedupe_zip_entries(path)
        if deduped:
            logger.info("ZIP 去重: 清理 %d 个重复条目", len(deduped))
        file_size = os.path.getsize(path) / (1024 * 1024)
        logger.info(f"已保存: {path} ({file_size:.1f}MB, 合并了 {self._merged_count} 个文件)")
        self.print_validation_report()

    @staticmethod
    def _dedupe_zip_entries(docx_path: str) -> List[str]:
        """清理 docx ZIP 中的重复条目，保留最后一次写入的版本。"""
        with zipfile.ZipFile(docx_path, "r") as src:
            infos = src.infolist()
            last_index = {info.filename: idx for idx, info in enumerate(infos)}
            duplicates = sorted(
                {
                    info.filename
                    for idx, info in enumerate(infos)
                    if last_index[info.filename] != idx
                }
            )
            if not duplicates:
                return []

            temp_path = f"{docx_path}.dedup"
            with zipfile.ZipFile(temp_path, "w") as dst:
                dst.comment = src.comment
                for idx, info in enumerate(infos):
                    if last_index[info.filename] != idx:
                        continue
                    data = src.read(info)
                    new_info = zipfile.ZipInfo(info.filename, date_time=info.date_time)
                    new_info.compress_type = info.compress_type
                    new_info.comment = info.comment
                    new_info.extra = info.extra
                    new_info.create_system = info.create_system
                    new_info.create_version = info.create_version
                    new_info.extract_version = info.extract_version
                    new_info.flag_bits = info.flag_bits
                    new_info.volume = info.volume
                    new_info.internal_attr = info.internal_attr
                    new_info.external_attr = info.external_attr
                    dst.writestr(new_info, data)

        os.replace(temp_path, docx_path)
        return duplicates


class AttachmentManager:
    """
    附件管理器

    管理投标文件模板目录中的所有大附件，提供按章节合并的高级接口。
    """

    # 标准章节映射：章节标识 → 可能的文件名模式
    SECTION_PATTERNS = {
        "资质": ["资质", "认证", "ISO"],
        "业绩_陆上": ["陆上", "6MW"],
        "业绩_海上": ["海上", "8MW"],
        "公司介绍": ["公司基本情况", "公司简介", "企业介绍"],
        "制造基地": ["制造基地"],
        "服务能力": ["服务能力"],
        "感谢信": ["感谢信"],
        "获奖情况": ["获奖"],
        "专利情况": ["专利"],
        "质量专题": ["质量专题", "质量管理"],
        "装机统计": ["装机", "统计"],
    }

    def __init__(self, template_dir: str):
        self.template_dir = template_dir
        self.attachments = self._scan_attachments()

    def _scan_attachments(self) -> Dict[str, List[str]]:
        """扫描模板目录，按章节分类所有附件文件"""
        result = {}
        if not os.path.isdir(self.template_dir):
            logger.warning(f"模板目录不存在: {self.template_dir}")
            return result

        docx_files = []
        png_files = []
        for f in os.listdir(self.template_dir):
            full_path = os.path.join(self.template_dir, f)
            if f.startswith('~$'):
                continue
            if f.lower().endswith('.docx'):
                docx_files.append(full_path)
            elif f.lower().endswith(('.png', '.jpg', '.jpeg')):
                png_files.append(full_path)

        for section_key, patterns in self.SECTION_PATTERNS.items():
            matched = []
            for fp in docx_files:
                fname = os.path.basename(fp)
                if any(p in fname for p in patterns):
                    matched.append(fp)
            for fp in png_files:
                fname = os.path.basename(fp)
                if any(p in fname for p in patterns):
                    matched.append(fp)
            if matched:
                result[section_key] = matched

        # 未分类的文件
        all_matched = set()
        for files in result.values():
            all_matched.update(files)

        unmatched = [f for f in docx_files if f not in all_matched]
        if unmatched:
            result["其他"] = unmatched

        return result

    def summary(self) -> str:
        """返回附件扫描摘要"""
        lines = [f"模板目录: {self.template_dir}", ""]
        for section, files in self.attachments.items():
            total_size = sum(os.path.getsize(f) for f in files) / (1024 * 1024)
            lines.append(f"  [{section}] {len(files)} 个文件, {total_size:.1f}MB")
            for f in files:
                fsize = os.path.getsize(f) / (1024 * 1024)
                lines.append(f"    - {os.path.basename(f)} ({fsize:.1f}MB)")
        return "\n".join(lines)

    def merge_all(self, merger: DocxMerger, sections: Optional[List[str]] = None,
                  section_titles: Optional[Dict[str, str]] = None):
        """
        将所有（或指定）章节的附件合并到主文档

        Args:
            merger: DocxMerger 实例
            sections: 要合并的章节列表，None 表示全部
            section_titles: 章节标题映射，如 {"资质": "7.1 资质证明文件"}
        """
        titles = section_titles or {}
        target_sections = sections or list(self.attachments.keys())

        for section in target_sections:
            files = self.attachments.get(section, [])
            if not files:
                logger.warning(f"章节 '{section}' 没有找到附件文件")
                continue

            title = titles.get(section, "")
            for fp in files:
                if fp.lower().endswith('.docx'):
                    merger.append_document(
                        fp,
                        section_title=title if title else None,
                        add_page_break=True
                    )
                    title = ""  # 只给第一个文件加标题
                elif fp.lower().endswith(('.png', '.jpg', '.jpeg')):
                    self._insert_image(merger, fp)

    @staticmethod
    def _insert_image(merger: DocxMerger, image_path: str):
        """向主文档追加一张图片"""
        try:
            para = merger.master_doc.add_paragraph()
            run = para.add_run()
            run.add_picture(image_path, width=Inches(5.5))
        except Exception as e:
            logger.warning(f"插入图片失败 {image_path}: {e}")


def merge_bid_attachments(
    master_docx: str,
    template_dir: str,
    output_path: str,
    sections: Optional[List[str]] = None,
    section_titles: Optional[Dict[str, str]] = None,
) -> str:
    """
    高级接口：将模板目录中的附件合并到主文档

    Args:
        master_docx: 主文档路径（AI生成的投标文件框架）
        template_dir: 模板素材目录
        output_path: 输出路径
        sections: 要合并的章节，None 表示全部
        section_titles: 章节标题

    Returns:
        输出文件路径
    """
    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')

    mgr = AttachmentManager(template_dir)
    logger.info(f"\n{mgr.summary()}")

    merger = DocxMerger(master_docx)
    mgr.merge_all(merger, sections=sections, section_titles=section_titles)
    merger.save(output_path)

    return output_path


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 4:
        print("用法: python merger.py <主文档.docx> <模板目录> <输出路径.docx>")
        print("示例: python merger.py 商务投标文件.docx 投标文件模板/ output/完整版.docx")
        sys.exit(1)

    merge_bid_attachments(sys.argv[1], sys.argv[2], sys.argv[3])
