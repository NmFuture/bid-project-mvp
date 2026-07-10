#!/usr/bin/env python3
"""附表 AI 填写金标评分（对比真实中标技术附表逐格计分）。

用法：python eval_fill_vs_golden.py <fill_results.json> <真实中标技术附表.docx> [出参json]
fill_results.json 每条：{appendixId, number(附表编号), title, blankDocx, outputFile}。
正式标书留在本地，不入库。

对齐口径演进——v4：v3（剔除 S1 越界表）基础上加行对齐（行数不一致时按行键序列对齐，消除投标人增删行导致的整段错位假阴性）。

v2 修了跨节错配；v3 修分母灌水——S1 切片系统性"错位一格"，每张附表 docx
末尾都带着下一张附表的标题段+第一张表。这些越界表不属于本附表，不该拿去和
本节答案比。判据自包含：blankDocx 内第二个（编号不同的）附表标题之后的表格
全部剔除，只对"自己的表"计分。
"""
from __future__ import annotations

import json
import re
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

HEADING_RE = re.compile(r"^(技术附表[A-Za-z]|附表[A-Za-z0-9]+(?:\.[0-9]+)*(?:-[0-9]+)?)\s*(.*)$")


def norm_text(value: Any) -> str:
    text = str(value or "").strip()
    return re.sub(r"\s+", "", text)


def norm_number(value: str) -> str:
    return re.sub(r"\s+", "", str(value or "")).rstrip(".")


def iter_block_items(doc: Document):
    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield ("P", Paragraph(child, doc))
        elif child.tag.endswith("}tbl"):
            yield ("T", Table(child, doc))


def build_heading_index(human_docx: Path) -> tuple[list[dict[str, Any]], Document]:
    doc = Document(str(human_docx))
    items = list(iter_block_items(doc))
    headings = []
    for idx, (kind, item) in enumerate(items):
        if kind != "P":
            continue
        text = item.text.strip()
        match = HEADING_RE.match(text)
        if match:
            headings.append({"blockIndex": idx, "number": norm_number(match.group(1)), "title": text})
    return items, headings


def tables_in_section(items: list, headings: list[dict[str, Any]], number: str) -> list[list[list[str]]]:
    target_number = norm_number(number)
    for h_idx, heading in enumerate(headings):
        if heading["number"] == target_number:
            start = heading["blockIndex"]
            end = headings[h_idx + 1]["blockIndex"] if h_idx + 1 < len(headings) else len(items)
            tables = []
            for kind, item in items[start + 1 : end]:
                if kind == "T":
                    rows = [[cell.text for cell in row.cells] for row in item.rows]
                    tables.append(rows)
            return tables
    return []


def table_cells(path: Path) -> list[list[list[str]]]:
    if not path.exists():
        return []
    doc = Document(str(path))
    return [[[cell.text for cell in row.cells] for row in table.rows] for table in doc.tables]


def own_table_count(blank_docx: Path) -> int:
    """blankDocx 里属于本附表的表数：第二个（编号不同的）附表标题之前的表。

    S1 切片系统性带上下一张附表的标题+首表；同编号重复标题（"续"）不算边界。
    没有第二个标题（如重建型 docx）→ 全部算本表。
    """
    if not blank_docx.exists():
        return 0
    doc = Document(str(blank_docx))
    first_number = ""
    count = 0
    for kind, item in iter_block_items(doc):
        if kind == "P":
            match = HEADING_RE.match(item.text.strip())
            if match:
                number = norm_number(match.group(1))
                if not first_number:
                    first_number = number
                elif number != first_number:
                    break
        else:
            count += 1
    return count


def table_static_tokens(table: list[list[str]]) -> set[str]:
    tokens: set[str] = set()
    for row in table:
        for cell in row:
            text = norm_text(cell)
            if len(text) >= 2:
                tokens.add(text)
    return tokens


def table_similarity(a: list[list[str]], b: list[list[str]]) -> float:
    a_tokens = table_static_tokens(a)
    b_tokens = table_static_tokens(b)
    if not a_tokens or not b_tokens:
        return 0.0
    overlap = len(a_tokens & b_tokens)
    return overlap / max(1, min(len(a_tokens), len(b_tokens)))


def best_table_in_section(blank_table: list[list[str]], section_tables: list[list[list[str]]]) -> tuple[int, float]:
    best_idx, best_score = -1, 0.0
    for idx, human_table in enumerate(section_tables):
        score = table_similarity(blank_table, human_table)
        if score > best_score:
            best_idx, best_score = idx, score
    return best_idx, best_score


def cell_at(table: list[list[str]], row: int, col: int) -> str:
    if row < 0 or row >= len(table):
        return ""
    if col < 0 or col >= len(table[row]):
        return ""
    return table[row][col]


def parse_simple_number(value: str) -> float | None:
    text = str(value or "").strip()
    if not text:
        return None
    normalized = text.replace(",", "").replace("，", "")
    if not re.fullmatch(r"~?\s*[-+]?\d+(?:\.\d+)?\s*", normalized):
        return None
    try:
        return float(normalized.replace("~", "").strip())
    except ValueError:
        return None


def tolerant_equal(left: str, right: str) -> bool:
    left_norm, right_norm = norm_text(left), norm_text(right)
    if left_norm == right_norm:
        return True
    left_num, right_num = parse_simple_number(str(left)), parse_simple_number(str(right))
    if left_num is None or right_num is None:
        return False
    tolerance = max(0.05, abs(right_num) * 0.0005)
    return abs(left_num - right_num) <= tolerance


# ---- v5 等价类（2026-07-09 差异审计确认：82/642 wrongCells 属"填对被判错"）----

_EMPTY_RESPONSE_CLASS = {"/", "\\", "—", "-", "－", "无", "n/a", "na", "不适用"}


def _norm_number_format(text: str) -> str:
    """千分位、结尾 .0、空格、全角逗号归一（10,000==10000、8.0rpm==8rpm）。"""
    t = norm_text(text).replace("，", ",")
    t = re.sub(r"(?<=\d),(?=\d{3})", "", t)
    t = re.sub(r"(\d+)\.0+(?=\D|$)", r"\1", t)
    t = re.sub(r"(\d+\.\d*?)0+(?=\D|$)", r"\1", t)
    return t.lower()


def _strip_separators(text: str) -> str:
    """换行/分隔符归一：docx 换行被 clean 成 " / " 的等价（B.1.1/D.5 实测 20 格）。"""
    return re.sub(r"[\s/\\|，,;；]+", "", str(text or ""))


def tolerant_equal_v5(left: str, right: str) -> bool:
    if tolerant_equal(left, right):
        return True
    l, r = norm_text(left).lower(), norm_text(right).lower()
    if l in _EMPTY_RESPONSE_CLASS and r in _EMPTY_RESPONSE_CLASS:
        return True
    if _norm_number_format(left) == _norm_number_format(right):
        return True
    ls, rs = _strip_separators(left), _strip_separators(right)
    if ls and ls == rs:
        return True
    return False


def _row_key(row: list[str]) -> str:
    for cell in row:
        text = norm_text(cell)
        if len(text) >= 2 and not re.fullmatch(r"[0-9.\-/]+", text):
            return text
    return norm_text(row[0]) if row else ""


def _row_mapping(human_table: list[list[str]], other_table: list[list[str]]) -> list[int | None]:
    """human 每行在 other（blank/agent）里的对应行号；行数一致时纯行号直连。"""
    if len(other_table) == len(human_table):
        return list(range(len(human_table)))
    import difflib

    human_keys = [_row_key(row) for row in human_table]
    other_keys = [_row_key(row) for row in other_table]
    mapping: list[int | None] = [None] * len(human_table)
    matcher = difflib.SequenceMatcher(None, human_keys, other_keys, autojunk=False)
    for block in matcher.get_matching_blocks():
        for offset in range(block.size):
            mapping[block.a + offset] = block.b + offset
    return mapping


def score_outputs(results: list[dict[str, Any]], human_docx: Path, min_align: float = 0.30) -> dict[str, Any]:
    items, headings = build_heading_index(human_docx)
    full_human_tables = table_cells(human_docx)  # 兜底用
    rows: list[dict[str, Any]] = []
    totals = {
        "targetTables": 0, "alignedTables": 0, "noSectionFound": 0, "bleedTablesDropped": 0,
        "humanChangedCells": 0, "agentCorrectCells": 0,
        "agentTolerantCorrectCells": 0, "agentTouchedCells": 0, "sameRowAttemptCells": 0,
    }
    for result in results:
        blank_path = Path(str(result.get("blankDocx") or ""))
        blank = table_cells(blank_path)
        agent = table_cells(Path(str(result.get("outputFile") or "")))
        number = result.get("number") or ""
        own = own_table_count(blank_path)
        totals["bleedTablesDropped"] += max(0, len(blank) - own)
        blank = blank[:own]
        section_tables = tables_in_section(items, headings, number)
        used_fallback = not section_tables
        pool = section_tables if section_tables else full_human_tables
        for table_idx, blank_table in enumerate(blank):
            totals["targetTables"] += 1
            agent_table = agent[table_idx] if table_idx < len(agent) else []
            human_idx, align_score = best_table_in_section(blank_table, pool)
            if human_idx < 0 or align_score < min_align:
                if used_fallback:
                    totals["noSectionFound"] += 1
                rows.append({
                    "appendixId": result.get("appendixId"), "number": number, "title": result.get("title"),
                    "table": table_idx, "aligned": False, "alignScore": round(align_score, 3),
                    "usedFallback": used_fallback,
                })
                continue
            totals["alignedTables"] += 1
            human_table = pool[human_idx]
            changed = correct = touched = same_row_attempt = tolerant_correct = tolerant_v5_correct = 0
            # v4 行对齐：投标人增删行会使纯行号比对从错位点起全部误判
            #（C.2 删 1 行 → 后半张表全算错）。行数不一致时按"行键"（首个非纯数字
            # 静态格文本）做序列对齐；行数一致保持纯行号（无噪声）。
            blank_of = _row_mapping(human_table, blank_table)
            agent_of = _row_mapping(human_table, agent_table)
            max_rows = len(human_table)
            wrong_cells = []
            for row_idx in range(max_rows):
                blank_row = blank_table[blank_of[row_idx]] if blank_of[row_idx] is not None else []
                agent_row = agent_table[agent_of[row_idx]] if agent_of[row_idx] is not None else []
                human_row = human_table[row_idx]
                max_cols = max(len(blank_row), len(human_row), len(agent_row))
                agent_row_values = {norm_text(c) for c in agent_row}
                for col_idx in range(max_cols):
                    blank_text = norm_text(blank_row[col_idx]) if col_idx < len(blank_row) else ""
                    human_text = norm_text(human_row[col_idx]) if col_idx < len(human_row) else ""
                    agent_raw = agent_row[col_idx] if col_idx < len(agent_row) else ""
                    agent_text = norm_text(agent_raw)
                    if human_text and human_text != blank_text:
                        changed += 1
                        is_correct = agent_text == human_text
                        is_tolerant = tolerant_equal(agent_raw, human_row[col_idx])
                        if is_correct:
                            correct += 1
                        if is_tolerant:
                            tolerant_correct += 1
                        elif human_text in agent_row_values:
                            same_row_attempt += 1
                        if tolerant_equal_v5(agent_raw, human_row[col_idx]):
                            tolerant_v5_correct += 1
                        if not is_tolerant:
                            wrong_cells.append({
                                "row": row_idx, "col": col_idx,
                                "expected": human_row[col_idx][:60],
                                "got": str(agent_raw)[:60],
                                "blank": (blank_row[col_idx] if col_idx < len(blank_row) else "")[:40],
                            })
                    if agent_text and agent_text != blank_text:
                        touched += 1
            totals["humanChangedCells"] += changed
            totals["agentCorrectCells"] += correct
            totals["agentTolerantCorrectCells"] += tolerant_correct
            totals["agentTolerantV5CorrectCells"] = totals.get("agentTolerantV5CorrectCells", 0) + tolerant_v5_correct
            totals["agentTouchedCells"] += touched
            totals["sameRowAttemptCells"] += same_row_attempt
            rows.append({
                "appendixId": result.get("appendixId"), "number": number, "title": result.get("title"),
                "table": table_idx, "aligned": True, "usedFallback": used_fallback,
                "humanTable": human_idx, "alignScore": round(align_score, 3),
                "humanChangedCells": changed, "agentCorrectCells": correct,
                "agentTolerantCorrectCells": tolerant_correct, "agentTolerantV5CorrectCells": tolerant_v5_correct, "agentTouchedCells": touched,
                "sameRowAttemptCells": same_row_attempt,
                "strictCoverage": round(correct / changed, 4) if changed else None,
                "tolerantCoverage": round(tolerant_correct / changed, 4) if changed else None,
                "wrongCells": wrong_cells,
            })
    totals["strictCoverage"] = round(totals["agentCorrectCells"] / totals["humanChangedCells"], 4) if totals["humanChangedCells"] else None
    totals["tolerantCoverage"] = round(totals["agentTolerantCorrectCells"] / totals["humanChangedCells"], 4) if totals["humanChangedCells"] else None
    totals["tolerantV5Coverage"] = round(totals.get("agentTolerantV5CorrectCells", 0) / totals["humanChangedCells"], 4) if totals["humanChangedCells"] else None
    totals["rowAttemptCoverage"] = round((totals["agentCorrectCells"] + totals["sameRowAttemptCells"]) / totals["humanChangedCells"], 4) if totals["humanChangedCells"] else None
    return {"summary": totals, "tables": rows, "headingsFound": len(headings)}


def main() -> None:
    results_path = Path(sys.argv[1])
    human_docx = Path(sys.argv[2])
    out_path = Path(sys.argv[3]) if len(sys.argv) > 3 else Path("score_report_v2.json")
    results = json.loads(results_path.read_text(encoding="utf-8"))
    report = score_outputs(results, human_docx)
    out_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report["summary"], ensure_ascii=False, indent=2))
    low = [r for r in report["tables"] if r.get("aligned") and (r.get("tolerantCoverage") or 0) < 0.6]
    print(f"\n低覆盖率表格({len(low)}个,tolerantCoverage<0.6):")
    for r in sorted(low, key=lambda x: -(x.get("humanChangedCells") or 0))[:30]:
        print(f"  {r['appendixId']:10} {r['number']:10} {r['title'][:26]:26} tolerant={r['tolerantCoverage']} changed={r['humanChangedCells']}")
    unaligned = [r for r in report["tables"] if not r.get("aligned")]
    print(f"\n未对齐 {len(unaligned)} 个 (含fallback失败 {report['summary']['noSectionFound']} 个)")


if __name__ == "__main__":
    main()
