#!/usr/bin/env python3
"""Run conservative technical-bid Word format cleaning from a manifest."""

from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.shared import Pt


SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
ASSEMBLER_SCRIPTS_DIR = SKILL_DIR.parent / "bid-tech-assembler" / "scripts"
ASSEMBLER_REFERENCES_DIR = SKILL_DIR.parent / "bid-tech-assembler" / "references"
sys.path.insert(0, str(ASSEMBLER_SCRIPTS_DIR))


def _drop_foreign_module(module_name: str, expected_path: Path) -> None:
    module = sys.modules.get(module_name)
    if module is None:
        return
    loaded_file = getattr(module, "__file__", None)
    if not loaded_file:
        return
    try:
        if Path(loaded_file).resolve() != expected_path.resolve():
            sys.modules.pop(module_name, None)
    except OSError:
        sys.modules.pop(module_name, None)


_drop_foreign_module("verify", ASSEMBLER_SCRIPTS_DIR / "verify.py")

from finalize import force_update_fields, insert_toc_field, reapply_heading_fonts, replace_header_text  # noqa: E402
from numbering_fixer import (  # noqa: E402
    enforce_no_auto_numbering_on_numbered_headings,
    strip_numPr_from_body,
    strip_numPr_from_heading_styles,
)
from verify import scan_docx  # noqa: E402


SCHEMA_VERSION = "bid-tech-format-clean-v1"
REQUIRED_FIELDS = ("inputFile", "outlineFile", "outputFile", "projectName")
TOC_BREAK_BOOKMARK = "_TECH_FORMAT_CLEANER_TOC_BREAK"


def run_manifest(manifest_path: str | Path, response: str = "summary") -> dict[str, Any]:
    path = Path(manifest_path)
    manifest = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(manifest, dict):
        raise ValueError("manifest must be a JSON object")
    _validate_manifest(manifest)

    input_path = _resolve_manifest_path(manifest["inputFile"], path)
    outline_path = _resolve_manifest_path(manifest["outlineFile"], path)
    output_path = _resolve_manifest_path(manifest["outputFile"], path)
    style_path = _resolve_style_path(manifest.get("styleSpecPath"), path)
    clean_result = clean_docx(
        input_file=input_path,
        outline_file=outline_path,
        output_file=output_path,
        project_name=str(manifest["projectName"]),
        style_spec_path=style_path,
    )
    report = verify_cleaned_docx(
        output_file=output_path,
        outline_file=outline_path,
        clean_result=clean_result,
    )

    summary = {
        "outlineCount": int(report["outlineCount"]),
        "matchedHeadingCount": int(report["matchedHeadingCount"]),
        "unmatchedHeadingCount": len(report["unmatchedHeadings"]),
        "internalHeadingCount": int(report["internalHeadingCount"]),
        "headingLevelCounts": report["headingLevelCounts"],
        "tocInserted": bool(clean_result["tocInserted"]),
        "tocPresent": bool(report["tocPresent"]),
        "headerCleaned": bool(report["headerCleaned"]),
        "placeholderCount": int(report["placeholderCount"]),
        "orientation": report["orientation"],
        "fontFamilies": clean_result["fontFamilies"],
        "riskCount": len(report["formatRisks"]),
        "warnings": report["warnings"],
    }
    result = {
        "schema_version": SCHEMA_VERSION,
        "inputFile": str(input_path),
        "outlineFile": str(outline_path),
        "outputFile": str(output_path),
        "reportFile": "",
        "summary": summary,
        "warnings": report["warnings"],
    }
    if response != "summary":
        result["details"] = {
            "cleanResult": clean_result,
            "report": report,
            "styleSpecPath": str(style_path),
        }
    return result


def clean_docx(
    *,
    input_file: str | Path,
    outline_file: str | Path,
    output_file: str | Path,
    project_name: str,
    style_spec_path: str | Path,
) -> dict[str, Any]:
    input_path = Path(input_file)
    outline_path = Path(outline_file)
    output_path = Path(output_file)
    style_path = Path(style_spec_path)

    if not input_path.exists():
        raise FileNotFoundError(f"inputFile does not exist: {input_path}")
    if not outline_path.exists():
        raise FileNotFoundError(f"outlineFile does not exist: {outline_path}")
    if not style_path.exists():
        raise FileNotFoundError(f"styleSpecPath does not exist: {style_path}")
    if input_path.resolve() == output_path.resolve():
        raise ValueError("inputFile must not be overwritten by outputFile")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(input_path, output_path)

    style_spec = _load_style_spec(style_path)
    outline_items = flatten_outline(load_outline(outline_path))

    doc = Document(str(output_path))
    toc_present_before = document_has_toc(doc)
    toc_inserted = False
    toc_cfg = style_spec.get("toc") if isinstance(style_spec.get("toc"), dict) else {}
    if not toc_present_before and toc_cfg.get("insert_when_missing", True) is True:
        insert_toc_field(doc)
        toc_inserted = True
    _apply_toc_page_break(
        doc,
        toc_cfg.get("page_break_after", True) is True,
        adopt_unmarked=toc_inserted,
    )

    heading_style_result = _configure_heading_styles(doc, style_spec)
    heading_result = _promote_existing_headings(doc, outline_items, style_spec)
    # 标题只来自现有 Heading/outline；Cleaner 不再根据加粗或文本形态猜标题。
    internal_heading_result = {"promoted_headings": []}
    _apply_document_page_format(doc, style_spec.get("page"))
    _apply_document_body_format(doc, style_spec)
    _apply_document_table_format(doc, style_spec.get("table_cell"))
    _apply_document_headers(doc, style_spec.get("header"), project_name)
    _reapply_heading_direct_formats(doc, style_spec)
    reapply_heading_fonts(doc, style_spec)
    strip_numPr_from_heading_styles(doc)
    strip_numPr_from_body(doc)
    # 不变量：已写入文本编号的 Heading 不得再有有效 Word 自动编号
    auto_numbering_fixed = enforce_no_auto_numbering_on_numbered_headings(doc)
    orientation_result = _normalize_section_orientations(doc)
    doc.save(str(output_path))

    if project_name:
        replace_header_text(output_path, project_name)
    force_update_fields(output_path)

    return {
        "inputFile": str(input_path),
        "outlineFile": str(outline_path),
        "outputFile": str(output_path),
        "styleSpecPath": str(style_path),
        "projectName": project_name,
        "outlineCount": len(outline_items),
        "matchedHeadingCount": len(heading_result["matched_headings"]),
        "unmatchedHeadingCount": len(heading_result["unmatched_headings"]),
        "matchedHeadings": heading_result["matched_headings"],
        "unmatchedHeadings": heading_result["unmatched_headings"],
        "internalHeadingCount": len(internal_heading_result["promoted_headings"]),
        "internalHeadings": internal_heading_result["promoted_headings"],
        "headingStylesConfigured": heading_style_result,
        "autoNumberingFixed": auto_numbering_fixed,
        "orientation": orientation_result,
        "tocInserted": toc_inserted,
        "tocPresent": toc_present_before or toc_inserted,
        "fontFamilies": _style_font_families(style_spec),
    }


def verify_cleaned_docx(
    *,
    output_file: str | Path,
    outline_file: str | Path,
    clean_result: dict[str, Any] | None = None,
) -> dict[str, Any]:
    output_path = Path(output_file)
    outline_path = Path(outline_file)
    scan = scan_docx(output_path)
    outline_items = flatten_outline(load_outline(outline_path))
    doc = Document(str(output_path))
    toc_present = document_has_toc(doc)
    header_cleaned = _header_cleaned(output_path)
    heading_level_counts = _heading_level_counts(scan)
    orientation_summary = dict((clean_result or {}).get("orientation") or _section_orientation_summary(doc))
    matched = list((clean_result or {}).get("matchedHeadings") or [])
    unmatched = list((clean_result or {}).get("unmatchedHeadings") or [])
    internal_heading_count = int((clean_result or {}).get("internalHeadingCount") or 0)
    if not matched and not unmatched:
        heading_texts = {text.strip() for _, text in scan.get("heading_list", []) if text.strip()}
        matched, unmatched = _infer_matches_from_doc(outline_items, heading_texts)

    risks = _structural_risks(
        scan=scan,
        toc_present=toc_present,
        header_cleaned=header_cleaned,
        outline_count=len(outline_items),
        unmatched_count=len(unmatched),
    )
    report = {
        "outlineCount": len(outline_items),
        "matchedHeadingCount": len(matched),
        "unmatchedHeadings": unmatched,
        "internalHeadingCount": internal_heading_count,
        "headingLevelCounts": heading_level_counts,
        "tocPresent": toc_present,
        "headerCleaned": header_cleaned,
        "placeholderCount": len(scan.get("placeholders") or []),
        "orientation": orientation_summary,
        "formatRisks": risks,
    }
    report["warnings"] = _build_warnings(report, scan)
    return report


def _validate_manifest(manifest: dict[str, Any]) -> None:
    missing = [field for field in REQUIRED_FIELDS if not manifest.get(field)]
    if missing:
        raise ValueError(f"manifest missing fields: {', '.join(missing)}")
    input_path = Path(manifest["inputFile"])
    output_path = Path(manifest["outputFile"])
    if input_path.resolve() == output_path.resolve():
        raise ValueError("inputFile must not equal outputFile")


def _resolve_style_path(style_path_value: Any, manifest_path: Path) -> Path:
    if style_path_value:
        return _resolve_manifest_path(style_path_value, manifest_path)
    return (ASSEMBLER_REFERENCES_DIR / "heading_style.json").resolve()


def _resolve_manifest_path(value: Any, manifest_path: Path) -> Path:
    candidate = Path(str(value))
    if not candidate.is_absolute():
        candidate = (manifest_path.parent / candidate).resolve()
    return candidate


def _load_style_spec(path: Path) -> dict[str, Any]:
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError("styleSpecPath must be a JSON object")
    data.setdefault("heading", {})
    return data


def _style_font_families(style_spec: dict[str, Any]) -> list[str]:
    families: list[str] = []

    def collect(value: Any) -> None:
        if isinstance(value, dict):
            for key, nested in value.items():
                if key in {"zh_font", "en_font"} and isinstance(nested, str) and nested not in families:
                    families.append(nested)
                else:
                    collect(nested)
        elif isinstance(value, list):
            for nested in value:
                collect(nested)

    collect(style_spec)
    return families


def load_outline(path: str | Path) -> dict[str, Any]:
    data = json.loads(Path(path).read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError("outlineFile must be a JSON object")
    schema = str(data.get("schema_version") or "")
    if schema and schema not in {"tech_bid_outline.v1", "business_bid_outline.v1"}:
        raise ValueError(f"unsupported outline schema_version: {schema}")
    return data


def flatten_outline(outline: dict[str, Any]) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []

    def visit(nodes: Any, inherited_level: int = 1) -> None:
        if not isinstance(nodes, list):
            return
        for index, node in enumerate(nodes, start=1):
            if not isinstance(node, dict):
                continue
            title = str(node.get("title") or "").strip()
            if not title:
                visit(node.get("children"), inherited_level + 1)
                continue
            level = _coerce_level(node.get("level"), inherited_level)
            items.append(
                {
                    "id": str(node.get("id") or f"outline-{len(items) + 1:04d}"),
                    "title": title,
                    "number": str(node.get("number") or "").strip(),
                    "level": level,
                    "order": len(items) + 1,
                    "source_index": index,
                }
            )
            visit(node.get("children"), level + 1)

    visit(outline.get("sections"), 1)
    return items


def _coerce_level(value: Any, default: int) -> int:
    try:
        level = int(value)
    except (TypeError, ValueError):
        level = default
    return max(1, min(level, 9))


def document_has_toc(doc: Document) -> bool:
    return any((node.text or "").upper().find("TOC") >= 0 for node in doc.element.iter(qn("w:instrText")))


def _configure_heading_styles(doc: Document, style_spec: dict[str, Any], max_level: int = 6) -> dict[str, Any]:
    configured: list[dict[str, Any]] = []
    heading_cfg = style_spec.get("heading", {})
    if not isinstance(heading_cfg, dict):
        return {"levels": configured}

    for level in range(1, max_level + 1):
        cfg = heading_cfg.get(str(level))
        if not isinstance(cfg, dict):
            continue
        style = _get_or_create_heading_style(doc, level)
        _apply_style_format(style, cfg)
        _ensure_style_outline_level(style, level)
        configured.append({"level": level, "styleName": style.name, "styleId": style.style_id})
    return {"levels": configured}


def _get_or_create_heading_style(doc: Document, level: int):
    for style_name in (f"Heading {level}", f"标题 {level}", f"标题{level}"):
        try:
            return doc.styles[style_name]
        except KeyError:
            continue
    style = doc.styles.add_style(f"Heading {level}", WD_STYLE_TYPE.PARAGRAPH)
    try:
        style.base_style = doc.styles["Normal"]
    except KeyError:
        pass
    return style


def _apply_style_format(style, cfg: dict[str, Any]) -> None:
    font = style.font
    if "en_font" in cfg:
        font.name = str(cfg["en_font"])
    if "size_pt" in cfg:
        font.size = Pt(float(cfg["size_pt"]))
    font.bold = bool(cfg.get("bold", False))

    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    zh_font = str(cfg.get("zh_font") or cfg.get("en_font") or "")
    en_font = str(cfg.get("en_font") or cfg.get("zh_font") or "")
    if zh_font or en_font:
        r_fonts.set(qn("w:eastAsia"), zh_font)
        r_fonts.set(qn("w:ascii"), en_font)
        r_fonts.set(qn("w:hAnsi"), en_font)
        r_fonts.set(qn("w:cs"), en_font)

    paragraph_format = style.paragraph_format
    align = str(cfg.get("align") or "").lower()
    if align == "center":
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align in {"both", "justify"}:
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "left":
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if "space_before_pt" in cfg:
        paragraph_format.space_before = Pt(float(cfg["space_before_pt"]))
    if "space_after_pt" in cfg:
        paragraph_format.space_after = Pt(float(cfg["space_after_pt"]))
    if "line_spacing" in cfg:
        paragraph_format.line_spacing = float(cfg["line_spacing"])
    if "left_indent_cm" in cfg:
        paragraph_format.left_indent = Cm(float(cfg["left_indent_cm"]))
    if "first_line_indent_chars" in cfg and "size_pt" in cfg:
        paragraph_format.first_line_indent = Pt(float(cfg["first_line_indent_chars"]) * float(cfg["size_pt"]))


def _apply_document_page_format(doc: Document, page_cfg: Any) -> None:
    if not isinstance(page_cfg, dict):
        return
    for section in doc.sections:
        for key, attribute in (
            ("top_cm", "top_margin"),
            ("bottom_cm", "bottom_margin"),
            ("left_cm", "left_margin"),
            ("right_cm", "right_margin"),
            ("header_top_cm", "header_distance"),
            ("footer_bottom_cm", "footer_distance"),
        ):
            if key in page_cfg:
                setattr(section, attribute, Cm(float(page_cfg[key])))


def _apply_document_body_format(doc: Document, style_spec: dict[str, Any]) -> None:
    body_cfg = style_spec.get("body")
    toc_protected = _toc_protected_paragraph_elements(doc)
    preserve_after_image = False
    for paragraph in doc.paragraphs:
        if paragraph._element in toc_protected:
            preserve_after_image = False
            continue
        if _paragraph_heading_level(paragraph):
            preserve_after_image = False
            continue
        if _paragraph_contains_visual(paragraph):
            preserve_after_image = True
            continue
        if preserve_after_image:
            if paragraph.text.strip():
                preserve_after_image = False
            continue
        if _is_preserved_layout_paragraph(paragraph):
            continue
        if isinstance(body_cfg, dict):
            _apply_paragraph_direct_format(paragraph, body_cfg)


def _apply_document_table_format(doc: Document, table_cfg: Any) -> None:
    if not isinstance(table_cfg, dict):
        return
    table_align = str(table_cfg.get("table_align") or "").lower()
    alignment = {
        "left": WD_TABLE_ALIGNMENT.LEFT,
        "center": WD_TABLE_ALIGNMENT.CENTER,
        "right": WD_TABLE_ALIGNMENT.RIGHT,
    }.get(table_align)
    for table in doc.tables:
        if _table_contains_visual(table):
            continue
        if alignment is not None:
            table.alignment = alignment
        for row in table.rows:
            for cell in row.cells:
                preserve_after_image = False
                for paragraph in cell.paragraphs:
                    if _paragraph_contains_visual(paragraph):
                        preserve_after_image = True
                        continue
                    if preserve_after_image:
                        if paragraph.text.strip():
                            preserve_after_image = False
                        continue
                    if _is_preserved_layout_paragraph(paragraph):
                        continue
                    _apply_paragraph_direct_format(paragraph, table_cfg)


def _apply_document_headers(doc: Document, header_cfg: Any, project_name: str) -> None:
    if not isinstance(header_cfg, dict):
        return
    template = str(header_cfg.get("text_template") or "")
    header_text = (
        template.replace("{project_name}", project_name).replace("{projectName}", project_name)
        if template
        else ""
    )
    seen: set[Any] = set()
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            marker = header._element
            if marker in seen:
                continue
            seen.add(marker)
            paragraphs = list(header.paragraphs)
            if not paragraphs:
                paragraphs = [header.add_paragraph()]
            target = next((paragraph for paragraph in paragraphs if paragraph.text.strip()), paragraphs[0])
            if header_text:
                text_runs = [run for run in target.runs if run.text]
                if text_runs:
                    text_runs[0].text = header_text
                    for run in text_runs[1:]:
                        run.text = ""
                else:
                    target.add_run(header_text)
            for paragraph in paragraphs:
                _apply_paragraph_direct_format(paragraph, header_cfg)


def _reapply_heading_direct_formats(doc: Document, style_spec: dict[str, Any]) -> None:
    for paragraph in doc.paragraphs:
        level = _paragraph_heading_level(paragraph)
        if not level:
            continue
        cfg = _heading_cfg(style_spec, level)
        if cfg:
            _apply_paragraph_properties(paragraph.paragraph_format, cfg)
            _apply_heading_run_format(paragraph, cfg)


def _apply_paragraph_direct_format(paragraph, cfg: dict[str, Any]) -> None:
    _apply_paragraph_properties(paragraph.paragraph_format, cfg)
    for run in paragraph.runs:
        _apply_run_format(run, cfg)


def _apply_paragraph_properties(paragraph_format, cfg: dict[str, Any]) -> None:
    align = str(cfg.get("align") or "").lower()
    paragraph_format.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(align, paragraph_format.alignment)
    if "space_before_pt" in cfg:
        paragraph_format.space_before = Pt(float(cfg["space_before_pt"]))
    if "space_after_pt" in cfg:
        paragraph_format.space_after = Pt(float(cfg["space_after_pt"]))
    if "line_spacing" in cfg:
        paragraph_format.line_spacing = float(cfg["line_spacing"])
    if "left_indent_cm" in cfg:
        paragraph_format.left_indent = Cm(float(cfg["left_indent_cm"]))
    if "first_line_indent_chars" in cfg and "size_pt" in cfg:
        paragraph_format.first_line_indent = Pt(float(cfg["first_line_indent_chars"]) * float(cfg["size_pt"]))


def _apply_run_format(run, cfg: dict[str, Any]) -> None:
    if "en_font" in cfg:
        run.font.name = str(cfg["en_font"])
    if "size_pt" in cfg:
        run.font.size = Pt(float(cfg["size_pt"]))
    if "bold" in cfg:
        # 正文规则的非粗体是默认值，不能覆盖素材 Run 上已有的局部强调。
        if bool(cfg["bold"]):
            run.font.bold = True
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    zh_font = str(cfg.get("zh_font") or cfg.get("en_font") or "")
    en_font = str(cfg.get("en_font") or cfg.get("zh_font") or "")
    if zh_font or en_font:
        r_fonts.set(qn("w:eastAsia"), zh_font)
        r_fonts.set(qn("w:ascii"), en_font)
        r_fonts.set(qn("w:hAnsi"), en_font)
        r_fonts.set(qn("w:cs"), en_font)


def _apply_toc_page_break(doc: Document, enabled: bool, *, adopt_unmarked: bool = False) -> None:
    anchor = _toc_result_anchor(doc)
    if anchor is None:
        return
    sibling = anchor.getnext()
    if adopt_unmarked and _is_dedicated_page_break(sibling) and not _is_cleaner_toc_break(sibling):
        next_sibling = sibling.getnext()
        sibling.getparent().remove(sibling)
        sibling = next_sibling
    while _is_cleaner_toc_break(sibling):
        next_sibling = sibling.getnext()
        sibling.getparent().remove(sibling)
        sibling = next_sibling
    if enabled:
        anchor.addnext(_make_cleaner_toc_break(doc))


def _toc_result_anchor(doc: Document):
    body = doc.element.body
    instruction = next(
        (node for node in body.iter(qn("w:instrText")) if _is_toc_instruction(node.text)),
        None,
    )
    if instruction is None:
        return None
    toc_paragraph = None
    outer_sdt = None
    current = instruction.getparent()
    while current is not None and current is not body:
        if current.tag == qn("w:p") and toc_paragraph is None:
            toc_paragraph = current
        if current.tag == qn("w:sdt"):
            outer_sdt = current
        current = current.getparent()
    if outer_sdt is not None:
        return outer_sdt
    if toc_paragraph is None:
        return None

    protected = _toc_protected_paragraph_elements(doc)
    anchor = toc_paragraph
    current = toc_paragraph
    while current is not None and current.tag == qn("w:p") and current in protected:
        anchor = current
        current = current.getnext()
    current = anchor.getnext()
    while _is_toc_result_paragraph(current):
        anchor = current
        current = current.getnext()
    return anchor


def _is_toc_result_paragraph(element) -> bool:
    if element is None or element.tag != qn("w:p"):
        return False
    p_pr = element.find(qn("w:pPr"))
    p_style = p_pr.find(qn("w:pStyle")) if p_pr is not None else None
    style_id = str(p_style.get(qn("w:val")) or "") if p_style is not None else ""
    return _is_toc_style_identifier(style_id)


def _is_toc_style_identifier(value: Any) -> bool:
    return bool(
        re.fullmatch(r"(?:TOC|目录)\s*[1-9]", str(value or "").strip(), flags=re.IGNORECASE)
    )


def _make_cleaner_toc_break(doc: Document):
    bookmark_id = str(
        max(
            (int(node.get(qn("w:id"))) for node in doc.element.iter(qn("w:bookmarkStart")) if str(node.get(qn("w:id")) or "").isdigit()),
            default=0,
        )
        + 1
    )
    page_break = OxmlElement("w:p")
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), bookmark_id)
    bookmark_start.set(qn("w:name"), TOC_BREAK_BOOKMARK)
    page_break.append(bookmark_start)
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    page_break.append(run)
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), bookmark_id)
    page_break.append(bookmark_end)
    return page_break


def _is_cleaner_toc_break(element) -> bool:
    if not _is_dedicated_page_break(element):
        return False
    return any(node.get(qn("w:name")) == TOC_BREAK_BOOKMARK for node in element.iter(qn("w:bookmarkStart")))


def _is_dedicated_page_break(element) -> bool:
    if element is None or element.tag != qn("w:p"):
        return False
    if any((node.text or "").strip() for node in element.iter(qn("w:t"))):
        return False
    if element.findall(".//" + qn("w:drawing")) or element.findall(".//" + qn("w:pict")):
        return False
    return any((node.get(qn("w:type")) or "page") == "page" for node in element.iter(qn("w:br")))


def _ensure_style_outline_level(style, level: int) -> None:
    p_pr = style.element.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(max(0, min(level - 1, 8))))


def _promote_existing_headings(
    doc: Document,
    outline_items: list[dict[str, Any]],
    style_spec: dict[str, Any],
) -> dict[str, Any]:
    paragraphs = list(doc.paragraphs)
    cursor = 0
    matched: list[dict[str, Any]] = []
    unmatched: list[dict[str, Any]] = []

    for item in outline_items:
        found = _find_outline_item(paragraphs, item, cursor)
        if found is None:
            unmatched.append(_outline_report_item(item))
            continue
        index, mode = found
        paragraph = paragraphs[index]
        level = _effective_outline_level(item)
        _set_heading_style(paragraph, level, doc)
        _strip_paragraph_numpr(paragraph)
        _clear_direct_outline_level(paragraph)
        _apply_heading_run_format(paragraph, _heading_cfg(style_spec, level))
        matched.append(
            {
                **_outline_report_item(item),
                "paragraphIndex": index,
                "matchMode": mode,
                "text": paragraph.text.strip(),
            }
        )
        cursor = index + 1

    return {"matched_headings": matched, "unmatched_headings": unmatched}


def _promote_body_internal_headings(doc: Document, style_spec: dict[str, Any]) -> dict[str, Any]:
    promoted: list[dict[str, Any]] = []
    current_outline_level = 1

    for index, paragraph in enumerate(doc.paragraphs):
        existing_level = _paragraph_heading_level(paragraph)
        if existing_level:
            current_outline_level = max(1, min(existing_level, 4))
            continue

        text = _clean_paragraph_text(paragraph.text)
        if not _looks_like_body_internal_heading(paragraph, text):
            continue

        level = _infer_body_internal_heading_level(text, current_outline_level)
        if not level:
            continue

        _set_heading_style(paragraph, level, doc)
        _strip_paragraph_numpr(paragraph)
        _clear_direct_outline_level(paragraph)
        _apply_heading_run_format(paragraph, _heading_cfg(style_spec, level))
        promoted.append(
            {
                "paragraphIndex": index,
                "level": level,
                "text": text,
                "matchMode": "body-internal",
            }
        )

    return {"promoted_headings": promoted}


def _paragraph_heading_level(paragraph) -> int:
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is not None:
        outline = p_pr.find(qn("w:outlineLvl"))
        if outline is not None:
            try:
                value = int(outline.get(qn("w:val")))
            except (TypeError, ValueError):
                value = -1
            if 0 <= value <= 8:
                return value + 1

    style = getattr(paragraph, "style", None)
    visited: set[str] = set()
    while style is not None:
        style_id = str(getattr(style, "style_id", "") or "")
        if style_id in visited:
            break
        visited.add(style_id)

        style_name = str(getattr(style, "name", "") or "")
        match = re.search(r"heading\s*(\d+)", style_name.strip().lower())
        if match:
            return max(1, min(int(match.group(1)), 9))
        match = re.search(r"标题\s*(\d+)", style_name)
        if match:
            return max(1, min(int(match.group(1)), 9))

        style_p_pr = style.element.find(qn("w:pPr"))
        if style_p_pr is not None:
            outline = style_p_pr.find(qn("w:outlineLvl"))
            if outline is not None:
                try:
                    value = int(outline.get(qn("w:val")))
                except (TypeError, ValueError):
                    value = -1
                if 0 <= value <= 8:
                    return value + 1
        style = getattr(style, "base_style", None)
    return 0


def _paragraph_contains_visual(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing|.//w:pict|.//w:object"))


def _table_contains_visual(table) -> bool:
    return bool(table._tbl.xpath(".//w:drawing|.//w:pict|.//w:object"))


def _toc_protected_paragraph_elements(doc: Document) -> set[Any]:
    body = doc.element.body
    paragraphs = list(body.iter(qn("w:p")))
    protected: set[Any] = set()
    toc_instruction_paragraphs: set[Any] = set()
    active_fields: list[dict[str, Any]] = []

    for paragraph in paragraphs:
        for field in active_fields:
            field["paragraphs"].add(paragraph)

        for node in paragraph.iter():
            if node.tag == qn("w:fldSimple") and _is_toc_instruction(node.get(qn("w:instr"))):
                protected.add(paragraph)
                toc_instruction_paragraphs.add(paragraph)
                continue

            if node.tag == qn("w:fldChar"):
                field_type = str(node.get(qn("w:fldCharType")) or "").lower()
                if field_type == "begin":
                    active_fields.append({"paragraphs": {paragraph}, "is_toc": False})
                elif field_type == "end" and active_fields:
                    field = active_fields.pop()
                    if field["is_toc"]:
                        protected.update(field["paragraphs"])
                continue

            if node.tag != qn("w:instrText") or not _is_toc_instruction(node.text):
                continue
            toc_instruction_paragraphs.add(paragraph)
            if active_fields:
                active_fields[-1]["is_toc"] = True
            else:
                protected.add(paragraph)

    for field in active_fields:
        if field["is_toc"]:
            protected.update(field["paragraphs"])

    paragraph_positions = {paragraph: index for index, paragraph in enumerate(paragraphs)}
    for paragraph in toc_instruction_paragraphs:
        index = paragraph_positions.get(paragraph, 0)
        while index > 0 and paragraphs[index - 1] in protected:
            index -= 1
        if index <= 0:
            continue
        title = paragraphs[index - 1]
        if _is_toc_heading_element(title, doc):
            protected.add(title)

    return protected


def _is_toc_instruction(value: Any) -> bool:
    return bool(re.search(r"\bTOC\b", str(value or ""), flags=re.IGNORECASE))


def _is_toc_heading_element(element, doc: Document) -> bool:
    p_pr = element.find(qn("w:pPr"))
    p_style = p_pr.find(qn("w:pStyle")) if p_pr is not None else None
    style_id = str(p_style.get(qn("w:val")) or "") if p_style is not None else ""
    identifiers = [style_id]
    if style_id:
        identifiers.extend(
            str(getattr(style, "name", "") or "")
            for style in doc.styles
            if str(getattr(style, "style_id", "") or "") == style_id
        )
    if any(_is_toc_heading_identifier(identifier) for identifier in identifiers):
        return True

    text = "".join(str(node.text or "") for node in element.iter(qn("w:t")))
    normalized = re.sub(r"\s+", "", text).casefold()
    return normalized in {"目录", "contents", "tableofcontents"}


def _is_toc_heading_identifier(value: Any) -> bool:
    normalized = re.sub(r"[\s_-]+", "", str(value or "")).casefold()
    return normalized in {"toc", "tocheading", "目录标题"}


def _is_preserved_layout_paragraph(paragraph) -> bool:
    if _paragraph_uses_toc_style(paragraph) or _paragraph_uses_toc_heading_style(paragraph):
        return True
    style = getattr(paragraph, "style", None)
    style_name = str(getattr(style, "name", "") or "").strip().lower()
    if any(marker in style_name for marker in ("图片", "图注", "照片", "caption", "figure", "image")):
        return True
    return _looks_like_caption_or_table_title(paragraph.text)


def _paragraph_uses_toc_style(paragraph) -> bool:
    style = getattr(paragraph, "style", None)
    visited: set[str] = set()
    while style is not None:
        style_id = str(getattr(style, "style_id", "") or "")
        style_name = str(getattr(style, "name", "") or "")
        marker = style_id or style_name
        if marker in visited:
            break
        visited.add(marker)
        if _is_toc_style_identifier(style_id) or _is_toc_style_identifier(style_name):
            return True
        style = getattr(style, "base_style", None)
    return False


def _paragraph_uses_toc_heading_style(paragraph) -> bool:
    p_pr = paragraph._element.find(qn("w:pPr"))
    p_style = p_pr.find(qn("w:pStyle")) if p_pr is not None else None
    style_id = str(p_style.get(qn("w:val")) or "") if p_style is not None else ""
    style = getattr(paragraph, "style", None)
    style_name = str(getattr(style, "name", "") or "")
    return _is_toc_heading_identifier(style_id) or _is_toc_heading_identifier(style_name)


def _clean_paragraph_text(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "").replace("\u3000", " ")).strip()


def _looks_like_body_internal_heading(paragraph, text: str) -> bool:
    if not text:
        return False
    if len(text) > 48:
        return False
    if text.startswith(("[", "【", "*")):
        return False
    if re.match(r"^(?:备注|注|说明)[:：]", text):
        return False
    if re.search(r"[。；;！!？?：:]$", text):
        return False
    if _looks_like_caption_or_table_title(text):
        return False

    explicit_level = _infer_level_from_number(text)
    if explicit_level >= 3:
        return True
    if _is_short_ordered_heading(text):
        return _bold_text_ratio(paragraph) >= 0.65 and _has_title_shape(text)
    return _bold_text_ratio(paragraph) >= 0.65 and _has_title_shape(text)


def _looks_like_caption_or_table_title(text: str) -> bool:
    clean = _clean_paragraph_text(text)
    compact = re.sub(r"\s+", "", clean)
    if re.match(r"^(?:图|表)\s+\S", clean):
        return True
    if re.match(r"^(?:图|表)[：:]\S", clean):
        return True
    if re.match(r"^(?:图|表)[A-Za-z]?[.-]?\d", compact):
        return True
    if re.match(r"^(?:图|表)[一二三四五六七八九十]+", compact):
        return True
    if re.match(r"^图[\u4e00-\u9fffA-Za-z0-9]{2,40}(?:示意图|结构图|流程图|布置图|接线图|曲线图|照片)$", compact):
        return True
    if "一览表" in compact or "统计表" in compact or "参数表" in compact:
        return True
    if compact.endswith(("表", "清单")) and len(compact) <= 28:
        return True
    return False


def _is_short_ordered_heading(text: str) -> bool:
    if len(text) > 34:
        return False
    return bool(
        re.match(r"^[（(][一二三四五六七八九十\d]+[）)]\s*\S", text)
        or re.match(r"^\d+[）)、.]\s*\S", text)
        or re.match(r"^[一二三四五六七八九十]+[、.]\s*\S", text)
    )


def _has_title_shape(text: str) -> bool:
    if len(text) < 2:
        return False
    if re.search(r"[，,。；;！!？?：:]", text):
        return False
    cjk_count = len(re.findall(r"[\u4e00-\u9fff]", text))
    return cjk_count >= 2


def _bold_text_ratio(paragraph) -> float:
    total = 0
    bold = 0
    for run in paragraph.runs:
        text = str(run.text or "")
        visible = len(re.sub(r"\s+", "", text))
        if not visible:
            continue
        total += visible
        if run.bold is True or getattr(run.font, "bold", None) is True:
            bold += visible
    return bold / total if total else 0.0


def _infer_body_internal_heading_level(text: str, current_outline_level: int) -> int:
    explicit_level = _infer_level_from_number(text)
    if explicit_level >= 3:
        return max(3, min(explicit_level, 4))
    if _is_short_ordered_heading(text):
        return max(3, min(current_outline_level + 2, 4))
    return max(2, min(current_outline_level + 1, 4))


def _find_outline_item(paragraphs, item: dict[str, Any], start_index: int) -> tuple[int, str] | None:
    full_key = _normalize_match_text(_format_heading_text(item)) if item.get("number") else ""
    title_key = _normalize_match_text(str(item.get("title") or ""))
    if not title_key:
        return None
    for index in range(start_index, len(paragraphs)):
        text_key = _normalize_match_text(paragraphs[index].text)
        if not text_key:
            continue
        if full_key and text_key == full_key:
            return index, "number+title"
        if text_key == title_key:
            return index, "title"
    return None


def _normalize_match_text(text: str) -> str:
    normalized = str(text or "").strip().replace("\u3000", " ")
    return re.sub(r"\s+", "", normalized)


def _format_heading_text(item: dict[str, Any]) -> str:
    title = str(item.get("title") or "").strip()
    number = str(item.get("number") or "").strip()
    return f"{number} {title}".strip() if number else title


def _outline_report_item(item: dict[str, Any]) -> dict[str, Any]:
    return {
        "id": item.get("id"),
        "number": item.get("number", ""),
        "title": item.get("title", ""),
        "level": _effective_outline_level(item),
        "order": item.get("order"),
    }


def _effective_outline_level(item: dict[str, Any]) -> int:
    try:
        level = int(item.get("level") or 1)
    except (TypeError, ValueError):
        level = 1
    inferred = max(
        (
            _infer_level_from_number(str(value or ""))
            for value in (
                item.get("number"),
                item.get("title"),
                _format_heading_text(item),
            )
        ),
        default=0,
    )
    if inferred:
        level = max(level, inferred)
    return max(1, min(level, 9))


def _infer_level_from_number(value: str) -> int:
    text = str(value or "").strip()
    if not text:
        return 0
    first_token = re.split(r"\s+", text, maxsplit=1)[0].strip("：:、")
    if re.fullmatch(r"第[一二三四五六七八九十百\d]+章", first_token):
        return 1
    if re.fullmatch(r"\d+(?:\.\d+)*", first_token):
        return first_token.count(".") + 1
    appendix = re.fullmatch(r"(?:技术)?附表\s*([A-Za-z])((?:[.-]\d+)*)", first_token)
    if appendix:
        suffix = appendix.group(2) or ""
        parts = [part for part in re.split(r"[.-]", suffix) if part]
        return 1 + len(parts)
    if re.fullmatch(r"附表\d+", first_token) or re.fullmatch(r"技术附表[A-Za-z]", first_token):
        return 1
    return 0


def _set_heading_style(paragraph, level: int, doc: Document) -> None:
    for style_name in (f"Heading {level}", f"标题 {level}", f"标题{level}"):
        try:
            paragraph.style = doc.styles[style_name]
            return
        except KeyError:
            continue


def _strip_paragraph_numpr(paragraph) -> bool:
    p_pr = paragraph._element.find(qn("w:pPr"))
    if p_pr is None:
        return False
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return False
    p_pr.remove(num_pr)
    return True


def _clear_direct_outline_level(paragraph) -> bool:
    p_pr = paragraph._element.find(qn("w:pPr"))
    if p_pr is None:
        return False
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        return False
    p_pr.remove(outline)
    return True


def _heading_cfg(style_spec: dict[str, Any], level: int) -> dict[str, Any]:
    heading = style_spec.get("heading", {})
    if str(level) in heading:
        return heading[str(level)]
    if heading:
        return heading[sorted(heading.keys(), key=lambda value: int(value))[-1]]
    return {}


def _apply_heading_run_format(paragraph, cfg: dict[str, Any]) -> None:
    if not cfg:
        return
    for run in paragraph.runs:
        if "en_font" in cfg:
            run.font.name = str(cfg["en_font"])
        if "size_pt" in cfg:
            run.font.size = Pt(float(cfg["size_pt"]))
        run.font.bold = bool(cfg.get("bold", False))
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        zh_font = str(cfg.get("zh_font") or cfg.get("en_font") or "")
        en_font = str(cfg.get("en_font") or cfg.get("zh_font") or "")
        if zh_font or en_font:
            r_fonts.set(qn("w:eastAsia"), zh_font)
            r_fonts.set(qn("w:ascii"), en_font)
            r_fonts.set(qn("w:hAnsi"), en_font)
            r_fonts.set(qn("w:cs"), en_font)


def _header_cleaned(docx_path: Path) -> bool:
    residual_terms = ("投标文件-技术卷",)
    text_parts: list[str] = []
    with zipfile.ZipFile(docx_path, "r") as archive:
        for name in archive.namelist():
            if not (name.startswith("word/header") and name.endswith(".xml")):
                continue
            text_parts.append(archive.read(name).decode("utf-8", errors="replace"))
    header_text = "\n".join(text_parts)
    return not any(term in header_text for term in residual_terms)


def _infer_matches_from_doc(outline_items: list[dict[str, Any]], heading_texts: set[str]):
    normalized_headings = {_normalize_match_text(text) for text in heading_texts if text}
    matched: list[dict[str, Any]] = []
    unmatched: list[dict[str, Any]] = []
    for item in outline_items:
        record = _outline_report_item(item)
        keys = {_normalize_match_text(_format_heading_text(item)), _normalize_match_text(str(item.get("title") or ""))}
        if normalized_headings & {key for key in keys if key}:
            matched.append(record)
        else:
            unmatched.append(record)
    return matched, unmatched


def _heading_level_counts(scan: dict[str, Any]) -> dict[str, int]:
    result: dict[str, int] = {}
    for key, value in dict(scan.get("heading_counts") or {}).items():
        level = _parse_heading_count_key(key)
        if level is None:
            continue
        result[str(level)] = int(value)
    return {str(level): result.get(str(level), 0) for level in range(1, 5)}


def _parse_heading_count_key(value: Any) -> int | None:
    match = re.search(r"(\d+)$", str(value or "").strip())
    if not match:
        return None
    try:
        return int(match.group(1))
    except ValueError:
        return None


def _normalize_section_orientations(doc: Document) -> dict[str, Any]:
    summary = {"sections": len(doc.sections), "portrait": 0, "landscape": 0, "normalized": 0}
    for section in doc.sections:
        changed = False
        width = int(section.page_width or 0)
        height = int(section.page_height or 0)
        declared = _section_declared_orientation(section)
        is_landscape = declared == "landscape" or (width > height > 0)

        if is_landscape:
            summary["landscape"] += 1
            if declared != "landscape":
                section.orientation = WD_ORIENT.LANDSCAPE
                changed = True
            if width < height:
                old_width = section.page_width
                old_height = section.page_height
                section.page_width = old_height
                section.page_height = old_width
                changed = True
        else:
            summary["portrait"] += 1
            if declared == "landscape":
                section.orientation = WD_ORIENT.PORTRAIT
                changed = True
            if width > height > 0:
                old_width = section.page_width
                old_height = section.page_height
                section.page_width = old_height
                section.page_height = old_width
                changed = True

        if changed:
            summary["normalized"] += 1
    return summary


def _section_orientation_summary(doc: Document) -> dict[str, Any]:
    summary = {"sections": len(doc.sections), "portrait": 0, "landscape": 0, "normalized": 0}
    for section in doc.sections:
        width = int(section.page_width or 0)
        height = int(section.page_height or 0)
        declared = _section_declared_orientation(section)
        if declared == "landscape" or (width > height > 0):
            summary["landscape"] += 1
        else:
            summary["portrait"] += 1
    return summary


def _section_declared_orientation(section) -> str:
    sect_pr = getattr(section, "_sectPr", None)
    if sect_pr is None:
        return ""
    pg_sz = sect_pr.find(qn("w:pgSz"))
    if pg_sz is None:
        return ""
    return str(pg_sz.get(qn("w:orient")) or "").lower()


def _structural_risks(
    *,
    scan: dict[str, Any],
    toc_present: bool,
    header_cleaned: bool,
    outline_count: int,
    unmatched_count: int,
) -> list[str]:
    risks: list[str] = []
    if outline_count == 0:
        risks.append("outline 为空，未能定位任何技术标标题。")
    if unmatched_count:
        risks.append("存在未匹配标题；技术标 cleaner 未插入或删除正文，请人工核对缺失章节。")
    if not toc_present:
        risks.append("未检测到 TOC 域，目录页需要人工补充。")
    if not header_cleaned:
        risks.append("页眉仍存在技术卷残留，请人工检查页眉。")
    if scan.get("placeholders"):
        risks.append("存在待填写或缺失占位符，请在共创阶段补齐。")
    if scan.get("dup_alerts"):
        risks.append("存在相邻重复标题，请检查素材首标题和目录标题是否重复。")
    if scan.get("invalid_h1"):
        risks.append("存在异常一级标题，请检查章节层级。")
    if scan.get("empty_leaf_headings"):
        risks.append("存在空叶子章节，请检查是否缺正文或素材未拼入。")
    return _dedupe(risks)


def _build_warnings(report: dict[str, Any], scan: dict[str, Any]) -> list[dict[str, Any]]:
    warnings: list[dict[str, Any]] = []

    def add(code: str, message: str, count: int) -> None:
        if count > 0:
            warnings.append({"code": code, "message": message, "count": count})

    if int(report.get("outlineCount") or 0) == 0:
        add("outline_empty", "outline 为空，未能定位任何技术标标题。", 1)
    add(
        "unmatched_heading",
        "存在未匹配标题；技术标 cleaner 未插入或删除正文，请人工核对缺失章节。",
        len(report.get("unmatchedHeadings") or []),
    )
    if not report.get("tocPresent", False):
        add("toc_missing", "未检测到 TOC 域，目录页需要人工补充。", 1)
    if not report.get("headerCleaned", False):
        add("header_not_cleaned", "页眉仍存在技术卷残留，请人工检查页眉。", 1)
    add(
        "placeholder",
        "存在待填写或缺失占位符，请在共创阶段补齐。",
        int(report.get("placeholderCount") or 0),
    )
    add(
        "duplicate_heading",
        "存在相邻重复标题，请检查素材首标题和目录标题是否重复。",
        len(scan.get("dup_alerts") or []),
    )
    add("invalid_h1", "存在异常一级标题，请检查章节层级。", len(scan.get("invalid_h1") or []))
    add(
        "empty_leaf_heading",
        "存在空叶子章节，请检查是否缺正文或素材未拼入。",
        len(scan.get("empty_leaf_headings") or []),
    )
    return warnings


def _dedupe(items: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for item in items:
        if item in seen:
            continue
        seen.add(item)
        result.append(item)
    return result


def _render_report(report: dict[str, Any], output_path: Path, outline_path: Path) -> str:
    lines = [
        "# 技术标格式清洗报告",
        "",
        f"- 输出文件：`{output_path}`",
        f"- outline 文件：`{outline_path}`",
        f"- outline 总数：{report['outlineCount']}",
        f"- 成功匹配标题数：{report['matchedHeadingCount']}",
        f"- 未匹配标题数：{len(report['unmatchedHeadings'])}",
        f"- 正文内部标题识别数：{report.get('internalHeadingCount', 0)}",
        f"- TOC 是否存在：{'是' if report['tocPresent'] else '否'}",
        f"- 页眉是否清理：{'是' if report['headerCleaned'] else '否'}",
        f"- 占位符数量：{report['placeholderCount']}",
        "- 标题层级统计："
        + " / ".join(
            f"{level}级 {count}"
            for level, count in sorted(report.get("headingLevelCounts", {}).items(), key=lambda item: int(item[0]))
        ),
        "- 横竖版小节："
        + f"竖版 {report.get('orientation', {}).get('portrait', 0)} / "
        + f"横版 {report.get('orientation', {}).get('landscape', 0)} / "
        + f"已规范化 {report.get('orientation', {}).get('normalized', 0)}",
        "",
        "## 未匹配标题清单",
        "",
    ]
    if report["unmatchedHeadings"]:
        for item in report["unmatchedHeadings"]:
            lines.append(f"- {_format_heading_text(item)}（id={item.get('id')}, level={item.get('level')}）")
    else:
        lines.append("- 无")
    lines.extend(["", "## 格式风险提示", ""])
    for risk in report["formatRisks"]:
        lines.append(f"- {risk}")
    lines.append("")
    return "\n".join(lines)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Run technical bid format cleaner from manifest")
    parser.add_argument("manifest")
    parser.add_argument("--response", choices=("summary", "details"), default="summary")
    args = parser.parse_args(argv)
    result = run_manifest(args.manifest, response=args.response)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
