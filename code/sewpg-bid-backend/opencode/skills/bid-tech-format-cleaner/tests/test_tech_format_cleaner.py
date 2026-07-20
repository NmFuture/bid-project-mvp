import base64
import copy
import importlib
import importlib.util
import json
import shutil
import sys
import tempfile
import unittest
from pathlib import Path
from types import ModuleType

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm


ROOT = Path(__file__).resolve().parents[1]
RUNNER_PATH = ROOT / "scripts" / "run_from_manifest.py"
SCRIPT_DIR = str(ROOT / "scripts")


def _run_manifest(manifest_path: Path, *, response: str = "summary"):
    module_name = "bid_tech_format_cleaner_run_from_manifest"
    inserted = False
    if SCRIPT_DIR not in sys.path:
        sys.path.insert(0, SCRIPT_DIR)
        inserted = True
    try:
        spec = importlib.util.spec_from_file_location(module_name, RUNNER_PATH)
        if spec is None or spec.loader is None:
            raise ImportError(f"cannot load runner: {RUNNER_PATH}")
        module = importlib.util.module_from_spec(spec)
        sys.modules[module_name] = module
        spec.loader.exec_module(module)
        return module.run_manifest(manifest_path, response=response)
    finally:
        sys.modules.pop(module_name, None)
        if inserted:
            sys.path.remove(SCRIPT_DIR)


def _write_style(path: Path) -> None:
    style = {
        "schema_version": "tech_heading_style.test",
        "page": {
            "top_cm": 2.54,
            "bottom_cm": 2.54,
            "left_cm": 3.18,
            "right_cm": 3.18,
            "header_top_cm": 1.0,
            "footer_bottom_cm": 0.6,
            "orientation": "portrait",
        },
        "heading": {
            str(level): {
                "zh_font": "等线",
                "en_font": "Times New Roman",
                "size_pt": 12 + max(0, 4 - level),
                "bold": True,
                "align": "left",
                "space_before_pt": 6,
                "space_after_pt": 6,
                "line_spacing": 1.5,
                "first_line_indent_chars": 0,
                "left_indent_cm": 0,
            }
            for level in range(1, 5)
        },
    }
    path.write_text(json.dumps(style, ensure_ascii=False, indent=2), encoding="utf-8")


def _write_outline(path: Path) -> None:
    outline = {
        "schema_version": "tech_bid_outline.v1",
        "sections": [
            {
                "id": "h1",
                "number": "1",
                "title": "一级标题",
                "level": 1,
                "children": [
                    {
                        "id": "h2",
                        "number": "1.1",
                        "title": "二级标题",
                        "level": 2,
                        "children": [
                            {
                                "id": "h3",
                                "number": "1.1.1",
                                "title": "三级标题",
                                "level": 3,
                                "children": [
                                    {
                                        "id": "h4",
                                        "number": "1.1.1.1",
                                        "title": "四级标题",
                                        "level": 4,
                                        "children": [],
                                    }
                                ],
                            }
                        ],
                    }
                ],
            }
        ],
    }
    path.write_text(json.dumps(outline, ensure_ascii=False, indent=2), encoding="utf-8")


def _write_deep_heading_docx(path: Path) -> None:
    doc = Document()
    for text in ("1 一级标题", "1.1 二级标题", "1.1.1 三级标题", "1.1.1.1 四级标题"):
        doc.add_paragraph(text)
        doc.add_paragraph("正文内容")
    doc.save(path)


def _write_portrait_and_landscape_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("1 一级标题")
    doc.add_paragraph("竖版正文")
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    doc.add_paragraph("1.1 二级标题")
    doc.add_paragraph("横版正文")
    doc.save(path)


def _write_manifest(
    path: Path,
    *,
    input_docx: Path,
    outline_path: Path,
    output_docx: Path,
    style_path: Path,
) -> None:
    path.write_text(
        json.dumps(
            {
                "inputFile": str(input_docx),
                "outlineFile": str(outline_path),
                "outputFile": str(output_docx),
                "projectName": "测试项目",
                "styleSpecPath": str(style_path),
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )


def _write_full_format_docx(path: Path) -> None:
    image_path = path.with_suffix(".png")
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
    )
    doc = Document()
    doc.add_paragraph("1 一级标题")
    doc.add_paragraph("正文内容")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "表格内容"
    doc.add_paragraph("表1 技术参数")
    doc.add_picture(str(image_path), width=Cm(1))
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    doc.add_paragraph("横版正文")
    doc.save(path)


def _write_full_style(
    path: Path,
    *,
    body_size: float,
    margin: float,
    insert_toc: bool = True,
    toc_page_break_after: bool = True,
) -> None:
    style = {
        "page": {
            "top_cm": margin,
            "bottom_cm": margin,
            "left_cm": margin,
            "right_cm": margin,
            "header_top_cm": 1,
            "footer_bottom_cm": 0.6,
        },
        "heading": {
            "1": {
                "zh_font": "等线",
                "en_font": "Arial",
                "size_pt": body_size + 4,
                "bold": True,
                "align": "center",
                "space_before_pt": 6,
                "space_after_pt": 6,
                "line_spacing": 1.5,
                "first_line_indent_chars": 0,
                "left_indent_cm": 0,
            }
        },
        "body": {
            "zh_font": "宋体",
            "en_font": "Arial",
            "size_pt": body_size,
            "bold": False,
            "align": "both",
            "space_before_pt": 0,
            "space_after_pt": 0,
            "line_spacing": 1.5,
            "first_line_indent_chars": 2,
        },
        "table_cell": {
            "zh_font": "宋体",
            "en_font": "Arial",
            "size_pt": body_size - 1,
            "bold": False,
            "align": "center",
            "line_spacing": 1,
            "table_align": "center",
        },
        "caption": {
            "zh_font": "等线",
            "en_font": "Arial",
            "size_pt": body_size,
            "bold": False,
            "align": "center",
            "line_spacing": 1.5,
        },
        "header": {
            "zh_font": "宋体",
            "en_font": "Arial",
            "size_pt": 8,
            "bold": False,
            "align": "right",
            "line_spacing": 1,
            "text_template": "自定义页眉-{projectName}",
        },
        "toc": {"insert_when_missing": insert_toc, "page_break_after": toc_page_break_after},
    }
    path.write_text(json.dumps(style, ensure_ascii=False, indent=2), encoding="utf-8")


def _toc_is_followed_by_page_break(path: Path) -> bool:
    doc = Document(str(path))
    for paragraph in doc.paragraphs:
        has_toc = any("TOC" in (node.text or "").upper() for node in paragraph._element.iter(qn("w:instrText")))
        if not has_toc:
            continue
        sibling = paragraph._element.getnext()
        if sibling is None or sibling.tag != qn("w:p"):
            return False
        return bool(sibling.findall(".//" + qn("w:br")))
    raise AssertionError("测试文档中未找到 TOC 域")


def _write_multi_paragraph_toc_docx(path: Path) -> None:
    doc = Document()
    doc.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
    doc.styles.add_style("TOC 2", WD_STYLE_TYPE.PARAGRAPH)
    field = doc.add_paragraph()
    for field_type in ("begin", "separate"):
        run = OxmlElement("w:r")
        node = OxmlElement("w:fldChar")
        node.set(qn("w:fldCharType"), field_type)
        run.append(node)
        field._element.append(run)
        if field_type == "begin":
            instruction_run = OxmlElement("w:r")
            instruction = OxmlElement("w:instrText")
            instruction.text = ' TOC \\o "1-3" '
            instruction_run.append(instruction)
            field._element.append(instruction_run)
    doc.add_paragraph("第一章 ................................ 1", style="TOC 1")
    toc_end = doc.add_paragraph("1.1 节 ................................ 2", style="TOC 2")
    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    toc_end._element.append(end_run)
    doc.add_paragraph("用户正文")
    user_break = doc.add_paragraph("用户分页")
    user_break.add_run().add_break(WD_BREAK.PAGE)
    doc.save(path)


def _page_break_count_after_toc_end(path: Path) -> int:
    doc = Document(str(path))
    for paragraph in doc.paragraphs:
        has_end = any(
            node.get(qn("w:fldCharType")) == "end" for node in paragraph._element.iter(qn("w:fldChar"))
        )
        if not has_end:
            continue
        count = 0
        sibling = paragraph._element.getnext()
        while sibling is not None and sibling.tag == qn("w:p") and sibling.findall(".//" + qn("w:br")):
            count += 1
            sibling = sibling.getnext()
        return count
    raise AssertionError("测试文档中未找到 TOC field end")


def _user_page_break_is_present(path: Path) -> bool:
    doc = Document(str(path))
    paragraph = next(item for item in doc.paragraphs if item.text == "用户分页")
    return bool(paragraph._element.findall(".//" + qn("w:br")))


def _write_direct_format_heading_docx(path: Path) -> None:
    doc = Document()
    heading = doc.add_paragraph(style="Heading 1")
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), "42")
    bookmark_start.set(qn("w:name"), "heading-anchor")
    heading._element.append(bookmark_start)
    for text in ("直接", "格式标题"):
        run = heading.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Cm(1.1)
        run.font.bold = False
    field_run = OxmlElement("w:r")
    field = OxmlElement("w:fldChar")
    field.set(qn("w:fldCharType"), "begin")
    field_run.append(field)
    heading._element.append(field_run)
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), "42")
    heading._element.append(bookmark_end)
    doc.add_paragraph("正文")
    doc.save(path)


def _write_empty_outline(path: Path) -> None:
    path.write_text(json.dumps({"schema_version": "tech_bid_outline.v1", "sections": []}), encoding="utf-8")


def _write_sdt_wrapped_toc_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("用户正文")
    body = doc.element.body
    sdt = OxmlElement("w:sdt")
    sdt_content = OxmlElement("w:sdtContent")
    sdt.append(sdt_content)

    field = OxmlElement("w:p")
    for field_type in ("begin", "separate"):
        run = OxmlElement("w:r")
        node = OxmlElement("w:fldChar")
        node.set(qn("w:fldCharType"), field_type)
        run.append(node)
        field.append(run)
        if field_type == "begin":
            instruction_run = OxmlElement("w:r")
            instruction = OxmlElement("w:instrText")
            instruction.text = ' TOC \\o "1-3" '
            instruction_run.append(instruction)
            field.append(instruction_run)
    sdt_content.append(field)

    for style_id, text, include_end in (
        ("TOC1", "第一章 ................................ 1", False),
        ("TOC2", "1.1 节 ................................ 2", True),
    ):
        paragraph = OxmlElement("w:p")
        p_pr = OxmlElement("w:pPr")
        p_style = OxmlElement("w:pStyle")
        p_style.set(qn("w:val"), style_id)
        p_pr.append(p_style)
        paragraph.append(p_pr)
        run = OxmlElement("w:r")
        text_node = OxmlElement("w:t")
        text_node.text = text
        run.append(text_node)
        paragraph.append(run)
        if include_end:
            end_run = OxmlElement("w:r")
            end = OxmlElement("w:fldChar")
            end.set(qn("w:fldCharType"), "end")
            end_run.append(end)
            paragraph.append(end_run)
        sdt_content.append(paragraph)

    first_body_child = next(child for child in body if child.tag != qn("w:sectPr"))
    first_body_child.addprevious(sdt)
    user_break = OxmlElement("w:p")
    user_run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    user_run.append(br)
    user_break.append(user_run)
    sdt.addnext(user_break)
    doc.save(path)


def _sdt_following_break_counts(path: Path) -> tuple[int, int]:
    doc = Document(str(path))
    sdt = doc.element.body.find(qn("w:sdt"))
    if sdt is None:
        raise AssertionError("测试文档中未找到 TOC sdt")
    marked = 0
    unmarked = 0
    sibling = sdt.getnext()
    while sibling is not None and sibling.tag == qn("w:p") and sibling.findall(".//" + qn("w:br")):
        is_marked = any(
            node.get(qn("w:name")) == "_TECH_FORMAT_CLEANER_TOC_BREAK"
            for node in sibling.iter(qn("w:bookmarkStart"))
        )
        marked += int(is_marked)
        unmarked += int(not is_marked)
        sibling = sibling.getnext()
    return marked, unmarked


class TechFormatCleanerTest(unittest.TestCase):
    def setUp(self) -> None:
        self.tmp_dir = Path(tempfile.mkdtemp(prefix="tech-format-cleaner-"))

    def tearDown(self) -> None:
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def test_promotes_level_three_and_four_headings(self):
        input_docx = self.tmp_dir / "input.docx"
        outline_path = self.tmp_dir / "outline.json"
        style_path = self.tmp_dir / "style.json"
        output_docx = self.tmp_dir / "output.docx"
        manifest_path = self.tmp_dir / "manifest.json"

        _write_deep_heading_docx(input_docx)
        _write_outline(outline_path)
        _write_style(style_path)
        _write_manifest(
            manifest_path,
            input_docx=input_docx,
            outline_path=outline_path,
            output_docx=output_docx,
            style_path=style_path,
        )

        result = _run_manifest(manifest_path, response="details")

        self.assertEqual(result["summary"]["matchedHeadingCount"], 4)
        self.assertEqual(result["summary"]["headingLevelCounts"]["3"], 1)
        self.assertEqual(result["summary"]["headingLevelCounts"]["4"], 1)

        output = Document(str(output_docx))
        style_by_text = {paragraph.text.strip(): paragraph.style.name for paragraph in output.paragraphs}
        self.assertEqual(style_by_text["1.1.1 三级标题"], "Heading 3")
        self.assertEqual(style_by_text["1.1.1.1 四级标题"], "Heading 4")

    def test_preserves_landscape_sections_during_cleaning(self):
        input_docx = self.tmp_dir / "input.docx"
        outline_path = self.tmp_dir / "outline.json"
        style_path = self.tmp_dir / "style.json"
        output_docx = self.tmp_dir / "output.docx"
        manifest_path = self.tmp_dir / "manifest.json"

        _write_portrait_and_landscape_docx(input_docx)
        _write_outline(outline_path)
        _write_style(style_path)
        _write_manifest(
            manifest_path,
            input_docx=input_docx,
            outline_path=outline_path,
            output_docx=output_docx,
            style_path=style_path,
        )

        result = _run_manifest(manifest_path, response="details")

        self.assertGreaterEqual(result["summary"]["orientation"]["portrait"], 1)
        self.assertGreaterEqual(result["summary"]["orientation"]["landscape"], 1)

        output = Document(str(output_docx))
        has_landscape = any(int(section.page_width) > int(section.page_height) for section in output.sections)
        has_portrait = any(int(section.page_height) >= int(section.page_width) for section in output.sections)
        self.assertTrue(has_landscape)
        self.assertTrue(has_portrait)

    def test_infers_deep_appendix_heading_levels_from_numbers(self):
        input_docx = self.tmp_dir / "input.docx"
        outline_path = self.tmp_dir / "outline.json"
        style_path = self.tmp_dir / "style.json"
        output_docx = self.tmp_dir / "output.docx"
        manifest_path = self.tmp_dir / "manifest.json"

        doc = Document()
        doc.add_paragraph("附表B.1.1 风力发电机组供货范围清单")
        doc.add_paragraph("三级附表内容")
        doc.add_paragraph("附表B.1.1.1 叶片供货范围")
        doc.add_paragraph("四级附表内容")
        doc.save(input_docx)
        outline_path.write_text(
            json.dumps(
                {
                    "schema_version": "tech_bid_outline.v1",
                    "sections": [
                        {
                            "id": "app-3",
                            "number": "附表B.1.1",
                            "title": "风力发电机组供货范围清单",
                            "level": 1,
                            "children": [],
                        },
                        {
                            "id": "app-4",
                            "number": "附表B.1.1.1",
                            "title": "叶片供货范围",
                            "level": 1,
                            "children": [],
                        },
                    ],
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        _write_style(style_path)
        _write_manifest(
            manifest_path,
            input_docx=input_docx,
            outline_path=outline_path,
            output_docx=output_docx,
            style_path=style_path,
        )

        result = _run_manifest(manifest_path, response="details")

        self.assertEqual(result["summary"]["headingLevelCounts"]["3"], 1)
        self.assertEqual(result["summary"]["headingLevelCounts"]["4"], 1)
        output = Document(str(output_docx))
        style_by_text = {paragraph.text.strip(): paragraph.style.name for paragraph in output.paragraphs}
        self.assertEqual(style_by_text["附表B.1.1 风力发电机组供货范围清单"], "Heading 3")
        self.assertEqual(style_by_text["附表B.1.1.1 叶片供货范围"], "Heading 4")

    def test_promotes_internal_material_headings_without_promoting_captions(self):
        input_docx = self.tmp_dir / "input.docx"
        outline_path = self.tmp_dir / "outline.json"
        style_path = self.tmp_dir / "style.json"
        output_docx = self.tmp_dir / "output.docx"
        manifest_path = self.tmp_dir / "manifest.json"

        doc = Document()
        doc.add_paragraph("1 一级标题")
        doc.add_paragraph("1.1 二级标题")
        h3 = doc.add_paragraph()
        h3.add_run("供应链能力保障").bold = True
        doc.add_paragraph("供应链正文内容")
        h4 = doc.add_paragraph("（1）签订产能协议")
        doc.add_paragraph("四级正文内容")
        caption = doc.add_paragraph()
        caption.add_run("表C.1 总体技术参数与规格").bold = True
        doc.save(input_docx)
        _write_outline(outline_path)
        _write_style(style_path)
        _write_manifest(
            manifest_path,
            input_docx=input_docx,
            outline_path=outline_path,
            output_docx=output_docx,
            style_path=style_path,
        )

        result = _run_manifest(manifest_path, response="details")

        self.assertEqual(result["summary"]["internalHeadingCount"], 1)
        output = Document(str(output_docx))
        style_by_text = {paragraph.text.strip(): paragraph.style.name for paragraph in output.paragraphs}
        self.assertEqual(style_by_text["供应链能力保障"], "Heading 3")
        self.assertNotEqual(style_by_text["（1）签订产能协议"], "Heading 4")
        self.assertNotEqual(style_by_text["表C.1 总体技术参数与规格"], "Heading 3")


    def test_applies_full_style_without_overwriting_input_or_losing_content(self):
        input_docx = self.tmp_dir / "input.docx"
        outline_path = self.tmp_dir / "outline.json"
        style_path = self.tmp_dir / "style.json"
        output_docx = self.tmp_dir / "output.docx"
        manifest_path = self.tmp_dir / "manifest.json"
        _write_full_format_docx(input_docx)
        _write_outline(outline_path)
        _write_full_style(style_path, body_size=16, margin=4, insert_toc=False)
        _write_manifest(
            manifest_path,
            input_docx=input_docx,
            outline_path=outline_path,
            output_docx=output_docx,
            style_path=style_path,
        )
        input_bytes = input_docx.read_bytes()

        result = _run_manifest(manifest_path, response="details")

        self.assertEqual(input_docx.read_bytes(), input_bytes)
        output = Document(str(output_docx))
        body = next(paragraph for paragraph in output.paragraphs if paragraph.text == "正文内容")
        self.assertAlmostEqual(body.runs[0].font.size.pt, 16)
        self.assertAlmostEqual(output.sections[0].top_margin.cm, 4, places=1)
        self.assertAlmostEqual(output.tables[0].cell(0, 0).paragraphs[0].runs[0].font.size.pt, 15)
        caption = next(paragraph for paragraph in output.paragraphs if paragraph.text == "表1 技术参数")
        self.assertAlmostEqual(caption.runs[0].font.size.pt, 16)
        self.assertEqual(len(output.inline_shapes), 1)
        self.assertTrue(any(int(section.page_width) > int(section.page_height) for section in output.sections))
        header_paragraph = next(
            paragraph
            for section in output.sections
            for paragraph in section.header.paragraphs
            if paragraph.text.strip()
        )
        self.assertEqual(header_paragraph.text, "自定义页眉-测试项目")
        self.assertAlmostEqual(header_paragraph.runs[0].font.size.pt, 8)
        self.assertFalse(result["summary"]["tocInserted"])
        self.assertFalse(result["summary"]["tocPresent"])
        self.assertIn("warnings", result["summary"])
        for warning in result["summary"]["warnings"]:
            self.assertEqual(set(warning), {"code", "message", "count"})

    def test_default_custom_default_reapplies_body_style(self):
        input_docx = self.tmp_dir / "chain-input.docx"
        outline_path = self.tmp_dir / "chain-outline.json"
        default_style = self.tmp_dir / "chain-default.json"
        custom_style = self.tmp_dir / "chain-custom.json"
        _write_full_format_docx(input_docx)
        _write_outline(outline_path)
        _write_full_style(default_style, body_size=12, margin=2.54)
        _write_full_style(custom_style, body_size=16, margin=4)

        current_input = input_docx
        for index, (style_path, expected_size) in enumerate(
            ((default_style, 12), (custom_style, 16), (default_style, 12)),
            start=1,
        ):
            output_docx = self.tmp_dir / f"chain-output-{index}.docx"
            manifest_path = self.tmp_dir / f"chain-manifest-{index}.json"
            _write_manifest(
                manifest_path,
                input_docx=current_input,
                outline_path=outline_path,
                output_docx=output_docx,
                style_path=style_path,
            )
            previous_bytes = current_input.read_bytes()
            _run_manifest(manifest_path)
            self.assertEqual(current_input.read_bytes(), previous_bytes)
            output = Document(str(output_docx))
            body = next(paragraph for paragraph in output.paragraphs if paragraph.text == "正文内容")
            self.assertAlmostEqual(body.runs[0].font.size.pt, expected_size)
            self.assertEqual(len(output.inline_shapes), 1)
            self.assertEqual(output.tables[0].cell(0, 0).text, "表格内容")
            current_input = output_docx

    def test_existing_toc_page_break_can_be_disabled_and_restored(self):
        input_docx = self.tmp_dir / "toc-input.docx"
        outline_path = self.tmp_dir / "toc-outline.json"
        _write_full_format_docx(input_docx)
        _write_outline(outline_path)
        initial = Document(str(input_docx))
        initial.paragraphs[0].style = "Heading 1"
        initial.save(input_docx)
        current_input = input_docx

        for index, expected_page_break in enumerate((True, False, True), start=1):
            style_path = self.tmp_dir / f"toc-style-{index}.json"
            output_docx = self.tmp_dir / f"toc-output-{index}.docx"
            manifest_path = self.tmp_dir / f"toc-manifest-{index}.json"
            _write_full_style(
                style_path,
                body_size=12,
                margin=2.54,
                toc_page_break_after=expected_page_break,
            )
            _write_manifest(
                manifest_path,
                input_docx=current_input,
                outline_path=outline_path,
                output_docx=output_docx,
                style_path=style_path,
            )

            _run_manifest(manifest_path)

            self.assertEqual(_toc_is_followed_by_page_break(output_docx), expected_page_break)
            current_input = output_docx

    def test_warning_counts_use_structured_scan_counts(self):
        module_name = "bid_tech_format_cleaner_warning_test"
        spec = importlib.util.spec_from_file_location(module_name, RUNNER_PATH)
        self.assertIsNotNone(spec)
        self.assertIsNotNone(spec.loader)
        module = importlib.util.module_from_spec(spec)
        sys.modules[module_name] = module
        try:
            spec.loader.exec_module(module)
            warnings = module._build_warnings(
                {
                    "formatRisks": [
                        "存在相邻重复标题，请检查素材首标题和目录标题是否重复。",
                        "存在异常一级标题，请检查章节层级。",
                        "存在空叶子章节，请检查是否缺正文或素材未拼入。",
                    ],
                    "unmatchedHeadings": [],
                    "placeholderCount": 0,
                },
                {
                    "dup_alerts": ["重复1", "重复2"],
                    "invalid_h1": ["异常1", "异常2", "异常3"],
                    "empty_leaf_headings": ["空1", "空2", "空3", "空4"],
                },
            )
        finally:
            sys.modules.pop(module_name, None)

        counts = {warning["code"]: warning["count"] for warning in warnings}
        self.assertEqual(counts["duplicate_heading"], 2)
        self.assertEqual(counts["invalid_h1"], 3)
        self.assertEqual(counts["empty_leaf_heading"], 4)

    def test_multi_paragraph_toc_break_is_managed_after_field_results(self):
        input_docx = self.tmp_dir / "multi-toc-input.docx"
        outline_path = self.tmp_dir / "multi-toc-outline.json"
        _write_multi_paragraph_toc_docx(input_docx)
        _write_empty_outline(outline_path)
        current_input = input_docx

        for index, enabled in enumerate((True, False, True, True), start=1):
            style_path = self.tmp_dir / f"multi-toc-style-{index}.json"
            output_docx = self.tmp_dir / f"multi-toc-output-{index}.docx"
            manifest_path = self.tmp_dir / f"multi-toc-manifest-{index}.json"
            _write_full_style(
                style_path,
                body_size=12,
                margin=2.54,
                toc_page_break_after=enabled,
            )
            _write_manifest(
                manifest_path,
                input_docx=current_input,
                outline_path=outline_path,
                output_docx=output_docx,
                style_path=style_path,
            )

            _run_manifest(manifest_path)

            self.assertEqual(_page_break_count_after_toc_end(output_docx), 1 if enabled else 0)
            self.assertTrue(_user_page_break_is_present(output_docx))
            current_input = output_docx

    def test_sdt_wrapped_toc_break_is_managed_outside_content_control(self):
        input_docx = self.tmp_dir / "sdt-toc-input.docx"
        outline_path = self.tmp_dir / "sdt-toc-outline.json"
        _write_sdt_wrapped_toc_docx(input_docx)
        _write_empty_outline(outline_path)
        current_input = input_docx

        for index, enabled in enumerate((True, False, True), start=1):
            style_path = self.tmp_dir / f"sdt-toc-style-{index}.json"
            output_docx = self.tmp_dir / f"sdt-toc-output-{index}.docx"
            manifest_path = self.tmp_dir / f"sdt-toc-manifest-{index}.json"
            _write_full_style(
                style_path,
                body_size=12,
                margin=2.54,
                toc_page_break_after=enabled,
            )
            _write_manifest(
                manifest_path,
                input_docx=current_input,
                outline_path=outline_path,
                output_docx=output_docx,
                style_path=style_path,
            )

            _run_manifest(manifest_path)

            self.assertEqual(_sdt_following_break_counts(output_docx), (1 if enabled else 0, 1))
            output = Document(str(output_docx))
            self.assertIn("用户正文", [paragraph.text for paragraph in output.paragraphs])
            current_input = output_docx

    def test_heading_direct_run_format_is_reapplied_without_rebuilding_xml(self):
        input_docx = self.tmp_dir / "heading-input.docx"
        outline_path = self.tmp_dir / "heading-outline.json"
        default_style = self.tmp_dir / "heading-default.json"
        custom_style = self.tmp_dir / "heading-custom.json"
        _write_direct_format_heading_docx(input_docx)
        _write_empty_outline(outline_path)
        _write_full_style(default_style, body_size=12, margin=2.54)
        _write_full_style(custom_style, body_size=15, margin=2.54)
        custom = json.loads(custom_style.read_text(encoding="utf-8"))
        custom["heading"]["1"].update(
            {"zh_font": "微软雅黑", "en_font": "Calibri", "size_pt": 19, "bold": False}
        )
        custom_style.write_text(json.dumps(custom, ensure_ascii=False, indent=2), encoding="utf-8")
        current_input = input_docx

        for index, (style_path, en_font, zh_font, size_pt, bold) in enumerate(
            (
                (default_style, "Arial", "等线", 16, True),
                (custom_style, "Calibri", "微软雅黑", 19, False),
                (default_style, "Arial", "等线", 16, True),
            ),
            start=1,
        ):
            output_docx = self.tmp_dir / f"heading-output-{index}.docx"
            manifest_path = self.tmp_dir / f"heading-manifest-{index}.json"
            _write_manifest(
                manifest_path,
                input_docx=current_input,
                outline_path=outline_path,
                output_docx=output_docx,
                style_path=style_path,
            )

            _run_manifest(manifest_path)

            output = Document(str(output_docx))
            heading = next(paragraph for paragraph in output.paragraphs if paragraph.text == "直接格式标题")
            self.assertEqual(len(heading._element.findall(".//" + qn("w:bookmarkStart"))), 1)
            self.assertEqual(len(heading._element.findall(".//" + qn("w:bookmarkEnd"))), 1)
            self.assertEqual(len(heading._element.findall(".//" + qn("w:fldChar"))), 1)
            for run in (run for run in heading.runs if run.text):
                self.assertEqual(run.font.name, en_font)
                self.assertAlmostEqual(run.font.size.pt, size_pt)
                self.assertEqual(run.font.bold, bold)
                fonts = run._element.get_or_add_rPr().find(qn("w:rFonts"))
                self.assertEqual(fonts.get(qn("w:eastAsia")), zh_font)
            current_input = output_docx

    def test_primary_first_and_even_headers_are_formatted_once_per_part(self):
        input_docx = self.tmp_dir / "headers-input.docx"
        outline_path = self.tmp_dir / "headers-outline.json"
        style_path = self.tmp_dir / "headers-style.json"
        _write_empty_outline(outline_path)
        _write_full_style(style_path, body_size=12, margin=2.54)
        doc = Document()
        doc.settings.odd_and_even_pages_header_footer = True
        first_section = doc.sections[0]
        first_section.different_first_page_header_footer = True
        first_section.header.paragraphs[0].text = "旧主页面眉"
        first_section.first_page_header.paragraphs[0].text = "旧首页页眉"
        first_section.even_page_header.paragraphs[0].text = ""
        second_section = doc.add_section(WD_SECTION.NEW_PAGE)
        second_section.header.is_linked_to_previous = True
        second_section.first_page_header.is_linked_to_previous = True
        second_section.even_page_header.is_linked_to_previous = True
        doc.add_paragraph("正文")
        doc.save(input_docx)
        current_input = input_docx

        for index in (1, 2):
            output_docx = self.tmp_dir / f"headers-output-{index}.docx"
            manifest_path = self.tmp_dir / f"headers-manifest-{index}.json"
            _write_manifest(
                manifest_path,
                input_docx=current_input,
                outline_path=outline_path,
                output_docx=output_docx,
                style_path=style_path,
            )

            _run_manifest(manifest_path)

            output = Document(str(output_docx))
            seen = set()
            for section in output.sections:
                for header in (section.header, section.first_page_header, section.even_page_header):
                    marker = id(header._element)
                    if marker in seen:
                        continue
                    seen.add(marker)
                    nonempty = [paragraph for paragraph in header.paragraphs if paragraph.text.strip()]
                    self.assertEqual([paragraph.text for paragraph in nonempty], ["自定义页眉-测试项目"])
                    self.assertAlmostEqual(nonempty[0].runs[0].font.size.pt, 8)
                    self.assertEqual(nonempty[0].alignment, WD_ALIGN_PARAGRAPH.RIGHT)
            current_input = output_docx

    def _load_format_services(self):
        definitions = {
            "app.services.onlyoffice_documents": ("document_path", lambda _project_id: None),
            "app.services.workspace_project_access": ("get_workspace_project_runtime_state", lambda *_a, **_k: {}),
            "app.services.workspace_artifacts": ("technical_workspace_stage_dir", lambda *_a, **_k: None),
        }
        stubs = {}
        for module_name, (attribute, value) in definitions.items():
            module = ModuleType(module_name)
            setattr(module, attribute, value)
            stubs[module_name] = module
        tracked_names = (*stubs, "app.services.technical_document_format")
        missing = object()
        originals = {name: sys.modules.get(name, missing) for name in tracked_names}
        try:
            sys.modules.update(stubs)
            sys.modules.pop("app.services.technical_document_format", None)
            format_module = importlib.import_module("app.services.technical_document_format")
            state_module = importlib.import_module("app.services.technical_document_state")
        finally:
            for name, original in originals.items():
                if original is missing:
                    sys.modules.pop(name, None)
                else:
                    sys.modules[name] = original
        return format_module, state_module

    def test_invalid_overrides_fall_back_to_defaults(self):
        module_names = (
            "app.services.onlyoffice_documents",
            "app.services.workspace_project_access",
            "app.services.workspace_artifacts",
            "app.services.technical_document_format",
        )
        before = {name: sys.modules.get(name) for name in module_names}
        format_module, _state_module = self._load_format_services()
        self.assertEqual({name: sys.modules.get(name) for name in module_names}, before)
        base_path = (
            Path(__file__).resolve().parents[4]
            / "opencode"
            / "skills"
            / "bid-tech-assembler"
            / "references"
            / "heading_style.json"
        )
        default_spec = json.loads(base_path.read_text(encoding="utf-8"))
        style_path = format_module._prepare_technical_format_style_spec(
            "custom",
            {"bodySizePt": 99, "pageTopCm": -1, "heading1Bold": "false", "insertToc": "false"},
            self.tmp_dir,
        )
        actual = json.loads(style_path.read_text(encoding="utf-8"))
        self.assertEqual(actual["body"]["size_pt"], default_spec["body"]["size_pt"])
        self.assertEqual(actual["page"]["top_cm"], default_spec["page"]["top_cm"])
        self.assertEqual(actual["heading"]["1"]["bold"], default_spec["heading"]["1"]["bold"])
        self.assertEqual(actual["toc"]["insert_when_missing"], default_spec["toc"]["insert_when_missing"])

        standard_path = format_module._prepare_technical_format_style_spec(
            "standard",
            {"bodySizePt": 16, "pageTopCm": 4},
            self.tmp_dir,
        )
        self.assertEqual(json.loads(standard_path.read_text(encoding="utf-8")), default_spec)

        valid_path = format_module._prepare_technical_format_style_spec(
            "custom",
            {"bodySizePt": 13, "pageTopCm": 3.5, "insertToc": False},
            self.tmp_dir,
        )
        valid = json.loads(valid_path.read_text(encoding="utf-8"))
        self.assertEqual(valid["body"]["size_pt"], 13)
        self.assertEqual(valid["page"]["top_cm"], 3.5)
        self.assertFalse(valid["toc"]["insert_when_missing"])

    def test_latest_format_selection_is_persisted_in_document_state(self):
        _format_module, state_module = self._load_format_services()
        project = {
            "id": "TECH-FORMAT",
            "document_state": {"version": 1, "onlyoffice": {}},
            "fill_state": {},
        }
        format_result = {
            "preset": "custom",
            "label": "自定义格式",
            "styleOverrides": {"bodySizePt": 13},
            "summary": {"matchedHeadingCount": 2, "warnings": []},
        }
        state = state_module.apply_technical_document_format_state(
            project,
            copy.deepcopy(format_result),
            updated_at="2026-07-20T00:00:00Z",
        )
        self.assertEqual(state["technicalFormatPreset"], "custom")
        self.assertEqual(state["technicalFormatStyleOverrides"], {"bodySizePt": 13})
        self.assertEqual(state["technicalFormatSummary"], {"matchedHeadingCount": 2, "warnings": []})


if __name__ == "__main__":
    unittest.main()
