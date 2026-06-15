#!/usr/bin/env python3
"""Business bid table fill runner.

This first production-safe version keeps the target Word structure, extracts
simple facts from the project fact table and selected source materials, and
fills obvious blank/placeholder table cells.
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt
from openpyxl import load_workbook


SCHEMA_VERSION = "bid-business-table-fill-v1"
FILL_HEADERS = ("投标人响应", "投标响应", "响应值", "响应内容", "填写内容", "填报内容", "报价", "金额", "说明", "备注", "内容")
FIELD_HEADERS = ("项目", "名称", "字段", "条款", "指标", "参数", "事项", "内容")
PLACEHOLDER_RE = re.compile(r"(\[待填写[:：]?[^\]]*\]|_{3,}|（\s*）|\(\s*\))")


def now_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def norm(text: Any) -> str:
    return re.sub(r"[\s　：:（）()【】\[\]<>《》]+", "", str(text or "")).strip().lower()


def clean(text: Any) -> str:
    return re.sub(r"\s+", " ", str(text or "")).strip()


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def fact_map(manifest: dict[str, Any]) -> dict[str, str]:
    result: dict[str, str] = {}
    facts = manifest.get("facts") if isinstance(manifest.get("facts"), dict) else {}
    for key, value in facts.items():
        if clean(value):
            result[norm(key)] = clean(value)
    table = manifest.get("projectFactTable") if isinstance(manifest.get("projectFactTable"), dict) else {}
    for field in table.get("fields") or []:
        if not isinstance(field, dict):
            continue
        label = field.get("label") or field.get("title") or field.get("id")
        value = clean(field.get("value"))
        if label and value:
            result[norm(label)] = value
    return result


def extract_docx_facts(path: Path) -> dict[str, str]:
    facts: dict[str, str] = {}
    doc = Document(str(path))
    for paragraph in doc.paragraphs:
        add_line_facts(facts, paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [clean(cell.text) for cell in row.cells]
            non_empty = [cell for cell in cells if cell]
            if len(non_empty) >= 2:
                key = non_empty[0]
                value = next((cell for cell in non_empty[1:] if cell and not PLACEHOLDER_RE.search(cell)), "")
                if key and value:
                    facts.setdefault(norm(key), value)
            for cell in cells:
                add_line_facts(facts, cell)
    return facts


def extract_xlsx_facts(path: Path) -> dict[str, str]:
    facts: dict[str, str] = {}
    wb = load_workbook(path, read_only=True, data_only=True)
    for ws in wb.worksheets[:8]:
        for row in ws.iter_rows(max_row=500, values_only=True):
            cells = [clean(value) for value in row]
            non_empty = [cell for cell in cells if cell]
            if len(non_empty) >= 2:
                facts.setdefault(norm(non_empty[0]), non_empty[1])
            for cell in non_empty:
                add_line_facts(facts, cell)
    wb.close()
    return facts


def add_line_facts(facts: dict[str, str], line: Any) -> None:
    text = clean(line)
    if not text or len(text) > 180:
        return
    for sep in ("：", ":", "=", "——", "-"):
        if sep in text:
            left, right = text.split(sep, 1)
            left = clean(left)
            right = clean(right)
            if 1 <= len(left) <= 40 and right:
                facts.setdefault(norm(left), right)
            return


def source_fact_map(manifest: dict[str, Any]) -> tuple[dict[str, str], list[dict[str, Any]]]:
    facts: dict[str, str] = {}
    evidence: list[dict[str, Any]] = []
    for material in manifest.get("sourceMaterials") or []:
        if not isinstance(material, dict):
            continue
        path = Path(str(material.get("path") or ""))
        if not path.exists():
            continue
        try:
            if path.suffix.lower() == ".docx":
                extracted = extract_docx_facts(path)
            elif path.suffix.lower() in {".xlsx", ".xlsm", ".xls"}:
                extracted = extract_xlsx_facts(path)
            else:
                extracted = {}
        except Exception:
            extracted = {}
        for key, value in extracted.items():
            facts.setdefault(key, value)
        if extracted:
            evidence.append({
                "materialId": material.get("materialId") or material.get("id") or "",
                "materialName": material.get("materialName") or material.get("name") or path.name,
                "factCount": len(extracted),
            })
    return facts, evidence


LABEL_DECOR_SUFFIXES = {"名称", "全称", "盖章", "公章", "签字", "签章", "印章", "签名", "盖公章", "盖单位章"}

BRACKET_FIELD_RE = re.compile(r"[（(]\s*([一-龥A-Za-z][一-龥A-Za-z0-9/、]{1,24}?)\s*([:：])?\s*[)）]")
UNDERLINE_FIELD_RE = re.compile(r"([一-龥A-Za-z/]{2,20})\s*[:：]?\s*[_＿]{3,}")
TRAILING_LABEL_RE = re.compile(r"([一-龥A-Za-z/（()）]{2,30})[:：]\s*$")
DATE_SKELETON_RE = re.compile(r"^(日期)[:：][\s年月日]*$")
INLINE_LABEL_BLANK_RE = re.compile(r"([一-龥A-Za-z/]{2,20})[:：]([ 　]{4,})(?=[一-龥（(A-Za-z]|$)")


def lookup(label: str, facts: dict[str, str]) -> str:
    key = norm(label)
    if not key:
        return ""
    if key in facts:
        return facts[key]
    for existing, value in facts.items():
        if key in existing:
            return value
    # 事实键比标签更泛时（如标签"投标人地址"含事实键"投标人"），残余部分必须是
    # 装饰性后缀才允许回填，避免地址/电话等更具体的字段被泛值污染
    for existing, value in facts.items():
        if not existing or existing not in key:
            continue
        remainder = key.replace(existing, "", 1)
        if remainder in LABEL_DECOR_SUFFIXES:
            return value
    return ""


def compact_fact_lines(facts: dict[str, str], *, limit: int = 8) -> list[str]:
    lines: list[str] = []
    for key, value in facts.items():
        text = clean(value)
        if not key or not text:
            continue
        lines.append(text if len(text) > 80 else f"{key}：{text}")
        if len(lines) >= limit:
            break
    return lines


def cell_is_blank_or_placeholder(text: str) -> bool:
    value = clean(text)
    return not value or bool(PLACEHOLDER_RE.search(value)) or value in {"待填写", "无", "-"}


def set_cell_text(cell: Any, value: str) -> None:
    cell.text = value


def highlight_cell(cell: Any) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def infer_label_from_row(cells: list[str], value_index: int) -> str:
    candidates = cells[:value_index]
    for candidate in reversed(candidates):
        text = clean(candidate)
        if text and not cell_is_blank_or_placeholder(text):
            return text
    return ""


def write_fill_report_docx(output_path: Path, manifest: dict[str, Any], facts: dict[str, str]) -> dict[str, Any]:
    doc = Document()
    title = doc.add_heading(clean(manifest.get("task", {}).get("title")) or "填表内容", level=1)
    for run in title.runs:
      run.font.size = Pt(16)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "字段"
    table.rows[0].cells[1].text = "内容"
    filled = 0
    labels = suggested_labels(manifest, facts)
    for label in labels:
        value = lookup(label, facts)
        if not value:
            continue
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value
        filled += 1
    if filled == 0:
        row = table.add_row().cells
        row[0].text = "填表依据"
        row[1].text = "\n".join(compact_fact_lines(facts, limit=12))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return {"filled": filled, "unfilledFields": [] if filled else [{"label": "未匹配到明确字段"}]}


def suggested_labels(manifest: dict[str, Any], facts: dict[str, str]) -> list[str]:
    labels: list[str] = []
    for field in (manifest.get("projectFactTable") or {}).get("fields") or []:
        if isinstance(field, dict):
            label = clean(field.get("label") or field.get("title"))
            if label:
                labels.append(label)
    task = manifest.get("task") if isinstance(manifest.get("task"), dict) else {}
    for token in re.split(r"[，,、；;\s]+", clean(task.get("title") or task.get("requirement"))):
        if 2 <= len(token) <= 24:
            labels.append(token)
    labels.extend(facts.keys())
    result: list[str] = []
    seen: set[str] = set()
    for label in labels:
        text = clean(label)
        key = norm(text)
        if not text or not key or key in seen:
            continue
        seen.add(key)
        result.append(text)
    return result[:24]


def fill_text_placeholders(text: str, facts: dict[str, str]) -> tuple[str, int, list[str]]:
    filled = 0
    unmatched: list[str] = []

    date_match = DATE_SKELETON_RE.match(text.strip())
    if date_match:
        value = lookup(date_match.group(1), facts)
        if value:
            return f"{date_match.group(1)}：{value}", 1, []
        return text, 0, [date_match.group(1)]

    def bracket_sub(match: re.Match[str]) -> str:
        nonlocal filled
        label, colon = match.group(1), match.group(2)
        value = lookup(label, facts)
        if not value:
            if colon:
                unmatched.append(label)
            return match.group(0)
        filled += 1
        if colon:
            return f"（{label}：{value}）"
        return value

    updated = BRACKET_FIELD_RE.sub(bracket_sub, text)

    def underline_sub(match: re.Match[str]) -> str:
        nonlocal filled
        label = match.group(1)
        value = lookup(label, facts)
        if not value:
            unmatched.append(label)
            return match.group(0)
        filled += 1
        return f"{label}：{value}"

    updated = UNDERLINE_FIELD_RE.sub(underline_sub, updated)

    def inline_blank_sub(match: re.Match[str]) -> str:
        nonlocal filled
        label = match.group(1)
        value = lookup(label, facts)
        if not value:
            unmatched.append(label)
            return match.group(0)
        filled += 1
        return f"{label}：{value}  "

    updated = INLINE_LABEL_BLANK_RE.sub(inline_blank_sub, updated)

    trailing_match = TRAILING_LABEL_RE.search(updated.rstrip())
    if trailing_match:
        label = trailing_match.group(1)
        value = lookup(label, facts)
        if value:
            prefix = updated.rstrip()[: trailing_match.start()]
            return f"{prefix}{label}：{value}", filled + 1, unmatched
        unmatched.append(label)
    return updated, filled, unmatched


def fill_docx_paragraph_placeholders(doc: Any, facts: dict[str, str]) -> tuple[int, list[dict[str, Any]]]:
    filled = 0
    unfilled: list[dict[str, Any]] = []

    def handle(paragraph: Any, location: str) -> None:
        nonlocal filled
        original = paragraph.text
        if not original or not original.strip() or not paragraph.runs:
            return
        updated, count, unmatched = fill_text_placeholders(original, facts)
        if count and updated != original:
            paragraph.runs[0].text = updated
            for run in paragraph.runs[1:]:
                run.text = ""
            filled += count
        if unmatched:
            for run in paragraph.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            for label in unmatched:
                unfilled.append({"label": label, "location": location, "text": original[:60]})

    for index, paragraph in enumerate(doc.paragraphs, start=1):
        handle(paragraph, f"段落{index}")
    for table_index, table in enumerate(doc.tables, start=1):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    handle(paragraph, f"表{table_index}单元格")
    return filled, unfilled


def is_template_body_missing(doc: Any) -> bool:
    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    return len(paragraphs) <= 2 and not doc.tables


def append_manual_notice(output_path: Path, message: str) -> None:
    doc = Document(str(output_path))
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(message)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    doc.save(str(output_path))


def fill_docx(target_path: Path, output_path: Path, facts: dict[str, str]) -> dict[str, Any]:
    doc = Document(str(target_path))
    filled = 0
    unfilled: list[dict[str, Any]] = []
    for table_index, table in enumerate(doc.tables, start=1):
        headers = [clean(cell.text) for cell in table.rows[0].cells] if table.rows else []
        value_indexes = [
            idx for idx, header in enumerate(headers)
            if any(marker in header for marker in FILL_HEADERS)
        ]
        if not value_indexes and table.rows:
            value_indexes = [len(table.rows[0].cells) - 1]
        for row_index, row in enumerate(table.rows[1:] if headers else table.rows, start=2 if headers else 1):
            cells = [clean(cell.text) for cell in row.cells]
            for value_index in value_indexes:
                if value_index < 0 or value_index >= len(row.cells):
                    continue
                target_cell = row.cells[value_index]
                if not cell_is_blank_or_placeholder(target_cell.text):
                    continue
                label = infer_label_from_row(cells, value_index)
                value = lookup(label, facts)
                if value:
                    set_cell_text(target_cell, value)
                    filled += 1
                elif label:
                    unfilled.append({"label": label, "tableIndex": table_index, "rowIndex": row_index})
                    highlight_cell(target_cell)
    paragraph_filled, paragraph_unfilled = fill_docx_paragraph_placeholders(doc, facts)
    filled += paragraph_filled
    unfilled.extend(paragraph_unfilled)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return {"filled": filled, "unfilledFields": unfilled}


def fill_xlsx(target_path: Path, output_path: Path, facts: dict[str, str]) -> dict[str, Any]:
    wb = load_workbook(target_path)
    filled = 0
    unfilled: list[dict[str, Any]] = []
    for ws in wb.worksheets[:8]:
        for row in ws.iter_rows(max_row=500):
            values = [clean(cell.value) for cell in row]
            for index, cell in enumerate(row):
                if not cell_is_blank_or_placeholder(cell.value):
                    continue
                label = infer_label_from_row(values, index)
                value = lookup(label, facts)
                if value:
                    cell.value = value
                    filled += 1
                elif label:
                    unfilled.append({"label": label, "sheet": ws.title, "rowIndex": cell.row, "columnIndex": cell.column})
    output_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(output_path)
    wb.close()
    return {"filled": filled, "unfilledFields": unfilled}


def resolve_target_path(target: dict[str, Any]) -> Path:
    for key in ("filePath", "path", "docxPath", "workspacePath"):
        value = str(target.get(key) or "")
        if value and Path(value).exists():
            return Path(value)
    raise ValueError("target 缺少可读取的 Word 路径")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--manifest", required=True)
    parser.add_argument("--response", default="summary")
    args = parser.parse_args()

    manifest_path = Path(args.manifest)
    manifest = read_json(manifest_path)
    output_path = Path(str(manifest.get("outputFile") or manifest_path.with_name("business-table-filled.docx")))
    target_path = resolve_target_path(manifest.get("target") if isinstance(manifest.get("target"), dict) else {})

    facts = fact_map(manifest)
    material_facts, evidence_refs = source_fact_map(manifest)
    merged_facts = {**material_facts, **facts}

    suffix = target_path.suffix.lower()
    if suffix == ".docx":
        fill_result = fill_docx(target_path, output_path, merged_facts)
        if is_template_body_missing(Document(str(target_path))):
            # 模板正文缺失时绝不注入通用事实表，保留原样并显式告警
            append_manual_notice(
                output_path,
                "【AI填写】模板正文缺失（解析切片仅含标题），未注入任何内容；请人工处理或重跑 S1 模板提取。",
            )
            fill_result["unfilledFields"] = [
                *fill_result["unfilledFields"],
                {"label": "模板正文", "location": "全文", "text": "模板正文缺失"},
            ]
            fill_result["templateBodyMissing"] = True
        elif fill_result["filled"] == 0:
            append_manual_notice(output_path, "【AI填写】未找到可自动填写的字段，已保留模板原样，请人工填写。")
    elif suffix in {".xlsx", ".xlsm"}:
        xlsx_output = output_path.with_suffix(target_path.suffix)
        fill_result = fill_xlsx(target_path, xlsx_output, merged_facts)
        output_path = xlsx_output
    else:
        fill_result = write_fill_report_docx(output_path, manifest, merged_facts)

    payload = {
        "schemaVersion": SCHEMA_VERSION,
        "outputFile": str(output_path),
        "filledAt": now_iso(),
        "fillReport": {
            "target": manifest.get("target", {}).get("fileName") or target_path.name,
            "sourceMaterialCount": len(manifest.get("sourceMaterials") or []),
            "factCount": len(merged_facts),
            "filledFieldCount": fill_result["filled"],
            "unfilledFieldCount": len(fill_result["unfilledFields"]),
            "templateBodyMissing": bool(fill_result.get("templateBodyMissing")),
        },
        "unfilledFields": fill_result["unfilledFields"],
        "evidenceRefs": evidence_refs,
    }
    print(json.dumps(payload, ensure_ascii=False))


if __name__ == "__main__":
    main()
