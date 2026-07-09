from __future__ import annotations

import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

try:
    import fitz  # type: ignore
except Exception:  # pragma: no cover
    fitz = None

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from . import nav_store


HEADING_NUMBER_RE = re.compile(
    r"^\s*(?:第[一二三四五六七八九十百]+[章节]|[一二三四五六七八九十]+[、.．]|"
    r"\d+(?:\.\d+){0,4}[、.．]?|[（(]\d+[）)])"
)


def clean(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "").replace("\u3000", " ")).strip()


def snippet(value: str, limit: int = 180) -> str:
    text = clean(value)
    return text if len(text) <= limit else text[: limit - 1] + "…"


@dataclass
class IndexedDocument:
    document: dict[str, Any]
    blocks: list[dict[str, Any]]
    tables: list[dict[str, Any]]
    headings: list[dict[str, Any]]


def _iter_docx_body(document: DocxDocument):
    body = document.element.body
    table_iter = iter(document.tables)
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            table = next(table_iter, None)
            if table is not None:
                yield table


def _paragraph_text(paragraph: Paragraph) -> str:
    return clean(paragraph.text)


def _paragraph_is_bold(paragraph: Paragraph) -> bool:
    runs = [run for run in paragraph.runs if clean(run.text)]
    if not runs:
        return False
    bold_count = sum(1 for run in runs if bool(run.bold))
    return bold_count >= max(1, len(runs) // 2)


def _heading_level(paragraph: Paragraph, text: str) -> int:
    if not text or len(text) > 120:
        return 0
    style_name = str(getattr(paragraph.style, "name", "") or "")
    match = re.search(r"Heading\s*(\d+)|标题\s*(\d+)", style_name, re.I)
    if match:
        return max(1, min(6, int(match.group(1) or match.group(2) or 1)))
    centered = paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    numbered = bool(HEADING_NUMBER_RE.search(text))
    bold = _paragraph_is_bold(paragraph)
    if text.startswith("第") and any(token in text for token in ("章", "节")):
        return 1
    if centered and len(text) <= 60 and (bold or numbered):
        return 1 if text.startswith("第") else 2
    if numbered and len(text) <= 80:
        dots = text.split(maxsplit=1)[0].count(".")
        return max(2, min(6, dots + 2))
    if bold and len(text) <= 50 and not text.endswith(("。", "；", ";")):
        return 3
    return 0


def _update_heading_stack(stack: list[tuple[int, str]], level: int, text: str) -> list[tuple[int, str]]:
    if level <= 0:
        return stack
    stack = [item for item in stack if item[0] < level]
    stack.append((level, text))
    return stack


def _heading_path(stack: list[tuple[int, str]]) -> str:
    return " > ".join(text for _, text in stack)


def _table_rows(table: Table) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([clean(cell.text) for cell in row.cells])
    return rows


def _table_preview(rows: list[list[str]], max_rows: int = 4) -> str:
    return " || ".join(" | ".join(cell for cell in row if cell) for row in rows[:max_rows])


def _index_docx(document_meta: dict[str, Any], source_path: Path) -> IndexedDocument:
    doc = Document(str(source_path))
    document_id = str(document_meta.get("id") or "DOC-1")
    blocks: list[dict[str, Any]] = []
    tables: list[dict[str, Any]] = []
    headings: list[dict[str, Any]] = []
    heading_stack: list[tuple[int, str]] = []
    table_no = 0
    for body_index, item in enumerate(_iter_docx_body(doc), start=1):
        if isinstance(item, Paragraph):
            text = _paragraph_text(item)
            level = _heading_level(item, text)
            if level:
                heading_stack = _update_heading_stack(heading_stack, level, text)
                headings.append(
                    {
                        "id": f"{document_id}:B{body_index:06d}",
                        "level": level,
                        "title": text,
                        "bodyIndex": body_index,
                    }
                )
            blocks.append(
                {
                    "id": f"{document_id}:B{body_index:06d}",
                    "documentId": document_id,
                    "bodyIndex": body_index,
                    "type": "paragraph",
                    "text": text,
                    "headingLevel": level,
                    "headingPath": _heading_path(heading_stack),
                }
            )
            continue
        table_no += 1
        table_id = f"{document_id}:T{table_no:04d}"
        rows = _table_rows(item)
        heading_path = _heading_path(heading_stack)
        title = ""
        for previous in reversed(blocks[-4:]):
            candidate = clean(previous.get("text"))
            if candidate:
                title = candidate
                break
        header_text = " | ".join(rows[0]) if rows else ""
        preview_text = _table_preview(rows)
        tables.append(
            {
                "id": table_id,
                "documentId": document_id,
                "bodyIndex": body_index,
                "title": title,
                "headingPath": heading_path,
                "rowCount": len(rows),
                "colCount": max((len(row) for row in rows), default=0),
                "headerText": header_text,
                "previewText": preview_text,
                "rows": rows,
            }
        )
        blocks.append(
            {
                "id": f"{document_id}:B{body_index:06d}",
                "documentId": document_id,
                "bodyIndex": body_index,
                "type": "table",
                "text": preview_text,
                "headingLevel": 0,
                "headingPath": heading_path,
                "tableId": table_id,
            }
        )
    _fill_neighbors(blocks)
    document = {
        **document_meta,
        "id": document_id,
        "sourcePath": str(source_path),
        "blockCount": len(blocks),
        "tableCount": len(tables),
    }
    return IndexedDocument(document=document, blocks=blocks, tables=tables, headings=headings)


def _plain_text_blocks(document_meta: dict[str, Any], source_path: Path, text: str) -> IndexedDocument:
    document_id = str(document_meta.get("id") or "DOC-1")
    blocks: list[dict[str, Any]] = []
    tables: list[dict[str, Any]] = []
    headings: list[dict[str, Any]] = []
    heading_stack: list[tuple[int, str]] = []
    table_no = 0
    body_index = 0
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i].strip()
        if raw.startswith("|") and raw.endswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [
                [clean(cell) for cell in line.strip("|").split("|")]
                for line in table_lines
                if not re.fullmatch(r"\s*\|?[-:\s|]+\|?\s*", line)
            ]
            body_index += 1
            table_no += 1
            table_id = f"{document_id}:T{table_no:04d}"
            preview_text = _table_preview(rows)
            title = clean(blocks[-1].get("text")) if blocks else ""
            tables.append(
                {
                    "id": table_id,
                    "documentId": document_id,
                    "bodyIndex": body_index,
                    "title": title,
                    "headingPath": _heading_path(heading_stack),
                    "rowCount": len(rows),
                    "colCount": max((len(row) for row in rows), default=0),
                    "headerText": " | ".join(rows[0]) if rows else "",
                    "previewText": preview_text,
                    "rows": rows,
                }
            )
            blocks.append(
                {
                    "id": f"{document_id}:B{body_index:06d}",
                    "documentId": document_id,
                    "bodyIndex": body_index,
                    "type": "table",
                    "text": preview_text,
                    "headingLevel": 0,
                    "headingPath": _heading_path(heading_stack),
                    "tableId": table_id,
                }
            )
            continue
        i += 1
        if not raw:
            continue
        body_index += 1
        text_line = clean(raw.lstrip("#").strip())
        level = 1 if raw.startswith("# ") else 2 if raw.startswith("## ") else (2 if HEADING_NUMBER_RE.search(text_line) and len(text_line) < 80 else 0)
        if level:
            heading_stack = _update_heading_stack(heading_stack, level, text_line)
            headings.append(
                {
                    "id": f"{document_id}:B{body_index:06d}",
                    "level": level,
                    "title": text_line,
                    "bodyIndex": body_index,
                }
            )
        blocks.append(
            {
                "id": f"{document_id}:B{body_index:06d}",
                "documentId": document_id,
                "bodyIndex": body_index,
                "type": "paragraph",
                "text": text_line,
                "headingLevel": level,
                "headingPath": _heading_path(heading_stack),
            }
        )
    _fill_neighbors(blocks)
    return IndexedDocument(
        document={**document_meta, "id": document_id, "sourcePath": str(source_path), "blockCount": len(blocks), "tableCount": len(tables)},
        blocks=blocks,
        tables=tables,
        headings=headings,
    )


def _fill_neighbors(blocks: list[dict[str, Any]]) -> None:
    for index, block in enumerate(blocks):
        block["prevText"] = snippet(blocks[index - 1].get("text") or "", 160) if index > 0 else ""
        block["nextText"] = snippet(blocks[index + 1].get("text") or "", 160) if index + 1 < len(blocks) else ""


def _extract_pdf_text(path: Path) -> str:
    if fitz is None:
        return ""
    parts: list[str] = []
    with fitz.open(str(path)) as pdf:
        for page in pdf:
            parts.append(page.get_text("text") or "")
    return "\n".join(parts)


def _read_document(document_meta: dict[str, Any]) -> IndexedDocument:
    source_path = Path(str(document_meta.get("sourcePath") or document_meta.get("textPath") or ""))
    if source_path.suffix.lower() == ".docx" and source_path.is_file():
        return _index_docx(document_meta, source_path)
    text_path = Path(str(document_meta.get("textPath") or ""))
    if text_path.is_file():
        text = text_path.read_text(encoding="utf-8", errors="replace")
        return _plain_text_blocks(document_meta, text_path, text)
    if source_path.suffix.lower() in {".md", ".txt"} and source_path.is_file():
        return _plain_text_blocks(document_meta, source_path, source_path.read_text(encoding="utf-8", errors="replace"))
    if source_path.suffix.lower() == ".pdf" and source_path.is_file():
        if str(document_meta.get("documentParseEngine") or "").strip().lower() == "docling":
            return _plain_text_blocks(document_meta, source_path, "")
        return _plain_text_blocks(document_meta, source_path, _extract_pdf_text(source_path))
    return _plain_text_blocks(document_meta, source_path, "")


def _output_dir(manifest_path: Path, manifest: dict[str, Any]) -> Path:
    output = Path(str(manifest.get("structuredResultPath") or manifest_path.with_name("s1_structured_result.json")))
    return output.parent


def nav_store_path(manifest_path: Path, manifest: dict[str, Any]) -> Path:
    value = str(manifest.get("navStorePath") or "").strip()
    return Path(value) if value else _output_dir(manifest_path, manifest) / "s1_nav.sqlite"


def document_map_path(manifest_path: Path, manifest: dict[str, Any]) -> Path:
    value = str(manifest.get("documentMapPath") or "").strip()
    return Path(value) if value else _output_dir(manifest_path, manifest) / "document_map.json"


def build_index(manifest_path: Path, manifest: dict[str, Any]) -> dict[str, Any]:
    store_path = nav_store_path(manifest_path, manifest)
    map_path = document_map_path(manifest_path, manifest)
    conn = nav_store.connect(store_path)
    indexed_docs: list[IndexedDocument] = []
    total_blocks = 0
    total_tables = 0
    try:
        nav_store.reset_db(conn)
        for raw_doc in manifest.get("documents") or []:
            if not isinstance(raw_doc, dict):
                continue
            indexed = _read_document(raw_doc)
            indexed_docs.append(indexed)
            nav_store.insert_document(conn, indexed.document)
            for block in indexed.blocks:
                nav_store.insert_block(conn, block)
            for table in indexed.tables:
                table_for_store = {key: value for key, value in table.items() if key != "rows"}
                nav_store.insert_table(conn, table_for_store)
                for row_index, row in enumerate(table.get("rows") or [], start=1):
                    row_id = f"{table['id']}:R{row_index:04d}"
                    nav_store.insert_table_row(
                        conn,
                        {
                            "id": row_id,
                            "tableId": table["id"],
                            "documentId": table["documentId"],
                            "bodyIndex": table["bodyIndex"],
                            "rowIndex": row_index,
                            "text": " | ".join(row),
                        },
                    )
                    for col_index, cell in enumerate(row, start=1):
                        nav_store.insert_table_cell(
                            conn,
                            {
                                "id": f"{row_id}:C{col_index:04d}",
                                "tableId": table["id"],
                                "documentId": table["documentId"],
                                "bodyIndex": table["bodyIndex"],
                                "rowIndex": row_index,
                                "colIndex": col_index,
                                "text": cell,
                            },
                        )
            total_blocks += len(indexed.blocks)
            total_tables += len(indexed.tables)
        nav_store.set_meta(conn, "manifestPath", str(manifest_path))
        nav_store.set_meta(conn, "documentMapPath", str(map_path))
        nav_store.set_meta(conn, "navStorePath", str(store_path))
        nav_store.set_meta(conn, "blockCount", total_blocks)
        nav_store.set_meta(conn, "tableCount", total_tables)
        conn.commit()
    finally:
        conn.close()
    document_map = {
        "schemaVersion": nav_store.SCHEMA_VERSION,
        "navStorePath": str(store_path),
        "documentCount": len(indexed_docs),
        "blockCount": total_blocks,
        "tableCount": total_tables,
        "documents": [
            {
                "id": item.document["id"],
                "name": item.document.get("name") or item.document["id"],
                "sourcePath": item.document.get("sourcePath") or "",
                "blockCount": item.document.get("blockCount") or 0,
                "tableCount": item.document.get("tableCount") or 0,
                "headings": item.headings[:80],
                "tables": [
                    {
                        "id": table["id"],
                        "bodyIndex": table["bodyIndex"],
                        "title": table.get("title") or "",
                        "headingPath": table.get("headingPath") or "",
                        "rowCount": table.get("rowCount") or 0,
                        "colCount": table.get("colCount") or 0,
                        "headerText": table.get("headerText") or "",
                        "previewText": snippet(table.get("previewText") or "", 240),
                    }
                    for table in item.tables[:120]
                ],
            }
            for item in indexed_docs
        ],
    }
    map_path.parent.mkdir(parents=True, exist_ok=True)
    map_path.write_text(json.dumps(document_map, ensure_ascii=False, indent=2), encoding="utf-8")
    return {
        "schemaVersion": nav_store.SCHEMA_VERSION,
        "stage": "prepared",
        "navStorePath": str(store_path),
        "documentMapPath": str(map_path),
        "documentCount": len(indexed_docs),
        "blockCount": total_blocks,
        "tableCount": total_tables,
    }

