#!/usr/bin/env python3
"""Run conservative technical-bid Word format cleaning from a manifest."""

from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
ASSEMBLER_SCRIPTS_DIR = SKILL_DIR.parent / "bid-tech-assembler" / "scripts"
ASSEMBLER_REFERENCES_DIR = SKILL_DIR.parent / "bid-tech-assembler" / "references"
sys.path.insert(0, str(ASSEMBLER_SCRIPTS_DIR))

from finalize import force_update_fields, insert_toc_field, reapply_heading_fonts, replace_header_text  # noqa: E402
from numbering_fixer import strip_numPr_from_body, strip_numPr_from_heading_styles  # noqa: E402
from verify import scan_docx  # noqa: E402


SCHEMA_VERSION = "bid-tech-format-clean-v1"
REQUIRED_FIELDS = ("inputFile", "outlineFile", "outputFile", "projectName")


def run_manifest(manifest_path: str | Path, response: str = "summary") -> dict[str, Any]:
    path = Path(manifest_path)
    manifest = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(manifest, dict):
        raise ValueError("manifest must be a JSON object")
    _validate_manifest(manifest)

    input_path = _resolve_manifest_path(manifest["inputFile"], path)
    outline_path = _resolve_manifest_path(manifest["outlineFile"], path)
    output_path = _resolve_manifest_path(manifest["outputFile"], path)
    style_path = _resolve_style_path(manifest.get("styleSpecPath"), path)
    report_path = output_path.with_name("tech_format_clean_report.md")

    clean_result = clean_docx(
        input_file=input_path,
        outline_file=outline_path,
        output_file=output_path,
        project_name=str(manifest["projectName"]),
        style_spec_path=style_path,
    )
    report = verify_cleaned_docx(
        output_file=output_path,
        outline_file=outline_path,
        report_file=report_path,
        clean_result=clean_result,
    )

    summary = {
        "outlineCount": int(report["outlineCount"]),
        "matchedHeadingCount": int(report["matchedHeadingCount"]),
        "unmatchedHeadingCount": len(report["unmatchedHeadings"]),
        "tocInserted": bool(clean_result["tocInserted"]),
        "tocPresent": bool(report["tocPresent"]),
        "headerCleaned": bool(report["headerCleaned"]),
        "placeholderCount": int(report["placeholderCount"]),
        "riskCount": len(report["formatRisks"]),
    }
    result = {
        "schema_version": SCHEMA_VERSION,
        "inputFile": str(input_path),
        "outlineFile": str(outline_path),
        "outputFile": str(output_path),
        "reportFile": str(report_path),
        "summary": summary,
    }
    if response != "summary":
        result["details"] = {
            "cleanResult": clean_result,
            "report": report,
            "styleSpecPath": str(style_path),
        }
    return result


def clean_docx(
    *,
    input_file: str | Path,
    outline_file: str | Path,
    output_file: str | Path,
    project_name: str,
    style_spec_path: str | Path,
) -> dict[str, Any]:
    input_path = Path(input_file)
    outline_path = Path(outline_file)
    output_path = Path(output_file)
    style_path = Path(style_spec_path)

    if not input_path.exists():
        raise FileNotFoundError(f"inputFile does not exist: {input_path}")
    if not outline_path.exists():
        raise FileNotFoundError(f"outlineFile does not exist: {outline_path}")
    if not style_path.exists():
        raise FileNotFoundError(f"styleSpecPath does not exist: {style_path}")
    if input_path.resolve() == output_path.resolve():
        raise ValueError("inputFile must not be overwritten by outputFile")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(input_path, output_path)

    style_spec = _load_style_spec(style_path)
    outline_items = flatten_outline(load_outline(outline_path))

    doc = Document(str(output_path))
    toc_present_before = document_has_toc(doc)
    toc_inserted = False
    if not toc_present_before:
        insert_toc_field(doc)
        toc_inserted = True

    heading_result = _promote_existing_headings(doc, outline_items, style_spec)
    reapply_heading_fonts(doc, style_spec)
    strip_numPr_from_heading_styles(doc)
    strip_numPr_from_body(doc)
    doc.save(str(output_path))

    if project_name:
        replace_header_text(output_path, project_name)
    force_update_fields(output_path)

    return {
        "inputFile": str(input_path),
        "outlineFile": str(outline_path),
        "outputFile": str(output_path),
        "styleSpecPath": str(style_path),
        "projectName": project_name,
        "outlineCount": len(outline_items),
        "matchedHeadingCount": len(heading_result["matched_headings"]),
        "unmatchedHeadingCount": len(heading_result["unmatched_headings"]),
        "matchedHeadings": heading_result["matched_headings"],
        "unmatchedHeadings": heading_result["unmatched_headings"],
        "tocInserted": toc_inserted,
        "tocPresent": toc_present_before or toc_inserted,
    }


def verify_cleaned_docx(
    *,
    output_file: str | Path,
    outline_file: str | Path,
    report_file: str | Path,
    clean_result: dict[str, Any] | None = None,
) -> dict[str, Any]:
    output_path = Path(output_file)
    outline_path = Path(outline_file)
    report_path = Path(report_file)
    scan = scan_docx(output_path)
    outline_items = flatten_outline(load_outline(outline_path))
    doc = Document(str(output_path))
    toc_present = document_has_toc(doc)
    header_cleaned = _header_cleaned(output_path)
    matched = list((clean_result or {}).get("matchedHeadings") or [])
    unmatched = list((clean_result or {}).get("unmatchedHeadings") or [])
    if not matched and not unmatched:
        heading_texts = {text.strip() for _, text in scan.get("heading_list", []) if text.strip()}
        matched, unmatched = _infer_matches_from_doc(outline_items, heading_texts)

    risks = _structural_risks(
        scan=scan,
        toc_present=toc_present,
        header_cleaned=header_cleaned,
        outline_count=len(outline_items),
        unmatched_count=len(unmatched),
    )
    report = {
        "outlineCount": len(outline_items),
        "matchedHeadingCount": len(matched),
        "unmatchedHeadings": unmatched,
        "tocPresent": toc_present,
        "headerCleaned": header_cleaned,
        "placeholderCount": len(scan.get("placeholders") or []),
        "formatRisks": risks,
    }
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text(_render_report(report, output_path, outline_path), encoding="utf-8")
    return report


def _validate_manifest(manifest: dict[str, Any]) -> None:
    missing = [field for field in REQUIRED_FIELDS if not manifest.get(field)]
    if missing:
        raise ValueError(f"manifest missing fields: {', '.join(missing)}")
    input_path = Path(manifest["inputFile"])
    output_path = Path(manifest["outputFile"])
    if input_path.resolve() == output_path.resolve():
        raise ValueError("inputFile must not equal outputFile")


def _resolve_style_path(style_path_value: Any, manifest_path: Path) -> Path:
    if style_path_value:
        return _resolve_manifest_path(style_path_value, manifest_path)
    return (ASSEMBLER_REFERENCES_DIR / "heading_style.json").resolve()


def _resolve_manifest_path(value: Any, manifest_path: Path) -> Path:
    candidate = Path(str(value))
    if not candidate.is_absolute():
        candidate = (manifest_path.parent / candidate).resolve()
    return candidate


def _load_style_spec(path: Path) -> dict[str, Any]:
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError("styleSpecPath must be a JSON object")
    data.setdefault("heading", {})
    return data


def load_outline(path: str | Path) -> dict[str, Any]:
    data = json.loads(Path(path).read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError("outlineFile must be a JSON object")
    schema = str(data.get("schema_version") or "")
    if schema and schema not in {"tech_bid_outline.v1", "business_bid_outline.v1"}:
        raise ValueError(f"unsupported outline schema_version: {schema}")
    return data


def flatten_outline(outline: dict[str, Any]) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []

    def visit(nodes: Any, inherited_level: int = 1) -> None:
        if not isinstance(nodes, list):
            return
        for index, node in enumerate(nodes, start=1):
            if not isinstance(node, dict):
                continue
            title = str(node.get("title") or "").strip()
            if not title:
                visit(node.get("children"), inherited_level + 1)
                continue
            level = _coerce_level(node.get("level"), inherited_level)
            items.append(
                {
                    "id": str(node.get("id") or f"outline-{len(items) + 1:04d}"),
                    "title": title,
                    "number": str(node.get("number") or "").strip(),
                    "level": level,
                    "order": len(items) + 1,
                    "source_index": index,
                }
            )
            visit(node.get("children"), level + 1)

    visit(outline.get("sections"), 1)
    return items


def _coerce_level(value: Any, default: int) -> int:
    try:
        level = int(value)
    except (TypeError, ValueError):
        level = default
    return max(1, min(level, 9))


def document_has_toc(doc: Document) -> bool:
    return any((node.text or "").upper().find("TOC") >= 0 for node in doc.element.iter(qn("w:instrText")))


def _promote_existing_headings(
    doc: Document,
    outline_items: list[dict[str, Any]],
    style_spec: dict[str, Any],
) -> dict[str, Any]:
    paragraphs = list(doc.paragraphs)
    cursor = 0
    matched: list[dict[str, Any]] = []
    unmatched: list[dict[str, Any]] = []

    for item in outline_items:
        found = _find_outline_item(paragraphs, item, cursor)
        if found is None:
            unmatched.append(_outline_report_item(item))
            continue
        index, mode = found
        paragraph = paragraphs[index]
        level = max(1, min(int(item["level"]), 9))
        _set_heading_style(paragraph, level, doc)
        _strip_paragraph_numpr(paragraph)
        _clear_direct_outline_level(paragraph)
        _apply_heading_run_format(paragraph, _heading_cfg(style_spec, level))
        matched.append(
            {
                **_outline_report_item(item),
                "paragraphIndex": index,
                "matchMode": mode,
                "text": paragraph.text.strip(),
            }
        )
        cursor = index + 1

    return {"matched_headings": matched, "unmatched_headings": unmatched}


def _find_outline_item(paragraphs, item: dict[str, Any], start_index: int) -> tuple[int, str] | None:
    full_key = _normalize_match_text(_format_heading_text(item)) if item.get("number") else ""
    title_key = _normalize_match_text(str(item.get("title") or ""))
    if not title_key:
        return None
    for index in range(start_index, len(paragraphs)):
        text_key = _normalize_match_text(paragraphs[index].text)
        if not text_key:
            continue
        if full_key and text_key == full_key:
            return index, "number+title"
        if text_key == title_key:
            return index, "title"
    return None


def _normalize_match_text(text: str) -> str:
    normalized = str(text or "").strip().replace("\u3000", " ")
    return re.sub(r"\s+", "", normalized)


def _format_heading_text(item: dict[str, Any]) -> str:
    title = str(item.get("title") or "").strip()
    number = str(item.get("number") or "").strip()
    return f"{number} {title}".strip() if number else title


def _outline_report_item(item: dict[str, Any]) -> dict[str, Any]:
    return {
        "id": item.get("id"),
        "number": item.get("number", ""),
        "title": item.get("title", ""),
        "level": item.get("level"),
        "order": item.get("order"),
    }


def _set_heading_style(paragraph, level: int, doc: Document) -> None:
    for style_name in (f"Heading {level}", f"标题 {level}", f"标题{level}"):
        try:
            paragraph.style = doc.styles[style_name]
            return
        except KeyError:
            continue


def _strip_paragraph_numpr(paragraph) -> bool:
    p_pr = paragraph._element.find(qn("w:pPr"))
    if p_pr is None:
        return False
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return False
    p_pr.remove(num_pr)
    return True


def _clear_direct_outline_level(paragraph) -> bool:
    p_pr = paragraph._element.find(qn("w:pPr"))
    if p_pr is None:
        return False
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        return False
    p_pr.remove(outline)
    return True


def _heading_cfg(style_spec: dict[str, Any], level: int) -> dict[str, Any]:
    heading = style_spec.get("heading", {})
    if str(level) in heading:
        return heading[str(level)]
    if heading:
        return heading[sorted(heading.keys(), key=lambda value: int(value))[-1]]
    return {}


def _apply_heading_run_format(paragraph, cfg: dict[str, Any]) -> None:
    if not cfg:
        return
    for run in paragraph.runs:
        if "en_font" in cfg:
            run.font.name = str(cfg["en_font"])
        if "size_pt" in cfg:
            run.font.size = Pt(float(cfg["size_pt"]))
        run.font.bold = bool(cfg.get("bold", False))
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        zh_font = str(cfg.get("zh_font") or cfg.get("en_font") or "")
        en_font = str(cfg.get("en_font") or cfg.get("zh_font") or "")
        if zh_font or en_font:
            r_fonts.set(qn("w:eastAsia"), zh_font)
            r_fonts.set(qn("w:ascii"), en_font)
            r_fonts.set(qn("w:hAnsi"), en_font)
            r_fonts.set(qn("w:cs"), en_font)


def _header_cleaned(docx_path: Path) -> bool:
    residual_terms = ("投标文件-技术卷",)
    text_parts: list[str] = []
    with zipfile.ZipFile(docx_path, "r") as archive:
        for name in archive.namelist():
            if not (name.startswith("word/header") and name.endswith(".xml")):
                continue
            text_parts.append(archive.read(name).decode("utf-8", errors="replace"))
    header_text = "\n".join(text_parts)
    return not any(term in header_text for term in residual_terms)


def _infer_matches_from_doc(outline_items: list[dict[str, Any]], heading_texts: set[str]):
    normalized_headings = {_normalize_match_text(text) for text in heading_texts if text}
    matched: list[dict[str, Any]] = []
    unmatched: list[dict[str, Any]] = []
    for item in outline_items:
        record = _outline_report_item(item)
        keys = {_normalize_match_text(_format_heading_text(item)), _normalize_match_text(str(item.get("title") or ""))}
        if normalized_headings & {key for key in keys if key}:
            matched.append(record)
        else:
            unmatched.append(record)
    return matched, unmatched


def _structural_risks(
    *,
    scan: dict[str, Any],
    toc_present: bool,
    header_cleaned: bool,
    outline_count: int,
    unmatched_count: int,
) -> list[str]:
    risks: list[str] = []
    if outline_count == 0:
        risks.append("outline 为空，未能定位任何技术标标题。")
    if unmatched_count:
        risks.append("存在未匹配标题；技术标 cleaner 未插入或删除正文，请人工核对缺失章节。")
    if not toc_present:
        risks.append("未检测到 TOC 域，目录页需要人工补充。")
    if not header_cleaned:
        risks.append("页眉仍存在技术卷残留，请人工检查页眉。")
    if scan.get("placeholders"):
        risks.append("存在待填写或缺失占位符，请在共创阶段补齐。")
    if scan.get("dup_alerts"):
        risks.append("存在相邻重复标题，请检查素材首标题和目录标题是否重复。")
    if scan.get("invalid_h1"):
        risks.append("存在异常一级标题，请检查章节层级。")
    if scan.get("empty_leaf_headings"):
        risks.append("存在空叶子章节，请检查是否缺正文或素材未拼入。")
    if not risks:
        risks.append("未发现脚本可识别的格式风险；仍建议在 Word/WPS 中打开后刷新目录并人工抽检。")
    return _dedupe(risks)


def _dedupe(items: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for item in items:
        if item in seen:
            continue
        seen.add(item)
        result.append(item)
    return result


def _render_report(report: dict[str, Any], output_path: Path, outline_path: Path) -> str:
    lines = [
        "# 技术标格式清洗报告",
        "",
        f"- 输出文件：`{output_path}`",
        f"- outline 文件：`{outline_path}`",
        f"- outline 总数：{report['outlineCount']}",
        f"- 成功匹配标题数：{report['matchedHeadingCount']}",
        f"- 未匹配标题数：{len(report['unmatchedHeadings'])}",
        f"- TOC 是否存在：{'是' if report['tocPresent'] else '否'}",
        f"- 页眉是否清理：{'是' if report['headerCleaned'] else '否'}",
        f"- 占位符数量：{report['placeholderCount']}",
        "",
        "## 未匹配标题清单",
        "",
    ]
    if report["unmatchedHeadings"]:
        for item in report["unmatchedHeadings"]:
            lines.append(f"- {_format_heading_text(item)}（id={item.get('id')}, level={item.get('level')}）")
    else:
        lines.append("- 无")
    lines.extend(["", "## 格式风险提示", ""])
    for risk in report["formatRisks"]:
        lines.append(f"- {risk}")
    lines.append("")
    return "\n".join(lines)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Run technical bid format cleaner from manifest")
    parser.add_argument("manifest")
    parser.add_argument("--response", choices=("summary", "details"), default="summary")
    args = parser.parse_args(argv)
    result = run_manifest(args.manifest, response=args.response)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
