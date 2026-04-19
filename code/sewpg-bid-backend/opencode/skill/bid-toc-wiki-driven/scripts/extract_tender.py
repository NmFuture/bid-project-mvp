#!/usr/bin/env python3
"""extract_tender.py — 从招标文件 docx 抽项目参数

用 python-docx 读全文文字，关键词扫描得：
- 业主 / 项目 / 编号
- 场址关键词命中表（site_flags）
- 机型约束（model_flags）
- 地块（plot_flags）
- 招标特殊要求（specials，碳纤维叶片/混塔/净空监测/等）

用法：
    python3 extract_tender.py <招标文件.docx>

输出 JSON：
    {"owner": "...", "project": "...", "code": "...",
     "site_flags": {"陆上": true, "沿海": false, ...},
     "model_flags": {"强制含液压": true, ...},
     "plot_flags": {"北区": false, "南区": false},
     "specials": [{"keyword": "碳纤维叶片", "hint_section": "叶片"}, ...]}
"""
import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx 未安装。pip install python-docx", file=sys.stderr)
    sys.exit(1)


# 关键词词典（可扩展）
SITE_KEYWORDS = {
    "陆上": ["陆上", "陆上风电"],
    "海上": ["海上风电", "海上机组"],
    "沿海": ["沿海", "海岸", "海岛"],
    "低温": ["低温", "寒冷", "严寒", "-20℃", "-30℃", "-40℃"],
    "高温": ["高温", "炎热"],
    "覆冰": ["覆冰", "结冰", "凝露"],
    "风沙": ["风沙", "沙尘", "戈壁", "沙漠", "扬尘"],
    "潮湿": ["潮湿", "高湿", "湿热"],
    "盐雾": ["盐雾", "沿海腐蚀"],
    "高海拔": ["高海拔", "高原"],
    "紫外": ["紫外", "UV", "抗老化"],
    "雷暴": ["雷暴", "多雷", "雷击"],
    "混塔": ["混塔", "混合塔", "混凝土-钢塔"],
}

MODEL_KEYWORDS = {
    "强制直驱": [r"必须.*?直驱", r"强制.*?直驱"],
    "强制双馈": [r"必须.*?双馈", r"强制.*?双馈"],
    "强制半直驱": [r"必须.*?半直驱"],
    "强制含液压": ["液压系统", "液压制动"],
    "强制含升降机": ["升降机", "电梯"],
}

PLOT_KEYWORDS = {
    "北区": ["北区"],
    "南区": ["南区"],
}

# 招标明文特殊要求 → 提示挂靠的 wiki 章节关键词
SPECIALS = [
    ("碳纤维叶片", "叶片"),
    ("净空监测", "净空"),
    ("叶尖净空", "净空"),
    ("混塔", "塔筒"),
    ("数字化智慧风场", "数字化"),
    ("智慧风场", "数字化"),
    ("碳排放", "碳排放"),
    ("自主可控", "自主可控"),
    ("状态监测", "状态监测"),
    ("源头管控", "状态监测"),
]


def read_all_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    parts = [p.text for p in doc.paragraphs if p.text]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def read_head_text(docx_path: Path, n_paragraphs: int = 80) -> str:
    """只读前 N 段（避免目录/正文大海里捞针误伤）。"""
    doc = Document(str(docx_path))
    parts = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        parts.append(t)
        if len(parts) >= n_paragraphs:
            break
    return "\n".join(parts)


def extract_owner(head_text: str, full_text: str) -> str:
    """业主：优先头部关键词群组识别，兜底全文关键词。"""
    OWNERS = ["华能", "大唐", "国家电投", "国电投", "三峡", "中广核", "国能", "华电", "国投", "龙源"]
    for k in OWNERS:
        if k in head_text:
            m = re.search(r"(" + k + r"[\u4e00-\u9fff]{1,25}(?:公司|集团|新能源|有限公司))", head_text)
            if m:
                return m.group(1)
            return k
    for k in OWNERS:
        if k in full_text:
            m = re.search(r"(" + k + r"[\u4e00-\u9fff]{1,25}(?:公司|集团|新能源|有限公司))", full_text)
            if m:
                return m.group(1)
            return k
    return ""


def extract_project(head_text: str, full_text: str):
    proj = ""
    code = ""
    # 项目名：头部优先
    for t in [head_text, full_text]:
        m = re.search(r"((?:[^\n（(]{4,60}?)?(?:万千瓦|MW|kW|千瓦)[^\n（(]{0,30}?(?:风电|风电场|项目|工程))", t)
        if m:
            proj = m.group(1).strip()
            break
    if not proj:
        m = re.search(r"([^\n]{4,60}?(?:项目|工程))", head_text)
        if m:
            proj = m.group(1).strip()
    # 编号
    m = re.search(r"(?:招标编号|项目编号)[:：\s]*([A-Za-z0-9\-_.]+)", full_text)
    if m:
        code = m.group(1).strip()
    return proj, code


def scan_keywords(text: str, keyword_map: dict, threshold: int = 2) -> dict:
    """返回 {分类: bool}，次数 >= threshold 才算真命中（降噪目录误触）。"""
    flags = {}
    for cat, kws in keyword_map.items():
        total = 0
        for kw in kws:
            if kw.startswith("^") or "*" in kw or "?" in kw or "+" in kw:
                try:
                    total += len(re.findall(kw, text))
                except re.error:
                    pass
            else:
                total += text.count(kw)
        flags[cat] = total >= threshold
    return flags


def extract_specials(text: str) -> list:
    out = []
    seen = set()
    for kw, hint in SPECIALS:
        if kw in text and kw not in seen:
            out.append({"keyword": kw, "hint_section": hint})
            seen.add(kw)
    return out


def extract(docx_path: Path) -> dict:
    full = read_all_text(docx_path)
    head = read_head_text(docx_path, 80)
    owner = extract_owner(head, full)
    project, code = extract_project(head, full)
    return {
        "owner": owner,
        "project": project,
        "code": code,
        "site_flags": scan_keywords(full, SITE_KEYWORDS, threshold=3),
        "model_flags": scan_keywords(full, MODEL_KEYWORDS, threshold=2),
        "plot_flags": scan_keywords(full, PLOT_KEYWORDS, threshold=3),
        "specials": extract_specials(full),
        "source": str(docx_path),
    }


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx")
    args = ap.parse_args()
    p = Path(args.docx).expanduser().resolve()
    if not p.exists():
        print(f"ERROR: {p} 不存在", file=sys.stderr)
        sys.exit(1)
    result = extract(p)
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
