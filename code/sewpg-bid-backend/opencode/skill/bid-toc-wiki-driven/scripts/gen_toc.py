#!/usr/bin/env python3
"""gen_toc.py — 按"投标文件格式要求.md"生成投标文件总目录 docx

用法：
    python3 gen_toc.py --plan plan.json --out 总目录.docx [--style-spec 投标文件格式要求.md]

plan.json 结构：
    {
      "title": "<完整标题>",
      "items": [
        {"level": 1, "number": "第一章", "text": "标前概述", "tag": "保留"},
        {"level": 2, "number": "1.1", "text": "技术评分标准索引表", "tag": "适配"},
        {"level": 1, "number": "附表", "text": "", "tag": "保留"},
        {"level": 2, "number": "附表 A", "text": "投标机型总体方案", "tag": "保留"},
        {"level": 3, "number": "A.1", "text": "...", "tag": "保留"}
      ]
    }
"""
import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx 未安装。请先 pip install python-docx", file=sys.stderr)
    sys.exit(1)


# 默认样式（兜底，等价于"投标文件格式要求.md"第八节速查表）
DEFAULT_STYLE = {
    "page": {"top_cm": 2.54, "bottom_cm": 2.54, "left_cm": 3.18, "right_cm": 3.18},
    "h1": {"font_cn": "等线 Light", "font_en": "Times New Roman", "size_pt": 15, "bold": True, "align": "center", "line_spacing": 1.75, "first_line_indent_ch": 0},
    "h2": {"font_cn": "等线 Light", "font_en": "Times New Roman", "size_pt": 14, "bold": True, "align": "left",   "line_spacing": 1.75, "first_line_indent_ch": 0},
    "h3": {"font_cn": "等线",      "font_en": "Times New Roman", "size_pt": 12, "bold": True, "align": "left",   "line_spacing": 1.75, "first_line_indent_ch": 0},
    "h4": {"font_cn": "等线 Light", "font_en": "Times New Roman", "size_pt": 12, "bold": True, "align": "justify","line_spacing": 1.75, "first_line_indent_ch": 0},
    "h5": {"font_cn": "等线",      "font_en": "Times New Roman", "size_pt": 12, "bold": True, "align": "left",   "line_spacing": 1.75, "first_line_indent_ch": 0},
    "h6": {"font_cn": "等线 Light", "font_en": "Times New Roman", "size_pt": 12, "bold": True, "align": "justify","line_spacing": 1.33, "first_line_indent_ch": 0, "left_indent_cm": 0.55},
    "normal": {"font_cn": "等线", "font_en": "Times New Roman", "size_pt": 12, "bold": False, "align": "justify", "line_spacing": 1.5, "first_line_indent_ch": 2},
}

ALIGN_MAP = {
    "left":    WD_ALIGN_PARAGRAPH.LEFT,
    "center":  WD_ALIGN_PARAGRAPH.CENTER,
    "right":   WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def parse_style_spec(md_path: Path) -> dict:
    """从"投标文件格式要求.md"解析样式（best-effort，缺字段回退到 DEFAULT_STYLE）。
    实现策略：现阶段不做完整 md 解析，直接返回 DEFAULT_STYLE；用户对样式有疑问可手工调 DEFAULT_STYLE。
    保留这个函数为未来扩展点。"""
    return DEFAULT_STYLE


def set_run_font(run, font_cn: str, font_en: str, size_pt: int, bold: bool):
    run.font.name = font_en
    run.font.size = Pt(size_pt)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_cn)
    rFonts.set(qn("w:ascii"), font_en)
    rFonts.set(qn("w:hAnsi"), font_en)


def apply_paragraph_format(paragraph, spec: dict):
    pf = paragraph.paragraph_format
    pf.line_spacing = spec.get("line_spacing", 1.5)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    if "first_line_indent_ch" in spec and spec["first_line_indent_ch"] > 0:
        pf.first_line_indent = Pt(spec["size_pt"] * spec["first_line_indent_ch"])
    if "left_indent_cm" in spec:
        pf.left_indent = Cm(spec["left_indent_cm"])
    paragraph.alignment = ALIGN_MAP.get(spec.get("align", "left"), WD_ALIGN_PARAGRAPH.LEFT)


def configure_styles(doc, style: dict):
    """覆盖 docx 默认 Heading1-6 + Normal 样式。"""
    sty = doc.styles
    for i in range(1, 7):
        spec = style[f"h{i}"]
        s = sty[f"Heading {i}"]
        s.font.name = spec["font_en"]
        s.font.size = Pt(spec["size_pt"])
        s.font.bold = spec["bold"]
        rPr = s.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), spec["font_cn"])
        rFonts.set(qn("w:ascii"), spec["font_en"])
        rFonts.set(qn("w:hAnsi"), spec["font_en"])

    nm = sty["Normal"]
    spec = style["normal"]
    nm.font.name = spec["font_en"]
    nm.font.size = Pt(spec["size_pt"])
    rPr = nm.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), spec["font_cn"])


def configure_page(doc, page_spec: dict):
    for section in doc.sections:
        section.top_margin = Cm(page_spec["top_cm"])
        section.bottom_margin = Cm(page_spec["bottom_cm"])
        section.left_margin = Cm(page_spec["left_cm"])
        section.right_margin = Cm(page_spec["right_cm"])


def add_heading_paragraph(doc, level: int, text: str, style: dict):
    """添加目录条目段落，应用对应级别样式。level 1-6。"""
    style_name = f"Heading {min(max(level, 1), 6)}"
    p = doc.add_paragraph(style=style_name)
    spec = style[f"h{min(max(level, 1), 6)}"]
    run = p.add_run(text)
    set_run_font(run, spec["font_cn"], spec["font_en"], spec["size_pt"], spec["bold"])
    apply_paragraph_format(p, spec)
    return p


def format_item_text(item: dict) -> str:
    n = item.get("number", "").strip()
    t = item.get("text", "").strip()
    tag = item.get("tag", "").strip()
    parts = []
    if n:
        parts.append(n)
    if t:
        parts.append(t)
    base = "  ".join(parts) if parts else ""
    if tag and tag not in ("保留", ""):
        base = f"{base}（{tag}）"
    return base


def build_doc(plan: dict, style: dict, out: Path):
    doc = Document()
    configure_page(doc, style["page"])
    configure_styles(doc, style)

    title = plan.get("title", "投标文件总目录")
    add_heading_paragraph(doc, 1, title, style)

    items = plan.get("items", [])
    for it in items:
        level = int(it.get("level", 3))
        text = format_item_text(it)
        if not text:
            continue
        add_heading_paragraph(doc, level, text, style)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return len(items) + 1  # +1 for title


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--plan", required=True, help="plan.json 路径")
    ap.add_argument("--out", required=True, help="输出 docx 路径")
    ap.add_argument("--style-spec", help="投标文件格式要求.md 路径（可选）")
    args = ap.parse_args()

    plan_path = Path(args.plan).expanduser().resolve()
    out_path = Path(args.out).expanduser().resolve()

    if not plan_path.exists():
        print(f"ERROR: plan 文件不存在: {plan_path}", file=sys.stderr)
        sys.exit(1)

    with plan_path.open(encoding="utf-8") as f:
        plan = json.load(f)

    if args.style_spec:
        style = parse_style_spec(Path(args.style_spec).expanduser().resolve())
    else:
        style = DEFAULT_STYLE

    n = build_doc(plan, style, out_path)
    print(json.dumps({"output": str(out_path), "paragraphs": n}, ensure_ascii=False))


if __name__ == "__main__":
    main()
