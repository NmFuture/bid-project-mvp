"""
Heading 编号处理器（方案 B）

两套能力：
1. strip_prefix / strip_heading_prefixes_in_doc — 剥离 Heading 文本里的章节号前缀
2. inject_prefix_to_headings — 按父章节号注入编号前缀（方案 B 核心）
3. strip_handwritten_numbering_in_body — 正文手写编号擦除（"7.10 xxx" 之类）

方案 B 下 Heading text 直接带章节号字符串，禁用 Word 多级列表。
"""

from __future__ import annotations

import re
from typing import Optional

# Heading 段落开头的编号前缀正则。按优先级从长到短匹配。
_PREFIX_PATTERNS = [
    # "第X章/节/篇/部分/卷"（中文数字或阿拉伯）
    re.compile(r"^\s*第[一二三四五六七八九十百千万零〇0-9]+[章节篇部分卷]\s*[::]?\s*"),
    # 阿拉伯数字点分前缀 "1.2.3 " 或 "1.2.3. "
    re.compile(r"^\s*\d+(?:\.\d+){0,6}[.．、\s]+"),
    # 括号中文序号 "（一）/（1）/(一)"
    re.compile(r"^\s*[（(][一二三四五六七八九十百千万\d]+[）)]\s*[::]?\s*"),
    # 中文数字序号 "一、/二．/三."
    re.compile(r"^\s*[一二三四五六七八九十百千万零〇]+[、.．:：]\s*"),
    # 罗马数字 + 点或括号
    re.compile(r"^\s*(?:[IVX]+|[ivx]+)[.．、)]\s*"),
    # 英文字母序号 "A./a)/A、"
    re.compile(r"^\s*[A-Za-z][.．、)]\s*"),
    # 前导"附"字（附录/附表类）
    re.compile(r"^\s*附\s*[:：]?\s*"),
]


def strip_prefix(text: str) -> str:
    """从单段 Heading 文本剥除所有可识别的章节号前缀。

    会连续尝试所有 pattern 直到没有匹配为止（处理 "第一章 1.1 xxx" 之类的复合前缀）。
    """
    if not text:
        return text
    prev = None
    current = text
    while prev != current:
        prev = current
        for pat in _PREFIX_PATTERNS:
            current = pat.sub("", current, count=1)
    return current.strip()


def strip_heading_prefixes_in_doc(doc, *, only_heading_styles: bool = True) -> int:
    """对 python-docx Document 中所有 Heading 段落剥除编号前缀。"""
    count = 0
    for para in doc.paragraphs:
        style_name = (para.style.name or "") if para.style else ""
        if only_heading_styles and not _is_heading_style(style_name):
            continue
        old = para.text
        new = strip_prefix(old)
        if new != old and new:
            _replace_paragraph_text_preserve_format(para, new)
            count += 1
    return count


def _is_heading_style(style_name: str) -> bool:
    """兼容英文 Heading X 和中文 标题 X。"""
    if not style_name:
        return False
    s = style_name.strip()
    if s.startswith("Heading ") or s.lower().startswith("heading "):
        return True
    if s.startswith("标题 ") or s == "标题":
        return True
    return False


def _heading_level(style_name: str) -> Optional[int]:
    """从样式名提取 Heading level（1-9）。返回 None 表示不是 Heading。"""
    if not style_name:
        return None
    s = style_name.strip()
    m = re.match(r"^(?:Heading|heading)\s+(\d+)$", s)
    if m:
        return int(m.group(1))
    m = re.match(r"^标题\s*(\d+)$", s)
    if m:
        return int(m.group(1))
    return None


def _replace_paragraph_text_preserve_format(para, new_text: str) -> None:
    """替换段落文本，尽量保留第一个 run 的格式（粗体/字体/字号）。"""
    if not para.runs:
        para.text = new_text
        return
    first_run = para.runs[0]
    first_run.text = new_text
    for extra_run in para.runs[1:]:
        extra_run.text = ""


# ---------- 方案 B 核心：前缀注入 ----------

def inject_prefix_to_headings(
    doc,
    parent_chapter_no: str,
    *,
    toc_title: Optional[str] = None,
    skip_first_if_match: bool = True,
    max_level: int = 6,
    l1_offset: int = 0,
) -> dict:
    """按父章节号给素材内部 Heading 注入编号前缀（方案 B 核心）。

    关键：用"**相对深度**"而非绝对 Heading level。素材起始 Heading 是 H3 还是 H1
    不影响结果——以素材内部**最小 level** 为 rel=1，按相对深度展开到父章节号下。

    Args:
        doc: python-docx Document
        parent_chapter_no: 父章节号（如 "5.8.7"）
        toc_title: toc 条目的纯标题（用于首 Heading 去重）
        skip_first_if_match: 首 Heading 文本匹配 toc_title 时是否剥除
        max_level: 只处理到这一级（绝对 Heading level，过深的忽略）

    Returns:
        {"injected": int, "skipped_first": bool, "removed": int, "min_level": int}
    """
    stats = {"injected": 0, "skipped_first": False, "removed": 0, "min_level": 0}

    # 第一遍：收集所有 Heading 段落（带 level）、识别首个需要 skip 的
    heading_entries: list[tuple] = []  # (para, lvl, pure)
    skip_idx: Optional[int] = None

    for para in doc.paragraphs:
        style_name = (para.style.name or "") if para.style else ""
        lvl = _heading_level(style_name)
        if lvl is None or lvl > max_level:
            continue
        pure = strip_prefix(para.text)
        if not pure:
            continue
        heading_entries.append((para, lvl, pure))

    if not heading_entries:
        return stats

    if skip_first_if_match and toc_title:
        first_para, first_lvl, first_pure = heading_entries[0]
        if _normalize_for_match(first_pure) == _normalize_for_match(toc_title):
            skip_idx = 0
            stats["skipped_first"] = True

    if not any(i != skip_idx for i in range(len(heading_entries))):
        # 整份素材只有一个首 Heading 且被 skip
        if skip_idx is not None:
            p = heading_entries[0][0]
            p._p.getparent().remove(p._p)
            stats["removed"] = 1
        return stats
    stats["min_level"] = min(lvl for i, (_, lvl, _) in enumerate(heading_entries) if i != skip_idx)

    parent_parts = [p for p in (parent_chapter_no or "").split(".") if p]
    parent_depth = len(parent_parts)
    # 目标起始 level：父章节深度 + 1（封顶 6）
    target_first_level = min(parent_depth + 1, 6) if parent_depth > 0 else 1
    stats["target_first_level"] = target_first_level

    counters = [0] * (max_level + 2)  # 1-based
    counters[1] = l1_offset  # 用于避免 appendix 与前序 MATCHED 撞号
    paras_to_remove = []
    level_stack: list[tuple[int, int]] = []

    for i, (para, lvl, pure) in enumerate(heading_entries):
        if i == skip_idx:
            paras_to_remove.append(para)
            continue

        # 用行文顺序维护相对层级，而不是用全局最小 Heading level。
        # 很多素材内部会混用 H1/H2/H5，若直接用绝对 level 差值会生成
        # 1.9.0.0.1 这类空洞编号。这里把当前 heading 挂到最近的上级
        # heading 下面，缺失的中间层级会自动压缩。
        while level_stack and level_stack[-1][0] >= lvl:
            level_stack.pop()
        rel = level_stack[-1][1] + 1 if level_stack else 1
        rel = max(1, min(rel, max_level))
        level_stack.append((lvl, rel))

        counters[rel] += 1
        for d in range(rel + 1, len(counters)):
            counters[d] = 0

        # 更新 level（同 shift 到 target_first_level + rel-1）
        new_level = max(1, min(6, target_first_level + rel - 1))
        _set_heading_style(para, new_level, doc)

        my_parts = parent_parts + [str(counters[k]) for k in range(1, rel + 1)]
        new_text = ".".join(my_parts) + "  " + pure
        _replace_paragraph_text_preserve_format(para, new_text)
        stats["injected"] += 1

    for p in paras_to_remove:
        p._p.getparent().remove(p._p)
        stats["removed"] += 1

    # 返回本次注入最终 L1（相对）计数器值，供 merger 追踪 parent 下已用编号
    stats["top_count"] = counters[1] - l1_offset
    stats["final_l1"] = counters[1]
    return stats


def _set_heading_style(para, level: int, doc) -> None:
    """把段落样式切换到 Heading {level}，兼容中文 "标题 N"。"""
    target = f"Heading {level}"
    try:
        para.style = doc.styles[target]
        return
    except KeyError:
        pass
    cn = f"标题 {level}"
    try:
        para.style = doc.styles[cn]
    except KeyError:
        pass


def _normalize_for_match(s: str) -> str:
    """归一化：去空白、去全半角差异、统一括号，便于比较首 heading 与 toc title。"""
    if not s:
        return ""
    s = strip_prefix(s)
    s = s.strip()
    s = re.sub(r"\s+", "", s)
    # 全角 → 半角常用
    s = s.replace("（", "(").replace("）", ")")
    s = s.replace("：", ":").replace("，", ",")
    return s


# ---------- 正文手写编号擦除 ----------

# 常见单位字（开头一个字就能判定是单位的）；后接其它字也不应被当章节号
_UNIT_CHARS = (
    "秒分时日天年月周"        # 时间
    "米厘毫千公英里"           # 长度
    "倍次个台件份块条根本页条次"  # 量词
    "度摄"                   # 温度
    "斤吨克"                 # 质量
    "瓦伏安"                 # 电功率/电压/电流
    "帕兆"                   # 压力/兆
    "微纳"                   # 微/纳
    "元万亿"                 # 货币
    "小"                     # "小时"
    "立方平方"                # 面积/体积（用"立"/"平"拦截）
    "赫摩尔焦"               # 物理单位
    "%％℃°"                 # 符号（顺便加进负 lookahead）
)

# 仅匹配段首多级数字前缀 + 空白 + 中文字符。要求：
# - 至少 X.Y（两级）
# - 紧邻后一个汉字非单位字
# - 再后还有 2+ 汉字（避免 "X.Y 秒" / "X.Y 倍" 误伤）
_HAND_NUMBER_PREFIX = re.compile(
    r"^(\s*)(\d+(?:\.\d+){1,6})([\.\s　]+)"
    rf"(?=[\u4e00-\u9fff])(?![{re.escape(_UNIT_CHARS)}])"
    r"(?=[\u4e00-\u9fff]{2,})"
)


def strip_handwritten_numbering_in_body(doc, *, only_normal_style: bool = True) -> int:
    """擦除正文段落首的手写多级编号（"5.1 xxx" → "xxx"）。

    只处理 Normal / 正文 / 空样式段落，不碰 Heading 样式。

    Returns:
        修改段落数
    """
    count = 0
    for para in doc.paragraphs:
        style_name = (para.style.name or "") if para.style else ""
        if only_normal_style:
            if _is_heading_style(style_name):
                continue
        old = para.text
        if not old:
            continue
        new = _HAND_NUMBER_PREFIX.sub(r"\1", old, count=1)
        if new != old:
            _replace_paragraph_text_preserve_format(para, new)
            count += 1
    return count


# ---------- 样式归一化 ----------

def normalize_heading_style_names(doc) -> int:
    """把段落样式名 '标题 N' 统一改为 'Heading N'。"""
    import re as _re
    pat = _re.compile(r"^标题\s*(\d+)$")
    count = 0
    for para in doc.paragraphs:
        st = para.style.name if para.style else ""
        m = pat.match(st or "")
        if not m:
            continue
        n = m.group(1)
        en_name = f"Heading {n}"
        try:
            para.style = doc.styles[en_name]
            count += 1
        except KeyError:
            pass
    return count


# ---------- numPr 剥离 ----------

def strip_numPr_from_body(doc, *, only_heading_styles: bool = True) -> int:
    """剥掉 Heading 样式段落的 w:numPr（破多级列表自动编号绑定）。

    默认 only_heading_styles=True：只对 Heading 样式段剥，**保留正文列表**
    （"1)xxx; 2)xxx; ..." 这种 Word 自动编号列表，用户期望保留）。
    """
    from docx.oxml.ns import qn
    count = 0

    def _should_strip(para) -> bool:
        if not only_heading_styles:
            return True
        st = para.style.name if para.style else ""
        return _is_heading_style(st)

    for para in doc.paragraphs:
        if not _should_strip(para):
            continue
        pPr = para._p.find(qn("w:pPr"))
        if pPr is None:
            continue
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            pPr.remove(numPr)
            count += 1
    # 表格里的段落也要处理
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if not _should_strip(para):
                        continue
                    pPr = para._p.find(qn("w:pPr"))
                    if pPr is None:
                        continue
                    numPr = pPr.find(qn("w:numPr"))
                    if numPr is not None:
                        pPr.remove(numPr)
                        count += 1
    return count


def strip_numPr_from_heading_styles(doc) -> int:
    """剥掉 Heading 样式自身绑定的 Word 自动编号。

    技术标正文使用"文本编号 + Heading 样式"的方案。如果母版 Heading
    style 仍绑定了多级列表，Word/OnlyOffice 会在显示层再自动加一次编号，
    形成 "1.7 1.7 标题"。此函数清掉样式级 numPr，保留字体、字号等样式。
    """
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn

    count = 0
    for style in doc.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        style_name = style.name or ""
        pPr = style.element.find(qn("w:pPr"))
        if pPr is None:
            continue
        outline = pPr.find(qn("w:outlineLvl"))
        is_outline_heading = False
        if outline is not None:
            try:
                is_outline_heading = 0 <= int(outline.get(qn("w:val"))) <= 8
            except (TypeError, ValueError):
                is_outline_heading = False
        if not (_is_heading_style(style_name) or is_outline_heading):
            continue
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            pPr.remove(numPr)
            count += 1
    return count


if __name__ == "__main__":
    # 前缀剥离 smoke test
    print("== strip_prefix ==")
    samples = [
        "1.1 发电机设计原理",
        "第一章 标前概述",
        "1.1.2 外壳结构",
        "一、总体方案",
        "（二）技术参数",
        "第一章  1.1  双编号",
        "纯标题",
        "附  校核报告",
    ]
    for s in samples:
        print(f"  {s!r:35s} → {strip_prefix(s)!r}")

    # 手写编号擦除 smoke test
    print("\n== strip_handwritten (regex only) ==")
    body_samples = [
        "7.10 符合招标公告及招标文件要求的业绩情况",  # 应擦
        "5.1 项目概况",                                # 应擦
        "1. 第一步",                                   # 不擦（单级列表）
        "7.10 m/s 额定风速",                           # 不擦（单位）
        "2.35 秒启动时间",                             # 不擦（数字+单位）
        "一、总体方案",                                # 不擦（中文序号）
        "普通段落",                                    # 不擦
        "1.5 倍超速保护",                              # 不擦（后接"倍"）
    ]
    for s in body_samples:
        new = _HAND_NUMBER_PREFIX.sub("", s, count=1)
        marker = "✗" if new == s else "✓"
        print(f"  {marker}  {s!r:40s} → {new!r}")
