#!/usr/bin/env python3
"""
素材 docx 预处理（方案 B）：对单份素材做

1. 归一化 Heading 样式名（标题 N → Heading N）
2. 剥掉所有段落的 w:numPr（破多级列表绑定，方案 B 必须）
3. 剥 Heading 文本原编号前缀（"1.1 xxx" → "xxx"）
4. 洗正文段落手写编号（"7.10 xxx中文" → "xxx中文"；有就洗，没有就不洗）
5. 去 (新增)/(适配)/(如有) 标签
6. 字段占位符替换 [FIELD] → project_params 值

**不做**编号前缀注入 — 那是 merger 的职责（需要父章节号）

输入输出都是 docx 文件路径，不改原文件。

用法：
    python3 preprocess.py <in.docx> <out.docx> [--params params.json]
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Optional

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent))
from numbering_fixer import (
    strip_heading_prefixes_in_doc,
    strip_handwritten_numbering_in_body,
    strip_numPr_from_body,
    normalize_heading_style_names,
    _replace_paragraph_text_preserve_format,
)


def _normalize_headings_by_outline_level(doc) -> int:
    """共性通用：凡 outlineLvl ∈ [0,5] 的段落 → 强制 style = Heading {lvl+1}。
    覆盖所有 WPS / Word 自定义 heading style（章标题、附件标题1、二级标题、专题标题等）。
    """
    from docx.oxml.ns import qn

    # 建 {styleId: outlineLvl(0-based)} map
    style_outline_map = {}
    styles_el = doc.part.styles._element
    for s in styles_el.findall(qn("w:style")):
        sid = s.get(qn("w:styleId"))
        if not sid:
            continue
        pPr = s.find(qn("w:pPr"))
        if pPr is None:
            continue
        ol = pPr.find(qn("w:outlineLvl"))
        if ol is None:
            continue
        try:
            style_outline_map[sid] = int(ol.get(qn("w:val")))
        except (ValueError, TypeError):
            pass

    count = 0
    for para in doc.paragraphs:
        pPr = para._p.find(qn("w:pPr"))
        if pPr is None:
            continue
        direct_ol = pPr.find(qn("w:outlineLvl"))
        effective_lvl = None
        if direct_ol is not None:
            try:
                effective_lvl = int(direct_ol.get(qn("w:val")))
            except (ValueError, TypeError):
                pass
        if effective_lvl is None:
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None:
                val = pStyle.get(qn("w:val")) or ""
                if val in style_outline_map:
                    effective_lvl = style_outline_map[val]
        if effective_lvl is None or effective_lvl < 0 or effective_lvl > 5:
            continue
        target_style_name = f"Heading {effective_lvl + 1}"
        cur_name = para.style.name if para.style else ""
        if cur_name == target_style_name:
            continue
        try:
            para.style = doc.styles[target_style_name]
            count += 1
        except KeyError:
            pass
    return count


# ---------- 占位符替换 ----------

# project_params 键 → docx 里出现的占位符形式
_PLACEHOLDER_KEY_MAP = {
    "project_name": ["[PROJECT_NAME]", "[项目名称]"],
    "project_short": ["[PROJECT_SHORT]", "[项目简称]"],
    "client_name": ["[CLIENT_NAME]", "[业主]", "[业主名称]"],
    "tender_no": ["[TENDER_NO]", "[招标编号]"],
    "model_no": ["[MODEL_NO]", "[机型号]"],
    "rated_power": ["[RATED_POWER]", "[额定功率]"],
    "rated_speed": ["[RATED_SPEED]", "[额定转速]"],
    "rotor_diameter": ["[ROTOR_DIAMETER]", "[风轮直径]"],
    "hub_height": ["[HUB_HEIGHT]", "[轮毂高度]"],
    "wind_class": ["[WIND_CLASS]", "[风区等级]"],
    "site_location": ["[SITE_LOCATION]", "[场址位置]"],
    "site_altitude": ["[SITE_ALTITUDE]", "[场址海拔]"],
    "delivery_date": ["[DELIVERY_DATE]", "[交货期]"],
    "warranty_years": ["[WARRANTY_YEARS]", "[质保年限]"],
    "cooling_type": ["[COOLING_TYPE]", "[冷却方式]"],
    "bid_date": ["[BID_DATE]", "[投标日期]"],
    "bid_date_cn": ["[BID_DATE_CN]", "[投标日期中文]"],
    "land_area": ["[LAND_AREA]", "[地块]"],
    "purchase_object": ["[PURCHASE_OBJECT]", "[采购对象]"],
}


def _iter_text_containers(doc):
    for para in doc.paragraphs:
        yield para
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


def replace_placeholders(doc, params: dict) -> int:
    if not params:
        return 0
    mapping: dict[str, str] = {}
    for key, aliases in _PLACEHOLDER_KEY_MAP.items():
        val = params.get(key)
        if val is None or val == "":
            continue
        val_str = str(val)
        for ph in aliases:
            mapping[ph] = val_str
    if not mapping:
        return 0
    count = 0
    for para in _iter_text_containers(doc):
        if not para.runs:
            continue
        full = para.text
        if not any(ph in full for ph in mapping):
            continue
        new_full = full
        for ph, val in mapping.items():
            if ph in new_full:
                new_full = new_full.replace(ph, val)
        if new_full == full:
            continue
        _replace_paragraph_text_preserve_format(para, new_full)
        count += 1
    return count


# ---------- 清除 (新增)/(适配)/(如有) 标签 ----------

_TAG_STRIP_PATTERN = re.compile(r"[（(](新增|适配|如有|可选|待定)[)）]")


def strip_tag_marks(doc) -> int:
    count = 0
    for para in _iter_text_containers(doc):
        if not para.runs:
            continue
        full = para.text
        if not _TAG_STRIP_PATTERN.search(full):
            continue
        new_full = _TAG_STRIP_PATTERN.sub("", full)
        if new_full == full:
            continue
        _replace_paragraph_text_preserve_format(para, new_full)
        count += 1
    return count


# ---------- Main ----------

def preprocess(
    in_path: Path,
    out_path: Path,
    params: Optional[dict] = None,
    *,
    verbose: bool = False,
) -> dict:
    doc = Document(str(in_path))

    stats = {
        "heading_by_outline": _normalize_headings_by_outline_level(doc),
        "heading_normalize": normalize_heading_style_names(doc),
        "numPr_stripped": strip_numPr_from_body(doc),
        "heading_prefix_strip": strip_heading_prefixes_in_doc(doc, only_heading_styles=True),
        "body_handwritten_strip": strip_handwritten_numbering_in_body(doc, only_normal_style=True),
        "tag_strip": strip_tag_marks(doc),
        "placeholder_replace": replace_placeholders(doc, params or {}),
    }

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

    if verbose:
        print(f"  preprocess {in_path.name}: {stats}", file=sys.stderr)

    return stats


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input", type=Path)
    ap.add_argument("output", type=Path)
    ap.add_argument("--params", type=Path, default=None)
    ap.add_argument("-v", "--verbose", action="store_true")
    args = ap.parse_args()

    params = {}
    if args.params and args.params.exists():
        params = json.loads(args.params.read_text(encoding="utf-8"))

    stats = preprocess(args.input, args.output, params, verbose=args.verbose)
    print(json.dumps(stats, ensure_ascii=False))


if __name__ == "__main__":
    main()
