#!/usr/bin/env python3
"""
对账与审计：

1. plan.json 中 in-scope 条目数 vs 输出 docx 中 Heading 数对账
2. 残留占位符扫描（[待填写：xxx] / [缺失：xxx] / [FIELD] / [PROJECT_NAME]...）
3. 可疑情况报告：相邻重复 Heading（合并时 toc title + 素材首标题 撞车）

输出：
  <output_dir>/assembly_report.md
  <output_dir>/needs_review.md

用法：
    python3 verify.py \\
        --docx 输出.docx \\
        --plan /tmp/assembly_plan.json \\
        --params project_params.json \\
        --report 工作目录/assembly_report.md \\
        --review 工作目录/needs_review.md
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from collections import Counter
from pathlib import Path

from docx import Document


PLACEHOLDER_PATTERNS = [
    re.compile(r"\[待填写[:：].*?\]"),
    re.compile(r"\[缺失[:：].*?\]"),
    re.compile(r"\[[A-Z_]{4,}\]"),  # [PROJECT_NAME] / [MODEL_NO] 等
]

# 方案 B 硬性检查
ALLOWED_H1_PATTERNS = [
    re.compile(r"^封面$"),
    re.compile(r"^前言(\s|$)"),
    re.compile(r"^第[一二三四五六七八九十]+章(\s|$)"),
    re.compile(r"^附表"),
    re.compile(r"^目录$"),  # finalize 插入的 TOC 标题
]

# 幽灵编号：Heading text 里出现"第 X 章"X > 6 或 > 10 的数字（真正章节最多 6）
GHOST_CHAPTER = re.compile(r"第\s*(\d+)\s*章")
# Heading text 是否以合法编号开头：前言 / 封面 / 第X章 / 数字.数字.. / 附
VALID_HEADING_START = re.compile(r"^(?:封面|前言|第[一二三四五六七八九十\d]+章|\d+(?:\.\d+){0,6}(?:\s|$)|附|附表|目录)")


def scan_docx(docx_path: Path) -> dict:
    from docx.oxml.ns import qn

    doc = Document(str(docx_path))

    heading_counts = Counter()
    heading_list: list[tuple[int, str]] = []
    placeholders: list[str] = []
    dup_alerts: list[str] = []
    ghost_chapters: list[str] = []
    invalid_h1: list[str] = []
    invalid_prefix: list[str] = []
    empty_leaf_headings: list[str] = []

    # Body 级遍历：既能看到段落也能看到表格，正确识别 heading 之间的"内容量"
    W_P = qn("w:p")
    W_TBL = qn("w:tbl")
    W_DRAWING = qn("w:drawing")
    W_PICT = qn("w:pict")

    style_level_map: dict[str, int] = {}
    styles_el = doc.part.styles._element
    for style_el in styles_el.findall(qn("w:style")):
        if style_el.get(qn("w:type")) != "paragraph":
            continue
        sid = style_el.get(qn("w:styleId")) or ""
        if not sid:
            continue
        lvl = None
        name_el = style_el.find(qn("w:name"))
        name = name_el.get(qn("w:val")) if name_el is not None else ""
        m = re.match(r"^(?:Heading|heading|标题)\s*(\d+)$", name or "")
        if m:
            lvl = int(m.group(1))
        if lvl is None:
            outline = style_el.find(qn("w:pPr") + "/" + qn("w:outlineLvl"))
            if outline is not None:
                try:
                    value = int(outline.get(qn("w:val")))
                    if 0 <= value <= 8:
                        lvl = value + 1
                except (TypeError, ValueError):
                    pass
        if lvl is not None:
            style_level_map[sid] = lvl

    def _para_style_level(el):
        pPr = el.find(qn("w:pPr"))
        if pPr is None:
            return None
        direct_outline = pPr.find(qn("w:outlineLvl"))
        if direct_outline is not None:
            try:
                value = int(direct_outline.get(qn("w:val")))
                if 0 <= value <= 8:
                    return value + 1
            except (TypeError, ValueError):
                pass
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is None:
            return None
        val = pStyle.get(qn("w:val")) or ""
        if val in style_level_map:
            return style_level_map[val]
        # styleId 数字 2-6 对应 Heading 2-6，10 对应 H1；也接受以 "Heading" 或 "标题" 开头
        if val in ("1", "10"):
            return 1
        if val in ("2", "3", "4", "5", "6"):
            return int(val)
        m = re.match(r"^(?:Heading|heading|标题)\s*(\d+)$", val)
        if m:
            return int(m.group(1))
        return None

    def _para_text(el) -> str:
        return "".join((t.text or "") for t in el.iter(qn("w:t")))

    def _has_image(el) -> bool:
        return el.find(".//" + W_DRAWING) is not None or el.find(".//" + W_PICT) is not None

    prev_heading: tuple[int, str] | None = None
    last_heading: tuple[int, str] | None = None
    content_for_last: int = 0  # >0 表示该 heading 有文字/表格/图片

    def _flush_last(next_level: int | None):
        """结算 last_heading 的内容计数，必要时登记空章节告警。

        next_level: 下一个 heading 的 level；None 表示文档末尾。
        只在"叶子"位置记录：下一个 heading level <= 本 heading level（或 EOF）。
        """
        nonlocal last_heading, content_for_last
        if last_heading is None:
            return
        is_leaf = (next_level is None) or (next_level <= last_heading[0])
        if is_leaf and content_for_last == 0:
            prev_lvl, prev_text = last_heading
            # 合法空：封面 / 前言 / 目录 / 纯章节标题（第X章 / 第X节）
            if not (
                prev_text in ("封面", "前言", "目录")
                or re.match(r"^第[一二三四五六七八九十\d]+[章节]", prev_text)
                or re.match(r"^\d+(\.\d+)*\s*$", prev_text.strip())
            ):
                empty_leaf_headings.append(f"L{prev_lvl} {prev_text}")

    for child in doc.element.body.iterchildren():
        tag = child.tag
        if tag == W_P:
            lvl = _para_style_level(child)
            text = _para_text(child).strip()

            # 占位符扫描
            for pat in PLACEHOLDER_PATTERNS:
                for m in pat.finditer(text):
                    placeholders.append(m.group(0))

            if lvl is not None:
                # 遇到新 heading：先结算上一个
                _flush_last(lvl)
                heading_counts[f"Heading {lvl}"] += 1
                heading_list.append((lvl, text))
                last_heading = (lvl, text)
                content_for_last = 0

                if prev_heading and prev_heading[1] == text and text:
                    dup_alerts.append(f"L{lvl} 相邻重复：{text}")
                prev_heading = (lvl, text)

                gm = GHOST_CHAPTER.search(text)
                if gm and int(gm.group(1)) > 6:
                    ghost_chapters.append(f"L{lvl} {text!r}")

                if lvl == 1 and text:
                    if not any(pat.search(text) for pat in ALLOWED_H1_PATTERNS):
                        invalid_h1.append(text)

                if text and not VALID_HEADING_START.match(text):
                    invalid_prefix.append(f"L{lvl} {text[:60]!r}")
            else:
                # 非 heading 段：算正文量（文字或图片/drawing）
                if text or _has_image(child):
                    content_for_last += max(1, len(text))
        elif tag == W_TBL:
            # 表格一律计为有内容
            content_for_last += 100
            # 同时扫表格里的占位符
            for t in child.iter(qn("w:t")):
                cell_text = t.text or ""
                for pat in PLACEHOLDER_PATTERNS:
                    for m in pat.finditer(cell_text):
                        placeholders.append(m.group(0))

    # 文档末尾结算
    _flush_last(None)

    return {
        "sections": len(doc.sections),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "heading_counts": dict(heading_counts),
        "heading_list": heading_list,
        "placeholders": placeholders,
        "dup_alerts": dup_alerts,
        "ghost_chapters": ghost_chapters,
        "invalid_h1": invalid_h1,
        "invalid_prefix": invalid_prefix,
        "empty_leaf_headings": empty_leaf_headings,
    }


def build_report(
    docx_path: Path,
    plan: list[dict],
    scan: dict,
    params: dict,
) -> str:
    lines = []
    lines.append(f"# 投标文件生成审计报告")
    lines.append("")
    lines.append(f"- **文件**：{docx_path}")
    lines.append(f"- **大小**：{docx_path.stat().st_size / 1024 / 1024:.1f} MB")
    lines.append(f"- **项目**：{params.get('project_name', '(未填)')}")
    lines.append(f"- **业主**：{params.get('client_name', '(未填)')}")
    lines.append(f"- **招标编号**：{params.get('tender_no', '(未填)')}")
    lines.append("")

    lines.append("## 结构统计")
    lines.append(f"- sections: {scan['sections']}")
    lines.append(f"- paragraphs: {scan['paragraphs']}")
    lines.append(f"- tables: {scan['tables']}")
    lines.append(f"- headings: {json.dumps(scan['heading_counts'], ensure_ascii=False)}")
    lines.append("")

    # plan status 分布
    st_counts = Counter(p["status"] for p in plan)
    lines.append("## 装配计划对账")
    lines.append(f"- plan 总条目：{len(plan)}")
    lines.append(f"- 状态分布：{json.dumps(dict(st_counts), ensure_ascii=False)}")
    in_scope = [p for p in plan if p["status"] != "OUT_OF_SCOPE"]
    lines.append(f"- 应生成 Heading 数（in-scope）：{len(in_scope)}")
    lines.append(f"- 实际 Heading 数：{sum(scan['heading_counts'].values())}")
    lines.append("")

    # 方案 B 硬性检查
    lines.append("## 硬性检查（方案 B）")
    lines.append(f"- 幽灵章节（第 7+ 章）：{len(scan['ghost_chapters'])}")
    lines.append(f"- 非法 H1（非封面/前言/第X章/附表）：{len(scan['invalid_h1'])}")
    lines.append(f"- 非法 Heading 前缀：{len(scan['invalid_prefix'])}")
    lines.append(f"- 相邻重复 Heading：{len(scan['dup_alerts'])}")
    lines.append("")

    if scan["ghost_chapters"]:
        lines.append(f"### ⚠ 幽灵章节编号（应为 0）")
        for g in scan["ghost_chapters"][:20]:
            lines.append(f"- {g}")
        lines.append("")

    if scan["invalid_h1"]:
        lines.append(f"### ⚠ 非法 H1（应为 0）")
        for t in scan["invalid_h1"][:20]:
            lines.append(f"- {t!r}")
        lines.append("")

    if scan["invalid_prefix"]:
        lines.append(f"### 非法 Heading 前缀（前 20）")
        for t in scan["invalid_prefix"][:20]:
            lines.append(f"- {t}")
        if len(scan["invalid_prefix"]) > 20:
            lines.append(f"- ... 还有 {len(scan['invalid_prefix']) - 20} 条")
        lines.append("")

    if scan["dup_alerts"]:
        lines.append(f"### 相邻重复 Heading (前 20)")
        for a in scan["dup_alerts"][:20]:
            lines.append(f"- {a}")
        if len(scan["dup_alerts"]) > 20:
            lines.append(f"- ... 还有 {len(scan['dup_alerts']) - 20} 条")
        lines.append("")

    if scan["placeholders"]:
        ph_counter = Counter(scan["placeholders"])
        lines.append(f"## 残留占位符 ({len(scan['placeholders'])})")
        for ph, cnt in ph_counter.most_common(30):
            lines.append(f"- `{ph}` × {cnt}")
        lines.append("")

    # 字段未填
    unfilled = [k for k, v in params.items() if v is None or v == ""]
    if unfilled:
        lines.append("## project_params 未填字段")
        for k in unfilled:
            lines.append(f"- `{k}`")
        lines.append("")

    return "\n".join(lines)


def build_needs_review(plan: list[dict], params: dict, scan: dict) -> str:
    lines = ["# 人工补齐清单", "", "> 一把出后需要你确认的全部项目集中在这里。", ""]

    # 1. project_params 未填字段（占位符 / null）
    def _is_placeholder(v) -> bool:
        return isinstance(v, str) and v.startswith("[待填写")

    ph_fields = [(k, v) for k, v in params.items() if _is_placeholder(v)]
    null_fields = [k for k, v in params.items() if v is None or v == ""]
    if ph_fields or null_fields:
        lines.append(f"## project_params 待补字段 ({len(ph_fields) + len(null_fields)})")
        lines.append("")
        lines.append("编辑 `project_params.json` 把占位符换成真值后重跑 assembler 即可。")
        lines.append("")
        for k, v in ph_fields:
            lines.append(f"- `{k}` = `{v}`")
        for k in null_fields:
            lines.append(f"- `{k}` = null（未抽到）")
        lines.append("")

    # 2. 残留占位符（从 docx 扫到的 [待填写：xx] / [FIELD] / [缺失：xx]）
    if scan.get("placeholders"):
        ph_counter = Counter(scan["placeholders"])
        lines.append(f"## 正文残留占位符 ({len(scan['placeholders'])})")
        lines.append("")
        for ph, cnt in ph_counter.most_common():
            lines.append(f"- `{ph}` × {cnt}")
        lines.append("")

    # 3. [新增] 条目
    needs = [p for p in plan if p["status"] == "NEEDS_REVIEW"]
    if needs:
        lines.append(f"## [新增] 条目 ({len(needs)}) — 需补素材")
        lines.append("")
        for p in needs:
            lines.append(f"- {p['chapter_no']} {p['title']}")
        lines.append("")

    # 4. [未匹配] 条目
    unmatched = [p for p in plan if p["status"] == "UNMATCHED"]
    if unmatched:
        lines.append(f"## [未匹配] 条目 ({len(unmatched)}) — wiki 未命中")
        lines.append("")
        for p in unmatched:
            lines.append(f"- {p['chapter_no']} {p['title']}  ← {p.get('note', '')}")
        lines.append("")

    # 5. [适配] 条目（字段替换已做，需核对）
    adapted = [p for p in plan if p["status"] == "ADAPTED"]
    if adapted:
        lines.append(f"## [适配] 条目 ({len(adapted)}) — 需核对占位符替换")
        lines.append("")
        for p in adapted:
            lines.append(f"- {p['chapter_no']} {p['title']}")
        lines.append("")

    # 6. 空章节（heading 后无正文 — 常暴露素材归位错误）
    empty = scan.get("empty_leaf_headings", [])
    if empty:
        lines.append(f"## 空章节告警 ({len(empty)}) — Heading 后无正文")
        lines.append("")
        lines.append("> 多为 wiki 卡片 `skeleton_section` 归位错或素材本身是空框架。")
        lines.append("")
        for h in empty[:50]:
            lines.append(f"- {h}")
        if len(empty) > 50:
            lines.append(f"- ... 还有 {len(empty) - 50} 条")
        lines.append("")

    # 7. 结构告警（幽灵章节 / 非法 H1 / 相邻重复）
    warn_parts = []
    if scan.get("ghost_chapters"):
        warn_parts.append(f"幽灵章节 {len(scan['ghost_chapters'])}")
    if scan.get("invalid_h1"):
        warn_parts.append(f"非法 H1 {len(scan['invalid_h1'])}")
    if scan.get("dup_alerts"):
        warn_parts.append(f"相邻重复 {len(scan['dup_alerts'])}")
    if warn_parts:
        lines.append(f"## 结构告警 — 详见 `assembly_report.md`")
        lines.append("")
        lines.append(f"- {' / '.join(warn_parts)}")
        lines.append("")

    if len(lines) <= 4:
        lines.append("_无待补项，可直接交付。_")

    return "\n".join(lines)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", type=Path, required=True)
    ap.add_argument("--plan", type=Path, required=True)
    ap.add_argument("--params", type=Path, required=True)
    ap.add_argument("--report", type=Path, required=True)
    ap.add_argument("--review", type=Path, required=True)
    args = ap.parse_args()

    plan = json.loads(args.plan.read_text(encoding="utf-8"))
    params = json.loads(args.params.read_text(encoding="utf-8"))
    scan = scan_docx(args.docx)

    report = build_report(args.docx, plan, scan, params)
    review = build_needs_review(plan, params, scan)

    args.report.parent.mkdir(parents=True, exist_ok=True)
    args.report.write_text(report, encoding="utf-8")
    args.review.write_text(review, encoding="utf-8")

    print(f"[OK] report → {args.report}")
    print(f"[OK] review → {args.review}")
    print()
    # stdout 摘要
    print("=== 摘要 ===")
    st_counts = Counter(p["status"] for p in plan)
    print(f"plan: {dict(st_counts)}")
    print(f"heading: {scan['heading_counts']}")
    print(f"placeholders: {len(scan['placeholders'])}")
    print(f"dup_alerts: {len(scan['dup_alerts'])}")


if __name__ == "__main__":
    main()
