#!/usr/bin/env python3
"""Fill a requested technical bid appendix/table from manifest data."""

from __future__ import annotations

import argparse
import json
import re
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document


SCHEMA_VERSION = "bid-tech-table-fill-v1"


def now_iso() -> str:
    return datetime.now(UTC).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def load_manifest(path: Path) -> dict[str, Any]:
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise RuntimeError("manifest must be a JSON object")
    return data


def object_items(value: Any) -> list[dict[str, Any]]:
    return [item for item in value if isinstance(item, dict)] if isinstance(value, list) else []


def string_items(value: Any) -> list[str]:
    if not isinstance(value, list):
        return []
    result: list[str] = []
    seen: set[str] = set()
    for item in value:
        text = str(item or "").strip()
        if not text or text in seen:
            continue
        seen.add(text)
        result.append(text)
    return result


def compact(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def material_label(material: dict[str, Any]) -> str:
    return compact(material.get("name") or material.get("id") or material.get("path") or material.get("folderPath"))


def field_label(field: dict[str, Any]) -> str:
    return compact(field.get("label") or field.get("title") or field.get("key") or field.get("id"))


def field_value(field: dict[str, Any]) -> str:
    return compact(field.get("value") or field.get("keyValue") or "")


def candidate_value_for_title(title: str, fields: list[dict[str, Any]], materials: list[dict[str, Any]], turbine: dict[str, Any]) -> tuple[str, str]:
    key = compact(title)
    normalized_key = re.sub(r"\W+", "", key.lower())
    for field in fields:
        label = field_label(field)
        if not label:
            continue
        normalized_label = re.sub(r"\W+", "", label.lower())
        value = field_value(field)
        if value and (normalized_label in normalized_key or normalized_key in normalized_label):
            return value, f"解析字段：{label}"

    turbine_map = {
        "投标机型": turbine.get("model"),
        "机型": turbine.get("model"),
        "平台": turbine.get("platform"),
        "单机容量": f"{turbine.get('ratedPowerKw')}kW" if turbine.get("ratedPowerKw") else "",
        "叶轮直径": f"{turbine.get('rotorDiameterM')}m" if turbine.get("rotorDiameterM") else "",
    }
    for label, value in turbine_map.items():
        if value and label in key:
            return str(value), f"项目投标机型参数：{label}"

    if materials:
        return f"参见{material_label(materials[0])}", f"参考素材：{materials[0].get('id') or material_label(materials[0])}"
    return "", ""


def load_blank_rows(blank_source: dict[str, Any], appendix_task: dict[str, Any]) -> tuple[list[list[str]], str]:
    for path_value in (
        blank_source.get("docxPath"),
        appendix_task.get("docxPath"),
        blank_source.get("workspacePath"),
        appendix_task.get("workspacePath"),
    ):
        text = str(path_value or "").strip()
        if not text:
            continue
        path = Path(text)
        if not path.is_absolute():
            continue
        if not path.exists() or path.suffix.lower() != ".docx":
            continue
        try:
            doc = Document(str(path))
        except Exception:
            continue
        for table in doc.tables:
            rows = [[compact(cell.text) for cell in row.cells] for row in table.rows]
            if rows:
                return rows, str(path)
    return [], ""


def fill_rows(
    rows: list[list[str]],
    fields: list[dict[str, Any]],
    materials: list[dict[str, Any]],
    turbine: dict[str, Any],
) -> tuple[list[list[str]], list[str], list[dict[str, Any]]]:
    if not rows:
        return [], ["未找到可复制的原始空表结构，已生成填充说明供人工补全"], []

    output: list[list[str]] = []
    unfilled: list[str] = []
    filled_fields: list[dict[str, Any]] = []
    max_columns = max(len(row) for row in rows)
    header = rows[0] if rows else []
    output.append(header)

    for row_index, row in enumerate(rows[1:], start=2):
        next_row = list(row) + [""] * max(0, max_columns - len(row))
        empty_indexes = [index for index, cell in enumerate(next_row) if not compact(cell)]
        title = next((cell for cell in next_row if compact(cell)), f"第{row_index}行")
        if empty_indexes:
            value, source = candidate_value_for_title(title, fields, materials, turbine)
            if value:
                next_row[empty_indexes[0]] = value
                filled_fields.append({"field": title, "value": value, "source": source, "row": row_index})
            else:
                marker = f"待人工补充：{title}"
                next_row[empty_indexes[0]] = marker
                unfilled.append(title)
        output.append(next_row[:max_columns])

    return output, unfilled, filled_fields


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "内容"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value


def add_rows_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index in range(column_count):
            table.cell(row_index, col_index).text = row[col_index] if col_index < len(row) else ""


def write_output_docx(
    output_file: Path,
    manifest: dict[str, Any],
    filled_rows: list[list[str]],
    unfilled_fields: list[str],
    filled_fields: list[dict[str, Any]],
    blank_docx_path: str,
) -> None:
    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    title = compact(manifest.get("title")) or "AI 填写产物"
    gap_item = manifest.get("gapItem") if isinstance(manifest.get("gapItem"), dict) else {}
    appendix_task = manifest.get("appendixTask") if isinstance(manifest.get("appendixTask"), dict) else {}
    turbine = manifest.get("projectTurbineModel") if isinstance(manifest.get("projectTurbineModel"), dict) else {}
    materials = object_items(manifest.get("referenceMaterials"))
    fields = object_items(manifest.get("parseFields"))

    doc.add_heading(title, level=1)
    doc.add_paragraph("本文件由 bid-tech-table-filler 根据 manifest 中限定的空表、参考素材、解析字段和投标机型参数生成。")
    add_key_value_table(
        doc,
        [
            ("缺口编号", compact(manifest.get("gapId"))),
            ("目录项", compact(f"{gap_item.get('number') or ''} {gap_item.get('title') or title}")),
            ("空副表", compact(appendix_task.get("title") or (manifest.get("blankSource") or {}).get("title") or "")),
            ("空表来源", blank_docx_path or compact(appendix_task.get("workspacePath") or "")),
            ("投标机型", " / ".join(compact(turbine.get(key)) for key in ("model", "platform") if compact(turbine.get(key))) or "未指定"),
        ],
    )

    doc.add_heading("一、填充结果", level=2)
    if filled_rows:
        add_rows_table(doc, filled_rows)
    else:
        doc.add_paragraph("未找到原始空表结构，已在下方列出可用素材与字段，需人工回填到招标空表。")

    doc.add_heading("二、参考素材", level=2)
    if materials:
        add_key_value_table(
            doc,
            [
                (
                    material.get("id") or material_label(material),
                    "；".join(
                        part
                        for part in (
                            material_label(material),
                            compact(material.get("folderPath")),
                            compact(material.get("usage")),
                        )
                        if part
                    ),
                )
                for material in materials
            ],
        )
    else:
        doc.add_paragraph("未指定参考素材。")

    doc.add_heading("三、解析字段", level=2)
    if fields:
        add_key_value_table(
            doc,
            [(field_label(field), field_value(field) or compact(field.get("evidence"))) for field in fields],
        )
    else:
        doc.add_paragraph("未指定解析字段。")

    doc.add_heading("四、未填字段", level=2)
    if unfilled_fields:
        for item in unfilled_fields:
            doc.add_paragraph(item, style=None)
    else:
        doc.add_paragraph("本次未发现必须留空的字段，仍需人工复核原始证明材料。")

    if filled_fields:
        doc.add_heading("五、填充依据", level=2)
        add_key_value_table(
            doc,
            [
                (str(item.get("field") or ""), f"{item.get('value') or ''}（{item.get('source') or '未标注'}）")
                for item in filled_fields
            ],
        )

    doc.save(output_file)


def run_from_manifest(manifest_path: Path) -> dict[str, Any]:
    manifest = load_manifest(manifest_path)
    output_file = Path(str(manifest.get("outputFile") or manifest_path.with_name("AI填写.docx")))
    blank_source = manifest.get("blankSource") if isinstance(manifest.get("blankSource"), dict) else {}
    appendix_task = manifest.get("appendixTask") if isinstance(manifest.get("appendixTask"), dict) else {}
    fields = object_items(manifest.get("parseFields"))
    materials = object_items(manifest.get("referenceMaterials"))
    turbine = manifest.get("projectTurbineModel") if isinstance(manifest.get("projectTurbineModel"), dict) else {}

    rows, blank_docx_path = load_blank_rows(blank_source, appendix_task)
    filled_rows, unfilled_fields, filled_fields = fill_rows(rows, fields, materials, turbine)
    write_output_docx(output_file, manifest, filled_rows, unfilled_fields, filled_fields, blank_docx_path)

    evidence_refs = [
        {"type": "material", "id": compact(material.get("id") or material.get("materialId")), "title": material_label(material)}
        for material in materials
        if compact(material.get("id") or material.get("materialId"))
    ]
    evidence_refs.extend(
        {"type": "parse_field", "id": compact(field.get("id") or field.get("key")), "title": field_label(field)}
        for field in fields
        if compact(field.get("id") or field.get("key"))
    )
    if blank_docx_path:
        evidence_refs.append({"type": "blank_source", "path": blank_docx_path, "title": compact(blank_source.get("title"))})

    return {
        "schema_version": SCHEMA_VERSION,
        "outputFile": str(output_file),
        "unfilledFields": unfilled_fields,
        "evidenceRefs": evidence_refs,
        "fillReport": {
            "filledFieldCount": len(filled_fields),
            "unfilledFieldCount": len(unfilled_fields),
            "referenceMaterialCount": len(materials),
            "parseFieldCount": len(fields),
            "blankDocxPath": blank_docx_path,
        },
        "filledAt": now_iso(),
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--manifest", required=True)
    parser.add_argument("--response", choices=("summary", "full"), default="summary")
    args = parser.parse_args()
    result = run_from_manifest(Path(args.manifest))
    if args.response == "summary":
        summary = {
            "schema_version": result["schema_version"],
            "outputFile": result["outputFile"],
            "unfilledFields": result["unfilledFields"],
            "evidenceRefs": result["evidenceRefs"],
            "fillReport": result["fillReport"],
        }
        print(json.dumps(summary, ensure_ascii=False, indent=2))
    else:
        print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
