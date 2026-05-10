#!/usr/bin/env python3
"""Run bid-business-assembler from a backend-prepared manifest."""

from __future__ import annotations

import argparse
import json
import mimetypes
import re
import shutil
import sys
import tempfile
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt

try:  # PyMuPDF is available in backend requirements.
    import fitz  # type: ignore
except Exception:  # pragma: no cover - dependency fallback
    fitz = None

try:
    from docxcompose.composer import Composer
except Exception:  # pragma: no cover - dependency fallback
    Composer = None

SCHEMA_VERSION = "bid-business-assembly-v1"
PLAN_SCHEMA_VERSION = "bid-business-assembly-plan-v1"
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
DOCX_SUFFIXES = {".docx"}
PDF_SUFFIXES = {".pdf"}
FIELD_PATTERN = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}|\[\[\s*([^\[\]]+?)\s*\]\]")
MAX_TEMPLATE_WHOLE_MERGE_BYTES = 50 * 1024 * 1024


class AssemblyContext:
    def __init__(self, manifest: dict[str, Any], manifest_path: Path) -> None:
        self.manifest = manifest
        self.manifest_path = manifest_path
        self.project_id = str(manifest.get("projectId") or "")
        self.project_name = str(manifest.get("projectName") or self.project_id or "商务标项目")
        self.work_dir = as_path(manifest.get("workDir"), required=False) or manifest_path.parent
        output_text = str(manifest.get("outputFile") or "").strip()
        if not output_text:
            raise RuntimeError("outputFile is required")
        self.output_file = Path(output_text).expanduser()
        self.plan_file = self.work_dir / "business_assembly_plan.json"
        self.report_file = self.work_dir / "business_assembly_report.md"
        self.review_file = self.work_dir / "business_needs_review.md"
        self.attachment_manifest_file = self.work_dir / "attachment_manifest.json"
        self.field_fill_report_file = self.work_dir / "field_fill_report.json"
        self.assets_dir = self.work_dir / "embedded-assets"
        self.review_items: list[dict[str, Any]] = []
        self.attachment_items: list[dict[str, Any]] = []
        self.field_events: list[dict[str, Any]] = []
        self.section_results: list[dict[str, Any]] = []
        self.used_artifact_keys: set[str] = set()
        self.counters: Counter[str] = Counter()

    def add_review(self, kind: str, message: str, *, section: str = "", source: str = "", level: str = "warning") -> None:
        key = (kind, message, section, source)
        existing = {(item["kind"], item["message"], item.get("section", ""), item.get("source", "")) for item in self.review_items}
        if key in existing:
            return
        self.review_items.append(
            {
                "kind": kind,
                "level": level,
                "message": message,
                "section": section,
                "source": source,
            }
        )


def main() -> None:
    parser = argparse.ArgumentParser(description="Assemble business bid Word from manifest.")
    parser.add_argument("manifest", nargs="?")
    parser.add_argument("--manifest", dest="manifest_option")
    parser.add_argument("--response", choices=("summary", "full"), default="summary")
    args = parser.parse_args()

    manifest_path = Path(str(args.manifest_option or args.manifest or "")).expanduser()
    if not manifest_path.exists():
        raise SystemExit(f"manifest not found: {manifest_path}")

    try:
        output = run_from_manifest(manifest_path)
    except Exception as exc:
        print(f"run_from_manifest failed: {exc}", file=sys.stderr)
        raise SystemExit(1) from exc

    payload = output if args.response == "full" else summary_payload(output)
    print(json.dumps(payload, ensure_ascii=False, indent=2))


def run_from_manifest(manifest_path: Path) -> dict[str, Any]:
    manifest = load_json(manifest_path)
    validate_manifest(manifest)
    ctx = AssemblyContext(manifest, manifest_path)
    ctx.work_dir.mkdir(parents=True, exist_ok=True)
    ctx.assets_dir.mkdir(parents=True, exist_ok=True)

    toc = load_json(as_path(manifest.get("tocJsonPath"), required=True, label="tocJsonPath"))
    gap_plan = load_json(as_path(manifest.get("businessGapPlanPath"), required=True, label="businessGapPlanPath"))
    fact_table = load_json(as_path(manifest.get("projectFactTablePath"), required=False) or ctx.work_dir / "missing_project_fact_table.json", default={})
    parse_result = load_json(as_path(manifest.get("parseResultPath"), required=False) or ctx.work_dir / "missing_parse_result.json", default={})

    toc_items = normalize_toc_items(toc)
    tasks = [task for task in gap_plan.get("tasks") or [] if isinstance(task, dict)]
    toc_refs = [ref for ref in gap_plan.get("tocRefs") or [] if isinstance(ref, dict)]
    facts = build_fact_values(fact_table, parse_result, manifest)
    scoring_artifacts = collect_scoring_artifacts(parse_result, tasks)
    parse_artifacts = collect_parse_artifacts(parse_result)
    material_templates = collect_material_templates(manifest)

    if str(fact_table.get("status") or "") != "confirmed":
        ctx.add_review("project_fact_table_unconfirmed", "项目事实表未确认，本次允许生成但需人工复核所有项目字段。")
    if has_unconfirmed_gap_tasks(tasks):
        ctx.add_review("business_gap_unconfirmed", "商务 S3 存在未完全确认任务，本次允许生成但需人工复核相关章节。")

    doc = Document()
    setup_base_styles(doc)
    doc.add_heading(f"{ctx.project_name} 商务投标文件", level=0)
    subtitle = doc.add_paragraph("由商务标 S4 生成标书链路自动装配，需在 S5 共创阶段复核签字、盖章、报价和证书有效期。")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    task_by_toc = group_tasks_by_toc(tasks)
    emitted_scoring = False
    assembled_count = 0
    placeholder_count = 0

    for item in toc_items:
        title = str(item.get("title") or "未命名章节")
        level = max(1, min(int(item.get("level") or 1), 6))
        heading = heading_text(item)
        doc.add_heading(heading, level=level)
        bound_tasks = bound_tasks_for_item(item, task_by_toc)
        section_status = "placeholder"
        section_artifacts: list[dict[str, Any]] = []
        risk_flags: list[str] = []
        section_emitted_scoring = False

        if not bound_tasks and is_scoring_section(title) and scoring_artifacts:
            section_status = "review_required"
        elif not bound_tasks:
            doc.add_paragraph(f"[待填写：{title}]")
            placeholder_count += 1
            risk_flags.append("no_business_gap_task")
            ctx.add_review("no_business_gap_task", f"目录章节未绑定商务 S3 任务：{title}", section=title)
        else:
            for task in bound_tasks:
                assembly_mode = task_assembly_mode(task)
                task_artifacts = task_artifacts_for_assembly(task)
                if not task_artifacts:
                    template_artifact = best_material_template_for_task(task, material_templates) if allows_template_fallback(task) else None
                    if template_artifact:
                        task_artifacts = [template_artifact]
                if not task_artifacts:
                    task_artifacts = best_parse_artifacts_for_task(task, parse_artifacts) if allows_parse_artifact_fallback(task) else []
                if str(task.get("status") or "") not in {"ready", "resolved"}:
                    risk_flags.append("business_gap_task_unconfirmed")
                    ctx.add_review(
                        "business_gap_task_unconfirmed",
                        f"任务未完全确认但已进入生成：{task.get('title') or title}",
                        section=title,
                        source=str(task.get("id") or ""),
                    )
                if task_artifacts:
                    for artifact in task_artifacts:
                        artifact = artifact_with_task_intent(task, artifact)
                        if emit_artifact(doc, artifact, facts, ctx, section_title=title):
                            assembled_count += 1
                            section_status = "assembled"
                            if is_business_scoring_artifact(artifact):
                                emitted_scoring = True
                                section_emitted_scoring = True
                        else:
                            placeholder_count += 1
                            section_status = "review_required"
                        section_artifacts.append(artifact)
                else:
                    generated = False
                    if should_emit_controlled_draft(task):
                        generated = emit_generated_draft(doc, task, facts, ctx, section_title=title)
                    if generated:
                        assembled_count += 1
                        section_status = "generated_draft"
                        risk_flags.append("generated_basic_draft")
                    else:
                        doc.add_paragraph(f"[待填写：{task.get('title') or title}]")
                        placeholder_count += 1
                        section_status = "placeholder"
                        risk_flags.append("missing_material")
                        if assembly_mode:
                            risk_flags.append(f"assembly_mode_{assembly_mode}")
                        ctx.add_review("missing_material", f"任务缺少可装配材料：{task.get('title') or title}", section=title)

        if is_scoring_section(title) and scoring_artifacts and not section_emitted_scoring:
            for artifact in scoring_artifacts:
                if artifact_key(artifact) in ctx.used_artifact_keys:
                    continue
                if emit_artifact(doc, artifact, facts, ctx, section_title=title):
                    emitted_scoring = True
                    section_emitted_scoring = True
                    assembled_count += 1
                    section_status = "assembled"
                    section_artifacts.append(artifact)

        ctx.section_results.append(
            {
                "tocNodeId": str(item.get("nodeId") or ""),
                "number": str(item.get("number") or ""),
                "title": title,
                "level": level,
                "status": section_status,
                "taskIds": [str(task.get("id") or "") for task in bound_tasks],
                "artifactRefs": [artifact_key(artifact) for artifact in section_artifacts],
                "riskFlags": sorted(set(risk_flags)),
            }
        )

    if scoring_artifacts and not emitted_scoring:
        doc.add_heading("商务评分标准", level=1)
        for artifact in scoring_artifacts:
            if emit_artifact(doc, artifact, facts, ctx, section_title="商务评分标准"):
                assembled_count += 1
                emitted_scoring = True
        ctx.section_results.append(
            {
                "tocNodeId": "BUSINESS-SCORING-AUTO",
                "number": "",
                "title": "商务评分标准",
                "level": 1,
                "status": "assembled" if emitted_scoring else "placeholder",
                "taskIds": [],
                "artifactRefs": [artifact_key(artifact) for artifact in scoring_artifacts],
                "riskFlags": [] if emitted_scoring else ["business_scoring_missing"],
            }
        )
    elif not scoring_artifacts:
        ctx.add_review("business_scoring_missing", "商务评分标准未找到可装配产物，需人工补充。")

    ctx.output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(ctx.output_file))

    summary = {
        "sectionCount": len(ctx.section_results),
        "assembledCount": assembled_count,
        "placeholderCount": placeholder_count,
        "reviewRequiredCount": len(ctx.review_items),
        "embeddedAttachmentCount": ctx.counters["embedded"],
        "referencedAttachmentCount": ctx.counters["referenced"],
    }
    write_plan(ctx, summary)
    write_reports(ctx, summary)

    return {
        "schema_version": SCHEMA_VERSION,
        "workDir": str(ctx.work_dir),
        "outputFile": str(ctx.output_file),
        "assemblyReport": str(ctx.report_file),
        "needsReview": str(ctx.review_file),
        "planFile": str(ctx.plan_file),
        "attachmentManifest": str(ctx.attachment_manifest_file),
        "fieldFillReport": str(ctx.field_fill_report_file),
        "summary": summary,
    }


def validate_manifest(manifest: dict[str, Any]) -> None:
    required = [
        "projectId",
        "projectName",
        "bidType",
        "workDir",
        "tocJsonPath",
        "businessGapPlanPath",
        "projectFactTablePath",
        "parseResultPath",
        "materialLibraryDir",
        "outputFile",
    ]
    missing = [key for key in required if key not in manifest]
    if missing:
        raise RuntimeError(f"manifest missing required keys: {', '.join(missing)}")
    if "商务" not in str(manifest.get("bidType") or ""):
        raise RuntimeError("bid-business-assembler only accepts bidType=商务标")
    for key in ("tocJsonPath", "businessGapPlanPath"):
        path = as_path(manifest.get(key), required=True, label=key)
        assert path is not None


def as_path(value: Any, *, required: bool = False, label: str = "path") -> Path | None:
    text = str(value or "").strip()
    if not text:
        if required:
            raise RuntimeError(f"{label} is required")
        return None
    path = Path(text).expanduser()
    if required and not path.exists():
        raise RuntimeError(f"{label} does not exist: {path}")
    return path


def load_json(path: Path | None, default: dict[str, Any] | None = None) -> dict[str, Any]:
    if path is None or not path.exists():
        return dict(default or {})
    data = json.loads(path.read_text(encoding="utf-8"))
    return data if isinstance(data, dict) else dict(default or {})


def normalize_toc_items(toc: dict[str, Any]) -> list[dict[str, Any]]:
    raw_items = toc.get("items") if isinstance(toc.get("items"), list) else []
    items: list[dict[str, Any]] = []
    stack: list[tuple[int, str]] = []
    counters: list[int] = []
    for index, raw in enumerate(raw_items, start=1):
        if not isinstance(raw, dict):
            continue
        level = max(1, min(int(raw.get("level") or 1), 8))
        while stack and stack[-1][0] >= level:
            stack.pop()
        counters = counters[:level]
        if len(counters) < level:
            counters.extend([0] * (level - len(counters)))
        counters[level - 1] += 1
        generated_node_id = "TOC-" + "-".join(str(part) for part in counters[:level] if part)
        parent_id = stack[-1][1] if stack else ""
        item = {
            **raw,
            "nodeId": str(raw.get("nodeId") or raw.get("itemId") or raw.get("id") or generated_node_id),
            "parentNodeId": str(raw.get("parentNodeId") or raw.get("parentId") or parent_id),
            "number": str(raw.get("number") or raw.get("tocNumber") or ""),
            "title": str(raw.get("title") or raw.get("name") or f"商务目录项{index}").strip(),
            "level": level,
            "order": int(raw.get("order") or index),
        }
        items.append(item)
        stack.append((level, item["nodeId"]))
    return items


def heading_text(item: dict[str, Any]) -> str:
    number = str(item.get("number") or "").strip()
    title = str(item.get("title") or "未命名章节").strip()
    return f"{number} {title}".strip()


def setup_base_styles(doc: Document) -> None:
    try:
        style = doc.styles["Normal"]
        style.font.name = "宋体"
        style.font.size = Pt(10.5)
    except Exception:
        pass


def build_fact_values(fact_table: dict[str, Any], parse_result: dict[str, Any], manifest: dict[str, Any]) -> dict[str, str]:
    values: dict[str, str] = {}

    def add(label: str, value: Any) -> None:
        clean_label = normalize_field_label(label)
        text = str(value or "").strip()
        if clean_label and text and clean_label not in values:
            values[clean_label] = text

    for field in fact_table.get("fields") or []:
        if not isinstance(field, dict):
            continue
        add(str(field.get("label") or field.get("fieldName") or field.get("key") or ""), field.get("value"))
    structured = parse_result.get("structured") if isinstance(parse_result.get("structured"), dict) else {}
    for field in structured.get("projectFactFields") or []:
        if not isinstance(field, dict):
            continue
        add(str(field.get("label") or field.get("fieldName") or field.get("fieldKey") or ""), field.get("value") or field.get("keyValue"))
    add("项目名称", manifest.get("projectName"))
    project = manifest.get("project") if isinstance(manifest.get("project"), dict) else {}
    add("招标编号", project.get("projectCode"))
    add("招标人", project.get("customerName") or project.get("owner"))
    add("投标人", project.get("bidderName") or "上海电气风电集团股份有限公司")
    return values


def normalize_field_label(label: str) -> str:
    text = re.sub(r"[\s:：，,。；;（）()【】\[\]]+", "", str(label or ""))
    aliases = {
        "招标项目名称": "项目名称",
        "工程名称": "项目名称",
        "项目编号": "招标编号",
        "招标项目编号": "招标编号",
        "标段编号": "招标编号",
        "采购人": "招标人",
        "项目单位": "招标人",
        "业主": "招标人",
        "投标单位": "投标人",
        "供应商": "投标人",
        "签署日期": "日期",
        "投标日期": "日期",
        "报价": "投标报价",
        "总价": "投标报价",
    }
    return aliases.get(text, text)


def fact_value(facts: dict[str, str], label: str) -> str:
    key = normalize_field_label(label)
    if key in facts:
        return facts[key]
    for existing, value in facts.items():
        if key and (key in existing or existing in key):
            return value
    return ""


def group_tasks_by_toc(tasks: list[dict[str, Any]]) -> dict[str, list[dict[str, Any]]]:
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for task in tasks:
        target = task.get("tocTarget") if isinstance(task.get("tocTarget"), dict) else {}
        for key in ("nodeId", "title", "number"):
            value = str(target.get(key) or "").strip()
            if value:
                grouped[value].append(task)
    return grouped


def bound_tasks_for_item(item: dict[str, Any], grouped: dict[str, list[dict[str, Any]]]) -> list[dict[str, Any]]:
    seen: set[int] = set()
    result: list[dict[str, Any]] = []
    for key in (str(item.get("nodeId") or ""), str(item.get("title") or ""), str(item.get("number") or "")):
        for task in grouped.get(key, []):
            marker = id(task)
            if marker in seen:
                continue
            seen.add(marker)
            result.append(task)
    return result


def task_artifacts_for_assembly(task: dict[str, Any]) -> list[dict[str, Any]]:
    artifacts = [item for item in task.get("resolvedArtifacts") or [] if isinstance(item, dict)]
    artifacts = [item for item in artifacts if str(item.get("filePath") or item.get("path") or "").strip()]
    return artifacts


def task_assembly_mode(task: dict[str, Any]) -> str:
    mode = str(task.get("assemblyMode") or "").strip()
    if mode:
        return mode
    fill_plan = task.get("fillPlan") if isinstance(task.get("fillPlan"), dict) else {}
    mode = str(fill_plan.get("mode") or "").strip()
    if mode:
        return mode
    task_type = str(task.get("taskType") or "")
    decision = str(task.get("decision") or "")
    if task_type == "table":
        return "table_fill_from_material"
    if task_type == "certificate":
        return "embed_scan_or_image"
    if task_type in {"attachment", "bundle"}:
        return "attach_whole_file"
    if decision == "fill_required":
        return "template_fill_docx"
    if should_emit_controlled_draft(task):
        return "ai_draft"
    return ""


def allows_template_fallback(task: dict[str, Any]) -> bool:
    return task_assembly_mode(task) in {"", "template_fill_docx", "table_fill_from_material", "ai_draft"}


def allows_parse_artifact_fallback(task: dict[str, Any]) -> bool:
    return task_assembly_mode(task) in {"", "template_fill_docx", "table_fill_from_material", "ai_draft", "attach_whole_file"}


def artifact_with_task_intent(task: dict[str, Any], artifact: dict[str, Any]) -> dict[str, Any]:
    output = dict(artifact)
    output.setdefault("assemblyMode", task_assembly_mode(task))
    output.setdefault("materialUsage", str(task.get("materialUsage") or ""))
    if not output.get("selectedEvidenceSegments") and isinstance(task.get("selectedEvidenceSegments"), list):
        output["selectedEvidenceSegments"] = task.get("selectedEvidenceSegments") or []
    if not output.get("fillPlan") and isinstance(task.get("fillPlan"), dict):
        output["fillPlan"] = task.get("fillPlan") or {}
    return output


def should_emit_controlled_draft(task: dict[str, Any]) -> bool:
    return (
        str(task.get("assigneeMode") or "") == "ai_draft"
        or str(task.get("decision") or "") == "ai_draft_required"
        or bool(task.get("allowS4AiDraft"))
    )


def collect_parse_artifacts(parse_result: dict[str, Any]) -> list[dict[str, Any]]:
    structured = parse_result.get("structured") if isinstance(parse_result.get("structured"), dict) else {}
    artifacts: list[dict[str, Any]] = []
    for key, artifact_type, source_mode in (
        ("appendices", "parse_appendix_template", "parsed_from_tender_attachment_template"),
        ("commitmentLetters", "parse_commitment_letter", "generated_by_s1_business_parser"),
    ):
        for index, item in enumerate(structured.get(key) or [], start=1):
            if not isinstance(item, dict):
                continue
            path = str(item.get("docxPath") or item.get("filePath") or item.get("path") or "").strip()
            if not path:
                continue
            title = str(item.get("title") or item.get("name") or Path(path).stem)
            artifacts.append(
                {
                    "artifactId": str(item.get("id") or f"{artifact_type}-{index}"),
                    "artifactType": artifact_type,
                    "title": title,
                    "fileName": str(item.get("fileName") or Path(path).name),
                    "filePath": path,
                    "sourceMode": source_mode,
                    "confirmed": str(item.get("assetReviewStatus") or item.get("reviewStatus") or item.get("status") or "") == "approved",
                    "reviewStatus": str(item.get("assetReviewStatus") or item.get("reviewStatus") or item.get("status") or "pending_review"),
                }
            )
    return artifacts


def collect_material_templates(manifest: dict[str, Any]) -> list[dict[str, Any]]:
    indexed = [
        template_artifact_from_index(item, index)
        for index, item in enumerate(manifest.get("templateIndex") or [], start=1)
        if isinstance(item, dict)
    ]
    indexed = [item for item in indexed if item]
    root = as_path(manifest.get("materialLibraryDir"), required=False)
    if not root or not root.exists():
        return indexed
    templates: list[dict[str, Any]] = list(indexed)
    for path in sorted(root.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in DOCX_SUFFIXES:
            continue
        text = normalize_text(str(path.relative_to(root)))
        # Keep this narrow: template/base-draft files are preferred for blank response documents.
        if not any(token in text for token in ["模板", "模版", "底稿", "空白", "格式", "商务响应文件"]):
            continue
        templates.append(
            {
                "artifactId": f"MAT-TEMPLATE-{len(templates) + 1:04d}",
                "artifactType": "material_template",
                "title": path.stem,
                "fileName": path.name,
                "filePath": str(path),
                "sourceMode": "business_material_template",
                "confirmed": True,
                "reviewStatus": "approved",
            }
        )
    return templates


def template_artifact_from_index(template: dict[str, Any], index: int) -> dict[str, Any]:
    path = Path(str(template.get("filePath") or "")).expanduser()
    if not path.exists() or path.suffix.lower() not in DOCX_SUFFIXES:
        return {}
    return {
        "artifactId": str(template.get("templateId") or f"BID-TEMPLATE-{index:04d}"),
        "artifactType": "bid_template",
        "title": str(template.get("templateName") or path.stem),
        "fileName": str(template.get("fileName") or path.name),
        "filePath": str(path),
        "sourceMode": str(template.get("sourceMode") or "bid_template"),
        "confirmed": True,
        "reviewStatus": "approved",
        "assemblyMode": "template_fill_docx",
        "materialUsage": "fill_template",
    }


def best_material_template_for_task(task: dict[str, Any], templates: list[dict[str, Any]]) -> dict[str, Any] | None:
    scored = [(artifact_match_score(task, item), item) for item in templates]
    scored = [(score, item) for score, item in scored if score >= 6]
    if not scored:
        return None
    scored.sort(key=lambda pair: pair[0], reverse=True)
    return scored[0][1]


def best_parse_artifacts_for_task(task: dict[str, Any], artifacts: list[dict[str, Any]]) -> list[dict[str, Any]]:
    scored = [(artifact_match_score(task, item), item) for item in artifacts]
    scored = [(score, item) for score, item in scored if score >= 5]
    scored.sort(key=lambda pair: pair[0], reverse=True)
    if not scored:
        return []
    best_score = scored[0][0]
    return [item for score, item in scored if score == best_score][:2]


def artifact_match_score(task: dict[str, Any], artifact: dict[str, Any]) -> int:
    title = normalize_text(str(task.get("title") or ""))
    task_key = normalize_text(str(task.get("taskKey") or ""))
    artifact_text = normalize_text(
        " ".join(
            [
                str(artifact.get("title") or ""),
                str(artifact.get("fileName") or ""),
                str(artifact.get("filePath") or ""),
                str(artifact.get("sourceMode") or ""),
            ]
        )
    )
    if not title or not artifact_text:
        return 0
    score = 0
    if title in artifact_text or artifact_text in title:
        score += 20
    for token in business_tokens(title):
        if token in artifact_text:
            score += max(2, len(token))
    for token in business_tokens(task_key):
        if token in artifact_text:
            score += 2
    if "承诺" in title and "承诺" in artifact_text:
        score += 6
    if "授权" in title and "授权" in artifact_text:
        score += 6
    if "投标函" in title and "投标函" in artifact_text:
        score += 8
    if "廉洁" in title and "廉洁" in artifact_text:
        score += 8
    return score


def business_tokens(value: str) -> list[str]:
    text = normalize_text(value)
    tokens = [
        "投标函",
        "授权",
        "廉洁",
        "承诺",
        "保证金",
        "保函",
        "履约",
        "报价",
        "价格",
        "规格",
        "偏差",
        "评分",
        "营业执照",
        "资质",
        "证书",
        "业绩",
        "协议",
        "其他说明",
    ]
    return [token for token in tokens if token in text]


def artifact_key(artifact: dict[str, Any]) -> str:
    return str(artifact.get("artifactId") or artifact.get("materialId") or artifact.get("filePath") or artifact.get("path") or artifact.get("fileName") or "")


def artifact_dedupe_key(artifact: dict[str, Any]) -> str:
    path = str(artifact.get("filePath") or artifact.get("path") or artifact.get("docxPath") or "").strip()
    if path:
        return f"path:{Path(path).expanduser()}"
    artifact_id = str(artifact.get("artifactId") or "").strip()
    if artifact_id:
        return f"id:{artifact_id}"
    if str(artifact.get("artifactType") or "") == "parse_business_scoring_json":
        return "structured-business-scoring"
    return str(artifact.get("fileName") or artifact.get("title") or "")


def is_business_scoring_artifact(artifact: dict[str, Any]) -> bool:
    artifact_type = str(artifact.get("artifactType") or "")
    source_mode = str(artifact.get("sourceMode") or "")
    text = normalize_text(
        " ".join(
            [
                artifact_type,
                source_mode,
                str(artifact.get("title") or ""),
                str(artifact.get("fileName") or ""),
                str(artifact.get("filePath") or artifact.get("path") or ""),
            ]
        )
    )
    return artifact_type in {"parse_business_scoring", "parse_business_scoring_json"} or any(
        token in text for token in ["商务评分", "评分标准", "评分索引", "评审表", "审查表", "符合性审查"]
    )


def emit_artifact(doc: Document, artifact: dict[str, Any], facts: dict[str, str], ctx: AssemblyContext, *, section_title: str) -> bool:
    key = artifact_key(artifact)
    if key and key in ctx.used_artifact_keys:
        return True
    if key:
        ctx.used_artifact_keys.add(key)
    assembly_mode = normalize_assembly_mode(str(artifact.get("assemblyMode") or ""))
    if str(artifact.get("artifactType") or "") == "parse_business_scoring_json":
        ok = append_business_scoring_table(doc, artifact, ctx, section_title=section_title)
        record_attachment(ctx, artifact, section_title, status="embedded" if ok else "missing", mode="structured_scoring_table")
        return ok
    path = as_path(artifact.get("filePath") or artifact.get("path") or artifact.get("docxPath"), required=False)
    file_name = str(artifact.get("fileName") or (path.name if path else "附件材料"))
    if not path or not path.exists():
        doc.add_paragraph(f"[待补充附件：{file_name}]")
        ctx.add_review("artifact_file_missing", f"附件文件不存在：{file_name}", section=section_title, source=str(path or ""), level="error")
        record_attachment(ctx, artifact, section_title, status="missing", mode="placeholder")
        return False

    suffix = path.suffix.lower()
    if assembly_mode == "template_fill_docx":
        if suffix not in DOCX_SUFFIXES:
            doc.add_paragraph(f"[待补充 Word 模板片段：{file_name}]")
            ctx.add_review("template_not_docx", f"模板填充仅支持 Word 模板，当前文件无法片段提取：{file_name}", section=section_title, source=str(path))
            record_attachment(ctx, artifact, section_title, status="referenced", mode="template_fill_docx_unsupported")
            return False
        ok = append_template_docx_fragment(doc, path, artifact, facts, ctx, section_title=section_title)
        record_attachment(ctx, artifact, section_title, status="embedded" if ok else "referenced", mode="template_fragment_fill")
        return ok

    if assembly_mode == "extract_segment":
        ok = append_evidence_segment_reference(doc, artifact, ctx, section_title=section_title)
        if ok:
            record_attachment(ctx, artifact, section_title, status="referenced", mode="extract_segment_reference")
            return True
        ctx.add_review("segment_location_missing", f"素材需要片段摘取但缺少证据片段定位：{file_name}", section=section_title, source=str(path))
        doc.add_paragraph(f"详见附件：{file_name}")
        ctx.counters["referenced"] += 1
        record_attachment(ctx, artifact, section_title, status="referenced", mode="extract_segment_fallback_reference")
        return False

    if assembly_mode == "attach_whole_file":
        if suffix in DOCX_SUFFIXES:
            ok = append_docx(doc, path, facts, ctx, section_title=section_title)
            record_attachment(ctx, artifact, section_title, status="embedded" if ok else "referenced", mode="docx_merge")
            return ok
        if suffix in IMAGE_SUFFIXES:
            ok = append_image(doc, path, ctx, section_title=section_title)
            record_attachment(ctx, artifact, section_title, status="embedded" if ok else "referenced", mode="whole_image_embed")
            return ok
        if suffix in PDF_SUFFIXES:
            ok = append_pdf(doc, path, ctx, section_title=section_title)
            record_attachment(ctx, artifact, section_title, status="embedded" if ok else "referenced", mode="whole_pdf_pages_embed")
            return ok
        doc.add_paragraph(f"详见附件：{file_name}")
        if artifact.get("evidenceSummary"):
            doc.add_paragraph(str(artifact.get("evidenceSummary") or ""))
        ctx.counters["referenced"] += 1
        record_attachment(ctx, artifact, section_title, status="referenced", mode="whole_file_reference")
        return True

    try:
        if suffix in DOCX_SUFFIXES:
            ok = append_docx(doc, path, facts, ctx, section_title=section_title)
            record_attachment(ctx, artifact, section_title, status="embedded" if ok else "referenced", mode=assembly_mode or "docx_merge")
            return ok
        if suffix in IMAGE_SUFFIXES:
            ok = append_image(doc, path, ctx, section_title=section_title)
            record_attachment(ctx, artifact, section_title, status="embedded" if ok else "referenced", mode=assembly_mode or "image_embed")
            return ok
        if suffix in PDF_SUFFIXES:
            ok = append_pdf(doc, path, ctx, section_title=section_title)
            record_attachment(ctx, artifact, section_title, status="embedded" if ok else "referenced", mode=assembly_mode or "pdf_pages_embed")
            return ok
    except Exception as exc:
        ctx.add_review("artifact_embed_failed", f"材料嵌入失败，已退化为附件引用：{file_name}；原因：{exc}", section=section_title, source=str(path))

    doc.add_paragraph(f"详见附件：{file_name}")
    ctx.counters["referenced"] += 1
    record_attachment(ctx, artifact, section_title, status="referenced", mode="fallback_reference")
    return False


def append_docx(doc: Document, path: Path, facts: dict[str, str], ctx: AssemblyContext, *, section_title: str) -> bool:
    if path.exists() and path.stat().st_size > MAX_TEMPLATE_WHOLE_MERGE_BYTES:
        doc.add_paragraph(f"详见附件：{path.name}")
        ctx.counters["referenced"] += 1
        ctx.add_review(
            "large_docx_reference_only",
            f"Word 材料超过 {MAX_TEMPLATE_WHOLE_MERGE_BYTES // 1024 // 1024}MB，未整份合入正文，已退化为附件引用：{path.name}",
            section=section_title,
            source=str(path),
        )
        return True
    if Composer is not None:
        temp_path: Path | None = None
        try:
            source = Document(str(path))
            replace_fields_in_doc(source, facts, ctx, section_title=section_title, source=str(path))
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                temp_path = Path(tmp.name)
            source.save(str(temp_path))
            composer = Composer(doc)
            composer.append(Document(str(temp_path)))
            ctx.counters["embedded"] += 1
            return True
        except Exception as exc:
            ctx.add_review("docx_merge_failed", f"Word 材料合并失败：{path.name}；原因：{exc}", section=section_title, source=str(path))
        finally:
            if temp_path is not None:
                temp_path.unlink(missing_ok=True)
    else:
        ctx.add_review("docxcompose_unavailable", f"docxcompose 不可用，将按文本方式合入 Word 材料：{path.name}", section=section_title, source=str(path))

    try:
        source = Document(str(path))
        replace_fields_in_doc(source, facts, ctx, section_title=section_title, source=str(path))
        if append_docx_text_content(doc, source):
            ctx.counters["embedded"] += 1
            ctx.add_review("docx_text_fallback", f"Word 材料已按文本方式合入，复杂格式需人工复核：{path.name}", section=section_title, source=str(path))
            return True
    except Exception as exc:
        ctx.add_review("docx_text_copy_failed", f"Word 材料文本合入失败：{path.name}；原因：{exc}", section=section_title, source=str(path))

    doc.add_paragraph(f"详见附件：{path.name}")
    ctx.counters["referenced"] += 1
    return False


def append_template_docx_fragment(
    doc: Document,
    path: Path,
    artifact: dict[str, Any],
    facts: dict[str, str],
    ctx: AssemblyContext,
    *,
    section_title: str,
) -> bool:
    try:
        source = Document(str(path))
    except Exception as exc:
        doc.add_paragraph(f"[待补充 Word 模板片段：{path.name}]")
        ctx.add_review("template_open_failed", f"模板 Word 打开失败：{path.name}；原因：{exc}", section=section_title, source=str(path))
        return False

    fragment = find_template_fragment(source, artifact, section_title)
    if not fragment and is_small_single_template_file(path, artifact, section_title):
        fragment = list(iter_doc_blocks(source))
        ctx.add_review("single_template_file_used", f"未找到内部章节标题，已按单一模板文件填充：{path.name}", section=section_title, source=str(path), level="info")
    if not fragment:
        doc.add_paragraph(f"[待提取模板片段：{section_title}]")
        doc.add_paragraph(f"模板文件：{path.name}")
        ctx.counters["referenced"] += 1
        ctx.add_review(
            "template_fragment_not_found",
            f"未在模板 Word 中定位到当前章节对应片段，未整份合入模板：{section_title}",
            section=section_title,
            source=str(path),
        )
        return False

    written = append_fragment_blocks(doc, fragment, facts, ctx, section_title=section_title, source=str(path))
    if written <= 0:
        doc.add_paragraph(f"[待提取模板片段：{section_title}]")
        ctx.add_review("template_fragment_empty", f"模板片段定位成功但内容为空：{section_title}", section=section_title, source=str(path))
        return False
    ctx.counters["embedded"] += 1
    ctx.add_review("template_fragment_filled", f"已从模板 Word 提取并填充片段：{section_title}", section=section_title, source=str(path), level="info")
    return True


def is_small_single_template_file(path: Path, artifact: dict[str, Any], section_title: str) -> bool:
    if not path.exists() or path.stat().st_size > MAX_TEMPLATE_WHOLE_MERGE_BYTES:
        return False
    terms = [term for term in template_target_terms(artifact, section_title) if len(term) >= 2]
    name_text = normalize_text(" ".join([path.stem, str(artifact.get("title") or ""), str(artifact.get("fileName") or "")]))
    if not name_text:
        return False
    return any(term in name_text or name_text in term for term in terms)


def find_template_fragment(source: Document, artifact: dict[str, Any], section_title: str) -> list[Any]:
    blocks = list(iter_doc_blocks(source))
    if not blocks:
        return []
    targets = template_target_terms(artifact, section_title)
    if not targets:
        return []
    start = find_template_fragment_start(blocks, targets)
    if start < 0:
        return []
    start_level = heading_rank(block_text(blocks[start])) or 9
    end = len(blocks)
    for index in range(start + 1, len(blocks)):
        text = block_text(blocks[index])
        if not text:
            continue
        rank = heading_rank(text)
        if rank and rank <= start_level:
            end = index
            break
    return blocks[start:end]


def iter_doc_blocks(doc: Document) -> list[Any]:
    body = doc.element.body
    blocks: list[Any] = []
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            blocks.append(Paragraph(child, doc))
        elif child.tag.endswith("}tbl"):
            blocks.append(Table(child, doc))
    return blocks


def template_target_terms(artifact: dict[str, Any], section_title: str) -> list[str]:
    values = [
        section_title,
        str(artifact.get("title") or ""),
        str(artifact.get("templateName") or ""),
        str(artifact.get("fileName") or ""),
    ]
    terms: list[str] = []
    for value in values:
        for term in expand_template_terms(value):
            if term and term not in terms:
                terms.append(term)
    return terms


def expand_template_terms(value: str) -> list[str]:
    text = str(value or "").strip()
    normalized = normalize_text(text)
    if not normalized:
        return []
    terms = [normalized]
    without_page = re.sub(r"\d+$", "", normalized).strip()
    if without_page and without_page not in terms:
        terms.append(without_page)
    without_number_prefix = re.sub(r"^(?:附件)?[一二三四五六七八九十百千万\d]+(?:[\.．、]?\d+)*", "", without_page).strip()
    if without_number_prefix and without_number_prefix not in terms:
        terms.append(without_number_prefix)
    bracketless = re.sub(r"格式.*$", "格式", without_number_prefix).strip()
    if bracketless and bracketless not in terms:
        terms.append(bracketless)
    keyword_terms = [
        "投标函",
        "法定代表人身份证明",
        "法定代表人授权",
        "授权委托书",
        "廉洁自律承诺",
        "投标专用章效力说明",
        "商务条款偏差",
        "商务偏差",
        "投标价格表",
        "开标价格表",
        "货物规格",
        "规格一览表",
        "投标保证金",
        "履约保证函",
        "业绩情况表",
        "保密承诺",
        "承诺书",
        "承诺函",
    ]
    for keyword in keyword_terms:
        if normalize_text(keyword) in normalized and normalize_text(keyword) not in terms:
            terms.append(normalize_text(keyword))
    return [term for term in terms if len(term) >= 2]


def find_template_fragment_start(blocks: list[Any], targets: list[str]) -> int:
    scored: list[tuple[int, int]] = []
    for index, block in enumerate(blocks):
        text = normalize_text(block_text(block))
        if not text:
            continue
        score = 0
        for term in targets:
            if term and term in text:
                score = max(score, len(term) + 20)
            elif term and text in term and len(text) >= 4:
                score = max(score, len(text) + 10)
        if score:
            rank = heading_rank(block_text(block))
            if rank:
                score += max(0, 10 - rank)
            scored.append((score, index))
    if not scored:
        return -1
    scored.sort(key=lambda item: (-item[0], item[1]))
    return scored[0][1]


def heading_rank(text: str) -> int:
    clean = str(text or "").strip()
    compact = normalize_text(clean)
    if not compact:
        return 0
    if re.match(r"^(?:第[一二三四五六七八九十百千万\d]+[章节条款部分]|[一二三四五六七八九十]+[、.．]|[0-9]+(?:[.．][0-9]+)*[、.．\\s])", clean):
        return min(max(compact.count(".") + compact.count("．") + 1, 1), 6)
    if len(compact) <= 36 and any(token in compact for token in ["附件", "投标函", "授权", "承诺", "价格表", "规格", "偏差表", "保证金", "业绩情况"]):
        return 3
    return 0


def block_text(block: Any) -> str:
    if isinstance(block, Paragraph):
        return block.text.strip()
    if isinstance(block, Table):
        return "\n".join(cell.text.strip() for row in block.rows for cell in row.cells if cell.text.strip())
    return ""


def append_fragment_blocks(
    doc: Document,
    blocks: list[Any],
    facts: dict[str, str],
    ctx: AssemblyContext,
    *,
    section_title: str,
    source: str,
) -> int:
    written = 0
    for block in blocks:
        if isinstance(block, Paragraph):
            text = replace_field_text(block.text, facts, ctx, section_title=section_title, source=source)
            if not text.strip():
                continue
            doc.add_paragraph(text)
            written += 1
        elif isinstance(block, Table):
            if append_table_block(doc, block, facts, ctx, section_title=section_title, source=source):
                written += 1
    return written


def append_table_block(doc: Document, source_table: Table, facts: dict[str, str], ctx: AssemblyContext, *, section_title: str, source: str) -> bool:
    rows = source_table.rows
    if not rows:
        return False
    max_cols = max((len(row.cells) for row in rows), default=0)
    if max_cols <= 0:
        return False
    target_table = doc.add_table(rows=len(rows), cols=max_cols)
    target_table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index in range(max_cols):
            value = row.cells[col_index].text if col_index < len(row.cells) else ""
            target_table.rows[row_index].cells[col_index].text = replace_field_text(value, facts, ctx, section_title=section_title, source=source)
    fill_fact_table_cells(target_table, facts, ctx, section_title=section_title, source=source)
    return True


def normalize_assembly_mode(value: str) -> str:
    mode = str(value or "").strip().lower().replace("-", "_")
    aliases = {
        "attach_whole": "attach_whole_file",
        "whole_file": "attach_whole_file",
        "full_attachment": "attach_whole_file",
        "embed_scan": "embed_scan_or_image",
        "image_embed": "embed_scan_or_image",
        "pdf_pages_embed": "embed_scan_or_image",
        "excerpt": "extract_segment",
        "partial_extract": "extract_segment",
        "fill_template": "template_fill_docx",
        "fill_table": "table_fill_from_material",
    }
    return aliases.get(mode, mode)


def append_evidence_segment_reference(doc: Document, artifact: dict[str, Any], ctx: AssemblyContext, *, section_title: str) -> bool:
    segments = artifact.get("selectedEvidenceSegments") if isinstance(artifact.get("selectedEvidenceSegments"), list) else []
    if not segments:
        segment_id = str(artifact.get("evidenceSegmentId") or "")
        if segment_id:
            segments = [
                {
                    "segmentId": segment_id,
                    "title": str(artifact.get("evidenceSegmentTitle") or ""),
                    "sourcePages": str(artifact.get("evidenceSourcePages") or ""),
                    "summary": str(artifact.get("evidenceSummary") or ""),
                }
            ]
    if not segments:
        return False
    doc.add_paragraph("已确认引用以下素材证据片段：")
    for index, segment in enumerate(segments, start=1):
        if not isinstance(segment, dict):
            continue
        title = str(segment.get("title") or artifact.get("materialName") or artifact.get("fileName") or f"证据片段{index}")
        pages = str(segment.get("sourcePages") or "")
        summary = str(segment.get("summary") or "")
        line = f"{index}. {title}"
        if pages:
            line += f"（{pages}）"
        doc.add_paragraph(line)
        if summary:
            doc.add_paragraph(summary)
    source_name = str(artifact.get("fileName") or artifact.get("materialName") or "")
    if source_name:
        doc.add_paragraph(f"原始证明材料详见附件：{source_name}")
    ctx.counters["referenced"] += 1
    ctx.add_review("extract_segment_reference", f"已按 S3 证据片段引用材料，需在共创阶段核对原文位置：{source_name}", section=section_title, source=str(artifact.get("filePath") or artifact.get("path") or ""))
    return True


def append_docx_text_content(doc: Document, source: Document) -> bool:
    written = 0
    for paragraph in source.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        doc.add_paragraph(text)
        written += 1
    for table in source.tables:
        rows = table.rows
        if not rows:
            continue
        max_cols = max((len(row.cells) for row in rows), default=0)
        if max_cols <= 0:
            continue
        target_table = doc.add_table(rows=len(rows), cols=max_cols)
        target_table.style = "Table Grid"
        for row_index, row in enumerate(rows):
            for col_index in range(max_cols):
                value = row.cells[col_index].text if col_index < len(row.cells) else ""
                target_table.rows[row_index].cells[col_index].text = value
        written += 1
    return written > 0


def replace_fields_in_doc(doc: Document, facts: dict[str, str], ctx: AssemblyContext, *, section_title: str, source: str) -> None:
    for paragraph in doc.paragraphs:
        replace_fields_in_paragraph(paragraph, facts, ctx, section_title=section_title, source=source)
    for table in doc.tables:
        fill_fact_table_cells(table, facts, ctx, section_title=section_title, source=source)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_fields_in_paragraph(paragraph, facts, ctx, section_title=section_title, source=source)


def fill_fact_table_cells(table: Any, facts: dict[str, str], ctx: AssemblyContext, *, section_title: str, source: str) -> None:
    for row in table.rows:
        cells = list(row.cells)
        for index, cell in enumerate(cells[:-1]):
            label = infer_fact_label_from_cell(cell.text)
            if not label:
                continue
            target = cells[index + 1]
            current = normalize_cell_text(target.text)
            if current and not is_blank_template_cell(current):
                continue
            value = fact_value(facts, label)
            if not value:
                ctx.field_events.append({"field": label, "value": "", "status": "placeholder", "section": section_title, "source": source})
                ctx.add_review("field_placeholder", f"模板表格字段未填写：{label}", section=section_title, source=source)
                continue
            set_cell_text(target, value)
            ctx.field_events.append({"field": label, "value": value, "status": "filled_table_cell", "section": section_title, "source": source})


def infer_fact_label_from_cell(text: str) -> str:
    raw = re.sub(r"\s+", "", str(text or ""))
    raw = raw.strip("：:，,。；;（）()【】[]")
    if not raw or len(raw) > 24:
        return ""
    candidates = [
        "项目名称",
        "招标项目名称",
        "工程名称",
        "招标编号",
        "项目编号",
        "招标项目编号",
        "标段编号",
        "招标人",
        "采购人",
        "项目单位",
        "业主",
        "投标人",
        "投标单位",
        "供应商",
        "法定代表人",
        "委托代理人",
        "授权代表",
        "投标报价",
        "日期",
        "投标日期",
        "签署日期",
    ]
    for label in candidates:
        if label in raw:
            return label
    return ""


def normalize_cell_text(text: str) -> str:
    return re.sub(r"\s+", "", str(text or "")).strip()


def is_blank_template_cell(text: str) -> bool:
    clean = normalize_cell_text(text)
    if not clean:
        return True
    return bool(re.fullmatch(r"[/\\_\-—─·.。]*|（\s*）|\(\s*\)|待填|待填写|填写", clean))


def set_cell_text(cell: Any, value: str) -> None:
    if cell.paragraphs:
        cell.paragraphs[0].clear()
        cell.paragraphs[0].add_run(value)
        for paragraph in cell.paragraphs[1:]:
            paragraph.clear()
    else:
        cell.text = value


def replace_fields_in_paragraph(paragraph: Any, facts: dict[str, str], ctx: AssemblyContext, *, section_title: str, source: str) -> None:
    text = paragraph.text
    if not text or ("{{" not in text and "[[" not in text):
        return
    new_text = replace_field_text(text, facts, ctx, section_title=section_title, source=source)
    if new_text != text:
        paragraph.clear()
        paragraph.add_run(new_text)


def replace_field_text(text: str, facts: dict[str, str], ctx: AssemblyContext, *, section_title: str, source: str) -> str:
    if not text or ("{{" not in text and "[[" not in text):
        return text

    def repl(match: re.Match[str]) -> str:
        label = str(match.group(1) or match.group(2) or "").strip()
        value = fact_value(facts, label)
        if value:
            ctx.field_events.append({"field": label, "value": value, "status": "filled", "section": section_title, "source": source})
            return value
        placeholder = f"[待填写：{label}]"
        ctx.field_events.append({"field": label, "value": "", "status": "placeholder", "section": section_title, "source": source})
        ctx.add_review("field_placeholder", f"字段未填写：{label}", section=section_title, source=source)
        return placeholder

    return FIELD_PATTERN.sub(repl, text)


def append_image(doc: Document, path: Path, ctx: AssemblyContext, *, section_title: str) -> bool:
    try:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(path), width=Inches(6.2))
        caption = doc.add_paragraph(path.name)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ctx.counters["embedded"] += 1
        return True
    except Exception as exc:
        ctx.add_review("image_embed_failed", f"图片嵌入失败：{path.name}；原因：{exc}", section=section_title, source=str(path))
        return False


def append_pdf(doc: Document, path: Path, ctx: AssemblyContext, *, section_title: str) -> bool:
    if fitz is None:
        ctx.add_review("pdf_renderer_unavailable", f"PyMuPDF 不可用，无法将 PDF 转图片嵌入：{path.name}", section=section_title, source=str(path))
        return False
    try:
        pdf = fitz.open(str(path))
    except Exception as exc:
        ctx.add_review("pdf_open_failed", f"PDF 打开失败：{path.name}；原因：{exc}", section=section_title, source=str(path))
        return False
    embedded = 0
    try:
        for page_index in range(len(pdf)):
            page = pdf[page_index]
            pix = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
            image_path = ctx.assets_dir / f"{safe_filename(path.stem, 'pdf')}-p{page_index + 1}.png"
            pix.save(str(image_path))
            if append_image(doc, image_path, ctx, section_title=section_title):
                embedded += 1
        return embedded > 0
    finally:
        pdf.close()


def emit_generated_draft(doc: Document, task: dict[str, Any], facts: dict[str, str], ctx: AssemblyContext, *, section_title: str) -> bool:
    title = str(task.get("title") or section_title)
    normalized = normalize_text(title)
    fields = ["项目名称", "招标编号", "招标人", "投标人", "日期"]
    if any(token in normalized for token in ["投标函", "授权", "廉洁", "承诺", "声明", "说明", "保证金", "保函", "履约"]):
        doc.add_paragraph(f"{title}")
        for field in fields:
            value = fact_value(facts, field)
            doc.add_paragraph(f"{field}：{value or f'[待填写：{field}]'}")
            if not value:
                ctx.add_review("field_placeholder", f"基础稿字段未填写：{field}", section=section_title)
        doc.add_paragraph("本文件为系统根据商务目录生成的基础稿，需人工核对并完成签字、盖章。")
        ctx.add_review("generated_basic_draft", f"已生成基础稿，需人工复核：{title}", section=section_title)
        return True
    return False


def append_business_scoring_table(doc: Document, artifact: dict[str, Any], ctx: AssemblyContext, *, section_title: str) -> bool:
    groups = artifact.get("scoringCriteria") if isinstance(artifact.get("scoringCriteria"), dict) else {}
    if not groups:
        ctx.add_review("business_scoring_json_empty", "商务评分标准结构化数据为空，无法生成评分表。", section=section_title)
        return False

    group_titles = {
        "business": "商务评分标准",
        "price": "投标报价评分标准",
        "compliance": "符合性审查标准",
    }
    written = 0
    ordered_keys = [key for key in ("business", "price", "compliance") if isinstance(groups.get(key), list) and groups.get(key)]
    ordered_keys.extend(key for key, rows in groups.items() if key not in ordered_keys and isinstance(rows, list) and rows)
    for key in ordered_keys:
        rows = groups.get(key)
        if not isinstance(rows, list) or not rows:
            continue
        doc.add_heading(group_titles.get(key, str(key)), level=2)
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = ["序号", "评分/审查项", "分值", "得分点/要求", "证明材料要求", "证据位置"]
        for index, header in enumerate(headers):
            table.cell(0, index).text = header
        for row_index, row in enumerate(rows, start=1):
            item = row if isinstance(row, dict) else {}
            cells = table.add_row().cells
            cells[0].text = str(item.get("order") or row_index)
            cells[1].text = str(item.get("scoringItem") or item.get("title") or "")
            cells[2].text = str(item.get("score") or "")
            cells[3].text = str(item.get("scorePoint") or item.get("requirement") or item.get("evidence") or "")
            cells[4].text = str(item.get("proofRequirement") or item.get("proof") or "")
            cells[5].text = str(item.get("evidenceLocation") or item.get("source") or "")
            written += 1
    if written <= 0:
        ctx.add_review("business_scoring_json_empty", "商务评分标准结构化数据无有效行，无法生成评分表。", section=section_title)
        return False
    ctx.counters["embedded"] += 1
    ctx.add_review("business_scoring_from_json", "商务评分标准由解析结构化数据生成，需人工复核格式与分值。", section=section_title)
    return True


def collect_scoring_artifacts(parse_result: dict[str, Any], tasks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    artifacts: list[dict[str, Any]] = []
    structured = parse_result.get("structured") if isinstance(parse_result.get("structured"), dict) else {}
    scoring = structured.get("businessScoringAsset") if isinstance(structured.get("businessScoringAsset"), dict) else {}
    if scoring:
        path = str(scoring.get("docxPath") or scoring.get("filePath") or scoring.get("path") or "")
        if path:
            artifacts.append(
                {
                    "artifactId": str(scoring.get("id") or "BIZ-SCORING-ASSET"),
                    "artifactType": "parse_business_scoring",
                    "fileName": str(scoring.get("fileName") or Path(path).name or "商务评分标准.docx"),
                    "filePath": path,
                    "sourceMode": "parsed_business_scoring",
                    "confirmed": str(scoring.get("reviewStatus") or scoring.get("status") or "") == "approved",
                }
            )
    scoring_groups = structured.get("scoringCriteria") if isinstance(structured.get("scoringCriteria"), dict) else {}
    if scoring_groups and not any(str(item.get("artifactType") or "") == "parse_business_scoring" for item in artifacts):
        artifacts.append(
            {
                "artifactId": "BIZ-SCORING-STRUCTURED",
                "artifactType": "parse_business_scoring_json",
                "fileName": "商务评分标准-结构化数据",
                "sourceMode": "parsed_business_scoring_json",
                "confirmed": False,
                "reviewStatus": str(scoring.get("reviewStatus") or scoring.get("status") or "pending_review"),
                "scoringCriteria": scoring_groups,
            }
        )
    for task in tasks:
        if not is_scoring_section(str(task.get("title") or "")):
            continue
        for artifact in task.get("resolvedArtifacts") or []:
            if isinstance(artifact, dict) and str(artifact.get("filePath") or artifact.get("path") or ""):
                artifacts.append(artifact)
    unique: dict[str, dict[str, Any]] = {}
    for artifact in artifacts:
        unique[artifact_dedupe_key(artifact)] = artifact
    return list(unique.values())


def is_scoring_section(title: str) -> bool:
    text = normalize_text(title)
    return any(token in text for token in ["商务评分", "评分标准", "评分索引", "评分表"])


def has_unconfirmed_gap_tasks(tasks: list[dict[str, Any]]) -> bool:
    return any(str(task.get("status") or "") not in {"ready", "resolved", "ignored"} for task in tasks)


def record_attachment(ctx: AssemblyContext, artifact: dict[str, Any], section_title: str, *, status: str, mode: str) -> None:
    ctx.attachment_items.append(
        {
            "artifactId": str(artifact.get("artifactId") or ""),
            "materialId": str(artifact.get("materialId") or ""),
            "fileName": str(artifact.get("fileName") or Path(str(artifact.get("filePath") or artifact.get("path") or "")).name),
            "filePath": str(artifact.get("filePath") or artifact.get("path") or artifact.get("docxPath") or ""),
            "sectionTitle": section_title,
            "status": status,
            "mode": mode,
            "sourceMode": str(artifact.get("sourceMode") or ""),
            "assemblyMode": str(artifact.get("assemblyMode") or ""),
            "materialUsage": str(artifact.get("materialUsage") or ""),
            "evidenceSegmentId": str(artifact.get("evidenceSegmentId") or ""),
        }
    )


def write_plan(ctx: AssemblyContext, summary: dict[str, Any]) -> None:
    payload = {
        "schemaVersion": PLAN_SCHEMA_VERSION,
        "projectId": ctx.project_id,
        "projectName": ctx.project_name,
        "generatedAt": now_iso(),
        "sections": ctx.section_results,
        "summary": summary,
    }
    ctx.plan_file.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def write_reports(ctx: AssemblyContext, summary: dict[str, Any]) -> None:
    report_lines = [
        "# 商务标 S4 装配报告",
        "",
        f"- 项目：{ctx.project_name}",
        f"- 输出文件：{ctx.output_file}",
        f"- 章节数：{summary['sectionCount']}",
        f"- 已装配项：{summary['assembledCount']}",
        f"- 占位项：{summary['placeholderCount']}",
        f"- 待复核项：{summary['reviewRequiredCount']}",
        f"- 嵌入附件数：{summary['embeddedAttachmentCount']}",
        f"- 引用附件数：{summary['referencedAttachmentCount']}",
        "",
        "## 章节结果",
    ]
    for section in ctx.section_results:
        report_lines.append(f"- {section.get('number') or ''} {section.get('title')}: {section.get('status')}")
    ctx.report_file.write_text("\n".join(report_lines).strip() + "\n", encoding="utf-8")

    review_lines = ["# 商务标 S4 待复核清单", ""]
    if not ctx.review_items:
        review_lines.append("暂无待复核项。")
    else:
        for index, item in enumerate(ctx.review_items, start=1):
            review_lines.append(f"{index}. [{item['level']}] {item['message']}")
            if item.get("section"):
                review_lines.append(f"   - 章节：{item['section']}")
            if item.get("source"):
                review_lines.append(f"   - 来源：{item['source']}")
    ctx.review_file.write_text("\n".join(review_lines).strip() + "\n", encoding="utf-8")

    ctx.attachment_manifest_file.write_text(
        json.dumps({"schemaVersion": "bid-business-attachment-manifest-v1", "items": ctx.attachment_items}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    ctx.field_fill_report_file.write_text(
        json.dumps({"schemaVersion": "bid-business-field-fill-report-v1", "items": ctx.field_events}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def summary_payload(output: dict[str, Any]) -> dict[str, Any]:
    return {
        "schema_version": output["schema_version"],
        "outputFile": output["outputFile"],
        "assemblyReport": output["assemblyReport"],
        "needsReview": output["needsReview"],
        "planFile": output["planFile"],
        "attachmentManifest": output["attachmentManifest"],
        "fieldFillReport": output["fieldFillReport"],
        "summary": output["summary"],
    }


def normalize_text(value: str) -> str:
    return re.sub(r"[\s:：，,。.！!？?\-_/'\"“”‘’·（）()【】\[\]/]+", "", str(value or ""))


def safe_filename(value: str, fallback: str) -> str:
    text = re.sub(r"[\\/:*?\"<>|]+", "-", str(value or "").strip())
    text = re.sub(r"\s+", " ", text).strip(" .")
    return text or fallback


def now_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")


if __name__ == "__main__":
    main()
