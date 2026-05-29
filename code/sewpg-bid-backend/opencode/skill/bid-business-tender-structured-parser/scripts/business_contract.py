from __future__ import annotations

import copy
import re
import sys
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any


CURRENT = Path(__file__).resolve()
TECHNICAL_PARSER_DIR = CURRENT.parents[2] / "bid-tech-tender-structured-parser" / "scripts"
if str(TECHNICAL_PARSER_DIR) not in sys.path:
    sys.path.insert(0, str(TECHNICAL_PARSER_DIR))

from parser_core import (  # type: ignore[import-not-found]
    _iter_docx_blocks,
    extract_docx_text,
    parse_manifest as parse_technical_manifest,
)


SKILL_NAME = "bid-business-tender-structured-parser"
SCHEMA_VERSION = "bid-business-tender-structured-v1"
MARKDOWN_TABLE_LINE_PATTERN = re.compile(r"^\s*\|.*\|\s*$")
LEADING_NUMBER_PATTERN = re.compile(r"^\s*(?:第?[一二三四五六七八九十百千0-9]+[章节条]?|[（(]?\d+[）)]?)\s*[、.．\s]+")
LABEL_VALUE_PATTERN = re.compile(r"^\s*(?P<label>[^:：]{2,80})\s*[:：]\s*(?P<value>.+?)\s*$")


@dataclass(frozen=True)
class FieldSpec:
    key: str
    label: str
    aliases: tuple[str, ...]


PROJECT_BASIC_FIELDS: tuple[FieldSpec, ...] = (
    FieldSpec("projectName", "项目名称", ("项目名称", "招标项目名称")),
    FieldSpec("tenderNo", "招标编号", ("招标编号", "项目编号", "招标文件编号")),
    FieldSpec("tenderer", "招标人", ("招标人", "业主", "建设单位", "项目单位")),
    FieldSpec("tenderAgency", "招标代理机构", ("招标代理机构", "代理机构")),
    FieldSpec("bidDeadline", "递交截止时间", ("递交截止时间", "投标截止时间", "投标文件递交截止时间", "开标时间")),
)

BUSINESS_RESPONSE_FIELDS: tuple[FieldSpec, ...] = (
    FieldSpec("bidLetterRequired", "投标函要求", ("投标函",)),
    FieldSpec("authorizationLetterRequired", "授权委托书要求", ("授权委托书", "法定代表人授权委托书")),
    FieldSpec("integrityCommitmentRequired", "廉洁承诺要求", ("廉洁", "廉洁自律承诺", "廉洁承诺")),
    FieldSpec("sealValidityStatementRequired", "投标专用章效力说明要求", ("投标专用章效力说明",)),
    FieldSpec("bidPriceTableRequired", "投标价格表要求", ("投标价格", "投标价格表")),
    FieldSpec("openingPriceTableRequired", "开标价格表要求", ("开标价格表",)),
    FieldSpec("specificationTableRequired", "货物规格表要求", ("货物规格", "规格表")),
    FieldSpec("commercialDeviationTableRequired", "商务偏差表要求", ("商务偏差", "偏差表", "偏离表")),
    FieldSpec("supplyScopeTableRequired", "供货范围表要求", ("供货范围", "供货范围表")),
    FieldSpec("bidSecurityRequired", "投标保证金要求", ("投标保证金", "保证金", "保函")),
    FieldSpec("performanceBondCommitmentRequired", "履约保证承诺要求", ("履约保证函", "履约承诺", "履约保证")),
    FieldSpec("attachment9Required", "附件9要求", ("附件9", "附件九")),
)

QUALIFICATION_SUPPORT_FIELDS: tuple[FieldSpec, ...] = (
    FieldSpec("qualificationDocumentRequired", "资格证明文件要求", ("资格证明", "合格投标人", "资格审查")),
    FieldSpec("performanceDocumentRequired", "业绩证明文件要求", ("业绩证明", "合同扫描件", "中标通知书", "验收报告", "试运行")),
    FieldSpec("financialDocumentRequired", "财务文件要求", ("审计报告", "财务报表", "财务状况")),
    FieldSpec("creditDocumentRequired", "资信诚信文件要求", ("资信证明", "信用中国", "纳税信用", "失信", "经营异常")),
    FieldSpec("certificationDocumentRequired", "证书文件要求", ("认证证书", "资质证书", "体系认证")),
    FieldSpec("customerSpecificProofRequired", "客户专项证明要求", ("战略协议", "框架协议", "评价信", "优秀供应商证明", "示范应用证明")),
)

COMMITMENT_REQUIREMENT_FIELDS: tuple[FieldSpec, ...] = (
    FieldSpec("generalCommitmentCount", "承诺线索识别数", ("承诺",)),
    FieldSpec("generatedCommitmentCount", "自动生成承诺文件数", ("承诺函", "承诺书", "不得存在下列情形")),
    FieldSpec("pendingCommitmentCount", "待确认承诺线索数", ("承诺",)),
    FieldSpec("disqualificationCommitmentRequired", "不得存在下列情形承诺要求", ("不得存在下列情形", "不得存在下列情形之一")),
    FieldSpec("otherCommitmentSectionRequired", "其他承诺章节要求", ("投标人需要说明的其他内容", "其他内容", "其他承诺")),
    FieldSpec("commitmentGenerationBasis", "承诺文件生成依据", ("承诺函", "承诺书", "不得存在下列情形")),
)

COMMERCIAL_REJECTION_KEYWORDS = (
    "否决",
    "废标",
    "无效投标",
    "不予受理",
    "★",
    "实质性响应",
    "投标人不得存在",
    "不得存在下列情形",
)


def _clean(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "").replace("\u3000", " ")).strip()


def _strip_leading_number(text: str) -> str:
    return LEADING_NUMBER_PATTERN.sub("", text).strip()


def _split_label_value(line: str, fallback_label: str) -> tuple[str, str]:
    normalized = _strip_leading_number(line)
    match = LABEL_VALUE_PATTERN.match(normalized)
    if match:
        label = _strip_leading_number(match.group("label")).strip()
        value = match.group("value").strip(" ；;。)）")
        return label or fallback_label, value or normalized
    return fallback_label, normalized


def _load_texts_by_id(documents: list[dict[str, Any]]) -> dict[str, str]:
    texts_by_id: dict[str, str] = {}
    for document in documents:
        document_id = str(document.get("id") or "")
        text = ""
        text_path_value = str(document.get("textPath") or "")
        source_path_value = str(document.get("sourcePath") or "")
        if text_path_value:
            text_path = Path(text_path_value)
            if text_path.exists() and text_path.is_file():
                text = text_path.read_text(encoding="utf-8", errors="replace")
        if not text and source_path_value:
            source_path = Path(source_path_value)
            if source_path.suffix.lower() in {".md", ".txt"} and source_path.exists() and source_path.is_file():
                text = source_path.read_text(encoding="utf-8", errors="replace")
            elif source_path.suffix.lower() == ".docx" and source_path.exists() and source_path.is_file():
                text = extract_docx_text(source_path)
        texts_by_id[document_id] = text
    return texts_by_id


def _copy_meta_fields(item: dict[str, Any]) -> dict[str, Any]:
    return {
        "sourceFile": str(item.get("sourceFile") or ""),
        "sourceDocumentId": str(item.get("sourceDocumentId") or ""),
        "section": str(item.get("section") or ""),
        "evidence": str(item.get("evidence") or ""),
        "evidenceLocation": str(item.get("evidenceLocation") or ""),
    }


def _business_field_from_item(spec: FieldSpec, item: dict[str, Any], *, value_override: str | None = None) -> dict[str, Any]:
    return {
        "key": spec.key,
        "label": spec.label,
        "value": (value_override if value_override is not None else str(item.get("value") or item.get("keyValue") or "")).strip(),
        "status": "found",
        **_copy_meta_fields(item),
        "confidence": float(item.get("confidence") or 0.86),
    }


def _empty_business_field(spec: FieldSpec, *, value: str = "") -> dict[str, Any]:
    return {
        "key": spec.key,
        "label": spec.label,
        "value": value,
        "status": "missing" if not value else "derived",
        "sourceFile": "",
        "sourceDocumentId": "",
        "section": "",
        "evidence": "",
        "evidenceLocation": "",
        "confidence": 0.0,
    }


def _find_business_item(items: list[dict[str, Any]], spec: FieldSpec) -> dict[str, Any] | None:
    for item in items:
        haystack = " ".join(str(item.get(key) or "") for key in ("title", "keyEntity", "value", "evidence", "section"))
        if any(alias in haystack for alias in spec.aliases):
            return item
    return None


def _is_reference_only_value(value: str) -> bool:
    normalized = re.sub(r"\s+", "", str(value or ""))
    if not normalized:
        return False
    return normalized in {
        "见投标人须知前附表",
        "详见招标公告",
        "详见技术规范书",
        "详见招标文件",
        "按招标文件要求",
    } or normalized.startswith("见投标人须知前附表")


def _block_number(location: str) -> int:
    match = re.match(r"B(\d+)", str(location or ""))
    return int(match.group(1)) if match else 0


def _business_core_field_score(item: dict[str, Any], spec: FieldSpec) -> int:
    value = str(item.get("value") or item.get("keyValue") or "").strip()
    section = str(item.get("section") or "")
    evidence = str(item.get("evidence") or "")
    location = str(item.get("evidenceLocation") or "")
    block_no = _block_number(location)
    score = 0
    if location.startswith("B"):
        score += 40
        if 0 < block_no <= 30:
            score += 25
    if section == "封面":
        score += 90
    if "投标人须知前附表" in section:
        score += 80
    if "招标公告" in section or "联系方式" in section:
        score += 50
    if _is_reference_only_value(value):
        score -= 220
    if re.search(r"\d{4}-\d{2}-\d{2}|20\d{2}年", value):
        score += 30
    if spec.key == "projectName":
        if "项目" in value and len(value) <= 120:
            score += 45
        if value.endswith("招标") or "采购" in value:
            score += 20
    if spec.key == "tenderNo":
        if re.search(r"[A-Z]{2,}.*\d", value):
            score += 70
        if value in {"招标编号", "项目编号", "招标文件编号"}:
            score -= 100
    if spec.key in {"tenderer", "tenderAgency"}:
        if len(value) <= 100:
            score += 35
        if "：" in value or ":" in value:
            score -= 10
    if len(value) > 180:
        score -= 50
    if not value:
        score -= 300
    if "目录" in evidence:
        score -= 60
    return score


def _normalize_bid_deadline(value: str) -> str:
    match = re.search(r"(20\d{2})\s*(?:年|[-/.])\s*(\d{1,2})\s*(?:月|[-/.])\s*(\d{1,2})\s*日?", str(value or ""))
    if not match:
        return str(value or "").strip()
    try:
        return date(int(match.group(1)), int(match.group(2)), int(match.group(3))).isoformat()
    except ValueError:
        return str(value or "").strip()


def _strip_core_party_contact_tail(value: str) -> str:
    text = _clean(value).strip(" ：:")
    tail_match = re.search(
        r"\s*(?:地\s*址|联\s*系\s*人|电\s*话|电子\s*邮\s*件|邮\s*箱|传\s*真|网\s*址)\s*[：:]",
        text,
    )
    if tail_match:
        text = text[: tail_match.start()]
    return text.strip(" ：:")


def _normalize_core_field_value(spec: FieldSpec, value: str) -> str:
    cleaned = _clean(value)
    if spec.key == "bidDeadline":
        return _normalize_bid_deadline(cleaned)
    stripped = cleaned
    for alias in sorted(spec.aliases, key=len, reverse=True):
        alias_pattern = r"\s*".join(re.escape(char) for char in alias)
        stripped_candidate = re.sub(rf"^\s*{alias_pattern}\s*[：:]\s*", "", cleaned, count=1)
        if stripped_candidate != cleaned:
            stripped = stripped_candidate.strip()
            break
    if spec.key in {"tenderer", "tenderAgency"}:
        return _strip_core_party_contact_tail(stripped)
    return stripped


def _build_business_project_basics(
    items: list[dict[str, Any]],
    project_dates: dict[str, Any] | None = None,
) -> list[dict[str, Any]]:
    project_dates = project_dates or {}
    fields: list[dict[str, Any]] = []
    for spec in PROJECT_BASIC_FIELDS:
        if spec.key == "bidDeadline":
            value = _normalize_bid_deadline(str(project_dates.get("endDate") or ""))
            if value:
                field = _empty_business_field(spec, value=value)
                field["status"] = "found"
                field["confidence"] = 0.78
                fields.append(field)
                continue
        candidates = [
            item
            for item in items
            if any(
                alias in " ".join(str(item.get(key) or "") for key in ("title", "keyEntity", "evidence", "section"))
                for alias in spec.aliases
            )
        ]
        matched = max(candidates, key=lambda item: _business_core_field_score(item, spec)) if candidates else None
        if matched and _business_core_field_score(matched, spec) > -50:
            value_override = _normalize_core_field_value(spec, str(matched.get("value") or matched.get("keyValue") or ""))
            fields.append(_business_field_from_item(spec, matched, value_override=value_override))
        else:
            fields.append(_empty_business_field(spec))
    return fields


def _build_business_response_fields(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [
        _business_field_from_item(spec, matched) if (matched := _find_business_item(items, spec)) else _empty_business_field(spec)
        for spec in BUSINESS_RESPONSE_FIELDS
    ]


def _build_qualification_support_fields(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [
        _business_field_from_item(spec, matched) if (matched := _find_business_item(items, spec)) else _empty_business_field(spec)
        for spec in QUALIFICATION_SUPPORT_FIELDS
    ]


def _new_docx_candidate_item(
    items: list[dict[str, Any]],
    *,
    document: dict[str, Any],
    label: str,
    value: str,
    section: str,
    evidence: str,
    location: str,
    confidence: float = 0.86,
) -> None:
    cleaned_label = _clean(label)
    cleaned_value = _clean(value)
    if not cleaned_label or not cleaned_value:
        return
    items.append(
        {
            "id": f"DOCX-{len(items) + 1:04d}",
            "type": "商务核心字段候选",
            "category": "business_core_candidate",
            "title": cleaned_label,
            "keyEntity": cleaned_label,
            "keyValue": cleaned_value,
            "value": cleaned_value,
            "sourceFile": str(document.get("name") or document.get("id") or "招标文件"),
            "sourceDocumentId": str(document.get("id") or ""),
            "section": section,
            "evidence": evidence,
            "evidenceLocation": location,
            "confidence": confidence,
        }
    )


def _is_docx_source(path: Path) -> bool:
    return path.is_file() and path.suffix.lower() == ".docx"


def _extract_docx_core_candidate_items(documents: list[dict[str, Any]]) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    for document in documents:
        source_path = Path(str(document.get("sourcePath") or ""))
        if not _is_docx_source(source_path):
            continue
        blocks = _iter_docx_blocks(source_path)
        current_section = "封面"
        for block_index, block in enumerate(blocks, start=1):
            if block.get("type") == "paragraph":
                text = _clean(block.get("text"))
                if text:
                    current_section = text
                continue
            if block.get("type") != "table":
                continue
            rows = block.get("rows") or []
            section = current_section or ("封面" if block_index <= 30 else "")
            for row_index, row in enumerate(rows, start=1):
                cells = [_clean(cell) for cell in row if _clean(cell)]
                if len(cells) < 2:
                    continue
                label = cells[0].strip(" ：:")
                value = cells[-1].strip(" ：:")
                evidence = " | ".join(cells)
                _new_docx_candidate_item(
                    items,
                    document=document,
                    label=label,
                    value=value,
                    section=section,
                    evidence=evidence,
                    location=f"B{block_index}/R{row_index}",
                    confidence=0.9 if block_index <= 30 or "投标人须知前附表" in section else 0.82,
                )
    return items


def _looks_like_bidder_instruction_table(title: str, rows: list[list[str]]) -> bool:
    title_text = _clean(title)
    header_text = "".join("".join(_clean(cell) for cell in row) for row in rows[:2])
    return "投标人须知前附表" in title_text and all(token in header_text for token in ("条款", "编列"))


def _parse_bidder_instruction_rows(
    rows: list[list[str]],
    *,
    document: dict[str, Any],
    section: str,
    block_index: int,
) -> list[dict[str, Any]]:
    cleaned_rows = [[_clean(cell) for cell in row] for row in rows if any(_clean(cell) for cell in row)]
    if len(cleaned_rows) <= 1:
        return []
    header = cleaned_rows[0]
    data_rows = cleaned_rows[1:]
    parsed: list[dict[str, Any]] = []
    for row_index, row in enumerate(data_rows, start=2):
        if len(row) < 3:
            continue
        clause_no = row[0]
        clause_name = row[1]
        content = " ".join(cell for cell in row[2:] if cell).strip()
        if not clause_no and not clause_name and not content:
            continue
        parsed.append(
            {
                "id": f"BIDDER-INST-{len(parsed) + 1:04d}",
                "clauseNo": clause_no,
                "clauseName": clause_name,
                "content": content,
                "sourceFile": str(document.get("name") or ""),
                "sourceDocumentId": str(document.get("id") or ""),
                "section": section,
                "evidence": "；".join(
                    f"{header[index]}：{cell}" if index < len(header) and header[index] else cell
                    for index, cell in enumerate(row)
                    if cell
                ),
                "evidenceLocation": f"B{block_index}/R{row_index}",
                "confidence": 0.9,
            }
        )
    return parsed


def _extract_bidder_instruction_rows(documents: list[dict[str, Any]]) -> list[dict[str, Any]]:
    rows_out: list[dict[str, Any]] = []
    for document in documents:
        source_path = Path(str(document.get("sourcePath") or ""))
        if not _is_docx_source(source_path):
            continue
        blocks = _iter_docx_blocks(source_path)
        current_section = ""
        for block_index, block in enumerate(blocks, start=1):
            if block.get("type") == "paragraph":
                text = _clean(block.get("text"))
                if text:
                    current_section = text
                continue
            if block.get("type") != "table":
                continue
            rows = block.get("rows") or []
            if _looks_like_bidder_instruction_table(current_section, rows):
                rows_out.extend(
                    _parse_bidder_instruction_rows(
                        rows,
                        document=document,
                        section=current_section,
                        block_index=block_index,
                    )
                )
                break
    return rows_out


def _build_qualification_requirements(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    keywords = ("投标人资格要求", "资格要求", "资格能力要求", "投标人资质条件", "合格投标人", "资格审查")
    matched: list[dict[str, Any]] = []
    seen: set[str] = set()
    for item in items:
        text = " ".join(str(item.get(key) or "") for key in ("title", "keyEntity", "value", "evidence", "section"))
        if not any(keyword in text for keyword in keywords):
            continue
        content = str(item.get("value") or item.get("evidence") or "").strip()
        if not content or content in seen:
            continue
        seen.add(content)
        matched.append(
            {
                "id": f"QUAL-{len(matched) + 1:04d}",
                "title": str(item.get("title") or "投标人资格要求"),
                "content": content,
                **_copy_meta_fields(item),
                "confidence": float(item.get("confidence") or 0.78),
            }
        )
    return matched[:12]


def _extract_commercial_rejection_clauses(
    documents: list[dict[str, Any]],
    texts_by_id: dict[str, str],
) -> list[dict[str, Any]]:
    clauses: list[dict[str, Any]] = []
    seen: set[str] = set()
    for document in documents:
        document_id = str(document.get("id") or "")
        source_file = str(document.get("name") or document_id or "招标文件")
        current_section = ""
        for line_number, raw_line in enumerate(str(texts_by_id.get(document_id) or "").splitlines(), start=1):
            line = _clean(raw_line)
            if not line:
                continue
            if _looks_like_section_heading(line):
                current_section = line
            if not any(keyword in line for keyword in COMMERCIAL_REJECTION_KEYWORDS):
                continue
            if line in seen:
                continue
            seen.add(line)
            matched_keywords = [keyword for keyword in COMMERCIAL_REJECTION_KEYWORDS if keyword in line]
            clauses.append(
                {
                    "id": f"REJECT-{len(clauses) + 1:04d}",
                    "title": current_section or "商务废标项",
                    "content": line,
                    "matchedKeywords": matched_keywords,
                    "riskLevel": "high"
                    if any(keyword in matched_keywords for keyword in ("否决", "废标", "无效投标", "不予受理"))
                    else "medium",
                    "sourceFile": source_file,
                    "sourceDocumentId": document_id,
                    "section": current_section,
                    "evidence": line,
                    "evidenceLocation": f"L{line_number}",
                    "confidence": 0.82,
                }
            )
    return clauses


def _line_has_explicit_commitment_obligation(text: str) -> bool:
    normalized = re.sub(r"\s+", "", str(text or ""))
    if not normalized:
        return False
    if "不得存在下列情形" in normalized:
        return True
    if any(hint in normalized for hint in COMMITMENT_GENERATION_HINTS):
        return True
    has_commitment = "承诺" in normalized
    has_obligation = any(token in normalized for token in ("须", "应", "需", "必须", "无条件"))
    has_doc_action = any(token in normalized for token in ("提供", "提交", "出具", "附", "另附", "递交"))
    has_doc_name = any(token in normalized for token in (*COMMITMENT_DOC_KEYWORDS, "书面承诺", "承诺材料"))
    return has_commitment and has_obligation and (has_doc_action or has_doc_name)


def _find_commitment_items(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    matched: list[dict[str, Any]] = []
    seen: set[str] = set()
    for item in items:
        text = " ".join(str(item.get(key) or "") for key in ("title", "value", "evidence", "section"))
        if "承诺" not in text and "不得存在下列情形" not in text:
            continue
        item_id = str(item.get("id") or "")
        if item_id in seen:
            continue
        seen.add(item_id)
        matched.append(item)
    return matched


COMMITMENT_DOC_KEYWORDS = ("承诺函", "承诺书")
COMMITMENT_GENERATION_HINTS = (
    "另附承诺函",
    "另附承诺书",
    "单独提供承诺函",
    "单独提供承诺书",
    "须提供承诺函",
    "须提供承诺书",
    "应提供承诺函",
    "应提供承诺书",
    "应出具承诺函",
    "应出具承诺书",
    "需提供承诺函",
    "需提供承诺书",
    "需提供书面承诺",
    "须出具承诺函",
    "须出具承诺书",
    "须无条件承诺",
)
COMMITMENT_REQUIREMENT_CONTEXT_HINTS = (
    "提供",
    "提交",
    "出具",
    "附",
    "另附",
    "单独",
    "递交",
    "响应",
    "按要求",
    "须",
    "应",
    "需",
    "必须",
)
COMMITMENT_IGNORE_KEYWORDS = (
    "技术承诺",
    "廉洁承诺",
    "廉洁自律承诺",
    "履约承诺",
    "履约保证承诺",
    "投标函",
    "授权委托书",
    "评分标准",
    "评分办法",
    "证明材料",
    "目录",
)
COMMITMENT_NON_REQUIREMENT_TITLE_HINTS = (
    "格式",
    "模板",
    "目录",
    "附件",
    "附录",
    "详见",
    "示例",
    "参考",
)
TECHNICAL_COMMITMENT_KEYWORDS = (
    "等效满负荷小时",
    "满负荷小时",
    "保证年等效",
    "保证小时",
    "发电小时",
    "发电量",
    "上网电量",
    "电量",
    "功率曲线",
    "功率",
    "可利用率",
    "涉网性能",
    "机组",
    "叶轮",
    "轮毂",
    "塔筒",
    "箱变",
    "风机",
    "载荷",
    "噪声",
    "振动",
    "发电性能",
    "性能保证",
    "技术指标",
)
COMMITMENT_TOPIC_TITLE_KEYWORDS: tuple[tuple[str, tuple[str, ...], str], ...] = (
    ("confidentiality", ("保密",), "保密承诺书"),
    ("disqualification", ("不得存在下列情形",), "投标人不存在下列情形之一承诺函"),
    ("certificate_obtainment", ("取得本条", "取得材料", "取得证书", "取得认证", "供货前取得"), "材料取得承诺书"),
    ("delivery", ("交货", "供货周期", "交付"), "交货周期承诺书"),
    ("quality", ("质量", "质保", "售后", "服务"), "质量服务承诺书"),
    ("security", ("投标保证金", "保函", "保证金"), "投标保证金承诺书"),
    ("compliance", ("合规", "守法", "违法", "违规", "信用"), "合规承诺书"),
)
COMMITMENT_TOPIC_KEYWORDS: tuple[tuple[str, tuple[str, ...]], ...] = (
    ("confidentiality", ("保密",)),
    ("compliance", ("合规", "守法", "违法", "违规", "信用")),
    ("security", ("投标保证金", "保函", "保证金")),
    ("disqualification", ("不得存在下列情形",)),
    ("certificate_obtainment", ("取得本条", "取得材料", "取得证书", "取得认证", "供货前取得")),
    ("integrity", ("廉洁",)),
    ("performance_bond", ("履约保证", "履约承诺")),
    ("delivery_commitment", ("交货", "工期", "供货周期")),
    ("quality_commitment", ("质量", "质保", "售后", "服务承诺")),
)


def _commitment_text(item: dict[str, Any]) -> str:
    return " ".join(str(item.get(key) or "") for key in ("title", "value", "evidence", "section"))


def _is_technical_commitment_item(item: dict[str, Any]) -> bool:
    normalized = re.sub(r"\s+", "", _commitment_text(item))
    if not normalized:
        return False
    return any(keyword in normalized for keyword in TECHNICAL_COMMITMENT_KEYWORDS)


def _preferred_commitment_title(item: dict[str, Any], trigger_text: str) -> str:
    raw_title = str(item.get("title") or "").strip()
    raw_value = str(item.get("value") or "").strip()
    raw_evidence = str(item.get("evidence") or "").strip()
    source_text = raw_evidence or raw_value or raw_title or trigger_text
    normalized_source = re.sub(r"\s+", "", source_text)

    for _, keywords, preferred_title in COMMITMENT_TOPIC_TITLE_KEYWORDS:
        if any(keyword in normalized_source for keyword in keywords):
            return preferred_title

    title = raw_title or trigger_text or "承诺文件"
    if title in COMMITMENT_DOC_KEYWORDS or title == "承诺":
        for keyword in COMMITMENT_GENERATION_HINTS + COMMITMENT_DOC_KEYWORDS:
            if keyword in source_text:
                prefix = _normalize_commitment_title_prefix(source_text.split(keyword, 1)[0])
                if prefix:
                    return f"{prefix}{keyword}"
                return keyword

    title = _normalize_commitment_title_prefix(title)
    if not any(keyword in title for keyword in COMMITMENT_DOC_KEYWORDS):
        title = f"{title}承诺书"
    return title


def _normalize_commitment_topic(text: str) -> str:
    normalized = re.sub(r"\s+", "", str(text or ""))
    for topic, keywords in COMMITMENT_TOPIC_KEYWORDS:
        if any(keyword in normalized for keyword in keywords):
            return topic
    if "承诺函" in normalized or "承诺书" in normalized:
        return normalized[:80]
    return normalized[:80] or "commitment"


def _normalize_commitment_title_prefix(text: str) -> str:
    normalized = str(text or "").strip(" ：:，,、。.；;（）()")
    for prefix in ("投标人应出具", "投标人应提供", "投标人须提供", "投标人需提供", "投标人须出具", "投标人另附", "应出具", "应提供", "须提供", "需提供", "须出具", "另附"):
        if normalized.startswith(prefix):
            normalized = normalized[len(prefix):].strip(" ：:，,、。.；;（）()")
            break
    return normalized


def _contains_commitment_requirement_context(text: str) -> bool:
    normalized = re.sub(r"\s+", "", str(text or ""))
    return any(hint in normalized for hint in COMMITMENT_REQUIREMENT_CONTEXT_HINTS)


def _looks_like_bare_commitment_title(item: dict[str, Any]) -> bool:
    evidence = str(item.get("evidence") or item.get("value") or item.get("title") or "").strip()
    normalized = re.sub(r"\s+", "", evidence)
    if not normalized:
        return False
    if not any(keyword in normalized for keyword in COMMITMENT_DOC_KEYWORDS):
        return False
    if _contains_commitment_requirement_context(normalized):
        return False
    if any(token in normalized for token in COMMITMENT_NON_REQUIREMENT_TITLE_HINTS):
        return True
    if _looks_like_section_heading(evidence):
        return True
    return len(normalized) <= 18 and normalized.endswith(COMMITMENT_DOC_KEYWORDS)


def _extract_commitment_trigger_phrase(text: str) -> str:
    normalized = str(text or "").strip()
    if "不得存在下列情形" in normalized:
        return "投标人不得存在下列情形之一"
    if "须无条件承诺" in normalized:
        return "须无条件承诺"
    for hint in COMMITMENT_GENERATION_HINTS:
        if hint in normalized:
            return hint
    for keyword in COMMITMENT_DOC_KEYWORDS:
        if keyword in normalized:
            return keyword
    return "承诺"


def _is_commitment_doc_required(item: dict[str, Any]) -> bool:
    text = _commitment_text(item)
    normalized = re.sub(r"\s+", "", text)
    if "不得存在下列情形" in normalized:
        return True
    if any(hint in normalized for hint in COMMITMENT_GENERATION_HINTS):
        return True
    if _line_has_explicit_commitment_obligation(normalized):
        return True
    if any(keyword in normalized for keyword in COMMITMENT_DOC_KEYWORDS):
        if _looks_like_bare_commitment_title(item):
            return False
        if _contains_commitment_requirement_context(normalized):
            return True
        return False
    if any(keyword in normalized for keyword in COMMITMENT_NON_REQUIREMENT_TITLE_HINTS):
        return False
    if _looks_like_section_heading(str(item.get("evidence") or item.get("title") or "")):
        return False
    if ("投标人" in normalized or "投标方" in normalized) and any(token in normalized for token in ("须承诺", "应承诺", "承诺如下")):
        return True
    return False


def _is_commitment_item_ignored(item: dict[str, Any]) -> bool:
    text = _commitment_text(item)
    normalized = re.sub(r"\s+", "", text)
    if "承诺" not in normalized and "不得存在下列情形" not in normalized:
        return True
    if _is_technical_commitment_item(item):
        return True
    if _is_commitment_doc_required(item):
        return False
    if any(keyword in normalized for keyword in COMMITMENT_IGNORE_KEYWORDS):
        if "不得存在下列情形" not in normalized:
            return True
    if any(token in normalized for token in ("评分", "得分", "分值", "证明材料要求")):
        return True
    return False


def _build_business_commitment_analysis(items: list[dict[str, Any]]) -> dict[str, list[dict[str, Any]]]:
    all_items = _find_commitment_items(items)
    clues: list[dict[str, Any]] = []
    letters: list[dict[str, Any]] = []
    generated_topics: set[str] = set()
    clue_topics: set[tuple[str, str]] = set()

    for item in all_items:
        if _is_commitment_item_ignored(item):
            continue
        text = _commitment_text(item)
        topic = _normalize_commitment_topic(text)
        trigger_text = _extract_commitment_trigger_phrase(text)
        topic_key = topic or "commitment"
        base = {
            **_copy_meta_fields(item),
            "topic": topic,
            "topicKey": topic_key,
            "triggerText": trigger_text,
            "triggerContext": str(item.get("evidence") or item.get("value") or "").strip(),
        }
        if _is_commitment_doc_required(item):
            if topic_key in generated_topics:
                continue
            generated_topics.add(topic_key)
            if topic_key == "disqualification":
                title = "投标人不存在下列情形之一承诺函"
                commitment_type = "disqualification"
            else:
                title = _preferred_commitment_title(item, trigger_text)
                commitment_type = "general_commitment"
            letters.append(
                {
                    "id": f"CL-{len(letters) + 1:04d}",
                    "artifactType": "commitment_letter",
                    "title": title,
                    "commitmentType": commitment_type,
                    "status": "pending_review",
                    **base,
                    "docxPath": "",
                    "workspacePath": "",
                    "placementHint": "投标人需要说明的其他内容",
                    "needsHumanReview": True,
                    "riskFlags": ["template_pending", "legal_wording_review_required"],
                    "previewType": "onlyoffice",
                }
            )
            continue

        normalized = re.sub(r"\s+", "", text)
        if any(keyword in normalized for keyword in COMMITMENT_DOC_KEYWORDS):
            continue

        clue_key = (topic_key, base["triggerContext"])
        if clue_key in clue_topics:
            continue
        clue_topics.add(clue_key)
        clues.append(
            {
                "id": f"CC-{len(clues) + 1:04d}",
                "artifactType": "commitment_clue",
                "title": str(item.get("title") or trigger_text or "承诺线索").strip() or "承诺线索",
                "clueType": "pending_manual_review",
                "status": "needs_review",
                **base,
                "recommendedAction": "暂不自动生成，请人工判断是否需要单独承诺函/承诺书。",
                "riskFlags": ["ambiguous_requirement"],
            }
        )

    return {"letters": letters, "clues": clues}


def _looks_like_section_heading(line: str) -> bool:
    text = str(line or "").strip()
    if not text:
        return False
    return bool(
        re.match(r"^(?:第[一二三四五六七八九十百千0-9]+[章节条]|[（(]?[一二三四五六七八九十0-9]+[）)])", text)
        or text.startswith("附件")
        or text.startswith("附表")
    )


def _scan_business_hint_items(documents: list[dict[str, Any]], texts_by_id: dict[str, str]) -> list[dict[str, Any]]:
    keywords = sorted(
        {
            *[alias for spec in PROJECT_BASIC_FIELDS for alias in spec.aliases],
            *[alias for spec in BUSINESS_RESPONSE_FIELDS for alias in spec.aliases],
            *[alias for spec in QUALIFICATION_SUPPORT_FIELDS for alias in spec.aliases],
            *[alias for spec in COMMITMENT_REQUIREMENT_FIELDS for alias in spec.aliases],
            "投标人不得存在下列情形之一",
            "不得存在下列情形",
            "投标人需要说明的其他内容",
            "书面承诺",
            "无条件承诺",
        },
        key=len,
        reverse=True,
    )
    items: list[dict[str, Any]] = []
    seen: set[tuple[str, str]] = set()
    for document in documents:
        document_id = str(document.get("id") or "")
        source_file = str(document.get("name") or document_id or "招标文件")
        lines = [line.strip() for line in str(texts_by_id.get(document_id) or "").splitlines()]
        current_section = ""
        for line_number, line in enumerate(lines, start=1):
            if not line:
                continue
            if line.startswith("#"):
                current_section = line.strip("# ").strip()
                continue
            if _looks_like_section_heading(line):
                current_section = line
            matched_keyword = next((keyword for keyword in keywords if keyword and keyword in line), "")
            is_commitment_obligation = _line_has_explicit_commitment_obligation(line)
            if not matched_keyword and not is_commitment_obligation:
                continue
            dedupe_key = (document_id, line)
            if dedupe_key in seen:
                continue
            seen.add(dedupe_key)
            label, value = _split_label_value(line, matched_keyword or "承诺要求")
            items.append(
                {
                    "id": f"RAW-{len(items) + 1:04d}",
                    "type": "商务提示",
                    "category": "business_hint",
                    "title": label or matched_keyword or "承诺要求",
                    "keyEntity": matched_keyword or "承诺要求",
                    "keyValue": value,
                    "value": value or line,
                    "sourceFile": source_file,
                    "sourceDocumentId": document_id,
                    "section": current_section,
                    "evidence": line,
                    "evidenceLocation": f"L{line_number}",
                    "confidence": 0.72,
                }
            )
    return items


def _build_commitment_requirement_fields(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    analysis = _build_business_commitment_analysis(items)
    commitment_items = _find_commitment_items(items)
    disqualification_item = next(
        (
            item
            for item in commitment_items
            if "不得存在下列情形" in " ".join(str(item.get(key) or "") for key in ("title", "value", "evidence"))
        ),
        None,
    )
    other_content_item = next(
        (
            item
            for item in items
            if "投标人需要说明的其他内容" in " ".join(str(item.get(key) or "") for key in ("title", "value", "evidence", "section"))
        ),
        None,
    )

    fields: list[dict[str, Any]] = []
    for spec in COMMITMENT_REQUIREMENT_FIELDS:
        if spec.key == "generalCommitmentCount":
            count_text = str(len(analysis["letters"]) + len(analysis["clues"]))
            matched = commitment_items[0] if commitment_items else None
            fields.append(
                _business_field_from_item(spec, matched, value_override=count_text)
                if matched
                else _empty_business_field(spec, value="0")
            )
            continue
        if spec.key == "generatedCommitmentCount":
            matched = analysis["letters"][0] if analysis["letters"] else None
            fields.append(
                _business_field_from_item(spec, matched, value_override=str(len(analysis["letters"])))
                if matched
                else _empty_business_field(spec, value="0")
            )
            continue
        if spec.key == "pendingCommitmentCount":
            matched = analysis["clues"][0] if analysis["clues"] else None
            fields.append(
                _business_field_from_item(spec, matched, value_override=str(len(analysis["clues"])))
                if matched
                else _empty_business_field(spec, value="0")
            )
            continue
        if spec.key == "disqualificationCommitmentRequired":
            fields.append(_business_field_from_item(spec, disqualification_item) if disqualification_item else _empty_business_field(spec))
            continue
        if spec.key == "otherCommitmentSectionRequired":
            fields.append(_business_field_from_item(spec, other_content_item) if other_content_item else _empty_business_field(spec))
            continue
        if spec.key == "commitmentGenerationBasis":
            matched = analysis["letters"][0] if analysis["letters"] else None
            value = "；".join(
                str(item.get("triggerContext") or "").strip()
                for item in analysis["letters"][:3]
                if str(item.get("triggerContext") or "").strip()
            ).strip("；")
            fields.append(
                _business_field_from_item(spec, matched, value_override=value or "未识别到明确的承诺函/承诺书生成依据")
                if matched
                else _empty_business_field(spec, value="未识别到明确的承诺函/承诺书生成依据")
            )
            continue
        fields.append(_empty_business_field(spec))
    return fields


def _unique_sources(records: list[dict[str, Any]]) -> list[dict[str, str]]:
    seen: set[tuple[str, str]] = set()
    sources: list[dict[str, str]] = []
    for record in records:
        source_file = str(record.get("sourceFile") or "")
        source_document_id = str(record.get("sourceDocumentId") or "")
        if not source_file and not source_document_id:
            continue
        key = (source_document_id, source_file)
        if key in seen:
            continue
        seen.add(key)
        sources.append({"sourceDocumentId": source_document_id, "sourceFile": source_file})
    return sources


def _business_presence_from_keywords(items: list[dict[str, Any]], *, keywords: tuple[str, ...]) -> dict[str, Any]:
    matched = [
        item
        for item in items
        if any(keyword in " ".join(str(item.get(key) or "") for key in ("title", "value", "evidence", "section")) for keyword in keywords)
    ]
    if not matched:
        return {
            "status": "missing",
            "summary": "招标文件中暂未识别到明确要求。",
            "evidences": [],
            "sources": [],
        }
    evidences = [{**_copy_meta_fields(item)} for item in matched[:8]]
    summary = "；".join(
        str(item.get("value") or item.get("evidence") or "").strip()
        for item in matched
        if str(item.get("value") or item.get("evidence") or "").strip()
    )
    return {
        "status": "present",
        "summary": summary[:800],
        "evidences": evidences,
        "sources": _unique_sources(evidences),
    }


def _build_business_requirement_presence(items: list[dict[str, Any]]) -> dict[str, Any]:
    return {
        "qualificationDocuments": _business_presence_from_keywords(items, keywords=("资格证明", "合格投标人", "资格审查")),
        "performanceDocuments": _business_presence_from_keywords(items, keywords=("业绩", "合同", "中标通知书", "验收报告", "试运行")),
        "deviationResponse": _business_presence_from_keywords(items, keywords=("商务偏差", "偏差表", "偏离表")),
        "bidSecurity": _business_presence_from_keywords(items, keywords=("投标保证金", "保证金", "保函")),
        "otherCommitments": _business_presence_from_keywords(items, keywords=("承诺", "投标人需要说明的其他内容", "履约保证")),
        "disqualificationClauses": _business_presence_from_keywords(items, keywords=("投标人不得存在下列情形之一", "不得存在下列情形")),
    }


def _filter_business_scoring(scoring: dict[str, Any]) -> dict[str, list[dict[str, Any]]]:
    if not isinstance(scoring, dict):
        return {"business": [], "price": [], "compliance": []}
    return {
        "business": copy.deepcopy(scoring.get("business") or []),
        "price": copy.deepcopy(scoring.get("price") or []),
        "compliance": copy.deepcopy(scoring.get("compliance") or []),
    }


def _build_business_coverage(field_groups: dict[str, Any], scoring: dict[str, list[dict[str, Any]]], presence: dict[str, Any]) -> list[dict[str, Any]]:
    checks = [
        ("商务评分要求", len(scoring.get("business") or []) > 0),
        (
            "报价与价格表",
            any(
                field.get("status") == "found"
                for field in field_groups.get("businessResponse") or []
                if field.get("key") in {"bidPriceTableRequired", "openingPriceTableRequired"}
            ),
        ),
        ("偏差响应", (presence.get("deviationResponse") or {}).get("status") == "present"),
        ("资格证明", (presence.get("qualificationDocuments") or {}).get("status") == "present"),
        ("业绩证明", (presence.get("performanceDocuments") or {}).get("status") == "present"),
        ("保证金", (presence.get("bidSecurity") or {}).get("status") == "present"),
        ("其他承诺", (presence.get("otherCommitments") or {}).get("status") == "present"),
    ]
    return [{"label": label, "status": "covered" if covered else "missing"} for label, covered in checks]


def _build_business_commitment_letters(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return _build_business_commitment_analysis(items)["letters"]



def _build_business_commitment_clues(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return _build_business_commitment_analysis(items)["clues"]


def _sanitize_docx_name(value: str, fallback: str) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|]+", "-", str(value or "").strip())
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" .")
    while len(cleaned.encode("utf-8")) > 120:
        cleaned = cleaned[:-1].strip(" .")
    return cleaned or fallback


def _commitment_letter_output_dir(manifest: dict[str, Any]) -> Path:
    structured_path = Path(str(manifest.get("structuredResultPath") or "")).expanduser().resolve()
    output_dir = structured_path.parent.parent / "s1_commitment_letters"
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir


def _write_commitment_letter_docx(path: Path, letter: dict[str, Any], *, project_name: str = "", tenderer_name: str = "") -> None:
    from docx import Document

    path.parent.mkdir(parents=True, exist_ok=True)
    title = str(letter.get("title") or "承诺函").strip() or "承诺函"
    trigger_context = str(letter.get("triggerContext") or letter.get("evidence") or letter.get("triggerText") or "").strip()
    commitment_type = str(letter.get("commitmentType") or "").strip()

    doc = Document()
    doc.add_heading(title, level=1)
    if project_name:
        doc.add_paragraph(f"项目名称：{project_name}")
    if tenderer_name:
        doc.add_paragraph(f"招标人：{tenderer_name}")
    if trigger_context:
        doc.add_paragraph(f"触发依据：{trigger_context}")
    doc.add_paragraph("说明：本文件为商务标解析阶段自动生成的承诺函草稿，请结合招标文件原文、法务要求和项目实际情况复核后使用。")
    if commitment_type == "disqualification":
        doc.add_paragraph("我方在参加本项目投标过程中，郑重承诺如下：")
        doc.add_paragraph("1. 我方不存在招标文件中列示的“投标人不得存在下列情形之一”所述任一情形。")
        doc.add_paragraph("2. 如本承诺与实际情况不符，我方愿意按照招标文件和相关规定承担相应责任。")
    else:
        doc.add_paragraph("根据招标文件中的承诺要求，我方郑重承诺如下：")
        doc.add_paragraph("1. 我方将严格按照招标文件要求对相关事项进行响应并履行承诺。")
        if trigger_context:
            doc.add_paragraph(f"2. 本承诺函重点对应的招标文件原文为：{trigger_context}")
        doc.add_paragraph("3. 如本承诺与实际情况不符，我方愿意按照招标文件和相关规定承担相应责任。")
    doc.add_paragraph("")
    doc.add_paragraph("投标人（盖章）：________________")
    doc.add_paragraph("法定代表人或授权代表（签字或盖章）：________________")
    doc.add_paragraph("日期：________________")
    doc.save(path)


def _business_project_name_from_fields(field_groups: dict[str, Any]) -> str:
    project_basics = field_groups.get("projectBasics") if isinstance(field_groups, dict) else []
    for field in project_basics if isinstance(project_basics, list) else []:
        if not isinstance(field, dict):
            continue
        if str(field.get("key") or "") != "projectName":
            continue
        value = str(field.get("value") or "").strip()
        if value:
            return value
    return ""


def _business_tenderer_name_from_fields(field_groups: dict[str, Any]) -> str:
    project_basics = field_groups.get("projectBasics") if isinstance(field_groups, dict) else []
    for field in project_basics if isinstance(project_basics, list) else []:
        if not isinstance(field, dict):
            continue
        if str(field.get("key") or "") != "tenderer":
            continue
        value = str(field.get("value") or "").strip()
        if value:
            return value
    return ""


def _materialize_commitment_letters(manifest: dict[str, Any], letters: list[dict[str, Any]], *, project_name: str = "", tenderer_name: str = "") -> list[dict[str, Any]]:
    output_dir = _commitment_letter_output_dir(manifest)
    prepared: list[dict[str, Any]] = []
    for letter in letters:
        item = copy.deepcopy(letter)
        letter_id = str(item.get("id") or "").strip() or "CL-0000"
        title = str(item.get("title") or item.get("triggerText") or "承诺函").strip() or "承诺函"
        path = output_dir / f"{letter_id}-{_sanitize_docx_name(title, '承诺函')}.docx"
        if not path.exists():
            _write_commitment_letter_docx(path, item, project_name=project_name, tenderer_name=tenderer_name)
        item["docxPath"] = str(path)
        item["workspacePath"] = f"s1_commitment_letters/{path.name}"
        item["status"] = "generated"
        item["previewType"] = "onlyoffice"
        prepared.append(item)
    return prepared


BUSINESS_PROJECT_FACT_FIELD_SPECS: tuple[dict[str, Any], ...] = (
    {"fieldKey": "projectName", "label": "项目名称", "required": True},
    {"fieldKey": "tenderNo", "label": "招标编号", "required": True},
    {"fieldKey": "tenderer", "label": "招标人", "required": True},
    {"fieldKey": "managementUnit", "label": "管理单位", "required": False},
    {"fieldKey": "bidSectionScale", "label": "标段规模", "required": False},
    {"fieldKey": "deliveryPeriod", "label": "交货周期", "required": False},
    {"fieldKey": "warrantyPeriod", "label": "质保期", "required": False},
)


def _build_business_project_fact_fields(
    field_groups: dict[str, Any],
    project_dates: dict[str, Any],
) -> list[dict[str, Any]]:
    project_basics = field_groups.get("projectBasics") if isinstance(field_groups.get("projectBasics"), list) else []
    fields_by_key = {
        str(field.get("key") or ""): field
        for field in project_basics
        if isinstance(field, dict)
    }
    fact_fields: list[dict[str, Any]] = []
    for spec in BUSINESS_PROJECT_FACT_FIELD_SPECS:
        field_key = str(spec["fieldKey"])
        source = fields_by_key.get(field_key) or {}
        value = str(source.get("value") or "").strip()
        fact_fields.append(
            {
                "fieldKey": field_key,
                "label": str(spec["label"]),
                "value": value,
                "category": "项目基础信息",
                "status": "found" if value else "missing",
                "required": bool(spec.get("required", False)),
                "confidence": float(source.get("confidence") or (0.86 if value else 0.0)),
                "sourceFile": str(source.get("sourceFile") or ""),
                "sourceDocumentId": str(source.get("sourceDocumentId") or ""),
                "section": str(source.get("section") or ""),
                "evidence": str(source.get("evidence") or ""),
                "evidenceLocation": str(source.get("evidenceLocation") or ""),
            }
        )
    for field_key, label, date_key, required in (
        ("bidStartDate", "投标起始日期", "startDate", False),
        ("bidDeadline", "投标截止日期", "endDate", True),
    ):
        value = str(project_dates.get(date_key) or "").strip()
        fact_fields.append(
            {
                "fieldKey": field_key,
                "label": label,
                "value": value,
                "category": "投标时间信息",
                "status": "found" if value else "missing",
                "required": required,
                "confidence": 0.78 if value else 0.0,
                "sourceFile": "",
                "sourceDocumentId": "",
                "section": "",
                "evidence": "",
                "evidenceLocation": "",
            }
        )
    return fact_fields


def _scoring_bucket_from_title(title: str) -> str:
    text = _clean(title)
    if not text:
        return ""
    if any(keyword in text for keyword in ("投标报价评分", "报价评分", "价格评分", "开标价格表", "报价表")):
        return "price"
    if any(keyword in text for keyword in ("符合性审查", "合规", "符合性", "审查标准")):
        return "compliance"
    if any(keyword in text for keyword in ("商务评分", "商务评审", "商务打分", "评标办法")):
        return "business"
    return ""


def _parse_markdown_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _is_markdown_separator_row(cells: list[str]) -> bool:
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def _collect_markdown_table(lines: list[str], start_index: int) -> tuple[list[tuple[int, str]], int]:
    rows: list[tuple[int, str]] = []
    index = start_index
    while index < len(lines):
        line = lines[index]
        if not MARKDOWN_TABLE_LINE_PATTERN.match(line.strip()):
            break
        rows.append((index + 1, line))
        index += 1
    return rows, index


def _parse_markdown_scoring_rows(
    *,
    bucket: str,
    rows: list[tuple[int, str]],
    document: dict[str, Any],
    section: str,
    start_index: int,
) -> list[dict[str, Any]]:
    parsed_rows: list[dict[str, Any]] = []
    table_rows = [(line_no, _parse_markdown_table_row(line)) for line_no, line in rows]
    filtered_rows = [(line_no, cells) for line_no, cells in table_rows if cells and not _is_markdown_separator_row(cells)]
    if len(filtered_rows) <= 1:
        return parsed_rows
    data_rows = filtered_rows[1:]
    order = start_index
    for line_no, cells in data_rows:
        if len(cells) < 4:
            continue
        scoring_item = cells[1].strip() if len(cells) > 1 else ""
        score = cells[2].strip() if len(cells) > 2 else ""
        score_point = cells[3].strip() if len(cells) > 3 else ""
        proof_requirement = cells[4].strip() if len(cells) > 4 else ""
        if not scoring_item and not score_point:
            continue
        parsed_rows.append(
            {
                "order": order,
                "scoringItem": scoring_item,
                "score": score,
                "scorePoint": score_point,
                "proofRequirement": proof_requirement,
                "sourceFile": str(document.get("name") or ""),
                "sourceDocumentId": str(document.get("id") or ""),
                "section": section,
                "evidence": " | ".join(cell for cell in cells if cell),
                "evidenceLocation": f"L{line_no}",
            }
        )
        order += 1
    return parsed_rows


def _extract_markdown_scoring(documents: list[dict[str, Any]], texts_by_id: dict[str, str]) -> dict[str, list[dict[str, Any]]]:
    scoring = {"business": [], "price": [], "compliance": []}
    for document in documents:
        text = str(texts_by_id.get(str(document.get("id") or "")) or "")
        if "|" not in text:
            continue
        lines = text.splitlines()
        current_section = ""
        pending_scoring_title = ""
        index = 0
        while index < len(lines):
            raw_line = lines[index]
            line = raw_line.strip()
            if not line:
                index += 1
                continue
            if line.startswith("#"):
                current_section = line.strip("# ").strip()
                pending_scoring_title = ""
                index += 1
                continue

            bucket_from_line = _scoring_bucket_from_title(line)
            if _looks_like_section_heading(line):
                current_section = line
                pending_scoring_title = line if bucket_from_line else ""
                index += 1
                continue

            if bucket_from_line:
                pending_scoring_title = line
                if index + 1 < len(lines) and MARKDOWN_TABLE_LINE_PATTERN.match(lines[index + 1].strip()):
                    rows, next_index = _collect_markdown_table(lines, index + 1)
                    scoring[bucket_from_line].extend(
                        _parse_markdown_scoring_rows(
                            bucket=bucket_from_line,
                            rows=rows,
                            document=document,
                            section=line,
                            start_index=len(scoring[bucket_from_line]) + 1,
                        )
                    )
                    pending_scoring_title = ""
                    index = next_index
                    continue

            if MARKDOWN_TABLE_LINE_PATTERN.match(line):
                bucket = _scoring_bucket_from_title(pending_scoring_title or current_section)
                if bucket:
                    rows, next_index = _collect_markdown_table(lines, index)
                    scoring[bucket].extend(
                        _parse_markdown_scoring_rows(
                            bucket=bucket,
                            rows=rows,
                            document=document,
                            section=pending_scoring_title or current_section,
                            start_index=len(scoring[bucket]) + 1,
                        )
                    )
                    pending_scoring_title = ""
                    index = next_index
                    continue

            index += 1
    return scoring


def _merge_business_scoring(base_scoring: dict[str, Any], markdown_scoring: dict[str, list[dict[str, Any]]]) -> dict[str, list[dict[str, Any]]]:
    merged = _filter_business_scoring(base_scoring)
    seen = {
        bucket: {
            (
                str(row.get("sourceDocumentId") or ""),
                str(row.get("evidenceLocation") or ""),
                str(row.get("scoringItem") or ""),
                str(row.get("score") or ""),
            )
            for row in merged.get(bucket) or []
            if isinstance(row, dict)
        }
        for bucket in merged
    }
    for bucket, rows in markdown_scoring.items():
        for row in rows:
            key = (
                str(row.get("sourceDocumentId") or ""),
                str(row.get("evidenceLocation") or ""),
                str(row.get("scoringItem") or ""),
                str(row.get("score") or ""),
            )
            if key in seen[bucket]:
                continue
            seen[bucket].add(key)
            merged[bucket].append(row)
    return merged


def build_business_result(manifest: dict[str, Any], *, mode: str = "opencode-skill") -> dict[str, Any]:
    base_result = parse_technical_manifest(manifest, mode=f"{mode}-technical-base")
    documents = [item for item in manifest.get("documents") or [] if isinstance(item, dict)]
    texts_by_id = _load_texts_by_id(documents)

    base_items = copy.deepcopy(base_result.get("items") if isinstance(base_result.get("items"), list) else [])
    structured = copy.deepcopy(base_result.get("structured") if isinstance(base_result.get("structured"), dict) else {})
    hint_items = _scan_business_hint_items(documents, texts_by_id)
    docx_candidate_items = _extract_docx_core_candidate_items(documents)
    merged_items = [*copy.deepcopy(base_items), *docx_candidate_items, *hint_items]

    scoring = _merge_business_scoring(structured.get("scoringCriteria") or {}, _extract_markdown_scoring(documents, texts_by_id))
    project_dates = structured.get("projectDates") if isinstance(structured.get("projectDates"), dict) else {}
    field_groups = {
        "projectBasics": _build_business_project_basics(merged_items, project_dates),
        "businessResponse": _build_business_response_fields(merged_items),
        "qualificationSupport": _build_qualification_support_fields(merged_items),
        "qualificationRequirements": _build_qualification_requirements(merged_items),
        "bidderInstructions": _extract_bidder_instruction_rows(documents),
        "commercialRejectionClauses": _extract_commercial_rejection_clauses(documents, texts_by_id),
        "commitmentRequirements": _build_commitment_requirement_fields(merged_items),
    }
    presence = _build_business_requirement_presence(merged_items)
    commitment_letters = _materialize_commitment_letters(
        manifest,
        _build_business_commitment_letters(merged_items),
        project_name=_business_project_name_from_fields(field_groups),
        tenderer_name=_business_tenderer_name_from_fields(field_groups),
    )
    commitment_clues = _build_business_commitment_clues(merged_items)
    project_fact_fields = _build_business_project_fact_fields(field_groups, project_dates)

    return {
        "items": merged_items,
        "structured": {
            "schemaVersion": SCHEMA_VERSION,
            "targetSkill": SKILL_NAME,
            "mode": mode,
            "sourceDocuments": copy.deepcopy(structured.get("sourceDocuments") or []),
            "scoringCriteria": scoring,
            "fieldGroups": field_groups,
            "requirementPresence": presence,
            "coverage": _build_business_coverage(field_groups, scoring, presence),
            "projectDates": {
                "startDate": str(project_dates.get("startDate") or ""),
                "endDate": str(project_dates.get("endDate") or ""),
            },
            "appendices": copy.deepcopy(structured.get("appendices") or []),
            "commitmentLetters": commitment_letters,
            "commitmentClues": commitment_clues,
            "projectFactFields": project_fact_fields,
            "categoryCounts": {
                "商务评分": len(scoring.get("business") or []),
                "报价评分": len(scoring.get("price") or []),
                "合规审查": len(scoring.get("compliance") or []),
                "商务附表": len(structured.get("appendices") or []),
                "承诺文件": len(commitment_letters),
                "待确认承诺线索": len(commitment_clues),
            },
        },
    }
