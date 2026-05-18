#!/usr/bin/env python3
"""商务标成稿 Word 格式清洗核心逻辑。"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import zipfile
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips

try:
    from outline_matcher import flatten_outline, load_outline
    from pagination_cleaner import clean_section_page_breaks
except ImportError:  # pragma: no cover - 支持包外直接导入
    from .outline_matcher import flatten_outline, load_outline
    from .pagination_cleaner import clean_section_page_breaks


def load_style_spec(path: str | Path) -> dict[str, Any]:
    data = json.loads(Path(path).read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError("styleSpecPath 必须是 JSON object")
    return data


def load_toc_style_spec(style_spec: dict[str, Any], style_spec_path: str | Path) -> dict[str, Any]:
    toc_cfg = style_spec.get("toc", {})
    configured_path = toc_cfg.get("style_spec_path") or toc_cfg.get("styleSpecPath")
    if configured_path:
        toc_path = Path(str(configured_path))
        if not toc_path.is_absolute():
            toc_path = Path(style_spec_path).parent / toc_path
        data = json.loads(toc_path.read_text(encoding="utf-8"))
        if not isinstance(data, dict):
            raise ValueError("toc style reference 必须是 JSON object")
        data["_resolved_path"] = str(toc_path.resolve())
        return data

    default_path = Path(style_spec_path).parent / "business_toc_style.json"
    if default_path.exists():
        data = json.loads(default_path.read_text(encoding="utf-8"))
        if not isinstance(data, dict):
            raise ValueError("business_toc_style.json 必须是 JSON object")
        data["_resolved_path"] = str(default_path.resolve())
        return data

    return {
        "schema_version": "business_toc_style.legacy",
        "title": {
            "text": str(toc_cfg.get("title", "")),
            "style_name": str(toc_cfg.get("title_style", "")),
        },
        "field": {
            "instruction": str(toc_cfg.get("field_instruction", "")),
            "placeholder": str(toc_cfg.get("placeholder", "")),
            "page_break_after": bool(toc_cfg.get("page_break_after", True)),
        },
        "entry_styles": {},
    }


def clean_docx(
    *,
    input_file: str | Path,
    outline_file: str | Path,
    output_file: str | Path,
    project_name: str,
    style_spec_path: str | Path,
) -> dict[str, Any]:
    """复制 inputFile 到 outputFile，并只在 outputFile 上执行格式清洗。"""
    input_path = Path(input_file)
    outline_path = Path(outline_file)
    output_path = Path(output_file)
    style_path = Path(style_spec_path)

    if not input_path.exists():
        raise FileNotFoundError(f"inputFile 不存在: {input_path}")
    if not outline_path.exists():
        raise FileNotFoundError(f"outlineFile 不存在: {outline_path}")
    if not style_path.exists():
        raise FileNotFoundError(f"styleSpecPath 不存在: {style_path}")
    if input_path.resolve() == output_path.resolve():
        raise ValueError("inputFile 不得被覆盖，outputFile 必须不同于 inputFile")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(input_path, output_path)

    style_spec = load_style_spec(style_path)
    toc_style_spec = load_toc_style_spec(style_spec, style_path)
    outline_items = flatten_outline(load_outline(outline_path))

    doc = Document(str(output_path))
    _apply_page_setup(doc, style_spec)
    toc_style_ids = _configure_document_styles(doc, style_spec, toc_style_spec)
    _strip_heading_style_numpr(doc)

    heading_result = _match_and_promote_headings(doc, outline_items, style_spec)
    _apply_body_format(doc, style_spec, set(heading_result["matched_paragraph_indexes"]))
    _apply_table_format(doc, style_spec)
    header_result = _apply_business_headers(doc, style_spec, project_name)

    toc_present_before = document_has_toc(doc)
    toc_inserted = False
    if not toc_present_before and style_spec.get("toc", {}).get("insert_when_missing", True):
        _insert_toc(doc, style_spec, toc_style_spec, toc_style_ids)
        toc_inserted = True
    toc_present = toc_present_before or toc_inserted

    doc.save(str(output_path))
    pagination_result = clean_section_page_breaks(output_path, outline_path)
    force_update_fields(output_path)

    risks = _collect_risks(
        outline_count=len(outline_items),
        unmatched_count=len(heading_result["unmatched_headings"]),
        toc_present=toc_present,
        header_cleaned=header_result["header_cleaned"],
    )
    return {
        "inputFile": str(input_path),
        "outlineFile": str(outline_path),
        "outputFile": str(output_path),
        "styleSpecPath": str(style_path),
        "projectName": project_name,
        "outlineCount": len(outline_items),
        "matchedHeadingCount": len(heading_result["matched_headings"]),
        "unmatchedHeadingCount": len(heading_result["unmatched_headings"]),
        "matchedHeadings": heading_result["matched_headings"],
        "unmatchedHeadings": heading_result["unmatched_headings"],
        "tocInserted": toc_inserted,
        "tocPresent": toc_present,
        "headerCleaned": header_result["header_cleaned"],
        "headerResidualFound": header_result["residual_found_before"],
        "pagination": pagination_result,
        "formatRisks": risks,
    }


def document_has_toc(doc: Document) -> bool:
    return any((node.text or "").upper().find("TOC") >= 0 for node in doc.element.iter(qn("w:instrText")))


def force_update_fields(docx_path: str | Path) -> None:
    path = Path(docx_path)
    with zipfile.ZipFile(path, "r") as zin:
        names = zin.namelist()
        data = {name: zin.read(name) for name in names}

    settings_name = "word/settings.xml"
    if settings_name not in data:
        return

    text = data[settings_name].decode("utf-8", errors="replace")
    if "<w:updateFields" in text:
        text = re.sub(r"<w:updateFields\b[^>]*/>", '<w:updateFields w:val="true"/>', text, count=1)
    else:
        text = re.sub(r"(<w:settings\b[^>]*>)", r'\1<w:updateFields w:val="true"/>', text, count=1)
    data[settings_name] = text.encode("utf-8")

    tmp = path.with_suffix(".fields.tmp.docx")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in names:
            zout.writestr(name, data[name])
    shutil.move(str(tmp), str(path))


def _apply_page_setup(doc: Document, style_spec: dict[str, Any]) -> None:
    page_cfg = style_spec["page"]
    for section in doc.sections:
        section.top_margin = Cm(float(page_cfg["top_cm"]))
        section.bottom_margin = Cm(float(page_cfg["bottom_cm"]))
        section.left_margin = Cm(float(page_cfg["left_cm"]))
        section.right_margin = Cm(float(page_cfg["right_cm"]))
        section.header_distance = Cm(float(page_cfg["header_top_cm"]))
        section.footer_distance = Cm(float(page_cfg["footer_bottom_cm"]))
        if section.start_type is None:
            section.start_type = WD_SECTION_START.NEW_PAGE


def _configure_document_styles(
    doc: Document,
    style_spec: dict[str, Any],
    toc_style_spec: dict[str, Any] | None = None,
) -> dict[str, str]:
    body_cfg = style_spec["body"]
    for style_name in ("Normal", "正文"):
        try:
            _apply_style_format(doc.styles[style_name], body_cfg)
        except KeyError:
            continue

    for level_text, cfg in style_spec.get("heading", {}).items():
        for style_name in (f"Heading {level_text}", f"标题 {level_text}", f"标题{level_text}"):
            try:
                _apply_style_format(doc.styles[style_name], cfg)
                break
            except KeyError:
                continue

    return _configure_toc_styles(doc, toc_style_spec or {})


def _configure_toc_styles(doc: Document, toc_style_spec: dict[str, Any]) -> dict[str, str]:
    style_ids: dict[str, str] = {}
    title_cfg = toc_style_spec.get("title") or {}
    if title_cfg:
        style_name = str(title_cfg.get("style_name") or title_cfg.get("styleName") or "TOC Heading")
        style = _get_or_create_paragraph_style(doc, style_name)
        _apply_style_format(style, title_cfg)
        style_ids["title"] = style.style_id

    entry_styles = toc_style_spec.get("entry_styles") or toc_style_spec.get("entryStyles") or {}
    if isinstance(entry_styles, dict):
        for level_text, cfg in entry_styles.items():
            if not isinstance(cfg, dict):
                continue
            style_name = str(cfg.get("style_name") or cfg.get("styleName") or f"TOC {level_text}")
            style = _get_or_create_paragraph_style(doc, style_name)
            _apply_style_format(style, cfg)
            style_ids[str(level_text)] = style.style_id
    return style_ids


def _get_or_create_paragraph_style(doc: Document, style_name: str):
    for candidate in (style_name, style_name.replace(" ", "")):
        try:
            return doc.styles[candidate]
        except KeyError:
            continue
    return doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)


def _apply_style_format(style, cfg: dict[str, Any]) -> None:
    _apply_font_to_style(style, cfg)
    if hasattr(style, "paragraph_format"):
        _apply_paragraph_format(style.paragraph_format, cfg)


def _apply_font_to_style(style, cfg: dict[str, Any]) -> None:
    font = style.font
    if "en_font" in cfg:
        font.name = str(cfg["en_font"])
    if "size_pt" in cfg:
        font.size = Pt(float(cfg["size_pt"]))
    font.bold = bool(cfg.get("bold", False))
    color = _font_color(cfg)
    if color:
        font.color.rgb = RGBColor.from_string(color)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    if "zh_font" in cfg or "en_font" in cfg:
        _set_rfonts(r_fonts, cfg)
    if color:
        _set_rpr_color(r_pr, color)


def _apply_paragraph_format(paragraph_format, cfg: dict[str, Any]) -> None:
    paragraph_format.alignment = _paragraph_alignment(cfg.get("align"))
    paragraph_format.space_before = Pt(float(cfg.get("space_before_pt", 0)))
    paragraph_format.space_after = Pt(float(cfg.get("space_after_pt", 0)))
    line_rule = str(cfg.get("line_spacing_rule") or cfg.get("lineSpacingRule") or "").strip().lower()
    if "line_spacing_twips" in cfg or "lineSpacingTwips" in cfg:
        paragraph_format.line_spacing = Twips(int(cfg.get("line_spacing_twips", cfg.get("lineSpacingTwips"))))
        if line_rule == "exact":
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        elif line_rule in {"at_least", "atleast"}:
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    else:
        paragraph_format.line_spacing = float(cfg.get("line_spacing", 1.0))
    if "left_indent_cm" in cfg:
        paragraph_format.left_indent = Cm(float(cfg.get("left_indent_cm", 0)))
    if "left_indent_twips" in cfg or "leftIndentTwips" in cfg:
        paragraph_format.left_indent = Twips(int(cfg.get("left_indent_twips", cfg.get("leftIndentTwips"))))
    if "first_line_indent_chars" in cfg:
        paragraph_format.first_line_indent = _char_indent(cfg)
    if "first_line_indent_twips" in cfg or "firstLineIndentTwips" in cfg:
        paragraph_format.first_line_indent = Twips(
            int(cfg.get("first_line_indent_twips", cfg.get("firstLineIndentTwips")))
        )
    if "left_indent_cm" not in cfg and "left_indent_twips" not in cfg and "leftIndentTwips" not in cfg:
        paragraph_format.left_indent = Pt(0)
    if "first_line_indent_chars" not in cfg and "first_line_indent_twips" not in cfg and "firstLineIndentTwips" not in cfg:
        paragraph_format.first_line_indent = Pt(0)
    _apply_tab_stops(paragraph_format, cfg)


def _char_indent(cfg: dict[str, Any]):
    chars = float(cfg.get("first_line_indent_chars", 0))
    if chars == 0:
        return Pt(0)
    return Pt(chars * float(cfg["size_pt"]))


def _paragraph_alignment(value: Any):
    key = str(value or "").strip().lower()
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(key, WD_ALIGN_PARAGRAPH.LEFT)


def _apply_tab_stops(paragraph_format, cfg: dict[str, Any]) -> None:
    tab_stops = cfg.get("tab_stops") or cfg.get("tabStops") or []
    if not isinstance(tab_stops, list):
        return
    paragraph_format.tab_stops.clear_all()
    for tab in tab_stops:
        if not isinstance(tab, dict):
            continue
        pos = tab.get("position_twips", tab.get("positionTwips"))
        if pos is None:
            continue
        paragraph_format.tab_stops.add_tab_stop(
            Twips(int(pos)),
            alignment=_tab_alignment(tab.get("alignment")),
            leader=_tab_leader(tab.get("leader")),
        )


def _tab_alignment(value: Any):
    key = str(value or "").strip().lower()
    mapping = {
        "left": WD_TAB_ALIGNMENT.LEFT,
        "center": WD_TAB_ALIGNMENT.CENTER,
        "right": WD_TAB_ALIGNMENT.RIGHT,
        "decimal": WD_TAB_ALIGNMENT.DECIMAL,
        "bar": WD_TAB_ALIGNMENT.BAR,
    }
    return mapping.get(key, WD_TAB_ALIGNMENT.LEFT)


def _tab_leader(value: Any):
    key = str(value or "").strip().lower()
    mapping = {
        "none": WD_TAB_LEADER.SPACES,
        "space": WD_TAB_LEADER.SPACES,
        "spaces": WD_TAB_LEADER.SPACES,
        "dot": WD_TAB_LEADER.DOTS,
        "dots": WD_TAB_LEADER.DOTS,
        "hyphen": WD_TAB_LEADER.DASHES,
        "dash": WD_TAB_LEADER.DASHES,
        "dashes": WD_TAB_LEADER.DASHES,
        "underscore": WD_TAB_LEADER.HEAVY,
        "heavy": WD_TAB_LEADER.HEAVY,
    }
    return mapping.get(key, WD_TAB_LEADER.SPACES)


def _table_alignment(value: Any):
    key = str(value or "").strip().lower()
    mapping = {
        "left": WD_TABLE_ALIGNMENT.LEFT,
        "center": WD_TABLE_ALIGNMENT.CENTER,
        "right": WD_TABLE_ALIGNMENT.RIGHT,
    }
    return mapping.get(key, WD_TABLE_ALIGNMENT.CENTER)


def _set_rfonts(r_fonts, cfg: dict[str, Any]) -> None:
    zh_font = str(cfg.get("zh_font") or cfg.get("en_font") or "")
    en_font = str(cfg.get("en_font") or cfg.get("zh_font") or "")
    if not zh_font and not en_font:
        return
    r_fonts.set(qn("w:eastAsia"), zh_font)
    r_fonts.set(qn("w:ascii"), en_font)
    r_fonts.set(qn("w:hAnsi"), en_font)
    r_fonts.set(qn("w:cs"), en_font)


def _font_color(cfg: dict[str, Any]) -> str | None:
    raw = cfg.get("font_color", cfg.get("fontColor", cfg.get("color")))
    if raw is None:
        return None
    value = str(raw).strip().lstrip("#").upper()
    if re.fullmatch(r"[0-9A-F]{6}", value):
        return value
    return None


def _set_rpr_color(r_pr, color: str) -> None:
    color_el = r_pr.find(qn("w:color"))
    if color_el is None:
        color_el = OxmlElement("w:color")
        r_pr.append(color_el)
    color_el.set(qn("w:val"), color)
    # Word built-in Heading styles often carry theme colors. Remove them so
    # OnlyOffice/WPS cannot render cleaned business headings as theme blue.
    for attr in ("themeColor", "themeTint", "themeShade"):
        color_el.attrib.pop(qn(f"w:{attr}"), None)


def _apply_run_format(run, cfg: dict[str, Any]) -> None:
    if "en_font" in cfg:
        run.font.name = str(cfg["en_font"])
    if "size_pt" in cfg:
        run.font.size = Pt(float(cfg["size_pt"]))
    run.font.bold = bool(cfg.get("bold", False))
    run.font.italic = False
    run.font.underline = False
    color = _font_color(cfg)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    if "zh_font" in cfg or "en_font" in cfg:
        _set_rfonts(r_fonts, cfg)
    if color:
        _set_rpr_color(r_pr, color)


def _apply_direct_paragraph_format(paragraph, cfg: dict[str, Any]) -> None:
    paragraph.alignment = _paragraph_alignment(cfg.get("align"))
    fmt = paragraph.paragraph_format
    _apply_paragraph_format(fmt, cfg)
    for run in paragraph.runs:
        _apply_run_format(run, cfg)


def _match_and_promote_headings(
    doc: Document,
    outline_items: list[dict[str, Any]],
    style_spec: dict[str, Any],
) -> dict[str, Any]:
    paragraphs = list(doc.paragraphs)
    cursor = 0
    matched: list[dict[str, Any]] = []
    unmatched: list[dict[str, Any]] = []
    matched_indexes: list[int] = []

    for item in outline_items:
        found = _find_outline_item(paragraphs, item, cursor)
        if found is None:
            found = _find_outline_item(paragraphs, item, cursor, allow_fuzzy=True)
        if found is None:
            unmatched.append(_outline_report_item(item))
            continue

        index, mode = found
        paragraph = paragraphs[index]
        level = max(1, min(int(item["level"]), 9))
        normalized_text = _format_heading_text(item)
        _replace_paragraph_text(paragraph, normalized_text)
        _set_heading_style(paragraph, level, doc)
        _strip_paragraph_numpr(paragraph)
        _clear_direct_outline_level(paragraph)
        _apply_direct_paragraph_format(paragraph, _heading_cfg(style_spec, level))
        matched_indexes.append(index)
        matched.append(
            {
                **_outline_report_item(item),
                "paragraphIndex": index,
                "matchMode": mode,
                "text": normalized_text,
            }
        )
        cursor = index + 1

    return {
        "matched_headings": matched,
        "unmatched_headings": unmatched,
        "matched_paragraph_indexes": matched_indexes,
    }


def _find_outline_item(
    paragraphs,
    item: dict[str, Any],
    start_index: int,
    *,
    end_index: int | None = None,
    allow_fuzzy: bool = False,
) -> tuple[int, str] | None:
    full_key = _normalize_match_text(_format_heading_text(item)) if item.get("number") else ""
    title_key = _normalize_match_text(str(item.get("title") or ""))
    number_key = _normalize_match_text(str(item.get("number") or ""))

    end = len(paragraphs) if end_index is None else max(0, min(end_index, len(paragraphs)))
    for index in range(start_index, end):
        text_key = _normalize_match_text(paragraphs[index].text)
        if not text_key:
            continue
        if full_key and text_key == full_key:
            return index, "number+title"
        if title_key and text_key == title_key:
            return index, "title"
        if allow_fuzzy and _is_probable_heading_match(text_key, title_key, number_key, full_key):
            return index, "fuzzy"
    return None


def _normalize_match_text(text: str) -> str:
    normalized = str(text or "").strip()
    normalized = normalized.replace("\u3000", " ")
    normalized = re.sub(r"^[一二三四五六七八九十]+[、.．]\s*", "", normalized)
    normalized = re.sub(r"^\d+(?:[.．]\d+)*[、.．]?\s*", "", normalized)
    normalized = re.sub(r"\s+", "", normalized)
    normalized = re.sub(r"[：:；;。,.，、（）()【】\\[\\]《》<>]+", "", normalized)
    return normalized


def _is_probable_heading_match(text_key: str, title_key: str, number_key: str, full_key: str) -> bool:
    if not text_key or len(text_key) > 80:
        return False
    if full_key and (text_key.startswith(full_key) or full_key.startswith(text_key)):
        return True
    if title_key and (text_key.startswith(title_key) or title_key.startswith(text_key)):
        return True
    if number_key and title_key and number_key in text_key and title_key in text_key:
        return True
    return False


def _format_heading_text(item: dict[str, Any]) -> str:
    title = str(item.get("title") or "").strip()
    number = str(item.get("number") or "").strip()
    return f"{number} {title}".strip() if number else title


def _outline_report_item(item: dict[str, Any]) -> dict[str, Any]:
    return {
        "id": item.get("id"),
        "number": item.get("number", ""),
        "title": item.get("title", ""),
        "level": item.get("level"),
        "order": item.get("order"),
    }


def _heading_cfg(style_spec: dict[str, Any], level: int) -> dict[str, Any]:
    heading = style_spec.get("heading", {})
    if str(level) in heading:
        return heading[str(level)]
    if heading:
        return heading[sorted(heading.keys(), key=lambda value: int(value))[-1]]
    return style_spec["body"]


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    text_nodes = [node for node in paragraph._element.iter(qn("w:t"))]
    if not text_nodes:
        paragraph.add_run(new_text)
        return
    text_nodes[0].text = new_text
    for node in text_nodes[1:]:
        node.text = ""


def _set_heading_style(paragraph, level: int, doc: Document) -> None:
    for style_name in (f"Heading {level}", f"标题 {level}", f"标题{level}"):
        try:
            paragraph.style = doc.styles[style_name]
            return
        except KeyError:
            continue


def _strip_heading_style_numpr(doc: Document) -> int:
    count = 0
    for level in range(1, 10):
        for style_name in (f"Heading {level}", f"标题 {level}", f"标题{level}"):
            try:
                style = doc.styles[style_name]
            except KeyError:
                continue
            p_pr = style.element.find(qn("w:pPr"))
            if p_pr is None:
                continue
            num_pr = p_pr.find(qn("w:numPr"))
            if num_pr is not None:
                p_pr.remove(num_pr)
                count += 1
    return count


def _strip_paragraph_numpr(paragraph) -> bool:
    p_pr = paragraph._element.find(qn("w:pPr"))
    if p_pr is None:
        return False
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return False
    p_pr.remove(num_pr)
    return True


def _clear_direct_outline_level(paragraph) -> bool:
    p_pr = paragraph._element.find(qn("w:pPr"))
    if p_pr is None:
        return False
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        return False
    p_pr.remove(outline)
    return True


def _apply_body_format(doc: Document, style_spec: dict[str, Any], matched_indexes: set[int]) -> None:
    body_cfg = style_spec["body"]
    for index, paragraph in enumerate(doc.paragraphs):
        if index in matched_indexes or _paragraph_has_toc_field(paragraph):
            continue
        # Whole-file mounted Word materials can carry arbitrary Heading/Title
        # styles from the source document. Non-outline paragraphs must be
        # normalized before direct formatting, otherwise style-level title
        # attributes may leak into the final business bid.
        _set_paragraph_style_safe(paragraph, "Normal")
        _apply_direct_paragraph_format(paragraph, body_cfg)


def _is_heading_paragraph(paragraph) -> bool:
    style_name = paragraph.style.name if paragraph.style else ""
    return bool(re.match(r"^(Heading|标题)\s*\d+$", style_name or "", flags=re.IGNORECASE))


def _set_paragraph_style_safe(paragraph, style_name: str) -> None:
    try:
        paragraph.style = style_name
    except Exception:
        pass


def _paragraph_has_toc_field(paragraph) -> bool:
    return any((node.text or "").upper().find("TOC") >= 0 for node in paragraph._element.iter(qn("w:instrText")))


def _apply_table_format(doc: Document, style_spec: dict[str, Any]) -> None:
    table_cfg = style_spec["table_cell"]
    for table in doc.tables:
        table.alignment = _table_alignment(table_cfg.get("table_align", table_cfg.get("align")))
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    _apply_direct_paragraph_format(paragraph, table_cfg)


def _apply_business_headers(doc: Document, style_spec: dict[str, Any], project_name: str) -> dict[str, Any]:
    header_cfg = style_spec["header"]
    header_text = _render_header_text(header_cfg, project_name)
    residual_terms = [str(term) for term in header_cfg.get("residual_terms", [])]
    residual_found_before = False
    residual_found_after = False

    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            try:
                header.is_linked_to_previous = False
            except ValueError:
                pass
            before = "\n".join(paragraph.text for paragraph in header.paragraphs)
            residual_found_before = residual_found_before or any(term in before for term in residual_terms)
            _clear_header_text_nodes(header)
            paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            run = paragraph.add_run(header_text)
            _apply_direct_paragraph_format(paragraph, header_cfg)
            _apply_run_format(run, header_cfg)
            after = "\n".join(paragraph.text for paragraph in header.paragraphs)
            residual_found_after = residual_found_after or any(term in after for term in residual_terms)

    return {
        "residual_found_before": residual_found_before,
        "residual_found_after": residual_found_after,
        "header_cleaned": not residual_found_after,
    }


def _render_header_text(header_cfg: dict[str, Any], project_name: str) -> str:
    template = str(header_cfg["text_template"])
    return template.replace("{projectName}", project_name).replace("{project_name}", project_name)


def _clear_header_text_nodes(header) -> None:
    for paragraph in header.paragraphs:
        for node in paragraph._element.iter(qn("w:t")):
            node.text = ""


def _insert_toc(
    doc: Document,
    style_spec: dict[str, Any],
    toc_style_spec: dict[str, Any],
    toc_style_ids: dict[str, str],
) -> None:
    toc_cfg = style_spec["toc"]
    title_cfg = toc_style_spec.get("title") or {}
    field_cfg = toc_style_spec.get("field") or {}
    title_text = str(title_cfg.get("text") or toc_cfg.get("title", ""))
    title_style_id = toc_style_ids.get("title") or str(title_cfg.get("style_name") or toc_cfg.get("title_style", ""))
    instruction = str(field_cfg.get("instruction") or toc_cfg.get("field_instruction", ""))
    placeholder = str(field_cfg.get("placeholder") or toc_cfg.get("placeholder", ""))
    field_style_id = (
        toc_style_ids.get(str(field_cfg.get("level", "1")))
        or toc_style_ids.get("1")
        or str(field_cfg.get("style_name") or "")
    )
    page_break_after = bool(field_cfg.get("page_break_after", toc_cfg.get("page_break_after", True)))
    body = doc.element.body
    first = next((child for child in list(body) if child.tag != qn("w:sectPr")), None)
    elements = [
        _make_text_paragraph(title_text, title_style_id),
        _make_toc_field_paragraph(instruction, placeholder, field_style_id),
    ]
    if page_break_after:
        elements.append(_make_toc_section_break_paragraph(doc, style_spec))

    insert_after = _find_cover_page_break_body_child(body)
    if insert_after is not None:
        for element in reversed(elements):
            insert_after.addnext(element)
        return

    for element in elements:
        if first is not None:
            first.addprevious(element)
        else:
            body.insert(len(body), element)


def _find_cover_page_break_body_child(body):
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        if _is_heading_body_child(child):
            return None
        if _element_contains_page_break(child):
            return child
    return None


def _is_heading_body_child(element) -> bool:
    if element.tag != qn("w:p"):
        return False
    p_style = element.find(qn("w:pPr") + "/" + qn("w:pStyle"))
    if p_style is None:
        return False
    value = p_style.get(qn("w:val")) or ""
    compact = value.replace(" ", "")
    return compact.startswith("Heading") and compact[7:].isdigit()


def _element_contains_page_break(element) -> bool:
    for br in element.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def _make_text_paragraph(text: str, style_name: str = ""):
    paragraph = OxmlElement("w:p")
    if style_name:
        p_pr = OxmlElement("w:pPr")
        p_style = OxmlElement("w:pStyle")
        p_style.set(qn("w:val"), style_name)
        p_pr.append(p_style)
        paragraph.append(p_pr)
    run = OxmlElement("w:r")
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    paragraph.append(run)
    return paragraph


def _make_toc_field_paragraph(instruction: str, placeholder: str, style_name: str = ""):
    paragraph = OxmlElement("w:p")
    if style_name:
        p_pr = OxmlElement("w:pPr")
        p_style = OxmlElement("w:pStyle")
        p_style.set(qn("w:val"), style_name)
        p_pr.append(p_style)
        paragraph.append(p_pr)
    field_parts = (
        ("begin", None),
        (None, instruction),
        ("separate", None),
        (None, placeholder),
        ("end", None),
    )
    for field_type, text in field_parts:
        run = OxmlElement("w:r")
        if field_type:
            fld_char = OxmlElement("w:fldChar")
            fld_char.set(qn("w:fldCharType"), field_type)
            run.append(fld_char)
        elif text == instruction:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = text
            run.append(instr)
        else:
            text_el = OxmlElement("w:t")
            text_el.text = text
            run.append(text_el)
        paragraph.append(run)
    return paragraph


def _make_page_break_paragraph():
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    paragraph.append(run)
    return paragraph


def _make_toc_section_break_paragraph(doc: Document, style_spec: dict[str, Any]):
    paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    sect_pr = _reference_section_properties(doc)
    _set_section_start_next_page(sect_pr)
    _set_portrait_page_size(sect_pr, doc, style_spec)
    _set_reference_page_margins(sect_pr, style_spec)
    p_pr.append(sect_pr)
    paragraph.append(p_pr)
    return paragraph


def _reference_section_properties(doc: Document):
    for section in doc.sections:
        sect_pr = getattr(section, "_sectPr", None)
        if sect_pr is not None and int(section.page_height) >= int(section.page_width):
            return deepcopy(sect_pr)
    if doc.sections:
        return deepcopy(doc.sections[0]._sectPr)
    return OxmlElement("w:sectPr")


def _set_section_start_next_page(sect_pr) -> None:
    type_el = sect_pr.find(qn("w:type"))
    if type_el is None:
        type_el = OxmlElement("w:type")
        pg_sz = sect_pr.find(qn("w:pgSz"))
        if pg_sz is not None:
            pg_sz.addprevious(type_el)
        else:
            sect_pr.insert(0, type_el)
    type_el.set(qn("w:val"), "nextPage")


def _set_portrait_page_size(sect_pr, doc: Document, style_spec: dict[str, Any]) -> None:
    width_twips, height_twips = _portrait_page_size_twips(doc, style_spec)
    pg_sz = sect_pr.find(qn("w:pgSz"))
    if pg_sz is None:
        pg_sz = OxmlElement("w:pgSz")
        type_el = sect_pr.find(qn("w:type"))
        if type_el is not None:
            type_el.addnext(pg_sz)
        else:
            sect_pr.insert(0, pg_sz)
    pg_sz.set(qn("w:w"), str(width_twips))
    pg_sz.set(qn("w:h"), str(height_twips))
    pg_sz.attrib.pop(qn("w:orient"), None)


def _portrait_page_size_twips(doc: Document, style_spec: dict[str, Any]) -> tuple[int, int]:
    page_cfg = style_spec.get("page", {})
    width_twips = _page_dimension_twips(
        page_cfg,
        ("width_twips", "page_width_twips", "portrait_width_twips"),
        ("width_cm", "page_width_cm", "portrait_width_cm"),
    )
    height_twips = _page_dimension_twips(
        page_cfg,
        ("height_twips", "page_height_twips", "portrait_height_twips"),
        ("height_cm", "page_height_cm", "portrait_height_cm"),
    )
    if width_twips is not None and height_twips is not None:
        return min(width_twips, height_twips), max(width_twips, height_twips)

    dimensions: list[tuple[int, int]] = []
    for section in doc.sections:
        width = int(section.page_width.twips)
        height = int(section.page_height.twips)
        if width and height:
            dimensions.append((width, height))
    for width, height in dimensions:
        if height >= width:
            return width, height
    if dimensions:
        width, height = dimensions[0]
        return min(width, height), max(width, height)
    raise ValueError("无法从 styleSpec 或文档 section 推导目录竖版页宽页高")


def _page_dimension_twips(
    cfg: dict[str, Any],
    twips_keys: tuple[str, ...],
    cm_keys: tuple[str, ...],
) -> int | None:
    for key in twips_keys:
        if key in cfg:
            return int(cfg[key])
    for key in cm_keys:
        if key in cfg:
            return _cm_to_twips(float(cfg[key]))
    return None


def _set_reference_page_margins(sect_pr, style_spec: dict[str, Any]) -> None:
    page_cfg = style_spec.get("page", {})
    pg_mar = sect_pr.find(qn("w:pgMar"))
    if pg_mar is None:
        pg_mar = OxmlElement("w:pgMar")
        pg_sz = sect_pr.find(qn("w:pgSz"))
        if pg_sz is not None:
            pg_sz.addnext(pg_mar)
        else:
            sect_pr.append(pg_mar)

    margin_keys = {
        "top": "top_cm",
        "bottom": "bottom_cm",
        "left": "left_cm",
        "right": "right_cm",
        "header": "header_top_cm",
        "footer": "footer_bottom_cm",
    }
    for attr, key in margin_keys.items():
        if key in page_cfg:
            pg_mar.set(qn(f"w:{attr}"), str(_cm_to_twips(float(page_cfg[key]))))
    if "gutter_cm" in page_cfg:
        pg_mar.set(qn("w:gutter"), str(_cm_to_twips(float(page_cfg["gutter_cm"]))))


def _cm_to_twips(value: float) -> int:
    return int(round(Cm(value).twips))


def _collect_risks(
    *,
    outline_count: int,
    unmatched_count: int,
    toc_present: bool,
    header_cleaned: bool,
) -> list[str]:
    risks: list[str] = []
    if outline_count == 0:
        risks.append("outline 为空，未能定位任何商务标标题。")
    if unmatched_count:
        risks.append("存在未匹配标题，脚本不会插入或删除正文，请人工核对缺失章节。")
    if not toc_present:
        risks.append("未检测到 TOC 域，目录页需要人工补充。")
    if not header_cleaned:
        risks.append("页眉仍疑似存在技术标残留语义，请人工检查页眉。")
    if not risks:
        risks.append("未发现脚本可识别的格式风险；仍建议在 Word/WPS 中打开后刷新目录并人工抽检。")
    return risks


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="清洗商务标成稿 docx 格式")
    parser.add_argument("--input", required=True)
    parser.add_argument("--outline", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--project-name", required=True)
    parser.add_argument("--style", required=True)
    args = parser.parse_args(argv)

    result = clean_docx(
        input_file=args.input,
        outline_file=args.outline,
        output_file=args.output,
        project_name=args.project_name,
        style_spec_path=args.style,
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
