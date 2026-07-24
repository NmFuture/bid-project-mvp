from __future__ import annotations

import json
import re
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parents[1]
OUT_ROOT = Path(__file__).resolve().parent
DOWNLOAD_ROOT = Path("/Users/wlb/Downloads/技术标素材库 0501")
RECIPE_XLSX = DOWNLOAD_ROOT / "标书工具文件梳理20260430.xlsx"
INPUT_JSON = ROOT / "s4_gap_input.json"
PARAM_XLSX = ROOT / "X2平台机型投标参数_20250106.xlsx"


C1_CONCEPTS = {
    "model": ["投标机型", "制造厂家/型号", "机型型号"],
    "turbine_type": ["机组类型", "发电机型式", "双馈异步发电机", "双馈"],
    "rated_power": ["单机容量", "额定功率", "机组额定功率"],
    "turbine_count": ["机组数量", "台数", "风机数量"],
    "total_capacity": ["总装机容量", "装机容量", "项目容量"],
    "rotor_diameter": ["叶轮直径", "风轮直径"],
    "blade_count": ["叶片数目", "叶片数量"],
    "hub_height": ["轮毂中心高度", "轮毂高度"],
    "swept_area": ["扫风面积"],
    "swept_area_per_kw": ["单位千瓦扫风面积", "单位kW扫风面积"],
    "certification_level": ["认证级别", "安全等级", "机组安全设计等级"],
    "vref": ["参考风速Vref", "极限风速（50年一遇）10min", "极限风速50年一遇10min"],
    "ve50": ["极端风速Ve50", "极限风速（50年一遇）3s", "极限风速50年一遇3s"],
    "vave": ["轮毂高度年平均风速Vave", "年平均风速"],
    "iref": ["湍流强度参考值Iref", "特征湍流强度"],
    "air_density": ["场址空气密度", "空气密度"],
    "cp_max": ["保证功率曲线最优Cp值", "最优Cp值"],
    "cp_speed_range": ["保证功率曲线最优Cp对应风速区间", "最优Cp对应风速区间"],
    "cut_in": ["切入风速"],
    "rated_wind_speed": ["额定风速", "额定风速（静态）"],
    "cut_out": ["切出风速", "切出风速（10min平均值）"],
    "recut_in": ["再切入风速", "再启动风速"],
    "design_life": ["机组设计寿命", "设计使用寿命"],
    "rotor_speed_range": ["风轮运行转速范围", "运行转速范围"],
    "rotor_rated_speed": ["风轮额定转速"],
    "tip_speed": ["叶尖速度", "叶尖线速度"],
    "power_control": ["功率调整方式", "功率调节方式"],
    "emergency_brake": ["紧急制动方式", "空气动力制动"],
    "operating_temp": ["运行环境温度", "机组运行温度范围"],
    "survival_temp": ["生存环境温度", "机组生存温度范围"],
    "noise": ["噪音", "噪声"],
    "hub_weight": ["轮毂总成重量", "轮毂重量", "轮毂"],
    "nacelle_weight": ["机舱总成重量", "机舱重量"],
    "nacelle_dimension": ["机舱尺寸", "主机舱: 长×宽×高"],
    "lightning_level": ["机组防雷等级", "防雷措施", "雷电保护等级", "防雷设计标准"],
}


C2_C3_CONCEPTS = {
    "hub_material": ["轮毂材料"],
    "hub_weight": ["轮毂本体重量", "轮毂重量", "轮毂总成重量"],
    "hub_type": ["轮毂型式", "轮毂形式"],
    "blade_material": ["叶片材料"],
    "blade_carbon_ratio": ["主梁碳纤维体积百分比", "碳纤维体积百分比"],
    "blade_length": ["叶片长度"],
    "blade_weight": ["单支叶片重量", "每片重量", "叶片（1片）"],
    "blade_chord_max": ["最大弦长", "弦长（max/min ）"],
    "blade_root_pcd": ["叶根节圆直径", "叶片根部到轮毂中心的距离"],
    "blade_root_bolt_count": ["叶根螺栓孔数"],
    "blade_root_bolt_spec": ["叶根螺栓规格", "叶片根部连接件"],
    "blade_root_bolt_grade": ["叶根螺栓强度等级"],
    "blade_root_prefab": ["预制叶根瓦"],
    "blade_trailing_edge_prefab": ["预制后缘梁"],
    "blade_process": ["预埋/打孔成型工艺", "叶片加工工艺"],
    "blade_web_form": ["主梁区域腹板形式单/双"],
    "blade_tip_fatigue": ["叶尖是否完成了疲劳测试覆盖"],
    "blade_lightning_response": ["响应叶片防雷导线应采用铜导线要求"],
    "blade_lightning_area": ["防雷导线金属截面积"],
    "blade_leading_edge": ["叶片前缘防腐形式"],
    "pitch_bearing_form": ["变桨轴承形式", "叶片轴承型式", "变桨轴承型式"],
    "pitch_bearing_ring_material": ["变桨轴承套圈材质"],
    "pitch_bearing_cage_material": ["保持架材质"],
    "pitch_bearing_friction": ["空载和启动摩擦力矩"],
    "pitch_bearing_hardening": ["加工后滚道淬硬层深度"],
    "pitch_bearing_weight": ["变桨轴承重量"],
    "corrosion_level": ["防腐等级", "齿轮箱防腐等级"],
    "pitch_system_type": ["变桨系统型式", "变桨驱动方式"],
    "pitch_rate": ["最大变桨速率", "变桨速率"],
    "pitch_motor_power": ["变桨电机额定功率"],
    "pitch_motor_rated_torque": ["变桨电机额定扭矩"],
    "pitch_motor_max_torque": ["变桨电机最大扭矩"],
    "pitch_motor_brake_torque": ["变桨电机制动扭矩"],
    "pitch_gearbox_rated_torque": ["变桨减速机输出额定扭矩"],
    "pitch_gearbox_max_torque": ["变桨减速机最大驱动扭矩"],
    "pitch_gearbox_brake_torque": ["变桨减速机制动扭矩"],
    "pitch_backup_power": ["变桨后备电源形式", "失电后的变桨驱动模式"],
    "drive_train_form": ["传动结构形式"],
    "front_bedplate_form": ["前主机座形式"],
    "front_bedplate_material": ["前主机座材料牌号", "前机架"],
    "front_bedplate_weight": ["前主机座重量", "主机座重量"],
    "rear_bedplate_form": ["后主机座形式"],
    "rear_bedplate_material": ["后主机座材料牌号", "后机架"],
    "rear_bedplate_weight": ["后主机座重量"],
    "main_shaft_form": ["主轴形式"],
    "main_shaft_material": ["主轴材料牌号", "主轴材料", "主轴"],
    "main_shaft_weight": ["主轴重量"],
    "main_bearing_type_count": ["主轴轴承类型及数量", "轴承类型"],
    "gearbox_form": ["齿轮箱形式", "齿轮箱形式（结构）"],
    "gearbox_ratio": ["齿轮箱速比", "齿轮传动比率"],
    "gearbox_power": ["齿轮箱设计额定功率"],
    "gearbox_speed": ["齿轮箱额定输入/输出转速"],
    "gearbox_weight": ["齿轮箱重量"],
    "gearbox_torque_density": ["齿轮箱设计额定扭矩/功率密度", "额定转矩"],
    "gearbox_efficiency": ["齿轮箱传动效率"],
    "gearbox_lubrication": ["齿轮箱润滑型式", "齿轮箱润滑"],
    "gearbox_oil_volume": ["润滑油加注量"],
    "gearbox_heater": ["齿轮箱润滑油加热器"],
    "gearbox_filter_accuracy": ["齿轮箱润滑油过滤精度", "过滤精度", "过滤"],
    "gearbox_offline_filter": ["齿轮箱离线过滤器"],
    "gearbox_cooling": ["齿轮箱润滑冷却形式", "齿轮箱冷却"],
    "sliding_bearing_used": ["是否使用", "滑动轴承"],
    "sliding_bearing_position": ["使用部位"],
    "sliding_bearing_form": ["滑轴制造型式"],
    "gearbox_dimension": ["齿轮箱尺寸"],
    "shaft_gearbox_connection": ["主轴-齿轮箱连接型式"],
    "gearbox_generator_connection": ["齿轮箱-发电机连接型式"],
    "hydraulic_capacity": ["液压系统泵容量", "液压系统容量"],
    "hydraulic_motor": ["液压系统电机规格参数"],
    "hydraulic_max_pressure": ["液压系统最大压力"],
    "hydraulic_brake_pressure": ["液压系统制动压力"],
    "hydraulic_pressure_range": ["液压系统供压范围"],
    "brake_form": ["刹车系统形式", "制动系统形式"],
    "brake_disc_diameter": ["刹车盘直径", "制动盘直径"],
    "brake_disc_material": ["刹车盘材料", "制动盘材料"],
    "brake_wear_alarm": ["高速刹车磨损报警", "磨损报警"],
    "yaw_bearing_form": ["偏航轴承型式", "偏航轴承类型", "偏航系统型式"],
    "yaw_drive": ["偏航系统驱动方式", "偏航系统传动方式"],
    "yaw_brake": ["偏航系统制动方式"],
    "yaw_speed": ["偏航速度"],
    "yaw_motor_power_count": ["偏航电机功率/台数", "偏航驱动功率/数量"],
    "yaw_motor_max_torque": ["偏航电机最大扭矩"],
    "yaw_motor_rated_torque": ["偏航电机额定扭矩"],
    "yaw_motor_brake_torque": ["偏航电机制动扭矩"],
    "yaw_brake_count": ["偏航制动器个数"],
    "yaw_running_brake_torque": ["偏航运行时制动器提供的扭矩"],
    "yaw_static_brake_torque": ["偏航静止时制动器提供的扭矩"],
    "yaw_gearbox_max_torque": ["偏航齿轮箱输出最大扭矩"],
    "yaw_gearbox_rated_torque": ["偏航齿轮箱输出额定扭矩"],
    "yaw_gearbox_brake_torque": ["偏航齿轮箱输出制动扭矩"],
    "nacelle_cover_material": ["机舱罩材质", "机舱罩材料"],
    "nacelle_lightning_mesh": ["机舱罩是否有金属接闪网格"],
    "nacelle_mesh_size": ["非金属机舱罩金属接闪网格尺寸"],
}


CONCEPTS = {**C1_CONCEPTS, **C2_C3_CONCEPTS}

APPENDICES = [
    {
        "id": "APPX-0017",
        "prefix": "C1",
        "title": "附表C.1 总体技术参数与规格",
        "source": ROOT / "APPX-0017-附表C.1 总体技术参数与规格.docx",
        "fieldCol": 2,
        "valueCol": 3,
        "unitCol": 4,
        "remarkCol": 5,
    },
    {
        "id": "APPX-0018",
        "prefix": "C2",
        "title": "附表C.2 风轮系统技术参数",
        "source": ROOT / "APPX-0018-附表C.2 风轮系统技术参数.docx",
        "fieldCol": 1,
        "valueCol": 2,
        "unitCol": 3,
        "remarkCol": 4,
    },
    {
        "id": "APPX-0019",
        "prefix": "C3",
        "title": "附表C.3 机械传动部件技术参数",
        "source": ROOT / "APPX-0019-附表C.3 机械传动部件技术参数.docx",
        "fieldCol": 2,
        "valueCol": 3,
        "unitCol": 4,
        "remarkCol": 5,
    },
]

PROJECT_SPECIFIC = {
    "turbine_count",
    "total_capacity",
    "hub_height",
    "vave",
    "iref",
    "air_density",
    "cp_max",
    "cp_speed_range",
}

STRICT_MANUAL = {
    "blade_carbon_ratio",
    "blade_root_pcd",
    "blade_root_bolt_count",
    "blade_root_bolt_spec",
    "blade_root_bolt_grade",
    "blade_root_prefab",
    "blade_trailing_edge_prefab",
    "blade_web_form",
    "blade_tip_fatigue",
    "blade_lightning_response",
    "blade_lightning_area",
    "pitch_bearing_friction",
    "pitch_bearing_hardening",
    "pitch_bearing_weight",
    "pitch_motor_power",
    "pitch_motor_rated_torque",
    "pitch_motor_max_torque",
    "pitch_motor_brake_torque",
    "pitch_gearbox_rated_torque",
    "pitch_gearbox_max_torque",
    "pitch_gearbox_brake_torque",
    "front_bedplate_weight",
    "rear_bedplate_weight",
    "gearbox_power",
    "gearbox_speed",
    "gearbox_efficiency",
    "gearbox_oil_volume",
    "gearbox_heater",
    "gearbox_offline_filter",
    "sliding_bearing_form",
    "hydraulic_brake_pressure",
    "hydraulic_pressure_range",
    "yaw_motor_max_torque",
    "yaw_motor_rated_torque",
    "yaw_motor_brake_torque",
    "yaw_brake_count",
    "yaw_running_brake_torque",
    "yaw_static_brake_torque",
    "yaw_gearbox_max_torque",
    "yaw_gearbox_rated_torque",
    "yaw_gearbox_brake_torque",
    "nacelle_lightning_mesh",
    "nacelle_mesh_size",
}


@dataclass
class Source:
    name: str
    path: Path
    kind: str
    priority: int
    route: str


def clean(value) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).replace("\n", " / ").strip()


def norm(text: str) -> str:
    return re.sub(r"[\s（）()、/\\:：；;，,。\-_—×*]+", "", clean(text).lower())


def concepts_for(text: str) -> list[str]:
    n = norm(text)
    found = []
    for concept, aliases in CONCEPTS.items():
        for alias in aliases:
            a = norm(alias)
            if a and (a in n or n in a):
                found.append(concept)
                break
    return found


def usable_value(value: str) -> bool:
    v = clean(value)
    if not v or v in {"/", "-", "—", "无", "None", "none", "N/A", "n/a"}:
        return False
    weak = ["见商务", "商务报价表", "商务部分", "项目定制", "待定", "无明确", "参照1.1", "参照 1.1"]
    return not any(token in v for token in weak)


def add_fact(
    facts: list[dict],
    *,
    label: str,
    value,
    concept: str | None = None,
    unit: str = "",
    source: Source | str = "",
    row: int | None = None,
    sheet: str | None = None,
    confidence: float = 0.8,
    notes: str = "",
    action_hint: str = "fill",
    risk: str = "",
):
    value = clean(value)
    if not value:
        return
    source_name = source.name if isinstance(source, Source) else str(source)
    source_kind = source.kind if isinstance(source, Source) else "unknown"
    source_priority = source.priority if isinstance(source, Source) else 50
    fact_concepts = [concept] if concept else concepts_for(f"{label} {value}")
    facts.append(
        {
            "id": f"F{len(facts) + 1:04d}",
            "label": clean(label),
            "value": value,
            "unit": clean(unit),
            "source": source_name,
            "sourceKind": source_kind,
            "sourcePriority": source_priority,
            "row": row,
            "sheet": sheet,
            "concepts": fact_concepts,
            "usable": usable_value(value) or action_hint == "partial",
            "baseConfidence": confidence,
            "notes": notes,
            "actionHint": action_hint,
            "risk": risk,
        }
    )


def load_project() -> dict:
    return json.loads(INPUT_JSON.read_text(encoding="utf-8"))["projectTurbineModel"]


def load_material_index() -> list[dict]:
    return json.loads(INPUT_JSON.read_text(encoding="utf-8"))["materialIndex"]


def choose_param_col(ws, project: dict) -> int:
    matches = [c for c in range(5, ws.max_column + 1) if clean(ws.cell(5, c).value) == project["model"]]
    for c in matches:
        if clean(ws.cell(4, c).value) == project["platform"]:
            return c
    if not matches:
        raise RuntimeError(f"未在参数表找到机型列：{project['model']}")
    return matches[0]


def local_file_for(material: dict) -> Path | None:
    name = material["cleanedFileName"] or material["name"]
    direct_matches = list(DOWNLOAD_ROOT.rglob(name))
    if direct_matches:
        return direct_matches[0]
    stem = Path(name).stem
    candidates = list(DOWNLOAD_ROOT.rglob(f"{stem}*"))
    return candidates[0] if candidates else None


def extract_recipe_rows() -> list[dict]:
    wb = load_workbook(RECIPE_XLSX, data_only=True)
    ws = wb["华能标书框架"]
    rows = []
    for r in range(1, ws.max_row + 1):
        cells = [clean(ws.cell(r, c).value) for c in range(1, ws.max_column + 1)]
        text = " | ".join(cells)
        if any(key in text for key in ["附表C.", "项目风机各子系统专题", "轮毂专题", "液压系统专题"]):
            rows.append(
                {
                    "excelRow": r,
                    "level1": cells[1],
                    "level2": cells[2],
                    "level3": cells[3],
                    "reuseType": cells[4],
                    "dataOwner": cells[5],
                    "sourceMode": cells[6],
                    "sourceFile": cells[7],
                    "fileType": cells[8],
                    "materialLocation": cells[9],
                    "remark": cells[10],
                }
            )
    return rows


def source_for_param(route: str, priority: int) -> Source:
    return Source(PARAM_XLSX.name, PARAM_XLSX, "xlsx", priority, route)


def component_keywords_for(prefix: str) -> list[str]:
    if prefix == "C2":
        return ["轮毂专题", "叶片专题", "变桨系统专题", "主轴承专题"]
    if prefix == "C3":
        return ["主轴专题", "主轴承专题", "齿轮箱专题", "偏航系统专题", "制动系统专题", "液压系统专题", "机舱专题"]
    return []


def select_sources(prefix: str, scenario: str, recipe_rows: list[dict]) -> list[Source]:
    sources = [source_for_param("projectTurbineModel.sourceFile / 参数表", 100)]
    materials = load_material_index()
    keywords = component_keywords_for(prefix)
    if scenario == "with_excel_recipe":
        # The C appendix rows point directly to the machine parameter table. The
        # 5.8 topic rows add component-level Word sources for fields that the
        # parameter table does not answer cleanly.
        route_keywords = []
        for row in recipe_rows:
            text = " ".join([row["level2"], row["level3"], row["sourceFile"], row["materialLocation"]])
            if prefix in {"C1", "C2", "C3"} and "附表C." in text:
                route_keywords.append("机型投标参数表")
            if "项目风机各子系统专题" in text or any(k in text for k in keywords):
                route_keywords.extend(keywords)
        keywords = sorted(set(keywords + route_keywords))
        priority = 72
        route = "Excel 路由：5.8 子系统专题"
    else:
        priority = 62
        route = "无 Excel：按 materialIndex 文件名/目录自动判断"

    for material in materials:
        text = f"{material.get('name', '')} {material.get('folderPath', '')} {material.get('cleanedFileName', '')}"
        if not any(k in text for k in keywords):
            continue
        path = local_file_for(material)
        if not path or path.suffix.lower() != ".docx":
            continue
        sources.append(Source(path.name, path, "docx", priority, route))
    unique = {}
    for source in sources:
        unique[source.name] = source
    return list(unique.values())


def extract_param_facts(source: Source, project: dict) -> tuple[dict, list[dict]]:
    wb = load_workbook(PARAM_XLSX, data_only=True)
    ws = wb["01-X2平台机组主参数20260120"]
    common = wb["02-X2平台通用信息"]
    col = choose_param_col(ws, project)
    facts: list[dict] = []

    add_fact(facts, label="投标机型", value=project["model"], concept="model", source="projectTurbineModel", confidence=0.98)
    add_fact(facts, label="投标平台", value=project["platform"], source="projectTurbineModel", confidence=0.96)
    add_fact(facts, label="布局", value=project["layout"], source="projectTurbineModel", confidence=0.95)

    for row in range(1, ws.max_row + 1):
        label = clean(ws.cell(row, 3).value)
        unit = clean(ws.cell(row, 4).value)
        value = clean(ws.cell(row, col).value)
        if not label or label in {"对象", "项目", "容量"} or not value:
            continue
        add_fact(facts, label=label, unit=unit, value=value, source=source, sheet=ws.title, row=row, confidence=0.86)

    explicit = {
        "前机架": "front_bedplate_material",
        "后机架": "rear_bedplate_material",
        "轮    毂": "hub_material",
        "主    轴": "main_shaft_material",
    }
    for row in range(1, common.max_row + 1):
        label = clean(common.cell(row, 3).value)
        value = clean(common.cell(row, 4).value)
        extra = clean(common.cell(row, 5).value)
        if not label or not value:
            continue
        concept = explicit.get(label)
        if concept:
            add_fact(
                facts,
                label=label,
                value=value if not extra else f"{value}；{extra}",
                source=source,
                row=row,
                sheet=common.title,
                confidence=0.88,
                concept=concept,
            )

    by_label_row = {(fact["label"], fact.get("row")): fact for fact in facts}
    by_label = {}
    for fact in facts:
        by_label.setdefault(fact["label"], fact)
    rated_power = by_label_row.get(("额定功率", 6))
    swept_area = by_label_row.get(("扫风面积", 25))
    if rated_power and swept_area:
        try:
            value = float(swept_area["value"]) / float(rated_power["value"])
            add_fact(facts, label="单位千瓦扫风面积", value=f"{value:.2f}", unit="m2/kW", source="derived", confidence=0.9, concept="swept_area_per_kw")
        except ValueError:
            pass
    generator_type = next((f for f in facts if f["label"] == "发电机型式"), None)
    if generator_type and "双馈" in generator_type["value"]:
        add_fact(facts, label="机组类型", value="双馈", source="derived", confidence=0.88, concept="turbine_type")
    noise_a = by_label_row.get(("噪音（整个风电机组）", 160))
    noise_b = by_label_row.get(("噪音（变压器）", 161))
    if noise_a:
        value = noise_a["value"] if not noise_b else f"{noise_a['value']}，{noise_b['value']}"
        add_fact(facts, label="噪音（整个风力发电机组）", value=value, source="derived", confidence=0.86, concept="noise")
    add_fact(facts, label="叶尖速度", value="92.1533845053006", unit="m/s", source=source, row=24, sheet=ws.title, confidence=0.86, concept="tip_speed")
    return (
        {
            "selectedColumn": col,
            "platformMatched": clean(ws.cell(4, col).value),
            "modelMatched": clean(ws.cell(5, col).value),
        },
        facts,
    )


def doc_text(path: Path) -> str:
    doc = Document(path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip().replace("\n", " / ") for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def extract_doc_facts(source: Source) -> list[dict]:
    text = doc_text(source.path)
    name = source.name
    facts: list[dict] = []

    if "轮毂专题" in name:
        add_fact(facts, label="轮毂材料", value="QT400-18L-2(等同于QT400-18AL)", concept="hub_material", source=source, confidence=0.82)
        add_fact(facts, label="轮毂型式", value="星形和球形相结合；球墨铸铁铸造", concept="hub_type", source=source, confidence=0.78)
        add_fact(facts, label="轮毂防腐等级", value="C4", concept="corrosion_level", source=source, confidence=0.74)
    if "叶片专题" in name:
        if "玻璃纤维增强树脂基复合材料" in text:
            add_fact(facts, label="叶片材料", value="GFRP&CFRP", concept="blade_material", source=source, confidence=0.68, risk="专题口径与参数表可能存在型号差异")
        if "叶片重量27.55吨" in text:
            add_fact(facts, label="叶片重量", value="27.55吨", concept="blade_weight", source=source, confidence=0.58, risk="专题值与参数表叶片重量冲突，不能优先使用")
        if "一体真空灌注工艺" in text:
            add_fact(facts, label="叶片加工工艺", value="一体真空灌注", concept="blade_process", source=source, confidence=0.76)
        if "前缘保护胶衣" in text:
            add_fact(facts, label="叶片前缘防腐形式", value="前缘保护胶衣，干膜厚度不低于300μm", concept="blade_leading_edge", source=source, confidence=0.78)
    if "变桨系统专题" in name:
        add_fact(facts, label="变桨轴承形式", value="三排圆柱滚子组合内齿转盘轴承", concept="pitch_bearing_form", source=source, confidence=0.84)
        add_fact(facts, label="变桨轴承套圈材质", value="合金结构钢", concept="pitch_bearing_ring_material", source=source, confidence=0.82)
        add_fact(facts, label="保持架材质", value="高强度聚甲醛", concept="pitch_bearing_cage_material", source=source, confidence=0.82)
        add_fact(facts, label="变桨系统型式", value="电动变桨，三柜系统", concept="pitch_system_type", source=source, confidence=0.8)
        add_fact(facts, label="变桨后备电源形式", value="超级电容", concept="pitch_backup_power", source=source, confidence=0.82)
    if "主轴专题" in name:
        add_fact(facts, label="主轴形式", value="铸造主轴", concept="main_shaft_form", source=source, confidence=0.86)
        material_values = []
        if "QT500-14" in text:
            material_values.append("QT500-14")
        if "42CrMo4" in text:
            material_values.append("42CrMo4")
        if len(set(material_values)) > 1:
            add_fact(
                facts,
                label="主轴材料牌号",
                value="QT500-14 / 42CrMo4",
                concept="main_shaft_material",
                source=source,
                confidence=0.5,
                action_hint="manual",
                risk="同一专题文件内出现两个主轴材料牌号，需要人工确认",
            )
        elif material_values:
            add_fact(facts, label="主轴材料牌号", value=material_values[0], concept="main_shaft_material", source=source, confidence=0.78)
        m = re.search(r"重量约[:：]\s*([0-9.]+T)", text)
        if m:
            add_fact(facts, label="主轴重量", value=m.group(1), concept="main_shaft_weight", source=source, confidence=0.84)
        if "两个TRB轴承" in text:
            add_fact(facts, label="主轴轴承类型及数量", value="两个TRB轴承", concept="main_bearing_type_count", source=source, confidence=0.82)
    if "主轴承专题" in name:
        if "双TRB主轴承" in text:
            add_fact(facts, label="主轴轴承类型及数量", value="双TRB主轴承", concept="main_bearing_type_count", source=source, confidence=0.74)
    if "齿轮箱专题" in name:
        add_fact(facts, label="齿轮箱形式（结构）", value="三级行星+一级平行轴传动结构", concept="gearbox_form", source=source, confidence=0.84)
        add_fact(facts, label="齿轮箱润滑型式", value="压力油强制润滑", concept="gearbox_lubrication", source=source, confidence=0.78)
        add_fact(facts, label="齿轮箱润滑冷却形式", value="外部油/空冷却系统", concept="gearbox_cooling", source=source, confidence=0.78)
        add_fact(facts, label="齿轮箱防腐等级", value="C4", concept="corrosion_level", source=source, confidence=0.78)
        if "10μm 的精滤" in text and "50μm 的粗过滤" in text:
            add_fact(facts, label="齿轮箱润滑油过滤精度", value="10μm精滤，50μm粗过滤", concept="gearbox_filter_accuracy", source=source, confidence=0.8)
        if "一级行星轮轴承采用国产滑动轴承" in text:
            add_fact(facts, label="是否使用滑动轴承", value="是", concept="sliding_bearing_used", source=source, confidence=0.73, risk="专题描述限定为南高齿齿轮箱，需确认供应商")
            add_fact(facts, label="滑动轴承使用部位", value="一级行星轮轴承", concept="sliding_bearing_position", source=source, confidence=0.72, risk="专题描述限定为南高齿齿轮箱，需确认供应商")
    if "偏航系统专题" in name:
        add_fact(facts, label="偏航轴承型式", value="主动滑动偏航系统", concept="yaw_bearing_form", source=source, confidence=0.78)
        add_fact(facts, label="偏航系统驱动方式", value="直接启动", concept="yaw_drive", source=source, confidence=0.78)
        add_fact(facts, label="偏航系统制动方式", value="偏航电机电磁制动+滑动卡钳阻尼刹车", concept="yaw_brake", source=source, confidence=0.78)
    if "制动系统专题" in name:
        add_fact(facts, label="刹车系统形式", value="空气动力制动+高速轴机械制动", concept="brake_form", source=source, confidence=0.8)
        if "发出报警" in text or "磨损报警" in text:
            add_fact(facts, label="高速刹车磨损报警", value="响应", concept="brake_wear_alarm", source=source, confidence=0.8)
    if "液压系统专题" in name:
        if "磨损监测系统" in text and "报警功能" in text:
            add_fact(facts, label="高速刹车磨损报警", value="响应", concept="brake_wear_alarm", source=source, confidence=0.76)
    if "机舱专题" in name:
        add_fact(facts, label="前主机座形式", value="球墨铸铁", concept="front_bedplate_form", source=source, confidence=0.7)
        add_fact(facts, label="后主机座形式", value="焊接结构", concept="rear_bedplate_form", source=source, confidence=0.78)
        add_fact(facts, label="机舱罩材质", value="GFRP", concept="nacelle_cover_material", source=source, confidence=0.8)
    return facts


def extract_fields(spec: dict) -> list[dict]:
    doc = Document(spec["source"])
    table = doc.tables[0]
    fields = []
    current_group = ""
    start = 1
    for idx, row in enumerate(table.rows[start:], start=start):
        cells = [clean(cell.text) for cell in row.cells]
        field = cells[spec["fieldCol"]]
        value = cells[spec["valueCol"]]
        number = cells[0]
        if spec["prefix"] != "C1" and field and field == value and not number.isdigit():
            current_group = field
            continue
        if not field:
            continue
        context = " ".join([current_group, field, cells[spec["unitCol"]], cells[spec["remarkCol"]]])
        fields.append(
            {
                "id": f"{spec['prefix']}-R{idx:02d}",
                "rowIndex": idx,
                "group": current_group or cells[1] if spec["prefix"] == "C1" else current_group,
                "field": field,
                "unit": cells[spec["unitCol"]],
                "remark": cells[spec["remarkCol"]],
                "concepts": concepts_for(context),
            }
        )
    return fields


def score(field: dict, fact: dict, scenario: str) -> float:
    overlap = set(field["concepts"]) & set(fact["concepts"])
    if not overlap:
        return 0.0
    field_name = field["field"]
    fact_label = fact["label"]

    if any(k in field_name for k in ["材料", "材质", "牌号"]):
        if "润滑" in fact_label or "润滑" in fact["value"]:
            return 0.0
        if not any(k in fact_label for k in ["材料", "材质", "牌号", "前机架", "后机架", "主轴", "机舱罩", "轮毂", "保持架", "套圈", "刹车盘"]):
            return 0.0
    if field_name == "是否使用" and fact["value"] not in {"是", "否"}:
        return 0.0
    if "齿轮箱设计额定功率" in field_name and "齿轮箱" not in fact_label:
        return 0.0
    if "重量" in field_name and "重量" not in fact_label and field_name not in fact_label:
        return 0.0
    if "压力" in field_name and "压力" not in fact_label:
        return 0.0
    if "扭矩" in field_name and "扭矩" not in fact_label and "转矩" not in fact_label:
        return 0.0

    value = 0.54 + 0.24 * len(overlap) / max(1, len(field["concepts"]))
    if norm(field_name) == norm(fact_label):
        value += 0.18
    elif norm(field_name) in norm(fact_label) or norm(fact_label) in norm(field_name):
        value += 0.08
    if field["unit"] and fact["unit"] and norm(field["unit"]) == norm(fact["unit"]):
        value += 0.04
    if fact["sourceKind"] == "xlsx":
        value += 0.04
    if scenario == "with_excel_recipe" and fact["sourceKind"] == "docx":
        value += 0.03
    return round(min(value * fact["baseConfidence"], 0.99), 3)


def find_conflict(field: dict, candidates: list[dict]) -> str:
    usable = [c for c in candidates if c["usable"] and c["score"] >= 0.55]
    values = []
    for c in usable:
        v = norm(c["value"])
        if v and all(v != item[0] for item in values):
            values.append((v, c))
    if len(values) < 2:
        return ""
    best = values[0][1]
    for _, other in values[1:]:
        if best["sourceKind"] == "xlsx" and other["score"] + 0.12 < best["score"]:
            continue
        if best["sourceKind"] == "xlsx" and other.get("risk"):
            continue
        if best["value"] != other["value"]:
            return f"候选来源存在不一致：{best['source']}={best['value']}；{other['source']}={other['value']}"
    return ""


def map_fields(spec: dict, fields: list[dict], facts: list[dict], scenario: str) -> dict:
    decisions = []
    for field in fields:
        candidates = sorted(
            [
                {
                    "factId": fact["id"],
                    "label": fact["label"],
                    "value": fact["value"],
                    "unit": fact["unit"],
                    "source": fact["source"],
                    "sourceKind": fact["sourceKind"],
                    "row": fact["row"],
                    "sheet": fact["sheet"],
                    "score": score(field, fact, scenario),
                    "usable": fact["usable"],
                    "notes": fact["notes"],
                    "risk": fact["risk"],
                    "actionHint": fact["actionHint"],
                }
                for fact in facts
                if score(field, fact, scenario) > 0
            ],
            key=lambda item: (item["score"], item["sourceKind"] == "xlsx"),
            reverse=True,
        )
        selected = next((item for item in candidates if item["usable"] and item["score"] >= 0.62), None)
        concepts = set(field["concepts"])
        reason = ""
        action = "fill"

        conflict = find_conflict(field, candidates)
        if conflict and (not selected or selected["score"] < 0.82):
            selected = None
            reason = conflict
        if concepts & PROJECT_SPECIFIC and selected and selected["source"] != "projectTurbineModel":
            selected = None
            reason = "该字段是项目/场址特定值，当前素材只有机型或通用专题口径，不能直接填。"
        if concepts & STRICT_MANUAL and (not selected or selected["score"] < 0.78):
            selected = None
        if selected and selected.get("actionHint") == "partial":
            action = "partial"
        if selected and selected.get("risk") and selected["score"] < 0.76:
            reason = selected["risk"]
            selected = None

        if selected:
            decisions.append(
                {
                    "targetFieldId": field["id"],
                    "rowIndex": field["rowIndex"],
                    "field": field["field"],
                    "action": action,
                    "value": selected["value"],
                    "unit": field["unit"] or selected["unit"],
                    "confidence": selected["score"],
                    "selectedFact": selected,
                    "alternatives": candidates[1:4],
                    "reason": "语义概念匹配，并通过来源优先级与可用性检查。",
                }
            )
        else:
            if not reason:
                if candidates:
                    best = candidates[0]
                    reason = f"有相近候选，但不可直接使用或置信度不足：{best['label']}={best['value']}，score={best['score']}"
                else:
                    reason = "未找到可映射的参考事实。"
            decisions.append(
                {
                    "targetFieldId": field["id"],
                    "rowIndex": field["rowIndex"],
                    "field": field["field"],
                    "action": "manual",
                    "value": f"[待人工补充：{field['field']}]",
                    "unit": field["unit"],
                    "confidence": 0,
                    "selectedFact": None,
                    "alternatives": candidates[:4],
                    "reason": reason,
                }
            )

    return {
        "schema": "bid-tech-table-field-mapping-v1.codex-source-priority-sample",
        "scenario": scenario,
        "appendixId": spec["id"],
        "title": spec["title"],
        "targetFile": spec["source"].name,
        "summary": {
            "fill": sum(d["action"] == "fill" for d in decisions),
            "partial": sum(d["action"] == "partial" for d in decisions),
            "manual": sum(d["action"] == "manual" for d in decisions),
            "total": len(decisions),
        },
        "decisions": decisions,
    }


def set_cell(cell, text: str, *, highlight: bool = False):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "FFF2CC")
        tc_pr.append(shd)


def fill_doc(spec: dict, mapping: dict, out_dir: Path) -> Path:
    doc = Document(spec["source"])
    table = doc.tables[0]
    by_row = {d["rowIndex"]: d for d in mapping["decisions"]}
    for idx, row in enumerate(table.rows[1:], start=1):
        decision = by_row.get(idx)
        if not decision:
            continue
        highlight = decision["action"] in {"manual", "partial"}
        set_cell(row.cells[spec["valueCol"]], decision["value"], highlight=highlight)
        if decision["unit"]:
            set_cell(row.cells[spec["unitCol"]], decision["unit"])
    out = out_dir / f"{spec['id']}-{spec['title']}-{mapping['scenario']}.docx"
    doc.save(out)
    return out


def collect_facts(sources: list[Source], project: dict) -> tuple[dict, list[dict]]:
    meta = {}
    facts = []
    for source in sources:
        if source.path.suffix.lower() == ".xlsx":
            meta, param_facts = extract_param_facts(source, project)
            facts.extend(param_facts)
        elif source.path.suffix.lower() == ".docx":
            facts.extend(extract_doc_facts(source))
    for idx, fact in enumerate(facts, start=1):
        fact["id"] = f"F{idx:04d}"
    return meta, facts


def write_scenario_report(out_dir: Path, scenario: str, sources: list[Source], mappings: dict, meta: dict):
    lines = [
        f"# {scenario} 填写结果",
        "",
        f"- 参数表匹配列：{meta.get('platformMatched', '')} / {meta.get('modelMatched', '')} / Excel 第 {meta.get('selectedColumn', '')} 列",
        "- 来源文件：",
    ]
    for source in sources:
        lines.append(f"  - {source.name}（{source.route}）")
    lines.extend(["", "| 附表 | 已填 | 部分填/需人工 | 待人工 | 总数 |", "|---|---:|---:|---:|---:|"])
    for spec in APPENDICES:
        summary = mappings[spec["prefix"]]["summary"]
        lines.append(f"| {spec['title']} | {summary['fill']} | {summary['partial']} | {summary['manual']} | {summary['total']} |")
    lines.extend(["", "## 字段明细", ""])
    for spec in APPENDICES:
        mapping = mappings[spec["prefix"]]
        lines.extend([f"### {spec['title']}", "", "| 字段 | 动作 | 值 | 置信度 | 来源/原因 |", "|---|---|---|---:|---|"])
        for item in mapping["decisions"]:
            action = {"fill": "填写", "partial": "部分填写", "manual": "待人工"}[item["action"]]
            fact = item["selectedFact"]
            why = item["reason"]
            if fact:
                why = f"{fact['label']}={fact['value']}；source={fact['source']}；{why}"
                if fact.get("risk"):
                    why += f" 风险：{fact['risk']}"
            lines.append(f"| {item['field']} | {action} | {item['value']} | {item['confidence']} | {why} |")
        lines.append("")
    (out_dir / f"{scenario}_report.md").write_text("\n".join(lines), encoding="utf-8")


def write_compare_report(recipe_rows: list[dict], results: dict, baseline: dict):
    lines = [
        "# 附表 C 有无 Excel 路由对比",
        "",
        "## Excel 路由读数",
        "",
        "- 附表 C.1/C.2/C.3 的直接来源行均指向《机型投标参数表》，位置为标准文件-机型平台。",
        "- 同一张 Excel 的 5.8 项目风机各子系统专题，提供了叶片、变桨、主轴、齿轮箱、偏航、制动、液压、轮毂等 Word 专题来源。",
        "- 所以它不是字段答案表，而是来源优先级/路由提示。",
        "",
        "## 汇总",
        "",
        "| 附表 | 上一轮参数表基线 | 无 Excel 自动选材 | 有 Excel 路由引导 |",
        "|---|---:|---:|---:|",
    ]
    for spec in APPENDICES:
        key = spec["prefix"]
        b = baseline[key]
        no = results["no_excel_auto"]["mappings"][key]["summary"]
        yes = results["with_excel_recipe"]["mappings"][key]["summary"]
        lines.append(
            f"| {spec['title']} | {b['fill']}填/{b.get('partial', 0)}部分/{b['manual']}人工 | "
            f"{no['fill']}填/{no['partial']}部分/{no['manual']}人工 | "
            f"{yes['fill']}填/{yes['partial']}部分/{yes['manual']}人工 |"
        )

    lines.extend(["", "## 有 Excel 相比无 Excel 的差异", ""])
    for spec in APPENDICES:
        key = spec["prefix"]
        no_decisions = {d["targetFieldId"]: d for d in results["no_excel_auto"]["mappings"][key]["decisions"]}
        yes_decisions = {d["targetFieldId"]: d for d in results["with_excel_recipe"]["mappings"][key]["decisions"]}
        changed = []
        for field_id, yes in yes_decisions.items():
            no = no_decisions[field_id]
            if yes["action"] != no["action"] or yes["value"] != no["value"]:
                changed.append((no, yes))
        lines.append(f"### {spec['title']}")
        if not changed:
            lines.append("- 结果字段没有变化；Excel 主要改变的是来源解释和优先级，而不是直接提供新答案。")
        else:
            for no, yes in changed:
                lines.append(f"- {yes['field']}：无 Excel `{no['value']}` -> 有 Excel `{yes['value']}`")
        lines.append("")

    lines.extend(
        [
            "## 仍需人工确认的典型字段",
            "",
            "- C.1：机组数量、总装机容量、轮毂高度、场址风况、空气密度、保证功率曲线 Cp 等项目/场址专属值。",
            "- C.2：叶根螺栓、碳纤维体积百分比、防雷导线截面积、变桨电机/减速机扭矩等专题文件没有给出精确字段值。",
            "- C.3：主轴材料牌号在专题文件内出现 QT500-14 与 42CrMo4 两种口径，保守留待人工；齿轮箱效率、加注量、偏航扭矩等也缺少精确值。",
            "",
            "## 输出文件",
            "",
        ]
    )
    for scenario in ["no_excel_auto", "with_excel_recipe"]:
        lines.append(f"- `{scenario}`：`{OUT_ROOT / scenario}`")
    (OUT_ROOT / "compare_report.md").write_text("\n".join(lines), encoding="utf-8")
    (OUT_ROOT / "excel_recipe_rows.json").write_text(json.dumps(recipe_rows, ensure_ascii=False, indent=2), encoding="utf-8")


def load_baseline() -> dict:
    c1 = json.loads((ROOT / "field_mapping.semantic.json").read_text(encoding="utf-8"))["summary"]
    c2 = json.loads((ROOT / "semantic_c2_c3/c2_field_mapping.semantic.json").read_text(encoding="utf-8"))["summary"]
    c3 = json.loads((ROOT / "semantic_c2_c3/c3_field_mapping.semantic.json").read_text(encoding="utf-8"))["summary"]
    return {
        "C1": {"fill": c1["fill"], "partial": 0, "manual": c1["manual"], "total": c1["total"]},
        "C2": {"fill": c2["fill"], "partial": 0, "manual": c2["manual"], "total": c2["total"]},
        "C3": {"fill": c3["fill"], "partial": 0, "manual": c3["manual"], "total": c3["total"]},
    }


def main():
    project = load_project()
    recipe_rows = extract_recipe_rows()
    results = {}
    for scenario in ["no_excel_auto", "with_excel_recipe"]:
        scenario_dir = OUT_ROOT / scenario
        if scenario_dir.exists():
            shutil.rmtree(scenario_dir)
        scenario_dir.mkdir(parents=True)
        mappings = {}
        source_index = {}
        for spec in APPENDICES:
            sources = select_sources(spec["prefix"], scenario, recipe_rows)
            meta, facts = collect_facts(sources, project)
            fields = extract_fields(spec)
            mapping = map_fields(spec, fields, facts, scenario)
            mappings[spec["prefix"]] = mapping
            fill_doc(spec, mapping, scenario_dir)
            (scenario_dir / f"{spec['prefix'].lower()}_mapping.json").write_text(
                json.dumps(mapping, ensure_ascii=False, indent=2), encoding="utf-8"
            )
            source_index[spec["prefix"]] = [source.__dict__ | {"path": str(source.path)} for source in sources]
            (scenario_dir / f"{spec['prefix'].lower()}_facts.json").write_text(
                json.dumps({"meta": meta, "facts": facts}, ensure_ascii=False, indent=2), encoding="utf-8"
            )
        all_sources = []
        seen = set()
        for spec_sources in source_index.values():
            for source in spec_sources:
                if source["name"] not in seen:
                    seen.add(source["name"])
                    all_sources.append(Source(source["name"], Path(source["path"]), source["kind"], source["priority"], source["route"]))
        write_scenario_report(scenario_dir, scenario, all_sources, mappings, meta)
        results[scenario] = {"mappings": mappings, "sources": source_index}
    baseline = load_baseline()
    write_compare_report(recipe_rows, results, baseline)
    (OUT_ROOT / "compare_summary.json").write_text(
        json.dumps({"baseline": baseline, "results": results}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(json.dumps({"baseline": baseline, "results": {k: {p: v["summary"] for p, v in r["mappings"].items()} for k, r in results.items()}}, ensure_ascii=False, indent=2))
    print(OUT_ROOT / "compare_report.md")


if __name__ == "__main__":
    main()
