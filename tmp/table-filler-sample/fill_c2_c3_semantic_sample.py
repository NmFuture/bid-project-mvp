from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parent
PARAM_XLSX = ROOT / "X2平台机型投标参数_20250106.xlsx"
INPUT_JSON = ROOT / "s4_gap_input.json"
OUT_DIR = ROOT / "semantic_c2_c3"


CONCEPTS = {
    "hub_material": ["轮毂材料", "轮毂"],
    "hub_weight": ["轮毂本体重量", "轮毂重量", "轮毂"],
    "hub_type": ["轮毂型式", "轮毂形式"],
    "blade_material": ["叶片材料"],
    "blade_carbon_ratio": ["主梁碳纤维体积百分比", "碳纤维体积百分比"],
    "blade_length": ["叶片长度"],
    "blade_weight": ["单支叶片重量", "每片重量", "叶片（1片）"],
    "blade_chord_max": ["最大弦长", "弦长（max/min ）"],
    "blade_root_pcd": ["叶根节圆直径", "叶片根部到轮毂中心的距离"],
    "blade_root_bolt_count": ["叶根螺栓孔数"],
    "blade_root_bolt_spec": ["叶根螺栓规格", "叶片根部连接件"],
    "blade_root_bolt_grade": ["叶根螺栓强度等级"],
    "blade_root_prefab": ["预制叶根瓦"],
    "blade_trailing_edge_prefab": ["预制后缘梁"],
    "blade_process": ["预埋/打孔成型工艺", "叶片加工工艺"],
    "blade_web_form": ["主梁区域腹板形式单/双"],
    "blade_tip_fatigue": ["叶尖是否完成了疲劳测试覆盖"],
    "blade_lightning_response": ["响应叶片防雷导线应采用铜导线要求"],
    "blade_lightning_area": ["防雷导线金属截面积"],
    "blade_leading_edge": ["叶片前缘防腐形式"],
    "pitch_bearing_form": ["变桨轴承形式", "叶片轴承型式", "变桨轴承型式"],
    "pitch_bearing_ring_material": ["变桨轴承套圈材质"],
    "pitch_bearing_cage_material": ["保持架材质"],
    "pitch_bearing_friction": ["空载和启动摩擦力矩"],
    "pitch_bearing_hardening": ["加工后滚道淬硬层深度"],
    "pitch_bearing_weight": ["变桨轴承重量"],
    "corrosion_level": ["防腐等级", "齿轮箱防腐等级"],
    "pitch_system_type": ["变桨系统型式", "变桨驱动方式"],
    "pitch_rate": ["最大变桨速率", "变桨速率"],
    "pitch_motor_power": ["变桨电机额定功率"],
    "pitch_motor_rated_torque": ["变桨电机额定扭矩"],
    "pitch_motor_max_torque": ["变桨电机最大扭矩"],
    "pitch_motor_brake_torque": ["变桨电机制动扭矩"],
    "pitch_gearbox_rated_torque": ["变桨减速机输出额定扭矩"],
    "pitch_gearbox_max_torque": ["变桨减速机最大驱动扭矩"],
    "pitch_gearbox_brake_torque": ["变桨减速机制动扭矩"],
    "pitch_backup_power": ["变桨后备电源形式", "失电后的变桨驱动模式"],
    "drive_train_form": ["传动结构形式"],
    "front_bedplate_form": ["前主机座形式"],
    "front_bedplate_material": ["前主机座材料牌号", "前机架"],
    "front_bedplate_weight": ["前主机座重量", "主机座重量"],
    "rear_bedplate_form": ["后主机座形式"],
    "rear_bedplate_material": ["后主机座材料牌号", "后机架"],
    "rear_bedplate_weight": ["后主机座重量"],
    "main_shaft_form": ["主轴形式"],
    "main_shaft_material": ["主轴材料牌号", "主轴"],
    "main_shaft_weight": ["主轴重量"],
    "main_bearing_type_count": ["主轴轴承类型及数量", "轴承类型"],
    "gearbox_form": ["齿轮箱形式", "齿轮箱形式（结构）"],
    "gearbox_ratio": ["齿轮箱速比", "齿轮传动比率"],
    "gearbox_power": ["齿轮箱设计额定功率"],
    "gearbox_speed": ["齿轮箱额定输入/输出转速"],
    "gearbox_weight": ["齿轮箱重量"],
    "gearbox_torque_density": ["齿轮箱设计额定扭矩/功率密度", "额定转矩"],
    "gearbox_efficiency": ["齿轮箱传动效率"],
    "gearbox_lubrication": ["齿轮箱润滑型式", "齿轮箱润滑"],
    "gearbox_oil_volume": ["润滑油加注量"],
    "gearbox_heater": ["齿轮箱润滑油加热器"],
    "gearbox_filter_accuracy": ["齿轮箱润滑油过滤精度"],
    "gearbox_offline_filter": ["齿轮箱离线过滤器"],
    "gearbox_cooling": ["齿轮箱润滑冷却形式", "齿轮箱冷却"],
    "sliding_bearing_used": ["是否使用", "滑动轴承"],
    "sliding_bearing_position": ["使用部位"],
    "sliding_bearing_form": ["滑轴制造型式"],
    "gearbox_dimension": ["齿轮箱尺寸"],
    "shaft_gearbox_connection": ["主轴-齿轮箱连接型式"],
    "gearbox_generator_connection": ["齿轮箱-发电机连接型式"],
    "hydraulic_capacity": ["液压系统泵容量", "液压系统容量"],
    "hydraulic_motor": ["液压系统电机规格参数"],
    "hydraulic_max_pressure": ["液压系统最大压力"],
    "hydraulic_brake_pressure": ["液压系统制动压力"],
    "hydraulic_pressure_range": ["液压系统供压范围"],
    "brake_form": ["刹车系统形式", "制动系统形式"],
    "brake_disc_diameter": ["刹车盘直径", "制动盘直径"],
    "brake_disc_material": ["刹车盘材料", "制动盘材料"],
    "brake_wear_alarm": ["高速刹车磨损报警"],
    "yaw_bearing_form": ["偏航轴承型式", "偏航轴承类型", "偏航系统型式"],
    "yaw_drive": ["偏航系统驱动方式", "偏航系统传动方式"],
    "yaw_brake": ["偏航系统制动方式"],
    "yaw_speed": ["偏航速度"],
    "yaw_motor_power_count": ["偏航电机功率/台数", "偏航驱动功率/数量"],
    "yaw_motor_max_torque": ["偏航电机最大扭矩"],
    "yaw_motor_rated_torque": ["偏航电机额定扭矩"],
    "yaw_motor_brake_torque": ["偏航电机制动扭矩"],
    "yaw_brake_count": ["偏航制动器个数"],
    "yaw_running_brake_torque": ["偏航运行时制动器提供的扭矩"],
    "yaw_static_brake_torque": ["偏航静止时制动器提供的扭矩"],
    "yaw_gearbox_max_torque": ["偏航齿轮箱输出最大扭矩"],
    "yaw_gearbox_rated_torque": ["偏航齿轮箱输出额定扭矩"],
    "yaw_gearbox_brake_torque": ["偏航齿轮箱输出制动扭矩"],
    "nacelle_cover_material": ["机舱罩材质", "机舱罩材料"],
    "nacelle_lightning_mesh": ["机舱罩是否有金属接闪网格"],
    "nacelle_mesh_size": ["非金属机舱罩金属接闪网格尺寸"],
}

MANUAL_CONCEPTS = {
    "blade_carbon_ratio", "blade_root_pcd", "blade_root_bolt_count",
    "blade_root_bolt_spec", "blade_root_bolt_grade", "blade_root_prefab",
    "blade_trailing_edge_prefab", "blade_web_form", "blade_tip_fatigue",
    "blade_lightning_response", "blade_lightning_area", "blade_leading_edge",
    "pitch_bearing_ring_material", "pitch_bearing_cage_material",
    "pitch_bearing_friction", "pitch_bearing_hardening", "pitch_bearing_weight",
    "pitch_motor_power", "pitch_motor_rated_torque", "pitch_motor_max_torque",
    "pitch_motor_brake_torque", "pitch_gearbox_rated_torque",
    "pitch_gearbox_max_torque", "pitch_gearbox_brake_torque",
    "front_bedplate_form", "front_bedplate_weight", "rear_bedplate_form",
    "rear_bedplate_weight", "main_shaft_weight", "gearbox_speed",
    "gearbox_efficiency", "gearbox_oil_volume", "gearbox_heater",
    "gearbox_filter_accuracy", "gearbox_offline_filter", "sliding_bearing_position",
    "sliding_bearing_form", "hydraulic_brake_pressure", "hydraulic_pressure_range",
    "brake_wear_alarm", "yaw_motor_max_torque", "yaw_motor_rated_torque",
    "yaw_motor_brake_torque", "yaw_brake_count", "yaw_running_brake_torque",
    "yaw_static_brake_torque", "yaw_gearbox_max_torque", "yaw_gearbox_rated_torque",
    "yaw_gearbox_brake_torque", "nacelle_lightning_mesh", "nacelle_mesh_size",
}

KNOWN_INFERENCES = {
    "blade_lightning_response": ("响应", "根据投标响应口径推断，需人工复核"),
}

APPENDICES = [
    {
        "id": "APPX-0018",
        "title": "附表C.2 风轮系统技术参数",
        "source": ROOT / "APPX-0018-附表C.2 风轮系统技术参数.docx",
        "fieldCol": 1,
        "valueCol": 2,
        "unitCol": 3,
        "remarkCol": 4,
        "prefix": "C2",
    },
    {
        "id": "APPX-0019",
        "title": "附表C.3 机械传动部件技术参数",
        "source": ROOT / "APPX-0019-附表C.3 机械传动部件技术参数.docx",
        "fieldCol": 2,
        "valueCol": 3,
        "unitCol": 4,
        "remarkCol": 5,
        "prefix": "C3",
    },
]


def clean(value):
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).replace("\n", " / ").strip()


def norm(text):
    return re.sub(r"[\s（）()、/\\:：；;，,。\-_—×*]+", "", clean(text).lower())


def concepts_for(text):
    n = norm(text)
    found = []
    for concept, aliases in CONCEPTS.items():
        for alias in aliases:
            a = norm(alias)
            if a and (a in n or n in a):
                found.append(concept)
                break
    return found


def usable_value(value):
    v = clean(value)
    if not v or v in {"/", "-", "—", "无", "None", "none", "N/A", "n/a"}:
        return False
    weak = ["见商务", "项目定制", "参照1.1", "待定", "无明确", "商务报价表", "商务部分", "随"]
    return not any(token in v for token in weak)


def set_cell(cell, text, *, highlight=False):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "FFF2CC")
        tc_pr.append(shd)


def load_project():
    return json.loads(INPUT_JSON.read_text(encoding="utf-8"))["projectTurbineModel"]


def choose_col(ws, project):
    matches = [c for c in range(5, ws.max_column + 1) if clean(ws.cell(5, c).value) == project["model"]]
    for c in matches:
        if clean(ws.cell(4, c).value) == project["platform"]:
            return c
    return matches[0]


def add_fact(facts, label, value, unit="", source="", row=None, sheet=None, confidence=0.82, derivation=None):
    value = clean(value)
    if not value:
        return
    facts.append({
        "id": f"F{len(facts)+1:04d}",
        "label": clean(label),
        "value": value,
        "unit": clean(unit),
        "source": source,
        "row": row,
        "sheet": sheet,
        "concepts": concepts_for(f"{label} {value}"),
        "usable": usable_value(value) if derivation is None else True,
        "baseConfidence": confidence,
        "derivation": derivation,
    })


def extract_facts(project):
    wb = load_workbook(PARAM_XLSX, data_only=True)
    ws = wb["01-X2平台机组主参数20260120"]
    common = wb["02-X2平台通用信息"]
    col = choose_col(ws, project)
    facts = []

    for row in range(1, ws.max_row + 1):
        label = clean(ws.cell(row, 3).value)
        value = clean(ws.cell(row, col).value)
        unit = clean(ws.cell(row, 4).value)
        if label and label not in {"对象", "项目", "容量"} and value:
            add_fact(facts, label, value, unit, PARAM_XLSX.name, row, ws.title)

    # Generic material table supplies a few material names that are not in the
    # model-specific column.
    for row in range(1, common.max_row + 1):
        label = clean(common.cell(row, 2).value)
        value = clean(common.cell(row, 3).value)
        extra = clean(common.cell(row, 4).value)
        if label and value and any(k in label for k in ["前机架", "后机架", "轮", "主", "轴", "偏航"]):
            add_fact(facts, label, value if not extra else f"{value}；{extra}", "", PARAM_XLSX.name, row, common.title)

    for concept, (value, why) in KNOWN_INFERENCES.items():
        add_fact(facts, next(iter(CONCEPTS[concept])), value, "", "derived", None, None, 0.72, why)

    payload = {
        "selectedMaterial": PARAM_XLSX.name,
        "selectedSheet": ws.title,
        "selectedColumn": col,
        "platformMatched": clean(ws.cell(4, col).value),
        "modelMatched": clean(ws.cell(5, col).value),
        "facts": facts,
    }
    OUT_DIR.mkdir(exist_ok=True)
    (OUT_DIR / "reference_facts.c2_c3.semantic.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return payload


def is_header_row(cells, spec):
    field = cells[spec["fieldCol"]]
    value = cells[spec["valueCol"]]
    number = cells[0]
    return bool(field and field == value and (not number.isdigit()))


def extract_fields(spec):
    doc = Document(spec["source"])
    fields = []
    current_group = ""
    for idx, row in enumerate(doc.tables[0].rows[1:], start=1):
        cells = [clean(cell.text) for cell in row.cells]
        if is_header_row(cells, spec):
            current_group = cells[spec["fieldCol"]]
            continue
        field = cells[spec["fieldCol"]]
        if not field:
            continue
        field_id = f"{spec['prefix']}-R{idx:02d}"
        context = " ".join([current_group, field, cells[spec["unitCol"]], cells[spec["remarkCol"]]])
        fields.append({
            "id": field_id,
            "rowIndex": idx,
            "group": current_group,
            "field": field,
            "unit": cells[spec["unitCol"]],
            "remark": cells[spec["remarkCol"]],
            "concepts": concepts_for(context),
        })
    (OUT_DIR / f"{spec['prefix'].lower()}_target_fields.semantic.json").write_text(
        json.dumps(fields, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return fields


def score(field, fact):
    overlap = set(field["concepts"]) & set(fact["concepts"])
    if not overlap:
        return 0.0
    field_name = field["field"]
    fact_label = fact["label"]
    if any(k in field_name for k in ["材料", "材质", "牌号"]):
        if "润滑" in fact_label or "润滑" in fact["value"]:
            return 0.0
        if "形式" in fact_label or "型式" in fact_label:
            return 0.0
        if not any(k in fact_label for k in ["材料", "材质", "牌号", "前机架", "后机架", "轮毂", "主轴", "刹车盘", "机舱罩"]):
            return 0.0
    if field_name == "是否使用" and fact["value"] not in {"是", "否"}:
        return 0.0
    if "齿轮箱设计额定功率" in field_name and "齿轮箱" not in fact_label:
        return 0.0
    if "重量" in field_name and "重量" not in fact_label and field_name not in fact_label:
        return 0.0
    if "压力" in field_name and "压力" not in fact_label:
        return 0.0
    if "扭矩" in field_name and "扭矩" not in fact_label and "转矩" not in fact_label:
        return 0.0
    value = 0.55 + 0.24 * len(overlap) / max(1, len(field["concepts"]))
    if norm(field["field"]) == norm(fact["label"]):
        value += 0.2
    elif norm(field["field"]) in norm(fact["label"]) or norm(fact["label"]) in norm(field["field"]):
        value += 0.08
    if field["unit"] and fact["unit"] and norm(field["unit"]) == norm(fact["unit"]):
        value += 0.04
    if fact["source"] == "derived":
        value -= 0.03
    return round(min(value * fact["baseConfidence"], 0.99), 3)


def map_fields(spec, fields, facts_payload):
    decisions = []
    facts = facts_payload["facts"]
    for field in fields:
        candidates = sorted(
            [
                {
                    "factId": fact["id"],
                    "label": fact["label"],
                    "value": fact["value"],
                    "unit": fact["unit"],
                    "source": fact["source"],
                    "row": fact["row"],
                    "score": score(field, fact),
                    "usable": fact["usable"],
                    "derivation": fact["derivation"],
                }
                for fact in facts
                if score(field, fact) > 0
            ],
            key=lambda item: item["score"],
            reverse=True,
        )
        selected = next((c for c in candidates if c["usable"] and c["score"] >= 0.62), None)
        concepts = set(field["concepts"])
        if concepts & MANUAL_CONCEPTS and (not selected or selected["score"] < 0.8):
            selected = None

        if selected:
            decisions.append({
                "targetFieldId": field["id"],
                "rowIndex": field["rowIndex"],
                "field": field["field"],
                "action": "fill",
                "value": selected["value"],
                "unit": field["unit"] or selected["unit"],
                "confidence": selected["score"],
                "selectedFact": selected,
                "alternatives": candidates[1:4],
                "reason": "语义概念匹配并通过可用性检查。",
            })
        else:
            reason = "未找到可映射的参考事实。"
            if candidates:
                best = candidates[0]
                reason = f"有相近候选，但不可直接使用或置信度不足：{best['label']}={best['value']}，score={best['score']}"
            decisions.append({
                "targetFieldId": field["id"],
                "rowIndex": field["rowIndex"],
                "field": field["field"],
                "action": "manual",
                "value": f"[待人工补充：{field['field']}]",
                "unit": field["unit"],
                "confidence": 0,
                "selectedFact": None,
                "alternatives": candidates[:4],
                "reason": reason,
            })

    payload = {
        "schema": "bid-tech-table-field-mapping-v1.codex-semantic-sample",
        "appendixId": spec["id"],
        "title": spec["title"],
        "targetFile": spec["source"].name,
        "referenceMaterial": PARAM_XLSX.name,
        "summary": {
            "fill": sum(d["action"] == "fill" for d in decisions),
            "manual": sum(d["action"] == "manual" for d in decisions),
            "total": len(decisions),
        },
        "decisions": decisions,
    }
    (OUT_DIR / f"{spec['prefix'].lower()}_field_mapping.semantic.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return payload


def fill_doc(spec, mapping):
    doc = Document(spec["source"])
    table = doc.tables[0]
    by_row = {d["rowIndex"]: d for d in mapping["decisions"]}
    for idx, row in enumerate(table.rows[1:], start=1):
        decision = by_row.get(idx)
        if not decision:
            continue
        set_cell(row.cells[spec["valueCol"]], decision["value"], highlight=decision["action"] == "manual")
        if decision["unit"]:
            set_cell(row.cells[spec["unitCol"]], decision["unit"])
    out = OUT_DIR / f"{spec['id']}-{spec['title']}-语义映射样例填充.docx"
    doc.save(out)
    return out


def write_report(spec, mapping, out_doc, facts_payload):
    lines = [
        f"# {spec['id']} {spec['title']} 语义映射样例填充报告",
        "",
        f"- 输出文件：{out_doc.name}",
        f"- 参考素材：{PARAM_XLSX.name}",
        f"- 匹配机型列：{facts_payload['platformMatched']} / {facts_payload['modelMatched']} / Excel 第 {facts_payload['selectedColumn']} 列",
        f"- 已填字段：{mapping['summary']['fill']}",
        f"- 待人工补充：{mapping['summary']['manual']}",
        "",
        "| 字段 | 动作 | 值 | 单位 | 置信度 | 依据/原因 |",
        "|---|---|---|---|---:|---|",
    ]
    for item in mapping["decisions"]:
        action = "填写" if item["action"] == "fill" else "待人工补充"
        fact = item["selectedFact"]
        why = item["reason"]
        if fact:
            why = f"{why} 选中事实：{fact['label']}={fact['value']}（row={fact['row']}, source={fact['source']}）。"
        lines.append(f"| {item['field']} | {action} | {item['value']} | {item['unit']} | {item['confidence']} | {why} |")
    report = OUT_DIR / f"{spec['prefix'].lower()}_semantic_fill_report.md"
    report.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return report


def main():
    OUT_DIR.mkdir(exist_ok=True)
    project = load_project()
    facts = extract_facts(project)
    summary = {}
    for spec in APPENDICES:
        fields = extract_fields(spec)
        mapping = map_fields(spec, fields, facts)
        out_doc = fill_doc(spec, mapping)
        report = write_report(spec, mapping, out_doc, facts)
        summary[spec["id"]] = {
            **mapping["summary"],
            "outputDocx": str(out_doc),
            "report": str(report),
        }
    (OUT_DIR / "c2_c3_semantic_summary.json").write_text(
        json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
