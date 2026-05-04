from __future__ import annotations

import json
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parent
SOURCE_DOCX = ROOT / "APPX-0017-附表C.1 总体技术参数与规格.docx"
PARAM_XLSX = ROOT / "X2平台机型投标参数_20250106.xlsx"
INPUT_JSON = ROOT / "s4_gap_input.json"
OUTPUT_DOCX = ROOT / "APPX-0017-附表C.1 总体技术参数与规格-Codex样例填充.docx"
REPORT_JSON = ROOT / "APPX-0017-Codex样例填充报告.json"
REPORT_MD = ROOT / "APPX-0017-Codex样例填充报告.md"


def clean(value):
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).replace("\n", " / ").strip()


def cell_text(row, col):
    return clean(row.cells[col].text)


def set_cell(cell, text, *, highlight=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "FFF2CC")
        tc_pr.append(shd)


def load_project():
    data = json.loads(INPUT_JSON.read_text(encoding="utf-8"))
    return data["projectTurbineModel"]


def load_param_col(project):
    wb = load_workbook(PARAM_XLSX, data_only=True)
    ws = wb["01-X2平台机组主参数20260120"]

    # The same model name appears twice. The project platform is X2E-A湍, so
    # choose the column whose platform row matches that value.
    model = project["model"]
    platform = project["platform"]
    matches = []
    for col in range(5, ws.max_column + 1):
        if clean(ws.cell(5, col).value) == model:
            matches.append(col)
    for col in matches:
        if clean(ws.cell(4, col).value) == platform:
            return ws, col
    return ws, matches[0]


def excel_value(ws, col, row, *, decimals=None):
    value = ws.cell(row, col).value
    if isinstance(value, float) and decimals is not None:
        return f"{value:.{decimals}f}"
    return clean(value)


def placeholder(name):
    return f"[待人工补充：{name}]"


def main():
    project = load_project()
    ws, col = load_param_col(project)

    rated_power = float(ws.cell(6, col).value)
    swept_area = float(ws.cell(25, col).value)
    area_per_kw = swept_area / rated_power

    mappings = {
        "投标机型": (project["model"], "", "projectTurbineModel.model"),
        "机组类型": ("双馈", "", "由发电机型式“" + excel_value(ws, col, 244) + "”推断"),
        "单机容量": (excel_value(ws, col, 6), "kW", "row 6 额定功率"),
        "机组数量": (None, "", "项目特定字段，当前 manifest 未提供"),
        "总装机容量": (None, "MW", "项目特定字段，当前 manifest 未提供"),
        "叶轮直径": (excel_value(ws, col, 8), "m", "row 8 风轮直径"),
        "叶片数目": (excel_value(ws, col, 164), "片", "row 164 叶片数目"),
        "轮毂中心高度": (None, "m", "项目定制；参数表仅给约束：" + excel_value(ws, col, 162)),
        "单位千瓦扫风面积": (f"{area_per_kw:.2f}", "m2/kW", "row 25 扫风面积 / row 6 额定功率"),
        "认证级别": (excel_value(ws, col, 146), "", "row 146 认证级别"),
        "参考风速Vref": (excel_value(ws, col, 147), "m/s", "row 147 极限风速50年一遇10min"),
        "极端风速Ve50": (excel_value(ws, col, 148), "m/s", "row 148 极限风速50年一遇3s"),
        "轮毂高度年平均风速Vave": (None, "m/s", "场址风资源字段，当前参考资料未提供"),
        "湍流强度参考值Iref": (None, "", "场址风资源字段，当前参考资料未提供"),
        "场址空气密度": (None, "kg/m3", "场址功率曲线字段，当前参考资料未提供"),
        "保证功率曲线最优Cp值": (None, "", "需功率曲线/性能文件，当前参考资料未提供"),
        "保证功率曲线最优Cp对应风速区间（m/s）": (None, "m/s", "需功率曲线/性能文件，当前参考资料未提供"),
        "切入风速": (excel_value(ws, col, 10), "m/s", "row 10 切入风速"),
        "额定风速": (excel_value(ws, col, 11), "m/s", "row 11 额定风速"),
        "切出风速": (excel_value(ws, col, 12), "m/s", "row 12 切出风速"),
        "再切入风速": (None, "m/s", "参数表仅写“参照1.1”，无明确数值"),
        "机组设计寿命": (excel_value(ws, col, 19), "年", "row 19 设计使用寿命"),
        "风轮运行转速范围": (None, "rpm", "当前参数表仅提供风轮额定转速，未提供运行范围"),
        "风轮额定转速": (excel_value(ws, col, 151), "rpm", "row 151 风轮额定转速"),
        "叶尖速度": (excel_value(ws, col, 24, decimals=2), "m/s", "row 24 叶尖线速度"),
        "功率调整方式": (excel_value(ws, col, 156), "", "row 156 功率调整方式"),
        "紧急制动方式": (excel_value(ws, col, 157), "", "row 157 紧急制动方式"),
        "运行环境温度": (excel_value(ws, col, 149), "℃", "row 149 运行环境温度"),
        "生存环境温度": (excel_value(ws, col, 18), "℃", "row 18 机组生存温度范围"),
        "噪音（整个风力发电机组）": ("符合IEC61400-11，≤110dB", "dB", "row 160-161 噪音"),
        "轮毂总成重量": (excel_value(ws, col, 129), "t", "row 129 轮毂"),
        "机舱总成重量": (excel_value(ws, col, 126), "t", "row 126 机舱总成"),
        "机舱尺寸": (excel_value(ws, col, 138), "m", "row 138 主机舱长宽高"),
        "机组防雷等级": (excel_value(ws, col, 116), "", "row 116 防雷措施"),
    }

    doc = Document(SOURCE_DOCX)
    table = doc.tables[0]
    records = []

    for row in table.rows[1:]:
        name = cell_text(row, 2)
        if not name:
            continue
        value, unit, evidence = mappings.get(name, (None, "", "未配置映射"))
        if value:
            set_cell(row.cells[3], value)
            if unit:
                set_cell(row.cells[4], unit)
            records.append({
                "field": name,
                "status": "filled",
                "value": value,
                "unit": unit,
                "evidence": evidence,
            })
        else:
            set_cell(row.cells[3], placeholder(name), highlight=True)
            if unit:
                set_cell(row.cells[4], unit)
            records.append({
                "field": name,
                "status": "unfilled",
                "value": placeholder(name),
                "unit": unit,
                "reason": evidence,
            })

    filled = sum(r["status"] == "filled" for r in records)
    unfilled = sum(r["status"] == "unfilled" for r in records)
    report = {
        "schema": "bid-tech-table-fill-v1.codex-sample",
        "target": {
            "appendixId": "APPX-0017",
            "title": "附表C.1 总体技术参数与规格",
            "sourceFile": str(SOURCE_DOCX),
            "outputFile": str(OUTPUT_DOCX),
        },
        "projectTurbineModel": project,
        "referenceMaterials": [{
            "id": "RAW-0437",
            "name": PARAM_XLSX.name,
            "sheet": ws.title,
            "selectedColumn": col,
            "platformMatched": clean(ws.cell(4, col).value),
            "modelMatched": clean(ws.cell(5, col).value),
        }],
        "summary": {
            "filledCount": filled,
            "unfilledCount": unfilled,
            "totalFields": len(records),
        },
        "fields": records,
    }

    doc.save(OUTPUT_DOCX)
    REPORT_JSON.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    lines = [
        "# APPX-0017 Codex 样例填充报告",
        "",
        f"- 目标文件：{SOURCE_DOCX.name}",
        f"- 输出文件：{OUTPUT_DOCX.name}",
        f"- 参考素材：{PARAM_XLSX.name}",
        f"- 匹配机型列：{clean(ws.cell(4, col).value)} / {clean(ws.cell(5, col).value)} / Excel 第 {col} 列",
        f"- 已填字段：{filled}",
        f"- 未填字段：{unfilled}",
        "",
        "## 字段结果",
        "",
        "| 字段 | 状态 | 值 | 单位 | 依据/原因 |",
        "|---|---|---|---|---|",
    ]
    for item in records:
        status = "已填" if item["status"] == "filled" else "待人工补充"
        why = item.get("evidence") or item.get("reason") or ""
        lines.append(f"| {item['field']} | {status} | {item['value']} | {item.get('unit','')} | {why} |")
    REPORT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")

    print(json.dumps(report["summary"], ensure_ascii=False))
    print(OUTPUT_DOCX)
    print(REPORT_JSON)


if __name__ == "__main__":
    main()
