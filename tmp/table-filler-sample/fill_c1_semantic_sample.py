from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parent
SOURCE_DOCX = ROOT / "APPX-0017-附表C.1 总体技术参数与规格.docx"
PARAM_XLSX = ROOT / "X2平台机型投标参数_20250106.xlsx"
INPUT_JSON = ROOT / "s4_gap_input.json"
OUT_DOCX = ROOT / "APPX-0017-附表C.1 总体技术参数与规格-语义映射样例填充.docx"
TARGET_JSON = ROOT / "target_fields.semantic.json"
FACTS_JSON = ROOT / "reference_facts.semantic.json"
MAPPING_JSON = ROOT / "field_mapping.semantic.json"
REPORT_MD = ROOT / "APPX-0017-语义映射样例填充报告.md"


CONCEPTS = {
    "model": ["投标机型", "制造厂家/型号", "机型型号"],
    "turbine_type": ["机组类型", "发电机型式", "双馈异步发电机", "双馈"],
    "rated_power": ["单机容量", "额定功率", "机组额定功率"],
    "turbine_count": ["机组数量", "台数", "风机数量"],
    "total_capacity": ["总装机容量", "装机容量", "项目容量"],
    "rotor_diameter": ["叶轮直径", "风轮直径"],
    "blade_count": ["叶片数目", "叶片数量"],
    "hub_height": ["轮毂中心高度", "轮毂高度"],
    "swept_area": ["扫风面积"],
    "swept_area_per_kw": ["单位千瓦扫风面积", "单位kW扫风面积"],
    "certification_level": ["认证级别", "安全等级", "机组安全设计等级"],
    "vref": ["参考风速Vref", "极限风速（50年一遇）10min", "极限风速50年一遇10min"],
    "ve50": ["极端风速Ve50", "极限风速（50年一遇）3s", "极限风速50年一遇3s", "3s最大值"],
    "vave": ["轮毂高度年平均风速Vave", "年平均风速"],
    "iref": ["湍流强度参考值Iref", "湍流强度参考值", "特征湍流强度"],
    "air_density": ["场址空气密度", "空气密度"],
    "cp_max": ["保证功率曲线最优Cp值", "最优Cp值"],
    "cp_speed_range": ["保证功率曲线最优Cp对应风速区间", "最优Cp对应风速区间"],
    "cut_in": ["切入风速"],
    "rated_wind_speed": ["额定风速", "额定风速（静态）"],
    "cut_out": ["切出风速", "切出风速（10min平均值）"],
    "recut_in": ["再切入风速", "再启动风速"],
    "design_life": ["机组设计寿命", "设计使用寿命"],
    "rotor_speed_range": ["风轮运行转速范围", "运行转速范围"],
    "rotor_rated_speed": ["风轮额定转速"],
    "tip_speed": ["叶尖速度", "叶尖线速度"],
    "power_control": ["功率调整方式", "功率调节方式"],
    "emergency_brake": ["紧急制动方式", "空气动力制动"],
    "operating_temp": ["运行环境温度", "机组运行温度范围"],
    "survival_temp": ["生存环境温度", "机组生存温度范围"],
    "noise": ["噪音", "噪声"],
    "hub_weight": ["轮毂总成重量", "轮毂重量", "轮毂"],
    "nacelle_weight": ["机舱总成重量", "机舱（总成）", "机舱重量"],
    "nacelle_dimension": ["机舱尺寸", "主机舱: 长×宽×高"],
    "lightning_level": ["机组防雷等级", "防雷措施", "雷电保护等级"],
}

PROJECT_SPECIFIC = {
    "turbine_count",
    "total_capacity",
    "hub_height",
    "vave",
    "iref",
    "air_density",
    "cp_max",
    "cp_speed_range",
}


def clean(value):
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).replace("\n", " / ").strip()


def norm(text):
    return re.sub(r"[\s（）()、/\\:：；;，,。\-_—×*]+", "", clean(text).lower())


def concepts_for(text):
    n = norm(text)
    found = []
    for concept, aliases in CONCEPTS.items():
        for alias in aliases:
            a = norm(alias)
            if a and (a in n or n in a):
                found.append(concept)
                break
    return found


def usable_value(value):
    v = clean(value)
    if not v:
        return False
    if v in {"/", "-", "—", "无", "None", "none", "N/A", "n/a"}:
        return False
    weak = ["见商务", "项目定制", "参照1.1", "参照 1.1", "待定", "无明确"]
    return not any(token in v for token in weak)


def set_cell(cell, text, *, highlight=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "FFF2CC")
        tc_pr.append(shd)


def load_project():
    data = json.loads(INPUT_JSON.read_text(encoding="utf-8"))
    return data["projectTurbineModel"]


def choose_param_col(ws, project):
    matches = [
        col for col in range(5, ws.max_column + 1)
        if clean(ws.cell(5, col).value) == project["model"]
    ]
    for col in matches:
        if clean(ws.cell(4, col).value) == project["platform"]:
            return col
    return matches[0]


def extract_target_fields():
    doc = Document(SOURCE_DOCX)
    table = doc.tables[0]
    fields = []
    for idx, row in enumerate(table.rows[1:], start=1):
        cells = [clean(cell.text) for cell in row.cells]
        if not cells[2]:
            continue
        field_id = f"C1-R{idx:02d}"
        field_text = cells[2]
        fields.append({
            "id": field_id,
            "rowIndex": idx,
            "number": cells[0],
            "group": cells[1],
            "field": field_text,
            "targetColumn": "技术参数与规格",
            "unit": cells[4],
            "remark": cells[5],
            # Use the row field itself as the strong matching signal. The group
            # often contains broad words such as 投标机型 or 场址空气密度 and can
            # contaminate unrelated fields in the same section.
            "concepts": concepts_for(" ".join([field_text, cells[4], cells[5]])),
        })
    TARGET_JSON.write_text(json.dumps(fields, ensure_ascii=False, indent=2), encoding="utf-8")
    return fields


def add_fact(facts, *, label, value, unit="", source, row=None, sheet=None, confidence=0.82, derivation=None):
    value = clean(value)
    if not value:
        return
    facts.append({
        "id": f"F{len(facts)+1:04d}",
        "label": clean(label),
        "value": value,
        "unit": clean(unit),
        "source": source,
        "sheet": sheet,
        "row": row,
        "concepts": concepts_for(label + " " + value),
        "baseConfidence": confidence,
        "usable": usable_value(value) if derivation is None else True,
        "derivation": derivation,
    })


def extract_reference_facts(project):
    wb = load_workbook(PARAM_XLSX, data_only=True)
    ws = wb["01-X2平台机组主参数20260120"]
    col = choose_param_col(ws, project)
    facts = []

    add_fact(facts, label="投标机型", value=project["model"], source="projectTurbineModel", confidence=0.98)
    add_fact(facts, label="投标平台", value=project["platform"], source="projectTurbineModel", confidence=0.96)
    add_fact(facts, label="布局", value=project["layout"], source="projectTurbineModel", confidence=0.95)

    for row in range(1, ws.max_row + 1):
        label = clean(ws.cell(row, 3).value)
        unit = clean(ws.cell(row, 4).value)
        value = clean(ws.cell(row, col).value)
        if not label or label in {"对象", "项目", "容量"} or not value:
            continue
        add_fact(
            facts,
            label=label,
            unit=unit,
            value=value,
            source=PARAM_XLSX.name,
            sheet=ws.title,
            row=row,
        )

    by_label = {}
    for fact in facts:
        by_label.setdefault(fact["label"], fact)
    rated_power = next((fact for fact in facts if fact["label"] == "额定功率" and fact["row"] == 6), None)
    swept_area = by_label.get("扫风面积")
    if rated_power and swept_area:
        try:
            value = float(swept_area["value"]) / float(rated_power["value"])
            add_fact(
                facts,
                label="单位千瓦扫风面积",
                value=f"{value:.2f}",
                unit="m2/kW",
                source="derived",
                confidence=0.9,
                derivation=f"{swept_area['id']} / {rated_power['id']}",
            )
        except ValueError:
            pass

    generator_type = next((f for f in facts if f["label"] == "发电机型式"), None)
    if generator_type and "双馈" in generator_type["value"]:
        add_fact(
            facts,
            label="机组类型",
            value="双馈",
            source="derived",
            confidence=0.88,
            derivation=f"由 {generator_type['id']} 发电机型式推断",
        )

    noise_a = next((f for f in facts if f["row"] == 160), None)
    noise_b = next((f for f in facts if f["row"] == 161), None)
    if noise_a and noise_b:
        add_fact(
            facts,
            label="噪音（整个风力发电机组）",
            value=f"{noise_a['value']}，{noise_b['value']}",
            unit="dB",
            source="derived",
            confidence=0.86,
            derivation=f"{noise_a['id']} + {noise_b['id']}",
        )

    payload = {
        "selectedMaterial": PARAM_XLSX.name,
        "selectedSheet": ws.title,
        "selectedColumn": col,
        "platformMatched": clean(ws.cell(4, col).value),
        "modelMatched": clean(ws.cell(5, col).value),
        "facts": facts,
    }
    FACTS_JSON.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return payload


def score(field, fact):
    target_concepts = set(field["concepts"])
    fact_concepts = set(fact["concepts"])
    if not target_concepts or not fact_concepts:
        return 0.0
    overlap = target_concepts & fact_concepts
    if not overlap:
        return 0.0
    score_value = 0.55 + 0.25 * (len(overlap) / max(len(target_concepts), 1))
    if norm(field["field"]) == norm(fact["label"]):
        score_value += 0.18
    elif norm(field["field"]) in norm(fact["label"]) or norm(fact["label"]) in norm(field["field"]):
        score_value += 0.08
    if field["unit"] and fact["unit"] and norm(field["unit"]) == norm(fact["unit"]):
        score_value += 0.05
    if fact["source"] == "projectTurbineModel":
        score_value += 0.04
    if fact["source"] == "derived":
        score_value += 0.03
    return min(score_value * float(fact.get("baseConfidence", 0.8)), 0.99)


def make_mapping(fields, facts_payload):
    facts = facts_payload["facts"]
    decisions = []
    for field in fields:
        target_concepts = set(field["concepts"])
        candidates = sorted(
            [
                {
                    "factId": fact["id"],
                    "label": fact["label"],
                    "value": fact["value"],
                    "unit": fact["unit"],
                    "source": fact["source"],
                    "row": fact["row"],
                    "score": round(score(field, fact), 3),
                    "usable": fact["usable"],
                    "derivation": fact["derivation"],
                }
                for fact in facts
                if score(field, fact) > 0
            ],
            key=lambda item: item["score"],
            reverse=True,
        )

        selected = next((item for item in candidates if item["usable"] and item["score"] >= 0.62), None)
        reason = ""
        if selected and target_concepts & PROJECT_SPECIFIC and selected["source"] != "projectTurbineModel":
            selected = None
            reason = "该字段是项目/场址特定值，当前参考资料只有机型参数或约束，不能直接填。"
        elif not selected:
            if candidates:
                best = candidates[0]
                reason = (
                    "有相近候选，但值不可直接使用或置信度不足："
                    f"{best['label']}={best['value']}，score={best['score']}"
                )
            else:
                reason = "未找到可映射的参考事实。"

        if selected:
            value = selected["value"]
            # The table expects a concise value, not explanatory text from a source row.
            decisions.append({
                "targetFieldId": field["id"],
                "field": field["field"],
                "action": "fill",
                "value": value,
                "unit": field["unit"] or selected["unit"],
                "confidence": selected["score"],
                "selectedFact": selected,
                "alternatives": candidates[1:4],
                "reason": "语义概念匹配并通过可用性检查。",
            })
        else:
            decisions.append({
                "targetFieldId": field["id"],
                "field": field["field"],
                "action": "manual",
                "value": f"[待人工补充：{field['field']}]",
                "unit": field["unit"] or "",
                "confidence": 0,
                "selectedFact": None,
                "alternatives": candidates[:4],
                "reason": reason,
            })

    payload = {
        "schema": "bid-tech-table-field-mapping-v1.codex-semantic-sample",
        "targetFile": SOURCE_DOCX.name,
        "referenceMaterial": PARAM_XLSX.name,
        "selectedColumn": facts_payload["selectedColumn"],
        "platformMatched": facts_payload["platformMatched"],
        "modelMatched": facts_payload["modelMatched"],
        "summary": {
            "fill": sum(d["action"] == "fill" for d in decisions),
            "manual": sum(d["action"] == "manual" for d in decisions),
            "total": len(decisions),
        },
        "decisions": decisions,
    }
    MAPPING_JSON.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return payload


def fill_docx(mapping):
    doc = Document(SOURCE_DOCX)
    table = doc.tables[0]
    by_id = {item["targetFieldId"]: item for item in mapping["decisions"]}
    for idx, row in enumerate(table.rows[1:], start=1):
        field_id = f"C1-R{idx:02d}"
        decision = by_id.get(field_id)
        if not decision:
            continue
        set_cell(row.cells[3], decision["value"], highlight=decision["action"] == "manual")
        if decision["unit"]:
            set_cell(row.cells[4], decision["unit"])
    doc.save(OUT_DOCX)


def write_report(mapping, facts_payload):
    lines = [
        "# APPX-0017 语义映射样例填充报告",
        "",
        f"- 目标文件：{SOURCE_DOCX.name}",
        f"- 输出文件：{OUT_DOCX.name}",
        f"- 参考素材：{PARAM_XLSX.name}",
        f"- 匹配机型列：{facts_payload['platformMatched']} / {facts_payload['modelMatched']} / Excel 第 {facts_payload['selectedColumn']} 列",
        f"- 已填字段：{mapping['summary']['fill']}",
        f"- 待人工补充：{mapping['summary']['manual']}",
        "",
        "## 中间产物",
        "",
        f"- `{TARGET_JSON.name}`：从业主空表自动抽取的字段。",
        f"- `{FACTS_JSON.name}`：从参考资料自动抽取的候选事实。",
        f"- `{MAPPING_JSON.name}`：字段语义映射、置信度、候选事实和人工补充原因。",
        "",
        "## 字段映射结果",
        "",
        "| 字段 | 动作 | 值 | 单位 | 置信度 | 映射依据/原因 |",
        "|---|---|---|---|---:|---|",
    ]
    for item in mapping["decisions"]:
        action = "填写" if item["action"] == "fill" else "待人工补充"
        fact = item["selectedFact"]
        why = item["reason"]
        if fact:
            why = f"{why} 选中事实：{fact['label']}={fact['value']}（source={fact['source']}, row={fact['row']}）。"
        lines.append(
            f"| {item['field']} | {action} | {item['value']} | {item['unit']} | {item['confidence']} | {why} |"
        )
    REPORT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main():
    project = load_project()
    fields = extract_target_fields()
    facts_payload = extract_reference_facts(project)
    mapping = make_mapping(fields, facts_payload)
    fill_docx(mapping)
    write_report(mapping, facts_payload)
    print(json.dumps(mapping["summary"], ensure_ascii=False))
    print(OUT_DOCX)
    print(TARGET_JSON)
    print(FACTS_JSON)
    print(MAPPING_JSON)


if __name__ == "__main__":
    main()
