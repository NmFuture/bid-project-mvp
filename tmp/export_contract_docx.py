from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/wlb/Agent/bid-project")
INPUT_MD = ROOT / "doc/10-甲方技术细议草案-合同预期最终交付版.md"
OUTPUT_DOCX = ROOT / "doc/10-甲方技术细议草案-合同预期最终交付版.docx"


def run_pandoc() -> None:
    subprocess.run(
        [
            "pandoc",
            "--from",
            "gfm+pipe_tables",
            "--to",
            "docx",
            "--output",
            str(OUTPUT_DOCX),
            str(INPUT_MD),
        ],
        check=True,
        cwd=ROOT,
    )


def set_run_fonts(run, east_asia: str, latin: str, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = latin
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def style_font(style, east_asia: str, latin: str, size: float, bold: bool = False) -> None:
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def ensure_title_style(doc: Document) -> None:
    styles = doc.styles
    title_style = styles["Title"]
    style_font(title_style, "SimHei", "Calibri", 18, bold=True)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    title_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    style_font(normal, "SimSun", "Times New Roman", 12)
    normal.paragraph_format.line_spacing = 1.4
    normal.paragraph_format.space_after = Pt(6)

    heading1 = styles["Heading 1"]
    style_font(heading1, "SimHei", "Calibri", 16, bold=True)
    heading1.paragraph_format.space_before = Pt(14)
    heading1.paragraph_format.space_after = Pt(6)

    heading2 = styles["Heading 2"]
    style_font(heading2, "SimHei", "Calibri", 14, bold=True)
    heading2.paragraph_format.space_before = Pt(10)
    heading2.paragraph_format.space_after = Pt(4)

    heading3 = styles["Heading 3"]
    style_font(heading3, "SimHei", "Calibri", 12, bold=True)
    heading3.paragraph_format.space_before = Pt(8)
    heading3.paragraph_format.space_after = Pt(4)

    if "Quote" in styles:
        quote = styles["Quote"]
        style_font(quote, "SimSun", "Times New Roman", 10.5)
        quote.paragraph_format.left_indent = Cm(0.7)
        quote.paragraph_format.right_indent = Cm(0.7)

    ensure_title_style(doc)


def configure_page(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)


def format_paragraphs(doc: Document) -> None:
    title_applied = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if not title_applied and text == "AI数智化投标平台技术细议草案":
            para.style = doc.styles["Title"]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                set_run_fonts(run, "SimHei", "Calibri", size=18, bold=True)
            title_applied = True
            continue

        if para.style.name == "Heading 1":
            for run in para.runs:
                set_run_fonts(run, "SimHei", "Calibri", size=16, bold=True)
        elif para.style.name == "Heading 2":
            for run in para.runs:
                set_run_fonts(run, "SimHei", "Calibri", size=14, bold=True)
        elif para.style.name == "Heading 3":
            for run in para.runs:
                set_run_fonts(run, "SimHei", "Calibri", size=12, bold=True)
        else:
            para.paragraph_format.line_spacing = 1.4
            para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                set_run_fonts(run, "SimSun", "Times New Roman", size=12)


def format_tables(doc: Document) -> None:
    table_style_name = None
    for candidate in ("Table Grid", "TableGrid", "Light Grid"):
        if candidate in doc.styles:
            table_style_name = candidate
            break

    for table in doc.tables:
        if table_style_name:
            table.style = table_style_name
        for r_idx, row in enumerate(table.rows):
            for cell in row.cells:
                if r_idx == 0:
                    set_cell_shading(cell, "D9E2F3")
                for para in cell.paragraphs:
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.line_spacing = 1.15
                    for run in para.runs:
                        set_run_fonts(
                            run,
                            "SimSun",
                            "Times New Roman",
                            size=10.5,
                            bold=True if r_idx == 0 else None,
                        )


def main() -> None:
    run_pandoc()
    doc = Document(OUTPUT_DOCX)
    configure_styles(doc)
    configure_page(doc)
    format_paragraphs(doc)
    format_tables(doc)
    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
